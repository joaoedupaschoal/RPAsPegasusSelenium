# -*- coding: utf-8 -*-
"""
QA Reporter — Relatório DOCX (Word) com gráficos
------------------------------------------------

Uso básico:
-----------
from qa_reporter import QAReporter

rep = QAReporter(out_dir="reports", environment="Homologação",
                 executor="Runner CLI", system_version="v2025.09")

rep.start_run(summary="Execução em massa — Cadastros Adicionais")

h = rep.start_scenario("Cadastro de Cor — Cenário 1", test_type="CADASTRO")
# ... rodar o cenário ...
rep.finish_scenario(h, status="SUCCESS", extra={"file_name": "cadastrodecor1ºcenario.py",
                                                "logfile": "reports/logs/cadastrodecor1ºcenario_01.log"})

rep.end_run()
docx_path = rep.save_docx("Relatorio_QA")
print("Gerado:", docx_path)
"""

from __future__ import annotations

from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional
import platform
import re

# --- Dependências opcionais (o reporter continua funcionando sem elas) ---
_HAS_MPL = False
_HAS_DOCX = False
try:
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    _HAS_MPL = True
except Exception:
    _HAS_MPL = False

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    _HAS_DOCX = True
except Exception:
    _HAS_DOCX = False


# ======================================================================
# Utilidades
# ======================================================================

def _fmt_dt(dt: Optional[datetime]) -> str:
    if not dt:
        return ""
    return dt.strftime("%d/%m/%Y %H:%M:%S")


def _ms_readable(ms: int) -> str:
    s = int(round(ms / 1000.0))
    h = s // 3600
    m = (s % 3600) // 60
    sec = s % 60
    parts = []
    if h:
        parts.append(f"{h}h")
    if m or (h and not m):
        parts.append(f"{m}m")
    parts.append(f"{sec}s")
    return " ".join(parts)


def safe_excerpt(text: str, max_chars=1200, head_lines=25, tail_lines=15) -> str:
    """Faz um resumo que preserva começo e fim do log (sem cortar feio)."""
    if not text:
        return ""
    if len(text) <= max_chars:
        return text
    lines = text.splitlines()
    if len(lines) > head_lines + tail_lines:
        meio = "\n[... log truncado para o relatório ...]\n"
        return "\n".join(lines[:head_lines]) + meio + "\n".join(lines[-tail_lines:])
    m = re.match(rf"^(.{{1,{max_chars}}})(?:\b|$)", text, flags=re.S)
    return (m.group(1).rstrip() + " ...")


# ======================================================================
# Estruturas
# ======================================================================

@dataclass
class ScenarioResult:
    name: str
    status: str
    started_at: datetime
    ended_at: datetime
    duration_ms: int
    test_type: str = "SCENARIO"
    error_message: Optional[str] = None
    screenshots: Optional[List[str]] = None
    extra: Optional[Dict[str, Any]] = None
    test_id: Optional[str] = None
    severity: str = "MEDIUM"
    file_name: Optional[str] = None  # duplicado em extra["file_name"] por comodidade


# ======================================================================
# Reporter
# ======================================================================

class QAReporter:
    def __init__(
        self,
        out_dir: str | Path = "reports",
        environment: str = "Produção",
        executor: str = "Pipeline CI/CD",
        system_version: str = "1.0.0",
        username: str = "Robô de Testes",
    ):
        self.out_dir = Path(out_dir)
        self.out_dir.mkdir(parents=True, exist_ok=True)
        self._results: List[ScenarioResult] = []
        self._run_started_at: Optional[datetime] = None
        self._run_ended_at: Optional[datetime] = None

        self.environment = environment
        self.executor = executor
        self.system_version = system_version
        self.username = username
        self.summary = ""
        self._stamp = datetime.now().strftime("%Y%m%d-%H%M%S")

    # ---------------- ciclo de execução ----------------

    def start_run(self, summary: str = ""):
        self._run_started_at = datetime.now()
        self.summary = summary or ""

    def end_run(self):
        self._run_ended_at = datetime.now()

    # ---------------- cenários ----------------

    def start_scenario(self, name: str, test_type: str = "SCENARIO", test_id: str | None = None) -> Dict[str, Any]:
        return {
            "name": name,
            "test_type": test_type,
            "test_id": test_id or f"{test_type}_{len(self._results)+1:04d}",
            "started_at": datetime.now(),
        }

    def finish_scenario(
        self,
        handle,
        status: Optional[str] = None,
        error_message: Optional[str] = None,
        screenshots: Optional[List[str]] = None,
        severity: Optional[str] = None,
        extra: Optional[Dict[str, Any]] = None,
    ):
        now = datetime.now()
        dur_ms = int((now - handle["started_at"]).total_seconds() * 1000)

        extra = extra or {}
        file_name = extra.get("file_name")

        self._results.append(
            ScenarioResult(
                name=handle["name"],
                status=str(status or "SUCCESS").upper(),
                started_at=handle["started_at"],
                ended_at=now,
                duration_ms=max(0, dur_ms),
                test_type=handle.get("test_type", "SCENARIO"),
                error_message=(error_message or None),
                screenshots=screenshots or [],
                extra=extra,
                test_id=handle.get("test_id"),
                severity=str(severity or "MEDIUM").upper(),
                file_name=file_name,
            )
        )

    # ---------------- métricas ----------------

    def _metrics(self) -> Dict[str, Any]:
        total = len(self._results)
        ok = sum(1 for r in self._results if r.status == "SUCCESS")
        err = sum(1 for r in self._results if r.status == "ERROR")
        alt = sum(1 for r in self._results if r.status == "ALERT")
        pass_rate = (ok / total * 100.0) if total else 0.0

        started = self._run_started_at or (min((r.started_at for r in self._results), default=None))
        ended = self._run_ended_at or (max((r.ended_at for r in self._results), default=None))
        window_ms = int((ended - started).total_seconds() * 1000) if (started and ended) else 0
        dur_total_ms = sum(r.duration_ms for r in self._results)
        dur_avg_ms = int(dur_total_ms / total) if total else 0

        cad = [r for r in self._results if r.test_type == "CADASTRO"]
        cad_total = len(cad)
        cad_ok = sum(1 for r in cad if r.status == "SUCCESS")
        cad_err = sum(1 for r in cad if r.status == "ERROR")
        cad_alt = sum(1 for r in cad if r.status == "ALERT")
        cad_pass = (cad_ok / cad_total * 100.0) if cad_total else 0.0
        cad_dur_total = sum(r.duration_ms for r in cad)
        cad_dur_avg = int(cad_dur_total / cad_total) if cad_total else 0

        return {
            "total": total,
            "success": ok,
            "error": err,
            "alert": alt,
            "pass_rate_pct": round(pass_rate, 2),
            "started_at": started,
            "ended_at": ended,
            "window_ms": window_ms,
            "duration_total_ms": dur_total_ms,
            "duration_avg_ms": dur_avg_ms,
            "cad_total": cad_total,
            "cad_success": cad_ok,
            "cad_error": cad_err,
            "cad_alert": cad_alt,
            "cad_pass_rate_pct": round(cad_pass, 2),
            "cad_duration_total_ms": cad_dur_total,
            "cad_duration_avg_ms": cad_dur_avg,
            "failures": [r for r in self._results if r.status == "ERROR"],
            "system_info": {
                "platform": platform.platform(),
                "python": platform.python_version(),
                "hostname": platform.node(),
                "processor": platform.processor(),
            },
        }

    # ---------------- gráficos ----------------

    def _make_pie(self, ok: int, err: int, alt: int, title: str,
                  size_in: float = 2.8, dpi: int = 180) -> Optional[Path]:
        if not _HAS_MPL:
            return None
        labels = ["Sucesso", "Erro", "Alerta"]
        sizes = [ok, err, alt]
        if sum(sizes) == 0:
            return None

        colors = ["#2ecc71", "#e74c3c", "#f1c40f"]  # verde, vermelho, amarelo

        def _autopct(pct: float) -> str:
            return f"{pct:.0f}%" if pct >= 1.0 else ""

        fig, ax = plt.subplots(figsize=(size_in, size_in), dpi=dpi)
        ax.pie(
            sizes,
            labels=[l if v > 0 else "" for l, v in zip(labels, sizes)],
            colors=colors,
            startangle=90,
            autopct=_autopct,
            pctdistance=0.72,
            labeldistance=1.08,
            wedgeprops={"linewidth": 0.8, "edgecolor": "white"},
            textprops={"fontsize": 8},
        )
        ax.axis("equal")
        ax.set_title(title, fontsize=10)

        out = self.out_dir / f"chart_{title.replace(' ', '_').lower()}_{self._stamp}.png"
        fig.savefig(out, bbox_inches="tight")
        plt.close(fig)
        return out

    # ---------------- DOCX ----------------

    def save_docx(self, base_name: str = "Relatorio_QA") -> Path:
        if not _HAS_DOCX:
            raise RuntimeError("python-docx não encontrado. Instale com: pip install python-docx")

        m = self._metrics()
        chart_all = self._make_pie(m["success"], m["error"], m["alert"], "Geral")
        chart_cad = self._make_pie(m["cad_success"], m["cad_error"], m["cad_alert"], "Cadastros")

        doc = Document()

        # Estilo global
        style = doc.styles["Normal"]
        style.font.name = "Arial"
        style.font.size = Pt(11)

        # Cabeçalho
        h = doc.add_heading("Relatório de Eficiência do Sistema — Execução de Testes", level=0)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f"Data/Hora do Relatório: {_fmt_dt(datetime.now())}")
        doc.add_paragraph(f"Período do Teste: {_fmt_dt(m['started_at'])} — {_fmt_dt(m['ended_at'])}")
        if self.summary:
            doc.add_paragraph(f"Resumo da Execução: {self.summary}")
        doc.add_paragraph("")

        # Resumo (Cadastros)
        doc.add_heading("Resumo da Execução", level=1)
        table = doc.add_table(rows=6, cols=3)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = table.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text = "Métrica", "Quantidade", "Percentual"
        for c in hdr:
            if c.paragraphs and c.paragraphs[0].runs:
                c.paragraphs[0].runs[0].bold = True

        data_rows = [
            ("Cadastros com Sucesso", str(m["cad_success"]), f"{m['cad_pass_rate_pct']:.2f}%"),
            ("Falhas (Erro)", str(m["cad_error"]),
             f"{(m['cad_error']/m['cad_total']*100.0 if m['cad_total'] else 0):.2f}%"),
            ("Total de Cadastros Executados", str(m["cad_total"]), "—"),
            ("Tempo Total de Execução", _ms_readable(m["cad_duration_total_ms"]), "—"),
            ("Média por Cadastro", _ms_readable(m["cad_duration_avg_ms"]), "—"),
        ]
        for i, (met, qtd, pct) in enumerate(data_rows, start=1):
            r = table.rows[i].cells
            r[0].text, r[1].text, r[2].text = met, qtd, pct

        doc.add_paragraph("")
        obs = doc.add_paragraph()
        obs.add_run("Obs.: ").bold = True
        obs.add_run('Registros com "alerta" não entram no cálculo de sucesso/erro.')
        doc.add_paragraph("")

        # Gráficos
        if chart_all or chart_cad:
            doc.add_heading("Gráficos", level=1)
        if chart_all and Path(chart_all).exists():
            doc.add_paragraph("Distribuição Geral")
            doc.add_picture(str(chart_all), width=Inches(2.8))
        if chart_cad and Path(chart_cad).exists():
            doc.add_paragraph("Distribuição — Cadastros")
            doc.add_picture(str(chart_cad), width=Inches(2.8))
        if chart_all or chart_cad:
            doc.add_paragraph("")

        # Listagem completa dos cenários (com regras de erro/ok)
        doc.add_heading("Detalhes dos Cenários", level=1)
        if not self._results:
            doc.add_paragraph("Nenhum cenário registrado.")
        else:
            for idx, r in enumerate(self._results, 1):
                p = doc.add_paragraph()
                p.add_run(f"{idx}. ").bold = True
                p.add_run(r.name)

                meta = []
                if r.test_type:
                    meta.append(f"Tipo: {r.test_type}")
                if r.test_id:
                    meta.append(f"ID: {r.test_id}")
                if r.severity:
                    meta.append(f"Severidade: {r.severity}")
                meta.append(f"Início: {_fmt_dt(r.started_at)}")
                meta.append(f"Duração: {_ms_readable(r.duration_ms)}")
                doc.add_paragraph("   • " + " | ".join(meta))

                if r.file_name or (r.extra and r.extra.get("file_name")):
                    fname = r.file_name or r.extra.get("file_name")
                    doc.add_paragraph(f"   • Arquivo: {fname}")

                # Só imprime "Erro:" quando realmente houve erro
                if r.status == "ERROR" and r.error_message:
                    doc.add_paragraph("   • Erro: " + safe_excerpt(r.error_message, max_chars=800))

                # Caminho do log completo (se veio do runner)
                if r.extra and r.extra.get("logfile"):
                    lf = Path(r.extra["logfile"])
                    doc.add_paragraph(f"   • Log completo salvo em: {lf.name}")

                if r.screenshots:
                    doc.add_paragraph("   • Evidências: " + ", ".join(r.screenshots))
                doc.add_paragraph("")

        # Seção: automações com ERRO (para reteste)
        fails = [r for r in self._results if r.status == "ERROR"]
        doc.add_heading("Automações com Erro (para reteste manual)", level=1)
        if not fails:
            doc.add_paragraph("Nenhuma automação falhou nesta execução.")
        else:
            doc.add_paragraph(f"Foram encontradas {len(fails)} automações com erro:")
            doc.add_paragraph("")
            for i, f in enumerate(fails, 1):
                p = doc.add_paragraph()
                p.add_run(f"{i}. ").bold = True
                p.add_run(f.name)
                first = (f.error_message or "").strip().splitlines()[0][:250]
                if first:
                    doc.add_paragraph(f"   • Erro: {first}")
                if f.extra and f.extra.get("logfile"):
                    doc.add_paragraph(f"   • Log completo: {Path(f.extra['logfile']).name}")
                if f.file_name:
                    doc.add_paragraph(f"   • Arquivo: {f.file_name}")
                doc.add_paragraph("")

        # Detalhes técnicos
        doc.add_heading("Detalhes Técnicos", level=1)
        doc.add_paragraph(f"• Ambiente: {self.environment}")
        doc.add_paragraph(f"• Executor: {self.executor}")
        doc.add_paragraph(f"• Versão do Sistema: {self.system_version}")
        doc.add_paragraph(f"• Plataforma: {m['system_info']['platform']}")
        doc.add_paragraph(f"• Hostname: {m['system_info']['hostname']}")
        doc.add_paragraph(f"• Python: {m['system_info']['python']}")
        doc.add_paragraph("")

        # TXT auxiliar com lista de falhas
        fail_txt = self.out_dir / f"automacoes_com_erro_{self._stamp}.txt"
        if fails:
            lines = []
            for f in fails:
                first_line = (f.error_message or "").strip().splitlines()[0][:200]
                lines.append(
                    " | ".join(
                        [
                            f.name,
                            f"Arquivo={f.file_name or f.extra.get('file_name', '-') if f.extra else (f.file_name or '-')}",
                            f"Tipo={f.test_type}",
                            f"ID={f.test_id or '-'}",
                            f"Sev={f.severity}",
                            f"Início={_fmt_dt(f.started_at)}",
                            f"Erro={first_line or '-'}",
                        ]
                    )
                )
            try:
                fail_txt.write_text("\n".join(lines), encoding="utf-8")
            except Exception:
                pass  # não quebrar o fluxo caso o disco esteja protegido

        out = self.out_dir / f"{base_name}_{self._stamp}.docx"
        doc.save(str(out))
        return out
