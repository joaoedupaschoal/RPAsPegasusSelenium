import sys
import os

# Adiciona a raiz do projeto ao sys.path
sys.path.append(
    os.path.abspath(
        os.path.join(os.path.dirname(__file__), "../../..")
    )
)


from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service
from webdriver_manager.chrome import ChromeDriverManager
from docx import Document
from datetime import datetime
import subprocess
import time
import os
import sys
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), '..', '..')))

from utils.actions import log, take_screenshot, safe_action, encontrar_mensagem_alerta, ajustar_zoom

URL = "http://localhost:8080/gs/index.xhtml"
LOGIN_EMAIL = "joaoeduardo.gold@outlook.com"
LOGIN_PASSWORD = "071999gs"

def main():
    doc = Document()
    doc.add_heading("RELATÓRIO DO TESTE", 0)
    doc.add_paragraph("Cadastro de Carrtórios.")
    doc.add_paragraph(f"🗕️ Data do teste: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")
    doc.add_paragraph("O robô preencherá os campos obrigatórios e opcionais e realizará o cadastro de um Cartório.")

    chrome_options = Options()
    chrome_options.add_argument("--start-maximized")
    driver = webdriver.Chrome(service=Service(ChromeDriverManager().install()), options=chrome_options)
    wait = WebDriverWait(driver, 10)

    def finalizar_relatorio():
        doc_name = f"relatorio_cartorios_cenario_2_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        doc.save(doc_name)
        log(doc, f"📄 Relatório salvo como: {doc_name}")
        try:
            subprocess.run(["start", "winword", doc_name], shell=True)
        except Exception as e:
            log(doc, f"Erro ao abrir o Word: {e}")
        driver.quit()

    def login():
        log(doc, "🔄 Iniciando login no sistema.")
        take_screenshot(driver, doc, "login_inicio")

        wait.until(EC.presence_of_element_located((By.ID, "j_id15:email"))).send_keys(LOGIN_EMAIL)
        wait.until(EC.presence_of_element_located((By.ID, "j_id15:senha"))).send_keys(LOGIN_PASSWORD, Keys.ENTER)
        wait.until(EC.presence_of_element_located((By.TAG_NAME, "body")))
        time.sleep(2)

        log(doc, "✅ Login realizado com sucesso.")
        take_screenshot(driver, doc, "login_concluido")

    def abrir_menu_cartorios():
        log(doc, "🔄 Abrindo o menu principal.")
        take_screenshot(driver, doc, "menu_inicio")

        driver.find_element(By.TAG_NAME, "body").send_keys(Keys.F2)
        time.sleep(1)

        campo = wait.until(EC.presence_of_element_located((By.XPATH, "//input[@placeholder='Busque um cadastro']")))
        campo.click()
        campo.send_keys("Cartórios", Keys.ENTER)
        time.sleep(2)

        log(doc, "✅ Menu Cartórios localizado e aberto.")
        take_screenshot(driver, doc, "menu_capelas_aberto")

    def clicar_em_cadastrar():
        log(doc, "🔄 Acessando o formulário de cadastro.")
        take_screenshot(driver, doc, "cadastro_inicio")

        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "#fmod_10029 > div.wdTelas > div > ul > li:nth-child(1) > a > span"))).click()
        time.sleep(2)

        log(doc, "✅ Formulário de cadastro aberto.")
        take_screenshot(driver, doc, "cadastro_aberto")

    def preencher_campos_e_cancelar():
        log(doc, "🔄 Preenchendo o campo 'Nome'.")
        campo_nome = wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR,
            "#fmod_10029 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div > div > div > div > input")))
        campo_nome.send_keys("TESTE CARTÓRIO SELENIUM")
        log(doc, "✅ Campo 'Nome' preenchido.")
        take_screenshot(driver, doc, "nome_preenchido")

        
        log(doc, "🔄 Clicando no botão 'Cancelar'.")
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR,
            "#fmod_10029 > div.wdTelas > div.telaCadastro.clearfix > div.btnHolder > a.btModel.btGray.btcancel"))).click()
        time.sleep(2)
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR,
            "#BtYes"))).click()
        log(doc, "✅ Cadastro cancelado com sucesso.")
        take_screenshot(driver, doc, "apos_clicar_salvar")

    def fechar_modal():
        log(doc, "🔄 Fechando o formulário.")
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR,
            "#fmod_10029 > div.wdTop.ui-draggable-handle > div.wdClose > a"))).click()
        time.sleep(1)
        log(doc, "✅ Formulário fechado.")
        take_screenshot(driver, doc, "formulario_fechado")

    # Execução do fluxo
    log(doc, "🔄 Acessando a URL do sistema.")
    driver.get(URL)
    take_screenshot(driver, doc, "url_acessada")

    login()
    ajustar_zoom(driver)
    abrir_menu_cartorios()
    clicar_em_cadastrar()
    preencher_campos_e_cancelar()

    # Mensagem de alerta
    _, tipo_alerta = encontrar_mensagem_alerta(driver, doc)
    if tipo_alerta == "sucesso":
        log(doc, "✅ Mensagem de sucesso exibida após o cadastro.")
    elif tipo_alerta == "alerta":
        log(doc, "⚠️ Alerta exibido após o cadastro.")
    elif tipo_alerta == "erro":
        log(doc, "❌ Erro exibido após o cadastro.")
    else:
        log(doc, "⚠️ Nenhuma mensagem foi exibida após o cadastro.")
    take_screenshot(driver, doc, "mensagem_final")

    fechar_modal()
    log(doc, "✅ Teste finalizado com sucesso.")
    finalizar_relatorio()

if __name__ == "__main__":
    main()
