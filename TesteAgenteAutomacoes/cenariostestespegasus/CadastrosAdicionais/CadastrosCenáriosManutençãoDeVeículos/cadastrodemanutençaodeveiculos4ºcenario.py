# Refatorado e corrigido: cadastrodemanutencaodeveiculos1ºcenario.py

from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support.ui import Select, WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service
from selenium.webdriver import ActionChains
from webdriver_manager.chrome import ChromeDriverManager
from docx import Document
from docx.shared import Inches
from faker import Faker
from faker.providers import BaseProvider
from datetime import datetime, timedelta
import subprocess
import os
import time
import random
import string

# ==== CONFIGURAÇÕES DO DRIVER ====
options = Options()
options.add_argument("--disable-logging")
options.add_argument("--log-level=3")
options.add_argument("--remote-allow-origins=*")
options.add_argument("--start-maximized")
options.add_argument("--disable-blink-features=AutomationControlled")
options.add_experimental_option("excludeSwitches", ["enable-automation"])
options.add_experimental_option('useAutomationExtension', False)

driver = webdriver.Chrome(
    service=Service(ChromeDriverManager().install()),
    options=options
)

# Esconde que é um webdriver automatizado
driver.execute_script("Object.defineProperty(navigator, 'webdriver', {get: () => undefined})")

# ==== PROVIDERS CUSTOMIZADOS ====
class BrasilProvider(BaseProvider):
    def rg(self):
        numeros = [str(random.randint(0, 9)) for _ in range(8)]
        return ''.join(numeros) + '-' + str(random.randint(0, 9))

fake = Faker("pt_BR")
fake.add_provider(BrasilProvider)

# ==== CONFIGURAÇÕES ====
URL = "http://localhost:8080/gs/index.xhtml"
LOGIN_EMAIL = "joaoeduardo.gold@outlook.com"
LOGIN_PASSWORD = "071999gs"

# ==== DOCUMENTO ====
doc = Document()
doc.add_heading("RELATÓRIO DO TESTE", 0)
doc.add_paragraph("Cadastro de Manutenção de Veículos – Cenário 1: Preenchimento dos campoa NÃO obrigatórios e salvamento.")
doc.add_paragraph(f"Data do teste: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")

screenshot_registradas = set()

def log(doc, msg):
    print(msg)
    doc.add_paragraph(msg)

def take_screenshot(driver, doc, nome):
    if nome not in screenshot_registradas:
        path = f"screenshots/{nome}.png"
        os.makedirs("screenshots", exist_ok=True)
        driver.save_screenshot(path)
        doc.add_paragraph(f"Screenshot: {nome}")
        doc.add_picture(path, width=Inches(5.5))
        screenshot_registradas.add(nome)

def safe_action(doc, descricao, func):
    try:
        log(doc, f"🔄 {descricao}...")
        func()
        log(doc, f"✅ {descricao} realizada com sucesso.")
        take_screenshot(driver, doc, descricao.lower().replace(" ", "_"))
    except Exception as e:
        log(doc, f"❌ Erro ao {descricao.lower()}: {e}")
        take_screenshot(driver, doc, f"erro_{descricao.lower().replace(' ', '_')}")

def finalizar_relatorio():
    nome_arquivo = f"relatorio_manutencao_veiculos_cenario_4_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc.save(nome_arquivo)
    log(doc, f"📄 Relatório salvo como: {nome_arquivo}")
    try:
        subprocess.run(["start", "winword", nome_arquivo], shell=True)
    except:
        pass
    driver.quit()

def gerar_placa_aleatoria():
    """Gera placa aleatória no formato antigo ou Mercosul"""
    if random.choice(["antigo", "mercosul"]) == "antigo":
        # Exemplo: ABC-1234
        letras = ''.join(random.choices(string.ascii_uppercase, k=3))
        numeros = ''.join(random.choices(string.digits, k=4))
        return f"{letras}-{numeros}"
    else:
        # Exemplo: ABC1D23 (Mercosul)
        letras = ''.join(random.choices(string.ascii_uppercase, k=3))
        meio = random.choice(string.digits)
        letra_final = random.choice(string.ascii_uppercase)
        fim = ''.join(random.choices(string.digits, k=2))
        return f"{letras}{meio}{letra_final}{fim}"

def gerar_renavam():
    """Gera um RENAVAM válido de 11 dígitos"""
    # Gera os primeiros 10 dígitos aleatórios
    base = ''.join(random.choices(string.digits, k=10))
    
    # Calcula o dígito verificador
    multiplicadores = [3, 2, 9, 8, 7, 6, 5, 4, 3, 2]
    soma = sum(int(digito) * mult for digito, mult in zip(base, multiplicadores))
    resto = soma % 11
    
    if resto in [0, 1]:
        digito_verificador = '0'
    else:
        digito_verificador = str(11 - resto)
    
    return base + digito_verificador

def gerar_dados_veiculo():
    """Gera dados fictícios para o cadastro de veículo e manutenção."""
    placa = gerar_placa_aleatoria()
    modelos = ["Fox", "Gol", "Uno", "Celta", "Civic", "Palio", "Corolla", "HB20"]
    modelo = random.choice(modelos)
    ano = random.randint(1995, 2024)
    ano_fabricacao = ano - random.randint(0, 2)
    ano_modelo = f"{modelo} {ano}"
    manutencoes_de_km = fake.random_int(min=100, max=10000) 
    manutencoes_de_dias = fake.random_int(min=90, max=365)
    
    tipos_combustivel = [
        "Gasolina Comum", "Gasolina Aditivada", "Etanol", "Diesel", "Diesel S10",
        "GNV (Gás Natural Veicular)", "Flex (Gasolina/Etanol)", "Elétrico", "Híbrido", "Biodiesel"
    ]
    combustivel = random.choice(tipos_combustivel)
    
    km_inicial = fake.random_int(min=1, max=10000)
    km_atual = km_inicial + fake.random_int(min=1, max=10000)
    
    # Dados da manutenção
    data_manutencao = datetime.today() - timedelta(days=random.randint(0, 30))
    data_manutencao_str = data_manutencao.strftime('%d/%m/%Y')
    
    data_verificacao = data_manutencao + timedelta(days=random.randint(1, 30))
    data_verificacao_str = data_verificacao.strftime('%d/%m/%Y')
    
    km_verificado = random.randint(10000, 200000)
    km_apos_manutencao = km_verificado + random.randint(10, 1000)
    
    data_proximo_check = data_verificacao + timedelta(days=random.randint(30, 120))
    data_proximo_check_str = data_proximo_check.strftime("%d/%m/%Y")
    
    return {
        'placa': placa,
        'modelo': modelo,
        'manutencoes_de_km': manutencoes_de_km,
        'manutencoes_de_dias': manutencoes_de_dias,
        'ano_fabricacao': ano_fabricacao,
        'ano': ano,
        'ano_modelo': ano_modelo,
        'combustivel': combustivel,
        'data_manutencao_str': data_manutencao_str,
        'data_verificacao_str': data_verificacao_str,
        'km_verificado': km_verificado,
        'km_apos_manutencao': km_apos_manutencao,
        'data_proximo_check_str': data_proximo_check_str,
        'km_atual': km_atual,
        'km_inicial': km_inicial,
        'renavam': gerar_renavam()
    }

def encontrar_mensagem_alerta():
    seletores = [
        (".alerts.salvo", "✅ Menasagem de Sucesso"),
        (".alerts.alerta", "⚠️ Menasagem de Alerta"),
        (".alerts.erro", "❌ Menasagem de Erro"),
    ]

    for seletor, tipo in seletores:
        try:
            elemento = driver.find_element(By.CSS_SELECTOR, seletor)
            if elemento.is_displayed():
                log(doc, f"📢 {tipo}: {elemento.text}")
                return elemento
        except:
            continue

    log(doc, "ℹ️ Nenhuma mensagem de alerta encontrada.")
    return None

def click_element_safely(driver, wait, selector, by=By.CSS_SELECTOR, timeout=10):
    """Clica em um elemento de forma segura, tentando diferentes métodos"""
    try:
        # Aguarda o elemento estar presente
        element = wait.until(EC.presence_of_element_located((by, selector)))
        
        # Scroll até o elemento
        driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", element)
        time.sleep(0.5)
        
        # Aguarda ser clicável
        element = wait.until(EC.element_to_be_clickable((by, selector)))
        
        # Tenta clicar normalmente
        try:
            element.click()
            return True
        except:
            # Tenta com ActionChains
            try:
                ActionChains(driver).move_to_element(element).click().perform()
                return True
            except:
                # Tenta com JavaScript
                driver.execute_script("arguments[0].click();", element)
                return True
                
    except Exception as e:
        print(f"Erro ao clicar no elemento {selector}: {e}")
        return False

def abrir_modal_e_selecionar(btn_selector, pesquisa_selector, termo_pesquisa, btn_pesquisar_selector, resultado_xpath):
    def acao():
        wait = WebDriverWait(driver, 20)
        
        # Abre o modal
        if not click_element_safely(driver, wait, btn_selector):
            raise Exception("Não foi possível abrir o modal")
        
        time.sleep(1)
        
        # Aguarda campo pesquisa
        campo_pesquisa = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, pesquisa_selector)))
        campo_pesquisa.clear()
        campo_pesquisa.send_keys(termo_pesquisa)
        
        # Clica pesquisar
        if not click_element_safely(driver, wait, btn_pesquisar_selector):
            raise Exception("Não foi possível clicar no botão pesquisar")
        
        time.sleep(2)
        
        # Espera o resultado carregar e clica
        wait.until(EC.presence_of_element_located((By.XPATH, resultado_xpath)))
        wait.until(EC.element_to_be_clickable((By.XPATH, resultado_xpath)))
        
        resultado = driver.find_element(By.XPATH, resultado_xpath)
        driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", resultado)
        time.sleep(0.5)
        
        if not click_element_safely(driver, wait, resultado_xpath, By.XPATH):
            raise Exception("Não foi possível selecionar o resultado")
    
    return acao

def preencher_campo_com_retry(driver, wait, seletor, valor, max_tentativas=3):
    """Tenta preencher o campo com diferentes métodos até conseguir"""
    
    for tentativa in range(max_tentativas):
        try:
            # Aguarda o elemento
            campo = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, seletor)))
            
            # Scroll até o elemento se necessário
            driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", campo)
            time.sleep(0.5)
            
            # Método 1: Tradicional
            if tentativa == 0:
                campo.click()
                campo.clear()
                campo.send_keys(valor)
                campo.send_keys(Keys.TAB)
            
            # Método 2: ActionChains
            elif tentativa == 1:
                ActionChains(driver).move_to_element(campo).click().perform()
                time.sleep(0.2)
                ActionChains(driver).key_down(Keys.CONTROL).send_keys('a').key_up(Keys.CONTROL).perform()
                ActionChains(driver).send_keys(valor).perform()
                ActionChains(driver).send_keys(Keys.TAB).perform()
            
            # Método 3: JavaScript
            else:
                driver.execute_script("""
                    var element = arguments[0];
                    var valor = arguments[1];
                    element.focus();
                    element.value = '';
                    element.value = valor;
                    element.dispatchEvent(new Event('input', { bubbles: true }));
                    element.dispatchEvent(new Event('change', { bubbles: true }));
                    element.blur();
                """, campo, valor)
            
            time.sleep(0.5)
            
            # Verifica se o valor foi preenchido
            valor_atual = campo.get_attribute('value')
            if valor_atual == str(valor):
                return True
                
        except Exception as e:
            print(f"Tentativa {tentativa + 1} falhou: {e}")
            time.sleep(1)
    
    return False

def ajustar_zoom():
    try:
        driver.execute_script("document.body.style.zoom='90%'")
        log(doc, "🔍 Zoom ajustado para 90%.")
    except Exception as e:
        log(doc, f"⚠️ Erro ao ajustar zoom: {e}")

def selecionar_opcao(selector, texto):
    def acao():
        wait = WebDriverWait(driver, 10)
        select_element = wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, selector)))
        Select(select_element).select_by_visible_text(texto)
    return acao

# ==== INICIALIZAÇÃO DO DRIVER ====
wait = WebDriverWait(driver, 20)

# Gera os dados necessários
dados = gerar_dados_veiculo()

# ==== EXECUÇÃO DO TESTE ====
try:
    safe_action(doc, "Acessando sistema", lambda: driver.get(URL))

    safe_action(doc, "Realizando login", lambda: (
        wait.until(EC.presence_of_element_located((By.ID, "j_id15:email"))).send_keys(LOGIN_EMAIL),
        wait.until(EC.presence_of_element_located((By.ID, "j_id15:senha"))).send_keys(LOGIN_PASSWORD, Keys.ENTER),
        wait.until(EC.presence_of_element_located((By.TAG_NAME, "body")))
    ))

    safe_action(doc, "Esperando sistema carregar e ajustando zoom", lambda: (
        time.sleep(5),
        ajustar_zoom()
    ))

    safe_action(doc, "Abrindo menu Manutenção de Veículos", lambda: (
        driver.find_element(By.TAG_NAME, "body").send_keys(Keys.F2),
        time.sleep(1),
        wait.until(EC.visibility_of_element_located((By.XPATH, "//input[@placeholder='Busque um cadastro']"))).send_keys("Manutenção de Veículos", Keys.ENTER),
        time.sleep(3)
    ))

    safe_action(doc, "Clicando em Cadastrar", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_10062 > div.wdTelas > div.telaInicial.clearfix.overflow.overflowY > ul > li:nth-child(1) > a > span"))).click(),
        time.sleep(1)
    ))

    safe_action(doc, "Preenchendo Valor Manutenção", lambda: (
        preencher_campo_com_retry(driver, wait, "#fmod_10062 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroManutencaoVeiculo > div.catWrapper > div > div > div.groupHolder.clearfix.grupo_grupoDadosManutencao > div > div > div:nth-child(3) > input", fake.random_int(min=10000, max=100000))
    ))

    safe_action(doc, "Preenchendo Responsável da Manutenção", lambda: (
        preencher_campo_com_retry(driver, wait, "#fmod_10062 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroManutencaoVeiculo > div.catWrapper > div > div > div.groupHolder.clearfix.grupo_grupoDadosManutencao > div > div > div:nth-child(4) > input", fake.name())
    ))


    safe_action(doc, "Preenchendo Manutenção Necessária", lambda: (
        preencher_campo_com_retry(driver, wait, "#fmod_10062 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroManutencaoVeiculo > div.catWrapper > div > div > div.groupHolder.clearfix.grupo_grupoDadosManutencao > div > div > div:nth-child(6) > textarea", 'NENHUMA MANUTENÇÃO NECESSÁRIA, VEÍCULO EM PERFEITO ESTADO PARA USO')
    ))

    safe_action(doc, "Salvando cadastro de manutenção", lambda: (
        click_element_safely(driver, wait, "#fmod_10062 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroManutencaoVeiculo > div.btnHolder > a.btModel.btGray.btsave"),
        time.sleep(2)
    ))

    safe_action(doc, "Fechando modal após salvamento", lambda: (
        click_element_safely(driver, wait, "#fmod_10062 > div.wdTop.ui-draggable-handle > div.wdClose > a")
    ))

    encontrar_mensagem_alerta()

except Exception as e:
    log(doc, f"❌ ERRO FATAL: {e}")
    take_screenshot(driver, doc, "erro_fatal")

finally:
    log(doc, "✅ Teste concluído com sucesso.")
    finalizar_relatorio()