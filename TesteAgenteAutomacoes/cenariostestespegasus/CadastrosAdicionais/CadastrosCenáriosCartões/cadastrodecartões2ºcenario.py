import sys
import os

# Adiciona a raiz do projeto ao sys.path
sys.path.append(
    os.path.abspath(
        os.path.join(os.path.dirname(__file__), "../../..")
    )
)


from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service
from webdriver_manager.chrome import ChromeDriverManager
from docx import Document
from datetime import datetime
import subprocess
import time
import os
import sys

sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), '..', '..')))
from utils.actions import log, take_screenshot, safe_action, encontrar_mensagem_alerta, ajustar_zoom

URL = "http://localhost:8080/gs/index.xhtml"
LOGIN_EMAIL = "joaoeduardo.gold@outlook.com"
LOGIN_PASSWORD = "071999gs"

screenshot_registradas = set()
def registrar_screenshot_unico(nome, driver, doc, descricao=None):
    if nome not in screenshot_registradas:
        if descricao:
            log(doc, f"📸 {descricao}")
        take_screenshot(driver, doc, nome)
        screenshot_registradas.add(nome)

def main():
    doc = Document()
    doc.add_heading("RELATÓRIO DO TESTE", 0)
    doc.add_paragraph("Cadastro de Cartões – Cenário 2: Preenchimento completo e cancelamento.")
    doc.add_paragraph(f"🗕️ Data do teste: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")

    chrome_options = Options()
    chrome_options.add_argument("--start-maximized")
    driver = webdriver.Chrome(service=Service(ChromeDriverManager().install()), options=chrome_options)
    wait = WebDriverWait(driver, 10)

    def finalizar_relatorio():
        doc_name = f"relatorio_cartoes_cenario_2_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        doc.save(doc_name)
        log(doc, f"📄 Relatório salvo como: {doc_name}")
        try:
            subprocess.run(["start", "winword", doc_name], shell=True)
        except Exception as e:
            log(doc, f"Erro ao abrir o Word: {e}")
        driver.quit()

    def login():
        wait.until(EC.presence_of_element_located((By.ID, "j_id15:email"))).send_keys(LOGIN_EMAIL)
        wait.until(EC.presence_of_element_located((By.ID, "j_id15:senha"))).send_keys(LOGIN_PASSWORD, Keys.ENTER)
        wait.until(EC.presence_of_element_located((By.TAG_NAME, "body")))
        time.sleep(2)

    def abrir_menu():
        driver.find_element(By.TAG_NAME, "body").send_keys(Keys.F2)
        campo = wait.until(EC.presence_of_element_located((By.XPATH, "//input[@placeholder='Busque um cadastro']")))
        campo.send_keys("Cartão")
        time.sleep(1)
        wait.until(EC.element_to_be_clickable((By.XPATH, "/html/body/div[17]/ul/li[6]/a"))).click()
        time.sleep(2)

    def acessar_formulario():
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR,
            "#fmod_10011 > div.wdTelas > div.telaInicial.clearfix.overflow.overflowY > ul > li:nth-child(1) > a > span"))).click()
        time.sleep(2)

    def selecionar_bandeira():
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR,
            "#fmod_10011 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div > div > div > div:nth-child(4) > div > a"))).click()
        tr = wait.until(EC.presence_of_element_located((By.XPATH,
            "//tr[td[contains(text(), 'MASTERCARD')]]")))
        driver.execute_script("arguments[0].scrollIntoView(true);", tr)
        time.sleep(0.5)
        tr.click()
        log(doc, "✅ Bandeira selecionada.")
        registrar_screenshot_unico("bandeira_selecionada", driver, doc)

    def preencher_campos():
        log(doc, "🔄 Preenchendo campos obrigatórios.")
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR,
            "#fmod_10011 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div > div > div > div:nth-child(2) > input"))).send_keys("1111222233334444")
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR,
            "#fmod_10011 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div > div > div > div:nth-child(3) > input"))).send_keys("João Teste")
        safe_action(doc, "Selecionando bandeira", selecionar_bandeira, driver, wait)
        wait.until(EC.visibility_of_element_located((By.XPATH,
            "//input[contains(@class, 'hasDatepicker') and contains(@class, 'mandatory')]"))).send_keys("01/01/2025")
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR,
            "#fmod_10011 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div > div > div > div:nth-child(7) > input"))).send_keys("3")
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR,
            "#fmod_10011 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div > div > div > div:nth-child(8) > input"))).send_keys("9999")
        log(doc, "✅ Campos preenchidos com sucesso.")
        registrar_screenshot_unico("campos_preenchidos", driver, doc)

    def cancelar():
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR,
            "#fmod_10011 > div.wdTelas > div.telaCadastro.clearfix > div.btnHolder > a.btModel.btGray.btcancel"))).click()
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR,
            "#BtYes"))).click()
        log(doc, "❌ Cadastro cancelado pelo usuário.")
        registrar_screenshot_unico("cadastro_cancelado", driver, doc)

    def fechar_modal():
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR,
            "#fmod_10011 > div.wdTop.ui-draggable-handle > div.wdClose > a"))).click()
        time.sleep(1)

    # FLUXO DO TESTE
    if not safe_action(doc, "Acessando o sistema", lambda: driver.get(URL), driver, wait)[0]: finalizar_relatorio(); return
    registrar_screenshot_unico("url_acessada", driver, doc)

    if not safe_action(doc, "Realizando login", login, driver, wait)[0]: finalizar_relatorio(); return
    registrar_screenshot_unico("login_concluido", driver, doc)

    if not safe_action(doc, "Ajustando zoom", lambda: ajustar_zoom(driver), driver, wait)[0]: finalizar_relatorio(); return
    if not safe_action(doc, "Abrindo menu Cartão", abrir_menu, driver, wait)[0]: finalizar_relatorio(); return
    registrar_screenshot_unico("menu_aberto", driver, doc)

    if not safe_action(doc, "Acessando formulário", acessar_formulario, driver, wait)[0]: finalizar_relatorio(); return
    registrar_screenshot_unico("formulario_aberto", driver, doc)

    if not safe_action(doc, "Preenchendo campos", preencher_campos, driver, wait)[0]: finalizar_relatorio(); return
    if not safe_action(doc, "Cancelando cadastro", cancelar, driver, wait)[0]: finalizar_relatorio(); return

    if not safe_action(doc, "Fechando formulário", fechar_modal, driver, wait)[0]: finalizar_relatorio(); return
    registrar_screenshot_unico("formulario_fechado", driver, doc)

    log(doc, "✅ Teste do cenário 2 concluído com sucesso (cancelamento).")
    finalizar_relatorio()

if __name__ == "__main__":
    main()
