# Refatorado no modelo de cadastrodepessoas1ºcenario.py

import sys
import os

# Adiciona a raiz do projeto ao sys.path
sys.path.append(
    os.path.abspath(
        os.path.join(os.path.dirname(__file__), "../../..")
    )
)


from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support.ui import Select, WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service
from selenium.webdriver import ActionChains
from webdriver_manager.chrome import ChromeDriverManager
from docx import Document
from docx.shared import Inches
from faker import Faker
from faker.providers import BaseProvider
from validate_docbr import CPF
from datetime import datetime
import subprocess
import os
import time
import random
import string

# ==== PROVIDERS CUSTOMIZADOS ====
class BrasilProvider(BaseProvider):
    def rg(self):
        numeros = [str(random.randint(0, 9)) for _ in range(8)]
        return ''.join(numeros) + '-' + str(random.randint(0, 9))

fake = Faker("pt_BR")
fake.add_provider(BrasilProvider)

# ==== CONFIGURAÇÕES ====
URL = "http://localhost:8080/gs/index.xhtml"
LOGIN_EMAIL = "joaoeduardo.gold@outlook.com"
LOGIN_PASSWORD = "071999gs"

# ==== DOCUMENTO ====
doc = Document()
doc.add_heading("RELATÓRIO DO TESTE", 0)
doc.add_paragraph("Cadastro de Cobrador – Cenário 1: Preenchimento completo e salvamento.")
doc.add_paragraph(f"Data do teste: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")

screenshot_registradas = set()

# ==== FUNÇÕES UTILITÁRIAS ====
def log(doc, msg):
    print(msg)
    doc.add_paragraph(msg)

def scroll_to_element_and_click(element):
    """Rola até o elemento e clica usando JavaScript"""
    driver.execute_script("arguments[0].scrollIntoView({behavior: 'smooth', block: 'center'});", element)
    time.sleep(1)
    driver.execute_script("arguments[0].click();", element)



vencimento_cnh = fake.date_between(start_date='today', end_date='+10y')
vencimento_cnh_str = vencimento_cnh.strftime('%d/%m/%Y')


# Gera uma data de admissão entre 10 anos atrás e hoje
data_admissao = fake.date_between(start_date='-10y', end_date='today')
data_admissao_str = data_admissao.strftime('%d/%m/%Y')





def abrir_modal_e_selecionar(btn_selector, pesquisa_selector, termo_pesquisa, btn_pesquisar_selector, resultado_xpath):
    """Abre modal e seleciona um item"""
    def acao():
        # Abre o modal
        open_lov = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, btn_selector)))
        open_lov.click()
        time.sleep(3)

        # Aguarda campo pesquisa
        campo_pesquisa = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, pesquisa_selector)))
        campo_pesquisa.clear()
        campo_pesquisa.send_keys(termo_pesquisa)

        # Clica pesquisar
        pesquisar = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, btn_pesquisar_selector)))
        pesquisar.click()
        time.sleep(3)
        pesquisar.click()

        # Espera o resultado e clica
        wait.until(EC.element_to_be_clickable((By.XPATH, resultado_xpath)))
        resultado = driver.find_element(By.XPATH, resultado_xpath)
        driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", resultado)
        time.sleep(0.2)
        resultado.click()

    return acao

def take_screenshot(driver, doc, nome):
    if nome not in screenshot_registradas:
        path = f"screenshots/{nome}.png"
        os.makedirs("screenshots", exist_ok=True)
        driver.save_screenshot(path)
        doc.add_paragraph(f"Screenshot: {nome}")
        doc.add_picture(path, width=Inches(5.5))
        screenshot_registradas.add(nome)

def safe_action(doc, descricao, func):
    try:
        log(doc, f"🔄 {descricao}...")
        func()
        log(doc, f"✅ {descricao} realizada com sucesso.")
        take_screenshot(driver, doc, descricao.lower().replace(" ", "_"))
    except Exception as e:
        log(doc, f"❌ Erro ao {descricao.lower()}: {e}")
        take_screenshot(driver, doc, f"erro_{descricao.lower().replace(' ', '_')}")

def finalizar_relatorio():
    nome_arquivo = f"relatorio_cobrador_cenario_1_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc.save(nome_arquivo)
    log(doc, f"📄 Relatório salvo como: {nome_arquivo}")
    subprocess.run(["start", "winword", nome_arquivo], shell=True)
    driver.quit()

def encontrar_mensagem_alerta():
    seletores = [
        (".alerts.salvo", "✅ Menasagem de Sucesso"),
        (".alerts.alerta", "⚠️ Menasagem de Alerta"),
        (".alerts.erro", "❌ Menasagem de Erro"),
    ]
    for seletor, tipo in seletores:
        try:
            elemento = driver.find_element(By.CSS_SELECTOR, seletor)
            if elemento.is_displayed():
                log(doc, f"📢 {tipo}: {elemento.text}")
                return elemento
        except:
            continue
    log(doc, "ℹ️ Nenhuma mensagem de alerta encontrada.")
    return None

def preencher_campo_data(selector, valor):
    def acao():
        campo = wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, selector)))
        campo.click()
        campo.send_keys(valor)
        time.sleep(0.2)
    return acao

def selecionar_opcao(selector, texto):
    def acao():
        select_element = wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, selector)))
        Select(select_element).select_by_visible_text(texto)
    return acao

# ==== GERAÇÃO DE DADOS ====
nome_completo = "COBRADOR TESTE SELENIUM AUTOMATIZADO"
cpf = CPF().generate()
rg = fake.rg()
data_expedicao = fake.date_of_birth(minimum_age=18, maximum_age=60).strftime("%d/%m/%Y")
data_nascimento = fake.date_of_birth(minimum_age=18, maximum_age=60).strftime("%d/%m/%Y")
email = fake.email()
telefone1 = fake.phone_number()
telefone2 = fake.phone_number()
telefone3 = fake.phone_number()
cidade = fake.city()
pais = fake.country()
nome_pai = fake.first_name()
nome_mae = fake.first_name()
profissao = fake.job()
data_admissao = fake.date_between(start_date='-10y', end_date='today').strftime('%d/%m/%Y')
carteira_trabalho = ''.join(str(random.randint(0, 9)) for _ in range(7))
pis = ''.join(str(random.randint(0, 9)) for _ in range(11))

# ==== DRIVER ====
options = Options()
options.add_argument("--start-maximized")
driver = webdriver.Chrome(service=Service(ChromeDriverManager().install()), options=options)
wait = WebDriverWait(driver, 20)

# ==== EXECUÇÃO ====
try:
    safe_action(doc, "Acessando sistema", lambda: driver.get(URL))

    safe_action(doc, "Realizando login", lambda: (
        wait.until(EC.presence_of_element_located((By.ID, "j_id15:email"))).send_keys(LOGIN_EMAIL),
        wait.until(EC.presence_of_element_located((By.ID, "j_id15:senha"))).send_keys(LOGIN_PASSWORD, Keys.ENTER),
        wait.until(EC.presence_of_element_located((By.TAG_NAME, "body")))
    ))

    safe_action(doc, "Aguardando carregamento e ajustando zoom", lambda: (
        time.sleep(5),
        driver.execute_script("document.body.style.zoom='90%'")
    ))

    safe_action(doc, "Abrindo menu Cobrador", lambda: (
        driver.find_element(By.TAG_NAME, "body").send_keys(Keys.F2),
        time.sleep(1),
        wait.until(EC.visibility_of_element_located((By.XPATH, "//input[@placeholder='Busque um cadastro']"))).send_keys("Cobrador", Keys.ENTER),
        time.sleep(2)
    ))

    safe_action(doc, "Clicando em Cadastrar", lambda: (
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "#fmod_10049 > div.wdTelas > div.telaInicial.clearfix.overflow.overflowY > ul > li:nth-child(1) > a > span"))).click(),
        time.sleep(3)
    ))

    safe_action(doc, "Abrindo LOV de Pessoa", lambda: (
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "#fmod_10049 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div > div:nth-child(1) > div > div:nth-child(2) > div > a"))).click(),
        time.sleep(1)
    ))

    safe_action(doc, "Criando nova Pessoa", lambda: (
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(3) > a"))).click(),
        time.sleep(1)
    ))

    safe_action(doc, "Preenchendo Nome Completo", lambda: (
        wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, "#cg_1 > div.wdTelas > div > div.catWrapper > div > div.cat_dadosPessoais.categoriaHolder > div > div > div:nth-child(2) > div:nth-child(2) > input"))).send_keys(nome_completo)
    ))

    safe_action(doc, "Selecionando Tipo de Pessoa", selecionar_opcao(
        "#cg_1 > div.wdTelas > div > div.catWrapper > div > div.cat_dadosPessoais.categoriaHolder > div > div > div:nth-child(2) > div:nth-child(3) > select",
        "Física"
    ))

    safe_action(doc, "Selecionando Tipo de Documento", selecionar_opcao(
        "#cg_1 > div.wdTelas > div > div.catWrapper > div > div.cat_dadosPessoais.categoriaHolder > div > div > div:nth-child(2) > div:nth-child(4) > select",
        "Carteira de Identidade Classista"
    ))

    safe_action(doc, "Preenchendo Número do Documento", lambda: (
        wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, "#cg_1 > div.wdTelas > div > div.catWrapper > div > div.cat_dadosPessoais.categoriaHolder > div > div > div.formRow.divPessoaFISICA > div:nth-child(1) > input"))).send_keys(rg)
    ))

    safe_action(doc, "Preenchendo Data de Expedição", preencher_campo_data("input.dataExpedicao", data_expedicao))

    safe_action(doc, "Preenchendo CPF", lambda: (
        wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, "#cg_1 > div.wdTelas > div > div.catWrapper > div > div.cat_dadosPessoais.categoriaHolder > div > div > div.formRow.divPessoaFISICA > div:nth-child(3) > input"))).click(),
        wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, "#cg_1 > div.wdTelas > div > div.catWrapper > div > div.cat_dadosPessoais.categoriaHolder > div > div > div.formRow.divPessoaFISICA > div:nth-child(3) > input"))).send_keys(cpf)

   ))

    safe_action(doc, "Acessando aba Dados Complementares", lambda: (
        wait.until(EC.element_to_be_clickable((By.LINK_TEXT, "Dados Complementares"))).click(),
        time.sleep(1)
    ))

    safe_action(doc, "Selecionando Estado Civil", selecionar_opcao(
        "#cg_1 > div.wdTelas > div > div.catWrapper > div > div.cat_dadosComplementares.categoriaHolder > div > div > div > div:nth-child(1) > select",
        "Solteiro"
    ))

    safe_action(doc, "Selecionando Sexo", selecionar_opcao(
        "#cg_1 > div.wdTelas > div > div.catWrapper > div > div.cat_dadosComplementares.categoriaHolder > div > div > div > div:nth-child(2) > select",
        "Feminino"
    ))

    safe_action(doc, "Preenchendo E-mail", lambda: (
        wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, "#cg_1 > div.wdTelas > div > div.catWrapper > div > div.cat_dadosComplementares.categoriaHolder > div > div > div > div:nth-child(9) > input"))).send_keys(email)
    ))

    safe_action(doc, "Preenchendo Data de Nascimento", preencher_campo_data("input.dataNascimento", data_nascimento))

    safe_action(doc, "Preenchendo Telefone 1", lambda: (
        wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, "#cg_1 > div.wdTelas > div > div.catWrapper > div > div.cat_dadosComplementares.categoriaHolder > div > div > div > div:nth-child(5) > input"))).send_keys(telefone1)
    ))

    safe_action(doc, "Preenchendo Telefone 2", lambda: (
        wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, "#cg_1 > div.wdTelas > div > div.catWrapper > div > div.cat_dadosComplementares.categoriaHolder > div > div > div > div:nth-child(6) > input"))).send_keys(telefone2)
    ))

    safe_action(doc, "Preenchendo Telefone 3", lambda: (
        wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, "#cg_1 > div.wdTelas > div > div.catWrapper > div > div.cat_dadosComplementares.categoriaHolder > div > div > div > div:nth-child(7) > input"))).send_keys(telefone3)
    ))

    safe_action(doc, "Preenchendo Cidade de Nascimento", lambda: (
        wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, "#cg_1 > div.wdTelas > div > div.catWrapper > div > div.cat_dadosComplementares.categoriaHolder > div > div > div > div:nth-child(11) > input"))).send_keys(cidade)
    ))

    safe_action(doc, "Preenchendo País de Nascimento", lambda: (
        wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, "#cg_1 > div.wdTelas > div > div.catWrapper > div > div.cat_dadosComplementares.categoriaHolder > div > div > div > div:nth-child(12) > input"))).send_keys(pais)
    ))

    safe_action(doc, "Preenchendo Nome do Pai", lambda: (
        wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, "#cg_1 > div.wdTelas > div > div.catWrapper > div > div.cat_dadosComplementares.categoriaHolder > div > div > div > div:nth-child(13) > input"))).send_keys(nome_pai)
    ))

    safe_action(doc, "Preenchendo Nome da Mãe", lambda: (
        wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, "#cg_1 > div.wdTelas > div > div.catWrapper > div > div.cat_dadosComplementares.categoriaHolder > div > div > div > div:nth-child(14) > input"))).send_keys(nome_mae)
    ))

    safe_action(doc, "Preenchendo Profissão", lambda: (
        wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, "#cg_1 > div.wdTelas > div > div.catWrapper > div > div.cat_dadosComplementares.categoriaHolder > div > div > div > div:nth-child(16) > input"))).send_keys(profissao)
    ))

    safe_action(doc, "Salvando cadastro", lambda: (
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "#cg_1 > div.wdTelas > div > div.btnHolder > a.btModel.btGray.btsave"))).click(),
        time.sleep(2)
    ))



    safe_action(doc, "Selecionando Coordenador", abrir_modal_e_selecionar(
        "#fmod_10049 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div > div:nth-child(1) > div > div:nth-child(6) > div > a",
        "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(2) > input.nomePesquisa",
        "COORDENADOR TESTE SELENIUM AUTOMATIZADO",
        "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(3) > a",
        "//td[contains(text(), 'COORDENADOR TESTE SELENIUM AUTOMATIZADO')]"
    ))


    safe_action(doc, "Selecionando Supervisor", abrir_modal_e_selecionar(
        "#fmod_10049 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div > div:nth-child(1) > div > div:nth-child(7) > div > a",
        "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(2) > input.nomePesquisa",
        "SUPERVISOR TESTE SELENIUM AUTOMATIZADO",
        "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(3) > a",
        "//td[contains(text(), 'SUPERVISOR TESTE SELENIUM AUTOMATIZADO')]"
    ))





    safe_action(doc, "Preenchendo a Data de Admissão ", lambda: (
        wait.until(EC.element_to_be_clickable((By.XPATH, "//input[@grupo='10074' and @ref='10066' and contains(@class, 'hasDatepicker')]"))).click(),
        wait.until(EC.element_to_be_clickable((By.XPATH, "//input[@grupo='10074' and @ref='10066' and contains(@class, 'hasDatepicker')]"))).send_keys(data_admissao_str),
        ))


    safe_action(doc, "Preenchendo Carteira de Trabalho", lambda: (
        wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, "#fmod_10049 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div > div:nth-child(1) > div > div:nth-child(4) > input"))).send_keys(carteira_trabalho)
    ))


    safe_action(doc, "Preenchendo PIS", lambda: (
        wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, "#fmod_10049 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div > div:nth-child(1) > div > div:nth-child(5) > input"))).send_keys(pis)
    ))

    safe_action(doc, "Preenchendo a Comissão em Porcentagem", lambda: (
        wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, "#fmod_10049 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div > div:nth-child(1) > div > div:nth-child(8) > input"))).send_keys(fake.random_int(min=10, max=500)
    )))


    safe_action(doc, "Preenchendo a Comissão em Reais", lambda: (
        wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, "#fmod_10049 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div > div:nth-child(1) > div > div:nth-child(9) > input"))).send_keys(fake.random_int(min=10, max=500)
    )))

    safe_action(doc, "Selecionando o Tipo Cobrador", selecionar_opcao(
        "#fmod_10049 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div > div:nth-child(1) > div > div:nth-child(10) > select",
        "Interno"
    ))


    safe_action(doc, "Selecionando o Tipo Contrato", selecionar_opcao(
        "#fmod_10049 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div > div:nth-child(1) > div > div:nth-child(12) > select",
        "Carteira Assinada"
    ))


    safe_action(doc, "Clicando no campo 'Código Referência' para o código ser gerado pelo sistema", lambda: (
        wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, "#fmod_10049 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div > div:nth-child(1) > div > div:nth-child(13) > input"))).click()
    ))



    safe_action(doc, "Selecionando Área do Cobrador", abrir_modal_e_selecionar(
        "#fmod_10049 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div > div:nth-child(1) > div > div:nth-child(14) > div > div > a",
        "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(2) > input",
        "0CEN",
        "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(3) > a",
        "//td[contains(text(), '0CEN')]"
    ))


    safe_action(doc, "Salvando cadastro", lambda: (
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "#fmod_10049 > div.wdTelas > div.telaCadastro.clearfix > div.btnHolder > a.btModel.btGray.btsave"))).click()
    ))

    encontrar_mensagem_alerta()

    safe_action(doc, "Fechando modal", lambda: (
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "#fmod_10049 > div.wdTop.ui-draggable-handle > div.wdClose > a"))).click()
    ))
    time.sleep(4)

    encontrar_mensagem_alerta()

except Exception as e:
    log(doc, f"❌ ERRO FATAL: {e}")
    take_screenshot(driver, doc, "erro_fatal")

finally:
    log(doc, "✅ Teste concluído com sucesso.")
    finalizar_relatorio()
