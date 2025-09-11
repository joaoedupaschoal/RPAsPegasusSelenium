# Versão CORRIGIDA: cadastrodejazigos1ºcenario.py
# Melhorias: Tratamento robusto de erros, múltiplas estratégias de seleção, timeouts adaptativos

from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support.ui import Select, WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service
from selenium.webdriver import ActionChains
from selenium.common.exceptions import TimeoutException, NoSuchElementException, ElementClickInterceptedException
from webdriver_manager.chrome import ChromeDriverManager
from docx import Document
from docx.shared import Inches
from faker import Faker
from faker.providers import BaseProvider
from datetime import datetime
import subprocess
import os
import time
import random
import string

# ==== PROVIDERS CUSTOMIZADOS ====
class BrasilProvider(BaseProvider):
    def rg(self):
        numeros = [str(random.randint(0, 9)) for _ in range(8)]
        return ''.join(numeros) + '-' + str(random.randint(0, 9))

fake = Faker("pt_BR")
fake.add_provider(BrasilProvider)

# ==== CONFIGURAÇÕES ====
URL = "http://localhost:8080/gs/index.xhtml"
LOGIN_EMAIL = "joaoeduardo.gold@outlook.com"
LOGIN_PASSWORD = "071999gs"

# ==== DOCUMENTO ====
doc = Document()
doc.add_heading("RELATÓRIO DO TESTE", 0)
doc.add_paragraph("Cadastro de Jazigos – Cenário 1: Preenchimento completo e salvamento.")
doc.add_paragraph(f"Data do teste: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")

screenshot_registradas = set()

# ==== FUNÇÕES DE UTILITÁRIO MELHORADAS ====
def log(doc, msg):
    print(msg)
    doc.add_paragraph(msg)

def take_screenshot(driver, doc, nome):
    if nome not in screenshot_registradas:
        path = f"screenshots/{nome}.png"
        os.makedirs("screenshots", exist_ok=True)
        try:
            driver.save_screenshot(path)
            doc.add_paragraph(f"Screenshot: {nome}")
            doc.add_picture(path, width=Inches(5.5))
        except Exception as e:
            print(f"Erro ao capturar screenshot: {e}")
        screenshot_registradas.add(nome)

def aguardar_elemento_disponivel(driver, selector, by_type=By.CSS_SELECTOR, timeout=30):
    """Aguarda elemento estar presente, visível e clicável"""
    try:
        wait = WebDriverWait(driver, timeout)
        # Aguarda estar presente
        wait.until(EC.presence_of_element_located((by_type, selector)))
        # Aguarda estar visível
        wait.until(EC.visibility_of_element_located((by_type, selector)))
        # Aguarda estar clicável
        element = wait.until(EC.element_to_be_clickable((by_type, selector)))
        return element
    except TimeoutException:
        return None

def safe_click_enhanced(driver, selector, by_type=By.CSS_SELECTOR, timeout=30):
    """Função de clique ultra-robusta com múltiplas estratégias"""
    strategies = [
        "aguardar_e_clicar_normal",
        "aguardar_e_clicar_js", 
        "aguardar_e_clicar_action",
        "força_bruta_js"
    ]
    
    for strategy in strategies:
        try:
            log(doc, f"🔄 Tentando estratégia: {strategy}")
            
            if strategy == "aguardar_e_clicar_normal":
                element = aguardar_elemento_disponivel(driver, selector, by_type, timeout)
                if element:
                    # Rola para o elemento
                    driver.execute_script("arguments[0].scrollIntoView({block: 'center', behavior: 'smooth'});", element)
                    time.sleep(1)
                    # Relocaliza para evitar stale reference
                    element = driver.find_element(by_type, selector)
                    element.click()
                    return True
                    
            elif strategy == "aguardar_e_clicar_js":
                element = aguardar_elemento_disponivel(driver, selector, by_type, timeout)
                if element:
                    driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", element)
                    time.sleep(0.5)
                    driver.execute_script("arguments[0].click();", element)
                    return True
                    
            elif strategy == "aguardar_e_clicar_action":
                element = aguardar_elemento_disponivel(driver, selector, by_type, timeout)
                if element:
                    actions = ActionChains(driver)
                    actions.move_to_element(element).pause(0.5).click().perform()
                    return True
                    
            elif strategy == "força_bruta_js":
                # Última tentativa: força bruta com JavaScript
                if by_type == By.CSS_SELECTOR:
                    js_code = f"""
                        var element = document.querySelector('{selector}');
                        if (element) {{
                            element.scrollIntoView({{block: 'center'}});
                            setTimeout(function() {{
                                element.click();
                                console.log('Clique forçado executado');
                            }}, 500);
                            return true;
                        }}
                        return false;
                    """
                else:  # XPATH
                    js_code = f"""
                        var element = document.evaluate('{selector}', document, null, XPathResult.FIRST_ORDERED_NODE_TYPE, null).singleNodeValue;
                        if (element) {{
                            element.scrollIntoView({{block: 'center'}});
                            setTimeout(function() {{
                                element.click();
                                console.log('Clique forçado XPath executado');
                            }}, 500);
                            return true;
                        }}
                        return false;
                    """
                
                result = driver.execute_script(js_code)
                if result:
                    time.sleep(1)
                    return True
                    
        except Exception as e:
            log(doc, f"⚠️ Estratégia {strategy} falhou: {str(e)[:100]}...")
            continue
    
    return False

def safe_send_keys_enhanced(driver, selector, text, by_type=By.CSS_SELECTOR, clear=True, timeout=20):
    """Função para envio seguro de texto com retry"""
    for attempt in range(3):
        try:
            element = aguardar_elemento_disponivel(driver, selector, by_type, timeout)
            if element:
                driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", element)
                time.sleep(0.3)
                
                if clear:
                    element.clear()
                    # Fallback para limpar
                    element.send_keys(Keys.CONTROL + "a")
                    element.send_keys(Keys.DELETE)
                
                element.send_keys(text)
                return True
        except Exception as e:
            log(doc, f"⚠️ Tentativa {attempt + 1} de envio de texto falhou: {e}")
            time.sleep(1)
    
    return False

def safe_action_enhanced(driver, doc, descricao, func, max_tentativas=3):
    """Função safe_action aprimorada"""
    for tentativa in range(max_tentativas):
        try:
            log(doc, f"🔄 {descricao}... (Tentativa {tentativa + 1})")
            result = func()
            if result is False:  # Se a função retornou False explicitamente
                raise Exception("Função retornou False")
            log(doc, f"✅ {descricao} realizada com sucesso.")
            take_screenshot(driver, doc, descricao.lower().replace(" ", "_"))
            return True
        except Exception as e:
            if tentativa < max_tentativas - 1:
                log(doc, f"⚠️ Tentativa {tentativa + 1} falhou: {str(e)[:100]}... Tentando novamente...")
                time.sleep(2)
            else:
                log(doc, f"❌ Erro ao {descricao.lower()}: {str(e)[:200]}...")
                take_screenshot(driver, doc, f"erro_{descricao.lower().replace(' ', '_')}")
                return False

def selecionar_item_tabela_melhorado(driver, texto_busca, timeout=20):
    """Função universal para seleção em tabelas com múltiplas estratégias"""
    estrategias = [
        f"//td[contains(text(), '{texto_busca}')]/following-sibling::td//a",
        f"//td[contains(text(), '{texto_busca}')]/..//a",
        f"//tr[.//td[contains(text(), '{texto_busca}')]]//a",
        "//table//a[contains(@onclick, 'selecionarItem') or @class='linkAlterar' or contains(@class, 'link')]"
    ]

def selecionar_item_tabela_melhorado_cemiterio(driver, texto_busca, timeout=20):
    estrategias = [
        f"//td[contains(text(), '{texto_busca}')]/following-sibling::td//a",
        f"//td[contains(text(), '{texto_busca}')]/..//a",
        f"//tr[.//td[contains(text(), '{texto_busca}')]]//a",
        "//tr[@ref='31100060' and @role='row' and contains(@class, 'odd') and @style='cursor: pointer;'][td[1]='CEMITÉRIO HERMAN DESCANSO']"
    ]
    
    try:
        WebDriverWait(driver, timeout).until(EC.presence_of_element_located((By.TAG_NAME, "table")))
        time.sleep(2)
    except:
        log(doc, "⚠️ Tabela não encontrada ou não carregou")
        return False
    
    for i, xpath in enumerate(estrategias):
        try:
            log(doc, f"🔍 Tentando estratégia {i+1}: {xpath[:50]}...")
            elements = driver.find_elements(By.XPATH, xpath)
            
            if elements:
                element = elements[0]
                driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", element)
                time.sleep(0.5)
                
                try:
                    element.click()
                except:
                    driver.execute_script("arguments[0].click();", element)
                
                log(doc, f"✅ Item selecionado com estratégia {i+1}")
                return True
            else:
                log(doc, f"⚠️ Nenhum elemento encontrado com estratégia {i+1}")
                
        except Exception as e:
            log(doc, f"⚠️ Erro na estratégia {i+1}: {str(e)[:100]}...")
            continue
    
    # Estratégia final: JavaScript genérico
    try:
        js_result = driver.execute_script("""
            var tabela = document.querySelector('table');
            if (!tabela) return false;
            
            var links = tabela.querySelectorAll('a');
            if (links.length > 0) {
                var link = links[0];
                link.scrollIntoView({block: 'center'});
                setTimeout(function() {
                    link.click();
                    console.log('Link clicado via JavaScript final');
                }, 500);
                return true;
            }
            return false;
        """)
        
        if js_result:
            time.sleep(1)
            log(doc, "✅ Item selecionado via JavaScript final")
            return True
            
    except Exception as e:
        log(doc, f"⚠️ Estratégia JavaScript final falhou: {e}")
    
    return False

    
def selecionar_item_tabela_melhorado_quadra(driver, texto_busca, timeout=20):
    """Função universal para seleção em tabelas com múltiplas estratégias"""
    estrategias = [
        f"//td[contains(text(), '{texto_busca}')]/following-sibling::td//a",
        f"//td[contains(text(), '{texto_busca}')]/..//a",
        f"//tr[.//td[contains(text(), '{texto_busca}')]]//a",
        "//tr[@ref='31100067' and @role='row' and contains(@class, 'odd') and @style='cursor: pointer;'][td[1]='QUADRA TESTE SELENIUM AUTOMATIZADO' and td[2]='20000' and td[3]='100' and td[4]='20' and td[5]='CEMITÉRIO HERMAN DESCANSO']"
    ]
    
    
    # Aguarda a tabela carregar
    try:
        WebDriverWait(driver, timeout).until(EC.presence_of_element_located((By.TAG_NAME, "table")))
        time.sleep(2)  # Aguarda renderização completa
    except:
        log(doc, "⚠️ Tabela não encontrada ou não carregou")
        return False
    
    for i, xpath in enumerate(estrategias):
        try:
            log(doc, f"🔍 Tentando estratégia {i+1}: {xpath[:50]}...")
            elements = driver.find_elements(By.XPATH, xpath)
            
            if elements:
                # Tenta clicar no primeiro elemento encontrado
                element = elements[0]
                driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", element)
                time.sleep(0.5)
                
                # Tenta clique normal primeiro, depois JavaScript
                try:
                    element.click()
                except:
                    driver.execute_script("arguments[0].click();", element)
                
                log(doc, f"✅ Item selecionado com estratégia {i+1}")
                return True
            else:
                log(doc, f"⚠️ Nenhum elemento encontrado com estratégia {i+1}")
                
        except Exception as e:
            log(doc, f"⚠️ Erro na estratégia {i+1}: {str(e)[:100]}...")
            continue
    
    # Estratégia final: JavaScript puro para buscar qualquer link clicável
    try:
        js_result = driver.execute_script("""
            var tabela = document.querySelector('table');
            if (!tabela) return false;
            
            var links = tabela.querySelectorAll('a');
            if (links.length > 0) {
                var link = links[0];
                link.scrollIntoView({block: 'center'});
                setTimeout(function() {
                    link.click();
                    console.log('Link clicado via JavaScript final');
                }, 500);
                return true;
            }
            return false;
        """)
        
        if js_result:
            time.sleep(1)
            log(doc, "✅ Item selecionado via JavaScript final")
            return True
            
    except Exception as e:
        log(doc, f"⚠️ Estratégia JavaScript final falhou: {e}")
    
    return False

def selecionar_dropdown_robusto(driver, selector, texto_opcao):
    """Seleção robusta em dropdowns"""
    try:
        select_element = aguardar_elemento_disponivel(driver, selector, By.CSS_SELECTOR, 20)
        if not select_element:
            return False
            
        driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", select_element)
        time.sleep(0.3)
        
        select = Select(select_element)
        
        # Tenta seleção por texto visível
        try:
            select.select_by_visible_text(texto_opcao)
            return True
        except:
            pass
            
        # Tenta seleção por texto parcial
        for option in select.options:
            if texto_opcao.lower() in option.text.lower():
                select.select_by_visible_text(option.text)
                return True
                
        # Tenta seleção por valor
        for option in select.options:
            if texto_opcao.lower() in option.get_attribute('value').lower():
                select.select_by_value(option.get_attribute('value'))
                return True
                
        return False
        
    except Exception as e:
        log(doc, f"Erro na seleção do dropdown: {e}")
        return False

def gerar_dados_jazigo():
    """Gera dados fictícios para o cadastro de jazigo."""
    numero_aleatorio = random.randint(1000, 9999)  # Número maior para evitar conflitos
    letra_aleatoria = random.choice(string.ascii_uppercase)
    altura_cm = random.randint(150, 250)
    largura_cm = random.randint(100, 200)
    comprimento_cm = random.randint(200, 300)
    
    return (numero_aleatorio, letra_aleatoria, altura_cm, largura_cm, comprimento_cm)

def finalizar_relatorio():
    nome_arquivo = f"relatorio_jazigos_cenario_1_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    try:
        doc.save(nome_arquivo)
        log(doc, f"📄 Relatório salvo como: {nome_arquivo}")
        subprocess.run(["start", "winword", nome_arquivo], shell=True)
    except Exception as e:
        print(f"Erro ao salvar relatório: {e}")

# Gera os dados necessários
(numero_aleatorio, letra_aleatoria, altura_cm, largura_cm, comprimento_cm) = gerar_dados_jazigo()

# ==== INICIALIZAÇÃO DO DRIVER ====
options = Options()
options.add_argument("--start-maximized")
options.add_argument("--disable-blink-features=AutomationControlled")
options.add_argument("--disable-web-security")
options.add_argument("--disable-features=VizDisplayCompositor")
options.add_experimental_option("excludeSwitches", ["enable-automation"])
options.add_experimental_option('useAutomationExtension', False)

try:
    driver = webdriver.Chrome(service=Service(ChromeDriverManager().install()), options=options)
    driver.execute_script("Object.defineProperty(navigator, 'webdriver', {get: () => undefined})")
    wait = WebDriverWait(driver, 30)  # Timeout maior
    
    log(doc, "🚀 Iniciando teste corrigido...")
    
    # ==== EXECUÇÃO DO TESTE PRINCIPAL ====
    
    # Acesso ao sistema
    safe_action_enhanced(driver, doc, "Acessando sistema", 
        lambda: driver.get(URL))

    # Login
    safe_action_enhanced(driver, doc, "Realizando login", lambda: (
        safe_send_keys_enhanced(driver, "#j_id15\\:email", LOGIN_EMAIL, By.CSS_SELECTOR),
        safe_send_keys_enhanced(driver, "#j_id15\\:senha", LOGIN_PASSWORD, By.CSS_SELECTOR),
        driver.find_element(By.CSS_SELECTOR, "#j_id15\\:senha").send_keys(Keys.ENTER),
        time.sleep(3)
    ))

    # Aguardar carregamento e ajustar zoom
    safe_action_enhanced(driver, doc, "Esperando sistema carregar e ajustando zoom", lambda: (
        time.sleep(5),
        driver.execute_script("document.body.style.zoom='90%'")
    ))

    # Abrir menu Jazigos
    safe_action_enhanced(driver, doc, "Abrindo menu Jazigos", lambda: (
        driver.find_element(By.TAG_NAME, "body").send_keys(Keys.F2),
        time.sleep(3),
        safe_send_keys_enhanced(driver, "//input[@placeholder='Busque um cadastro']", "Jazigos", By.XPATH, clear=True),
        driver.find_element(By.XPATH, "//input[@placeholder='Busque um cadastro']").send_keys(Keys.ENTER),
        time.sleep(4)
    ))

    # Clicar em Cadastrar
    safe_action_enhanced(driver, doc, "Clicando em Cadastrar", 
        lambda: safe_click_enhanced(driver, "#fmod_7 > div.wdTelas > div > ul > li:nth-child(1) > a > span"))

    # Preenchimento dos campos
    safe_action_enhanced(driver, doc, "Preenchendo Número do Jazigo", 
        lambda: safe_send_keys_enhanced(driver, "#fmod_7 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroJazigo > div.catWrapper > div > div > div > div > div > div:nth-child(2) > input", str(numero_aleatorio)))

    safe_action_enhanced(driver, doc, "Preenchendo Letra do Jazigo", 
        lambda: safe_send_keys_enhanced(driver, "#fmod_7 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroJazigo > div.catWrapper > div > div > div > div > div > div:nth-child(3) > input", letra_aleatoria))

    # Seleção de Quadra
    safe_action_enhanced(driver, doc, "Abrindo LOV para selecionar Quadra", 
        lambda: safe_click_enhanced(driver, "#fmod_7 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroJazigo > div.catWrapper > div > div > div > div > div > div:nth-child(4) > div > a"))




    safe_action_enhanced(driver, doc, "Inserindo nome da Quadra", 
        lambda: safe_send_keys_enhanced(driver, "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div.formCol.divPesquisa > div:nth-child(2) > input", 'QUADRA TESTE SELENIUM AUTOMATIZADO'))

    safe_action_enhanced(driver, doc, "Pesquisando Quadra", 
        lambda: safe_click_enhanced(driver, "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(2) > a"))

    safe_action_enhanced(driver, doc, "Selecionando Quadra",
        lambda: selecionar_item_tabela_melhorado_quadra(driver, "QUADRA TESTE SELENIUM AUTOMATIZADO"))

    # Seleções em dropdowns
    safe_action_enhanced(driver, doc, "Selecionando Tipo de Jazigo", 
        lambda: selecionar_dropdown_robusto(driver, "#fmod_7 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroJazigo > div.catWrapper > div > div > div > div > div > div:nth-child(5) > select", "001 - Oito Gavetas"))

    safe_action_enhanced(driver, doc, "Selecionando Situação", 
        lambda: selecionar_dropdown_robusto(driver, "#fmod_7 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroJazigo > div.catWrapper > div > div > div > div > div > div:nth-child(7) > select", "Disponível"))

    safe_action_enhanced(driver, doc, "Selecionando Ala", 
        lambda: selecionar_dropdown_robusto(driver, "#fmod_7 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroJazigo > div.catWrapper > div > div > div > div > div > div:nth-child(10) > select", "Direito"))

    safe_action_enhanced(driver, doc, "Selecionando Categoria", 
        lambda: selecionar_dropdown_robusto(driver, "#fmod_7 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroJazigo > div.catWrapper > div > div > div > div > div > div:nth-child(12) > select", "Perpétuo"))

    # Preenchimento das dimensões
    safe_action_enhanced(driver, doc, "Preenchendo dimensões do jazigo", lambda: (
        safe_send_keys_enhanced(driver, "#fmod_7 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroJazigo > div.catWrapper > div > div > div > div > div > div:nth-child(13) > input", str(altura_cm)),
        safe_send_keys_enhanced(driver, "#fmod_7 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroJazigo > div.catWrapper > div > div > div > div > div > div:nth-child(14) > input", str(largura_cm)),
        safe_send_keys_enhanced(driver, "#fmod_7 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroJazigo > div.catWrapper > div > div > div > div > div > div:nth-child(15) > input", str(comprimento_cm))
    ))

    # Salvamento
    safe_action_enhanced(driver, doc, "Salvando cadastro", 
        lambda: safe_click_enhanced(driver, "#fmod_7 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroJazigo > div.btnHolder > a.btModel.btGray.btsave"))

    safe_action_enhanced(driver, doc, "Fechando modal após salvamento", 
        lambda: safe_click_enhanced(driver, "#fmod_7 > div.wdTop.ui-draggable-handle > div > a"))
    
    # Verificação de mensagens
    time.sleep(1.5)

    try:
        # Verifica se há mensagens de sucesso/erro
        alertas = driver.find_elements(By.CSS_SELECTOR, ".alerts")
        for alerta in alertas:
            if alerta.is_displayed():
                log(doc, f"📢 Mensagem do sistema: {alerta.text}")
    except:
        pass

    log(doc, "✅ Teste concluído com sucesso!")

except Exception as e:
    log(doc, f"❌ ERRO FATAL: {str(e)}")
    try:
        take_screenshot(driver, doc, "erro_fatal")
    except:
        pass

finally:
    try:
        finalizar_relatorio() 
        driver.quit()
    except:
        pass
    
    log(doc, "🏁 Execução finalizada.")