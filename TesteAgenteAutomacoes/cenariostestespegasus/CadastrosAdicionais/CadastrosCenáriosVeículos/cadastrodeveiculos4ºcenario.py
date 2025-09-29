from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support.ui import Select, WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service
from webdriver_manager.chrome import ChromeDriverManager
from docx import Document
from docx.shared import Inches
from faker import Faker
from faker.providers import BaseProvider
from validate_docbr import CPF
from selenium.webdriver import ActionChains
from datetime import datetime, timedelta
import subprocess
import os
import time
import random
import string

# ==== PROVIDERS CUSTOMIZADOS ====
class BrasilProvider(BaseProvider):
    def rg(self):
        numeros = [str(random.randint(0, 9)) for _ in range(8)]
        return ''.join(numeros) + '-' + str(random.randint(0, 9))

fake = Faker("pt_BR")
fake.add_provider(BrasilProvider)

# ==== CONFIGURAÇÕES ====
URL = "http://localhost:8080/gs/index.xhtml"
LOGIN_EMAIL = "joaoeduardo.gold@outlook.com"
LOGIN_PASSWORD = "071999gs"

# ==== DOCUMENTO ====
doc = Document()
doc.add_heading("RELATÓRIO DO TESTE", 0)
doc.add_paragraph("Cadastro de Veículos - Cenário 4: Preenchimento dos campos NÃO obrigatórios e salvamento")
doc.add_paragraph(f"Data do teste: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")

screenshot_registradas = set()

# ==== FUNÇÕES DE UTILITÁRIO ====
def log(doc, msg):
    print(msg)
    doc.add_paragraph(msg)

def take_screenshot(driver, doc, nome):
    if nome not in screenshot_registradas:
        path = f"screenshots/{nome}.png"
        os.makedirs("screenshots", exist_ok=True)
        driver.save_screenshot(path)
        doc.add_paragraph(f"Screenshot: {nome}")
        try:
            doc.add_picture(path, width=Inches(5.5))
        except:
            doc.add_paragraph("Erro ao adicionar imagem ao relatório")
        screenshot_registradas.add(nome)

def safe_action(doc, descricao, func):
    try:
        log(doc, f"🔄 {descricao}...")
        func()
        log(doc, f"✅ {descricao} realizada com sucesso.")
        take_screenshot(driver, doc, descricao.lower().replace(" ", "_"))
    except Exception as e:
        log(doc, f"❌ Erro ao {descricao.lower()}: {e}")
        take_screenshot(driver, doc, f"erro_{descricao.lower().replace(' ', '_')}")

def finalizar_relatorio():
    nome_arquivo = f"relatorio_veiculo_cenario_4_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc.save(nome_arquivo)
    log(doc, f"📄 Relatório salvo como: {nome_arquivo}")
    try:
        subprocess.run(["start", "winword", nome_arquivo], shell=True)
    except:
        log(doc, "⚠️ Não foi possível abrir o Word automaticamente")

def encontrar_mensagem_alerta():
    seletores = [
        (".alerts.salvo", "✅ Mensagem de Sucesso"),
        (".alerts.alerta", "⚠️ Mensagem de Alerta"),
        (".alerts.erro", "❌ Mensagem de Erro"),
    ]

    for seletor, tipo in seletores:
        try:
            elemento = driver.find_element(By.CSS_SELECTOR, seletor)
            if elemento.is_displayed():
                log(doc, f"📢 {tipo}: {elemento.text}")
                return elemento
        except:
            continue

    log(doc, "ℹ️ Nenhuma mensagem de alerta encontrada.")
    return None

def click_element_safely(driver, wait, selector, by=By.CSS_SELECTOR, timeout=10):
    """Clica em um elemento de forma segura, tentando diferentes métodos"""
    try:
        element = wait.until(EC.presence_of_element_located((by, selector)))
        driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", element)
        time.sleep(0.5)
        
        element = wait.until(EC.element_to_be_clickable((by, selector)))
        
        try:
            element.click()
            return True
        except:
            try:
                ActionChains(driver).move_to_element(element).click().perform()
                return True
            except:
                driver.execute_script("arguments[0].click();", element)
                return True
                
    except Exception as e:
        print(f"Erro ao clicar no elemento {selector}: {e}")
        return False

def gerar_placa_aleatoria():
    """Gera placa aleatória no formato antigo ou Mercosul"""
    if random.choice(["antigo", "mercosul"]) == "antigo":
        letras = ''.join(random.choices(string.ascii_uppercase, k=3))
        numeros = ''.join(random.choices(string.digits, k=4))
        return f"{letras}-{numeros}"
    else:
        letras = ''.join(random.choices(string.ascii_uppercase, k=3))
        meio = random.choice(string.digits)
        letra_final = random.choice(string.ascii_uppercase)
        fim = ''.join(random.choices(string.digits, k=2))
        return f"{letras}{meio}{letra_final}{fim}"

def gerar_renavam():
    """Gera um RENAVAM válido de 11 dígitos"""
    base = ''.join(random.choices(string.digits, k=10))
    
    multiplicadores = [3, 2, 9, 8, 7, 6, 5, 4, 3, 2]
    soma = sum(int(digito) * mult for digito, mult in zip(base, multiplicadores))
    resto = soma % 11
    
    if resto in [0, 1]:
        digito_verificador = '0'
    else:
        digito_verificador = str(11 - resto)
    
    return base + digito_verificador

def gerar_dados_veiculo():
    """Gera dados fictícios para o cadastro de veículo"""
    modelos = ["Fox", "Gol", "Uno", "Celta", "Civic", "Palio", "Corolla", "HB20"]
    modelo = random.choice(modelos)
    ano = random.randint(2000, 2024)
    ano_fabricacao = ano - random.randint(0, 2)
    
    manutencoes_de_km = random.randint(5000, 15000)
    manutencoes_de_dias = random.randint(90, 365)
    notificar_faltando_km = random.randint(500, 2000)
    notificar_faltando_dias = random.randint(5, 30)
    
    km_inicial = random.randint(0, 10000)
    km_atual = km_inicial + random.randint(1000, 50000)
    
    chassi = ''.join(random.choices(string.ascii_uppercase + string.digits, k=17))
    
    return {
        'placa': gerar_placa_aleatoria(),
        'modelo': modelo,
        'ano_fabricacao': ano_fabricacao,
        'ano': ano,
        'manutencoes_de_km': manutencoes_de_km,
        'manutencoes_de_dias': manutencoes_de_dias,
        'notificar_km': notificar_faltando_km,
        'notificar_dias': notificar_faltando_dias,
        'km_atual': km_atual,
        'km_inicial': km_inicial,
        'renavam': gerar_renavam(),
        'chassi': chassi,
        'motorizacao': f"{random.randint(10, 50)}.{random.randint(0, 9)}",
        'valor_locacao': random.randint(50000, 200000)
    }

def preencher_campo_com_retry(driver, wait, seletor, valor, max_tentativas=3):
    """Tenta preencher o campo com diferentes métodos até conseguir"""
    for tentativa in range(max_tentativas):
        try:
            campo = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, seletor)))
            driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", campo)
            time.sleep(0.5)
            
            if tentativa == 0:
                campo.click()
                campo.clear()
                campo.send_keys(valor)
                campo.send_keys(Keys.TAB)
            elif tentativa == 1:
                ActionChains(driver).move_to_element(campo).click().perform()
                time.sleep(0.2)
                ActionChains(driver).key_down(Keys.CONTROL).send_keys('a').key_up(Keys.CONTROL).perform()
                ActionChains(driver).send_keys(valor).perform()
                ActionChains(driver).send_keys(Keys.TAB).perform()
            else:
                driver.execute_script("""
                    var element = arguments[0];
                    var valor = arguments[1];
                    element.focus();
                    element.value = '';
                    element.value = valor;
                    element.dispatchEvent(new Event('input', { bubbles: true }));
                    element.dispatchEvent(new Event('change', { bubbles: true }));
                    element.blur();
                """, campo, valor)
            
            time.sleep(0.5)
            valor_atual = campo.get_attribute('value')
            if valor_atual == str(valor):
                return True
                
        except Exception as e:
            print(f"Tentativa {tentativa + 1} falhou: {e}")
            time.sleep(1)
    
    return False

def abrir_modal_e_selecionar(btn_selector, pesquisa_selector, termo_pesquisa, btn_pesquisar_selector, resultado_xpath):
    def acao():
        wait = WebDriverWait(driver, 20)
        
        if not click_element_safely(driver, wait, btn_selector):
            raise Exception("Não foi possível abrir o modal")
        
        time.sleep(1)
        
        campo_pesquisa = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, pesquisa_selector)))
        campo_pesquisa.clear()
        campo_pesquisa.send_keys(termo_pesquisa)
        
        if not click_element_safely(driver, wait, btn_pesquisar_selector):
            raise Exception("Não foi possível clicar no botão pesquisar")
        
        time.sleep(2)
        
        wait.until(EC.presence_of_element_located((By.XPATH, resultado_xpath)))
        wait.until(EC.element_to_be_clickable((By.XPATH, resultado_xpath)))
        
        resultado = driver.find_element(By.XPATH, resultado_xpath)
        driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", resultado)
        time.sleep(0.5)
        
        if not click_element_safely(driver, wait, resultado_xpath, By.XPATH):
            raise Exception("Não foi possível selecionar o resultado")
    
    return acao

def selecionar_opcao(selector, texto):
    def acao():
        select_element = wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, selector)))
        Select(select_element).select_by_visible_text(texto)
    return acao

def ajustar_zoom():
    try:
        driver.execute_script("document.body.style.zoom='90%'")
        log(doc, "🔍 Zoom ajustado para 90%.")
    except Exception as e:
        log(doc, f"⚠️ Erro ao ajustar zoom: {e}")

# ==== INICIALIZAÇÃO DO DRIVER ====
options = Options()
options.add_argument("--start-maximized")
options.add_argument("--disable-blink-features=AutomationControlled")
options.add_experimental_option("excludeSwitches", ["enable-automation"])
options.add_experimental_option('useAutomationExtension', False)

driver = webdriver.Chrome(service=Service(ChromeDriverManager().install()), options=options)
driver.execute_script("Object.defineProperty(navigator, 'webdriver', {get: () => undefined})")
wait = WebDriverWait(driver, 20)

# Gerar dados do veículo
dados_veiculo = gerar_dados_veiculo()

# ==== EXECUÇÃO DO TESTE ====
try:
    safe_action(doc, "Acessando sistema", lambda: driver.get(URL))

    safe_action(doc, "Realizando login", lambda: (
        wait.until(EC.presence_of_element_located((By.ID, "j_id15:email"))).send_keys(LOGIN_EMAIL),
        wait.until(EC.presence_of_element_located((By.ID, "j_id15:senha"))).send_keys(LOGIN_PASSWORD, Keys.ENTER),
        wait.until(EC.presence_of_element_located((By.TAG_NAME, "body")))
    ))

    safe_action(doc, "Esperando sistema carregar e ajustando zoom", lambda: (
        time.sleep(5),
        ajustar_zoom()
    ))

    safe_action(doc, "Abrindo menu Veículos", lambda: (
        driver.find_element(By.TAG_NAME, "body").send_keys(Keys.F2),
        time.sleep(1),
        wait.until(EC.visibility_of_element_located((By.XPATH, "//input[@placeholder='Busque um cadastro']"))).send_keys("Veículos", Keys.ENTER),
        wait.until(EC.visibility_of_element_located((By.XPATH, "/html/body/div[17]/ul/li[54]/a"))).click()
    ))

    safe_action(doc, "Clicando em Cadastrar", lambda: (
        time.sleep(2),
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_10004 > div.wdTelas > div > ul > li:nth-child(1) > a > span"))).click()
    ))

    safe_action(doc, "Preenchendo o campo Renavam", lambda: 
        preencher_campo_com_retry(
            driver, wait,
            "#fmod_10004 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroVeiculo > div.catWrapper > div > div.cat_categoriaDadosVeiculo.categoriaHolder > div > div > div > div:nth-child(7) > input",
            dados_veiculo['renavam']
        )
    )

    safe_action(doc, "Preenchendo Chassi", lambda: 
        preencher_campo_com_retry(
            driver, wait,
            "#fmod_10004 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroVeiculo > div.catWrapper > div > div.cat_categoriaDadosVeiculo.categoriaHolder > div > div > div > div:nth-child(8) > input",
            dados_veiculo['chassi']
        )
    )

    safe_action(doc, "Preenchendo Motorização", lambda: 
        preencher_campo_com_retry(
            driver, wait,
            "#fmod_10004 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroVeiculo > div.catWrapper > div > div.cat_categoriaDadosVeiculo.categoriaHolder > div > div > div > div:nth-child(10) > input",
            dados_veiculo['motorizacao']
        )
    )

    safe_action(doc, "Preenchendo o campo Km Inicial", lambda: 
        preencher_campo_com_retry(
            driver, wait,
            "#fmod_10004 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroVeiculo > div.catWrapper > div > div.cat_categoriaDadosVeiculo.categoriaHolder > div > div > div > div:nth-child(11) > input",
            dados_veiculo['km_inicial']
        )
    )


    safe_action(doc, "Preenchendo Ano Modelo", lambda: 
        preencher_campo_com_retry(
            driver, wait,
            "#fmod_10004 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroVeiculo > div.catWrapper > div > div.cat_categoriaDadosVeiculo.categoriaHolder > div > div > div > div:nth-child(13) > input",
            dados_veiculo['ano']
        )
    )

    safe_action(doc, "Preenchendo Ano Fabricação", lambda: 
        preencher_campo_com_retry(
            driver, wait,
            "#fmod_10004 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroVeiculo > div.catWrapper > div > div.cat_categoriaDadosVeiculo.categoriaHolder > div > div > div > div:nth-child(14) > input",
            dados_veiculo['ano_fabricacao']
        )
    )

    safe_action(doc, "Preenchendo Valor da locação", lambda: 
        preencher_campo_com_retry(
            driver, wait,
            "#fmod_10004 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroVeiculo > div.catWrapper > div > div.cat_categoriaDadosVeiculo.categoriaHolder > div > div > div > div:nth-child(15) > input",
            dados_veiculo['valor_locacao']
        )
    )

    safe_action(doc, "Selecionando Cor", abrir_modal_e_selecionar(
        "#fmod_10004 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroVeiculo > div.catWrapper > div > div.cat_categoriaDadosVeiculo.categoriaHolder > div > div > div > div:nth-child(16) > div > a",
        "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div.formCol.divPesquisa > div:nth-child(1) > input",
        "VERMELHO",
        "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(3) > a",
        "//td[contains(text(), 'VERMELHO')]"
    ))

    safe_action(doc, "Preenchendo Modelo", lambda: 
        preencher_campo_com_retry(
            driver, wait,
            "#fmod_10004 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroVeiculo > div.catWrapper > div > div.cat_categoriaDadosVeiculo.categoriaHolder > div > div > div > div:nth-child(17) > input",
            dados_veiculo['modelo']
        )
    )

    safe_action(doc, "Acessando aba Manutenção preventiva", lambda: (
        wait.until(EC.element_to_be_clickable((By.LINK_TEXT, "Manunteção preventiva"))).click(),
        time.sleep(1)
    ))

    safe_action(doc, "Preenchendo o campo Manutenções de (km)", lambda: 
        preencher_campo_com_retry(
            driver, wait,
            "#fmod_10004 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroVeiculo > div.catWrapper > div > div.cat_categoriaManutencaoPreventiva.categoriaHolder > div.groupHolder.clearfix.grupo_grupoNotificacaoManutencaoPreventiva > div > div > div:nth-child(1) > input",
            dados_veiculo['manutencoes_de_km']
        )
    )

    safe_action(doc, "Preenchendo o campo Manutenções de (dias)", lambda: 
        preencher_campo_com_retry(
            driver, wait,
            "#fmod_10004 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroVeiculo > div.catWrapper > div > div.cat_categoriaManutencaoPreventiva.categoriaHolder > div.groupHolder.clearfix.grupo_grupoNotificacaoManutencaoPreventiva > div > div > div:nth-child(2) > input",
            dados_veiculo['manutencoes_de_dias']
        )
    )

    safe_action(doc, "Salvando cadastro", lambda: driver.find_element(
        By.CSS_SELECTOR, "#fmod_10004 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroVeiculo > div.btnHolder > a.btModel.btGray.btsave"
    ).click())

    time.sleep(1.5)  # Aguarda processamento

    safe_action(doc, "Fechando modal após salvamento", lambda: wait.until(EC.element_to_be_clickable((
        By.CSS_SELECTOR, "#fmod_10004 > div.wdTop.ui-draggable-handle > div.wdClose > a"
    ))).click())

    encontrar_mensagem_alerta()

except Exception as e:
    log(doc, f"❌ ERRO FATAL: {e}")
    take_screenshot(driver, doc, "erro_fatal")

finally:
    log(doc, "✅ Teste concluído.")
    finalizar_relatorio()
    driver.quit()