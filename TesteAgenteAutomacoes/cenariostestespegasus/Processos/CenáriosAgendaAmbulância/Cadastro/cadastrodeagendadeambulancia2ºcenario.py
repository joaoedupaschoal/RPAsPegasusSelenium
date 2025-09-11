# Corrigido: cadastroregistrodeobito1cenario.py

from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support.ui import Select, WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service
from selenium.common.exceptions import TimeoutException, NoSuchElementException, StaleElementReferenceException
from webdriver_manager.chrome import ChromeDriverManager
from docx import Document
from docx.shared import Inches
from faker import Faker
from faker.providers import BaseProvider
from validate_docbr import CPF
from datetime import datetime, timedelta
from selenium.webdriver import ActionChains
import subprocess
import os
import time
import random

# ==== PROVIDERS CUSTOMIZADOS ====
class BrasilProvider(BaseProvider):
    def rg(self):
        numeros = [str(random.randint(0, 9)) for _ in range(8)]
        return ''.join(numeros) + '-' + str(random.randint(0, 9))

fake = Faker("pt_BR")
fake.add_provider(BrasilProvider)

def gerar_datas_validas():
    """Gera datas coerentes para nascimento, falecimento e sepultamento dentro de um intervalo válido."""
    hoje = datetime.today().date()
    dez_anos_atras = hoje - timedelta(days=3650)  # Limite máximo de 10 anos atrás
    
    # Gera uma data de falecimento entre 10 anos atrás e hoje
    data_falecimento = fake.date_between(start_date=dez_anos_atras, end_date=hoje)

    # Garante que a pessoa tenha no mínimo 18 anos na data do falecimento
    idade_minima = 18
    idade_maxima = 110
    data_nascimento = data_falecimento - timedelta(days=random.randint(idade_minima * 365, idade_maxima * 365))

    # Sepultamento entre 1 e 10 dias após o falecimento
    data_sepultamento = data_falecimento + timedelta(days=random.randint(1, 10))

    # Registro entre 1 e 10 dias após o sepultamento
    data_registro = data_sepultamento + timedelta(days=random.randint(1, 10))

    data_velorio = fake.date_between(start_date=data_falecimento, end_date=data_sepultamento)

    return (
        data_nascimento.strftime("%d/%m/%Y"),
        data_falecimento.strftime("%d/%m/%Y"),
        data_sepultamento.strftime("%d/%m/%Y"),
        data_velorio.strftime("%d/%m/%Y"),
        data_registro.strftime("%d/%m/%Y")
    )

# Gera os valores corretos
data_nascimento, data_falecimento, data_sepultamento, data_velorio, data_registro = gerar_datas_validas()
hora_falecimento = fake.time(pattern="%H:%M")
hora_sepultamento = fake.time(pattern="%H:%M")
localizacao = fake.city()

# ==== CONFIGURAÇÕES ====
URL = "http://localhost:8080/gs/index.xhtml"
LOGIN_EMAIL = "joaoeduardo.gold@outlook.com"
LOGIN_PASSWORD = "071999gs"

# ==== INICIALIZAÇÃO DE VARIÁVEIS GLOBAIS ====
doc = Document()
doc.add_heading("RELATÓRIO DO TESTE", 0)
doc.add_paragraph("Cadastro de Agenda de Ambulância – Cenário 2: Preenchimento completo e cancelamento.")
doc.add_paragraph(f"Data do teste: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")

screenshot_registradas = set()
driver = None
wait = None

# ==== FUNÇÕES DE UTILITÁRIO MELHORADAS ====
def log(doc, msg):
    print(msg)
    doc.add_paragraph(msg)

def take_screenshot(driver, doc, nome):
    if driver is None:
        return
    if nome not in screenshot_registradas:
        path = f"screenshots/{nome}.png"
        os.makedirs("screenshots", exist_ok=True)
        try:
            driver.save_screenshot(path)
            doc.add_paragraph(f"Screenshot: {nome}")
            doc.add_picture(path, width=Inches(5.5))
            screenshot_registradas.add(nome)
        except Exception as e:
            log(doc, f"⚠️ Erro ao tirar screenshot {nome}: {e}")

def safe_action(doc, descricao, func, max_retries=3):
    """Executa uma ação com retry e tratamento robusto de erros"""
    global driver
    
    for attempt in range(max_retries):
        try:
            if attempt == 0:
                log(doc, f"🔄 {descricao}...")
            else:
                log(doc, f"🔄 {descricao}... (Tentativa {attempt + 1})")
            
            func()
            log(doc, f"✅ {descricao} realizada com sucesso.")
            take_screenshot(driver, doc, descricao.lower().replace(" ", "_"))
            return True
            
        except (TimeoutException, NoSuchElementException, StaleElementReferenceException) as e:
            if attempt < max_retries - 1:
                log(doc, f"⚠️ Tentativa {attempt + 1} falhou para {descricao}, tentando novamente...")
                time.sleep(2)
                continue
            else:
                log(doc, f"❌ Erro ao {descricao.lower()} após {max_retries} tentativas: {e}")
                take_screenshot(driver, doc, f"erro_{descricao.lower().replace(' ', '_')}")
                return False
        except Exception as e:
            log(doc, f"❌ Erro inesperado ao {descricao.lower()}: {e}")
            take_screenshot(driver, doc, f"erro_{descricao.lower().replace(' ', '_')}")
            return False

def finalizar_relatorio():
    global driver, doc
    
    nome_arquivo = f"relatorio_agenda_ambulancia_cenario_2_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    try:
        doc.save(nome_arquivo)
        log(doc, f"📄 Relatório salvo como: {nome_arquivo}")
        try:
            subprocess.run(["start", "winword", nome_arquivo], shell=True)
        except:
            pass
    except Exception as e:
        print(f"Erro ao salvar relatório: {e}")
    
    if driver:
        try:
            driver.quit()
        except:
            pass

def encontrar_mensagem_alerta():
    global driver, doc
    
    if driver is None:
        return None
        
    seletores = [
        (".alerts.salvo", "✅ Sucesso"),
        (".alerts.alerta", "⚠️ Alerta"),
        (".alerts.erro", "❌ Erro"),
    ]

    for seletor, tipo in seletores:
        try:
            elemento = driver.find_element(By.CSS_SELECTOR, seletor)
            if elemento.is_displayed():
                log(doc, f"📢 {tipo}: {elemento.text}")
                return elemento
        except:
            continue

    log(doc, "ℹ️ Nenhuma mensagem de alerta encontrada.")
    return None

def ajustar_zoom():
    global driver, doc
    
    if driver is None:
        return
        
    try:
        driver.execute_script("document.body.style.zoom='90%'")
        log(doc, "🔍 Zoom ajustado para 90%.")
    except Exception as e:
        log(doc, f"⚠️ Erro ao ajustar zoom: {e}")

def safe_scroll_and_interact(selector, action_type="click", value=None, timeout=10, by_xpath=False):
    """Função robusta para rolar até elemento e interagir com ele"""
    global driver, doc
    
    if driver is None:
        return None
        
    try:
        # Escolhe o tipo de seletor
        by_type = By.XPATH if by_xpath else By.CSS_SELECTOR
        
        # Aguarda o elemento estar presente
        element = WebDriverWait(driver, timeout).until(
            EC.presence_of_element_located((by_type, selector))
        )
        
        # Rola até o elemento
        driver.execute_script("arguments[0].scrollIntoView({block: 'center', behavior: 'smooth'});", element)
        time.sleep(0.5)
        
        # Aguarda o elemento estar clicável se necessário
        if action_type in ["click", "send_keys"]:
            element = WebDriverWait(driver, timeout).until(
                EC.element_to_be_clickable((by_type, selector))
            )
        
        # Executa a ação
        if action_type == "click":
            element.click()
        elif action_type == "send_keys" and value:
            element.clear()
            element.send_keys(value)
        elif action_type == "select" and value:
            Select(element).select_by_visible_text(value)
            
        return element
        
    except Exception as e:
        log(doc, f"❌ Erro ao interagir com elemento {selector}: {e}")
        return None

def preencher_campo_robusto(selector, valor, clear_first=True):
    """Preenche campo de forma robusta"""
    def acao():
        element = safe_scroll_and_interact(selector, "send_keys", valor)
        if element is None:
            raise Exception(f"Não foi possível encontrar o elemento: {selector}")
    return acao

def selecionar_opcao_robusta(selector, texto):
    """Seleciona opção de forma robusta"""
    def acao():
        element = safe_scroll_and_interact(selector, "select", texto)
        if element is None:
            raise Exception(f"Não foi possível encontrar o select: {selector}")
    return acao

def clicar_elemento_robusto(selector):
    """Clica em elemento de forma robusta"""
    def acao():
        element = safe_scroll_and_interact(selector, "click")
        if element is None:
            raise Exception(f"Não foi possível clicar no elemento: {selector}")
    return acao

def preencher_campo_com_retry(driver, wait, seletor, valor, max_tentativas=3):
    """Tenta preencher o campo com diferentes métodos até conseguir"""
    global doc
    
    if driver is None or wait is None:
        return False
    
    for tentativa in range(max_tentativas):
        try:
            # Aguarda o elemento
            campo = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, seletor)))
            
            # Scroll até o elemento se necessário
            driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", campo)
            time.sleep(0.5)
            
            # Método 1: Tradicional
            if tentativa == 0:
                campo.click()
                campo.clear()
                campo.send_keys(valor)
                campo.send_keys(Keys.TAB)
            
            # Método 2: ActionChains
            elif tentativa == 1:
                ActionChains(driver).move_to_element(campo).click().perform()
                time.sleep(0.2)
                ActionChains(driver).key_down(Keys.CONTROL).send_keys('a').key_up(Keys.CONTROL).perform()
                ActionChains(driver).send_keys(valor).perform()
                ActionChains(driver).send_keys(Keys.TAB).perform()
            
            # Método 3: JavaScript
            else:
                driver.execute_script("""
                    var element = arguments[0];
                    var valor = arguments[1];
                    element.focus();
                    element.value = '';
                    element.value = valor;
                    element.dispatchEvent(new Event('input', { bubbles: true }));
                    element.dispatchEvent(new Event('change', { bubbles: true }));
                    element.blur();
                """, campo, valor)
            
            time.sleep(0.5)
            
            # Verifica se o valor foi preenchido
            valor_atual = campo.get_attribute('value')
            if valor_atual == valor:
                return True
                
        except Exception as e:
            log(doc, f"⚠️ Tentativa {tentativa + 1} falhou: {e}")
            time.sleep(1)
    
    return False

def abrir_modal_e_selecionar_robusto(btn_selector, pesquisa_selector, termo_pesquisa, btn_pesquisar_selector, resultado_xpath):
    """Versão robusta da função de modal"""
    global driver, wait, doc
    
    def acao():
        if driver is None or wait is None:
            raise Exception("Driver ou wait não inicializados")
            
        # Abre o modal
        safe_scroll_and_interact(btn_selector, "click")
        time.sleep(1)

        # Aguarda e preenche campo pesquisa
        campo_pesquisa = wait.until(
            EC.element_to_be_clickable((By.CSS_SELECTOR, pesquisa_selector))
        )
        campo_pesquisa.clear()
        campo_pesquisa.send_keys(termo_pesquisa)
        time.sleep(0.5)

        # Clica pesquisar
        pesquisar = wait.until(
            EC.element_to_be_clickable((By.CSS_SELECTOR, btn_pesquisar_selector))
        )
        pesquisar.click()
        time.sleep(2)
        
        # Aguarda resultado e clica
        resultado = wait.until(
            EC.element_to_be_clickable((By.XPATH, resultado_xpath))
        )
        driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", resultado)
        time.sleep(0.5)
        resultado.click()
        time.sleep(1)

    return acao

def inicializar_driver():
    """Inicializa o driver do Chrome"""
    global driver, wait
    
    try:
        options = Options()
        options.add_argument("--start-maximized")
        options.add_argument("--disable-blink-features=AutomationControlled")
        options.add_experimental_option("excludeSwitches", ["enable-automation"])
        options.add_experimental_option('useAutomationExtension', False)

        driver = webdriver.Chrome(service=Service(ChromeDriverManager().install()), options=options)
        driver.execute_script("Object.defineProperty(navigator, 'webdriver', {get: () => undefined})")
        wait = WebDriverWait(driver, 20)
        
        return True
    except Exception as e:
        log(doc, f"❌ Erro ao inicializar driver: {e}")
        return False

# ==== EXECUÇÃO DO TESTE ====
def executar_teste():
    global driver, wait, doc
    
    try:
        # Inicializa o driver
        if not inicializar_driver():
            return False

        safe_action(doc, "Acessando sistema", lambda: driver.get(URL))

        safe_action(doc, "Realizando login", lambda: (
            wait.until(EC.presence_of_element_located((By.ID, "j_id15:email"))).send_keys(LOGIN_EMAIL),
            wait.until(EC.presence_of_element_located((By.ID, "j_id15:senha"))).send_keys(LOGIN_PASSWORD, Keys.ENTER),
            time.sleep(5)
        ))

        safe_action(doc, "Ajustando zoom e abrindo menu", lambda: (
            ajustar_zoom(),
            driver.find_element(By.TAG_NAME, "body").send_keys(Keys.F3)
        ))

        safe_action(doc, "Acessando Agenda Ambulância", lambda: (
            wait.until(EC.element_to_be_clickable((By.XPATH, '/html/body/div[15]/ul/li[27]/img'))).click()
        ))

        safe_action(doc, "Clicando em Cadastrar", lambda: (
            wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, '#gsAgendaAmbulancia > div.wdTelas > div > ul > li:nth-child(1) > a > span'))).click()
        ))


        safe_action(doc, "Preenchendo Número do Contrato", preencher_campo_robusto(
            '#gsAgendaAmbulancia > div.wdTelas > div.wdWizard.clearfix.telaConsulta > div.wizardHolder > div > div.step1 > div > div.formCol.numero-contrato > div > input.fc.mandatory',
            '113056'
        ))



        safe_action(doc, "Abrindo lov do Solicitante", lambda: WebDriverWait(driver, 10).until(
            EC.element_to_be_clickable((By.XPATH, '/html/body/div[16]/div[2]/div[2]/div[2]/div[2]/div/div[1]/div/div[3]/div/div/a'))
        ).click())


        xpath = "//tr[td[1][text()='TESTE TITULAR'] and td[2][text()='Física'] and td[3][text()='947.011.330-60'] and td[7][text()='04/08/2025']]"

        safe_action(doc, "Clicando na linha do 'TESTE TITULAR'", lambda: WebDriverWait(driver, 10).until(
            EC.element_to_be_clickable((By.XPATH, xpath))
        ).click())

        safe_action(doc, "Abrindo lov do Paciente", lambda: WebDriverWait(driver, 10).until(
            EC.element_to_be_clickable((By.XPATH, '/html/body/div[16]/div[2]/div[2]/div[2]/div[2]/div/div[1]/div/div[4]/div/div/a'))
        ).click())


        xpath = "//tr[td[1][text()='TESTE TITULAR'] and td[2][text()='Física'] and td[3][text()='947.011.330-60'] and td[7][text()='04/08/2025']]"

        safe_action(doc, "Clicando na linha do 'TESTE TITULAR'", lambda: WebDriverWait(driver, 10).until(
            EC.element_to_be_clickable((By.XPATH, xpath))
        ).click())


        safe_action(doc, "Preenchendo Telefone", preencher_campo_robusto(
            '#gsAgendaAmbulancia > div.wdTelas > div.wdWizard.clearfix.telaConsulta > div.wizardHolder > div > div.step1 > div > div:nth-child(5) > input[type=text]',
            fake.phone_number()
        ))


        safe_action(doc, "Preenchendo Complemento", preencher_campo_robusto(
            '#gsAgendaAmbulancia > div.wdTelas > div.wdWizard.clearfix.telaConsulta > div.wizardHolder > div > div.step1 > div > div.dadosEnderecoPaciente.group_pacienter > div:nth-child(5) > input',
            'Casa'
        ))


   
        safe_action(doc, "Avançando para próxima etapa", lambda: (
            wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, '#gsAgendaAmbulancia > div.wdTelas > div.wdWizard.clearfix.telaConsulta > div.btnHolder > a:nth-child(3)'))).click()
        ))

        safe_action(doc, "Clicando no Lov para consultar um motorista", lambda: (
            wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, '#gsAgendaAmbulancia > div.wdTelas > div.wdWizard.clearfix.telaConsulta > div.wizardHolder > div > div.step2 > div:nth-child(1) > div:nth-child(1) > a'))).click()
        ))

        safe_action(doc, "Selecionando Motorista", lambda: WebDriverWait(driver, 10).until(
            EC.element_to_be_clickable((By.XPATH, "//td[normalize-space(text())='MOTORISTA TREINAMENTO']"))
        ).click())

        safe_action(doc, "Preenchendo Nome do Acompanhante", preencher_campo_robusto(
            '#gsAgendaAmbulancia > div.wdTelas > div.wdWizard.clearfix.telaConsulta > div.wizardHolder > div > div.step2 > div:nth-child(2) > div:nth-child(6) > input[type=text]',
            fake.name()
        ))

        safe_action(doc, "Preenchendo Local Transporte", preencher_campo_robusto(
            '#gsAgendaAmbulancia > div.wdTelas > div.wdWizard.clearfix.telaConsulta > div.wizardHolder > div > div.step2 > div:nth-child(2) > div:nth-child(7) > input[type=text]',
            "Rua das Flores 123, Bairro Jardim, Cidade Exemplo"
        ))

        safe_action(doc, "Preenchendo Cidade Transporte", preencher_campo_robusto(
             '#gsAgendaAmbulancia > div.wdTelas > div.wdWizard.clearfix.telaConsulta > div.wizardHolder > div > div.step2 > div:nth-child(2) > div:nth-child(8) > input',
             fake.city()
        ))


        safe_action(doc, "Preenchendo Retorno", preencher_campo_robusto(
             '#gsAgendaAmbulancia > div.wdTelas > div.wdWizard.clearfix.telaConsulta > div.wizardHolder > div > div.step2 > div:nth-child(2) > div:nth-child(9) > input[type=text]',
             'TESTE RETORNO SELENIUM AUTOMATIZADO'
        ))

        safe_action(doc, "Preenchendo Observação", preencher_campo_robusto(
             '#gsAgendaAmbulancia > div.wdTelas > div.wdWizard.clearfix.telaConsulta > div.wizardHolder > div > div.step2 > div:nth-child(3) > div:nth-child(1) > textarea',
             'TESTE OBSERVAÇÃO SELENIUM AUTOMATIZADO'
        ))

        safe_action(doc, "Preenchendo Situação do Paciente", preencher_campo_robusto(
             '#gsAgendaAmbulancia > div.wdTelas > div.wdWizard.clearfix.telaConsulta > div.wizardHolder > div > div.step2 > div:nth-child(3) > div:nth-child(2) > textarea',
             'TESTE OBSERVAÇÃO SELENIUM AUTOMATIZADO: O PACIENTE ESTÁ ESTÁVEL E PRONTO PARA O TRANSPORTE, PORÉM, ESTÁ SERIAMENTE MACHUCADO.'
        ))

        safe_action(doc, "Avançando para próxima etapa", lambda: (
            wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, '#gsAgendaAmbulancia > div.wdTelas > div.wdWizard.clearfix.telaConsulta > div.btnHolder > a:nth-child(3)'))).click()
        ))

        time.sleep(1)  # Pequena pausa para garantir que a tela carregue

        safe_action(doc, "Avançando para finalização", lambda: (
            wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, '#gsAgendaAmbulancia > div.wdTelas > div.wdWizard.clearfix.telaConsulta > div.btnHolder > a:nth-child(3)'))).click()
        ))


        safe_action(doc, "Cancelando cadastro", clicar_elemento_robusto(
            '#gsAgendaAmbulancia > div.wdTelas > div.wdWizard.clearfix.telaConsulta > div.btnHolder > a:nth-child(1)'
        ))


        safe_action(doc, "Fechando modal após o cancelamento", clicar_elemento_robusto(
            '#gsAgendaAmbulancia > div.wdTop.ui-draggable-handle > div > a'
        ))

        log(doc, "🔍 Verificando mensagens de alerta...")
        encontrar_mensagem_alerta()



        return True

    except Exception as e:
        log(doc, f"❌ ERRO FATAL: {e}")
        take_screenshot(driver, doc, "erro_fatal")
        return False

    finally:
        log(doc, "✅ Teste concluído.")

# ==== FUNÇÃO PRINCIPAL ====
def main():
    """Função principal que executa o teste"""
    global doc
    
    try:
        log(doc, "🚀 Iniciando teste de cadastro de Agenda de Ambulância...")        
        sucesso = executar_teste()
        
        if sucesso:
            log(doc, "✅ Teste executado com sucesso!")
        else:
            log(doc, "❌ Teste finalizado com erros.")
            
    except Exception as e:
        log(doc, f"❌ Erro na execução principal: {e}")
        
    finally:
        finalizar_relatorio()

# ==== EXECUÇÃO ====
if __name__ == "__main__":
    main()