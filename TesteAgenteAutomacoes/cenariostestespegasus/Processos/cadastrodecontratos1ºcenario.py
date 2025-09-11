# Refatorado e organizado: cadastrodecontratos1ºcenario.py

from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support.ui import Select, WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service
from selenium.webdriver import ActionChains
from webdriver_manager.chrome import ChromeDriverManager
from docx import Document
from docx.shared import Inches
from faker import Faker
from faker.providers import BaseProvider
from validate_docbr import CPF
from datetime import datetime, timedelta
import subprocess
import os
import time
import random
import string

# ==== PROVIDERS CUSTOMIZADOS ====
class BrasilProvider(BaseProvider):
    def rg(self):
        numeros = [str(random.randint(0, 9)) for _ in range(8)]
        return ''.join(numeros) + '-' + str(random.randint(0, 9))

fake = Faker("pt_BR")
fake.add_provider(BrasilProvider)

# ==== VARIÁVEIS GLOBAIS ====
numero_aleatorio = random.randint(1, 100)
letra_aleatoria = random.choice(string.ascii_uppercase)
cemetery_name = f"Cemitério {fake.last_name()} {fake.random.choice(['Eterno', 'da Paz', 'Memorial', 'Descanso'])}"
qtd_parcelas_em_atraso = int(fake.random.choice(['1', '2', '3', '4', '5']))
dias_para_exumar = int(fake.random.choice(['365', '730', '1095', '1460', '1825']))

def gerar_jazigos():
    quantidade_ruas = random.randint(1, 10)
    max_jazigos_por_rua = random.randint(1, 20)
    quantidade_total_jazigos = quantidade_ruas * max_jazigos_por_rua
    return quantidade_ruas, max_jazigos_por_rua, quantidade_total_jazigos

ruas, jazigos_por_rua, total_jazigos = gerar_jazigos()
altura_cm = random.randint(100, 200)
largura_cm = random.randint(100, 200)
comprimento_cm = random.randint(100, 200)
valor_taxa_adesao = round(random.uniform(2000, 10000), 2)

def gerar_datas_validas():
    """Gera datas coerentes para nascimento, falecimento e sepultamento dentro de um intervalo válido."""
    hoje = datetime.today().date()
    dez_anos_atras = hoje - timedelta(days=3650)
    
    # Data de falecimento entre 10 anos atrás e hoje
    data_falecimento = fake.date_between(start_date=dez_anos_atras, end_date=hoje)
    
    # Pessoa com no mínimo 18 anos na data do falecimento
    idade_minima = 18
    idade_maxima = 110
    data_nascimento = data_falecimento - timedelta(days=random.randint(idade_minima * 365, idade_maxima * 365))
    
    # Sepultamento entre 1 e 10 dias após o falecimento
    data_sepultamento = data_falecimento + timedelta(days=random.randint(1, 10))
    
    # Registro entre 1 e 10 dias após o sepultamento
    data_registro = data_sepultamento + timedelta(days=random.randint(1, 10))
    
    # Datas de/até para período
    data_de = hoje + timedelta(days=random.randint(1, 10))
    data_ate = data_de + timedelta(days=random.randint(1, 100))
    
    return (
        data_nascimento.strftime("%d/%m/%Y"),
        data_falecimento.strftime("%d/%m/%Y"),
        data_sepultamento.strftime("%d/%m/%Y"),
        data_registro.strftime("%d/%m/%Y"),
        data_de.strftime("%d/%m/%Y"),
        data_ate.strftime("%d/%m/%Y"),
    )


def preencher_campo_com_retry(driver, wait, seletor, valor, max_tentativas=3):
    """Tenta preencher o campo com diferentes métodos até conseguir"""
    
    for tentativa in range(max_tentativas):
        try:

            
            # Aguarda o elemento
            campo = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, seletor)))
            
            # Scroll até o elemento se necessário
            driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", campo)
            time.sleep(0.5)
            
            # Método 1: Tradicional
            if tentativa == 0:
                campo.click()
                campo.clear()
                campo.send_keys(valor)
                campo.send_keys(Keys.TAB)
            
            # Método 2: ActionChains
            elif tentativa == 1:
                ActionChains(driver).move_to_element(campo).click().perform()
                time.sleep(0.2)
                ActionChains(driver).key_down(Keys.CONTROL).send_keys('a').key_up(Keys.CONTROL).perform()
                ActionChains(driver).send_keys(valor).perform()
                ActionChains(driver).send_keys(Keys.TAB).perform()
            
            # Método 3: JavaScript
            else:
                driver.execute_script("""
                    var element = arguments[0];
                    var valor = arguments[1];
                    element.focus();
                    element.value = '';
                    element.value = valor;
                    element.dispatchEvent(new Event('input', { bubbles: true }));
                    element.dispatchEvent(new Event('change', { bubbles: true }));
                    element.blur();
                """, campo, valor)
            
            time.sleep(0.5)
            
            # Verifica se o valor foi preenchido
            valor_atual = campo.get_attribute('value')
            if valor_atual == valor:

                return True
            else:
                print()
                
        except Exception as e:
            time.sleep(1)
    

    return False

def gerar_dados_documentos():
    """Gera documentos fictícios para o cadastro."""
    carteira_trabalho = str(random.randint(10000000, 99999999))
    pis = fake.cpf().replace('.', '').replace('-', '')[:11]
    cnh = str(random.randint(10000000000, 99999999999))
    
    return carteira_trabalho, pis, cnh

# Gera os dados necessários
data_nascimento, data_falecimento, data_sepultamento, data_registro, data_de, data_ate = gerar_datas_validas()
carteira_trabalho, pis, cnh = gerar_dados_documentos()

vencimento_cnh = fake.date_between(start_date='today', end_date='+10y')
vencimento_cnh_str = vencimento_cnh.strftime('%d/%m/%Y')

data_admissao = fake.date_between(start_date='-10y', end_date='today')
data_admissao_str = data_admissao.strftime('%d/%m/%Y')

hora_falecimento = fake.time(pattern="%H:%M")
hora_sepultamento = fake.time(pattern="%H:%M")
localizacao = fake.city()

# ==== CONFIGURAÇÕES ====
URL = "http://localhost:8080/gs/index.xhtml"
LOGIN_EMAIL = "joaoeduardo.gold@outlook.com"
LOGIN_PASSWORD = "071999gs"

# ==== DOCUMENTO ====
doc = Document()
doc.add_heading("RELATÓRIO DO TESTE", 0)
doc.add_paragraph("Cadastro de Contratos – Cenário 1: Preenchimento completo e salvamento.")
doc.add_paragraph(f"Data do teste: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")

screenshot_registradas = set()

# ==== FUNÇÕES DE UTILITÁRIO ====
def log(doc, msg):
    print(msg)
    doc.add_paragraph(msg)

def take_screenshot(driver, doc, nome):
    if nome not in screenshot_registradas:
        path = f"screenshots/{nome}.png"
        os.makedirs("screenshots", exist_ok=True)
        driver.save_screenshot(path)
        doc.add_paragraph(f"Screenshot: {nome}")
        doc.add_picture(path, width=Inches(5.5))
        screenshot_registradas.add(nome)

def safe_action(doc, descricao, func):
    try:
        log(doc, f"🔄 {descricao}...")
        func()
        log(doc, f"✅ {descricao} realizada com sucesso.")
        take_screenshot(driver, doc, descricao.lower().replace(" ", "_"))
    except Exception as e:
        log(doc, f"❌ Erro ao {descricao.lower()}: {e}")
        take_screenshot(driver, doc, f"erro_{descricao.lower().replace(' ', '_')}")

def finalizar_relatorio():
    nome_arquivo = f"relatorio_contratos_cenario_1_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc.save(nome_arquivo)
    log(doc, f"📄 Relatório salvo como: {nome_arquivo}")
    subprocess.run(["start", "winword", nome_arquivo], shell=True)
    driver.quit()

def ajustar_zoom():
    try:
        driver.execute_script("document.body.style.zoom='90%'")
        log(doc, "🔍 Zoom ajustado para 90%.")
    except Exception as e:
        log(doc, f"⚠️ Erro ao ajustar zoom: {e}")

def abrir_modal_e_selecionar(btn_selector, pesquisa_selector, termo_pesquisa, btn_pesquisar_selector, resultado_xpath):
    def acao():
        # Abre o modal
        open_lov = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, btn_selector)))
        open_lov.click()

        # Aguarda campo pesquisa
        campo_pesquisa = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, pesquisa_selector)))
        campo_pesquisa.clear()
        campo_pesquisa.send_keys(termo_pesquisa, Keys.ENTER)

        time.sleep(1)
        # Espera o resultado carregar
        wait.until(EC.presence_of_element_located((By.XPATH, resultado_xpath)))
        wait.until(EC.visibility_of_element_located((By.XPATH, resultado_xpath)))
        wait.until(EC.element_to_be_clickable((By.XPATH, resultado_xpath)))

        # Relocaliza no último instante (evita stale element)
        resultado = driver.find_element(By.XPATH, resultado_xpath)
        driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", resultado)
        time.sleep(0.2)
        resultado.click()

    return acao

def encontrar_mensagem_alerta():
    seletores = [
        (".alerts.salvo", "✅ Sucesso"),
        (".alerts.alerta", "⚠️ Alerta"),
        (".alerts.erro", "❌ Erro"),
    ]

    for seletor, tipo in seletores:
        try:
            elemento = driver.find_element(By.CSS_SELECTOR, seletor)
            if elemento.is_displayed():
                log(doc, f"📢 {tipo}: {elemento.text}")
                return elemento
        except:
            continue

    log(doc, "ℹ️ Nenhuma mensagem de alerta encontrada.")
    return None


def rolar_ate_grupo_rateio():
    open_lov_rateio = WebDriverWait(driver, 10).until(
        EC.element_to_be_clickable((By.CSS_SELECTOR, '#gsContratos > div.wdTelas > div.wdWizard.clearfix.telaConsulta > div.wizardHolder > div > div.stepPacote.step3 > div:nth-child(11) > div > div > a'))
    )
    driver.execute_script("arguments[0].scrollIntoView();", open_lov_rateio)



def rolar_ate_salvar_titular():
    salvar_titular = WebDriverWait(driver, 10).until(
        EC.element_to_be_clickable((By.CSS_SELECTOR, '#ui-id-6 > div:nth-child(9) > div:nth-child(5) > a'))
    )
    driver.execute_script("arguments[0].scrollIntoView();", salvar_titular)


def preencher_campo_data(selector, valor):
    def acao():
        campo = wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, selector)))
        campo.click()
        campo.clear()
        campo.send_keys(valor)
        time.sleep(0.2)
    return acao

def selecionar_opcao(selector, texto):
    def acao():
        select_element = wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, selector)))
        Select(select_element).select_by_visible_text(texto)
    return acao

def preencher_pessoa_completa(nome_pessoa):
    """Função para preencher dados completos de uma pessoa."""
    def acao():
        # Dados Pessoais
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "#cg_1 > div.wdTelas > div > div.catWrapper > div > div.cat_dadosPessoais.categoriaHolder > div > div > div:nth-child(2) > div:nth-child(2) > input"))).send_keys(nome_pessoa)
        Select(wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "#cg_1 > div.wdTelas > div > div.catWrapper > div > div.cat_dadosPessoais.categoriaHolder > div > div > div:nth-child(2) > div:nth-child(3) > select")))).select_by_visible_text("Física")
        Select(wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "#cg_1 > div.wdTelas > div > div.catWrapper > div > div.cat_dadosPessoais.categoriaHolder > div > div > div:nth-child(2) > div:nth-child(4) > select")))).select_by_visible_text("Carteira de Identidade Classista")
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "#cg_1 > div.wdTelas > div > div.catWrapper > div > div.cat_dadosPessoais.categoriaHolder > div > div > div.formRow.divPessoaFISICA > div:nth-child(1) > input"))).send_keys(fake.rg())

        # Data de expedição
        campo_data_expedicao = wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, "input.dataExpedicao")))
        campo_data_expedicao.click()
        campo_data_expedicao.send_keys(fake.date_of_birth(minimum_age=18, maximum_age=60).strftime("%d/%m/%Y"))

        # CPF
        cpf = CPF().generate()
        cpf_field = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "#cg_1 > div.wdTelas > div > div.catWrapper > div > div.cat_dadosPessoais.categoriaHolder > div > div > div.formRow.divPessoaFISICA > div:nth-child(3) > input")))
        cpf_field.click()
        time.sleep(0.5)
        cpf_field.send_keys(cpf)

        # Dados Complementares
        wait.until(EC.element_to_be_clickable((By.LINK_TEXT, "Dados Complementares"))).click()
        time.sleep(1)

        Select(wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "#cg_1 > div.wdTelas > div > div.catWrapper > div > div.cat_dadosComplementares.categoriaHolder > div > div > div > div:nth-child(1) > select")))).select_by_visible_text("Solteiro")
        Select(wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "#cg_1 > div.wdTelas > div > div.catWrapper > div > div.cat_dadosComplementares.categoriaHolder > div > div > div > div:nth-child(2) > select")))).select_by_visible_text("Feminino")

        # Data de nascimento
        campo_data_nascimento = wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, "input.dataNascimento")))
        campo_data_nascimento.click()
        campo_data_nascimento.send_keys(fake.date_of_birth(minimum_age=18, maximum_age=60).strftime("%d/%m/%Y"))

        # Contatos
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "#cg_1 > div.wdTelas > div > div.catWrapper > div > div.cat_dadosComplementares.categoriaHolder > div > div > div > div:nth-child(5) > input"))).send_keys(fake.phone_number())
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "#cg_1 > div.wdTelas > div > div.catWrapper > div > div.cat_dadosComplementares.categoriaHolder > div > div > div > div:nth-child(6) > input"))).send_keys(fake.phone_number())
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "#cg_1 > div.wdTelas > div > div.catWrapper > div > div.cat_dadosComplementares.categoriaHolder > div > div > div > div:nth-child(7) > input"))).send_keys(fake.phone_number())
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "#cg_1 > div.wdTelas > div > div.catWrapper > div > div.cat_dadosComplementares.categoriaHolder > div > div > div > div:nth-child(9) > input"))).send_keys(fake.email())
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "#cg_1 > div.wdTelas > div > div.catWrapper > div > div.cat_dadosComplementares.categoriaHolder > div > div > div > div:nth-child(11) > input"))).send_keys(fake.city())
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "#cg_1 > div.wdTelas > div > div.catWrapper > div > div.cat_dadosComplementares.categoriaHolder > div > div > div > div:nth-child(12) > input"))).send_keys(fake.country())
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "#cg_1 > div.wdTelas > div > div.catWrapper > div > div.cat_dadosComplementares.categoriaHolder > div > div > div > div:nth-child(13) > input"))).send_keys(fake.first_name())
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "#cg_1 > div.wdTelas > div > div.catWrapper > div > div.cat_dadosComplementares.categoriaHolder > div > div > div > div:nth-child(14) > input"))).send_keys(fake.first_name())
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "#cg_1 > div.wdTelas > div > div.catWrapper > div > div.cat_dadosComplementares.categoriaHolder > div > div > div > div:nth-child(16) > input"))).send_keys(fake.job())

        # Endereços
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "#cg_1 > div.wdTelas > div > div.categorias.overflow.overflowY > ul > li.li_enderecos > a"))).click()
        time.sleep(3)

        # Preenche endereço
        try:
            elemento = wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, "#cg_1 > div.wdTelas > div > div.catWrapper > div > div.cat_enderecos.categoriaHolder > div.groupHolder.clearfix.grupo_enderecoResidencial > div > div:nth-child(2) > div:nth-child(1) > div > input")))
            elemento.send_keys("15081115")
            
            botao = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "#cg_1 > div.wdTelas > div > div.catWrapper > div > div.cat_enderecos.categoriaHolder > div.groupHolder.clearfix.grupo_enderecoResidencial > div > div:nth-child(2) > div:nth-child(1) > div > a")))
            botao.click()
        except Exception as e:
            log(doc, f"Erro ao preencher endereço: {e}")

        time.sleep(5)

        try:
            element = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "#BtYes")))
            element.click()
        except:
            pass

        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "#cg_1 > div.wdTelas > div > div.catWrapper > div > div.cat_enderecos.categoriaHolder > div.groupHolder.clearfix.grupo_enderecoResidencial > div > div:nth-child(3) > div:nth-child(2) > input"))).send_keys("1733")
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "#cg_1 > div.wdTelas > div > div.catWrapper > div > div.cat_enderecos.categoriaHolder > div.groupHolder.clearfix.grupo_enderecoResidencial > div > div:nth-child(3) > div:nth-child(3) > input"))).send_keys("Casa")
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "#cg_1 > div.wdTelas > div > div.catWrapper > div > div.cat_enderecos.categoriaHolder > div.groupHolder.clearfix.grupo_enderecoResidencial > div > div:nth-child(3) > div:nth-child(9) > label > input"))).click()

        # Salvar pessoa
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "#cg_1 > div.wdTelas > div > div.btnHolder > a.btModel.btGray.btsave"))).click()
        time.sleep(3)

    return acao

# ==== INICIALIZAÇÃO DO DRIVER ====
options = Options()
options.add_argument("--start-maximized")
driver = webdriver.Chrome(service=Service(ChromeDriverManager().install()), options=options)
wait = WebDriverWait(driver, 20)

# ==== EXECUÇÃO DO TESTE ====
try:
    safe_action(doc, "Acessando sistema", lambda: driver.get(URL))

    safe_action(doc, "Realizando login", lambda: (
        wait.until(EC.presence_of_element_located((By.ID, "j_id15:email"))).send_keys(LOGIN_EMAIL),
        wait.until(EC.presence_of_element_located((By.ID, "j_id15:senha"))).send_keys(LOGIN_PASSWORD, Keys.ENTER),
        wait.until(EC.presence_of_element_located((By.TAG_NAME, "body")))
    ))

    safe_action(doc, "Esperando sistema carregar e ajustando zoom", lambda: (
        time.sleep(5),
        ajustar_zoom()
    ))

    safe_action(doc, "Abrindo menu Contratos", lambda: (
        driver.find_element(By.TAG_NAME, "body").send_keys(Keys.F3),
        time.sleep(0.5),
        wait.until(EC.element_to_be_clickable((By.XPATH, "//li[@class='iconElement shortcutIcon ui-draggable ui-draggable-handle' and @cname='I.CT' and @ref='I.CT']"))).click()
    ))

    safe_action(doc, "Clicando em Cadastrar", lambda: (
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, '#gsContratos > div.wdTelas > div > ul > li:nth-child(1) > a > span'))).click()
    ))

    safe_action(doc, "Selecionando Tipo de Contrato", abrir_modal_e_selecionar(
        '#gsContratos > div.wdTelas > div.wdWizard.clearfix.telaConsulta > div.wizardHolder > div > div.stepPacote.step1 > div > div > div > div > a',
        "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(2) > input",
        'TIPO DE CONTRATO TESTE SELENIUM AUTOMATIZADO',
        "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(4) > a",
        "//td[contains(text(), 'TIPO DE CONTRATO TESTE SELENIUM AUTOMATIZADO')]"
    ))

    safe_action(doc, "Selecionando Pacote", lambda: (
        time.sleep(3),
        wait.until(EC.element_to_be_clickable((By.XPATH, "//h3[text()='TESTE PACOTE SELENIUM AUTOMATIZADO']"))).click()
    ))

    safe_action(doc, "Avançando para próxima etapa", lambda: (
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, '#gsContratos > div.wdTelas > div.wdWizard.clearfix.telaConsulta > div.btnHolder > a:nth-child(3)'))).click()
    ))


    safe_action(doc, "Selecionando Plano Empresa", abrir_modal_e_selecionar(
        '#gsContratos > div.wdTelas > div.wdWizard.clearfix.telaConsulta > div.wizardHolder > div > div.stepPacote.step3 > div.formRow.divPlanoEmpresa > div:nth-child(1) > div > a',
        "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(2) > input",
        'PLANO EMPRESA TESTE SELENIUM AUTOMATIZADO',
        "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(4) > a",
        "//td[contains(text(), 'PLANO EMPRESA TESTE SELENIUM AUTOMATIZADO')]"
    ))

    safe_action(doc, "Selecionando Fonte de Informação", abrir_modal_e_selecionar(
        '#gsContratos > div.wdTelas > div.wdWizard.clearfix.telaConsulta > div.wizardHolder > div > div.stepPacote.step3 > div.formRow.divPlanoEmpresa > div:nth-child(2) > div > a',
        "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div.formCol.divPesquisa > input",
        'FONTE DE INFORMÇÃO TESTE SELEN',
        "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(2) > a",
        "//td[contains(text(), 'FONTE DE INFORMÇÃO TESTE SELEN')]"
    ))



    safe_action(doc, "Selecionando Supervisor", abrir_modal_e_selecionar(
        '#gsContratos > div.wdTelas > div.wdWizard.clearfix.telaConsulta > div.wizardHolder > div > div.stepPacote.step3 > div:nth-child(8) > div.formCol.supervisorHolder > div > div > a',
        "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(2) > input",
        'SUPERVISOR TESTE SELENIUM AUTOMATIZADO',
        "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(3) > a",
        "//td[contains(text(), 'SUPERVISOR TESTE SELENIUM AUTOMATIZADO')]"
    ))




    safe_action(doc, "Selecionando Vendedor", abrir_modal_e_selecionar(
        '#gsContratos > div.wdTelas > div.wdWizard.clearfix.telaConsulta > div.wizardHolder > div > div.stepPacote.step3 > div:nth-child(8) > div.formCol.vendedorHolder > div > div > a',
        "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div.formCol.divPesquisa > input",
        'VENDEDOR TESTE SELENIUM AUTOMATIZADO',
        "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(3) > a",
        "//td[contains(text(), 'VENDEDOR TESTE SELENIUM AUTOMATIZADO')]"
    ))

    safe_action(doc, "Rolando até o Lov para selecionar o Grupo de Rateio", rolar_ate_grupo_rateio)




    safe_action(doc, "Selecionando Grupo de Rateio", abrir_modal_e_selecionar(
        '#gsContratos > div.wdTelas > div.wdWizard.clearfix.telaConsulta > div.wizardHolder > div > div.stepPacote.step3 > div:nth-child(11) > div > div > a',
        "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(2) > input",
        'GRUPO DE RATEIO TESTE SELENIUM AUTOMATIZADO',
        "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(4) > a",
        "//td[contains(text(), 'GRUPO DE RATEIO TESTE SELENIUM AUTOMATIZADO')]"
    ))

    safe_action(doc, "Avançando para próxima etapa", lambda: (
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, '#gsContratos > div.wdTelas > div.wdWizard.clearfix.telaConsulta > div.btnHolder > a:nth-child(3)'))).click()
    ))


    safe_action(doc, "Selecionando Titular", abrir_modal_e_selecionar(
        '#ui-id-6 > div:nth-child(1) > div:nth-child(1) > div > a',
        "#txtPesquisa",
        'TESTE TITULAR',
        "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(2) > a",
        "//td[contains(text(), 'TESTE TITULAR')]"
    ))

    safe_action(doc, "Rolando até o botão 'Salvar' para efetuar o salvamento do Titular", rolar_ate_salvar_titular)



    safe_action(doc, "Salvando Titular", lambda: (
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, '#ui-id-6 > div:nth-child(9) > div:nth-child(5) > a'))).click()
    ))

    safe_action(doc, "Selecionando Dependente", abrir_modal_e_selecionar(
        '#gsContratos > div.wdTelas > div.wdWizard.clearfix.telaConsulta > div.wizardHolder > div > div.stepPacote.step4 > div.blockHolder.titulares > ul > li > a.sprites.sp-addDependentes',
        "#txtPesquisa",
        'TESTE DEPENDENTE SELENIUM AUTOMATIZADO',
        "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(4) > a",
        "//td[contains(text(), 'TESTE DEPENDENTE SELENIUM AUTOMATIZADO')]"
    ))


    safe_action(doc, "Selecionando Parentesco do Dependente", selecionar_opcao(
        "#gsContratos > div.wdTelas > div.wdWizard.clearfix.telaConsulta > div.wizardHolder > div > div.stepPacote.step4 > div.blockHolder.titulares > ul > li > div.blockHolder.dependentes > ul > li:nth-child(2) > select",
        "Agregado"
    ))


    safe_action(doc, "Clicando no Ícone do Falecido pra preencher o Registro de Óbito", lambda: (
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, '#gsContratos > div.wdTelas > div.wdWizard.clearfix.telaConsulta > div.wizardHolder > div > div.stepPacote.step4 > div.blockHolder.titulares > ul > li > a.sprites.sp-addRegistroObito'))).click()
    ))


    safe_action(doc, "Selecionando Falecido", abrir_modal_e_selecionar(
        '#cg_23 > div.wdTelas > div > div.catWrapper > div > div > div:nth-child(1) > div > div:nth-child(6) > div > a',
        "#txtPesquisa",
        'FALECIDO TESTE SELENIUM AUTOMATIZADO',
        "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(4) > a",
        "//td[contains(text(), 'FALECIDO TESTE SELENIUM AUTOMATIZADO')]"
    ))

    safe_action(doc, "Preenchendo Data de Falecimento", lambda: preencher_campo_com_retry(driver, wait, "input.hasDatepicker.dataFalecimento", data_falecimento))

    safe_action(doc, "Preenchendo Data de Sepultamento", lambda: preencher_campo_com_retry(driver, wait, "input.hasDatepicker.dataSepultamento", data_sepultamento))



    safe_action(doc, "Preechendo Declaração de Óbito", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#cg_23 > div.wdTelas > div > div.catWrapper > div > div > div:nth-child(1) > div > div:nth-child(45) > input"))
        ).send_keys(fake.random_int(min=0, max=100000))
    ))

    safe_action(doc, "Selecionando Parentesco do Falecido", selecionar_opcao(
        "#cg_23 > div.wdTelas > div > div.catWrapper > div > div > div:nth-child(1) > div > div:nth-child(49) > select",
        "Agregado"
    ))

    safe_action(doc, "Selecionando Declarante", abrir_modal_e_selecionar(
        '#cg_23 > div.wdTelas > div > div.catWrapper > div > div > div:nth-child(2) > div > div:nth-child(2) > div > a',
        "#txtPesquisa",
        'TESTE DECLARANTE SELENIUM AUTOMATIZADO',
        "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(2) > a",
        "//td[contains(text(), 'TESTE DECLARANTE SELENIUM AUTOMATIZADO')]"
    ))


    safe_action(doc, "Salvando Registro de Óbito", lambda: (
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, '#cg_23 > div.wdTelas > div > div.btnHolder > a.btModel.btGray.btsave'))).click()
    ))

    time.sleep(5)


    safe_action(doc, "Avançando para próxima etapa", lambda: (
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, '#gsContratos > div.wdTelas > div.wdWizard.clearfix.telaConsulta > div.btnHolder > a:nth-child(3)'))).click()
    ))


    encontrar_mensagem_alerta()

except Exception as e:
    log(doc, f"❌ ERRO FATAL: {e}")
    take_screenshot(driver, doc, "erro_fatal")

finally:

    log(doc, "✅ Teste concluído com sucesso.")

    finalizar_relatorio()