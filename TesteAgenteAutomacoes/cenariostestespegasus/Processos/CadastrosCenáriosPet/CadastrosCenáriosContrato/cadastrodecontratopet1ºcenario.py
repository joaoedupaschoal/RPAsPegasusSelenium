
# ==== IMPORTS (sem conflitos) ====
from datetime import datetime, timedelta
from datetime import time as dt_time  # usar para objetos de hora
import time                           # usar para time.sleep(...)

from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support.ui import Select, WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service
from selenium.common.exceptions import TimeoutException, NoSuchElementException, StaleElementReferenceException, ElementClickInterceptedException
from webdriver_manager.chrome import ChromeDriverManager
from docx import Document
from docx.shared import Inches
from faker import Faker
from faker.providers import BaseProvider
from validate_docbr import CPF
from selenium.webdriver import ActionChains
import subprocess
import os
import random
import re
from selenium.webdriver.common.by import By
from selenium.webdriver.support import expected_conditions as EC
from selenium.common.exceptions import ElementClickInterceptedException, StaleElementReferenceException
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.remote.webelement import WebElement
def scroll_and_click_robusto(selector, timeout=10, by=By.CSS_SELECTOR):
    timeout = _sanitize_timeout(timeout)
    el = WebDriverWait(driver, timeout).until(EC.presence_of_element_located((by, selector)))
    driver.execute_script("arguments[0].scrollIntoView({block:'center'});", el)
    try:
        WebDriverWait(driver, timeout).until(EC.element_to_be_clickable((by, selector)))
        el.click()
    except Exception:
        driver.execute_script("arguments[0].click();", el)
# ==== PROVIDERS CUSTOMIZADOS ====
class BrasilProvider(BaseProvider):
    def rg(self):
        numeros = [str(random.randint(0, 9)) for _ in range(8)]
        return ''.join(numeros) + '-' + str(random.randint(0, 9))

fake = Faker("pt_BR")
fake.add_provider(BrasilProvider)

# ==== GERAÇÃO DE DATAS ====
def gerar_datas_validas(hora_padrao="00:00", dias_fim=0):
    """Gera datas coerentes. data_inicio/data_fim no formato 'dd/MM/yyyy HH:mm'."""
    hoje_date = datetime.today().date()
    dez_anos_atras = hoje_date - timedelta(days=3650)

    # Falecimento entre 10 anos atrás e hoje
    data_falecimento = fake.date_between(start_date=dez_anos_atras, end_date=hoje_date)

    # Nascimento (entre 18 e 110 anos antes do falecimento)
    idade_minima, idade_maxima = 18, 110
    data_nascimento = data_falecimento - timedelta(days=random.randint(idade_minima * 365, idade_maxima * 365))

    # Sepultamento 1..10 dias após o falecimento
    data_sepultamento = data_falecimento + timedelta(days=random.randint(1, 10))

    # Registro 1..10 dias após o sepultamento
    data_registro = data_sepultamento + timedelta(days=random.randint(1, 10))

    # Velório entre o falecimento e o sepultamento
    data_velorio = fake.date_between(start_date=data_falecimento, end_date=data_sepultamento)

    # Início entre 2 e 30 dias no futuro, com hora escolhida
    data_inicio_date = hoje_date + timedelta(days=random.randint(2, 30))

    # Monta datetime com hora escolhida (ex: "00:00")
    h, m = map(int, hora_padrao.split(":"))
    dt_inicio = datetime.combine(data_inicio_date, dt_time(h, m))

    # Fim: mesmo dia por padrão (dias_fim=0). Ajuste se quiser +N dias.
    dt_fim = dt_inicio + timedelta(days=dias_fim)

    fmt_data = "%d/%m/%Y"
    fmt_dt = "%d/%m/%Y %H:%M"

    return (
        data_nascimento.strftime(fmt_data),
        data_falecimento.strftime(fmt_data),
        data_sepultamento.strftime(fmt_data),
        data_velorio.strftime(fmt_data),
        dt_inicio.strftime(fmt_dt),   # data_inicio com hora
        dt_fim.strftime(fmt_dt),      # data_fim com hora
        data_registro.strftime(fmt_data),
        hoje_date.strftime(fmt_data),
    )

(data_nascimento, data_falecimento, data_sepultamento,
 data_velorio, data_inicio, data_fim, data_registro, hoje) = gerar_datas_validas(
    hora_padrao="08:50",
    dias_fim=0
)

hora_falecimento = fake.time(pattern="%H:%M")
hora_sepultamento = fake.time(pattern="%H:%M")
localizacao = fake.city()

# ==== CONFIGURAÇÕES ====
URL = "http://localhost:8080/gs/index.xhtml"
LOGIN_EMAIL = "joaoeduardo.gold@outlook.com"
LOGIN_PASSWORD = "071999gs"

# ==== INICIALIZAÇÃO DE VARIÁVEIS GLOBAIS ====
doc = Document()
doc.add_heading("RELATÓRIO DO TESTE", 0)
doc.add_paragraph("Cadastro de Contrato Pet – Cenário 1: Preenchimento completo e salvamento.")
doc.add_paragraph(f"Data do teste: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")

screenshot_registradas = set()
driver = None
wait = None

# ==== UTILITÁRIOS ====
def log(doc, msg):
    print(msg)
    doc.add_paragraph(msg)

def _sanitize_filename(name: str) -> str:
    name = name.strip().lower()
    name = re.sub(r"[<>:\"/\\|?*']", "_", name)
    name = re.sub(r"_+", "_", name)
    return name[:120]

def take_screenshot(driver, doc, nome):
    if driver is None:
        return
    nome = _sanitize_filename(nome)
    if nome not in screenshot_registradas:
        path = f"screenshots/{nome}.png"
        os.makedirs("screenshots", exist_ok=True)
        try:
            driver.save_screenshot(path)
            doc.add_paragraph(f"Screenshot: {nome}")
            doc.add_picture(path, width=Inches(5.5))
            screenshot_registradas.add(nome)
        except Exception as e:
            log(doc, f"⚠️ Erro ao tirar screenshot {nome}: {e}")

import traceback

def safe_action(doc, descricao, func):
    try:
        log(doc, f"🔄 {descricao}...")
        func()
        log(doc, f"✅ {descricao} realizada com sucesso.")
        take_screenshot(driver, doc, descricao.lower().replace(" ", "_"))
    except Exception as e:
        take_screenshot(driver, doc, f"erro_{descricao.lower().replace(' ', '_')}")
        log(doc, f"❌ {descricao} falhou: {type(e).__name__}: {e}")
        log(doc, "— stacktrace —")
        log(doc, traceback.format_exc())

def finalizar_cadastro():
    seletor_css_finalizar = '#gsPet > div.wdTelas > div.wdWizard.clearfix.telaConsulta > div.btnHolder > div'
    clicar_elemento_robusto(driver, wait, seletor_css_finalizar)
    time.sleep(6)


def abrir_modal_e_selecionar_robusto(btn_selector, pesquisa_selector, termo_pesquisa, btn_pesquisar_selector, resultado_xpath):
    """Versão robusta da função de modal"""
    global driver, wait, doc
    
    def acao():
        if driver is None or wait is None:
            raise Exception("Driver ou wait não inicializados")
            
        # Abre o modal
        safe_scroll_and_interact(btn_selector, "click")
        time.sleep(1)

        # Aguarda e preenche campo pesquisa
        campo_pesquisa = wait.until(
            EC.element_to_be_clickable((By.CSS_SELECTOR, pesquisa_selector))
        )
        campo_pesquisa.clear()
        campo_pesquisa.send_keys(termo_pesquisa)
        time.sleep(0.5)

        # Clica pesquisar
        pesquisar = wait.until(
            EC.element_to_be_clickable((By.CSS_SELECTOR, btn_pesquisar_selector))
        )
        pesquisar.click()
        time.sleep(2)
        
        # Aguarda resultado e clica
        resultado = wait.until(
            EC.element_to_be_clickable((By.XPATH, resultado_xpath))
        )
        driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", resultado)
        time.sleep(0.5)
        resultado.click()
        time.sleep(1)

    return acao


def finalizar_relatorio():
    global driver, doc
    nome_arquivo = f"relatorio_cadastro_contrato_pet_cenario_1_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    try:
        doc.save(nome_arquivo)
        log(doc, f"📄 Relatório salvo como: {nome_arquivo}")
        try:
            subprocess.run(["start", "winword", nome_arquivo], shell=True)
        except:
            pass
    except Exception as e:
        print(f"Erro ao salvar relatório: {e}")
    if driver:
        try:
            driver.quit()
        except:
            pass

def encontrar_mensagem_alerta():
    global driver, doc
    if driver is None:
        return None
    seletores = [
        (".alerts.salvo", "✅ Sucesso"),
        (".alerts.alerta", "⚠️ Alerta"),
        (".alerts.erro", "❌ Erro"),
    ]
    for seletor, tipo in seletores:
        try:
            elemento = driver.find_element(By.CSS_SELECTOR, seletor)
            if elemento.is_displayed():
                log(doc, f"📢 {tipo}: {elemento.text}")
                return elemento
        except:
            continue
    log(doc, "ℹ️ Nenhuma mensagem de alerta encontrada.")
    return None

def ajustar_zoom():
    global driver, doc
    if driver is None:
        return
    try:
        driver.execute_script("document.body.style.zoom='90%'")
        log(doc, "🔍 Zoom ajustado para 90%.")
    except Exception as e:
        log(doc, f"⚠️ Erro ao ajustar zoom: {e}")

from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.common.by import By
from selenium.webdriver.support import expected_conditions as EC

def _sanitize_timeout(t):
    # garante número > 0
    return 10 if not isinstance(t, (int, float)) or t <= 0 else t

def safe_scroll_and_interact(selector, action_type="click", value=None, timeout=10, by_xpath=False):
    timeout = _sanitize_timeout(timeout)
    by_type = By.XPATH if by_xpath else By.CSS_SELECTOR

    element = WebDriverWait(driver, timeout).until(EC.presence_of_element_located((by_type, selector)))
    driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", element)

    if action_type in ("click", "send_keys"):
        WebDriverWait(driver, timeout).until(EC.element_to_be_clickable((by_type, selector)))

    if action_type == "click":
        try:
            element.click()
        except Exception:
            driver.execute_script("arguments[0].click();", element)
    elif action_type == "send_keys":
        element.clear()
        element.send_keys("" if value is None else str(value))
    elif action_type == "select" and value is not None:
        Select(element).select_by_visible_text(str(value))
    return element

# ==== CLICAR ROBUSTO ====
def clicar_elemento_robusto(driver, wait, seletor_css, timeout=10):
    global doc
    try:
        elem = wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, seletor_css)))
        try:
            driver.execute_script("""
                document.querySelectorAll('.modal, .overlay, .blockUI, .toast, .tooltip, [role="dialog"], [data-overlay]')
                .forEach(e => { if (getComputedStyle(e).position === 'fixed') e.style.display = 'none'; });
            """)
        except Exception:
            pass
        driver.execute_script("arguments[0].scrollIntoView({block:'center', inline:'center'});", elem)
        time.sleep(0.2)
        try:
            elem = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, seletor_css)))
            elem.click()
            return True
        except (TimeoutException, ElementClickInterceptedException, StaleElementReferenceException):
            pass
        try:
            elem = driver.find_element(By.CSS_SELECTOR, seletor_css)
            ActionChains(driver).move_to_element(elem).pause(0.05).click().perform()
            return True
        except Exception:
            pass
        try:
            elem = driver.find_element(By.CSS_SELECTOR, seletor_css)
            driver.execute_script("arguments[0].click();", elem)
            return True
        except Exception:
            pass
        try:
            elem = driver.find_element(By.CSS_SELECTOR, seletor_css)
            driver.execute_script("""
                const el = arguments[0];
                function fire(type){ el.dispatchEvent(new MouseEvent(type,{bubbles:true,cancelable:true,view:window})); }
                el.focus(); fire('mouseover'); fire('mousedown'); fire('mouseup'); fire('click');
            """, elem)
            return True
        except Exception:
            pass
        try:
            elem = driver.find_element(By.CSS_SELECTOR, seletor_css)
            ActionChains(driver).move_to_element_with_offset(elem, 1, 1).click().perform()
            return True
        except Exception:
            pass
        log(doc, f"❌ Não foi possível clicar em: {seletor_css}")
        return False
    except Exception as e:
        log(doc, f"❌ Erro ao clicar robusto: {e}")
        return False

# ==== DATEPICKER ====
def encontrar_campos_datepicker():
    global driver
    seletores = [
        "input.hasDatepicker",
        "input[id^='dp']",
        "input[class*='datepicker']",
        ".hasDatepicker",
        "[data-provide='datepicker']"
    ]
    campos = []
    for seletor in seletores:
        try:
            for elemento in driver.find_elements(By.CSS_SELECTOR, seletor):
                if elemento.is_displayed():
                    cid = elemento.get_attribute('id')
                    if not any(c.get('id') == cid for c in campos):
                        campos.append({'elemento': elemento, 'id': cid})
        except:
            continue
    log(doc, f"📊 Encontrados {len(campos)} campos datepicker")
    return campos

def _datepicker_jquery(campo_id, data_valor):
    global driver
    resultado = driver.execute_script("""
        var campoId = arguments[0], valor = arguments[1];
        if (typeof jQuery === 'undefined') return 'jQuery não disponível';
        var $campo = $('#' + campoId);
        if (!$campo.length) return 'Campo não encontrado: ' + campoId;
        try {
            if ($campo.hasClass('hasDatepicker')) { $campo.datepicker('setDate', valor); }
            else { $campo.val(valor); }
            $campo.trigger('input').trigger('change').trigger('blur');
            return $campo.val();
        } catch(e) { return 'Erro: ' + e.message; }
    """, campo_id, data_valor)
    if isinstance(resultado, str) and 'Erro' in resultado:
        raise Exception(f"jQuery falhou: {resultado}")

def _datepicker_javascript(elemento, data_valor):
    global driver
    driver.execute_script("""
        var campo = arguments[0], valor = arguments[1];
        campo.focus(); campo.value = ''; campo.value = valor;
        ['input','change','blur','keyup'].forEach(ev => campo.dispatchEvent(new Event(ev,{bubbles:true})));
    """, elemento, data_valor)

def _datepicker_actionchains(elemento, data_valor):
    global driver
    driver.execute_script("arguments[0].scrollIntoView({block:'center'});", elemento)
    ActionChains(driver).move_to_element(elemento).click().perform()
    time.sleep(0.3)
    ActionChains(driver).key_down(Keys.CONTROL).send_keys('a').key_up(Keys.CONTROL).perform()
    ActionChains(driver).send_keys(Keys.DELETE).perform()
    time.sleep(0.2)
    for ch in data_valor:
        ActionChains(driver).send_keys(ch).perform()
        time.sleep(0.03)
    ActionChains(driver).send_keys(Keys.TAB).perform()

def _datepicker_tradicional(elemento, data_valor):
    driver.execute_script("arguments[0].scrollIntoView({block:'center'});", elemento)
    elemento.click(); time.sleep(0.2); elemento.clear()
    elemento.send_keys(data_valor); elemento.send_keys(Keys.TAB)

def preencher_datepicker_persistente(indice_campo, data_valor, max_tentativas=10, timeout=30):
    """Preenche datepicker com várias estratégias e valida."""
    inicio_tempo = time.time()
    tentativa = 0

    def validar_data_preenchida(el, data_esperada):
        try:
            val = (el.get_attribute('value') or '').strip()
            if not val: return False
            if val == data_esperada or data_esperada in val: return True
            formatos = [
                '%d/%m/%Y %H:%M', '%d/%m/%Y %H:%M:%S',  # <- COM HORA
                '%d/%m/%Y', '%m/%d/%Y', '%Y-%m-%d', '%d-%m-%Y'
            ]
            for f in formatos:
                try:
                    d1 = datetime.strptime(val, f)
                    d2 = datetime.strptime(data_esperada, f)
                    if d1 == d2: return True
                except:
                    continue
            return False
        except:
            return False

    while tentativa < max_tentativas and (time.time() - inicio_tempo) < timeout:
        tentativa += 1
        try:
            log(doc, f"🔄 Tentativa {tentativa}/{max_tentativas} para campo {indice_campo}")
            campos = encontrar_campos_datepicker()
            if not campos:
                time.sleep(1); continue
            if indice_campo >= len(campos):
                time.sleep(1); continue

            info = campos[indice_campo]
            elemento, campo_id = info['elemento'], info['id']

            if validar_data_preenchida(elemento, data_valor):
                log(doc, f"✅ Campo {indice_campo} já está preenchido corretamente!")
                return True

            estrategias = [
                lambda: _datepicker_jquery(campo_id, data_valor),
                lambda: _datepicker_javascript(elemento, data_valor),
                lambda: _datepicker_actionchains(elemento, data_valor),
                lambda: _datepicker_tradicional(elemento, data_valor),
            ]
            for acao in estrategias:
                try:
                    acao(); time.sleep(0.4)
                    if validar_data_preenchida(elemento, data_valor):
                        log(doc, f"✅ Campo {indice_campo} preenchido!")
                        return True
                except Exception as e:
                    log(doc, f"   ⚠️ Estratégia falhou: {e}")

            log(doc, "❌ Tentativa falhou; tentando novamente...")
            time.sleep(1.2)

        except Exception as e:
            log(doc, f"❌ Erro na tentativa {tentativa}: {e}")
            time.sleep(1)

    raise Exception(f"Falha ao preencher datepicker {indice_campo} após {tentativa} tentativas em {int(time.time()-inicio_tempo)}s")

def preencher_campo_xpath_com_retry(driver, wait, xpath, valor, max_tentativas=3):
    global doc
    if driver is None or wait is None:
        return False
    for tentativa in range(max_tentativas):
        try:
            campo = wait.until(EC.element_to_be_clickable((By.XPATH, xpath)))
            driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", campo)
            time.sleep(0.3)
            if tentativa == 0:
                campo.click(); campo.clear(); campo.send_keys(valor); campo.send_keys(Keys.TAB)
            elif tentativa == 1:
                ActionChains(driver).move_to_element(campo).click().pause(0.1).key_down(Keys.CONTROL).send_keys('a').key_up(Keys.CONTROL).send_keys(valor).send_keys(Keys.TAB).perform()
            else:
                driver.execute_script("""
                    var el = arguments[0], v = arguments[1];
                    el.focus(); el.value = ''; el.value = v;
                    el.dispatchEvent(new Event('input', { bubbles: true }));
                    el.dispatchEvent(new Event('change', { bubbles: true }));
                    el.blur();
                """, campo, valor)
            time.sleep(0.3)
            if (campo.get_attribute('value') or '').strip():
                return True
        except Exception as e:
            log(doc, f"⚠️ Tentativa {tentativa + 1} falhou: {e}")
            time.sleep(0.8)
    return False

# ==== WIZARD (clique por ícone) ====
def get_xpath(we):
    js = """
    function absoluteXPath(element) {
      if (element.tagName.toLowerCase() === 'html') return '/html';
      if (element===document.body) return '/html/body';
      var ix=0; var siblings=element.parentNode.children;
      var same = 0;
      for (var i=0; i<siblings.length; i++) {
        var sib=siblings[i];
        if (sib.tagName===element.tagName) {
          same++;
          if (sib===element) {
            var path = absoluteXPath(element.parentNode) + '/' + element.tagName.toLowerCase();
            if (same>1) path += '['+same+']';
            return path;
          }
        }
      }
      return absoluteXPath(element.parentNode) + '/' + element.tagName.toLowerCase();
    }
    return absoluteXPath(arguments[0]);
    """
    return we.parent.execute_script(js, we)

def click_wizard_by_icon(driver, wait, icon_class, expect_selector=None, timeout=12):
    span = wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, f"#gsCrm .btnHolder .sprites.{icon_class}")))
    driver.execute_script("arguments[0].scrollIntoView({block:'center'});", span)
    alvo = driver.execute_script("""
        const span = arguments[0]; let el = span;
        while (el && el !== document && !/^(A|DIV)$/i.test(el.tagName)) el = el.parentElement;
        return el;
    """, span)
    if not alvo:
        raise Exception(f"Não achei ancestral clicável para o ícone .{icon_class}")
    try:
        xp = get_xpath(alvo)
        wait.until(EC.element_to_be_clickable((By.XPATH, xp))).click()
    except Exception:
        try:
            driver.execute_script("arguments[0].click();", alvo)
        except Exception:
            driver.execute_script("""
                const el=arguments[0];
                function fire(t){el.dispatchEvent(new MouseEvent(t,{bubbles:true,cancelable:true,view:window}));}
                el.focus(); fire('mouseover'); fire('mousedown'); fire('mouseup'); fire('click');
            """, alvo)
    if expect_selector:
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, expect_selector)))




# ==== DRIVER ====
def inicializar_driver():
    global driver, wait
    try:
        options = Options()
        options.add_argument("--start-maximized")
        options.add_argument("--disable-blink-features=AutomationControlled")
        options.add_experimental_option("excludeSwitches", ["enable-automation"])
        options.add_experimental_option('useAutomationExtension', False)

        driver = webdriver.Chrome(service=Service(ChromeDriverManager().install()), options=options)
        driver.execute_script("Object.defineProperty(navigator, 'webdriver', {get: () => undefined})")
        wait = WebDriverWait(driver, 20)
        return True
    except Exception as e:
        log(doc, f"❌ Erro ao inicializar driver: {e}")
        return False

# ==== EXECUÇÃO DO TESTE ====
def executar_teste():
    global driver, wait, doc
    try:
        if not inicializar_driver():
            return False

        # Sanidade: garantir que time é módulo
        assert hasattr(time, "sleep"), f"time virou {time!r}"

        safe_action(doc, "Acessando sistema", lambda: driver.get(URL))

        safe_action(doc, "Realizando login", lambda: (
            wait.until(EC.presence_of_element_located((By.ID, "j_id15:email"))).send_keys(LOGIN_EMAIL),
            wait.until(EC.presence_of_element_located((By.ID, "j_id15:senha"))).send_keys(LOGIN_PASSWORD, Keys.ENTER),
            time.sleep(5)
        ))

        safe_action(doc, "Ajustando zoom e abrindo menu", lambda: (
            ajustar_zoom(),
            driver.find_element(By.TAG_NAME, "body").send_keys(Keys.F3)
        ))

        safe_action(doc, "Acessando PET", lambda: (
            driver.execute_script("arguments[0].scrollIntoView({block: 'center'});",
                                  wait.until(EC.presence_of_element_located((By.XPATH, '/html/body/div[15]/ul/li[39]/img')))),
            wait.until(EC.element_to_be_clickable((By.XPATH, '/html/body/div[15]/ul/li[39]/img'))).click()
        ))

        safe_action(doc, "Clicando em Cadastrar Contrato", lambda:
            wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, '#gsPet > div.wdTelas > div > ul > li:nth-child(1) > a > span'))).click()
        )

        time.sleep(2)



        safe_action(doc, "Selecionando Pacote", lambda: (
            time.sleep(3),
            wait.until(EC.element_to_be_clickable((By.XPATH, "//h3[text()='PACOTE PET TESTE SELENIUM AUTOMATIZADO']"))).click()
        ))



        safe_action(doc, "Avançando para aba: 'Dados do Contrato'", lambda:
            clicar_elemento_robusto(driver, wait, '#gsPet > div.wdTelas > div.wdWizard.clearfix.telaConsulta > div.btnHolder > a:nth-child(3)')
        )


        safe_action(doc, "Selecionando Supervisor", abrir_modal_e_selecionar_robusto(
            '#gsPet > div.wdTelas > div.wdWizard.clearfix.telaConsulta > div.wizardHolder > div > div.stepDadosContrato.step2 > div.formRow > div.formCol.supervisorHolder > div > div > a',
            'body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(2) > input',
            '153',
            'body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(3) > a',
            "//td[contains(text(), 'SUPERVISOR TESTE SELENIUM AUTOMATIZADO')]"
        ))

        safe_action(doc, "Selecionando Vendedor", abrir_modal_e_selecionar_robusto(
            '#gsPet > div.wdTelas > div.wdWizard.clearfix.telaConsulta > div.wizardHolder > div > div.stepDadosContrato.step2 > div.formRow > div.formCol.vendedorHolder > div > div > a',
            'body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div.formCol.divPesquisa > input',
            'VENDEDOR TESTE SELENIUM AUTOMATIZADO',
            'body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(3) > a',
            "//td[contains(text(), 'VENDEDOR TESTE SELENIUM AUTOMATIZADO')]"
        ))

        safe_action(doc, "Avançando para aba: 'Tutor'", lambda:
            clicar_elemento_robusto(driver, wait, '#gsPet > div.wdTelas > div.wdWizard.clearfix.telaConsulta > div.btnHolder > a:nth-child(3)')
        )
        time.sleep(2)

        safe_action(doc, "Selecionando Tutor", abrir_modal_e_selecionar_robusto(
            '#ui-id-2 > div:nth-child(1) > div:nth-child(1) > div > a',
            '#txtPesquisa',
            'TESTE TITULAR 233',
            'body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(2) > a',
            "//td[contains(text(), 'TESTE TITULAR 233')]"
        ))

        safe_action(doc, "Salvando Tutor", lambda: scroll_and_click_robusto('#ui-id-2 > div.formRow.fRight > div > a'))


        safe_action(doc, "Selecionando Pet", abrir_modal_e_selecionar_robusto(
            '#gsPet > div.wdTelas > div.wdWizard.clearfix.telaConsulta > div.wizardHolder > div > div.stepTutor.step3 > div.blockHolder.titulares > ul > li > a',
            'body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div.formCol.divPesquisa > input',
            'PET SELENIUM',
            'body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(3) > a',
            "//td[contains(text(), 'PET SELENIUM')]"
        ))

        safe_action(doc, "Avançando para aba: 'Pagamento'", lambda:
            clicar_elemento_robusto(driver, wait, '#gsPet > div.wdTelas > div.wdWizard.clearfix.telaConsulta > div.btnHolder > a:nth-child(3)')
        )

        safe_action(doc, "Avançando para aba: 'Resumo'", lambda:
            clicar_elemento_robusto(driver, wait, '#gsPet > div.wdTelas > div.wdWizard.clearfix.telaConsulta > div.btnHolder > a:nth-child(3)')
        )


        safe_action(doc, "Finalizando cadastro", finalizar_cadastro)





        safe_action(doc, "Fechando modal após o salvamento", lambda:
            clicar_elemento_robusto(driver, wait, '#gsPet > div.wdTop.ui-draggable-handle > div > a')
        )


        log(doc, "🔍 Verificando mensagens de alerta...")
        encontrar_mensagem_alerta()

        return True

    except Exception as e:
        log(doc, f"❌ ERRO FATAL: {e}")
        take_screenshot(driver, doc, "erro_fatal")
        return False
    finally:
        log(doc, "✅ Teste concluído.")

# ==== MAIN ====
def main():
    global doc
    try:
        log(doc, "🚀 Iniciando teste de cadastro de Contrato Pet")
        sucesso = executar_teste()
        if sucesso:
            log(doc, "✅ Teste executado com sucesso!")
        else:
            log(doc, "❌ Teste finalizado com erros.")
    except Exception as e:
        log(doc, f"❌ Erro na execução principal: {e}")
    finally:
        finalizar_relatorio()

if __name__ == "__main__":
    main()





