# Código corrigido: cadastro_controle_cremacao_cenario_1.py

from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support.ui import Select, WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service
from webdriver_manager.chrome import ChromeDriverManager
from docx import Document
from docx.shared import Inches
from faker import Faker
from faker.providers import BaseProvider
from validate_docbr import CPF
from datetime import datetime, timedelta
import subprocess
import os
import time
import random
import sys
import re

sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))

# ==== CONFIGURAÇÕES ====
URL = "http://localhost:8080/gs/index.xhtml"
LOGIN_EMAIL = "joaoeduardo.gold@outlook.com"
LOGIN_PASSWORD = "071999gs"

# ==== PROVIDERS CUSTOMIZADOS ====
class BrasilProvider(BaseProvider):
    def rg(self):
        numeros = [str(random.randint(0, 9)) for _ in range(8)]
        return ''.join(numeros) + '-' + str(random.randint(0, 9))

fake = Faker("pt_BR")
fake.add_provider(BrasilProvider)

# Controle de screenshots únicas
screenshot_registradas = set()

# ==== FUNÇÕES DE UTILITÁRIO ====
def log(doc, msg):
    """Registra mensagem no console e no documento"""
    print(msg)
    doc.add_paragraph(msg)

def sanitize_filename(nome):
    """Sanitiza nome do arquivo para evitar caracteres inválidos"""
    nome = nome.lower()
    nome = re.sub(r'[^a-z0-9_]', '_', nome)
    return nome

def take_screenshot(driver, doc, nome):
    """Captura screenshot e adiciona ao documento"""
    nome_limpo = sanitize_filename(nome)
    if nome_limpo not in screenshot_registradas:
        os.makedirs("screenshots", exist_ok=True)
        path = f"screenshots/{nome_limpo}.png"
        driver.save_screenshot(path)
        doc.add_paragraph(f"Screenshot: {nome}")
        doc.add_picture(path, width=Inches(5.5))
        screenshot_registradas.add(nome_limpo)

def safe_action(doc, descricao, func, driver, wait):
    """Executa ação com tratamento de erro e screenshot"""
    try:
        log(doc, f"🔄 {descricao}...")
        func()
        log(doc, f"✅ {descricao} realizada com sucesso.")
        take_screenshot(driver, doc, descricao.lower().replace(" ", "_"))
        return True, "sucesso"
    except Exception as e:
        log(doc, f"❌ Erro ao {descricao.lower()}: {e}")
        take_screenshot(driver, doc, f"erro_{descricao.lower().replace(' ', '_')}")
        return False, "erro"

def registrar_screenshot_unico(nome, driver, doc, descricao=None):
    """Registra screenshot único"""
    if nome not in screenshot_registradas:
        if descricao:
            log(doc, f"📸 {descricao}")
        take_screenshot(driver, doc, nome)
        screenshot_registradas.add(nome)

def encontrar_mensagem_alerta(driver, doc):
    """Procura por mensagens de alerta no sistema"""
    seletores = [
        (".alerts.salvo", "sucesso"),
        (".alerts.alerta", "alerta"),
        (".alerts.erro", "erro"),
    ]

    for seletor, tipo in seletores:
        try:
            elemento = driver.find_element(By.CSS_SELECTOR, seletor)
            if elemento.is_displayed():
                if tipo == "sucesso":
                    log(doc, f"📢 ✅ SUCESSO: {elemento.text}")
                elif tipo == "alerta":
                    log(doc, f"📢 ⚠️ ALERTA: {elemento.text}")
                else:
                    log(doc, f"📢 ❌ ERRO: {elemento.text}")
                return elemento, tipo
        except:
            continue

    log(doc, "ℹ️ Nenhuma mensagem de alerta encontrada.")
    return None, None

def ajustar_zoom(driver):
    """Ajusta o zoom da página para melhor visualização"""
    try:
        driver.execute_script("document.body.style.zoom='90%'")
        time.sleep(1)
    except Exception as e:
        raise Exception(f"Erro ao ajustar zoom: {e}")

def gerar_datas_validas():
    """Gera datas coerentes para admissão, início e fim da escala, e vencimento da CNH."""
    hoje = datetime.today().date()
    
    # Data de admissão entre 10 anos atrás e hoje
    data_admissao = fake.date_between(start_date=hoje - timedelta(days=3650), end_date=hoje)
    
    # Data de início da escala entre hoje e 1 ano no futuro
    data_inicio = fake.date_between(start_date=hoje, end_date=hoje + timedelta(days=365))
    
    # Data fim entre 1 e 180 dias após a data de início
    data_fim = data_inicio + timedelta(days=random.randint(1, 180))
    
    # Vencimento CNH entre hoje e 10 anos no futuro
    vencimento_cnh = fake.date_between(start_date=hoje, end_date=hoje + timedelta(days=3650))
    
    return (data_admissao.strftime('%d/%m/%Y'), 
            data_inicio.strftime('%d/%m/%Y'), 
            data_fim.strftime('%d/%m/%Y'), 
            vencimento_cnh.strftime('%d/%m/%Y'))

def gerar_dados_documentos():
    """Gera documentos fictícios para o cadastro."""
    carteira_trabalho = str(random.randint(10000000, 99999999))
    pis = fake.cpf().replace('.', '').replace('-', '')[:11]
    cnh = str(random.randint(10000000000, 99999999999))
    cpf = CPF().generate()
    
    return carteira_trabalho, pis, cnh, cpf

def main():
    """Função principal do teste"""
    global screenshot_registradas
    screenshot_registradas = set()
    
    # Inicialização do documento
    doc = Document()
    doc.add_heading("RELATÓRIO DO TESTE", 0)
    doc.add_paragraph("Cadastro de Controle de Cremação - 2º Cenário.")
    doc.add_paragraph(f"🗓️ Data do teste: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")
    doc.add_paragraph("Neste teste, o robô preencherá todos os campos e cancelará o cadastro de um novo Controle de Cremação.")

    # Configuração do Chrome
    chrome_options = Options()
    chrome_options.add_argument("--start-maximized")
    driver = webdriver.Chrome(service=Service(ChromeDriverManager().install()), options=chrome_options)
    wait = WebDriverWait(driver, 10)

    def finalizar_relatorio():
        """Finaliza e salva o relatório"""
        doc_name = f"relatorio_controle_cremacao_cenario_2_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        doc.save(doc_name)
        log(doc, f"📄 Relatório salvo como: {doc_name}")
        try:
            subprocess.run(["start", "winword", doc_name], shell=True)
        except Exception as e:
            log(doc, f"Erro ao abrir o Word: {e}")
        driver.quit()

    def login():
        """Realiza login no sistema"""
        wait.until(EC.presence_of_element_located((By.ID, "j_id15:email"))).send_keys(LOGIN_EMAIL)
        wait.until(EC.presence_of_element_located((By.ID, "j_id15:senha"))).send_keys(LOGIN_PASSWORD, Keys.ENTER)
        wait.until(EC.presence_of_element_located((By.TAG_NAME, "body")))
        time.sleep(4)

    def abrir_menu():
        """Abre o menu de Controle Cremação"""
        driver.find_element(By.TAG_NAME, "body").send_keys(Keys.F2)
        campo = wait.until(EC.presence_of_element_located((By.XPATH, "//input[@placeholder='Busque um cadastro']")))
        campo.click()
        campo.send_keys("Controle Cremação", Keys.ENTER)
        time.sleep(3)

    def acessar_formulario():
        """Acessa o formulário de cadastro"""
        cadastrar = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, 
            "#fmod_200010 > div.wdTelas > div.telaInicial.clearfix.overflow.overflowY > ul > li:nth-child(1) > a > span")))
        cadastrar.click()
        time.sleep(2)

    def abrir_lov_falecido():
        """Abre LOV de Falecido"""
        open_lov_falecido = wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, 
            "#fmod_200010 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div > div > div > div:nth-child(2) > div > a")))
        open_lov_falecido.click()
        time.sleep(2)

    def pesquisar_e_selecionar_falecido():
        """Pesquisa e seleciona falecido"""
        campo_pesquisa = wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#txtPesquisa")))
        campo_pesquisa.clear()
        campo_pesquisa.send_keys('GUSTAVO VIEIRA', Keys.ENTER)
        time.sleep(3)
        
        # Aguarda o resultado da pesquisa e clica
        falecido = wait.until(EC.element_to_be_clickable((By.XPATH, "//td[contains(text(), 'GUSTAVO VIEIRA')]")))
        falecido.click()
        time.sleep(2)

    def preencher_numero_os():
        """Preenche campo Número OS"""
        numero_OS = wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, 
            "#fmod_200010 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div > div > div > div:nth-child(3) > input")))
        numero_OS.clear()
        numero_OS.send_keys('201680')
        time.sleep(1)

    def selecionar_status_aguardando():
        """Seleciona status 'Aguardando'"""
        select_aguardando = wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR,
            "#fmod_200010 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div > div > div > div:nth-child(4) > select")))
        Select(select_aguardando).select_by_visible_text("Aguardando")
        time.sleep(1)

    def selecionar_forno():
        """Seleciona 'Forno 1'"""
        select_forno = wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR,
            "#fmod_200010 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div > div > div > div:nth-child(5) > select")))
        Select(select_forno).select_by_visible_text("Forno 1")
        time.sleep(1)


    def cancelar():
        cancelar_btn = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR,
            "#fmod_200010 > div.wdTelas > div.telaCadastro.clearfix > div.btnHolder > a.btModel.btGray.btcancel")))
        cancelar_btn.click()
        time.sleep(2)



    def confirmar_cancelamento():
        confirmar_btn = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR,
            "#BtYes")))
        confirmar_btn.click()
        time.sleep(2)         

    def fechar_modal():
        """Fecha o modal"""
        x_btn = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR,
            "#fmod_200010 > div.wdTop.ui-draggable-handle > div.wdClose > a")))
        x_btn.click()
        time.sleep(1)

    # ==== EXECUÇÃO DO TESTE ====
    try:
        # 1. Acessar o sistema
        if not safe_action(doc, "Acessando o sistema", lambda: driver.get(URL), driver, wait)[0]:
            finalizar_relatorio()
            return
        registrar_screenshot_unico("url_acessada", driver, doc, "Sistema acessado.")

        # 2. Realizar login
        if not safe_action(doc, "Realizando login", login, driver, wait)[0]:
            finalizar_relatorio()
            return
        registrar_screenshot_unico("login_concluido", driver, doc, "Login realizado.")

        # 3. Ajustar zoom
        if not safe_action(doc, "Ajustando zoom", lambda: ajustar_zoom(driver), driver, wait)[0]:
            finalizar_relatorio()
            return

        # 4. Abrir menu
        if not safe_action(doc, "Abrindo menu Controle Cremação", abrir_menu, driver, wait)[0]:
            finalizar_relatorio()
            return
        registrar_screenshot_unico("menu_aberto", driver, doc, "Menu 'Controle Cremação' aberto.")

        # 5. Acessar formulário
        if not safe_action(doc, "Acessando formulário de cadastro", acessar_formulario, driver, wait)[0]:
            finalizar_relatorio()
            return
        registrar_screenshot_unico("formulario_aberto", driver, doc, "Formulário de cadastro aberto.")

        # 6. Preenchimento dos campos
        log(doc, "🔧 Iniciando preenchimento dos campos...")

        if not safe_action(doc, "Abrindo LOV de Falecido", abrir_lov_falecido, driver, wait)[0]:
            finalizar_relatorio()
            return

        if not safe_action(doc, "Pesquisando e selecionando falecido", pesquisar_e_selecionar_falecido, driver, wait)[0]:
            finalizar_relatorio()
            return

        if not safe_action(doc, "Preenchendo Número OS", preencher_numero_os, driver, wait)[0]:
            finalizar_relatorio()
            return

        if not safe_action(doc, "Selecionando status Aguardando", selecionar_status_aguardando, driver, wait)[0]:
            finalizar_relatorio()
            return

        if not safe_action(doc, "Selecionando Forno", selecionar_forno, driver, wait)[0]:
            finalizar_relatorio()
            return

        registrar_screenshot_unico("campos_preenchidos", driver, doc, "Todos os campos preenchidos.")

        # SALVANDO O CADASTRO
        if not safe_action(doc, "Clicando no botão Cancelar", cancelar, driver, wait)[0]:
            finalizar_relatorio()
            return
        registrar_screenshot_unico("apos_cancelar", driver, doc, "Clique no botão Cancelar realizado.")

        if not safe_action(doc, "Clicando no botão Sim", confirmar_cancelamento, driver, wait)[0]:
            finalizar_relatorio()
            return
        registrar_screenshot_unico("apos_cancelar", driver, doc, "Clique no botão Sim realizado.")


        # 8. Verificar mensagem de retorno
        time.sleep(2)
        _, tipo_alerta = encontrar_mensagem_alerta(driver, doc)
        if tipo_alerta == "sucesso":
            log(doc, "✅ Mensagem de sucesso exibida.")
        elif tipo_alerta == "alerta":
            log(doc, "⚠️ Mensagem de alerta exibida.")
        elif tipo_alerta == "erro":
            log(doc, "❌ Mensagem de erro exibida.")
        else:
            log(doc, "⚠️ Nenhuma mensagem encontrada após salvar.")
        registrar_screenshot_unico("mensagem_final", driver, doc, "Mensagem exibida após salvar.")

        # 9. Fechar formulário
        if not safe_action(doc, "Fechando formulário", fechar_modal, driver, wait)[0]:
            finalizar_relatorio()
            return
        registrar_screenshot_unico("formulario_fechado", driver, doc, "Formulário fechado.")

        log(doc, "✅ Teste de cadastro de Controle de Cremação concluído com sucesso.")

    except Exception as e:
        log(doc, f"❌ Erro geral durante a execução do teste: {e}")
        take_screenshot(driver, doc, "erro_geral")
    
    finally:
        finalizar_relatorio()

if __name__ == "__main__":
    main()