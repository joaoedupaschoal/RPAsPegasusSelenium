# Código corrigido: cadastro_cesta_basica_cenario_3.py

from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support.ui import Select, WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service
from webdriver_manager.chrome import ChromeDriverManager
from docx import Document
from docx.shared import Inches
from faker import Faker
from faker.providers import BaseProvider
from validate_docbr import CPF
from datetime import datetime, timedelta
import subprocess
from selenium.webdriver import ActionChains
import os, sys
import time
import random
import re

sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))

# ==== PROVIDERS CUSTOMIZADOS ====
class BrasilProvider(BaseProvider):
    def rg(self):
        numeros = [str(random.randint(0, 9)) for _ in range(8)]
        return ''.join(numeros) + '-' + str(random.randint(0, 9))

fake = Faker("pt_BR")
fake.add_provider(BrasilProvider)

# ==== CONFIGURAÇÕES ====
URL = "http://localhost:8080/gs/index.xhtml"
LOGIN_EMAIL = "joaoeduardo.gold@outlook.com"
LOGIN_PASSWORD = "071999gs"

# Controle de screenshots únicas
screenshot_registradas = set()

# ==== FUNÇÕES DE UTILITÁRIO ====
def log(doc, msg):
    """Registra mensagem no console e no documento"""
    print(msg)
    doc.add_paragraph(msg)

def sanitize_filename(nome):
    """Sanitiza nome do arquivo para evitar caracteres inválidos"""
    nome = nome.lower()
    nome = re.sub(r'[^a-z0-9_]', '_', nome)
    return nome

def take_screenshot(driver, doc, nome):
    """Captura screenshot e adiciona ao documento"""
    nome_limpo = sanitize_filename(nome)
    if nome_limpo not in screenshot_registradas:
        os.makedirs("screenshots", exist_ok=True)
        path = f"screenshots/{nome_limpo}.png"
        driver.save_screenshot(path)
        doc.add_paragraph(f"Screenshot: {nome}")
        doc.add_picture(path, width=Inches(5.5))
        screenshot_registradas.add(nome_limpo)

def safe_action(doc, descricao, func, driver, wait):
    """Executa ação com tratamento de erro e screenshot"""
    try:
        log(doc, f"🔄 {descricao}...")
        func()
        log(doc, f"✅ {descricao} realizada com sucesso.")
        take_screenshot(driver, doc, descricao.lower().replace(" ", "_"))
        return True, "sucesso"
    except Exception as e:
        log(doc, f"❌ Erro ao {descricao.lower()}: {e}")
        take_screenshot(driver, doc, f"erro_{descricao.lower().replace(' ', '_')}")
        return False, "erro"

def encontrar_mensagem_alerta(driver, doc):
    """Procura por mensagens de alerta no sistema"""
    seletores = [
        (".alerts.salvo", "sucesso"),
        (".alerts.alerta", "alerta"),
        (".alerts.erro", "erro"),
    ]

    for seletor, tipo in seletores:
        try:
            elemento = driver.find_element(By.CSS_SELECTOR, seletor)
            if elemento.is_displayed():
                if tipo == "sucesso":
                    log(doc, f"📢 ✅ Sucesso: {elemento.text}")
                elif tipo == "alerta":
                    log(doc, f"📢 ⚠️ Alerta: {elemento.text}")
                else:
                    log(doc, f"📢 ❌ Erro: {elemento.text}")
                return elemento, tipo
        except:
            continue

    log(doc, "ℹ️ Nenhuma mensagem de alerta encontrada.")
    return None, None

def ajustar_zoom(driver):
    """Ajusta o zoom da página para melhor visualização"""
    try:
        driver.execute_script("document.body.style.zoom='90%'")
        time.sleep(1)
    except Exception as e:
        raise Exception(f"Erro ao ajustar zoom: {e}")

def registrar_screenshot_unico(nome, driver, doc, descricao=None):
    """Registra screenshot único"""
    if nome not in screenshot_registradas:
        if descricao:
            log(doc, f"📸 {descricao}")
        take_screenshot(driver, doc, nome)
        screenshot_registradas.add(nome)

def main():
    # Inicialização do documento
    doc = Document()
    doc.add_heading("RELATÓRIO DO TESTE", 0)
    doc.add_paragraph("Cadastro de Cesta Básica - 3º Cenário.")
    doc.add_paragraph(f"🗓️ Data do teste: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")
    doc.add_paragraph("Neste teste, o robô NÃO preencherá os campos obrigatórios e tentará salvar o cadastro de uma nova Cesta Básica para validar as mensagens de erro.")

    # Configuração do Chrome
    chrome_options = Options()
    chrome_options.add_argument("--start-maximized")
    driver = webdriver.Chrome(service=Service(ChromeDriverManager().install()), options=chrome_options)
    wait = WebDriverWait(driver, 10)

    def finalizar_relatorio():
        """Finaliza e salva o relatório"""
        doc_name = f"relatorio_cesta_basica_cenario_3_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        doc.save(doc_name)
        log(doc, f"📄 Relatório salvo como: {doc_name}")
        try:
            subprocess.run(["start", "winword", doc_name], shell=True)
        except Exception as e:
            log(doc, f"Erro ao abrir o Word: {e}")
        driver.quit()

    def login():
        """Realiza login no sistema"""
        wait.until(EC.presence_of_element_located((By.ID, "j_id15:email"))).send_keys(LOGIN_EMAIL)
        wait.until(EC.presence_of_element_located((By.ID, "j_id15:senha"))).send_keys(LOGIN_PASSWORD, Keys.ENTER)
        wait.until(EC.presence_of_element_located((By.TAG_NAME, "body")))
        time.sleep(2)

    def abrir_menu():
        """Abre o menu de Cesta Básica"""
        driver.find_element(By.TAG_NAME, "body").send_keys(Keys.F2)
        campo = wait.until(EC.presence_of_element_located((By.XPATH, "//input[@placeholder='Busque um cadastro']")))
        campo.click()
        campo.send_keys("Cesta Básica", Keys.ENTER)
        time.sleep(2)

    def acessar_formulario():
        """Acessa o formulário de cadastro"""
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR,
            "#fmod_200009 > div.wdTelas > div.telaInicial.clearfix.overflow.overflowY > ul > li:nth-child(1) > a > span"))).click()
        time.sleep(2)

    def salvar():
        """Clica no botão salvar"""
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR,
            "#fmod_200009 > div.wdTelas > div.telaCadastro.clearfix > div.btnHolder > a.btModel.btGray.btsave"))).click()
        time.sleep(3)  # Aumentado o tempo para aguardar a mensagem aparecer

    def fechar_modal():
        """Fecha o modal"""
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR,
            "#fmod_200009 > div.wdTop.ui-draggable-handle > div.wdClose > a"))).click()
        time.sleep(1)

    # ==== EXECUÇÃO DO TESTE ====
    try:
        # 1. Acessar o sistema
        if not safe_action(doc, "Acessando o sistema", lambda: driver.get(URL), driver, wait)[0]:
            finalizar_relatorio()
            return
        registrar_screenshot_unico("url_acessada", driver, doc, "Sistema acessado.")

        # 2. Realizar login
        if not safe_action(doc, "Realizando login", login, driver, wait)[0]:
            finalizar_relatorio()
            return
        registrar_screenshot_unico("login_concluido", driver, doc, "Login realizado.")

        # 3. Ajustar zoom
        if not safe_action(doc, "Ajustando zoom", lambda: ajustar_zoom(driver), driver, wait)[0]:
            finalizar_relatorio()
            return

        # 4. Abrir menu
        if not safe_action(doc, "Abrindo menu Cesta Básica", abrir_menu, driver, wait)[0]:
            finalizar_relatorio()
            return
        registrar_screenshot_unico("menu_aberto", driver, doc, "Menu 'Cesta Básica' aberto.")

        # 5. Acessar formulário
        if not safe_action(doc, "Acessando formulário", acessar_formulario, driver, wait)[0]:
            finalizar_relatorio()
            return
        registrar_screenshot_unico("formulario_aberto", driver, doc, "Formulário de cadastro aberto.")

        # 6. Tentar salvar sem preencher campos (cenário de teste de validação)
        log(doc, "ℹ️ Tentando salvar formulário sem preencher campos obrigatórios...")
        if not safe_action(doc, "Clicando em Salvar", salvar, driver, wait)[0]:
            finalizar_relatorio()
            return
        registrar_screenshot_unico("apos_salvar", driver, doc, "Clique no botão Salvar realizado.")

        # 7. Verificar mensagens de alerta/erro
        time.sleep(2)  # Aguarda um pouco mais para a mensagem aparecer
        _, tipo_alerta = encontrar_mensagem_alerta(driver, doc)
        if tipo_alerta == "sucesso":
            log(doc, "✅ Mensagem de sucesso exibida.")
        elif tipo_alerta == "alerta":
            log(doc, "⚠️ Mensagem de alerta exibida.")
        elif tipo_alerta == "erro":
            log(doc, "❌ Mensagem de erro exibida.")
        else:
            log(doc, "⚠️ Nenhuma mensagem de validação encontrada.")
        registrar_screenshot_unico("mensagem_final", driver, doc, "Status após tentativa de salvamento.")

        # 8. Fechar modal
        if not safe_action(doc, "Fechando formulário", fechar_modal, driver, wait)[0]:
            finalizar_relatorio()
            return
        registrar_screenshot_unico("formulario_fechado", driver, doc, "Formulário fechado.")

        log(doc, "✅ Teste concluído com sucesso.")
        
    except Exception as e:
        log(doc, f"❌ Erro geral durante a execução do teste: {e}")
        take_screenshot(driver, doc, "erro_geral")
    
    finally:
        finalizar_relatorio()

if __name__ == "__main__":
    main()