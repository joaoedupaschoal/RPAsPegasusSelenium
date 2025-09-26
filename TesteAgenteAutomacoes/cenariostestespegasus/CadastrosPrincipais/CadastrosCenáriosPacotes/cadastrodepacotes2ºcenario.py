# Refatorado e organizado: cadastrodeescalamotorista2Âºcenario.py

from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support.ui import Select, WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service
from webdriver_manager.chrome import ChromeDriverManager
from docx import Document
from docx.shared import Inches
from faker import Faker
from faker.providers import BaseProvider
from validate_docbr import CPF
from datetime import datetime, timedelta
import subprocess
import os
import time
import random

# ==== PROVIDERS CUSTOMIZADOS ====
class BrasilProvider(BaseProvider):
    def rg(self):
        numeros = [str(random.randint(0, 9)) for _ in range(8)]
        return ''.join(numeros) + '-' + str(random.randint(0, 9))

fake = Faker("pt_BR")
fake.add_provider(BrasilProvider)
qtd_parcelas_em_atraso = int(fake.random.choice(['1', '2', '3', '4', '5']))


idade_de = fake.random_int(min=0, max=5)
idade_atÃ© = idade_de + fake.random_int(min=0, max=5)

valor_taxa_adesao = round(random.uniform(2000, 10000), 2)

def gerar_datas_validas():
    """Gera datas coerentes para admissÃ£o, inÃ­cio e fim da escala, e vencimento da CNH."""
    hoje = datetime.today().date()
    
    # Data de admissÃ£o entre 10 anos atrÃ¡s e hoje
    data_admissao = fake.date_between(start_date=hoje - timedelta(days=3650), end_date=hoje)
    
    # Data de inÃ­cio da escala entre hoje e 1 ano no futuro
    data_inicio = fake.date_between(start_date=hoje, end_date=hoje + timedelta(days=365))
    
    # Data fim entre 1 e 180 dias apÃ³s a data de inÃ­cio
    data_fim = data_inicio + timedelta(days=random.randint(1, 180))
    
    # Vencimento CNH entre hoje e 10 anos no futuro
    vencimento_cnh = fake.date_between(start_date=hoje, end_date=hoje + timedelta(days=3650))
    
    return (data_admissao.strftime('%d/%m/%Y'), 
            data_inicio.strftime('%d/%m/%Y'), 
            data_fim.strftime('%d/%m/%Y'), 
            vencimento_cnh.strftime('%d/%m/%Y'))

def gerar_dados_documentos():
    """Gera documentos fictÃ­cios para o cadastro."""
    carteira_trabalho = str(random.randint(10000000, 99999999))
    pis = fake.cpf().replace('.', '').replace('-', '')[:11]
    cnh = str(random.randint(10000000000, 99999999999))
    cpf = CPF().generate()
    
    return carteira_trabalho, pis, cnh, cpf

# Gera os dados necessÃ¡rios
data_admissao, data_inicio, data_fim, vencimento_cnh = gerar_datas_validas()
carteira_trabalho, pis, cnh, cpf_valido = gerar_dados_documentos()

# ==== CONFIGURAÃ‡Ã•ES ====
URL = "http://localhost:8080/gs/index.xhtml"
LOGIN_EMAIL = "joaoeduardo.gold@outlook.com"
LOGIN_PASSWORD = "071999gs"

# ==== DOCUMENTO ====
doc = Document()
doc.add_heading("RELATÃ“RIO DO TESTE", 0)
doc.add_paragraph("Cadastro de Pacotes â€“ CenÃ¡rio 2: Preenchimento completo e cancelamento.")
doc.add_paragraph(f"Data do teste: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")

screenshot_registradas = set()

# ==== FUNÃ‡Ã•ES DE UTILITÃRIO ====
def log(doc, msg):
    print(msg)
    doc.add_paragraph(msg)

def take_screenshot(driver, doc, nome):
    if nome not in screenshot_registradas:
        path = f"screenshots/{nome}.png"
        os.makedirs("screenshots", exist_ok=True)
        driver.save_screenshot(path)
        doc.add_paragraph(f"Screenshot: {nome}")
        doc.add_picture(path, width=Inches(5.5))
        screenshot_registradas.add(nome)

def safe_action(doc, descricao, func):
    try:
        log(doc, f"ðŸ”„ {descricao}...")
        func()
        log(doc, f"âœ… {descricao} realizada com sucesso.")
        take_screenshot(driver, doc, descricao.lower().replace(" ", "_"))
    except Exception as e:
        log(doc, f"âŒ Erro ao {descricao.lower()}: {e}")
        take_screenshot(driver, doc, f"erro_{descricao.lower().replace(' ', '_')}")

def finalizar_relatorio():
    nome_arquivo = f"relatorio_pacotes_cenario_2_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc.save(nome_arquivo)
    log(doc, f"ðŸ“„ RelatÃ³rio salvo como: {nome_arquivo}")
    subprocess.run(["start", "winword", nome_arquivo], shell=True)
    driver.quit()

def rolar_ate_dependentes_contemplados():
    campo_dependentes_contemplados = WebDriverWait(driver, 10).until(
        EC.element_to_be_clickable((By.CSS_SELECTOR, '#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_parametrosGerais.categoriaHolder > div.groupHolder.clearfix.grupo_dependentesContemplados > div.group_dependentesContemplados.clearfix.grupoHolder.lista > div > div > select'))
    )
    driver.execute_script("arguments[0].scrollIntoView();", campo_dependentes_contemplados)


def rolar_ate_formas_pagamento_mensalidade():
    campo_formas_pagamento_mensalidade = WebDriverWait(driver, 10).until(
        EC.element_to_be_clickable((By.CSS_SELECTOR, '#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_formasPagamento.categoriaHolder > div.groupHolder.clearfix.grupo_formaPagamentoMensalidade > div.group_formaPagamentoMensalidade.clearfix.grupoHolder.lista > div > div:nth-child(1) > select'))
    )
    driver.execute_script("arguments[0].scrollIntoView();", campo_formas_pagamento_mensalidade)


def rolar_ate_adicionar_formas_pagamento_mensalidade():
    adicionar = WebDriverWait(driver, 10).until(
        EC.element_to_be_clickable((By.CSS_SELECTOR, '#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_formasPagamento.categoriaHolder > div.groupHolder.clearfix.grupo_formaPagamentoMensalidade > div.btnListHolder > a.btAddGroup'))
    )
    driver.execute_script("arguments[0].scrollIntoView();", adicionar)
    adicionar.click()

def abrir_modal_e_selecionar(btn_selector, pesquisa_selector, termo_pesquisa, btn_pesquisar_selector, resultado_xpath):
    """Abre modal e seleciona um item"""
    def acao():
        # Abre o modal
        open_lov = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, btn_selector)))
        open_lov.click()
        time.sleep(3)

        # Aguarda campo pesquisa
        campo_pesquisa = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, pesquisa_selector)))
        campo_pesquisa.clear()
        campo_pesquisa.send_keys(termo_pesquisa)

        # Clica pesquisar
        pesquisar = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, btn_pesquisar_selector)))
        pesquisar.click()
        time.sleep(3)
        pesquisar.click()

        # Espera o resultado e clica
        wait.until(EC.element_to_be_clickable((By.XPATH, resultado_xpath)))
        resultado = driver.find_element(By.XPATH, resultado_xpath)
        driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", resultado)
        time.sleep(0.2)
        resultado.click()

    return acao

def encontrar_mensagem_alerta():
    seletores = [
        (".alerts.salvo", "âœ… Menasagem de Sucesso"),
        (".alerts.alerta", "âš ï¸ Menasagem de Alerta"),
        (".alerts.erro", "âŒ Menasagem de Erro"),
    ]

    for seletor, tipo in seletores:
        try:
            elemento = driver.find_element(By.CSS_SELECTOR, seletor)
            if elemento.is_displayed():
                log(doc, f"ðŸ“¢ {tipo}: {elemento.text}")
                return elemento
        except:
            continue

    log(doc, "â„¹ï¸ Nenhuma mensagem de alerta encontrada.")
    return None

def ajustar_zoom():
    try:
        driver.execute_script("document.body.style.zoom='90%'")
        log(doc, "ðŸ” Zoom ajustado para 90%.")
    except Exception as e:
        log(doc, f"âš ï¸ Erro ao ajustar zoom: {e}")

def preencher_campo_data(selector, valor):
    def acao():
        campo = wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, selector)))
        campo.click()
        campo.clear()
        campo.send_keys(valor)
        time.sleep(0.2)
    return acao

def selecionar_opcao(selector, texto):
    def acao():
        select_element = wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, selector)))
        Select(select_element).select_by_visible_text(texto)
    return acao

# ==== INICIALIZAÃ‡ÃƒO DO DRIVER ====
options = Options()
options.add_argument("--start-maximized")
driver = webdriver.Chrome(service=Service(ChromeDriverManager().install()), options=options)
wait = WebDriverWait(driver, 20)

# ==== EXECUÃ‡ÃƒO DO TESTE ====
try:
    safe_action(doc, "Acessando sistema", lambda: driver.get(URL))

    safe_action(doc, "Realizando login", lambda: (
        wait.until(EC.presence_of_element_located((By.ID, "j_id15:email"))).send_keys(LOGIN_EMAIL),
        wait.until(EC.presence_of_element_located((By.ID, "j_id15:senha"))).send_keys(LOGIN_PASSWORD, Keys.ENTER),
        wait.until(EC.presence_of_element_located((By.TAG_NAME, "body")))
    ))

    safe_action(doc, "Esperando sistema carregar e ajustando zoom", lambda: (
        time.sleep(5),
        ajustar_zoom()
    ))

    safe_action(doc, "Abrindo menu Pacotes", lambda: (
        driver.find_element(By.TAG_NAME, "body").send_keys(Keys.F2),
        time.sleep(1),
        wait.until(EC.visibility_of_element_located((By.XPATH, "//input[@placeholder='Busque um cadastro']"))).send_keys("Pacotes", Keys.ENTER),
        wait.until(EC.visibility_of_element_located((By.XPATH, "/html/body/div[17]/div[2]/ul/li[28]/a"))).click()

    ))

    safe_action(doc, "Clicando em Cadastrar", lambda: (
        time.sleep(2),
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_2 > div.wdTelas > div.telaInicial.clearfix.overflow.overflowY > ul > li:nth-child(1) > a > span"))).click()
    ))

    safe_action(doc, "Preechendo o Nome do Pacote", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_parametrosGerais.categoriaHolder > div.groupHolder.clearfix.grupo_parametrosGerais > div > div > div:nth-child(2) > input"))
        ).send_keys('TESTE PACOTE SELENIUM AUTOMATIZADO')
    ))

    safe_action(doc, "Preechendo a Quantidade MÃ¡xima de Titulares", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_parametrosGerais.categoriaHolder > div.groupHolder.clearfix.grupo_parametrosGerais > div > div > div:nth-child(4) > input"))
        ).send_keys(fake.random_int(min=1, max=3))
    ))

    safe_action(doc, "Preechendo a Quantidade MÃ¡xima de Dependentes", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_parametrosGerais.categoriaHolder > div.groupHolder.clearfix.grupo_parametrosGerais > div > div > div:nth-child(5) > input"))
        ).send_keys(fake.random_int(min=1, max=99))
    ))


    safe_action(doc, "Preechendo a VigÃªncia (em dias)", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_parametrosGerais.categoriaHolder > div.groupHolder.clearfix.grupo_parametrosGerais > div > div > div:nth-child(6) > input"))
        ).send_keys(fake.random_int(min=30, max=365))
    ))


    safe_action(doc, "Preechendo a CarÃªncia (em dias) para utilizaÃ§Ã£o", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_parametrosGerais.categoriaHolder > div.groupHolder.clearfix.grupo_parametrosGerais > div > div > div:nth-child(8) > input"))
        ).send_keys(fake.random_int(min=0, max=15))
    ))


    safe_action(doc, "Preechendo a CarÃªncia (em dias) para mensalidade", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_parametrosGerais.categoriaHolder > div.groupHolder.clearfix.grupo_parametrosGerais > div > div > div:nth-child(10) > input"))
        ).send_keys(fake.random_int(min=0, max=15))
    ))



    safe_action(doc, "Preechendo a CarÃªncia (em dias) para adesÃ£o", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_parametrosGerais.categoriaHolder > div.groupHolder.clearfix.grupo_parametrosGerais > div > div > div:nth-child(11) > input"))
        ).send_keys(fake.random_int(min=0, max=15))
    ))



    safe_action(doc, "Selecionando a opÃ§Ã£o para o Pacote possuir Jazigo", selecionar_opcao(
        "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_parametrosGerais.categoriaHolder > div.groupHolder.clearfix.grupo_parametrosGerais > div > div > div:nth-child(12) > select",
        "Com Jazigo"
    ))


    safe_action(doc, "Selecionando a opÃ§Ã£o para o Pacote nÃ£o vincular a quadra", selecionar_opcao(
        "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_parametrosGerais.categoriaHolder > div.groupHolder.clearfix.grupo_parametrosGerais > div > div > div:nth-child(13) > select",
        "NÃ£o"
    ))

    
    safe_action(doc, "Selecionando a opÃ§Ã£o para o Pacote possuir Sepultamento", selecionar_opcao(
        "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_parametrosGerais.categoriaHolder > div.groupHolder.clearfix.grupo_parametrosGerais > div > div > div:nth-child(14) > select",
        "Com Sepultamento"
    ))


    safe_action(doc, "Selecionando Tipo de Contrato", abrir_modal_e_selecionar(
        "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_parametrosGerais.categoriaHolder > div.groupHolder.clearfix.grupo_parametrosGerais > div > div > div:nth-child(15) > div > a",
        "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(2) > input",
        "TIPO DE CONTRATO TESTE SELENIUM AUTOMATIZADO",
        "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(4) > a",
        "//td[contains(text(), 'TIPO DE CONTRATO TESTE SELENIUM AUTOMATIZADO')]/a[contains(@class, 'linkAlterar')]"
    ))



    safe_action(doc, "Selecionando a opÃ§Ã£o para o Tipo de Pacote na OS ser Por Contrato", selecionar_opcao(
        "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_parametrosGerais.categoriaHolder > div.groupHolder.clearfix.grupo_parametrosGerais > div > div > div:nth-child(16) > select",
        "Por Contrato"
    ))

    safe_action(doc, "Selecionando a opÃ§Ã£o para LanÃ§ar valor por Quantidade de Dependentes", selecionar_opcao(
        "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_parametrosGerais.categoriaHolder > div.groupHolder.clearfix.grupo_parametrosGerais > div > div > div:nth-child(17) > select",
        "Sim"
    ))

    safe_action(doc, "Selecionando a opÃ§Ã£o para o Pacote possuir CremaÃ§Ã£o", selecionar_opcao(
        "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_parametrosGerais.categoriaHolder > div.groupHolder.clearfix.grupo_parametrosGerais > div > div > div:nth-child(18) > select",
        "Sim"
    ))


    safe_action(doc, "Selecionando a opÃ§Ã£o para o Registro de Ã“bito ser obrigatÃ³rio na CremaÃ§Ã£o", selecionar_opcao(
        "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_parametrosGerais.categoriaHolder > div.groupHolder.clearfix.grupo_parametrosGerais > div > div > div:nth-child(19) > select",
        "Sim"
    ))

    safe_action(doc, "Selecionando a opÃ§Ã£o para o Pacote lanÃ§ar manutenÃ§Ã£o apÃ³s o 1Âº Ã“bito", selecionar_opcao(
        "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_parametrosGerais.categoriaHolder > div.groupHolder.clearfix.grupo_parametrosGerais > div > div > div:nth-child(20) > select",
        "Sim"
    ))

    safe_action(doc, "Selecionando a opÃ§Ã£o para o Pacote LanÃ§ar TÃ­tulos no Cadastro de Contrato", selecionar_opcao(
        "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_parametrosGerais.categoriaHolder > div.groupHolder.clearfix.grupo_parametrosGerais > div > div > div:nth-child(22) > select",
        "Sim"
    ))


    safe_action(doc, "Rolando atÃ© o Select de Dependentes Contemplados", rolar_ate_dependentes_contemplados)

    safe_action(doc, "Selecionando a opÃ§Ã£o 'Agregado'", selecionar_opcao(
        "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_parametrosGerais.categoriaHolder > div.groupHolder.clearfix.grupo_dependentesContemplados > div.group_dependentesContemplados.clearfix.grupoHolder.lista > div > div > select",
        "Agregado"
    ))

    safe_action(doc, "Adicionado Dependente", rolar_ate_adicionar_formas_pagamento_mensalidade)



    safe_action(doc, "Adicionado Dependente", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_parametrosGerais.categoriaHolder > div.groupHolder.clearfix.grupo_dependentesContemplados > div.btnListHolder > a.btAddGroup"))
        ).click()
    ))



    safe_action(doc, "Acessando aba Taxas do Pacote", lambda: (
        wait.until(EC.element_to_be_clickable((By.LINK_TEXT, "Taxas do Pacote"))).click()
    ))


    safe_action(doc, "Preenchendo a AdesÃ£o Total do pacote", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_taxasPacote.categoriaHolder > div.groupHolder.clearfix.grupo_taxasAdesao > div > div > div:nth-child(1) > input"))
        ).send_keys(str(valor_taxa_adesao).replace('.', ','))
    ))


    safe_action(doc, "Selecionando Tipo de Mensalidade AdesÃ£o", abrir_modal_e_selecionar(
        "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_taxasPacote.categoriaHolder > div.groupHolder.clearfix.grupo_taxasAdesao > div > div > div:nth-child(2) > div > a",
        "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(2) > input",
        "TIPO MENSALIDADE JOÃƒO TESTE",
        "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(4) > a",
        "//td[contains(text(), 'TIPO MENSALIDADE JOÃƒO TESTE')]"
    ))



    safe_action(doc, "Selecionando a opÃ§Ã£o 'Por Periodicidade' no Tipo de Taxa para ManutenÃ§Ã£o ", selecionar_opcao(
        "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_taxasPacote.categoriaHolder > div.groupHolder.clearfix.grupo_taxasManutencao > div > div > div:nth-child(1) > select",
        "Por periodicidade"
    ))


    safe_action(doc, "Preenchendo a Taxa Mensal", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_taxasPacote.categoriaHolder > div.groupHolder.clearfix.grupo_taxasManutencao > div > div > div:nth-child(3) > input"))
        ).send_keys(fake.random_int(min=10, max=500))
    ))



    safe_action(doc, "Preenchendo a Taxa Bimestral", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_taxasPacote.categoriaHolder > div.groupHolder.clearfix.grupo_taxasManutencao > div > div > div:nth-child(4) > input"))
        ).send_keys(fake.random_int(min=1000, max=50000))
    ))



    safe_action(doc, "Preenchendo a Taxa Trimestral", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_taxasPacote.categoriaHolder > div.groupHolder.clearfix.grupo_taxasManutencao > div > div > div:nth-child(5) > input"))
        ).send_keys(fake.random_int(min=1000, max=50000))
    ))

    safe_action(doc, "Preenchendo a Taxa Quadrimestral", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_taxasPacote.categoriaHolder > div.groupHolder.clearfix.grupo_taxasManutencao > div > div > div:nth-child(6) > input"))
        ).send_keys(fake.random_int(min=1000, max=50000))
    ))


    safe_action(doc, "Preenchendo a Taxa Semestral", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_taxasPacote.categoriaHolder > div.groupHolder.clearfix.grupo_taxasManutencao > div > div > div:nth-child(7) > input"))
        ).send_keys(fake.random_int(min=1000, max=50000))
    ))

    safe_action(doc, "Preenchendo a Taxa Anual", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_taxasPacote.categoriaHolder > div.groupHolder.clearfix.grupo_taxasManutencao > div > div > div:nth-child(8) > input"))
        ).send_keys(fake.random_int(min=1000, max=50000))
    ))



    safe_action(doc, "Selecionando Tipo de Mensalidade ManutenÃ§Ã£o", abrir_modal_e_selecionar(
        "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_taxasPacote.categoriaHolder > div.groupHolder.clearfix.grupo_taxasManutencao > div > div > div:nth-child(9) > div > a",
        "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(2) > input",
        "TIPO MENSALIDADE MANUTENÃ‡ÃƒO",
        "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(4) > a",
        "//td[contains(text(), 'TIPO MENSALIDADE MANUTENÃ‡ÃƒO')]"
    ))


    safe_action(doc, "Preenchendo a o Valor da TransferÃªncia de Titularidade", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_taxasPacote.categoriaHolder > div.groupHolder.clearfix.grupo_outras > div > div > div:nth-child(1) > input"))
        ).send_keys(fake.random_int(min=1000, max=50000))
    ))

    safe_action(doc, "Preenchendo a o Valor requerido para Beneficiar pessoas fora do Contrato", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_taxasPacote.categoriaHolder > div.groupHolder.clearfix.grupo_outras > div > div > div:nth-child(3) > input"))
        ).send_keys(fake.random_int(min=1000, max=50000))
    ))




    safe_action(doc, "Acessando aba Formas de Pagamento", lambda: (
        wait.until(EC.element_to_be_clickable((By.LINK_TEXT, "Formas de Pagamento"))).click()
    ))


    safe_action(doc, "Selecionando a opÃ§Ã£o 'Boleto' em Formas de pagamento da AdesÃ£o", selecionar_opcao(
        "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_formasPagamento.categoriaHolder > div.groupHolder.clearfix.grupo_formaPagamentoAdesao > div.group_formaPagamentoAdesao.clearfix.grupoHolder.lista > div > div:nth-child(1) > select",
        "Boleto"
    ))

    safe_action(doc, "Preenchendo o campo de Parcelamento MÃ¡ximo ", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_formasPagamento.categoriaHolder > div.groupHolder.clearfix.grupo_formaPagamentoAdesao > div.group_formaPagamentoAdesao.clearfix.grupoHolder.lista > div > div:nth-child(2) > input"))
        ).send_keys(fake.random_element([2, 3, 4, 5, 10, 6, 12, 24]))
    ))



    safe_action(doc, "Adicionando Forma de Pagamento", lambda: driver.find_element(
        By.CSS_SELECTOR, "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_formasPagamento.categoriaHolder > div.groupHolder.clearfix.grupo_formaPagamentoAdesao > div.btnListHolder > a.btAddGroup"
    ).click())



    safe_action(doc, "Selecionando a opÃ§Ã£o 'CartÃ£o de CrÃ©dito' em Formas de pagamento da AdesÃ£o", selecionar_opcao(
        "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_formasPagamento.categoriaHolder > div.groupHolder.clearfix.grupo_formaPagamentoAdesao > div.group_formaPagamentoAdesao.clearfix.grupoHolder.lista > div > div:nth-child(1) > select",
        "CartÃ£o de CrÃ©dito"
    ))

    safe_action(doc, "Preenchendo o campo de Parcelamento MÃ¡ximo ", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_formasPagamento.categoriaHolder > div.groupHolder.clearfix.grupo_formaPagamentoAdesao > div.group_formaPagamentoAdesao.clearfix.grupoHolder.lista > div > div:nth-child(2) > input"))
        ).send_keys(fake.random_element([2, 3, 4, 5, 10, 6, 12, 24]))
    ))



    safe_action(doc, "Adicionando Forma de Pagamento", lambda: driver.find_element(
        By.CSS_SELECTOR, "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_formasPagamento.categoriaHolder > div.groupHolder.clearfix.grupo_formaPagamentoAdesao > div.btnListHolder > a.btAddGroup"
    ).click())


    safe_action(doc, "Selecionando a opÃ§Ã£o 'CartÃ£o de DÃ©bito' em Formas de pagamento da AdesÃ£o", selecionar_opcao(
        "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_formasPagamento.categoriaHolder > div.groupHolder.clearfix.grupo_formaPagamentoAdesao > div.group_formaPagamentoAdesao.clearfix.grupoHolder.lista > div > div:nth-child(1) > select",
        "CartÃ£o de DÃ©bito"
    ))

    safe_action(doc, "Preenchendo o campo de Parcelamento MÃ¡ximo ", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_formasPagamento.categoriaHolder > div.groupHolder.clearfix.grupo_formaPagamentoAdesao > div.group_formaPagamentoAdesao.clearfix.grupoHolder.lista > div > div:nth-child(2) > input"))
        ).send_keys(fake.random_element([2, 3, 4, 5, 10, 6, 12, 24]))
    ))



    safe_action(doc, "Adicionando Forma de Pagamento", lambda: driver.find_element(
        By.CSS_SELECTOR, "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_formasPagamento.categoriaHolder > div.groupHolder.clearfix.grupo_formaPagamentoAdesao > div.btnListHolder > a.btAddGroup"
    ).click())



    safe_action(doc, "Selecionando a opÃ§Ã£o 'Cheque' em Formas de pagamento da AdesÃ£o", selecionar_opcao(
        "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_formasPagamento.categoriaHolder > div.groupHolder.clearfix.grupo_formaPagamentoAdesao > div.group_formaPagamentoAdesao.clearfix.grupoHolder.lista > div > div:nth-child(1) > select",
        "Cheque"
    ))

    safe_action(doc, "Preenchendo o campo de Parcelamento MÃ¡ximo ", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_formasPagamento.categoriaHolder > div.groupHolder.clearfix.grupo_formaPagamentoAdesao > div.group_formaPagamentoAdesao.clearfix.grupoHolder.lista > div > div:nth-child(2) > input"))
        ).send_keys(fake.random_element([2, 3, 4, 5, 10, 6, 12, 24]))
    ))



    safe_action(doc, "Adicionando Forma de Pagamento", lambda: driver.find_element(
        By.CSS_SELECTOR, "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_formasPagamento.categoriaHolder > div.groupHolder.clearfix.grupo_formaPagamentoAdesao > div.btnListHolder > a.btAddGroup"
    ).click())



    safe_action(doc, "Selecionando a opÃ§Ã£o 'DepÃ³sito' em Formas de pagamento da AdesÃ£o", selecionar_opcao(
        "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_formasPagamento.categoriaHolder > div.groupHolder.clearfix.grupo_formaPagamentoAdesao > div.group_formaPagamentoAdesao.clearfix.grupoHolder.lista > div > div:nth-child(1) > select",
        "DepÃ³sito"
    ))

    safe_action(doc, "Preenchendo o campo de Parcelamento MÃ¡ximo ", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_formasPagamento.categoriaHolder > div.groupHolder.clearfix.grupo_formaPagamentoAdesao > div.group_formaPagamentoAdesao.clearfix.grupoHolder.lista > div > div:nth-child(2) > input"))
        ).send_keys(fake.random_element([2, 3, 4, 5, 10, 6, 12, 24]))
    ))



    safe_action(doc, "Adicionando Forma de Pagamento", lambda: driver.find_element(
        By.CSS_SELECTOR, "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_formasPagamento.categoriaHolder > div.groupHolder.clearfix.grupo_formaPagamentoAdesao > div.btnListHolder > a.btAddGroup"
    ).click())


    safe_action(doc, "Selecionando a opÃ§Ã£o 'Dinheiro' em Formas de pagamento da AdesÃ£o", selecionar_opcao(
        "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_formasPagamento.categoriaHolder > div.groupHolder.clearfix.grupo_formaPagamentoAdesao > div.group_formaPagamentoAdesao.clearfix.grupoHolder.lista > div > div:nth-child(1) > select",
        "Dinheiro"
    ))

    safe_action(doc, "Preenchendo o campo de Parcelamento MÃ¡ximo ", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_formasPagamento.categoriaHolder > div.groupHolder.clearfix.grupo_formaPagamentoAdesao > div.group_formaPagamentoAdesao.clearfix.grupoHolder.lista > div > div:nth-child(2) > input"))
        ).send_keys(fake.random_element([2, 3, 4, 5, 10, 6, 12, 24]))
    ))



    safe_action(doc, "Adicionando Forma de Pagamento", lambda: driver.find_element(
        By.CSS_SELECTOR, "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_formasPagamento.categoriaHolder > div.groupHolder.clearfix.grupo_formaPagamentoAdesao > div.btnListHolder > a.btAddGroup"
    ).click())


    safe_action(doc, "Selecionando a opÃ§Ã£o 'DÃ©bito em Conta' em Formas de pagamento da AdesÃ£o", selecionar_opcao(
        "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_formasPagamento.categoriaHolder > div.groupHolder.clearfix.grupo_formaPagamentoAdesao > div.group_formaPagamentoAdesao.clearfix.grupoHolder.lista > div > div:nth-child(1) > select",
        "DÃ©bito em Conta"
    ))

    safe_action(doc, "Preenchendo o campo de Parcelamento MÃ¡ximo ", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_formasPagamento.categoriaHolder > div.groupHolder.clearfix.grupo_formaPagamentoAdesao > div.group_formaPagamentoAdesao.clearfix.grupoHolder.lista > div > div:nth-child(2) > input"))
        ).send_keys(fake.random_element([2, 3, 4, 5, 10, 6, 12, 24]))
    ))



    safe_action(doc, "Adicionando Forma de Pagamento", lambda: driver.find_element(
        By.CSS_SELECTOR, "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_formasPagamento.categoriaHolder > div.groupHolder.clearfix.grupo_formaPagamentoAdesao > div.btnListHolder > a.btAddGroup"
    ).click())


    safe_action(doc, "Selecionando a opÃ§Ã£o 'TransferÃªncia BancÃ¡ria' em Formas de pagamento da AdesÃ£o", selecionar_opcao(
        "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_formasPagamento.categoriaHolder > div.groupHolder.clearfix.grupo_formaPagamentoAdesao > div.group_formaPagamentoAdesao.clearfix.grupoHolder.lista > div > div:nth-child(1) > select",
        "TransferÃªncia BancÃ¡ria"
    ))

    safe_action(doc, "Preenchendo o campo de Parcelamento MÃ¡ximo ", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_formasPagamento.categoriaHolder > div.groupHolder.clearfix.grupo_formaPagamentoAdesao > div.group_formaPagamentoAdesao.clearfix.grupoHolder.lista > div > div:nth-child(2) > input"))
        ).send_keys(fake.random_element([2, 3, 4, 5, 10, 6, 12, 24]))
    ))



    safe_action(doc, "Adicionando Forma de Pagamento", lambda: driver.find_element(
        By.CSS_SELECTOR, "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_formasPagamento.categoriaHolder > div.groupHolder.clearfix.grupo_formaPagamentoAdesao > div.btnListHolder > a.btAddGroup"
    ).click())



    safe_action(doc, "Rolando atÃ© o Select de Formas de Pagamanto da Mensalidade", rolar_ate_formas_pagamento_mensalidade)



    safe_action(doc, "Selecionando a opÃ§Ã£o 'Boleto' em Formas de pagamento da Mensalidade", selecionar_opcao(
        "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_formasPagamento.categoriaHolder > div.groupHolder.clearfix.grupo_formaPagamentoMensalidade > div.group_formaPagamentoMensalidade.clearfix.grupoHolder.lista > div > div:nth-child(1) > select",
        "Boleto"
    ))

    safe_action(doc, "Preenchendo o campo de Parcelamento MÃ¡ximo ", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_formasPagamento.categoriaHolder > div.groupHolder.clearfix.grupo_formaPagamentoMensalidade > div.group_formaPagamentoMensalidade.clearfix.grupoHolder.lista > div > div:nth-child(2) > input"))
        ).send_keys(fake.random_element([2, 3, 4, 5, 10, 6, 12, 24]))
    ))

    safe_action(doc, "Adicionando Forma de Pagamento", rolar_ate_adicionar_formas_pagamento_mensalidade)


    safe_action(doc, "Selecionando a opÃ§Ã£o 'CarnÃª' em Formas de pagamento da Mensalidade", selecionar_opcao(
        "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_formasPagamento.categoriaHolder > div.groupHolder.clearfix.grupo_formaPagamentoMensalidade > div.group_formaPagamentoMensalidade.clearfix.grupoHolder.lista > div > div:nth-child(1) > select",
        "CarnÃª"
    ))

    safe_action(doc, "Preenchendo o campo de Parcelamento MÃ¡ximo ", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_formasPagamento.categoriaHolder > div.groupHolder.clearfix.grupo_formaPagamentoMensalidade > div.group_formaPagamentoMensalidade.clearfix.grupoHolder.lista > div > div:nth-child(2) > input"))
        ).send_keys(fake.random_element([2, 3, 4, 5, 10, 6, 12, 24]))
    ))


    safe_action(doc, "Adicionando Forma de Pagamento", rolar_ate_adicionar_formas_pagamento_mensalidade)



    safe_action(doc, "Selecionando a opÃ§Ã£o 'Pix' em Formas de pagamento da Mensalidade", selecionar_opcao(
        "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_formasPagamento.categoriaHolder > div.groupHolder.clearfix.grupo_formaPagamentoMensalidade > div.group_formaPagamentoMensalidade.clearfix.grupoHolder.lista > div > div:nth-child(1) > select",
        "Pix"
    ))

    safe_action(doc, "Preenchendo o campo de Parcelamento MÃ¡ximo ", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_formasPagamento.categoriaHolder > div.groupHolder.clearfix.grupo_formaPagamentoMensalidade > div.group_formaPagamentoMensalidade.clearfix.grupoHolder.lista > div > div:nth-child(2) > input"))
        ).send_keys(fake.random_element([2, 3, 4, 5, 10, 6, 12, 24]))
    ))



    safe_action(doc, "Adicionando Forma de Pagamento", rolar_ate_adicionar_formas_pagamento_mensalidade)

    safe_action(doc, "Acessando aba Produtos", lambda: (
        wait.until(EC.element_to_be_clickable((By.LINK_TEXT, "Produtos"))).click()
    ))


    safe_action(doc, "Selecionando o Produto PRODUTO TESTE SELENIUM", abrir_modal_e_selecionar(
        "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_produtos.categoriaHolder > div > div.group_produtos.clearfix.grupoHolder.lista > div > div:nth-child(1) > div > a",
        "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(2) > input",
        "PRODUTO TESTE SELENIUM",
        "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(4) > a",
        "//td[contains(text(), 'PRODUTO TESTE SELENIUM')]"
    ))



    safe_action(doc, "Preenchendo a quantidade do produto", lambda: (
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_produtos.categoriaHolder > div > div.group_produtos.clearfix.grupoHolder.lista > div > div:nth-child(2) > input"))
        ).send_keys(fake.random_choices(elements=[1,2,3,4,5,6,7,8,9])[0])
    ))


    safe_action(doc, "Adicionando Produto", lambda: (
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_produtos.categoriaHolder > div > div.btnListHolder > a.btAddGroup"))).click()
    ))


    safe_action(doc, "Acessando aba ServiÃ§os", lambda: (
        wait.until(EC.element_to_be_clickable((By.LINK_TEXT, "ServiÃ§os"))).click()
    ))


    safe_action(doc, "Selecionando o ServiÃ§o SERVIÃ‡O TESTE SELENIUM", abrir_modal_e_selecionar(
        "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_servicos.categoriaHolder > div > div.group_servicos.clearfix.grupoHolder.lista > div > div:nth-child(1) > div > a",
        "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(2) > input",
        "SERVIÃ‡O TESTE SELENIUM",
        "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(4) > a",
        "//td[contains(text(), 'SERVIÃ‡O TESTE SELENIUM')]"
    ))


    safe_action(doc, "Adicionando ServiÃ§o", lambda: (
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_servicos.categoriaHolder > div > div.btnListHolder > a.btAddGroup"))).click()
    ))




    safe_action(doc, "Acessando aba InadimplÃªncia", lambda: (
        wait.until(EC.element_to_be_clickable((By.LINK_TEXT, "InadimplÃªncia"))).click()
    ))


    safe_action(doc, "Preenchendo a Quantidade de parcelas em atraso para o Bloqueio da Ordem de ServiÃ§o ", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_inadimplencia.categoriaHolder > div.groupHolder.clearfix.grupo_inadimplenciaOS > div > div > div:nth-child(2) > input"))
        ).send_keys(qtd_parcelas_em_atraso)
    ))

    safe_action(doc, "Preenchendo a CarÃªncia (em dias) apÃ³s pagamento ", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_inadimplencia.categoriaHolder > div.groupHolder.clearfix.grupo_inadimplenciaOS > div > div > div:nth-child(3) > input"))
        ).send_keys(fake.random_int(min=0, max=15))
    ))

    safe_action(doc, "Preenchendo a Quantidade de parcelas em atraso para o Bloqueio da LocaÃ§Ã£o de Equipamento ", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_inadimplencia.categoriaHolder > div.groupHolder.clearfix.grupo_inadimplenciaLocacaoEquipamento > div > div > div:nth-child(2) > input"))
        ).send_keys(qtd_parcelas_em_atraso)
    ))

    safe_action(doc, "Acessando aba CarÃªncia Dependente", lambda: (
        wait.until(EC.element_to_be_clickable((By.LINK_TEXT, "CarÃªncia Dependente"))).click()
    ))



    safe_action(doc, "Preenchendo o campo 'Idade de'", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_carenciaDependente.categoriaHolder > div > div.group_carenciaDependente.clearfix.grupoHolder.lista > div > div:nth-child(1) > input"))
        ).send_keys(idade_de)
    ))

    safe_action(doc, "Preenchendo o campo 'Idade atÃ©'", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_carenciaDependente.categoriaHolder > div > div.group_carenciaDependente.clearfix.grupoHolder.lista > div > div:nth-child(2) > input"))
        ).send_keys(idade_atÃ©)

    ))



    safe_action(doc, "Preenchendo a CarÃªncia (em dias) para utilizaÃ§Ã£o ", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_carenciaDependente.categoriaHolder > div > div.group_carenciaDependente.clearfix.grupoHolder.lista > div > div:nth-child(3) > input"))
        ).send_keys(fake.random_int(min=0, max=15))
    ))

    safe_action(doc, "Adicionando parÃ¢mtros da CarÃªncia Dependente", lambda: (
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_carenciaDependente.categoriaHolder > div > div.btnListHolder > a.btAddGroup"))).click()
    ))


    safe_action(doc, "Acessando aba AlteraÃ§Ã£o AutomÃ¡tica de Status do Contrato", lambda: (
        wait.until(EC.element_to_be_clickable((By.LINK_TEXT, "AlteraÃ§Ã£o AutomÃ¡tica de Status do Contrato"))).click()
    ))

    safe_action(doc, "Preenchendo a Quantidade parcelas vencidas para cancelar ", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_alteracaoStatusContrato.categoriaHolder > div.groupHolder.clearfix.grupo_parametros > div > div > div:nth-child(1) > input"))
        ).send_keys(fake.random_int(min=0, max=15))
    ))


    safe_action(doc, "Selecionando o Motivo Cancelamento CLIENTE SOLICITOU CANCELAMENTO", abrir_modal_e_selecionar(
        "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_alteracaoStatusContrato.categoriaHolder > div.groupHolder.clearfix.grupo_parametros > div > div > div:nth-child(2) > div > a",
        "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(2) > input",
        "CLIENTE SOLICITOU CANCELAMENTO",
        "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(4) > a",
        "//td[contains(text(), 'CLIENTE SOLICITOU CANCELAMENTO')]"
    ))


    safe_action(doc, "Preenchendo a Quantidade parcelas vencidas para suspender ", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_alteracaoStatusContrato.categoriaHolder > div.groupHolder.clearfix.grupo_parametros > div > div > div:nth-child(3) > input"))
        ).send_keys(fake.random_int(min=0, max=15))
    ))


    safe_action(doc, "Selecionando o Motivo SuspensÃ£o SEM COMUNICAÃ‡ÃƒO", abrir_modal_e_selecionar(
        "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_alteracaoStatusContrato.categoriaHolder > div.groupHolder.clearfix.grupo_parametros > div > div > div:nth-child(4) > div > a",
        "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(2) > input",
        "SEM COMUNICAÃ‡ÃƒO",
        "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(4) > a",
        "//td[contains(text(), 'SEM COMUNICAÃ‡ÃƒO')]"
    ))

    safe_action(doc, "Preenchendo a Quantidade parcelas vencidas para reativar ", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_alteracaoStatusContrato.categoriaHolder > div.groupHolder.clearfix.grupo_parametros > div > div > div:nth-child(5) > input"))
        ).send_keys(fake.random_int(min=0, max=15))
    ))


    safe_action(doc, "Selecionando o Motivo ReativaÃ§Ã£o PAGOU AS PARCELAS", abrir_modal_e_selecionar(
        "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_alteracaoStatusContrato.categoriaHolder > div.groupHolder.clearfix.grupo_parametros > div > div > div:nth-child(6) > div > a",
        "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(2) > input",
        "PAGOU AS PARCELAS",
        "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(4) > a",
        "//td[contains(text(), 'PAGOU AS PARCELAS')]"
    ))

    safe_action(doc, "Selecionando a opÃ§Ã£o 'Sim' para Cancelar a RecorrÃªncia Vindi na AlteraÃ§Ã£o AutomÃ¡tica de Status do Contrato", selecionar_opcao(
        "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_alteracaoStatusContrato.categoriaHolder > div.groupHolder.clearfix.grupo_parametros > div > div > div:nth-child(7) > select",
        "Sim"
    ))

    safe_action(doc, "Selecionando o Motivo Cancelamento RecorrÃªncia CANCELAMENTO POR FALTA DE PAGAMENTO", abrir_modal_e_selecionar(
        "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_alteracaoStatusContrato.categoriaHolder > div.groupHolder.clearfix.grupo_parametros > div > div > div:nth-child(8) > div > a",
        "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(2) > input",
        "CANCELAMENTO POR FALTA DE PAGAMENTO",
        "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(4) > a",
        "//td[contains(text(), 'CANCELAMENTO POR FALTA DE PAGAMENTO')]"
    ))


    safe_action(doc, "Selecionando o Tipo de Mensalidade TIPO MENSALIDADE JOÃƒO TESTE", abrir_modal_e_selecionar(
        "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_alteracaoStatusContrato.categoriaHolder > div.groupHolder.clearfix.grupo_tipoMensalidade > div.group_tipoMensalidade.clearfix.grupoHolder.lista > div > div > div > a",
        "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(2) > input",
        "TIPO MENSALIDADE JOÃƒO TESTE",
        "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(4) > a",
        "//td[contains(text(), 'TIPO MENSALIDADE JOÃƒO TESTE')]"
    ))



    safe_action(doc, "Adicionando Tipo de Mensalidade ", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_alteracaoStatusContrato.categoriaHolder > div.groupHolder.clearfix.grupo_tipoMensalidade > div.btnListHolder > a.btAddGroup"))
        ).click()
    ))
    
    safe_action(doc, "Cancelando cadastro", lambda: driver.find_element(
        By.CSS_SELECTOR, "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.btnHolder > a.btModel.btGray.btcancel"
    ).click())





    safe_action(doc, "Fechando modal apÃ³s cancelamento", lambda: wait.until(EC.element_to_be_clickable((
        By.CSS_SELECTOR, "#fmod_2 > div.wdTop.ui-draggable-handle > div.wdClose > a"
        ))
    ).click())



    encontrar_mensagem_alerta()

except Exception as e:
    log(doc, f"âŒ ERRO FATAL: {e}")
    take_screenshot(driver, doc, "erro_fatal")

finally:

    log(doc, "âœ… Teste concluÃ­do com sucesso.")

    finalizar_relatorio()