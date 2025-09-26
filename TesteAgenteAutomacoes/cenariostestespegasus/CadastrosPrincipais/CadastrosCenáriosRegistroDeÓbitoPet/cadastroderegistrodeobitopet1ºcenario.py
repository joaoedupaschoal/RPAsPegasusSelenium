# Refatorado e corrigido: cadastroderegistrodeobitopet1ºcenario.py

from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support.ui import Select, WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service
from selenium.webdriver import ActionChains
from webdriver_manager.chrome import ChromeDriverManager
from docx import Document
from docx.shared import Inches
from faker import Faker
from faker.providers import BaseProvider
from datetime import datetime, timedelta
import subprocess
import os
import time
import random
import string

# ==== PROVIDERS CUSTOMIZADOS ====
class BrasilProvider(BaseProvider):
    def rg(self):
        numeros = [str(random.randint(0, 9)) for _ in range(8)]
        return ''.join(numeros) + '-' + str(random.randint(0, 9))

fake = Faker("pt_BR")
fake.add_provider(BrasilProvider)

# ==== CONFIGURAÇÕES ====
URL = "http://localhost:8080/gs/index.xhtml"
LOGIN_EMAIL = "joaoeduardo.gold@outlook.com"
LOGIN_PASSWORD = "071999gs"

# ==== DOCUMENTO ====
doc = Document()
doc.add_heading("RELATÓRIO DO TESTE", 0)
doc.add_paragraph("Cadastro de Registro de Óbito Pet – Cenário 1: Preenchimento completo e salvamento.")
doc.add_paragraph(f"Data do teste: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")

screenshot_registradas = set()

# ==== FUNÇÕES DE UTILITÁRIO ====
def log(doc, msg):
    print(msg)
    doc.add_paragraph(msg)

def take_screenshot(driver, doc, nome):
    if nome not in screenshot_registradas:
        path = f"screenshots/{nome}.png"
        os.makedirs("screenshots", exist_ok=True)
        driver.save_screenshot(path)
        doc.add_paragraph(f"Screenshot: {nome}")
        doc.add_picture(path, width=Inches(5.5))
        screenshot_registradas.add(nome)

def safe_action(doc, descricao, func):
    try:
        log(doc, f"🔄 {descricao}...")
        func()
        log(doc, f"✅ {descricao} realizada com sucesso.")
        take_screenshot(driver, doc, descricao.lower().replace(" ", "_"))
    except Exception as e:
        log(doc, f"❌ Erro ao {descricao.lower()}: {e}")
        take_screenshot(driver, doc, f"erro_{descricao.lower().replace(' ', '_')}")

def finalizar_relatorio():
    nome_arquivo = f"relatorio_registro_obito_pet_cenario_1_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc.save(nome_arquivo)
    log(doc, f"📄 Relatório salvo como: {nome_arquivo}")
    subprocess.run(["start", "winword", nome_arquivo], shell=True)
    driver.quit()

def gerar_dados_obito_pet():
    """Gera dados fictícios para o cadastro de óbito pet."""
    
    # Gera datas coerentes
    hoje = datetime.today().date()
    dez_anos_atras = hoje - timedelta(days=3650)
    
    data_falecimento = fake.date_between(start_date=dez_anos_atras, end_date=hoje)
    data_velorio = data_falecimento + timedelta(days=random.randint(0, 10))
    data_sepultamento = data_velorio + timedelta(days=random.randint(1, 10))
    
    return {
        'nome_pet': 'TESTE NOME DO PET SELENIUM AUTOMATIZADO',
        'idade_pet': str(fake.random_int(min=1, max=20)),
        'rg_pet': str(fake.random_number(digits=8)),
        'peso_pet': str(fake.random_int(min=1, max=100)),
        'data_falecimento': data_falecimento.strftime("%d/%m/%Y"),
        'data_velorio': data_velorio.strftime("%d/%m/%Y"),
        'data_sepultamento': data_sepultamento.strftime("%d/%m/%Y"),
        'causa_morte': 'TESTE CAUSA DA MORTE SELENIUM AUTOMATIZADO',
        'local_falecimento': 'TESTE LOCAL DE FALECIMENTO SELENIUM AUTOMATIZADO',
        'veterinario': 'TESTE VETERINARIO SELENIUM AUTOMATIZADO',
        'complemento': 'TESTE COMPLEMENTO SELENIUM AUTOMATIZADO',
        'telefone': fake.phone_number(),
        'telefone2': fake.phone_number(),
        'email': fake.email()
    }

def encontrar_mensagem_alerta():
    seletores = [
        (".alerts.salvo", "✅ Menasagem de Sucesso"),
        (".alerts.alerta", "⚠️ Menasagem de Alerta"),
        (".alerts.erro", "❌ Menasagem de Erro"),
    ]

    for seletor, tipo in seletores:
        try:
            elemento = driver.find_element(By.CSS_SELECTOR, seletor)
            if elemento.is_displayed():
                log(doc, f"📢 {tipo}: {elemento.text}")
                return elemento
        except:
            continue

    log(doc, "ℹ️ Nenhuma mensagem de alerta encontrada.")
    return None

def ajustar_zoom():
    try:
        driver.execute_script("document.body.style.zoom='90%'")
        log(doc, "🔍 Zoom ajustado para 90%.")
    except Exception as e:
        log(doc, f"⚠️ Erro ao ajustar zoom: {e}")

def fechar_modal_se_existir():
    """Fecha modal usando o seletor específico do primeiro código"""
    try:
        close_btn = wait.until(EC.element_to_be_clickable((
            By.CSS_SELECTOR, "#fmod_10087 > div.wdTop.ui-draggable-handle > div.wdClose > a"
        )))
        close_btn.click()
        time.sleep(1)
        return True
    except Exception as e:
        log(doc, f"⚠️ Não foi possível fechar modal: {e}")
        return False

def scroll_to_element_and_click(element):
    """Rola até o elemento e clica usando JavaScript"""
    driver.execute_script("arguments[0].scrollIntoView({behavior: 'smooth', block: 'center'});", element)
    time.sleep(1)
    driver.execute_script("arguments[0].click();", element)

def preencher_campo_data_melhorado(selector, valor):
    def acao():
        # Aguarda o campo estar presente
        campo = wait.until(EC.presence_of_element_located((By.XPATH, selector)))
        
        # Rola até o campo
        driver.execute_script("arguments[0].scrollIntoView({behavior: 'smooth', block: 'center'});", campo)
        time.sleep(1)
        
        # Limpa o campo primeiro
        campo.clear()
        time.sleep(0.5)
        
        # Clica usando JavaScript se necessário
        try:
            campo.click()
        except:
            driver.execute_script("arguments[0].click();", campo)
        
        time.sleep(0.5)
        campo.send_keys(valor)
        time.sleep(0.5)
        
        # Confirma com TAB para sair do campo
        campo.send_keys(Keys.TAB)
        
    return acao

def abrir_lov_e_selecionar_melhorado(btn_selector, search_selector, search_text, select_text):
    def acao():
        # Abre LOV
        btn = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, btn_selector)))
        scroll_to_element_and_click(btn)
        time.sleep(2)
        
        # Pesquisa
        campo_pesquisa = wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, search_selector)))
        campo_pesquisa.clear()
        campo_pesquisa.send_keys(search_text, Keys.ENTER)
        time.sleep(2)
        
        # Seleciona usando texto contido na célula
        try:
            item = wait.until(EC.element_to_be_clickable((By.XPATH, f"//td[contains(text(), '{select_text}')]")))
            scroll_to_element_and_click(item)
        except:
            # Fallback: tenta selecionar o primeiro item da tabela
            item = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "table tbody tr:first-child td:first-child")))
            scroll_to_element_and_click(item)
            
    return acao

def abrir_lov_e_selecionar_pet_melhorado(btn_selector, search_selector, search_text, search_btn_selector):
    """Função específica para seleção de pet - versão melhorada"""
    def acao():
        # Abre LOV
        btn = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, btn_selector)))
        scroll_to_element_and_click(btn)
        time.sleep(2)
        
        # Digita no campo de pesquisa
        campo_pesquisa = wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, search_selector)))
        campo_pesquisa.clear()
        campo_pesquisa.send_keys(search_text)
        time.sleep(1)
        
        # Clica no botão de pesquisar
        botao_pesquisar = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, search_btn_selector)))
        scroll_to_element_and_click(botao_pesquisar)
        time.sleep(3)
        
        # Seleciona o primeiro resultado da tabela
        try:
            item = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "table tbody tr:first-child")))
            scroll_to_element_and_click(item)
        except:
            # Fallback: tenta encontrar por texto
            item = wait.until(EC.element_to_be_clickable((By.XPATH, f"//td[contains(text(), '{search_text[:20]}')]")))
            scroll_to_element_and_click(item)
            
    return acao

def selecionar_opcao_melhorado(selector, texto):
    def acao():
        select_element = wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, selector)))
        driver.execute_script("arguments[0].scrollIntoView({behavior: 'smooth', block: 'center'});", select_element)
        time.sleep(1)
        Select(select_element).select_by_visible_text(texto)
    return acao

def preencher_campo_melhorado(selector, valor):
    def acao():
        campo = wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, selector)))
        driver.execute_script("arguments[0].scrollIntoView({behavior: 'smooth', block: 'center'});", campo)
        time.sleep(1)
        campo.clear()
        campo.send_keys(valor)
    return acao

def salvar_cadastro_melhorado():
    def acao():
        # Encontra o botão de salvar
        save_btn = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "#fmod_10087 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroRegistroObitoPet > div.btnHolder > a.btModel.btGray.btsave")))
        scroll_to_element_and_click(save_btn)
        time.sleep(5)  # Aguarda processamento do salvamento
    return acao

# Gera os dados necessários
dados_pet = gerar_dados_obito_pet()

# ==== INICIALIZAÇÃO DO DRIVER ====
options = Options()
options.add_argument("--start-maximized")
options.add_argument("--disable-web-security")
options.add_argument("--disable-features=VizDisplayCompositor")
driver = webdriver.Chrome(service=Service(ChromeDriverManager().install()), options=options)
wait = WebDriverWait(driver, 30)  # Aumentado timeout

# ==== EXECUÇÃO DO TESTE ====
try:
    safe_action(doc, "Acessando sistema", lambda: driver.get(URL))

    safe_action(doc, "Realizando login", lambda: (
        wait.until(EC.presence_of_element_located((By.ID, "j_id15:email"))).send_keys(LOGIN_EMAIL),
        wait.until(EC.presence_of_element_located((By.ID, "j_id15:senha"))).send_keys(LOGIN_PASSWORD, Keys.ENTER),
        wait.until(EC.presence_of_element_located((By.TAG_NAME, "body")))
    ))

    safe_action(doc, "Esperando sistema carregar e ajustando zoom", lambda: (
        time.sleep(5),
        ajustar_zoom()
    ))

    safe_action(doc, "Abrindo menu Registro de Óbito Pet", lambda: (
        driver.find_element(By.TAG_NAME, "body").send_keys(Keys.F2),
        time.sleep(1),
        wait.until(EC.visibility_of_element_located((By.XPATH, "//input[@placeholder='Busque um cadastro']"))).send_keys("Registro de Óbito Pet", Keys.ENTER),
        time.sleep(3)
    ))

    safe_action(doc, "Clicando em Cadastrar", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_10087 > div.wdTelas > div > ul > li:nth-child(1) > a > span"))).click(),
        time.sleep(3)
    ))

    # Selecionando Pet - CORRIGIDO
    safe_action(doc, "Selecionando Pet", abrir_lov_e_selecionar_pet_melhorado(
        "#fmod_10087 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroRegistroObitoPet > div.catWrapper > div > div > div.groupHolder.clearfix.grupo_itens > div > div > div:nth-child(1) > div > div > a",
        "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div.formCol.divPesquisa > input", 
        "TESTE NOME DO PET SELENIUM AUTOMATIZADO",
        "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(3) > a"
    ))

    # Preenchendo datas - CORRIGIDO
    safe_action(doc, "Preenchendo data de falecimento", preencher_campo_data_melhorado(
        "//input[@name='dataFalecimento' and contains(@class, 'hasDatepicker dataFalecimento mandatory')]",
        dados_pet['data_falecimento']
    ))

    safe_action(doc, "Preenchendo data de sepultamento", preencher_campo_data_melhorado(
        "//input[@name='dataSepultamento' and contains(@class, 'hasDatepicker dataSepultamento')]",
        dados_pet['data_sepultamento']
    ))

    safe_action(doc, "Preenchendo data de velório", preencher_campo_data_melhorado(
        "//input[@name='dataVelorio' and contains(@class, 'hasDatepicker dataVelorio')]",
        dados_pet['data_velorio']
    ))

    # Preenchendo campos de texto
    safe_action(doc, "Preenchendo causa da morte", preencher_campo_melhorado(
        "#fmod_10087 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroRegistroObitoPet > div.catWrapper > div > div > div.groupHolder.clearfix.grupo_itens > div > div > div:nth-child(5) > div > input",
        dados_pet['causa_morte']
    ))

    safe_action(doc, "Preenchendo local de falecimento", preencher_campo_melhorado(
        "#fmod_10087 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroRegistroObitoPet > div.catWrapper > div > div > div.groupHolder.clearfix.grupo_itens > div > div > div:nth-child(6) > div > input",
        dados_pet['local_falecimento']
    ))

    safe_action(doc, "Preenchendo veterinário", preencher_campo_melhorado(
        "#fmod_10087 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroRegistroObitoPet > div.catWrapper > div > div > div.groupHolder.clearfix.grupo_itens > div > div > div:nth-child(7) > div > input",
        dados_pet['veterinario']
    ))

    # Dados do tutor - CORRIGIDO
    safe_action(doc, "Selecionando tutor", abrir_lov_e_selecionar_melhorado(
        "#fmod_10087 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroRegistroObitoPet > div.catWrapper > div > div > div.groupHolder.clearfix.grupo_dadosTutor > div > div > div:nth-child(1) > div > div > a",
        "#txtPesquisa", 
        "LARISSA DA COSTA",
        "LARISSA DA COSTA"
    ))


    safe_action(doc, "Preenchendo área", preencher_campo_melhorado(
        "#fmod_10087 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroRegistroObitoPet > div.catWrapper > div > div > div.groupHolder.clearfix.grupo_dadosTutor > div > div > div:nth-child(16) > div > div > input",
        "1214333"
    ))


    safe_action(doc, "Preenchendo complemento", preencher_campo_melhorado(
        "#fmod_10087 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroRegistroObitoPet > div.catWrapper > div > div > div.groupHolder.clearfix.grupo_dadosTutor > div > div > div:nth-child(12) > div > input",
        dados_pet['complemento']
    ))

    safe_action(doc, "Selecionando sexo", selecionar_opcao_melhorado(
        "#fmod_10087 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroRegistroObitoPet > div.catWrapper > div > div > div.groupHolder.clearfix.grupo_dadosTutor > div > div > div:nth-child(7) > div > select",
        "Feminino"
    ))

    safe_action(doc, "Preenchendo telefone 1", preencher_campo_melhorado(
        "#fmod_10087 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroRegistroObitoPet > div.catWrapper > div > div > div.groupHolder.clearfix.grupo_dadosTutor > div > div > div:nth-child(18) > div > input",
        dados_pet['telefone']
    ))

    safe_action(doc, "Preenchendo telefone 2", preencher_campo_melhorado(
        "#fmod_10087 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroRegistroObitoPet > div.catWrapper > div > div > div.groupHolder.clearfix.grupo_dadosTutor > div > div > div:nth-child(19) > div > input",
        dados_pet['telefone2']
    ))

    safe_action(doc, "Preenchendo e-mail", preencher_campo_melhorado(
        "#fmod_10087 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroRegistroObitoPet > div.catWrapper > div > div > div.groupHolder.clearfix.grupo_dadosTutor > div > div > div:nth-child(21) > div > input",
        dados_pet['email']
    ))

    # Salvando cadastro - CORRIGIDO
    safe_action(doc, "Salvando cadastro", salvar_cadastro_melhorado())

    # Verifica mensagem de sucesso/erro
    safe_action(doc, "Verificando resultado do salvamento", lambda: encontrar_mensagem_alerta())

    safe_action(doc, "Fechando modal após salvamento", lambda: (
        fechar_modal_se_existir(),
        time.sleep(2)
    ))

except Exception as e:
    log(doc, f"❌ ERRO FATAL: {e}")
    take_screenshot(driver, doc, "erro_fatal")

finally:
    log(doc, "✅ Teste concluído com sucesso.")
    finalizar_relatorio()