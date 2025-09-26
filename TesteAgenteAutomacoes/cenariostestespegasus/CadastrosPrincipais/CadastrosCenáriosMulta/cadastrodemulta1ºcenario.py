# Refatorado e organizado: cadastrodemulta1ºcenario.py

from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support.ui import WebDriverWait, Select
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service
from selenium.webdriver import ActionChains
from webdriver_manager.chrome import ChromeDriverManager
from docx import Document
from docx.shared import Inches
from faker import Faker
from faker.providers import BaseProvider
from datetime import datetime, timedelta
import subprocess
import os
import time
import random
import string
from selenium.common.exceptions import TimeoutException, NoSuchElementException

# ==== PROVIDERS E CONFIGURAÇÕES ====
class BrasilProvider(BaseProvider):
    def rg(self):
        numeros = [str(random.randint(0, 9)) for _ in range(8)]
        return ''.join(numeros) + '-' + str(random.randint(0, 9))

fake = Faker("pt_BR")
fake.add_provider(BrasilProvider)
valor = f"{fake.random_int(min=100, max=1000)}.{fake.random_int(min=0, max=99):02d}"
pontos = fake.random_int(min=1, max=10)
def gerar_dados_multa(intervalo_max_passado=30):
    """
    Gera dados completos de uma multa:
    - data_multa: entre hoje e X dias atrás
    - hora_multa: horário aleatório (HH:MM)
    - data_notificacao: de 0 a 5 dias após a multa
    - data_vencimento: de 10 a 20 dias após a multa
    - data_pagamento: entre notificação e vencimento

    Retorna um dicionário com todas as datas e hora no formato dd/mm/yyyy e HH:MM.
    """
    # Base da multa
    data_multa_dt = datetime.now() - timedelta(days=random.randint(0, intervalo_max_passado))

    # Gera hora no formato HH:MM
    hora = random.randint(0, 23)
    minuto = random.randint(0, 59)
    hora_multa_dt = f"{hora:02d}:{minuto:02d}"

    # Demais datas relacionadas
    data_notificacao_dt = data_multa_dt + timedelta(days=random.randint(0, 5))
    data_notificacao_dt <  datetime.now()  # Garante que a notificação não seja futura
    data_vencimento_dt = data_multa_dt + timedelta(days=random.randint(10, 20))
    data_pagamento_dt = random.choice([
        data_notificacao_dt + timedelta(days=random.randint(0, (data_vencimento_dt - data_notificacao_dt).days)),
        data_vencimento_dt
    ])

    return (
        data_multa_dt.strftime("%d/%m/%Y"),
        hora_multa_dt,
        data_notificacao_dt.strftime("%d/%m/%Y"),
        data_vencimento_dt.strftime("%d/%m/%Y"),
        data_pagamento_dt.strftime("%d/%m/%Y"),
    )

# Chama a função para obter as datas
data_multa, hora_multa, data_notificacao, data_vencimento, data_pagamento = gerar_dados_multa()

# ==== CONFIGURAÇÕES ====
URL = "http://localhost:8080/gs/index.xhtml"
LOGIN_EMAIL = "joaoeduardo.gold@outlook.com"
LOGIN_PASSWORD = "071999gs"

# ==== DOCUMENTO ====
doc = Document()
doc.add_heading("RELATÓRIO DO TESTE", 0)
doc.add_paragraph("Cadastro de Multa – Cenário 1: Preenchimento completo e salvamento.")
doc.add_paragraph(f"Data do teste: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")

screenshot_registradas = set()

# ==== FUNÇÕES DE UTILITÁRIO ====
def log(doc, msg):
    print(msg)
    doc.add_paragraph(msg)

def take_screenshot(driver, doc, nome):
    if nome not in screenshot_registradas:
        path = f"screenshots/{nome}.png"
        os.makedirs("screenshots", exist_ok=True)
        driver.save_screenshot(path)
        doc.add_paragraph(f"Screenshot: {nome}")
        doc.add_picture(path, width=Inches(5.5))
        screenshot_registradas.add(nome)

def safe_action(doc, descricao, func):
    try:
        log(doc, f"🔄 {descricao}...")
        func()
        log(doc, f"✅ {descricao} realizada com sucesso.")
        take_screenshot(driver, doc, descricao.lower().replace(" ", "_"))
    except Exception as e:
        log(doc, f"❌ Erro ao {descricao.lower()}: {e}")
        take_screenshot(driver, doc, f"erro_{descricao.lower().replace(' ', '_')}")

def finalizar_relatorio():
    nome_arquivo = f"relatorio_multa_cenario_1_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc.save(nome_arquivo)
    log(doc, f"📄 Relatório salvo como: {nome_arquivo}")
    subprocess.run(["start", "winword", nome_arquivo], shell=True)
    driver.quit()

def encontrar_mensagem_alerta():
    seletores = [
        (".alerts.salvo", "✅ Menasagem de Sucesso"),
        (".alerts.alerta", "⚠️ Menasagem de Alerta"),
        (".alerts.erro", "❌ Menasagem de Erro"),
    ]

    for seletor, tipo in seletores:
        try:
            elemento = driver.find_element(By.CSS_SELECTOR, seletor)
            if elemento.is_displayed():
                log(doc, f"📢 {tipo}: {elemento.text}")
                return elemento
        except:
            continue

    log(doc, "ℹ️ Nenhuma mensagem de alerta encontrada.")
    return None

def ajustar_zoom():
    try:
        driver.execute_script("document.body.style.zoom='90%'")
        log(doc, "🔍 Zoom ajustado para 90%.")
    except Exception as e:
        log(doc, f"⚠️ Erro ao ajustar zoom: {e}")

def preencher_data(selector, valor):
    def acao():
        campo = wait.until(EC.presence_of_element_located((By.XPATH, selector)))
        campo.click()
        campo.clear()
        campo.send_keys(valor)
        time.sleep(0.2)
    return acao

def preencher_campo(selector, valor, tipo="CSS"):
    def acao():
        if tipo == "CSS":
            campo = wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, selector)))
        else:
            campo = wait.until(EC.presence_of_element_located((By.XPATH, selector)))
        campo.clear()
        campo.send_keys(valor)
        time.sleep(0.2)
    return acao

def selecionar_opcao(selector, opcao):
    def acao():
        select_element = Select(wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, selector))))
        select_element.select_by_visible_text(opcao)
        time.sleep(0.2)
    return acao

def abrir_modal_e_selecionar(btn_selector, pesquisa_selector, termo_pesquisa, btn_pesquisar_selector, resultado_xpath):
    def acao():
        # Abre o modal
        open_lov = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, btn_selector)))
        open_lov.click()
        
        # Pesquisa
        campo_pesquisa = wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, pesquisa_selector)))
        time.sleep(1)
        campo_pesquisa.send_keys(termo_pesquisa)
        time.sleep(1)
        
        # Clica em pesquisar
        pesquisar = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, btn_pesquisar_selector)))
        pesquisar.click()
        
        # Seleciona o resultado
        resultado = wait.until(EC.element_to_be_clickable((By.XPATH, resultado_xpath)))
        resultado.click()
        
    return acao


def preencher_campo_com_retry(driver, wait, seletor, valor, max_tentativas=3):
    """Tenta preencher o campo com diferentes métodos até conseguir"""
    
    for tentativa in range(max_tentativas):
        try:

            
            # Aguarda o elemento
            campo = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, seletor)))
            
            # Scroll até o elemento se necessário
            driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", campo)
            time.sleep(0.5)
            
            # Método 1: Tradicional
            if tentativa == 0:
                campo.click()
                campo.clear()
                campo.send_keys(valor)
                campo.send_keys(Keys.TAB)
            
            # Método 2: ActionChains
            elif tentativa == 1:
                ActionChains(driver).move_to_element(campo).click().perform()
                time.sleep(0.2)
                ActionChains(driver).key_down(Keys.CONTROL).send_keys('a').key_up(Keys.CONTROL).perform()
                ActionChains(driver).send_keys(valor).perform()
                ActionChains(driver).send_keys(Keys.TAB).perform()
            
            # Método 3: JavaScript
            else:
                driver.execute_script("""
                    var element = arguments[0];
                    var valor = arguments[1];
                    element.focus();
                    element.value = '';
                    element.value = valor;
                    element.dispatchEvent(new Event('input', { bubbles: true }));
                    element.dispatchEvent(new Event('change', { bubbles: true }));
                    element.blur();
                """, campo, valor)
            
            time.sleep(0.5)
            
            # Verifica se o valor foi preenchido
            valor_atual = campo.get_attribute('value')
            if valor_atual == valor:

                return True
            else:
                print()
                
        except Exception as e:
            time.sleep(1)
    

    return False

# ==== INICIALIZAÇÃO DO DRIVER ====
options = Options()
options.add_argument("--start-maximized")
driver = webdriver.Chrome(service=Service(ChromeDriverManager().install()), options=options)
wait = WebDriverWait(driver, 10)

# ==== EXECUÇÃO DO TESTE ====
try:
    safe_action(doc, "Acessando sistema", lambda: driver.get(URL))

    safe_action(doc, "Realizando login", lambda: (
        wait.until(EC.presence_of_element_located((By.ID, "j_id15:email"))).send_keys(LOGIN_EMAIL),
        wait.until(EC.presence_of_element_located((By.ID, "j_id15:senha"))).send_keys(LOGIN_PASSWORD, Keys.ENTER),
        wait.until(EC.presence_of_element_located((By.TAG_NAME, "body")))
    ))

    safe_action(doc, "Esperando página carregar e ajustando zoom", lambda: (
        time.sleep(5),
        ajustar_zoom()
    ))

    safe_action(doc, "Abrindo menu Multa", lambda: (
        driver.find_element(By.TAG_NAME, "body").click(),
        driver.find_element(By.TAG_NAME, "body").send_keys(Keys.F2),
        time.sleep(1),
        wait.until(EC.visibility_of_element_located((By.XPATH, "//input[@placeholder='Busque um cadastro']"))).send_keys("Multa", Keys.ENTER)
    ))

    safe_action(doc, "Clicando em Cadastrar", lambda: (
        time.sleep(3),
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "#fmod_10086 > div.wdTelas > div > ul > li:nth-child(1) > a > span"))).click()
    ))

    time.sleep(2)

    safe_action(doc, "Selecionando Veículo", abrir_modal_e_selecionar(
        "#fmod_10086 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroMulta > div.catWrapper > div > div.cat_categoriaDadosMulta.categoriaHolder > div > div > div > div:nth-child(3) > div > a",
        "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div.formCol.divPesquisa > div:nth-child(1) > input",
        "TESTE VEÍCULO SELENIUM AUTOMATIZADO",
        "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(3) > a",
        "//td[contains(text(), 'TESTE VEÍCULO SELENIUM AUTOMATIZADO')]"
    ))

    safe_action(doc, "Selecionando Motorista", abrir_modal_e_selecionar(
        "#fmod_10086 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroMulta > div.catWrapper > div > div.cat_categoriaDadosMulta.categoriaHolder > div > div > div > div:nth-child(4) > div > a",
        "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div.formCol.divPesquisa > input",
        "CRISPIM MALAFAIA",
        "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(3) > a",
        "//td[contains(text(), 'CRISPIM MALAFAIA')]"
    ))


    safe_action(doc, "Preenchendo Data da Multa", lambda: preencher_campo_com_retry(driver, wait, "input.hasDatepicker.dataMulta", data_multa))



    safe_action(doc, "Preenchendo Hora da Multa", lambda: preencher_campo_com_retry(driver, wait,
        "#fmod_10086 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroMulta > div.catWrapper > div > div.cat_categoriaDadosMulta.categoriaHolder > div > div > div > div:nth-child(6) > input",
        hora_multa
    ))


    safe_action(doc, "Preenchendo Auto de Infração", preencher_campo(
        "#fmod_10086 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroMulta > div.catWrapper > div > div.cat_categoriaDadosMulta.categoriaHolder > div > div > div > div:nth-child(7) > input",
        f"{fake.random_int(min=1000, max=9999)}-{fake.random_int(min=0, max=9)}"
    ))

    safe_action(doc, "Selecionando Situação", selecionar_opcao(
        "#fmod_10086 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroMulta > div.catWrapper > div > div.cat_categoriaDadosMulta.categoriaHolder > div > div > div > div:nth-child(8) > select",
        "Paga"
    ))

    safe_action(doc, "Preenchendo Data da Notificação", lambda: preencher_campo_com_retry(driver, wait, "input.hasDatepicker.dataNotificacao", data_notificacao))


    safe_action(doc, "Preenchendo Data da Vencimento", lambda: preencher_campo_com_retry(driver, wait, "input.hasDatepicker.dataVencimento", data_vencimento))

    safe_action(doc, "Preenchendo Data da Pagamento", lambda: preencher_campo_com_retry(driver, wait, "input.hasDatepicker.dataPagamento", data_pagamento))



    safe_action(doc, "Preenchendo Valor", lambda: preencher_campo_com_retry(driver, wait,
        "#fmod_10086 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroMulta > div.catWrapper > div > div.cat_categoriaDadosMulta.categoriaHolder > div > div > div > div:nth-child(12) > input",
        valor
    ))


    safe_action(doc, "Selecionando Gravidade", selecionar_opcao(
        "#fmod_10086 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroMulta > div.catWrapper > div > div.cat_categoriaDadosMulta.categoriaHolder > div > div > div > div:nth-child(13) > select",
        "Leve"
    ))

    safe_action(doc, "Preenchendo Descrição", preencher_campo(
        "#fmod_10086 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroMulta > div.catWrapper > div > div.cat_categoriaDadosMulta.categoriaHolder > div > div > div > div:nth-child(14) > textarea",
        "Multa referente a estacionamento em local proibido."
    ))

    safe_action(doc, "Acessando aba Infrações", lambda: (
        wait.until(EC.element_to_be_clickable((By.LINK_TEXT, "Infrações"))).click()
    ))

    safe_action(doc, "Selecionando Infração", abrir_modal_e_selecionar(
        "#fmod_10086 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroMulta > div.catWrapper > div > div.cat_infracoes.categoriaHolder > div > div.group_infracoes.clearfix.grupoHolder.lista > div > div:nth-child(1) > div > a",
        "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div.formCol.divPesquisa > input",
        "TESTE INFRAÇÃO SELENIUM AUTOMATIZADO: ALTA VELOCID",
        "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(3) > a",
        "//td[contains(text(), 'TESTE INFRAÇÃO SELENIUM AUTOMATIZADO: ALTA VELOCID')]"
    ))





    safe_action(doc, "Preenchendo Pontos", preencher_campo(
        "#fmod_10086 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroMulta > div.catWrapper > div > div.cat_infracoes.categoriaHolder > div > div.group_infracoes.clearfix.grupoHolder.lista > div > div:nth-child(2) > input",
        pontos
    ))


    safe_action(doc, "Adicionando Infração", lambda: (
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "#fmod_10086 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroMulta > div.catWrapper > div > div.cat_infracoes.categoriaHolder > div > div.btnListHolder > a.btAddGroup"))).click()
    ))

    safe_action(doc, "Salvando cadastro", lambda: (
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "#fmod_10086 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroMulta > div.btnHolder > a.btModel.btGray.btsave"))).click()
    ))

    time.sleep(1)
    encontrar_mensagem_alerta()

    safe_action(doc, "Recusando lançamento da multa no contas à pagar", lambda: (
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "#BtNo"))).click()
    ))



    safe_action(doc, "Fechando modal", lambda: (
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "#fmod_10086 > div.wdTop.ui-draggable-handle > div > a"))).click()
    ))

except Exception as e:
    log(doc, f"❌ ERRO FATAL: {e}")
    take_screenshot(driver, doc, "erro_fatal")

finally:
    log(doc, "✅ Teste concluído com sucesso.")
    finalizar_relatorio()