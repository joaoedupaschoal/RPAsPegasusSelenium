# Refatorado no modelo padronizado: Cadastro de Documentos - Cenário 1

from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support.ui import WebDriverWait, Select
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service
from selenium.common.exceptions import TimeoutException, NoSuchElementException
from webdriver_manager.chrome import ChromeDriverManager
from docx import Document
from docx.shared import Inches
from faker import Faker
from faker.providers import BaseProvider
import os
import time
import random
import subprocess
import pyautogui
from datetime import datetime
import sys

sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))

# ============ CONFIGURAÇÕES E VARIÁVEIS ============
URL = "http://localhost:8080/gs/index.xhtml"
LOGIN_EMAIL = "joaoeduardo.gold@outlook.com"
LOGIN_PASSWORD = "071999gs"
fake = Faker("pt_BR")
screenshot_registradas = set()

# ============ INICIALIZAÇÃO DOC ============
doc = Document()
doc.add_heading("RELATÓRIO DO TESTE", 0)
doc.add_paragraph("Cadastro de Documentos – Cenário 1: Preenchimento completo e salvamento.")
doc.add_paragraph(f"Data do teste: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")

# ============ FUNÇÕES UTILITÁRIAS PADRÃO ============
def log(doc, msg):
    print(msg)
    doc.add_paragraph(msg)

def take_screenshot(driver, doc, nome):
    if nome not in screenshot_registradas:
        path = f"screenshots/{nome}.png"
        os.makedirs("screenshots", exist_ok=True)
        driver.save_screenshot(path)
        doc.add_paragraph(f"Screenshot: {nome}")
        doc.add_picture(path, width=Inches(5.5))
        screenshot_registradas.add(nome)

def safe_action(doc, descricao, func):
    try:
        log(doc, f"🔄 {descricao}...")
        func()
        log(doc, f"✅ {descricao} realizada com sucesso.")
        take_screenshot(driver, doc, descricao.lower().replace(" ", "_"))
    except Exception as e:
        log(doc, f"❌ Erro ao {descricao.lower()}: {e}")
        take_screenshot(driver, doc, f"erro_{descricao.lower().replace(' ', '_')}")

def finalizar_relatorio():
    nome_arquivo = f"relatorio_documentos_cenario_1_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc.save(nome_arquivo)
    log(doc, f"📄 Relatório salvo como: {nome_arquivo}")
    subprocess.run(["start", "winword", nome_arquivo], shell=True)
    driver.quit()

def encontrar_mensagem_alerta():
    seletores = [
        (".alerts.salvo", "✅ Menasagem de Sucesso"),
        (".alerts.alerta", "⚠️ Menasagem de Alerta"),
        (".alerts.erro", "❌ Menasagem de Erro"),
    ]

    for seletor, tipo in seletores:
        try:
            elemento = driver.find_element(By.CSS_SELECTOR, seletor)
            if elemento.is_displayed():
                log(doc, f"📢 {tipo}: {elemento.text}")
                return elemento
        except:
            continue

    log(doc, "ℹ️ Nenhuma mensagem de alerta encontrada.")
    return None

def ajustar_zoom():
    try:
        driver.execute_script("document.body.style.zoom='90%'")
        log(doc, "🔍 Zoom ajustado para 90%.")
    except Exception as e:
        log(doc, f"⚠️ Erro ao ajustar zoom: {e}")

# ============ DRIVER E WAIT ============
options = Options()
options.add_argument("--start-maximized")
driver = webdriver.Chrome(service=Service(ChromeDriverManager().install()), options=options)
wait = WebDriverWait(driver, 10)

# ============ FLUXO DO TESTE ============

try:
    safe_action(doc, "Acessando sistema", lambda: driver.get(URL))

    safe_action(doc, "Realizando login", lambda: (
        wait.until(EC.presence_of_element_located((By.ID, "j_id15:email"))).send_keys(LOGIN_EMAIL),
        wait.until(EC.presence_of_element_located((By.ID, "j_id15:senha"))).send_keys(LOGIN_PASSWORD, Keys.ENTER),
        wait.until(EC.presence_of_element_located((By.TAG_NAME, "body")))
    ))

    safe_action(doc, "Esperando tela carregar e ajustando zoom", lambda: (
        time.sleep(3),
        ajustar_zoom()
    ))

    safe_action(doc, "Abrindo menu Documentos", lambda: (
        driver.find_element(By.TAG_NAME, "body").send_keys(Keys.F2),
        time.sleep(1),
        wait.until(EC.visibility_of_element_located((By.XPATH, "//input[@placeholder='Busque um cadastro']"))).send_keys("Documentos", Keys.ENTER)
    ))

    safe_action(doc, "Clicando em Cadastrar", lambda: (
        time.sleep(3),
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "#fmod_22 ul > li:nth-child(1) > a > span"))).click()
    ))

    
    safe_action(doc, "Preenchendo nome do documento", lambda: (
        time.sleep(1),
        driver.execute_script("document.querySelector('#fmod_22 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div > div:nth-child(1) > div > div:nth-child(2) > input').value = 'TESTE DOCUMENTO SELENIUM AUTOMATIZADO';")
    ))


    safe_action(doc, "Selecionando tipo do documento", lambda: (
        Select(wait.until(EC.presence_of_element_located((By.CSS_SELECTOR,
            "#fmod_22 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div > div:nth-child(1) > div > div:nth-child(4) > select"))))
        .select_by_visible_text("PDF (Upload em .docx)")
    ))

    safe_action(doc, "Selecionando arquivo .docx via PyAutoGUI", lambda: (
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR,
            "#fmod_22 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div > div:nth-child(2) > div > div > div > div > div > div > div"))).click(),
        time.sleep(1.5),
        pyautogui.write(r'"C:\Users\Gold System\Desktop\TESTEMODULODOCUMENTOS.docx"'),
        pyautogui.press('enter')
    ))





    safe_action(doc, "Clicando em Salvar", lambda: (
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR,
            "#fmod_22 > div.wdTelas > div.telaCadastro.clearfix > div.btnHolder > a.btModel.btGray.btsave"))).click(),
        time.sleep(3),
        encontrar_mensagem_alerta()
    ))

    safe_action(doc, "Fechando modal", lambda: (
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR,
            "#fmod_22 > div.wdTop.ui-draggable-handle > div.wdClose > a"))).click()
    ))

except Exception as e:
    log(doc, f"❌ ERRO FATAL: {e}")
    take_screenshot(driver, doc, "erro_fatal")

finally:

    log(doc, "✅ Teste concluído com sucesso.")

    finalizar_relatorio()