# Refatorado e organizado: cadastrodeescalamotorista2¬∫cenario.py

from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support.ui import Select, WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service
from webdriver_manager.chrome import ChromeDriverManager
from docx import Document
from docx.shared import Inches
from faker import Faker
from faker.providers import BaseProvider
from validate_docbr import CPF
from datetime import datetime, timedelta
import subprocess
import os
import time
import random
import sys 

sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))


# ==== PROVIDERS CUSTOMIZADOS ====
class BrasilProvider(BaseProvider):
    def rg(self):
        numeros = [str(random.randint(0, 9)) for _ in range(8)]
        return ''.join(numeros) + '-' + str(random.randint(0, 9))

fake = Faker("pt_BR")
fake.add_provider(BrasilProvider)

def gerar_datas_validas():
    """Gera datas coerentes para admiss√£o, in√≠cio e fim da escala, e vencimento da CNH."""
    hoje = datetime.today().date()
    
    # Data de admiss√£o entre 10 anos atr√°s e hoje
    data_admissao = fake.date_between(start_date=hoje - timedelta(days=3650), end_date=hoje)
    
    # Data de in√≠cio da escala entre hoje e 1 ano no futuro
    data_inicio = fake.date_between(start_date=hoje, end_date=hoje + timedelta(days=365))
    
    # Data fim entre 1 e 180 dias ap√≥s a data de in√≠cio
    data_fim = data_inicio + timedelta(days=random.randint(1, 180))
    
    # Vencimento CNH entre hoje e 10 anos no futuro
    vencimento_cnh = fake.date_between(start_date=hoje, end_date=hoje + timedelta(days=3650))
    
    return (data_admissao.strftime('%d/%m/%Y'), 
            data_inicio.strftime('%d/%m/%Y'), 
            data_fim.strftime('%d/%m/%Y'), 
            vencimento_cnh.strftime('%d/%m/%Y'))

def gerar_dados_documentos():
    """Gera documentos fict√≠cios para o cadastro."""
    carteira_trabalho = str(random.randint(10000000, 99999999))
    pis = fake.cpf().replace('.', '').replace('-', '')[:11]
    cnh = str(random.randint(10000000000, 99999999999))
    cpf = CPF().generate()
    
    return carteira_trabalho, pis, cnh, cpf

# Gera os dados necess√°rios
data_admissao, data_inicio, data_fim, vencimento_cnh = gerar_datas_validas()
carteira_trabalho, pis, cnh, cpf_valido = gerar_dados_documentos()

# ==== CONFIGURA√á√ïES ====
URL = "http://localhost:8080/gs/index.xhtml"
LOGIN_EMAIL = "joaoeduardo.gold@outlook.com"
LOGIN_PASSWORD = "071999gs"

# ==== DOCUMENTO ====
doc = Document()
doc.add_heading("RELAT√ìRIO DO TESTE", 0)
doc.add_paragraph("Cadastro de Especialidades ‚Äì Cen√°rio 1: Preenchimento completo e salvamento.")
doc.add_paragraph(f"Data do teste: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")

screenshot_registradas = set()

# ==== FUN√á√ïES DE UTILIT√ÅRIO ====
def log(doc, msg):
    print(msg)
    doc.add_paragraph(msg)

def take_screenshot(driver, doc, nome):
    if nome not in screenshot_registradas:
        path = f"screenshots/{nome}.png"
        os.makedirs("screenshots", exist_ok=True)
        driver.save_screenshot(path)
        doc.add_paragraph(f"Screenshot: {nome}")
        doc.add_picture(path, width=Inches(5.5))
        screenshot_registradas.add(nome)

def safe_action(doc, descricao, func):
    try:
        log(doc, f"üîÑ {descricao}...")
        func()
        log(doc, f"‚úÖ {descricao} realizada com sucesso.")
        take_screenshot(driver, doc, descricao.lower().replace(" ", "_"))
    except Exception as e:
        log(doc, f"‚ùå Erro ao {descricao.lower()}: {e}")
        take_screenshot(driver, doc, f"erro_{descricao.lower().replace(' ', '_')}")

def finalizar_relatorio():
    nome_arquivo = f"relatorio_especialidades_cenario_1_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc.save(nome_arquivo)
    log(doc, f"üìÑ Relat√≥rio salvo como: {nome_arquivo}")
    subprocess.run(["start", "winword", nome_arquivo], shell=True)
    driver.quit()

def encontrar_mensagem_alerta():
    seletores = [
        (".alerts.salvo", "‚úÖ Sucesso"),
        (".alerts.alerta", "‚ö†Ô∏è Alerta"),
        (".alerts.erro", "‚ùå Erro"),
    ]

    for seletor, tipo in seletores:
        try:
            elemento = driver.find_element(By.CSS_SELECTOR, seletor)
            if elemento.is_displayed():
                log(doc, f"üì¢ {tipo}: {elemento.text}")
                return elemento
        except:
            continue

    log(doc, "‚ÑπÔ∏è Nenhuma mensagem de alerta encontrada.")
    return None

def ajustar_zoom():
    try:
        driver.execute_script("document.body.style.zoom='90%'")
        log(doc, "üîç Zoom ajustado para 90%.")
    except Exception as e:
        log(doc, f"‚ö†Ô∏è Erro ao ajustar zoom: {e}")

# ==== INICIALIZA√á√ÉO DO DRIVER ====
options = Options()
options.add_argument("--start-maximized")
driver = webdriver.Chrome(service=Service(ChromeDriverManager().install()), options=options)
wait = WebDriverWait(driver, 20)


URL = "http://localhost:8080/gs/index.xhtml"
LOGIN_EMAIL = "joaoeduardo.gold@outlook.com"
LOGIN_PASSWORD = "071999gs"

# Configura√ß√£o do Faker
class BrasilProvider(BaseProvider):
    def rg(self):
        numeros = [str(random.randint(0, 9)) for _ in range(8)]
        return ''.join(numeros) + '-' + str(random.randint(0, 9))

fake = Faker("pt_BR")
fake.add_provider(BrasilProvider)

# Gera√ß√£o de dados aleat√≥rios
def gerar_datas_cronograma():
    hoje = datetime.today()
    
    # Gera uma data de leitura aleat√≥ria neste m√™s (entre dia 1 e 28)
    dia_leitura = random.randint(1, 28)
    data_leitura = hoje.replace(day=dia_leitura)

    # Gera uma data de entrega entre 5 e 10 dias depois da leitura
    dias_apos_leitura = random.randint(5, 10)
    data_entrega = data_leitura + timedelta(days=dias_apos_leitura)

    # Retorna as datas formatadas
    return {
        "data_leitura": data_leitura.strftime("%d/%m/%Y"),
        "data_entrega_conta": data_entrega.strftime("%d/%m/%Y")
    }

datas = gerar_datas_cronograma()

# Controle de screenshots √∫nicas
screenshot_registradas = set()
def registrar_screenshot_unico(nome, driver, doc, descricao=None):
    if nome not in screenshot_registradas:
        if descricao:
            log(doc, f"üì∏ {descricao}")
        take_screenshot(driver, doc, nome)
        screenshot_registradas.add(nome)

def main():
    doc = Document()
    doc.add_heading("RELAT√ìRIO DO TESTE", 0)
    doc.add_paragraph("Cadastro de Cronograma de Faturamento Teste.")
    doc.add_paragraph(f"üóïÔ∏è Data do teste: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")
    doc.add_paragraph("Neste teste, o rob√¥ preencher√° APENAS os campos obrigat√≥rios e salvar√° o cadastro de um novo Cronograma de Faturamento.")

    chrome_options = Options()
    chrome_options.add_argument("--start-maximized")
    driver = webdriver.Chrome(service=Service(ChromeDriverManager().install()), options=chrome_options)
    wait = WebDriverWait(driver, 10)

    def finalizar_relatorio():
        doc_name = f"relatorio_cronograma_faturamento_cenario_3_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        doc.save(doc_name)
        log(doc, f"üìÑ Relat√≥rio salvo como: {doc_name}")
        try:
            subprocess.run(["start", "winword", doc_name], shell=True)
        except Exception as e:
            log(doc, f"Erro ao abrir o Word: {e}")
        driver.quit()

    def login():
        wait.until(EC.presence_of_element_located((By.ID, "j_id15:email"))).send_keys(LOGIN_EMAIL)
        wait.until(EC.presence_of_element_located((By.ID, "j_id15:senha"))).send_keys(LOGIN_PASSWORD, Keys.ENTER)
        wait.until(EC.presence_of_element_located((By.TAG_NAME, "body")))
        time.sleep(2)

    def abrir_menu():
        driver.find_element(By.TAG_NAME, "body").send_keys(Keys.F2)
        campo = wait.until(EC.presence_of_element_located((By.XPATH, "//input[@placeholder='Busque um cadastro']")))
        campo.click()
        campo.send_keys("Cronograma de Faturamento", Keys.ENTER)
        time.sleep(3)

    def acessar_formulario():
        cadastrar = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, 
            "#fmod_10053 > div.wdTelas > div.telaInicial.clearfix.overflow.overflowY > ul > li:nth-child(1) > a > span")))
        cadastrar.click()
        time.sleep(2)

    def preencher_exercicio():
        log(doc, "üîÑ Preenchendo campo 'Exerc√≠cio'.")
        campo_exercicio = wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, 
            "#fmod_10053 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_10080.categoriaHolder > div > div > div:nth-child(2) > input")))
        campo_exercicio.send_keys(fake.random_int(min=2020, max=2030))
        log(doc, "‚úÖ Campo 'Exerc√≠cio' preenchido.")

    def selecionar_concessionaria():
        log(doc, "üîÑ Selecionando Concession√°ria.")
        open_lov_concessionaria = wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, 
            "#fmod_10053 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_10080.categoriaHolder > div > div > div:nth-child(3) > div > a")))
        open_lov_concessionaria.click()
        
        campo_pesquisa = wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, 
            "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(2) > input.nomePesquisa")))
        time.sleep(1)
        campo_pesquisa.send_keys('TESTE CONCESSION√ÅRIA DE ENERGIA SELENIUM AUTOMATIZ', Keys.ENTER)
        time.sleep(1)

        concessionaria = wait.until(EC.element_to_be_clickable((By.XPATH, 
            "//td[contains(text(), 'TESTE CONCESSION√ÅRIA DE ENERGIA SELENIUM AUTOMATIZ')]")))
        concessionaria.click()
        log(doc, "‚úÖ Concession√°ria selecionada.")



    def salvar():
        salvar_btn = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR,
            "#fmod_10053 > div.wdTelas > div.telaCadastro.clearfix > div.btnHolder > a.btModel.btGray.btsave")))
        salvar_btn.click()
        time.sleep(2)

    def fechar_modal():
        x_btn = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR,
            "#fmod_10053 > div.wdTop.ui-draggable-handle > div.wdClose > a")))
        x_btn.click()
        time.sleep(1)

    # EXECU√á√ÉO COM safe_action INDIVIDUAL PARA CADA A√á√ÉO
    if not safe_action(doc, "Acessando o sistema", lambda: driver.get(URL), driver, wait)[0]:
        finalizar_relatorio()
        return
    registrar_screenshot_unico("url_acessada", driver, doc, "Sistema acessado.")

    if not safe_action(doc, "Realizando login", login, driver, wait)[0]:
        finalizar_relatorio()
        return
    registrar_screenshot_unico("login_concluido", driver, doc, "Login realizado.")

    if not safe_action(doc, "Ajustando zoom", lambda: ajustar_zoom(driver), driver, wait)[0]:
        finalizar_relatorio()
        return

    if not safe_action(doc, "Abrindo menu Cronograma de Faturamento", abrir_menu, driver, wait)[0]:
        finalizar_relatorio()
        return
    registrar_screenshot_unico("menu_aberto", driver, doc, "Menu 'Cronograma de Faturamento' aberto.")

    if not safe_action(doc, "Acessando formul√°rio de cadastro", acessar_formulario, driver, wait)[0]:
        finalizar_relatorio()
        return
    registrar_screenshot_unico("formulario_aberto", driver, doc, "Formul√°rio de cadastro aberto.")

    # PREENCHIMENTO DOS CAMPOS - safe_action individual para cada campo
    if not safe_action(doc, "Preenchendo Exerc√≠cio", preencher_exercicio, driver, wait)[0]:
        finalizar_relatorio()
        return

    if not safe_action(doc, "Selecionando Concession√°ria", selecionar_concessionaria, driver, wait)[0]:
        finalizar_relatorio()
        return


    registrar_screenshot_unico("campos_preenchidos", driver, doc, "Todos os campos preenchidos.")

    # SALVANDO O CADASTRO
    if not safe_action(doc, "Clicando no bot√£o Salvar", salvar, driver, wait)[0]:
        finalizar_relatorio()
        return
    registrar_screenshot_unico("apos_salvar", driver, doc, "Clique no bot√£o Salvar realizado.")

    # VERIFICANDO MENSAGEM DE RETORNO
    _, tipo_alerta = encontrar_mensagem_alerta(driver, doc)
    if tipo_alerta == "sucesso":
        log(doc, "‚úÖ Mensagem de sucesso exibida.")
    elif tipo_alerta == "alerta":
        log(doc, "‚ö†Ô∏è Mensagem de alerta exibida.")
    elif tipo_alerta == "erro":
        log(doc, "‚ùå Mensagem de erro exibida.")
    else:
        log(doc, "‚ö†Ô∏è Nenhuma mensagem encontrada ap√≥s salvar.")
    registrar_screenshot_unico("mensagem_final", driver, doc, "Mensagem exibida ap√≥s salvar.")

    # FECHANDO O FORMUL√ÅRIO
    if not safe_action(doc, "Fechando formul√°rio", fechar_modal, driver, wait)[0]:
        finalizar_relatorio()
        return
    registrar_screenshot_unico("formulario_fechado", driver, doc, "Formul√°rio fechado.")

    log(doc, "‚úÖ Teste de cadastro de Cronograma de Faturamento conclu√≠do com sucesso.")
    finalizar_relatorio()

if __name__ == "__main__":
    main()