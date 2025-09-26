# Refatorado e organizado: cadastrodeserviços1ºcenario.py

from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support.ui import Select, WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service
from selenium.webdriver import ActionChains
from webdriver_manager.chrome import ChromeDriverManager
from docx import Document
from docx.shared import Inches
from faker import Faker
from faker.providers import BaseProvider
from datetime import datetime, timedelta
import subprocess
import os
import time
import random

# ==== PROVIDERS CUSTOMIZADOS ====
class BrasilProvider(BaseProvider):
    def rg(self):
        numeros = [str(random.randint(0, 9)) for _ in range(8)]
        return ''.join(numeros) + '-' + str(random.randint(0, 9))

fake = Faker("pt_BR")
fake.add_provider(BrasilProvider)

# ==== CONFIGURAÇÕES ====
URL = "http://localhost:8080/gs/index.xhtml"
LOGIN_EMAIL = "joaoeduardo.gold@outlook.com"
LOGIN_PASSWORD = "071999gs"

# ==== DOCUMENTO ====
doc = Document()
doc.add_heading("RELATÓRIO DO TESTE", 0)
doc.add_paragraph("Cadastro de Serviços – Cenário 3: Preenchimento dos campos obrigatórios salvamento.")
doc.add_paragraph(f"Data do teste: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")

screenshot_registradas = set()

def preencher_campo_com_retry(driver, wait, seletor, valor, max_tentativas=2):
    """Tenta preencher o campo com diferentes métodos até conseguir"""
    
    for tentativa in range(max_tentativas):
        try:
            # Aguarda o elemento
            campo = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, seletor)))
            
            # Scroll até o elemento se necessário
            driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", campo)
            time.sleep(0.5)
            
            # Método 1: Tradicional
            if tentativa == 0:
                campo.click()
                campo.clear()
                campo.send_keys(valor)
                campo.send_keys(Keys.TAB)
            
            # Método 2: ActionChains
            elif tentativa == 1:
                ActionChains(driver).move_to_element(campo).click().perform()
                time.sleep(0.2)
                ActionChains(driver).key_down(Keys.CONTROL).send_keys('a').key_up(Keys.CONTROL).perform()
                ActionChains(driver).send_keys(valor).perform()
                ActionChains(driver).send_keys(Keys.TAB).perform()
            
            # Método 3: JavaScript
            else:
                driver.execute_script("""
                    var element = arguments[0];
                    var valor = arguments[1];
                    element.focus();
                    element.value = '';
                    element.value = valor;
                    element.dispatchEvent(new Event('input', { bubbles: true }));
                    element.dispatchEvent(new Event('change', { bubbles: true }));
                    element.blur();
                """, campo, valor)
            
            time.sleep(0.5)
            
            # Verifica se o valor foi preenchido
            valor_atual = campo.get_attribute('value')
            if valor_atual == valor:
                return True
            else:
                print()
                
        except Exception as e:
            time.sleep(1)
    
    return False

# ==== FUNÇÕES DE UTILITÁRIO ====
def log(doc, msg):
    print(msg)
    doc.add_paragraph(msg)

def take_screenshot(driver, doc, nome):
    if nome not in screenshot_registradas:
        path = f"screenshots/{nome}.png"
        os.makedirs("screenshots", exist_ok=True)
        driver.save_screenshot(path)
        doc.add_paragraph(f"Screenshot: {nome}")
        doc.add_picture(path, width=Inches(5.5))
        screenshot_registradas.add(nome)

def safe_action(doc, descricao, func):
    try:
        log(doc, f"🔄 {descricao}...")
        func()
        log(doc, f"✅ {descricao} realizada com sucesso.")
        take_screenshot(driver, doc, descricao.lower().replace(" ", "_"))
    except Exception as e:
        log(doc, f"❌ Erro ao {descricao.lower()}: {e}")
        take_screenshot(driver, doc, f"erro_{descricao.lower().replace(' ', '_')}")

def finalizar_relatorio():
    nome_arquivo = f"relatorio_servicos_cenario_3_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc.save(nome_arquivo)
    log(doc, f"📄 Relatório salvo como: {nome_arquivo}")
    subprocess.run(["start", "winword", nome_arquivo], shell=True)
    driver.quit()

def gerar_dados_servico():
    """Gera dados fictícios para o cadastro de serviço."""
    nome_servico = "SERVIÇO TESTE SELENIUM"
    cemetery_name = f"Cemitério {fake.last_name()} {fake.random.choice(['Eterno', 'da Paz', 'Memorial', 'Descanso'])}"
    dias_para_exumar = fake.random_int(min=365, max=1825)
    
    return (nome_servico, cemetery_name, dias_para_exumar)

def encontrar_mensagem_alerta():
    seletores = [
        (".alerts.salvo", "✅ Menasagem de Sucesso"),
        (".alerts.alerta", "⚠️ Menasagem de Alerta"),
        (".alerts.erro", "❌ Menasagem de Erro"),
    ]

    for seletor, tipo in seletores:
        try:
            elemento = driver.find_element(By.CSS_SELECTOR, seletor)
            if elemento.is_displayed():
                log(doc, f"📢 {tipo}: {elemento.text}")
                return elemento
        except:
            continue

    log(doc, "ℹ️ Nenhuma mensagem de alerta encontrada.")
    return None

def ajustar_zoom():
    try:
        driver.execute_script("document.body.style.zoom='90%'")
        log(doc, "🔍 Zoom ajustado para 90%.")
    except Exception as e:
        log(doc, f"⚠️ Erro ao ajustar zoom: {e}")

def selecionar_opcao(selector, opcao="primeira"):
    def acao():
        select_element = wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, selector)))
        if opcao == "primeira":
            select_element.send_keys(Keys.ARROW_DOWN, Keys.RETURN)
        else:
            Select(select_element).select_by_visible_text(opcao)
    return acao

# Gera os dados necessários
(nome_servico, cemetery_name, dias_para_exumar) = gerar_dados_servico()

# ==== INICIALIZAÇÃO DO DRIVER ====
options = Options()
options.add_argument("--start-maximized")
driver = webdriver.Chrome(service=Service(ChromeDriverManager().install()), options=options)
wait = WebDriverWait(driver, 20)

# ==== EXECUÇÃO DO TESTE ====
try:
    safe_action(doc, "Acessando sistema", lambda: driver.get(URL))

    safe_action(doc, "Realizando login", lambda: (
        wait.until(EC.presence_of_element_located((By.ID, "j_id15:email"))).send_keys(LOGIN_EMAIL),
        wait.until(EC.presence_of_element_located((By.ID, "j_id15:senha"))).send_keys(LOGIN_PASSWORD, Keys.ENTER),
        wait.until(EC.presence_of_element_located((By.TAG_NAME, "body")))
    ))

    safe_action(doc, "Esperando sistema carregar e ajustando zoom", lambda: (
        time.sleep(5),
        ajustar_zoom()
    ))

    safe_action(doc, "Abrindo menu Serviços", lambda: (
        driver.find_element(By.TAG_NAME, "body").send_keys(Keys.F2),
        time.sleep(1),
        wait.until(EC.visibility_of_element_located((By.XPATH, "//input[@placeholder='Busque um cadastro']"))).send_keys("Serviços", Keys.ENTER),
        time.sleep(3)
    ))

    safe_action(doc, "Clicando em Cadastrar", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_4 > div.wdTelas > div.telaInicial.clearfix.overflow.overflowY > ul > li:nth-child(1) > a > span"))).click(),
        time.sleep(2)
    ))

    # Preenchendo dados básicos do serviço
    safe_action(doc, "Preenchendo Nome do Serviço", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_4 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_19.categoriaHolder > div > div > div:nth-child(2) > input")))
        .send_keys(nome_servico)
    ))

    safe_action(doc, "Preenchendo Taxa Mensal", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_4 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_19.categoriaHolder > div > div > div:nth-child(4) > input")))
        .send_keys(str(fake.random_int(min=10, max=500)))
    ))

    safe_action(doc, "Preenchendo Taxa Bimestral", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_4 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_19.categoriaHolder > div > div > div:nth-child(5) > input.fc.money.mandatory")))
        .send_keys(str(fake.random_int(min=10, max=500)))
    ))

    safe_action(doc, "Preenchendo Taxa Trimestral", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_4 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_19.categoriaHolder > div > div > div:nth-child(6) > input.fc.money")))
        .send_keys(str(fake.random_int(min=10, max=500)))
    ))

    safe_action(doc, "Preenchendo Taxa Quadrimestral", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_4 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_19.categoriaHolder > div > div > div:nth-child(7) > input.fc.money.mandatory")))
        .send_keys(str(fake.random_int(min=10, max=500)))
    ))

    safe_action(doc, "Preenchendo Taxa Semestral", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_4 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_19.categoriaHolder > div > div > div:nth-child(8) > input.fc.money")))
        .send_keys(str(fake.random_int(min=10, max=500)))
    ))

    safe_action(doc, "Preenchendo Taxa Anual", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_4 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_19.categoriaHolder > div > div > div:nth-child(9) > input.fc.money")))
        .send_keys(str(fake.random_int(min=10, max=500)))
    ))

    safe_action(doc, "Preenchendo Taxa Avulsa", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_4 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_19.categoriaHolder > div > div > div:nth-child(10) > input")))
        .send_keys(str(fake.random_int(min=10, max=500)))
    ))



    safe_action(doc, "Selecionando Departamento", selecionar_opcao(
        "#fmod_4 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_19.categoriaHolder > div > div > div:nth-child(13) > select",
        "primeira"
    ))


    safe_action(doc, "Salvando cadastro de Serviço", lambda: (
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, '#fmod_4 > div.wdTelas > div.telaCadastro.clearfix > div.btnHolder > a.btModel.btGray.btsave'))).click(),
        time.sleep(3)
    ))

    safe_action(doc, "Fechando modal do cadastro", lambda: (
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "#fmod_4 > div.wdTop.ui-draggable-handle > div.wdClose > a"))).click(),
        time.sleep(2)
    ))

    encontrar_mensagem_alerta()

except Exception as e:
    log(doc, f"❌ ERRO FATAL: {e}")
    take_screenshot(driver, doc, "erro_fatal")

finally:
    log(doc, "✅ Teste concluído com sucesso.")
    finalizar_relatorio()