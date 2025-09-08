# Código corrigido e organizado: cadastro_conta_bancaria.py

from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support.ui import Select, WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service
from webdriver_manager.chrome import ChromeDriverManager
from docx import Document
from docx.shared import Inches
from faker import Faker
from faker.providers import BaseProvider
from validate_docbr import CPF
from datetime import datetime, timedelta
import subprocess
from selenium.webdriver import ActionChains
import os, sys
import time
import random
import re

sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))


# ==== PROVIDERS CUSTOMIZADOS ====
class BrasilProvider(BaseProvider):
    def rg(self):
        numeros = [str(random.randint(0, 9)) for _ in range(8)]
        return ''.join(numeros) + '-' + str(random.randint(0, 9))

fake = Faker("pt_BR")
fake.add_provider(BrasilProvider)

# ==== CONFIGURAÇÕES ====
URL = "http://localhost:8080/gs/index.xhtml"
LOGIN_EMAIL = "joaoeduardo.gold@outlook.com"
LOGIN_PASSWORD = "071999gs"

# ==== DOCUMENTO ====
doc = Document()
doc.add_heading("RELATÓRIO DO TESTE", 0)
doc.add_paragraph("Cadastro de Conta Bancária – Cenário: Preenchimento completo e salvamento.")
doc.add_paragraph(f"Data do teste: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")

screenshot_registradas = set()

# ==== FUNÇÕES DE UTILITÁRIO ====
def log(doc, msg):
    """Registra mensagem no console e no documento"""
    print(msg)
    doc.add_paragraph(msg)

def sanitize_filename(nome):
    """Sanitiza nome do arquivo para evitar caracteres inválidos"""
    nome = nome.lower()
    nome = re.sub(r'[^a-z0-9_]', '_', nome)
    return nome

def take_screenshot(driver, doc, nome):
    """Captura screenshot e adiciona ao documento"""
    nome_limpo = sanitize_filename(nome)
    if nome_limpo not in screenshot_registradas:
        os.makedirs("screenshots", exist_ok=True)
        path = f"screenshots/{nome_limpo}.png"
        driver.save_screenshot(path)
        doc.add_paragraph(f"Screenshot: {nome}")
        doc.add_picture(path, width=Inches(5.5))
        screenshot_registradas.add(nome_limpo)

def safe_action(doc, descricao, func):
    """Executa ação com tratamento de erro e screenshot"""
    try:
        log(doc, f"🔄 {descricao}...")
        func()
        log(doc, f"✅ {descricao} realizada com sucesso.")
        take_screenshot(driver, doc, descricao.lower().replace(" ", "_"))
    except Exception as e:
        log(doc, f"❌ Erro ao {descricao.lower()}: {e}")
        take_screenshot(driver, doc, f"erro_{descricao.lower().replace(' ', '_')}")

def finalizar_relatorio():
    """Finaliza e salva o relatório"""
    nome_arquivo = f"relatorio_conta_bancaria_cenario_1_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc.save(nome_arquivo)
    log(doc, f"📄 Relatório salvo como: {nome_arquivo}")
    try:
        subprocess.run(["start", "winword", nome_arquivo], shell=True)
    except:
        print("Não foi possível abrir o Word automaticamente")
    driver.quit()

def preencher_campo_com_retry(driver, wait, seletor, valor, max_tentativas=3):
    """Tenta preencher campo com diferentes métodos até conseguir"""
    for tentativa in range(max_tentativas):
        try:
            # Aguarda o elemento
            campo = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, seletor)))
            
            # Scroll até o elemento
            driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", campo)
            time.sleep(0.5)
            
            if tentativa == 0:
                # Método 1: Tradicional
                campo.click()
                campo.clear()
                campo.send_keys(valor)
                campo.send_keys(Keys.TAB)
            elif tentativa == 1:
                # Método 2: ActionChains
                ActionChains(driver).move_to_element(campo).click().perform()
                time.sleep(0.2)
                ActionChains(driver).key_down(Keys.CONTROL).send_keys('a').key_up(Keys.CONTROL).perform()
                ActionChains(driver).send_keys(valor).perform()
                ActionChains(driver).send_keys(Keys.TAB).perform()
            else:
                # Método 3: JavaScript
                driver.execute_script("""
                    var element = arguments[0];
                    var valor = arguments[1];
                    element.focus();
                    element.value = '';
                    element.value = valor;
                    element.dispatchEvent(new Event('input', { bubbles: true }));
                    element.dispatchEvent(new Event('change', { bubbles: true }));
                    element.blur();
                """, campo, valor)
            
            time.sleep(0.5)
            
            # Verifica se o valor foi preenchido
            valor_atual = campo.get_attribute('value')
            if valor_atual == valor:
                return True
                
        except Exception as e:
            print(f"Tentativa {tentativa + 1} falhou: {e}")
            time.sleep(1)
    
    return False

def encontrar_mensagem_alerta():
    """Procura por mensagens de alerta no sistema"""
    seletores = [
        (".alerts.salvo", "✅ Sucesso"),
        (".alerts.alerta", "⚠️ Alerta"),
        (".alerts.erro", "❌ Erro"),
    ]

    for seletor, tipo in seletores:
        try:
            elemento = driver.find_element(By.CSS_SELECTOR, seletor)
            if elemento.is_displayed():
                log(doc, f"📢 {tipo}: {elemento.text}")
                return elemento
        except:
            continue

    log(doc, "ℹ️ Nenhuma mensagem de alerta encontrada.")
    return None

def ajustar_zoom():
    """Ajusta o zoom da página para melhor visualização"""
    try:
        driver.execute_script("document.body.style.zoom='90%'")
        log(doc, "🔍 Zoom ajustado para 90%.")
    except Exception as e:
        log(doc, f"⚠️ Erro ao ajustar zoom: {e}")

def abrir_modal_e_selecionar(btn_selector, pesquisa_selector, termo_pesquisa, btn_pesquisar_selector, resultado_xpath):
    """Abre modal e seleciona um item"""
    def acao():
        # Abre o modal
        open_lov = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, btn_selector)))
        open_lov.click()
        time.sleep(3)

        # Aguarda campo pesquisa
        campo_pesquisa = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, pesquisa_selector)))
        campo_pesquisa.clear()
        campo_pesquisa.send_keys(termo_pesquisa)

        # Clica pesquisar
        pesquisar = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, btn_pesquisar_selector)))
        pesquisar.click()
        time.sleep(1)

        # Espera o resultado e clica
        wait.until(EC.element_to_be_clickable((By.XPATH, resultado_xpath)))
        resultado = driver.find_element(By.XPATH, resultado_xpath)
        driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", resultado)
        time.sleep(0.2)
        resultado.click()

    return acao

def selecionar_opcao(selector, texto):
    """Seleciona opção em dropdown"""
    def acao():
        select_element = wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, selector)))
        Select(select_element).select_by_visible_text(texto)
    return acao

def preencher_campo_simples(selector, valor):
    """Função auxiliar para preenchimento simples de campos"""
    def acao():
        campo = wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, selector)))
        campo.clear()
        campo.send_keys(valor)
    return acao

# ==== INICIALIZAÇÃO DO DRIVER ====
options = Options()
options.add_argument("--start-maximized")
driver = webdriver.Chrome(service=Service(ChromeDriverManager().install()), options=options)
wait = WebDriverWait(driver, 20)

# ==== EXECUÇÃO DO TESTE ====
try:
    # Login no sistema
    safe_action(doc, "Acessando sistema", lambda: driver.get(URL))

    safe_action(doc, "Realizando login", lambda: (
        wait.until(EC.presence_of_element_located((By.ID, "j_id15:email"))).send_keys(LOGIN_EMAIL),
        wait.until(EC.presence_of_element_located((By.ID, "j_id15:senha"))).send_keys(LOGIN_PASSWORD, Keys.ENTER),
        wait.until(EC.presence_of_element_located((By.TAG_NAME, "body")))
    ))

    safe_action(doc, "Esperando sistema carregar e ajustando zoom", lambda: (
        time.sleep(5),
        ajustar_zoom()
    ))

    # Navegação para o cadastro
    safe_action(doc, "Abrindo menu Conta Bancária", lambda: (
        driver.find_element(By.TAG_NAME, "body").send_keys(Keys.F2),
        time.sleep(1),
        wait.until(EC.visibility_of_element_located((By.XPATH, "//input[@placeholder='Busque um cadastro']"))).send_keys("Conta Bancária", Keys.ENTER)
    ))

    safe_action(doc, "Clicando em Cadastrar", lambda: (
        time.sleep(2),
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_10 > div.wdTelas > div.telaInicial.clearfix.overflow.overflowY > ul > li:nth-child(1) > a > span"))).click()
    ))

    # Preenchimento dos campos principais
    safe_action(doc, "Preenchendo Nome", lambda: preencher_campo_com_retry(
        driver, wait,
        "#fmod_10 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_22.categoriaHolder > div > div > div:nth-child(2) > input",
        "TESTE CONTA BANCÁRIA SELENIUM AUTOMATIZADO"
    ))

    safe_action(doc, "Selecionando Banco", selecionar_opcao(
        "#fmod_10 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_22.categoriaHolder > div > div > div:nth-child(3) > select",
        "001 - Banco do Brasil S.A."
    ))

    safe_action(doc, "Preenchendo Nome Cedente", lambda: preencher_campo_com_retry(
        driver, wait,
        "#fmod_10 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_22.categoriaHolder > div > div > div:nth-child(5) > input",
        "TESTE CEDENTE SELENIUM AUTOMATIZADO"
    ))

    safe_action(doc, "Selecionando Tipo de Pessoa", selecionar_opcao(
        "#fmod_10 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_22.categoriaHolder > div > div > div:nth-child(6) > select",
        "Física"
    ))

    safe_action(doc, "Preenchendo CPF", lambda: preencher_campo_com_retry(
        driver, wait,
        "#fmod_10 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_22.categoriaHolder > div > div > div:nth-child(7) > input",
        fake.cpf().replace('.', '').replace('-', '')
    ))

    safe_action(doc, "Preenchendo Agência-DV / Código Cedente-DV", lambda: preencher_campo_com_retry(
        driver, wait,
        "#fmod_10 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_22.categoriaHolder > div > div > div:nth-child(9) > input",
        "TESTE COD"
    ))

    safe_action(doc, "Preenchendo Taxa de Antecipação", lambda: preencher_campo_com_retry(
        driver, wait,
        "#fmod_10 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_22.categoriaHolder > div > div > div:nth-child(10) > input",
        str(fake.random_int(min=1000, max=100000))
    ))

    safe_action(doc, "Selecionando Aceite", selecionar_opcao(
        "#fmod_10 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_22.categoriaHolder > div > div > div:nth-child(11) > select",
        "Sim"
    ))

    safe_action(doc, "Preenchendo Espécie Documento", lambda: preencher_campo_com_retry(
        driver, wait,
        "#fmod_10 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_22.categoriaHolder > div > div > div:nth-child(12) > input",
        "DM"
    ))

    safe_action(doc, "Preenchendo Carteira", lambda: preencher_campo_com_retry(
        driver, wait,
        "#fmod_10 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_22.categoriaHolder > div > div > div:nth-child(13) > input",
        str(fake.random_int(min=0, max=999))
    ))

    safe_action(doc, "Selecionando Tipo de Carteira", selecionar_opcao(
        "#fmod_10 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_22.categoriaHolder > div > div > div:nth-child(14) > select",
        "Com Registro"
    ))

    safe_action(doc, "Preenchendo Código Empresa Banco", lambda: preencher_campo_com_retry(
        driver, wait,
        "#fmod_10 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_22.categoriaHolder > div > div > div:nth-child(15) > input",
        str(fake.random_int(min=10000000000000, max=999999999999999))
    ))

    safe_action(doc, "Preenchendo Identificador Cedente", lambda: preencher_campo_com_retry(
        driver, wait,
        "#fmod_10 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_22.categoriaHolder > div > div > div:nth-child(16) > input",
        str(fake.random_int(min=100000000000000, max=999999999999999))
    ))

    # Preenchimento das instruções
    safe_action(doc, "Preenchendo Instrução 1", preencher_campo_simples(
        "#fmod_10 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_22.categoriaHolder > div > div > div:nth-child(17) > input",
        'TESTE INSTRUÇÃO 1 SELENIUM AUTOMATIZADO'
    ))

    safe_action(doc, "Preenchendo Instrução 2", preencher_campo_simples(
        "#fmod_10 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_22.categoriaHolder > div > div > div:nth-child(18) > input",
        'TESTE INSTRUÇÃO 2 SELENIUM AUTOMATIZADO'
    ))

    safe_action(doc, "Preenchendo Instrução 3", preencher_campo_simples(
        "#fmod_10 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_22.categoriaHolder > div > div > div:nth-child(19) > input",
        'TESTE INSTRUÇÃO 3 SELENIUM AUTOMATIZADO'
    ))

    safe_action(doc, "Preenchendo Instrução 4", preencher_campo_simples(
        "#fmod_10 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_22.categoriaHolder > div > div > div:nth-child(20) > input",
        'TESTE INSTRUÇÃO 4 SELENIUM AUTOMATIZADO'
    ))

    safe_action(doc, "Preenchendo Endereço Completo", preencher_campo_simples(
        "#fmod_10 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_22.categoriaHolder > div > div > div:nth-child(22) > input",
        'Rua Manoel Caldeira Filho, 1735, Cidade Jardim, São Paulo - SP, 15081-115'
    ))

    # Seleção de Conta Débito
    safe_action(doc, "Selecionando Conta Débito", abrir_modal_e_selecionar(
        "#fmod_10 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_22.categoriaHolder > div > div > div:nth-child(23) > div > div > a",
        "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(2) > input",
        "774",
        "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(3) > a",
        "//tr[td[3][text()='774']]//a[contains(@class, 'linkAlterar')]"
    ))

    # Continuação do preenchimento
    safe_action(doc, "Preenchendo Local Pagamento", preencher_campo_simples(
        "#fmod_10 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_22.categoriaHolder > div > div > div:nth-child(24) > input",
        'TESTE LOCAL PAGAMENTO SELENIUM AUTOMATIZADO'
    ))

    safe_action(doc, "Selecionando Cobrança Boleto Pré Impresso", selecionar_opcao(
        "#fmod_10 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_22.categoriaHolder > div > div > div:nth-child(25) > select",
        "Não"
    ))

    safe_action(doc, "Preenchendo Quantidade de Dias para Devolução", lambda: preencher_campo_com_retry(
        driver, wait,
        "#fmod_10 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_22.categoriaHolder > div > div > div:nth-child(26) > input",
        str(fake.random_int(min=1, max=30))
    ))

    safe_action(doc, "Selecionando Tipo de Emissão", selecionar_opcao(
        "#fmod_10 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_22.categoriaHolder > div > div > div:nth-child(27) > select",
        "Banco"
    ))

    safe_action(doc, "Preenchendo Campo Auxiliar", preencher_campo_simples(
        "#fmod_10 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_22.categoriaHolder > div > div > div:nth-child(29) > input",
        'TESTE CAMPO AUXILIAR SELENIUM AUTOMATIZADO'
    ))

    safe_action(doc, "Preenchendo Multa", lambda: preencher_campo_com_retry(
        driver, wait,
        "#fmod_10 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_22.categoriaHolder > div > div > div:nth-child(32) > input",
        str(fake.random_int(min=1000, max=999999))
    ))

    safe_action(doc, "Preenchendo Juros de Mora", lambda: preencher_campo_com_retry(
        driver, wait,
        "#fmod_10 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_22.categoriaHolder > div > div > div:nth-child(35) > input",
        str(fake.random_int(min=1000, max=999999))
    ))

    safe_action(doc, "Preenchendo Nome Sacador Avalista", preencher_campo_simples(
        "#fmod_10 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_22.categoriaHolder > div > div > div:nth-child(36) > input",
        fake.name()
    ))

    safe_action(doc, "Selecionando Tipo de Desconto", selecionar_opcao(
        "#fmod_10 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_22.categoriaHolder > div > div > div:nth-child(38) > select",
        "Moeda"
    ))

    safe_action(doc, "Preenchendo Desconto", lambda: preencher_campo_com_retry(
        driver, wait,
        "#fmod_10 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_22.categoriaHolder > div > div > div:nth-child(39) > input",
        str(fake.random_int(min=1000, max=999999))
    ))

    safe_action(doc, "Preenchendo Quantidade de Dias para Desconto", lambda: preencher_campo_com_retry(
        driver, wait,
        "#fmod_10 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_22.categoriaHolder > div > div > div:nth-child(40) > input",
        str(fake.random_int(min=0, max=30))
    ))

    safe_action(doc, "Selecionando Impedir Remessa Inconsistente", selecionar_opcao(
        "#fmod_10 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_22.categoriaHolder > div > div > div:nth-child(49) > select",
        "Sim"
    ))

    # Aba Taxas
    safe_action(doc, "Acessando aba Taxas", lambda: 
        wait.until(EC.element_to_be_clickable((By.LINK_TEXT, "Taxas"))).click()
    )

    safe_action(doc, "Preenchendo Taxa para Liquidação", lambda: preencher_campo_com_retry(
        driver, wait,
        "#fmod_10 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_10091.categoriaHolder > div > div > div:nth-child(2) > input",
        str(fake.random_int(min=1000, max=999999))
    ))

    safe_action(doc, "Preenchendo Taxa para Registro", lambda: preencher_campo_com_retry(
        driver, wait,
        "#fmod_10 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_10091.categoriaHolder > div > div > div:nth-child(3) > input",
        str(fake.random_int(min=1000, max=999999))
    ))

    safe_action(doc, "Preenchendo Taxa para Alteração", lambda: preencher_campo_com_retry(
        driver, wait,
        "#fmod_10 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_10091.categoriaHolder > div > div > div:nth-child(4) > input",
        str(fake.random_int(min=1000, max=999999))
    ))

    # Aba Pix
    safe_action(doc, "Acessando aba Pix", lambda: 
        wait.until(EC.element_to_be_clickable((By.LINK_TEXT, "Pix"))).click()
    )

    safe_action(doc, "Selecionando Tipo de Chave PIX", selecionar_opcao(
        "#fmod_10 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_10092.categoriaHolder > div > div > div:nth-child(2) > select",
        "Celular"
    ))

    safe_action(doc, "Preenchendo Chave PIX", preencher_campo_simples(
        "#fmod_10 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_10092.categoriaHolder > div > div > div:nth-child(3) > input",
        fake.phone_number()
    ))

    # Salvamento
    safe_action(doc, "Salvando Cadastro", lambda: 
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, 
            "#fmod_10 > div.wdTelas > div.telaCadastro.clearfix > div.btnHolder > a.btModel.btGray.btsave"
        ))).click()
    )


    safe_action(doc, "Fechando modal após salvamento", lambda: 
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, 
            "#fmod_10 > div.wdTop.ui-draggable-handle > div.wdClose > a"
        ))).click()
    )

    time.sleep(6)

    encontrar_mensagem_alerta()

except Exception as e:
    log(doc, f"❌ ERRO FATAL: {e}")
    take_screenshot(driver, doc, "erro_fatal")

finally:
    log(doc, "✅ Teste concluído.")
    finalizar_relatorio()