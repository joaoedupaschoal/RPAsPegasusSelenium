from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.chrome.options import Options
from selenium.common.exceptions import TimeoutException, NoSuchElementException
from selenium.webdriver.chrome.service import Service
from webdriver_manager.chrome import ChromeDriverManager
from docx import Document
from datetime import datetime
import subprocess
import time
import os
import re

# ==========================================
#  CENÁRIO: Cesta Básica (cenário 1)
#  AJUSTES NESTA VERSÃO
#  - Corrigido ValueError em encontrar_mensagem_alerta (tuplas agora têm 3 valores: seletor, tipo, emoji)
#  - Unificado padrão do safe_action(doc, descricao, func) -> (ok, err)
#  - Removida dependência de utils.actions (tudo local e coeso)
#  - Mantido fluxo e screenshots "únicos"
# ==========================================

URL = "http://localhost:8080/gs/index.xhtml"
LOGIN_EMAIL = "joaoeduardo.gold@outlook.com"
LOGIN_PASSWORD = "071999gs"

doc = None
driver = None
wait = None

# ==== LOG & SCREENSHOT ====
screenshot_registradas = set()

def log(doc: Document, msg: str) -> None:
    print(msg)
    doc.add_paragraph(msg)


def sanitize_filename(nome: str) -> str:
    nome = nome.lower()
    nome = re.sub(r"[^a-z0-9_]", "_", nome)
    return nome


def take_screenshot(driver: webdriver.Chrome, doc: Document, nome: str) -> None:
    nome_limpo = sanitize_filename(nome)
    os.makedirs("screenshots", exist_ok=True)
    path = f"screenshots/{nome_limpo}.png"
    driver.save_screenshot(path)
    doc.add_paragraph(f"Screenshot: {nome}")
    try:
        from docx.shared import Inches
        doc.add_picture(path, width=Inches(5.5))
    except Exception:
        pass


def registrar_screenshot_unico(nome: str, driver: webdriver.Chrome, doc: Document, descricao: str | None = None) -> None:
    if nome not in screenshot_registradas:
        if descricao:
            log(doc, f"📸 {descricao}")
        take_screenshot(driver, doc, nome)
        screenshot_registradas.add(nome)

# ==== SAFE ACTION ====

def safe_action(doc: Document, descricao: str, func) -> tuple[bool, str | None]:
    global driver
    try:
        log(doc, f"🔄 {descricao}...")
        func()
        log(doc, f"✅ {descricao} realizada com sucesso.")
        take_screenshot(driver, doc, descricao.lower().replace(" ", "_"))
        return True, None
    except Exception as e:
        err = str(e)
        log(doc, f"❌ Erro ao {descricao.lower()}: {err}")
        take_screenshot(driver, doc, f"erro_{descricao.lower().replace(' ', '_')}")
        return False, err

# ==== AÇÕES DE TELA ====

def ajustar_zoom() -> None:
    global driver
    driver.execute_script("document.body.style.zoom='90%'")


def login() -> None:
    global wait
    wait.until(EC.presence_of_element_located((By.ID, "j_id15:email"))).send_keys(LOGIN_EMAIL)
    wait.until(EC.presence_of_element_located((By.ID, "j_id15:senha"))).send_keys(LOGIN_PASSWORD, Keys.ENTER)
    wait.until(EC.presence_of_element_located((By.TAG_NAME, "body")))
    time.sleep(2)


def abrir_menu() -> None:
    global driver, wait
    driver.find_element(By.TAG_NAME, "body").send_keys(Keys.F2)
    campo = wait.until(EC.presence_of_element_located((By.XPATH, "//input[@placeholder='Busque um cadastro']")))
    campo.click()
    campo.send_keys("Cesta Básica", Keys.ENTER)
    time.sleep(2)


def acessar_formulario() -> None:
    global wait
    wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR,
        "#fmod_200009 > div.wdTelas > div.telaInicial.clearfix.overflow.overflowY > ul > li:nth-child(1) > a > span"
    ))).click()
    time.sleep(2)


def preencher_campos() -> None:
    global wait, doc
    log(doc, "🔄 Preenchendo o campo 'TITULAR CB'.")
    campo_titular_cb = wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR,
        "#fmod_200009 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div > div > div > div.formCol > input"
    )))
    campo_titular_cb.clear()
    campo_titular_cb.send_keys("TESTE TITULAR CB SELENIUM AUTOMATIZADO")
    log(doc, "✅ Campo 'TITULAR CB' preenchido.")
    registrar_screenshot_unico("titular_cb_preenchido", driver, doc)

    adicionar_cb = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR,
        "#fmod_200009 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div > div > div > div.btnListHolder > a.btAddGroup"
    )))
    adicionar_cb.click()


def salvar() -> None:
    global wait
    wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR,
        "#fmod_200009 > div.wdTelas > div.telaCadastro.clearfix > div.btnHolder > a.btModel.btGray.btsave"
    ))).click()
    time.sleep(2)


def fechar_modal() -> None:
    global wait
    wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR,
        "#fmod_200009 > div.wdTop.ui-draggable-handle > div.wdClose > a"
    ))).click()
    time.sleep(1)

# ==== ALERTAS ====

def encontrar_mensagem_alerta():
    """Retorna (elemento, tipo) onde tipo ∈ {"sucesso", "alerta", "erro"} ou None."""
    seletores = [
        (".alerts.salvo",  "sucesso", "✅"),
        (".alerts.alerta", "alerta",  "⚠️"),
        (".alerts.erro",   "erro",    "❌"),
    ]
    for seletor, tipo, emoji in seletores:
        try:
            elemento = driver.find_element(By.CSS_SELECTOR, seletor)
            if elemento.is_displayed():
                try:
                    texto = elemento.find_element(By.TAG_NAME, "p").text.strip()
                except NoSuchElementException:
                    texto = elemento.text.strip()
                log(doc, f"{emoji} Mensagem de {tipo}: {texto}")
                take_screenshot(driver, doc, f"mensagem_{tipo}")
                return elemento, tipo
        except NoSuchElementException:
            continue
    log(doc, "ℹ️ Nenhuma mensagem de alerta encontrada.")
    take_screenshot(driver, doc, "mensagem_nao_encontrada")
    return None, None

# ==== MAIN ====

def main() -> None:
    global doc, driver, wait

    # Documento
    doc = Document()
    doc.add_heading("RELATÓRIO DO TESTE", 0)
    doc.add_paragraph("Cadastro de Cesta Básica.")
    doc.add_paragraph(f"🗕️ Data do teste: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")
    doc.add_paragraph("Neste teste, o robô preencherá os campos obrigatórios e salvará o cadastro de uma nova Cesta Básica.")

    # Browser
    chrome_options = Options()
    chrome_options.add_argument("--start-maximized")
    driver = webdriver.Chrome(service=Service(ChromeDriverManager().install()), options=chrome_options)
    wait = WebDriverWait(driver, 10)

    def finalizar_relatorio():
        doc_name = f"relatorio_cesta_basica_cenario_1_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        doc.save(doc_name)
        log(doc, f"📄 Relatório salvo como: {doc_name}")
        try:
            subprocess.run(["start", "winword", doc_name], shell=True)
        except Exception as e:
            log(doc, f"Erro ao abrir o Word: {e}")
        driver.quit()

    # Fluxo
    ok, _ = safe_action(doc, "Acessando o sistema", lambda: driver.get(URL))
    if not ok:
        finalizar_relatorio(); return
    registrar_screenshot_unico("url_acessada", driver, doc, "Sistema acessado.")

    ok, _ = safe_action(doc, "Realizando login", login)
    if not ok:
        finalizar_relatorio(); return
    registrar_screenshot_unico("login_concluido", driver, doc, "Login realizado.")

    ok, _ = safe_action(doc, "Ajustando zoom", ajustar_zoom)
    if not ok:
        finalizar_relatorio(); return

    ok, _ = safe_action(doc, "Abrindo menu Cesta Básica", abrir_menu)
    if not ok:
        finalizar_relatorio(); return
    registrar_screenshot_unico("menu_aberto", driver, doc, "Menu 'Cesta Básica' aberto.")

    ok, _ = safe_action(doc, "Acessando formulário", acessar_formulario)
    if not ok:
        finalizar_relatorio(); return
    registrar_screenshot_unico("formulario_aberto", driver, doc, "Formulário de cadastro aberto.")

    ok, _ = safe_action(doc, "Preenchendo campos", preencher_campos)
    if not ok:
        finalizar_relatorio(); return

    ok, _ = safe_action(doc, "Clicando em Salvar", salvar)
    if not ok:
        finalizar_relatorio(); return
    registrar_screenshot_unico("apos_salvar", driver, doc, "Clique no botão Salvar realizado.")

    _, tipo = encontrar_mensagem_alerta()
    if tipo == "sucesso":
        log(doc, "✅ Mensagem de sucesso exibida.")
    elif tipo == "alerta":
        log(doc, "⚠️ Mensagem de alerta exibida.")
    elif tipo == "erro":
        log(doc, "❌ Mensagem de erro exibida.")
    else:
        log(doc, "⚠️ Nenhuma mensagem exibida.")

    registrar_screenshot_unico("mensagem_final", driver, doc, "Mensagem exibida após salvar.")

    ok, _ = safe_action(doc, "Fechando formulário", fechar_modal)
    if not ok:
        finalizar_relatorio(); return
    registrar_screenshot_unico("formulario_fechado", driver, doc, "Formulário fechado.")

    log(doc, "✅ Teste concluído com sucesso.")
    finalizar_relatorio()


if __name__ == "__main__":
    main()
