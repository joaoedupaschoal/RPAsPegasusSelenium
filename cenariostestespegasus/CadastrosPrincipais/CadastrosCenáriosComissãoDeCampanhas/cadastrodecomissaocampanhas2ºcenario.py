from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service
from webdriver_manager.chrome import ChromeDriverManager
from docx import Document
from datetime import datetime, timedelta
import subprocess
import time
import os
import sys
from faker import Faker
from faker.providers import BaseProvider
import random
import string
from selenium.common.exceptions import TimeoutException, NoSuchElementException


sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))


from utils.actions import log, take_screenshot, safe_action, ajustar_zoom

URL = "http://localhost:8080/gs/index.xhtml"
LOGIN_EMAIL = "joaoeduardo.gold@outlook.com"
LOGIN_PASSWORD = "071999gs"

# Configuração do Faker
class BrasilProvider(BaseProvider):
    def rg(self):
        numeros = [str(random.randint(0, 9)) for _ in range(8)]
        return ''.join(numeros) + '-' + str(random.randint(0, 9))

fake = Faker("pt_BR")
fake.add_provider(BrasilProvider)

# Geração de dados aleatórios
def gerar_datas_validas(dias_adiante=30):
    hoje = datetime.today()
    data_inicial = hoje.strftime('%d/%m/%Y')
    data_final = (hoje + timedelta(days=dias_adiante)).strftime('%d/%m/%Y')
    return data_inicial, data_final

data_inicio, data_fim = gerar_datas_validas()

# Controle de screenshots únicas
screenshot_registradas = set()
def registrar_screenshot_unico(nome, driver, doc, descricao=None):
    if nome not in screenshot_registradas:
        if descricao:
            log(doc, f"📸 {descricao}")
        take_screenshot(driver, doc, nome)
        screenshot_registradas.add(nome)

def main():
    doc = Document()
    doc.add_heading("RELATÓRIO DO TESTE", 0)
    doc.add_paragraph("Cadastro de Comissão Campanhas Teste.")
    doc.add_paragraph(f"🗕️ Data do teste: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")
    doc.add_paragraph("Neste teste, o robô preencherá os campos obrigatórios e cancelará o cadastro de uma nova Comissão Campanha.")

    chrome_options = Options()
    chrome_options.add_argument("--start-maximized")
    driver = webdriver.Chrome(service=Service(ChromeDriverManager().install()), options=chrome_options)
    wait = WebDriverWait(driver, 10)

    def finalizar_relatorio():
        doc_name = f"relatorio_comissao_campanhas_cenario_2_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        doc.save(doc_name)
        log(doc, f"📄 Relatório salvo como: {doc_name}")
        try:
            subprocess.run(["start", "winword", doc_name], shell=True)
        except Exception as e:
            log(doc, f"Erro ao abrir o Word: {e}")
        driver.quit()

    def login():
        wait.until(EC.presence_of_element_located((By.ID, "j_id15:email"))).send_keys(LOGIN_EMAIL)
        wait.until(EC.presence_of_element_located((By.ID, "j_id15:senha"))).send_keys(LOGIN_PASSWORD, Keys.ENTER)
        wait.until(EC.presence_of_element_located((By.TAG_NAME, "body")))
        time.sleep(2)

    def abrir_menu():
        driver.find_element(By.TAG_NAME, "body").send_keys(Keys.F2)
        campo = wait.until(EC.presence_of_element_located((By.XPATH, "//input[@placeholder='Busque um cadastro']")))
        campo.click()
        campo.send_keys("Comissão Campanhas", Keys.ENTER)
        time.sleep(3)

    def acessar_formulario():
        cadastrar = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, 
            "#fmod_200031 > div.wdTelas > div.telaInicial.clearfix.overflow.overflowY > ul > li:nth-child(1) > a > span")))
        cadastrar.click()
        time.sleep(2)

    def preencher_nome_campanha():
        log(doc, "🔄 Preenchendo campo 'Nome da Campanha'.")
        campo_nome_campanha = wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, 
            "#fmod_200031 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div > div:nth-child(1) > div > div:nth-child(2) > input")))
        campo_nome_campanha.send_keys('TESTE CAMPANHA SELENIUM AUTOMATIZADO')
        log(doc, "✅ Campo 'Nome da Campanha' preenchido.")

    def preencher_data_inicial():
        log(doc, "🔄 Preenchendo campo 'Data Inicial'.")
        comissao_data_inicial = wait.until(EC.element_to_be_clickable((By.XPATH,
            "//input[@grupo='100062' and @ref='100212' and contains(@class, 'mandatory')]")))
        comissao_data_inicial.click()
        comissao_data_inicial.send_keys(data_inicio)
        comissao_data_inicial.send_keys(data_inicio)
        log(doc, "✅ Campo 'Data Inicial' preenchido.")

    def preencher_data_final():
        log(doc, "🔄 Preenchendo campo 'Data Final'.")
        comissao_data_final = wait.until(EC.element_to_be_clickable((By.XPATH,
            "//input[@grupo='100062' and @ref='100213' and contains(@class, 'mandatory')]")))
        comissao_data_final.click()
        comissao_data_final.send_keys(data_fim)
        comissao_data_final.send_keys(data_fim)
        log(doc, "✅ Campo 'Data Final' preenchido.")

    def preencher_comissao_meta_vendedor():
        log(doc, "🔄 Preenchendo campo 'Comissão Meta Vendedor'.")
        campo_comissao_meta_vendedor = wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, 
            "#fmod_200031 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div > div:nth-child(2) > div > div:nth-child(2) > input")))
        campo_comissao_meta_vendedor.send_keys(fake.random_int(min=10, max=500))
        log(doc, "✅ Campo 'Comissão Meta Vendedor' preenchido.")

    def preencher_comissao_porcentagem_vendedor():
        log(doc, "🔄 Preenchendo campo 'Comissão Porcentagem Vendedor'.")
        campo_comissao_porcentagem_vendedor = wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, 
            "#fmod_200031 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div > div:nth-child(2) > div > div:nth-child(3) > input")))
        campo_comissao_porcentagem_vendedor.send_keys(fake.random_int(min=10, max=500))
        log(doc, "✅ Campo 'Comissão Porcentagem Vendedor' preenchido.")

    def preencher_comissao_meta_supervisor():
        log(doc, "🔄 Preenchendo campo 'Comissão Meta Supervisor'.")
        campo_comissao_meta_supervisor = wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, 
            "#fmod_200031 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div > div:nth-child(3) > div > div:nth-child(2) > input")))
        campo_comissao_meta_supervisor.send_keys(fake.random_int(min=10, max=500))
        log(doc, "✅ Campo 'Comissão Meta Supervisor' preenchido.")

    def preencher_comissao_porcentagem_supervisor():
        log(doc, "🔄 Preenchendo campo 'Comissão Porcentagem Supervisor'.")
        campo_comissao_porcentagem_supervisor = wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, 
            "#fmod_200031 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div > div:nth-child(3) > div > div:nth-child(3) > input")))
        campo_comissao_porcentagem_supervisor.send_keys(fake.random_int(min=10, max=500))
        log(doc, "✅ Campo 'Comissão Porcentagem Supervisor' preenchido.")

    def cancelar():
        cancelar_btn = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR,
            "#fmod_200031 > div.wdTelas > div.telaCadastro.clearfix > div.btnHolder > a.btModel.btGray.btcancel")))
        cancelar_btn.click()
        time.sleep(2)

    def confirmar_cancelamento():
        confirmar_btn = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR,
            "#BtYes")))
        confirmar_btn.click()
        time.sleep(2)  


    def encontrar_mensagem_alerta():
        seletores = [
            (".alerts.salvo", "✅ Mensagem de Sucesso"),
            (".alerts.alerta", "⚠️ Mensagem de Alerta"),
            (".alerts.erro", "❌ Mensagem de Erro"),
        ]

        for seletor, tipo in seletores:
            try:
                elemento = driver.find_element(By.CSS_SELECTOR, seletor)
                if elemento.is_displayed():
                    log(doc, f"📢 {tipo}: {elemento.text}")
                    return elemento
            except:
                continue

        log(doc, "ℹ️ Nenhuma mensagem de alerta encontrada.")
        return None

    def fechar_modal():
        x_btn = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR,
            "#fmod_200031 > div.wdTop.ui-draggable-handle > div.wdClose > a")))
        x_btn.click()
        time.sleep(1)

    # EXECUÇÃO COM safe_action INDIVIDUAL PARA CADA AÇÃO
    if not safe_action(doc, "Acessando o sistema", lambda: driver.get(URL), driver, wait)[0]:
        finalizar_relatorio()
        return
    registrar_screenshot_unico("url_acessada", driver, doc, "Sistema acessado.")

    if not safe_action(doc, "Realizando login", login, driver, wait)[0]:
        finalizar_relatorio()
        return
    registrar_screenshot_unico("login_concluido", driver, doc, "Login realizado.")

    if not safe_action(doc, "Ajustando zoom", lambda: ajustar_zoom(driver), driver, wait)[0]:
        finalizar_relatorio()
        return

    if not safe_action(doc, "Abrindo menu Comissão Campanhas", abrir_menu, driver, wait)[0]:
        finalizar_relatorio()
        return
    registrar_screenshot_unico("menu_aberto", driver, doc, "Menu 'Comissão Campanhas' aberto.")

    if not safe_action(doc, "Acessando formulário de cadastro", acessar_formulario, driver, wait)[0]:
        finalizar_relatorio()
        return
    registrar_screenshot_unico("formulario_aberto", driver, doc, "Formulário de cadastro aberto.")

    # PREENCHIMENTO DOS CAMPOS - safe_action individual para cada campo
    if not safe_action(doc, "Preenchendo Nome da Campanha", preencher_nome_campanha, driver, wait)[0]:
        finalizar_relatorio()
        return

    if not safe_action(doc, "Preenchendo Data Inicial", preencher_data_inicial, driver, wait)[0]:
        finalizar_relatorio()
        return

    if not safe_action(doc, "Preenchendo Data Final", preencher_data_final, driver, wait)[0]:
        finalizar_relatorio()
        return

    if not safe_action(doc, "Preenchendo Comissão Meta Vendedor", preencher_comissao_meta_vendedor, driver, wait)[0]:
        finalizar_relatorio()
        return

    if not safe_action(doc, "Preenchendo Comissão Porcentagem Vendedor", preencher_comissao_porcentagem_vendedor, driver, wait)[0]:
        finalizar_relatorio()
        return

    if not safe_action(doc, "Preenchendo Comissão Meta Supervisor", preencher_comissao_meta_supervisor, driver, wait)[0]:
        finalizar_relatorio()
        return

    if not safe_action(doc, "Preenchendo Comissão Porcentagem Supervisor", preencher_comissao_porcentagem_supervisor, driver, wait)[0]:
        finalizar_relatorio()
        return

    registrar_screenshot_unico("campos_preenchidos", driver, doc, "Todos os campos preenchidos.")

    if not safe_action(doc, "Clicando no botão Cancelar", cancelar, driver, wait)[0]:
        finalizar_relatorio()
        return
    registrar_screenshot_unico("apos_cancelar", driver, doc, "Clique no botão Cancelar realizado.")

    if not safe_action(doc, "Clicando no botão Sim", confirmar_cancelamento, driver, wait)[0]:
        finalizar_relatorio()
        return
    registrar_screenshot_unico("apos_cancelar", driver, doc, "Clique no botão Sim realizado.")


    encontrar_mensagem_alerta()

    # FECHANDO O FORMULÁRIO
    if not safe_action(doc, "Fechando formulário", fechar_modal, driver, wait)[0]:
        finalizar_relatorio()
        return
    registrar_screenshot_unico("formulario_fechado", driver, doc, "Formulário fechado.")

    log(doc, "✅ Teste de cadastro de Comissão Campanhas concluído com sucesso.")
    finalizar_relatorio()

if __name__ == "__main__":
    main()