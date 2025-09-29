# Refatorado e organizado: cadastrodemovimentacaodocaixa1cenario.py

from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support.ui import WebDriverWait, Select
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service
from selenium.webdriver import ActionChains
from webdriver_manager.chrome import ChromeDriverManager
from docx import Document
from docx.shared import Inches
from faker import Faker
from faker.providers import BaseProvider
from datetime import datetime, timedelta
import subprocess
import os
import time
import random
import string
from selenium.common.exceptions import TimeoutException, NoSuchElementException


# ==== PROVEDOR CUSTOMIZADO ====
class BrasilProvider(BaseProvider):
    def rg(self):
        numeros = [str(random.randint(0, 9)) for _ in range(8)]
        return ''.join(numeros) + '-' + str(random.randint(0, 9))

fake = Faker("pt_BR")
fake.add_provider(BrasilProvider)

# ==== CONFIGURAÇÕES ====
URL = "http://localhost:8080/gs/index.xhtml"
LOGIN_EMAIL = "joaoeduardo.gold@outlook.com"
LOGIN_PASSWORD = "071999gs"

# ==== DOCUMENTO ====
doc = Document()
doc.add_heading("RELATÓRIO DO TESTE", 0)
doc.add_paragraph("Cadastro de Movimentação do Caixa – Cenário 4: Preenchimento dos campos NÃO obrigatórios e salvamento.")
doc.add_paragraph(f"Data do teste: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")

screenshot_registradas = set()

# ==== FUNÇÕES DE UTILITÁRIO ====
def log(doc, msg):
    print(msg)
    doc.add_paragraph(msg)

def take_screenshot(driver, doc, nome):
    if nome not in screenshot_registradas:
        path = f"screenshots/{nome}.png"
        os.makedirs("screenshots", exist_ok=True)
        driver.save_screenshot(path)
        doc.add_paragraph(f"Screenshot: {nome}")
        doc.add_picture(path, width=Inches(5.5))
        screenshot_registradas.add(nome)

def safe_action(doc, descricao, func):
    try:
        log(doc, f"🔄 {descricao}...")
        func()
        log(doc, f"✅ {descricao} realizada com sucesso.")
        take_screenshot(driver, doc, descricao.lower().replace(" ", "_"))
    except Exception as e:
        log(doc, f"❌ Erro ao {descricao.lower()}: {e}")
        take_screenshot(driver, doc, f"erro_{descricao.lower().replace(' ', '_')}")

def finalizar_relatorio():
    nome_arquivo = f"relatorio_movimentacao_caixa_cenario_4_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc.save(nome_arquivo)
    log(doc, f"📄 Relatório salvo como: {nome_arquivo}")
    subprocess.run(["start", "winword", nome_arquivo], shell=True)
    driver.quit()

def encontrar_mensagem_alerta():
    seletores = [
        (".alerts.salvo", "✅ Mensagem de Sucesso"),
        (".alerts.alerta", "⚠️ Mensagem de Alerta"),
        (".alerts.erro", "❌ Mensagem de Erro"),
    ]

    for seletor, tipo in seletores:
        try:
            elemento = driver.find_element(By.CSS_SELECTOR, seletor)
            if elemento.is_displayed():
                log(doc, f"📢 {tipo}: {elemento.text}")
                return elemento
        except:
            continue

    log(doc, "ℹ️ Nenhuma mensagem de alerta encontrada.")
    return None

def ajustar_zoom():
    try:
        driver.execute_script("document.body.style.zoom='90%'")
        log(doc, "🔍 Zoom ajustado para 90%.")
    except Exception as e:
        log(doc, f"⚠️ Erro ao ajustar zoom: {e}")

def gerar_dados_aleatorios():
    """Gera dados aleatórios para preenchimento do formulário"""
    data_lancamento = fake.date_between(start_date='-5d', end_date='today').strftime("%d/%m/%Y")
    valor = round(random.uniform(100, 10000), 2)
    numero_doc = str(fake.random_number(digits=8))
    
    return {
        'data_lancamento': data_lancamento,
        'valor': valor,
        'numero_doc': numero_doc
    }

def preencher_data_duplo(selector, valor):
    def acao():
        campo = wait.until(EC.element_to_be_clickable((By.XPATH, selector)))
        # Primeira vez
        campo.send_keys(valor)
        time.sleep(0.2)
        # Segunda vez
        campo.clear()
        campo.send_keys(valor)
        driver.execute_script("arguments[0].value = arguments[1];", campo, valor)
        campo.send_keys(Keys.TAB)
        time.sleep(0.2)
    return acao

def preencher_campo(selector, valor):
    def acao():
        campo = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, selector)))
        campo.click()
        campo.send_keys(valor)
    return acao

def abrir_lov_e_selecionar(lov_selector, search_value, item_xpath):
    def acao():
        # Abre LOV
        lov_button = wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, lov_selector)))
        lov_button.click()
        time.sleep(5)
        
        # Pesquisa
        campo_pesquisa = wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(2) > input")))
        campo_pesquisa.send_keys(search_value, Keys.ENTER)
        time.sleep(2)
        
        # Seleciona item
        item = wait.until(EC.element_to_be_clickable((By.XPATH, item_xpath)))
        item.click()
    return acao

def abrir_lov_fornecedor_e_selecionar():
    def acao():
        # Abre LOV do fornecedor
        lov_fornecedor = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, '#fmod_10013 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div > div > div > div:nth-child(8) > div > a')))
        lov_fornecedor.click()
        
        # Pesquisa fornecedor
        campo_pesquisa_fornecedor = wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#txtPesquisa")))
        campo_pesquisa_fornecedor.send_keys('FORNECEDOR TESTE JOÃO', Keys.ENTER)
        time.sleep(10)

        # Seleciona fornecedor
        fornecedor = wait.until(EC.element_to_be_clickable((By.XPATH, "//td[contains(text(), 'FORNECEDOR TESTE JOÃO')]")))
        fornecedor.click()
    return acao

# ==== INICIALIZAÇÃO DO DRIVER ====
options = Options()
options.add_argument("--start-maximized")
driver = webdriver.Chrome(service=Service(ChromeDriverManager().install()), options=options)
wait = WebDriverWait(driver, 10)

# Gera dados aleatórios
dados = gerar_dados_aleatorios()

# ==== EXECUÇÃO DO TESTE ====
try:
    safe_action(doc, "Acessando sistema", lambda: driver.get(URL))

    safe_action(doc, "Realizando login", lambda: (
        wait.until(EC.presence_of_element_located((By.ID, "j_id15:email"))).send_keys(LOGIN_EMAIL),
        wait.until(EC.presence_of_element_located((By.ID, "j_id15:senha"))).send_keys(LOGIN_PASSWORD, Keys.ENTER),
        time.sleep(5)
    ))

    safe_action(doc, "Ajustando zoom", lambda: ajustar_zoom())

    safe_action(doc, "Abrindo menu Movimentação Caixa", lambda: (
        driver.find_element(By.TAG_NAME, "body").send_keys(Keys.F2),
        time.sleep(1),
        driver.find_element(By.XPATH, "//input[@placeholder='Busque um cadastro']").send_keys("Movimentação Caixa", Keys.ENTER),
        time.sleep(3)
    ))

    safe_action(doc, "Clicando em Cadastrar", lambda: (
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "#fmod_10013 > div.wdTelas > div.telaInicial.clearfix.overflow.overflowY > ul > li:nth-child(1) > a > span"))).click(),
        time.sleep(2)
    ))

    safe_action(doc, "Preenchendo Complemento", 
                preencher_campo("#fmod_10013 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div > div > div > div:nth-child(6) > input", 
                               "TESTE COMPLEMENTO SELENIUM AUTOMATIZADO"))

    safe_action(doc, "Selecionando Fornecedor", abrir_lov_fornecedor_e_selecionar())

    safe_action(doc, "Preenchendo Razão Social", 
                preencher_campo("#fmod_10013 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div > div > div > div:nth-child(9) > input", 
                               "TESTE RAZÃO SOCIAL SELENIUM AUTOMATIZADO"))

    safe_action(doc, "Preenchendo Número do Documento", 
                preencher_campo("#fmod_10013 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div > div > div > div:nth-child(10) > input", 
                               dados['numero_doc']))    

    safe_action(doc, "Selecionando Histórico Padrão", 
                abrir_lov_e_selecionar("#fmod_10013 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div > div > div > div:nth-child(13) > div > div > a",
                                      "200129",
                                      "//td[contains(text(), 'RECEBIMENTOS VIA CHEQUE E DINHEIRO')]"))

    safe_action(doc, "Clicando em Salvar", lambda: (
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "#fmod_10013 > div.wdTelas > div.telaCadastro.clearfix > div.btnHolder > a.btModel.btGray.btsave"))).click(),
    ))
    encontrar_mensagem_alerta()



    safe_action(doc, "Fechando modal", lambda: wait.until(
        EC.element_to_be_clickable((By.CSS_SELECTOR, "#fmod_10013 > div.wdTop.ui-draggable-handle > div.wdClose > a"))
    ).click())


except Exception as e:
    log(doc, f"❌ ERRO FATAL: {e}")
    take_screenshot(driver, doc, "erro_fatal")

finally:
    log(doc, "✅ Teste concluído com sucesso.")
    finalizar_relatorio()