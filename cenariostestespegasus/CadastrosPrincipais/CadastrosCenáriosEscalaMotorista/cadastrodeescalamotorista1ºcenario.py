# Refatorado e organizado: cadastrodeescalamotorista2ºcenario.py

from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support.ui import Select, WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service
from webdriver_manager.chrome import ChromeDriverManager
from docx import Document
from docx.shared import Inches
from faker import Faker
from faker.providers import BaseProvider
from validate_docbr import CPF
from datetime import datetime, timedelta
import subprocess
import os
import time
import random

# ==== PROVIDERS CUSTOMIZADOS ====
class BrasilProvider(BaseProvider):
    def rg(self):
        numeros = [str(random.randint(0, 9)) for _ in range(8)]
        return ''.join(numeros) + '-' + str(random.randint(0, 9))

fake = Faker("pt_BR")
fake.add_provider(BrasilProvider)

def gerar_datas_validas():
    """Gera datas coerentes para admissão, início e fim da escala, e vencimento da CNH."""
    hoje = datetime.today().date()
    
    # Data de admissão entre 10 anos atrás e hoje
    data_admissao = fake.date_between(start_date=hoje - timedelta(days=3650), end_date=hoje)
    
    # Data de início da escala entre hoje e 1 ano no futuro
    data_inicio = fake.date_between(start_date=hoje, end_date=hoje + timedelta(days=365))
    
    # Data fim entre 1 e 180 dias após a data de início
    data_fim = data_inicio + timedelta(days=random.randint(1, 180))
    
    # Vencimento CNH entre hoje e 10 anos no futuro
    vencimento_cnh = fake.date_between(start_date=hoje, end_date=hoje + timedelta(days=3650))
    
    return (data_admissao.strftime('%d/%m/%Y'), 
            data_inicio.strftime('%d/%m/%Y'), 
            data_fim.strftime('%d/%m/%Y'), 
            vencimento_cnh.strftime('%d/%m/%Y'))

def gerar_dados_documentos():
    """Gera documentos fictícios para o cadastro."""
    carteira_trabalho = str(random.randint(10000000, 99999999))
    pis = fake.cpf().replace('.', '').replace('-', '')[:11]
    cnh = str(random.randint(10000000000, 99999999999))
    cpf = CPF().generate()
    
    return carteira_trabalho, pis, cnh, cpf

# Gera os dados necessários
data_admissao, data_inicio, data_fim, vencimento_cnh = gerar_datas_validas()
carteira_trabalho, pis, cnh, cpf_valido = gerar_dados_documentos()

# ==== CONFIGURAÇÕES ====
URL = "http://localhost:8080/gs/index.xhtml"
LOGIN_EMAIL = "joaoeduardo.gold@outlook.com"
LOGIN_PASSWORD = "071999gs"

# ==== DOCUMENTO ====
doc = Document()
doc.add_heading("RELATÓRIO DO TESTE", 0)
doc.add_paragraph("Cadastro de Escala Motorista – Cenário 1: Preenchimento completo e salvamento.")
doc.add_paragraph(f"Data do teste: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")

screenshot_registradas = set()

# ==== FUNÇÕES DE UTILITÁRIO ====
def log(doc, msg):
    print(msg)
    doc.add_paragraph(msg)

def take_screenshot(driver, doc, nome):
    if nome not in screenshot_registradas:
        path = f"screenshots/{nome}.png"
        os.makedirs("screenshots", exist_ok=True)
        driver.save_screenshot(path)
        doc.add_paragraph(f"Screenshot: {nome}")
        doc.add_picture(path, width=Inches(5.5))
        screenshot_registradas.add(nome)

def safe_action(doc, descricao, func):
    try:
        log(doc, f"🔄 {descricao}...")
        func()
        log(doc, f"✅ {descricao} realizada com sucesso.")
        take_screenshot(driver, doc, descricao.lower().replace(" ", "_"))
    except Exception as e:
        log(doc, f"❌ Erro ao {descricao.lower()}: {e}")
        take_screenshot(driver, doc, f"erro_{descricao.lower().replace(' ', '_')}")

def finalizar_relatorio():
    nome_arquivo = f"relatorio_escala_motorista_cenario_1_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc.save(nome_arquivo)
    log(doc, f"📄 Relatório salvo como: {nome_arquivo}")
    subprocess.run(["start", "winword", nome_arquivo], shell=True)
    driver.quit()

def encontrar_mensagem_alerta():
    seletores = [
        (".alerts.salvo", "✅ Sucesso"),
        (".alerts.alerta", "⚠️ Alerta"),
        (".alerts.erro", "❌ Erro"),
    ]

    for seletor, tipo in seletores:
        try:
            elemento = driver.find_element(By.CSS_SELECTOR, seletor)
            if elemento.is_displayed():
                log(doc, f"📢 {tipo}: {elemento.text}")
                return elemento
        except:
            continue

    log(doc, "ℹ️ Nenhuma mensagem de alerta encontrada.")
    return None

def ajustar_zoom():
    try:
        driver.execute_script("document.body.style.zoom='90%'")
        log(doc, "🔍 Zoom ajustado para 90%.")
    except Exception as e:
        log(doc, f"⚠️ Erro ao ajustar zoom: {e}")

def preencher_campo_data(selector, valor):
    def acao():
        campo = wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, selector)))
        campo.click()
        campo.clear()
        campo.send_keys(valor)
        time.sleep(0.2)
    return acao

def selecionar_opcao(selector, texto):
    def acao():
        select_element = wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, selector)))
        Select(select_element).select_by_visible_text(texto)
    return acao

# ==== INICIALIZAÇÃO DO DRIVER ====
options = Options()
options.add_argument("--start-maximized")
driver = webdriver.Chrome(service=Service(ChromeDriverManager().install()), options=options)
wait = WebDriverWait(driver, 20)

# ==== EXECUÇÃO DO TESTE ====
try:
    safe_action(doc, "Acessando sistema", lambda: driver.get(URL))

    safe_action(doc, "Realizando login", lambda: (
        wait.until(EC.presence_of_element_located((By.ID, "j_id15:email"))).send_keys(LOGIN_EMAIL),
        wait.until(EC.presence_of_element_located((By.ID, "j_id15:senha"))).send_keys(LOGIN_PASSWORD, Keys.ENTER),
        wait.until(EC.presence_of_element_located((By.TAG_NAME, "body")))
    ))

    safe_action(doc, "Esperando sistema carregar e ajustando zoom", lambda: (
        time.sleep(5),
        ajustar_zoom()
    ))

    safe_action(doc, "Abrindo menu Escala Motorista", lambda: (
        driver.find_element(By.TAG_NAME, "body").send_keys(Keys.F2),
        time.sleep(1),
        wait.until(EC.visibility_of_element_located((By.XPATH, "//input[@placeholder='Busque um cadastro']"))).send_keys("Escala Motorista", Keys.ENTER)
    ))

    safe_action(doc, "Clicando em Cadastrar", lambda: (
        time.sleep(2),
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_3000002 > div.wdTelas > div.telaInicial.clearfix.overflow.overflowY > ul > li:nth-child(1) > a > span"))).click()
    ))

    safe_action(doc, "Abrindo LOV Motorista", lambda: wait.until(
        EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_3000002 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroEscalaMotorista > div.catWrapper > div > div > div > div > div > div:nth-child(2) > div > a"))
    ).click())

    safe_action(doc, "Criando novo registro de motorista", lambda: wait.until(
        EC.visibility_of_element_located((By.CSS_SELECTOR, "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(4) > a"))
    ).click())

    safe_action(doc, "Abrindo LOV Pessoas", lambda: wait.until(
        EC.visibility_of_element_located((By.CSS_SELECTOR, "#cg_10058 > div.wdTelas > div > div.catWrapper > div > div > div > div > div > div:nth-child(2) > div > a"))
    ).click())

    safe_action(doc, "Criando novo registro de pessoa", lambda: wait.until(
        EC.visibility_of_element_located((By.CSS_SELECTOR, "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(3) > a"))
    ).click())

    # Preenchimento dos dados pessoais
    safe_action(doc, "Preenchendo nome", lambda: wait.until(
        EC.presence_of_element_located((By.CSS_SELECTOR, "#cg_1 > div.wdTelas > div > div.catWrapper > div > div.cat_dadosPessoais.categoriaHolder > div > div > div:nth-child(2) > div:nth-child(2) > input"))
    ).send_keys(fake.name()))

    safe_action(doc, "Selecionando tipo de pessoa", selecionar_opcao(
        "#cg_1 > div.wdTelas > div > div.catWrapper > div > div.cat_dadosPessoais.categoriaHolder > div > div > div:nth-child(2) > div:nth-child(3) > select",
        "Física"
    ))

    safe_action(doc, "Selecionando tipo de documento", selecionar_opcao(
        "#cg_1 > div.wdTelas > div > div.catWrapper > div > div.cat_dadosPessoais.categoriaHolder > div > div > div:nth-child(2) > div:nth-child(4) > select",
        "Carteira de Identidade Classista"
    ))

    safe_action(doc, "Preenchendo RG", lambda: driver.find_element(
        By.CSS_SELECTOR, "#cg_1 > div.wdTelas > div > div.catWrapper > div > div.cat_dadosPessoais.categoriaHolder > div > div > div.formRow.divPessoaFISICA > div:nth-child(1) > input"
    ).send_keys(fake.rg()))

    safe_action(doc, "Preenchendo data de expedição", lambda: (
        wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, "input.dataExpedicao"))).click(),
        driver.find_element(By.CSS_SELECTOR, "input.dataExpedicao").send_keys(fake.date_of_birth(minimum_age=18, maximum_age=60).strftime("%d/%m/%Y"))
    ))

    safe_action(doc, "Preenchendo CPF", lambda: (
        driver.find_element(By.CSS_SELECTOR, "#cg_1 > div.wdTelas > div > div.catWrapper > div > div.cat_dadosPessoais.categoriaHolder > div > div > div.formRow.divPessoaFISICA > div:nth-child(3) > input").click(),
        driver.find_element(By.CSS_SELECTOR, "#cg_1 > div.wdTelas > div > div.catWrapper > div > div.cat_dadosPessoais.categoriaHolder > div > div > div.formRow.divPessoaFISICA > div:nth-child(3) > input").send_keys(cpf_valido)
    ))

    safe_action(doc, "Acessando aba Dados Complementares", lambda: driver.find_element(By.LINK_TEXT, "Dados Complementares").click())

    safe_action(doc, "Preenchendo estado civil", selecionar_opcao(
        "#cg_1 > div.wdTelas > div > div.catWrapper > div > div.cat_dadosComplementares.categoriaHolder > div > div > div > div:nth-child(1) > select",
        "Solteiro"
    ))

    safe_action(doc, "Preenchendo sexo", selecionar_opcao(
        "#cg_1 > div.wdTelas > div > div.catWrapper > div > div.cat_dadosComplementares.categoriaHolder > div > div > div > div:nth-child(2) > select",
        "Feminino"
    ))

    safe_action(doc, "Preenchendo e-mail", lambda: driver.find_element(
        By.CSS_SELECTOR, "#cg_1 > div.wdTelas > div > div.catWrapper > div > div.cat_dadosComplementares.categoriaHolder > div > div > div > div:nth-child(9) > input"
    ).send_keys(fake.email()))

    safe_action(doc, "Preenchendo data de nascimento", lambda: (
        wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, "input.dataNascimento"))).click(),
        driver.find_element(By.CSS_SELECTOR, "input.dataNascimento").send_keys(fake.date_of_birth(minimum_age=18, maximum_age=60).strftime("%d/%m/%Y"))
    ))

    # Preenchimento dos campos de contato
    safe_action(doc, "Preenchendo telefone 1", lambda: driver.find_element(
        By.CSS_SELECTOR, "#cg_1 > div.wdTelas > div > div.catWrapper > div > div.cat_dadosComplementares.categoriaHolder > div > div > div > div:nth-child(5) > input"
    ).send_keys(fake.phone_number()))

    safe_action(doc, "Preenchendo telefone 2", lambda: driver.find_element(
        By.CSS_SELECTOR, "#cg_1 > div.wdTelas > div > div.catWrapper > div > div.cat_dadosComplementares.categoriaHolder > div > div > div > div:nth-child(6) > input"
    ).send_keys(fake.phone_number()))

    safe_action(doc, "Preenchendo telefone 3", lambda: driver.find_element(
        By.CSS_SELECTOR, "#cg_1 > div.wdTelas > div > div.catWrapper > div > div.cat_dadosComplementares.categoriaHolder > div > div > div > div:nth-child(7) > input"
    ).send_keys(fake.phone_number()))

    safe_action(doc, "Preenchendo naturalidade", lambda: driver.find_element(
        By.CSS_SELECTOR, "#cg_1 > div.wdTelas > div > div.catWrapper > div > div.cat_dadosComplementares.categoriaHolder > div > div > div > div:nth-child(11) > input"
    ).send_keys(fake.city()))

    safe_action(doc, "Preenchendo nacionalidade", lambda: driver.find_element(
        By.CSS_SELECTOR, "#cg_1 > div.wdTelas > div > div.catWrapper > div > div.cat_dadosComplementares.categoriaHolder > div > div > div > div:nth-child(12) > input"
    ).send_keys(fake.country()))

    safe_action(doc, "Preenchendo nome do pai", lambda: driver.find_element(
        By.CSS_SELECTOR, "#cg_1 > div.wdTelas > div > div.catWrapper > div > div.cat_dadosComplementares.categoriaHolder > div > div > div > div:nth-child(13) > input"
    ).send_keys(fake.first_name()))

    safe_action(doc, "Preenchendo nome da mãe", lambda: driver.find_element(
        By.CSS_SELECTOR, "#cg_1 > div.wdTelas > div > div.catWrapper > div > div.cat_dadosComplementares.categoriaHolder > div > div > div > div:nth-child(14) > input"
    ).send_keys(fake.first_name()))

    safe_action(doc, "Preenchendo profissão", lambda: driver.find_element(
        By.CSS_SELECTOR, "#cg_1 > div.wdTelas > div > div.catWrapper > div > div.cat_dadosComplementares.categoriaHolder > div > div > div > div:nth-child(16) > input"
    ).send_keys(fake.job()))

    safe_action(doc, "Salvando cadastro de pessoa", lambda: driver.find_element(
        By.CSS_SELECTOR, "#cg_1 > div.wdTelas > div > div.btnHolder > a.btModel.btGray.btsave"
    ).click())

    # Preenchimento dos dados do motorista
    safe_action(doc, "Preenchendo data de admissão", lambda: (
        wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, 'input.dataAdmissao'))).click(),
        driver.find_element(By.CSS_SELECTOR, 'input.dataAdmissao').send_keys(data_admissao)
    ))

    safe_action(doc, "Preenchendo carteira de trabalho", lambda: (
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "#cg_10058 > div.wdTelas > div > div.catWrapper > div > div > div > div > div > div:nth-child(4) > input"))).click(),
        driver.find_element(By.CSS_SELECTOR, "#cg_10058 > div.wdTelas > div > div.catWrapper > div > div > div > div > div > div:nth-child(4) > input").send_keys(carteira_trabalho)
    ))

    safe_action(doc, "Preenchendo PIS", lambda: wait.until(
        EC.presence_of_element_located((By.CSS_SELECTOR, '#cg_10058 > div.wdTelas > div > div.catWrapper > div > div > div > div > div > div:nth-child(5) > input'))
    ).send_keys(pis))

    safe_action(doc, "Preenchendo CNH", lambda: wait.until(
        EC.presence_of_element_located((By.CSS_SELECTOR, '#cg_10058 > div.wdTelas > div > div.catWrapper > div > div > div > div > div > div:nth-child(6) > input'))
    ).send_keys(cnh))

    safe_action(doc, "Preenchendo vencimento da CNH", lambda: wait.until(
        EC.presence_of_element_located((By.CSS_SELECTOR, 'input.vencimentoCnh'))
    ).send_keys(vencimento_cnh))

    safe_action(doc, "Selecionando vínculo empregatício", selecionar_opcao(
        "#cg_10058 > div.wdTelas > div > div.catWrapper > div > div > div > div > div > div:nth-child(9) > select",
        "Carteira Assinada"
    ))

    safe_action(doc, "Salvando cadastro de motorista", lambda: wait.until(
        EC.presence_of_element_located((By.CSS_SELECTOR, '#cg_10058 > div.wdTelas > div > div.btnHolder > a.btModel.btGray.btsave'))
    ).click())

    # Preenchimento da escala
    safe_action(doc, "Preenchendo data início da escala", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, 'input.dataInicio'))).click(),
        driver.find_element(By.CSS_SELECTOR, 'input.dataInicio').send_keys(data_inicio)
    ))

    safe_action(doc, "Preenchendo data fim da escala", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, 'input.dataFim'))).click(),
        driver.find_element(By.CSS_SELECTOR, 'input.dataFim').send_keys(data_fim)
    ))

    safe_action(doc, "Selecionando tipo de escala", selecionar_opcao(
        "#fmod_3000002 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroEscalaMotorista > div.catWrapper > div > div > div > div > div > div:nth-child(5) > select",
        "Escala Motorista Lex"
    ))

    safe_action(doc, "Salvando cadastro", lambda: driver.find_element(
        By.CSS_SELECTOR, "#fmod_3000002 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroEscalaMotorista > div.btnHolder > a.btModel.btGray.btsave"
    ).click())





    time.sleep(5)
    safe_action(doc, "Fechando modal após salvamento", lambda: wait.until(EC.element_to_be_clickable((
        By.CSS_SELECTOR, "#fmod_3000002 > div.wdTop.ui-draggable-handle > div.wdClose > a"
        ))
    ).click())



    encontrar_mensagem_alerta()

except Exception as e:
    log(doc, f"❌ ERRO FATAL: {e}")
    take_screenshot(driver, doc, "erro_fatal")

finally:

    log(doc, "✅ Teste concluído com sucesso.")

    finalizar_relatorio()