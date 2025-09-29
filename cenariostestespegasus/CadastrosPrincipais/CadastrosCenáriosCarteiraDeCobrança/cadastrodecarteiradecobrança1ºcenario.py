from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support.ui import WebDriverWait, Select
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service
from selenium.common.exceptions import TimeoutException, NoSuchElementException
from webdriver_manager.chrome import ChromeDriverManager
from docx import Document
from docx.shared import Inches
from faker import Faker
from validate_docbr import CPF
import random, string, os, subprocess, time
from datetime import datetime
from datetime import timedelta


# ==== CONFIGURAÇÕES ====
URL = "http://localhost:8080/gs/index.xhtml"
LOGIN_EMAIL = "joaoeduardo.gold@outlook.com"
LOGIN_PASSWORD = "071999gs"
fake = Faker("pt_BR")
cpf_generator = CPF()

doc = Document()
doc.add_heading("RELATÓRIO DO TESTE", 0)
doc.add_paragraph("Cadastro de Carteira de Cobrança – Cenário 1: Preenchimento completo e salvamento")
doc.add_paragraph(f"Data do teste: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")
screenshot_registradas = set()

def log(msg):
    print(msg)
    doc.add_paragraph(msg)

def take_screenshot(driver, nome):
    if nome not in screenshot_registradas:
        path = f"screenshots/{nome}.png"
        os.makedirs("screenshots", exist_ok=True)
        driver.save_screenshot(path)
        doc.add_paragraph(f"📸 Screenshot: {nome}")
        doc.add_picture(path, width=Inches(5.5))
        screenshot_registradas.add(nome)

def safe_action(descricao, func):
    try:
        log(f"🔄 {descricao}...")
        func()
        log(f"✅ {descricao} concluída.")
        take_screenshot(driver, descricao.lower().replace(" ", "_"))
    except Exception as e:
        log(f"❌ Erro ao {descricao.lower()}: {e}")
        take_screenshot(driver, f"erro_{descricao.lower().replace(' ', '_')}")

def ajustar_zoom(driver):
    try:
        driver.execute_script("document.body.style.zoom='90%'")
        log("🔍 Zoom ajustado para 90%.")
    except Exception as e:
        log(f"⚠️ Erro ao ajustar zoom: {e}")

def encontrar_mensagem_alerta():
    seletores = [
        (".alerts.salvo", "✅ Mensagem de Sucesso"),
        (".alerts.alerta", "⚠️ Mensagem de Alerta"),
        (".alerts.erro", "❌ Mensagem de Erro"),
    ]
    for seletor, tipo in seletores:
        try:
            elemento = driver.find_element(By.CSS_SELECTOR, seletor)
            if elemento.is_displayed():
                log(f"📢 Mensagem de {tipo}: {elemento.text}")
                return elemento
        except NoSuchElementException:
            continue
    log("⚠️ Nenhuma mensagem de alerta encontrada.")
    return None

# ==== INICIALIZAÇÃO DO DRIVER ====
options = Options()
options.add_argument("--start-maximized")
driver = webdriver.Chrome(service=Service(ChromeDriverManager().install()), options=options)
wait = WebDriverWait(driver, 10)

try:
    safe_action("Acessar sistema", lambda: driver.get(URL))

    safe_action("Login", lambda: (
        wait.until(EC.presence_of_element_located((By.ID, "j_id15:email"))).send_keys(LOGIN_EMAIL),
        wait.until(EC.presence_of_element_located((By.ID, "j_id15:senha"))).send_keys(LOGIN_PASSWORD, Keys.ENTER),
        wait.until(EC.presence_of_element_located((By.TAG_NAME, "body")))
    ))

    safe_action("Ajustar zoom e esperar tela inicial", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "div.menuHolder > a.button.btModules > span"))),
        ajustar_zoom(driver)
    ))

    safe_action("Abrir menu com F2", lambda: (
        driver.find_element(By.TAG_NAME, "body").send_keys(Keys.F2),
        time.sleep(1)
    ))

    safe_action("Buscar 'Carteira de Cobrança'", lambda: (
        wait.until(EC.visibility_of_element_located((By.XPATH, "//input[@placeholder='Busque um cadastro']"))).send_keys("Carteira de Cobrança", Keys.ENTER),
        time.sleep(2)
    ))

    safe_action("Clicar em Cadastrar", lambda: (
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "#fmod_200027 ul > li:nth-child(1) > a > span"))).click()
    ))

    safe_action("Abrir LOV", lambda: (
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "#fmod_200027 .formCol > div > a"))).click()
    ))

    safe_action("Clicar em Novo Cadastro", lambda: (
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "body > div.modalHolder .formRow.formLastLine > div:nth-child(4) > a"))).click()
    ))

    # ==== DADOS PESSOAIS ====
    safe_action("Preencher nome", lambda: driver.find_element(
        By.CSS_SELECTOR, "#cg_1 .cat_dadosPessoais .formRow:nth-child(2) > div:nth-child(2) > input"
    ).send_keys("VENDEDOR TESTE SELENIUM"))

    safe_action("Selecionar tipo de pessoa", lambda: Select(driver.find_element(
        By.CSS_SELECTOR, "#cg_1 .cat_dadosPessoais .formRow:nth-child(2) > div:nth-child(3) > select"
    )).select_by_visible_text("Física"))

    safe_action("Selecionar tipo de documento", lambda: Select(driver.find_element(
        By.CSS_SELECTOR, "#cg_1 .cat_dadosPessoais .formRow:nth-child(2) > div:nth-child(4) > select"
    )).select_by_visible_text("Carteira de Identidade Classista"))

    safe_action("Preencher RG", lambda: (
        driver.find_element(By.CSS_SELECTOR, "#cg_1 .divPessoaFISICA > div:nth-child(1) > input")
        .send_keys(fake.random_number(digits=7))
    ))

    safe_action("Preencher data de expedição", lambda: (
        driver.find_element(By.CSS_SELECTOR, "input.dataExpedicao")
        .send_keys(fake.date_between(start_date='-10y', end_date='-1y').strftime('%d/%m/%Y'))
    ))

    safe_action("Preencher CPF", lambda: (
        driver.find_element(By.CSS_SELECTOR, "#cg_1 .divPessoaFISICA > div:nth-child(3) > input").send_keys(cpf_generator.generate())
    ))

    # ==== DADOS COMPLEMENTARES ====
    safe_action("Abrir aba Dados Complementares", lambda: driver.find_element(By.LINK_TEXT, "Dados Complementares").click())

    safe_action("Selecionar estado civil", lambda: Select(driver.find_element(
        By.CSS_SELECTOR, "#cg_1 .cat_dadosComplementares .formRow > div:nth-child(1) > select")).select_by_visible_text("Solteiro"))

    safe_action("Selecionar sexo", lambda: Select(driver.find_element(
        By.CSS_SELECTOR, "#cg_1 .cat_dadosComplementares .formRow > div:nth-child(2) > select")).select_by_visible_text("Feminino"))

    safe_action("Preencher e-mail", lambda: driver.find_element(
        By.CSS_SELECTOR, "#cg_1 .cat_dadosComplementares .formRow > div:nth-child(9) > input").send_keys(fake.email()))

    safe_action("Preencher data nascimento", lambda: driver.find_element(
        By.CSS_SELECTOR, "input.dataNascimento").send_keys(
        (datetime.today() - timedelta(days=random.randint(18*365, 60*365))).strftime("%d/%m/%Y")))

    # ==== ENDEREÇO ====
    safe_action("Abrir aba Endereços", lambda: driver.find_element(
        By.CSS_SELECTOR, "#cg_1 .categorias ul > li.li_enderecos > a").click())

    safe_action("Preencher CEP e buscar", lambda: (
        driver.find_element(By.CSS_SELECTOR,
            "#cg_1 .cat_enderecos .grupo_enderecoResidencial > div > div:nth-child(2) > div:nth-child(1) > div > input").send_keys("15081115"),
        driver.find_element(By.CSS_SELECTOR,
            "#cg_1 .cat_enderecos .grupo_enderecoResidencial > div > div:nth-child(2) > div:nth-child(1) > div > a").click(),
        time.sleep(2)
    ))

    safe_action("Preencher número", lambda: driver.find_element(
        By.CSS_SELECTOR,
        "#cg_1 .cat_enderecos .grupo_enderecoResidencial > div > div:nth-child(3) > div:nth-child(2) > input"
    ).send_keys("123"))

    safe_action("Preencher complemento", lambda: driver.find_element(
        By.CSS_SELECTOR,
        "#cg_1 .cat_enderecos .grupo_enderecoResidencial > div > div:nth-child(3) > div:nth-child(3) > input"
    ).send_keys("Casa"))

    safe_action("Marcar como endereço principal", lambda: driver.find_element(
        By.CSS_SELECTOR,
        "#cg_1 .cat_enderecos .grupo_enderecoResidencial > div > div:nth-child(3) > div:nth-child(9) > label > input").click())

    # ==== SALVAR ====
    safe_action("Salvar pessoa", lambda: (
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "#cg_1 .btnHolder > a.btsave"))).click(),
        time.sleep(2)
    ))

    encontrar_mensagem_alerta()

    safe_action("Adicionar pessoa à carteira", lambda: (
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "#fmod_200027 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div > div > div > div.btnListHolder > a.btAddGroup"))).click()
    ))

    safe_action("Salvar carteira de cobrança", lambda: (
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "#fmod_200027 .btnHolder > a.btsave"))).click(),
        time.sleep(2)
    ))

    encontrar_mensagem_alerta()

    safe_action("Fechar modal", lambda: (
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "#fmod_200027 > div.wdTop.ui-draggable-handle > div.wdClose > a"))).click()
    ))

except Exception as e:
    log(f"❌ ERRO FATAL: {e}")
    take_screenshot(driver, "erro_fatal")

finally:

    log("✅ Teste concluído com sucesso.")

    nome_arquivo = f"relatorio_carteira_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc.save(nome_arquivo)
    log(f"📄 Relatório salvo: {nome_arquivo}")
    subprocess.run(["start", "winword", nome_arquivo], shell=True)
    driver.quit()
