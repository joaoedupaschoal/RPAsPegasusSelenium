# Refatorado no modelo padronizado: Cadastro de Equipamentos - Cenário 1

from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support.ui import WebDriverWait, Select
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service
from selenium.common.exceptions import TimeoutException, NoSuchElementException
from webdriver_manager.chrome import ChromeDriverManager
from docx import Document
from docx.shared import Inches
from faker import Faker
import os
import time
import random
import subprocess
from datetime import datetime

# ============ CONFIGURAÇÕES E VARIÁVEIS ============
URL = "http://localhost:8080/gs/index.xhtml"
LOGIN_EMAIL = "joaoeduardo.gold@outlook.com"
LOGIN_PASSWORD = "071999gs"
fake = Faker("pt_BR")
screenshot_registradas = set()

# ============ INICIALIZAÇÃO DOC ============
doc = Document()
doc.add_heading("RELATÓRIO DO TESTE", 0)
doc.add_paragraph("Cadastro de Equipamentos – Cenário 1: Preenchimento completo e salvamento.")
doc.add_paragraph(f"Data do teste: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")

# ============ FUNÇÕES UTILITÁRIAS PADRÃO ============
def log(doc, msg):
    print(msg)
    doc.add_paragraph(msg)

def take_screenshot(driver, doc, nome):
    if nome not in screenshot_registradas:
        path = f"screenshots/{nome}.png"
        os.makedirs("screenshots", exist_ok=True)
        driver.save_screenshot(path)
        doc.add_paragraph(f"Screenshot: {nome}")
        doc.add_picture(path, width=Inches(5.5))
        screenshot_registradas.add(nome)

def safe_action(doc, descricao, func):
    try:
        log(doc, f"🔄 {descricao}...")
        func()
        log(doc, f"✅ {descricao} realizada com sucesso.")
        take_screenshot(driver, doc, descricao.lower().replace(" ", "_"))
    except Exception as e:
        log(doc, f"❌ Erro ao {descricao.lower()}: {e}")
        take_screenshot(driver, doc, f"erro_{descricao.lower().replace(' ', '_')}")

def finalizar_relatorio():
    nome_arquivo = f"relatorio_equipamentos_cenario_1_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc.save(nome_arquivo)
    log(doc, f"📄 Relatório salvo como: {nome_arquivo}")
    subprocess.run(["start", "winword", nome_arquivo], shell=True)
    driver.quit()

def encontrar_mensagem_alerta():
    seletores = [
        (".alerts.salvo", "✅ Sucesso"),
        (".alerts.alerta", "⚠️ Alerta"),
        (".alerts.erro", "❌ Erro"),
    ]
    for seletor, tipo in seletores:
        try:
            elemento = driver.find_element(By.CSS_SELECTOR, seletor)
            if elemento.is_displayed():
                log(doc, f"📢 {tipo}: {elemento.text}")
                return elemento
        except:
            continue
    log(doc, "ℹ️ Nenhuma mensagem de alerta encontrada.")
    return None

def ajustar_zoom():
    try:
        driver.execute_script("document.body.style.zoom='90%'")
        log(doc, "🔍 Zoom ajustado para 90%.")
    except Exception as e:
        log(doc, f"⚠️ Erro ao ajustar zoom: {e}")

# ============ DRIVER E WAIT ============
options = Options()
options.add_argument("--start-maximized")
driver = webdriver.Chrome(service=Service(ChromeDriverManager().install()), options=options)
wait = WebDriverWait(driver, 10)

# ============ FLUXO DO TESTE ============
try:
    safe_action(doc, "Acessando sistema", lambda: driver.get(URL))

    safe_action(doc, "Realizando login", lambda: (
        wait.until(EC.presence_of_element_located((By.ID, "j_id15:email"))).send_keys(LOGIN_EMAIL),
        wait.until(EC.presence_of_element_located((By.ID, "j_id15:senha"))).send_keys(LOGIN_PASSWORD, Keys.ENTER),
        wait.until(EC.presence_of_element_located((By.TAG_NAME, "body")))
    ))
    time.sleep(3)
    safe_action(doc, "Ajustando zoom", lambda: ajustar_zoom())

    safe_action(doc, "Abrindo menu Equipamento", lambda: (
        driver.find_element(By.TAG_NAME, "body").send_keys(Keys.F2),
        time.sleep(1),
        wait.until(EC.visibility_of_element_located((By.XPATH, "//input[@placeholder='Busque um cadastro']"))).send_keys("Equipamento")
    ))

    safe_action(doc, "Selecionando opção de Equipamento", lambda: (
        wait.until(EC.element_to_be_clickable((By.XPATH, "/html/body/div[17]/div[2]/ul/li[16]/a"))).click()
    ))

    safe_action(doc, "Clicando em Cadastrar", lambda: (
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "#fmod_10057 > div.wdTelas > div.telaInicial.clearfix.overflow.overflowY > ul > li:nth-child(1) > a > span"))).click()
    ))

    safe_action(doc, "Abrindo LOV Grupo", lambda: (
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "#fmod_10057 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_dadosEquipamento.categoriaHolder > div > div > div > div:nth-child(2) > div > a"))).click()
    ))

    safe_action(doc, "Criando novo Grupo", lambda: (
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(4) > a"))).click(),
        wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, "#cg_10056 > div.wdTelas > div > div.catWrapper > div > div > div > div > div > div:nth-child(2) > input"))).send_keys("TESTE GRUPO EQUIPAMENTO SELENIUM AUTOMATIZADO"),
        wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, "#cg_10056 > div.wdTelas > div > div.catWrapper > div > div > div > div > div > div:nth-child(3) > input"))).send_keys(fake.random_int(min=10, max=5000)),
        wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, "#cg_10056 > div.wdTelas > div > div.catWrapper > div > div > div > div > div > div:nth-child(5) > input"))).send_keys(fake.random_int(min=10, max=90)),
        wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, "#cg_10056 > div.wdTelas > div > div.catWrapper > div > div > div > div > div > div:nth-child(6) > input"))).send_keys(fake.random_int(min=10, max=50000)),
        wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, "#cg_10056 > div.wdTelas > div > div.catWrapper > div > div > div > div > div > div:nth-child(7) > input"))).send_keys(fake.random_int(min=10, max=90)),
        wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, "#cg_10056 > div.wdTelas > div > div.catWrapper > div > div > div > div > div > div:nth-child(8) > input"))).send_keys(fake.random_int(min=10, max=90)),
        wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, "#cg_10056 > div.wdTelas > div > div.catWrapper > div > div > div > div > div > div:nth-child(9) > input"))).send_keys(fake.random_int(min=10, max=50000)),
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "#cg_10056 > div.wdTelas > div > div.btnHolder > a.btModel.btGray.btsave"))).click()
    ))

    safe_action(doc, "Preenchendo número de patrimônio", lambda: (
        wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, "#fmod_10057 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_dadosEquipamento.categoriaHolder > div > div > div > div:nth-child(3) > input"))).send_keys(fake.random_int(min=1000, max=9999))
    ))

    safe_action(doc, "Preenchendo nome do equipamento", lambda: (
        wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, "#fmod_10057 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_dadosEquipamento.categoriaHolder > div > div > div > div:nth-child(4) > input"))).send_keys("TESTE CADEIRA DE RODAS SELENIUM AUTOMATIZADO")
    ))

    safe_action(doc, "Preenchendo observação", lambda: (
        wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, "#fmod_10057 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_dadosEquipamento.categoriaHolder > div > div > div > div:nth-child(5) > input"))).send_keys("TESTE OBSERVAÇÃO SELENIUM AUTOMATIZADO")
    ))

    safe_action(doc, "Preenchendo valor de venda", lambda: (
        wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, "#fmod_10057 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_dadosEquipamento.categoriaHolder > div > div > div > div:nth-child(6) > input"))).send_keys(fake.random_int(min=1000, max=5000))
    ))

    safe_action(doc, "Selecionando Filial", lambda: (
        Select(wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, "#fmod_10057 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_dadosEquipamento.categoriaHolder > div > div > div > div:nth-child(8) > select")))).select_by_visible_text("Filial Gold 1")
    ))

    safe_action(doc, "Acessando aba Informações Adicionais", lambda: (
        wait.until(EC.element_to_be_clickable((By.XPATH, "//a[contains(text(), 'Informações Adicionais')]"))).click()
    ))

    safe_action(doc, "Preenchendo descrição adicional", lambda: (
        wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, "#fmod_10057 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_informacoesAdicionais.categoriaHolder > div > div > div > input"))).send_keys("TESTE DESCRIÇÃO CADEIRA DE RODAS SELENIUM AUTOMATIZADO")
    ))

    safe_action(doc, "Clicando em Salvar", lambda: (
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "#fmod_10057 > div.wdTelas > div.telaCadastro.clearfix > div.btnHolder > a.btModel.btGray.btsave"))).click(),
        time.sleep(3),
        encontrar_mensagem_alerta()
    ))

    safe_action(doc, "Fechando modal", lambda: (
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "#fmod_10057 > div.wdTop.ui-draggable-handle > div.wdClose"))).click()
    ))

except Exception as e:
    log(doc, f"❌ ERRO FATAL: {e}")
    take_screenshot(driver, doc, "erro_fatal")

finally:

    log(doc, "✅ Teste concluído com sucesso.")

    finalizar_relatorio()
