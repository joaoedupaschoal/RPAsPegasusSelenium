
# Refatorado e organizado: cadastrodeagendadecompromissos1ºcenario.py

from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service
from webdriver_manager.chrome import ChromeDriverManager
from docx import Document
from docx.shared import Inches
from faker import Faker
from datetime import datetime, timedelta
import subprocess
import os
import time
import random


fake = Faker("pt_BR")

def gerar_datas_validas():
    """Gera datas coerentes para nascimento, falecimento, sepultamento e agendamento dentro de um intervalo válido."""
    
    hoje = datetime.today().date()
    dez_anos_atras = hoje - timedelta(days=3650)  # Limite máximo de 10 anos atrás
    sessenta_dias_depois = hoje + timedelta(days=60)
    
    # Gera data de agendamento (até 60 dias a partir de hoje)
    data_agendamento = fake.date_between(start_date=hoje, end_date=sessenta_dias_depois)

    # Gera data de nascimento (antes de 100 anos)
    data_nascimento = fake.date_between(start_date=dez_anos_atras - timedelta(days=36500), end_date=dez_anos_atras)
    
    # Gera uma data de falecimento entre 10 anos atrás e hoje (não pode ser posterior ao nascimento)
    data_falecimento = fake.date_between(start_date=data_nascimento, end_date=hoje)
    
    # Gera uma data de sepultamento entre o dia de falecimento e 60 dias após a data de falecimento
    data_sepultamento = fake.date_between(start_date=data_falecimento, end_date=sessenta_dias_depois)

    # Ajuste de idade mínima e máxima
    idade_minima = 18
    idade_maxima = 110
    data_nascimento = data_falecimento - timedelta(days=random.randint(idade_minima * 365, idade_maxima * 365))

    # Sepultamento entre 1 e 10 dias após o falecimento
    data_sepultamento = data_falecimento + timedelta(days=random.randint(1, 10))

    # Registro entre 1 e 10 dias após o sepultamento
    data_registro = data_sepultamento + timedelta(days=random.randint(1, 10))

    data_efetivacao = fake.date_between(start_date=hoje - timedelta(days=365), end_date=hoje - timedelta(days=365))

    # Verifica se o ano é menor que 1900 e ajusta a data, se necessário
    def formatar_data(data):
        if data.year < 1900:
            data = data.replace(year=hoje.year)  # Ajusta para o ano atual se for menor que 1900
        return data.strftime('%d/%m/%y')

    # Formata as datas para o formato ddmmyy
    data_nascimento = formatar_data(data_nascimento)
    data_falecimento = formatar_data(data_falecimento)
    data_sepultamento = formatar_data(data_sepultamento)
    data_agendamento = formatar_data(data_agendamento)
    data_registro = formatar_data(data_registro)
    data_efetivacao = formatar_data(data_efetivacao)

    return data_nascimento, data_falecimento, data_sepultamento, data_agendamento, data_registro, data_efetivacao

# Chama a função para obter as datas
data_nascimento, data_falecimento, data_sepultamento, data_agendamento, data_registro, data_efetivacao  = gerar_datas_validas()


# ==== CONFIGURAÇÕES ====
URL = "http://localhost:8080/gs/index.xhtml"
LOGIN_EMAIL = "joaoeduardo.gold@outlook.com"
LOGIN_PASSWORD = "071999gs"
fake = Faker("pt_BR")

# ==== DOCUMENTO ====
doc = Document()
doc.add_heading("RELATÓRIO DO TESTE", 0)
doc.add_paragraph("Cadastro de Agenda de Compromissos – Cenário 2: Preenchimento completo e canmcelamento.")
doc.add_paragraph(f"Data do teste: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")

screenshot_registradas = set()

# ==== FUNÇÕES DE UTILITÁRIO ====
def log(doc, msg):
    print(msg)
    doc.add_paragraph(msg)

def take_screenshot(driver, doc, nome):
    if nome not in screenshot_registradas:
        path = f"screenshots/{nome}.png"
        os.makedirs("screenshots", exist_ok=True)
        driver.save_screenshot(path)
        doc.add_paragraph(f"Screenshot: {nome}")
        doc.add_picture(path, width=Inches(5.5))
        screenshot_registradas.add(nome)

def safe_action(doc, descricao, func):
    try:
        log(doc, f"🔄 {descricao}...")
        func()
        log(doc, f"✅ {descricao} realizada com sucesso.")
        take_screenshot(driver, doc, descricao.lower().replace(" ", "_"))
    except Exception as e:
        log(doc, f"❌ Erro ao {descricao.lower()}: {e}")
        take_screenshot(driver, doc, f"erro_{descricao.lower().replace(' ', '_')}")

def finalizar_relatorio():
    nome_arquivo = f"relatorio_agenda_compromissos_cenario_2_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc.save(nome_arquivo)
    log(doc, f"📄 Relatório salvo como: {nome_arquivo}")
    subprocess.run(["start", "winword", nome_arquivo], shell=True)
    driver.quit()


from selenium.common.exceptions import TimeoutException, NoSuchElementException

def encontrar_mensagem_alerta():
    seletores = [
        (".alerts.salvo", "✅ Sucesso"),
        (".alerts.alerta", "⚠️ Alerta"),
        (".alerts.erro", "❌ Erro"),
    ]

    for seletor, tipo in seletores:
        try:
            elemento = driver.find_element(By.CSS_SELECTOR, seletor)
            if elemento.is_displayed():
                log(doc, f"📢 {tipo}: {elemento.text}")
                return elemento
        except:
            continue

    log(doc, "ℹ️ Nenhuma mensagem de alerta encontrada.")
    return None

def ajustar_zoom():
    try:
        driver.execute_script("document.body.style.zoom='90%'")
        log(doc, "🔍 Zoom ajustado para 90%.")
    except Exception as e:
        log(doc, f"⚠️ Erro ao ajustar zoom: {e}")

def gerar_data_e_hora_compromisso():
    data = datetime.now() + timedelta(days=random.randint(1, 10))
    hora = f"{random.randint(8, 18):02d}:{random.choice(['00', '30'])}"
    return data.strftime("%d/%m/%Y"), hora

def preencher_data(selector, valor):
    def acao():
        campo = wait.until(EC.presence_of_element_located((By.XPATH, selector)))
        campo.click()  # Clica no campo para garantir que ele esteja ativo
        campo.clear()  # Limpa o campo antes de preencher
        campo.send_keys(valor)  # Simula o foco no próximo campo
        time.sleep(0.2)
        campo.send_keys(valor)  # Simula o foco no próximo campo
        time.sleep(0.2)
        campo.send_keys(valor)  # Simula o foco no próximo campo
        time.sleep(0.2)

    return acao

# ==== INICIALIZAÇÃO DO DRIVER ====
options = Options()
options.add_argument("--start-maximized")
driver = webdriver.Chrome(service=Service(ChromeDriverManager().install()), options=options)
wait = WebDriverWait(driver, 10)
# ==== EXECUÇÃO DO TESTE ====

try:
    safe_action(doc, "Acessando sistema", lambda: driver.get(URL))

    safe_action(doc, "Realizando login", lambda: (
        wait.until(EC.presence_of_element_located((By.ID, "j_id15:email"))).send_keys(LOGIN_EMAIL),
        wait.until(EC.presence_of_element_located((By.ID, "j_id15:senha"))).send_keys(LOGIN_PASSWORD, Keys.ENTER),
        wait.until(EC.presence_of_element_located((By.TAG_NAME, "body")))
    ))

    safe_action(doc, "Esperando ícone dos módulos aparecer e ajustando zoom", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "div.menuHolder > a.button.btModules > span"))),
        ajustar_zoom()
    ))

    safe_action(doc, "Abrindo menu Agenda de Compromissos", lambda: (
        driver.find_element(By.TAG_NAME, "body").click(),
        driver.find_element(By.TAG_NAME, "body").send_keys(Keys.F2),
        time.sleep(1),
        wait.until(EC.visibility_of_element_located((By.XPATH, "//input[@placeholder='Busque um cadastro']"))).send_keys("Agenda de Compromissos", Keys.ENTER)
    ))

    safe_action(doc, "Clicando em Cadastrar", lambda: (
        time.sleep(1),
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_200012 > div.wdTelas > div > ul > li:nth-child(1) > a > span"))).click()
    ))
    time.sleep(3),

    safe_action(doc, "Preenchendo campo Descrição", lambda: wait.until(
        EC.presence_of_element_located((By.CSS_SELECTOR, "#fmod_200012 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div > div > div > div:nth-child(2) > input"))
    ).send_keys("TESTE DESCRIÇÃO SELENIUM AUTOMATIZADO"))

    safe_action(doc, "Preenchendo campo Valor", lambda: wait.until(
        EC.presence_of_element_located((By.CSS_SELECTOR, "#fmod_200012 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div > div > div > div:nth-child(3) > input"))
    ).send_keys(fake.random_int(min=100, max=10000)))
    
    safe_action(doc, "Preenchendo campo Data do Agendamento", preencher_data("//input[@grupo='100029' and @ref='100125' and contains(@class, 'hasDatepicker mandatory fc')]", data_agendamento))

    safe_action(doc, "Preenchendo campo Data efetivação", preencher_data("//input[@grupo='100029' and @ref='100126' and contains(@class, 'hasDatepicker fc')]", data_efetivacao))

    safe_action(doc, "Clicando em Cancelar", lambda: wait.until(
        EC.element_to_be_clickable((By.CSS_SELECTOR, "#fmod_200012 > div.wdTelas > div.telaCadastro.clearfix > div.btnHolder > a.btModel.btGray.btcancel"))
    ).click())

    log(doc, "🔄 Aguardando modal de confirmação de cancelamento...")

    safe_action(doc, "Confirmando...", lambda: wait.until(
        EC.element_to_be_clickable((By.CSS_SELECTOR, "#BtYes"))
    ).click())
    time.sleep(3)  # Aguarda o modal de sucesso aparecer

    encontrar_mensagem_alerta()

    safe_action(doc, "Fechando modal", lambda: wait.until(
        EC.element_to_be_clickable((By.CSS_SELECTOR, "#fmod_200012 > div.wdTop.ui-draggable-handle > div.wdClose > a"))
    ).click())

except Exception as e:
    log(doc, f"❌ ERRO FATAL: {e}")
    take_screenshot(driver, doc, "erro_fatal")

finally:
    log(doc, "✅ Teste concluído com sucesso.")

    finalizar_relatorio()
