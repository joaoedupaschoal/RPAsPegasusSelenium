# Refatorado e organizado: cadastrodepetcores1¬∫cenario.py

from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support.ui import Select, WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service
from webdriver_manager.chrome import ChromeDriverManager
from docx import Document
from docx.shared import Inches
from faker import Faker
from faker.providers import BaseProvider
from validate_docbr import CPF
from datetime import datetime, timedelta
import subprocess
import os
import time
import random

# ==== PROVIDERS CUSTOMIZADOS ====
class BrasilProvider(BaseProvider):
    def rg(self):
        numeros = [str(random.randint(0, 9)) for _ in range(8)]
        return ''.join(numeros) + '-' + str(random.randint(0, 9))

fake = Faker("pt_BR")
fake.add_provider(BrasilProvider)

# Vari√°veis espec√≠ficas do contexto original
numero_aleatorio = random.randint(1, 100)
qtd_parcelas_em_atraso = int(fake.random.choice(['1', '2', '3', '4', '5']))
dias_para_exumar = int(fake.random.choice(['365', '730', '1095', '1460', '1825']))
valor_taxa_adesao = round(random.uniform(2000, 10000), 2)
cemetery_name = f"Cemit√©rio {fake.last_name()} {fake.random.choice(['Eterno', 'da Paz', 'Memorial', 'Descanso'])}"

def gerar_jazigos():
    quantidade_ruas = random.randint(1, 10)
    max_jazigos_por_rua = random.randint(1, 20)
    quantidade_total_jazigos = quantidade_ruas * max_jazigos_por_rua
    return quantidade_ruas, max_jazigos_por_rua, quantidade_total_jazigos

ruas, jazigos_por_rua, total_jazigos = gerar_jazigos()
altura_cm = random.randint(100, 200)
largura_cm = random.randint(100, 200)
comprimento_cm = random.randint(100, 200)

# ==== CONFIGURA√á√ïES ====
URL = "http://localhost:8080/gs/index.xhtml"
LOGIN_EMAIL = "joaoeduardo.gold@outlook.com"
LOGIN_PASSWORD = "071999gs"

# ==== DOCUMENTO ====
doc = Document()
doc.add_heading("RELAT√ìRIO DO TESTE", 0)
doc.add_paragraph("Cadastro de PET - Ra√ßas ‚Äì Cen√°rio 1: Preenchimento completo e salvamento.")
doc.add_paragraph(f"Data do teste: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")

screenshot_registradas = set()

# ==== FUN√á√ïES DE UTILIT√ÅRIO ====
def log(doc, msg):
    print(msg)
    doc.add_paragraph(msg)

def take_screenshot(driver, doc, nome):
    if nome not in screenshot_registradas:
        path = f"screenshots/{nome}.png"
        os.makedirs("screenshots", exist_ok=True)
        driver.save_screenshot(path)
        doc.add_paragraph(f"Screenshot: {nome}")
        doc.add_picture(path, width=Inches(5.5))
        screenshot_registradas.add(nome)

def safe_action(doc, descricao, func):
    try:
        log(doc, f"üîÑ {descricao}...")
        func()
        log(doc, f"‚úÖ {descricao} realizada com sucesso.")
        take_screenshot(driver, doc, descricao.lower().replace(" ", "_"))
    except Exception as e:
        log(doc, f"‚ùå Erro ao {descricao.lower()}: {e}")
        take_screenshot(driver, doc, f"erro_{descricao.lower().replace(' ', '_')}")

def finalizar_relatorio():
    nome_arquivo = f"relatorio_pet_racas_cenario_1_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc.save(nome_arquivo)
    log(doc, f"üìÑ Relat√≥rio salvo como: {nome_arquivo}")
    subprocess.run(["start", "winword", nome_arquivo], shell=True)
    driver.quit()

def encontrar_mensagem_alerta():
    seletores = [
        (".alerts.salvo", "‚úÖ Mensagem de Sucesso"),
        (".alerts.alerta", "‚ö†Ô∏è Mensagem de Alerta"),
        (".alerts.erro", "‚ùå Mensagem de Erro"),
    ]

    for seletor, tipo in seletores:
        try:
            elemento = driver.find_element(By.CSS_SELECTOR, seletor)
            if elemento.is_displayed():
                log(doc, f"üì¢ {tipo}: {elemento.text}")
                return elemento
        except:
            continue

    log(doc, "‚ÑπÔ∏è Nenhuma mensagem de alerta encontrada.")
    return None

def ajustar_zoom():
    try:
        driver.execute_script("document.body.style.zoom='90%'")
        log(doc, "üîç Zoom ajustado para 90%.")
    except Exception as e:
        log(doc, f"‚ö†Ô∏è Erro ao ajustar zoom: {e}")

# ==== INICIALIZA√á√ÉO DO DRIVER ====
options = Options()
options.add_argument("--start-maximized")
driver = webdriver.Chrome(service=Service(ChromeDriverManager().install()), options=options)
wait = WebDriverWait(driver, 20)

# ==== EXECU√á√ÉO DO TESTE ====
try:
    safe_action(doc, "Acessando sistema", lambda: driver.get(URL))

    safe_action(doc, "Realizando login", lambda: (
        wait.until(EC.presence_of_element_located((By.ID, "j_id15:email"))).send_keys(LOGIN_EMAIL),
        wait.until(EC.presence_of_element_located((By.ID, "j_id15:senha"))).send_keys(LOGIN_PASSWORD, Keys.ENTER),
        wait.until(EC.presence_of_element_located((By.TAG_NAME, "body")))
    ))

    safe_action(doc, "Esperando sistema carregar e ajustando zoom", lambda: (
        time.sleep(5),
        ajustar_zoom()
    ))

    safe_action(doc, "Abrindo menu PET - Ra√ßas", lambda: (
        driver.find_element(By.TAG_NAME, "body").send_keys(Keys.F2),
        time.sleep(1),
        wait.until(EC.visibility_of_element_located((By.XPATH, "//input[@placeholder='Busque um cadastro']"))).send_keys("PET - Ra√ßas", Keys.ENTER),
    ))

    safe_action(doc, "Clicando em Cadastrar", lambda: (
        time.sleep(3),
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_200005 > div.wdTelas > div > ul > li:nth-child(1) > a > span"))).click()
    ))

    safe_action(doc, "Preenchendo o Nome da Ra√ßa do PET", lambda: (
        time.sleep(2),
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_200005 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div > div > div > div > input")))
        .send_keys('TESTE RA√áA DO PET SELENIUM AUTOMATIZADO' + str(fake.random_int(min=1, max=1000)))
    ))

    safe_action(doc, "Salvando cadastro", lambda: driver.find_element(
        By.CSS_SELECTOR, "#fmod_200005 > div.wdTelas > div.telaCadastro.clearfix > div.btnHolder > a.btModel.btGray.btsave"
    ).click())
    encontrar_mensagem_alerta()

    safe_action(doc, "Fechando modal ap√≥s salvamento", lambda: wait.until(EC.element_to_be_clickable((
        By.CSS_SELECTOR, "#fmod_200005 > div.wdTop.ui-draggable-handle > div > a"
        ))
    ).click())

    encontrar_mensagem_alerta()

except Exception as e:
    log(doc, f"‚ùå ERRO FATAL: {e}")
    take_screenshot(driver, doc, "erro_fatal")

finally:
    log(doc, "‚úÖ Teste conclu√≠do com sucesso.")
    finalizar_relatorio()