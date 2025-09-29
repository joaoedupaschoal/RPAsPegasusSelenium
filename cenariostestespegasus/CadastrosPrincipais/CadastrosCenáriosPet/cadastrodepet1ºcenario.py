# Refatorado e organizado: cadastrodepet1¬∫cenario.py

from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support.ui import Select, WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service
from webdriver_manager.chrome import ChromeDriverManager
from docx import Document
from docx.shared import Inches
from faker import Faker
from faker.providers import BaseProvider
from validate_docbr import CPF
from datetime import datetime, timedelta
import subprocess
import os
import time
import random
import string

# ==== PROVIDERS CUSTOMIZADOS ====
class BrasilProvider(BaseProvider):
    def rg(self):
        numeros = [str(random.randint(0, 9)) for _ in range(8)]
        return ''.join(numeros) + '-' + str(random.randint(0, 9))

fake = Faker("pt_BR")
fake.add_provider(BrasilProvider)

numero_aleatorio = random.randint(1, 100)
letra_aleatoria = random.choice(string.ascii_uppercase)

def gerar_dados_pet():
    """Gera dados fict√≠cios para o cadastro do pet."""
    nome_pet = f'TESTE NOME DO PET SELENIUM AUTOMATIZADO {numero_aleatorio}{letra_aleatoria}'
    idade = fake.random_int(min=1, max=20)
    rg_pet = fake.random_number(digits=8)
    peso = fake.random_int(min=1, max=100)
    
    return nome_pet, idade, rg_pet, peso

# Gera os dados necess√°rios
nome_pet, idade_pet, rg_pet, peso_pet = gerar_dados_pet()

# ==== CONFIGURA√á√ïES ====
URL = "http://localhost:8080/gs/index.xhtml"
LOGIN_EMAIL = "joaoeduardo.gold@outlook.com"
LOGIN_PASSWORD = "071999gs"

# ==== DOCUMENTO ====
doc = Document()
doc.add_heading("RELAT√ìRIO DO TESTE", 0)
doc.add_paragraph("Cadastro de Pet ‚Äì Cen√°rio 1: Preenchimento completo e salvamento.")
doc.add_paragraph(f"Data do teste: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")

screenshot_registradas = set()

# ==== FUN√á√ïES DE UTILIT√ÅRIO ====
def log(doc, msg):
    print(msg)
    doc.add_paragraph(msg)

def take_screenshot(driver, doc, nome):
    if nome not in screenshot_registradas:
        path = f"screenshots/{nome}.png"
        os.makedirs("screenshots", exist_ok=True)
        driver.save_screenshot(path)
        doc.add_paragraph(f"Screenshot: {nome}")
        doc.add_picture(path, width=Inches(5.5))
        screenshot_registradas.add(nome)

def safe_action(doc, descricao, func):
    try:
        log(doc, f"üîÑ {descricao}...")
        func()
        log(doc, f"‚úÖ {descricao} realizada com sucesso.")
        take_screenshot(driver, doc, descricao.lower().replace(" ", "_"))
    except Exception as e:
        log(doc, f"‚ùå Erro ao {descricao.lower()}: {e}")
        take_screenshot(driver, doc, f"erro_{descricao.lower().replace(' ', '_')}")

def finalizar_relatorio():
    nome_arquivo = f"relatorio_pet_cenario_1_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc.save(nome_arquivo)
    log(doc, f"üìÑ Relat√≥rio salvo como: {nome_arquivo}")
    subprocess.run(["start", "winword", nome_arquivo], shell=True)
    driver.quit()

def abrir_modal_e_selecionar(btn_selector, pesquisa_selector, termo_pesquisa, btn_pesquisar_selector, resultado_xpath):
    def acao():
        # Abre o modal
        open_lov = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, btn_selector)))
        open_lov.click()

        # Aguarda campo pesquisa
        campo_pesquisa = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, pesquisa_selector)))
        campo_pesquisa.clear()
        campo_pesquisa.send_keys(termo_pesquisa)
        
        # Clica pesquisar ou pressiona ENTER
        if btn_pesquisar_selector:
            pesquisar = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, btn_pesquisar_selector)))
            pesquisar.click()
        else:
            campo_pesquisa.send_keys(Keys.ENTER)
        
        time.sleep(1)
        
        # Espera o resultado carregar
        wait.until(EC.presence_of_element_located((By.XPATH, resultado_xpath)))
        wait.until(EC.visibility_of_element_located((By.XPATH, resultado_xpath)))
        wait.until(EC.element_to_be_clickable((By.XPATH, resultado_xpath)))

        # Relocaliza no √∫ltimo instante (evita stale element)
        resultado = driver.find_element(By.XPATH, resultado_xpath)
        driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", resultado)
        time.sleep(0.2)
        resultado.click()

    return acao

def encontrar_mensagem_alerta():
    seletores = [
        (".alerts.salvo", "‚úÖ Mensagem de Sucesso"),
        (".alerts.alerta", "‚ö†Ô∏è Mensagem de Alerta"),
        (".alerts.erro", "‚ùå Mensagem de Erro"),
    ]

    for seletor, tipo in seletores:
        try:
            elemento = driver.find_element(By.CSS_SELECTOR, seletor)
            if elemento.is_displayed():
                log(doc, f"üì¢ {tipo}: {elemento.text}")
                return elemento
        except:
            continue

    log(doc, "‚ÑπÔ∏è Nenhuma mensagem de alerta encontrada.")
    return None

def ajustar_zoom():
    try:
        driver.execute_script("document.body.style.zoom='90%'")
        log(doc, "üîç Zoom ajustado para 90%.")
    except Exception as e:
        log(doc, f"‚ö†Ô∏è Erro ao ajustar zoom: {e}")

# ==== INICIALIZA√á√ÉO DO DRIVER ====
options = Options()
options.add_argument("--start-maximized")
driver = webdriver.Chrome(service=Service(ChromeDriverManager().install()), options=options)
wait = WebDriverWait(driver, 20)

# ==== EXECU√á√ÉO DO TESTE ====
try:
    safe_action(doc, "Acessando sistema", lambda: driver.get(URL))

    safe_action(doc, "Realizando login", lambda: (
        wait.until(EC.presence_of_element_located((By.ID, "j_id15:email"))).send_keys(LOGIN_EMAIL),
        wait.until(EC.presence_of_element_located((By.ID, "j_id15:senha"))).send_keys(LOGIN_PASSWORD, Keys.ENTER),
        wait.until(EC.presence_of_element_located((By.TAG_NAME, "body")))
    ))

    safe_action(doc, "Esperando sistema carregar e ajustando zoom", lambda: (
        time.sleep(5),
        ajustar_zoom()
    ))

    safe_action(doc, "Abrindo menu Pet", lambda: (
        driver.find_element(By.TAG_NAME, "body").send_keys(Keys.F2),
        time.sleep(1),
        wait.until(EC.visibility_of_element_located((By.XPATH, "//input[@placeholder='Busque um cadastro']"))).send_keys("Pet"),
        time.sleep(1),
        wait.until(EC.element_to_be_clickable((By.XPATH, "/html/body/div[17]/div[2]/ul/li[35]/a"))).click()
    ))

    safe_action(doc, "Clicando em Cadastrar", lambda: (
        time.sleep(3),
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_10085 > div.wdTelas > div > ul > li:nth-child(1) > a > span"))).click()
    ))

    safe_action(doc, "Preenchendo o Nome do Pet", lambda: (
        time.sleep(2),
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_10085 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPet > div.catWrapper > div > div > div > div > div > div:nth-child(1) > div > input")))
        .send_keys(nome_pet)
    ))

    safe_action(doc, "Preenchendo a Idade do Pet", lambda: (
        time.sleep(2),
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_10085 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPet > div.catWrapper > div > div > div > div > div > div:nth-child(2) > input[type=text]")))
        .send_keys(str(idade_pet))
    ))

    safe_action(doc, "Preenchendo o RG do Pet", lambda: (
        time.sleep(2),
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_10085 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPet > div.catWrapper > div > div > div > div > div > div:nth-child(4) > div > input")))
        .send_keys(str(rg_pet))
    ))

    safe_action(doc, "Selecionando a Ra√ßa TESTE RA√áA SELENIUM AUTOMATIZADO", abrir_modal_e_selecionar(
        "#fmod_10085 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPet > div.catWrapper > div > div > div > div > div > div:nth-child(5) > div > a",
        "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(2) > input",
        "TESTE RA√áA SELENIUM AUTOMATIZADO",
        None,  # Usa ENTER ao inv√©s de bot√£o
        "//td[contains(text(), 'TESTE RA√áA SELENIUM AUTOMATIZADO')]"
    ))

    safe_action(doc, "Selecionando a Cor TESTE COR SELENIUM AUTOMATIZADO", abrir_modal_e_selecionar(
        "#fmod_10085 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPet > div.catWrapper > div > div > div > div > div > div:nth-child(6) > div > a",
        "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div.formCol.divPesquisa > div:nth-child(1) > input",
        "TESTE COR SELENIUM AUTOMATIZADO",
        None,  # Usa ENTER ao inv√©s de bot√£o
        "//td[contains(text(), 'TESTE COR SELENIUM AUTOMATIZADO')]"
    ))

    safe_action(doc, "Preenchendo o Peso do Pet", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_10085 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPet > div.catWrapper > div > div > div > div > div > div:nth-child(7) > input[type=text]")))
        .send_keys(str(peso_pet))
    ))

    safe_action(doc, "Salvando cadastro", lambda: driver.find_element(
        By.CSS_SELECTOR, "#fmod_10085 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPet > div.btnHolder > a.btModel.btGray.btsave"
    ).click())

    safe_action(doc, "Fechando modal ap√≥s salvamento", lambda: wait.until(EC.element_to_be_clickable((
        By.CSS_SELECTOR, "#fmod_10085 > div.wdTop.ui-draggable-handle > div.wdClose > a"
        ))
    ).click())

    encontrar_mensagem_alerta()

except Exception as e:
    log(doc, f"‚ùå ERRO FATAL: {e}")
    take_screenshot(driver, doc, "erro_fatal")

finally:
    log(doc, "‚úÖ Teste conclu√≠do com sucesso.")
    finalizar_relatorio()