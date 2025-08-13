from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service
from webdriver_manager.chrome import ChromeDriverManager
from docx import Document
from datetime import datetime, timedelta
import subprocess
import time
import os
import sys
from faker import Faker
from faker.providers import BaseProvider
import random
import string
from selenium.common.exceptions import TimeoutException, NoSuchElementException

sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), '..', '..')))

from utils.actions import log, take_screenshot, safe_action, encontrar_mensagem_alerta, ajustar_zoom

URL = "http://localhost:8080/gs/index.xhtml"
LOGIN_EMAIL = "joaoeduardo.gold@outlook.com"
LOGIN_PASSWORD = "071999gs"

# Configura√ß√£o do Faker
class BrasilProvider(BaseProvider):
    def rg(self):
        numeros = [str(random.randint(0, 9)) for _ in range(8)]
        return ''.join(numeros) + '-' + str(random.randint(0, 9))

fake = Faker("pt_BR")
fake.add_provider(BrasilProvider)

# Gera√ß√£o de dados aleat√≥rios
def gerar_datas_validas(dias_adiante=30):
    hoje = datetime.today()
    data_inicial = hoje.strftime('%d/%m/%Y')
    data_final = (hoje + timedelta(days=dias_adiante)).strftime('%d/%m/%Y')
    return data_inicial, data_final

data_inicio, data_fim = gerar_datas_validas()

# Controle de screenshots √∫nicas
screenshot_registradas = set()
def registrar_screenshot_unico(nome, driver, doc, descricao=None):
    if nome not in screenshot_registradas:
        if descricao:
            log(doc, f"üì∏ {descricao}")
        take_screenshot(driver, doc, nome)
        screenshot_registradas.add(nome)

def main():
    doc = Document()
    doc.add_heading("RELAT√ìRIO DO TESTE", 0)
    doc.add_paragraph("Cadastro de Comiss√£o Teste.")
    doc.add_paragraph(f"üóïÔ∏è Data do teste: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")
    doc.add_paragraph("Neste teste, o rob√¥ preencher√° os campos N√ÉO obrigat√≥rios e salvar√° o cadastro de uma nova Comiss√£o.")

    chrome_options = Options()
    chrome_options.add_argument("--start-maximized")
    driver = webdriver.Chrome(service=Service(ChromeDriverManager().install()), options=chrome_options)
    wait = WebDriverWait(driver, 10)

    def finalizar_relatorio():
        doc_name = f"relatorio_comissao_cenario_3_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        doc.save(doc_name)
        log(doc, f"üìÑ Relat√≥rio salvo como: {doc_name}")
        try:
            subprocess.run(["start", "winword", doc_name], shell=True)
        except Exception as e:
            log(doc, f"Erro ao abrir o Word: {e}")
        driver.quit()

    def login():
        wait.until(EC.presence_of_element_located((By.ID, "j_id15:email"))).send_keys(LOGIN_EMAIL)
        wait.until(EC.presence_of_element_located((By.ID, "j_id15:senha"))).send_keys(LOGIN_PASSWORD, Keys.ENTER)
        wait.until(EC.presence_of_element_located((By.TAG_NAME, "body")))
        time.sleep(2)

    def abrir_menu():
        driver.find_element(By.TAG_NAME, "body").send_keys(Keys.F2)
        campo = wait.until(EC.presence_of_element_located((By.XPATH, "//input[@placeholder='Busque um cadastro']")))
        campo.click()
        campo.send_keys("Comiss√£o")
        
        elemento = wait.until(EC.element_to_be_clickable((By.XPATH, "/html/body/div[17]/div[2]/ul/li[8]/a")))
        elemento.click()
        time.sleep(3)

    def acessar_formulario():
        cadastrar = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, 
            "#fmod_200030 > div.wdTelas > div.telaInicial.clearfix.overflow.overflowY > ul > li:nth-child(1) > a > span")))
        cadastrar.click()
        time.sleep(2)


    def salvar():
        salvar_btn = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR,
            "#fmod_200030 > div.wdTelas > div.telaCadastro.clearfix > div.btnHolder > a.btModel.btGray.btsave")))
        salvar_btn.click()
        time.sleep(2)

    def fechar_modal():
        x_btn = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR,
            "#fmod_200030 > div.wdTop.ui-draggable-handle > div.wdClose > a")))
        x_btn.click()
        time.sleep(1)

    # EXECU√á√ÉO COM safe_action INDIVIDUAL PARA CADA A√á√ÉO
    if not safe_action(doc, "Acessando o sistema", lambda: driver.get(URL), driver, wait)[0]:
        finalizar_relatorio()
        return
    registrar_screenshot_unico("url_acessada", driver, doc, "Sistema acessado.")

    if not safe_action(doc, "Realizando login", login, driver, wait)[0]:
        finalizar_relatorio()
        return
    registrar_screenshot_unico("login_concluido", driver, doc, "Login realizado.")

    if not safe_action(doc, "Ajustando zoom", lambda: ajustar_zoom(driver), driver, wait)[0]:
        finalizar_relatorio()
        return

    if not safe_action(doc, "Abrindo menu Comiss√£o", abrir_menu, driver, wait)[0]:
        finalizar_relatorio()
        return
    registrar_screenshot_unico("menu_aberto", driver, doc, "Menu 'Comiss√£o' aberto.")

    if not safe_action(doc, "Acessando formul√°rio de cadastro", acessar_formulario, driver, wait)[0]:
        finalizar_relatorio()
        return
    registrar_screenshot_unico("formulario_aberto", driver, doc, "Formul√°rio de cadastro aberto.")


    # SALVANDO O CADASTRO
    if not safe_action(doc, "Clicando no bot√£o Salvar", salvar, driver, wait)[0]:
        finalizar_relatorio()
        return
    registrar_screenshot_unico("apos_salvar", driver, doc, "Clique no bot√£o Salvar realizado.")

    # VERIFICANDO MENSAGEM DE RETORNO
    _, tipo_alerta = encontrar_mensagem_alerta(driver, doc)
    if tipo_alerta == "sucesso":
        log(doc, "‚úÖ Mensagem de sucesso exibida.")
    elif tipo_alerta == "alerta":
        log(doc, "‚ö†Ô∏è Mensagem de alerta exibida.")
    elif tipo_alerta == "erro":
        log(doc, "‚ùå Mensagem de erro exibida.")
    else:
        log(doc, "‚ö†Ô∏è Nenhuma mensagem encontrada ap√≥s salvar.")
    registrar_screenshot_unico("mensagem_final", driver, doc, "Mensagem exibida ap√≥s salvar.")

    # FECHANDO O FORMUL√ÅRIO
    if not safe_action(doc, "Fechando formul√°rio", fechar_modal, driver, wait)[0]:
        finalizar_relatorio()
        return
    registrar_screenshot_unico("formulario_fechado", driver, doc, "Formul√°rio fechado.")

    log(doc, "‚úÖ Teste de cadastro de Comiss√£o conclu√≠do com sucesso.")
    finalizar_relatorio()

if __name__ == "__main__":
    main()