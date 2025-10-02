from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support.ui import WebDriverWait, Select
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service
from webdriver_manager.chrome import ChromeDriverManager
from docx import Document
from datetime import datetime
import subprocess
import time
import os
import sys
from faker import Faker
from faker.providers import BaseProvider
import random
import string
import pyautogui
from selenium.common.exceptions import TimeoutException, NoSuchElementException


sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))


from utils.actions import log, take_screenshot, safe_action, ajustar_zoom

URL = "http://localhost:8080/gs/index.xhtml"
LOGIN_EMAIL = "joaoeduardo.gold@outlook.com"
LOGIN_PASSWORD = "071999gs"

# Configuração do Faker
class BrasilProvider(BaseProvider):
    def rg(self):
        numeros = [str(random.randint(0, 9)) for _ in range(8)]
        return ''.join(numeros) + '-' + str(random.randint(0, 9))

fake = Faker("pt_BR")
fake.add_provider(BrasilProvider)

# Geração de dados aleatórios
numero_aleatorio = random.randint(1, 100)
letra_aleatoria = random.choice(string.ascii_uppercase)
cemetery_name = f"Cemitério {fake.last_name()} {random.choice(['Eterno', 'da Paz', 'Memorial', 'Descanso'])}"
qtd_parcelas_em_atraso = int(random.choice(['1', '2', '3', '4', '5']))
dias_para_exumar = int(random.choice(['365', '730', '1095', '1460', '1825']))
valor_taxa_adesao = round(random.uniform(2000, 10000), 2)

def gerar_jazigos():
    ruas = random.randint(1, 10)
    jazigos_por_rua = random.randint(1, 20)
    total = ruas * jazigos_por_rua
    return ruas, jazigos_por_rua, total

ruas, jazigos_por_rua, total_jazigos = gerar_jazigos()
altura_cm = random.randint(100, 200)
largura_cm = random.randint(100, 200)
comprimento_cm = random.randint(100, 200)

# Controle de screenshots únicas
screenshot_registradas = set()
def registrar_screenshot_unico(nome, driver, doc, descricao=None):
    if nome not in screenshot_registradas:
        if descricao:
            log(doc, f"📸 {descricao}")
        take_screenshot(driver, doc, nome)
        screenshot_registradas.add(nome)

def main():
    doc = Document()
    doc.add_heading("RELATÓRIO DO TESTE", 0)
    doc.add_paragraph("Cadastro de Conciliação Bancária Teste.")
    doc.add_paragraph(f"🗕️ Data do teste: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")
    doc.add_paragraph("Neste teste, o robô preencherá todos os dados e salvará o cadastro de uma nova Conciliação Bancária.")

    chrome_options = Options()
    chrome_options.add_argument("--start-maximized")
    driver = webdriver.Chrome(service=Service(ChromeDriverManager().install()), options=chrome_options)
    wait = WebDriverWait(driver, 10)

    def finalizar_relatorio():
        doc_name = f"relatorio_conciliacao_bancaria_cenario_1_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        doc.save(doc_name)
        log(doc, f"📄 Relatório salvo como: {doc_name}")
        try:
            subprocess.run(["start", "winword", doc_name], shell=True)
        except Exception as e:
            log(doc, f"Erro ao abrir o Word: {e}")
        driver.quit()

    def login():
        wait.until(EC.presence_of_element_located((By.ID, "j_id15:email"))).send_keys(LOGIN_EMAIL)
        wait.until(EC.presence_of_element_located((By.ID, "j_id15:senha"))).send_keys(LOGIN_PASSWORD, Keys.ENTER)
        wait.until(EC.presence_of_element_located((By.TAG_NAME, "body")))
        time.sleep(2)

    def abrir_menu():
        driver.find_element(By.TAG_NAME, "body").send_keys(Keys.F2)
        campo = wait.until(EC.presence_of_element_located((By.XPATH, "//input[@placeholder='Busque um cadastro']")))
        campo.click()
        campo.send_keys("Conciliação Bancária", Keys.ENTER)
        time.sleep(3)

    def acessar_formulario():
        cadastrar = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, 
            "#fmod_10008 ul li:nth-child(1) > a > span")))
        cadastrar.click()
        time.sleep(2)

    def fazer_upload_arquivo():
        log(doc, "🔄 Realizando upload do arquivo.")
        botao_selecione = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, 
            "#fmod_10008 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div > div > div > div:nth-child(2) > div > div > div > div > div")))
        botao_selecione.click()
        time.sleep(1)
        pyautogui.write(r'C:\Users\Gold System\Downloads\TESTECONCILIACAOBANCARIA.docx')
        pyautogui.press('enter')
        log(doc, "✅ Upload do arquivo realizado.")

    def selecionar_banco():
        log(doc, "🔄 Selecionando banco.")
        select_bancos_conciliacao = wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_10008 select")))
        Select(select_bancos_conciliacao).select_by_visible_text("Banco Bradesco")
        time.sleep(0.5)
        log(doc, "✅ Banco selecionado.")

    def salvar():
        salvar_btn = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR,
            "#fmod_10008 .btsave")))
        salvar_btn.click()
        time.sleep(2)


    def encontrar_mensagem_alerta():
        seletores = [
            (".alerts.salvo", "✅ Mensagem de Sucesso"),
            (".alerts.alerta", "⚠️ Mensagem de Alerta"),
            (".alerts.erro", "❌ Mensagem de Erro"),
        ]

        for seletor, tipo in seletores:
            try:
                elemento = driver.find_element(By.CSS_SELECTOR, seletor)
                if elemento.is_displayed():
                    log(doc, f"📢 {tipo}: {elemento.text}")
                    return elemento
            except:
                continue

        log(doc, "ℹ️ Nenhuma mensagem de alerta encontrada.")
        return None

    def fechar_modal():
        x_btn = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR,
            "#fmod_10008 .wdClose a")))
        x_btn.click()
        time.sleep(1)

    # EXECUÇÃO COM safe_action INDIVIDUAL PARA CADA AÇÃO
    if not safe_action(doc, "Acessando o sistema", lambda: driver.get(URL), driver, wait)[0]:
        finalizar_relatorio()
        return
    registrar_screenshot_unico("url_acessada", driver, doc, "Sistema acessado.")

    if not safe_action(doc, "Realizando login", login, driver, wait)[0]:
        finalizar_relatorio()
        return
    registrar_screenshot_unico("login_concluido", driver, doc, "Login realizado.")

    if not safe_action(doc, "Ajustando zoom", lambda: ajustar_zoom(driver), driver, wait)[0]:
        finalizar_relatorio()
        return

    if not safe_action(doc, "Abrindo menu Conciliação Bancária", abrir_menu, driver, wait)[0]:
        finalizar_relatorio()
        return
    registrar_screenshot_unico("menu_aberto", driver, doc, "Menu 'Conciliação Bancária' aberto.")

    if not safe_action(doc, "Acessando formulário de cadastro", acessar_formulario, driver, wait)[0]:
        finalizar_relatorio()
        return
    registrar_screenshot_unico("formulario_aberto", driver, doc, "Formulário de cadastro aberto.")

    # PREENCHIMENTO DOS CAMPOS
    if not safe_action(doc, "Fazendo upload do arquivo", fazer_upload_arquivo, driver, wait)[0]:
        finalizar_relatorio()
        return

    if not safe_action(doc, "Selecionando banco", selecionar_banco, driver, wait)[0]:
        finalizar_relatorio()
        return

    registrar_screenshot_unico("campos_preenchidos", driver, doc, "Campos preenchidos.")

    # SALVANDO O CADASTRO
    if not safe_action(doc, "Clicando no botão Salvar", salvar, driver, wait)[0]:
        finalizar_relatorio()
        return
    registrar_screenshot_unico("apos_salvar", driver, doc, "Clique no botão Salvar realizado.")

    encontrar_mensagem_alerta()

    # FECHANDO O FORMULÁRIO
    if not safe_action(doc, "Fechando formulário", fechar_modal, driver, wait)[0]:
        finalizar_relatorio()
        return
    registrar_screenshot_unico("formulario_fechado", driver, doc, "Formulário fechado.")

    log(doc, "✅ Teste de cadastro de Conciliação Bancária concluído com sucesso.")
    finalizar_relatorio()

if __name__ == "__main__":
    main()