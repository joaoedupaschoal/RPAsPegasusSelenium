# Refatorado e corrigido: cadastrodecronogramadefaturamento1ºcenario.py

from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support.ui import Select, WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service
from webdriver_manager.chrome import ChromeDriverManager
from docx import Document
from docx.shared import Inches
from faker import Faker
from faker.providers import BaseProvider
from datetime import datetime, timedelta
import subprocess
import os
import time
import random
import sys

sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))

# ==== CONFIGURAÇÕES ====
URL = "http://localhost:8080/gs/index.xhtml"
LOGIN_EMAIL = "joaoeduardo.gold@outlook.com"
LOGIN_PASSWORD = "071999gs"

# ==== PROVIDERS CUSTOMIZADOS ====
class BrasilProvider(BaseProvider):
    def rg(self):
        numeros = [str(random.randint(0, 9)) for _ in range(8)]
        return ''.join(numeros) + '-' + str(random.randint(0, 9))

fake = Faker("pt_BR")
fake.add_provider(BrasilProvider)

# ==== GERAÇÃO DE DADOS ====
def gerar_datas_cronograma():
    """Gera datas coerentes para o cronograma de faturamento."""
    hoje = datetime.today()
    
    # Gera uma data de leitura aleatória neste mês (entre dia 1 e 28)
    dia_leitura = random.randint(1, 28)
    data_leitura = hoje.replace(day=dia_leitura)

    # Gera uma data de entrega entre 5 e 10 dias depois da leitura
    dias_apos_leitura = random.randint(5, 10)
    data_entrega = data_leitura + timedelta(days=dias_apos_leitura)

    # Retorna as datas formatadas
    return {
        "data_leitura": data_leitura.strftime("%d/%m/%Y"),
        "data_entrega_conta": data_entrega.strftime("%d/%m/%Y")
    }

# Gera as datas
datas = gerar_datas_cronograma()

# Controle de screenshots únicas
screenshot_registradas = set()


def main():
    # ==== DOCUMENTO ====
    doc = Document()
    doc.add_heading("RELATÓRIO DO TESTE", 0)
    doc.add_paragraph("Cadastro de Cronograma de Faturamento - Cenário 4")
    doc.add_paragraph(f"🗓️ Data do teste: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")
    doc.add_paragraph("Neste teste, o robô preencherá os campos NÃO obrigatórios e salvará o cadastro de um novo Cronograma de Faturamento.")

    # ==== FUNÇÕES DE UTILIDADE ====
    def log(msg):
        print(msg)
        doc.add_paragraph(msg)

    def take_screenshot(nome):
        if nome not in screenshot_registradas:
            path = f"screenshots/{nome}.png"
            os.makedirs("screenshots", exist_ok=True)
            driver.save_screenshot(path)
            doc.add_paragraph(f"Screenshot: {nome}")
            try:
                doc.add_picture(path, width=Inches(5.5))
            except Exception as e:
                log(f"⚠️ Erro ao adicionar imagem: {e}")
            screenshot_registradas.add(nome)

    def safe_action(descricao, func):
        """Executa uma ação com tratamento de erro e retorna tupla (sucesso, resultado)"""
        try:
            log(f"📄 {descricao}...")
            resultado = func()
            log(f"✅ {descricao} realizada com sucesso.")
            take_screenshot(descricao.lower().replace(" ", "_"))
            return (True, resultado)
        except Exception as e:
            log(f"❌ Erro ao {descricao.lower()}: {e}")
            take_screenshot(f"erro_{descricao.lower().replace(' ', '_')}")
            return (False, None)

    def finalizar_relatorio():
        doc_name = f"relatorio_cronograma_faturamento_cenario_4_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        doc.save(doc_name)
        log(f"📄 Relatório salvo como: {doc_name}")
        try:
            subprocess.run(["start", "winword", doc_name], shell=True)
        except Exception as e:
            log(f"Erro ao abrir o Word: {e}")
        driver.quit()

    # ==== INICIALIZAÇÃO DO DRIVER ====
    chrome_options = Options()
    chrome_options.add_argument("--start-maximized")
    driver = webdriver.Chrome(service=Service(ChromeDriverManager().install()), options=chrome_options)
    wait = WebDriverWait(driver, 20)

    # ==== FUNÇÕES DE AÇÃO ====
    def ajustar_zoom():
        driver.execute_script("document.body.style.zoom='90%'")
        log("🔍 Zoom ajustado para 90%.")

    def login():
        wait.until(EC.presence_of_element_located((By.ID, "j_id15:email"))).send_keys(LOGIN_EMAIL)
        wait.until(EC.presence_of_element_located((By.ID, "j_id15:senha"))).send_keys(LOGIN_PASSWORD, Keys.ENTER)
        wait.until(EC.presence_of_element_located((By.TAG_NAME, "body")))
        time.sleep(2)

    def abrir_menu():
        driver.find_element(By.TAG_NAME, "body").send_keys(Keys.F2)
        campo = wait.until(EC.presence_of_element_located((By.XPATH, "//input[@placeholder='Busque um cadastro']")))
        campo.click()
        campo.send_keys("Cronograma de Faturamento", Keys.ENTER)
        time.sleep(3)

    def acessar_formulario():
        cadastrar = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, 
            "#fmod_10053 > div.wdTelas > div.telaInicial.clearfix.overflow.overflowY > ul > li:nth-child(1) > a > span")))
        cadastrar.click()
        time.sleep(2)

    def preencher_exercicio():
        log("📄 Preenchendo campo 'Exercício'.")
        campo_exercicio = wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, 
            "#fmod_10053 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_10080.categoriaHolder > div > div > div:nth-child(2) > input")))
        campo_exercicio.send_keys(fake.random_int(min=2020, max=2030))
        log("✅ Campo 'Exercício' preenchido.")

    def selecionar_concessionaria():
        log("📄 Selecionando Concessionária.")
        open_lov_concessionaria = wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, 
            "#fmod_10053 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_10080.categoriaHolder > div > div > div:nth-child(3) > div > a")))
        open_lov_concessionaria.click()
        
        campo_pesquisa = wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, 
            "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(2) > input.nomePesquisa")))
        time.sleep(1)
        campo_pesquisa.send_keys('TESTE CONCESSIONÁRIA DE ENERGIA SELENIUM AUTOMATIZ', Keys.ENTER)
        time.sleep(1)

        concessionaria = wait.until(EC.element_to_be_clickable((By.XPATH, 
            "//td[contains(text(), 'TESTE CONCESSIONÁRIA DE ENERGIA SELENIUM AUTOMATIZ')]")))
        concessionaria.click()
        log("✅ Concessionária selecionada.")

    def entrar_na_outra_aba():
        time.sleep(2)
        aba_cronograma = wait.until(EC.element_to_be_clickable((By.LINK_TEXT, 
            "Cronograma Faturamento")))
        aba_cronograma.click()
        time.sleep(2)

    def preencher_lote():
        log("📄 Preenchendo campo 'Lote'.")
        campo_lote = wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, 
            "#fmod_10053 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_10081.categoriaHolder > div > div > div:nth-child(3) > input")))
        campo_lote.send_keys(fake.random_int(min=1, max=99))
        log("✅ Campo 'Lote' preenchido.")

    def preencher_data_leitura():
        log("📄 Preenchendo campo 'Data de Leitura'.")
        campo_data_leitura = wait.until(EC.element_to_be_clickable((By.XPATH,
            "//input[@grupo='10087' and @ref='10257' and contains(@class, 'hasDatepicker isList mandatory fc')]")))
        campo_data_leitura.send_keys(datas["data_leitura"])
        driver.execute_script("arguments[0].value = arguments[1];", campo_data_leitura, datas["data_leitura"])
        log("✅ Campo 'Data de Leitura' preenchido.")

    def selecionar_dia_semana():
        log("📄 Selecionando dia da semana.")
        Select(driver.find_element(By.CSS_SELECTOR, 
            "#fmod_10053 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_10081.categoriaHolder > div > div > div:nth-child(5) > select")).select_by_visible_text("Quarta-feira")
        log("✅ Dia da semana selecionado.")

    def preencher_data_entrega_conta():
        log("📄 Preenchendo campo 'Data de Entrega da Conta'.")
        campo_data_entrega_conta = wait.until(EC.element_to_be_clickable((By.XPATH,
            "//input[@grupo='10087' and @ref='10259' and contains(@class, 'hasDatepicker isList mandatory fc')]")))
        campo_data_entrega_conta.send_keys(datas["data_entrega_conta"])
        driver.execute_script("arguments[0].value = arguments[1];", campo_data_entrega_conta, datas["data_entrega_conta"])
        log("✅ Campo 'Data de Entrega da Conta' preenchido.")

    def preencher_pfc():
        log("📄 Preenchendo campo 'PFC'.")
        campo_pfc = wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, 
            "#fmod_10053 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_10081.categoriaHolder > div > div > div:nth-child(7) > input")))
        campo_pfc.send_keys(fake.random_int(min=1, max=99))
        log("✅ Campo 'PFC' preenchido.")

    def selecionar_mes_inicio():
        log("📄 Selecionando mês de início.")
        Select(driver.find_element(By.CSS_SELECTOR, 
            "#fmod_10053 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_10081.categoriaHolder > div > div > div:nth-child(8) > select")).select_by_visible_text("01-Janeiro")
        log("✅ Mês de início selecionado.")

    def preencher_dia_vencimento():
        log("📄 Preenchendo campo 'Dia de Vencimento'.")
        campo_dia_vencimento = wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, 
            "#fmod_10053 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_10081.categoriaHolder > div > div > div:nth-child(9) > input")))
        campo_dia_vencimento.send_keys(fake.random_int(min=15, max=30))
        log("✅ Campo 'Dia de Vencimento' preenchido.")

    def selecionar_mes_fevereiro():
        log("📄 Selecionando mês de fevereiro.")
        Select(driver.find_element(By.CSS_SELECTOR, 
            "#fmod_10053 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_10081.categoriaHolder > div > div > div:nth-child(10) > select")).select_by_visible_text("02-Fevereiro")
        log("✅ Mês de fevereiro selecionado.")

    def selecionar_mes_marco():
        log("📄 Selecionando mês de março.")
        Select(driver.find_element(By.CSS_SELECTOR, 
            "#fmod_10053 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_10081.categoriaHolder > div > div > div:nth-child(11) > select")).select_by_visible_text("03-Março")
        log("✅ Mês de março selecionado.")

    def adicionar_item():
        log("📄 Clicando no botão 'Adicionar'.")
        adicionar = wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, 
            "#fmod_10053 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_10081.categoriaHolder > div > div > div.btnListHolder > a.btAddGroup")))
        adicionar.click()
        log("✅ Item adicionado.")

    def salvar():
        salvar_btn = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR,
            "#fmod_10053 > div.wdTelas > div.telaCadastro.clearfix > div.btnHolder > a.btModel.btGray.btsave")))
        salvar_btn.click()
        time.sleep(2)

    def encontrar_mensagem_alerta():
        seletores = [
            (".alerts.salvo", "✅ Mensagem de Sucesso"),
            (".alerts.alerta", "⚠️ Mensagem de Alerta"),
            (".alerts.erro", "❌ Mensagem de Erro"),
        ]

        for seletor, tipo in seletores:
            try:
                elemento = driver.find_element(By.CSS_SELECTOR, seletor)
                if elemento.is_displayed():
                    log(f"📢 {tipo}: {elemento.text}")
                    return elemento
            except:
                continue

        log("ℹ️ Nenhuma mensagem de alerta encontrada.")
        return None

    def fechar_modal():
        x_btn = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR,
            "#fmod_10053 > div.wdTop.ui-draggable-handle > div.wdClose > a")))
        x_btn.click()
        time.sleep(1)

    def registrar_screenshot_unico(nome, descricao=None):
        if nome not in screenshot_registradas:
            if descricao:
                log(f"📸 {descricao}")
            take_screenshot(nome)

    # ==== EXECUÇÃO DO TESTE ====
    sucesso, _ = safe_action("Acessando o sistema", lambda: driver.get(URL))
    if not sucesso:
        finalizar_relatorio()
        return

    sucesso, _ = safe_action("Realizando Login", login)
    if not sucesso:
        finalizar_relatorio()
        return

    sucesso, _ = safe_action("Ajustando zoom", ajustar_zoom)
    if not sucesso:
        finalizar_relatorio()
        return

    sucesso, _ = safe_action("Abrindo menu do Cronograma de Faturamento", abrir_menu)
    if not sucesso:
        finalizar_relatorio()
        return

    sucesso, _ = safe_action("Acessando Formulário de cadastro", acessar_formulario)
    if not sucesso:
        finalizar_relatorio()
        return

    sucesso, _ = safe_action("Entrando na aba 'Cronograma Faturamento'", entrar_na_outra_aba)
    if not sucesso:
        finalizar_relatorio()
        return


    sucesso, _ = safe_action("Preenchendo Lote", preencher_lote)
    if not sucesso:
        finalizar_relatorio()
        return

    sucesso, _ = safe_action("Preenchendo Data de Leitura", preencher_data_leitura)
    if not sucesso:
        finalizar_relatorio()
        return

    sucesso, _ = safe_action("Selecionando Dia da Semana", selecionar_dia_semana)
    if not sucesso:
        finalizar_relatorio()
        return

    sucesso, _ = safe_action("Preenchendo Data de Entrega da Conta", preencher_data_entrega_conta)
    if not sucesso:
        finalizar_relatorio()
        return

    sucesso, _ = safe_action("Preenchendo PFC", preencher_pfc)
    if not sucesso:
        finalizar_relatorio()
        return

    sucesso, _ = safe_action("Selecionando Mês de Início", selecionar_mes_inicio)
    if not sucesso:
        finalizar_relatorio()
        return

    sucesso, _ = safe_action("Preenchendo Dia de Vencimento", preencher_dia_vencimento)
    if not sucesso:
        finalizar_relatorio()
        return

    sucesso, _ = safe_action("Selecionando Mês de Fevereiro", selecionar_mes_fevereiro)
    if not sucesso:
        finalizar_relatorio()
        return

    sucesso, _ = safe_action("Selecionando Mês de Março", selecionar_mes_marco)
    if not sucesso:
        finalizar_relatorio()
        return

    sucesso, _ = safe_action("Adicionando Item", adicionar_item)
    if not sucesso:
        finalizar_relatorio()
        return

    registrar_screenshot_unico("campos_preenchidos", "Todos os campos preenchidos.")

    # SALVANDO O CADASTRO
    sucesso, _ = safe_action("Clicando no botão Salvar", salvar)
    if not sucesso:
        finalizar_relatorio()
        return
    
    registrar_screenshot_unico("apos_salvar", "Clique no botão Salvar realizado.")

    encontrar_mensagem_alerta()

    # FECHANDO O FORMULÁRIO
    sucesso, _ = safe_action("Fechando formulário", fechar_modal)
    if not sucesso:
        finalizar_relatorio()
        return
    
    registrar_screenshot_unico("formulario_fechado", "Formulário fechado.")

    log("✅ Teste de cadastro de Cronograma de Faturamento concluído com sucesso.")
    finalizar_relatorio()


if __name__ == "__main__":
    main()