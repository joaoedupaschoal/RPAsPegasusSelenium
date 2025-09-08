from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support.ui import WebDriverWait, Select
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service
from selenium.webdriver import ActionChains
from webdriver_manager.chrome import ChromeDriverManager
from docx import Document
from datetime import datetime, timedelta
import subprocess
import time
import os
import sys
from faker import Faker
from faker.providers import BaseProvider
import random
import string
from selenium.common.exceptions import TimeoutException, NoSuchElementException

sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))

from utils.actions import log, take_screenshot, safe_action, encontrar_mensagem_alerta, ajustar_zoom

URL = "http://localhost:8080/gs/index.xhtml"
LOGIN_EMAIL = "joaoeduardo.gold@outlook.com"
LOGIN_PASSWORD = "071999gs"

# Configuração do Faker
class BrasilProvider(BaseProvider):
    def rg(self):
        numeros = [str(random.randint(0, 9)) for _ in range(8)]
        return ''.join(numeros) + '-' + str(random.randint(0, 9))

fake = Faker("pt_BR")
fake.add_provider(BrasilProvider)

# Geração de dados aleatórios
def gerar_datas_cronograma():
    hoje = datetime.today()
    
    # Gera uma data de leitura aleatória neste mês (entre dia 1 e 28)
    dia_leitura = random.randint(1, 28)
    data_leitura = hoje.replace(day=dia_leitura)

    # Gera uma data de entrega entre 5 e 10 dias depois da leitura
    dias_apos_leitura = random.randint(5, 10)
    data_entrega = data_leitura + timedelta(days=dias_apos_leitura)

    # Retorna as datas formatadas
    return {
        "data_leitura": data_leitura.strftime("%d/%m/%Y"),
        "data_entrega_conta": data_entrega.strftime("%d/%m/%Y")
    }

datas = gerar_datas_cronograma()

# Controle de screenshots únicas
screenshot_registradas = set()
def registrar_screenshot_unico(nome, driver, doc, descricao=None):
    if nome not in screenshot_registradas:
        if descricao:
            log(doc, f"📸 {descricao}")
        take_screenshot(driver, doc, nome)
        screenshot_registradas.add(nome)

def main():
    doc = Document()
    doc.add_heading("RELATÓRIO DO TESTE", 0)
    doc.add_paragraph("Cadastro de Cronograma de Faturamento Teste.")
    doc.add_paragraph(f"🗕️ Data do teste: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")
    doc.add_paragraph("Neste teste, o robô preencherá os campos obrigatórios e salvará o cadastro de um novo Cronograma de Faturamento.")

    chrome_options = Options()
    chrome_options.add_argument("--start-maximized")
    driver = webdriver.Chrome(service=Service(ChromeDriverManager().install()), options=chrome_options)
    wait = WebDriverWait(driver, 10)

    def finalizar_relatorio():
        doc_name = f"relatorio_cronograma_faturamento_cenario_1_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        doc.save(doc_name)
        log(doc, f"📄 Relatório salvo como: {doc_name}")
        try:
            subprocess.run(["start", "winword", doc_name], shell=True)
        except Exception as e:
            log(doc, f"Erro ao abrir o Word: {e}")
        driver.quit()

    def login():
        wait.until(EC.presence_of_element_located((By.ID, "j_id15:email"))).send_keys(LOGIN_EMAIL)
        wait.until(EC.presence_of_element_located((By.ID, "j_id15:senha"))).send_keys(LOGIN_PASSWORD, Keys.ENTER)
        wait.until(EC.presence_of_element_located((By.TAG_NAME, "body")))
        time.sleep(2)

    def abrir_menu():
        driver.find_element(By.TAG_NAME, "body").send_keys(Keys.F2)
        campo = wait.until(EC.presence_of_element_located((By.XPATH, "//input[@placeholder='Busque um cadastro']")))
        campo.click()
        campo.send_keys("Cronograma de Faturamento", Keys.ENTER)
        time.sleep(3)

    def acessar_formulario():
        cadastrar = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, 
            "#fmod_10053 > div.wdTelas > div.telaInicial.clearfix.overflow.overflowY > ul > li:nth-child(1) > a > span")))
        cadastrar.click()
        time.sleep(2)

    def preencher_exercicio():
        log(doc, "🔄 Preenchendo campo 'Exercício'.")
        campo_exercicio = wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, 
            "#fmod_10053 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_10080.categoriaHolder > div > div > div:nth-child(2) > input")))
        campo_exercicio.send_keys(fake.random_int(min=2020, max=2030))
        log(doc, "✅ Campo 'Exercício' preenchido.")

    def selecionar_concessionaria():
        log(doc, "🔄 Selecionando Concessionária.")
        open_lov_concessionaria = wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, 
            "#fmod_10053 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_10080.categoriaHolder > div > div > div:nth-child(3) > div > a")))
        open_lov_concessionaria.click()
        
        campo_pesquisa = wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, 
            "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(2) > input.nomePesquisa")))
        time.sleep(1)
        campo_pesquisa.send_keys('TESTE CONCESSIONÁRIA DE ENERGIA SELENIUM AUTOMATIZ', Keys.ENTER)
        time.sleep(1)

        concessionaria = wait.until(EC.element_to_be_clickable((By.XPATH, 
            "//td[contains(text(), 'TESTE CONCESSIONÁRIA DE ENERGIA SELENIUM AUTOMATIZ')]")))
        concessionaria.click()
        log(doc, "✅ Concessionária selecionada.")

    def preencher_lote():
        log(doc, "🔄 Preenchendo campo 'Lote'.")
        campo_lote = wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, 
            "#fmod_10053 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_10081.categoriaHolder > div > div > div:nth-child(3) > input")))
        campo_lote.send_keys(fake.random_int(min=1, max=99))
        log(doc, "✅ Campo 'Lote' preenchido.")

    def preencher_data_leitura():
        log(doc, "🔄 Preenchendo campo 'Data de Leitura'.")
        campo_data_leitura = wait.until(EC.element_to_be_clickable((By.XPATH,
            "//input[@grupo='10087' and @ref='10257' and contains(@class, 'hasDatepicker isList mandatory fc')]")))
        campo_data_leitura.send_keys(datas["data_leitura"])
        driver.execute_script("arguments[0].value = arguments[1];", campo_data_leitura, datas["data_leitura"])
        log(doc, "✅ Campo 'Data de Leitura' preenchido.")

    def selecionar_dia_semana():
        log(doc, "🔄 Selecionando dia da semana.")
        Select(driver.find_element(By.CSS_SELECTOR, 
            "#fmod_10053 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_10081.categoriaHolder > div > div > div:nth-child(5) > select")).select_by_visible_text("Quarta-feira")
        log(doc, "✅ Dia da semana selecionado.")

    def preencher_data_entrega_conta():
        log(doc, "🔄 Preenchendo campo 'Data de Entrega da Conta'.")
        campo_data_entrega_conta = wait.until(EC.element_to_be_clickable((By.XPATH,
            "//input[@grupo='10087' and @ref='10259' and contains(@class, 'hasDatepicker isList mandatory fc')]")))
        campo_data_entrega_conta.send_keys(datas["data_entrega_conta"])
        driver.execute_script("arguments[0].value = arguments[1];", campo_data_entrega_conta, datas["data_entrega_conta"])
        log(doc, "✅ Campo 'Data de Entrega da Conta' preenchido.")

    def preencher_pfc():
        log(doc, "🔄 Preenchendo campo 'PFC'.")
        campo_pfc = wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, 
            "#fmod_10053 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_10081.categoriaHolder > div > div > div:nth-child(7) > input")))
        campo_pfc.send_keys(fake.random_int(min=1, max=99))
        log(doc, "✅ Campo 'PFC' preenchido.")

    def selecionar_mes_inicio():
        log(doc, "🔄 Selecionando mês de início.")
        Select(driver.find_element(By.CSS_SELECTOR, 
            "#fmod_10053 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_10081.categoriaHolder > div > div > div:nth-child(8) > select")).select_by_visible_text("01-Janeiro")
        log(doc, "✅ Mês de início selecionado.")

    def preencher_dia_vencimento():
        log(doc, "🔄 Preenchendo campo 'Dia de Vencimento'.")
        campo_dia_vencimento = wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, 
            "#fmod_10053 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_10081.categoriaHolder > div > div > div:nth-child(9) > input")))
        campo_dia_vencimento.send_keys(fake.random_int(min=15, max=30))
        log(doc, "✅ Campo 'Dia de Vencimento' preenchido.")

    def selecionar_mes_fevereiro():
        log(doc, "🔄 Selecionando mês de fevereiro.")
        Select(driver.find_element(By.CSS_SELECTOR, 
            "#fmod_10053 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_10081.categoriaHolder > div > div > div:nth-child(10) > select")).select_by_visible_text("02-Fevereiro")
        log(doc, "✅ Mês de fevereiro selecionado.")

    def selecionar_mes_marco():
        log(doc, "🔄 Selecionando mês de março.")
        Select(driver.find_element(By.CSS_SELECTOR, 
            "#fmod_10053 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_10081.categoriaHolder > div > div > div:nth-child(11) > select")).select_by_visible_text("03-Março")
        log(doc, "✅ Mês de março selecionado.")

    def adicionar_item():
        log(doc, "🔄 Clicando no botão 'Adicionar'.")
        adicionar = wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, 
            "#fmod_10053 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_10081.categoriaHolder > div > div > div.btnListHolder > a.btAddGroup")))
        adicionar.click()
        log(doc, "✅ Item adicionado.")

    def salvar():
        salvar_btn = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR,
            "#fmod_10053 > div.wdTelas > div.telaCadastro.clearfix > div.btnHolder > a.btModel.btGray.btsave")))
        salvar_btn.click()
        time.sleep(2)

    def fechar_modal():
        x_btn = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR,
            "#fmod_10053 > div.wdTop.ui-draggable-handle > div.wdClose > a")))
        x_btn.click()
        time.sleep(1)

    # EXECUÇÃO COM safe_action INDIVIDUAL PARA CADA AÇÃO
    if not safe_action(doc, "Acessando o sistema", lambda: driver.get(URL), driver, wait)[0]:
        finalizar_relatorio()
        return
    registrar_screenshot_unico("url_acessada", driver, doc, "Sistema acessado.")

    if not safe_action(doc, "Realizando login", login, driver, wait)[0]:
        finalizar_relatorio()
        return
    registrar_screenshot_unico("login_concluido", driver, doc, "Login realizado.")

    if not safe_action(doc, "Ajustando zoom", lambda: ajustar_zoom(driver), driver, wait)[0]:
        finalizar_relatorio()
        return

    if not safe_action(doc, "Abrindo menu Cronograma de Faturamento", abrir_menu, driver, wait)[0]:
        finalizar_relatorio()
        return
    registrar_screenshot_unico("menu_aberto", driver, doc, "Menu 'Cronograma de Faturamento' aberto.")

    if not safe_action(doc, "Acessando formulário de cadastro", acessar_formulario, driver, wait)[0]:
        finalizar_relatorio()
        return
    registrar_screenshot_unico("formulario_aberto", driver, doc, "Formulário de cadastro aberto.")

    # PREENCHIMENTO DOS CAMPOS - safe_action individual para cada campo
    if not safe_action(doc, "Preenchendo Exercício", preencher_exercicio, driver, wait)[0]:
        finalizar_relatorio()
        return

    if not safe_action(doc, "Selecionando Concessionária", selecionar_concessionaria, driver, wait)[0]:
        finalizar_relatorio()
        return

    if not safe_action(doc, "Preenchendo Lote", preencher_lote, driver, wait)[0]:
        finalizar_relatorio()
        return

    if not safe_action(doc, "Preenchendo Data de Leitura", preencher_data_leitura, driver, wait)[0]:
        finalizar_relatorio()
        return

    if not safe_action(doc, "Selecionando Dia da Semana", selecionar_dia_semana, driver, wait)[0]:
        finalizar_relatorio()
        return

    if not safe_action(doc, "Preenchendo Data de Entrega da Conta", preencher_data_entrega_conta, driver, wait)[0]:
        finalizar_relatorio()
        return

    if not safe_action(doc, "Preenchendo PFC", preencher_pfc, driver, wait)[0]:
        finalizar_relatorio()
        return

    if not safe_action(doc, "Selecionando Mês de Início", selecionar_mes_inicio, driver, wait)[0]:
        finalizar_relatorio()
        return

    if not safe_action(doc, "Preenchendo Dia de Vencimento", preencher_dia_vencimento, driver, wait)[0]:
        finalizar_relatorio()
        return

    if not safe_action(doc, "Selecionando Mês de Fevereiro", selecionar_mes_fevereiro, driver, wait)[0]:
        finalizar_relatorio()
        return

    if not safe_action(doc, "Selecionando Mês de Março", selecionar_mes_marco, driver, wait)[0]:
        finalizar_relatorio()
        return

    if not safe_action(doc, "Adicionando Item", adicionar_item, driver, wait)[0]:
        finalizar_relatorio()
        return

    registrar_screenshot_unico("campos_preenchidos", driver, doc, "Todos os campos preenchidos.")

    # SALVANDO O CADASTRO
    if not safe_action(doc, "Clicando no botão Salvar", salvar, driver, wait)[0]:
        finalizar_relatorio()
        return
    registrar_screenshot_unico("apos_salvar", driver, doc, "Clique no botão Salvar realizado.")

    # VERIFICANDO MENSAGEM DE RETORNO
    _, tipo_alerta = encontrar_mensagem_alerta(driver, doc)
    if tipo_alerta == "sucesso":
        log(doc, "✅ Mensagem de sucesso exibida.")
    elif tipo_alerta == "alerta":
        log(doc, "⚠️ Mensagem de alerta exibida.")
    elif tipo_alerta == "erro":
        log(doc, "❌ Mensagem de erro exibida.")
    else:
        log(doc, "⚠️ Nenhuma mensagem encontrada após salvar.")
    registrar_screenshot_unico("mensagem_final", driver, doc, "Mensagem exibida após salvar.")

    # FECHANDO O FORMULÁRIO
    if not safe_action(doc, "Fechando formulário", fechar_modal, driver, wait)[0]:
        finalizar_relatorio()
        return
    registrar_screenshot_unico("formulario_fechado", driver, doc, "Formulário fechado.")

    log(doc, "✅ Teste de cadastro de Cronograma de Faturamento concluído com sucesso.")
    finalizar_relatorio()

if __name__ == "__main__":
    main()