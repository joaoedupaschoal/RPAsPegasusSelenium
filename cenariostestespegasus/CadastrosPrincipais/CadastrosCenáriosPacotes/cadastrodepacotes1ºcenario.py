# Refatorado e organizado: cadastrodeescalamotorista2ºcenario.py

from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support.ui import Select, WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service
from webdriver_manager.chrome import ChromeDriverManager
from docx import Document
from docx.shared import Inches
from faker import Faker
from faker.providers import BaseProvider
from validate_docbr import CPF
from datetime import datetime, timedelta
import subprocess
import os
import time
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service
from webdriver_manager.chrome import ChromeDriverManager
from docx import Document
from docx.shared import Inches
from faker import Faker
from datetime import datetime
import subprocess
import os
import time
import random

# ==== PROVIDERS CUSTOMIZADOS ====
class BrasilProvider(BaseProvider):
    def rg(self):
        numeros = [str(random.randint(0, 9)) for _ in range(8)]
        return ''.join(numeros) + '-' + str(random.randint(0, 9))

fake = Faker("pt_BR")
fake.add_provider(BrasilProvider)
qtd_parcelas_em_atraso = int(fake.random.choice(['1', '2', '3', '4', '5']))


idade_de = fake.random_int(min=0, max=5)
idade_até = idade_de + fake.random_int(min=0, max=5)

valor_taxa_adesao = round(random.uniform(2000, 10000), 2)

def gerar_datas_validas():
    """Gera datas coerentes para admissão, início e fim da escala, e vencimento da CNH."""
    hoje = datetime.today().date()
    
    # Data de admissão entre 10 anos atrás e hoje
    data_admissao = fake.date_between(start_date=hoje - timedelta(days=3650), end_date=hoje)
    
    # Data de início da escala entre hoje e 1 ano no futuro
    data_inicio = fake.date_between(start_date=hoje, end_date=hoje + timedelta(days=365))
    
    # Data fim entre 1 e 180 dias após a data de início
    data_fim = data_inicio + timedelta(days=random.randint(1, 180))
    
    # Vencimento CNH entre hoje e 10 anos no futuro
    vencimento_cnh = fake.date_between(start_date=hoje, end_date=hoje + timedelta(days=3650))
    
    return (data_admissao.strftime('%d/%m/%Y'), 
            data_inicio.strftime('%d/%m/%Y'), 
            data_fim.strftime('%d/%m/%Y'), 
            vencimento_cnh.strftime('%d/%m/%Y'))

def gerar_dados_documentos():
    """Gera documentos fictícios para o cadastro."""
    carteira_trabalho = str(random.randint(10000000, 99999999))
    pis = fake.cpf().replace('.', '').replace('-', '')[:11]
    cnh = str(random.randint(10000000000, 99999999999))
    cpf = CPF().generate()
    
    return carteira_trabalho, pis, cnh, cpf

# Gera os dados necessários
data_admissao, data_inicio, data_fim, vencimento_cnh = gerar_datas_validas()
carteira_trabalho, pis, cnh, cpf_valido = gerar_dados_documentos()

# ==== CONFIGURAÇÕES ====
URL = "http://localhost:8080/gs/index.xhtml"
LOGIN_EMAIL = "joaoeduardo.gold@outlook.com"
LOGIN_PASSWORD = "071999gs"

# ==== DOCUMENTO ====
doc = Document()
doc.add_heading("RELATÓRIO DO TESTE", 0)
doc.add_paragraph("Cadastro de Pacotes – Cenário 1: Preenchimento completo e salvamento.")
doc.add_paragraph(f"Data do teste: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")

screenshot_registradas = set()

# ==== FUNÇÕES DE UTILITÁRIO ====
def log(doc, msg):
    print(msg)
    doc.add_paragraph(msg)

def take_screenshot(driver, doc, nome):
    if nome not in screenshot_registradas:
        path = f"screenshots/{nome}.png"
        os.makedirs("screenshots", exist_ok=True)
        driver.save_screenshot(path)
        doc.add_paragraph(f"Screenshot: {nome}")
        doc.add_picture(path, width=Inches(5.5))
        screenshot_registradas.add(nome)

def safe_action(doc, descricao, func):
    try:
        log(doc, f"🔄 {descricao}...")
        func()
        log(doc, f"✅ {descricao} realizada com sucesso.")
        take_screenshot(driver, doc, descricao.lower().replace(" ", "_"))
    except Exception as e:
        log(doc, f"❌ Erro ao {descricao.lower()}: {e}")
        take_screenshot(driver, doc, f"erro_{descricao.lower().replace(' ', '_')}")

def finalizar_relatorio():
    nome_arquivo = f"relatorio_pacotes_cenario_1_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc.save(nome_arquivo)
    log(doc, f"📄 Relatório salvo como: {nome_arquivo}")
    subprocess.run(["start", "winword", nome_arquivo], shell=True)
    driver.quit()

def rolar_ate_dependentes_contemplados():
    campo_dependentes_contemplados = WebDriverWait(driver, 10).until(
        EC.element_to_be_clickable((By.CSS_SELECTOR, '#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_parametrosGerais.categoriaHolder > div.groupHolder.clearfix.grupo_dependentesContemplados > div.group_dependentesContemplados.clearfix.grupoHolder.lista > div > div > select'))
    )
    driver.execute_script("arguments[0].scrollIntoView();", campo_dependentes_contemplados)


def rolar_ate_formas_pagamento_mensalidade():
    campo_formas_pagamento_mensalidade = WebDriverWait(driver, 10).until(
        EC.element_to_be_clickable((By.CSS_SELECTOR, '#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_formasPagamento.categoriaHolder > div.groupHolder.clearfix.grupo_formaPagamentoMensalidade > div.group_formaPagamentoMensalidade.clearfix.grupoHolder.lista > div > div:nth-child(1) > select'))
    )
    driver.execute_script("arguments[0].scrollIntoView();", campo_formas_pagamento_mensalidade)


def rolar_ate_adicionar_formas_pagamento_mensalidade():
    adicionar = WebDriverWait(driver, 10).until(
        EC.element_to_be_clickable((By.CSS_SELECTOR, '#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_formasPagamento.categoriaHolder > div.groupHolder.clearfix.grupo_formaPagamentoMensalidade > div.btnListHolder > a.btAddGroup'))
    )
    driver.execute_script("arguments[0].scrollIntoView();", adicionar)
    adicionar.click()


def abrir_modal_e_selecionar(btn_selector, pesquisa_selector, termo_pesquisa, btn_pesquisar_selector, resultado_xpath):
    def acao():
        # Abre o modal
        open_lov = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, btn_selector)))
        open_lov.click()

        # Aguarda campo pesquisa
        campo_pesquisa = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, pesquisa_selector)))
        campo_pesquisa.clear()
        campo_pesquisa.send_keys(termo_pesquisa)

        # Clica pesquisar
        pesquisar = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, btn_pesquisar_selector)))
        pesquisar.click()
        time.sleep(1)
        # Espera o resultado carregar
        wait.until(EC.presence_of_element_located((By.XPATH, resultado_xpath)))
        wait.until(EC.visibility_of_element_located((By.XPATH, resultado_xpath)))
        wait.until(EC.element_to_be_clickable((By.XPATH, resultado_xpath)))

        # Relocaliza no último instante (evita stale element)
        resultado = driver.find_element(By.XPATH, resultado_xpath)
        driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", resultado)
        time.sleep(0.2)
        resultado.click()

    return acao

def encontrar_mensagem_alerta():
    seletores = [
        (".alerts.salvo", "✅ Sucesso"),
        (".alerts.alerta", "⚠️ Alerta"),
        (".alerts.erro", "❌ Erro"),
    ]

    for seletor, tipo in seletores:
        try:
            elemento = driver.find_element(By.CSS_SELECTOR, seletor)
            if elemento.is_displayed():
                log(doc, f"📢 {tipo}: {elemento.text}")
                return elemento
        except:
            continue

    log(doc, "ℹ️ Nenhuma mensagem de alerta encontrada.")
    return None

def ajustar_zoom():
    try:
        driver.execute_script("document.body.style.zoom='90%'")
        log(doc, "🔍 Zoom ajustado para 90%.")
    except Exception as e:
        log(doc, f"⚠️ Erro ao ajustar zoom: {e}")

def preencher_campo_data(selector, valor):
    def acao():
        campo = wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, selector)))
        campo.click()
        campo.clear()
        campo.send_keys(valor)
        time.sleep(0.2)
    return acao

def selecionar_opcao(selector, texto):
    def acao():
        select_element = wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, selector)))
        Select(select_element).select_by_visible_text(texto)
    return acao

# ==== INICIALIZAÇÃO DO DRIVER ====
options = Options()
options.add_argument("--start-maximized")
driver = webdriver.Chrome(service=Service(ChromeDriverManager().install()), options=options)
wait = WebDriverWait(driver, 20)

# ==== EXECUÇÃO DO TESTE ====
try:
    safe_action(doc, "Acessando sistema", lambda: driver.get(URL))

    safe_action(doc, "Realizando login", lambda: (
        wait.until(EC.presence_of_element_located((By.ID, "j_id15:email"))).send_keys(LOGIN_EMAIL),
        wait.until(EC.presence_of_element_located((By.ID, "j_id15:senha"))).send_keys(LOGIN_PASSWORD, Keys.ENTER),
        wait.until(EC.presence_of_element_located((By.TAG_NAME, "body")))
    ))

    safe_action(doc, "Esperando sistema carregar e ajustando zoom", lambda: (
        time.sleep(5),
        ajustar_zoom()
    ))

    safe_action(doc, "Abrindo menu Pacotes", lambda: (
        driver.find_element(By.TAG_NAME, "body").send_keys(Keys.F2),
        time.sleep(1),
        wait.until(EC.visibility_of_element_located((By.XPATH, "//input[@placeholder='Busque um cadastro']"))).send_keys("Pacotes", Keys.ENTER),
        wait.until(EC.visibility_of_element_located((By.XPATH, "/html/body/div[17]/div[2]/ul/li[28]/a"))).click()

#

    ))

    safe_action(doc, "Clicando em Cadastrar", lambda: (
        time.sleep(2),
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_2 > div.wdTelas > div.telaInicial.clearfix.overflow.overflowY > ul > li:nth-child(1) > a > span"))).click()
    ))

    safe_action(doc, "Preechendo o Nome do Pacote", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_parametrosGerais.categoriaHolder > div.groupHolder.clearfix.grupo_parametrosGerais > div > div > div:nth-child(2) > input"))
        ).send_keys('TESTE PACOTE SELENIUM AUTOMATIZADO')
    ))

    safe_action(doc, "Preechendo a Quantidade Máxima de Titulares", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_parametrosGerais.categoriaHolder > div.groupHolder.clearfix.grupo_parametrosGerais > div > div > div:nth-child(4) > input"))
        ).send_keys(fake.random_int(min=1, max=3))
    ))

    safe_action(doc, "Preechendo a Quantidade Máxima de Dependentes", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_parametrosGerais.categoriaHolder > div.groupHolder.clearfix.grupo_parametrosGerais > div > div > div:nth-child(5) > input"))
        ).send_keys(fake.random_int(min=1, max=99))
    ))


    safe_action(doc, "Preechendo a Vigência (em dias)", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_parametrosGerais.categoriaHolder > div.groupHolder.clearfix.grupo_parametrosGerais > div > div > div:nth-child(6) > input"))
        ).send_keys(fake.random_int(min=30, max=365))
    ))


    safe_action(doc, "Preechendo a Carência (em dias) para utilização", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_parametrosGerais.categoriaHolder > div.groupHolder.clearfix.grupo_parametrosGerais > div > div > div:nth-child(8) > input"))
        ).send_keys(fake.random_int(min=0, max=15))
    ))


    safe_action(doc, "Preechendo a Carência (em dias) para mensalidade", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_parametrosGerais.categoriaHolder > div.groupHolder.clearfix.grupo_parametrosGerais > div > div > div:nth-child(10) > input"))
        ).send_keys(fake.random_int(min=0, max=15))
    ))



    safe_action(doc, "Preechendo a Carência (em dias) para adesão", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_parametrosGerais.categoriaHolder > div.groupHolder.clearfix.grupo_parametrosGerais > div > div > div:nth-child(11) > input"))
        ).send_keys(fake.random_int(min=0, max=15))
    ))



    safe_action(doc, "Selecionando a opção para o Pacote possuir Jazigo", selecionar_opcao(
        "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_parametrosGerais.categoriaHolder > div.groupHolder.clearfix.grupo_parametrosGerais > div > div > div:nth-child(12) > select",
        "Com Jazigo"
    ))


    safe_action(doc, "Selecionando a opção para o Pacote não vincular a quadra", selecionar_opcao(
        "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_parametrosGerais.categoriaHolder > div.groupHolder.clearfix.grupo_parametrosGerais > div > div > div:nth-child(13) > select",
        "Não"
    ))

    
    safe_action(doc, "Selecionando a opção para o Pacote possuir Sepultamento", selecionar_opcao(
        "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_parametrosGerais.categoriaHolder > div.groupHolder.clearfix.grupo_parametrosGerais > div > div > div:nth-child(14) > select",
        "Com Sepultamento"
    ))


    safe_action(doc, "Selecionando Tipo de Contrato", abrir_modal_e_selecionar(
        "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_parametrosGerais.categoriaHolder > div.groupHolder.clearfix.grupo_parametrosGerais > div > div > div:nth-child(15) > div > a",
        "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(2) > input",
        "TIPO DE CONTRATO TESTE SELENIUM AUTOMATIZADO",
        "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(4) > a",
        "//td[contains(text(), 'TIPO DE CONTRATO TESTE SELENIUM AUTOMATIZADO')]/a[contains(@class, 'linkAlterar')]"
    ))



    safe_action(doc, "Selecionando a opção para o Tipo de Pacote na OS ser Por Contrato", selecionar_opcao(
        "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_parametrosGerais.categoriaHolder > div.groupHolder.clearfix.grupo_parametrosGerais > div > div > div:nth-child(16) > select",
        "Por Contrato"
    ))

    safe_action(doc, "Selecionando a opção para Lançar valor por Quantidade de Dependentes", selecionar_opcao(
        "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_parametrosGerais.categoriaHolder > div.groupHolder.clearfix.grupo_parametrosGerais > div > div > div:nth-child(17) > select",
        "Sim"
    ))

    safe_action(doc, "Selecionando a opção para o Pacote possuir Cremação", selecionar_opcao(
        "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_parametrosGerais.categoriaHolder > div.groupHolder.clearfix.grupo_parametrosGerais > div > div > div:nth-child(18) > select",
        "Sim"
    ))


    safe_action(doc, "Selecionando a opção para o Registro de Óbito ser obrigatório na Cremação", selecionar_opcao(
        "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_parametrosGerais.categoriaHolder > div.groupHolder.clearfix.grupo_parametrosGerais > div > div > div:nth-child(19) > select",
        "Sim"
    ))

    safe_action(doc, "Selecionando a opção para o Pacote lançar manutenção após o 1º Óbito", selecionar_opcao(
        "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_parametrosGerais.categoriaHolder > div.groupHolder.clearfix.grupo_parametrosGerais > div > div > div:nth-child(20) > select",
        "Sim"
    ))

    safe_action(doc, "Selecionando a opção para o Pacote Lançar Títulos no Cadastro de Contrato", selecionar_opcao(
        "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_parametrosGerais.categoriaHolder > div.groupHolder.clearfix.grupo_parametrosGerais > div > div > div:nth-child(22) > select",
        "Sim"
    ))


    safe_action(doc, "Rolando até o Select de Dependentes Contemplados", rolar_ate_dependentes_contemplados)

    safe_action(doc, "Selecionando a opção 'Agregado'", selecionar_opcao(
        "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_parametrosGerais.categoriaHolder > div.groupHolder.clearfix.grupo_dependentesContemplados > div.group_dependentesContemplados.clearfix.grupoHolder.lista > div > div > select",
        "Agregado"
    ))

    safe_action(doc, "Adicionado Dependente", rolar_ate_adicionar_formas_pagamento_mensalidade)



    safe_action(doc, "Adicionado Dependente", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_parametrosGerais.categoriaHolder > div.groupHolder.clearfix.grupo_dependentesContemplados > div.btnListHolder > a.btAddGroup"))
        ).click()
    ))



    safe_action(doc, "Acessando aba Taxas do Pacote", lambda: (
        wait.until(EC.element_to_be_clickable((By.LINK_TEXT, "Taxas do Pacote"))).click()
    ))


    safe_action(doc, "Preenchendo a Adesão Total do pacote", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_taxasPacote.categoriaHolder > div.groupHolder.clearfix.grupo_taxasAdesao > div > div > div:nth-child(1) > input"))
        ).send_keys(str(valor_taxa_adesao).replace('.', ','))
    ))


    safe_action(doc, "Selecionando Tipo de Mensalidade Adesão", abrir_modal_e_selecionar(
        "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_taxasPacote.categoriaHolder > div.groupHolder.clearfix.grupo_taxasAdesao > div > div > div:nth-child(2) > div > a",
        "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(2) > input",
        "TIPO MENSALIDADE JOÃO TESTE",
        "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(4) > a",
        "//td[contains(text(), 'TIPO MENSALIDADE JOÃO TESTE')]"
    ))



    safe_action(doc, "Selecionando a opção 'Por Periodicidade' no Tipo de Taxa para Manutenção ", selecionar_opcao(
        "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_taxasPacote.categoriaHolder > div.groupHolder.clearfix.grupo_taxasManutencao > div > div > div:nth-child(1) > select",
        "Por periodicidade"
    ))


    safe_action(doc, "Preenchendo a Taxa Mensal", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_taxasPacote.categoriaHolder > div.groupHolder.clearfix.grupo_taxasManutencao > div > div > div:nth-child(3) > input"))
        ).send_keys(fake.random_int(min=10, max=500))
    ))



    safe_action(doc, "Preenchendo a Taxa Bimestral", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_taxasPacote.categoriaHolder > div.groupHolder.clearfix.grupo_taxasManutencao > div > div > div:nth-child(4) > input"))
        ).send_keys(fake.random_int(min=1000, max=50000))
    ))



    safe_action(doc, "Preenchendo a Taxa Trimestral", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_taxasPacote.categoriaHolder > div.groupHolder.clearfix.grupo_taxasManutencao > div > div > div:nth-child(5) > input"))
        ).send_keys(fake.random_int(min=1000, max=50000))
    ))

    safe_action(doc, "Preenchendo a Taxa Quadrimestral", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_taxasPacote.categoriaHolder > div.groupHolder.clearfix.grupo_taxasManutencao > div > div > div:nth-child(6) > input"))
        ).send_keys(fake.random_int(min=1000, max=50000))
    ))


    safe_action(doc, "Preenchendo a Taxa Semestral", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_taxasPacote.categoriaHolder > div.groupHolder.clearfix.grupo_taxasManutencao > div > div > div:nth-child(7) > input"))
        ).send_keys(fake.random_int(min=1000, max=50000))
    ))

    safe_action(doc, "Preenchendo a Taxa Anual", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_taxasPacote.categoriaHolder > div.groupHolder.clearfix.grupo_taxasManutencao > div > div > div:nth-child(8) > input"))
        ).send_keys(fake.random_int(min=1000, max=50000))
    ))



    safe_action(doc, "Selecionando Tipo de Mensalidade Manutenção", abrir_modal_e_selecionar(
        "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_taxasPacote.categoriaHolder > div.groupHolder.clearfix.grupo_taxasManutencao > div > div > div:nth-child(9) > div > a",
        "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(2) > input",
        "TIPO MENSALIDADE MANUTENÇÃO",
        "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(4) > a",
        "//td[contains(text(), 'TIPO MENSALIDADE MANUTENÇÃO')]"
    ))


    safe_action(doc, "Preenchendo a o Valor da Transferência de Titularidade", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_taxasPacote.categoriaHolder > div.groupHolder.clearfix.grupo_outras > div > div > div:nth-child(1) > input"))
        ).send_keys(fake.random_int(min=1000, max=50000))
    ))

    safe_action(doc, "Preenchendo a o Valor requerido para Beneficiar pessoas fora do Contrato", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_taxasPacote.categoriaHolder > div.groupHolder.clearfix.grupo_outras > div > div > div:nth-child(3) > input"))
        ).send_keys(fake.random_int(min=1000, max=50000))
    ))




    safe_action(doc, "Acessando aba Formas de Pagamento", lambda: (
        wait.until(EC.element_to_be_clickable((By.LINK_TEXT, "Formas de Pagamento"))).click()
    ))


    safe_action(doc, "Selecionando a opção 'Boleto' em Formas de pagamento da Adesão", selecionar_opcao(
        "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_formasPagamento.categoriaHolder > div.groupHolder.clearfix.grupo_formaPagamentoAdesao > div.group_formaPagamentoAdesao.clearfix.grupoHolder.lista > div > div:nth-child(1) > select",
        "Boleto"
    ))

    safe_action(doc, "Preenchendo o campo de Parcelamento Máximo ", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_formasPagamento.categoriaHolder > div.groupHolder.clearfix.grupo_formaPagamentoAdesao > div.group_formaPagamentoAdesao.clearfix.grupoHolder.lista > div > div:nth-child(2) > input"))
        ).send_keys(fake.random_element([2, 3, 4, 5, 10, 6, 12, 24]))
    ))



    safe_action(doc, "Adicionando Forma de Pagamento", lambda: driver.find_element(
        By.CSS_SELECTOR, "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_formasPagamento.categoriaHolder > div.groupHolder.clearfix.grupo_formaPagamentoAdesao > div.btnListHolder > a.btAddGroup"
    ).click())



    safe_action(doc, "Selecionando a opção 'Cartão de Crédito' em Formas de pagamento da Adesão", selecionar_opcao(
        "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_formasPagamento.categoriaHolder > div.groupHolder.clearfix.grupo_formaPagamentoAdesao > div.group_formaPagamentoAdesao.clearfix.grupoHolder.lista > div > div:nth-child(1) > select",
        "Cartão de Crédito"
    ))

    safe_action(doc, "Preenchendo o campo de Parcelamento Máximo ", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_formasPagamento.categoriaHolder > div.groupHolder.clearfix.grupo_formaPagamentoAdesao > div.group_formaPagamentoAdesao.clearfix.grupoHolder.lista > div > div:nth-child(2) > input"))
        ).send_keys(fake.random_element([2, 3, 4, 5, 10, 6, 12, 24]))
    ))



    safe_action(doc, "Adicionando Forma de Pagamento", lambda: driver.find_element(
        By.CSS_SELECTOR, "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_formasPagamento.categoriaHolder > div.groupHolder.clearfix.grupo_formaPagamentoAdesao > div.btnListHolder > a.btAddGroup"
    ).click())


    safe_action(doc, "Selecionando a opção 'Cartão de Débito' em Formas de pagamento da Adesão", selecionar_opcao(
        "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_formasPagamento.categoriaHolder > div.groupHolder.clearfix.grupo_formaPagamentoAdesao > div.group_formaPagamentoAdesao.clearfix.grupoHolder.lista > div > div:nth-child(1) > select",
        "Cartão de Débito"
    ))

    safe_action(doc, "Preenchendo o campo de Parcelamento Máximo ", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_formasPagamento.categoriaHolder > div.groupHolder.clearfix.grupo_formaPagamentoAdesao > div.group_formaPagamentoAdesao.clearfix.grupoHolder.lista > div > div:nth-child(2) > input"))
        ).send_keys(fake.random_element([2, 3, 4, 5, 10, 6, 12, 24]))
    ))



    safe_action(doc, "Adicionando Forma de Pagamento", lambda: driver.find_element(
        By.CSS_SELECTOR, "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_formasPagamento.categoriaHolder > div.groupHolder.clearfix.grupo_formaPagamentoAdesao > div.btnListHolder > a.btAddGroup"
    ).click())



    safe_action(doc, "Selecionando a opção 'Cheque' em Formas de pagamento da Adesão", selecionar_opcao(
        "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_formasPagamento.categoriaHolder > div.groupHolder.clearfix.grupo_formaPagamentoAdesao > div.group_formaPagamentoAdesao.clearfix.grupoHolder.lista > div > div:nth-child(1) > select",
        "Cheque"
    ))

    safe_action(doc, "Preenchendo o campo de Parcelamento Máximo ", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_formasPagamento.categoriaHolder > div.groupHolder.clearfix.grupo_formaPagamentoAdesao > div.group_formaPagamentoAdesao.clearfix.grupoHolder.lista > div > div:nth-child(2) > input"))
        ).send_keys(fake.random_element([2, 3, 4, 5, 10, 6, 12, 24]))
    ))



    safe_action(doc, "Adicionando Forma de Pagamento", lambda: driver.find_element(
        By.CSS_SELECTOR, "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_formasPagamento.categoriaHolder > div.groupHolder.clearfix.grupo_formaPagamentoAdesao > div.btnListHolder > a.btAddGroup"
    ).click())



    safe_action(doc, "Selecionando a opção 'Depósito' em Formas de pagamento da Adesão", selecionar_opcao(
        "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_formasPagamento.categoriaHolder > div.groupHolder.clearfix.grupo_formaPagamentoAdesao > div.group_formaPagamentoAdesao.clearfix.grupoHolder.lista > div > div:nth-child(1) > select",
        "Depósito"
    ))

    safe_action(doc, "Preenchendo o campo de Parcelamento Máximo ", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_formasPagamento.categoriaHolder > div.groupHolder.clearfix.grupo_formaPagamentoAdesao > div.group_formaPagamentoAdesao.clearfix.grupoHolder.lista > div > div:nth-child(2) > input"))
        ).send_keys(fake.random_element([2, 3, 4, 5, 10, 6, 12, 24]))
    ))



    safe_action(doc, "Adicionando Forma de Pagamento", lambda: driver.find_element(
        By.CSS_SELECTOR, "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_formasPagamento.categoriaHolder > div.groupHolder.clearfix.grupo_formaPagamentoAdesao > div.btnListHolder > a.btAddGroup"
    ).click())


    safe_action(doc, "Selecionando a opção 'Dinheiro' em Formas de pagamento da Adesão", selecionar_opcao(
        "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_formasPagamento.categoriaHolder > div.groupHolder.clearfix.grupo_formaPagamentoAdesao > div.group_formaPagamentoAdesao.clearfix.grupoHolder.lista > div > div:nth-child(1) > select",
        "Dinheiro"
    ))

    safe_action(doc, "Preenchendo o campo de Parcelamento Máximo ", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_formasPagamento.categoriaHolder > div.groupHolder.clearfix.grupo_formaPagamentoAdesao > div.group_formaPagamentoAdesao.clearfix.grupoHolder.lista > div > div:nth-child(2) > input"))
        ).send_keys(fake.random_element([2, 3, 4, 5, 10, 6, 12, 24]))
    ))



    safe_action(doc, "Adicionando Forma de Pagamento", lambda: driver.find_element(
        By.CSS_SELECTOR, "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_formasPagamento.categoriaHolder > div.groupHolder.clearfix.grupo_formaPagamentoAdesao > div.btnListHolder > a.btAddGroup"
    ).click())


    safe_action(doc, "Selecionando a opção 'Débito em Conta' em Formas de pagamento da Adesão", selecionar_opcao(
        "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_formasPagamento.categoriaHolder > div.groupHolder.clearfix.grupo_formaPagamentoAdesao > div.group_formaPagamentoAdesao.clearfix.grupoHolder.lista > div > div:nth-child(1) > select",
        "Débito em Conta"
    ))

    safe_action(doc, "Preenchendo o campo de Parcelamento Máximo ", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_formasPagamento.categoriaHolder > div.groupHolder.clearfix.grupo_formaPagamentoAdesao > div.group_formaPagamentoAdesao.clearfix.grupoHolder.lista > div > div:nth-child(2) > input"))
        ).send_keys(fake.random_element([2, 3, 4, 5, 10, 6, 12, 24]))
    ))



    safe_action(doc, "Adicionando Forma de Pagamento", lambda: driver.find_element(
        By.CSS_SELECTOR, "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_formasPagamento.categoriaHolder > div.groupHolder.clearfix.grupo_formaPagamentoAdesao > div.btnListHolder > a.btAddGroup"
    ).click())


    safe_action(doc, "Selecionando a opção 'Transferência Bancária' em Formas de pagamento da Adesão", selecionar_opcao(
        "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_formasPagamento.categoriaHolder > div.groupHolder.clearfix.grupo_formaPagamentoAdesao > div.group_formaPagamentoAdesao.clearfix.grupoHolder.lista > div > div:nth-child(1) > select",
        "Transferência Bancária"
    ))

    safe_action(doc, "Preenchendo o campo de Parcelamento Máximo ", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_formasPagamento.categoriaHolder > div.groupHolder.clearfix.grupo_formaPagamentoAdesao > div.group_formaPagamentoAdesao.clearfix.grupoHolder.lista > div > div:nth-child(2) > input"))
        ).send_keys(fake.random_element([2, 3, 4, 5, 10, 6, 12, 24]))
    ))



    safe_action(doc, "Adicionando Forma de Pagamento", lambda: driver.find_element(
        By.CSS_SELECTOR, "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_formasPagamento.categoriaHolder > div.groupHolder.clearfix.grupo_formaPagamentoAdesao > div.btnListHolder > a.btAddGroup"
    ).click())



    safe_action(doc, "Rolando até o Select de Formas de Pagamanto da Mensalidade", rolar_ate_formas_pagamento_mensalidade)



    safe_action(doc, "Selecionando a opção 'Boleto' em Formas de pagamento da Mensalidade", selecionar_opcao(
        "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_formasPagamento.categoriaHolder > div.groupHolder.clearfix.grupo_formaPagamentoMensalidade > div.group_formaPagamentoMensalidade.clearfix.grupoHolder.lista > div > div:nth-child(1) > select",
        "Boleto"
    ))

    safe_action(doc, "Preenchendo o campo de Parcelamento Máximo ", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_formasPagamento.categoriaHolder > div.groupHolder.clearfix.grupo_formaPagamentoMensalidade > div.group_formaPagamentoMensalidade.clearfix.grupoHolder.lista > div > div:nth-child(2) > input"))
        ).send_keys(fake.random_element([2, 3, 4, 5, 10, 6, 12, 24]))
    ))

    safe_action(doc, "Adicionando Forma de Pagamento", rolar_ate_adicionar_formas_pagamento_mensalidade)


    safe_action(doc, "Selecionando a opção 'Carnê' em Formas de pagamento da Mensalidade", selecionar_opcao(
        "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_formasPagamento.categoriaHolder > div.groupHolder.clearfix.grupo_formaPagamentoMensalidade > div.group_formaPagamentoMensalidade.clearfix.grupoHolder.lista > div > div:nth-child(1) > select",
        "Carnê"
    ))

    safe_action(doc, "Preenchendo o campo de Parcelamento Máximo ", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_formasPagamento.categoriaHolder > div.groupHolder.clearfix.grupo_formaPagamentoMensalidade > div.group_formaPagamentoMensalidade.clearfix.grupoHolder.lista > div > div:nth-child(2) > input"))
        ).send_keys(fake.random_element([2, 3, 4, 5, 10, 6, 12, 24]))
    ))


    safe_action(doc, "Adicionando Forma de Pagamento", rolar_ate_adicionar_formas_pagamento_mensalidade)



    safe_action(doc, "Selecionando a opção 'Pix' em Formas de pagamento da Mensalidade", selecionar_opcao(
        "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_formasPagamento.categoriaHolder > div.groupHolder.clearfix.grupo_formaPagamentoMensalidade > div.group_formaPagamentoMensalidade.clearfix.grupoHolder.lista > div > div:nth-child(1) > select",
        "Pix"
    ))

    safe_action(doc, "Preenchendo o campo de Parcelamento Máximo ", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_formasPagamento.categoriaHolder > div.groupHolder.clearfix.grupo_formaPagamentoMensalidade > div.group_formaPagamentoMensalidade.clearfix.grupoHolder.lista > div > div:nth-child(2) > input"))
        ).send_keys(fake.random_element([2, 3, 4, 5, 10, 6, 12, 24]))
    ))



    safe_action(doc, "Adicionando Forma de Pagamento", rolar_ate_adicionar_formas_pagamento_mensalidade)

    safe_action(doc, "Acessando aba Produtos", lambda: (
        wait.until(EC.element_to_be_clickable((By.LINK_TEXT, "Produtos"))).click()
    ))


    safe_action(doc, "Selecionando o Produto PRODUTO TESTE SELENIUM", abrir_modal_e_selecionar(
        "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_produtos.categoriaHolder > div > div.group_produtos.clearfix.grupoHolder.lista > div > div:nth-child(1) > div > a",
        "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(2) > input",
        "PRODUTO TESTE SELENIUM",
        "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(4) > a",
        "//td[contains(text(), 'PRODUTO TESTE SELENIUM')]"
    ))



    safe_action(doc, "Preenchendo a quantidade do produto", lambda: (
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_produtos.categoriaHolder > div > div.group_produtos.clearfix.grupoHolder.lista > div > div:nth-child(2) > input"))
        ).send_keys(fake.random_choices(elements=[1,2,3,4,5,6,7,8,9])[0])
    ))


    safe_action(doc, "Adicionando Produto", lambda: (
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_produtos.categoriaHolder > div > div.btnListHolder > a.btAddGroup"))).click()
    ))


    safe_action(doc, "Acessando aba Serviços", lambda: (
        wait.until(EC.element_to_be_clickable((By.LINK_TEXT, "Serviços"))).click()
    ))


    safe_action(doc, "Selecionando o Serviço SERVIÇO TESTE SELENIUM", abrir_modal_e_selecionar(
        "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_servicos.categoriaHolder > div > div.group_servicos.clearfix.grupoHolder.lista > div > div:nth-child(1) > div > a",
        "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(2) > input",
        "SERVIÇO TESTE SELENIUM",
        "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(4) > a",
        "//td[contains(text(), 'SERVIÇO TESTE SELENIUM')]"
    ))


    safe_action(doc, "Adicionando Serviço", lambda: (
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_servicos.categoriaHolder > div > div.btnListHolder > a.btAddGroup"))).click()
    ))




    safe_action(doc, "Acessando aba Inadimplência", lambda: (
        wait.until(EC.element_to_be_clickable((By.LINK_TEXT, "Inadimplência"))).click()
    ))


    safe_action(doc, "Preenchendo a Quantidade de parcelas em atraso para o Bloqueio da Ordem de Serviço ", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_inadimplencia.categoriaHolder > div.groupHolder.clearfix.grupo_inadimplenciaOS > div > div > div:nth-child(2) > input"))
        ).send_keys(qtd_parcelas_em_atraso)
    ))

    safe_action(doc, "Preenchendo a Carência (em dias) após pagamento ", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_inadimplencia.categoriaHolder > div.groupHolder.clearfix.grupo_inadimplenciaOS > div > div > div:nth-child(3) > input"))
        ).send_keys(fake.random_int(min=0, max=15))
    ))

    safe_action(doc, "Preenchendo a Quantidade de parcelas em atraso para o Bloqueio da Locação de Equipamento ", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_inadimplencia.categoriaHolder > div.groupHolder.clearfix.grupo_inadimplenciaLocacaoEquipamento > div > div > div:nth-child(2) > input"))
        ).send_keys(qtd_parcelas_em_atraso)
    ))

    safe_action(doc, "Acessando aba Carência Dependente", lambda: (
        wait.until(EC.element_to_be_clickable((By.LINK_TEXT, "Carência Dependente"))).click()
    ))



    safe_action(doc, "Preenchendo o campo 'Idade de'", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_carenciaDependente.categoriaHolder > div > div.group_carenciaDependente.clearfix.grupoHolder.lista > div > div:nth-child(1) > input"))
        ).send_keys(idade_de)
    ))

    safe_action(doc, "Preenchendo o campo 'Idade até'", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_carenciaDependente.categoriaHolder > div > div.group_carenciaDependente.clearfix.grupoHolder.lista > div > div:nth-child(2) > input"))
        ).send_keys(idade_até)

    ))



    safe_action(doc, "Preenchendo a Carência (em dias) para utilização ", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_carenciaDependente.categoriaHolder > div > div.group_carenciaDependente.clearfix.grupoHolder.lista > div > div:nth-child(3) > input"))
        ).send_keys(fake.random_int(min=0, max=15))
    ))

    safe_action(doc, "Adicionando parâmtros da Carência Dependente", lambda: (
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_carenciaDependente.categoriaHolder > div > div.btnListHolder > a.btAddGroup"))).click()
    ))


    safe_action(doc, "Acessando aba Alteração Automática de Status do Contrato", lambda: (
        wait.until(EC.element_to_be_clickable((By.LINK_TEXT, "Alteração Automática de Status do Contrato"))).click()
    ))

    safe_action(doc, "Preenchendo a Quantidade parcelas vencidas para salvar ", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_alteracaoStatusContrato.categoriaHolder > div.groupHolder.clearfix.grupo_parametros > div > div > div:nth-child(1) > input"))
        ).send_keys(fake.random_int(min=0, max=15))
    ))


    safe_action(doc, "Selecionando o Motivo Salvamento CLIENTE SOLICITOU CANCELAMENTO", abrir_modal_e_selecionar(
        "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_alteracaoStatusContrato.categoriaHolder > div.groupHolder.clearfix.grupo_parametros > div > div > div:nth-child(2) > div > a",
        "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(2) > input",
        "CLIENTE SOLICITOU CANCELAMENTO",
        "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(4) > a",
        "//td[contains(text(), 'CLIENTE SOLICITOU CANCELAMENTO')]"
    ))


    safe_action(doc, "Preenchendo a Quantidade parcelas vencidas para suspender ", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_alteracaoStatusContrato.categoriaHolder > div.groupHolder.clearfix.grupo_parametros > div > div > div:nth-child(3) > input"))
        ).send_keys(fake.random_int(min=0, max=15))
    ))


    safe_action(doc, "Selecionando o Motivo Suspensão SEM COMUNICAÇÃO", abrir_modal_e_selecionar(
        "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_alteracaoStatusContrato.categoriaHolder > div.groupHolder.clearfix.grupo_parametros > div > div > div:nth-child(4) > div > a",
        "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(2) > input",
        "SEM COMUNICAÇÃO",
        "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(4) > a",
        "//td[contains(text(), 'SEM COMUNICAÇÃO')]"
    ))

    safe_action(doc, "Preenchendo a Quantidade parcelas vencidas para reativar ", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_alteracaoStatusContrato.categoriaHolder > div.groupHolder.clearfix.grupo_parametros > div > div > div:nth-child(5) > input"))
        ).send_keys(fake.random_int(min=0, max=15))
    ))


    safe_action(doc, "Selecionando o Motivo Reativação PAGOU AS PARCELAS", abrir_modal_e_selecionar(
        "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_alteracaoStatusContrato.categoriaHolder > div.groupHolder.clearfix.grupo_parametros > div > div > div:nth-child(6) > div > a",
        "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(2) > input",
        "PAGOU AS PARCELAS",
        "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(4) > a",
        "//td[contains(text(), 'PAGOU AS PARCELAS')]"
    ))

    safe_action(doc, "Selecionando a opção 'Sim' para Salvar a Recorrência Vindi na Alteração Automática de Status do Contrato", selecionar_opcao(
        "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_alteracaoStatusContrato.categoriaHolder > div.groupHolder.clearfix.grupo_parametros > div > div > div:nth-child(7) > select",
        "Sim"
    ))

    safe_action(doc, "Selecionando o Motivo Salvamento Recorrência CANCELAMENTO POR FALTA DE PAGAMENTO", abrir_modal_e_selecionar(
        "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_alteracaoStatusContrato.categoriaHolder > div.groupHolder.clearfix.grupo_parametros > div > div > div:nth-child(8) > div > a",
        "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(2) > input",
        "CANCELAMENTO POR FALTA DE PAGAMENTO",
        "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(4) > a",
        "//td[contains(text(), 'CANCELAMENTO POR FALTA DE PAGAMENTO')]"
    ))


    safe_action(doc, "Selecionando o Tipo de Mensalidade TIPO MENSALIDADE JOÃO TESTE", abrir_modal_e_selecionar(
        "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_alteracaoStatusContrato.categoriaHolder > div.groupHolder.clearfix.grupo_tipoMensalidade > div.group_tipoMensalidade.clearfix.grupoHolder.lista > div > div > div > a",
        "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(2) > input",
        "TIPO MENSALIDADE JOÃO TESTE",
        "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(4) > a",
        "//td[contains(text(), 'TIPO MENSALIDADE JOÃO TESTE')]"
    ))



    safe_action(doc, "Adicionando Tipo de Mensalidade ", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.catWrapper > div > div.cat_alteracaoStatusContrato.categoriaHolder > div.groupHolder.clearfix.grupo_tipoMensalidade > div.btnListHolder > a.btAddGroup"))
        ).click()
    ))
    
    safe_action(doc, "Salvando cadastro", lambda: driver.find_element(
        By.CSS_SELECTOR, "#fmod_2 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroPacote > div.btnHolder > a.btModel.btGray.btsave"
    ).click())





    safe_action(doc, "Fechando modal após salvamento", lambda: wait.until(EC.element_to_be_clickable((
        By.CSS_SELECTOR, "#fmod_2 > div.wdTop.ui-draggable-handle > div.wdClose > a"
        ))
    ).click())



    encontrar_mensagem_alerta()

except Exception as e:
    log(doc, f"❌ ERRO FATAL: {e}")
    take_screenshot(driver, doc, "erro_fatal")

finally:

    log(doc, "✅ Teste concluído com sucesso.")

    finalizar_relatorio()