# Refatorado e organizado: cadastrodeescalamotorista2ºcenario.py

from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support.ui import Select, WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service
from webdriver_manager.chrome import ChromeDriverManager
from docx import Document
from docx.shared import Inches
from faker import Faker
from faker.providers import BaseProvider
from validate_docbr import CPF
from datetime import datetime, timedelta
import subprocess
import os
import time
import random
import sys 

sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))

# ==== CONFIGURAÇÕES ====
URL = "http://localhost:8080/gs/index.xhtml"
LOGIN_EMAIL = "joaoeduardo.gold@outlook.com"
LOGIN_PASSWORD = "071999gs"

# ==== PROVIDERS CUSTOMIZADOS ====
class BrasilProvider(BaseProvider):
    def rg(self):
        numeros = [str(random.randint(0, 9)) for _ in range(8)]
        return ''.join(numeros) + '-' + str(random.randint(0, 9))

fake = Faker("pt_BR")
fake.add_provider(BrasilProvider)

def gerar_datas_validas():
    """Gera datas coerentes para admissão, início e fim da escala, e vencimento da CNH."""
    hoje = datetime.today().date()
    
    # Data de admissão entre 10 anos atrás e hoje
    data_admissao = fake.date_between(start_date=hoje - timedelta(days=3650), end_date=hoje)
    
    # Data de início da escala entre hoje e 1 ano no futuro
    data_inicio = fake.date_between(start_date=hoje, end_date=hoje + timedelta(days=365))
    
    # Data fim entre 1 e 180 dias após a data de início
    data_fim = data_inicio + timedelta(days=random.randint(1, 180))
    
    # Vencimento CNH entre hoje e 10 anos no futuro
    vencimento_cnh = fake.date_between(start_date=hoje, end_date=hoje + timedelta(days=3650))
    
    return (data_admissao.strftime('%d/%m/%Y'), 
            data_inicio.strftime('%d/%m/%Y'), 
            data_fim.strftime('%d/%m/%Y'), 
            vencimento_cnh.strftime('%d/%m/%Y'))

def gerar_dados_documentos():
    """Gera documentos fictícios para o cadastro."""
    carteira_trabalho = str(random.randint(10000000, 99999999))
    pis = fake.cpf().replace('.', '').replace('-', '')[:11]
    cnh = str(random.randint(10000000000, 99999999999))
    cpf = CPF().generate()
    
    return carteira_trabalho, pis, cnh, cpf

# Controle de screenshots únicas
screenshot_registradas = set()

# ==== FUNÇÕES DE UTILITÁRIO ====
def log(doc, msg):
    print(msg)
    doc.add_paragraph(msg)

def take_screenshot(driver, doc, nome):
    if nome not in screenshot_registradas:
        path = f"screenshots/{nome}.png"
        os.makedirs("screenshots", exist_ok=True)
        driver.save_screenshot(path)
        doc.add_paragraph(f"Screenshot: {nome}")
        doc.add_picture(path, width=Inches(5.5))
        screenshot_registradas.add(nome)

def safe_action(doc, descricao, func, driver, wait):
    try:
        log(doc, f"🔄 {descricao}...")
        func()
        log(doc, f"✅ {descricao} realizada com sucesso.")
        take_screenshot(driver, doc, descricao.lower().replace(" ", "_"))
        return True, "sucesso"
    except Exception as e:
        log(doc, f"❌ Erro ao {descricao.lower()}: {e}")
        take_screenshot(driver, doc, f"erro_{descricao.lower().replace(' ', '_')}")
        return False, "erro"

def registrar_screenshot_unico(nome, driver, doc, descricao=None):
    if nome not in screenshot_registradas:
        if descricao:
            log(doc, f"📸 {descricao}")
        take_screenshot(driver, doc, nome)
        screenshot_registradas.add(nome)

def encontrar_mensagem_alerta(driver, doc):
    seletores = [
        (".alerts.salvo", "✅ Menasagem de Sucesso"),
        (".alerts.alerta", "⚠️ Menasagem de Alerta"),
        (".alerts.erro", "❌ Menasagem de Erro"),
    ]

    for seletor, tipo in seletores:
        try:
            elemento = driver.find_element(By.CSS_SELECTOR, seletor)
            if elemento.is_displayed():
                log(doc, f"📢 {tipo.upper()}: {elemento.text}")
                return elemento, tipo
        except:
            continue

    log(doc, "ℹ️ Nenhuma mensagem de alerta encontrada.")
    return None, None


def main():
    global screenshot_registradas
    screenshot_registradas = set()
    
    doc = Document()
    doc.add_heading("RELATÓRIO DO TESTE", 0)
    doc.add_paragraph("Cadastro de Controle de Cremação Teste.")
    doc.add_paragraph(f"🗕️ Data do teste: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")
    doc.add_paragraph("Neste teste, o robô preencherá APENAS os campos NÃO obrigatórios e salvará o cadastro de um novo Controle de Cremação.")

    chrome_options = Options()
    chrome_options.add_argument("--start-maximized")
    driver = webdriver.Chrome(service=Service(ChromeDriverManager().install()), options=chrome_options)
    wait = WebDriverWait(driver, 10)

    def ajustar_zoom(driver):
        try:
            driver.execute_script("document.body.style.zoom='90%'")
            log(doc, "🔍 Zoom ajustado para 90%.")
        except Exception as e:
            log(doc, f"⚠️ Erro ao ajustar zoom: {e}")


    def finalizar_relatorio():
        doc_name = f"relatorio_controle_cremacao_cenario_4_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        doc.save(doc_name)
        log(doc, f"📄 Relatório salvo como: {doc_name}")
        try:
            subprocess.run(["start", "winword", doc_name], shell=True)
        except Exception as e:
            log(doc, f"Erro ao abrir o Word: {e}")
        driver.quit()

    def login():
        wait.until(EC.presence_of_element_located((By.ID, "j_id15:email"))).send_keys(LOGIN_EMAIL)
        wait.until(EC.presence_of_element_located((By.ID, "j_id15:senha"))).send_keys(LOGIN_PASSWORD, Keys.ENTER)
        wait.until(EC.presence_of_element_located((By.TAG_NAME, "body")))
        time.sleep(4)

    def abrir_menu():
        driver.find_element(By.TAG_NAME, "body").send_keys(Keys.F2)
        campo = wait.until(EC.presence_of_element_located((By.XPATH, "//input[@placeholder='Busque um cadastro']")))
        campo.click()
        campo.send_keys("Controle Cremação", Keys.ENTER)
        time.sleep(3)

    def acessar_formulario():
        cadastrar = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, 
            "#fmod_200010 > div.wdTelas > div.telaInicial.clearfix.overflow.overflowY > ul > li:nth-child(1) > a > span")))
        cadastrar.click()
        time.sleep(2)

    def abrir_lov_falecido():
        log(doc, "🔄 Abrindo LOV de Falecido.")
        open_lov_falecido = wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, 
            "#fmod_200010 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div > div > div > div:nth-child(2) > div > a")))
        open_lov_falecido.click()
        log(doc, "✅ LOV de Falecido aberto.")

    def pesquisar_e_selecionar_falecido():
        log(doc, "🔄 Pesquisando e selecionando falecido.")
        campo_pesquisa = wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#txtPesquisa")))
        campo_pesquisa.send_keys('GUSTAVO VIEIRA', Keys.ENTER)
        time.sleep(2)
        falecido = wait.until(EC.element_to_be_clickable((By.XPATH, "//td[contains(text(), 'GUSTAVO VIEIRA')]")))
        falecido.click()
        log(doc, "✅ Falecido selecionado.")

    def preencher_numero_os():
        log(doc, "🔄 Preenchendo campo 'Número OS'.")
        numero_OS = wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, 
            "#fmod_200010 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div > div > div > div:nth-child(3) > input")))
        numero_OS.send_keys('201680')
        log(doc, "✅ Campo 'Número OS' preenchido.")

    def selecionar_status_aguardando():
        log(doc, "🔄 Selecionando status 'Aguardando'.")
        select_aguardando = wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR,
            "#fmod_200010 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div > div > div > div:nth-child(4) > select")))
        Select(select_aguardando).select_by_visible_text("Aguardando")
        log(doc, "✅ Status 'Aguardando' selecionado.")

    def selecionar_forno():
        log(doc, "🔄 Selecionando 'Forno 1'.")
        select_forno = wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR,
            "#fmod_200010 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div > div > div > div:nth-child(5) > select")))
        Select(select_forno).select_by_visible_text("Forno 1")
        log(doc, "✅ 'Forno 1' selecionado.")

    def salvar():
        salvar_btn = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR,
            "#fmod_200010 > div.wdTelas > div.telaCadastro.clearfix > div.btnHolder > a.btModel.btGray.btsave")))
        salvar_btn.click()
        time.sleep(2)

    def fechar_modal():
        x_btn = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR,
            "#fmod_200010 > div.wdTop.ui-draggable-handle > div.wdClose > a")))
        x_btn.click()
        time.sleep(1)

    # EXECUÇÃO COM safe_action INDIVIDUAL PARA CADA AÇÃO
    if not safe_action(doc, "Acessando o sistema", lambda: driver.get(URL), driver, wait)[0]:
        finalizar_relatorio()
        return
    registrar_screenshot_unico("url_acessada", driver, doc, "Sistema acessado.")

    if not safe_action(doc, "Realizando login", login, driver, wait)[0]:
        finalizar_relatorio()
        return
    registrar_screenshot_unico("login_concluido", driver, doc, "Login realizado.")

    if not safe_action(doc, "Ajustando zoom", lambda: ajustar_zoom(driver), driver, wait)[0]:
        finalizar_relatorio()
        return

    if not safe_action(doc, "Abrindo menu Controle Cremação", abrir_menu, driver, wait)[0]:
        finalizar_relatorio()
        return
    registrar_screenshot_unico("menu_aberto", driver, doc, "Menu 'Controle Cremação' aberto.")

    if not safe_action(doc, "Acessando formulário de cadastro", acessar_formulario, driver, wait)[0]:
        finalizar_relatorio()
        return
    registrar_screenshot_unico("formulario_aberto", driver, doc, "Formulário de cadastro aberto.")

    # PREENCHIMENTO DOS CAMPOS - safe_action individual para cada campo
    if not safe_action(doc, "Abrindo LOV de Falecido", abrir_lov_falecido, driver, wait)[0]:
        finalizar_relatorio()
        return

    if not safe_action(doc, "Pesquisando e selecionando falecido", pesquisar_e_selecionar_falecido, driver, wait)[0]:
        finalizar_relatorio()
        return



    if not safe_action(doc, "Selecionando status Aguardando", selecionar_status_aguardando, driver, wait)[0]:
        finalizar_relatorio()
        return

    if not safe_action(doc, "Selecionando Forno", selecionar_forno, driver, wait)[0]:
        finalizar_relatorio()
        return

    registrar_screenshot_unico("campos_preenchidos", driver, doc, "Todos os campos preenchidos.")

    # SALVANDO O CADASTRO
    if not safe_action(doc, "Clicando no botão Salvar", salvar, driver, wait)[0]:
        finalizar_relatorio()
        return
    registrar_screenshot_unico("apos_salvar", driver, doc, "Clique no botão Salvar realizado.")

    # VERIFICANDO MENSAGEM DE RETORNO
    _, tipo_alerta = encontrar_mensagem_alerta(driver, doc)
    if tipo_alerta == "sucesso":
        log(doc, "✅ Mensagem de sucesso exibida.")
    elif tipo_alerta == "alerta":
        log(doc, "⚠️ Mensagem de alerta exibida.")
    elif tipo_alerta == "erro":
        log(doc, "❌ Mensagem de erro exibida.")
    else:
        log(doc, "⚠️ Nenhuma mensagem encontrada após salvar.")
    registrar_screenshot_unico("mensagem_final", driver, doc, "Mensagem exibida após salvar.")

    # FECHANDO O FORMULÁRIO
    if not safe_action(doc, "Fechando formulário", fechar_modal, driver, wait)[0]:
        finalizar_relatorio()
        return
    registrar_screenshot_unico("formulario_fechado", driver, doc, "Formulário fechado.")

    log(doc, "✅ Teste de cadastro de Controle de Cremação concluído com sucesso.")
    finalizar_relatorio()

if __name__ == "__main__":
    main()