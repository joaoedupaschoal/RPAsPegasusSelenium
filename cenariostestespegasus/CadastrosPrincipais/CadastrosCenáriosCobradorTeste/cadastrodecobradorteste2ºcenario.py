from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service
from webdriver_manager.chrome import ChromeDriverManager
from docx import Document
from docx.shared import Inches
from datetime import datetime
from selenium.webdriver import ActionChains
import subprocess
import os
import sys
import time
import random
import re

# ==========================================
#  AJUSTES PRINCIPAIS NESTE ARQUIVO
#  - Corrigido TypeError do safe_action (assinaturas unificadas)
#  - Corrigido uso de encontrar_mensagem_alerta (passa doc e retorna tipo)
#  - Removidos parâmetros extras nas chamadas de safe_action
#  - Corrigida chamada de ajustar_zoom (sem parâmetros)
#  - Evitado duplicidade de constantes e funções
#  - Mantida a lógica do cenário (preencher 'Cobrador Teste' e salvar)
# ==========================================

# ==== CONFIGURAÇÕES GERAIS ====
URL = "http://localhost:8080/gs/index.xhtml"
LOGIN_EMAIL = "joaoeduardo.gold@outlook.com"
LOGIN_PASSWORD = "071999gs"

# ==== DOCX E DRIVER GLOBAIS (instanciados no main) ====
doc = None
driver = None
wait = None

# ==== SUPORTE: LOG / SCREENSHOT ====
screenshot_registradas = set()

def log(doc: Document, msg: str) -> None:
    print(msg)
    doc.add_paragraph(msg)


def sanitize_filename(nome: str) -> str:
    nome = nome.lower()
    nome = re.sub(r"[^a-z0-9_]", "_", nome)
    return nome


def take_screenshot(driver: webdriver.Chrome, doc: Document, nome: str) -> None:
    nome_limpo = sanitize_filename(nome)
    os.makedirs("screenshots", exist_ok=True)
    path = f"screenshots/{nome_limpo}.png"
    driver.save_screenshot(path)
    doc.add_paragraph(f"Screenshot: {nome}")
    doc.add_picture(path, width=Inches(5.5))


def registrar_screenshot_unico(nome: str, driver: webdriver.Chrome, doc: Document, descricao: str | None = None) -> None:
    if nome not in screenshot_registradas:
        if descricao:
            log(doc, f"📸 {descricao}")
        take_screenshot(driver, doc, nome)
        screenshot_registradas.add(nome)

# ==== SAFE ACTION PADRÃO (RETORNA TUPLA) ====
# Uso: ok, err = safe_action(doc, "descrição", lambda: ...)
# Todas as chamadas devem seguir este padrão.
def safe_action(doc: Document, descricao: str, func) -> tuple[bool, str | None]:
    global driver
    try:
        log(doc, f"🔄 {descricao}...")
        func()
        log(doc, f"✅ {descricao} realizada com sucesso.")
        take_screenshot(driver, doc, descricao.lower().replace(" ", "_"))
        return True, None
    except Exception as e:
        err = str(e)
        log(doc, f"❌ Erro ao {descricao.lower()}: {err}")
        take_screenshot(driver, doc, f"erro_{descricao.lower().replace(' ', '_')}")
        return False, err

# ==== OUTRAS FUNÇÕES DE UTILIDADE ====

def ajustar_zoom() -> None:
    global driver
    driver.execute_script("document.body.style.zoom='90%'")
    # NÃO precisa de screenshot aqui; safe_action já tira.


def encontrar_mensagem_alerta(doc: Document) -> tuple[object | None, str | None]:
    """
    Procura mensagens de feedback na tela.
    Retorna (elemento_encontrado, tipo) onde tipo ∈ {"sucesso", "alerta", "erro"} ou None.
    """
    global driver
    mapeamento = [
        (".alerts.salvo", "sucesso"),
        (".alerts.alerta", "alerta"),
        (".alerts.erro", "erro"),
    ]
    for seletor, tipo in mapeamento:
        try:
            elemento = driver.find_element(By.CSS_SELECTOR, seletor)
            if elemento.is_displayed():
                log(doc, f"📢 Mensagem de {tipo}: {elemento.text}")
                return elemento, tipo
        except Exception:
            continue
    log(doc, "ℹ️ Nenhuma mensagem de alerta encontrada.")
    return None, None

# ==== FLUXO DE TELA ====

def login() -> None:
    global wait
    wait.until(EC.presence_of_element_located((By.ID, "j_id15:email"))).send_keys(LOGIN_EMAIL)
    wait.until(EC.presence_of_element_located((By.ID, "j_id15:senha"))).send_keys(LOGIN_PASSWORD, Keys.ENTER)
    wait.until(EC.presence_of_element_located((By.TAG_NAME, "body")))
    time.sleep(2)


def abrir_menu() -> None:
    global driver, wait
    driver.find_element(By.TAG_NAME, "body").send_keys(Keys.F2)
    campo = wait.until(EC.presence_of_element_located((By.XPATH, "//input[@placeholder='Busque um cadastro']")))
    campo.click()
    campo.send_keys("Cobrado teste", Keys.ENTER)
    time.sleep(2)


def acessar_formulario() -> None:
    global wait
    wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR,
        "#fmod_200022 > div.wdTelas > div.telaInicial.clearfix.overflow.overflowY > ul > li:nth-child(1) > a > span"
    ))).click()
    time.sleep(2)


def preencher_campos(doc: Document) -> None:
    global wait
    log(doc, "🔄 Preenchendo o campo 'Nome do Cobrador'.")
    campo = wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR,
        "#fmod_200022 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div > div > div > div > textarea"
    )))
    ActionChains(driver).move_to_element(campo).click().perform()
    campo.clear()
    campo.send_keys("TESTE COBRADOR SELENIUM AUTOMATIZADO")
    log(doc, "✅ Campo 'Nome do Cobrador' preenchido.")
    registrar_screenshot_unico("cobrador_preenchido", driver, doc)


def cancelar() -> None:
    global wait
    wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR,
        "#fmod_200022 > div.wdTelas > div.telaCadastro.clearfix > div.btnHolder > a.btModel.btGray.btcancel"
    ))).click()
    time.sleep(2)


def confirmar_cancelamento() -> None:
    global wait
    wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR,
        "#BtYes"
    ))).click()
    time.sleep(2)


def fechar_modal() -> None:
    global wait
    wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR,
        "#fmod_200022 > div.wdTop.ui-draggable-handle > div.wdClose > a"
    ))).click()
    time.sleep(1)

# ==== MAIN ====

def main() -> None:
    global doc, driver, wait

    # Documento
    doc = Document()
    doc.add_heading("RELATÓRIO DO TESTE", 0)
    doc.add_paragraph("Cadastro de Cobrador Teste.")
    doc.add_paragraph(f"🗕️ Data do teste: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")
    doc.add_paragraph("Neste teste, o robô preencherá os campos obrigatórios e cancelará o cadastro de um novo Cobrador.")

    # Driver
    chrome_options = Options()
    chrome_options.add_argument("--start-maximized")
    driver = webdriver.Chrome(service=Service(ChromeDriverManager().install()), options=chrome_options)
    wait = WebDriverWait(driver, 10)

    def finalizar_relatorio():
        doc_name = f"relatorio_cobrador_teste_cenario_2_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        doc.save(doc_name)
        log(doc, f"📄 Relatório salvo como: {doc_name}")
        try:
            subprocess.run(["start", "winword", doc_name], shell=True)
        except Exception as e:
            log(doc, f"Erro ao abrir o Word: {e}")
        driver.quit()

    # Execução segura passo a passo
    ok, _ = safe_action(doc, "Acessando o sistema", lambda: driver.get(URL))
    if not ok:
        finalizar_relatorio(); return
    registrar_screenshot_unico("url_acessada", driver, doc, "Sistema acessado.")

    ok, _ = safe_action(doc, "Realizando login", login)
    if not ok:
        finalizar_relatorio(); return
    registrar_screenshot_unico("login_concluido", driver, doc, "Login realizado.")

    ok, _ = safe_action(doc, "Ajustando zoom", ajustar_zoom)
    if not ok:
        finalizar_relatorio(); return

    ok, _ = safe_action(doc, "Abrindo menu Cobrado Teste", abrir_menu)
    if not ok:
        finalizar_relatorio(); return
    registrar_screenshot_unico("menu_aberto", driver, doc, "Menu 'Cobrado Teste' aberto.")

    ok, _ = safe_action(doc, "Acessando formulário", acessar_formulario)
    if not ok:
        finalizar_relatorio(); return
    registrar_screenshot_unico("formulario_aberto", driver, doc, "Formulário de cadastro aberto.")

    ok, _ = safe_action(doc, "Preenchendo campos", lambda: preencher_campos(doc))
    if not ok:
        finalizar_relatorio(); return

    ok, _ = safe_action(doc, "Clicando em Cancelar", cancelar)
    if not ok:
        finalizar_relatorio(); return
    registrar_screenshot_unico("apos_cancelar", driver, doc, "Clique no botão Cancelar realizado.")

    ok, _ = safe_action(doc, "Confirmando cancelamento", confirmar_cancelamento)
    if not ok:
        finalizar_relatorio(); return
    registrar_screenshot_unico("confirmando_cancelamento", driver, doc, "Confirmando cancelamento realizado.")

    # Mensagem final
    _, tipo_alerta = encontrar_mensagem_alerta(doc)
    if tipo_alerta == "sucesso":
        log(doc, "✅ Mensagem de sucesso exibida.")
    elif tipo_alerta == "alerta":
        log(doc, "⚠️ Mensagem de alerta exibida.")
    elif tipo_alerta == "erro":
        log(doc, "❌ Mensagem de erro exibida.")
    else:
        log(doc, "⚠️ Nenhuma mensagem exibida.")
    registrar_screenshot_unico("mensagem_final", driver, doc, "Mensagem exibida após salvar.")

    ok, _ = safe_action(doc, "Fechando formulário", fechar_modal)
    if not ok:
        finalizar_relatorio(); return
    registrar_screenshot_unico("formulario_fechado", driver, doc, "Formulário fechado.")

    log(doc, "✅ Teste concluído com sucesso.")
    finalizar_relatorio()


if __name__ == "__main__":
    main()
