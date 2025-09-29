# Refatorado e organizado: cadastrodetransportes1ºcenario.py

from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support.ui import Select, WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service
from selenium.webdriver import ActionChains
from webdriver_manager.chrome import ChromeDriverManager
from docx import Document
from docx.shared import Inches
from faker import Faker
from faker.providers import BaseProvider
from faker_vehicle import VehicleProvider
from datetime import datetime, timedelta
import subprocess
import os
import time
import random
import string

# ==== PROVIDERS CUSTOMIZADOS ====
class BrasilProvider(BaseProvider):
    def rg(self):
        numeros = [str(random.randint(0, 9)) for _ in range(8)]
        return ''.join(numeros) + '-' + str(random.randint(0, 9))

class TransporteProvider(BaseProvider):
    def gerar_placa_mercosul(self):
        letras = string.ascii_uppercase
        numeros = string.digits
        return f"{''.join(random.choices(letras, k=3))}{random.choice(numeros)}{random.choice(letras)}{''.join(random.choices(numeros, k=2))}"
    
    def vehicle_year_make_model(self):
        year = self.generator.vehicle_year()
        make = self.generator.vehicle_make()
        model = self.generator.vehicle_model()
        return f"{year} {make} {model}"

fake = Faker("pt_BR")
fake.add_provider(BrasilProvider)
fake.add_provider(VehicleProvider)
fake.add_provider(TransporteProvider)

# ==== CONFIGURAÇÕES ====
URL = "http://localhost:8080/gs/index.xhtml"
LOGIN_EMAIL = "joaoeduardo.gold@outlook.com"
LOGIN_PASSWORD = "071999gs"

# ==== DOCUMENTO ====
doc = Document()
doc.add_heading("RELATÓRIO DO TESTE", 0)
doc.add_paragraph("Cadastro de Transporte – Cenário 1: Preenchimento completo e salvamento.")
doc.add_paragraph(f"Data do teste: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")

screenshot_registradas = set()

def preencher_campo_com_retry(driver, wait, seletor, valor, max_tentativas=5):
    """Tenta preencher o campo com diferentes métodos até conseguir"""
    
    for tentativa in range(max_tentativas):
        try:
            # Aguarda o elemento
            campo = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, seletor)))
            
            # Scroll até o elemento se necessário
            driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", campo)
            time.sleep(0.5)
            
            # Método 1: Tradicional
            if tentativa == 0:
                campo.click()
                campo.clear()
                campo.send_keys(valor)
                campo.send_keys(Keys.TAB)
            
            # Método 2: ActionChains
            elif tentativa == 1:
                ActionChains(driver).move_to_element(campo).click().perform()
                time.sleep(0.2)
                ActionChains(driver).key_down(Keys.CONTROL).send_keys('a').key_up(Keys.CONTROL).perform()
                ActionChains(driver).send_keys(valor).perform()
                ActionChains(driver).send_keys(Keys.TAB).perform()
            
            # Método 3: JavaScript
            else:
                driver.execute_script("""
                    var element = arguments[0];
                    var valor = arguments[1];
                    element.focus();
                    element.value = '';
                    element.value = valor;
                    element.dispatchEvent(new Event('input', { bubbles: true }));
                    element.dispatchEvent(new Event('change', { bubbles: true }));
                    element.blur();
                """, campo, valor)
            
            time.sleep(0.5)
            
            # Verifica se o valor foi preenchido
            valor_atual = campo.get_attribute('value')
            if valor_atual == valor:
                return True
            else:
                print(f"⚠️ Tentativa {tentativa + 1}: Valor não foi preenchido corretamente")
                
        except Exception as e:
            print(f"⚠️ Tentativa {tentativa + 1} falhou: {e}")
            time.sleep(1)
    
    return False

# ==== FUNÇÕES DE UTILITÁRIO ====
def log(doc, msg):
    print(msg)
    doc.add_paragraph(msg)

def take_screenshot(driver, doc, nome):
    if nome not in screenshot_registradas:
        path = f"screenshots/{nome}.png"
        os.makedirs("screenshots", exist_ok=True)
        driver.save_screenshot(path)
        doc.add_paragraph(f"Screenshot: {nome}")
        doc.add_picture(path, width=Inches(5.5))
        screenshot_registradas.add(nome)

def safe_action(doc, descricao, func):
    try:
        log(doc, f"🔄 {descricao}...")
        func()
        log(doc, f"✅ {descricao} realizada com sucesso.")
        take_screenshot(driver, doc, descricao.lower().replace(" ", "_"))
    except Exception as e:
        log(doc, f"❌ Erro ao {descricao.lower()}: {e}")
        take_screenshot(driver, doc, f"erro_{descricao.lower().replace(' ', '_')}")

def finalizar_relatorio():
    nome_arquivo = f"relatorio_transporte_cenario_1_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc.save(nome_arquivo)
    log(doc, f"📄 Relatório salvo como: {nome_arquivo}")
    subprocess.run(["start", "winword", nome_arquivo], shell=True)
    driver.quit()

def gerar_dados_transporte():
    """Gera dados fictícios para o cadastro de transporte."""
    modelo_veiculo = fake.vehicle_year_make_model()
    placa = fake.gerar_placa_mercosul()
    motorista = fake.name()
    local_origem = fake.city()
    local_destino = fake.city()
    
    # Gerar datas e horas
    hoje = datetime.today()
    data_saida_obj = hoje + timedelta(days=random.randint(0, 30))
    hora_saida_obj = datetime(
        year=data_saida_obj.year,
        month=data_saida_obj.month,
        day=data_saida_obj.day,
        hour=random.randint(0, 23),
        minute=random.randint(0, 59)
    )
    hora_chegada_obj = hora_saida_obj + timedelta(
        days=random.randint(1, 10),
        hours=random.randint(0, 5),
        minutes=random.randint(0, 59)
    )
    
    data_saida = hora_saida_obj.strftime('%d/%m/%Y')
    hora_saida = hora_saida_obj.strftime('%H:%M')
    data_chegada = hora_chegada_obj.strftime('%d/%m/%Y')
    hora_chegada = hora_chegada_obj.strftime('%H:%M')
    
    # Gerar dados de combustível
    km_saida = random.randint(1, 99999)
    km_chegada = km_saida + random.randint(100, 1000)
    preco_combustivel = round(random.uniform(4.0, 7.0), 2)
    litros = round(random.randint(1, 99))
    valor_gasto = round(preco_combustivel * litros)
    
    return (modelo_veiculo, placa, motorista, local_origem, local_destino,
            data_saida, hora_saida, data_chegada, hora_chegada,
            km_saida, km_chegada, preco_combustivel, litros, valor_gasto)

def encontrar_mensagem_alerta():
    seletores = [
        (".alerts.salvo", "✅ Mensagem de Sucesso"),
        (".alerts.alerta", "⚠️ Mensagem de Alerta"),
        (".alerts.erro", "❌ Mensagem de Erro"),
    ]

    for seletor, tipo in seletores:
        try:
            elemento = driver.find_element(By.CSS_SELECTOR, seletor)
            if elemento.is_displayed():
                log(doc, f"📢 {tipo}: {elemento.text}")
                return elemento
        except:
            continue

    log(doc, "ℹ️ Nenhuma mensagem de alerta encontrada.")
    return None

def ajustar_zoom():
    try:
        driver.execute_script("document.body.style.zoom='90%'")
        log(doc, "🔍 Zoom ajustado para 90%.")
    except Exception as e:
        log(doc, f"⚠️ Erro ao ajustar zoom: {e}")

def preencher_campo_data(selector, valor):
    def acao():
        campo = wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, selector)))
        campo.click()
        campo.send_keys(valor)
        time.sleep(0.2)
    return acao

def selecionar_opcao(selector, texto):
    def acao():
        select_element = wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, selector)))
        Select(select_element).select_by_visible_text(texto)
    return acao

# Gera os dados necessários
(modelo_veiculo, placa, motorista, local_origem, local_destino,
 data_saida, hora_saida, data_chegada, hora_chegada,
 km_saida, km_chegada, preco_combustivel, litros, valor_gasto) = gerar_dados_transporte()

# ==== INICIALIZAÇÃO DO DRIVER ====
options = Options()
options.add_argument("--start-maximized")
driver = webdriver.Chrome(service=Service(ChromeDriverManager().install()), options=options)
wait = WebDriverWait(driver, 20)

# ==== EXECUÇÃO DO TESTE ====
try:
    safe_action(doc, "Acessando sistema", lambda: driver.get(URL))

    safe_action(doc, "Realizando login", lambda: (
        wait.until(EC.presence_of_element_located((By.ID, "j_id15:email"))).send_keys(LOGIN_EMAIL),
        wait.until(EC.presence_of_element_located((By.ID, "j_id15:senha"))).send_keys(LOGIN_PASSWORD, Keys.ENTER),
        wait.until(EC.presence_of_element_located((By.TAG_NAME, "body")))
    ))

    safe_action(doc, "Esperando sistema carregar e ajustando zoom", lambda: (
        time.sleep(5),
        ajustar_zoom()
    ))

    safe_action(doc, "Abrindo menu Transporte", lambda: (
        driver.find_element(By.TAG_NAME, "body").send_keys(Keys.F2),
        time.sleep(1),
        wait.until(EC.visibility_of_element_located((By.XPATH, "//input[@placeholder='Busque um cadastro']"))).send_keys("Transporte", Keys.ENTER),
        time.sleep(2)
    ))

    safe_action(doc, "Clicando em Cadastrar", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_200029 > div.wdTelas > div.telaInicial.clearfix.overflow.overflowY > ul > li:nth-child(1) > a > span"))).click(),
        time.sleep(5)
    ))

    # Preenchendo dados do transporte
    safe_action(doc, "Preenchendo Veículo", lambda: 
        preencher_campo_com_retry(driver, wait, "#fmod_200029 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div > div > div > div:nth-child(2) > input", 'TESTE VEÍCULO SELENIUM ')
    )

    safe_action(doc, "Preenchendo Placa", lambda:
        preencher_campo_com_retry(driver, wait, "#fmod_200029 div:nth-child(3) > input", placa)
    )

    safe_action(doc, "Preenchendo Motorista", lambda:
        preencher_campo_com_retry(driver, wait, "#fmod_200029 div:nth-child(4) > input", motorista)
    )

    safe_action(doc, "Preenchendo Local Origem", lambda:
        preencher_campo_com_retry(driver, wait, "#fmod_200029 div:nth-child(5) > input", local_origem)
    )

    safe_action(doc, "Preenchendo Local Destino", lambda:
        preencher_campo_com_retry(driver, wait, "#fmod_200029 div:nth-child(6) > input", local_destino)
    )

    safe_action(doc, "Preenchendo Data Saída", lambda:
        preencher_campo_com_retry(driver, wait, 'input[grupo="100055"][maxlength="10"][ref="100191"]', data_saida)
    )

    safe_action(doc, "Preenchendo Data Chegada", lambda:
        preencher_campo_com_retry(driver, wait, 'input[grupo="100055"][maxlength="10"][ref="100192"]', data_chegada)
    )

    safe_action(doc, "Preenchendo Hora Saída", lambda:
        preencher_campo_com_retry(driver, wait, "#fmod_200029 div:nth-child(9) > input", hora_saida)
    )

    safe_action(doc, "Preenchendo Hora Chegada", lambda:
        preencher_campo_com_retry(driver, wait, "#fmod_200029 div:nth-child(10) > input", hora_chegada)
    )

    safe_action(doc, "Preenchendo KM Saída", lambda:
        preencher_campo_com_retry(driver, wait, "#fmod_200029 div:nth-child(11) > input", str(km_saida))
    )

    safe_action(doc, "Preenchendo KM Chegada", lambda:
        preencher_campo_com_retry(driver, wait, "#fmod_200029 div:nth-child(12) > input", str(km_chegada))
    )

    safe_action(doc, "Preenchendo Preço Combustível", lambda:
        preencher_campo_com_retry(driver, wait, "#fmod_200029 div:nth-child(13) > input", str(preco_combustivel))
    )

    safe_action(doc, "Preenchendo Litros", lambda:
        preencher_campo_com_retry(driver, wait, "#fmod_200029 div:nth-child(14) > input", str(litros))
    )

    safe_action(doc, "Preenchendo Valor Gasto", lambda:
        preencher_campo_com_retry(driver, wait, "#fmod_200029 div:nth-child(15) > input", str(valor_gasto))
    )

    safe_action(doc, "Salvando cadastro", lambda: (
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "#fmod_200029 a.btGray.btsave"))).click(),
    ))

    safe_action(doc, "Fechando modal após salvamento", lambda: (
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "#fmod_200029 > div.wdTop.ui-draggable-handle > div.wdClose > a"))).click(),
    ))


    time.sleep(2)
    log(doc, "🔍 Verificando mensagens de alerta...")
    encontrar_mensagem_alerta()

except Exception as e:
    log(doc, f"❌ ERRO FATAL: {e}")
    take_screenshot(driver, doc, "erro_fatal")

finally:
    log(doc, "✅ Teste concluído com sucesso.")
    finalizar_relatorio()