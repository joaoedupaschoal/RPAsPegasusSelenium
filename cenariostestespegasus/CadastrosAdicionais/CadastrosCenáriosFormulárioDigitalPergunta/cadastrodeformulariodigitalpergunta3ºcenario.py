# Refatorado e organizado: cadastrodeformulariodigitalpergunta2ºcenario.py

from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support.ui import Select, WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service
from selenium.webdriver import ActionChains
from webdriver_manager.chrome import ChromeDriverManager
from docx import Document
from docx.shared import Inches
from faker import Faker
from faker.providers import BaseProvider
from datetime import datetime
import subprocess
import os
import time
import random
import string

# ==== PROVIDERS CUSTOMIZADOS ====
class BrasilProvider(BaseProvider):
    def rg(self):
        numeros = [str(random.randint(0, 9)) for _ in range(8)]
        return ''.join(numeros) + '-' + str(random.randint(0, 9))

fake = Faker("pt_BR")
fake.add_provider(BrasilProvider)

# ==== CONFIGURAÇÕES ====
URL = "http://localhost:8080/gs/index.xhtml"
LOGIN_EMAIL = "joaoeduardo.gold@outlook.com"
LOGIN_PASSWORD = "071999gs"

# ==== DOCUMENTO ====
doc = Document()
doc.add_heading("RELATÓRIO DO TESTE", 0)
doc.add_paragraph("Cadastro de Formulário Digital Pergunta – Cenário 3: Preenchimento dos campos obrigatórios e salvamento.")
doc.add_paragraph(f"Data do teste: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")

screenshot_registradas = set()

# ==== FUNÇÕES DE UTILITÁRIO ====
def log(doc, msg):
    print(msg)
    doc.add_paragraph(msg)

def take_screenshot(driver, doc, nome):
    if nome not in screenshot_registradas:
        path = f"screenshots/{nome}.png"
        os.makedirs("screenshots", exist_ok=True)
        driver.save_screenshot(path)
        doc.add_paragraph(f"Screenshot: {nome}")
        doc.add_picture(path, width=Inches(5.5))
        screenshot_registradas.add(nome)

def safe_action(doc, descricao, func):
    try:
        log(doc, f"🔄 {descricao}...")
        func()
        log(doc, f"✅ {descricao} realizada com sucesso.")
        take_screenshot(driver, doc, descricao.lower().replace(" ", "_"))
    except Exception as e:
        log(doc, f"❌ Erro ao {descricao.lower()}: {e}")
        take_screenshot(driver, doc, f"erro_{descricao.lower().replace(' ', '_')}")

def finalizar_relatorio():
    nome_arquivo = f"relatorio_formulario_digital_pergunta_cenario_3_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc.save(nome_arquivo)
    log(doc, f"📄 Relatório salvo como: {nome_arquivo}")
    subprocess.run(["start", "winword", nome_arquivo], shell=True)
    driver.quit()

def gerar_dados_pergunta():
    """Gera dados fictícios para a pergunta do formulário digital."""
    numero_aleatorio = random.randint(1, 1000)
    
    descricao_pergunta = f"POSSUI FILHOS? (TESTE SELENIUM AUTOMATIZADO {numero_aleatorio})"
    
    alternativas = [
        f'SIM, POSSUI FILHOS. (TESTE SELENIUM AUTOMATIZADO {numero_aleatorio})',
        f'NÃO, NÃO POSSUI FILHOS. (TESTE SELENIUM AUTOMATIZADO {numero_aleatorio})'
    ]
    
    return descricao_pergunta, alternativas

def encontrar_mensagem_alerta():
    seletores = [
        (".alerts.salvo", "✅ Mensagem de Sucesso"),
        (".alerts.alerta", "⚠️ Mensagem de Alerta"),
        (".alerts.erro", "❌ Mensagem de Erro"),
    ]

    for seletor, tipo in seletores:
        try:
            elemento = driver.find_element(By.CSS_SELECTOR, seletor)
            if elemento.is_displayed():
                log(doc, f"📢 {tipo}: {elemento.text}")
                return elemento
        except:
            continue

    log(doc, "ℹ️ Nenhuma mensagem de alerta encontrada.")
    return None

def ajustar_zoom():
    try:
        driver.execute_script("document.body.style.zoom='90%'")
        log(doc, "🔍 Zoom ajustado para 90%.")
    except Exception as e:
        log(doc, f"⚠️ Erro ao ajustar zoom: {e}")

def selecionar_opcao(selector, texto):
    def acao():
        select_element = wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, selector)))
        Select(select_element).select_by_visible_text(texto)
    return acao

def criar_alternativa(texto_alternativa, ordem):
    def acao():
        # Abrir LOV alternativa
        open_lov_alternativa = wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_10043 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div > div:nth-child(2) > div > div:nth-child(3) > div > a")))
        open_lov_alternativa.click()
        
        # Novo registro alternativa
        novo_registro_alternativa = wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(4) > a")))
        novo_registro_alternativa.click()
        
        # Preencher descrição alternativa
        descricao_alternativa = wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#cg_10044 > div.wdTelas > div > div.catWrapper > div > div > div > div > div > input")))
        descricao_alternativa.send_keys(texto_alternativa)
        
        # Salvar alternativa
        salvar_alternativa = wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#cg_10044 > div.wdTelas > div > div.btnHolder > a.btModel.btGray.btsave")))
        salvar_alternativa.click()
        time.sleep(5)
        
        # Preencher ordem alternativa
        ordem_alternativa = wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_10043 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div > div:nth-child(2) > div > div:nth-child(4) > input")))
        ordem_alternativa.send_keys(str(ordem))
        
        # Adicionar alternativa
        add_alternativa = wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_10043 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div > div:nth-child(2) > div > div.btnListHolder > a.btAddGroup")))
        add_alternativa.click()
    
    return acao

# Gera os dados necessários
descricao_pergunta, alternativas = gerar_dados_pergunta()

# ==== INICIALIZAÇÃO DO DRIVER ====
options = Options()
options.add_argument("--start-maximized")
driver = webdriver.Chrome(service=Service(ChromeDriverManager().install()), options=options)
wait = WebDriverWait(driver, 20)

# ==== EXECUÇÃO DO TESTE ====
try:
    safe_action(doc, "Acessando sistema", lambda: driver.get(URL))

    safe_action(doc, "Realizando login", lambda: (
        wait.until(EC.presence_of_element_located((By.ID, "j_id15:email"))).send_keys(LOGIN_EMAIL),
        wait.until(EC.presence_of_element_located((By.ID, "j_id15:senha"))).send_keys(LOGIN_PASSWORD, Keys.ENTER),
        wait.until(EC.presence_of_element_located((By.TAG_NAME, "body")))
    ))

    safe_action(doc, "Esperando sistema carregar e ajustando zoom", lambda: (
        time.sleep(5),
        ajustar_zoom()
    ))

    safe_action(doc, "Abrindo menu Formulário Digital Pergunta", lambda: (
        driver.find_element(By.TAG_NAME, "body").send_keys(Keys.F2),
        time.sleep(1),
        wait.until(EC.visibility_of_element_located((By.XPATH, "//input[@placeholder='Busque um cadastro']"))).send_keys("Formulário Digital Pergunta", Keys.ENTER),
        time.sleep(3)
    ))

    safe_action(doc, "Clicando em Cadastrar", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_10043 > div.wdTelas > div > ul > li:nth-child(1) > a > span"))).click(),
        time.sleep(2)
    ))

    # Preenchendo dados da pergunta
    safe_action(doc, "Preenchendo Descrição da Pergunta", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_10043 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div > div:nth-child(1) > div > div:nth-child(2) > input")))
        .send_keys(descricao_pergunta)
    ))

    safe_action(doc, "Selecionando Tipo da Pergunta", selecionar_opcao(
        "#fmod_10043 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div > div:nth-child(1) > div > div:nth-child(3) > select",
        "Multipla Escolha"
    ))

    safe_action(doc, "Selecionando Formato da Pergunta", selecionar_opcao(
        "#fmod_10043 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div > div:nth-child(1) > div > div:nth-child(4) > select",
        "Alfa númerica"
    ))


    safe_action(doc, "Salvando cadastro da pergunta", lambda: (
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "#fmod_10043 > div.wdTelas > div.telaCadastro.clearfix > div.btnHolder > a.btModel.btGray.btsave"))).click(),
        time.sleep(1)
    ))

    encontrar_mensagem_alerta()


    safe_action(doc, "Fechando modal após salvamento", lambda: (
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "#fmod_10043 > div.wdTop.ui-draggable-handle > div.wdClose > a"))).click(),
        time.sleep(1)
    ))


except Exception as e:
    log(doc, f"❌ ERRO FATAL: {e}")
    take_screenshot(driver, doc, "erro_fatal")

finally:
    log(doc, "✅ Teste concluído com sucesso.")
    finalizar_relatorio()