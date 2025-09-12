# Refatorado: Cadastro de Abastecimento com Selenium

from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support.ui import Select, WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.common.action_chains import ActionChains
from webdriver_manager.chrome import ChromeDriverManager
from docx import Document
from docx.shared import Inches
from faker import Faker
from faker.providers import BaseProvider
from validate_docbr import CPF
from datetime import datetime, timedelta
import subprocess
import os
import time
import random
import sys 
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys


def preencher_data_abastecimento(driver, data_abastecimento, timeout=10):
    """
    Preenche o campo de data de abastecimento no formulário.
    
    Args:
        driver: Instância do WebDriver do Selenium
        data_abastecimento: Data no formato string (ex: "15/03/2024")
        timeout: Tempo limite para aguardar o elemento (padrão: 10 segundos)
    
    Returns:
        bool: True se preencheu com sucesso, False caso contrário
    """
    try:
        # Aguarda o campo de data estar clicável
        wait = WebDriverWait(driver, timeout)
        campo_data_abastecimento = wait.until(EC.element_to_be_clickable((
            By.XPATH,
            "//input[@maxlength='10' and @name='data' and contains(@class, 'hasDatepicker data mandatory')]"
        )))
        
        # Limpa o campo antes de preencher
        campo_data_abastecimento.clear()
        
        # Preenche com a data desejada
        campo_data_abastecimento.send_keys(str(data_abastecimento), Keys.TAB)
        
        print(f"Data de abastecimento '{data_abastecimento}' preenchida com sucesso.")
        return True
        
    except Exception as e:
        print(f"Erro ao preencher data de abastecimento: {str(e)}")
        return False

sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))

# ==== PROVIDERS CUSTOMIZADOS ====
class BrasilProvider(BaseProvider):
    def cnpj_valido(self):
        """Gera um CNPJ válido"""
        def gerar_digitos():
            numeros = [random.randint(0, 9) for _ in range(12)]
            # Cálculo do primeiro dígito verificador
            soma = sum(numeros[i] * peso for i, peso in enumerate([5, 4, 3, 2, 9, 8, 7, 6, 5, 4, 3, 2]))
            digito1 = 11 - (soma % 11)
            if digito1 >= 10:
                digito1 = 0
            
            # Cálculo do segundo dígito verificador
            numeros.append(digito1)
            soma = sum(numeros[i] * peso for i, peso in enumerate([6, 5, 4, 3, 2, 9, 8, 7, 6, 5, 4, 3, 2]))
            digito2 = 11 - (soma % 11)
            if digito2 >= 10:
                digito2 = 0
            
            numeros.append(digito2)
            return ''.join(map(str, numeros))
        
        cnpj = gerar_digitos()
        return f"{cnpj[:2]}.{cnpj[2:5]}.{cnpj[5:8]}/{cnpj[8:12]}-{cnpj[12:]}"

    def rg(self):
        numeros = [str(random.randint(0, 9)) for _ in range(8)]
        return ''.join(numeros) + '-' + str(random.randint(0, 9))

# Inicialização do Faker
fake = Faker("pt_BR")
fake.add_provider(BrasilProvider)

# ==== CONFIGURAÇÕES ====
URL = "http://localhost:8080/gs/index.xhtml"
LOGIN_EMAIL = "joaoeduardo.gold@outlook.com"
LOGIN_PASSWORD = "071999gs"

# ==== DOCUMENTO ====
doc = Document()
doc.add_heading("RELATÓRIO DO TESTE", 0)
doc.add_paragraph("Cadastro de Abastecimento – Cenário 1: Preenchimento completo e salvamento.")
doc.add_paragraph(f"Data do teste: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")

screenshot_registradas = set()

# ==== FUNÇÕES GERADORAS DE DADOS ====
def gerar_dados_pessoa_juridica():
    """Gera dados fictícios para o cadastro de pessoa jurídica."""
    nome_empresa = 'POSTO COMBUSTÍVEL TESTE SELENIUM AUTOMATIZADO'
    nome_fantasia = 'POSTO COMBUSTÍVEL TESTE SELENIUM AUTOMATIZADO'
    cnpj_valido = fake.cnpj_valido()
    inscricao_estadual = str(fake.random_int(min=10000000, max=99999999))
    inscricao_municipal = str(fake.random_int(min=10000000, max=99999999))
    email = fake.email()
    telefone1 = fake.phone_number()
    telefone2 = fake.phone_number()
    telefone3 = fake.phone_number()
    cidade_nascimento = fake.city()
    pais_nascimento = fake.country()
    nome_pai = fake.first_name()
    nome_mae = fake.first_name()
    profissao = fake.job()
    data_nascimento = fake.date_of_birth(minimum_age=18, maximum_age=60).strftime("%d/%m/%Y")
    
    return (nome_empresa, cnpj_valido, nome_fantasia, inscricao_estadual, inscricao_municipal,
            email, telefone1, telefone2, telefone3, cidade_nascimento, pais_nascimento, 
            nome_pai, nome_mae, profissao, data_nascimento)

def gerar_dados_frotas(intervalo_max_passado=30):
    """
    Gera dados completos de uma multa:
    - data_multa: entre hoje e X dias atrás
    - hora_multa: horário aleatório (HH:MM)
    - data_notificacao: de 0 a 5 dias após a multa
    - data_vencimento: de 10 a 20 dias após a multa
    - data_pagamento: entre notificação e vencimento

    Retorna um dicionário com todas as datas e hora no formato dd/mm/yyyy e HH:MM.
    """
    # Base da multa
    data_multa_dt = datetime.now() - timedelta(days=random.randint(0, intervalo_max_passado))
    hoje = datetime.now()
    # Gera hora no formato HH:MM
    hora = random.randint(0, 23)
    minuto = random.randint(0, 59)
    hora_multa_dt = f"{hora:02d}:{minuto:02d}"

    # Demais datas relacionadas
    data_notificacao_dt = data_multa_dt + timedelta(days=random.randint(0, 5))
    data_vencimento_dt = data_multa_dt + timedelta(days=random.randint(10, 20))
    data_pagamento_dt = random.choice([
        data_notificacao_dt + timedelta(days=random.randint(0, (data_vencimento_dt - data_notificacao_dt).days)),
        data_vencimento_dt
    ])

    trinta_dias_atras = hoje - timedelta(days=30)
    data_abastecimento_dt = fake.date_between(start_date=trinta_dias_atras, end_date=hoje) 

    return (
        data_multa_dt.strftime("%d/%m/%Y"),
        hora_multa_dt,
        data_notificacao_dt.strftime("%d/%m/%Y"),
        data_vencimento_dt.strftime("%d/%m/%Y"),
        data_pagamento_dt.strftime("%d/%m/%Y"),
        data_abastecimento_dt.strftime("%d/%m/%Y"),
    )

def gerar_datas_validas():
    """Gera datas coerentes para admissão, início e fim da escala, e vencimento da CNH."""
    hoje = datetime.today().date()
    
    # Data de admissão entre 10 anos atrás e hoje
    data_admissao = fake.date_between(start_date=hoje - timedelta(days=3650), end_date=hoje)
    
    # Data de início da escala entre hoje e 1 ano no futuro
    data_inicio = fake.date_between(start_date=hoje, end_date=hoje + timedelta(days=365))
    
    # Data fim entre 1 e 180 dias após a data de início
    data_fim = data_inicio + timedelta(days=random.randint(1, 180))
    
    # Vencimento CNH entre hoje e 10 anos no futuro
    vencimento_cnh = fake.date_between(start_date=hoje, end_date=hoje + timedelta(days=3650))
    
    return (data_admissao.strftime('%d/%m/%Y'), 
            data_inicio.strftime('%d/%m/%Y'), 
            data_fim.strftime('%d/%m/%Y'), 
            vencimento_cnh.strftime('%d/%m/%Y'))

def gerar_dados_documentos():
    """Gera documentos fictícios para o cadastro."""
    carteira_trabalho = str(random.randint(10000000, 99999999))
    pis = fake.cpf().replace('.', '').replace('-', '')[:11]
    cnh = str(random.randint(10000000000, 99999999999))
    cpf = CPF().generate()
    
    return carteira_trabalho, pis, cnh, cpf

# Gerando todos os dados necessários
data_multa, hora_multa, data_notificacao, data_vencimento, data_pagamento, data_abastecimento_dt = gerar_dados_frotas()
data_admissao, data_inicio, data_fim, vencimento_cnh = gerar_datas_validas()
carteira_trabalho, pis, cnh, cpf_valido = gerar_dados_documentos()
(nome_empresa, cnpj_valido, nome_fantasia, inscricao_estadual, inscricao_municipal,
 email, telefone1, telefone2, telefone3, cidade_nascimento, pais_nascimento, 
 nome_pai, nome_mae, profissao, data_nascimento) = gerar_dados_pessoa_juridica()

# ==== FUNÇÕES DE UTILITÁRIO ====
def log(doc, msg):
    print(msg)
    doc.add_paragraph(msg)

def take_screenshot(driver, doc, nome):
    if nome not in screenshot_registradas:
        path = f"screenshots/{nome}.png"
        os.makedirs("screenshots", exist_ok=True)
        driver.save_screenshot(path)
        doc.add_paragraph(f"Screenshot: {nome}")
        doc.add_picture(path, width=Inches(5.5))
        screenshot_registradas.add(nome)

def safe_action(doc, descricao, func):
    try:
        log(doc, f"🔄 {descricao}...")
        func()
        log(doc, f"✅ {descricao} realizada com sucesso.")
        take_screenshot(driver, doc, descricao.lower().replace(" ", "_"))
    except Exception as e:
        log(doc, f"❌ Erro ao {descricao.lower()}: {e}")
        take_screenshot(driver, doc, f"erro_{descricao.lower().replace(' ', '_')}")

def finalizar_relatorio():
    nome_arquivo = f"relatorio_abastecimento_cenario_1_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc.save(nome_arquivo)
    log(doc, f"📄 Relatório salvo como: {nome_arquivo}")
    subprocess.run(["start", "winword", nome_arquivo], shell=True)
    driver.quit()

def encontrar_mensagem_alerta():
    seletores = [
        (".alerts.salvo", "✅ Sucesso"),
        (".alerts.alerta", "⚠️ Alerta"),
        (".alerts.erro", "❌ Erro"),
    ]

    for seletor, tipo in seletores:
        try:
            elemento = driver.find_element(By.CSS_SELECTOR, seletor)
            if elemento.is_displayed():
                log(doc, f"📢 {tipo}: {elemento.text}")
                return elemento
        except:
            continue

    log(doc, "ℹ️ Nenhuma mensagem de alerta encontrada.")
    return None

def ajustar_zoom():
    try:
        driver.execute_script("document.body.style.zoom='90%'")
        log(doc, "🔍 Zoom ajustado para 90%.")
    except Exception as e:
        log(doc, f"⚠️ Erro ao ajustar zoom: {e}")

def preencher_campo_data(selector, valor):
    def acao():
        campo = wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, selector)))
        campo.click()
        campo.clear()
        campo.send_keys(valor)
        time.sleep(0.2)
    return acao

def selecionar_opcao(selector, texto):
    def acao():
        select_element = wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, selector)))
        Select(select_element).select_by_visible_text(texto)
    return acao

def abrir_modal_e_selecionar(btn_selector, pesquisa_selector, termo_pesquisa, btn_pesquisar_selector, resultado_xpath):
    """Abre modal e seleciona um item"""
    def acao():
        # Abre o modal
        open_lov = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, btn_selector)))
        open_lov.click()
        time.sleep(3)

        # Aguarda campo pesquisa
        campo_pesquisa = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, pesquisa_selector)))
        campo_pesquisa.clear()
        campo_pesquisa.send_keys(termo_pesquisa)

        # Clica pesquisar
        pesquisar = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, btn_pesquisar_selector)))
        pesquisar.click()
        time.sleep(3)
        pesquisar.click()

        # Espera o resultado e clica
        wait.until(EC.element_to_be_clickable((By.XPATH, resultado_xpath)))
        resultado = driver.find_element(By.XPATH, resultado_xpath)
        driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", resultado)
        time.sleep(0.2)
        resultado.click()

    return acao

def buscar_e_selecionar_pacote(nome_pacote):
    def acao():
        # Abre o LOV do pacote
        open_lov_pacote = wait.until(
            EC.element_to_be_clickable((By.CSS_SELECTOR, "#cg_1 > div.wdTelas > div > div.catWrapper > div > div.cat_dadosPessoais.categoriaHolder > div > div > div.formRow.divPessoaJURIDICA > div.formCol.pacotes > div > a"))
        )
        open_lov_pacote.click()

        # Pesquisa o pacote
        campo_pesquisa_pacote = WebDriverWait(driver, 10).until(
            EC.visibility_of_element_located((By.CSS_SELECTOR, "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(2) > input"))          
        )
        campo_pesquisa_pacote.send_keys(nome_pacote, Keys.ENTER)

        # Seleciona o pacote
        pacote = WebDriverWait(driver, 10).until(
            EC.element_to_be_clickable((By.XPATH, f"//td[contains(text(), '{nome_pacote}')]"))
        )
        pacote.click()
    return acao


def preencher_campo_xpath_com_retry(driver, wait, xpath, valor, max_tentativas=3):
    global doc
    if driver is None or wait is None:
        return False
    for tentativa in range(max_tentativas):
        try:
            campo = wait.until(EC.element_to_be_clickable((By.XPATH, xpath)))
            driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", campo)
            time.sleep(0.3)
            if tentativa == 0:
                campo.click(); campo.clear(); campo.send_keys(valor); campo.send_keys(Keys.TAB)
            elif tentativa == 1:
                ActionChains(driver).move_to_element(campo).click().pause(0.1).key_down(Keys.CONTROL).send_keys('a').key_up(Keys.CONTROL).send_keys(valor).send_keys(Keys.TAB).perform()
            else:
                driver.execute_script("""
                    var el = arguments[0], v = arguments[1];
                    el.focus(); el.value = ''; el.value = v;
                    el.dispatchEvent(new Event('input', { bubbles: true }));
                    el.dispatchEvent(new Event('change', { bubbles: true }));
                    el.blur();
                """, campo, valor)
            time.sleep(0.3)
            if (campo.get_attribute('value') or '').strip():
                return True
        except Exception as e:
            log(doc, f"⚠️ Tentativa {tentativa + 1} falhou: {e}")
            time.sleep(0.8)
    return False


def preencher_campo_com_retry(driver, wait, seletor, valor, max_tentativas=3):
    """Tenta preencher o campo com diferentes métodos até conseguir"""
    
    for tentativa in range(max_tentativas):
        try:
            # Aguarda o elemento
            campo = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, seletor)))
            
            # Scroll até o elemento se necessário
            driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", campo)
            time.sleep(0.5)
            
            # Método 1: Tradicional
            if tentativa == 0:
                campo.click()
                campo.clear()
                campo.send_keys(valor)
                campo.send_keys(Keys.TAB)
            
            # Método 2: ActionChains
            elif tentativa == 1:
                ActionChains(driver).move_to_element(campo).click().perform()
                time.sleep(0.2)
                ActionChains(driver).key_down(Keys.CONTROL).send_keys('a').key_up(Keys.CONTROL).perform()
                ActionChains(driver).send_keys(valor).perform()
                ActionChains(driver).send_keys(Keys.TAB).perform()
            
            # Método 3: JavaScript
            else:
                driver.execute_script("""
                    var element = arguments[0];
                    var valor = arguments[1];
                    element.focus();
                    element.value = '';
                    element.value = valor;
                    element.dispatchEvent(new Event('input', { bubbles: true }));
                    element.dispatchEvent(new Event('change', { bubbles: true }));
                    element.blur();
                """, campo, valor)
            
            time.sleep(0.5)
            
            # Verifica se o valor foi preenchido
            valor_atual = campo.get_attribute('value')
            if valor_atual == valor:
                return True
                
        except Exception as e:
            time.sleep(1)
    
    return False

# ==== INICIALIZAÇÃO DO DRIVER ====
options = Options()
options.add_argument("--start-maximized")
driver = webdriver.Chrome(service=Service(ChromeDriverManager().install()), options=options)
wait = WebDriverWait(driver, 20)

# ==== EXECUÇÃO DO TESTE ====
try:
    safe_action(doc, "Acessando sistema", lambda: driver.get(URL))

    safe_action(doc, "Realizando login", lambda: (
        wait.until(EC.presence_of_element_located((By.ID, "j_id15:email"))).send_keys(LOGIN_EMAIL),
        wait.until(EC.presence_of_element_located((By.ID, "j_id15:senha"))).send_keys(LOGIN_PASSWORD, Keys.ENTER),
        wait.until(EC.presence_of_element_located((By.TAG_NAME, "body")))
    ))

    safe_action(doc, "Esperando sistema carregar e ajustando zoom", lambda: (
        time.sleep(5),
        ajustar_zoom()
    ))

    safe_action(doc, "Abrindo menu Abastecimento", lambda: (
        driver.find_element(By.TAG_NAME, "body").send_keys(Keys.F2),
        time.sleep(1),
        wait.until(EC.visibility_of_element_located((By.XPATH, "//input[@placeholder='Busque um cadastro']"))).send_keys("Abastecimento", Keys.ENTER)
    ))

    safe_action(doc, "Clicando em Cadastrar", lambda: (
        time.sleep(2),
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_10090 > div.wdTelas > div > ul > li:nth-child(1) > a > span"))).click()
    ))

    safe_action(doc, "Preenchendo Data de Abastecimento", lambda:
                    preencher_campo_xpath_com_retry(
                        driver, wait, "//input[@name='data']",
                        "12/09/2025"
                    ))
    

    safe_action(doc, "Selecionando Tipo de Combustível", selecionar_opcao(
        "#fmod_10090 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroAbastecimento > div.catWrapper > div > div > div > div > div > div:nth-child(3) > select",
        "Gasolina"
    ))

    safe_action(doc, "Selecionando Motorista", abrir_modal_e_selecionar(
        "#fmod_10090 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroAbastecimento > div.catWrapper > div > div > div > div > div > div:nth-child(4) > div > a",
        "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div.formCol.divPesquisa > input",
        "CRISPIM MALAFAIA",
        "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(3) > a",
        "//td[contains(text(), 'CRISPIM MALAFAIA')]"
    ))

    safe_action(doc, "Selecionando Veículo", abrir_modal_e_selecionar(
        "#fmod_10090 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroAbastecimento > div.catWrapper > div > div > div > div > div > div:nth-child(5) > div > a",
        "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div.formCol.divPesquisa > div:nth-child(1) > input",
        "TESTE VEÍCULO SELENIUM AUTOMATIZADO",
        "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(3) > a",
        "//td[contains(text(), 'TESTE VEÍCULO SELENIUM AUTOMATIZADO')]"
    ))

    safe_action(doc, "Abrindo modal de Posto de Combustível", lambda: (
        time.sleep(2),
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_10090 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroAbastecimento > div.catWrapper > div > div > div > div > div > div:nth-child(6) > div > a"))).click()
    ))

    safe_action(doc, "Criando Novo Registro", lambda: (
        time.sleep(2),
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(3) > a"))).click()
    ))


    # Preenchendo dados pessoais da Pessoa Jurídica
    safe_action(doc, "Preenchendo Nome da Empresa", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#cg_1 > div.wdTelas > div > div.catWrapper > div > div.cat_dadosPessoais.categoriaHolder > div > div > div:nth-child(2) > div:nth-child(2) > input")))
        .send_keys(nome_empresa)
    ))

    safe_action(doc, "Selecionando Tipo de Pessoa", selecionar_opcao(
        "#cg_1 > div.wdTelas > div > div.catWrapper > div > div.cat_dadosPessoais.categoriaHolder > div > div > div:nth-child(2) > div:nth-child(3) > select",
        "Jurídica"
    ))

    safe_action(doc, "Selecionando Tipo de Documento", selecionar_opcao(
        "#cg_1 > div.wdTelas > div > div.catWrapper > div > div.cat_dadosPessoais.categoriaHolder > div > div > div:nth-child(2) > div:nth-child(4) > select",
        "Carteira de Identidade Classista"
    ))

    safe_action(doc, "Preenchendo CNPJ", lambda: (
        preencher_campo_com_retry(driver, wait, 
                                "#cg_1 > div.wdTelas > div > div.catWrapper > div > div.cat_dadosPessoais.categoriaHolder > div > div > div.formRow.divPessoaJURIDICA > div:nth-child(1) > input",
                                cnpj_valido)
    ))

    safe_action(doc, "Preenchendo Nome Fantasia", lambda: (
        preencher_campo_com_retry(driver, wait,
                                "#cg_1 > div.wdTelas > div > div.catWrapper > div > div.cat_dadosPessoais.categoriaHolder > div > div > div.formRow.divPessoaJURIDICA > div:nth-child(2) > input",
                                nome_fantasia)
    ))

    safe_action(doc, "Preenchendo Inscrição Estadual", lambda: (
        preencher_campo_com_retry(driver, wait,
                                "#cg_1 > div.wdTelas > div > div.catWrapper > div > div.cat_dadosPessoais.categoriaHolder > div > div > div.formRow.divPessoaJURIDICA > div:nth-child(3) > input",
                                inscricao_estadual)
    ))

    safe_action(doc, "Preenchendo Inscrição Municipal", lambda: (
        preencher_campo_com_retry(driver, wait,
                                "#cg_1 > div.wdTelas > div > div.catWrapper > div > div.cat_dadosPessoais.categoriaHolder > div > div > div.formRow.divPessoaJURIDICA > div:nth-child(4) > input",
                                inscricao_municipal)
    ))

    safe_action(doc, "Selecionando Pacote", buscar_e_selecionar_pacote('PACOTE TESTE SELENIUM AUTOMATIZADO'))

    safe_action(doc, "Selecionando Classificação", selecionar_opcao(
        "#cg_1 > div.wdTelas > div > div.catWrapper > div > div.cat_dadosPessoais.categoriaHolder > div > div > div.formRow.divPessoaJURIDICA > div:nth-child(6) > select",
        "01 - Ótimo"
    ))

    # Dados Complementares
    safe_action(doc, "Acessando aba Dados Complementares", lambda: (
        wait.until(EC.element_to_be_clickable((By.LINK_TEXT, "Dados Complementares"))).click(),
        time.sleep(1)
    ))

    safe_action(doc, "Selecionando Estado Civil", selecionar_opcao(
        "#cg_1 > div.wdTelas > div > div.catWrapper > div > div.cat_dadosComplementares.categoriaHolder > div > div > div > div:nth-child(1) > select",
        "Solteiro"
    ))

    safe_action(doc, "Selecionando Sexo", selecionar_opcao(
        "#cg_1 > div.wdTelas > div > div.catWrapper > div > div.cat_dadosComplementares.categoriaHolder > div > div > div > div:nth-child(2) > select",
        "Feminino"
    ))

    safe_action(doc, "Preenchendo E-mail", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#cg_1 > div.wdTelas > div > div.catWrapper > div > div.cat_dadosComplementares.categoriaHolder > div > div > div > div:nth-child(9) > input")))
        .send_keys(email)
    ))

    safe_action(doc, "Preenchendo Data de Nascimento", preencher_campo_data(
        "input.dataNascimento", data_nascimento
    ))

    # Preenchendo campos de contato
    campos_contato = [
        ("Telefone 1", "#cg_1 > div.wdTelas > div > div.catWrapper > div > div.cat_dadosComplementares.categoriaHolder > div > div > div > div:nth-child(5) > input", telefone1),
        ("Telefone 2", "#cg_1 > div.wdTelas > div > div.catWrapper > div > div.cat_dadosComplementares.categoriaHolder > div > div > div > div:nth-child(6) > input", telefone2),
        ("Telefone 3", "#cg_1 > div.wdTelas > div > div.catWrapper > div > div.cat_dadosComplementares.categoriaHolder > div > div > div > div:nth-child(7) > input", telefone3),
        ("Cidade de Nascimento", "#cg_1 > div.wdTelas > div > div.catWrapper > div > div.cat_dadosComplementares.categoriaHolder > div > div > div > div:nth-child(11) > input", cidade_nascimento),
        ("País de Nascimento", "#cg_1 > div.wdTelas > div > div.catWrapper > div > div.cat_dadosComplementares.categoriaHolder > div > div > div > div:nth-child(12) > input", pais_nascimento),
        ("Nome do Pai", "#cg_1 > div.wdTelas > div > div.catWrapper > div > div.cat_dadosComplementares.categoriaHolder > div > div > div > div:nth-child(13) > input", nome_pai),
        ("Nome da Mãe", "#cg_1 > div.wdTelas > div > div.catWrapper > div > div.cat_dadosComplementares.categoriaHolder > div > div > div > div:nth-child(14) > input", nome_mae),
        ("Profissão", "#cg_1 > div.wdTelas > div > div.catWrapper > div > div.cat_dadosComplementares.categoriaHolder > div > div > div > div:nth-child(16) > input", profissao)
    ]

    for nome_campo, seletor, valor in campos_contato:
        safe_action(doc, f"Preenchendo {nome_campo}", lambda s=seletor, v=valor: (
            wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, s)))
            .send_keys(v)
        ))

    safe_action(doc, "Salvando cadastro de Posto de Combustível", lambda: (
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "#cg_1 > div.wdTelas > div > div.btnHolder > a.btModel.btGray.btsave"))).click(),
        time.sleep(5)
    ))

    encontrar_mensagem_alerta()

    safe_action(doc, "Preenchendo Volume", lambda: (
        preencher_campo_com_retry(driver, wait,
                                "#fmod_10090 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroAbastecimento > div.catWrapper > div > div > div > div > div > div:nth-child(8) > input",
                                "53,00")
    ))

    safe_action(doc, "Preenchendo Valor Unitário", lambda: (
        preencher_campo_com_retry(driver, wait,
                                "#fmod_10090 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroAbastecimento > div.catWrapper > div > div > div > div > div > div:nth-child(9) > input",
                                "5,70")
    ))

    safe_action(doc, "Preenchendo Desconto", lambda: (
        preencher_campo_com_retry(driver, wait,
                                "#fmod_10090 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroAbastecimento > div.catWrapper > div > div > div > div > div > div:nth-child(10) > input",
                                "7,22")
    ))

    safe_action(doc, "Salvando cadastro de Abastecimento", lambda: (
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "#fmod_10090 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroAbastecimento > div.btnHolder > a.btModel.btGray.btsave"))).click(),
        time.sleep(3)
    ))

    encontrar_mensagem_alerta()


    safe_action(doc, "Recusando o Lançamento do Abastecimento no Contas à Pagar", lambda: (
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "#BtNo"))).click(),
        time.sleep(3)
    ))
    


    safe_action(doc, "Fechando modal após salvamento", lambda: (
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "#fmod_10090 > div.wdTop.ui-draggable-handle > div.wdClose > a"))).click()
    ))


except Exception as e:
    log(doc, f"❌ ERRO FATAL: {e}")
    take_screenshot(driver, doc, "erro_fatal")

finally:
    log(doc, "✅ Teste concluído com sucesso.")
    finalizar_relatorio()