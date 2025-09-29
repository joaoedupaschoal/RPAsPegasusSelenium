import sys
import os

# Adiciona a raiz do projeto ao sys.path
sys.path.append(
    os.path.abspath(
        os.path.join(os.path.dirname(__file__), "../../..")
    )
)


from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support.ui import WebDriverWait, Select
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service
from webdriver_manager.chrome import ChromeDriverManager
from docx import Document
from datetime import datetime
import subprocess
import time
import os
import sys
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), '..', '..')))
from utils.actions import log, take_screenshot, safe_action,  ajustar_zoom

URL = "http://localhost:8080/gs/index.xhtml"
LOGIN_EMAIL = "joaoeduardo.gold@outlook.com"
LOGIN_PASSWORD = "071999gs"

screenshot_registradas = set()
def registrar_screenshot_unico(nome, driver, doc, descricao=None):
    if nome not in screenshot_registradas:
        if descricao:
            log(doc, f"📸 {descricao}")
        take_screenshot(driver, doc, nome)
        screenshot_registradas.add(nome)


  
def main():
    doc = Document()
    doc.add_heading("RELATÓRIO DO TESTE", 0)
    doc.add_paragraph("Cadastro de Carteira – Cenário 3: Apenas campos não obrigatórios.")
    doc.add_paragraph(f"🗕️ Data do teste: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")

    chrome_options = Options()
    chrome_options.add_argument("--start-maximized")
    driver = webdriver.Chrome(service=Service(ChromeDriverManager().install()), options=chrome_options)
    wait = WebDriverWait(driver, 10)

    def encontrar_mensagem_alerta():
        seletores = [
            (".alerts.salvo", "✅ Mensagem de Sucesso"),
            (".alerts.alerta", "⚠️ Mensagem de Alerta"),
            (".alerts.erro", "❌ Mensagem de Erro"),
        ]

        for seletor, tipo in seletores:
            try:
                elemento = driver.find_element(By.CSS_SELECTOR, seletor)
                if elemento.is_displayed():
                    log(doc, f"📢 {tipo}: {elemento.text}")
                    return elemento
            except:
                continue

        log(doc, "ℹ️ Nenhuma mensagem de alerta encontrada.")
        return None


    def finalizar_relatorio():
        doc_name = f"relatorio_carteira_cenario_3_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        doc.save(doc_name)
        log(doc, f"📄 Relatório salvo como: {doc_name}")
        try:
            subprocess.run(["start", "winword", doc_name], shell=True)
        except Exception as e:
            log(doc, f"Erro ao abrir o Word: {e}")
        driver.quit()

    def login():
        wait.until(EC.presence_of_element_located((By.ID, "j_id15:email"))).send_keys(LOGIN_EMAIL)
        wait.until(EC.presence_of_element_located((By.ID, "j_id15:senha"))).send_keys(LOGIN_PASSWORD, Keys.ENTER)
        wait.until(EC.presence_of_element_located((By.TAG_NAME, "body")))
        time.sleep(2)

    def abrir_menu():
        driver.find_element(By.TAG_NAME, "body").send_keys(Keys.F2)
        campo = wait.until(EC.presence_of_element_located((By.XPATH, "//input[@placeholder='Busque um cadastro']")))
        campo.click()
        campo.send_keys("Carteira")
        wait.until(EC.element_to_be_clickable((By.XPATH, "/html/body/div[17]/ul/li[8]/a"))).click()
        time.sleep(2)

    def acessar_formulario():
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR,
            "#fmod_10021 > div.wdTelas > div > ul > li:nth-child(1) > a > span"))).click()
        time.sleep(2)

  

    def salvar():
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR,
            "#fmod_10021 > div.wdTelas > div.telaCadastro.clearfix > div.btnHolder > a.btModel.btGray.btsave"))).click()
        time.sleep(2)

    def fechar_modal():
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR,
            "#fmod_10021 > div.wdTop.ui-draggable-handle > div.wdClose > a"))).click()
        time.sleep(1)

    if not safe_action(doc, "Acessando o sistema", lambda: driver.get(URL), driver, wait)[0]: finalizar_relatorio(); return
    registrar_screenshot_unico("url_acessada", driver, doc)

    if not safe_action(doc, "Realizando login", login, driver, wait)[0]: finalizar_relatorio(); return
    registrar_screenshot_unico("login_concluido", driver, doc)

    if not safe_action(doc, "Ajustando zoom", lambda: ajustar_zoom(driver), driver, wait)[0]: finalizar_relatorio(); return
    if not safe_action(doc, "Abrindo menu Carteira", abrir_menu, driver, wait)[0]: finalizar_relatorio(); return
    registrar_screenshot_unico("menu_aberto", driver, doc)

    if not safe_action(doc, "Acessando formulário", acessar_formulario, driver, wait)[0]: finalizar_relatorio(); return
    registrar_screenshot_unico("formulario_aberto", driver, doc)

    # NÃO preencher campos obrigatórios de propósito

    if not safe_action(doc, "Clicando em Salvar", salvar, driver, wait)[0]: finalizar_relatorio(); return
    registrar_screenshot_unico("apos_salvar", driver, doc, "Clique em salvar mesmo sem preencher obrigatórios.")
 
 

    encontrar_mensagem_alerta()
    

    fechar_modal()
    log(doc, "✅ Teste finalizado com sucesso.")
    finalizar_relatorio()
if __name__ == "__main__":
    main()

