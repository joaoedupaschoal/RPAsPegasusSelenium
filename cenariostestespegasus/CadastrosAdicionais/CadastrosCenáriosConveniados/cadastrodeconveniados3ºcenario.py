import sys
import os

# Adiciona a raiz do projeto ao sys.path
sys.path.append(
    os.path.abspath(
        os.path.join(os.path.dirname(__file__), "../../..")
    )
)
# Refatorado: cadastrodeconveniados1ºcenario.py

from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service
from webdriver_manager.chrome import ChromeDriverManager
from docx import Document
from docx.shared import Inches
from faker import Faker
from datetime import datetime
import subprocess
import os
import time
from selenium.webdriver.support.ui import Select
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service
from webdriver_manager.chrome import ChromeDriverManager
from docx import Document
from docx.shared import Inches
from faker import Faker
from datetime import datetime
import subprocess
import os
import time
import re


# ==== CONFIG ====
URL = "http://localhost:8080/gs/index.xhtml"
LOGIN_EMAIL = "joaoeduardo.gold@outlook.com"
LOGIN_PASSWORD = "071999gs"
fake = Faker("pt_BR")



# ==== DOCUMENTO ====
doc = Document()
doc.add_heading("RELATÓRIO DO TESTE", 0)
doc.add_paragraph("Cadastro de Conveniados – Cenário 3: Preenchimento dos campos obrigatórios e salvamento.")
doc.add_paragraph(f"Data do teste: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")

# ==== LOG E PRINT ====
screenshot_registradas = set()
def log(doc, msg):
    print(msg)
    doc.add_paragraph(msg)

def take_screenshot(driver, doc, nome):
    if nome not in screenshot_registradas:
        path = f"screenshots/{nome}.png"
        os.makedirs("screenshots", exist_ok=True)
        driver.save_screenshot(path)
        doc.add_paragraph(f"Screenshot: {nome}")
        doc.add_picture(path, width=Inches(5.5))
        screenshot_registradas.add(nome)

def abrir_modal_e_selecionar(btn_selector, pesquisa_selector, termo_pesquisa, btn_pesquisar_selector, resultado_xpath):
    """Abre modal e seleciona um item"""
    def acao():
        # Abre o modal
        open_lov = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, btn_selector)))
        open_lov.click()
        time.sleep(3)

        # Aguarda campo pesquisa
        campo_pesquisa = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, pesquisa_selector)))
        campo_pesquisa.clear()
        campo_pesquisa.send_keys(termo_pesquisa)

        # Clica pesquisar
        pesquisar = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, btn_pesquisar_selector)))
        pesquisar.click()
        time.sleep(3)
        pesquisar.click()
        time.sleep(3)

        # Espera o resultado e clica
        wait.until(EC.element_to_be_clickable((By.XPATH, resultado_xpath)))
        resultado = driver.find_element(By.XPATH, resultado_xpath)
        driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", resultado)
        time.sleep(0.2)
        resultado.click()

    return acao

def selecionar_opcao(selector, texto):
    def acao():
        select_element = wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, selector)))
        Select(select_element).select_by_visible_text(texto)
    return acao

def safe_action(doc, descricao, func):
    try:
        log(doc, f"🔄 {descricao}...")
        func()
        log(doc, f"✅ {descricao} realizada com sucesso.")
        take_screenshot(driver, doc, descricao.lower().replace(" ", "_"))
    except Exception as e:
        log(doc, f"❌ Erro ao {descricao.lower()}: {e}")
        take_screenshot(driver, doc, f"erro_{descricao.lower().replace(' ', '_')}")

def finalizar_relatorio():
    nome_arquivo = f"relatorio_conveniados_cenario_3_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc.save(nome_arquivo)
    log(doc, f"📄 Relatório salvo como: {nome_arquivo}")
    subprocess.run(["start", "winword", nome_arquivo], shell=True)
    driver.quit()



def take_screenshot(driver, doc, nome):
    nome = re.sub(r'[^a-zA-Z0-9_\-]', '_', str(nome).lower().replace(" ", "_"))
    if nome not in screenshot_registradas:
        path = f"screenshots/{nome}.png"
        os.makedirs("screenshots", exist_ok=True)
        driver.save_screenshot(path)
        doc.add_paragraph(f"Screenshot: {nome}")
        doc.add_picture(path, width=Inches(5.5))
        screenshot_registradas.add(nome)





def safe_action(doc, descricao, func):
    try:
        log(doc, f"🔀 {descricao}...")
        func()
        log(doc, f"✅ {descricao} realizada com sucesso.")
        take_screenshot(driver, doc, descricao)
    except Exception as e:
        log(doc, f"❌ Erro ao {descricao.lower()}: {e}")
        take_screenshot(driver, doc, f"erro_{descricao}")



def encontrar_mensagem_alerta():
    seletores = [
        (".alerts.salvo", "✅ Mensagem de Sucesso"),
        (".alerts.alerta", "⚠️ Mensagem de Alerta"),
        (".alerts.erro", "❌ Mensagem de Erro"),
    ]
    for seletor, tipo in seletores:
        try:
            elemento = driver.find_element(By.CSS_SELECTOR, seletor)
            if elemento.is_displayed():
                log(doc, f"📳 {tipo}: {elemento.text}")
                return elemento
        except:
            continue
    log(doc, "ℹ️ Nenhuma mensagem de alerta encontrada.")
    return None

def ajustar_zoom():
    try:
        driver.execute_script("document.body.style.zoom='90%'")
        log(doc, "🔍 Zoom ajustado para 90%.")
    except Exception as e:
        log(doc, f"⚠️ Erro ao ajustar zoom: {e}")


def preencher_email():
    campo = wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, "#fmod_10063 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroConveniado > div.catWrapper > div > div.cat_contatos.categoriaHolder > div > div.group_contatos.clearfix.grupoHolder.lista > div > div:nth-child(3) > input")))
    campo.click()
    campo.clear()
    campo.send_keys(fake.email())

def preencher_telefone():
    campo = wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, "#fmod_10063 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroConveniado > div.catWrapper > div > div.cat_contatos.categoriaHolder > div > div.group_contatos.clearfix.grupoHolder.lista > div > div:nth-child(3) > input")))
    campo.click()
    campo.send_keys(fake.phone_number())

def preencher_celular():
    campo = wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, "#fmod_10063 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroConveniado > div.catWrapper > div > div.cat_contatos.categoriaHolder > div > div.group_contatos.clearfix.grupoHolder.lista > div > div:nth-child(3) > input")))
    campo.click()
    campo.send_keys(fake.phone_number())

def preencher_escala_trabalho():
    def preencher_linha(dia_index, manha_inicio, manha_fim, tarde_inicio, tarde_fim, noite_inicio, noite_fim, tempo_servico, intervalo):
        base_xpath = f"//table[contains(@class, 'escalaTrabalho')]//tr[{dia_index + 1}]"
        campos = [
            (f"{base_xpath}/td[2]/input", manha_inicio),
            (f"{base_xpath}/td[3]/input", manha_fim),
            (f"{base_xpath}/td[4]/input", tarde_inicio),
            (f"{base_xpath}/td[5]/input", tarde_fim),
            (f"{base_xpath}/td[6]/input", noite_inicio),
            (f"{base_xpath}/td[7]/input", noite_fim),
            (f"{base_xpath}/td[8]/input", tempo_servico),
            (f"{base_xpath}/td[9]/input", intervalo),
        ]
        for xpath, valor in campos:
            try:
                campo = wait.until(EC.presence_of_element_located((By.XPATH, xpath)))
                campo.click()
                campo.clear()
                campo.send_keys(valor)
            except Exception as e:
                log(doc, f"❌ Erro ao preencher campo {xpath}: {e}")

    log(doc, "🔀 Preenchendo Escala de Trabalho...")
    try:
        preencher_linha(
            dia_index=0,
            manha_inicio="08:00",
            manha_fim="12:00",
            tarde_inicio="13:00",
            tarde_fim="17:00",
            noite_inicio="18:00",
            noite_fim="20:00",
            tempo_servico="30",
            intervalo="10"
        )
        log(doc, "✅ Escala de Trabalho preenchida com sucesso.")
        take_screenshot(driver, doc, "Escala de Trabalho")
    except Exception as e:
        log(doc, f"❌ Erro ao preencher Escala de Trabalho: {e}")
        take_screenshot(driver, doc, "erro_escala_trabalho")

# ==== INICIALIZAÇÃO DO DRIVER ====
options = Options()
options.add_argument("--start-maximized")
driver = webdriver.Chrome(service=Service(ChromeDriverManager().install()), options=options)
wait = WebDriverWait(driver, 20)

# ==== EXECUÇÃO DO TESTE ====
try:
    safe_action(doc, "Acessando sistema", lambda: driver.get(URL))

    safe_action(doc, "Realizando login", lambda: (
        wait.until(EC.presence_of_element_located((By.ID, "j_id15:email"))).send_keys(LOGIN_EMAIL),
        wait.until(EC.presence_of_element_located((By.ID, "j_id15:senha"))).send_keys(LOGIN_PASSWORD, Keys.ENTER),
        wait.until(EC.presence_of_element_located((By.TAG_NAME, "body")))
    ))

    safe_action(doc, "Esperando sistema carregar e ajustando zoom", lambda: (
        time.sleep(5),
        ajustar_zoom()
    ))

    # Navega ao menu Conveniados
    safe_action(doc, "Abrindo menu Conveniados", lambda: (
        driver.find_element(By.TAG_NAME, "body").send_keys(Keys.F2),
        time.sleep(2),
        driver.find_element(By.XPATH, "//input[@placeholder='Busque um cadastro']").send_keys("Conveniados", Keys.ENTER)
    ))

    # Clica em cadastrar
    safe_action(doc, "Clicando em Cadastrar", lambda: wait.until(
        EC.element_to_be_clickable((By.CSS_SELECTOR, "#fmod_10063 > div.wdTelas > div > ul > li:nth-child(1) > a > span"))
    ).click())

    # Preenche nome
    safe_action(doc, "Preenchendo campo Nome do Conveniado", lambda: wait.until(
        EC.presence_of_element_located((By.CSS_SELECTOR, "#fmod_10063 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroConveniado > div.catWrapper > div > div.cat_geral.categoriaHolder > div > div > div > div:nth-child(2) > input"))
    ).send_keys(fake.name()))



    safe_action(doc, "Preenchendo Valor Particular", lambda: wait.until(
        EC.presence_of_element_located((By.CSS_SELECTOR, "#fmod_10063 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroConveniado > div.catWrapper > div > div.cat_geral.categoriaHolder > div > div > div > div:nth-child(5) > input"))
    ).send_keys(fake.random_int(min=10000, max=100000)))



    safe_action(doc, "Preenchendo Valor convênio", lambda: wait.until(
        EC.presence_of_element_located((By.CSS_SELECTOR, "#fmod_10063 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroConveniado > div.catWrapper > div > div.cat_geral.categoriaHolder > div > div > div > div:nth-child(6) > div > input"))
    ).send_keys(fake.random_int(min=10000, max=100000)))


    # Salvar
    safe_action(doc, "Clicando em Salvar", lambda: wait.until(
        EC.element_to_be_clickable((By.CSS_SELECTOR, "#fmod_10063 > div.wdTelas > div.telaCadastro.clearfix.telaCadastroConveniado > div.btnHolder > a.btModel.btGray.btsave"))
    ).click())

    encontrar_mensagem_alerta()

    # Fechar modal
    safe_action(doc, "Fechando modal", lambda: wait.until(
        EC.element_to_be_clickable((By.CSS_SELECTOR, "#fmod_10063 > div.wdTop.ui-draggable-handle > div.wdClose > a"))
    ).click())

except Exception as e:
    log(doc, f"❌ ERRO FATAL: {e}")
    take_screenshot(driver, doc, "erro_fatal")

finally:
    finalizar_relatorio()

