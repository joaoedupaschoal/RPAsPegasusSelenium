import sys
import os

# Adiciona a raiz do projeto ao sys.path
sys.path.append(
    os.path.abspath(
        os.path.join(os.path.dirname(__file__), "../../..")
    )
)
# Refatorado com logs detalhados e preenchimento via JavaScript - VERSÃO CORRIGIDA

from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support.ui import Select, WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service
from selenium.webdriver import ActionChains
from webdriver_manager.chrome import ChromeDriverManager
from docx import Document
from docx.shared import Inches
from faker import Faker
from faker.providers import BaseProvider
from datetime import datetime
import subprocess
import os
import time
import random
import string

# ==== PROVIDERS CUSTOMIZADOS ====
class BrasilProvider(BaseProvider):
    def rg(self):
        numeros = [str(random.randint(0, 9)) for _ in range(8)]
        return ''.join(numeros) + '-' + str(random.randint(0, 9))

fake = Faker("pt_BR")
fake.add_provider(BrasilProvider)

# ==== CONFIGURAÇÕES ====
URL = "http://localhost:8080/gs/index.xhtml"
LOGIN_EMAIL = "joaoeduardo.gold@outlook.com"
LOGIN_PASSWORD = "071999gs"

# ==== DOCUMENTO ====
doc = Document()
doc.add_heading("RELATÓRIO DO TESTE", 0)
doc.add_paragraph("Cadastro de Emitente Nota Fiscal – Cenário 3: Preenchimento dos campos obrigatórios e salvamento")
doc.add_paragraph(f"Data do teste: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")

screenshot_registradas = set()

def log(doc, msg, nivel="INFO"):
    """
    Função para registrar logs detalhados - CORRIGIDA
    """
    timestamp = datetime.now().strftime('%H:%M:%S')
    
    # Emojis para diferentes níveis
    emojis = {
        "INFO": "ℹ️",
        "WARN": "⚠️", 
        "ERROR": "❌",
        "SUCCESS": "✅"
    }
    
    emoji = emojis.get(nivel, "ℹ️")
    log_msg = f"{emoji} {msg}"
    print(log_msg)
    doc.add_paragraph(log_msg)

def take_screenshot(driver, doc, nome):
    """
    Captura screenshot com log detalhado - CORRIGIDA
    """
    try:
        if nome not in screenshot_registradas:
            path = f"screenshots/{nome}.png"
            os.makedirs("screenshots", exist_ok=True)
            driver.save_screenshot(path)
            log(doc, f"Screenshot capturada: {nome} -> {path}")
            doc.add_paragraph(f"Screenshot: {nome}")
            doc.add_picture(path, width=Inches(5.5))
            screenshot_registradas.add(nome)
        else:
            log(doc, f"Screenshot já registrada: {nome}", "WARN")
    except Exception as e:
        log(doc, f"Erro ao capturar screenshot {nome}: {e}", "ERROR")

def preencher_campo_js(driver, wait, seletor, valor, nome_campo="Campo"):
    """
    Preenche campo usando JavaScript com logs detalhados
    """
    try:
        log(doc, f"Iniciando preenchimento do campo '{nome_campo}' com valor: {valor}")
        
        # Aguarda o elemento estar presente
        log(doc, f"Aguardando elemento '{nome_campo}' estar disponível...")
        campo = wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, seletor)))
        log(doc, f"Elemento '{nome_campo}' encontrado")
        
        # Scroll para o elemento
        log(doc, f"Fazendo scroll para o campo '{nome_campo}'...")
        driver.execute_script("arguments[0].scrollIntoView({behavior: 'smooth', block: 'center'});", campo)
        time.sleep(0.5)
        
        # Preenchimento via JavaScript
        log(doc, f"Preenchendo campo '{nome_campo}' via JavaScript...")
        script = """
            var element = arguments[0];
            var valor = arguments[1];
            
            // Foca no elemento
            element.focus();
            
            // Limpa o campo
            element.value = '';
            
            // Define o novo valor
            element.value = valor;
            
            // Dispara eventos necessários
            element.dispatchEvent(new Event('input', { bubbles: true }));
            element.dispatchEvent(new Event('change', { bubbles: true }));
            element.dispatchEvent(new Event('blur', { bubbles: true }));
            
            return element.value;
        """
        
        resultado = driver.execute_script(script, campo, str(valor))
        
        # Verifica se o preenchimento foi bem-sucedido
        if str(resultado) == str(valor):
            log(doc, f"Campo '{nome_campo}' preenchido com sucesso: {resultado}", "SUCCESS")
            return True
        else:
            log(doc, f"Falha no preenchimento do campo '{nome_campo}'. Esperado: {valor}, Obtido: {resultado}", "WARN")
            return False
            
    except Exception as e:
        log(doc, f"Erro ao preencher campo '{nome_campo}': {e}", "ERROR")
        return False

def selecionar_opcao_js(driver, wait, seletor, texto, nome_select="Select"):
    """
    Seleciona opção em select usando JavaScript com logs detalhados
    """
    try:
        log(doc, f"Iniciando seleção em '{nome_select}' com texto: {texto}")
        
        # Aguarda o elemento
        log(doc, f"Aguardando select '{nome_select}' estar disponível...")
        select_element = wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, seletor)))
        log(doc, f"Select '{nome_select}' encontrado")
        
        # Scroll para o elemento
        log(doc, f"Fazendo scroll para o select '{nome_select}'...")
        driver.execute_script("arguments[0].scrollIntoView({behavior: 'smooth', block: 'center'});", select_element)
        time.sleep(0.5)
        
        # Seleção via JavaScript
        log(doc, f"Selecionando opção '{texto}' no select '{nome_select}'...")
        script = """
            var select = arguments[0];
            var texto = arguments[1];
            
            // Procura pela opção com o texto
            for (var i = 0; i < select.options.length; i++) {
                if (select.options[i].text === texto) {
                    select.selectedIndex = i;
                    select.dispatchEvent(new Event('change', { bubbles: true }));
                    return select.options[i].text;
                }
            }
            return null;
        """
        
        resultado = driver.execute_script(script, select_element, texto)
        
        if resultado:
            log(doc, f"Opção selecionada com sucesso em '{nome_select}': {resultado}", "SUCCESS")
            return True
        else:
            log(doc, f"Falha ao selecionar opção '{texto}' em '{nome_select}'", "WARN")
            return False
            
    except Exception as e:
        log(doc, f"Erro ao selecionar opção em '{nome_select}': {e}", "ERROR")
        return False

def clicar_elemento_js(driver, wait, seletor, nome_elemento="Elemento"):
    """
    Clica em elemento usando JavaScript com logs detalhados
    """
    try:
        log(doc, f"Iniciando clique no elemento '{nome_elemento}'")
        
        # Aguarda o elemento
        log(doc, f"Aguardando elemento '{nome_elemento}' estar clicável...")
        elemento = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, seletor)))
        log(doc, f"Elemento '{nome_elemento}' encontrado e clicável")
        
        # Scroll para o elemento
        log(doc, f"Fazendo scroll para o elemento '{nome_elemento}'...")
        driver.execute_script("arguments[0].scrollIntoView({behavior: 'smooth', block: 'center'});", elemento)
        time.sleep(0.5)
        
        # Clique via JavaScript
        log(doc, f"Clicando no elemento '{nome_elemento}' via JavaScript...")
        driver.execute_script("arguments[0].click();", elemento)
        
        log(doc, f"Clique realizado com sucesso no elemento '{nome_elemento}'", "SUCCESS")
        return True
        
    except Exception as e:
        log(doc, f"Erro ao clicar no elemento '{nome_elemento}': {e}", "ERROR")
        return False

def safe_action(doc, descricao, func):
    """
    Executa ação com tratamento de erro e logs detalhados - CORRIGIDA
    """
    try:
        log(doc, f"🔄 INICIANDO: {descricao}")
        resultado = func()
        log(doc, f"✅ SUCESSO: {descricao}", "SUCCESS")
        take_screenshot(driver, doc, descricao.lower().replace(" ", "_"))
        return resultado
    except Exception as e:
        log(doc, f"❌ ERRO: {descricao} - {str(e)}", "ERROR")
        take_screenshot(driver, doc, f"erro_{descricao.lower().replace(' ', '_')}")
        raise e

def gerar_dados_emitente():
    """
    Gera dados fictícios para o emitente com logs
    """

    
    dados = {
        'razao_social': 'TESTE EMITENTE NOTA FISCAL SELENIUM AUTOMATIZADO',
        'cnpj': fake.cnpj(),
        'cep': '15081115',
        'numero_endereco': '1735',
        'complemento': 'Casa',
        'cidade': fake.city(),
        'cod_municipio': fake.random_int(min=1000000, max=9999999),
        'inscricao_municipal': fake.random_int(min=111111111111111, max=999999999999999),
        'serie': fake.random_int(min=11111, max=99999),
        'cod_regime_tributario': fake.random_int(min=0, max=9),
        'natureza_operacao': fake.random_int(min=0, max=9),
        'frase_secreta': 'TESTE FRASE SECRETA SELENIUM AUTOMATIZADO',
        'senha': 'TESTE SENHA SECRETA SELENIUM AUTOMATIZADO',
        'outras_informacoes': 'TESTE OUTRAS INFORMAÇÕES SELENIUM AUTOMATIZADO',
        'percentual_DAS': fake.random_int(min=1, max=10000),
        'url_webservice': 'testewebservice@seleniumautomatizado.org.br'
    }
    

    return dados

def encontrar_mensagem_alerta():
    seletores = [
        (".alerts.salvo", "✅ Mensagem de Sucesso"),
        (".alerts.alerta", "⚠️ Mensagem de Alerta"),
        (".alerts.erro", "❌ Mensagem de Erro"),
    ]
    for seletor, tipo in seletores:
        try:
            elemento = driver.find_element(By.CSS_SELECTOR, seletor)
            if elemento.is_displayed():
                log(doc, f"📢 {tipo}: {elemento.text}")
                return elemento
        except:
            continue

    log(doc, "ℹ️ Nenhuma mensagem de alerta encontrada.")
    return None

def ajustar_zoom():
    try:
        driver.execute_script("document.body.style.zoom='90%'")
        log(doc, "🔍 Zoom ajustado para 90%.")
    except Exception as e:
        log(doc, f"⚠️ Erro ao ajustar zoom: {e}")

def preencher_campo_data(selector, valor):
    def acao():
        campo = wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, selector)))
        campo.click()
        campo.send_keys(valor)
        time.sleep(0.2)
    return acao

def ajustar_zoom():
    """
    Ajusta zoom da página com logs
    """
    try:
        log(doc, "Ajustando zoom da página para 90%...")
        driver.execute_script("document.body.style.zoom='90%'")
        log(doc, "🔍 Zoom ajustado para 90% com sucesso", "SUCCESS")
    except Exception as e:
        log(doc, f"⚠️ Erro ao ajustar zoom: {e}", "ERROR")

def selecionar_opcao_select(driver, wait, selector, texto, nome_select="Select"):
    """
    Função para selecionar opção em select - CORRIGIDA
    """
    try:
        log(doc, f"Selecionando '{texto}' em '{nome_select}'...")
        select_element = wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, selector)))
        
        # Scroll para o elemento
        driver.execute_script("arguments[0].scrollIntoView({behavior: 'smooth', block: 'center'});", select_element)
        time.sleep(0.5)
        
        # Seleciona a opção
        Select(select_element).select_by_visible_text(texto)
        log(doc, f"Opção '{texto}' selecionada com sucesso em '{nome_select}'", "SUCCESS")
        return True
    except Exception as e:
        log(doc, f"Erro ao selecionar '{texto}' em '{nome_select}': {e}", "ERROR")
        return False

def preencher_campo_com_retry(driver, wait, seletor, valor, max_tentativas=5):
    """Tenta preencher o campo com diferentes métodos até conseguir"""
    
    for tentativa in range(max_tentativas):
        try:

            
            # Aguarda o elemento
            campo = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, seletor)))
            
            # Scroll até o elemento se necessário
            driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", campo)
            time.sleep(0.5)
            
            # Método 1: Tradicional
            if tentativa == 0:
                campo.click()
                campo.send_keys(valor)
                campo.send_keys(Keys.TAB)
            
            # Método 2: ActionChains
            elif tentativa == 1:
                ActionChains(driver).move_to_element(campo).click().perform()
                time.sleep(0.2)
                ActionChains(driver).key_down(Keys.CONTROL).send_keys('a').key_up(Keys.CONTROL).perform()
                ActionChains(driver).send_keys(valor).perform()
                ActionChains(driver).send_keys(Keys.TAB).perform()
            
            # Método 3: JavaScript
            else:
                driver.execute_script("""
                    var element = arguments[0];
                    var valor = arguments[1];
                    element.focus();
                    element.value = '';
                    element.value = valor;
                    element.dispatchEvent(new Event('input', { bubbles: true }));
                    element.dispatchEvent(new Event('change', { bubbles: true }));
                    element.blur();
                """, campo, valor)
            
            time.sleep(0.5)
            
            # Verifica se o valor foi preenchido
            valor_atual = campo.get_attribute('value')
            if valor_atual == valor:

                return True
            else:
                print()
                
        except Exception as e:
            time.sleep(1)
    

    return False

def finalizar_relatorio():
    """
    Finaliza o relatório com logs
    """
    try:
        log(doc, "Finalizando relatório...")
        nome_arquivo = f"relatorio_emitente_nf_cenario_3_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        doc.save(nome_arquivo)
        log(doc, f"📄 Relatório salvo como: {nome_arquivo}", "SUCCESS")
        
        # Tenta abrir o arquivo
        try:
            subprocess.run(["start", "winword", nome_arquivo], shell=True)
            log(doc, "📄 Relatório aberto no Word", "SUCCESS")
        except Exception as e:
            log(doc, f"Não foi possível abrir o Word: {e}", "WARN")
        
        log(doc, "Fechando navegador...")
        driver.quit()
        log(doc, "✅ Teste finalizado com sucesso", "SUCCESS")
        
    except Exception as e:
        log(doc, f"Erro ao finalizar relatório: {e}", "ERROR")

# ==== INICIALIZAÇÃO ====
log(doc, "🚀 INICIANDO TESTE DE CADASTRO DE EMITENTE NOTA FISCAL")

# Configuração do driver
log(doc, "Configurando driver do Chrome...")
options = Options()
options.add_argument("--start-maximized")
options.add_argument("--disable-dev-shm-usage")
options.add_argument("--no-sandbox")

log(doc, "Iniciando navegador Chrome...")
driver = webdriver.Chrome(service=Service(ChromeDriverManager().install()), options=options)
wait = WebDriverWait(driver, 20)

log(doc, "Navegador iniciado com sucesso", "SUCCESS")

# Geração de dados
dados_emitente = gerar_dados_emitente()

# ==== EXECUÇÃO PRINCIPAL ====
try:
    # Acesso ao sistema
    safe_action(doc, "Acessando sistema", lambda: (
        log(doc, f"Navegando para URL: {URL}"),
        driver.get(URL),
        log(doc, "Página carregada com sucesso", "SUCCESS")
    ))

    # Login
    safe_action(doc, "Realizando login", lambda: (
        log(doc, "Preenchendo email de login..."),
        preencher_campo_js(driver, wait, "#j_id15\\:email", LOGIN_EMAIL, "Email"),
        log(doc, "Preenchendo senha de login..."),
        preencher_campo_js(driver, wait, "#j_id15\\:senha", LOGIN_PASSWORD, "Senha"),
        log(doc, "Enviando formulário de login..."),
        driver.find_element(By.ID, "j_id15:senha").send_keys(Keys.ENTER),
        log(doc, "Aguardando carregamento da página principal..."),
        wait.until(EC.presence_of_element_located((By.TAG_NAME, "body")))
    ))

    # Aguardar carregamento e ajustar zoom
    safe_action(doc, "Esperando sistema carregar e ajustando zoom", lambda: (
        log(doc, "Aguardando 5 segundos para carregamento completo..."),
        time.sleep(5),
        ajustar_zoom()
    ))

    # Abrir menu
    safe_action(doc, "Abrindo menu Emitente Nota Fiscal", lambda: (
        log(doc, "Pressionando F2 para abrir busca..."),
        driver.find_element(By.TAG_NAME, "body").send_keys(Keys.F2),
        time.sleep(1),
        log(doc, "Preenchendo campo de busca..."),
        wait.until(EC.visibility_of_element_located((By.XPATH, "//input[@placeholder='Busque um cadastro']"))).send_keys("Emitente Nota Fiscal"),
        log(doc, "Enviando busca..."),
        driver.find_element(By.XPATH, "//input[@placeholder='Busque um cadastro']").send_keys(Keys.ENTER),
        log(doc, "Aguardando resultado da busca..."),
        time.sleep(3)
    ))

    # Clicar em Cadastrar
    safe_action(doc, "Clicando em Cadastrar", lambda: (
        clicar_elemento_js(driver, wait, "#fmod_10046 > div.wdTelas > div.telaInicial.clearfix.overflow.overflowY > ul > li:nth-child(1) > a > span", "Botão Cadastrar"),
        log(doc, "Aguardando carregamento do formulário..."),
        time.sleep(2)
    ))

    # Preenchimento dos campos - PARTE 1
    campos_parte1 = [
        ("#fmod_10046 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_10056.categoriaHolder > div > div > div:nth-child(2) > input", dados_emitente['razao_social'], "Razão Social"),
        ("#fmod_10046 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_10056.categoriaHolder > div > div > div:nth-child(3) > input", dados_emitente['cnpj'], "CNPJ"),
        ("#fmod_10046 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_10056.categoriaHolder > div > div > div:nth-child(4) > div > input", dados_emitente['cep'], "CEP"),
        ("#fmod_10046 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_10056.categoriaHolder > div > div > div:nth-child(7) > input", dados_emitente['numero_endereco'], "Número Endereço"),
        ("#fmod_10046 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_10056.categoriaHolder > div > div > div:nth-child(11) > input", str(dados_emitente['cod_municipio']), "Código Município"),
        ("#fmod_10046 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_10056.categoriaHolder > div > div > div:nth-child(12) > input", str(dados_emitente['inscricao_municipal']), "Inscrição Municipal"),
    ]

    for seletor, valor, nome in campos_parte1:
        safe_action(doc, f"Preenchendo campo {nome}", lambda s=seletor, v=valor, n=nome: (
            preencher_campo_js(driver, wait, s, v, n),
            time.sleep(0.5)
        ))

    # Selects da primeira parte
    selects_parte1 = [
        ("#fmod_10046 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_10056.categoriaHolder > div > div > div:nth-child(13) > select", "Sim", "Incentivador Cultural"),
        ("#fmod_10046 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_10056.categoriaHolder > div > div > div:nth-child(14) > select", "Sim", "Optante Simples Nacional"),
        ("#fmod_10046 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_10056.categoriaHolder > div > div > div:nth-child(15) > select", "Sim", "ISS Retido"),
    ]

    for seletor, valor, nome in selects_parte1:
        safe_action(doc, f"Selecionando {nome}", lambda s=seletor, v=valor, n=nome: 
            selecionar_opcao_select(driver, wait, s, v, n)
        )

    # Preenchimento dos campos - PARTE 2
    campos_parte2 = [
        ("#fmod_10046 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_10056.categoriaHolder > div > div > div:nth-child(16) > input", str(dados_emitente['serie']), "Série"),
        ("#fmod_10046 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_10056.categoriaHolder > div > div > div:nth-child(17) > input", str(dados_emitente['cod_regime_tributario']), "Código Regime Tributário"),
        ("#fmod_10046 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_10056.categoriaHolder > div > div > div:nth-child(18) > input", str(dados_emitente['natureza_operacao']), "Natureza Operação"),
    ]

    for seletor, valor, nome in campos_parte2:
        safe_action(doc, f"Preenchendo campo {nome}", lambda s=seletor, v=valor, n=nome: (
            preencher_campo_js(driver, wait, s, v, n),
            time.sleep(0.5)
        ))

    # Select Calcular DAS
    safe_action(doc, "Selecionando Calcular DAS", lambda: 
        selecionar_opcao_select(driver, wait, 
            "#fmod_10046 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_10056.categoriaHolder > div > div > div:nth-child(22) > select", 
            "Sim", "Calcular DAS")
    )

    # Preenchimento dos campos finais
    campos_finais = [
        ("#fmod_10046 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_10056.categoriaHolder > div > div > div:nth-child(23) > input", str(dados_emitente['percentual_DAS']), "Percentual DAS"),
        ("#fmod_10046 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_10056.categoriaHolder > div > div > div:nth-child(24) > input", dados_emitente['url_webservice'], "URL WebService"),
    ]


    for seletor, valor, nome in campos_finais:
        print(f" Preenchendo campo {nome} com valor: '{valor}'")
        preencher_campo_com_retry(driver, wait, seletor, valor)




        # Salvamento
    safe_action(doc, "Salvando cadastro", lambda: (
        log(doc, "Clicando no botão salvar..."),
        clicar_elemento_js(driver, wait, "#fmod_10046 div.btnHolder > a.btModel.btGray.btsave", "Botão Salvar"),
    ))

    # Salvamento
    safe_action(doc, "Fechando modal", lambda: (
        log(doc, "Clicando no X..."),
        clicar_elemento_js(driver, wait, "#fmod_10046 > div.wdTop.ui-draggable-handle > div.wdClose > a", "X"),
        time.sleep(5)
    ))


    # Verificar mensagem de alerta
    safe_action(doc, "Verificando resultado do salvamento", lambda: (
        encontrar_mensagem_alerta()
    ))

except Exception as e:
    log(doc, f"❌ ERRO FATAL: {str(e)}", "ERROR")
    take_screenshot(driver, doc, "erro_fatal")
    raise e

finally:
    log(doc, "✅ TESTE CONCLUÍDO", "SUCCESS")
    finalizar_relatorio()

    