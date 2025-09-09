# Refatorado e organizado: cadastrodeescalamotorista2ºcenario.py

from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support.ui import Select, WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service
from webdriver_manager.chrome import ChromeDriverManager
from docx import Document
from docx.shared import Inches
from faker import Faker
from faker.providers import BaseProvider
from validate_docbr import CPF
from datetime import datetime, timedelta
import subprocess
import os
import time
import random

# ==== PROVIDERS CUSTOMIZADOS ====
class BrasilProvider(BaseProvider):
    def rg(self):
        numeros = [str(random.randint(0, 9)) for _ in range(8)]
        return ''.join(numeros) + '-' + str(random.randint(0, 9))

fake = Faker("pt_BR")
fake.add_provider(BrasilProvider)

def gerar_datas_validas():
    """Gera datas coerentes para admissão, início e fim da escala, e vencimento da CNH."""
    hoje = datetime.today().date()
    
    # Data de admissão entre 10 anos atrás e hoje
    data_admissao = fake.date_between(start_date=hoje - timedelta(days=3650), end_date=hoje)
    
    # Data de início da escala entre hoje e 1 ano no futuro
    data_inicio = fake.date_between(start_date=hoje, end_date=hoje + timedelta(days=365))
    
    # Data fim entre 1 e 180 dias após a data de início
    data_fim = data_inicio + timedelta(days=random.randint(1, 180))
    
    # Vencimento CNH entre hoje e 10 anos no futuro
    vencimento_cnh = fake.date_between(start_date=hoje, end_date=hoje + timedelta(days=3650))
    
    return (data_admissao.strftime('%d/%m/%Y'), 
            data_inicio.strftime('%d/%m/%Y'), 
            data_fim.strftime('%d/%m/%Y'), 
            vencimento_cnh.strftime('%d/%m/%Y'))

def gerar_dados_documentos():
    """Gera documentos fictícios para o cadastro."""
    carteira_trabalho = str(random.randint(10000000, 99999999))
    pis = fake.cpf().replace('.', '').replace('-', '')[:11]
    cnh = str(random.randint(10000000000, 99999999999))
    cpf = CPF().generate()
    
    return carteira_trabalho, pis, cnh, cpf

# Gera os dados necessários
data_admissao, data_inicio, data_fim, vencimento_cnh = gerar_datas_validas()
carteira_trabalho, pis, cnh, cpf_valido = gerar_dados_documentos()

# ==== CONFIGURAÇÕES ====
URL = "http://localhost:8080/gs/index.xhtml"
LOGIN_EMAIL = "joaoeduardo.gold@outlook.com"
LOGIN_PASSWORD = "071999gs"

# ==== DOCUMENTO ====
doc = Document()
doc.add_heading("RELATÓRIO DO TESTE", 0)
doc.add_paragraph("Cadastro de Grupos de Rateio – Cenário 2: Preenchimento completo e cancelamento.")
doc.add_paragraph(f"Data do teste: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")

screenshot_registradas = set()

# ==== FUNÇÕES DE UTILITÁRIO ====
def log(doc, msg):
    print(msg)
    doc.add_paragraph(msg)

def take_screenshot(driver, doc, nome):
    if nome not in screenshot_registradas:
        path = f"screenshots/{nome}.png"
        os.makedirs("screenshots", exist_ok=True)
        driver.save_screenshot(path)
        doc.add_paragraph(f"Screenshot: {nome}")
        doc.add_picture(path, width=Inches(5.5))
        screenshot_registradas.add(nome)

def safe_action(doc, descricao, func):
    try:
        log(doc, f"🔄 {descricao}...")
        func()
        log(doc, f"✅ {descricao} realizada com sucesso.")
        take_screenshot(driver, doc, descricao.lower().replace(" ", "_"))
    except Exception as e:
        log(doc, f"❌ Erro ao {descricao.lower()}: {e}")
        take_screenshot(driver, doc, f"erro_{descricao.lower().replace(' ', '_')}")

def finalizar_relatorio():
    nome_arquivo = f"relatorio_grupo_rateio_cenario_2_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc.save(nome_arquivo)
    log(doc, f"📄 Relatório salvo como: {nome_arquivo}")
    subprocess.run(["start", "winword", nome_arquivo], shell=True)
    driver.quit()

def encontrar_mensagem_alerta():
    seletores = [
        (".alerts.salvo", "✅ Sucesso"),
        (".alerts.alerta", "⚠️ Alerta"),
        (".alerts.erro", "❌ Erro"),
    ]

    for seletor, tipo in seletores:
        try:
            elemento = driver.find_element(By.CSS_SELECTOR, seletor)
            if elemento.is_displayed():
                log(doc, f"📢 {tipo}: {elemento.text}")
                return elemento
        except:
            continue

    log(doc, "ℹ️ Nenhuma mensagem de alerta encontrada.")
    return None

def ajustar_zoom():
    try:
        driver.execute_script("document.body.style.zoom='90%'")
        log(doc, "🔍 Zoom ajustado para 90%.")
    except Exception as e:
        log(doc, f"⚠️ Erro ao ajustar zoom: {e}")

def preencher_campo_data(selector, valor):
    def acao():
        campo = wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, selector)))
        campo.click()
        campo.clear()
        campo.send_keys(valor)
        time.sleep(0.2)
    return acao

def abrir_modal_e_selecionar(btn_selector, pesquisa_selector, termo_pesquisa, btn_pesquisar_selector, resultado_xpath):
    """Abre modal e seleciona um item"""
    def acao():
        # Abre o modal
        open_lov = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, btn_selector)))
        open_lov.click()
        time.sleep(3)

        # Aguarda campo pesquisa
        campo_pesquisa = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, pesquisa_selector)))
        campo_pesquisa.clear()
        campo_pesquisa.send_keys(termo_pesquisa)

        # Clica pesquisar
        pesquisar = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, btn_pesquisar_selector)))
        pesquisar.click()
        time.sleep(3)
        pesquisar.click()

        # Espera o resultado e clica
        wait.until(EC.element_to_be_clickable((By.XPATH, resultado_xpath)))
        resultado = driver.find_element(By.XPATH, resultado_xpath)
        driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", resultado)
        time.sleep(0.2)
        resultado.click()

    return acao


def selecionar_opcao(selector, texto):
    def acao():
        select_element = wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, selector)))
        Select(select_element).select_by_visible_text(texto)
    return acao

# ==== INICIALIZAÇÃO DO DRIVER ====
options = Options()
options.add_argument("--start-maximized")
driver = webdriver.Chrome(service=Service(ChromeDriverManager().install()), options=options)
wait = WebDriverWait(driver, 20)

# ==== EXECUÇÃO DO TESTE ====
try:
    safe_action(doc, "Acessando sistema", lambda: driver.get(URL))

    safe_action(doc, "Realizando login", lambda: (
        wait.until(EC.presence_of_element_located((By.ID, "j_id15:email"))).send_keys(LOGIN_EMAIL),
        wait.until(EC.presence_of_element_located((By.ID, "j_id15:senha"))).send_keys(LOGIN_PASSWORD, Keys.ENTER),
        wait.until(EC.presence_of_element_located((By.TAG_NAME, "body")))
    ))

    safe_action(doc, "Esperando sistema carregar e ajustando zoom", lambda: (
        time.sleep(5),
        ajustar_zoom()
    ))

    safe_action(doc, "Abrindo menu Grupo de Rateio", lambda: (
        driver.find_element(By.TAG_NAME, "body").send_keys(Keys.F2),
        time.sleep(1),
        wait.until(EC.visibility_of_element_located((By.XPATH, "//input[@placeholder='Busque um cadastro']"))).send_keys("Grupo de Rateio", Keys.ENTER),

    ))

    safe_action(doc, "Clicando em Cadastrar", lambda: (
        time.sleep(2),
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_10051 > div.wdTelas > div.telaInicial.clearfix.overflow.overflowY > ul > li:nth-child(1) > a > span"))).click()
    ))

    safe_action(doc, "Preechendo o Nome do Grupo de Rateio", lambda: (
        time.sleep(2),
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_10051 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div > div > div > div:nth-child(2) > input"))
        ).send_keys('TESTE GRUPO DE RATEIO SELENIUM AUTOMATIZADO 3')
    ))

    safe_action(doc, "Preechendo a Quantidade Máxima de Contratos", lambda: (
        time.sleep(2),
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_10051 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div > div > div > div:nth-child(4) > input"))
        ).send_keys(fake.random_int(min=1, max=20))
    ))

    safe_action(doc, "Selecionando a Forma de Cobrança como Apenas Primeiro Sepultamento", selecionar_opcao(
        "#fmod_10051 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div > div > div > div:nth-child(6) > select",
        "Apenas Primeiro Sepultamento"
    ))

    safe_action(doc, "Preechendo o Valor", lambda: (
        time.sleep(2),
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_10051 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div > div > div > div:nth-child(7) > input"))
        ).send_keys(fake.random_int(min=1000, max=100000))
    ))


    safe_action(doc, "Selecionando Tipo de Mensalidade", abrir_modal_e_selecionar(
        "#fmod_10051 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div > div > div > div:nth-child(8) > div > div > a",
        "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(2) > input",
        "§",
        "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(3) > a",
        "//td[contains(text(), 'TIPO MENSALIDADE JOÃO TESTE')]/a[contains(@class, 'linkAlterar')]"
    ))


    safe_action(doc, "Cancelando cadastro", lambda: driver.find_element(
        By.CSS_SELECTOR, "#fmod_10051 > div.wdTelas > div.telaCadastro.clearfix > div.btnHolder > a.btModel.btGray.btcancel"
    ).click())



    safe_action(doc, "Confirmando cancelamento", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#BtYes"))).click()
    ))




    safe_action(doc, "Fechando modal após cancelamento", lambda: wait.until(EC.element_to_be_clickable((
        By.CSS_SELECTOR, "#fmod_10051 > div.wdTop.ui-draggable-handle > div.wdClose > a"
        ))
    ).click())

    time.sleep(1.5)

    encontrar_mensagem_alerta()

except Exception as e:
    log(doc, f"❌ ERRO FATAL: {e}")
    take_screenshot(driver, doc, "erro_fatal")

finally:

    log(doc, "✅ Teste concluído com sucesso.")

    finalizar_relatorio()