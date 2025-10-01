import sys
import os
import time
import subprocess
from datetime import datetime

# Adiciona a raiz do projeto ao sys.path (para utils.actions)
sys.path.append(
    os.path.abspath(
        os.path.join(os.path.dirname(__file__), "../../..")
    )
)

from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service
from selenium.common.exceptions import (
    TimeoutException,
    NoSuchElementException,
    ElementClickInterceptedException,
    ElementNotInteractableException,
)
from webdriver_manager.chrome import ChromeDriverManager
from docx import Document

# Utilidades padronizadas do projeto
from utils.actions import log, take_screenshot, safe_action, ajustar_zoom

URL = "http://localhost:8080/gs/index.xhtml"
LOGIN_EMAIL = "joaoeduardo.gold@outlook.com"
LOGIN_PASSWORD = "071999gs"

screenshot_registradas = set()

def registrar_screenshot_unico(nome, driver, doc, descricao=None):
    if nome not in screenshot_registradas:
        if descricao:
            log(doc, f"📸 {descricao}")
        take_screenshot(driver, doc, nome)
        screenshot_registradas.add(nome)

# ==========================
# Utilitários robustos
# Mantêm compatibilidade com safe_action(doc, desc, func, driver, wait)
# ==========================

def find_visible(driver, by, locator):
    els = driver.find_elements(by, locator)
    for el in els:
        try:
            if el.is_displayed() and el.is_enabled():
                return el
        except Exception:
            continue
    return None


def force_click(driver, el):
    try:
        el.click()
        return
    except (ElementClickInterceptedException, ElementNotInteractableException):
        pass
    driver.execute_script("arguments[0].scrollIntoView({block:'center'});", el)
    time.sleep(0.1)
    try:
        el.click()
        return
    except (ElementClickInterceptedException, ElementNotInteractableException):
        driver.execute_script("arguments[0].click();", el)


def fill_input(driver, wait, selector, value, clear_first=True):
    el = wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, selector)))
    driver.execute_script("arguments[0].scrollIntoView({block:'center'});", el)
    if clear_first:
        try:
            el.clear()
        except Exception:
            pass
    try:
        el.send_keys(value)
    except ElementNotInteractableException:
        driver.execute_script("arguments[0].focus();", el)
        el.send_keys(value)


def click_css(driver, wait, selector):
    el = wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, selector)))
    visible = el if (el.is_displayed() and el.is_enabled()) else find_visible(driver, By.CSS_SELECTOR, selector)
    if visible is None:
        # último recurso: espera clicável pelo mesmo seletor
        try:
            visible = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, selector)))
        except TimeoutException:
            visible = el
    driver.execute_script("arguments[0].scrollIntoView({block:'center'});", visible)
    force_click(driver, visible)


def abrir_lov(driver, wait, css_selector_lov):
    click_css(driver, wait, css_selector_lov)
    time.sleep(0.3)
    # campo de busca do LOV (modal)
    campo = wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, "body > div.modalHolder input[type='text']")))
    return campo


def selecionar_lov(driver, wait, css_selector_lov, termo_busca, texto_tr):
    campo = abrir_lov(driver, wait, css_selector_lov)
    try:
        campo.clear()
    except Exception:
        pass
    try:
        campo.send_keys(termo_busca + Keys.ENTER)
    except ElementNotInteractableException:
        driver.execute_script("arguments[0].focus();", campo)
        campo.send_keys(termo_busca + Keys.ENTER)
    time.sleep(0.6)
    # escolhe a linha
    linha = wait.until(EC.element_to_be_clickable((By.XPATH, f"//tr[contains(., '{texto_tr}')]")))
    driver.execute_script("arguments[0].scrollIntoView({block:'center'});", linha)
    force_click(driver, linha)


def selecionar_select_por_tecla(driver, wait, selector, valor):
    # Combo simples <select>: envia texto e ENTER
    el = wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, selector)))
    driver.execute_script("arguments[0].scrollIntoView({block:'center'});", el)
    try:
        el.send_keys(valor)
    except ElementNotInteractableException:
        driver.execute_script("arguments[0].focus();", el)
        el.send_keys(valor)

# ==========================
# Fluxo do cenário
# ==========================

def main():
    doc = Document()
    doc.add_heading("RELATÓRIO DO TESTE", 0)
    doc.add_paragraph("Cadastro de Cartões - Bandeira – Cenário 2: Preenchimento completo e **cancelamento**.")
    doc.add_paragraph(f"📅 Data do teste: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")

    chrome_options = Options()
    chrome_options.add_argument("--start-maximized")
    driver = webdriver.Chrome(service=Service(ChromeDriverManager().install()), options=chrome_options)
    wait = WebDriverWait(driver, 10)

    def finalizar_relatorio():
        doc_name = f"relatorio_cartoes_bandeira_cenario_2_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        doc.save(doc_name)
        log(doc, f"📄 Relatório salvo como: {doc_name}")
        try:
            subprocess.run(["start", "winword", doc_name], shell=True)
        except Exception as e:
            log(doc, f"Erro ao abrir o Word: {e}")
        driver.quit()

    def encontrar_mensagem_alerta():
        seletores = [
            (".alerts.salvo",  "✅ Mensagem de Sucesso"),
            (".alerts.alerta", "⚠️ Mensagem de Alerta"),
            (".alerts.erro",   "❌ Mensagem de Erro"),
        ]
        for seletor, rotulo in seletores:
            try:
                elemento = driver.find_element(By.CSS_SELECTOR, seletor)
                if elemento.is_displayed():
                    try:
                        texto = elemento.find_element(By.TAG_NAME, "p").text.strip()
                    except NoSuchElementException:
                        texto = elemento.text.strip()
                    log(doc, f"📢 {rotulo}: {texto}")
                    take_screenshot(driver, doc, f"mensagem_{rotulo.split()[-1].lower()}")
                    return elemento
            except Exception:
                continue
        log(doc, "ℹ️ Nenhuma mensagem de alerta encontrada.")
        return None

    # ============= Ações do cenário =============
    def login():
        wait.until(EC.presence_of_element_located((By.ID, "j_id15:email"))).send_keys(LOGIN_EMAIL)
        wait.until(EC.presence_of_element_located((By.ID, "j_id15:senha"))).send_keys(LOGIN_PASSWORD, Keys.ENTER)
        wait.until(EC.presence_of_element_located((By.TAG_NAME, "body")))
        time.sleep(2)

    def abrir_menu():
        body = wait.until(EC.presence_of_element_located((By.TAG_NAME, "body")))
        body.send_keys(Keys.F2)
        time.sleep(0.3)  # debounce overlay
        # pode haver mais de um input com mesmo placeholder: pega o visível
        campos = driver.find_elements(By.XPATH, "//input[@placeholder='Busque um cadastro']")
        campo = None
        for c in campos:
            if c.is_displayed() and c.is_enabled():
                campo = c
                break
        if not campo:
            campo = wait.until(EC.element_to_be_clickable((By.XPATH, "//input[@placeholder='Busque um cadastro']")))
        force_click(driver, campo)
        try:
            campo.clear()
        except Exception:
            pass
        try:
            campo.send_keys("Cartão - Bandeira")
        except ElementNotInteractableException:
            driver.execute_script("arguments[0].focus();", campo)
            campo.send_keys("Cartão - Bandeira")
        campo.send_keys(Keys.ENTER)
        time.sleep(1.2)

    def acessar_formulario():
        click_css(driver, wait, "#fmod_10010 > div.wdTelas > div > ul > li:nth-child(1) > a > span")

    def preencher_campos_completos():
        campos = [
            ("Preenchendo nome",  
             "#fmod_10010 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_10010.categoriaHolder > div:nth-child(1) > div > div:nth-child(2) > input",
             "BANDEIRA TESTE"),
            ("Preenchendo tarifa",
             "#fmod_10010 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_10010.categoriaHolder > div:nth-child(1) > div > div:nth-child(3) > input",
             "3,25"),
            ("Preenchendo dias",
             "#fmod_10010 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_10010.categoriaHolder > div:nth-child(1) > div > div:nth-child(4) > input",
             "30"),
        ]
        for descricao, seletor, valor in campos:
            safe_action(doc, descricao, lambda s=seletor, v=valor: fill_input(driver, wait, s, v), driver, wait)

        selects = [
            ("Selecionando tipo cartão",
             "#fmod_10010 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_10010.categoriaHolder > div:nth-child(1) > div > div:nth-child(5) > select",
             "Crédito"),
            ("Selecionando tipo cobrança",
             "#fmod_10010 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_10010.categoriaHolder > div:nth-child(1) > div > div:nth-child(7) > select",
             "Boleto"),
        ]
        for descricao, seletor, valor in selects:
            safe_action(doc, descricao, lambda s=seletor, v=valor: selecionar_select_por_tecla(driver, wait, s, v), driver, wait)

        # LOVs
        safe_action(doc, "Selecionando conta crédito 1", lambda: selecionar_lov(
            driver, wait,
            "#fmod_10010 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_10010.categoriaHolder > div:nth-child(1) > div > div:nth-child(8) > div > div > a",
            "745", "ITAU"), driver, wait)

        safe_action(doc, "Selecionando conta crédito 2", lambda: selecionar_lov(
            driver, wait,
            "#fmod_10010 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_10010.categoriaHolder > div:nth-child(2) > div > div:nth-child(2) > div > div > a",
            "745", "ITAU"), driver, wait)

        safe_action(doc, "Selecionando centro de custo", lambda: selecionar_lov(
            driver, wait,
            "#fmod_10010 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_10010.categoriaHolder > div:nth-child(2) > div > div:nth-child(3) > div > div > a",
            "80.79.4703", "TESTE CENTRO DE CUSTO SELENIUM AUTOMATIZADO"), driver, wait)

        safe_action(doc, "Selecionando histórico padrão", lambda: selecionar_lov(
            driver, wait,
            "#fmod_10010 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_10010.categoriaHolder > div:nth-child(2) > div > div:nth-child(4) > div > div > a",
            "200116", "DESPESA COM TARIFAS DE CARTÃO"), driver, wait)

        # Aba Parametrização
        safe_action(doc, "Acessando aba Parametrização", lambda: click_css(
            driver, wait, "#fmod_10010 a[href='#cat_10094']"), driver, wait)

        safe_action(doc, "Preenchendo campo De", lambda: fill_input(driver, wait,
            "#fmod_10010 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_10094.categoriaHolder > div > div > div:nth-child(3) > input", "1"), driver, wait)

        safe_action(doc, "Preenchendo campo Até", lambda: fill_input(driver, wait,
            "#fmod_10010 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_10094.categoriaHolder > div > div > div:nth-child(4) > input", "31"), driver, wait)

        safe_action(doc, "Preenchendo Tarifa em %", lambda: fill_input(driver, wait,
            "#fmod_10010 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_10094.categoriaHolder > div > div > div:nth-child(5) > input", "2,5"), driver, wait)

        safe_action(doc, "Clicando em Adicionar", lambda: click_css(driver, wait,
            "#fmod_10010 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_10094.categoriaHolder > div > div > div.btnListHolder > a.btAddGroup"), driver, wait)

        registrar_screenshot_unico("preenchimento_completo", driver, doc)

    def cancelar():
        # clica Cancelar e confirma
        safe_action(doc, "Clicando em Cancelar", lambda: click_css(driver, wait,
            "#fmod_10010 > div.wdTelas > div.telaCadastro.clearfix > div.btnHolder > a.btModel.btGray.btcancel"), driver, wait)
        safe_action(doc, "Confirmando cancelamento", lambda: click_css(driver, wait, "#BtYes"), driver, wait)
        log(doc, "❌ Cadastro cancelado pelo usuário.")
        registrar_screenshot_unico("cadastro_cancelado", driver, doc)

    def fechar_modal():
        # fecha o modal se ainda estiver aberto
        try:
            modal = driver.find_element(By.CSS_SELECTOR, "#fmod_10010")
            if modal.is_displayed():
                safe_action(doc, "Fechando formulário", lambda: click_css(driver, wait,
                    "#fmod_10010 > div.wdTop.ui-draggable-handle > div.wdClose > a"), driver, wait)
        except NoSuchElementException:
            pass

    # Execução
    if not safe_action(doc, "Acessando o sistema", lambda: driver.get(URL), driver, wait)[0]: finalizar_relatorio(); return
    if not safe_action(doc, "Realizando login", login, driver, wait)[0]: finalizar_relatorio(); return
    if not safe_action(doc, "Ajustando zoom", lambda: ajustar_zoom(driver), driver, wait)[0]: finalizar_relatorio(); return
    if not safe_action(doc, "Abrindo menu Cartão Bandeira", abrir_menu, driver, wait)[0]: finalizar_relatorio(); return
    if not safe_action(doc, "Acessando formulário", acessar_formulario, driver, wait)[0]: finalizar_relatorio(); return
    if not safe_action(doc, "Preenchendo campos completos", preencher_campos_completos, driver, wait)[0]: finalizar_relatorio(); return
    if not safe_action(doc, "Cancelando cadastro", cancelar, driver, wait)[0]: finalizar_relatorio(); return

    # Mensagens (opcional neste cenário de cancelamento)
    encontrar_mensagem_alerta()

    if not safe_action(doc, "Fechando formulário", fechar_modal, driver, wait)[0]:
        finalizar_relatorio(); return

    log(doc, "✅ Teste concluído com sucesso.")
    finalizar_relatorio()


if __name__ == "__main__":
    main()
