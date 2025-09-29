import sys
import os

# Adiciona a raiz do projeto ao sys.path
sys.path.append(
    os.path.abspath(
        os.path.join(os.path.dirname(__file__), "../../..")
    )
)


from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service
from webdriver_manager.chrome import ChromeDriverManager
from docx import Document
from datetime import datetime
import subprocess
import time
import os
import sys

sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), '..', '..')))
from utils.actions import log, take_screenshot, safe_action,  ajustar_zoom

URL = "http://localhost:8080/gs/index.xhtml"
LOGIN_EMAIL = "joaoeduardo.gold@outlook.com"
LOGIN_PASSWORD = "071999gs"

screenshot_registradas = set()
def registrar_screenshot_unico(nome, driver, doc, descricao=None):
    if nome not in screenshot_registradas:
        if descricao:
            log(doc, f"📸 {descricao}")
        take_screenshot(driver, doc, nome)
        screenshot_registradas.add(nome)

def main():
    doc = Document()
    doc.add_heading("RELATÓRIO DO TESTE", 0)
    doc.add_paragraph("Cadastro de Cartões - Bandeira – Cenário 4: Preenchimento dos campos NÃO obrigatórios e salvamento.")
    doc.add_paragraph(f"📅 Data do teste: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")

    chrome_options = Options()
    chrome_options.add_argument("--start-maximized")
    driver = webdriver.Chrome(service=Service(ChromeDriverManager().install()), options=chrome_options)
    wait = WebDriverWait(driver, 10)

    def finalizar_relatorio():
        doc_name = f"relatorio_cartoes_bandeira_cenario_4_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        doc.save(doc_name)
        log(doc, f"📄 Relatório salvo como: {doc_name}")
        try:
            subprocess.run(["start", "winword", doc_name], shell=True)
        except Exception as e:
            log(doc, f"Erro ao abrir o Word: {e}")
        driver.quit()

    def encontrar_mensagem_alerta():
        seletores = [
            (".alerts.salvo", "✅ Mensagem de Sucesso"),
            (".alerts.alerta", "⚠️ Mensagem de Alerta"),
            (".alerts.erro", "❌ Mensagem de Erro"),
        ]

        for seletor, tipo in seletores:
            try:
                elemento = driver.find_element(By.CSS_SELECTOR, seletor)
                if elemento.is_displayed():
                    log(doc, f"📢 {tipo}: {elemento.text}")
                    return elemento
            except:
                continue

        log(doc, "ℹ️ Nenhuma mensagem de alerta encontrada.")
        return None

    def login():
        wait.until(EC.presence_of_element_located((By.ID, "j_id15:email"))).send_keys(LOGIN_EMAIL)
        wait.until(EC.presence_of_element_located((By.ID, "j_id15:senha"))).send_keys(LOGIN_PASSWORD, Keys.ENTER)
        wait.until(EC.presence_of_element_located((By.TAG_NAME, "body")))
        time.sleep(2)

    def abrir_menu():
        driver.find_element(By.TAG_NAME, "body").send_keys(Keys.F2)
        campo = wait.until(EC.presence_of_element_located((By.XPATH, "//input[@placeholder='Busque um cadastro']")))
        campo.clear()
        campo.send_keys("Cartão - Bandeira", Keys.ENTER)

    def acessar_formulario():
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR,
            "#fmod_10010 > div.wdTelas > div > ul > li:nth-child(1) > a > span"))).click()

    def selecionar_lov(css_selector_lov, termo_busca, texto_tr):
        driver.find_element(By.CSS_SELECTOR, css_selector_lov).click()
        time.sleep(2)
        campo_pesquisa = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "body > div.modalHolder input[type='text']")))
        campo_pesquisa.clear()
        campo_pesquisa.send_keys(termo_busca + Keys.ENTER)
        time.sleep(2)
        wait.until(EC.element_to_be_clickable((By.XPATH, f"//tr[contains(., '{texto_tr}')]"))).click()

    def preencher_campos_obrigatórios():
        time.sleep(2)
        safe_action(doc, "Selecionando conta crédito 2", lambda: selecionar_lov(
            "#fmod_10010 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_10010.categoriaHolder > div:nth-child(2) > div > div:nth-child(2) > div > div > a", "745", "ITAU"), driver, wait)

        safe_action(doc, "Selecionando centro de custo", lambda: selecionar_lov(
            "#fmod_10010 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_10010.categoriaHolder > div:nth-child(2) > div > div:nth-child(3) > div > div > a", "80.79.4703", "TESTE CENTRO DE CUSTO SELENIUM AUTOMATIZADO"), driver, wait)

        safe_action(doc, "Selecionando histórico padrão", lambda: selecionar_lov(
            "#fmod_10010 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_10010.categoriaHolder > div:nth-child(2) > div > div:nth-child(4) > div > div > a", "200116", "DESPESA COM TARIFAS DE CARTÃO"), driver, wait)

        safe_action(doc, "Acessando aba Parametrização", lambda: wait.until(
            EC.element_to_be_clickable((By.LINK_TEXT, "Parametrização"))).click(), driver, wait)

        safe_action(doc, "Preenchendo campo De", lambda: driver.find_element(By.CSS_SELECTOR,
            "#fmod_10010 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_10094.categoriaHolder > div > div > div:nth-child(3) > input").send_keys("1"), driver, wait)

        safe_action(doc, "Preenchendo campo Até", lambda: driver.find_element(By.CSS_SELECTOR,
            "#fmod_10010 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_10094.categoriaHolder > div > div > div:nth-child(4) > input").send_keys("31"), driver, wait)

        safe_action(doc, "Preenchendo Tarifa em %", lambda: driver.find_element(By.CSS_SELECTOR,
            "#fmod_10010 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_10094.categoriaHolder > div > div > div:nth-child(5) > input").send_keys("2,5"), driver, wait)

        safe_action(doc, "Clicando em Adicionar", lambda: driver.find_element(By.CSS_SELECTOR,
            "#fmod_10010 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_10094.categoriaHolder > div > div > div.btnListHolder > a.btAddGroup").click(), driver, wait)

        registrar_screenshot_unico("preenchimento_completo", driver, doc)

    def salvar():
        safe_action(doc, "Clicando em Salvar", lambda: wait.until(EC.element_to_be_clickable((
            By.CSS_SELECTOR, "#fmod_10010 > div.wdTelas > div.telaCadastro.clearfix > div.btnHolder > a.btModel.btGray.btsave"))).click(), driver, wait)

    def fechar_modal():
        safe_action(doc, "Fechando modal", lambda: wait.until(EC.element_to_be_clickable((
            By.CSS_SELECTOR, "#fmod_10010 > div.wdTop.ui-draggable-handle > div.wdClose > a"))).click(), driver, wait)

    # Execução
    if not safe_action(doc, "Acessando o sistema", lambda: driver.get(URL), driver, wait)[0]: finalizar_relatorio(); return
    if not safe_action(doc, "Realizando login", login, driver, wait)[0]: finalizar_relatorio(); return
    if not safe_action(doc, "Ajustando zoom", lambda: ajustar_zoom(driver), driver, wait)[0]: finalizar_relatorio(); return
    if not safe_action(doc, "Abrindo menu Cartão Bandeira", abrir_menu, driver, wait)[0]: finalizar_relatorio(); return
    if not safe_action(doc, "Acessando formulário", acessar_formulario, driver, wait)[0]: finalizar_relatorio(); return
    if not safe_action(doc, "Preenchendo campos completos", preencher_campos_obrigatórios, driver, wait)[0]: finalizar_relatorio(); return
    if not safe_action(doc, "Salvando dados", salvar, driver, wait)[0]: finalizar_relatorio(); return
    time.sleep(1)  # Espera para garantir que a mensagem de sucesso seja exibida    
    
    

    encontrar_mensagem_alerta()
    

    if not safe_action(doc, "Fechando formulário", fechar_modal, driver, wait)[0]: finalizar_relatorio(); return
    log(doc, "✅ Teste concluído com sucesso.")
    finalizar_relatorio()

if __name__ == "__main__":
    main()
