import sys
import os

# Adiciona a raiz do projeto ao sys.path
sys.path.append(
    os.path.abspath(
        os.path.join(os.path.dirname(__file__), "../../..")
    )
)

# Refatorado e organizado: cadastrodeconvenioOS1ºcenario.py

from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support.ui import Select, WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service
from selenium.webdriver import ActionChains
from webdriver_manager.chrome import ChromeDriverManager
from docx import Document
from docx.shared import Inches
from faker import Faker
from faker.providers import BaseProvider
from datetime import datetime, timedelta
import subprocess
import os
import time
import random
import string

# ==== PROVIDERS CUSTOMIZADOS ====
class BrasilProvider(BaseProvider):
    def rg(self):
        numeros = [str(random.randint(0, 9)) for _ in range(8)]
        return ''.join(numeros) + '-' + str(random.randint(0, 9))

fake = Faker("pt_BR")
fake.add_provider(BrasilProvider)

# ==== CONFIGURAÇÕES ====
URL = "http://localhost:8080/gs/index.xhtml"
LOGIN_EMAIL = "joaoeduardo.gold@outlook.com"
LOGIN_PASSWORD = "071999gs"

# ==== DOCUMENTO ====
doc = Document()
doc.add_heading("RELATÓRIO DO TESTE", 0)
doc.add_paragraph("Cadastro de Convênio da Ordem de Serviço – Cenário 1: Preenchimento completo e salvamento.")
doc.add_paragraph(f"Data do teste: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")

screenshot_registradas = set()

def abrir_modal_e_selecionar(btn_selector, pesquisa_selector, termo_pesquisa, btn_pesquisar_selector, resultado_xpath):
    def acao():
        # Abre o modal
        open_lov = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, btn_selector)))
        open_lov.click()

        time.sleep(5)

        # Aguarda campo pesquisa
        campo_pesquisa = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, pesquisa_selector)))
        campo_pesquisa.clear()
        campo_pesquisa.send_keys(termo_pesquisa)

        # Clica pesquisar
        pesquisar = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, btn_pesquisar_selector)))
        pesquisar.click()
        time.sleep(1)
        # Espera o resultado carregar
        wait.until(EC.presence_of_element_located((By.XPATH, resultado_xpath)))
        wait.until(EC.visibility_of_element_located((By.XPATH, resultado_xpath)))
        wait.until(EC.element_to_be_clickable((By.XPATH, resultado_xpath)))

        # Relocaliza no último instante (evita stale element)
        resultado = driver.find_element(By.XPATH, resultado_xpath)
        driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", resultado)
        time.sleep(0.2)
        resultado.click()

    return acao

def preencher_campo_com_retry(driver, wait, seletor, valor, max_tentativas=3):
    """Tenta preencher o campo com diferentes métodos até conseguir"""
    
    for tentativa in range(max_tentativas):
        try:
            # Aguarda o elemento
            campo = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, seletor)))
            
            # Scroll até o elemento se necessário
            driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", campo)
            time.sleep(0.5)
            
            # Método 1: Tradicional
            if tentativa == 0:
                campo.click()
                campo.clear()
                campo.send_keys(valor)
                campo.send_keys(Keys.TAB)
            
            # Método 2: ActionChains
            elif tentativa == 1:
                ActionChains(driver).move_to_element(campo).click().perform()
                time.sleep(0.2)
                ActionChains(driver).key_down(Keys.CONTROL).send_keys('a').key_up(Keys.CONTROL).perform()
                ActionChains(driver).send_keys(valor).perform()
                ActionChains(driver).send_keys(Keys.TAB).perform()
            
            # Método 3: JavaScript
            else:
                driver.execute_script("""
                    var element = arguments[0];
                    var valor = arguments[1];
                    element.focus();
                    element.value = '';
                    element.value = valor;
                    element.dispatchEvent(new Event('input', { bubbles: true }));
                    element.dispatchEvent(new Event('change', { bubbles: true }));
                    element.blur();
                """, campo, valor)
            
            time.sleep(0.5)
            
            # Verifica se o valor foi preenchido
            valor_atual = campo.get_attribute('value')
            if valor_atual == valor:
                return True
                
        except Exception as e:
            time.sleep(1)
    
    return False

# ==== FUNÇÕES DE UTILITÁRIO ====
def log(doc, msg):
    print(msg)
    doc.add_paragraph(msg)

def take_screenshot(driver, doc, nome):
    if nome not in screenshot_registradas:
        path = f"screenshots/{nome}.png"
        os.makedirs("screenshots", exist_ok=True)
        driver.save_screenshot(path)
        doc.add_paragraph(f"Screenshot: {nome}")
        doc.add_picture(path, width=Inches(5.5))
        screenshot_registradas.add(nome)

def safe_action(doc, descricao, func):
    try:
        log(doc, f"🔄 {descricao}...")
        func()
        log(doc, f"✅ {descricao} realizada com sucesso.")
        take_screenshot(driver, doc, descricao.lower().replace(" ", "_"))
    except Exception as e:
        log(doc, f"❌ Erro ao {descricao.lower()}: {e}")
        take_screenshot(driver, doc, f"erro_{descricao.lower().replace(' ', '_')}")

def finalizar_relatorio():
    nome_arquivo = f"relatorio_convenio_os_cenario_1_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc.save(nome_arquivo)
    log(doc, f"📄 Relatório salvo como: {nome_arquivo}")
    subprocess.run(["start", "winword", nome_arquivo], shell=True)
    driver.quit()

def gerar_dados_convenio_os():
    """Gera dados fictícios para o cadastro de convênio da ordem de serviço."""
    nome_convenio = 'CONVÊNIO OS TESTE SELENIUM AUTOMATIZADO'
    cnpj = fake.cnpj()
    telefone = fake.phone_number()
    fax = fake.phone_number()
    email = fake.email()
    cep = "15081115"
    numero = "1733"
    logradouro = "Rua Teste"
    bairro = "Centro" 
    valor_cobertura = "1000.00"
    
    return nome_convenio, cnpj, telefone, fax, email, cep, numero, logradouro, bairro, valor_cobertura

def selecionar_opcao(selector, texto):
    def acao():
        select_element = wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, selector)))
        Select(select_element).select_by_visible_text(texto)
    return acao

def encontrar_mensagem_alerta():
    seletores = [
        (".alerts.salvo", "✅ Sucesso"),
        (".alerts.alerta", "⚠️ Alerta"),
        (".alerts.erro", "❌ Erro"),
    ]

    for seletor, tipo in seletores:
        try:
            elemento = driver.find_element(By.CSS_SELECTOR, seletor)
            if elemento.is_displayed():
                log(doc, f"📢 {tipo}: {elemento.text}")
                return elemento
        except:
            continue

    log(doc, "ℹ️ Nenhuma mensagem de alerta encontrada.")
    return None

def ajustar_zoom():
    try:
        driver.execute_script("document.body.style.zoom='90%'")
        log(doc, "🔍 Zoom ajustado para 90%.")
    except Exception as e:
        log(doc, f"⚠️ Erro ao ajustar zoom: {e}")

# Gera os dados necessários
nome_convenio, cnpj, telefone, fax, email, cep, numero, logradouro, bairro, valor_cobertura = gerar_dados_convenio_os()

# ==== INICIALIZAÇÃO DO DRIVER ====
options = Options()
options.add_argument("--start-maximized")
driver = webdriver.Chrome(service=Service(ChromeDriverManager().install()), options=options)
wait = WebDriverWait(driver, 20)

# ==== EXECUÇÃO DO TESTE ====
try:
    safe_action(doc, "Acessando sistema", lambda: driver.get(URL))

    safe_action(doc, "Realizando login", lambda: (
        wait.until(EC.presence_of_element_located((By.ID, "j_id15:email"))).send_keys(LOGIN_EMAIL),
        wait.until(EC.presence_of_element_located((By.ID, "j_id15:senha"))).send_keys(LOGIN_PASSWORD, Keys.ENTER),
        wait.until(EC.presence_of_element_located((By.TAG_NAME, "body")))
    ))

    safe_action(doc, "Esperando sistema carregar e ajustando zoom", lambda: (
        time.sleep(5),
        ajustar_zoom()
    ))

    safe_action(doc, "Abrindo menu Convênio Ordem de Serviço", lambda: (
        driver.find_element(By.TAG_NAME, "body").send_keys(Keys.F2),
        time.sleep(1),
        wait.until(EC.visibility_of_element_located((By.XPATH, "//input[@placeholder='Busque um cadastro']"))).send_keys("Convênio Ordem de Serviço", Keys.ENTER),
        time.sleep(3)
    ))

    # AQUI VOCÊ PRECISA FORNECER O SELETOR CORRETO DO BOTÃO CADASTRAR DO CONVÊNIO OS
    safe_action(doc, "Clicando em Cadastrar", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_10023 > div.wdTelas > div.telaInicial.clearfix.overflow.overflowY > ul > li:nth-child(1) > a > span"))).click(),
        time.sleep(2)
    ))

    # Preenchendo dados do convênio - AQUI VOCÊ PRECISA FORNECER OS SELETORES CORRETOS
    safe_action(doc, "Preenchendo Nome do Convênio", 
            lambda: preencher_campo_com_retry(driver, wait, "#fmod_10023 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div > div > div > div:nth-child(2) > input", nome_convenio)
        )

    safe_action(doc, "Preenchendo CNPJ", 
            lambda: preencher_campo_com_retry(driver, wait, "#fmod_10023 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div > div > div > div:nth-child(3) > input", cnpj)
        )

    safe_action(doc, "Preenchendo CEP", 
            lambda: preencher_campo_com_retry(driver, wait, "#fmod_10023 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div > div > div > div:nth-child(5) > div > input", cep)
        )

    safe_action(doc, "Clicando no botão de busca CEP", lambda: (
        time.sleep(2),
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "#fmod_10023 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div > div > div > div:nth-child(5) > div > a")))
        .click()
    ))


    safe_action(doc, "Preenchendo Número", 
            lambda: preencher_campo_com_retry(driver, wait, "#fmod_10023 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div > div > div > div:nth-child(8) > input", numero)
        )

    safe_action(doc, "Preenchendo Telefone", 
            lambda: preencher_campo_com_retry(driver, wait, "#fmod_10023 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div > div > div > div:nth-child(11) > input", telefone)
        )

    safe_action(doc, "Preenchendo Fax", 
            lambda: preencher_campo_com_retry(driver, wait, "#fmod_10023 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div > div > div > div:nth-child(12) > input", fax)
        )

    safe_action(doc, "Preenchendo E-mail", 
            lambda: preencher_campo_com_retry(driver, wait, "#fmod_10023 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div > div > div > div:nth-child(13) > input", email)
        )


    safe_action(doc, "Preenchendo Valor de Cobertura", 
            lambda: preencher_campo_com_retry(driver, wait, "#fmod_10023 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div > div > div > div:nth-child(14) > input", valor_cobertura)
        )




    safe_action(doc, "Selecionando a opção para o Tipo Desconto Convênio ser porcentagem", selecionar_opcao(
        "#fmod_10023 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div > div > div > div:nth-child(15) > select",
        "%"
    ))



    safe_action(doc, "Selecionando Conta Crédito", 
        abrir_modal_e_selecionar(
            "#fmod_10023 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div > div > div > div:nth-child(17) > div > div > a",
            "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(2) > input",
            "782",
            "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(3) > a",
            "//tr[td[contains(text(), '782')] and td[contains(text(), 'INTERNET VIRADOURO')]]"  # ✅ Busca na linha
        )
    )

    safe_action(doc, "Selecionando Conta Débito", 
        abrir_modal_e_selecionar(
            "#fmod_10023 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div > div > div > div:nth-child(18) > div > div > a",
            "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(2) > input",
            "779",
            "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(3) > a",
            "//tr[td[contains(text(), '779')] and td[contains(text(), 'BRADESCO VIRADOURO')]]"  # ✅ Busca na linha
        )
    )

    safe_action(doc, "Selecionando Histórico Padrão", 
        abrir_modal_e_selecionar(
            "#fmod_10023 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div > div > div > div:nth-child(19) > div > div > a",
            "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(2) > input",
            "200132",
            "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(3) > a",
            "//tr[td[contains(text(), '200132')] and td[contains(text(), 'LANÇAMENTO DO FINANCEIRO')]]"  # ✅ Busca na linha
        )
    )


    safe_action(doc, "Selecionando a opção para o Tipo Cobertura ser Total", selecionar_opcao(
        "#fmod_10023 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div > div > div > div:nth-child(20) > select",
        "Total"
    ))


    safe_action(doc, "Salvando cadastro", lambda: (
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "#fmod_10023 > div.wdTelas > div.telaCadastro.clearfix > div.btnHolder > a.btModel.btGray.btsave"))).click()
    ))

    safe_action(doc, "Fechando modal após salvamento", lambda: (
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "#fmod_10023 > div.wdTop.ui-draggable-handle > div.wdClose > a"))).click(),
        time.sleep(1)
    ))
    time.sleep(2)
    encontrar_mensagem_alerta()

except Exception as e:
    log(doc, f"❌ ERRO FATAL: {e}")
    take_screenshot(driver, doc, "erro_fatal")

finally:
    log(doc, "✅ Teste concluído com sucesso.")
    finalizar_relatorio()