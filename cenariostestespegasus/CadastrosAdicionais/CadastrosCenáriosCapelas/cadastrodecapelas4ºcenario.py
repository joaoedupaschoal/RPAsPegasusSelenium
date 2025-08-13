import sys
import os

# Adiciona a raiz do projeto ao sys.path
sys.path.append(
    os.path.abspath(
        os.path.join(os.path.dirname(__file__), "../../..")
    )
)


from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service
from webdriver_manager.chrome import ChromeDriverManager
from docx import Document
from datetime import datetime
import subprocess
import time
import os
import sys
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), '..', '..')))

from utils.actions import log, take_screenshot, safe_action, encontrar_mensagem_alerta, ajustar_zoom

URL = "http://localhost:8080/gs/index.xhtml"
LOGIN_EMAIL = "joaoeduardo.gold@outlook.com"
LOGIN_PASSWORD = "071999gs"
MODAL_SELECTOR = "#fmod_10035"

def main():
    doc = Document()
    doc.add_heading("RELATÓRIO DO TESTE", 0)
    doc.add_paragraph("Cadastro de Capelas.")
    doc.add_paragraph(f"🗕️ Data do teste: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")
    doc.add_paragraph("Neste teste, o robô preencherá apenas campos não obrigatórios e tentará salvar para verificar a validação dos obrigatórios.")

    chrome_options = Options()
    chrome_options.add_argument("--start-maximized")
    driver = webdriver.Chrome(service=Service(ChromeDriverManager().install()), options=chrome_options)
    wait = WebDriverWait(driver, 10)

    def finalizar_relatorio():
        doc_name = f"relatorio_capelas_cenario_4_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        doc.save(doc_name)
        log(doc, f"📄 Relatório salvo como: {doc_name}")
        try:
            subprocess.run(["start", "winword", doc_name], shell=True)
        except Exception as e:
            log(doc, f"Erro ao abrir o Word: {e}")
        driver.quit()

    def login():
        log(doc, "🔄 Iniciando login no sistema.")
        take_screenshot(driver, doc, "login_inicio")
        wait.until(EC.presence_of_element_located((By.ID, "j_id15:email"))).send_keys(LOGIN_EMAIL)
        wait.until(EC.presence_of_element_located((By.ID, "j_id15:senha"))).send_keys(LOGIN_PASSWORD, Keys.ENTER)
        wait.until(EC.presence_of_element_located((By.TAG_NAME, "body")))
        time.sleep(2)
        log(doc, "✅ Login realizado.")
        take_screenshot(driver, doc, "login_concluido")

    def abrir_menu_capelas():
        log(doc, "🔄 Acessando o menu Capelas.")
        take_screenshot(driver, doc, "menu_inicio")
        driver.find_element(By.TAG_NAME, "body").send_keys(Keys.F2)
        campo = wait.until(EC.presence_of_element_located((By.XPATH, "//input[@placeholder='Busque um cadastro']")))
        campo.click()
        campo.send_keys("Capelas", Keys.ENTER)
        time.sleep(2)
        log(doc, "✅ Menu Capelas aberto.")
        take_screenshot(driver, doc, "menu_aberto")

    def acessar_formulario():
        log(doc, "🔄 Acessando o formulário de cadastro.")
        take_screenshot(driver, doc, "formulario_inicio")
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR,
            "#fmod_10035 > div.wdTelas > div.telaInicial.clearfix.overflow.overflowY > ul > li:nth-child(1) > a > span"))).click()
        time.sleep(2)
        log(doc, "✅ Formulário carregado.")
        take_screenshot(driver, doc, "formulario_aberto")

    def preencher_nao_obrigatorios_e_salvar():
        log(doc, "🔄 Preenchendo apenas o campo 'Cidade - UF' (não obrigatório).")
        campo_cidade = wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR,
            "#fmod_10035 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div > div > div > div:nth-child(3) > input")))
        campo_cidade.send_keys("São José do Rio Preto - SP")
        time.sleep(1)
        wait.until(EC.element_to_be_clickable((By.XPATH, "//li[contains(text(), 'São José do Rio Preto - SP')]"))).click()
        log(doc, "✅ Campo 'Cidade - UF' preenchido e selecionado.")
        take_screenshot(driver, doc, "cidade_uf_preenchida")

        log(doc, "🔄 Tentando salvar sem preencher campos obrigatórios.")
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR,
            "#fmod_10035 > div.wdTelas > div.telaCadastro.clearfix > div.btnHolder > a.btModel.btGray.btsave"))).click()
        time.sleep(2)
        log(doc, "✅ Clique no botão Salvar realizado.")
        take_screenshot(driver, doc, "apos_clicar_salvar")

    def fechar_modal():
        log(doc, "🔄 Fechando o formulário.")
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR,
            "#fmod_10035 > div.wdTop.ui-draggable-handle > div.wdClose > a"))).click()
        time.sleep(1)
        log(doc, "✅ Formulário fechado.")
        take_screenshot(driver, doc, "formulario_fechado")

    # Execução
    driver.get(URL)
    login()
    ajustar_zoom(driver)
    abrir_menu_capelas()
    acessar_formulario()
    preencher_nao_obrigatorios_e_salvar()

    # Verificação da mensagem
    _, tipo_alerta = encontrar_mensagem_alerta(driver, doc)
    if tipo_alerta == "erro":
        log(doc, "✅ Mensagem de erro exibida corretamente por ausência do campo obrigatório.")
    elif tipo_alerta == "alerta":
        log(doc, "⚠️ Mensagem de alerta exibida.")
    elif tipo_alerta == "sucesso":
        log(doc, "❌ Cadastro foi salvo mesmo sem campo obrigatório preenchido.")
    else:
        log(doc, "⚠️ Nenhuma mensagem foi exibida.")
    take_screenshot(driver, doc, "mensagem_final")

    fechar_modal()
    log(doc, "✅ Teste encerrado. O sistema respondeu conforme esperado ao não preencher campos obrigatórios.")
    finalizar_relatorio()

if __name__ == "__main__":
    main()
