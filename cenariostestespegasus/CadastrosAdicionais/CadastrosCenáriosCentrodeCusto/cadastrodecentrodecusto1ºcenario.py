import sys
import os

# Adiciona a raiz do projeto ao sys.path
sys.path.append(
    os.path.abspath(
        os.path.join(os.path.dirname(__file__), "../../..")
    )
)


from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service
from webdriver_manager.chrome import ChromeDriverManager
from docx import Document
from docx.shared import Inches
from faker import Faker
from datetime import datetime
import subprocess
import time
import os

# ==== CONFIGURA√á√ïES ====
URL = "http://localhost:8080/gs/index.xhtml"
LOGIN_EMAIL = "joaoeduardo.gold@outlook.com"
LOGIN_PASSWORD = "071999gs"
fake = Faker("pt_BR")

# ==== DOCUMENTO ====
doc = Document()
doc.add_heading("RELAT√ìRIO DO TESTE", 0)
doc.add_paragraph("Cadastro de Centro de Custo ‚Äì Cen√°rio 1: Preenchimento completo e salvamento.")
doc.add_paragraph(f"üóïÔ∏è Data do teste: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")

# ==== LOG E SCREENSHOT ====
screenshot_registradas = set()
def log(doc, msg):
    print(msg)
    doc.add_paragraph(msg)

def registrar_screenshot_unico(nome, driver, doc, descricao=None):
    os.makedirs("screenshots", exist_ok=True)
    if nome not in screenshot_registradas:
        if descricao:
            log(doc, f"üì∏ {descricao}")
        path = f"screenshots/{nome}.png"
        driver.save_screenshot(path)
        doc.add_paragraph(f"üñºÔ∏è Screenshot: {nome}")
        doc.add_picture(path, width=Inches(5.5))
        screenshot_registradas.add(nome)

# ==== CHROME DRIVER ====
options = Options()
options.add_argument("--start-maximized")
driver = webdriver.Chrome(service=Service(ChromeDriverManager().install()), options=options)
wait = WebDriverWait(driver, 10)

# ==== FUN√á√ïES ====
def finalizar_relatorio():
    doc_name = f"relatorio_centro_custo_cenario_1_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc.save(doc_name)
    log(doc, f"üìÑ Relat√≥rio salvo como: {doc_name}")
    subprocess.run(["start", "winword", doc_name], shell=True)
    driver.quit()

def safe_action(doc, descricao, func, driver=None, wait=None):
    try:
        log(doc, f"üîÑ {descricao}...")
        func()
        log(doc, f"‚úÖ {descricao} realizado com sucesso.")
        if driver:
            registrar_screenshot_unico(descricao.lower().replace(" ", "_"), driver, doc)
        return True, None
    except Exception as e:
        log(doc, f"‚ùå Erro ao {descricao.lower()}: {e}")
        if driver:
            registrar_screenshot_unico(f"erro_{descricao.lower().replace(' ', '_')}", driver, doc)
        return False, str(e)



def login():
    wait.until(EC.presence_of_element_located((By.ID, "j_id15:email"))).send_keys(LOGIN_EMAIL)
    wait.until(EC.presence_of_element_located((By.ID, "j_id15:senha"))).send_keys(LOGIN_PASSWORD, Keys.ENTER)
    wait.until(EC.presence_of_element_located((By.TAG_NAME, "body")))

def ajustar_zoom(driver):
    wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, "body > div.sideMenu.animate > div.menuHolder > a.button.btModules > span")))
    driver.execute_script("document.body.style.zoom='90%'")

# ==== EXECU√á√ÉO ====
safe_action(doc, "Acessando o sistema", lambda: driver.get(URL), driver, wait)
safe_action(doc, "Realizando login", login, driver, wait)
safe_action(doc, "Ajustando zoom", lambda: ajustar_zoom(driver), driver, wait)

safe_action(doc, "Abrindo menu Centro de Custo", lambda: (
    driver.find_element(By.TAG_NAME, "body").send_keys(Keys.F2),
    time.sleep(2),
    driver.find_element(By.XPATH, "//input[@placeholder='Busque um cadastro']").send_keys("Centro de Custo", Keys.ENTER)
), driver, wait)

safe_action(doc, "Acessando op√ß√£o Cadastrar", lambda: wait.until(
    EC.element_to_be_clickable((By.CSS_SELECTOR, "#fmod_10005 > div.wdTelas > div > ul > li:nth-child(1) > a > span"))
).click(), driver, wait)

safe_action(
    doc,
    "Preenchendo C√≥digo do Centro de Custo",
    lambda: wait.until(
        EC.element_to_be_clickable((By.CSS_SELECTOR,
        "#fmod_10005 > div.wdTelas > div > div.catWrapper > div > div > div:nth-child(1) > div > div:nth-child(2) > input"))
    ).send_keys(str(fake.random_int(min=10000000, max=99999999))),
    driver,
    wait
)


safe_action(doc, "Clicando no campo Nome do Centro de Custo", lambda: wait.until(
    EC.element_to_be_clickable((By.CSS_SELECTOR, "#fmod_10005 > div.wdTelas > div > div.catWrapper > div > div > div:nth-child(1) > div > div:nth-child(3) > input"))
).click(), driver, wait)

safe_action(doc, "Preenchendo Nome do Centro de Custo", lambda: wait.until(
    EC.presence_of_element_located((By.CSS_SELECTOR, "#fmod_10005 > div.wdTelas > div > div.catWrapper > div > div > div:nth-child(1) > div > div:nth-child(3) > input"))
).send_keys("TESTE CENTRO DE CUSTO SELENIUM AUTOMATIZADO"), driver, wait)

safe_action(doc, "Clicando no campo Telefone", lambda: wait.until(
    EC.element_to_be_clickable((By.CSS_SELECTOR, "#fmod_10005 > div.wdTelas > div > div.catWrapper > div > div > div:nth-child(1) > div > div:nth-child(5) > input"))
).click(), driver, wait)

safe_action(doc, "Preenchendo o campo Telefone", lambda: wait.until(
    EC.presence_of_element_located((By.CSS_SELECTOR, "#fmod_10005 > div.wdTelas > div > div.catWrapper > div > div > div:nth-child(1) > div > div:nth-child(5) > input"))
).send_keys(fake.phone_number()), driver, wait)

safe_action(doc, "Clicando no campo CEP", lambda: WebDriverWait(driver, 1).until(
    EC.element_to_be_clickable((By.CSS_SELECTOR,
    "#fmod_10005 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div > div:nth-child(1) > div > div:nth-child(8) > div > input"))
).click())

safe_action(doc, "Preenchendo o campo CEP", lambda: wait.until(
    EC.presence_of_element_located((By.CSS_SELECTOR,
    "#fmod_10005 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div > div:nth-child(1) > div > div:nth-child(8) > div > input"))
).send_keys("15081115"), driver, wait)

safe_action(doc, "Clicando no bot√£o de busca do CEP", lambda: wait.until(
    EC.element_to_be_clickable((By.CSS_SELECTOR,
    "#fmod_10005 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div > div:nth-child(1) > div > div:nth-child(8) > div > a"))
).click(), driver, wait)

safe_action(doc, "Preenchendo o n√∫mero do endere√ßo", lambda: wait.until(
    EC.element_to_be_clickable((By.CSS_SELECTOR,
    "#fmod_10005 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div > div:nth-child(1) > div > div:nth-child(9) > input"))
).send_keys("1733"), driver, wait)

safe_action(doc, "Clicando em Salvar", lambda: wait.until(
    EC.element_to_be_clickable((By.CSS_SELECTOR,
    "#fmod_10005 > div.wdTelas > div.telaCadastro.clearfix > div.btnHolder > a.btModel.btGray.btsave"))
).click(), driver, wait)

safe_action(doc, "Fechando modal", lambda: wait.until(
    EC.element_to_be_clickable((By.CSS_SELECTOR,
    "#fmod_10005 > div.wdTop.ui-draggable-handle > div.wdClose > a"))
).click(), driver, wait)


    # Mensagem de alerta
# Fun√ß√£o para encontrar mensagem de alerta
def encontrar_mensagem_alerta():
    seletores = [
        (".alerts.salvo", "‚úÖ Mensagem de Sucesso"),
        (".alerts.alerta", "‚ö†Ô∏è Mensagem de Alerta"),
        (".alerts.erro", "‚ùå Mensagem de Erro"),
    ]

    for seletor, tipo in seletores:
        try:
            elemento = driver.find_element(By.CSS_SELECTOR, seletor)
            if elemento.is_displayed():
                log(doc, f"üì¢ {tipo}: {elemento.text}")
                return elemento
        except:
            continue

    log(doc, "‚ÑπÔ∏è Nenhuma mensagem de alerta encontrada.")
    return None

# Fun√ß√£o para tirar screenshot
def take_screenshot(driver, doc, nome):
    registrar_screenshot_unico(nome, driver, doc, f"Screenshot: {nome}")


# Bloco final do teste

    encontrar_mensagem_alerta()
    


log(doc, "‚úÖ Teste executado com sucesso.")
registrar_screenshot_unico("finalizacao_teste", driver, doc, "Tela final ap√≥s o cadastro")
finalizar_relatorio()
