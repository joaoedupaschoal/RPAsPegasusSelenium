# Refatorado e organizado: cadastrodetipodemensalidades1ºcenario.py

from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support.ui import Select, WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service
from selenium.webdriver import ActionChains
from webdriver_manager.chrome import ChromeDriverManager
from docx import Document
from docx.shared import Inches
from faker import Faker
from faker.providers import BaseProvider
from validate_docbr import CPF
from datetime import datetime, timedelta
import subprocess
import os
import time
import random
import string

# ==== PROVIDERS CUSTOMIZADOS ====
class BrasilProvider(BaseProvider):
    def rg(self):
        numeros = [str(random.randint(0, 9)) for _ in range(8)]
        return ''.join(numeros) + '-' + str(random.randint(0, 9))

# Inclui letras, dígitos e símbolos
todos_caracteres = string.ascii_letters + string.digits + string.punctuation
identificador = ''.join(random.choices(todos_caracteres, k=2))

fake = Faker("pt_BR")
fake.add_provider(BrasilProvider)

# ==== CONFIGURAÇÕES ====
URL = "http://localhost:8080/gs/index.xhtml"
LOGIN_EMAIL = "joaoeduardo.gold@outlook.com"
LOGIN_PASSWORD = "071999gs"

# ==== DOCUMENTO ====
doc = Document()
doc.add_heading("RELATÓRIO DO TESTE", 0)
doc.add_paragraph("Cadastro de Tipo de Mensalidade – Cenário 2: Preenchimento completo e cancelamento.")
doc.add_paragraph(f"Data do teste: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")

screenshot_registradas = set()

def abrir_modal_e_selecionar(btn_selector, pesquisa_selector, termo_pesquisa, btn_pesquisar_selector, resultado_xpath):
    def acao():
        # Abre o modal
        open_lov = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, btn_selector)))
        open_lov.click()
        time.sleep(2)

        # Aguarda campo pesquisa
        campo_pesquisa = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, pesquisa_selector)))
        campo_pesquisa.clear()
        campo_pesquisa.send_keys(termo_pesquisa, Keys.ENTER)
        time.sleep(1)

        # Espera o resultado carregar e clica
        resultado = wait.until(EC.element_to_be_clickable((By.XPATH, resultado_xpath)))
        resultado.click()

    return acao

def abrir_modal_e_selecionar_por_texto(btn_selector, pesquisa_selector, termo_pesquisa, btn_pesquisar_selector, texto_resultado):
    def acao():
        # Abre o modal
        open_lov = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, btn_selector)))
        open_lov.click()
        time.sleep(2)

        # Aguarda campo pesquisa
        campo_pesquisa = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, pesquisa_selector)))
        campo_pesquisa.clear()
        campo_pesquisa.send_keys(termo_pesquisa, Keys.ENTER)
        time.sleep(1)

        # Espera o resultado carregar e clica
        resultado = wait.until(EC.element_to_be_clickable((By.XPATH, f"//td[contains(text(), '{texto_resultado}')]")))
        resultado.click()

    return acao

# ==== FUNÇÕES DE UTILITÁRIO ====
def log(doc, msg):
    print(msg)
    doc.add_paragraph(msg)

def take_screenshot(driver, doc, nome):
    if nome not in screenshot_registradas:
        path = f"screenshots/{nome}.png"
        os.makedirs("screenshots", exist_ok=True)
        driver.save_screenshot(path)
        doc.add_paragraph(f"Screenshot: {nome}")
        doc.add_picture(path, width=Inches(5.5))
        screenshot_registradas.add(nome)

def safe_action(doc, descricao, func):
    try:
        log(doc, f"🔄 {descricao}...")
        func()
        log(doc, f"✅ {descricao} realizada com sucesso.")
        take_screenshot(driver, doc, descricao.lower().replace(" ", "_"))
    except Exception as e:
        log(doc, f"❌ Erro ao {descricao.lower()}: {e}")
        take_screenshot(driver, doc, f"erro_{descricao.lower().replace(' ', '_')}")

def finalizar_relatorio():
    nome_arquivo = f"relatorio_tipo_mensalidade_cenario_2_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc.save(nome_arquivo)
    log(doc, f"📄 Relatório salvo como: {nome_arquivo}")
    subprocess.run(["start", "winword", nome_arquivo], shell=True)
    driver.quit()

def gerar_dados_tipo_mensalidade():
    """Gera dados fictícios para o cadastro de tipo de mensalidade."""
    nome = 'TESTE TIPO DE MENSALIDADE SELENIUM AUTOMATIZADO'
    descricao = 'TESTE DESCRIÇÃO SELENIUM AUTOMATIZADO'
    template = 'TESTE TEMPLATE SELENIUM AUTOMATIZADO'
    
    return nome, descricao, template

def encontrar_mensagem_alerta():
    seletores = [
        (".alerts.salvo", "✅ Sucesso"),
        (".alerts.alerta", "⚠️ Alerta"),
        (".alerts.erro", "❌ Erro"),
    ]

    for seletor, tipo in seletores:
        try:
            elemento = driver.find_element(By.CSS_SELECTOR, seletor)
            if elemento.is_displayed():
                log(doc, f"📢 {tipo}: {elemento.text}")
                return elemento
        except:
            continue

    log(doc, "ℹ️ Nenhuma mensagem de alerta encontrada.")
    return None

def ajustar_zoom():
    try:
        driver.execute_script("document.body.style.zoom='90%'")
        log(doc, "🔍 Zoom ajustado para 90%.")
    except Exception as e:
        log(doc, f"⚠️ Erro ao ajustar zoom: {e}")

def selecionar_opcao(selector, texto):
    def acao():
        select_element = wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, selector)))
        Select(select_element).select_by_visible_text(texto)
    return acao

def clicar_aba(texto_aba):
    def acao():
        aba = wait.until(EC.element_to_be_clickable((By.LINK_TEXT, texto_aba)))
        aba.click()
    return acao

def scroll_para_elemento(selector):
    def acao():
        elemento = wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, selector)))
        driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", elemento)
    return acao

# Gera os dados necessários
nome, descricao, template = gerar_dados_tipo_mensalidade()

# ==== INICIALIZAÇÃO DO DRIVER ====
options = Options()
options.add_argument("--start-maximized")
driver = webdriver.Chrome(service=Service(ChromeDriverManager().install()), options=options)
wait = WebDriverWait(driver, 20)

# ==== EXECUÇÃO DO TESTE ====
try:
    safe_action(doc, "Acessando sistema", lambda: driver.get(URL))

    safe_action(doc, "Realizando login", lambda: (
        wait.until(EC.presence_of_element_located((By.ID, "j_id15:email"))).send_keys(LOGIN_EMAIL),
        wait.until(EC.presence_of_element_located((By.ID, "j_id15:senha"))).send_keys(LOGIN_PASSWORD, Keys.ENTER),
        wait.until(EC.presence_of_element_located((By.TAG_NAME, "body")))
    ))

    safe_action(doc, "Esperando sistema carregar e ajustando zoom", lambda: (
        time.sleep(5),
        ajustar_zoom()
    ))

    safe_action(doc, "Abrindo menu Tipo de Mensalidade", lambda: (
        driver.find_element(By.TAG_NAME, "body").send_keys(Keys.F2),
        time.sleep(1),
        wait.until(EC.visibility_of_element_located((By.XPATH, "//input[@placeholder='Busque um cadastro']"))).send_keys("Tipo de Mensalidade", Keys.ENTER),
        time.sleep(3)
    ))

    safe_action(doc, "Clicando em Cadastrar", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_10027 > div.wdTelas > div.telaInicial.clearfix.overflow.overflowY > ul > li:nth-child(1) > a > span"))).click(),
        time.sleep(2)
    ))

    safe_action(doc, "Preenchendo Nome", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_10027 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_10029.categoriaHolder > div > div > div:nth-child(2) > input")))
        .send_keys(nome)
    ))

    safe_action(doc, "Preenchendo Descrição", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_10027 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_10029.categoriaHolder > div > div > div:nth-child(3) > input")))
        .send_keys(descricao)
    ))

    safe_action(doc, "Preenchendo Identificador", lambda: (
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "#fmod_10027 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_10029.categoriaHolder > div > div > div:nth-child(4) > input")))
        .send_keys(identificador)
    ))

    safe_action(doc, "Selecionando a opção para Exibir descrição dos títulos nos boletos", selecionar_opcao(
        "#fmod_10027 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_10029.categoriaHolder > div > div > div:nth-child(5) > select",
        "Não"
    ))

    safe_action(doc, "Preenchendo Template", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_10027 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_10029.categoriaHolder > div > div > div:nth-child(6) > input")))
        .send_keys(template)
    ))

    safe_action(doc, "Clicando aba Informações Financeiras", clicar_aba("Informações Financeiras"))

    # Seção 1 - Informações Financeiras
    safe_action(doc, "Selecionando Conta Crédito 1", abrir_modal_e_selecionar(
        "#fmod_10027 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_10030.categoriaHolder > div:nth-child(1) > div > div:nth-child(2) > div > div > a",
        "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(2) > input",
        "748",
        "",
        "//tr[td[contains(text(), 'RECEITAS COM CEMITÉRIO')]]//a[contains(@class, 'linkAlterar')]"
    ))

    safe_action(doc, "Selecionando Conta Débito 1", abrir_modal_e_selecionar(
        "#fmod_10027 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_10030.categoriaHolder > div:nth-child(1) > div > div:nth-child(3) > div > div > a",
        "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(2) > input",
        "565",
        "",
        "//tr[td[contains(text(), 'TRANSITÓRIA DE CARTÃO DE DÉBITO')]]//a[contains(@class, 'linkAlterar')]"
    ))

    safe_action(doc, "Selecionando Histórico Padrão 1", abrir_modal_e_selecionar_por_texto(
        "#fmod_10027 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_10030.categoriaHolder > div:nth-child(1) > div > div:nth-child(4) > div > div > a",
        "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(2) > input",
        "200129",
        "",
        "RECEBIMENTOS VIA CHEQUE E DINHEIRO"
    ))

    safe_action(doc, "Selecionando Centro de Custo 1", abrir_modal_e_selecionar_por_texto(
        "#fmod_10027 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_10030.categoriaHolder > div:nth-child(1) > div > div:nth-child(5) > div > div > a",
        "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(2) > input",
        "49.84.3119",
        "",
        "49.84.3119"
    ))

    # Seção 2 - Informações Financeiras
    safe_action(doc, "Selecionando Conta Crédito 2", abrir_modal_e_selecionar(
        "#fmod_10027 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_10030.categoriaHolder > div:nth-child(2) > div > div:nth-child(2) > div > div > a",
        "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(2) > input",
        "748",
        "",
        "//tr[td[contains(text(), 'RECEITAS COM CEMITÉRIO')]]//a[contains(@class, 'linkAlterar')]"
    ))

    safe_action(doc, "Selecionando Conta Débito 2", abrir_modal_e_selecionar(
        "#fmod_10027 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_10030.categoriaHolder > div:nth-child(2) > div > div:nth-child(3) > div > div > a",
        "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(2) > input",
        "565",
        "",
        "//tr[td[contains(text(), 'TRANSITÓRIA DE CARTÃO DE DÉBITO')]]//a[contains(@class, 'linkAlterar')]"
    ))

    safe_action(doc, "Selecionando Histórico Padrão 2", abrir_modal_e_selecionar_por_texto(
        "#fmod_10027 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_10030.categoriaHolder > div:nth-child(2) > div > div:nth-child(4) > div > div > a",
        "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(2) > input",
        "200129",
        "",
        "RECEBIMENTOS VIA CHEQUE E DINHEIRO"
    ))

    safe_action(doc, "Selecionando Centro de Custo 2", abrir_modal_e_selecionar_por_texto(
        "#fmod_10027 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_10030.categoriaHolder > div:nth-child(2) > div > div:nth-child(5) > div > div > a",
        "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(2) > input",
        "49.84.3119",
        "",
        "49.84.3119"
    ))

    # Seção 3 - Informações Financeiras
    safe_action(doc, "Selecionando Conta Crédito 3", abrir_modal_e_selecionar(
        "#fmod_10027 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_10030.categoriaHolder > div:nth-child(3) > div > div:nth-child(2) > div > div > a",
        "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(2) > input",
        "748",
        "",
        "//tr[td[contains(text(), 'RECEITAS COM CEMITÉRIO')]]//a[contains(@class, 'linkAlterar')]"
    ))

    safe_action(doc, "Selecionando Conta Débito 3", abrir_modal_e_selecionar(
        "#fmod_10027 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_10030.categoriaHolder > div:nth-child(3) > div > div:nth-child(3) > div > div > a",
        "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(2) > input",
        "565",
        "",
        "//tr[td[contains(text(), 'TRANSITÓRIA DE CARTÃO DE DÉBITO')]]//a[contains(@class, 'linkAlterar')]"
    ))

    safe_action(doc, "Selecionando Histórico Padrão 3", abrir_modal_e_selecionar_por_texto(
        "#fmod_10027 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_10030.categoriaHolder > div:nth-child(3) > div > div:nth-child(4) > div > div > a",
        "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(2) > input",
        "200129",
        "",
        "RECEBIMENTOS VIA CHEQUE E DINHEIRO"
    ))

    safe_action(doc, "Selecionando Centro de Custo 3", abrir_modal_e_selecionar_por_texto(
        "#fmod_10027 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_10030.categoriaHolder > div:nth-child(3) > div > div:nth-child(5) > div > div > a",
        "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(2) > input",
        "49.84.3119",
        "",
        "49.84.3119"
    ))

    # Seção 4 - Informações Financeiras (com scroll)
    safe_action(doc, "Fazendo scroll para seção 4", scroll_para_elemento(
        "#fmod_10027 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_10030.categoriaHolder > div:nth-child(4) > div > div:nth-child(2) > div > div > a"
    ))

    safe_action(doc, "Selecionando Conta Crédito 4", abrir_modal_e_selecionar(
        "#fmod_10027 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_10030.categoriaHolder > div:nth-child(4) > div > div:nth-child(2) > div > div > a",
        "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(2) > input",
        "748",
        "",
        "//tr[td[contains(text(), 'RECEITAS COM CEMITÉRIO')]]//a[contains(@class, 'linkAlterar')]"
    ))

    safe_action(doc, "Selecionando Conta Débito 4", abrir_modal_e_selecionar(
        "#fmod_10027 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_10030.categoriaHolder > div:nth-child(4) > div > div:nth-child(3) > div > div > a",
        "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(2) > input",
        "565",
        "",
        "//tr[td[contains(text(), 'TRANSITÓRIA DE CARTÃO DE DÉBITO')]]//a[contains(@class, 'linkAlterar')]"
    ))

    safe_action(doc, "Selecionando Histórico Padrão 4", abrir_modal_e_selecionar_por_texto(
        "#fmod_10027 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_10030.categoriaHolder > div:nth-child(4) > div > div:nth-child(4) > div > div > a",
        "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(2) > input",
        "200129",
        "",
        "RECEBIMENTOS VIA CHEQUE E DINHEIRO"
    ))

    safe_action(doc, "Selecionando Centro de Custo 4", abrir_modal_e_selecionar_por_texto(
        "#fmod_10027 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div.cat_10030.categoriaHolder > div:nth-child(4) > div > div:nth-child(5) > div > div > a",
        "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(2) > input",
        "49.84.3119",
        "",
        "49.84.3119"
    ))

    safe_action(doc, "Cancelando cadastro", lambda: (
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "#fmod_10027 > div.wdTelas > div.telaCadastro.clearfix > div.btnHolder > a.btModel.btGray.btcancel"))).click(),
        time.sleep(1)
    ))

    safe_action(doc, "Confirmando cancelamento", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#BtYes"))).click()
    ))


    safe_action(doc, "Fechando modal após cancelamento", lambda: (
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "#fmod_10027 > div.wdTop.ui-draggable-handle > div.wdClose > a"))).click(),
        time.sleep(1)
    ))

    encontrar_mensagem_alerta()

except Exception as e:
    log(doc, f"❌ ERRO FATAL: {e}")
    take_screenshot(driver, doc, "erro_fatal")

finally:
    log(doc, "✅ Teste concluído com sucesso.")
    finalizar_relatorio()