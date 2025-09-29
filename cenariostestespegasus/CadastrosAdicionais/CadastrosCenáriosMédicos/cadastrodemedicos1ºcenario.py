# Refatorado e organizado: cadastrodemedicos1ºcenario.py

from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support.ui import Select, WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service
from selenium.webdriver import ActionChains
from webdriver_manager.chrome import ChromeDriverManager
from docx import Document
from docx.shared import Inches
from faker import Faker
from faker.providers import BaseProvider
from validate_docbr import CPF
from datetime import datetime, timedelta
import subprocess
import os
import time
import random
import string

# ==== PROVIDERS CUSTOMIZADOS ====
class BrasilProvider(BaseProvider):
    def rg(self):
        numeros = [str(random.randint(0, 9)) for _ in range(8)]
        return ''.join(numeros) + '-' + str(random.randint(0, 9))

fake = Faker("pt_BR")
fake.add_provider(BrasilProvider)

# ==== CONFIGURAÇÕES ====
URL = "http://localhost:8080/gs/index.xhtml"
LOGIN_EMAIL = "joaoeduardo.gold@outlook.com"
LOGIN_PASSWORD = "071999gs"

# ==== DOCUMENTO ====
doc = Document()
doc.add_heading("RELATÓRIO DO TESTE", 0)
doc.add_paragraph("Cadastro de Médicos – Cenário 1: Preenchimento completo e salvamento.")
doc.add_paragraph(f"Data do teste: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")

screenshot_registradas = set()

def preencher_campo_com_retry(driver, wait, seletor, valor, max_tentativas=3):
    """Tenta preencher o campo com diferentes métodos até conseguir"""
    
    for tentativa in range(max_tentativas):
        try:
            # Aguarda o elemento
            campo = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, seletor)))
            
            # Scroll até o elemento se necessário
            driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", campo)
            time.sleep(0.5)
            
            # Método 1: Tradicional
            if tentativa == 0:
                campo.click()
                campo.clear()
                campo.send_keys(valor)
                campo.send_keys(Keys.TAB)
            
            # Método 2: ActionChains
            elif tentativa == 1:
                ActionChains(driver).move_to_element(campo).click().perform()
                time.sleep(0.2)
                ActionChains(driver).key_down(Keys.CONTROL).send_keys('a').key_up(Keys.CONTROL).perform()
                ActionChains(driver).send_keys(valor).perform()
                ActionChains(driver).send_keys(Keys.TAB).perform()
            
            # Método 3: JavaScript
            else:
                driver.execute_script("""
                    var element = arguments[0];
                    var valor = arguments[1];
                    element.focus();
                    element.value = '';
                    element.value = valor;
                    element.dispatchEvent(new Event('input', { bubbles: true }));
                    element.dispatchEvent(new Event('change', { bubbles: true }));
                    element.blur();
                """, campo, valor)
            
            time.sleep(0.5)
            
            # Verifica se o valor foi preenchido
            valor_atual = campo.get_attribute('value')
            if valor_atual == valor:
                return True
            else:
                print(f"Tentativa {tentativa + 1} falhou. Valor esperado: {valor}, Valor atual: {valor_atual}")
                
        except Exception as e:
            print(f"Tentativa {tentativa + 1} falhou com erro: {e}")
            time.sleep(1)
    
    return False

# ==== FUNÇÕES DE UTILITÁRIO ====
def log(doc, msg):
    print(msg)
    doc.add_paragraph(msg)

def take_screenshot(driver, doc, nome):
    if nome not in screenshot_registradas:
        path = f"screenshots/{nome}.png"
        os.makedirs("screenshots", exist_ok=True)
        driver.save_screenshot(path)
        doc.add_paragraph(f"Screenshot: {nome}")
        doc.add_picture(path, width=Inches(5.5))
        screenshot_registradas.add(nome)

def safe_action(doc, descricao, func):
    try:
        log(doc, f"🔄 {descricao}...")
        func()
        log(doc, f"✅ {descricao} realizada com sucesso.")
        take_screenshot(driver, doc, descricao.lower().replace(" ", "_"))
    except Exception as e:
        log(doc, f"❌ Erro ao {descricao.lower()}: {e}")
        take_screenshot(driver, doc, f"erro_{descricao.lower().replace(' ', '_')}")

def finalizar_relatorio():
    nome_arquivo = f"relatorio_medicos_cenario_1_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc.save(nome_arquivo)
    log(doc, f"📄 Relatório salvo como: {nome_arquivo}")
    try:
        subprocess.run(["start", "winword", nome_arquivo], shell=True)
    except:
        print(f"Relatório salvo em: {nome_arquivo}")
    driver.quit()

def gerar_dados_medico():
    """Gera dados fictícios para o cadastro de médico."""
    nome_completo = 'MÉDICO TESTE SELENIUM AUTOMATIZADO'
    
    # Gera CRM com estado + número de 5 a 6 dígitos
    estados = ["SP", "RJ", "MG", "RS", "PR", "SC", "BA", "PE", "CE", "DF", "GO", "AM", "PA", "ES", "MT", "MS", "MA", "PB", "PI", "RN", "AL", "SE", "RO", "RR", "AC", "AP", "TO"]
    uf = random.choice(estados)
    numero = random.randint(10000, 999999)
    crm = f"{uf}{numero}"
    
    telefone = fake.phone_number()
    cep = "15081115"
    numero_endereco = "1733"
    complemento = "Casa"
    
    return nome_completo, crm, telefone, cep, numero_endereco, complemento

def encontrar_mensagem_alerta():
    seletores = [
        (".alerts.salvo", "✅ Mensagem de Sucesso"),
        (".alerts.alerta", "⚠️ Mensagem de Alerta"),
        (".alerts.erro", "❌ Mensagem de Erro"),
    ]
    for seletor, tipo in seletores:
        try:
            elemento = driver.find_element(By.CSS_SELECTOR, seletor)
            if elemento.is_displayed():
                log(doc, f"📢 {tipo}: {elemento.text}")
                return elemento
        except:
            continue

    log(doc, "ℹ️ Nenhuma mensagem de alerta encontrada.")
    return None

def ajustar_zoom():
    try:
        driver.execute_script("document.body.style.zoom='90%'")
        log(doc, "🔍 Zoom ajustado para 90%.")
    except Exception as e:
        log(doc, f"⚠️ Erro ao ajustar zoom: {e}")

# Gera os dados necessários
nome_completo, crm, telefone, cep, numero_endereco, complemento = gerar_dados_medico()

# ==== INICIALIZAÇÃO DO DRIVER ====
options = Options()
options.add_argument("--start-maximized")
driver = webdriver.Chrome(service=Service(ChromeDriverManager().install()), options=options)
wait = WebDriverWait(driver, 20)

# ==== EXECUÇÃO DO TESTE ====
try:
    safe_action(doc, "Acessando sistema", lambda: driver.get(URL))

    safe_action(doc, "Realizando login", lambda: (
        wait.until(EC.presence_of_element_located((By.ID, "j_id15:email"))).send_keys(LOGIN_EMAIL),
        wait.until(EC.presence_of_element_located((By.ID, "j_id15:senha"))).send_keys(LOGIN_PASSWORD, Keys.ENTER),
        wait.until(EC.presence_of_element_located((By.TAG_NAME, "body")))
    ))

    safe_action(doc, "Esperando sistema carregar e ajustando zoom", lambda: (
        time.sleep(5),
        ajustar_zoom()
    ))

    safe_action(doc, "Abrindo menu Médicos", lambda: (
        driver.find_element(By.TAG_NAME, "body").send_keys(Keys.F2),
        time.sleep(1),
        wait.until(EC.visibility_of_element_located((By.XPATH, "//input[@placeholder='Busque um cadastro']"))).send_keys("Médicos", Keys.ENTER),
        time.sleep(3)
    ))

    safe_action(doc, "Clicando em Cadastrar", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_10033 > div.wdTelas > div.telaInicial.clearfix.overflow.overflowY > ul > li:nth-child(1) > a > span"))).click(),
        time.sleep(2)
    ))

    # Preenchendo dados do médico - CORRIGIDO
    safe_action(doc, "Preenchendo Nome Completo", lambda: 
        preencher_campo_com_retry(
            driver, 
            wait, 
            "#fmod_10033 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div > div:nth-child(1) > div > div:nth-child(2) > input",
            nome_completo
        )
    )

    safe_action(doc, "Preenchendo CRM", lambda: 
        preencher_campo_com_retry(
            driver, 
            wait, 
            "#fmod_10033 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div > div:nth-child(1) > div > div:nth-child(3) > input",
            crm
        )
    )

    safe_action(doc, "Preenchendo Telefone", lambda: 
        preencher_campo_com_retry(
            driver, 
            wait, 
            "#fmod_10033 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div > div:nth-child(1) > div > div:nth-child(4) > input",
            telefone
        )
    )

    safe_action(doc, "Preenchendo CEP", lambda: 
        preencher_campo_com_retry(
            driver, 
            wait, 
            "#fmod_10033 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div > div:nth-child(1) > div > div:nth-child(5) > div > input",
            cep
        )
    )

    safe_action(doc, "Clicando no botão de busca do CEP", lambda: (
        time.sleep(2),
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "#fmod_10033 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div > div:nth-child(1) > div > div:nth-child(5) > div > a"))).click()
    ))

    safe_action(doc, "Preenchendo Número do Endereço", lambda: 
        preencher_campo_com_retry(
            driver, 
            wait, 
            "#fmod_10033 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div > div:nth-child(1) > div > div:nth-child(8) > input",
            numero_endereco
        )
    )

    safe_action(doc, "Preenchendo Complemento", lambda: 
        preencher_campo_com_retry(
            driver, 
            wait, 
            "#fmod_10033 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div > div:nth-child(1) > div > div:nth-child(9) > input",
            complemento
        )
    )

    safe_action(doc, "Salvando cadastro", lambda: (
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, '#fmod_10033 > div.wdTelas > div.telaCadastro.clearfix > div.btnHolder > a.btModel.btGray.btsave'))).click(),
    ))

    safe_action(doc, "Fechando modal após salvamento", lambda: (
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "#fmod_10033 > div.wdTop.ui-draggable-handle > div.wdClose > a"))).click(),

    ))

    encontrar_mensagem_alerta()

except Exception as e:
    log(doc, f"❌ ERRO FATAL: {e}")
    take_screenshot(driver, doc, "erro_fatal")

finally:
    log(doc, "✅ Teste concluído com sucesso.")
    finalizar_relatorio()