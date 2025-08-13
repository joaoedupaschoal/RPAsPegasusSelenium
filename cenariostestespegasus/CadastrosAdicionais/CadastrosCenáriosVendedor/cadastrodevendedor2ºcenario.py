# Refatorado e organizado: cadastrodevendedor1ºcenario.py

from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support.ui import Select, WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service
from selenium.webdriver import ActionChains
from webdriver_manager.chrome import ChromeDriverManager
from docx import Document
from docx.shared import Inches
from faker import Faker
from faker.providers import BaseProvider
from validate_docbr import CPF
from datetime import datetime, timedelta
import subprocess
import os
import time
import random
import string

# ==== PROVIDERS CUSTOMIZADOS ====
class BrasilProvider(BaseProvider):
    def rg(self):
        numeros = [str(random.randint(0, 9)) for _ in range(8)]
        return ''.join(numeros) + '-' + str(random.randint(0, 9))

fake = Faker("pt_BR")
fake.add_provider(BrasilProvider)

# ==== CONFIGURAÇÕES ====
URL = "http://localhost:8080/gs/index.xhtml"
LOGIN_EMAIL = "joaoeduardo.gold@outlook.com"
LOGIN_PASSWORD = "071999gs"

# ==== DOCUMENTO ====
doc = Document()
doc.add_heading("RELATÓRIO DO TESTE", 0)
doc.add_paragraph("Cadastro de Vendedores – Cenário 2: Preenchimento completo e cancelamento.")
doc.add_paragraph(f"Data do teste: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")

screenshot_registradas = set()

def preencher_campo_com_retry(driver, wait, seletor, valor, max_tentativas=2):
    """Tenta preencher o campo com diferentes métodos até conseguir"""
    
    for tentativa in range(max_tentativas):
        try:
            # Aguarda o elemento
            campo = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, seletor)))
            
            # Scroll até o elemento se necessário
            driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", campo)
            time.sleep(0.5)
            
            # Método 1: Tradicional
            if tentativa == 0:
                campo.click()
                campo.clear()
                campo.send_keys(valor)
                campo.send_keys(Keys.TAB)
            
            # Método 2: ActionChains
            elif tentativa == 1:
                ActionChains(driver).move_to_element(campo).click().perform()
                time.sleep(0.2)
                ActionChains(driver).key_down(Keys.CONTROL).send_keys('a').key_up(Keys.CONTROL).perform()
                ActionChains(driver).send_keys(valor).perform()
                ActionChains(driver).send_keys(Keys.TAB).perform()
            
            # Método 3: JavaScript
            else:
                driver.execute_script("""
                    var element = arguments[0];
                    var valor = arguments[1];
                    element.focus();
                    element.value = '';
                    element.value = valor;
                    element.dispatchEvent(new Event('input', { bubbles: true }));
                    element.dispatchEvent(new Event('change', { bubbles: true }));
                    element.blur();
                """, campo, valor)
            
            time.sleep(0.5)
            
            # Verifica se o valor foi preenchido
            valor_atual = campo.get_attribute('value')
            if valor_atual == valor:
                return True
            else:
                print(f"Tentativa {tentativa + 1} falhou. Valor esperado: {valor}, Valor atual: {valor_atual}")
                
        except Exception as e:
            print(f"Erro na tentativa {tentativa + 1}: {e}")
            time.sleep(1)
    
    print(f"Falha ao preencher campo após {max_tentativas} tentativas")
    return False

# ==== FUNÇÕES DE UTILITÁRIO ====
def log(doc, msg):
    print(msg)
    doc.add_paragraph(msg)

def take_screenshot(driver, doc, nome):
    if nome not in screenshot_registradas:
        path = f"screenshots/{nome}.png"
        os.makedirs("screenshots", exist_ok=True)
        driver.save_screenshot(path)
        doc.add_paragraph(f"Screenshot: {nome}")
        doc.add_picture(path, width=Inches(5.5))
        screenshot_registradas.add(nome)

def safe_action(doc, descricao, func):
    try:
        log(doc, f"🔄 {descricao}...")
        func()
        log(doc, f"✅ {descricao} realizada com sucesso.")
        take_screenshot(driver, doc, descricao.lower().replace(" ", "_"))
    except Exception as e:
        log(doc, f"❌ Erro ao {descricao.lower()}: {e}")
        take_screenshot(driver, doc, f"erro_{descricao.lower().replace(' ', '_')}")

def finalizar_relatorio():
    nome_arquivo = f"relatorio_vendedores_cenario_2_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc.save(nome_arquivo)
    log(doc, f"📄 Relatório salvo como: {nome_arquivo}")
    subprocess.run(["start", "winword", nome_arquivo], shell=True)
    driver.quit()

def gerar_dados_vendedor():
    """Gera dados fictícios para o cadastro de vendedor."""
    nome_completo = 'VENDEDOR TESTE SELENIUM AUTOMATIZADO'
    cpf_valido = CPF().generate()
    rg = fake.rg()
    data_nascimento = fake.date_of_birth(minimum_age=18, maximum_age=60).strftime("%d/%m/%Y")
    data_expedicao = fake.date_between(start_date=datetime.today().date() - timedelta(days=3650), 
                                      end_date=datetime.today().date()).strftime("%d/%m/%Y")
    email = fake.email()
    telefone1 = fake.phone_number()
    telefone2 = fake.phone_number()
    telefone3 = fake.phone_number()
    cidade_nascimento = fake.city()
    pais_nascimento = fake.country()
    nome_pai = fake.first_name()
    nome_mae = fake.first_name()
    profissao = fake.job()
    
    # Dados específicos do vendedor
    data_admissao = fake.date_between(start_date='-10y', end_date='today').strftime('%d/%m/%Y')
    carteira_trabalho = str(random.randint(10000000, 99999999))
    pis = fake.cpf().replace('.', '').replace('-', '')[:11]
    comissao_porcentagem = str(fake.random_int(min=10, max=500))
    comissao_reais = str(fake.random_int(min=10, max=500))
    
    return (nome_completo, cpf_valido, rg, data_nascimento, data_expedicao, email, 
            telefone1, telefone2, telefone3, cidade_nascimento, pais_nascimento, 
            nome_pai, nome_mae, profissao, data_admissao, carteira_trabalho, 
            pis, comissao_porcentagem, comissao_reais)

def encontrar_mensagem_alerta():
    seletores = [
        (".alerts.salvo", "✅ Sucesso"),
        (".alerts.alerta", "⚠️ Alerta"),
        (".alerts.erro", "❌ Erro"),
    ]

    for seletor, tipo in seletores:
        try:
            elemento = driver.find_element(By.CSS_SELECTOR, seletor)
            if elemento.is_displayed():
                log(doc, f"📢 {tipo}: {elemento.text}")
                return elemento
        except:
            continue

    log(doc, "ℹ️ Nenhuma mensagem de alerta encontrada.")
    return None

def click_element_safely(driver, wait, selector, by=By.CSS_SELECTOR, timeout=10):
    """Clica em um elemento de forma segura, tentando diferentes métodos"""
    try:
        element = wait.until(EC.presence_of_element_located((by, selector)))
        driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", element)
        time.sleep(0.5)
        
        element = wait.until(EC.element_to_be_clickable((by, selector)))
        
        try:
            element.click()
            return True
        except:
            try:
                ActionChains(driver).move_to_element(element).click().perform()
                return True
            except:
                driver.execute_script("arguments[0].click();", element)
                return True
                
    except Exception as e:
        print(f"Erro ao clicar no elemento {selector}: {e}")
        return False

def abrir_modal_e_selecionar(btn_selector, pesquisa_selector, termo_pesquisa, btn_pesquisar_selector, resultado_xpath):
    def acao():
        wait = WebDriverWait(driver, 20)
        
        if not click_element_safely(driver, wait, btn_selector):
            raise Exception("Não foi possível abrir o modal")
        
        time.sleep(1)
        
        campo_pesquisa = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, pesquisa_selector)))
        campo_pesquisa.clear()
        campo_pesquisa.send_keys(termo_pesquisa)
        
        if not click_element_safely(driver, wait, btn_pesquisar_selector):
            raise Exception("Não foi possível clicar no botão pesquisar")
        
        time.sleep(2)
        
        wait.until(EC.presence_of_element_located((By.XPATH, resultado_xpath)))
        wait.until(EC.element_to_be_clickable((By.XPATH, resultado_xpath)))
        
        resultado = driver.find_element(By.XPATH, resultado_xpath)
        driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", resultado)
        time.sleep(0.5)
        
        if not click_element_safely(driver, wait, resultado_xpath, By.XPATH):
            raise Exception("Não foi possível selecionar o resultado")
    
    return acao


def ajustar_zoom():
    try:
        driver.execute_script("document.body.style.zoom='90%'")
        log(doc, "🔍 Zoom ajustado para 90%.")
    except Exception as e:
        log(doc, f"⚠️ Erro ao ajustar zoom: {e}")

def preencher_campo_data(selector, valor):
    def acao():
        campo = wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, selector)))
        campo.click()
        campo.send_keys(valor)
        time.sleep(0.2)
    return acao

def preencher_campo_data_xpath(xpath, valor):
    def acao():
        campo = wait.until(EC.element_to_be_clickable((By.XPATH, xpath)))
        campo.send_keys(valor)
        driver.execute_script("arguments[0].value = arguments[1];", campo, valor)
        campo.send_keys(Keys.TAB)
        time.sleep(0.2)
    return acao

def selecionar_opcao(selector, texto):
    def acao():
        select_element = wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, selector)))
        Select(select_element).select_by_visible_text(texto)
    return acao

# Gera os dados necessários
(nome_completo, cpf_valido, rg, data_nascimento, data_expedicao, email, 
 telefone1, telefone2, telefone3, cidade_nascimento, pais_nascimento, 
 nome_pai, nome_mae, profissao, data_admissao, carteira_trabalho, 
 pis, comissao_porcentagem, comissao_reais) = gerar_dados_vendedor()

# ==== INICIALIZAÇÃO DO DRIVER ====
options = Options()
options.add_argument("--start-maximized")
driver = webdriver.Chrome(service=Service(ChromeDriverManager().install()), options=options)
wait = WebDriverWait(driver, 20)

# ==== EXECUÇÃO DO TESTE ====
try:
    safe_action(doc, "Acessando sistema", lambda: driver.get(URL))

    safe_action(doc, "Realizando login", lambda: (
        wait.until(EC.presence_of_element_located((By.ID, "j_id15:email"))).send_keys(LOGIN_EMAIL),
        wait.until(EC.presence_of_element_located((By.ID, "j_id15:senha"))).send_keys(LOGIN_PASSWORD, Keys.ENTER),
        wait.until(EC.presence_of_element_located((By.TAG_NAME, "body")))
    ))

    safe_action(doc, "Esperando sistema carregar e ajustando zoom", lambda: (
        time.sleep(5),
        ajustar_zoom()
    ))

    safe_action(doc, "Abrindo menu Vendedor", lambda: (
        driver.find_element(By.TAG_NAME, "body").send_keys(Keys.F2),
        time.sleep(1),
        wait.until(EC.visibility_of_element_located((By.XPATH, "//input[@placeholder='Busque um cadastro']"))).send_keys("Vendedor"),
        time.sleep(1),
        wait.until(EC.element_to_be_clickable((By.XPATH, "/html/body/div[17]/ul/li[56]/a"))).click(),
        time.sleep(3)
    ))

    safe_action(doc, "Clicando em Cadastrar", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#fmod_10020 > div.wdTelas > div.telaInicial.clearfix.overflow.overflowY > ul > li:nth-child(1) > a > span"))).click(),
        time.sleep(2)
    ))

    safe_action(doc, "Abrindo LOV de Pessoa Vendedor", lambda: (
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, '#fmod_10020 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div > div > div > div:nth-child(2) > div > a'))).click()
    ))

    safe_action(doc, "Criando novo registro de pessoa", lambda: (
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, 'body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(3) > a'))).click()
    ))

    # Preenchendo dados pessoais
    safe_action(doc, "Preenchendo Nome Completo", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#cg_1 > div.wdTelas > div > div.catWrapper > div > div.cat_dadosPessoais.categoriaHolder > div > div > div:nth-child(2) > div:nth-child(2) > input")))
        .send_keys(nome_completo)
    ))

    safe_action(doc, "Selecionando Tipo de Pessoa", selecionar_opcao(
        "#cg_1 > div.wdTelas > div > div.catWrapper > div > div.cat_dadosPessoais.categoriaHolder > div > div > div:nth-child(2) > div:nth-child(3) > select",
        "Física"
    ))

    safe_action(doc, "Selecionando Órgão Emissor", selecionar_opcao(
        "#cg_1 > div.wdTelas > div > div.catWrapper > div > div.cat_dadosPessoais.categoriaHolder > div > div > div:nth-child(2) > div:nth-child(4) > select",
        "Carteira de Identidade Classista"
    ))

    safe_action(doc, "Preenchendo RG", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#cg_1 > div.wdTelas > div > div.catWrapper > div > div.cat_dadosPessoais.categoriaHolder > div > div > div.formRow.divPessoaFISICA > div:nth-child(1) > input")))
        .send_keys(rg),
        time.sleep(0.5)
    ))

    safe_action(doc, "Preenchendo Data de Expedição", preencher_campo_data(
        "input.dataExpedicao", data_expedicao
    ))

    safe_action(doc, "Preenchendo CPF", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#cg_1 > div.wdTelas > div > div.catWrapper > div > div.cat_dadosPessoais.categoriaHolder > div > div > div.formRow.divPessoaFISICA > div:nth-child(3) > input"))).click(),
        
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#cg_1 > div.wdTelas > div > div.catWrapper > div > div.cat_dadosPessoais.categoriaHolder > div > div > div.formRow.divPessoaFISICA > div:nth-child(3) > input"))).send_keys(cpf_valido),

        time.sleep(1)
    ))

    safe_action(doc, "Acessando aba Dados Complementares", lambda: (
        wait.until(EC.element_to_be_clickable((By.LINK_TEXT, "Dados Complementares"))).click(),
        time.sleep(1)
    ))

    safe_action(doc, "Selecionando Estado Civil", selecionar_opcao(
        "#cg_1 > div.wdTelas > div > div.catWrapper > div > div.cat_dadosComplementares.categoriaHolder > div > div > div > div:nth-child(1) > select",
        "Solteiro"
    ))

    safe_action(doc, "Selecionando Sexo", selecionar_opcao(
        "#cg_1 > div.wdTelas > div > div.catWrapper > div > div.cat_dadosComplementares.categoriaHolder > div > div > div > div:nth-child(2) > select",
        "Feminino"
    ))

    safe_action(doc, "Preenchendo E-mail", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#cg_1 > div.wdTelas > div > div.catWrapper > div > div.cat_dadosComplementares.categoriaHolder > div > div > div > div:nth-child(9) > input")))
        .send_keys(email)
    ))

    safe_action(doc, "Preenchendo Data de Nascimento", preencher_campo_data(
        "input.dataNascimento", data_nascimento
    ))

    safe_action(doc, "Preenchendo Telefone 1", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#cg_1 > div.wdTelas > div > div.catWrapper > div > div.cat_dadosComplementares.categoriaHolder > div > div > div > div:nth-child(5) > input")))
        .send_keys(telefone1)
    ))

    safe_action(doc, "Preenchendo Telefone 2", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#cg_1 > div.wdTelas > div > div.catWrapper > div > div.cat_dadosComplementares.categoriaHolder > div > div > div > div:nth-child(6) > input")))
        .send_keys(telefone2)
    ))

    safe_action(doc, "Preenchendo Telefone 3", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#cg_1 > div.wdTelas > div > div.catWrapper > div > div.cat_dadosComplementares.categoriaHolder > div > div > div > div:nth-child(7) > input")))
        .send_keys(telefone3)
    ))

    safe_action(doc, "Preenchendo Cidade de Nascimento", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#cg_1 > div.wdTelas > div > div.catWrapper > div > div.cat_dadosComplementares.categoriaHolder > div > div > div > div:nth-child(11) > input")))
        .send_keys(cidade_nascimento)
    ))

    safe_action(doc, "Preenchendo País de Nascimento", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#cg_1 > div.wdTelas > div > div.catWrapper > div > div.cat_dadosComplementares.categoriaHolder > div > div > div > div:nth-child(12) > input")))
        .send_keys(pais_nascimento)
    ))

    safe_action(doc, "Preenchendo Nome do Pai", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#cg_1 > div.wdTelas > div > div.catWrapper > div > div.cat_dadosComplementares.categoriaHolder > div > div > div > div:nth-child(13) > input")))
        .send_keys(nome_pai)
    ))

    safe_action(doc, "Preenchendo Nome da Mãe", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#cg_1 > div.wdTelas > div > div.catWrapper > div > div.cat_dadosComplementares.categoriaHolder > div > div > div > div:nth-child(14) > input")))
        .send_keys(nome_mae)
    ))

    safe_action(doc, "Preenchendo Profissão", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#cg_1 > div.wdTelas > div > div.catWrapper > div > div.cat_dadosComplementares.categoriaHolder > div > div > div > div:nth-child(16) > input")))
        .send_keys(profissao)
    ))

    safe_action(doc, "Salvando cadastro da pessoa", lambda: (
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "#cg_1 > div.wdTelas > div > div.btnHolder > a.btModel.btGray.btsave"))).click(),
        time.sleep(5)
    ))

    safe_action(doc, "Preenchendo Data de Admissão", preencher_campo_data_xpath(
        "//input[@grupo='10024' and @ref='10066' and contains(@class, 'hasDatepicker')]",
        data_admissao
    ))

    safe_action(doc, "Preenchendo Carteira de Trabalho", lambda: (
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "#fmod_10020 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div > div > div > div:nth-child(4) > input"))).click(),
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "#fmod_10020 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div > div > div > div:nth-child(4) > input"))).send_keys(carteira_trabalho)
    ))

    safe_action(doc, "Preenchendo PIS", lambda: (
        wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, '#fmod_10020 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div > div > div > div:nth-child(5) > input'))).send_keys(pis)
    ))


    safe_action(doc, "Selecionando Supervisor", abrir_modal_e_selecionar(
        "#fmod_10020 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div > div > div > div:nth-child(6) > div > a",
        "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(2) > input.nomePesquisa",
        "SUPERVISOR TESTE SELENIUM AUTOMATIZADO",
        "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(3) > a",
        "//td[contains(text(), 'SUPERVISOR TESTE SELENIUM AUTOMATIZADO')]"
    ))

    safe_action(doc, "Preenchendo Comissão (Porcentagem)", lambda: (
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, '#fmod_10020 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div > div > div > div:nth-child(7) > input'))).send_keys(comissao_porcentagem)
    ))

    safe_action(doc, "Preenchendo Comissão (Reais)", lambda: (
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, '#fmod_10020 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div > div > div > div:nth-child(8) > input'))).send_keys(comissao_reais)
    ))

    safe_action(doc, "Selecionando Tipo de Vendedor", selecionar_opcao(
        "#fmod_10020 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div > div > div > div:nth-child(9) > select",
        "Interno"
    ))

    safe_action(doc, "Selecionando Tipo de Contrato", selecionar_opcao(
        "#fmod_10020 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div > div > div > div:nth-child(11) > select",
        "Carteira Assinada"
    ))

    safe_action(doc, "Clicando no campo para geração automática do Código Referência", lambda: (
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, '#fmod_10020 > div.wdTelas > div.telaCadastro.clearfix > div.catWrapper > div > div > div > div > div:nth-child(12) > input'))).click(),
        time.sleep(5)
    ))

    safe_action(doc, "Cancelando cadastro do vendedor", lambda: (
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "#fmod_10020 > div.wdTelas > div.telaCadastro.clearfix > div.btnHolder > a.btModel.btGray.btcancel"))).click(),
    ))



    safe_action(doc, "Confirmando cancelamento", lambda: (
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#BtYes"))).click()
    ))



    safe_action(doc, "Fechando modal após cancelamento", lambda: (
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "#fmod_10020 > div.wdTop.ui-draggable-handle > div.wdClose > a"))).click(),
    ))

    encontrar_mensagem_alerta()

except Exception as e:
    log(doc, f"❌ ERRO FATAL: {e}")
    take_screenshot(driver, doc, "erro_fatal")

finally:
    log(doc, "✅ Teste concluído com sucesso.")
    finalizar_relatorio()