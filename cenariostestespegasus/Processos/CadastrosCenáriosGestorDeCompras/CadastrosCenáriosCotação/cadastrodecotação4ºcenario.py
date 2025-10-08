# ==== IMPORTS (sem conflitos) ====
from datetime import datetime, timedelta
from datetime import time as dt_time  # usar para objetos de hora
import time                           # usar para time.sleep(...)
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver import ActionChains
from selenium.common.exceptions import TimeoutException, StaleElementReferenceException, JavascriptException
from selenium.webdriver.support import expected_conditions as EC
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support.ui import Select, WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service
from selenium.common.exceptions import TimeoutException, NoSuchElementException, StaleElementReferenceException, ElementClickInterceptedException
from webdriver_manager.chrome import ChromeDriverManager
from docx import Document
from docx.shared import Inches
from faker import Faker
from faker.providers import BaseProvider
from validate_docbr import CPF
from selenium.webdriver import ActionChains
import subprocess
import os
import random
import re

# ==== PROVIDERS CUSTOMIZADOS ====
class BrasilProvider(BaseProvider):
    def rg(self):
        numeros = [str(random.randint(0, 9)) for _ in range(8)]
        return ''.join(numeros) + '-' + str(random.randint(0, 9))

fake = Faker("pt_BR")
fake.add_provider(BrasilProvider)

# ==== GERAÇÃO DE DATAS ====
def gerar_datas_validas(hora_padrao="00:00", dias_fim=0):
    """Gera datas coerentes. data_inicio/data_fim no formato 'dd/MM/yyyy HH:mm'."""
    hoje_date = datetime.today().date()
    dez_anos_atras = hoje_date - timedelta(days=3650)

    # Falecimento entre 10 anos atrás e hoje
    data_falecimento = fake.date_between(start_date=dez_anos_atras, end_date=hoje_date)

    # Nascimento (entre 18 e 110 anos antes do falecimento)
    idade_minima, idade_maxima = 18, 110
    data_nascimento = data_falecimento - timedelta(days=random.randint(idade_minima * 365, idade_maxima * 365))

    # Sepultamento 1..10 dias após o falecimento
    data_sepultamento = data_falecimento + timedelta(days=random.randint(1, 10))

    # Registro 1..10 dias após o sepultamento
    data_registro = data_sepultamento + timedelta(days=random.randint(1, 10))

    # Velório entre o falecimento e o sepultamento
    data_velorio = fake.date_between(start_date=data_falecimento, end_date=data_sepultamento)

    # Início entre 2 e 30 dias no futuro, com hora escolhida
    data_inicio_date = hoje_date + timedelta(days=random.randint(2, 30))

    # Monta datetime com hora escolhida (ex: "00:00")
    h, m = map(int, hora_padrao.split(":"))
    dt_inicio = datetime.combine(data_inicio_date, dt_time(h, m))

    # Fim: mesmo dia por padrão (dias_fim=0). Ajuste se quiser +N dias.
    dt_fim = dt_inicio + timedelta(days=dias_fim)

    fmt_data = "%d/%m/%Y"
    fmt_dt = "%d/%m/%Y %H:%M"

    return (
        data_nascimento.strftime(fmt_data),
        data_falecimento.strftime(fmt_data),
        data_sepultamento.strftime(fmt_data),
        data_velorio.strftime(fmt_data),
        dt_inicio.strftime(fmt_dt),   # data_inicio com hora
        dt_fim.strftime(fmt_dt),      # data_fim com hora
        data_registro.strftime(fmt_data),
        hoje_date.strftime(fmt_data),
    )

(data_nascimento, data_falecimento, data_sepultamento,
 data_velorio, data_inicio, data_fim, data_registro, hoje) = gerar_datas_validas(
    hora_padrao="08:50",
    dias_fim=0
)

hora_falecimento = fake.time(pattern="%H:%M")
hora_sepultamento = fake.time(pattern="%H:%M")
localizacao = fake.city()

# ==== CONFIGURAÇÕES ====
URL = "http://localhost:8080/gs/index.xhtml"
LOGIN_EMAIL = "joaoeduardo.gold@outlook.com"
LOGIN_PASSWORD = "071999gs"

# ==== INICIALIZAÇÃO DE VARIÁVEIS GLOBAIS ====
doc = Document()
doc.add_heading("RELATÓRIO DO TESTE", 0)
doc.add_paragraph("Cadastro de Cotação - Gestor de Compras – Cenário 4: Preenchimento completo e salvamento de várias cotações")
doc.add_paragraph(f"Data do teste: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")

screenshot_registradas = set()
driver = None
wait = None

# ==== UTILITÁRIOS ====
def log(doc, msg):
    print(msg)
    doc.add_paragraph(msg)

def _sanitize_filename(name: str) -> str:
    name = name.strip().lower()
    name = re.sub(r"[<>:\"/\\|?*']", "_", name)
    name = re.sub(r"_+", "_", name)
    return name[:120]

def take_screenshot(driver, doc, nome):
    if driver is None:
        return
    nome = _sanitize_filename(nome)
    if nome not in screenshot_registradas:
        path = f"screenshots/{nome}.png"
        os.makedirs("screenshots", exist_ok=True)
        try:
            driver.save_screenshot(path)
            doc.add_paragraph(f"Screenshot: {nome}")
            doc.add_picture(path, width=Inches(5.5))
            screenshot_registradas.add(nome)
        except Exception as e:
            log(doc, f"⚠️ Erro ao tirar screenshot {nome}: {e}")

def safe_action(doc, descricao, func, max_retries=3):
    global driver
    for attempt in range(max_retries):
        try:
            log(doc, f"🔄 {descricao}..." if attempt == 0 else f"🔄 {descricao}... (Tentativa {attempt + 1})")
            func()
            log(doc, f"✅ {descricao} realizada com sucesso.")
            take_screenshot(driver, doc, _sanitize_filename(descricao))
            return True
        except (TimeoutException, NoSuchElementException, StaleElementReferenceException) as e:
            if attempt < max_retries - 1:
                log(doc, f"⚠️ Tentativa {attempt + 1} falhou para {descricao}, tentando novamente...")
                time.sleep(2)
                continue
            else:
                log(doc, f"❌ Erro ao {descricao.lower()} após {max_retries} tentativas: {e}")
                take_screenshot(driver, doc, _sanitize_filename(f"erro_{descricao}"))
                return False
        except Exception as e:
            log(doc, f"❌ Erro inesperado ao {descricao.lower()}: {e}")
            take_screenshot(driver, doc, _sanitize_filename(f"erro_{descricao}"))
            return False


def clicar_primeiro_checkbox(driver, wait, timeout=10, max_tentativas=5):
    """
    Clica no primeiro <input type='checkbox'> visível encontrado na tela.
    Usa scroll, clique normal e fallback via JavaScript.
    """
    tentativa = 0
    while tentativa < max_tentativas:
        tentativa += 1
        try:
            checkboxes = wait.until(
                EC.presence_of_all_elements_located((By.CSS_SELECTOR, "input[type='checkbox']"))
            )

            if not checkboxes:
                log(doc, f"⚠️ Nenhum checkbox encontrado (tentativa {tentativa}/{max_tentativas}).")
                time.sleep(1)
                continue

            # Filtra apenas os visíveis e habilitados
            checkbox_visiveis = [cb for cb in checkboxes if cb.is_displayed() and cb.is_enabled()]
            if not checkbox_visiveis:
                log(doc, f"⚠️ Nenhum checkbox visível/habilitado (tentativa {tentativa}/{max_tentativas}).")
                time.sleep(1)
                continue

            checkbox = checkbox_visiveis[0]
            driver.execute_script("arguments[0].scrollIntoView({block:'center'});", checkbox)
            time.sleep(0.3)

            try:
                checkbox.click()
            except Exception:
                # fallback via JavaScript
                driver.execute_script("arguments[0].click();", checkbox)

            log(doc, f"✅ Checkbox clicado com sucesso na tentativa {tentativa}.")
            return True

        except Exception as e:
            log(doc, f"⚠️ Erro na tentativa {tentativa}: {e}")
            time.sleep(1)

    raise Exception(f"❌ Falha ao clicar no primeiro checkbox após {max_tentativas} tentativas.")


def selecionar_opcao(selector, texto):
    def acao():
        select_element = wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, selector)))
        Select(select_element).select_by_visible_text(texto)
    return acao

def abrir_modal_e_selecionar_robusto(btn_selector, pesquisa_selector, termo_pesquisa, btn_pesquisar_selector, resultado_xpath):
    """Versão robusta da função de modal"""
    global driver, wait, doc
    
    def acao():
        if driver is None or wait is None:
            raise Exception("Driver ou wait não inicializados")
        # Abre o modal
        safe_scroll_and_interact(btn_selector, "click")
        time.sleep(1)
        # Aguarda e preenche campo pesquisa
        campo_pesquisa = wait.until(
            EC.element_to_be_clickable((By.CSS_SELECTOR, pesquisa_selector))
        )
        campo_pesquisa.clear()
        campo_pesquisa.send_keys(termo_pesquisa)
        time.sleep(0.5)
        # Clica pesquisar
        pesquisar = wait.until(
            EC.element_to_be_clickable((By.CSS_SELECTOR, btn_pesquisar_selector))
        )
        pesquisar.click()
        time.sleep(2)
        # Aguarda resultado e clica
        resultado = wait.until(
            EC.element_to_be_clickable((By.XPATH, resultado_xpath))
        )
        driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", resultado)
        time.sleep(0.5)
        resultado.click()
        time.sleep(1)

    return acao


def finalizar_relatorio():
    global driver, doc
    nome_arquivo = f"relatorio_cadastro_de_cotacao_cenario_4_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    try:
        doc.save(nome_arquivo)
        log(doc, f"📄 Relatório salvo como: {nome_arquivo}")
        try:
            subprocess.run(["start", "winword", nome_arquivo], shell=True)
        except:
            pass
    except Exception as e:
        print(f"Erro ao salvar relatório: {e}")
    if driver:
        try:
            driver.quit()
        except:
            pass

def encontrar_mensagem_alerta():
    global driver, doc
    if driver is None:
        return None
    seletores = [
        (".alerts.salvo", "✅ Sucesso"),
        (".alerts.alerta", "⚠️ Alerta"),
        (".alerts.erro", "❌ Erro"),
    ]
    for seletor, tipo in seletores:
        try:
            elemento = driver.find_element(By.CSS_SELECTOR, seletor)
            if elemento.is_displayed():
                log(doc, f"📢 {tipo}: {elemento.text}")
                return elemento
        except:
            continue
    log(doc, "ℹ️ Nenhuma mensagem de alerta encontrada.")
    return None

def ajustar_zoom():
    global driver, doc
    if driver is None:
        return
    try:
        driver.execute_script("document.body.style.zoom='90%'")
        log(doc, "🔍 Zoom ajustado para 90%.")
    except Exception as e:
        log(doc, f"⚠️ Erro ao ajustar zoom: {e}")

def safe_scroll_and_interact(selector, action_type="click", value=None, timeout=10, by_xpath=False):
    """Rola até o elemento e interage com ele de forma robusta."""
    global driver, doc
    if driver is None:
        return None
    try:
        by_type = By.XPATH if by_xpath else By.CSS_SELECTOR
        element = WebDriverWait(driver, timeout).until(EC.presence_of_element_located((by_type, selector)))
        driver.execute_script("arguments[0].scrollIntoView({block: 'center', behavior: 'smooth'});", element)
        time.sleep(0.5)
        if action_type in ["click", "send_keys"]:
            element = WebDriverWait(driver, timeout).until(EC.element_to_be_clickable((by_type, selector)))
        if action_type == "click":
            element.click()
        elif action_type == "send_keys" and value:
            element.clear()
            element.send_keys(value)
        elif action_type == "select" and value:
            Select(element).select_by_visible_text(value)
        return element
    except Exception as e:
        log(doc, f"❌ Erro ao interagir com elemento {selector}: {e}")
        return None

# ==== CLICAR ROBUSTO ====
def clicar_elemento_robusto(driver, wait, seletor_css, timeout=10):
    global doc
    try:
        elem = wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, seletor_css)))
        try:
            driver.execute_script("""
                document.querySelectorAll('.modal, .overlay, .blockUI, .toast, .tooltip, [role="dialog"], [data-overlay]')
                .forEach(e => { if (getComputedStyle(e).position === 'fixed') e.style.display = 'none'; });
            """)
        except Exception:
            pass
        driver.execute_script("arguments[0].scrollIntoView({block:'center', inline:'center'});", elem)
        time.sleep(0.2)
        try:
            elem = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, seletor_css)))
            elem.click()
            return True
        except (TimeoutException, ElementClickInterceptedException, StaleElementReferenceException):
            pass
        try:
            elem = driver.find_element(By.CSS_SELECTOR, seletor_css)
            ActionChains(driver).move_to_element(elem).pause(0.05).click().perform()
            return True
        except Exception:
            pass
        try:
            elem = driver.find_element(By.CSS_SELECTOR, seletor_css)
            driver.execute_script("arguments[0].click();", elem)
            return True
        except Exception:
            pass
        try:
            elem = driver.find_element(By.CSS_SELECTOR, seletor_css)
            driver.execute_script("""
                const el = arguments[0];
                function fire(type){ el.dispatchEvent(new MouseEvent(type,{bubbles:true,cancelable:true,view:window})); }
                el.focus(); fire('mouseover'); fire('mousedown'); fire('mouseup'); fire('click');
            """, elem)
            return True
        except Exception:
            pass
        try:
            elem = driver.find_element(By.CSS_SELECTOR, seletor_css)
            ActionChains(driver).move_to_element_with_offset(elem, 1, 1).click().perform()
            return True
        except Exception:
            pass
        log(doc, f"❌ Não foi possível clicar em: {seletor_css}")
        return False
    except Exception as e:
        log(doc, f"❌ Erro ao clicar robusto: {e}")
        return False

# ==== DATEPICKER ====
def encontrar_campos_datepicker():
    global driver
    seletores = [
        "input.hasDatepicker",
        "input[id^='dp']",
        "input[class*='datepicker']",
        ".hasDatepicker",
        "[data-provide='datepicker']"
    ]
    campos = []
    for seletor in seletores:
        try:
            for elemento in driver.find_elements(By.CSS_SELECTOR, seletor):
                if elemento.is_displayed():
                    cid = elemento.get_attribute('id')
                    if not any(c.get('id') == cid for c in campos):
                        campos.append({'elemento': elemento, 'id': cid})
        except:
            continue
    log(doc, f"📊 Encontrados {len(campos)} campos datepicker")
    return campos

def _datepicker_jquery(campo_id, data_valor):
    global driver
    resultado = driver.execute_script("""
        var campoId = arguments[0], valor = arguments[1];
        if (typeof jQuery === 'undefined') return 'jQuery não disponível';
        var $campo = $('#' + campoId);
        if (!$campo.length) return 'Campo não encontrado: ' + campoId;
        try {
            if ($campo.hasClass('hasDatepicker')) { $campo.datepicker('setDate', valor); }
            else { $campo.val(valor); }
            $campo.trigger('input').trigger('change').trigger('blur');
            return $campo.val();
        } catch(e) { return 'Erro: ' + e.message; }
    """, campo_id, data_valor)
    if isinstance(resultado, str) and 'Erro' in resultado:
        raise Exception(f"jQuery falhou: {resultado}")

def _datepicker_javascript(elemento, data_valor):
    global driver
    driver.execute_script("""
        var campo = arguments[0], valor = arguments[1];
        campo.focus(); campo.value = ''; campo.value = valor;
        ['input','change','blur','keyup'].forEach(ev => campo.dispatchEvent(new Event(ev,{bubbles:true})));
    """, elemento, data_valor)

def _datepicker_actionchains(elemento, data_valor):
    global driver
    driver.execute_script("arguments[0].scrollIntoView({block:'center'});", elemento)
    ActionChains(driver).move_to_element(elemento).click().perform()
    time.sleep(0.3)
    ActionChains(driver).key_down(Keys.CONTROL).send_keys('a').key_up(Keys.CONTROL).perform()
    ActionChains(driver).send_keys(Keys.DELETE).perform()
    time.sleep(0.2)
    for ch in data_valor:
        ActionChains(driver).send_keys(ch).perform()
        time.sleep(0.03)
    ActionChains(driver).send_keys(Keys.TAB).perform()

def _datepicker_tradicional(elemento, data_valor):
    driver.execute_script("arguments[0].scrollIntoView({block:'center'});", elemento)
    elemento.click(); time.sleep(0.2); elemento.clear()
    elemento.send_keys(data_valor); elemento.send_keys(Keys.TAB)

def preencher_datepicker_persistente(indice_campo, data_valor, max_tentativas=10, timeout=30):
    """Preenche datepicker com várias estratégias e valida."""
    inicio_tempo = time.time()
    tentativa = 0

    def validar_data_preenchida(el, data_esperada):
        try:
            val = (el.get_attribute('value') or '').strip()
            if not val: return False
            if val == data_esperada or data_esperada in val: return True
            formatos = [
                '%d/%m/%Y %H:%M', '%d/%m/%Y %H:%M:%S',
                '%d/%m/%Y', '%m/%d/%Y', '%Y-%m-%d', '%d-%m-%Y'
            ]
            for f in formatos:
                try:
                    d1 = datetime.strptime(val, f)
                    d2 = datetime.strptime(data_esperada, f)
                    if d1 == d2: return True
                except:
                    continue
            return False
        except:
            return False

    while tentativa < max_tentativas and (time.time() - inicio_tempo) < timeout:
        tentativa += 1
        try:
            log(doc, f"🔄 Tentativa {tentativa}/{max_tentativas} para campo {indice_campo}")
            campos = encontrar_campos_datepicker()
            if not campos:
                time.sleep(1); continue
            if indice_campo >= len(campos):
                time.sleep(1); continue

            info = campos[indice_campo]
            elemento, campo_id = info['elemento'], info['id']

            if validar_data_preenchida(elemento, data_valor):
                log(doc, f"✅ Campo {indice_campo} já está preenchido corretamente!")
                return True

            estrategias = [
                lambda: _datepicker_jquery(campo_id, data_valor),
                lambda: _datepicker_javascript(elemento, data_valor),
                lambda: _datepicker_actionchains(elemento, data_valor),
                lambda: _datepicker_tradicional(elemento, data_valor),
            ]
            for acao in estrategias:
                try:
                    acao(); time.sleep(0.4)
                    if validar_data_preenchida(elemento, data_valor):
                        log(doc, f"✅ Campo {indice_campo} preenchido!")
                        return True
                except Exception as e:
                    log(doc, f"   ⚠️ Estratégia falhou: {e}")

            log(doc, "❌ Tentativa falhou; tentando novamente...")
            time.sleep(1.2)

        except Exception as e:
            log(doc, f"❌ Erro na tentativa {tentativa}: {e}")
            time.sleep(1)

    raise Exception(f"Falha ao preencher datepicker {indice_campo} após {tentativa} tentativas em {int(time.time()-inicio_tempo)}s")

def preencher_campo_xpath_com_retry(driver, wait, xpath, valor, max_tentativas=3):
    global doc
    if driver is None or wait is None:
        return False
    for tentativa in range(max_tentativas):
        try:
            campo = wait.until(EC.element_to_be_clickable((By.XPATH, xpath)))
            driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", campo)
            time.sleep(0.3)
            if tentativa == 0:
                campo.click(); campo.clear(); campo.send_keys(valor); campo.send_keys(Keys.TAB)
            elif tentativa == 1:
                ActionChains(driver).move_to_element(campo).click().pause(0.1).key_down(Keys.CONTROL).send_keys('a').key_up(Keys.CONTROL).send_keys(valor).send_keys(Keys.TAB).perform()
            else:
                driver.execute_script("""
                    var el = arguments[0], v = arguments[1];
                    el.focus(); el.value = ''; el.value = v;
                    el.dispatchEvent(new Event('input', { bubbles: true }));
                    el.dispatchEvent(new Event('change', { bubbles: true }));
                    el.blur();
                """, campo, valor)
            time.sleep(0.3)
            if (campo.get_attribute('value') or '').strip():
                return True
        except Exception as e:
            log(doc, f"⚠️ Tentativa {tentativa + 1} falhou: {e}")
            time.sleep(0.8)
    return False

# ==== WIZARD (clique por ícone) ====
def get_xpath(we):
    js = """
    function absoluteXPath(element) {
      if (element.tagName.toLowerCase() === 'html') return '/html';
      if (element===document.body) return '/html/body';
      var ix=0; var siblings=element.parentNode.children;
      var same = 0;
      for (var i=0; i<siblings.length; i++) {
        var sib=siblings[i];
        if (sib.tagName===element.tagName) {
          same++;
          if (sib===element) {
            var path = absoluteXPath(element.parentNode) + '/' + element.tagName.toLowerCase();
            if (same>1) path += '['+same+']';
            return path;
          }
        }
      }
      return absoluteXPath(element.parentNode) + '/' + element.tagName.toLowerCase();
    }
    return absoluteXPath(arguments[0]);
    """
    return we.parent.execute_script(js, we)

def _prepare_focus_and_clear(elemento, limpar_primeiro=True):
    driver.execute_script("arguments[0].scrollIntoView({block:'center'});", elemento)
    try:
        elemento.click()
    except Exception:
        driver.execute_script("arguments[0].focus();", elemento)

    if limpar_primeiro:
        try:
            elemento.clear()
        except Exception:
            ActionChains(driver)\
                .move_to_element(elemento).click()\
                .key_down(Keys.CONTROL).send_keys("a").key_up(Keys.CONTROL)\
                .send_keys(Keys.DELETE).perform()

def _textarea_tradicional(elemento, texto, limpar_primeiro=True):
    _prepare_focus_and_clear(elemento, limpar_primeiro)
    elemento.send_keys(texto)
    elemento.send_keys(Keys.TAB)

def _textarea_actionchains(elemento, texto, limpar_primeiro=True):
    _prepare_focus_and_clear(elemento, limpar_primeiro)
    ac = ActionChains(driver)
    ac.move_to_element(elemento).click().perform()
    for chunk_start in range(0, len(texto), 400):
        ac.send_keys(texto[chunk_start:chunk_start+400]).perform()
        time.sleep(0.05)
    ac.send_keys(Keys.TAB).perform()

def _textarea_js_setvalue(elemento, texto):
    driver.execute_script("""
        const el = arguments[0];
        const val = arguments[1];
        el.value = val;
        el.dispatchEvent(new Event('input',  {bubbles:true}));
        el.dispatchEvent(new Event('change', {bubbles:true}));
        el.dispatchEvent(new KeyboardEvent('keyup', {bubbles:true}));
        el.dispatchEvent(new Event('blur',   {bubbles:true}));
    """, elemento, texto)

def _textarea_js_react_input(elemento, texto):
    driver.execute_script("""
        const el = arguments[0];
        const val = arguments[1];

        const desc = Object.getOwnPropertyDescriptor(window.HTMLTextAreaElement.prototype, 'value');
        if (desc && desc.set) {
            desc.set.call(el, val);
        } else {
            el.value = val;
        }
        el.dispatchEvent(new Event('input', {bubbles: true}));
        el.dispatchEvent(new Event('change', {bubbles: true}));
        el.dispatchEvent(new Event('blur', {bubbles: true}));
    """, elemento, texto)

def preencher_textarea_por_indice(indice_campo, texto, max_tentativas=5, limpar_primeiro=True):
    """Preenche um <textarea> pelo índice (ordem no DOM) usando estratégias múltiplas"""
    def acao():
        if not isinstance(indice_campo, int) or indice_campo < 0:
            raise ValueError(f"Índice inválido: {indice_campo}")
        if texto is None or not isinstance(texto, str):
            raise ValueError(f"Texto inválido: {texto!r}")

        tentativa = 0
        while tentativa < max_tentativas:
            tentativa += 1
            try:
                campos = encontrar_campos_textarea()
                if not campos:
                    if tentativa < max_tentativas:
                        log(doc, f"⚠️ Nenhuma <textarea> encontrada (tentativa {tentativa}/{max_tentativas})", "WARN")
                        time.sleep(1.5)
                        continue
                    raise Exception("Nenhuma <textarea> foi encontrada na página.")

                if indice_campo >= len(campos):
                    raise Exception(f"Índice {indice_campo} inválido. Encontradas {len(campos)} textareas.")

                campo_info = campos[indice_campo]
                elemento   = campo_info["elemento"]
                campo_id   = campo_info.get("id") or "(sem id)"
                campo_name = campo_info.get("name") or "(sem name)"

                log(doc, f"🎯 Tentativa {tentativa}: Preenchendo textarea {indice_campo} (ID: {campo_id}, name: {campo_name}) com {len(texto)} caracteres")

                if validar_textarea_preenchida(elemento, texto):
                    log(doc, f"✅ Textarea {indice_campo} já está com o valor desejado.")
                    return True

                estrategias = [
                    lambda: _textarea_tradicional(elemento, texto, limpar_primeiro),
                    lambda: _textarea_actionchains(elemento, texto, limpar_primeiro),
                    lambda: _textarea_js_setvalue(elemento, texto),
                    lambda: _textarea_js_react_input(elemento, texto),
                ]

                for i, estrategia in enumerate(estrategias, 1):
                    try:
                        log(doc, f"   ▶️ Estratégia {i}…")
                        estrategia()
                        time.sleep(0.8)
                        if validar_textarea_preenchida(elemento, texto):
                            val = (elemento.get_attribute("value") or "").strip()
                            log(doc, f"✅ Preenchido com sucesso pela estratégia {i}: '{val[:60]}{'…' if len(val) > 60 else ''}'")
                            return True
                        else:
                            log(doc, f"⚠️ Estratégia {i} não refletiu o valor esperado.", "WARN")
                    except (StaleElementReferenceException, JavascriptException, TimeoutException) as e:
                        log(doc, f"⚠️ Estratégia {i} falhou: {e}", "WARN")
                        try:
                            campos = encontrar_campos_textarea()
                            elemento = campos[indice_campo]["elemento"]
                        except Exception:
                            pass
                        continue

                if tentativa < max_tentativas:
                    log(doc, f"⚠️ Tentativa {tentativa} falhou; reintentando em 1.5s…", "WARN")
                    time.sleep(1.5)
                    continue
            except Exception as e:
                if tentativa < max_tentativas:
                    log(doc, f"⚠️ Erro na tentativa {tentativa}: {e}. Retentando…", "WARN")
                    time.sleep(1.5)
                    continue
                else:
                    raise

        raise Exception(f"Falha ao preencher textarea {indice_campo} após {max_tentativas} tentativas.")
    return acao



from selenium.webdriver.common.by import By
from selenium.webdriver.common.action_chains import ActionChains
from selenium.common.exceptions import StaleElementReferenceException, JavascriptException, TimeoutException
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.common.keys import Keys
import time

def _clear_resistente(el):
    try:
        el.clear()
    except Exception:
        pass
    el.send_keys(Keys.CONTROL, "a")
    el.send_keys(Keys.DELETE)

def _validar_input_preenchido(el, texto):
    try:
        val = (el.get_attribute("value") or "").strip()
        return val == (texto or "").strip()
    except Exception:
        return False

def _input_tradicional(el, texto, limpar_primeiro):
    el.click()
    if limpar_primeiro:
        _clear_resistente(el)
    el.send_keys(texto)
    # força blur para muitos autocompletes/mascaras
    el.send_keys(Keys.TAB)

def _input_actionchains(driver, el, texto, limpar_primeiro):
    ActionChains(driver).move_to_element(el).click().perform()
    if limpar_primeiro:
        _clear_resistente(el)
    el.send_keys(texto)
    ActionChains(driver).send_keys(Keys.TAB).perform()

def _input_js_setvalue(driver, el, texto):
    driver.execute_script(
        """
        const el = arguments[0], val = arguments[1];
        const setter = Object.getOwnPropertyDescriptor(window.HTMLInputElement.prototype, 'value').set;
        setter.call(el, val);
        el.dispatchEvent(new Event('input', {bubbles:true}));
        el.dispatchEvent(new Event('change', {bubbles:true}));
        """,
        el, texto
    )

def _input_js_react(driver, el, texto):
    driver.execute_script(
        """
        const el = arguments[0], val = arguments[1];
        if (!el) return;
        const last = el._valueTracker;
        el.value = val;
        if (last) { last.setValue(val); }
        el.dispatchEvent(new Event('input', {bubbles:true}));
        el.dispatchEvent(new Event('change', {bubbles:true}));
        """,
        el, texto
    )

def _coletar_inputs(driver, css_selector):
    # retorna lista de dicts com o elemento e metadados úteis
    elementos = driver.find_elements(By.CSS_SELECTOR, css_selector)
    out = []
    for e in elementos:
        out.append({
            "elemento": e,
            "id": e.get_attribute("id"),
            "name": e.get_attribute("name"),
            "class": e.get_attribute("class"),
            "placeholder": e.get_attribute("placeholder"),
        })
    return out

def preencher_input_por_indice(driver, wait, css_selector, indice, texto,
                               max_tentativas=5, limpar_primeiro=True, timeout_vis=8):
    """
    Preenche um <input> identificado por um CSS selector, escolhendo pelo índice (0-based).
    Usa múltiplas estratégias (click/send_keys, ActionChains, JS value setter, React dispatch).
    """
    if not isinstance(indice, int) or indice < 0:
        raise ValueError(f"Índice inválido: {indice}")
    if texto is None or not isinstance(texto, str):
        raise ValueError(f"Texto inválido: {texto!r}")

    tentativa = 0
    ultimo_erro = None

    while tentativa < max_tentativas:
        tentativa += 1
        try:
            # aguarda existir ao menos 1 que case com o seletor
            wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, css_selector)))
            campos = _coletar_inputs(driver, css_selector)
            if not campos:
                raise Exception("Nenhum input encontrado para o seletor: " + css_selector)

            if indice >= len(campos):
                raise Exception(f"Índice {indice} inválido. Encontrados {len(campos)} inputs para {css_selector}.")

            el = campos[indice]["elemento"]
            # aguarda ficar visível/ativado
            wait.until(EC.element_to_be_clickable(el))

            if _validar_input_preenchido(el, texto):
                return True

            estrategias = [
                lambda: _input_tradicional(el, texto, limpar_primeiro),
                lambda: _input_actionchains(driver, el, texto, limpar_primeiro),
                lambda: _input_js_setvalue(driver, el, texto),
                lambda: _input_js_react(driver, el, texto),
            ]

            for strat in estrategias:
                try:
                    strat()
                    time.sleep(0.6)
                    if _validar_input_preenchido(el, texto):
                        return True
                except (StaleElementReferenceException, JavascriptException, TimeoutException) as e:
                    ultimo_erro = e
                    # recaptura o elemento se estagnou
                    try:
                        campos = _coletar_inputs(driver, css_selector)
                        el = campos[indice]["elemento"]
                    except Exception:
                        pass
                    continue

            # re-tentativa completa
            time.sleep(1.0)
            continue

        except Exception as e:
            ultimo_erro = e
            time.sleep(1.0)
            continue

    raise Exception(f"Falha ao preencher input (seletor={css_selector}, índice={indice}) após {max_tentativas} tentativas. Último erro: {ultimo_erro}")

# --------- Fachada para os campos do fornecedor (por índice) ---------

def preencher_dados_fornecedor_por_indices(
    driver, wait,
    idx_contato: int,
    idx_telefone: int,
    idx_email: int,
    idx_condicao_pagamento: int,
    idx_forma_pagamento: int,
    contato="1799999999",
    telefone="1799999999",
    email="teste@teste.selenium.com",
    condicao="100",
    forma="Teste"
):
    """
    Preenche os campos do fornecedor por ÍNDICE.
    - idx_contato: índice do input .contato
    - idx_telefone: índice do input .telefone
    - idx_email: índice do input .email
    - idx_condicao_pagamento: índice do input .formaPagamento (condição)
    - idx_forma_pagamento: índice do input .tipoPagamento (forma)
    """
    # CSS max estrito, igual aos seus XPaths (tipo + maxlength + width) — ajuste se necessário
    css_contato   = "input.contato[type='text'][maxlength='30'][style*='width: 105px']"
    css_telefone  = "input.telefone[type='text'][maxlength='13'][style*='width:80px']"
    css_email     = "input.email[type='text'][maxlength='50'][style*='width: 132px']"

    # Para os autocompletes de pagamento, use a classe parcial como você já fazia:
    css_condicao  = "input.formaPagamento, input[class*='formaPagamento']"
    css_forma     = "input.tipoPagamento,  input[class*='tipoPagamento']"

    # Preenche cada um
    preencher_input_por_indice(driver, wait, css_contato,  idx_contato,  contato)
    preencher_input_por_indice(driver, wait, css_telefone, idx_telefone, telefone)
    preencher_input_por_indice(driver, wait, css_email,    idx_email,    email)
    preencher_input_por_indice(driver, wait, css_condicao, idx_condicao_pagamento, condicao)
    preencher_input_por_indice(driver, wait, css_forma,    idx_forma_pagamento,    forma)
    return True


# =========================
# Helpers usados pela função
# =========================

def encontrar_campos_textarea(timeout=10):
    """Retorna uma lista de dicts com metadados de cada <textarea> visível e interativa."""
    elementos = []
    try:
        wait.until(lambda d: len(d.find_elements(By.TAG_NAME, "textarea")) >= 0)
        textareas = driver.find_elements(By.TAG_NAME, "textarea")
    except Exception:
        textareas = []

    for el in textareas:
        try:
            if not el.is_displayed() or not el.is_enabled():
                continue
            driver.execute_script("arguments[0].scrollIntoView({block:'center'});", el)
            elementos.append({
                "elemento": el,
                "id": el.get_attribute("id"),
                "name": el.get_attribute("name"),
            })
        except Exception:
            continue

    return elementos

def normalizar_texto(txt):
    if txt is None:
        return ""
    return txt.replace("\r\n", "\n").replace("\r", "\n").strip()

def validar_textarea_preenchida(elemento, texto_esperado):
    """Confere se o valor atual da textarea bate com o texto esperado (normalizado)."""
    try:
        atual = elemento.get_attribute("value")
        if atual is None or atual == "":
            atual = (elemento.text or "")
        return normalizar_texto(atual) == normalizar_texto(texto_esperado)
    except StaleElementReferenceException:
        return False

def click_wizard_by_icon(driver, wait, icon_class, expect_selector=None, timeout=12):
    span = wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, f"#gsCrm .btnHolder .sprites.{icon_class}")))
    driver.execute_script("arguments[0].scrollIntoView({block:'center'});", span)
    alvo = driver.execute_script("""
        const span = arguments[0]; let el = span;
        while (el && el !== document && !/^(A|DIV)$/i.test(el.tagName)) el = el.parentElement;
        return el;
    """, span)
    if not alvo:
        raise Exception(f"Não achei ancestral clicável para o ícone .{icon_class}")
    try:
        xp = get_xpath(alvo)
        wait.until(EC.element_to_be_clickable((By.XPATH, xp))).click()
    except Exception:
        try:
            driver.execute_script("arguments[0].click();", alvo)
        except Exception:
            driver.execute_script("""
                const el=arguments[0];
                function fire(t){el.dispatchEvent(new MouseEvent(t,{bubbles:true,cancelable:true,view:window}));}
                el.focus(); fire('mouseover'); fire('mousedown'); fire('mouseup'); fire('click');
            """, alvo)
    if expect_selector:
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, expect_selector)))



def clicar_lov_por_indice(indice_lov: int, max_tentativas: int = 5, timeout: int = 10, scroll: bool = True):
    """
    Clica no ícone de LOV <a class="sprites sp-openLov"> pelo índice (ordem no DOM).
    Retorna uma função 'acao' para ser usada com safe_action(..., lambda: ...).

    Estratégias de clique:
      1) Espera 'clickable' + click nativo
      2) ScrollIntoView + click nativo
      3) JavaScript click
      4) ActionChains move_to_element + click
    """
    def acao():
        if not isinstance(indice_lov, int) or indice_lov < 0:
            raise ValueError(f"Índice inválido: {indice_lov}")

        tentativa = 0
        while tentativa < max_tentativas:
            tentativa += 1
            try:
                log(doc, f"🔎 Tentativa {tentativa}: Localizando ícones LOV ('sp-openLov')...")
                # Coleta atual dos elementos
                elementos = driver.find_elements(By.CSS_SELECTOR, "a.sprites.sp-openLov")

                if not elementos:
                    if tentativa < max_tentativas:
                        log(doc, f"⚠️ Nenhum ícone LOV encontrado (tentativa {tentativa}/{max_tentativas}). Reintentando...", "WARN")
                        time.sleep(1.2)
                        continue
                    raise Exception("Nenhum ícone LOV ('a.sprites.sp-openLov') foi encontrado na página.")

                if indice_lov >= len(elementos):
                    raise Exception(f"Índice {indice_lov} inválido. Encontrados {len(elementos)} ícones LOV.")

                # Mantém um localizador estável por índice para waits/refresh
                locator_xpath = f"(//a[contains(@class,'sp-openLov')])[{indice_lov + 1}]"

                # Reaponta o elemento pelo locator (evita stale)
                elemento = driver.find_element(By.XPATH, locator_xpath)

                log(doc, f"🎯 Preparando clique no LOV de índice {indice_lov} (total: {len(elementos)}).")

                def _wait_clickable():
                    wait.until(EC.element_to_be_clickable((By.XPATH, locator_xpath)))

                estrategias = [
                    # 1) Espera 'clickable' + click nativo
                    lambda: (_wait_clickable(), elemento.click()),
                    # 2) ScrollIntoView + click nativo
                    lambda: (
                        driver.execute_script("arguments[0].scrollIntoView({block:'center', inline:'nearest'});", elemento) if scroll else None,
                        time.sleep(0.2),
                        elemento.click()
                    ),
                    # 3) JavaScript click
                    lambda: driver.execute_script("arguments[0].click();", elemento),
                    # 4) ActionChains
                    lambda: (ActionChains(driver).move_to_element(elemento).pause(0.1).click().perform())
                ]

                for i, estrategia in enumerate(estrategias, 1):
                    try:
                        log(doc, f"   ▶️ Estratégia {i} de clique no LOV...")
                        estrategia()
                        time.sleep(0.3)
                        log(doc, f"✅ Clique no LOV (índice {indice_lov}) realizado com sucesso (estratégia {i}).")
                        return True
                    except (ElementClickInterceptedException, StaleElementReferenceException, JavascriptException, TimeoutException) as e:
                        log(doc, f"⚠️ Estratégia {i} falhou: {e}", "WARN")
                        # Tenta re-obter o elemento em caso de stale/interceptação
                        try:
                            _ = driver.find_elements(By.CSS_SELECTOR, "a.sprites.sp-openLov")
                            elemento = driver.find_element(By.XPATH, locator_xpath)
                        except Exception:
                            pass
                        continue

                if tentativa < max_tentativas:
                    log(doc, f"⚠️ Tentativa {tentativa} não conseguiu clicar no LOV. Reintentando em 1.2s…", "WARN")
                    time.sleep(1.2)
                    continue

            except Exception as e:
                if tentativa < max_tentativas:
                    log(doc, f"⚠️ Erro na tentativa {tentativa}: {e}. Reintentando em 1.2s…", "WARN")
                    time.sleep(1.2)
                    continue
                raise

        raise Exception(f"Falha ao clicar no LOV de índice {indice_lov} após {max_tentativas} tentativas.")

    return acao


def abrir_modal_e_selecionar_robusto_xpath(
    btn_xpath,
    pesquisa_xpath,
    termo_pesquisa,
    btn_pesquisar_xpath,
    resultado_xpath,
    timeout=12,
    max_tentativas=3,
    iframe_xpath=None,       # se o LOV abrir em iframe, informe o xpath aqui
    indice_lov=None,         # >>> usa clicar_lov_por_indice se informado
    **_ignorar_kwargs        # >>> evita "unexpected keyword argument" se algo extra vier
):
    """
    Abre o modal (LOV), pesquisa pelo termo e clica no resultado.
    - Opcional: abre o LOV clicando pelo índice (clicar_lov_por_indice)
    - Usa retries, fallback JS e limpeza resistente do input
    - Opcional: troca para iframe do modal
    """

    def _js_click(el):
        driver.execute_script("arguments[0].click();", el)

    def _clear_resistente(el):
        try:
            el.clear()
            el.send_keys(Keys.CONTROL, "a")
            el.send_keys(Keys.DELETE)
        except Exception:
            driver.execute_script("arguments[0].value='';", el)

    def _aguardar_ajax_overlay():
        t0 = time.time()
        while time.time() - t0 < 8:
            try:
                ready = driver.execute_script("return document.readyState")
                ajax_ok = driver.execute_script("return window.jQuery ? jQuery.active === 0 : true")
                if ready == "complete" and ajax_ok:
                    break
            except Exception:
                pass
            time.sleep(0.2)

    def acao():
        tentativa = 0
        while tentativa < max_tentativas:
            tentativa += 1
            try:
                # 1) Abrir o modal (por índice de LOV ou por XPATH do botão)
                if indice_lov is not None:
                    log(doc, f"🧭 Abrindo LOV pelo índice {indice_lov}…")
                    clicar_lov_por_indice(indice_lov, max_tentativas=3, timeout=timeout)()
                else:
                    log(doc, "🧭 Abrindo LOV pelo XPATH do botão…")
                    btn = WebDriverWait(driver, timeout).until(
                        EC.element_to_be_clickable((By.XPATH, btn_xpath))
                    )
                    try:
                        btn.click()
                    except Exception:
                        _js_click(btn)

                # 2) (Opcional) Entrar no iframe do modal
                if iframe_xpath:
                    WebDriverWait(driver, timeout).until(
                        EC.frame_to_be_available_and_switch_to_it((By.XPATH, iframe_xpath))
                    )

                # 3) Campo de pesquisa
                campo = WebDriverWait(driver, timeout).until(
                    EC.visibility_of_element_located((By.XPATH, pesquisa_xpath))
                )
                _clear_resistente(campo)
                campo.send_keys(termo_pesquisa)

                # 4) Botão Pesquisar
                btn_pesq = WebDriverWait(driver, timeout).until(
                    EC.element_to_be_clickable((By.XPATH, btn_pesquisar_xpath))
                )
                try:
                    btn_pesq.click()
                except Exception:
                    _js_click(btn_pesq)

                _aguardar_ajax_overlay()
                time.sleep(0.4)

                # 5) Clicar no resultado
                resultado = WebDriverWait(driver, timeout).until(
                    EC.element_to_be_clickable((By.XPATH, resultado_xpath))
                )
                driver.execute_script("arguments[0].scrollIntoView({block:'center'});", resultado)
                try:
                    resultado.click()
                except Exception:
                    _js_click(resultado)

                # 6) Sair do iframe (se entrou)
                if iframe_xpath:
                    driver.switch_to.default_content()

                time.sleep(0.6)
                return True

            except Exception as e:
                try:
                    driver.switch_to.default_content()
                except Exception:
                    pass

                if tentativa < max_tentativas:
                    log(doc, f"⚠️ Tentativa {tentativa}/{max_tentativas} falhou: {e}. Retentando…")
                    time.sleep(1.0 + 0.3 * tentativa)
                    continue
                else:
                    raise

    return acao



def preencher_todos_valores_e_quantidades(
    driver, wait,
    valor_padrao="10,00",
    quantidade_padrao="1",
    valores_por_indice=None,         # ex.: ["10,00","12,50","9,90"]
    quantidades_por_indice=None,     # ex.: ["2","1","3"]
    pausa=0.3
):
    """
    Procura APENAS input.vu (valor unitário) e input.innerQtdField (quantidade),
    loga a quantidade e preenche TODOS os pares encontrados.

    - Se 'valores_por_indice' / 'quantidades_por_indice' forem passados, usa-os por índice.
      Caso a lista não tenha valor para algum índice, usa os padrões.
    """
    try:
        valores = driver.find_elements(By.CSS_SELECTOR, "input.vu")
        quantds = driver.find_elements(By.CSS_SELECTOR, "input.innerQtdField")

        n_val = len(valores)
        n_qtd = len(quantds)
        n = min(n_val, n_qtd)

        if n == 0:
            raise Exception(f"Nenhum par de campos encontrado. vu={n_val}, innerQtdField={n_qtd}")

        log(doc, f"🔎 Encontrados {n} pares de campos input.vu e input.innerQtdField, preenchendo...")

        if n_val != n_qtd:
            log(doc, f"⚠️ Tamanhos diferentes detectados (vu={n_val}, qtd={n_qtd}). Vou preencher até {n}.",)

        def _clear_and_type(el, txt):
            driver.execute_script("arguments[0].scrollIntoView({block:'center'});", el)
            try:
                el.clear()
            except Exception:
                driver.execute_script("arguments[0].value='';", el)
            el.send_keys(txt)
            el.send_keys(Keys.TAB)

        for i in range(n):
            v = (valores_por_indice[i] if (isinstance(valores_por_indice, list) and i < len(valores_por_indice)) else valor_padrao)
            q = (quantidades_por_indice[i] if (isinstance(quantidades_por_indice, list) and i < len(quantidades_por_indice)) else quantidade_padrao)

            try:
                _clear_and_type(valores[i], v);  time.sleep(pausa)
                _clear_and_type(quantds[i], q);  time.sleep(pausa)
                log(doc, f"✅ Produto {i+1}: Valor={v} | Quantidade={q}")
            except Exception as e:
                log(doc, f"❌ Falha ao preencher item {i+1}: {e}")

        return True

    except Exception as e:
        log(doc, f"❌ Erro ao preencher todos os valores/quantidades: {e}")
        return False

def fechar_modal_cotacoes():
    try:
        el = wait.until(EC.presence_of_element_located(
            (By.XPATH, "//a[@class='sprites sp-fecharGrande' and @title='Sair']")))
        driver.execute_script("arguments[0].scrollIntoView({block:'center'});", el)
        driver.execute_script("arguments[0].click();", el)
        log(doc, "✅ Modal de Cotações fechado.")
        return True
    except Exception as e:
        log(doc, f"⚠️ Falha ao fechar Modal de Cotações: {e}")
        return False

def fechar_modal_gestor_compras():
    try:
        clicar_elemento_robusto(driver, wait, '#gsCompras > div.wdTop.ui-draggable-handle > div > a')
        log(doc, "✅ Gestor de Compras fechado.")
        return True
    except Exception as e:
        log(doc, f"⚠️ Falha ao fechar Gestor de Compras: {e}")
        return False

def tratar_resultado_cotacoes():
    """
    Verifica se há cotações após clicar em 'Buscar'.
    - Se não houver: loga, fecha modais e retorna False.
    - Se houver: clica no primeiro checkbox e retorna True.
    """
    try:
        # Aguarda um resultado OU conclui que não há (timeout curto)
        time.sleep(0.6)  # pequeno respiro para render
        # tente encontrar checkboxes dentro da listagem
        checkboxes = driver.find_elements(By.CSS_SELECTOR, "table tbody tr input[type='checkbox']")
        if not checkboxes:
            # fallback mais amplo (caso a tabela use outro markup)
            checkboxes = driver.find_elements(By.CSS_SELECTOR, "input[type='checkbox']")

        # filtra visíveis e habilitados
        checkboxes = [cb for cb in checkboxes if cb.is_displayed() and cb.is_enabled()]

        if not checkboxes:
            log(doc, "❌ Nenhuma Cotação encontrada")
            # fecha os modais conforme solicitado
            fechar_modal_cotacoes()
            fechar_modal_gestor_compras()
            return False

        # clica no primeiro checkbox
        cb = checkboxes[0]
        driver.execute_script("arguments[0].scrollIntoView({block:'center'});", cb)
        try:
            cb.click()
        except Exception:
            driver.execute_script("arguments[0].click();", cb)
        log(doc, "✅ Primeira cotação marcada (checkbox).")
        return True

    except Exception as e:
        log(doc, f"❌ Erro ao tratar resultado de cotações: {e}")
        return False



# ==== DRIVER ====
def inicializar_driver():
    global driver, wait
    try:
        options = Options()
        options.add_argument("--start-maximized")
        options.add_argument("--disable-blink-features=AutomationControlled")
        options.add_experimental_option("excludeSwitches", ["enable-automation"])
        options.add_experimental_option('useAutomationExtension', False)

        driver = webdriver.Chrome(service=Service(ChromeDriverManager().install()), options=options)
        driver.execute_script("Object.defineProperty(navigator, 'webdriver', {get: () => undefined})")
        wait = WebDriverWait(driver, 20)
        return True
    except Exception as e:
        log(doc, f"❌ Erro ao inicializar driver: {e}")
        return False

# ==== EXECUÇÃO DO TESTE ====
def executar_teste():
    global driver, wait, doc
    try:
        if not inicializar_driver():
            return False

        # Sanidade: garantir que time é módulo
        assert hasattr(time, "sleep"), f"time virou {time!r}"

        # ---------- Funções compostas (substituem lambdas com várias expressões) ----------
        def realizar_login():
            wait.until(EC.presence_of_element_located((By.ID, "j_id15:email"))).send_keys(LOGIN_EMAIL)
            wait.until(EC.presence_of_element_located((By.ID, "j_id15:senha"))).send_keys(LOGIN_PASSWORD, Keys.ENTER)
            time.sleep(5)

        def ajustar_zoom_e_abrir_menu():
            ajustar_zoom()
            driver.find_element(By.TAG_NAME, "body").send_keys(Keys.F3)

        def acessar_gestor_de_compras():
            driver.execute_script(
                "arguments[0].scrollIntoView({block: 'center'});",
                wait.until(EC.presence_of_element_located((By.XPATH, '/html/body/div[15]/ul/li[3]/img')))
            )
            wait.until(EC.element_to_be_clickable((By.XPATH, '/html/body/div[15]/ul/li[3]/img'))).click()

        def avancar_para_servicos():
            time.sleep(2)
            wait.until(EC.element_to_be_clickable((By.XPATH, '/html/body/div[16]/div[2]/div[2]/div[2]/div[3]/a[3]'))).click()

        def abrir_lov_servicos():
            time.sleep(1)
            wait.until(EC.element_to_be_clickable((
                By.XPATH, "//*[@id='gsCompras']/div[2]/div[2]/div[2]/div/div[3]/div/div[1]/ul/li[1]/div/div/a"
            ))).click()
        # -------------------------------------------------------------------------------

        safe_action(doc, "Acessando sistema", lambda: driver.get(URL))
        safe_action(doc, "Realizando login", realizar_login)
        safe_action(doc, "Ajustando zoom e abrindo menu", ajustar_zoom_e_abrir_menu)
        safe_action(doc, "Acessando Gestor de Compras", acessar_gestor_de_compras)

        safe_action(doc, "Clicando em Cotações", lambda:
            wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, '#gsCompras > div.wdTelas > div > ul > li:nth-child(4) > a > span'))).click()
        )

        time.sleep(5)

        safe_action(doc, "Preenchendo Data Aprovação inicial", lambda: preencher_campo_xpath_com_retry(driver, wait, "//input[@type='text' and @class='hasDatepicker dataI' and @maxlength='10' and @style='width: 100px;' and @grupo='']", "06/03/2025"))


        safe_action(doc, "Buscando cotações", lambda:
            wait.until(EC.element_to_be_clickable((By.XPATH, "//a[contains(normalize-space(.), 'Buscar')]"))).click()
        )
        time.sleep(20)
        # Após clicar em "Buscar"

 
        resultado_ok = safe_action(doc, "Verificando resultado da busca de cotações", tratar_resultado_cotacoes)

        if not resultado_ok:
            return False
        else:
            log(doc, "➡️ Cotação encontrada. Prosseguindo com as próximas ações...")


        safe_action(doc, "Adicionando Cotação", lambda:
            wait.until(EC.element_to_be_clickable((By.XPATH, "//a[contains(normalize-space(.), 'Adicionar cotação')]"))).click()
        )



        safe_action(doc, "Selecionando Fornecedor",
            abrir_modal_e_selecionar_robusto_xpath(
                btn_xpath=None,  # ignorado quando usar indice_lov
                pesquisa_xpath="//input[@id='txtPesquisa']",
                termo_pesquisa="FORNECEDOR TESTE JOÃO",
                btn_pesquisar_xpath="//a[contains(@class,'lpFind') and contains(normalize-space(.),'Pesquisar')]",
                resultado_xpath="//tr[td[normalize-space(.)='FORNECEDOR TESTE JOÃO']]",
                indice_lov=0 
            )
        )



        log(doc, "🔍 Verificando mensagens de alerta...")
        encontrar_mensagem_alerta()

        safe_action(doc, "Preenchendo dados do Fornecedor por índice", lambda:
        preencher_dados_fornecedor_por_indices(
            driver, wait,
            idx_contato=0,
            idx_telefone=0,
            idx_email=0,
            idx_condicao_pagamento=0,
            idx_forma_pagamento=0,
            contato="1799999999",
            telefone="1799999999",
            email="teste@teste.selenium.com",
            condicao="100",
            forma="Teste"
        )
    )

        safe_action(doc, "Preenchendo todos os valores e quantidades (padrão)", lambda:
            preencher_todos_valores_e_quantidades(driver, wait, valor_padrao="10,00", quantidade_padrao="2")
        )



        safe_action(doc, "Adicionando Nova Cotação", lambda:
            wait.until(EC.element_to_be_clickable((By.XPATH, "//a[@class='btModel btGray btsave fRight'][span[@class='sprites sp-addVerde'] and contains(., 'Nova Cotação')]"))).click()
        )



        safe_action(doc, "Selecionando Fornecedor",
            abrir_modal_e_selecionar_robusto_xpath(
                btn_xpath=None,  # ignorado quando usar indice_lov
                pesquisa_xpath="//input[@id='txtPesquisa']",
                termo_pesquisa="FORNECEDOR TESTE LA PACE",
                btn_pesquisar_xpath="//a[contains(@class,'lpFind') and contains(normalize-space(.),'Pesquisar')]",
                resultado_xpath="//tr[td[normalize-space(.)='FORNECEDOR TESTE LA PACE']]",
                indice_lov=1
            )
        )


        safe_action(doc, "Preenchendo dados do Fornecedor por índice", lambda:
            preencher_dados_fornecedor_por_indices(
                driver, wait,
                idx_contato=1,
                idx_telefone=1,
                idx_email=1,
                idx_condicao_pagamento=1,
                idx_forma_pagamento=1,
                contato="1799999999",
                telefone="1799999999",
                email="teste@teste.selenium.com",
                condicao="100",
                forma="Teste"
            )
        )


        safe_action(doc, "Preenchendo todos os valores e quantidades (padrão)", lambda:
            preencher_todos_valores_e_quantidades(driver, wait, valor_padrao="30,00", quantidade_padrao="3")
        )



        safe_action(doc, "Adicionando Nova Cotação", lambda:
            wait.until(EC.element_to_be_clickable((By.XPATH, "//a[@class='btModel btGray btsave fRight'][span[@class='sprites sp-addVerde'] and contains(., 'Nova Cotação')]"))).click()
        )

        safe_action(doc, "Selecionando Fornecedor",
            abrir_modal_e_selecionar_robusto_xpath(
                btn_xpath=None,  # ignorado quando usar indice_lov
                pesquisa_xpath="//input[@id='txtPesquisa']",
                termo_pesquisa="FORNECEDOR JOÃO TESTE 2",
                btn_pesquisar_xpath="//a[contains(@class,'lpFind') and contains(normalize-space(.),'Pesquisar')]",
                resultado_xpath="//tr[td[normalize-space(.)='FORNECEDOR JOÃO TESTE 2']]",
                indice_lov=2
            )
        )

        safe_action(doc, "Preenchendo dados do Fornecedor por índice", lambda:
            preencher_dados_fornecedor_por_indices(
                driver, wait,
                idx_contato=2,
                idx_telefone=2,
                idx_email=2,
                idx_condicao_pagamento=2,
                idx_forma_pagamento=2,
                contato="1799999999",
                telefone="1799999999",
                email="teste@teste.selenium.com",
                condicao="100",
                forma="Teste"
            )
        )
        safe_action(doc, "Preenchendo todos os valores e quantidades (padrão)", lambda:
            preencher_todos_valores_e_quantidades(driver, wait, valor_padrao="5,00", quantidade_padrao="5")
        )


        safe_action(doc, "Finalizando", lambda:
            driver.execute_script(
                "arguments[0].click();",
                driver.find_element(By.XPATH, "//a[contains(@class,'btsave') and .//span[contains(@class,'sp-salvar')]]")
            )
        )
        time.sleep(2)


        safe_action(doc, "Fechando Modal de Cotações", lambda:
            driver.execute_script(
                "arguments[0].click();",
                driver.find_element(By.XPATH, "//a[@class='sprites sp-fecharGrande' and @title='Sair']")
            )
        )
        
        log(doc, "🔍 Verificando mensagens de alerta...")
        encontrar_mensagem_alerta()

        time.sleep(2)


        safe_action(doc, "Fechando modal Gestor de Compras", lambda:
            clicar_elemento_robusto(driver, wait, '#gsCompras > div.wdTop.ui-draggable-handle > div > a')
        )


        return True

    except Exception as e:
        log(doc, f"❌ ERRO FATAL: {e}")
        take_screenshot(driver, doc, "erro_fatal")
        return False
    finally:
        log(doc, "✅ Teste concluído.")

# ==== MAIN ====
def main():
    global doc
    try:
        log(doc, "🚀 Iniciando teste de Cadastro de Cotação")
        sucesso = executar_teste()
        if sucesso:
            log(doc, "✅ Teste executado com sucesso!")
        else:
            log(doc, "❌ Teste finalizado com erros.")
    except Exception as e:
        log(doc, f"❌ Erro na execução principal: {e}")
    finally:
        finalizar_relatorio()

if __name__ == "__main__":
    main()
