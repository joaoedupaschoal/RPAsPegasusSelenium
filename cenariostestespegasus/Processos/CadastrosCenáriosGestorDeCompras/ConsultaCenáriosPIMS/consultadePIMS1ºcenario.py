from datetime import datetime, timedelta
from datetime import time as dt_time
import time
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver import ActionChains
from selenium.common.exceptions import TimeoutException, StaleElementReferenceException, JavascriptException, NoSuchElementException, ElementClickInterceptedException
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.support.ui import Select, WebDriverWait
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service
from webdriver_manager.chrome import ChromeDriverManager
from docx import Document
from docx.shared import Inches
from faker import Faker
from faker.providers import BaseProvider
import os
import re
import pyautogui

# Inicializações
doc = Document()
doc.add_heading("RELATÓRIO DO TESTE", 0)
doc.add_paragraph("Consulta de PIMS - Gestor de Compras – Cenário 1: Preenchimento completo e salvamento")
doc.add_paragraph(f"Data do teste: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")
driver = None
wait = None
screenshot_registradas = set()

def log(doc, msg):
    print(msg)
    doc.add_paragraph(msg)

def _sanitize_filename(name: str) -> str:
    name = name.strip().lower()
    name = re.sub(r"[<>:\"/\\|?*']", "_", name)
    name = re.sub(r"_+", "_", name)
    return name[:120]

def take_screenshot(driver, doc, nome):
    if driver is None:
        return
    nome = _sanitize_filename(nome)
    if nome not in screenshot_registradas:
        path = f"screenshots/{nome}.png"
        os.makedirs("screenshots", exist_ok=True)
        try:
            driver.save_screenshot(path)
            doc.add_paragraph(f"Screenshot: {nome}")
            doc.add_picture(path, width=Inches(5.5))
            screenshot_registradas.add(nome)
        except Exception as e:
            log(doc, f"⚠️ Erro ao tirar screenshot {nome}: {e}")

def aguardar_overlay_sumir(timeout=15):
    t0 = time.time()
    while time.time() - t0 < timeout:
        try:
            # se existir jQuery, espera AJAX zerar
            ajax_ok = driver.execute_script("return !!window.jQuery ? (jQuery.active === 0) : true")
            # espera o overlay sair/ficar invisível
            overlay_visivel = driver.execute_script("""
                const o = document.querySelector('div.blockScreen');
                if (!o) return false;
                const st = window.getComputedStyle(o);
                const op = parseFloat(st.opacity || '1');
                const disp = st.display !== 'none';
                const vis = st.visibility !== 'hidden';
                // visível se display!=none e opacity>0
                return disp && vis && op > 0.01;
            """)
            if ajax_ok and not overlay_visivel:
                return True
        except Exception:
            pass
        time.sleep(0.2)
    return True  # não falha o fluxo, só segue

def selecionar_linha_item_por_indice(indice_base_1):
    # Tenta selecionar a <li> do item (ancora a seleção na linha)
    driver.execute_script("""
        const i = arguments[0];
        const li = document.querySelector(`#gsCompras div.consultarPims li:nth-child(${i})`);
        if (!li) throw new Error("Linha do item não encontrada: " + i);
        // clique na linha para manter seleção/FOCUS do item
        li.dispatchEvent(new MouseEvent('mousedown', {bubbles:true}));
        li.dispatchEvent(new MouseEvent('mouseup', {bubbles:true}));
        li.click();
    """, indice_base_1)

def safe_action(doc, descricao, func, max_retries=3):
    global driver
    for attempt in range(max_retries):
        try:
            log(doc, f"🔄 {descricao}..." if attempt == 0 else f"🔄 {descricao}... (Tentativa {attempt + 1})")
            func()
            log(doc, f"✅ {descricao} realizada com sucesso.")
            take_screenshot(driver, doc, _sanitize_filename(descricao))
            return True
        except (TimeoutException, NoSuchElementException, StaleElementReferenceException) as e:
            if attempt < max_retries - 1:
                log(doc, f"⚠️ Tentativa {attempt + 1} falhou para {descricao}, tentando novamente...")
                time.sleep(2)
                continue
            else:
                log(doc, f"❌ Erro ao {descricao.lower()} após {max_retries} tentativas: {e}")
                take_screenshot(driver, doc, _sanitize_filename(f"erro_{descricao}"))
                return False
        except Exception as e:
            log(doc, f"❌ Erro inesperado ao {descricao.lower()}: {e}")
            take_screenshot(driver, doc, _sanitize_filename(f"erro_{descricao}"))
            return False

def clicar_elemento_robusto(driver, wait, seletor_css, timeout=10):
    global doc
    try:
        elem = wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, seletor_css)))
        driver.execute_script("arguments[0].scrollIntoView({block:'center', inline:'center'});", elem)
        time.sleep(0.2)
        try:
            elem = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, seletor_css)))
            elem.click()
            return True
        except (TimeoutException, ElementClickInterceptedException, StaleElementReferenceException):
            pass
        # Try alternate click methods:
        try:
            elem = driver.find_element(By.CSS_SELECTOR, seletor_css)
            ActionChains(driver).move_to_element(elem).pause(0.05).click().perform()
            return True
        except Exception:
            pass
        try:
            elem = driver.find_element(By.CSS_SELECTOR, seletor_css)
            driver.execute_script("arguments[0].click();", elem)
            return True
        except Exception:
            pass
        try:
            elem = driver.find_element(By.CSS_SELECTOR, seletor_css)
            ActionChains(driver).move_to_element_with_offset(elem, 1, 1).click().perform()
            return True
        except Exception:
            pass
        log(doc, f"❌ Não foi possível clicar em: {seletor_css}")
        return False
    except Exception as e:
        log(doc, f"❌ Erro ao clicar robusto: {e}")
        return False

def preencher_campo_xpath_com_retry(driver, wait, xpath, valor, max_tentativas=3):
    global doc
    if driver is None or wait is None:
        return False
    for tentativa in range(max_tentativas):
        try:
            campo = wait.until(EC.element_to_be_clickable((By.XPATH, xpath)))
            driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", campo)
            time.sleep(0.3)
            if tentativa == 0:
                campo.click()
                campo.clear()
                campo.clear()
                campo.send_keys(valor)
                campo.send_keys(Keys.TAB)
            elif tentativa == 1:
                ActionChains(driver).move_to_element(campo).click().pause(0.1).key_down(Keys.CONTROL).send_keys('a').key_up(Keys.CONTROL).send_keys(valor).send_keys(Keys.TAB).perform()
            else:
                driver.execute_script("""
                var el = arguments[0], v = arguments[1];
                el.focus(); el.value = ''; el.value = v;
                el.dispatchEvent(new Event('input', { bubbles: true }));
                el.dispatchEvent(new Event('change', { bubbles: true }));
                el.blur();
                """, campo, valor)
            time.sleep(0.3)
            if (campo.get_attribute('value') or '').strip():
                return True
        except Exception as e:
            log(doc, f"⚠️ Tentativa {tentativa + 1} falhou: {e}")
            time.sleep(0.8)
    return False

def encontrar_campos_textarea(timeout=10):
    elementos = []
    try:
        wait.until(lambda d: len(d.find_elements(By.TAG_NAME, "textarea")) >= 0)
        textareas = driver.find_elements(By.TAG_NAME, "textarea")
    except Exception:
        textareas = []
    for el in textareas:
        try:
            if not el.is_displayed() or not el.is_enabled():
                continue
            driver.execute_script("arguments[0].scrollIntoView({block:'center'});", el)
            elementos.append({
                "elemento": el,
                "id": el.get_attribute("id"),
                "name": el.get_attribute("name"),
            })
        except Exception:
            continue
    return elementos

def validar_textarea_preenchida(elemento, texto_esperado):
    try:
        atual = elemento.get_attribute("value")
        if atual is None or atual == "":
            atual = (elemento.text or "")
        return (atual or "").strip() == texto_esperado.strip()
    except StaleElementReferenceException:
        return False

def _prepare_focus_and_clear(elemento, limpar_primeiro=True):
    driver.execute_script("arguments[0].scrollIntoView({block:'center'});", elemento)
    try:
        elemento.click()
    except Exception:
        driver.execute_script("arguments[0].focus();", elemento)
    if limpar_primeiro:
        try:
            elemento.clear()
        except Exception:
            ActionChains(driver)\
                .move_to_element(elemento).click()\
                .key_down(Keys.CONTROL).send_keys("a").key_up(Keys.CONTROL)\
                .send_keys(Keys.DELETE).perform()

def _textarea_tradicional(elemento, texto, limpar_primeiro=True):
    _prepare_focus_and_clear(elemento, limpar_primeiro)
    elemento.send_keys(texto)
    elemento.send_keys(Keys.TAB)

def _textarea_actionchains(elemento, texto, limpar_primeiro=True):
    _prepare_focus_and_clear(elemento, limpar_primeiro)
    ac = ActionChains(driver)
    ac.move_to_element(elemento).click().perform()
    for chunk_start in range(0, len(texto), 400):
        ac.send_keys(texto[chunk_start:chunk_start+400]).perform()
        time.sleep(0.05)
    ac.send_keys(Keys.TAB).perform()

def _textarea_js_setvalue(elemento, texto):
    driver.execute_script("""
    const el = arguments[0];
    const val = arguments[1];
    el.value = val;
    el.dispatchEvent(new Event('input', {bubbles:true}));
    el.dispatchEvent(new Event('change', {bubbles:true}));
    el.dispatchEvent(new KeyboardEvent('keyup', {bubbles:true}));
    el.dispatchEvent(new Event('blur', {bubbles:true}));
    """, elemento, texto)

def validar_registros_encontrados():
    global driver
    try:
        log(doc, "🔍 Validando registros encontrados...")
        time.sleep(5)
        seletores_tabela = [
            '#DataTables_Table_0 tbody tr',
            '#DataTables_Table_1 tbody tr',
            'table[id*="DataTables"] tbody tr',
            '.wdGrid table tbody tr',
            'table tbody tr'
        ]
        for seletor in seletores_tabela:
            try:
                linhas = driver.find_elements(By.CSS_SELECTOR, seletor)
                linhas_validas = [
                    linha for linha in linhas 
                    if linha.is_displayed() and 
                    linha.text.strip() and 
                    len(linha.text.strip()) > 5 and
                    not any(termo in linha.text.lower() for termo in [
                        'nenhum registro', 'sem dados', 'no data', 'vazio', 'empty'
                    ])
                ]
                if linhas_validas:
                    quantidade = len(linhas_validas)
                    log(doc, f"✅ {quantidade} registro(s) encontrado(s)")
                    for i, linha in enumerate(linhas_validas[:3], 1):
                        texto = linha.text.strip()[:100]
                        log(doc, f"   Registro {i}: {texto}...")
                    return {
                        'encontrou_registros': True,
                        'quantidade_registros': quantidade,
                        'detalhes': [linha.text.strip()[:100] for linha in linhas_validas[:5]]
                    }
            except Exception as e:
                log(doc, f"⚠️ Erro ao verificar tabela com {seletor}: {e}")
                continue
        mensagens_vazio = [
            "Nenhum registro encontrado",
            "Não foram encontrados registros",
            "Nenhum resultado",
            "0 registros encontrados"
        ]
        for mensagem in mensagens_vazio:
            try:
                elem = driver.find_element(By.XPATH, f"//*[contains(text(), '{mensagem}')]")
                if elem.is_displayed():
                    log(doc, f"ℹ️ Sistema informou: {elem.text.strip()}")
                    return {
                        'encontrou_registros': False,
                        'quantidade_registros': 0,
                        'mensagem': f"Sistema informou: {elem.text.strip()}"
                    }
            except:
                continue
        log(doc, "⚠️ Não foi possível determinar se há registros")
        return {
            'encontrou_registros': True,
            'quantidade_registros': 1,
            'mensagem': "Validação inconclusiva, continuando teste"
        }
    except Exception as e:
        log(doc, f"❌ Erro na validação: {e}")
        return {
            'encontrou_registros': True,
            'quantidade_registros': 1,
            'mensagem': f"Erro na validação, continuando: {e}"
        }

def preencher_textarea_por_indice(indice_campo, texto, max_tentativas=5, limpar_primeiro=True):
    def acao():
        if not isinstance(indice_campo, int) or indice_campo < 0:
            raise ValueError(f"Índice inválido: {indice_campo}")
        if texto is None or not isinstance(texto, str):
            raise ValueError(f"Texto inválido: {texto!r}")
        tentativa = 0
        while tentativa < max_tentativas:
            tentativa += 1
            try:
                campos = encontrar_campos_textarea()
                if not campos:
                    if tentativa < max_tentativas:
                        log(doc, f"⚠️ Nenhuma textarea encontrada (tentativa {tentativa}/{max_tentativas})")
                        time.sleep(1.5)
                        continue
                    raise Exception("Nenhuma textarea foi encontrada na página.")
                if indice_campo >= len(campos):
                    raise Exception(f"Índice {indice_campo} inválido. Encontradas {len(campos)} textareas.")
                campo_info = campos[indice_campo]
                elemento = campo_info["elemento"]
                if validar_textarea_preenchida(elemento, texto):
                    log(doc, f"✅ Textarea {indice_campo} já está com o valor desejado.")
                    return True
                estrategias = [
                    lambda: _textarea_tradicional(elemento, texto, limpar_primeiro),
                    lambda: _textarea_actionchains(elemento, texto, limpar_primeiro),
                    lambda: _textarea_js_setvalue(elemento, texto),
                ]
                for i, estrategia in enumerate(estrategias, 1):
                    try:
                        log(doc, f" ▶️ Estratégia {i}…")
                        estrategia()
                        time.sleep(0.8)
                        if validar_textarea_preenchida(elemento, texto):
                            val = (elemento.get_attribute("value") or "").strip()
                            log(doc, f"✅ Preenchido com sucesso pela estratégia {i}: '{val[:60]}{'…' if len(val) > 60 else ''}'")
                            return True
                        else:
                            log(doc, f"⚠️ Estratégia {i} não refletiu o valor esperado.")
                    except Exception as e:
                        log(doc, f"⚠️ Estratégia {i} falhou: {e}")
                        continue
                if tentativa < max_tentativas:
                    log(doc, f"⚠️ Tentativa {tentativa} falhou; reintentando em 1.5s…")
                    time.sleep(1.5)
                    continue
            except Exception as e:
                if tentativa < max_tentativas:
                    log(doc, f"⚠️ Erro na tentativa {tentativa}: {e}. Retentando…")
                    time.sleep(1.5)
                    continue
                else:
                    raise
        raise Exception(f"Falha ao preencher textarea {indice_campo} após {max_tentativas} tentativas.")
    return acao

from selenium.webdriver.common.by import By
from selenium.webdriver.common.action_chains import ActionChains
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.support.ui import WebDriverWait

def clicar_link_btsave_por_indice(indice, max_tentativas=5):
    """
    Retorna uma ação callable para ser usada em safe_action.
    Clica no elemento <a class='btModel btGray btsave'> pelo índice informado.
    """
    def acao():
        if not isinstance(indice, int) or indice < 0:
            raise ValueError(f"Índice inválido: {indice}")
        tentativa = 0
        while tentativa < max_tentativas:
            tentativa += 1
            try:
                botoes = driver.find_elements(By.CSS_SELECTOR, "a.btModel.btGray.btsave")
                if not botoes:
                    if tentativa < max_tentativas:
                        log(doc, f"⚠️ Nenhum botão 'Salvar' encontrado (tentativa {tentativa}/{max_tentativas})")
                        time.sleep(1.5)
                        continue
                    raise Exception("Nenhum botão 'Salvar' encontrado na página.")

                if indice >= len(botoes):
                    raise Exception(f"Índice {indice} inválido. Encontrados {len(botoes)} elementos 'Salvar'.")

                elemento = botoes[indice]

                estrategias = [
                    lambda: (WebDriverWait(driver, 5).until(EC.element_to_be_clickable((By.CSS_SELECTOR, "a.btModel.btGray.btsave"))), elemento.click()),
                    lambda: (driver.execute_script("arguments[0].scrollIntoView({block:'center'});", elemento), elemento.click()),
                    lambda: driver.execute_script("arguments[0].click();", elemento),
                    lambda: ActionChains(driver).move_to_element(elemento).click().perform(),
                    lambda: driver.execute_script("arguments[0].dispatchEvent(new MouseEvent('click', {view: window, bubbles: true, cancelable: true}));", elemento),
                ]

                for i, estrategia in enumerate(estrategias, 1):
                    try:
                        log(doc, f"▶️ Estratégia {i} para clicar no botão índice {indice}…")
                        estrategia()
                        time.sleep(0.4)
                        log(doc, f"✅ Clique no botão índice {indice} realizado com sucesso.")
                        return True
                    except Exception as e:
                        log(doc, f"⚠️ Estratégia {i} falhou: {e}")
                        continue

                if tentativa < max_tentativas:
                    log(doc, f"⚠️ Tentativa {tentativa} falhou; reintentando em 1.5s…")
                    time.sleep(1.5)
                    continue
                else:
                    raise Exception("Nenhuma estratégia de clique funcionou.")

            except Exception as e:
                if tentativa < max_tentativas:
                    log(doc, f"⚠️ Erro na tentativa {tentativa}: {e}. Retentando…")
                    time.sleep(1.5)
                    continue
                else:
                    raise
        raise Exception(f"Falha ao clicar no botão índice {indice} após {max_tentativas} tentativas.")
    return acao


def inicializar_driver():
    global driver, wait
    try:
        options = Options()
        options.add_argument("--start-maximized")
        options.add_argument("--disable-blink-features=AutomationControlled")
        options.add_experimental_option("excludeSwitches", ["enable-automation"])
        options.add_experimental_option('useAutomationExtension', False)
        driver = webdriver.Chrome(service=Service(ChromeDriverManager().install()), options=options)
        driver.execute_script("Object.defineProperty(navigator, 'webdriver', {get: () => undefined})")
        wait = WebDriverWait(driver, 20)
        return True
    except Exception as e:
        log(doc, f"❌ Erro ao inicializar driver: {e}")
        return False

def executar_teste():
    global driver, wait, doc
    try:
        if not inicializar_driver():
            return False

        safe_action(doc, "Acessando sistema", lambda: driver.get("http://localhost:8080/gs/index.xhtml"))

        # Função para login
        def realizar_login():
            wait.until(EC.presence_of_element_located((By.ID, "j_id15:email"))).send_keys("joaoeduardo.gold@outlook.com")
            wait.until(EC.presence_of_element_located((By.ID, "j_id15:senha"))).send_keys("071999gs", Keys.ENTER)
            time.sleep(5)

        safe_action(doc, "Realizando login", realizar_login)

        def ajustar_zoom():
            driver.execute_script("document.body.style.zoom='90%'")
            log(doc, "🔍 Zoom ajustado para 90%.")

        safe_action(doc, "Ajustando zoom e abrindo menu", lambda: (ajustar_zoom(), driver.find_element(By.TAG_NAME, "body").send_keys(Keys.F3)))

        safe_action(doc, "Acessando Gestor de Compras", lambda: wait.until(EC.element_to_be_clickable((By.XPATH, '/html/body/div[15]/ul/li[3]/img'))).click())
        safe_action(doc, "Clicando em Consulta de PIMS", lambda: wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, '#gsCompras > div.wdTelas > div.telaInicial.clearfix.overflow.overflowY > ul > li:nth-child(2) > a > span'))).click())

        time.sleep(5)

        safe_action(doc, "Preenchendo Número do PIMS", lambda: preencher_campo_xpath_com_retry(driver, wait, "//*[@id='gsCompras']/div[2]/div[2]/div[1]/div/div[1]/input", "200848"))
        safe_action(doc, "Preenchendo Data inicial", lambda: preencher_campo_xpath_com_retry(driver, wait, "//input[@type='text' and @class='hasDatepicker dataI' and @maxlength='10' and @style='width: 100px;' and @grupo='']", "06/10/2025"))
        safe_action(doc, "Selecionando Centro de Custo", lambda: Select(wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, "#gsCompras > div.wdTelas > div.consultarPims.telaConsulta > div:nth-child(1) > div > div:nth-child(4) > select")))).select_by_visible_text("CENTRO DE CUSTO JOÃO"))
        safe_action(doc, "Selecionando Departamento", lambda: Select(wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, "#gsCompras > div.wdTelas > div.consultarPims.telaConsulta > div:nth-child(1) > div > div:nth-child(5) > select")))).select_by_visible_text("Teste"))
        safe_action(doc, "Selecionando Status", lambda: Select(wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, "#gsCompras > div.wdTelas > div.consultarPims.telaConsulta > div:nth-child(1) > div > div:nth-child(6) > select")))).select_by_visible_text("Solicitado"))


        safe_action(doc, "Realizando consulta", lambda:
            wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, '#gsCompras > div.wdTelas > div.consultarPims.telaConsulta > div:nth-child(1) > div > div:nth-child(7) > a'))).click()
        )

        log(doc, "🔄 Validando resultados da consulta...")
        resultado_validacao = validar_registros_encontrados()

        if not resultado_validacao or not resultado_validacao.get('encontrou_registros', False):
            log(doc, "Nenhum registro encontrado - Finalizando consulta")
            safe_action(doc, "Fechando tela", lambda:
                clicar_elemento_robusto(driver, wait, "#gsApoioOrtopedico > div.wdTop.ui-draggable-handle > div.wdClose > a")
            )
            return True

        quantidade = resultado_validacao.get('quantidade_registros', 0)
        log(doc, f"Processando {quantidade} registro(s) encontrado(s)")

        safe_action(doc, "Realizando Download do PIMS", lambda: (
            clicar_elemento_robusto(driver, wait, "span.sprites.sp-download"),
            time.sleep(5),
            pyautogui.press('enter'),
        ))

        safe_action(doc, "Alterando PIMS", lambda:
            clicar_elemento_robusto(driver, wait, "span[title='Alterar PIMS']")
        )
        time.sleep(2)

        # === ITEM 1: selecionar, preencher e salvar ===
        safe_action(doc, "Selecionando item 1", lambda: selecionar_linha_item_por_indice(1))

        safe_action(doc, "Preenchendo quantidade do item 1", lambda:
            driver.execute_script("""
                const el = document.querySelector("#gsCompras div.consultarPims li:nth-child(1) input.qtdProd");
                if (!el) throw new Error("Input do item 1 não encontrado");
                el.focus();
                el.value = "1000";
                // Evite input/change se o front replica p/ todos; deixe o submit ler o value
                el.blur();
            """)
        )
        
        time.sleep(2)


        safe_action(doc, "Desviando foco do campo", lambda:
            clicar_elemento_robusto(driver, wait, "//*[@id='gsCompras']/div[2]/div[2]/div[1]/div/div[1]/input")
        )
        time.sleep(2)


        safe_action(doc, "Salvar alterações do item 1", clicar_link_btsave_por_indice(0))

        # Reabrir modo de edição caso a tela volte para visualização
        safe_action(doc, "Reabrir Alterar PIMS", lambda: clicar_elemento_robusto(driver, wait, "span[title='Alterar PIMS']"))
        time.sleep(0.8)

        # === ITEM 2: selecionar, preencher e salvar ===
        safe_action(doc, "Selecionando item 2", lambda: selecionar_linha_item_por_indice(2))

        safe_action(doc, "Preenchendo quantidade do item 2", lambda:
            driver.execute_script("""
                const el = document.querySelector("#gsCompras div.consultarPims li:nth-child(2) input.qtdProd");
                if (!el) throw new Error("Input do item 2 não encontrado");
                el.focus();
                el.value = "1000";
                el.blur();
            """)
        )

        safe_action(doc, "Desviando foco do campo", lambda:
            clicar_elemento_robusto(driver, wait, "//*[@id='gsCompras']/div[2]/div[2]/div[1]/div/div[1]/input")
        )
        time.sleep(2)

        safe_action(doc, "Salvar alterações do item 2", clicar_link_btsave_por_indice(0))

        # (Opcional) conferir que 1º e 2º estão corretos
        safe_action(doc, "Conferindo valores finais", lambda:
            log(doc, "Valores: " + str(driver.execute_script("""
                const arr = Array.from(document.querySelectorAll("input.qtdProd")).map(e => e.value);
                return arr;
            """)))
        )


        log(doc, "🔍 Verificando mensagens de alerta...")
        # Coloque uma função para verificar mensagens


        safe_action(doc, "Alterando PIMS", lambda:
            clicar_elemento_robusto(driver, wait, "span[title='Alterar PIMS']")
        )
        time.sleep(2)

        # Mensagem no serviço
        safe_action(doc, "Adicionando mensagem no Serviço", lambda: clicar_elemento_robusto(driver, wait, "#gsCompras div.consultarPims li:nth-child(1) ul li:nth-child(5) a:nth-child(1)"))
        time.sleep(2)
        safe_action(doc, "Preenchendo Mensagem no Serviço", preencher_textarea_por_indice(0, "TESTE MENSAGEM SELENIUM AUTOMATIZADO (Automação de Testes): Teste de Mensagem longa para validar o campo textarea no sistema.")())
        safe_action(doc, "Salvando mensagem no Serviço", clicar_link_btsave_por_indice(1))

        time.sleep(2)


        # Mensagem no produto
        safe_action(doc, "Adicionando mensagem no Produto", lambda:
            driver.execute_script(
                "arguments[0].click();",
                driver.find_element(By.CSS_SELECTOR, "#gsCompras div.consultarPims li:nth-child(2) ul li:nth-child(5) a:nth-child(1)")
            )
        )
        time.sleep(2)
        safe_action(doc, "Preenchendo Mensagem no Produto", preencher_textarea_por_indice(0, "TESTE MENSAGEM SELENIUM AUTOMATIZADO (Automação de Testes): Teste de Mensagem longa para validar o campo textarea no sistema.")())
        safe_action(doc, "Salvando mensagem no Produto", clicar_link_btsave_por_indice(1))





        # Fechar modal Gestor de Compras
        safe_action(doc, "Fechando modal Gestor de Compras", lambda:
            driver.execute_script("arguments[0].click();", driver.find_element(By.CSS_SELECTOR, "#gsCompras > div.wdTop.ui-draggable-handle > div > a"))
        )
        time.sleep(2)

        log(doc, "🔍 Verificando mensagens de alerta...")
        # Coloque uma função para verificar mensagens

        return True
    except Exception as e:
        log(doc, f"❌ ERRO FATAL: {e}")
        take_screenshot(driver, doc, "erro_fatal")
        return False
    finally:
        log(doc, "✅ Teste concluído.")

def finalizar_relatorio():
    global driver, doc
    nome_arquivo = f"relatorio_consulta_de_PIMS_cenario_1_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    try:
        doc.save(nome_arquivo)
        log(doc, f"📄 Relatório salvo como: {nome_arquivo}")
    except Exception as e:
        print(f"Erro ao salvar relatório: {e}")
    if driver:
        try:
            driver.quit()
        except:
            pass

def main():
    global doc
    try:
        log(doc, "🚀 Iniciando teste de Consulta de PIMS")
        sucesso = executar_teste()
        if sucesso:
            log(doc, "✅ Teste executado com sucesso!")
        else:
            log(doc, "❌ Teste finalizado com erros.")
    except Exception as e:
        log(doc, f"❌ Erro na execução principal: {e}")
    finally:
        finalizar_relatorio()

if __name__ == "__main__":
    main()
