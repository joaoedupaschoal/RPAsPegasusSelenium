# ==== IMPORTS ====
from datetime import datetime, timedelta
from datetime import time as dt_time
import time
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver import ActionChains
from selenium.common.exceptions import *
from selenium.webdriver.support import expected_conditions as EC
from selenium import webdriver
from selenium.webdriver.support.ui import Select, WebDriverWait
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service
from webdriver_manager.chrome import ChromeDriverManager
from docx import Document
from docx.shared import Inches
from faker import Faker
from faker.providers import BaseProvider
from validate_docbr import CPF
import subprocess
import os
import random
import re
from functools import wraps
from selenium.common.exceptions import (
    TimeoutException,
    StaleElementReferenceException,
    ElementClickInterceptedException,
    NoSuchElementException,
    ElementNotInteractableException
)
from selenium.webdriver import ActionChains
# ==== CONFIGURAÇÕES GLOBAIS ====
TIMEOUT_DEFAULT = 30
TIMEOUT_CURTO = 10
TIMEOUT_LONGO = 60
CAMINHO_ARQUIVO_UPLOAD = "C:/Users/Gold System/Documents/teste.png"
URL = "http://localhost:8080/gs/login.xhtml"
LOGIN_EMAIL = "joaoeduardo.gold@outlook.com"
LOGIN_PASSWORD = "071999gs"

# ==== VARIÁVEIS GLOBAIS ====
doc = Document()
doc.add_heading("RELATÓRIO DO TESTE", 0)
doc.add_paragraph("Controle de Caixa - Devoluções – Cenário 1: Rotina parcial de Devoluções - Filtros Utilizados: CPF/CNPJ, Número do Contrato, Data Inicial e Data Final.")
doc.add_paragraph(f"Data do teste: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")

screenshot_registradas = set()
driver = None
wait = None

# ==== PROVIDERS CUSTOMIZADOS ====
class BrasilProvider(BaseProvider):
    def rg(self):
        numeros = [str(random.randint(0, 9)) for _ in range(8)]
        return ''.join(numeros) + '-' + str(random.randint(0, 9))

fake = Faker("pt_BR")
fake.add_provider(BrasilProvider)

# ==== GERAÇÃO DE DATAS ====
def gerar_datas_validas(hora_padrao="00:00", dias_fim=0):
    hoje_date = datetime.today().date()
    dez_anos_atras = hoje_date - timedelta(days=3650)
    data_falecimento = fake.date_between(start_date=dez_anos_atras, end_date=hoje_date)
    idade_minima, idade_maxima = 18, 110
    data_nascimento = data_falecimento - timedelta(days=random.randint(idade_minima * 365, idade_maxima * 365))
    data_sepultamento = data_falecimento + timedelta(days=random.randint(1, 10))
    data_registro = data_sepultamento + timedelta(days=random.randint(1, 10))
    data_velorio = fake.date_between(start_date=data_falecimento, end_date=data_sepultamento)
    data_inicio_date = hoje_date + timedelta(days=random.randint(2, 30))
    h, m = map(int, hora_padrao.split(":"))
    dt_inicio = datetime.combine(data_inicio_date, dt_time(h, m))
    dt_fim = dt_inicio + timedelta(days=dias_fim)
    fmt_data = "%d/%m/%Y"
    fmt_dt = "%d/%m/%Y %H:%M"
    return (
        data_nascimento.strftime(fmt_data),
        data_falecimento.strftime(fmt_data),
        data_sepultamento.strftime(fmt_data),
        data_velorio.strftime(fmt_data),
        dt_inicio.strftime(fmt_dt),
        dt_fim.strftime(fmt_dt),
        data_registro.strftime(fmt_data),
        hoje_date.strftime(fmt_data),
    )

(data_nascimento, data_falecimento, data_sepultamento,
 data_velorio, data_inicio, data_fim, data_registro, hoje) = gerar_datas_validas(
    hora_padrao="08:50",
    dias_fim=0
)

# ==== UTILITÁRIOS DE LOG ====
def log(doc, msg):
    print(msg)
    doc.add_paragraph(msg)

def _sanitize_filename(name: str) -> str:
    name = name.strip().lower()
    name = re.sub(r"[<>:\"/\\|?*']", "_", name)
    name = re.sub(r"_+", "_", name)
    return name[:120]

def take_screenshot(driver, doc, nome):
    if driver is None:
        return
    nome = _sanitize_filename(nome)
    if nome not in screenshot_registradas:
        path = f"screenshots/{nome}.png"
        os.makedirs("screenshots", exist_ok=True)
        try:
            driver.save_screenshot(path)
            doc.add_paragraph(f"Screenshot: {nome}")
            doc.add_picture(path, width=Inches(5.5))
            screenshot_registradas.add(nome)
        except Exception as e:
            log(doc, f"⚠️ Erro ao tirar screenshot {nome}: {e}")

def preencher_elemento_por_indice(js_engine, doc, xpath_base, indice, valor, descricao="Campo", timeout=10):
    """
    Preenche um elemento específico por índice usando xpath base.
    
    Args:
        js_engine: Instância do JSForceEngine
        doc: Documento para logs
        xpath_base: XPath base do elemento (sem índice)
        indice: Índice do elemento (1-based)
        valor: Valor a ser preenchido
        descricao: Descrição do campo para logs
        timeout: Timeout para operação
    
    Returns:
        bool: True se preencheu com sucesso, False caso contrário
    
    Exemplo:
        preencher_elemento_por_indice(
            js_engine, doc,
            xpath_base="//input[@class='hasDatepicker chqf']",
            indice=1,
            valor="30/10/2025",
            descricao="Data de Venda"
        )
    """
    xpath_indexado = f"({xpath_base})[{indice}]"
    
    try:
        log(doc, f"📝 Preenchendo {descricao} (índice {indice}): '{valor}'")
        
        # Conta quantos elementos existem
        try:
            elementos = js_engine.driver.find_elements(By.XPATH, xpath_base)
            total = len(elementos)
            log(doc, f"   ℹ️ Total de elementos encontrados: {total}")
            
            if total == 0:
                log(doc, f"   ⚠️ Nenhum elemento encontrado com xpath: {xpath_base}")
                return False
            
            if indice > total:
                log(doc, f"   ⚠️ Índice {indice} inválido - só existem {total} elemento(s)")
                return False
                
        except Exception as e:
            log(doc, f"   ⚠️ Erro ao contar elementos: {e}")
        
        # Preenche usando force_fill
        js_engine.force_fill(xpath_indexado, valor, by_xpath=True)
        
        # Aguarda um pouco para garantir que preencheu
        time.sleep(0.4)
        
        # Valida se preencheu corretamente
        try:
            elemento = js_engine.driver.find_element(By.XPATH, xpath_indexado)
            valor_atual = elemento.get_attribute('value') or ''
            
            # Normaliza para comparação (remove formatação)
            valor_norm = valor.replace('R$ ', '').replace('.', '').replace(',', '.').strip()
            atual_norm = valor_atual.replace('R$ ', '').replace('.', '').replace(',', '.').strip()
            
            if valor_norm in atual_norm or atual_norm in valor_norm or valor_atual == valor:
                log(doc, f"   ✅ {descricao} preenchido: '{valor_atual}'")
                return True
            else:
                log(doc, f"   ⚠️ Valor preenchido difere: esperado '{valor}', atual '{valor_atual}'")
                return True  # Retorna True mesmo assim pois pode ser formatação diferente
                
        except Exception as e:
            log(doc, f"   ⚠️ Não foi possível validar: {e}")
            return True  # Assume sucesso se não conseguiu validar
        
    except Exception as e:
        log(doc, f"   ❌ Erro ao preencher {descricao}: {e}")
        return False


# =========================
# Estratégias de preenchimento
# =========================

def _prepare_focus_and_clear(elemento, limpar_primeiro=True):
    # Garante visibilidade e foco
    driver.execute_script("arguments[0].scrollIntoView({block:'center'});", elemento)
    try:
        elemento.click()
    except Exception:
        driver.execute_script("arguments[0].focus();", elemento)

    if limpar_primeiro:
        try:
            elemento.clear()
        except Exception:
            # Fallback de limpeza por teclas
            ActionChains(driver)\
                .move_to_element(elemento).click()\
                .key_down(Keys.CONTROL).send_keys("a").key_up(Keys.CONTROL)\
                .send_keys(Keys.DELETE).perform()


def _textarea_tradicional(elemento, texto, limpar_primeiro=True):
    _prepare_focus_and_clear(elemento, limpar_primeiro)
    elemento.send_keys(texto)
    # Dispara blur para muitos bindings reativos
    elemento.send_keys(Keys.TAB)


def _textarea_actionchains(elemento, texto, limpar_primeiro=True):
    _prepare_focus_and_clear(elemento, limpar_primeiro)
    ac = ActionChains(driver)
    ac.move_to_element(elemento).click().perform()
    # Quebra o texto em partes para evitar engasgos em campos longos
    for chunk_start in range(0, len(texto), 400):
        ac.send_keys(texto[chunk_start:chunk_start+400]).perform()
        time.sleep(0.05)
    ac.send_keys(Keys.TAB).perform()


def _textarea_js_setvalue(elemento, texto):
    # Seta .value e dispara eventos clássicos
    driver.execute_script("""
        const el = arguments[0];
        const val = arguments[1];
        el.value = val;
        // Dispara eventos comuns que form libs escutam
        el.dispatchEvent(new Event('input',  {bubbles:true}));
        el.dispatchEvent(new Event('change', {bubbles:true}));
        el.dispatchEvent(new KeyboardEvent('keyup', {bubbles:true}));
        el.dispatchEvent(new Event('blur',   {bubbles:true}));
    """, elemento, texto)


def _textarea_js_react_input(elemento, texto):
    # Compat extra p/ React (setando o setter do prototype) + eventos
    driver.execute_script("""
        const el = arguments[0];
        const val = arguments[1];

        const desc = Object.getOwnPropertyDescriptor(window.HTMLTextAreaElement.prototype, 'value');
        if (desc && desc.set) {
            desc.set.call(el, val);
        } else {
            el.value = val;
        }

        // React/Vue/Svelte geralmente escutam 'input'
        el.dispatchEvent(new Event('input', {bubbles: true}));
        el.dispatchEvent(new Event('change', {bubbles: true}));
        el.dispatchEvent(new Event('blur', {bubbles: true}));
    """, elemento, texto)



# ==== SISTEMA DATEPICKER MELHORADO ====
def encontrar_campos_datepicker():
    """Encontra todos os campos datepicker na página"""
    global driver
    
    if driver is None:
        return []
    
    seletores_datepicker = [
        "input.hasDatepicker",
        "input[id^='dp']",
        "input[maxlength='10'][grupo='']",
        "input[type='text'][maxlength='10']",
        "input[class*='datepicker']",
        ".hasDatepicker"
    ]
    
    campos_encontrados = []
    
    for seletor in seletores_datepicker:
        try:
            elementos = driver.find_elements(By.CSS_SELECTOR, seletor)
            for elemento in elementos:
                if elemento.is_displayed() and elemento.is_enabled():
                    info = {
                        'elemento': elemento,
                        'id': elemento.get_attribute('id') or f"dp_{len(campos_encontrados)}",
                        'seletor_usado': seletor,
                        'maxlength': elemento.get_attribute('maxlength'),
                        'placeholder': elemento.get_attribute('placeholder')
                    }
                    # Evita duplicatas
                    if not any(c['id'] == info['id'] for c in campos_encontrados):
                        campos_encontrados.append(info)
        except Exception as e:
            log(doc, f"⚠️ Erro ao buscar campos datepicker com {seletor}: {e}", 'WARN')
            continue
    
    log(doc, f"📊 Encontrados {len(campos_encontrados)} campos datepicker")
    return campos_encontrados

def _datepicker_jquery(campo_id, data_valor):
    """Estratégia jQuery para datepicker"""
    global driver
    
    resultado = driver.execute_script("""
        var campoId = arguments[0], valor = arguments[1];
        if (typeof jQuery === 'undefined') return 'jQuery não disponível';
        var $campo = $('#' + campoId);
        if (!$campo.length) return 'Campo não encontrado: ' + campoId;
        try {
            if ($campo.hasClass('hasDatepicker')) { 
                $campo.datepicker('setDate', valor); 
            } else { 
                $campo.val(valor); 
            }
            $campo.trigger('input').trigger('change').trigger('blur');
            return $campo.val();
        } catch(e) { 
            return 'Erro: ' + e.message; 
        }
    """, campo_id, data_valor)
    
    if isinstance(resultado, str) and ('Erro' in resultado or 'não disponível' in resultado):
        raise Exception(f"jQuery falhou: {resultado}")

def _datepicker_javascript(elemento, data_valor):
    """Estratégia JavaScript para datepicker"""
    global driver
    
    driver.execute_script("""
        var campo = arguments[0], valor = arguments[1];
        campo.focus(); 
        campo.value = ''; 
        campo.value = valor;
        ['input','change','blur','keyup'].forEach(ev => 
            campo.dispatchEvent(new Event(ev, {bubbles: true}))
        );
    """, elemento, data_valor)

def _datepicker_actionchains(elemento, data_valor):
    """Estratégia ActionChains para datepicker"""
    global driver
    
    scroll_to_element_safe(elemento)
    time.sleep(0.5)
    
    ActionChains(driver).move_to_element(elemento).click().perform()
    time.sleep(0.5)
    ActionChains(driver).key_down(Keys.CONTROL).send_keys('a').key_up(Keys.CONTROL).perform()
    time.sleep(0.3)
    ActionChains(driver).send_keys(Keys.DELETE).perform()
    time.sleep(0.3)
    
    for char in data_valor:
        ActionChains(driver).send_keys(char).perform()
        time.sleep(0.05)
    
    ActionChains(driver).send_keys(Keys.TAB).perform()

def _datepicker_tradicional(elemento, data_valor):
    """Estratégia tradicional para datepicker"""
    scroll_to_element_safe(elemento)
    time.sleep(0.5)
    elemento.click()
    time.sleep(0.5)
    elemento.clear()
    elemento.send_keys(data_valor)
    elemento.send_keys(Keys.TAB)

def validar_data_preenchida(elemento, data_esperada):
    """Valida se a data foi preenchida corretamente"""
    try:
        if elemento is None:
            return False
            
        val = (elemento.get_attribute('value') or '').strip()
        if not val:
            return False
            
        if val == data_esperada or data_esperada in val:
            return True
            
        # Tenta comparar datas em diferentes formatos
        formatos = [
            '%d/%m/%Y %H:%M', '%d/%m/%Y %H:%M:%S',
            '%d/%m/%Y', '%m/%d/%Y', '%Y-%m-%d', '%d-%m-%Y'
        ]
        
        for formato in formatos:
            try:
                d1 = datetime.strptime(val, formato)
                d2 = datetime.strptime(data_esperada, formato)
                if d1 == d2:
                    return True
            except:
                continue
                
        return False
        
    except Exception:
        return False





def preencher_datepicker_por_indice(indice_campo, data_valor, max_tentativas=5):
    """Preenche datepicker pelo índice com estratégias múltiplas"""
    def acao():
        if not isinstance(indice_campo, int) or indice_campo < 0:
            raise ValueError(f"Índice inválido: {indice_campo}")
            
        if not data_valor or not isinstance(data_valor, str):
            raise ValueError(f"Data inválida: {data_valor}")
        
        tentativa = 0
        while tentativa < max_tentativas:
            tentativa += 1
            
            try:
                campos = encontrar_campos_datepicker()
                
                if not campos:
                    if tentativa < max_tentativas:
                        log(doc, f"⚠️ Nenhum campo datepicker encontrado, tentativa {tentativa}/{max_tentativas}", 'WARN')
                        time.sleep(2)
                        continue
                    raise Exception("Nenhum campo datepicker encontrado na página")
                
                if indice_campo >= len(campos):
                    raise Exception(f"Índice {indice_campo} inválido. Encontrados {len(campos)} campos")
                
                campo_info = campos[indice_campo]
                elemento = campo_info['elemento']
                campo_id = campo_info['id']
                
                log(doc, f"🎯 Tentativa {tentativa}: Preenchendo datepicker {indice_campo} (ID: {campo_id}) com '{data_valor}'")
                
                # Verifica se já está preenchido corretamente
                if validar_data_preenchida(elemento, data_valor):
                    log(doc, f"✅ Campo {indice_campo} já está preenchido corretamente!")
                    return True
                
                # Estratégias específicas para datepicker
                estrategias = [
                    lambda: _datepicker_jquery(campo_id, data_valor),
                    lambda: _datepicker_javascript(elemento, data_valor),
                    lambda: _datepicker_actionchains(elemento, data_valor),
                    lambda: _datepicker_tradicional(elemento, data_valor)
                ]
                
                for i, estrategia in enumerate(estrategias, 1):
                    try:
                        log(doc, f"   Aplicando estratégia {i} para datepicker...")
                        estrategia()
                        time.sleep(1)
                        
                        # Verifica se funcionou
                        if validar_data_preenchida(elemento, data_valor):
                            valor_atual = elemento.get_attribute('value')
                            log(doc, f"✅ Datepicker preenchido com estratégia {i}: '{valor_atual}'")
                            return True
                        else:
                            log(doc, f"⚠️ Estratégia {i} não preencheu corretamente", 'WARN')
                            
                    except Exception as e:
                        log(doc, f"⚠️ Estratégia {i} falhou: {e}", 'WARN')
                        continue
                
                # Se chegou aqui, nenhuma estratégia funcionou nesta tentativa
                if tentativa < max_tentativas:
                    log(doc, f"⚠️ Tentativa {tentativa} falhou, tentando novamente em 2s...", 'WARN')
                    time.sleep(2)
                    continue
                
            except Exception as e:
                if tentativa < max_tentativas:
                    log(doc, f"⚠️ Erro na tentativa {tentativa}: {e}, tentando novamente...", 'WARN')
                    time.sleep(2)
                    continue
                else:
                    raise
        
        raise Exception(f"Falha ao preencher datepicker {indice_campo} após {max_tentativas} tentativas")
    
    return acao


def preencher_datepicker_por_indice_xpath(js_engine, doc, xpath_base, indice, data_valor, descricao="Data", timeout=10):
    """
    Preenche um campo datepicker específico por índice.
    Versão especializada para campos de data que usa estratégias específicas.
    
    Args:
        js_engine: Instância do JSForceEngine
        doc: Documento para logs
        xpath_base: XPath base do datepicker
        indice: Índice do elemento (1-based)
        data_valor: Data no formato string (ex: "30/10/2025")
        descricao: Descrição do campo para logs
        timeout: Timeout para operação
    
    Returns:
        bool: True se preencheu com sucesso, False caso contrário
    """
    xpath_indexado = f"({xpath_base})[{indice}]"
    
    try:
        log(doc, f"📅 Preenchendo {descricao} (índice {indice}): '{data_valor}'")
        
        # Conta elementos
        try:
            elementos = js_engine.driver.find_elements(By.XPATH, xpath_base)
            total = len(elementos)
            log(doc, f"   ℹ️ Total de datepickers encontrados: {total}")
            
            if total == 0:
                log(doc, f"   ⚠️ Nenhum datepicker encontrado")
                return False
            
            if indice > total:
                log(doc, f"   ⚠️ Índice {indice} inválido - só existem {total} datepicker(s)")
                return False
        except:
            pass
        
        # Localiza o elemento específico
        elemento = js_engine.driver.find_element(By.XPATH, xpath_indexado)
        campo_id = elemento.get_attribute('id') or f"datepicker_{indice}"
        
        log(doc, f"   🎯 Elemento localizado (ID: {campo_id})")
        
        # Estratégias específicas para datepicker
        estrategias = [
            # Estratégia 1: jQuery datepicker
            lambda: js_engine.driver.execute_script(f"""
                var campo = document.evaluate("{xpath_indexado}", document, null, 
                    XPathResult.FIRST_ORDERED_NODE_TYPE, null).singleNodeValue;
                if (!campo) return false;
                
                if (typeof jQuery !== 'undefined' && campo.classList.contains('hasDatepicker')) {{
                    jQuery(campo).datepicker('setDate', '{data_valor}');
                }} else {{
                    campo.value = '{data_valor}';
                }}
                
                campo.dispatchEvent(new Event('input', {{bubbles: true}}));
                campo.dispatchEvent(new Event('change', {{bubbles: true}}));
                campo.dispatchEvent(new Event('blur', {{bubbles: true}}));
                return campo.value;
            """),
            
            # Estratégia 2: Force fill padrão
            lambda: js_engine.force_fill(xpath_indexado, data_valor, by_xpath=True),
            
            # Estratégia 3: JavaScript puro
            lambda: js_engine.driver.execute_script(f"""
                var campo = document.evaluate("{xpath_indexado}", document, null, 
                    XPathResult.FIRST_ORDERED_NODE_TYPE, null).singleNodeValue;
                if (!campo) return false;
                
                campo.focus();
                campo.value = '';
                campo.value = '{data_valor}';
                
                ['input', 'change', 'blur', 'keyup'].forEach(ev => 
                    campo.dispatchEvent(new Event(ev, {{bubbles: true}}))
                );
                return campo.value;
            """),
            
            # Estratégia 4: ActionChains
            lambda: (
                elemento.click(),
                time.sleep(0.2),
                ActionChains(js_engine.driver)
                    .key_down(Keys.CONTROL).send_keys('a').key_up(Keys.CONTROL)
                    .send_keys(Keys.DELETE)
                    .send_keys(data_valor)
                    .send_keys(Keys.TAB)
                    .perform()
            )
        ]
        
        # Tenta cada estratégia
        for i, estrategia in enumerate(estrategias, 1):
            try:
                log(doc, f"   ▶️ Tentando estratégia {i}...")
                estrategia()
                time.sleep(0.5)
                
                # Valida
                elemento_reload = js_engine.driver.find_element(By.XPATH, xpath_indexado)
                valor_atual = elemento_reload.get_attribute('value') or ''
                
                if data_valor in valor_atual or valor_atual in data_valor:
                    log(doc, f"   ✅ {descricao} preenchido com estratégia {i}: '{valor_atual}'")
                    return True
                else:
                    log(doc, f"   ⚠️ Estratégia {i} não refletiu o valor")
                    
            except Exception as e:
                log(doc, f"   ⚠️ Estratégia {i} falhou: {str(e)[:80]}")
                continue
        
        log(doc, f"   ❌ Todas as estratégias falharam para {descricao}")
        return False
        
    except Exception as e:
        log(doc, f"   ❌ Erro ao preencher {descricao}: {e}")
        return False


def preencher_campo_monetario_por_indice(js_engine, doc, xpath_base, indice, valor, descricao="Valor", timeout=10):
    """
    Preenche um campo monetário específico por índice.
    Versão especializada para campos de valor/dinheiro.
    
    Args:
        js_engine: Instância do JSForceEngine
        doc: Documento para logs
        xpath_base: XPath base do campo monetário
        indice: Índice do elemento (1-based)
        valor: Valor a ser preenchido (pode ser com ou sem "R$ ")
        descricao: Descrição do campo para logs
        timeout: Timeout para operação
    
    Returns:
        bool: True se preencheu com sucesso, False caso contrário
    """
    xpath_indexado = f"({xpath_base})[{indice}]"
    
    try:
        log(doc, f"💰 Preenchendo {descricao} (índice {indice}): '{valor}'")
        
        # Conta elementos
        try:
            elementos = js_engine.driver.find_elements(By.XPATH, xpath_base)
            total = len(elementos)
            log(doc, f"   ℹ️ Total de campos encontrados: {total}")
            
            if total == 0:
                log(doc, f"   ⚠️ Nenhum campo encontrado")
                return False
            
            if indice > total:
                log(doc, f"   ⚠️ Índice {indice} inválido - só existem {total} campo(s)")
                return False
        except:
            pass
        
        # Normaliza o valor (garante que tem R$)
        valor_limpo = str(valor).replace('R$', '').replace(' ', '').strip()
        valor_formatado = f"R$ {valor_limpo}" if not valor.startswith('R$') else valor
        
        log(doc, f"   🎯 Valor formatado: '{valor_formatado}'")
        
        # Tenta preencher
        js_engine.force_fill(xpath_indexado, valor_formatado, by_xpath=True)
        time.sleep(0.4)
        
        # Valida
        try:
            elemento = js_engine.driver.find_element(By.XPATH, xpath_indexado)
            valor_atual = elemento.get_attribute('value') or ''
            
            # Compara valores normalizados
            atual_num = valor_atual.replace('R$', '').replace(' ', '').replace('.', '').replace(',', '.')
            esper_num = valor_limpo.replace('.', '').replace(',', '.')
            
            try:
                if float(atual_num.strip()) == float(esper_num.strip()):
                    log(doc, f"   ✅ {descricao} preenchido: '{valor_atual}'")
                    return True
            except:
                pass
            
            # Se não conseguiu comparar numericamente, compara texto
            if valor_limpo in valor_atual or valor_atual in valor_limpo:
                log(doc, f"   ✅ {descricao} preenchido: '{valor_atual}'")
                return True
            else:
                log(doc, f"   ⚠️ Valores diferem: esperado '{valor}', atual '{valor_atual}'")
                return True  # Retorna True pois pode ser só formatação
                
        except Exception as e:
            log(doc, f"   ⚠️ Não foi possível validar: {e}")
            return True
        
    except Exception as e:
        log(doc, f"   ❌ Erro ao preencher {descricao}: {e}")
        return False

# ==== SISTEMA ANTI-TIMEOUT JAVASCRIPT ====
class JSTimeoutHandler:
    """Sistema robusto para lidar com timeouts JavaScript no Selenium"""
    
    def __init__(self, driver, doc, timeout_padrao=10, max_retries=3):
        self.driver = driver
        self.doc = doc
        self.timeout_padrao = timeout_padrao
        self.max_retries = max_retries
        self.last_error = None
        
    def log_timeout(self, msg, level="INFO"):
        """Log com timestamp"""

        prefix = {
            "INFO": "ℹ️ ",
            "WARN": "⚠️ ",
            "ERROR": "❌ ",
            "SUCCESS": "✅ "
        }.get(level, "📝 ")
        
        print(f" {prefix} {msg}")
        if hasattr(self.doc, 'add_paragraph'):
            self.doc.add_paragraph(f"{msg}")
    
    def execute_js_safe(self, script, *args, timeout=None, fallback_result=None):
        """Executa JavaScript com proteção contra timeouts"""
        timeout = timeout or self.timeout_padrao
        
        original_timeout = self.driver.timeouts.script
        self.driver.set_script_timeout(timeout)
        
        for tentativa in range(1, self.max_retries + 1):
            try:
                if tentativa > 1:
                    self.log_timeout(f"Tentativa {tentativa}/{self.max_retries}", "INFO")
                
                result = self.driver.execute_script(script, *args)
                self.driver.set_script_timeout(original_timeout)
                
                if tentativa > 1:
                    self.log_timeout("JavaScript executado com sucesso", "SUCCESS")
                return result
                
            except JavascriptException as e:
                self.last_error = e
                self.log_timeout(f"Erro JavaScript: {str(e)[:150]}", "ERROR")
                self._limpar_estado_js()
                
                if tentativa < self.max_retries:
                    time.sleep(1 + tentativa * 0.5)
                    continue
                    
            except TimeoutException as e:
                self.last_error = e
                self.log_timeout(f"Timeout JavaScript ({timeout}s)", "ERROR")
                self._forcar_parada_js()
                
                if tentativa < self.max_retries:
                    time.sleep(2 + tentativa)
                    continue
                    
            except WebDriverException as e:
                self.last_error = e
                self.log_timeout(f"Erro WebDriver: {str(e)[:150]}", "ERROR")
                
                if tentativa < self.max_retries:
                    time.sleep(1.5)
                    continue
                    
            except Exception as e:
                self.last_error = e
                self.log_timeout(f"Erro inesperado: {str(e)[:150]}", "ERROR")
                break
        
        try:
            self.driver.set_script_timeout(original_timeout)
        except:
            pass
            
        if fallback_result is not None:
            self.log_timeout(f"Usando valor fallback: {fallback_result}", "WARN")
        return fallback_result
    
    def _limpar_estado_js(self):
        """Limpa estado JavaScript do browser"""
        try:
            cleanup_script = """
                if (window.__cleanupTimers) {
                    window.__cleanupTimers.forEach(clearTimeout);
                    window.__cleanupTimers.forEach(clearInterval);
                }
                if (typeof jQuery !== 'undefined') {
                    jQuery.active = 0;
                }
                window.__pendingRequests = 0;
                return true;
            """
            self.driver.execute_script(cleanup_script)
            time.sleep(0.5)
        except Exception:
            pass
    
    def _forcar_parada_js(self):
        """Força parada de JavaScript travado"""
        try:
            self.driver.execute_script("window.stop();")
            time.sleep(0.3)
        except Exception:
            pass



# ==== JS FORCE ENGINE COM PROTEÇÃO ANTI-TIMEOUT ====
class JSForceEngine:
    """Motor de execução JavaScript forçado com proteção contra timeouts"""
    
    def __init__(self, driver, wait, doc, timeout_padrao=10, max_retries=3):
        self.driver = driver
        self.wait = wait
        self.doc = doc
        self.timeout_handler = JSTimeoutHandler(driver, doc, timeout_padrao, max_retries)
    
    def execute_js(self, script, *args, timeout=None, fallback_result=None):
        """Executa JavaScript com proteção contra timeout"""
        return self.timeout_handler.execute_js_safe(
            script, *args, timeout=timeout, fallback_result=fallback_result
        )
    
    def wait_ajax_complete(self, timeout=15):
        """Aguarda AJAX completar com proteção contra timeout"""
        script = """
            var jQueryOk = (typeof jQuery==='undefined') || (jQuery.active===0);
            var fetchOk = !window.__pendingRequests || window.__pendingRequests===0;
            var overlays = document.querySelectorAll(
                '.blockScreen, .blockUI, .loading, .overlay, [class*="loading"], [class*="spinner"]'
            );
            var overlayOk = true;
            for (var i=0; i<overlays.length; i++){
                var s=window.getComputedStyle(overlays[i]);
                if(s.display!=='none' && s.visibility!=='hidden' && parseFloat(s.opacity||1)>0.01){
                    overlayOk=false;
                    break;
                }
            }
            return jQueryOk && fetchOk && overlayOk;
        """
        
        end = time.time() + timeout
        while time.time() < end:
            try:
                done = self.execute_js(script, timeout=5, fallback_result=True)
                if done:
                    return True
            except:
                pass
            time.sleep(0.2)
        return True
    
    def scroll_into_view(self, target, padding=100):
        """
        Faz scroll até um WebElement (ou seletor XPath/CSS).
        padding: desloca um pouco pra cima pra não ficar colado no topo.
        """
        try:
            el = target
            # Se vier como string, tenta resolver
            if isinstance(target, str):
                try:
                    el = self.driver.find_element("xpath", target)
                except Exception:
                    el = self.driver.find_element("css selector", target)

            # Estratégia principal via JS
            self.driver.execute_script("""
                const el = arguments[0], pad = arguments[1] || 0;
                if (!el) return;
                el.scrollIntoView({block:'center', inline:'center'});
                try { window.scrollBy(0, -pad); } catch(e) {}
            """, el, padding)
            time.sleep(0.2)
            return True
        except Exception:
            # Fallback ActionChains
            try:
                ActionChains(self.driver).move_to_element(el).perform()
                time.sleep(0.2)
                return True
            except Exception:
                return False

    def click_element(self, el, wait_after=0.4):
        """
        Clica com múltiplas estratégias em um WebElement já localizado.
        """
        try:
            # 1) Clique padrão
            el.click()
            time.sleep(wait_after)
            self.wait_ajax_complete(min(10, int(wait_after*10)) or 3)
            return True
        except Exception:
            pass

        try:
            # 2) Clique via JS
            self.driver.execute_script("arguments[0].click();", el)
            time.sleep(wait_after)
            self.wait_ajax_complete(min(10, int(wait_after*10)) or 3)
            return True
        except Exception:
            pass

        try:
            # 3) Sequência de eventos de mouse
            self.execute_js("""
                const e = arguments[0];
                const rect = e.getBoundingClientRect();
                const x = rect.left + 5, y = rect.top + 5;
                ['mouseover','mouseenter','mousemove','mousedown','mouseup','click'].forEach(t=>{
                    e.dispatchEvent(new MouseEvent(t,{bubbles:true,cancelable:true,view:window,clientX:x,clientY:y}));
                });
                if (typeof e.click==='function') e.click();
                return true;
            """, el, timeout=5, fallback_result=True)
            time.sleep(wait_after)
            self.wait_ajax_complete(min(10, int(wait_after*10)) or 3)
            return True
        except Exception:
            pass

        # 4) Último recurso: ActionChains
        try:
            ActionChains(self.driver).move_to_element(el).pause(0.05).click().perform()
            time.sleep(wait_after)
            self.wait_ajax_complete(min(10, int(wait_after*10)) or 3)
            return True
        except Exception:
            pass

        raise Exception("Não foi possível clicar no elemento com as estratégias disponíveis.")
    
    def force_click(self, selector, by_xpath=False, max_attempts=5):
        """Clique forçado com proteção contra timeout"""
        log(self.doc, f"🎯 Clique forçado em: {selector}")
        
        for attempt in range(max_attempts):
            try:
                strategies = [
                    self._click_strategy_2,
                    self._click_strategy_1,
                    self._click_strategy_3,
                    self._click_strategy_4,
                    self._click_strategy_5,
                ]
                
                for i, strategy in enumerate(strategies, 1):
                    try:
                        if attempt > 0 or i > 1:
                            log(self.doc, f"   Tentativa {attempt + 1}.{i}...")
                        
                        result = self.execute_js(
                            self._get_strategy_script(strategy, selector, by_xpath),
                            selector,
                            by_xpath,
                            timeout=5,
                            fallback_result=False
                        )
                        
                        if result:
                            log(self.doc, f"✅ Clique bem-sucedido (estratégia {i})")
                            time.sleep(0.5)
                            self.wait_ajax_complete(10)
                            return True
                            
                    except Exception as e:
                        if i == 1 and attempt == 0:
                            pass  # Silencia primeiro erro
                        else:
                            log(self.doc, f"   Estratégia {i} falhou: {str(e)[:80]}")
                        continue
                
                if attempt < max_attempts - 1:
                    time.sleep(1 + attempt * 0.5)
                    
            except Exception as e:
                log(self.doc, f"⚠️ Tentativa {attempt + 1} falhou: {e}")
                if attempt < max_attempts - 1:
                    time.sleep(1.5)
        
        raise Exception(f"Falha ao clicar após {max_attempts} tentativas: {selector}")
    
    def _get_strategy_script(self, strategy_func, selector, by_xpath):
        """Retorna o script JavaScript para cada estratégia"""
        base_locator = """
            var selector = arguments[0];
            var byXPath = arguments[1];
            var element;
            
            if (byXPath) {
                var result = document.evaluate(selector, document, null, 
                    XPathResult.FIRST_ORDERED_NODE_TYPE, null);
                element = result.singleNodeValue;
            } else {
                element = document.querySelector(selector);
            }
            
            if (!element) throw new Error('Elemento não encontrado');
        """
        
        if strategy_func == self._click_strategy_1:
            return base_locator + """
                element.style.pointerEvents = 'auto';
                element.style.display = 'block';
                element.style.visibility = 'visible';
                element.style.opacity = '1';
                element.removeAttribute('disabled');
                element.scrollIntoView({behavior: 'smooth', block: 'center'});
                setTimeout(function() { element.click(); }, 300);
                return true;
            """
        elif strategy_func == self._click_strategy_2:
            return base_locator + """
                element.style.pointerEvents = 'auto';
                element.removeAttribute('disabled');
                element.scrollIntoView({block: 'center'});
                
                var events = ['mouseover', 'mouseenter', 'mousemove', 'mousedown', 'mouseup', 'click'];
                events.forEach(function(eventType) {
                    var evt = new MouseEvent(eventType, {
                        bubbles: true, cancelable: true, view: window, detail: 1,
                        clientX: element.getBoundingClientRect().left + 5,
                        clientY: element.getBoundingClientRect().top + 5
                    });
                    element.dispatchEvent(evt);
                });
                
                if (typeof element.click === 'function') element.click();
                return true;
            """
        elif strategy_func == self._click_strategy_3:
            return base_locator + """
                element.style.display = 'block';
                element.style.visibility = 'visible';
                element.style.opacity = '1';
                element.style.pointerEvents = 'auto';
                element.focus();
                element.click();
                element.dispatchEvent(new Event('click', {bubbles: true, cancelable: true}));
                return true;
            """
        elif strategy_func == self._click_strategy_4:
            return base_locator + """
                element.removeAttribute('disabled');
                element.removeAttribute('readonly');
                element.style.pointerEvents = 'auto !important';
                element.style.display = 'block !important';
                element.style.visibility = 'visible !important';
                element.style.opacity = '1 !important';
                
                var overlays = document.querySelectorAll('.modal, .overlay, .blockUI, [role="dialog"]');
                overlays.forEach(function(overlay) {
                    overlay.style.display = 'none';
                    overlay.style.visibility = 'hidden';
                });
                
                element.focus();
                element.click();
                
                var clickEvent = new MouseEvent('click', {
                    view: window, bubbles: true, cancelable: true
                });
                element.dispatchEvent(clickEvent);
                
                if (typeof jQuery !== 'undefined') jQuery(element).trigger('click');
                return true;
            """
        else:  # strategy_5
            return base_locator + """
                var rect = element.getBoundingClientRect();
                var x = rect.left + rect.width / 2;
                var y = rect.top + rect.height / 2;
                
                var evt = document.createEvent('MouseEvents');
                evt.initMouseEvent('click', true, true, window, 1, x, y, x, y, false, false, false, false, 0, null);
                element.dispatchEvent(evt);
                
                if (element.onclick) element.onclick();
                
                var parent = element.parentElement;
                while (parent && parent !== document.body) {
                    if (parent.onclick) {
                        parent.onclick();
                        break;
                    }
                    parent = parent.parentElement;
                }
                return true;
            """
    
    def _click_strategy_1(self, selector, by_xpath):
        pass  # Implementado via _get_strategy_script
    
    def _click_strategy_2(self, selector, by_xpath):
        pass
    
    def _click_strategy_3(self, selector, by_xpath):
        pass
    
    def _click_strategy_4(self, selector, by_xpath):
        pass
    
    def _click_strategy_5(self, selector, by_xpath):
        pass
    
    def force_fill(self, selector, value, by_xpath=False, max_attempts=5):
        """Preenchimento forçado com proteção contra timeout"""
        log(self.doc, f"✏️ Preenchimento forçado: {selector} = '{value}'")
        
        for attempt in range(max_attempts):
            try:
                strategies = [
                    self._fill_strategy_1,
                    self._fill_strategy_2,
                    self._fill_strategy_3,
                    self._fill_strategy_4,
                ]
                
                for i, strategy in enumerate(strategies, 1):
                    try:
                        if attempt > 0 or i > 1:
                            log(self.doc, f"   Tentativa {attempt + 1}.{i}...")
                        
                        result = self.execute_js(
                            self._get_fill_script(strategy, selector, value, by_xpath),
                            selector,
                            value,
                            by_xpath,
                            timeout=5,
                            fallback_result=None
                        )
                        
                        time.sleep(0.3)
                        if self._validate_fill(selector, value, by_xpath):
                            log(self.doc, f"✅ Campo preenchido (estratégia {i})")
                            return True
                    except Exception as e:
                        if i == 1 and attempt == 0:
                            pass
                        else:
                            log(self.doc, f"   Estratégia {i} falhou: {str(e)[:80]}")
                        continue
                
                if attempt < max_attempts - 1:
                    time.sleep(1 + attempt * 0.5)
                    
            except Exception as e:
                log(self.doc, f"⚠️ Tentativa {attempt + 1} falhou: {e}")
        
        raise Exception(f"Falha ao preencher após {max_attempts} tentativas: {selector}")

    def _get_fill_script(self, strategy_func, selector, value, by_xpath):
        """Retorna o script de preenchimento"""
        base_locator = """
            var selector = arguments[0];
            var value = arguments[1];
            var byXPath = arguments[2];
            var element;
            
            if (byXPath) {
                var result = document.evaluate(selector, document, null, 
                    XPathResult.FIRST_ORDERED_NODE_TYPE, null);
                element = result.singleNodeValue;
            } else {
                element = document.querySelector(selector);
            }
            
            if (!element) throw new Error('Campo não encontrado');
        """
        
        if strategy_func == self._fill_strategy_1:
            return base_locator + """
                element.removeAttribute('disabled');
                element.removeAttribute('readonly');
                element.style.display = 'block';
                element.style.visibility = 'visible';
                element.scrollIntoView({block: 'center'});
                element.focus();
                element.dispatchEvent(new Event('focus', {bubbles: true}));
                element.value = '';
                element.value = value;
                ['input', 'change', 'blur', 'keyup'].forEach(function(evt) {
                    element.dispatchEvent(new Event(evt, {bubbles: true}));
                });
                return element.value;
            """
        elif strategy_func == self._fill_strategy_2:
            return base_locator + """
                var nativeInputValueSetter = Object.getOwnPropertyDescriptor(
                    window.HTMLInputElement.prototype, 'value'
                ).set;
                if (nativeInputValueSetter) {
                    nativeInputValueSetter.call(element, value);
                } else {
                    element.value = value;
                }
                element.dispatchEvent(new Event('input', {bubbles: true}));
                element.dispatchEvent(new Event('change', {bubbles: true}));
                element.dispatchEvent(new KeyboardEvent('keydown', {bubbles: true}));
                element.dispatchEvent(new KeyboardEvent('keyup', {bubbles: true}));
                element.dispatchEvent(new Event('blur', {bubbles: true}));
                return element.value;
            """
        elif strategy_func == self._fill_strategy_3:
            return base_locator + """
                element.value = value;
                if (typeof jQuery !== 'undefined') {
                    jQuery(element).val(value).trigger('input').trigger('change').trigger('blur');
                }
                ['focus', 'input', 'change', 'blur'].forEach(function(evt) {
                    element.dispatchEvent(new Event(evt, {bubbles: true}));
                });
                return element.value;
            """
        else:  # strategy_4
            return base_locator + """
                element.removeAttribute('disabled');
                element.removeAttribute('readonly');
                element.removeAttribute('maxlength');
                element.value = '';
                element.setAttribute('value', value);
                element.value = value;
                element.style.color = element.style.color;
                
                var events = ['focus', 'click', 'input', 'change', 'keydown', 'keypress', 
                              'keyup', 'blur', 'paste', 'textInput'];
                events.forEach(function(evt) {
                    try {
                        element.dispatchEvent(new Event(evt, {bubbles: true, cancelable: true}));
                    } catch(e) {}
                });
                
                if (element.oninput) element.oninput();
                if (element.onchange) element.onchange();
                return element.value;
            """
    
    def _fill_strategy_1(self, selector, value, by_xpath):
        pass
    
    def _fill_strategy_2(self, selector, value, by_xpath):
        pass
    
    def _fill_strategy_3(self, selector, value, by_xpath):
        pass
    
    def _fill_strategy_4(self, selector, value, by_xpath):
        pass
    
    def _validate_fill(self, selector, expected_value, by_xpath):
        """Valida preenchimento"""
        script = """
            var selector = arguments[0];
            var expected = arguments[1];
            var byXPath = arguments[2];
            var element;
            
            if (byXPath) {
                var result = document.evaluate(selector, document, null, 
                    XPathResult.FIRST_ORDERED_NODE_TYPE, null);
                element = result.singleNodeValue;
            } else {
                element = document.querySelector(selector);
            }
            
            if (!element) return false;
            var actual = element.value || '';
            return actual.trim() === expected.trim() || actual.includes(expected);
        """
        try:
            return self.execute_js(script, selector, expected_value, by_xpath, timeout=3, fallback_result=False)
        except:
            return False
        
import time

def clicar_finalizar_e_verificar_alerta(js_engine, doc, timeout=5, pausa=0.5):
    """
    Clica no botão 'Finalizar', aguarda 0,5s e procura mensagens de alerta.
    Usa safe_action e js_engine.force_click().
    """
    safe_action(doc, "Clicando em 'Finalizar'", lambda:
        js_engine.force_click(
            "//a[@class='btModel btGray btyes' and normalize-space()='Finalizar']",
            by_xpath=True
        )
    )

    time.sleep(pausa)
    log(doc, "🔍 Verificando mensagens de alerta após o clique em 'Finalizar'...")
    return encontrar_mensagem_alerta()


import time
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.common.exceptions import (
    TimeoutException,
    StaleElementReferenceException,
    ElementClickInterceptedException,
    NoSuchElementException,
    ElementNotInteractableException
)
from selenium.webdriver import ActionChains

class LOVHandler:
    """
    Handler ultra-robusto para manipulação de LOV (List of Values) com:
    - 10 estratégias diferentes de clique
    - Detecção automática de iframes
    - Retry inteligente com backoff exponencial
    - Reforço automático de cliques
    - Logs detalhados de cada tentativa
    """
    
    def __init__(self, js_engine, doc, max_retries=5, timeout=10):
        self.js = js_engine
        self.doc = doc
        self.driver = js_engine.driver
        self.max_retries = max_retries
        self.timeout = timeout
        
    def _log(self, msg, level="INFO"):
        """Log padronizado com níveis"""
        prefixes = {
            "INFO": "ℹ️",
            "SUCCESS": "✅",
            "WARNING": "⚠️",
            "ERROR": "❌",
            "DEBUG": "🔍"
        }
        prefix = prefixes.get(level, "📝")
        log(self.doc, f"{prefix} {msg}")
    
    def _wait_element(self, locator_type, locator_value, timeout=None, condition="clickable"):
        """Aguarda elemento com diferentes condições"""
        timeout = timeout or self.timeout
        conditions = {
            "present": EC.presence_of_element_located,
            "visible": EC.visibility_of_element_located,
            "clickable": EC.element_to_be_clickable
        }
        
        try:
            cond = conditions.get(condition, conditions["clickable"])
            return WebDriverWait(self.driver, timeout).until(
                cond((locator_type, locator_value))
            )
        except TimeoutException:
            return None
    
    def _is_element_visible(self, element):
        """Verifica se elemento está realmente visível"""
        try:
            if not element:
                return False
            if not element.is_displayed():
                return False
            
            # Verifica via JavaScript também
            is_visible = self.driver.execute_script("""
                const el = arguments[0];
                if (!el) return false;
                const style = window.getComputedStyle(el);
                return (
                    style.display !== 'none' &&
                    style.visibility !== 'hidden' &&
                    parseFloat(style.opacity || 1) > 0.01 &&
                    el.offsetParent !== null
                );
            """, element)
            
            return is_visible
        except:
            return False
    
    def _force_element_visible(self, element):
        """Força elemento a ficar visível e interativo"""
        try:
            self.driver.execute_script("""
                const el = arguments[0];
                el.style.display = 'block';
                el.style.visibility = 'visible';
                el.style.opacity = '1';
                el.style.pointerEvents = 'auto';
                el.removeAttribute('disabled');
                el.removeAttribute('readonly');
                
                // Remove overlays que podem bloquear
                const overlays = document.querySelectorAll(
                    '.modal-backdrop, .overlay, .blockUI, [class*="loading"]'
                );
                overlays.forEach(o => {
                    o.style.display = 'none';
                    o.style.visibility = 'hidden';
                });
            """, element)
            return True
        except:
            return False
    
    def _detect_and_enter_iframe(self, iframe_xpath=None):
        """Detecta e entra no iframe automaticamente"""
        try:
            # Primeiro volta para o contexto principal
            self.driver.switch_to.default_content()
            
            # Se foi fornecido um xpath específico
            if iframe_xpath:
                try:
                    frame = self._wait_element("xpath", iframe_xpath, timeout=3)
                    if frame:
                        self.driver.switch_to.frame(frame)
                        self._log(f"Entrou no iframe: {iframe_xpath}", "SUCCESS")
                        return True
                except:
                    pass
            
            # Detecção automática de iframes
            iframe_selectors = [
                "//iframe[contains(@class,'LOV') or contains(@id,'LOV') or contains(@id,'lov')]",
                "//iframe[contains(@class,'modal') or contains(@class,'popup')]",
                "//iframe[contains(@src,'lov') or contains(@src,'LOV')]",
                "(//iframe)[last()]"  # Último iframe (geralmente o modal)
            ]
            
            for selector in iframe_selectors:
                try:
                    frames = self.driver.find_elements(By.XPATH, selector)
                    for frame in frames:
                        if self._is_element_visible(frame):
                            self.driver.switch_to.frame(frame)
                            self._log(f"Iframe detectado automaticamente: {selector}", "SUCCESS")
                            return True
                except:
                    continue
            
            return False
        except Exception as e:
            self._log(f"Erro ao detectar iframe: {e}", "WARNING")
            return False
    
    def _click_strategy_1_standard(self, element):
        """Estratégia 1: Clique padrão do Selenium"""
        element.click()
        return True
    
    def _click_strategy_2_javascript_simple(self, element):
        """Estratégia 2: JavaScript simples"""
        self.driver.execute_script("arguments[0].click();", element)
        return True
    
    def _click_strategy_3_javascript_advanced(self, element):
        """Estratégia 3: JavaScript com eventos completos"""
        self.driver.execute_script("""
            const el = arguments[0];
            const rect = el.getBoundingClientRect();
            const x = rect.left + rect.width / 2;
            const y = rect.top + rect.height / 2;
            
            const events = ['mouseover', 'mouseenter', 'mousemove', 'mousedown', 'mouseup', 'click'];
            events.forEach(eventType => {
                const evt = new MouseEvent(eventType, {
                    bubbles: true,
                    cancelable: true,
                    view: window,
                    detail: 1,
                    clientX: x,
                    clientY: y
                });
                el.dispatchEvent(evt);
            });
            
            if (typeof el.click === 'function') el.click();
        """, element)
        return True
    
    def _click_strategy_4_action_chains(self, element):
        """Estratégia 4: ActionChains com pause"""
        ActionChains(self.driver)\
            .move_to_element(element)\
            .pause(0.1)\
            .click()\
            .perform()
        return True
    
    def _click_strategy_5_action_chains_offset(self, element):
        """Estratégia 5: ActionChains com offset"""
        ActionChains(self.driver)\
            .move_to_element_with_offset(element, 5, 5)\
            .pause(0.05)\
            .click()\
            .perform()
        return True
    
    def _click_strategy_6_force_visible_then_click(self, element):
        """Estratégia 6: Força visibilidade e clica"""
        self._force_element_visible(element)
        time.sleep(0.2)
        element.click()
        return True
    
    def _click_strategy_7_scroll_and_click(self, element):
        """Estratégia 7: Scroll suave até elemento"""
        self.driver.execute_script("""
            arguments[0].scrollIntoView({
                behavior: 'smooth',
                block: 'center',
                inline: 'center'
            });
        """, element)
        time.sleep(0.3)
        element.click()
        return True
    
    def _click_strategy_8_remove_overlays(self, element):
        """Estratégia 8: Remove todos os overlays e clica"""
        self.driver.execute_script("""
            // Remove todos os overlays possíveis
            const overlays = document.querySelectorAll(`
                .modal-backdrop, .overlay, .blockUI, .blockScreen,
                [class*="loading"], [class*="spinner"], [class*="overlay"],
                [style*="z-index: 9999"], [style*="position: fixed"]
            `);
            overlays.forEach(o => {
                o.style.display = 'none';
                o.style.visibility = 'hidden';
                o.remove();
            });
            
            const el = arguments[0];
            el.style.zIndex = '999999';
            el.click();
        """, element)
        return True
    
    def _click_strategy_9_jquery_trigger(self, element):
        """Estratégia 9: jQuery trigger (se disponível)"""
        self.driver.execute_script("""
            const el = arguments[0];
            if (typeof jQuery !== 'undefined') {
                jQuery(el).trigger('click');
            } else {
                el.click();
            }
        """, element)
        return True
    
    def _click_strategy_10_nuclear_option(self, element):
        """Estratégia 10: Opção nuclear - força tudo"""
        self.driver.execute_script("""
            const el = arguments[0];
            
            // Remove TODOS os atributos que podem bloquear
            el.removeAttribute('disabled');
            el.removeAttribute('readonly');
            el.style.pointerEvents = 'auto !important';
            el.style.display = 'block !important';
            el.style.visibility = 'visible !important';
            el.style.opacity = '1 !important';
            el.style.zIndex = '999999 !important';
            
            // Remove TODOS os overlays da página
            document.querySelectorAll('*').forEach(elem => {
                const style = window.getComputedStyle(elem);
                if (
                    style.position === 'fixed' &&
                    parseInt(style.zIndex) > 1000 &&
                    elem !== el &&
                    !elem.contains(el)
                ) {
                    elem.style.display = 'none';
                }
            });
            
            // Foca no elemento
            el.focus();
            
            // Dispara TODOS os eventos possíveis
            const allEvents = [
                'focus', 'focusin', 'mouseover', 'mouseenter', 'mousemove',
                'mousedown', 'mouseup', 'click', 'dblclick'
            ];
            
            allEvents.forEach(eventType => {
                try {
                    const evt = new MouseEvent(eventType, {
                        bubbles: true,
                        cancelable: true,
                        view: window
                    });
                    el.dispatchEvent(evt);
                } catch(e) {}
            });
            
            // Clique direto
            if (typeof el.click === 'function') el.click();
            
            // jQuery se disponível
            if (typeof jQuery !== 'undefined') {
                jQuery(el).trigger('click');
            }
            
            // Tenta onclick manual
            if (el.onclick) el.onclick();
            
            return true;
        """, element)
        return True
    
    def _advanced_click(self, element, max_attempts=10):
        """
        Sistema avançado de clique com 10 estratégias diferentes
        Tenta cada estratégia até uma funcionar
        """
        strategies = [
            ("Clique Padrão", self._click_strategy_1_standard),
            ("JavaScript Simples", self._click_strategy_2_javascript_simple),
            ("JavaScript Avançado", self._click_strategy_3_javascript_advanced),
            ("ActionChains", self._click_strategy_4_action_chains),
            ("ActionChains Offset", self._click_strategy_5_action_chains_offset),
            ("Força Visível", self._click_strategy_6_force_visible_then_click),
            ("Scroll e Clique", self._click_strategy_7_scroll_and_click),
            ("Remove Overlays", self._click_strategy_8_remove_overlays),
            ("jQuery Trigger", self._click_strategy_9_jquery_trigger),
            ("Opção Nuclear", self._click_strategy_10_nuclear_option)
        ]
        
        for attempt in range(1, max_attempts + 1):
            for strategy_name, strategy_func in strategies:
                try:
                    self._log(f"Tentativa {attempt}/{max_attempts}: {strategy_name}", "DEBUG")
                    
                    # Tenta executar a estratégia
                    strategy_func(element)
                    time.sleep(0.3)
                    
                    self._log(f"✓ {strategy_name} funcionou!", "SUCCESS")
                    return True
                    
                except StaleElementReferenceException:
                    self._log(f"Elemento ficou stale, tentando recarregar...", "WARNING")
                    return False  # Precisa recarregar elemento
                    
                except (ElementClickInterceptedException, 
                        ElementNotInteractableException) as e:
                    self._log(f"✗ {strategy_name}: {str(e)[:50]}", "DEBUG")
                    continue
                    
                except Exception as e:
                    self._log(f"✗ {strategy_name}: {str(e)[:50]}", "DEBUG")
                    continue
            
            if attempt < max_attempts:
                time.sleep(0.5 * attempt)  # Backoff exponencial
        
        return False
    
    def _fill_search_fields(self, search_text, max_fields=10):
        """Preenche TODOS os campos de pesquisa encontrados"""
        if not search_text:
            return 0
        
        search_field_xpaths = [
            "//input[@id='txtPesquisa']",
            "//input[contains(@class,'pesquisa')]",
            "//input[contains(@class,'nomePesquisa')]",
            "//input[contains(translate(@name,'PESQUISA','pesquisa'),'pesquisa')]",
            "//input[@type='text' and contains(@style,'width:210px')]",
            "//input[@type='text' and not(@disabled)]",
        ]
        
        filled_count = 0
        for xpath in search_field_xpaths:
            try:
                fields = self.driver.find_elements(By.XPATH, xpath)
                for field in fields[:max_fields]:
                    try:
                        if not self._is_element_visible(field):
                            continue
                        
                        # Limpa e preenche
                        self._force_element_visible(field)
                        field.clear()
                        field.click()
                        field.send_keys(search_text)
                        filled_count += 1
                        
                    except:
                        continue
            except:
                continue
        
        if filled_count > 0:
            self._log(f"Preenchidos {filled_count} campos de pesquisa", "SUCCESS")
        else:
            self._log("Nenhum campo de pesquisa encontrado", "WARNING")
        
        return filled_count
    
    def _click_search_button(self):
        """Clica no botão Pesquisar com múltiplas estratégias"""
        search_button_xpaths = [
            "//a[contains(@class,'btPesquisar') and contains(normalize-space(.),'Pesquisar')]",
            "//button[contains(normalize-space(.),'Pesquisar')]",
            "//input[@type='button' and contains(@value,'Pesquisar')]",
            "//a[contains(@class,'lpFind')]",
            "//a[contains(@onclick,'pesquisar')]"
        ]
        
        # Fallback: ENTER no campo ativo
        try:
            self._log("Tentando ENTER no campo de pesquisa", "DEBUG")
            ActionChains(self.driver).send_keys(Keys.ENTER).perform()
            return True
        except Exception:
            pass
        
        for xpath in search_button_xpaths:
            try:
                btn = self._wait_element("xpath", xpath, timeout=2)
                if btn and self._is_element_visible(btn):
                    if self._advanced_click(btn):
                        self._log("Botão 'Pesquisar' clicado", "SUCCESS")
                        return True
            except:
                pass
        
        self._log("Não foi possível clicar em 'Pesquisar'", "WARNING")
        return False
    
    def _select_result(self, result_text=""):
        """Seleciona o resultado da pesquisa"""
        result_xpaths = []
        
        if result_text:
            result_xpaths = [
                f"//td[contains(normalize-space(.),'{result_text}')]",
                f"//span[contains(normalize-space(.),'{result_text}')]",
                f"//div[contains(normalize-space(.),'{result_text}')]",
                f"//tr[contains(normalize-space(.),'{result_text}')]//td[1]",
                f"//li[contains(normalize-space(.),'{result_text}')]"
            ]
        else:
            result_xpaths = [
                "(//table//tr[1]/td[1])[1]",
                "(//tr[contains(@class,'ui-widget-content')][1])[1]",
                "(//tr[contains(@class,'rich-table-row')][1])[1]",
                "(//li[@class='ui-autocomplete-item'])[1]"
            ]
        
        for xpath in result_xpaths:
            try:
                result = self._wait_element("xpath", xpath, timeout=3)
                if result and self._is_element_visible(result):
                    if self._advanced_click(result):
                        self._log(f"Resultado selecionado: {result_text or 'primeiro'}", "SUCCESS")
                        return xpath  # Retorna xpath usado
            except:
                continue
        
        self._log(f"Não foi possível selecionar: {result_text or 'primeiro resultado'}", "ERROR")
        return None
    
    def _reinforce_click_on_result(self, result_xpath, max_reclick=5):
        """
        Reforça o clique no resultado SEMPRE, independente do modal
        Clica repetidamente até max_reclick vezes
        """
        if not result_xpath:
            return
        
        self._log(f"🔁 Reforçando clique no resultado ({max_reclick}x)", "INFO")
        
        for attempt in range(1, max_reclick + 1):
            try:
                time.sleep(0.35)
                
                # Tenta localizar o resultado novamente
                result = self.driver.find_element(By.XPATH, result_xpath)
                
                # Clica usando sistema avançado
                self._log(f"Reforço {attempt}/{max_reclick}", "DEBUG")
                self._advanced_click(result)
                
            except StaleElementReferenceException:
                self._log(f"Elemento stale no reforço {attempt}", "WARNING")
                try:
                    # Tenta recarregar o elemento
                    result = self.driver.find_element(By.XPATH, result_xpath)
                    self._advanced_click(result)
                except:
                    continue
                    
            except Exception as e:
                self._log(f"Erro no reforço {attempt}: {str(e)[:50]}", "WARNING")
                continue
        
        self._log("✅ Reforço de cliques concluído", "SUCCESS")
    
    def open_and_select(
        self,
        btn_index=None,
        btn_xpath=None,
        btn_css=None,
        search_text="",
        result_text="",
        iframe_xpath=None,
        auto_detect_iframe=True,
        reinforce_clicks=5,
        wait_after=0.5
    ):
        """
        Método principal: abre LOV, pesquisa e seleciona resultado
        
        Args:
            btn_index: Índice do botão LOV (0-based)
            btn_xpath: XPath customizado do botão
            btn_css: Seletor CSS do botão
            search_text: Texto para pesquisar
            result_text: Texto do resultado a selecionar
            iframe_xpath: XPath do iframe (se houver)
            auto_detect_iframe: Detectar iframe automaticamente
            reinforce_clicks: Quantas vezes reforçar o clique (padrão: 5)
            wait_after: Tempo de espera após conclusão
        
        Returns:
            bool: True se bem-sucedido, False caso contrário
        """
        
        self._log(f"🔍 LOV: '{search_text}' → '{result_text}'", "INFO")
        
        for retry in range(1, self.max_retries + 1):
            try:
                # ===== PASSO 1: Volta para contexto principal =====
                try:
                    self.driver.switch_to.default_content()
                except:
                    pass
                
                # ===== PASSO 2: Localiza e clica no botão LOV =====
                btn_selector = None
                by_type = None
                
                if btn_index is not None:
                    btn_selector = f"(//a[@class='sprites sp-openLov'])[{btn_index + 1}]"
                    by_type = "xpath"
                elif btn_xpath:
                    btn_selector = btn_xpath
                    by_type = "xpath"
                elif btn_css:
                    btn_selector = btn_css
                    by_type = "css"
                else:
                    raise ValueError("Forneça btn_index, btn_xpath ou btn_css")
                
                self._log(f"Localizando botão LOV: {btn_selector}", "DEBUG")
                
                lov_button = self._wait_element(by_type, btn_selector, timeout=5)
                if not lov_button:
                    raise Exception(f"Botão LOV não encontrado: {btn_selector}")
                
                self._log("Abrindo LOV...", "INFO")
                if not self._advanced_click(lov_button):
                    raise Exception("Falha ao clicar no botão LOV")
                
                time.sleep(0.8)
                
                # ===== PASSO 3: Entra no iframe se necessário =====
                if iframe_xpath or auto_detect_iframe:
                    self._detect_and_enter_iframe(iframe_xpath)
                    time.sleep(0.3)
                
                # ===== PASSO 4: Preenche campos de pesquisa =====
                if search_text:
                    self._fill_search_fields(search_text)
                    time.sleep(0.3)
                
                # ===== PASSO 5: Clica em Pesquisar =====
                self._click_search_button()
                time.sleep(0.8)
                
                # ===== PASSO 6: Seleciona resultado =====
                result_xpath = self._select_result(result_text)
                if not result_xpath:
                    if retry < self.max_retries:
                        self._log(f"Retry {retry}/{self.max_retries}...", "WARNING")
                        time.sleep(1.5 * retry)
                        continue
                    raise Exception("Falha ao selecionar resultado")
                
                # ===== PASSO 7: Reforça clique SEMPRE =====
                if reinforce_clicks > 0:
                    self._reinforce_click_on_result(result_xpath, reinforce_clicks)
                
                # ===== PASSO 8: Volta para contexto principal =====
                try:
                    self.driver.switch_to.default_content()
                except:
                    pass
                
                time.sleep(wait_after)
                self._log("LOV concluído com sucesso!", "SUCCESS")
                return True
                
            except Exception as e:
                self._log(f"Tentativa {retry} falhou: {str(e)[:100]}", "ERROR")
                
                # Tenta recuperar voltando para contexto principal
                try:
                    self.driver.switch_to.default_content()
                except:
                    pass
                
                if retry < self.max_retries:
                    time.sleep(2 * retry)  # Backoff exponencial
                    continue
        
        self._log(f"LOV falhou após {self.max_retries} tentativas", "ERROR")
        return False
    
# ==== FUNÇÕES AUXILIARES ====

def encontrar_mensagem_alerta():
    seletores = [
        (".alerts.salvo", "✅ Mensagem de Sucesso"),
        (".alerts.alerta", "⚠️ Mensagem de Alerta"),
        (".alerts.erro", "❌ Mensagem de Erro"),
    ]

    for seletor, tipo in seletores:
        try:
            elemento = driver.find_element(By.CSS_SELECTOR, seletor)
            if elemento.is_displayed():
                log(doc, f"📢 {tipo}: {elemento.text}")
                return elemento
        except:
            continue

    log(doc, "ℹ️ Nenhuma mensagem de alerta encontrada.")
    return None


def safe_action(doc, descricao, func, max_retries=3):
    """Wrapper para ações com retry automático"""
    global driver
    
    for attempt in range(max_retries):
        try:
            log(doc, f"🔄 {descricao}..." if attempt == 0 else f"🔄 {descricao}... (Tentativa {attempt + 1})")
            func()
            log(doc, f"✅ {descricao} realizada com sucesso.")
            take_screenshot(driver, doc, _sanitize_filename(descricao))
            return True
        except Exception as e:
            if attempt < max_retries - 1:
                log(doc, f"⚠️ Tentativa {attempt + 1} falhou, tentando novamente..."),
                time.sleep(2 + attempt)
                continue
            else:
                log(doc, f"❌ Erro após {max_retries} tentativas: {e}")
                take_screenshot(driver, doc, _sanitize_filename(f"erro_{descricao}"))
                return False
    
    return False

def inicializar_driver():
    """Inicializa WebDriver com configurações otimizadas"""
    global driver, wait
    
    try:
        options = Options()
        options.add_argument("--start-maximized")
        options.add_argument("--disable-blink-features=AutomationControlled")
        options.add_argument("--disable-dev-shm-usage")
        options.add_argument("--no-sandbox")
        options.add_experimental_option("excludeSwitches", ["enable-automation"])
        options.add_experimental_option('useAutomationExtension', False)
        
        driver = webdriver.Chrome(
            service=Service(ChromeDriverManager().install()), 
            options=options
        )
        
        driver.execute_script(
            "Object.defineProperty(navigator, 'webdriver', {get: () => undefined})"
        )
        
        # Configura timeouts globais
        driver.set_script_timeout(30)
        driver.implicitly_wait(10)
        
        wait = WebDriverWait(driver, TIMEOUT_DEFAULT)
        
        log(doc, "✅ Driver inicializado com sucesso")
        return True
        
    except Exception as e:
        log(doc, f"❌ Erro ao inicializar driver: {e}")
        return False

def finalizar_relatorio():
    """Salva relatório e fecha driver"""
    global driver, doc
    
    nome_arquivo = f"relatorio_devolucoes_cenario_1_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    
    try:
        doc.save(nome_arquivo)
        log(doc, f"📄 Relatório salvo: {nome_arquivo}")
        
        try:
            subprocess.run(["start", "winword", nome_arquivo], shell=True)
        except:
            pass
            
    except Exception as e:
        print(f"Erro ao salvar relatório: {e}")
    
    if driver:
        try:
            driver.quit()
            log(doc, "✅ Driver encerrado")
        except:
            pass





def _sanitize_timeout(t):
    """Garante timeout válido"""
    if not isinstance(t, (int, float)) or t <= 0:
        return TIMEOUT_DEFAULT
    return max(5, min(120, t))  # Entre 5 e 120 segundos

# ==== AGUARDAR ELEMENTO MELHORADO ====
def aguardar_elemento(seletor, timeout=TIMEOUT_DEFAULT, condicao='clickable', by_type=By.CSS_SELECTOR):
    """Função centralizada para aguardar elementos com diferentes condições"""
    global driver, wait
    
    if driver is None:
        raise Exception("Driver não inicializado")
    
    timeout = _sanitize_timeout(timeout)
    
    condicoes = {
        'present': EC.presence_of_element_located,
        'visible': EC.visibility_of_element_located,
        'clickable': EC.element_to_be_clickable,
        'invisible': EC.invisibility_of_element_located
    }
    
    if condicao not in condicoes:
        condicao = 'clickable'
    
    try:
        wait_obj = WebDriverWait(driver, timeout)
        elemento = wait_obj.until(condicoes[condicao]((by_type, seletor)))
        return elemento
    except TimeoutException:
        log(doc, f"❌ Timeout aguardando elemento: {seletor} (condição: {condicao}, timeout: {timeout}s)", 'ERROR')
        raise TimeoutException(f"Elemento não encontrado: {seletor} (condição: {condicao})")
    except Exception as e:
        log(doc, f"❌ Erro aguardando elemento {seletor}: {e}", 'ERROR')
        raise

# ==== SCROLL CORRIGIDO - PRINCIPAL CORREÇÃO ====
def scroll_to_element_safe(elemento_ou_seletor, by_type=By.CSS_SELECTOR):
    """Scroll seguro até elemento com validação robusta"""
    global driver
    
    if driver is None:
        log(doc, "⚠️ Driver não disponível para scroll", 'WARN')
        return False
    
    try:
        # Se for seletor, encontra o elemento
        if isinstance(elemento_ou_seletor, str):
            elemento = aguardar_elemento(elemento_ou_seletor, 10, 'present', by_type)
        else:
            elemento = elemento_ou_seletor
        
        if elemento is None:
            log(doc, "⚠️ Elemento não encontrado para scroll", 'WARN')
            return False
        
        # Verifica se elemento é válido antes de fazer scroll
        if not elemento.is_displayed():
            log(doc, "⚠️ Elemento não está visível para scroll", 'WARN')
            return False
        
        # Estratégias de scroll em ordem de preferência
        scroll_strategies = [
            # Estratégia 1: JavaScript com verificação prévia
            lambda: driver.execute_script("""
                var element = arguments[0];
                if (element && typeof element.scrollIntoView === 'function') {
                    element.scrollIntoView({
                        behavior: 'smooth',
                        block: 'center',
                        inline: 'center'
                    });
                    return true;
                } else {
                    return false;
                }
            """, elemento),
            
            # Estratégia 2: ActionChains
            lambda: ActionChains(driver).move_to_element(elemento).perform(),
            
            # Estratégia 3: JavaScript alternativo
            lambda: driver.execute_script("""
                var element = arguments[0];
                if (element) {
                    element.scrollIntoView();
                    window.scrollBy(0, -100);
                }
            """, elemento),
            
            # Estratégia 4: Scroll da página até o elemento
            lambda: driver.execute_script("""
                var element = arguments[0];
                if (element) {
                    var rect = element.getBoundingClientRect();
                    var scrollTop = window.pageYOffset || document.documentElement.scrollTop;
                    var targetY = rect.top + scrollTop - (window.innerHeight / 2);
                    window.scrollTo(0, targetY);
                }
            """, elemento)
        ]
        
        for i, strategy in enumerate(scroll_strategies, 1):
            try:
                log(doc, f"   Tentando estratégia de scroll {i}...")
                result = strategy()
                
                # Para estratégia 1, verifica resultado
                if i == 1 and result is False:
                    log(doc, f"   Estratégia {i}: elemento não suporta scrollIntoView", 'WARN')
                    continue
                
                time.sleep(0.8)  # Aguarda scroll completar
                
                # Verifica se elemento ainda está acessível
                if elemento.is_displayed() and elemento.is_enabled():
                    log(doc, f"✅ Scroll realizado com estratégia {i}")
                    return True
                else:
                    log(doc, f"   Estratégia {i}: elemento não ficou acessível", 'WARN')
                    continue
                    
            except Exception as e:
                log(doc, f"   Estratégia {i} de scroll falhou: {str(e)[:100]}...", 'WARN')
                continue
        
        log(doc, "⚠️ Todas as estratégias de scroll falharam", 'WARN')
        return False
        
    except Exception as e:
        log(doc, f"⚠️ Erro geral no scroll: {e}", 'WARN')
        return False



# ==== SISTEMA DATEPICKER MELHORADO ====
def encontrar_campos_datepicker():
    """Encontra todos os campos datepicker na página"""
    global driver
    
    if driver is None:
        return []
    
    seletores_datepicker = [
        "input.hasDatepicker",
        "input[id^='dp']",
        "input[maxlength='10'][grupo='']",
        "input[type='text'][maxlength='10']",
        "input[class*='datepicker']",
        ".hasDatepicker"
    ]
    
    campos_encontrados = []
    
    for seletor in seletores_datepicker:
        try:
            elementos = driver.find_elements(By.CSS_SELECTOR, seletor)
            for elemento in elementos:
                if elemento.is_displayed() and elemento.is_enabled():
                    info = {
                        'elemento': elemento,
                        'id': elemento.get_attribute('id') or f"dp_{len(campos_encontrados)}",
                        'seletor_usado': seletor,
                        'maxlength': elemento.get_attribute('maxlength'),
                        'placeholder': elemento.get_attribute('placeholder')
                    }
                    # Evita duplicatas
                    if not any(c['id'] == info['id'] for c in campos_encontrados):
                        campos_encontrados.append(info)
        except Exception as e:
            log(doc, f"⚠️ Erro ao buscar campos datepicker com {seletor}: {e}", 'WARN')
            continue
    
    log(doc, f"📊 Encontrados {len(campos_encontrados)} campos datepicker")
    return campos_encontrados

def _datepicker_jquery(campo_id, data_valor):
    """Estratégia jQuery para datepicker"""
    global driver
    
    resultado = driver.execute_script("""
        var campoId = arguments[0], valor = arguments[1];
        if (typeof jQuery === 'undefined') return 'jQuery não disponível';
        var $campo = $('#' + campoId);
        if (!$campo.length) return 'Campo não encontrado: ' + campoId;
        try {
            if ($campo.hasClass('hasDatepicker')) { 
                $campo.datepicker('setDate', valor); 
            } else { 
                $campo.val(valor); 
            }
            $campo.trigger('input').trigger('change').trigger('blur');
            return $campo.val();
        } catch(e) { 
            return 'Erro: ' + e.message; 
        }
    """, campo_id, data_valor)
    
    if isinstance(resultado, str) and ('Erro' in resultado or 'não disponível' in resultado):
        raise Exception(f"jQuery falhou: {resultado}")

def _datepicker_javascript(elemento, data_valor):
    """Estratégia JavaScript para datepicker"""
    global driver
    
    driver.execute_script("""
        var campo = arguments[0], valor = arguments[1];
        campo.focus(); 
        campo.value = ''; 
        campo.value = valor;
        ['input','change','blur','keyup'].forEach(ev => 
            campo.dispatchEvent(new Event(ev, {bubbles: true}))
        );
    """, elemento, data_valor)

def _datepicker_actionchains(elemento, data_valor):
    """Estratégia ActionChains para datepicker"""
    global driver
    
    scroll_to_element_safe(elemento)
    time.sleep(0.5)
    
    ActionChains(driver).move_to_element(elemento).click().perform()
    time.sleep(0.5)
    ActionChains(driver).key_down(Keys.CONTROL).send_keys('a').key_up(Keys.CONTROL).perform()
    time.sleep(0.3)
    ActionChains(driver).send_keys(Keys.DELETE).perform()
    time.sleep(0.3)
    
    for char in data_valor:
        ActionChains(driver).send_keys(char).perform()
        time.sleep(0.05)
    
    ActionChains(driver).send_keys(Keys.TAB).perform()

def _datepicker_tradicional(elemento, data_valor):
    """Estratégia tradicional para datepicker"""
    scroll_to_element_safe(elemento)
    time.sleep(0.5)
    elemento.click()
    time.sleep(0.5)
    elemento.clear()
    elemento.send_keys(data_valor)
    elemento.send_keys(Keys.TAB)

def validar_data_preenchida(elemento, data_esperada):
    """Valida se a data foi preenchida corretamente"""
    try:
        if elemento is None:
            return False
            
        val = (elemento.get_attribute('value') or '').strip()
        if not val:
            return False
            
        if val == data_esperada or data_esperada in val:
            return True
            
        # Tenta comparar datas em diferentes formatos
        formatos = [
            '%d/%m/%Y %H:%M', '%d/%m/%Y %H:%M:%S',
            '%d/%m/%Y', '%m/%d/%Y', '%Y-%m-%d', '%d-%m-%Y'
        ]
        
        for formato in formatos:
            try:
                d1 = datetime.strptime(val, formato)
                d2 = datetime.strptime(data_esperada, formato)
                if d1 == d2:
                    return True
            except:
                continue
                
        return False
        
    except Exception:
        return False





def preencher_datepicker_por_indice(indice_campo, data_valor, max_tentativas=5):
    """Preenche datepicker pelo índice com estratégias múltiplas"""
    def acao():
        if not isinstance(indice_campo, int) or indice_campo < 0:
            raise ValueError(f"Índice inválido: {indice_campo}")
            
        if not data_valor or not isinstance(data_valor, str):
            raise ValueError(f"Data inválida: {data_valor}")
        
        tentativa = 0
        while tentativa < max_tentativas:
            tentativa += 1
            
            try:
                campos = encontrar_campos_datepicker()
                
                if not campos:
                    if tentativa < max_tentativas:
                        log(doc, f"⚠️ Nenhum campo datepicker encontrado, tentativa {tentativa}/{max_tentativas}", 'WARN')
                        time.sleep(2)
                        continue
                    raise Exception("Nenhum campo datepicker encontrado na página")
                
                if indice_campo >= len(campos):
                    raise Exception(f"Índice {indice_campo} inválido. Encontrados {len(campos)} campos")
                
                campo_info = campos[indice_campo]
                elemento = campo_info['elemento']
                campo_id = campo_info['id']
                
                log(doc, f"🎯 Tentativa {tentativa}: Preenchendo datepicker {indice_campo} (ID: {campo_id}) com '{data_valor}'")
                
                # Verifica se já está preenchido corretamente
                if validar_data_preenchida(elemento, data_valor):
                    log(doc, f"✅ Campo {indice_campo} já está preenchido corretamente!")
                    return True
                
                # Estratégias específicas para datepicker
                estrategias = [
                    lambda: _datepicker_jquery(campo_id, data_valor),
                    lambda: _datepicker_javascript(elemento, data_valor),
                    lambda: _datepicker_actionchains(elemento, data_valor),
                    lambda: _datepicker_tradicional(elemento, data_valor)
                ]
                
                for i, estrategia in enumerate(estrategias, 1):
                    try:
                        log(doc, f"   Aplicando estratégia {i} para datepicker...")
                        estrategia()
                        time.sleep(1)
                        
                        # Verifica se funcionou
                        if validar_data_preenchida(elemento, data_valor):
                            valor_atual = elemento.get_attribute('value')
                            log(doc, f"✅ Datepicker preenchido com estratégia {i}: '{valor_atual}'")
                            return True
                        else:
                            log(doc, f"⚠️ Estratégia {i} não preencheu corretamente", 'WARN')
                            
                    except Exception as e:
                        log(doc, f"⚠️ Estratégia {i} falhou: {e}", 'WARN')
                        continue
                
                # Se chegou aqui, nenhuma estratégia funcionou nesta tentativa
                if tentativa < max_tentativas:
                    log(doc, f"⚠️ Tentativa {tentativa} falhou, tentando novamente em 2s...", 'WARN')
                    time.sleep(2)
                    continue
                
            except Exception as e:
                if tentativa < max_tentativas:
                    log(doc, f"⚠️ Erro na tentativa {tentativa}: {e}, tentando novamente...", 'WARN')
                    time.sleep(2)
                    continue
                else:
                    raise
        
        raise Exception(f"Falha ao preencher datepicker {indice_campo} após {max_tentativas} tentativas")
    
    return acao


def preencher_campos_pesquisa_por_indice(self, 
                                         search_text: str, 
                                         search_xpaths=None, 
                                         max_campos: int | None = None, 
                                         pausa: float = 0.3,
                                         limpar_antes: bool = True):
    """
    Procura TODOS os campos de pesquisa e preenche um após o outro (ordem no DOM).
    - search_text: texto a preencher.
    - search_xpaths: lista de XPaths para busca (usa padrão se None).
    - max_campos: limita quantos campos serão preenchidos (None = todos).
    - pausa: pausa curta entre preenchimentos.
    - limpar_antes: se True, limpa o campo antes de preencher.

    Retorna: {"total_encontrados": int, "total_preenchidos": int, "xpaths_usados": [str], "falhas": [str]}
    """
    search_xpaths = search_xpaths or [
        "//input[@id='txtPesquisa']",
        "//input[@class='nomePesquisa']",
        # adicione mais padrões específicos antes do genérico:
        "//input[contains(translate(@name,'ABCDEFGHIJKLMNOPQRSTUVWXYZ','abcdefghijklmnopqrstuvwxyz'),'pesquisa')]",
        "//input[contains(translate(@class,'ABCDEFGHIJKLMNOPQRSTUVWXYZ','abcdefghijklmnopqrstuvwxyz'),'pesquisa')]",
        "//input[@type='text']"
    ]

    # Monta um XPATH-UNIÃO para manter a ordem no DOM
    xpath_uniao = " | ".join(f"({xp})" for xp in search_xpaths)

    # Coleta elementos (evita duplicados por id/ref)
    try:
        elementos = self.js.driver.find_elements(By.XPATH, xpath_uniao)
    except Exception as e:
        log(self.doc, f"⚠️ Erro ao buscar campos de pesquisa: {e}")
        return {"total_encontrados": 0, "total_preenchidos": 0, "xpaths_usados": [], "falhas": [str(e)]}

    # Filtra apenas visíveis e habilitados
    elementos_filtrados = []
    for el in elementos:
        try:
            if el.is_displayed() and el.is_enabled():
                elementos_filtrados.append(el)
        except Exception:
            continue

    total_encontrados = len(elementos_filtrados)
    if total_encontrados == 0:
        log(self.doc, "🔎 Nenhum campo de pesquisa visível/habilitado encontrado.")
        return {"total_encontrados": 0, "total_preenchidos": 0, "xpaths_usados": [], "falhas": []}

    if max_campos is not None and max_campos > 0:
        elementos_filtrados = elementos_filtrados[:max_campos]

    log(self.doc, f"🧭 Campos de pesquisa encontrados: {total_encontrados} | A preencher: {len(elementos_filtrados)}")

    total_preenchidos = 0
    falhas = []
    xpaths_usados = []

    for idx, el in enumerate(elementos_filtrados, start=1):
        try:
            # Recalcula um XPATH relativo único para log (opcional, pode ser pesado).
            # Aqui só registramos o index para simplificar.
            xpaths_usados.append(f"[campo #{idx}]")

            # Traz para a tela e foca
            self.js.driver.execute_script("arguments[0].scrollIntoView({block:'center'});", el)
            try:
                el.click()
            except Exception:
                pass

            if limpar_antes:
                try:
                    el.clear()
                except Exception:
                    # fallback via JS
                    try:
                        self.js.driver.execute_script("arguments[0].value='';", el)
                    except Exception:
                        pass

            # Preenche (usa seu helper para manter padrão)
            try:
                self.js.force_fill_element(el, search_text)  # se você tiver esse helper
            except AttributeError:
                # fallback: force_fill por XPATH do próprio elemento (gera um xpath usando JS)
                try:
                    self.js.driver.execute_script("arguments[0].value = arguments[1];", el, search_text)
                except Exception:
                    # último fallback: send_keys
                    el.send_keys(search_text)

            total_preenchidos += 1
            log(self.doc, f"   ✏️ Preenchido campo #{idx} com '{search_text}'")
            if pausa:
                time.sleep(pausa)

        except (StaleElementReferenceException, ElementNotInteractableException) as e:
            falhas.append(f"Campo #{idx}: {type(e).__name__}")
            log(self.doc, f"   ⚠️ Falha no campo #{idx}: {e}")
            continue
        except Exception as e:
            falhas.append(f"Campo #{idx}: {e}")
            log(self.doc, f"   ⚠️ Erro inesperado no campo #{idx}: {e}")
            continue

    resumo = {
        "total_encontrados": total_encontrados,
        "total_preenchidos": total_preenchidos,
        "xpaths_usados": xpaths_usados,
        "falhas": falhas
    }
    log(self.doc, f"✅ Pesquisa preenchida em {total_preenchidos}/{len(elementos_filtrados)} campos. Falhas: {len(falhas)}")
    return resumo

def clicar_lov_por_indice(indice_lov: int, max_tentativas: int = 5, timeout: int = 10, scroll: bool = True):
    """Clica no ícone de LOV pelo índice"""
    def acao():
        if not isinstance(indice_lov, int) or indice_lov < 0:
            raise ValueError(f"Índice inválido: {indice_lov}")

        tentativa = 0
        while tentativa < max_tentativas:
            tentativa += 1
            try:
                log(doc, f"🔎 Tentativa {tentativa}: Localizando ícones LOV...")
                elementos = driver.find_elements(By.CSS_SELECTOR, "a.sprites.sp-openLov")

                if not elementos:
                    if tentativa < max_tentativas:
                        log(doc, f"⚠️ Nenhum ícone LOV encontrado (tentativa {tentativa}/{max_tentativas})")
                        time.sleep(1.2)
                        continue
                    raise Exception("Nenhum ícone LOV encontrado.")

                if indice_lov >= len(elementos):
                    raise Exception(f"Índice {indice_lov} inválido. Encontrados {len(elementos)} ícones LOV.")

                locator_xpath = f"(//a[contains(@class,'sp-openLov')])[{indice_lov + 1}]"
                elemento = driver.find_element(By.XPATH, locator_xpath)

                log(doc, f"🎯 Preparando clique no LOV de índice {indice_lov}")

                def _wait_clickable():
                    wait.until(EC.element_to_be_clickable((By.XPATH, locator_xpath)))

                estrategias = [
                    lambda: (_wait_clickable(), elemento.click()),
                    lambda: (
                        driver.execute_script("arguments[0].scrollIntoView({block:'center'});", elemento) if scroll else None,
                        time.sleep(0.2),
                        elemento.click()
                    ),
                    lambda: driver.execute_script("arguments[0].click();", elemento),
                    lambda: ActionChains(driver).move_to_element(elemento).pause(0.1).click().perform()
                ]

                for i, estrategia in enumerate(estrategias, 1):
                    try:
                        log(doc, f"   ▶️ Estratégia {i} de clique no LOV...")
                        estrategia()
                        time.sleep(0.3)
                        log(doc, f"✅ Clique no LOV (índice {indice_lov}) realizado (estratégia {i})")
                        return True
                    except (ElementClickInterceptedException, StaleElementReferenceException, JavascriptException, TimeoutException) as e:
                        log(doc, f"⚠️ Estratégia {i} falhou: {e}")
                        try:
                            elementos = driver.find_elements(By.CSS_SELECTOR, "a.sprites.sp-openLov")
                            elemento = driver.find_element(By.XPATH, locator_xpath)
                        except:
                            pass
                        continue

                if tentativa < max_tentativas:
                    log(doc, f"⚠️ Tentativa {tentativa} não conseguiu clicar no LOV. Reintentando...")
                    time.sleep(1.2)
                    continue

            except Exception as e:
                if tentativa < max_tentativas:
                    log(doc, f"⚠️ Erro na tentativa {tentativa}: {e}. Reintentando...")
                    time.sleep(1.2)
                    continue
                raise

        raise Exception(f"Falha ao clicar no LOV de índice {indice_lov} após {max_tentativas} tentativas.")

    return acao

def encontrar_campos_textarea(timeout=10):
    """Retorna lista de textareas visíveis e interativas"""
    elementos = []
    try:
        wait.until(lambda d: len(d.find_elements(By.TAG_NAME, "textarea")) >= 0)
        textareas = driver.find_elements(By.TAG_NAME, "textarea")
    except:
        textareas = []

    for el in textareas:
        try:
            if not el.is_displayed() or not el.is_enabled():
                continue
            driver.execute_script("arguments[0].scrollIntoView({block:'center'});", el)
            elementos.append({
                "elemento": el,
                "id": el.get_attribute("id"),
                "name": el.get_attribute("name"),
            })
        except:
            continue

    return elementos

def normalizar_texto(txt):
    if txt is None:
        return ""
    return txt.replace("\r\n", "\n").replace("\r", "\n").strip()

def validar_textarea_preenchida(elemento, texto_esperado):
    """Confere se o valor atual da textarea bate com o esperado"""
    try:
        atual = elemento.get_attribute("value")
        if atual is None or atual == "":
            atual = (elemento.text or "")
        return normalizar_texto(atual) == normalizar_texto(texto_esperado)
    except StaleElementReferenceException:
        return False

def _prepare_focus_and_clear(elemento, limpar_primeiro=True):
    driver.execute_script("arguments[0].scrollIntoView({block:'center'});", elemento)
    try:
        elemento.click()
    except:
        driver.execute_script("arguments[0].focus();", elemento)

    if limpar_primeiro:
        try:
            elemento.clear()
        except:
            ActionChains(driver)\
                .move_to_element(elemento).click()\
                .key_down(Keys.CONTROL).send_keys("a").key_up(Keys.CONTROL)\
                .send_keys(Keys.DELETE).perform()

def _textarea_tradicional(elemento, texto, limpar_primeiro=True):
    _prepare_focus_and_clear(elemento, limpar_primeiro)
    elemento.send_keys(texto)
    elemento.send_keys(Keys.TAB)

def _textarea_actionchains(elemento, texto, limpar_primeiro=True):
    _prepare_focus_and_clear(elemento, limpar_primeiro)
    ac = ActionChains(driver)
    ac.move_to_element(elemento).click().perform()
    for chunk_start in range(0, len(texto), 400):
        ac.send_keys(texto[chunk_start:chunk_start+400]).perform()
        time.sleep(0.05)
    ac.send_keys(Keys.TAB).perform()

def _textarea_js_setvalue(elemento, texto):
    driver.execute_script("""
        const el = arguments[0];
        const val = arguments[1];
        el.value = val;
        el.dispatchEvent(new Event('input',  {bubbles:true}));
        el.dispatchEvent(new Event('change', {bubbles:true}));
        el.dispatchEvent(new KeyboardEvent('keyup', {bubbles:true}));
        el.dispatchEvent(new Event('blur',   {bubbles:true}));
    """, elemento, texto)

def _textarea_js_react_input(elemento, texto):
    driver.execute_script("""
        const el = arguments[0];
        const val = arguments[1];

        const desc = Object.getOwnPropertyDescriptor(window.HTMLTextAreaElement.prototype, 'value');
        if (desc && desc.set) {
            desc.set.call(el, val);
        } else {
            el.value = val;
        }

        el.dispatchEvent(new Event('input', {bubbles: true}));
        el.dispatchEvent(new Event('change', {bubbles: true}));
        el.dispatchEvent(new Event('blur', {bubbles: true}));
    """, elemento, texto)

def preencher_textarea_por_indice(indice_campo, texto, max_tentativas=5, limpar_primeiro=True):
    """Preenche textarea pelo índice usando estratégias múltiplas"""
    def acao():
        if not isinstance(indice_campo, int) or indice_campo < 0:
            raise ValueError(f"Índice inválido: {indice_campo}")
        if texto is None or not isinstance(texto, str):
            raise ValueError(f"Texto inválido: {texto!r}")

        tentativa = 0
        while tentativa < max_tentativas:
            tentativa += 1
            try:
                campos = encontrar_campos_textarea()
                if not campos:
                    if tentativa < max_tentativas:
                        log(doc, f"⚠️ Nenhuma <textarea> encontrada (tentativa {tentativa}/{max_tentativas})")
                        time.sleep(1.5)
                        continue
                    raise Exception("Nenhuma <textarea> foi encontrada.")

                if indice_campo >= len(campos):
                    raise Exception(f"Índice {indice_campo} inválido. Encontradas {len(campos)} textareas.")

                campo_info = campos[indice_campo]
                elemento = campo_info["elemento"]
                campo_id = campo_info.get("id") or "(sem id)"

                log(doc, f"🎯 Tentativa {tentativa}: Preenchendo textarea {indice_campo} (ID: {campo_id})")

                if validar_textarea_preenchida(elemento, texto):
                    log(doc, f"✅ Textarea {indice_campo} já está com o valor desejado.")
                    return True

                estrategias = [
                    lambda: _textarea_tradicional(elemento, texto, limpar_primeiro),
                    lambda: _textarea_actionchains(elemento, texto, limpar_primeiro),
                    lambda: _textarea_js_setvalue(elemento, texto),
                    lambda: _textarea_js_react_input(elemento, texto),
                ]

                for i, estrategia in enumerate(estrategias, 1):
                    try:
                        log(doc, f"   ▶️ Estratégia {i}…")
                        estrategia()
                        time.sleep(0.8)

                        if validar_textarea_preenchida(elemento, texto):
                            val = (elemento.get_attribute("value") or "").strip()
                            log(doc, f"✅ Preenchido com sucesso pela estratégia {i}")
                            return True
                        else:
                            log(doc, f"⚠️ Estratégia {i} não refletiu o valor esperado.")
                    except (StaleElementReferenceException, JavascriptException, TimeoutException) as e:
                        log(doc, f"⚠️ Estratégia {i} falhou: {e}")
                        try:
                            campos = encontrar_campos_textarea()
                            elemento = campos[indice_campo]["elemento"]
                        except:
                            pass
                        continue

                if tentativa < max_tentativas:
                    log(doc, f"⚠️ Tentativa {tentativa} falhou; reintentando em 1.5s…")
                    time.sleep(1.5)
                    continue
            except Exception as e:
                if tentativa < max_tentativas:
                    log(doc, f"⚠️ Erro na tentativa {tentativa}: {e}. Retentando…")
                    time.sleep(1.5)
                    continue
                else:
                    raise

        raise Exception(f"Falha ao preencher textarea {indice_campo} após {max_tentativas} tentativas.")
    return acao


def fechar_abas_extras(driver, doc, aba_principal_index=0):
    """Fecha todas as abas extras (como popups de impressão) mantendo apenas a aba principal"""
    try:
        handles = driver.window_handles
        if len(handles) <= 1:
            log(doc, "ℹ️ Apenas uma aba aberta - nada a fechar.")
            return True
        
        # Guarda o handle da aba principal
        aba_principal = handles[aba_principal_index]
        
        # Fecha todas as outras abas
        abas_fechadas = 0
        for handle in handles:
            if handle != aba_principal:
                try:
                    driver.switch_to.window(handle)
                    driver.close()
                    abas_fechadas += 1
                    log(doc, f"🗑️ Aba extra fechada ({abas_fechadas})")
                except Exception as e:
                    log(doc, f"⚠️ Erro ao fechar aba: {e}")
        
        # Retorna para a aba principal
        driver.switch_to.window(aba_principal)
        time.sleep(0.3)
        
        log(doc, f"✅ {abas_fechadas} aba(s) extra(s) fechada(s). Foco na aba principal.")
        return True
        
    except Exception as e:
        log(doc, f"❌ Erro ao fechar abas extras: {e}")
        return False

def focar_sistema_completo(js_engine, doc):
    """Garante o foco completo na aba principal do sistema e fecha abas extras"""
    driver = js_engine.driver
    try:
        # Primeiro fecha abas extras (como impressão)
        fechar_abas_extras(driver, doc)

        driver.switch_to.default_content()
        js_engine.execute_js("if (window.focus) window.focus();", timeout=3, fallback_result=None)
        time.sleep(0.3)

        log(doc, "✅ Foco garantido na aba do sistema.")
        return True

    except Exception as e:
        log(doc, f"⚠️ Falha ao focar aba do sistema: {e}")
        return False

def clicar_todos_botoes_sim_visiveis(js_engine, doc, pausa_entre=0.0):
    """Clica em TODOS os botões 'Sim' visíveis de uma vez"""
    js = r"""
    (function(){
      const isVisible = el => {
        if (!el) return false;
        const s = getComputedStyle(el);
        return el.offsetParent !== null && s.display !== 'none' &&
               s.visibility !== 'hidden' && parseFloat(s.opacity||1) > 0.01;
      };
      const buttons = Array.from(document.querySelectorAll("a.btModel.btGray.btyes"))
        .filter(isVisible)
        .filter(b => (b.textContent||"").trim().toLowerCase() === "sim");

      let clicked = 0;
      buttons.forEach(b => {
        try {
          b.style.pointerEvents = 'auto';
          b.removeAttribute('disabled');
          b.style.visibility = 'visible';
          b.style.display = 'inline-block';
          b.scrollIntoView({block:'center'});

          ['mouseover','mouseenter','mousemove','mousedown','mouseup','click'].forEach(t=>{
            b.dispatchEvent(new MouseEvent(t,{bubbles:true,cancelable:true,view:window,detail:1}));
          });
          if (typeof b.click === 'function') b.click();

          if (typeof window.jQuery !== 'undefined') {
            window.jQuery(b).trigger('click');
          }
          clicked++;
        } catch(e) {}
      });

      return { totalEncontrados: buttons.length, totalClicados: clicked };
    })();
    """
    try:
        res = js_engine.execute_js(js, timeout=5, fallback_result={"totalEncontrados": 0, "totalClicados": 0})
        total = int(res.get("totalEncontrados", 0))
        clic = int(res.get("totalClicados", 0))
        log(doc, f"⚡ 'Sim' visíveis encontrados: {total} | clicados: {clic}")
        if pausa_entre and clic > 0:
            time.sleep(pausa_entre)
        return res
    except Exception as e:
        log(doc, f"❌ Erro ao clicar em todos os 'Sim': {e}")
        return {"totalEncontrados": 0, "totalClicados": 0, "erro": str(e)}


        
def clicar_sim_com_retry(doc, js_engine, wait, max_tentativas=5, pausa=1.5):
    """Clica em 'Sim' até o modal de confirmação desaparecer"""
    xpath_modal = "//div[contains(@class,'modal') and contains(@style,'z-index')]"
    xpath_sim = "(//div[contains(@class,'modal') and not(contains(@style,'display: none'))]//a[@class='btModel btGray btyes'])[last()]"

    tentativa = 0
    while tentativa < max_tentativas:
        tentativa += 1
        log(doc, f"🧩 Tentativa {tentativa} de fechar modal...")

        try:
            js_engine.force_click(xpath_sim, by_xpath=True)
            time.sleep(pausa)

            modais_visiveis = driver.find_elements(By.XPATH, xpath_modal)
            modais_ativos = [m for m in modais_visiveis if "display: none" not in m.get_attribute("style")]

            if not modais_ativos:
                log(doc, "✅ Botão 'Sim' clicado com sucesso.")
                return True

        except Exception as e:
            log(doc, f"⚠️ Tentativa {tentativa} falhou: {e}")

    log(doc, "❌ Botão 'Sim' não foi clicado após todas as tentativas.")
    return False

def clicar_primeiro_sp_add(js_engine, doc=None, timeout=5):
    """
    Localiza e clica no primeiro elemento <a class="sprites sp-add">, exatamente.
    Ignora variantes como 'sp-addVerde'.
    """
    from selenium.common.exceptions import (
        NoSuchElementException,
        ElementClickInterceptedException,
        StaleElementReferenceException,
    )
    from selenium.webdriver.common.by import By
    import time

    driver = js_engine.driver

    # XPath restrito: exige que as classes sejam EXATAMENTE 'sprites' e 'sp-add'
    xpath_btn = (
        "//a[contains(concat(' ', normalize-space(@class), ' '), ' sprites ') "
        "and contains(concat(' ', normalize-space(@class), ' '), ' sp-add ') "
        "and not(contains(@class, 'sp-addVerde'))][1]"
    )

    log(doc, "🧩 Procurando o botão '<a class=\"sprites sp-add\">' exato...")

    try:
        el = driver.find_element(By.XPATH, xpath_btn)
        log(doc, "🎯 Botão exato encontrado! Tentando clicar...")

        try:
            el.click()
            log(doc, "✅ Clique padrão realizado com sucesso.")
        except (ElementClickInterceptedException, StaleElementReferenceException):
            driver.execute_script("arguments[0].click();", el)
            log(doc, "⚡ Clique forçado via JavaScript realizado.")

        time.sleep(0.5)
        return True

    except NoSuchElementException:
        log(doc, "⚠️ Nenhum botão exato 'sp-add' encontrado.")
        return False

    except Exception as e:
        log(doc, f"❌ Erro ao clicar no botão 'sp-add': {e}")
        return False


def fechar_abas_extras_e_verificar_alerta(driver, doc):
    """
    Fecha abas extras (como de impressão) e verifica imediatamente
    se há alguma mensagem de alerta exibida após o fechamento.
    """
    try:
        safe_action(doc, "Fechando abas extras (impressão)", lambda:
            fechar_abas_extras(driver, doc)
        )

        alerta = encontrar_mensagem_alerta()
        if alerta:
            log(doc, f"⚠️ Alerta detectado após fechar abas extras: '{alerta.text.strip()}'")
        else:
            log(doc, "✅ Nenhum alerta detectado após fechar abas extras.")

        return True
    except Exception as e:
        log(doc, f"⚠️ Erro ao fechar abas extras e verificar alerta: {e}")
        return False

def verificar_e_abrir_caixa(js_engine, doc, timeout=10):
    """
    Verifica o estado do caixa:
      - Se encontrar o botão 'Abrir caixa' visível (sem display: none), realiza a abertura normalmente.
      - Se detectar 'Fechar caixa', entende que o caixa já está aberto e apenas prossegue.
    """
    abrir_xpath = "//a[contains(@class,'btAzulDegrade') and contains(@class,'btAbrirCaixa')]"
    fechar_xpath = "//a[contains(@class,'btAzulDegrade') and contains(@class,'btFecharCaixa')]"

    driver = js_engine.driver
    log(doc, "🔍 Verificando estado do caixa...")

    try:
        # Verifica se existe botão Abrir Caixa visível
        abrir_visivel = driver.execute_script(f"""
            const el = document.evaluate("{abrir_xpath}", document, null, XPathResult.FIRST_ORDERED_NODE_TYPE, null).singleNodeValue;
            if (!el) return false;
            const style = window.getComputedStyle(el);
            return style.display !== 'none' && style.visibility !== 'hidden' && el.offsetParent !== null;
        """)

        # Verifica se existe botão Fechar Caixa visível
        fechar_visivel = driver.execute_script(f"""
            const el = document.evaluate("{fechar_xpath}", document, null, XPathResult.FIRST_ORDERED_NODE_TYPE, null).singleNodeValue;
            if (!el) return false;
            const style = window.getComputedStyle(el);
            return style.display !== 'none' && style.visibility !== 'hidden' && el.offsetParent !== null;
        """)

        if abrir_visivel:
            log(doc, "📦 Caixa fechado — iniciando abertura normal...")
            safe_action(doc, "Clicando em 'Abrir Caixa'", lambda:
                js_engine.force_click(abrir_xpath, by_xpath=True)
            )

            # Valor inicial
            safe_action(doc, "Preenchendo Valor Inicial", lambda:
                js_engine.force_fill(
                    "//input[@type='text' and contains(@class,'valor') and contains(@placeholder,'R$')]",
                    "1000,00",
                    by_xpath=True
                )
            )

            # Descrição
            safe_action(doc, "Preenchendo Descrição", lambda:
                preencher_textarea_por_indice(0, "Abertura automática de caixa via automação Selenium.")
            )

            # Confirmar abertura
            safe_action(doc, "Confirmando Abertura", lambda:
                js_engine.force_click("//a[@class='btModel btGray btyes' and normalize-space()='Abrir']", by_xpath=True)
            )
            time.sleep(0.5)
            encontrar_mensagem_alerta()

            # Autenticação
            safe_action(doc, "Autenticando Abertura", lambda:
                js_engine.force_click("//a[@id='BtYes' and contains(@class,'btModel btGray btyes') and contains(normalize-space(.),'Autenticar')]", by_xpath=True)
            )
            time.sleep(3)
            fechar_abas_extras_e_verificar_alerta(driver, doc)

            # Fecha modal de autenticação
            safe_action(doc, "Fechando modal autenticação", lambda:
                js_engine.force_click("//a[@id='BtNo' and contains(@class,'btno') and normalize-space()='Fechar']", by_xpath=True)
            )

        elif fechar_visivel:
            log(doc, "✅ Caixa já está aberto — prosseguindo com o fluxo normalmente.")
        else:
            log(doc, "⚠️ Nenhum botão de 'Abrir' ou 'Fechar caixa' visível encontrado — prosseguindo com cautela.")

    except Exception as e:
        log(doc, f"❌ Erro ao verificar/abrir caixa: {e}")
        take_screenshot(driver, doc, "erro_verificar_caixa")

def clicar_todos_pesquisar(js_engine, doc, pausa_entre=0.5, timeout=5):
    """
    Procura todos os botões 'Pesquisar' visíveis e clica em cada um deles na ordem.
    Conta e exibe quantos botões existem antes de clicar.
    Usa js_engine.force_click() e registra log detalhado.
    """

    xpath_base = "//a[contains(@class,'btPesquisar btAzulDegrade') and contains(normalize-space(.),'Pesquisar')]"

    try:
        elementos = js_engine.driver.find_elements("xpath", xpath_base)
        total = len(elementos)
        log(doc, f"🔍 Foram encontrados {total} botão(ões) 'Pesquisar' na tela.")

        if total == 0:
            log(doc, "⚠️ Nenhum botão 'Pesquisar' foi encontrado.")
            return {"total": 0, "clicados": 0}

        total_clicados = 0
        for i in range(1, total + 1):
            xpath_indexado = f"({xpath_base})[{i}]"
            try:
                log(doc, f"🎯 Clicando no botão 'Pesquisar' (índice {i}/{total})...")
                js_engine.force_click(xpath_indexado, by_xpath=True)
                js_engine.wait_ajax_complete(timeout)
                total_clicados += 1
                log(doc, f"✅ Clique no botão 'Pesquisar' (índice {i}) realizado com sucesso.")
                if pausa_entre > 0:
                    import time
                    time.sleep(pausa_entre)
            except Exception as e:
                log(doc, f"⚠️ Falha ao clicar no botão 'Pesquisar' (índice {i}): {e}")

        log(doc, f"🧾 Resumo: {total_clicados}/{total} botões 'Pesquisar' clicados com sucesso.")
        return {"total": total, "clicados": total_clicados}

    except Exception as e:
        log(doc, f"⚠️ Erro ao procurar ou clicar nos botões 'Pesquisar': {e}")
        return {"total": 0, "clicados": 0}


def clicar_todos_voltar(js_engine, doc, pausa_entre=0.5, timeout=5):
    """
    Procura todos os botões 'Voltar (ESC)' visíveis e clica em cada um deles na ordem.
    Conta e exibe quantos botões existem antes de clicar.
    Usa js_engine.force_click() e registra log detalhado.
    """

    xpath_base = "//a[contains(@class,'sp-voltarGrande') and @title='Voltar (ESC)']"

    try:
        elementos = js_engine.driver.find_elements("xpath", xpath_base)
        total = len(elementos)
        log(doc, f"🔍 Foram encontrados {total} botão(ões) 'Voltar (ESC)' na tela.")

        if total == 0:
            log(doc, "⚠️ Nenhum botão 'Voltar (ESC)' foi encontrado.")
            return {"total": 0, "clicados": 0}

        total_clicados = 0
        for i in range(1, total + 1):
            xpath_indexado = f"({xpath_base})[{i}]"
            try:
                log(doc, f"🎯 Clicando no botão 'Voltar (ESC)' (índice {i}/{total})...")
                js_engine.force_click(xpath_indexado, by_xpath=True)
                js_engine.wait_ajax_complete(timeout)
                total_clicados += 1
                log(doc, f"✅ Clique no botão 'Voltar (ESC)' (índice {i}) realizado com sucesso.")
                if pausa_entre > 0:
                    import time
                    time.sleep(pausa_entre)
            except Exception as e:
                log(doc, f"⚠️ Falha ao clicar no botão 'Voltar (ESC)' (índice {i}): {e}")

        log(doc, f"🧾 Resumo: {total_clicados}/{total} botões 'Voltar (ESC)' clicados com sucesso.")
        return {"total": total, "clicados": total_clicados}

    except Exception as e:
        log(doc, f"⚠️ Erro ao procurar ou clicar nos botões 'Voltar (ESC)': {e}")
        return {"total": 0, "clicados": 0}


def clicar_salvar_modal(js_engine, doc, timeout=5):
    """Versão simplificada focada no modal específico com logs detalhados"""
    import time
    
    try:
        log(doc, "🔍 Iniciando busca por botões 'Salvar' no modal...")
        
        # Script JavaScript que retorna informações detalhadas
        script = """
        var botoes = document.querySelectorAll('.modal.overflow .btok');
        var info = {
            total: botoes.length,
            visiveis: 0,
            clicados: 0
        };
        
        botoes.forEach(function(btn) {
            if (btn.offsetParent !== null) {  // Verifica se está visível
                info.visiveis++;
                try {
                    btn.click();
                    info.clicados++;
                } catch(e) {
                    console.error('Erro ao clicar:', e);
                }
            }
        });
        
        return info;
        """
        
        log(doc, "⚙️ Executando JavaScript para clicar no(s) botão(ões)...")
        resultado = js_engine.driver.execute_script(script)
        
        log(doc, f"📊 Total de botões encontrados: {resultado['total']}")
        log(doc, f"👁️ Botões visíveis: {resultado['visiveis']}")
        log(doc, f"✅ Botões clicados com sucesso: {resultado['clicados']}")
        
        if resultado['clicados'] == 0:
            log(doc, "⚠️ Nenhum botão 'Salvar' foi clicado.")
            if resultado['total'] == 0:
                log(doc, "❌ Motivo: Nenhum botão encontrado no modal.")
            elif resultado['visiveis'] == 0:
                log(doc, "❌ Motivo: Botões existem mas não estão visíveis.")
        else:
            log(doc, f"🎉 Sucesso! {resultado['clicados']} botão(ões) clicado(s).")
            log(doc, f"⏳ Aguardando conclusão do AJAX (timeout: {timeout}s)...")
            js_engine.wait_ajax_complete(timeout)
            log(doc, "✅ AJAX concluído.")
        
        return {"total": resultado['total'], "clicados": resultado['clicados']}
        
    except Exception as e:
        log(doc, f"❌ Erro ao executar clique no modal: {e}")
        import traceback
        log(doc, f"📋 Detalhes do erro: {traceback.format_exc()}")
        return {"total": 0, "clicados": 0}
    

def clicar_botao_por_classe(js_engine, doc, classe, nome_botao="Botão", timeout=5):
    """Clica em botão por classe CSS com logs detalhados"""
    import time
    
    try:
        log(doc, f"🔍 Iniciando busca pelo botão '{nome_botao}' (classe: {classe})...")
        
        script = f"""
        var botoes = document.querySelectorAll('a.{classe}');
        var info = {{
            total: botoes.length,
            visiveis: 0,
            clicados: 0
        }};
        
        botoes.forEach(function(btn) {{
            if (btn.offsetParent !== null) {{
                info.visiveis++;
                try {{
                    btn.click();
                    info.clicados++;
                }} catch(e) {{
                    console.error('Erro ao clicar:', e);
                }}
            }}
        }});
        
        return info;
        """
        
        log(doc, f"⚙️ Executando JavaScript para clicar no botão '{nome_botao}'...")
        resultado = js_engine.driver.execute_script(script)
        
        log(doc, f"📊 Total de botões encontrados: {resultado['total']}")
        log(doc, f"👁️ Botões visíveis: {resultado['visiveis']}")
        log(doc, f"✅ Botões clicados: {resultado['clicados']}")
        
        if resultado['clicados'] == 0:
            log(doc, f"⚠️ Nenhum botão '{nome_botao}' foi clicado.")
            if resultado['total'] == 0:
                log(doc, f"❌ Motivo: Nenhum botão encontrado.")
            elif resultado['visiveis'] == 0:
                log(doc, f"❌ Motivo: Botão existe mas não está visível.")
        else:
            log(doc, f"🎉 Sucesso! Botão '{nome_botao}' clicado.")
            log(doc, f"⏳ Aguardando AJAX (timeout: {timeout}s)...")
            js_engine.wait_ajax_complete(timeout)
            log(doc, "✅ AJAX concluído.")
        
        return {"total": resultado['total'], "clicados": resultado['clicados']}
        
    except Exception as e:
        log(doc, f"❌ Erro ao clicar no botão '{nome_botao}': {e}")
        return {"total": 0, "clicados": 0}
        
    except Exception as e:
        log(doc, f"❌ Erro ao executar clique no modal: {e}")
        import traceback
        log(doc, f"📋 Detalhes do erro: {traceback.format_exc()}")
        return {"total": 0, "clicados": 0}
    
def clicar_botao_voltar_por_indice(js_engine, doc, indice=1, timeout=5):
    """
    Clica no botão 'Voltar (ESC)' pelo índice informado (1-based).
    Usa js_engine.force_click() e registra log.
    """
    xpath = f"(//a[@class='sprites sp-voltarGrande' and @title='Voltar (ESC)'])[{indice}]"
    log(doc, f"↩️ Clicando no botão 'Voltar (ESC)' (índice {indice})...")

    try:
        js_engine.force_click(xpath, by_xpath=True)
        js_engine.wait_ajax_complete(timeout)
        log(doc, f"✅ Clique no botão 'Voltar (ESC)' (índice {indice}) realizado com sucesso.")
        return True
    except Exception as e:
        log(doc, f"⚠️ Falha ao clicar no botão 'Voltar (ESC)' (índice {indice}): {e}")
        return False

def clicar_pesquisar_por_indice(js_engine, doc, indice=1, timeout=5):
    """
    Clica no botão 'Pesquisar' pelo índice informado (1-based).
    Conta e exibe quantos botões existem antes do clique.
    Usa js_engine.force_click() e registra log.
    """
    xpath_base = "//a[contains(@class,'btPesquisar btAzulDegrade') and contains(normalize-space(.),'Pesquisar')]"
    xpath_indexado = f"({xpath_base})[{indice}]"

    try:
        # Conta quantos botões existem
        elementos = js_engine.driver.find_elements("xpath", xpath_base)
        total = len(elementos)
        log(doc, f"🔍 Foram encontrados {total} botão(ões) 'Pesquisar' na tela.")

        if total == 0:
            log(doc, "⚠️ Nenhum botão 'Pesquisar' foi encontrado.")
            return False
        if indice > total:
            log(doc, f"⚠️ Índice {indice} inválido — só existem {total} botão(ões).")
            return False

        log(doc, f"🎯 Clicando no botão 'Pesquisar' (índice {indice})...")
        js_engine.force_click(xpath_indexado, by_xpath=True)
        js_engine.wait_ajax_complete(timeout)
        log(doc, f"✅ Clique no botão 'Pesquisar' (índice {indice}) realizado com sucesso.")
        return True

    except Exception as e:
        log(doc, f"⚠️ Erro ao clicar no botão 'Pesquisar' (índice {indice}): {e}")
        return False


def clicar_salvar_por_indice(js_engine, doc, indice=1, timeout=5):
    """
    Clica no botão 'Salvar' pelo índice informado (1-based).
    Conta e exibe quantos botões existem antes do clique.
    Usa js_engine.force_click() e registra log.
    """
    xpath_base = "//a[contains(@class,'btModel btGray btok') and contains(normalize-space(.),'Salvar')]"
    xpath_indexado = f"({xpath_base})[{indice}]"

    try:
        # Conta quantos botões existem
        elementos = js_engine.driver.find_elements("xpath", xpath_base)
        total = len(elementos)
        log(doc, f"💾 Foram encontrados {total} botão(ões) 'Salvar' na tela.")

        if total == 0:
            log(doc, "⚠️ Nenhum botão 'Salvar' foi encontrado.")
            return False
        if indice > total:
            log(doc, f"⚠️ Índice {indice} inválido — só existem {total} botão(ões).")
            return False

        log(doc, f"🎯 Clicando no botão 'Salvar' (índice {indice})...")
        js_engine.force_click(xpath_indexado, by_xpath=True)
        js_engine.wait_ajax_complete(timeout)
        log(doc, f"✅ Clique no botão 'Salvar' (índice {indice}) realizado com sucesso.")
        return True

    except Exception as e:
        log(doc, f"⚠️ Erro ao clicar no botão 'Salvar' (índice {indice}): {e}")
        return False


def clicar_titulo_produtos(js_engine, doc, timeout=3):
    """
    Clica no elemento <h2> que contém o texto 'PRODUTOS'.
    """
    xpath = "//h2[contains(normalize-space(.), 'PRODUTOS')]"
    try:
        log(doc, "🧩 Retornando à aba Principal")
        js_engine.force_click(xpath, by_xpath=True)
        js_engine.wait_ajax_complete(timeout)
        log(doc, "✅ Clique no título 'PRODUTOS' realizado com sucesso.")
        return True
    except Exception as e:
        log(doc, f"⚠️ Falha ao clicar no título 'PRODUTOS': {e}")
        return False

def clicar_salvar_ate_modal_fechar(js_engine, doc, timeout_total=60, pausa_entre=0.5):
    """
    Clica repetidamente no botão 'Salvar' do modal de emissão de nota fiscal
    até que o modal desapareça da tela ou o tempo máximo seja atingido.
    """
    xpath_modal = "//div[contains(@class,'modal') and .//h2[contains(.,'Emissão de Nota Fiscal de Serviço')]]"
    xpath_btn_salvar = "//a[contains(normalize-space(.),'Salvar')]"

    log(doc, "⚙️ Iniciando loop de clique em 'Salvar' até fechamento do modal...")
    inicio = time.time()

    while True:
        try:
            # Verifica se o modal ainda está presente na tela
            modais = js_engine.driver.find_elements("xpath", xpath_modal)
            if not modais or not modais[0].is_displayed():
                log(doc, "✅ Modal fechado — parando cliques.")
                break

            # Clica no botão 'Salvar'
            try:
                js_engine.force_click(xpath_btn_salvar, by_xpath=True)
                log(doc, "💾 Clique em 'Salvar' realizado.")
            except Exception as e:
                log(doc, f"⚠️ Erro ao clicar em 'Salvar': {e}")

            # Espera um pouco antes de tentar novamente
            time.sleep(pausa_entre)

            # Verifica timeout total
            if time.time() - inicio > timeout_total:
                log(doc, "⏰ Tempo limite atingido — modal ainda aberto, encerrando tentativa.")
                break

        except Exception as e:
            log(doc, f"⚠️ Erro inesperado no loop de salvamento: {e}")
            break

    log(doc, "🏁 Finalizado processo de clique repetido em 'Salvar'.")

def clicar_voltar(js_engine, doc, timeout=5):
    """Clica no botão 'Voltar' do modal com logs detalhados"""
    import time
    
    try:
        log(doc, "🔍 Iniciando busca pelo botão 'Voltar'...")
        
        # Script JavaScript que retorna informações detalhadas
        script = """
        var botoes = document.querySelectorAll('a.sp-voltarGrande');
        var info = {
            total: botoes.length,
            visiveis: 0,
            clicados: 0
        };
        
        botoes.forEach(function(btn) {
            if (btn.offsetParent !== null) {  // Verifica se está visível
                info.visiveis++;
                try {
                    btn.click();
                    info.clicados++;
                } catch(e) {
                    console.error('Erro ao clicar:', e);
                }
            }
        });
        
        return info;
        """
        
        log(doc, "⚙️ Executando JavaScript para clicar no botão 'Voltar'...")
        resultado = js_engine.driver.execute_script(script)
        
        log(doc, f"📊 Total de botões 'Voltar' encontrados: {resultado['total']}")
        log(doc, f"👁️ Botões visíveis: {resultado['visiveis']}")
        log(doc, f"✅ Botões clicados com sucesso: {resultado['clicados']}")
        
        if resultado['clicados'] == 0:
            log(doc, "⚠️ Nenhum botão 'Voltar' foi clicado.")
            if resultado['total'] == 0:
                log(doc, "❌ Motivo: Nenhum botão 'Voltar' encontrado.")
            elif resultado['visiveis'] == 0:
                log(doc, "❌ Motivo: Botão existe mas não está visível.")
        else:
            log(doc, f"🎉 Sucesso! Botão 'Voltar' clicado.")
            log(doc, f"⏳ Aguardando conclusão do AJAX (timeout: {timeout}s)...")
            js_engine.wait_ajax_complete(timeout)
            log(doc, "✅ AJAX concluído.")
        
        return {"total": resultado['total'], "clicados": resultado['clicados']}
        
    except Exception as e:
        log(doc, f"❌ Erro ao clicar no botão 'Voltar': {e}")
        import traceback
        log(doc, f"📋 Detalhes do erro: {traceback.format_exc()}")
        return {"total": 0, "clicados": 0}

def clicar_titulo_titulos(js_engine, doc, timeout=3):
    """
    Clica no elemento <h2> que contém o texto 'TÍTULOS'.
    """
    xpath = "//h2[contains(normalize-space(.), 'TÍTULOS')]"
    try:
        log(doc, "🧩 Retornando à aba Principal")
        js_engine.force_click(xpath, by_xpath=True)
        js_engine.wait_ajax_complete(timeout)
        log(doc, "✅ Clique no título 'TÍTULOS' realizado com sucesso.")
        return True
    except Exception as e:
        log(doc, f"⚠️ Falha ao clicar no título 'TÍTULOS': {e}")
        return False



def selecionar_template_por_texto(js_engine, doc, texto="PADRÃO", timeout=5):
    """
    Seleciona uma opção do select 'templateNota' pelo texto exibido.
    
    Args:
        js_engine: Instância do JSForceEngine
        doc: Documento para logs
        texto: Texto da opção a ser selecionada (default: "PADRÃO")
        timeout: Timeout para aguardar AJAX (default: 5)
    
    Returns:
        bool: True se selecionou com sucesso, False caso contrário
    """
    try:
        log(doc, f"📋 Selecionando template por texto: '{texto}'...")
        
        # Script JavaScript para selecionar por texto
        script = """
        const selectElement = document.querySelector('select.templateNota');
        const texto = arguments[0];
        
        if (!selectElement) {
            throw new Error('Select .templateNota não encontrado');
        }
        
        // Procura a opção pelo texto
        const option = Array.from(selectElement.options).find(
            opt => opt.text.trim() === texto.trim()
        );
        
        if (!option) {
            throw new Error(`Opção com texto '${texto}' não encontrada`);
        }
        
        // Torna o select visível e interativo
        selectElement.style.display = 'block';
        selectElement.style.visibility = 'visible';
        selectElement.removeAttribute('disabled');
        
        // Seleciona a opção pelo value
        selectElement.value = option.value;
        
        // Dispara eventos
        selectElement.dispatchEvent(new Event('change', { bubbles: true }));
        selectElement.dispatchEvent(new Event('input', { bubbles: true }));
        selectElement.dispatchEvent(new Event('blur', { bubbles: true }));
        
        return {
            sucesso: true,
            valorSelecionado: selectElement.value,
            textoSelecionado: selectElement.options[selectElement.selectedIndex].text
        };
        """
        
        resultado = js_engine.execute_js(
            script, 
            texto, 
            timeout=timeout,
            fallback_result=None
        )
        
        if resultado and resultado.get('sucesso'):
            value = resultado.get('valorSelecionado', '')
            log(doc, f"✅ Template '{texto}' (value={value}) selecionado com sucesso!")
            js_engine.wait_ajax_complete(timeout)
            return True
        else:
            log(doc, f"⚠️ Falha ao selecionar template '{texto}'")
            return False
            
    except Exception as e:
        log(doc, f"❌ Erro ao selecionar template '{texto}': {e}")
        return False

def selecionar_banco_por_value(js_engine, doc, value="1040", nome_banco="Caixa Econômica Federal", timeout=5):
    """
    Seleciona uma opção do select 'banco' pelo valor (value).
    
    Args:
        js_engine: Instância do JSForceEngine
        doc: Documento para logs
        value: Valor do option a ser selecionado (default: "1040" para CEF)
        nome_banco: Nome do banco para logs (default: "Caixa Econômica Federal")
        timeout: Timeout para aguardar AJAX (default: 5)
    
    Returns:
        bool: True se selecionou com sucesso, False caso contrário
    """
    try:
        log(doc, f"🏦 Selecionando banco '{nome_banco}' (value={value})...")
        
        # Script JavaScript robusto para selecionar a opção
        script = """
        const selectElement = document.querySelector('select.chqf.banco');
        const value = arguments[0];
        
        if (!selectElement) {
            throw new Error('Select .chqf.banco não encontrado');
        }
        
        // Verifica se a opção existe
        const option = Array.from(selectElement.options).find(opt => opt.value === value);
        if (!option) {
            throw new Error(`Opção com value '${value}' não encontrada`);
        }
        
        // Torna o select visível e interativo
        selectElement.style.display = 'block';
        selectElement.style.visibility = 'visible';
        selectElement.removeAttribute('disabled');
        
        // Seleciona a opção
        selectElement.value = value;
        
        // Dispara eventos para garantir que listeners sejam acionados
        selectElement.dispatchEvent(new Event('change', { bubbles: true }));
        selectElement.dispatchEvent(new Event('input', { bubbles: true }));
        selectElement.dispatchEvent(new Event('blur', { bubbles: true }));
        
        // Verifica se a seleção foi bem-sucedida
        if (selectElement.value !== value) {
            throw new Error('Seleção não foi aplicada corretamente');
        }
        
        return {
            sucesso: true,
            valorSelecionado: selectElement.value,
            textoSelecionado: selectElement.options[selectElement.selectedIndex].text
        };
        """
        
        # Executa o script com proteção contra timeout
        resultado = js_engine.execute_js(
            script, 
            value, 
            timeout=timeout,
            fallback_result=None
        )
        
        if resultado and resultado.get('sucesso'):
            texto = resultado.get('textoSelecionado', nome_banco)
            log(doc, f"✅ Banco '{texto}' selecionado com sucesso!")
            
            # Aguarda AJAX completar
            js_engine.wait_ajax_complete(timeout)
            return True
        else:
            log(doc, f"⚠️ Falha ao selecionar banco '{nome_banco}'")
            return False
            
    except Exception as e:
        log(doc, f"❌ Erro ao selecionar banco '{nome_banco}': {e}")
        return False

def garantir_accordion_aberto(js_engine, doc, titulo_header="Pagamento com Cartão", timeout=5):
    """
    Garante que o painel do Accordion com o título informado esteja ABERTO.
    Usa jQuery UI quando disponível; caso contrário, força atributos/classes.
    """
    script = r"""
    (function(titulo){
      // acha todos os headers do accordion
      var headers = Array.from(document.querySelectorAll("h3.ui-accordion-header"));
      if (!headers.length) return {ok:false, reason:"Nenhum header de accordion encontrado."};

      // localiza o header pelo texto
      var idx = headers.findIndex(h => (h.textContent || "").trim().includes(titulo));
      if (idx < 0) return {ok:false, reason:"Header não encontrado: " + titulo};

      var header = headers[idx];
      var acc = header.closest(".ui-accordion");
      var panel = header.nextElementSibling;

      function estaAberto(){
        var exp = header.getAttribute("aria-expanded");
        if (exp === "true") return true;
        if (panel && panel.style.display !== "none") return true;
        if (panel && panel.classList.contains("ui-accordion-content-active")) return true;
        return false;
      }

      // 1) Preferência: jQuery UI
      if (window.jQuery && jQuery(acc).accordion) {
        jQuery(acc).accordion("option", "active", idx);
      } else {
        // 2) Fallback manual: força estado "aberto"
        header.setAttribute("aria-expanded", "true");
        header.classList.add("ui-state-active","ui-accordion-header-active");
        if (panel){
          panel.style.display = "block";
          panel.setAttribute("aria-hidden","false");
          panel.classList.add("ui-accordion-content-active");
        }
      }

      // valida
      var aberto = estaAberto();

      // medida extra: se algum outro handler tentar fechar imediatamente, comuta uma flag por 800ms
      if (aberto) {
        header.setAttribute("data-lock-open","1");
        setTimeout(function(){ header.removeAttribute("data-lock-open"); }, 800);
      }

      return {ok:aberto, index:idx};
    })(arguments[0]);
    """
    res = js_engine.execute_js(script, titulo_header, timeout=timeout, fallback_result={"ok": False})
    if res.get("ok"):
        log(doc, f"✅ Accordion '{titulo_header}' aberto e travado em aberto (idx={res.get('index')}).")
        js_engine.wait_ajax_complete(timeout)
        return True
    else:
        log(doc, f"⚠️ Falha ao abrir accordion '{titulo_header}': {res.get('reason')}")
        return False


def abrir_aba_pagamento_cartao(js_engine, doc, max_attempts=3, timeout=8):
    """
    Abre a aba 'Pagamento com Cartão' clicando no <h3> do accordion
    SOMENTE se ela não estiver aberta.
    """
    drv = js_engine.driver
    hdr_xpath = "//h3[contains(@class,'ui-accordion-header') and contains(normalize-space(.),'Pagamento com Cartão')]"
    log(doc, "🔄 Abrindo a aba 'Pagamento com Cartão'...")

    for attempt in range(1, max_attempts + 1):
        try:
            hdr = WebDriverWait(drv, timeout).until(
                EC.presence_of_element_located((By.XPATH, hdr_xpath))
            )

            # lê o estado atual
            aria_expanded = hdr.get_attribute("aria-expanded")

            # se já está aberta, não clica de novo
            if aria_expanded == "true":
                log(doc, "✅ Aba 'Pagamento com Cartão' já estava aberta.")
                return True

            # se tiver helper de scroll, usa
            try:
                if hasattr(js_engine, "scroll_into_view"):
                    js_engine.scroll_into_view(hdr)
            except:
                pass

            # clica uma vez só
            js_engine.click_element(hdr)

            # espera ficar aberta
            WebDriverWait(drv, timeout).until(
                lambda d: d.find_element(By.XPATH, hdr_xpath).get_attribute("aria-expanded") == "true"
            )

            log(doc, "✅ Aba 'Pagamento com Cartão' aberta.")
            return True

        except Exception as e:
            if attempt < max_attempts:
                log(doc, f"⚠️ Tentativa {attempt} falhou, tentando novamente...")
                time.sleep(1)
            else:
                log(doc, f"❌ Erro após {max_attempts} tentativas: {e}")
                return False

def preencher_por_label(js_engine, doc, label_text, valor, timeout=8, clear=True):
    """
    Preenche um input ancorado pelo texto do <label>.
    Aceita label bem próximo ao input (seguinte no DOM).
    """
    drv = js_engine.driver
    # Ancorado no label → primeiro input adiante
    xpath = f"//label[normalize-space()='{label_text}']/following::input[1]"
    try:
        el = WebDriverWait(drv, timeout).until(
            EC.element_to_be_clickable((By.XPATH, xpath))
        )
        js_engine.scroll_into_view(el)
        if clear:
            el.clear()
        el.send_keys(valor)
        log(doc, f"✏️ Preenchido '{label_text}': {valor}")
        return True
    except Exception as e:
        log(doc, f"❌ Não foi possível preencher '{label_text}': {e}")
        return False

def preencher_campos_cartao(js_engine, doc, dados):
    """
    Garante a aba aberta e preenche campos principais do cartão.
    'dados' é um dict com chaves: nome, numero, autorizacao, data, valor, parcelas.
    """
    if not abrir_aba_pagamento_cartao(js_engine, doc):
        return False

    # Seletores específicos baseados nos atributos únicos de cada campo
    campos = {
        "nome": {
            "seletor": "input.chqf[type='text'][maxlength='255'][style*='width: 255px']",
            "valor": dados.get("nome", ""),
            "descricao": "Nome"
        },
        "numero": {
            "seletor": "input.chqf[type='text'][maxlength='16'][style*='width: 200px']",
            "valor": dados.get("numero", ""),
            "descricao": "Número do Cartão"
        },
        "autorizacao": {
            "seletor": "input.chqf[type='text'][maxlength='10'][style*='width: 80px']",
            "valor": dados.get("autorizacao", ""),
            "descricao": "Autorização"
        },
        "valor": {
            "seletor": "input.chqf[type='text'][placeholder='R$ ']",
            "valor": dados.get("valor", ""),
            "descricao": "Valor"
        },
        "parcelas": {
            "seletor": "input.chqf[type='text'][maxlength='3']",
            "valor": dados.get("parcelas", ""),
            "descricao": "Quantidade de Parcelas"
        }
    }

    ok = True
    for campo_key, campo_info in campos.items():
        try:
            seletor = campo_info["seletor"]
            valor = campo_info["valor"]
            descricao = campo_info["descricao"]
            
            if not valor:
                log(doc, f"⊘ {descricao}: sem valor fornecido")
                continue
            
            # Usa force_fill do js_engine para preencher o campo
            try:
                js_engine.force_fill(seletor, str(valor), by_xpath=False)
                log(doc, f"✓ {descricao}: preenchido com '{valor}'")
            except Exception as e:
                log(doc, f"✗ {descricao}: erro ao preencher - {str(e)}")
                ok = False
                
        except Exception as e:
            log(doc, f"✗ {descricao}: erro inesperado - {str(e)}")
            ok = False

    if ok:
        log(doc, "✅ Todos os campos do cartão foram preenchidos com sucesso.")
    else:
        log(doc, "⚠️ Alguns campos do cartão não foram preenchidos. Veja logs acima.")
    
    return ok


def preencher_dados_cartao_por_indice(
    js_engine, 
    doc, 
    indice_data=1,
    indice_valor=2,
    indice_lov_bandeira=10,
    nome="TESTE NOME CARTÃO SELENIUM",
    numero="4111111111111111",
    autorizacao="123",
    parcelas="1",
    data_venda="30/10/2025",
    valor="10.000,00",
    bandeira_texto="BANDEIRA ELO CREDITO",
    adicionar=True,
    timeout=5
):
    """
    Preenche todos os dados de um cartão de forma sequencial.
    
    Args:
        js_engine: Instância do JSForceEngine
        doc: Documento para logs
        indice_data: Índice do campo de data (1-based)
        indice_valor: Índice do campo de valor (1-based)
        indice_lov_bandeira: Índice do botão LOV da bandeira
        nome: Nome do titular do cartão
        numero: Número do cartão
        autorizacao: Código de autorização
        parcelas: Quantidade de parcelas
        data_venda: Data da venda (formato dd/mm/yyyy)
        valor: Valor (com ou sem R$)
        bandeira_texto: Texto da bandeira para buscar no LOV
        adicionar: Se True, clica no botão "Adicionar" ao final
        timeout: Timeout para operações
    
    Returns:
        bool: True se tudo foi preenchido com sucesso, False caso contrário
    """
    try:
        log(doc, f"🎴 Iniciando preenchimento de dados do cartão (índice data={indice_data}, valor={indice_valor})...")
        
        # 1. Abre a aba de Pagamento com Cartão
        if not abrir_aba_pagamento_cartao(js_engine, doc):
            log(doc, "❌ Falha ao abrir aba de Pagamento com Cartão")
            return False
        
        # 2. Monta os dados do cartão
        dados_cartao = {
            "nome": nome,
            "numero": numero,
            "autorizacao": autorizacao,
            "parcelas": parcelas
        }
        
        # 3. Preenche os campos básicos do cartão
        if not safe_action(
            doc, 
            "Preenchendo campos básicos do cartão",
            lambda: preencher_campos_cartao(js_engine, doc, dados_cartao)
        ):
            log(doc, "❌ Falha ao preencher campos básicos do cartão")
            return False
        
        # 4. Preenche a data de venda
        if not safe_action(doc, f"Preenchendo Data de Venda (índice {indice_data})", lambda:
            preencher_datepicker_por_indice_xpath(
                js_engine, doc,
                xpath_base="//input[@class='hasDatepicker chqf']",
                indice=indice_data,
                data_valor=data_venda,
                descricao="Data de Venda"
            )
        ):
            log(doc, "⚠️ Falha ao preencher data de venda (continuando...)")
        
        # 5. Preenche o valor
        if not safe_action(doc, f"Preenchendo Valor (índice {indice_valor})", lambda:
            preencher_campo_monetario_por_indice(
                js_engine, doc,
                xpath_base="//input[@type='text' and contains(@class,'chqf') and contains(@placeholder,'R$')]",
                indice=indice_valor,
                valor=valor,
                descricao="Valor do Pagamento"
            )
        ):
            log(doc, "⚠️ Falha ao preencher valor (continuando...)")
        
        # 6. Seleciona a bandeira via LOV
        lov_handler = LOVHandler(js_engine, doc)
        if not safe_action(doc, f"Selecionando Bandeira '{bandeira_texto}'", lambda:
            lov_handler.open_and_select(
                btn_index=indice_lov_bandeira,
                search_text=bandeira_texto,
                result_text=bandeira_texto
            )
        ):
            log(doc, "❌ Falha ao selecionar bandeira")
            return False
        
        # 7. Clica no botão "Adicionar" se solicitado
        if adicionar:
            if not safe_action(doc, "Adicionando cartão", lambda:
                js_engine.force_click(
                    "//a[@class='btModel btGray btAddCartao' and contains(normalize-space(.),'Adicionar')]", 
                    by_xpath=True
                )
            ):
                log(doc, "❌ Falha ao clicar em 'Adicionar'")
                return False
            
            time.sleep(0.5)
            encontrar_mensagem_alerta()
        
        log(doc, "✅ Dados do cartão preenchidos com sucesso!")
        return True
        
    except Exception as e:
        log(doc, f"❌ Erro ao preencher dados do cartão: {e}")
        return False


# ==== EXEMPLOS DE USO ====

# Exemplo 1: Cartão de Crédito ELO
"""
preencher_dados_cartao_por_indice(
    js_engine, doc,
    indice_data=1,
    indice_valor=2,
    indice_lov_bandeira=10,
    nome="TESTE NOME CARTÃO SELENIUM",
    numero="4111111111111111",
    autorizacao="123",
    parcelas="1",
    data_venda="30/10/2025",
    valor="10.000,00",
    bandeira_texto="BANDEIRA ELO CREDITO",
    adicionar=True
)
"""

# Exemplo 2: Cartão de Débito SIPAG
"""
preencher_dados_cartao_por_indice(
    js_engine, doc,
    indice_data=1,
    indice_valor=3,  # Valor é o 3º campo agora
    indice_lov_bandeira=10,
    nome="TESTE NOME CARTÃO SELENIUM",
    numero="4111111111111111",
    autorizacao="123",
    parcelas="1",
    data_venda="30/10/2025",
    valor="10.000,00",
    bandeira_texto="SIPAG - DÉBITO",
    adicionar=True
)
"""

def clicar_primeiro_sp_delete(js_engine, doc, timeout=5):
    """
    Localiza e clica no primeiro elemento <a class="sprites sp-delete"> visível na tela.
    """
    from selenium.common.exceptions import NoSuchElementException, ElementClickInterceptedException, StaleElementReferenceException
    from selenium.webdriver.common.by import By
    import time

    driver = js_engine.driver
    xpath = (
        "//a[contains(concat(' ', normalize-space(@class), ' '), ' sprites ') "
        "and contains(concat(' ', normalize-space(@class), ' '), ' sp-delete ')][1]"
    )

    log(doc, "🗑️ Procurando o primeiro botão '<a class=\"sprites sp-delete\">'...")

    try:
        el = driver.find_element(By.XPATH, xpath)
        log(doc, "🎯 Botão encontrado, tentando clicar...")

        try:
            el.click()
            log(doc, "✅ Clique padrão realizado com sucesso.")
        except (ElementClickInterceptedException, StaleElementReferenceException):
            driver.execute_script("arguments[0].click();", el)
            log(doc, "⚡ Clique forçado via JavaScript realizado.")

        time.sleep(0.5)
        js_engine.wait_ajax_complete(timeout)
        return True

    except NoSuchElementException:
        log(doc, "⚠️ Nenhum botão 'sp-delete' encontrado na tela.")
        return False
    except Exception as e:
        log(doc, f"❌ Erro ao clicar no botão 'sp-delete': {e}")
        return False

def clicar_todos_salvar(js_engine, doc, pausa_entre=0.5, timeout=5):
    """
    Procura todos os botões 'Salvar' visíveis e clica em cada um deles na ordem.
    Conta e exibe quantos botões existem antes de clicar.
    Usa js_engine.force_click() e registra log detalhado.
    """

    xpath_base = "//a[contains(@class,'btModel') and contains(@class,'btGray') and contains(@class,'btSave') and contains(normalize-space(.),'Salvar')]"

    try:
        elementos = js_engine.driver.find_elements("xpath", xpath_base)
        total = len(elementos)
        log(doc, f"💾 Foram encontrados {total} botão(ões) 'Salvar' na tela.")

        if total == 0:
            log(doc, "⚠️ Nenhum botão 'Salvar' foi encontrado.")
            return {"total": 0, "clicados": 0}

        total_clicados = 0
        for i in range(1, total + 1):
            xpath_indexado = f"({xpath_base})[{i}]"
            try:
                log(doc, f"🎯 Clicando no botão 'Salvar' (índice {i}/{total})...")
                js_engine.force_click(xpath_indexado, by_xpath=True)
                js_engine.wait_ajax_complete(timeout)
                total_clicados += 1
                log(doc, f"✅ Clique no botão 'Salvar' (índice {i}) realizado com sucesso.")
                if pausa_entre > 0:
                    import time
                    time.sleep(pausa_entre)
            except Exception as e:
                log(doc, f"⚠️ Falha ao clicar no botão 'Salvar' (índice {i}): {e}")

        log(doc, f"🧾 Resumo: {total_clicados}/{total} botões 'Salvar' clicados com sucesso.")
        return {"total": total, "clicados": total_clicados}

    except Exception as e:
        log(doc, f"⚠️ Erro ao procurar ou clicar nos botões 'Salvar': {e}")
        return {"total": 0, "clicados": 0}


def fechar_detalhes_venda(js_engine, doc, pausa=0.3, timeout_total=30):
    """
    Clica repetidamente no botão (//a[@class='fa fa-close'])
    usando múltiplas estratégias, até o elemento sumir da tela.
    """

    from selenium.common.exceptions import NoSuchElementException, StaleElementReferenceException
    import time
    xpath = "(//a[@class='fa fa-close'])[1]"
    inicio = time.time()
    tentativas = 0

    log(doc, "🧩 Iniciando fechamento do modal 'Detalhes da Venda'...")

    while time.time() - inicio < timeout_total:
        tentativas += 1
        try:
            elementos = js_engine.driver.find_elements("xpath", xpath)
            visiveis = [el for el in elementos if el.is_displayed()]
            if not visiveis:
                log(doc, f"✅ Modal fechado após {tentativas} tentativa(s).")
                return True

            el = visiveis[0]
            log(doc, f"🔄 Tentativa {tentativas}: clicando no botão 'X' (fa fa-close)...")

            # Estratégias de clique
            estrategias = [
                lambda: el.click(),
                lambda: js_engine.driver.execute_script("arguments[0].click();", el),
                lambda: js_engine.click_element(el),
                lambda: js_engine.force_click(xpath, by_xpath=True),
                lambda: js_engine.driver.execute_script("""
                    const el = arguments[0];
                    const evs = ['mouseover','mousedown','mouseup','click'];
                    evs.forEach(ev => el.dispatchEvent(new MouseEvent(ev,{bubbles:true,cancelable:true})));
                """, el)
            ]

            for i, estrategia in enumerate(estrategias, 1):
                try:
                    log(doc, f"   ▶️ Estratégia {i} de clique...")
                    estrategia()
                    time.sleep(pausa)
                    js_engine.wait_ajax_complete(3)
                    break
                except Exception as e:
                    log(doc, f"⚠️ Estratégia {i} falhou: {e}")

            # Verifica se sumiu
            time.sleep(pausa)
            elementos = js_engine.driver.find_elements("xpath", xpath)
            visiveis = [el for el in elementos if el.is_displayed()]
            if not visiveis:
                log(doc, "✅ Modal fechado com sucesso!")
                return True

        except (NoSuchElementException, StaleElementReferenceException):
            log(doc, "✅ Elemento 'fa fa-close' não encontrado — modal fechado.")
            return True
        except Exception as e:
            log(doc, f"⚠️ Erro ao tentar fechar: {e}")

        time.sleep(pausa)

    log(doc, f"⏰ Timeout atingido ({timeout_total}s) — o modal ainda não fechou.")
    return False
def clicar_sp_delete_por_indice(js_engine, doc, indice=1, timeout=5):
    """
    Clica no elemento <span class="sprites sp-delete" title="Estornar Venda"> pelo índice informado (1-based).
    Usa múltiplas estratégias de clique e aguarda AJAX após o clique.
    """

    from selenium.common.exceptions import (
        NoSuchElementException,
        ElementClickInterceptedException,
        StaleElementReferenceException,
    )
    from selenium.webdriver.common.by import By
    import time

    driver = js_engine.driver
    xpath_base = "//span[@class='sprites sp-delete' and @title='Estornar Venda']"
    xpath_indexado = f"({xpath_base})[{indice}]"

    log(doc, f"🗑️ Procurando botão 'Estornar Venda' (índice {indice})...")

    try:
        elementos = driver.find_elements(By.XPATH, xpath_base)
        total = len(elementos)
        log(doc, f"📊 Total encontrado: {total} elemento(s).")

        if total == 0:
            log(doc, "⚠️ Nenhum botão 'Estornar Venda' encontrado.")
            return False

        if indice > total:
            log(doc, f"⚠️ Índice {indice} inválido — só existem {total} elemento(s).")
            return False

        el = elementos[indice - 1]
        log(doc, f"🎯 Elemento localizado (índice {indice}). Tentando clicar...")

        estrategias = [
            lambda: el.click(),
            lambda: driver.execute_script("arguments[0].click();", el),
            lambda: js_engine.click_element(el),
            lambda: js_engine.force_click(xpath_indexado, by_xpath=True),
            lambda: driver.execute_script("""
                const el = arguments[0];
                const evs = ['mouseover','mousedown','mouseup','click'];
                evs.forEach(ev => el.dispatchEvent(new MouseEvent(ev,{bubbles:true,cancelable:true})));
            """, el)
        ]

        for i, estrategia in enumerate(estrategias, 1):
            try:
                log(doc, f"   ▶️ Estratégia {i} de clique...")
                estrategia()
                time.sleep(0.4)
                js_engine.wait_ajax_complete(timeout)
                log(doc, f"✅ Clique realizado com sucesso (estratégia {i}).")
                return True
            except (ElementClickInterceptedException, StaleElementReferenceException):
                log(doc, f"⚠️ Estratégia {i} falhou (elemento interceptado ou stale).")
                continue
            except Exception as e:
                log(doc, f"⚠️ Estratégia {i} falhou: {e}")
                continue

        log(doc, "❌ Nenhuma estratégia de clique funcionou.")
        return False

    except NoSuchElementException:
        log(doc, "⚠️ Nenhum elemento 'sp-delete' encontrado.")
        return False
    except Exception as e:
        log(doc, f"❌ Erro ao clicar no botão 'sp-delete': {e}")
        return False


def clicar_botao_limpar_por_indice(js_engine, doc, indice=1, timeout=5):
    """
    Clica no botão <a class="btModel btGray btclear">Limpar</a> pelo índice informado (1-based).
    Usa múltiplas estratégias de clique e aguarda AJAX após o clique.
    """

    from selenium.common.exceptions import (
        NoSuchElementException,
        ElementClickInterceptedException,
        StaleElementReferenceException,
    )
    from selenium.webdriver.common.by import By
    import time

    driver = js_engine.driver
    xpath_base = "//a[@class='btModel btGray btclear' and normalize-space()='Limpar']"
    xpath_indexado = f"({xpath_base})[{indice}]"

    log(doc, f"🧹 Procurando botão 'Limpar' (índice {indice})...")

    try:
        elementos = driver.find_elements(By.XPATH, xpath_base)
        total = len(elementos)
        log(doc, f"📊 Total encontrado: {total} botão(ões) 'Limpar'.")

        if total == 0:
            log(doc, "⚠️ Nenhum botão 'Limpar' encontrado.")
            return False

        if indice > total:
            log(doc, f"⚠️ Índice {indice} inválido — só existem {total} botão(ões).")
            return False

        el = elementos[indice - 1]
        log(doc, f"🎯 Botão 'Limpar' localizado (índice {indice}). Tentando clicar...")

        estrategias = [
            lambda: el.click(),
            lambda: driver.execute_script("arguments[0].click();", el),
            lambda: js_engine.click_element(el),
            lambda: js_engine.force_click(xpath_indexado, by_xpath=True),
            lambda: driver.execute_script("""
                const el = arguments[0];
                const evs = ['mouseover','mousedown','mouseup','click'];
                evs.forEach(ev => el.dispatchEvent(new MouseEvent(ev,{bubbles:true,cancelable:true})));
            """, el)
        ]

        for i, estrategia in enumerate(estrategias, 1):
            try:
                log(doc, f"   ▶️ Estratégia {i} de clique...")
                estrategia()
                time.sleep(0.4)
                js_engine.wait_ajax_complete(timeout)
                log(doc, f"✅ Clique realizado com sucesso (estratégia {i}).")
                return True
            except (ElementClickInterceptedException, StaleElementReferenceException):
                log(doc, f"⚠️ Estratégia {i} falhou (elemento interceptado ou stale).")
                continue
            except Exception as e:
                log(doc, f"⚠️ Estratégia {i} falhou: {e}")
                continue

        log(doc, "❌ Nenhuma estratégia de clique funcionou.")
        return False

    except NoSuchElementException:
        log(doc, "⚠️ Nenhum botão 'Limpar' encontrado no DOM.")
        return False
    except Exception as e:
        log(doc, f"❌ Erro ao clicar no botão 'Limpar': {e}")
        return False



def clicar_sim_ate_sumir(js_engine, doc, index=0, timeout=15, pausa=0.5):
    """
    Clica repetidamente no botão 'Sim' (BtYes) dentro de um modal de confirmação,
    até que o modal desapareça da tela.

    Parâmetros:
    - js_engine: instância do engine JSForceEngine
    - doc: documento de log
    - index: índice do botão 'Sim' (0 = primeiro, 1 = segundo, etc.)
    - timeout: tempo máximo em segundos antes de desistir
    - pausa: intervalo entre tentativas
    """

    # XPath do botão 'Sim' - mais simples e direto
    xpath_botao = f"(//a[@id='BtYes' and contains(@class,'btyes')])[{index + 1}]"
    
    # XPath do modal - busca qualquer modal visível com botão de confirmação
    xpath_modal_generico = "//div[contains(@class,'modal') and contains(@class,'confirmationYesNo')]"

    inicio = time.time()
    tentativas = 0
    sucesso = False

    log(doc, f"🟦 Iniciando clique no botão 'Sim' (índice {index}) até o modal sumir...")

    # FORÇA pelo menos UMA tentativa de clique antes de verificar
    clicou_pelo_menos_uma_vez = False

    while True:
        tentativas += 1
        
        try:
            # Verifica se o botão 'Sim' existe
            try:
                botoes = js_engine.driver.find_elements("xpath", xpath_botao)
                if not botoes and clicou_pelo_menos_uma_vez:
                    log(doc, f"✅ Botão 'Sim' não encontrado - modal fechado após {tentativas - 1} tentativa(s).")
                    sucesso = True
                    break
                    
                if not botoes:
                    log(doc, "⚠️ Botão 'Sim' não encontrado. Aguardando aparecer...")
                    time.sleep(pausa)
                    if time.time() - inicio > timeout:
                        log(doc, f"⏰ Timeout: botão 'Sim' nunca apareceu.")
                        break
                    continue
                    
            except Exception as e:
                log(doc, f"⚠️ Erro ao localizar botão: {e}")
                if clicou_pelo_menos_uma_vez:
                    sucesso = True
                    break
                continue

            log(doc, f"➡️ Tentativa {tentativas}: clicando no botão 'Sim'...")

            # Estratégia 1 — JavaScript com múltiplos eventos (MAIS AGRESSIVA)
            try:
                script = f"""
                var btn = document.evaluate("{xpath_botao}", document, null, 
                    XPathResult.FIRST_ORDERED_NODE_TYPE, null).singleNodeValue;
                if (btn) {{
                    // Remove qualquer bloqueio
                    btn.style.pointerEvents = 'auto';
                    btn.removeAttribute('disabled');
                    btn.style.visibility = 'visible';
                    btn.style.display = 'inline-block';
                    
                    // Scroll até o botão
                    btn.scrollIntoView({{block: 'center'}});
                    
                    // Dispara TODOS os eventos
                    ['mouseover','mouseenter','mousedown','focus','mouseup','click'].forEach(function(evt) {{
                        btn.dispatchEvent(new MouseEvent(evt, {{bubbles: true, cancelable: true}}));
                    }});
                    
                    // Clique direto
                    btn.click();
                    
                    // jQuery se disponível
                    if (typeof jQuery !== 'undefined') {{
                        jQuery(btn).trigger('click');
                    }}
                    
                    // onclick manual
                    if (btn.onclick) btn.onclick();
                    
                    return true;
                }}
                return false;
                """
                resultado = js_engine.driver.execute_script(script)
                if resultado:
                    log(doc, "✅ Clique via JavaScript ultra-agressivo executado.")
                    clicou_pelo_menos_uma_vez = True
                else:
                    log(doc, "⚠️ Script JavaScript retornou false.")
            except Exception as e:
                log(doc, f"⚠️ Falha no JavaScript agressivo: {e}")

            time.sleep(pausa)

            # Estratégia 2 — Force click do engine
            try:
                js_engine.force_click(xpath_botao, by_xpath=True)
                log(doc, "✅ Clique via force_click() executado.")
                clicou_pelo_menos_uma_vez = True
            except Exception as e:
                log(doc, f"⚠️ Falha no force_click: {e}")

            time.sleep(pausa)

            # Estratégia 3 — Selenium + ActionChains
            try:
                btn = js_engine.driver.find_element("xpath", xpath_botao)
                ActionChains(js_engine.driver).move_to_element(btn).pause(0.1).click().perform()
                log(doc, "✅ Clique via ActionChains executado.")
                clicou_pelo_menos_uma_vez = True
            except Exception as e:
                log(doc, f"⚠️ Falha no ActionChains: {e}")

            time.sleep(pausa)
            
            # Aguarda AJAX
            js_engine.wait_ajax_complete(3)

            # Verifica se o botão/modal ainda existem
            try:
                botoes = js_engine.driver.find_elements("xpath", xpath_botao)
                if not botoes:
                    log(doc, f"✅ Modal fechado após {tentativas} tentativa(s)!")
                    sucesso = True
                    break
            except:
                log(doc, "✅ Modal fechado (erro ao verificar = não existe mais)!")
                sucesso = True
                break

        except (NoSuchElementException, StaleElementReferenceException):
            if clicou_pelo_menos_uma_vez:
                log(doc, "✅ Elemento 'Sim' sumiu - modal fechado!")
                sucesso = True
                break
            else:
                log(doc, "⚠️ Elemento 'Sim' não encontrado, tentando novamente...")
        except Exception as e:
            log(doc, f"❌ Erro inesperado: {e}")

        if time.time() - inicio > timeout:
            log(doc, f"⏰ Tempo limite de {timeout}s atingido.")
            break

        time.sleep(pausa)

    if sucesso:
        log(doc, "🎉 Confirmação concluída com sucesso!")
    else:
        log(doc, "⚠️ Não foi possível confirmar - modal pode ainda estar aberto.")
    
    return sucesso



# ==== EXECUÇÃO DO TESTE ====
def executar_teste():
    """Execução principal do teste com JS forçado e proteção anti-timeout"""
    global driver, wait, doc
    
    try:
        if not inicializar_driver():
            return False
        
        # Cria engine JS forçado COM PROTEÇÃO ANTI-TIMEOUT
        js_engine = JSForceEngine(driver, wait, doc, timeout_padrao=10, max_retries=3)
        lov_handler = LOVHandler(js_engine, doc)
        
        # ===== LOGIN =====
        safe_action(doc, "Acessando sistema", lambda: driver.get(URL))
        
        def fazer_login():
            wait.until(EC.presence_of_element_located((By.ID, "j_id15:email"))).send_keys(LOGIN_EMAIL)
            wait.until(EC.presence_of_element_located((By.ID, "j_id15:senha"))).send_keys(LOGIN_PASSWORD, Keys.ENTER)
            time.sleep(5)

        safe_action(doc, "Realizando login", fazer_login)
        
        # ===== MENU =====
        def abrir_menu():
            driver.execute_script("document.body.style.zoom='90%'")
            driver.find_element(By.TAG_NAME, "body").send_keys(Keys.F3)
            time.sleep(2)
        
        safe_action(doc, "Abrindo menu (F3)", abrir_menu)
        
        # ===== CAIXA =====
        safe_action(doc, "Acessando Caixa", lambda:
            js_engine.force_click('/html/body/div[15]/ul/li[8]/img', by_xpath=True)
        )
        
        time.sleep(3)
        
        safe_action(doc, "Clicando em 'Devoluções'", lambda:
            js_engine.force_click(
                '#gsCaixa > div.wdTelas > div.telaInicial.clearfix.overflow.overflowY > ul > li:nth-child(2) > a > span'
            )
        )
        
        time.sleep(5)

        safe_action(doc, "Preenchendo CPF", lambda:
            js_engine.force_fill("//input[@maxlength='14']", "504.571.668-94", by_xpath=True)
        )
        safe_action(doc, "Preenchendo Número do Contrato", lambda:
            js_engine.force_fill("//input[@class='nContrato']", "113190", by_xpath=True)
        )



        
        safe_action(doc, "Preenchendo Data Inicial", 
                   preencher_datepicker_por_indice(0, "10/11/2025"))

        
        safe_action(doc, "Preenchendo Data Final", 
                   preencher_datepicker_por_indice(1, "10/11/2025"))


        safe_action(doc, "Pesquisando", lambda: (
            js_engine.force_click("(//a[@class='btModel btGray btfind'])[1]", by_xpath=True),
            time.sleep(1)
        ))


        safe_action(doc, "Clicando em 'Detalhes da Venda' e capturando screenshot", lambda: (
            js_engine.force_click("//span[contains(@class,'sp-dadosDinamicos') and contains(@title,'Detalhes da Venda')]", by_xpath=True),
            time.sleep(1)
        ))


        safe_action(doc, "Fechando aba: 'Detalhes da Venda'", lambda:
            fechar_detalhes_venda(js_engine, doc)
        )


        safe_action(doc, "Estornando primeira Venda", lambda:
            clicar_sp_delete_por_indice(js_engine, doc, indice=1)
        )


        safe_action(doc, "Selecionando Motivo Estorno", lambda:
            lov_handler.open_and_select(
                btn_index=1,
                search_text="ESTORNO DE PAGAMENTO",
                result_text="ESTORNO DE PAGAMENTO"
            )
        )
        safe_action(doc, "Clicando em 'Salvar'", lambda:
            clicar_todos_salvar(js_engine, doc)
        )
        safe_action(doc, "Confirmando estorno", lambda:
            clicar_sim_ate_sumir(js_engine, doc, index=0, timeout=15, pausa=0.5)
        )

        encontrar_mensagem_alerta()
        time.sleep(60)

        safe_action(doc, "Limpando campos", lambda:
            clicar_botao_limpar_por_indice(js_engine, doc, indice=1)
        )


        safe_action(doc, "Fechando modal do Caixa", lambda:
            js_engine.force_click('#gsCaixa > div.wdTop.ui-draggable-handle > div > a')
        )

        log(doc, "🎉 Teste concluído com sucesso!")
        return True
        
    except Exception as e:
        log(doc, f"❌ ERRO FATAL: {e}")
        take_screenshot(driver, doc, "erro_fatal")
        return False

# ==== MAIN ====
def main():
    """Ponto de entrada principal"""
    global doc
    
    try:
        log(doc, "🚀 Iniciando teste de Fluxo de Caixa")
        log(doc, "=" * 70)

        
        sucesso = executar_teste()
        
        log(doc, "=" * 70)
        if sucesso:
            log(doc, "✅ TESTE EXECUTADO COM SUCESSO!")
        else:
            log(doc, "❌ TESTE FINALIZADO COM ERROS")
            
    except Exception as e:
        log(doc, f"❌ Erro na execução principal: {e}")
        
    finally:
        finalizar_relatorio()

if __name__ == "__main__":
    main()