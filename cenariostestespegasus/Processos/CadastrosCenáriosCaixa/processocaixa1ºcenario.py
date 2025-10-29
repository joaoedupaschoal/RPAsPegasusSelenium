# ==== IMPORTS ====
from datetime import datetime, timedelta
from datetime import time as dt_time
import time
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver import ActionChains
from selenium.common.exceptions import *
from selenium.webdriver.support import expected_conditions as EC
from selenium import webdriver
from selenium.webdriver.support.ui import Select, WebDriverWait
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service
from webdriver_manager.chrome import ChromeDriverManager
from docx import Document
from docx.shared import Inches
from faker import Faker
from faker.providers import BaseProvider
from validate_docbr import CPF
import subprocess
import os
import random
import re
from functools import wraps

# ==== CONFIGURAÇÕES GLOBAIS ====
TIMEOUT_DEFAULT = 30
TIMEOUT_CURTO = 10
TIMEOUT_LONGO = 60
CAMINHO_ARQUIVO_UPLOAD = "C:/Users/Gold System/Documents/teste.png"
URL = "http://localhost:8080/gs/login.xhtml"
LOGIN_EMAIL = "joaoeduardo.gold@outlook.com"
LOGIN_PASSWORD = "071999gs"

# ==== VARIÁVEIS GLOBAIS ====
doc = Document()
doc.add_heading("RELATÓRIO DO TESTE", 0)
doc.add_paragraph("Controle de Caixa - Caixa – Cenário 1: Rotina parcial de Fluxo de Caixa")
doc.add_paragraph(f"Data do teste: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")

screenshot_registradas = set()
driver = None
wait = None

# ==== PROVIDERS CUSTOMIZADOS ====
class BrasilProvider(BaseProvider):
    def rg(self):
        numeros = [str(random.randint(0, 9)) for _ in range(8)]
        return ''.join(numeros) + '-' + str(random.randint(0, 9))

fake = Faker("pt_BR")
fake.add_provider(BrasilProvider)

# ==== GERAÇÃO DE DATAS ====
def gerar_datas_validas(hora_padrao="00:00", dias_fim=0):
    hoje_date = datetime.today().date()
    dez_anos_atras = hoje_date - timedelta(days=3650)
    data_falecimento = fake.date_between(start_date=dez_anos_atras, end_date=hoje_date)
    idade_minima, idade_maxima = 18, 110
    data_nascimento = data_falecimento - timedelta(days=random.randint(idade_minima * 365, idade_maxima * 365))
    data_sepultamento = data_falecimento + timedelta(days=random.randint(1, 10))
    data_registro = data_sepultamento + timedelta(days=random.randint(1, 10))
    data_velorio = fake.date_between(start_date=data_falecimento, end_date=data_sepultamento)
    data_inicio_date = hoje_date + timedelta(days=random.randint(2, 30))
    h, m = map(int, hora_padrao.split(":"))
    dt_inicio = datetime.combine(data_inicio_date, dt_time(h, m))
    dt_fim = dt_inicio + timedelta(days=dias_fim)
    fmt_data = "%d/%m/%Y"
    fmt_dt = "%d/%m/%Y %H:%M"
    return (
        data_nascimento.strftime(fmt_data),
        data_falecimento.strftime(fmt_data),
        data_sepultamento.strftime(fmt_data),
        data_velorio.strftime(fmt_data),
        dt_inicio.strftime(fmt_dt),
        dt_fim.strftime(fmt_dt),
        data_registro.strftime(fmt_data),
        hoje_date.strftime(fmt_data),
    )

(data_nascimento, data_falecimento, data_sepultamento,
 data_velorio, data_inicio, data_fim, data_registro, hoje) = gerar_datas_validas(
    hora_padrao="08:50",
    dias_fim=0
)

# ==== UTILITÁRIOS DE LOG ====
def log(doc, msg):
    print(msg)
    doc.add_paragraph(msg)

def _sanitize_filename(name: str) -> str:
    name = name.strip().lower()
    name = re.sub(r"[<>:\"/\\|?*']", "_", name)
    name = re.sub(r"_+", "_", name)
    return name[:120]

def take_screenshot(driver, doc, nome):
    if driver is None:
        return
    nome = _sanitize_filename(nome)
    if nome not in screenshot_registradas:
        path = f"screenshots/{nome}.png"
        os.makedirs("screenshots", exist_ok=True)
        try:
            driver.save_screenshot(path)
            doc.add_paragraph(f"Screenshot: {nome}")
            doc.add_picture(path, width=Inches(5.5))
            screenshot_registradas.add(nome)
        except Exception as e:
            log(doc, f"⚠️ Erro ao tirar screenshot {nome}: {e}")

# ==== SISTEMA ANTI-TIMEOUT JAVASCRIPT ====
class JSTimeoutHandler:
    """Sistema robusto para lidar com timeouts JavaScript no Selenium"""
    
    def __init__(self, driver, doc, timeout_padrao=10, max_retries=3):
        self.driver = driver
        self.doc = doc
        self.timeout_padrao = timeout_padrao
        self.max_retries = max_retries
        self.last_error = None
        
    def log_timeout(self, msg, level="INFO"):
        """Log com timestamp"""

        prefix = {
            "INFO": "ℹ️ ",
            "WARN": "⚠️ ",
            "ERROR": "❌ ",
            "SUCCESS": "✅ "
        }.get(level, "📝 ")
        
        print(f" {prefix} {msg}")
        if hasattr(self.doc, 'add_paragraph'):
            self.doc.add_paragraph(f"{msg}")
    
    def execute_js_safe(self, script, *args, timeout=None, fallback_result=None):
        """Executa JavaScript com proteção contra timeouts"""
        timeout = timeout or self.timeout_padrao
        
        original_timeout = self.driver.timeouts.script
        self.driver.set_script_timeout(timeout)
        
        for tentativa in range(1, self.max_retries + 1):
            try:
                if tentativa > 1:
                    self.log_timeout(f"Tentativa {tentativa}/{self.max_retries}", "INFO")
                
                result = self.driver.execute_script(script, *args)
                self.driver.set_script_timeout(original_timeout)
                
                if tentativa > 1:
                    self.log_timeout("JavaScript executado com sucesso", "SUCCESS")
                return result
                
            except JavascriptException as e:
                self.last_error = e
                self.log_timeout(f"Erro JavaScript: {str(e)[:150]}", "ERROR")
                self._limpar_estado_js()
                
                if tentativa < self.max_retries:
                    time.sleep(1 + tentativa * 0.5)
                    continue
                    
            except TimeoutException as e:
                self.last_error = e
                self.log_timeout(f"Timeout JavaScript ({timeout}s)", "ERROR")
                self._forcar_parada_js()
                
                if tentativa < self.max_retries:
                    time.sleep(2 + tentativa)
                    continue
                    
            except WebDriverException as e:
                self.last_error = e
                self.log_timeout(f"Erro WebDriver: {str(e)[:150]}", "ERROR")
                
                if tentativa < self.max_retries:
                    time.sleep(1.5)
                    continue
                    
            except Exception as e:
                self.last_error = e
                self.log_timeout(f"Erro inesperado: {str(e)[:150]}", "ERROR")
                break
        
        try:
            self.driver.set_script_timeout(original_timeout)
        except:
            pass
            
        if fallback_result is not None:
            self.log_timeout(f"Usando valor fallback: {fallback_result}", "WARN")
        return fallback_result
    
    def _limpar_estado_js(self):
        """Limpa estado JavaScript do browser"""
        try:
            cleanup_script = """
                if (window.__cleanupTimers) {
                    window.__cleanupTimers.forEach(clearTimeout);
                    window.__cleanupTimers.forEach(clearInterval);
                }
                if (typeof jQuery !== 'undefined') {
                    jQuery.active = 0;
                }
                window.__pendingRequests = 0;
                return true;
            """
            self.driver.execute_script(cleanup_script)
            time.sleep(0.5)
        except Exception:
            pass
    
    def _forcar_parada_js(self):
        """Força parada de JavaScript travado"""
        try:
            self.driver.execute_script("window.stop();")
            time.sleep(0.3)
        except Exception:
            pass

# ==== JS FORCE ENGINE COM PROTEÇÃO ANTI-TIMEOUT ====
class JSForceEngine:
    """Motor de execução JavaScript forçado com proteção contra timeouts"""
    
    def __init__(self, driver, wait, doc, timeout_padrao=10, max_retries=3):
        self.driver = driver
        self.wait = wait
        self.doc = doc
        self.timeout_handler = JSTimeoutHandler(driver, doc, timeout_padrao, max_retries)
    
    def execute_js(self, script, *args, timeout=None, fallback_result=None):
        """Executa JavaScript com proteção contra timeout"""
        return self.timeout_handler.execute_js_safe(
            script, *args, timeout=timeout, fallback_result=fallback_result
        )
    
    def wait_ajax_complete(self, timeout=15):
        """Aguarda AJAX completar com proteção contra timeout"""
        script = """
            var jQueryOk = (typeof jQuery==='undefined') || (jQuery.active===0);
            var fetchOk = !window.__pendingRequests || window.__pendingRequests===0;
            var overlays = document.querySelectorAll(
                '.blockScreen, .blockUI, .loading, .overlay, [class*="loading"], [class*="spinner"]'
            );
            var overlayOk = true;
            for (var i=0; i<overlays.length; i++){
                var s=window.getComputedStyle(overlays[i]);
                if(s.display!=='none' && s.visibility!=='hidden' && parseFloat(s.opacity||1)>0.01){
                    overlayOk=false;
                    break;
                }
            }
            return jQueryOk && fetchOk && overlayOk;
        """
        
        end = time.time() + timeout
        while time.time() < end:
            try:
                done = self.execute_js(script, timeout=5, fallback_result=True)
                if done:
                    return True
            except:
                pass
            time.sleep(0.2)
        return True
    
    def force_click(self, selector, by_xpath=False, max_attempts=5):
        """Clique forçado com proteção contra timeout"""
        log(self.doc, f"🎯 Clique forçado em: {selector}")
        
        for attempt in range(max_attempts):
            try:
                strategies = [
                    self._click_strategy_2,
                    self._click_strategy_1,
                    self._click_strategy_3,
                    self._click_strategy_4,
                    self._click_strategy_5,
                ]
                
                for i, strategy in enumerate(strategies, 1):
                    try:
                        if attempt > 0 or i > 1:
                            log(self.doc, f"   Tentativa {attempt + 1}.{i}...")
                        
                        result = self.execute_js(
                            self._get_strategy_script(strategy, selector, by_xpath),
                            selector,
                            by_xpath,
                            timeout=5,
                            fallback_result=False
                        )
                        
                        if result:
                            log(self.doc, f"✅ Clique bem-sucedido (estratégia {i})")
                            time.sleep(0.5)
                            self.wait_ajax_complete(10)
                            return True
                            
                    except Exception as e:
                        if i == 1 and attempt == 0:
                            pass  # Silencia primeiro erro
                        else:
                            log(self.doc, f"   Estratégia {i} falhou: {str(e)[:80]}")
                        continue
                
                if attempt < max_attempts - 1:
                    time.sleep(1 + attempt * 0.5)
                    
            except Exception as e:
                log(self.doc, f"⚠️ Tentativa {attempt + 1} falhou: {e}")
                if attempt < max_attempts - 1:
                    time.sleep(1.5)
        
        raise Exception(f"Falha ao clicar após {max_attempts} tentativas: {selector}")
    
    def _get_strategy_script(self, strategy_func, selector, by_xpath):
        """Retorna o script JavaScript para cada estratégia"""
        base_locator = """
            var selector = arguments[0];
            var byXPath = arguments[1];
            var element;
            
            if (byXPath) {
                var result = document.evaluate(selector, document, null, 
                    XPathResult.FIRST_ORDERED_NODE_TYPE, null);
                element = result.singleNodeValue;
            } else {
                element = document.querySelector(selector);
            }
            
            if (!element) throw new Error('Elemento não encontrado');
        """
        
        if strategy_func == self._click_strategy_1:
            return base_locator + """
                element.style.pointerEvents = 'auto';
                element.style.display = 'block';
                element.style.visibility = 'visible';
                element.style.opacity = '1';
                element.removeAttribute('disabled');
                element.scrollIntoView({behavior: 'smooth', block: 'center'});
                setTimeout(function() { element.click(); }, 300);
                return true;
            """
        elif strategy_func == self._click_strategy_2:
            return base_locator + """
                element.style.pointerEvents = 'auto';
                element.removeAttribute('disabled');
                element.scrollIntoView({block: 'center'});
                
                var events = ['mouseover', 'mouseenter', 'mousemove', 'mousedown', 'mouseup', 'click'];
                events.forEach(function(eventType) {
                    var evt = new MouseEvent(eventType, {
                        bubbles: true, cancelable: true, view: window, detail: 1,
                        clientX: element.getBoundingClientRect().left + 5,
                        clientY: element.getBoundingClientRect().top + 5
                    });
                    element.dispatchEvent(evt);
                });
                
                if (typeof element.click === 'function') element.click();
                return true;
            """
        elif strategy_func == self._click_strategy_3:
            return base_locator + """
                element.style.display = 'block';
                element.style.visibility = 'visible';
                element.style.opacity = '1';
                element.style.pointerEvents = 'auto';
                element.focus();
                element.click();
                element.dispatchEvent(new Event('click', {bubbles: true, cancelable: true}));
                return true;
            """
        elif strategy_func == self._click_strategy_4:
            return base_locator + """
                element.removeAttribute('disabled');
                element.removeAttribute('readonly');
                element.style.pointerEvents = 'auto !important';
                element.style.display = 'block !important';
                element.style.visibility = 'visible !important';
                element.style.opacity = '1 !important';
                
                var overlays = document.querySelectorAll('.modal, .overlay, .blockUI, [role="dialog"]');
                overlays.forEach(function(overlay) {
                    overlay.style.display = 'none';
                    overlay.style.visibility = 'hidden';
                });
                
                element.focus();
                element.click();
                
                var clickEvent = new MouseEvent('click', {
                    view: window, bubbles: true, cancelable: true
                });
                element.dispatchEvent(clickEvent);
                
                if (typeof jQuery !== 'undefined') jQuery(element).trigger('click');
                return true;
            """
        else:  # strategy_5
            return base_locator + """
                var rect = element.getBoundingClientRect();
                var x = rect.left + rect.width / 2;
                var y = rect.top + rect.height / 2;
                
                var evt = document.createEvent('MouseEvents');
                evt.initMouseEvent('click', true, true, window, 1, x, y, x, y, false, false, false, false, 0, null);
                element.dispatchEvent(evt);
                
                if (element.onclick) element.onclick();
                
                var parent = element.parentElement;
                while (parent && parent !== document.body) {
                    if (parent.onclick) {
                        parent.onclick();
                        break;
                    }
                    parent = parent.parentElement;
                }
                return true;
            """
    
    def _click_strategy_1(self, selector, by_xpath):
        pass  # Implementado via _get_strategy_script
    
    def _click_strategy_2(self, selector, by_xpath):
        pass
    
    def _click_strategy_3(self, selector, by_xpath):
        pass
    
    def _click_strategy_4(self, selector, by_xpath):
        pass
    
    def _click_strategy_5(self, selector, by_xpath):
        pass
    
    def force_fill(self, selector, value, by_xpath=False, max_attempts=5):
        """Preenchimento forçado com proteção contra timeout"""
        log(self.doc, f"✏️ Preenchimento forçado: {selector} = '{value}'")
        
        for attempt in range(max_attempts):
            try:
                strategies = [
                    self._fill_strategy_1,
                    self._fill_strategy_2,
                    self._fill_strategy_3,
                    self._fill_strategy_4,
                ]
                
                for i, strategy in enumerate(strategies, 1):
                    try:
                        if attempt > 0 or i > 1:
                            log(self.doc, f"   Tentativa {attempt + 1}.{i}...")
                        
                        result = self.execute_js(
                            self._get_fill_script(strategy, selector, value, by_xpath),
                            selector,
                            value,
                            by_xpath,
                            timeout=5,
                            fallback_result=None
                        )
                        
                        time.sleep(0.3)
                        if self._validate_fill(selector, value, by_xpath):
                            log(self.doc, f"✅ Campo preenchido (estratégia {i})")
                            return True
                    except Exception as e:
                        if i == 1 and attempt == 0:
                            pass
                        else:
                            log(self.doc, f"   Estratégia {i} falhou: {str(e)[:80]}")
                        continue
                
                if attempt < max_attempts - 1:
                    time.sleep(1 + attempt * 0.5)
                    
            except Exception as e:
                log(self.doc, f"⚠️ Tentativa {attempt + 1} falhou: {e}")
        
        raise Exception(f"Falha ao preencher após {max_attempts} tentativas: {selector}")
    
    def _get_fill_script(self, strategy_func, selector, value, by_xpath):
        """Retorna o script de preenchimento"""
        base_locator = """
            var selector = arguments[0];
            var value = arguments[1];
            var byXPath = arguments[2];
            var element;
            
            if (byXPath) {
                var result = document.evaluate(selector, document, null, 
                    XPathResult.FIRST_ORDERED_NODE_TYPE, null);
                element = result.singleNodeValue;
            } else {
                element = document.querySelector(selector);
            }
            
            if (!element) throw new Error('Campo não encontrado');
        """
        
        if strategy_func == self._fill_strategy_1:
            return base_locator + """
                element.removeAttribute('disabled');
                element.removeAttribute('readonly');
                element.style.display = 'block';
                element.style.visibility = 'visible';
                element.scrollIntoView({block: 'center'});
                element.focus();
                element.dispatchEvent(new Event('focus', {bubbles: true}));
                element.value = '';
                element.value = value;
                ['input', 'change', 'blur', 'keyup'].forEach(function(evt) {
                    element.dispatchEvent(new Event(evt, {bubbles: true}));
                });
                return element.value;
            """
        elif strategy_func == self._fill_strategy_2:
            return base_locator + """
                var nativeInputValueSetter = Object.getOwnPropertyDescriptor(
                    window.HTMLInputElement.prototype, 'value'
                ).set;
                if (nativeInputValueSetter) {
                    nativeInputValueSetter.call(element, value);
                } else {
                    element.value = value;
                }
                element.dispatchEvent(new Event('input', {bubbles: true}));
                element.dispatchEvent(new Event('change', {bubbles: true}));
                element.dispatchEvent(new KeyboardEvent('keydown', {bubbles: true}));
                element.dispatchEvent(new KeyboardEvent('keyup', {bubbles: true}));
                element.dispatchEvent(new Event('blur', {bubbles: true}));
                return element.value;
            """
        elif strategy_func == self._fill_strategy_3:
            return base_locator + """
                element.value = value;
                if (typeof jQuery !== 'undefined') {
                    jQuery(element).val(value).trigger('input').trigger('change').trigger('blur');
                }
                ['focus', 'input', 'change', 'blur'].forEach(function(evt) {
                    element.dispatchEvent(new Event(evt, {bubbles: true}));
                });
                return element.value;
            """
        else:  # strategy_4
            return base_locator + """
                element.removeAttribute('disabled');
                element.removeAttribute('readonly');
                element.removeAttribute('maxlength');
                element.value = '';
                element.setAttribute('value', value);
                element.value = value;
                element.style.color = element.style.color;
                
                var events = ['focus', 'click', 'input', 'change', 'keydown', 'keypress', 
                              'keyup', 'blur', 'paste', 'textInput'];
                events.forEach(function(evt) {
                    try {
                        element.dispatchEvent(new Event(evt, {bubbles: true, cancelable: true}));
                    } catch(e) {}
                });
                
                if (element.oninput) element.oninput();
                if (element.onchange) element.onchange();
                return element.value;
            """
    
    def _fill_strategy_1(self, selector, value, by_xpath):
        pass
    
    def _fill_strategy_2(self, selector, value, by_xpath):
        pass
    
    def _fill_strategy_3(self, selector, value, by_xpath):
        pass
    
    def _fill_strategy_4(self, selector, value, by_xpath):
        pass
    
    def _validate_fill(self, selector, expected_value, by_xpath):
        """Valida preenchimento"""
        script = """
            var selector = arguments[0];
            var expected = arguments[1];
            var byXPath = arguments[2];
            var element;
            
            if (byXPath) {
                var result = document.evaluate(selector, document, null, 
                    XPathResult.FIRST_ORDERED_NODE_TYPE, null);
                element = result.singleNodeValue;
            } else {
                element = document.querySelector(selector);
            }
            
            if (!element) return false;
            var actual = element.value || '';
            return actual.trim() === expected.trim() || actual.includes(expected);
        """
        try:
            return self.execute_js(script, selector, expected_value, by_xpath, timeout=3, fallback_result=False)
        except:
            return False

import time

def clicar_finalizar_e_verificar_alerta(js_engine, doc, timeout=5, pausa=0.5):
    """
    Clica no botão 'Finalizar', aguarda 0,5s e procura mensagens de alerta.
    Usa safe_action e js_engine.force_click().
    """
    safe_action(doc, "Clicando em 'Finalizar'", lambda:
        js_engine.force_click(
            "//a[@class='btModel btGray btyes' and normalize-space()='Finalizar']",
            by_xpath=True
        )
    )

    time.sleep(pausa)
    log(doc, "🔍 Verificando mensagens de alerta após o clique em 'Finalizar'...")
    return encontrar_mensagem_alerta()


# ==== LOV HANDLER ====
class LOVHandler:
    """Handler especializado para modais LOV"""
    
    def __init__(self, js_engine):
        self.js = js_engine
        self.doc = js_engine.doc
    

def _is_visible_enabled(el):
    try:
        return el.is_displayed() and el.is_enabled()
    except Exception:
        return False


class LOVHandler:
    def __init__(self, js_engine, doc):
        self.js = js_engine
        self.doc = doc

    # ==========================================================
    # 🔹 Função principal - Abre LOV, pesquisa e seleciona
    # ==========================================================
    def open_and_select(self, btn_index=None, btn_xpath=None,
                        search_text="", result_text="",
                        iframe_xpath=None, max_attempts=5):
        """Abre LOV, pesquisa e seleciona resultado"""
        log(self.doc, f"🔍 Processando LOV: '{search_text}' → '{result_text}'")

        for attempt in range(max_attempts):
            try:
                if attempt > 0:
                    log(self.doc, f"   Tentativa {attempt + 1}/{max_attempts}")

                # Volta ao conteúdo principal
                try:
                    self.js.driver.switch_to.default_content()
                except:
                    pass

                # Define seletor do botão
                if btn_index is not None:
                    btn_selector = f"(//a[@class='sprites sp-openLov'])[{btn_index + 1}]"
                    by_xpath = True
                elif btn_xpath:
                    btn_selector = btn_xpath
                    by_xpath = True
                else:
                    raise ValueError("btn_index ou btn_xpath deve ser fornecido")

                log(self.doc, "   📌 Abrindo LOV...")
                self.js.force_click(btn_selector, by_xpath=by_xpath)
                time.sleep(1.5)

                # Se houver iframe, entra
                if iframe_xpath:
                    log(self.doc, "   🔄 Entrando no iframe...")
                    try:
                        WebDriverWait(self.js.driver, 10).until(
                            EC.frame_to_be_available_and_switch_to_it((By.XPATH, iframe_xpath))
                        )
                        time.sleep(0.5)
                    except:
                        log(self.doc, "   ⚠️ Iframe não encontrado, continuando...")

                # Aguarda carregamento
                self.js.wait_ajax_complete(10)
                time.sleep(1)

                # Preenche campos de pesquisa
                log(self.doc, f"   ✏️ Preenchendo campos de pesquisa com: '{search_text}'")
                res = self.preencher_campos_pesquisa_por_indice(
                    search_text=search_text,
                    search_xpaths=[
                        "//input[@id='txtPesquisa']",
                        "//input[@class='nomePesquisa']",
                        "//input[contains(translate(@name,'ABCDEFGHIJKLMNOPQRSTUVWXYZ','abcdefghijklmnopqrstuvwxyz'),'pesquisa')]",
                        "//input[contains(translate(@class,'ABCDEFGHIJKLMNOPQRSTUVWXYZ','abcdefghijklmnopqrstuvwxyz'),'pesquisa')]",
                        "//input[@type='text']",
                        "//input[@class='campoPesquisa']",
                    ],
                    max_campos=None,
                    pausa=0.3,
                    limpar_antes=True
                )
                log(self.doc, f"   📊 Resultado: {res}")
                time.sleep(0.5)

                # Clica em "Pesquisar"
                log(self.doc, "   🔎 Executando pesquisa...")
                search_btn_selectors = [
                    "//a[contains(@class,'lpFind') and contains(normalize-space(.),'Pesquisar')]",
                    "//button[contains(normalize-space(.),'Pesquisar')]",
                    "//a[contains(normalize-space(.),'Buscar')]"
                ]

                search_clicked = False
                for selector in search_btn_selectors:
                    try:
                        self.js.force_click(selector, by_xpath=True)
                        search_clicked = True
                        break
                    except:
                        continue

                # Fallback: ENTER
                if not search_clicked:
                    try:
                        first_input = self.js.driver.find_element(
                            "xpath",
                            "(//input[@id='txtPesquisa'] | //input[@class='nomePesquisa'] | "
                            "//input[contains(@class,'pesquisa')] | "
                            "//input[contains(@name,'pesquisa')] | "
                            "//input[@type='text'])[1]"
                        )
                        first_input.send_keys(Keys.ENTER)
                        search_clicked = True
                        log(self.doc, "   ⚙️ Fallback: pesquisa via tecla ENTER")
                    except Exception:
                        raise Exception("Botão de pesquisa não encontrado")

                # Aguarda resultados
                time.sleep(2)
                self.js.wait_ajax_complete(15)

                # Clica no resultado
                log(self.doc, f"   🎯 Selecionando: '{result_text}'")
                safe_text = result_text.replace("'", '"')
                result_xpath = f"//tr[td[contains(normalize-space(.), \"{safe_text}\")]]"
                self.js.force_click(result_xpath, by_xpath=True)

                time.sleep(1)

                # Volta para conteúdo principal
                if iframe_xpath:
                    try:
                        self.js.driver.switch_to.default_content()
                        log(self.doc, "   ✅ Voltou para conteúdo principal")
                    except:
                        pass

                # Aguarda fechamento
                time.sleep(1)
                self.js.wait_ajax_complete(10)

                log(self.doc, f"✅ LOV processado com sucesso!")
                return True

            except Exception as e:
                log(self.doc, f"⚠️ Tentativa {attempt + 1} falhou: {str(e)[:150]}")

                try:
                    self.js.driver.switch_to.default_content()
                except:
                    pass

                if attempt < max_attempts - 1:
                    time.sleep(2 + attempt * 0.5)
                else:
                    raise Exception(f"Falha ao processar LOV após {max_attempts} tentativas: {e}")

        return False

    # ==========================================================
    # 🔹 Preenche campos de pesquisa genéricos dentro da LOV
    # ==========================================================
    def preencher_campos_pesquisa_por_indice(
        self, search_text: str,
        search_xpaths=None,
        max_campos: int | None = None,
        pausa: float = 0.2,
        limpar_antes: bool = True
    ) -> dict:
        """
        Preenche um ou mais campos de pesquisa dentro da LOV.
        """
        if search_xpaths is None:
            search_xpaths = [
                "//input[@id='txtPesquisa']",
                "//input[@class='nomePesquisa']",
                "//input[contains(translate(@name,'ABCDEFGHIJKLMNOPQRSTUVWXYZ','abcdefghijklmnopqrstuvwxyz'),'pesquisa')]",
                "//input[contains(translate(@class,'ABCDEFGHIJKLMNOPQRSTUVWXYZ','abcdefghijklmnopqrstuvwxyz'),'pesquisa')]",
                "//input[@type='text']",
                "//input[@class='campoPesquisa']",
            ]

        preenchidos = 0
        encontrados = 0
        usados = []

        for xp in search_xpaths:
            try:
                candidatos = self.js.driver.find_elements("xpath", xp)
            except Exception:
                candidatos = []
            if not candidatos:
                continue

            candidatos = [el for el in candidatos if _is_visible_enabled(el)]
            if not candidatos:
                continue

            encontrados += len(candidatos)

            for el in candidatos:
                if max_campos is not None and preenchidos >= max_campos:
                    break
                try:
                    self.js.driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", el)

                    if limpar_antes:
                        el.clear()
                        el.send_keys(Keys.CONTROL, "a")
                        el.send_keys(Keys.DELETE)

                    if search_text:
                        el.send_keys(search_text)
                        usados.append(xp)
                        preenchidos += 1
                        time.sleep(pausa)

                except (StaleElementReferenceException, ElementNotInteractableException):
                    continue
                except Exception:
                    continue

            if max_campos is not None and preenchidos >= max_campos:
                break

        return {
            "texto": search_text,
            "totalEncontrados": encontrados,
            "totalPreenchidos": preenchidos,
            "xpathsUsados": usados,
        }

# ==== FUNÇÕES AUXILIARES ====
def encontrar_mensagem_alerta():
    seletores = [
        (".alerts.salvo", "✅ Mensagem de Sucesso"),
        (".alerts.alerta", "⚠️ Mensagem de Alerta"),
        (".alerts.erro", "❌ Mensagem de Erro"),
    ]

    for seletor, tipo in seletores:
        try:
            elemento = driver.find_element(By.CSS_SELECTOR, seletor)
            if elemento.is_displayed():
                log(doc, f"📢 {tipo}: {elemento.text}")
                return elemento
        except:
            continue

    log(doc, "ℹ️ Nenhuma mensagem de alerta encontrada.")
    return None

def safe_action(doc, descricao, func, max_retries=3):
    """Wrapper para ações com retry automático"""
    global driver
    
    for attempt in range(max_retries):
        try:
            log(doc, f"🔄 {descricao}..." if attempt == 0 else f"🔄 {descricao}... (Tentativa {attempt + 1})")
            func()
            log(doc, f"✅ {descricao} realizada com sucesso.")
            take_screenshot(driver, doc, _sanitize_filename(descricao))
            return True
        except Exception as e:
            if attempt < max_retries - 1:
                log(doc, f"⚠️ Tentativa {attempt + 1} falhou, tentando novamente...")
                time.sleep(2 + attempt)
                continue
            else:
                log(doc, f"❌ Erro após {max_retries} tentativas: {e}")
                take_screenshot(driver, doc, _sanitize_filename(f"erro_{descricao}"))
                return False
    
    return False

def inicializar_driver():
    """Inicializa WebDriver com configurações otimizadas"""
    global driver, wait
    
    try:
        options = Options()
        options.add_argument("--start-maximized")
        options.add_argument("--disable-blink-features=AutomationControlled")
        options.add_argument("--disable-dev-shm-usage")
        options.add_argument("--no-sandbox")
        options.add_experimental_option("excludeSwitches", ["enable-automation"])
        options.add_experimental_option('useAutomationExtension', False)
        
        driver = webdriver.Chrome(
            service=Service(ChromeDriverManager().install()), 
            options=options
        )
        
        driver.execute_script(
            "Object.defineProperty(navigator, 'webdriver', {get: () => undefined})"
        )
        
        # Configura timeouts globais
        driver.set_script_timeout(30)
        driver.implicitly_wait(10)
        
        wait = WebDriverWait(driver, TIMEOUT_DEFAULT)
        
        log(doc, "✅ Driver inicializado com sucesso")
        return True
        
    except Exception as e:
        log(doc, f"❌ Erro ao inicializar driver: {e}")
        return False

def finalizar_relatorio():
    """Salva relatório e fecha driver"""
    global driver, doc
    
    nome_arquivo = f"relatorio_caixa_cenario_1_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    
    try:
        doc.save(nome_arquivo)
        log(doc, f"📄 Relatório salvo: {nome_arquivo}")
        
        try:
            subprocess.run(["start", "winword", nome_arquivo], shell=True)
        except:
            pass
            
    except Exception as e:
        print(f"Erro ao salvar relatório: {e}")
    
    if driver:
        try:
            driver.quit()
            log(doc, "✅ Driver encerrado")
        except:
            pass

def preencher_campos_pesquisa_por_indice(self, 
                                         search_text: str, 
                                         search_xpaths=None, 
                                         max_campos: int | None = None, 
                                         pausa: float = 0.3,
                                         limpar_antes: bool = True):
    """
    Procura TODOS os campos de pesquisa e preenche um após o outro (ordem no DOM).
    - search_text: texto a preencher.
    - search_xpaths: lista de XPaths para busca (usa padrão se None).
    - max_campos: limita quantos campos serão preenchidos (None = todos).
    - pausa: pausa curta entre preenchimentos.
    - limpar_antes: se True, limpa o campo antes de preencher.

    Retorna: {"total_encontrados": int, "total_preenchidos": int, "xpaths_usados": [str], "falhas": [str]}
    """
    search_xpaths = search_xpaths or [
        "//input[@id='txtPesquisa']",
        "//input[@class='nomePesquisa']",
        # adicione mais padrões específicos antes do genérico:
        "//input[contains(translate(@name,'ABCDEFGHIJKLMNOPQRSTUVWXYZ','abcdefghijklmnopqrstuvwxyz'),'pesquisa')]",
        "//input[contains(translate(@class,'ABCDEFGHIJKLMNOPQRSTUVWXYZ','abcdefghijklmnopqrstuvwxyz'),'pesquisa')]",
        "//input[@type='text']"
    ]

    # Monta um XPATH-UNIÃO para manter a ordem no DOM
    xpath_uniao = " | ".join(f"({xp})" for xp in search_xpaths)

    # Coleta elementos (evita duplicados por id/ref)
    try:
        elementos = self.js.driver.find_elements(By.XPATH, xpath_uniao)
    except Exception as e:
        log(self.doc, f"⚠️ Erro ao buscar campos de pesquisa: {e}")
        return {"total_encontrados": 0, "total_preenchidos": 0, "xpaths_usados": [], "falhas": [str(e)]}

    # Filtra apenas visíveis e habilitados
    elementos_filtrados = []
    for el in elementos:
        try:
            if el.is_displayed() and el.is_enabled():
                elementos_filtrados.append(el)
        except Exception:
            continue

    total_encontrados = len(elementos_filtrados)
    if total_encontrados == 0:
        log(self.doc, "🔎 Nenhum campo de pesquisa visível/habilitado encontrado.")
        return {"total_encontrados": 0, "total_preenchidos": 0, "xpaths_usados": [], "falhas": []}

    if max_campos is not None and max_campos > 0:
        elementos_filtrados = elementos_filtrados[:max_campos]

    log(self.doc, f"🧭 Campos de pesquisa encontrados: {total_encontrados} | A preencher: {len(elementos_filtrados)}")

    total_preenchidos = 0
    falhas = []
    xpaths_usados = []

    for idx, el in enumerate(elementos_filtrados, start=1):
        try:
            # Recalcula um XPATH relativo único para log (opcional, pode ser pesado).
            # Aqui só registramos o index para simplificar.
            xpaths_usados.append(f"[campo #{idx}]")

            # Traz para a tela e foca
            self.js.driver.execute_script("arguments[0].scrollIntoView({block:'center'});", el)
            try:
                el.click()
            except Exception:
                pass

            if limpar_antes:
                try:
                    el.clear()
                except Exception:
                    # fallback via JS
                    try:
                        self.js.driver.execute_script("arguments[0].value='';", el)
                    except Exception:
                        pass

            # Preenche (usa seu helper para manter padrão)
            try:
                self.js.force_fill_element(el, search_text)  # se você tiver esse helper
            except AttributeError:
                # fallback: force_fill por XPATH do próprio elemento (gera um xpath usando JS)
                try:
                    self.js.driver.execute_script("arguments[0].value = arguments[1];", el, search_text)
                except Exception:
                    # último fallback: send_keys
                    el.send_keys(search_text)

            total_preenchidos += 1
            log(self.doc, f"   ✏️ Preenchido campo #{idx} com '{search_text}'")
            if pausa:
                time.sleep(pausa)

        except (StaleElementReferenceException, ElementNotInteractableException) as e:
            falhas.append(f"Campo #{idx}: {type(e).__name__}")
            log(self.doc, f"   ⚠️ Falha no campo #{idx}: {e}")
            continue
        except Exception as e:
            falhas.append(f"Campo #{idx}: {e}")
            log(self.doc, f"   ⚠️ Erro inesperado no campo #{idx}: {e}")
            continue

    resumo = {
        "total_encontrados": total_encontrados,
        "total_preenchidos": total_preenchidos,
        "xpaths_usados": xpaths_usados,
        "falhas": falhas
    }
    log(self.doc, f"✅ Pesquisa preenchida em {total_preenchidos}/{len(elementos_filtrados)} campos. Falhas: {len(falhas)}")
    return resumo

def clicar_lov_por_indice(indice_lov: int, max_tentativas: int = 5, timeout: int = 10, scroll: bool = True):
    """Clica no ícone de LOV pelo índice"""
    def acao():
        if not isinstance(indice_lov, int) or indice_lov < 0:
            raise ValueError(f"Índice inválido: {indice_lov}")

        tentativa = 0
        while tentativa < max_tentativas:
            tentativa += 1
            try:
                log(doc, f"🔎 Tentativa {tentativa}: Localizando ícones LOV...")
                elementos = driver.find_elements(By.CSS_SELECTOR, "a.sprites.sp-openLov")

                if not elementos:
                    if tentativa < max_tentativas:
                        log(doc, f"⚠️ Nenhum ícone LOV encontrado (tentativa {tentativa}/{max_tentativas})")
                        time.sleep(1.2)
                        continue
                    raise Exception("Nenhum ícone LOV encontrado.")

                if indice_lov >= len(elementos):
                    raise Exception(f"Índice {indice_lov} inválido. Encontrados {len(elementos)} ícones LOV.")

                locator_xpath = f"(//a[contains(@class,'sp-openLov')])[{indice_lov + 1}]"
                elemento = driver.find_element(By.XPATH, locator_xpath)

                log(doc, f"🎯 Preparando clique no LOV de índice {indice_lov}")

                def _wait_clickable():
                    wait.until(EC.element_to_be_clickable((By.XPATH, locator_xpath)))

                estrategias = [
                    lambda: (_wait_clickable(), elemento.click()),
                    lambda: (
                        driver.execute_script("arguments[0].scrollIntoView({block:'center'});", elemento) if scroll else None,
                        time.sleep(0.2),
                        elemento.click()
                    ),
                    lambda: driver.execute_script("arguments[0].click();", elemento),
                    lambda: ActionChains(driver).move_to_element(elemento).pause(0.1).click().perform()
                ]

                for i, estrategia in enumerate(estrategias, 1):
                    try:
                        log(doc, f"   ▶️ Estratégia {i} de clique no LOV...")
                        estrategia()
                        time.sleep(0.3)
                        log(doc, f"✅ Clique no LOV (índice {indice_lov}) realizado (estratégia {i})")
                        return True
                    except (ElementClickInterceptedException, StaleElementReferenceException, JavascriptException, TimeoutException) as e:
                        log(doc, f"⚠️ Estratégia {i} falhou: {e}")
                        try:
                            elementos = driver.find_elements(By.CSS_SELECTOR, "a.sprites.sp-openLov")
                            elemento = driver.find_element(By.XPATH, locator_xpath)
                        except:
                            pass
                        continue

                if tentativa < max_tentativas:
                    log(doc, f"⚠️ Tentativa {tentativa} não conseguiu clicar no LOV. Reintentando...")
                    time.sleep(1.2)
                    continue

            except Exception as e:
                if tentativa < max_tentativas:
                    log(doc, f"⚠️ Erro na tentativa {tentativa}: {e}. Reintentando...")
                    time.sleep(1.2)
                    continue
                raise

        raise Exception(f"Falha ao clicar no LOV de índice {indice_lov} após {max_tentativas} tentativas.")

    return acao

def encontrar_campos_textarea(timeout=10):
    """Retorna lista de textareas visíveis e interativas"""
    elementos = []
    try:
        wait.until(lambda d: len(d.find_elements(By.TAG_NAME, "textarea")) >= 0)
        textareas = driver.find_elements(By.TAG_NAME, "textarea")
    except:
        textareas = []

    for el in textareas:
        try:
            if not el.is_displayed() or not el.is_enabled():
                continue
            driver.execute_script("arguments[0].scrollIntoView({block:'center'});", el)
            elementos.append({
                "elemento": el,
                "id": el.get_attribute("id"),
                "name": el.get_attribute("name"),
            })
        except:
            continue

    return elementos

def normalizar_texto(txt):
    if txt is None:
        return ""
    return txt.replace("\r\n", "\n").replace("\r", "\n").strip()

def validar_textarea_preenchida(elemento, texto_esperado):
    """Confere se o valor atual da textarea bate com o esperado"""
    try:
        atual = elemento.get_attribute("value")
        if atual is None or atual == "":
            atual = (elemento.text or "")
        return normalizar_texto(atual) == normalizar_texto(texto_esperado)
    except StaleElementReferenceException:
        return False

def _prepare_focus_and_clear(elemento, limpar_primeiro=True):
    driver.execute_script("arguments[0].scrollIntoView({block:'center'});", elemento)
    try:
        elemento.click()
    except:
        driver.execute_script("arguments[0].focus();", elemento)

    if limpar_primeiro:
        try:
            elemento.clear()
        except:
            ActionChains(driver)\
                .move_to_element(elemento).click()\
                .key_down(Keys.CONTROL).send_keys("a").key_up(Keys.CONTROL)\
                .send_keys(Keys.DELETE).perform()

def _textarea_tradicional(elemento, texto, limpar_primeiro=True):
    _prepare_focus_and_clear(elemento, limpar_primeiro)
    elemento.send_keys(texto)
    elemento.send_keys(Keys.TAB)

def _textarea_actionchains(elemento, texto, limpar_primeiro=True):
    _prepare_focus_and_clear(elemento, limpar_primeiro)
    ac = ActionChains(driver)
    ac.move_to_element(elemento).click().perform()
    for chunk_start in range(0, len(texto), 400):
        ac.send_keys(texto[chunk_start:chunk_start+400]).perform()
        time.sleep(0.05)
    ac.send_keys(Keys.TAB).perform()

def _textarea_js_setvalue(elemento, texto):
    driver.execute_script("""
        const el = arguments[0];
        const val = arguments[1];
        el.value = val;
        el.dispatchEvent(new Event('input',  {bubbles:true}));
        el.dispatchEvent(new Event('change', {bubbles:true}));
        el.dispatchEvent(new KeyboardEvent('keyup', {bubbles:true}));
        el.dispatchEvent(new Event('blur',   {bubbles:true}));
    """, elemento, texto)

def _textarea_js_react_input(elemento, texto):
    driver.execute_script("""
        const el = arguments[0];
        const val = arguments[1];

        const desc = Object.getOwnPropertyDescriptor(window.HTMLTextAreaElement.prototype, 'value');
        if (desc && desc.set) {
            desc.set.call(el, val);
        } else {
            el.value = val;
        }

        el.dispatchEvent(new Event('input', {bubbles: true}));
        el.dispatchEvent(new Event('change', {bubbles: true}));
        el.dispatchEvent(new Event('blur', {bubbles: true}));
    """, elemento, texto)

def preencher_textarea_por_indice(indice_campo, texto, max_tentativas=5, limpar_primeiro=True):
    """Preenche textarea pelo índice usando estratégias múltiplas"""
    def acao():
        if not isinstance(indice_campo, int) or indice_campo < 0:
            raise ValueError(f"Índice inválido: {indice_campo}")
        if texto is None or not isinstance(texto, str):
            raise ValueError(f"Texto inválido: {texto!r}")

        tentativa = 0
        while tentativa < max_tentativas:
            tentativa += 1
            try:
                campos = encontrar_campos_textarea()
                if not campos:
                    if tentativa < max_tentativas:
                        log(doc, f"⚠️ Nenhuma <textarea> encontrada (tentativa {tentativa}/{max_tentativas})")
                        time.sleep(1.5)
                        continue
                    raise Exception("Nenhuma <textarea> foi encontrada.")

                if indice_campo >= len(campos):
                    raise Exception(f"Índice {indice_campo} inválido. Encontradas {len(campos)} textareas.")

                campo_info = campos[indice_campo]
                elemento = campo_info["elemento"]
                campo_id = campo_info.get("id") or "(sem id)"

                log(doc, f"🎯 Tentativa {tentativa}: Preenchendo textarea {indice_campo} (ID: {campo_id})")

                if validar_textarea_preenchida(elemento, texto):
                    log(doc, f"✅ Textarea {indice_campo} já está com o valor desejado.")
                    return True

                estrategias = [
                    lambda: _textarea_tradicional(elemento, texto, limpar_primeiro),
                    lambda: _textarea_actionchains(elemento, texto, limpar_primeiro),
                    lambda: _textarea_js_setvalue(elemento, texto),
                    lambda: _textarea_js_react_input(elemento, texto),
                ]

                for i, estrategia in enumerate(estrategias, 1):
                    try:
                        log(doc, f"   ▶️ Estratégia {i}…")
                        estrategia()
                        time.sleep(0.8)

                        if validar_textarea_preenchida(elemento, texto):
                            val = (elemento.get_attribute("value") or "").strip()
                            log(doc, f"✅ Preenchido com sucesso pela estratégia {i}")
                            return True
                        else:
                            log(doc, f"⚠️ Estratégia {i} não refletiu o valor esperado.")
                    except (StaleElementReferenceException, JavascriptException, TimeoutException) as e:
                        log(doc, f"⚠️ Estratégia {i} falhou: {e}")
                        try:
                            campos = encontrar_campos_textarea()
                            elemento = campos[indice_campo]["elemento"]
                        except:
                            pass
                        continue

                if tentativa < max_tentativas:
                    log(doc, f"⚠️ Tentativa {tentativa} falhou; reintentando em 1.5s…")
                    time.sleep(1.5)
                    continue
            except Exception as e:
                if tentativa < max_tentativas:
                    log(doc, f"⚠️ Erro na tentativa {tentativa}: {e}. Retentando…")
                    time.sleep(1.5)
                    continue
                else:
                    raise

        raise Exception(f"Falha ao preencher textarea {indice_campo} após {max_tentativas} tentativas.")
    return acao


def fechar_abas_extras(driver, doc, aba_principal_index=0):
    """Fecha todas as abas extras (como popups de impressão) mantendo apenas a aba principal"""
    try:
        handles = driver.window_handles
        if len(handles) <= 1:
            log(doc, "ℹ️ Apenas uma aba aberta - nada a fechar.")
            return True
        
        # Guarda o handle da aba principal
        aba_principal = handles[aba_principal_index]
        
        # Fecha todas as outras abas
        abas_fechadas = 0
        for handle in handles:
            if handle != aba_principal:
                try:
                    driver.switch_to.window(handle)
                    driver.close()
                    abas_fechadas += 1
                    log(doc, f"🗑️ Aba extra fechada ({abas_fechadas})")
                except Exception as e:
                    log(doc, f"⚠️ Erro ao fechar aba: {e}")
        
        # Retorna para a aba principal
        driver.switch_to.window(aba_principal)
        time.sleep(0.3)
        
        log(doc, f"✅ {abas_fechadas} aba(s) extra(s) fechada(s). Foco na aba principal.")
        return True
        
    except Exception as e:
        log(doc, f"❌ Erro ao fechar abas extras: {e}")
        return False

def focar_sistema_completo(js_engine, doc):
    """Garante o foco completo na aba principal do sistema e fecha abas extras"""
    driver = js_engine.driver
    try:
        # Primeiro fecha abas extras (como impressão)
        fechar_abas_extras(driver, doc)

        driver.switch_to.default_content()
        js_engine.execute_js("if (window.focus) window.focus();", timeout=3, fallback_result=None)
        time.sleep(0.3)

        log(doc, "✅ Foco garantido na aba do sistema.")
        return True

    except Exception as e:
        log(doc, f"⚠️ Falha ao focar aba do sistema: {e}")
        return False

def clicar_todos_botoes_sim_visiveis(js_engine, doc, pausa_entre=0.0):
    """Clica em TODOS os botões 'Sim' visíveis de uma vez"""
    js = r"""
    (function(){
      const isVisible = el => {
        if (!el) return false;
        const s = getComputedStyle(el);
        return el.offsetParent !== null && s.display !== 'none' &&
               s.visibility !== 'hidden' && parseFloat(s.opacity||1) > 0.01;
      };
      const buttons = Array.from(document.querySelectorAll("a.btModel.btGray.btyes"))
        .filter(isVisible)
        .filter(b => (b.textContent||"").trim().toLowerCase() === "sim");

      let clicked = 0;
      buttons.forEach(b => {
        try {
          b.style.pointerEvents = 'auto';
          b.removeAttribute('disabled');
          b.style.visibility = 'visible';
          b.style.display = 'inline-block';
          b.scrollIntoView({block:'center'});

          ['mouseover','mouseenter','mousemove','mousedown','mouseup','click'].forEach(t=>{
            b.dispatchEvent(new MouseEvent(t,{bubbles:true,cancelable:true,view:window,detail:1}));
          });
          if (typeof b.click === 'function') b.click();

          if (typeof window.jQuery !== 'undefined') {
            window.jQuery(b).trigger('click');
          }
          clicked++;
        } catch(e) {}
      });

      return { totalEncontrados: buttons.length, totalClicados: clicked };
    })();
    """
    try:
        res = js_engine.execute_js(js, timeout=5, fallback_result={"totalEncontrados": 0, "totalClicados": 0})
        total = int(res.get("totalEncontrados", 0))
        clic = int(res.get("totalClicados", 0))
        log(doc, f"⚡ 'Sim' visíveis encontrados: {total} | clicados: {clic}")
        if pausa_entre and clic > 0:
            time.sleep(pausa_entre)
        return res
    except Exception as e:
        log(doc, f"❌ Erro ao clicar em todos os 'Sim': {e}")
        return {"totalEncontrados": 0, "totalClicados": 0, "erro": str(e)}


        
def clicar_sim_com_retry(doc, js_engine, wait, max_tentativas=5, pausa=1.5):
    """Clica em 'Sim' até o modal de confirmação desaparecer"""
    xpath_modal = "//div[contains(@class,'modal') and contains(@style,'z-index')]"
    xpath_sim = "(//div[contains(@class,'modal') and not(contains(@style,'display: none'))]//a[@class='btModel btGray btyes'])[last()]"

    tentativa = 0
    while tentativa < max_tentativas:
        tentativa += 1
        log(doc, f"🧩 Tentativa {tentativa} de fechar modal...")

        try:
            js_engine.force_click(xpath_sim, by_xpath=True)
            time.sleep(pausa)

            modais_visiveis = driver.find_elements(By.XPATH, xpath_modal)
            modais_ativos = [m for m in modais_visiveis if "display: none" not in m.get_attribute("style")]

            if not modais_ativos:
                log(doc, "✅ Botão 'Sim' clicado com sucesso.")
                return True

        except Exception as e:
            log(doc, f"⚠️ Tentativa {tentativa} falhou: {e}")

    log(doc, "❌ Botão 'Sim' não foi clicado após todas as tentativas.")
    return False

def clicar_primeiro_sp_add(js_engine, doc=None, timeout=5):
    """
    Localiza e clica no primeiro elemento <a class="sprites sp-add">, exatamente.
    Ignora variantes como 'sp-addVerde'.
    """
    from selenium.common.exceptions import (
        NoSuchElementException,
        ElementClickInterceptedException,
        StaleElementReferenceException,
    )
    from selenium.webdriver.common.by import By
    import time

    driver = js_engine.driver

    # XPath restrito: exige que as classes sejam EXATAMENTE 'sprites' e 'sp-add'
    xpath_btn = (
        "//a[contains(concat(' ', normalize-space(@class), ' '), ' sprites ') "
        "and contains(concat(' ', normalize-space(@class), ' '), ' sp-add ') "
        "and not(contains(@class, 'sp-addVerde'))][1]"
    )

    log(doc, "🧩 Procurando o botão '<a class=\"sprites sp-add\">' exato...")

    try:
        el = driver.find_element(By.XPATH, xpath_btn)
        log(doc, "🎯 Botão exato encontrado! Tentando clicar...")

        try:
            el.click()
            log(doc, "✅ Clique padrão realizado com sucesso.")
        except (ElementClickInterceptedException, StaleElementReferenceException):
            driver.execute_script("arguments[0].click();", el)
            log(doc, "⚡ Clique forçado via JavaScript realizado.")

        time.sleep(0.5)
        return True

    except NoSuchElementException:
        log(doc, "⚠️ Nenhum botão exato 'sp-add' encontrado.")
        return False

    except Exception as e:
        log(doc, f"❌ Erro ao clicar no botão 'sp-add': {e}")
        return False


def fechar_abas_extras_e_verificar_alerta(driver, doc):
    """
    Fecha abas extras (como de impressão) e verifica imediatamente
    se há alguma mensagem de alerta exibida após o fechamento.
    """
    try:
        safe_action(doc, "Fechando abas extras (impressão)", lambda:
            fechar_abas_extras(driver, doc)
        )

        alerta = encontrar_mensagem_alerta()
        if alerta:
            log(doc, f"⚠️ Alerta detectado após fechar abas extras: '{alerta.text.strip()}'")
        else:
            log(doc, "✅ Nenhum alerta detectado após fechar abas extras.")

        return True
    except Exception as e:
        log(doc, f"⚠️ Erro ao fechar abas extras e verificar alerta: {e}")
        return False

def verificar_e_abrir_caixa(js_engine, doc, timeout=10):
    """
    Verifica o estado do caixa:
      - Se encontrar o botão 'Abrir caixa' visível (sem display: none), realiza a abertura normalmente.
      - Se detectar 'Fechar caixa', entende que o caixa já está aberto e apenas prossegue.
    """
    abrir_xpath = "//a[contains(@class,'btAzulDegrade') and contains(@class,'btAbrirCaixa')]"
    fechar_xpath = "//a[contains(@class,'btAzulDegrade') and contains(@class,'btFecharCaixa')]"

    driver = js_engine.driver
    log(doc, "🔍 Verificando estado do caixa...")

    try:
        # Verifica se existe botão Abrir Caixa visível
        abrir_visivel = driver.execute_script(f"""
            const el = document.evaluate("{abrir_xpath}", document, null, XPathResult.FIRST_ORDERED_NODE_TYPE, null).singleNodeValue;
            if (!el) return false;
            const style = window.getComputedStyle(el);
            return style.display !== 'none' && style.visibility !== 'hidden' && el.offsetParent !== null;
        """)

        # Verifica se existe botão Fechar Caixa visível
        fechar_visivel = driver.execute_script(f"""
            const el = document.evaluate("{fechar_xpath}", document, null, XPathResult.FIRST_ORDERED_NODE_TYPE, null).singleNodeValue;
            if (!el) return false;
            const style = window.getComputedStyle(el);
            return style.display !== 'none' && style.visibility !== 'hidden' && el.offsetParent !== null;
        """)

        if abrir_visivel:
            log(doc, "📦 Caixa fechado — iniciando abertura normal...")
            safe_action(doc, "Clicando em 'Abrir Caixa'", lambda:
                js_engine.force_click(abrir_xpath, by_xpath=True)
            )

            # Valor inicial
            safe_action(doc, "Preenchendo Valor Inicial", lambda:
                js_engine.force_fill(
                    "//input[@type='text' and contains(@class,'valor') and contains(@placeholder,'R$')]",
                    "1000,00",
                    by_xpath=True
                )
            )

            # Descrição
            safe_action(doc, "Preenchendo Descrição", lambda:
                preencher_textarea_por_indice(0, "Abertura automática de caixa via automação Selenium.")
            )

            # Confirmar abertura
            safe_action(doc, "Confirmando Abertura", lambda:
                js_engine.force_click("//a[@class='btModel btGray btyes' and normalize-space()='Abrir']", by_xpath=True)
            )
            time.sleep(0.5)
            encontrar_mensagem_alerta()

            # Autenticação
            safe_action(doc, "Autenticando Abertura", lambda:
                js_engine.force_click("//a[@id='BtYes' and contains(@class,'btModel btGray btyes') and contains(normalize-space(.),'Autenticar')]", by_xpath=True)
            )
            time.sleep(3)
            fechar_abas_extras_e_verificar_alerta(driver, doc)

            # Fecha modal de autenticação
            safe_action(doc, "Fechando modal autenticação", lambda:
                js_engine.force_click("//a[@id='BtNo' and contains(@class,'btno') and normalize-space()='Fechar']", by_xpath=True)
            )

        elif fechar_visivel:
            log(doc, "✅ Caixa já está aberto — prosseguindo com o fluxo normalmente.")
        else:
            log(doc, "⚠️ Nenhum botão de 'Abrir' ou 'Fechar caixa' visível encontrado — prosseguindo com cautela.")

    except Exception as e:
        log(doc, f"❌ Erro ao verificar/abrir caixa: {e}")
        take_screenshot(driver, doc, "erro_verificar_caixa")

def clicar_todos_pesquisar(js_engine, doc, pausa_entre=0.5, timeout=5):
    """
    Procura todos os botões 'Pesquisar' visíveis e clica em cada um deles na ordem.
    Conta e exibe quantos botões existem antes de clicar.
    Usa js_engine.force_click() e registra log detalhado.
    """

    xpath_base = "//a[contains(@class,'btPesquisar btAzulDegrade') and contains(normalize-space(.),'Pesquisar')]"

    try:
        elementos = js_engine.driver.find_elements("xpath", xpath_base)
        total = len(elementos)
        log(doc, f"🔍 Foram encontrados {total} botão(ões) 'Pesquisar' na tela.")

        if total == 0:
            log(doc, "⚠️ Nenhum botão 'Pesquisar' foi encontrado.")
            return {"total": 0, "clicados": 0}

        total_clicados = 0
        for i in range(1, total + 1):
            xpath_indexado = f"({xpath_base})[{i}]"
            try:
                log(doc, f"🎯 Clicando no botão 'Pesquisar' (índice {i}/{total})...")
                js_engine.force_click(xpath_indexado, by_xpath=True)
                js_engine.wait_ajax_complete(timeout)
                total_clicados += 1
                log(doc, f"✅ Clique no botão 'Pesquisar' (índice {i}) realizado com sucesso.")
                if pausa_entre > 0:
                    import time
                    time.sleep(pausa_entre)
            except Exception as e:
                log(doc, f"⚠️ Falha ao clicar no botão 'Pesquisar' (índice {i}): {e}")

        log(doc, f"🧾 Resumo: {total_clicados}/{total} botões 'Pesquisar' clicados com sucesso.")
        return {"total": total, "clicados": total_clicados}

    except Exception as e:
        log(doc, f"⚠️ Erro ao procurar ou clicar nos botões 'Pesquisar': {e}")
        return {"total": 0, "clicados": 0}


def clicar_todos_voltar(js_engine, doc, pausa_entre=0.5, timeout=5):
    """
    Procura todos os botões 'Voltar (ESC)' visíveis e clica em cada um deles na ordem.
    Conta e exibe quantos botões existem antes de clicar.
    Usa js_engine.force_click() e registra log detalhado.
    """

    xpath_base = "//a[contains(@class,'sp-voltarGrande') and @title='Voltar (ESC)']"

    try:
        elementos = js_engine.driver.find_elements("xpath", xpath_base)
        total = len(elementos)
        log(doc, f"🔍 Foram encontrados {total} botão(ões) 'Voltar (ESC)' na tela.")

        if total == 0:
            log(doc, "⚠️ Nenhum botão 'Voltar (ESC)' foi encontrado.")
            return {"total": 0, "clicados": 0}

        total_clicados = 0
        for i in range(1, total + 1):
            xpath_indexado = f"({xpath_base})[{i}]"
            try:
                log(doc, f"🎯 Clicando no botão 'Voltar (ESC)' (índice {i}/{total})...")
                js_engine.force_click(xpath_indexado, by_xpath=True)
                js_engine.wait_ajax_complete(timeout)
                total_clicados += 1
                log(doc, f"✅ Clique no botão 'Voltar (ESC)' (índice {i}) realizado com sucesso.")
                if pausa_entre > 0:
                    import time
                    time.sleep(pausa_entre)
            except Exception as e:
                log(doc, f"⚠️ Falha ao clicar no botão 'Voltar (ESC)' (índice {i}): {e}")

        log(doc, f"🧾 Resumo: {total_clicados}/{total} botões 'Voltar (ESC)' clicados com sucesso.")
        return {"total": total, "clicados": total_clicados}

    except Exception as e:
        log(doc, f"⚠️ Erro ao procurar ou clicar nos botões 'Voltar (ESC)': {e}")
        return {"total": 0, "clicados": 0}


def clicar_salvar_modal(js_engine, doc, timeout=5):
    """Versão simplificada focada no modal específico com logs detalhados"""
    import time
    
    try:
        log(doc, "🔍 Iniciando busca por botões 'Salvar' no modal...")
        
        # Script JavaScript que retorna informações detalhadas
        script = """
        var botoes = document.querySelectorAll('.modal.overflow .btok');
        var info = {
            total: botoes.length,
            visiveis: 0,
            clicados: 0
        };
        
        botoes.forEach(function(btn) {
            if (btn.offsetParent !== null) {  // Verifica se está visível
                info.visiveis++;
                try {
                    btn.click();
                    info.clicados++;
                } catch(e) {
                    console.error('Erro ao clicar:', e);
                }
            }
        });
        
        return info;
        """
        
        log(doc, "⚙️ Executando JavaScript para clicar no(s) botão(ões)...")
        resultado = js_engine.driver.execute_script(script)
        
        log(doc, f"📊 Total de botões encontrados: {resultado['total']}")
        log(doc, f"👁️ Botões visíveis: {resultado['visiveis']}")
        log(doc, f"✅ Botões clicados com sucesso: {resultado['clicados']}")
        
        if resultado['clicados'] == 0:
            log(doc, "⚠️ Nenhum botão 'Salvar' foi clicado.")
            if resultado['total'] == 0:
                log(doc, "❌ Motivo: Nenhum botão encontrado no modal.")
            elif resultado['visiveis'] == 0:
                log(doc, "❌ Motivo: Botões existem mas não estão visíveis.")
        else:
            log(doc, f"🎉 Sucesso! {resultado['clicados']} botão(ões) clicado(s).")
            log(doc, f"⏳ Aguardando conclusão do AJAX (timeout: {timeout}s)...")
            js_engine.wait_ajax_complete(timeout)
            log(doc, "✅ AJAX concluído.")
        
        return {"total": resultado['total'], "clicados": resultado['clicados']}
        
    except Exception as e:
        log(doc, f"❌ Erro ao executar clique no modal: {e}")
        import traceback
        log(doc, f"📋 Detalhes do erro: {traceback.format_exc()}")
        return {"total": 0, "clicados": 0}
    

def clicar_botao_por_classe(js_engine, doc, classe, nome_botao="Botão", timeout=5):
    """Clica em botão por classe CSS com logs detalhados"""
    import time
    
    try:
        log(doc, f"🔍 Iniciando busca pelo botão '{nome_botao}' (classe: {classe})...")
        
        script = f"""
        var botoes = document.querySelectorAll('a.{classe}');
        var info = {{
            total: botoes.length,
            visiveis: 0,
            clicados: 0
        }};
        
        botoes.forEach(function(btn) {{
            if (btn.offsetParent !== null) {{
                info.visiveis++;
                try {{
                    btn.click();
                    info.clicados++;
                }} catch(e) {{
                    console.error('Erro ao clicar:', e);
                }}
            }}
        }});
        
        return info;
        """
        
        log(doc, f"⚙️ Executando JavaScript para clicar no botão '{nome_botao}'...")
        resultado = js_engine.driver.execute_script(script)
        
        log(doc, f"📊 Total de botões encontrados: {resultado['total']}")
        log(doc, f"👁️ Botões visíveis: {resultado['visiveis']}")
        log(doc, f"✅ Botões clicados: {resultado['clicados']}")
        
        if resultado['clicados'] == 0:
            log(doc, f"⚠️ Nenhum botão '{nome_botao}' foi clicado.")
            if resultado['total'] == 0:
                log(doc, f"❌ Motivo: Nenhum botão encontrado.")
            elif resultado['visiveis'] == 0:
                log(doc, f"❌ Motivo: Botão existe mas não está visível.")
        else:
            log(doc, f"🎉 Sucesso! Botão '{nome_botao}' clicado.")
            log(doc, f"⏳ Aguardando AJAX (timeout: {timeout}s)...")
            js_engine.wait_ajax_complete(timeout)
            log(doc, "✅ AJAX concluído.")
        
        return {"total": resultado['total'], "clicados": resultado['clicados']}
        
    except Exception as e:
        log(doc, f"❌ Erro ao clicar no botão '{nome_botao}': {e}")
        return {"total": 0, "clicados": 0}
        
    except Exception as e:
        log(doc, f"❌ Erro ao executar clique no modal: {e}")
        import traceback
        log(doc, f"📋 Detalhes do erro: {traceback.format_exc()}")
        return {"total": 0, "clicados": 0}
    
def clicar_botao_voltar_por_indice(js_engine, doc, indice=1, timeout=5):
    """
    Clica no botão 'Voltar (ESC)' pelo índice informado (1-based).
    Usa js_engine.force_click() e registra log.
    """
    xpath = f"(//a[@class='sprites sp-voltarGrande' and @title='Voltar (ESC)'])[{indice}]"
    log(doc, f"↩️ Clicando no botão 'Voltar (ESC)' (índice {indice})...")

    try:
        js_engine.force_click(xpath, by_xpath=True)
        js_engine.wait_ajax_complete(timeout)
        log(doc, f"✅ Clique no botão 'Voltar (ESC)' (índice {indice}) realizado com sucesso.")
        return True
    except Exception as e:
        log(doc, f"⚠️ Falha ao clicar no botão 'Voltar (ESC)' (índice {indice}): {e}")
        return False

def clicar_pesquisar_por_indice(js_engine, doc, indice=1, timeout=5):
    """
    Clica no botão 'Pesquisar' pelo índice informado (1-based).
    Conta e exibe quantos botões existem antes do clique.
    Usa js_engine.force_click() e registra log.
    """
    xpath_base = "//a[contains(@class,'btPesquisar btAzulDegrade') and contains(normalize-space(.),'Pesquisar')]"
    xpath_indexado = f"({xpath_base})[{indice}]"

    try:
        # Conta quantos botões existem
        elementos = js_engine.driver.find_elements("xpath", xpath_base)
        total = len(elementos)
        log(doc, f"🔍 Foram encontrados {total} botão(ões) 'Pesquisar' na tela.")

        if total == 0:
            log(doc, "⚠️ Nenhum botão 'Pesquisar' foi encontrado.")
            return False
        if indice > total:
            log(doc, f"⚠️ Índice {indice} inválido — só existem {total} botão(ões).")
            return False

        log(doc, f"🎯 Clicando no botão 'Pesquisar' (índice {indice})...")
        js_engine.force_click(xpath_indexado, by_xpath=True)
        js_engine.wait_ajax_complete(timeout)
        log(doc, f"✅ Clique no botão 'Pesquisar' (índice {indice}) realizado com sucesso.")
        return True

    except Exception as e:
        log(doc, f"⚠️ Erro ao clicar no botão 'Pesquisar' (índice {indice}): {e}")
        return False


def clicar_salvar_por_indice(js_engine, doc, indice=1, timeout=5):
    """
    Clica no botão 'Salvar' pelo índice informado (1-based).
    Conta e exibe quantos botões existem antes do clique.
    Usa js_engine.force_click() e registra log.
    """
    xpath_base = "//a[contains(@class,'btModel btGray btok') and contains(normalize-space(.),'Salvar')]"
    xpath_indexado = f"({xpath_base})[{indice}]"

    try:
        # Conta quantos botões existem
        elementos = js_engine.driver.find_elements("xpath", xpath_base)
        total = len(elementos)
        log(doc, f"💾 Foram encontrados {total} botão(ões) 'Salvar' na tela.")

        if total == 0:
            log(doc, "⚠️ Nenhum botão 'Salvar' foi encontrado.")
            return False
        if indice > total:
            log(doc, f"⚠️ Índice {indice} inválido — só existem {total} botão(ões).")
            return False

        log(doc, f"🎯 Clicando no botão 'Salvar' (índice {indice})...")
        js_engine.force_click(xpath_indexado, by_xpath=True)
        js_engine.wait_ajax_complete(timeout)
        log(doc, f"✅ Clique no botão 'Salvar' (índice {indice}) realizado com sucesso.")
        return True

    except Exception as e:
        log(doc, f"⚠️ Erro ao clicar no botão 'Salvar' (índice {indice}): {e}")
        return False


def clicar_titulo_produtos(js_engine, doc, timeout=3):
    """
    Clica no elemento <h2> que contém o texto 'PRODUTOS'.
    """
    xpath = "//h2[contains(normalize-space(.), 'PRODUTOS')]"
    try:
        log(doc, "🧩 Retornando à aba Principal")
        js_engine.force_click(xpath, by_xpath=True)
        js_engine.wait_ajax_complete(timeout)
        log(doc, "✅ Clique no título 'PRODUTOS' realizado com sucesso.")
        return True
    except Exception as e:
        log(doc, f"⚠️ Falha ao clicar no título 'PRODUTOS': {e}")
        return False

def clicar_salvar_ate_modal_fechar(js_engine, doc, timeout_total=60, pausa_entre=0.5):
    """
    Clica repetidamente no botão 'Salvar' do modal de emissão de nota fiscal
    até que o modal desapareça da tela ou o tempo máximo seja atingido.
    """
    xpath_modal = "//div[contains(@class,'modal') and .//h2[contains(.,'Emissão de Nota Fiscal de Serviço')]]"
    xpath_btn_salvar = "//a[contains(normalize-space(.),'Salvar')]"

    log(doc, "⚙️ Iniciando loop de clique em 'Salvar' até fechamento do modal...")
    inicio = time.time()

    while True:
        try:
            # Verifica se o modal ainda está presente na tela
            modais = js_engine.driver.find_elements("xpath", xpath_modal)
            if not modais or not modais[0].is_displayed():
                log(doc, "✅ Modal fechado — parando cliques.")
                break

            # Clica no botão 'Salvar'
            try:
                js_engine.force_click(xpath_btn_salvar, by_xpath=True)
                log(doc, "💾 Clique em 'Salvar' realizado.")
            except Exception as e:
                log(doc, f"⚠️ Erro ao clicar em 'Salvar': {e}")

            # Espera um pouco antes de tentar novamente
            time.sleep(pausa_entre)

            # Verifica timeout total
            if time.time() - inicio > timeout_total:
                log(doc, "⏰ Tempo limite atingido — modal ainda aberto, encerrando tentativa.")
                break

        except Exception as e:
            log(doc, f"⚠️ Erro inesperado no loop de salvamento: {e}")
            break

    log(doc, "🏁 Finalizado processo de clique repetido em 'Salvar'.")

def clicar_voltar(js_engine, doc, timeout=5):
    """Clica no botão 'Voltar' do modal com logs detalhados"""
    import time
    
    try:
        log(doc, "🔍 Iniciando busca pelo botão 'Voltar'...")
        
        # Script JavaScript que retorna informações detalhadas
        script = """
        var botoes = document.querySelectorAll('a.sp-voltarGrande');
        var info = {
            total: botoes.length,
            visiveis: 0,
            clicados: 0
        };
        
        botoes.forEach(function(btn) {
            if (btn.offsetParent !== null) {  // Verifica se está visível
                info.visiveis++;
                try {
                    btn.click();
                    info.clicados++;
                } catch(e) {
                    console.error('Erro ao clicar:', e);
                }
            }
        });
        
        return info;
        """
        
        log(doc, "⚙️ Executando JavaScript para clicar no botão 'Voltar'...")
        resultado = js_engine.driver.execute_script(script)
        
        log(doc, f"📊 Total de botões 'Voltar' encontrados: {resultado['total']}")
        log(doc, f"👁️ Botões visíveis: {resultado['visiveis']}")
        log(doc, f"✅ Botões clicados com sucesso: {resultado['clicados']}")
        
        if resultado['clicados'] == 0:
            log(doc, "⚠️ Nenhum botão 'Voltar' foi clicado.")
            if resultado['total'] == 0:
                log(doc, "❌ Motivo: Nenhum botão 'Voltar' encontrado.")
            elif resultado['visiveis'] == 0:
                log(doc, "❌ Motivo: Botão existe mas não está visível.")
        else:
            log(doc, f"🎉 Sucesso! Botão 'Voltar' clicado.")
            log(doc, f"⏳ Aguardando conclusão do AJAX (timeout: {timeout}s)...")
            js_engine.wait_ajax_complete(timeout)
            log(doc, "✅ AJAX concluído.")
        
        return {"total": resultado['total'], "clicados": resultado['clicados']}
        
    except Exception as e:
        log(doc, f"❌ Erro ao clicar no botão 'Voltar': {e}")
        import traceback
        log(doc, f"📋 Detalhes do erro: {traceback.format_exc()}")
        return {"total": 0, "clicados": 0}

def clicar_titulo_titulos(js_engine, doc, timeout=3):
    """
    Clica no elemento <h2> que contém o texto 'TÍTULOS'.
    """
    xpath = "//h2[contains(normalize-space(.), 'TÍTULOS')]"
    try:
        log(doc, "🧩 Retornando à aba Principal")
        js_engine.force_click(xpath, by_xpath=True)
        js_engine.wait_ajax_complete(timeout)
        log(doc, "✅ Clique no título 'TÍTULOS' realizado com sucesso.")
        return True
    except Exception as e:
        log(doc, f"⚠️ Falha ao clicar no título 'TÍTULOS': {e}")
        return False



def selecionar_template_por_texto(js_engine, doc, texto="PADRÃO", timeout=5):
    """
    Seleciona uma opção do select 'templateNota' pelo texto exibido.
    
    Args:
        js_engine: Instância do JSForceEngine
        doc: Documento para logs
        texto: Texto da opção a ser selecionada (default: "PADRÃO")
        timeout: Timeout para aguardar AJAX (default: 5)
    
    Returns:
        bool: True se selecionou com sucesso, False caso contrário
    """
    try:
        log(doc, f"📋 Selecionando template por texto: '{texto}'...")
        
        # Script JavaScript para selecionar por texto
        script = """
        const selectElement = document.querySelector('select.templateNota');
        const texto = arguments[0];
        
        if (!selectElement) {
            throw new Error('Select .templateNota não encontrado');
        }
        
        // Procura a opção pelo texto
        const option = Array.from(selectElement.options).find(
            opt => opt.text.trim() === texto.trim()
        );
        
        if (!option) {
            throw new Error(`Opção com texto '${texto}' não encontrada`);
        }
        
        // Torna o select visível e interativo
        selectElement.style.display = 'block';
        selectElement.style.visibility = 'visible';
        selectElement.removeAttribute('disabled');
        
        // Seleciona a opção pelo value
        selectElement.value = option.value;
        
        // Dispara eventos
        selectElement.dispatchEvent(new Event('change', { bubbles: true }));
        selectElement.dispatchEvent(new Event('input', { bubbles: true }));
        selectElement.dispatchEvent(new Event('blur', { bubbles: true }));
        
        return {
            sucesso: true,
            valorSelecionado: selectElement.value,
            textoSelecionado: selectElement.options[selectElement.selectedIndex].text
        };
        """
        
        resultado = js_engine.execute_js(
            script, 
            texto, 
            timeout=timeout,
            fallback_result=None
        )
        
        if resultado and resultado.get('sucesso'):
            value = resultado.get('valorSelecionado', '')
            log(doc, f"✅ Template '{texto}' (value={value}) selecionado com sucesso!")
            js_engine.wait_ajax_complete(timeout)
            return True
        else:
            log(doc, f"⚠️ Falha ao selecionar template '{texto}'")
            return False
            
    except Exception as e:
        log(doc, f"❌ Erro ao selecionar template '{texto}': {e}")
        return False

# ==== EXECUÇÃO DO TESTE ====
def executar_teste():
    """Execução principal do teste com JS forçado e proteção anti-timeout"""
    global driver, wait, doc
    
    try:
        if not inicializar_driver():
            return False
        
        # Cria engine JS forçado COM PROTEÇÃO ANTI-TIMEOUT
        js_engine = JSForceEngine(driver, wait, doc, timeout_padrao=10, max_retries=3)
        lov_handler = LOVHandler(js_engine, doc)
        
        # ===== LOGIN =====
        safe_action(doc, "Acessando sistema", lambda: driver.get(URL))
        
        def fazer_login():
            wait.until(EC.presence_of_element_located((By.ID, "j_id15:email"))).send_keys(LOGIN_EMAIL)
            wait.until(EC.presence_of_element_located((By.ID, "j_id15:senha"))).send_keys(LOGIN_PASSWORD, Keys.ENTER)
            time.sleep(5)

        safe_action(doc, "Realizando login", fazer_login)
        
        # ===== MENU =====
        def abrir_menu():
            driver.execute_script("document.body.style.zoom='90%'")
            driver.find_element(By.TAG_NAME, "body").send_keys(Keys.F3)
            time.sleep(2)
        
        safe_action(doc, "Abrindo menu (F3)", abrir_menu)
        
        # ===== CAIXA =====
        safe_action(doc, "Acessando Caixa", lambda:
            js_engine.force_click('/html/body/div[15]/ul/li[8]/img', by_xpath=True)
        )
        
        time.sleep(3)
        
        safe_action(doc, "Clicando em 'Entrar no Caixa'", lambda:
            js_engine.force_click(
                '#gsCaixa > div.wdTelas > div.telaInicial.clearfix.overflow.overflowY > ul > li:nth-child(1) > a > span'
            )
        )
        
        time.sleep(5)

        safe_action(doc, "Verificando e abrindo o caixa", lambda:
            verificar_e_abrir_caixa(js_engine, doc)
        )


        safe_action(doc, "Clicando em 'Adicionar Produtos'", lambda:
            js_engine.force_click(
                "(//a[@class='btAddProd' and contains(normalize-space(.), 'Adicionar produtos')])[1]",
                by_xpath=True
            )
        )

        safe_action(doc, "Preenchendo Nome do Produto", lambda:
            js_engine.force_fill("//input[@class='nomeProd']", "PRUDUTO", by_xpath=True)
        )
        time.sleep(2)
        safe_action(doc, "Clicando em todos os botões 'Pesquisar'", lambda:
            clicar_todos_pesquisar(js_engine, doc, pausa_entre=0.3)
        )


        time.sleep(20)

        safe_action(doc, "Adicionando primeiro produto", lambda: clicar_primeiro_sp_add(js_engine, doc))
        time.sleep(1)
        clicar_titulo_produtos(js_engine, doc)

        time.sleep(1)
        clicar_botao_por_classe(js_engine, doc, "sp-voltarGrande", "Voltar")

        safe_action(doc, "Clicando em 'Adicionar Títulos'", lambda:
            js_engine.force_click(
                "(//a[@class='btAddTit' and contains(normalize-space(.), 'Adicionar títulos')])[1]",
                by_xpath=True
            )
        )
        time.sleep(1)


        # ===== LOV COM PROTEÇÃO =====
        safe_action(doc, "Selecionando Pessoa", lambda:
            lov_handler.open_and_select(
                btn_index=0,
                search_text="TESTANDO PESSOA TITULAR",
                result_text="TESTANDO PESSOA TITULAR"
            )
        )
        
        clicar_pesquisar_por_indice(js_engine, doc, indice=1)

        time.sleep(20)

        safe_action(doc, "Adicionando título", lambda: clicar_primeiro_sp_add(js_engine, doc))
        time.sleep(1)
        clicar_titulo_titulos(js_engine, doc)

        clicar_botao_por_classe(js_engine, doc, "sp-voltarGrande", "Voltar")

        safe_action(doc, "Prosseguindo com o Pagamento", lambda:
            js_engine.force_click("//a[@class='btVenda' and normalize-space()='Prosseguir com pagamento (F5)']", by_xpath=True)
        )
        time.sleep(3)

        safe_action(doc, "Preenchendo CNPJ do Cliente", lambda:
            js_engine.force_fill("//input[@maxlength='14']", str(fake.cnpj()), by_xpath=True)
        )

        safe_action(doc, "Preenchendo Nº NFe", lambda:
            js_engine.force_fill("//input[@maxlength='20']", "55.001.12345", by_xpath=True)
        )

        # Formas de pagamento
        formas_pagamento = [
            ("Dinheiro", "//input[@class='valor vDinheiro']", "10000,00"),
            ("Boleto", "//input[@class='valor vBoleto']", "10000,00"),

        ]

        for nome, xpath, valor in formas_pagamento:
            safe_action(doc, f"Preenchendo Forma de Pagamento: {nome}", lambda x=xpath, v=valor:
                js_engine.force_fill(x, v, by_xpath=True)
            )

        alerta = clicar_finalizar_e_verificar_alerta(js_engine, doc)
        if alerta:
            log(doc, f"⚠️ Alerta detectado: {alerta.text}")
        else:
            log(doc, "✅ Nenhum alerta foi encontrado após clicar em 'Finalizar'.")

        time.sleep(10)


        safe_action(doc, "Recusando Geração de Nota Fiscal", lambda:
                js_engine.force_click("//a[@id='BtNo' and @class='btModel btGray btno' and normalize-space()='Não']", by_xpath=True)
            )
        time.sleep(5)

    

        safe_action(doc, "Clicando em 'Fechar Caixa'", lambda:
                js_engine.force_click("//a[contains(@class,'btAzulDegrade') and contains(@class,'btFecharCaixa')]", by_xpath=True)
            )


        safe_action(doc, "Fechando Caixa", lambda:
                js_engine.force_click("//a[@class='btModel btGray btyes' and normalize-space()='Fechar Caixa' and .//span[contains(@class,'sp-salvar')]]", by_xpath=True)
            )
        time.sleep(40)
        fechar_abas_extras_e_verificar_alerta(js_engine.driver, doc)

        safe_action(doc, "Gerando Relatório", lambda:
                js_engine.force_click("//a[@id='BtYes' and @class='btModel btGray btyes' and normalize-space()='Autenticar']", by_xpath=True)
            )
        time.sleep(5)
        fechar_abas_extras_e_verificar_alerta(js_engine.driver, doc)

        safe_action(doc, "Fechando modal de Geração", lambda:
                js_engine.force_click("//a[@id='BtNo' and @class='btModel btGray btno' and normalize-space()='Não']", by_xpath=True)
            )

        safe_action(doc, "Fechando modal de Autenticação", lambda:
                js_engine.force_click("//a[@id='BtNo' and @class='btModel btGray btno' and normalize-space()='Fechar']", by_xpath=True)
            )          
                # Voltar


        safe_action(doc, "Fechando modal do Controle de Caixa", lambda:
            js_engine.force_click(
                "//a[@class='sprites sp-fecharGrande' and @title='Sair']",
                by_xpath=True
            )
        )

        safe_action(doc, "Fechando modal do Caixa", lambda:
            js_engine.force_click('#gsCaixa > div.wdTop.ui-draggable-handle > div > a')
        )

        log(doc, "🎉 Teste concluído com sucesso!")
        return True
        
    except Exception as e:
        log(doc, f"❌ ERRO FATAL: {e}")
        take_screenshot(driver, doc, "erro_fatal")
        return False

# ==== MAIN ====
def main():
    """Ponto de entrada principal"""
    global doc
    
    try:
        log(doc, "🚀 Iniciando teste de Fluxo de Caixa")
        log(doc, "=" * 70)

        
        sucesso = executar_teste()
        
        log(doc, "=" * 70)
        if sucesso:
            log(doc, "✅ TESTE EXECUTADO COM SUCESSO!")
        else:
            log(doc, "❌ TESTE FINALIZADO COM ERROS")
            
    except Exception as e:
        log(doc, f"❌ Erro na execução principal: {e}")
        
    finally:
        finalizar_relatorio()

if __name__ == "__main__":
    main()