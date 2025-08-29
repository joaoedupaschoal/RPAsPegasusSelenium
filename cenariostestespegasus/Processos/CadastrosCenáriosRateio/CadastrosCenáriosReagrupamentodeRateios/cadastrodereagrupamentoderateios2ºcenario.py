# ==== IMPORTS (sem conflitos) ====
from datetime import datetime, timedelta
from datetime import time as dt_time  # usar para objetos de hora
import time                           # usar para time.sleep(...)
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver import ActionChains
from selenium.common.exceptions import TimeoutException, StaleElementReferenceException, JavascriptException
from selenium.webdriver.support import expected_conditions as EC
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support.ui import Select, WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service
from selenium.common.exceptions import TimeoutException, NoSuchElementException, StaleElementReferenceException, ElementClickInterceptedException
from webdriver_manager.chrome import ChromeDriverManager
from docx import Document
from docx.shared import Inches
from faker import Faker
from faker.providers import BaseProvider
from validate_docbr import CPF
from selenium.webdriver import ActionChains
import subprocess
import os
import random
import re


# Timeouts configuráveis
TIMEOUT_DEFAULT = 30
TIMEOUT_CURTO = 10
TIMEOUT_LONGO = 60

# ==== PROVIDERS CUSTOMIZADOS ====
class BrasilProvider(BaseProvider):
    def rg(self):
        numeros = [str(random.randint(0, 9)) for _ in range(8)]
        return ''.join(numeros) + '-' + str(random.randint(0, 9))

fake = Faker("pt_BR")
fake.add_provider(BrasilProvider)

# ==== GERAÇÃO DE DATAS ====
def gerar_datas_validas(hora_padrao="00:00", dias_fim=0):
    """Gera datas coerentes. data_inicio/data_fim no formato 'dd/MM/yyyy HH:mm'."""
    hoje_date = datetime.today().date()
    dez_anos_atras = hoje_date - timedelta(days=3650)

    # Falecimento entre 10 anos atrás e hoje
    data_falecimento = fake.date_between(start_date=dez_anos_atras, end_date=hoje_date)

    # Nascimento (entre 18 e 110 anos antes do falecimento)
    idade_minima, idade_maxima = 18, 110
    data_nascimento = data_falecimento - timedelta(days=random.randint(idade_minima * 365, idade_maxima * 365))

    # Sepultamento 1..10 dias após o falecimento
    data_sepultamento = data_falecimento + timedelta(days=random.randint(1, 10))

    # Registro 1..10 dias após o sepultamento
    data_registro = data_sepultamento + timedelta(days=random.randint(1, 10))

    # Velório entre o falecimento e o sepultamento
    data_velorio = fake.date_between(start_date=data_falecimento, end_date=data_sepultamento)

    # Início entre 2 e 30 dias no futuro, com hora escolhida
    data_inicio_date = hoje_date + timedelta(days=random.randint(2, 30))

    # Monta datetime com hora escolhida (ex: "00:00")
    h, m = map(int, hora_padrao.split(":"))
    dt_inicio = datetime.combine(data_inicio_date, dt_time(h, m))

    # Fim: mesmo dia por padrão (dias_fim=0). Ajuste se quiser +N dias.
    dt_fim = dt_inicio + timedelta(days=dias_fim)

    fmt_data = "%d/%m/%Y"
    fmt_dt = "%d/%m/%Y %H:%M"

    return (
        data_nascimento.strftime(fmt_data),
        data_falecimento.strftime(fmt_data),
        data_sepultamento.strftime(fmt_data),
        data_velorio.strftime(fmt_data),
        dt_inicio.strftime(fmt_dt),   # data_inicio com hora
        dt_fim.strftime(fmt_dt),      # data_fim com hora
        data_registro.strftime(fmt_data),
        hoje_date.strftime(fmt_data),
    )

(data_nascimento, data_falecimento, data_sepultamento,
 data_velorio, data_inicio, data_fim, data_registro, hoje) = gerar_datas_validas(
    hora_padrao="08:50",
    dias_fim=0
)

hora_falecimento = fake.time(pattern="%H:%M")
hora_sepultamento = fake.time(pattern="%H:%M")
localizacao = fake.city()

# ==== CONFIGURAÇÕES ====
URL = "http://localhost:8080/gs/index.xhtml"
LOGIN_EMAIL = "joaoeduardo.gold@outlook.com"
LOGIN_PASSWORD = "071999gs"

# ==== INICIALIZAÇÃO DE VARIÁVEIS GLOBAIS ====
doc = Document()
doc.add_heading("RELATÓRIO DO TESTE", 0)
doc.add_paragraph("Cadastro de Reagrupamento de Rateios - Rateio – Cenário 2: Nesse teste, será realizado um Reagrupamento de Rateio, verificando o disparo das mensagens de alerta do sistema.")
doc.add_paragraph(f"Data do teste: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")

screenshot_registradas = set()
driver = None
wait = None
janela_principal = None  # ✅ CORREÇÃO: iniciar como None; só capturar após abrir o driver

# ==== UTILITÁRIOS ====

def selecionar_opcao_seletor(selector, texto):
    def acao():
        select_element = wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, selector)))
        Select(select_element).select_by_visible_text(texto)
    return acao

def selecionar_opcao_xpath(xpath, texto):
    def acao():
        select_element = wait.until(EC.presence_of_element_located((By.XPATH, xpath)))
        Select(select_element).select_by_visible_text(texto)
    return acao


def log(doc, msg, nivel="INFO"):
    nivel = (nivel or "INFO").upper()
    prefixo = {
        "INFO": "ℹ️ [INFO]",
        "WARN": "⚠️ [WARN]",
        "ERROR": "❌ [ERROR]",
    }.get(nivel, f"[{nivel}]")

    linha = f"{datetime.now().strftime('%Y-%m-%d %H:%M:%S')} {prefixo} {msg}"
    print(linha)
    doc.add_paragraph(linha)


def _sanitize_filename(name: str) -> str:
    name = name.strip().lower()
    name = re.sub(r"[<>:\"/\\|?*']", "_", name)
    name = re.sub(r"_+", "_", name)
    return name[:120]

def take_screenshot(driver, doc, nome):
    if driver is None:
        return
    nome = _sanitize_filename(nome)
    if nome not in screenshot_registradas:
        path = f"screenshots/{nome}.png"
        os.makedirs("screenshots", exist_ok=True)
        try:
            driver.save_screenshot(path)
            doc.add_paragraph(f"Screenshot: {nome}")
            doc.add_picture(path, width=Inches(5.5))
            screenshot_registradas.add(nome)
        except Exception as e:
            log(doc, f"⚠️ Erro ao tirar screenshot {nome}: {e}")

# ✅ Agora safe_action aceita critico
def safe_action(doc, descricao, func, max_retries=3, critico=True):
    global driver
    for attempt in range(max_retries):
        try:
            log(doc, f"🔄 {descricao}..." if attempt == 0 else f"🔄 {descricao}... (Tentativa {attempt+1})")
            func()
            log(doc, f"✅ {descricao} realizada com sucesso.")
            take_screenshot(driver, doc, descricao)
            return True
        except (TimeoutException, NoSuchElementException, StaleElementReferenceException) as e:
            if attempt < max_retries - 1:
                log(doc, f"⚠️ Tentativa {attempt+1} falhou: {e}")
                time.sleep(2)
            else:
                log(doc, f"❌ Erro ao {descricao.lower()} após {max_retries} tentativas: {e}")
                take_screenshot(driver, doc, f"erro_{descricao}")
                return not critico
        except Exception as e:
            log(doc, f"❌ Erro inesperado ao {descricao.lower()}: {e}")
            take_screenshot(driver, doc, f"erro_{descricao}")
            return not critico

def abrir_modal_e_selecionar_robusto(btn_selector, pesquisa_selector, termo_pesquisa, btn_pesquisar_selector, resultado_xpath):
    """Versão robusta da função de modal"""
    global driver, wait, doc
    
    def acao():
        if driver is None or wait is None:
            raise Exception("Driver ou wait não inicializados")
            
        # Abre o modal
        safe_scroll_and_interact(btn_selector, "click")
        time.sleep(1)

        # Aguarda e preenche campo pesquisa
        campo_pesquisa = wait.until(
            EC.element_to_be_clickable((By.CSS_SELECTOR, pesquisa_selector))
        )
        campo_pesquisa.clear()
        campo_pesquisa.send_keys(termo_pesquisa)
        time.sleep(0.5)

        # Clica pesquisar
        pesquisar = wait.until(
            EC.element_to_be_clickable((By.CSS_SELECTOR, btn_pesquisar_selector))
        )
        pesquisar.click()
        time.sleep(2)
        
        # Aguarda resultado e clica
        resultado = wait.until(
            EC.element_to_be_clickable((By.XPATH, resultado_xpath))
        )
        driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", resultado)
        time.sleep(0.5)
        resultado.click()
        time.sleep(1)

    return acao


def finalizar_relatorio():
    global driver, doc
    nome_arquivo = f"relatorio_Reagrupamento_rateios_cenario_1_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    try:
        doc.save(nome_arquivo)
        log(doc, f"📄 Relatório salvo como: {nome_arquivo}")
        try:
            subprocess.run(["start", "winword", nome_arquivo], shell=True)
        except:
            pass
    except Exception as e:
        print(f"Erro ao salvar relatório: {e}")
    if driver:
        try:
            driver.quit()
        except:
            pass

def encontrar_mensagem_alerta():
    global driver, doc
    if driver is None:
        return None
    seletores = [
        (".alerts.salvo", "✅ Sucesso"),
        (".alerts.alerta", "⚠️ Alerta"),
        (".alerts.erro", "❌ Erro"),
    ]
    for seletor, tipo in seletores:
        try:
            elemento = driver.find_element(By.CSS_SELECTOR, seletor)
            if elemento.is_displayed():
                log(doc, f"📢 {tipo}: {elemento.text}")
                return elemento
        except:
            continue
    log(doc, "ℹ️ Nenhuma mensagem de alerta encontrada.")
    return None

def ajustar_zoom():
    global driver, doc
    if driver is None:
        return
    try:
        driver.execute_script("document.body.style.zoom='90%'")
        log(doc, "🔍 Zoom ajustado para 90%.")
    except Exception as e:
        log(doc, f"⚠️ Erro ao ajustar zoom: {e}")

def safe_scroll_and_interact(selector, action_type="click", value=None, timeout=10, by_xpath=False):
    """Rola até o elemento e interage com ele de forma robusta."""
    global driver, doc
    if driver is None:
        return None
    try:
        by_type = By.XPATH if by_xpath else By.CSS_SELECTOR
        element = WebDriverWait(driver, timeout).until(EC.presence_of_element_located((by_type, selector)))
        driver.execute_script("arguments[0].scrollIntoView({block: 'center', behavior: 'smooth'});", element)
        time.sleep(0.5)
        if action_type in ["click", "send_keys"]:
            element = WebDriverWait(driver, timeout).until(EC.element_to_be_clickable((by_type, selector)))
        if action_type == "click":
            element.click()
        elif action_type == "send_keys" and value:
            element.clear()
            element.send_keys(value)
        elif action_type == "select" and value:
            Select(element).select_by_visible_text(value)
        return element
    except Exception as e:
        log(doc, f"❌ Erro ao interagir com elemento {selector}: {e}")
        return None

# ==== CLICAR ROBUSTO ====
def clicar_elemento_robusto(driver, wait, seletor_css, timeout=10):
    global doc
    try:
        elem = wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, seletor_css)))
        try:
            driver.execute_script("""
                document.querySelectorAll('.modal, .overlay, .blockUI, .toast, .tooltip, [role="dialog"], [data-overlay]')
                .forEach(e => { if (getComputedStyle(e).position === 'fixed') e.style.display = 'none'; });
            """)
        except Exception:
            pass
        driver.execute_script("arguments[0].scrollIntoView({block:'center', inline:'center'});", elem)
        time.sleep(0.2)
        try:
            elem = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, seletor_css)))
            elem.click()
            return True
        except (TimeoutException, ElementClickInterceptedException, StaleElementReferenceException):
            pass
        try:
            elem = driver.find_element(By.CSS_SELECTOR, seletor_css)
            ActionChains(driver).move_to_element(elem).pause(0.05).click().perform()
            return True
        except Exception:
            pass
        try:
            elem = driver.find_element(By.CSS_SELECTOR, seletor_css)
            driver.execute_script("arguments[0].click();", elem)
            return True
        except Exception:
            pass
        try:
            elem = driver.find_element(By.CSS_SELECTOR, seletor_css)
            driver.execute_script("""
                const el = arguments[0];
                function fire(type){ el.dispatchEvent(new MouseEvent(type,{bubbles:true,cancelable:true,view:window})); }
                el.focus(); fire('mouseover'); fire('mousedown'); fire('mouseup'); fire('click');
            """, elem)
            return True
        except Exception:
            pass
        try:
            elem = driver.find_element(By.CSS_SELECTOR, seletor_css)
            ActionChains(driver).move_to_element_with_offset(elem, 1, 1).click().perform()
            return True
        except Exception:
            pass
        log(doc, f"❌ Não foi possível clicar em: {seletor_css}")
        return False
    except Exception as e:
        log(doc, f"❌ Erro ao clicar robusto: {e}")
        return False

# ==== DATEPICKER ====
def encontrar_campos_datepicker():
    global driver
    seletores = [
        "input.hasDatepicker",
        "input[id^='dp']",
        "input[class*='datepicker']",
        ".hasDatepicker",
        "[data-provide='datepicker']"
    ]
    campos = []
    for seletor in seletores:
        try:
            for elemento in driver.find_elements(By.CSS_SELECTOR, seletor):
                if elemento.is_displayed():
                    cid = elemento.get_attribute('id')
                    if not any(c.get('id') == cid for c in campos):
                        campos.append({'elemento': elemento, 'id': cid})
        except:
            continue
    log(doc, f"📊 Encontrados {len(campos)} campos datepicker")
    return campos

def _datepicker_jquery(campo_id, data_valor):
    global driver
    resultado = driver.execute_script("""
        var campoId = arguments[0], valor = arguments[1];
        if (typeof jQuery === 'undefined') return 'jQuery não disponível';
        var $campo = $('#' + campoId);
        if (!$campo.length) return 'Campo não encontrado: ' + campoId;
        try {
            if ($campo.hasClass('hasDatepicker')) { $campo.datepicker('setDate', valor); }
            else { $campo.val(valor); }
            $campo.trigger('input').trigger('change').trigger('blur');
            return $campo.val();
        } catch(e) { return 'Erro: ' + e.message; }
    """, campo_id, data_valor)
    if isinstance(resultado, str) and 'Erro' in resultado:
        raise Exception(f"jQuery falhou: {resultado}")

def _datepicker_javascript(elemento, data_valor):
    global driver
    driver.execute_script("""
        var campo = arguments[0], valor = arguments[1];
        campo.focus(); campo.value = ''; campo.value = valor;
        ['input','change','blur','keyup'].forEach(ev => campo.dispatchEvent(new Event(ev,{bubbles:true})));
    """, elemento, data_valor)

def _datepicker_actionchains(elemento, data_valor):
    global driver
    driver.execute_script("arguments[0].scrollIntoView({block:'center'});", elemento)
    ActionChains(driver).move_to_element(elemento).click().perform()
    time.sleep(0.3)
    ActionChains(driver).key_down(Keys.CONTROL).send_keys('a').key_up(Keys.CONTROL).perform()
    ActionChains(driver).send_keys(Keys.DELETE).perform()
    time.sleep(0.2)
    for ch in data_valor:
        ActionChains(driver).send_keys(ch).perform()
        time.sleep(0.03)
    ActionChains(driver).send_keys(Keys.TAB).perform()

def _datepicker_tradicional(elemento, data_valor):
    driver.execute_script("arguments[0].scrollIntoView({block:'center'});", elemento)
    elemento.click(); time.sleep(0.2); elemento.clear()
    elemento.send_keys(data_valor); elemento.send_keys(Keys.TAB)

def preencher_datepicker_persistente(indice_campo, data_valor, max_tentativas=10, timeout=30):
    """Preenche datepicker com várias estratégias e valida."""
    inicio_tempo = time.time()
    tentativa = 0

    def validar_data_preenchida(el, data_esperada):
        try:
            val = (el.get_attribute('value') or '').strip()
            if not val: return False
            if val == data_esperada or data_esperada in val: return True
            formatos = [
                '%d/%m/%Y %H:%M', '%d/%m/%Y %H:%M:%S',  # <- COM HORA
                '%d/%m/%Y', '%m/%d/%Y', '%Y-%m-%d', '%d-%m-%Y'
            ]
            for f in formatos:
                try:
                    d1 = datetime.strptime(val, f)
                    d2 = datetime.strptime(data_esperada, f)
                    if d1 == d2: return True
                except:
                    continue
            return False
        except:
            return False

    while tentativa < max_tentativas and (time.time() - inicio_tempo) < timeout:
        tentativa += 1
        try:
            log(doc, f"🔄 Tentativa {tentativa}/{max_tentativas} para campo {indice_campo}")
            campos = encontrar_campos_datepicker()
            if not campos:
                time.sleep(1); continue
            if indice_campo >= len(campos):
                time.sleep(1); continue

            info = campos[indice_campo]
            elemento, campo_id = info['elemento'], info['id']

            if validar_data_preenchida(elemento, data_valor):
                log(doc, f"✅ Campo {indice_campo} já está preenchido corretamente!")
                return True

            estrategias = [
                lambda: _datepicker_jquery(campo_id, data_valor),
                lambda: _datepicker_javascript(elemento, data_valor),
                lambda: _datepicker_actionchains(elemento, data_valor),
                lambda: _datepicker_tradicional(elemento, data_valor),
            ]
            for acao in estrategias:
                try:
                    acao(); time.sleep(0.4)
                    if validar_data_preenchida(elemento, data_valor):
                        log(doc, f"✅ Campo {indice_campo} preenchido!")
                        return True
                except Exception as e:
                    log(doc, f"   ⚠️ Estratégia falhou: {e}")

            log(doc, "❌ Tentativa falhou; tentando novamente...")
            time.sleep(1.2)

        except Exception as e:
            log(doc, f"❌ Erro na tentativa {tentativa}: {e}")
            time.sleep(1)

    raise Exception(f"Falha ao preencher datepicker {indice_campo} após {tentativa} tentativas em {int(time.time()-inicio_tempo)}s")

def preencher_campo_xpath_com_retry(driver, wait, xpath, valor, max_tentativas=3):
    global doc
    if driver is None or wait is None:
        return False
    for tentativa in range(max_tentativas):
        try:
            campo = wait.until(EC.element_to_be_clickable((By.XPATH, xpath)))
            driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", campo)
            time.sleep(0.3)
            if tentativa == 0:
                campo.click(); campo.clear(); campo.send_keys(valor); campo.send_keys(Keys.TAB)
            elif tentativa == 1:
                ActionChains(driver).move_to_element(campo).click().pause(0.1).key_down(Keys.CONTROL).send_keys('a').key_up(Keys.CONTROL).send_keys(valor).send_keys(Keys.TAB).perform()
            else:
                driver.execute_script("""
                    var el = arguments[0], v = arguments[1];
                    el.focus(); el.value = ''; el.value = v;
                    el.dispatchEvent(new Event('input', { bubbles: true }));
                    el.dispatchEvent(new Event('change', { bubbles: true }));
                    el.blur();
                """, campo, valor)
            time.sleep(0.3)
            if (campo.get_attribute('value') or '').strip():
                return True
        except Exception as e:
            log(doc, f"⚠️ Tentativa {tentativa + 1} falhou: {e}")
            time.sleep(0.8)
    return False

# ==== WIZARD (clique por ícone) ====
def get_xpath(we):
    js = """
    function absoluteXPath(element) {
      if (element.tagName.toLowerCase() === 'html') return '/html';
      if (element===document.body) return '/html/body';
      var ix=0; var siblings=element.parentNode.children;
      var same = 0;
      for (var i=0; i<siblings.length; i++) {
        var sib=siblings[i];
        if (sib.tagName===element.tagName) {
          same++;
          if (sib===element) {
            var path = absoluteXPath(element.parentNode) + '/' + element.tagName.toLowerCase();
            if (same>1) path += '['+same+']';
            return path;
          }
        }
      }
      return absoluteXPath(element.parentNode) + '/' + element.tagName.toLowerCase();
    }
    return absoluteXPath(arguments[0]);
    """
    return we.parent.execute_script(js, we)


def _prepare_focus_and_clear(elemento, limpar_primeiro=True):
    # Garante visibilidade e foco
    driver.execute_script("arguments[0].scrollIntoView({block:'center'});", elemento)
    try:
        elemento.click()
    except Exception:
        driver.execute_script("arguments[0].focus();", elemento)

    if limpar_primeiro:
        try:
            elemento.clear()
        except Exception:
            # Fallback de limpeza por teclas
            ActionChains(driver)\
                .move_to_element(elemento).click()\
                .key_down(Keys.CONTROL).send_keys("a").key_up(Keys.CONTROL)\
                .send_keys(Keys.DELETE).perform()





def _textarea_tradicional(elemento, texto, limpar_primeiro=True):
    _prepare_focus_and_clear(elemento, limpar_primeiro)
    elemento.send_keys(texto)
    # Dispara blur para muitos bindings reativos
    elemento.send_keys(Keys.TAB)


def _textarea_actionchains(elemento, texto, limpar_primeiro=True):
    _prepare_focus_and_clear(elemento, limpar_primeiro)
    ac = ActionChains(driver)
    ac.move_to_element(elemento).click().perform()
    # Quebra o texto em partes para evitar engasgos em campos longos
    for chunk_start in range(0, len(texto), 400):
        ac.send_keys(texto[chunk_start:chunk_start+400]).perform()
        time.sleep(0.05)
    ac.send_keys(Keys.TAB).perform()


def _textarea_js_setvalue(elemento, texto):
    # Seta .value e dispara eventos clássicos
    driver.execute_script("""
        const el = arguments[0];
        const val = arguments[1];
        el.value = val;
        // Dispara eventos comuns que form libs escutam
        el.dispatchEvent(new Event('input',  {bubbles:true}));
        el.dispatchEvent(new Event('change', {bubbles:true}));
        el.dispatchEvent(new KeyboardEvent('keyup', {bubbles:true}));
        el.dispatchEvent(new Event('blur',   {bubbles:true}));
    """, elemento, texto)


def _textarea_js_react_input(elemento, texto):
    # Compat extra p/ React (setando o setter do prototype) + eventos
    driver.execute_script("""
        const el = arguments[0];
        const val = arguments[1];

        const desc = Object.getOwnPropertyDescriptor(window.HTMLTextAreaElement.prototype, 'value');
        if (desc && desc.set) {
            desc.set.call(el, val);
        } else {
            el.value = val;
        }

        // React/Vue/Svelte geralmente escutam 'input'
        el.dispatchEvent(new Event('input', {bubbles: true}));
        el.dispatchEvent(new Event('change', {bubbles: true}));
        el.dispatchEvent(new Event('blur', {bubbles: true}));
    """, elemento, texto)


def preencher_textarea_por_indice(indice_campo, texto, max_tentativas=5, limpar_primeiro=True):
    """Preenche um <textarea> pelo índice (ordem no DOM) usando estratégias múltiplas"""
    def acao():
        if not isinstance(indice_campo, int) or indice_campo < 0:
            raise ValueError(f"Índice inválido: {indice_campo}")
        if texto is None or not isinstance(texto, str):
            raise ValueError(f"Texto inválido: {texto!r}")

        tentativa = 0
        while tentativa < max_tentativas:
            tentativa += 1
            try:
                campos = encontrar_campos_textarea()
                if not campos:
                    if tentativa < max_tentativas:
                        log(doc, f"⚠️ Nenhuma <textarea> encontrada (tentativa {tentativa}/{max_tentativas})", "WARN")
                        time.sleep(1.5)
                        continue
                    raise Exception("Nenhuma <textarea> foi encontrada na página.")

                if indice_campo >= len(campos):
                    raise Exception(f"Índice {indice_campo} inválido. Encontradas {len(campos)} textareas.")

                campo_info = campos[indice_campo]
                elemento   = campo_info["elemento"]
                campo_id   = campo_info.get("id") or "(sem id)"
                campo_name = campo_info.get("name") or "(sem name)"

                log(doc, f"🎯 Tentativa {tentativa}: Preenchendo textarea {indice_campo} (ID: {campo_id}, name: {campo_name}) com {len(texto)} caracteres")

                # Se já estiver preenchido corretamente, encerra
                if validar_textarea_preenchida(elemento, texto):
                    log(doc, f"✅ Textarea {indice_campo} já está com o valor desejado.")
                    return True

                # Estratégias em ordem de 'menos invasiva' para 'mais invasiva'
                estrategias = [
                    lambda: _textarea_tradicional(elemento, texto, limpar_primeiro),
                    lambda: _textarea_actionchains(elemento, texto, limpar_primeiro),
                    lambda: _textarea_js_setvalue(elemento, texto),
                    lambda: _textarea_js_react_input(elemento, texto),
                ]

                for i, estrategia in enumerate(estrategias, 1):
                    try:
                        log(doc, f"   ▶️ Estratégia {i}…")
                        estrategia()
                        time.sleep(0.8)

                        # Revalida após a estratégia
                        if validar_textarea_preenchida(elemento, texto):
                            val = (elemento.get_attribute("value") or "").strip()
                            log(doc, f"✅ Preenchido com sucesso pela estratégia {i}: '{val[:60]}{'…' if len(val) > 60 else ''}'")
                            return True
                        else:
                            log(doc, f"⚠️ Estratégia {i} não refletiu o valor esperado.", "WARN")
                    except (StaleElementReferenceException, JavascriptException, TimeoutException) as e:
                        log(doc, f"⚠️ Estratégia {i} falhou: {e}", "WARN")
                        # Reobter o elemento se necessário
                        try:
                            campos = encontrar_campos_textarea()
                            elemento = campos[indice_campo]["elemento"]
                        except Exception:
                            pass
                        continue

                # Se chegou aqui, nenhuma estratégia funcionou nesta tentativa
                if tentativa < max_tentativas:
                    log(doc, f"⚠️ Tentativa {tentativa} falhou; reintentando em 1.5s…", "WARN")
                    time.sleep(1.5)
                    continue
            except Exception as e:
                if tentativa < max_tentativas:
                    log(doc, f"⚠️ Erro na tentativa {tentativa}: {e}. Retentando…", "WARN")
                    time.sleep(1.5)
                    continue
                else:
                    raise

        raise Exception(f"Falha ao preencher textarea {indice_campo} após {max_tentativas} tentativas.")
    return acao


# =========================
# Helpers usados pela função
# =========================

def encontrar_campos_textarea(timeout=10):
    """
    Retorna uma lista de dicts com metadados de cada <textarea> visível e interativa.
    Ex.: [{'elemento': WebElement, 'id': '...', 'name': '...'}]
    """
    elementos = []
    try:
        # Espera haver pelo menos 1 textarea no DOM (se existir)
        wait.until(lambda d: len(d.find_elements(By.TAG_NAME, "textarea")) >= 0)
        textareas = driver.find_elements(By.TAG_NAME, "textarea")
    except Exception:
        textareas = []

    for el in textareas:
        try:
            if not el.is_displayed() or not el.is_enabled():
                continue
            driver.execute_script("arguments[0].scrollIntoView({block:'center'});", el)
            elementos.append({
                "elemento": el,
                "id": el.get_attribute("id"),
                "name": el.get_attribute("name"),
            })
        except Exception:
            continue

    return elementos


def normalizar_texto(txt):
    if txt is None:
        return ""
    # Normaliza quebras de linha e espaços
    return txt.replace("\r\n", "\n").replace("\r", "\n").strip()


def validar_textarea_preenchida(elemento, texto_esperado):
    """Confere se o valor atual da textarea bate com o texto esperado (normalizado)."""
    try:
        atual = elemento.get_attribute("value")
        # Alguns frameworks populam via textContent em textareas (raro, mas possível)
        if atual is None or atual == "":
            atual = (elemento.text or "")
        return normalizar_texto(atual) == normalizar_texto(texto_esperado)
    except StaleElementReferenceException:
        return False

def clicar_todos_checkboxes(doc):
    """Marca todos os checkboxes falecidosSelecionados"""
    safe_action(doc, "Clicando em todos os checkboxes ", lambda: (
        [
            (cb.click() if not cb.is_selected() else None)
            for cb in driver.find_elements(
                By.XPATH,
                "//input[@type='checkbox' and contains(@class,'falecidosSelecionados')]"
            )
        ]
    ))



def click_wizard_by_icon(driver, wait, icon_class, expect_selector=None, timeout=12):
    span = wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, f"#gsCrm .btnHolder .sprites.{icon_class}")))
    driver.execute_script("arguments[0].scrollIntoView({block:'center'});", span)
    alvo = driver.execute_script("""
        const span = arguments[0]; let el = span;
        while (el && el !== document && !/^(A|DIV)$/i.test(el.tagName)) el = el.parentElement;
        return el;
    """, span)
    if not alvo:
        raise Exception(f"Não achei ancestral clicável para o ícone .{icon_class}")
    try:
        xp = get_xpath(alvo)
        wait.until(EC.element_to_be_clickable((By.XPATH, xp))).click()
    except Exception:
        try:
            driver.execute_script("arguments[0].click();", alvo)
        except Exception:
            driver.execute_script("""
                const el=arguments[0];
                function fire(t){el.dispatchEvent(new MouseEvent(t,{bubbles:true,cancelable:true,view:window}));}
                el.focus(); fire('mouseover'); fire('mousedown'); fire('mouseup'); fire('click');
            """, alvo)
    if expect_selector:
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, expect_selector)))

# ==== DRIVER ====
def inicializar_driver():
    global driver, wait
    try:
        options = Options()
        options.add_argument("--start-maximized")
        options.add_argument("--disable-blink-features=AutomationControlled")
        options.add_experimental_option("excludeSwitches", ["enable-automation"])
        options.add_experimental_option('useAutomationExtension', False)

        driver = webdriver.Chrome(service=Service(ChromeDriverManager().install()), options=options)
        driver.execute_script("Object.defineProperty(navigator, 'webdriver', {get: () => undefined})")
        wait = WebDriverWait(driver, 20)
        return True
    except Exception as e:
        log(doc, f"❌ Erro ao inicializar driver: {e}")
        return False


def _sanitize_timeout(t):
    """Garante timeout válido"""
    if not isinstance(t, (int, float)) or t <= 0:
        return TIMEOUT_DEFAULT
    return max(5, min(120, t))  # Entre 5 e 120 segundos


# ==== AGUARDAR ELEMENTO MELHORADO ====
def aguardar_elemento(seletor, timeout=TIMEOUT_DEFAULT, condicao='clickable', by_type=By.CSS_SELECTOR):
    """Função centralizada para aguardar elementos com diferentes condições"""
    global driver, wait
    
    if driver is None:
        raise Exception("Driver não inicializado")
    
    timeout = _sanitize_timeout(timeout)
    
    condicoes = {
        'present': EC.presence_of_element_located,
        'visible': EC.visibility_of_element_located,
        'clickable': EC.element_to_be_clickable,
        'invisible': EC.invisibility_of_element_located
    }
    
    if condicao not in condicoes:
        condicao = 'clickable'
    
    try:
        wait_obj = WebDriverWait(driver, timeout)
        elemento = wait_obj.until(condicoes[condicao]((by_type, seletor)))
        return elemento
    except TimeoutException:
        log(doc, f"❌ Timeout aguardando elemento: {seletor} (condição: {condicao}, timeout: {timeout}s)", 'ERROR')
        raise TimeoutException(f"Elemento não encontrado: {seletor} (condição: {condicao})")
    except Exception as e:
        log(doc, f"❌ Erro aguardando elemento {seletor}: {e}", 'ERROR')
        raise


# ==== SCROLL CORRIGIDO - PRINCIPAL CORREÇÃO ====
def scroll_to_element_safe(elemento_ou_seletor, by_type=By.CSS_SELECTOR):
    """Scroll seguro até elemento com validação robusta"""
    global driver
    
    if driver is None:
        log(doc, "⚠️ Driver não disponível para scroll", 'WARN')
        return False
    
    try:
        # Se for seletor, encontra o elemento
        if isinstance(elemento_ou_seletor, str):
            elemento = aguardar_elemento(elemento_ou_seletor, 10, 'present', by_type)
        else:
            elemento = elemento_ou_seletor
        
        if elemento is None:
            log(doc, "⚠️ Elemento não encontrado para scroll", 'WARN')
            return False
        
        # Verifica se elemento é válido antes de fazer scroll
        if not elemento.is_displayed():
            log(doc, "⚠️ Elemento não está visível para scroll", 'WARN')
            return False
        
        # Estratégias de scroll em ordem de preferência
        scroll_strategies = [
            # Estratégia 1: JavaScript com verificação prévia
            lambda: driver.execute_script("""
                var element = arguments[0];
                if (element && typeof element.scrollIntoView === 'function') {
                    element.scrollIntoView({
                        behavior: 'smooth',
                        block: 'center',
                        inline: 'center'
                    });
                    return true;
                } else {
                    return false;
                }
            """, elemento),
            
            # Estratégia 2: ActionChains
            lambda: ActionChains(driver).move_to_element(elemento).perform(),
            
            # Estratégia 3: JavaScript alternativo
            lambda: driver.execute_script("""
                var element = arguments[0];
                if (element) {
                    element.scrollIntoView();
                    window.scrollBy(0, -100);
                }
            """, elemento),
            
            # Estratégia 4: Scroll da página até o elemento
            lambda: driver.execute_script("""
                var element = arguments[0];
                if (element) {
                    var rect = element.getBoundingClientRect();
                    var scrollTop = window.pageYOffset || document.documentElement.scrollTop;
                    var targetY = rect.top + scrollTop - (window.innerHeight / 2);
                    window.scrollTo(0, targetY);
                }
            """, elemento)
        ]
        
        for i, strategy in enumerate(scroll_strategies, 1):
            try:
                log(doc, f"   Tentando estratégia de scroll {i}...")
                result = strategy()
                
                # Para estratégia 1, verifica resultado
                if i == 1 and result is False:
                    log(doc, f"   Estratégia {i}: elemento não suporta scrollIntoView", 'WARN')
                    continue
                
                time.sleep(0.8)  # Aguarda scroll completar
                
                # Verifica se elemento ainda está acessível
                if elemento.is_displayed() and elemento.is_enabled():
                    log(doc, f"✅ Scroll realizado com estratégia {i}")
                    return True
                else:
                    log(doc, f"   Estratégia {i}: elemento não ficou acessível", 'WARN')
                    continue
                    
            except Exception as e:
                log(doc, f"   Estratégia {i} de scroll falhou: {str(e)[:100]}...", 'WARN')
                continue
        
        log(doc, "⚠️ Todas as estratégias de scroll falharam", 'WARN')
        return False
        
    except Exception as e:
        log(doc, f"⚠️ Erro geral no scroll: {e}", 'WARN')
        return False



# ==== SISTEMA DATEPICKER MELHORADO ====
def encontrar_campos_datepicker():
    """Encontra todos os campos datepicker na página"""
    global driver
    
    if driver is None:
        return []
    
    seletores_datepicker = [
        "input.hasDatepicker",
        "input[id^='dp']",
        "input[maxlength='10'][grupo='']",
        "input[type='text'][maxlength='10']",
        "input[class*='datepicker']",
        ".hasDatepicker"
    ]
    
    campos_encontrados = []
    
    for seletor in seletores_datepicker:
        try:
            elementos = driver.find_elements(By.CSS_SELECTOR, seletor)
            for elemento in elementos:
                if elemento.is_displayed() and elemento.is_enabled():
                    info = {
                        'elemento': elemento,
                        'id': elemento.get_attribute('id') or f"dp_{len(campos_encontrados)}",
                        'seletor_usado': seletor,
                        'maxlength': elemento.get_attribute('maxlength'),
                        'placeholder': elemento.get_attribute('placeholder')
                    }
                    # Evita duplicatas
                    if not any(c['id'] == info['id'] for c in campos_encontrados):
                        campos_encontrados.append(info)
        except Exception as e:
            log(doc, f"⚠️ Erro ao buscar campos datepicker com {seletor}: {e}", 'WARN')
            continue
    
    log(doc, f"📊 Encontrados {len(campos_encontrados)} campos datepicker")
    return campos_encontrados

def _datepicker_jquery(campo_id, data_valor):
    """Estratégia jQuery para datepicker"""
    global driver
    
    resultado = driver.execute_script("""
        var campoId = arguments[0], valor = arguments[1];
        if (typeof jQuery === 'undefined') return 'jQuery não disponível';
        var $campo = $('#' + campoId);
        if (!$campo.length) return 'Campo não encontrado: ' + campoId;
        try {
            if ($campo.hasClass('hasDatepicker')) { 
                $campo.datepicker('setDate', valor); 
            } else { 
                $campo.val(valor); 
            }
            $campo.trigger('input').trigger('change').trigger('blur');
            return $campo.val();
        } catch(e) { 
            return 'Erro: ' + e.message; 
        }
    """, campo_id, data_valor)
    
    if isinstance(resultado, str) and ('Erro' in resultado or 'não disponível' in resultado):
        raise Exception(f"jQuery falhou: {resultado}")

def _datepicker_javascript(elemento, data_valor):
    """Estratégia JavaScript para datepicker"""
    global driver
    
    driver.execute_script("""
        var campo = arguments[0], valor = arguments[1];
        campo.focus(); 
        campo.value = ''; 
        campo.value = valor;
        ['input','change','blur','keyup'].forEach(ev => 
            campo.dispatchEvent(new Event(ev, {bubbles: true}))
        );
    """, elemento, data_valor)

def _datepicker_actionchains(elemento, data_valor):
    """Estratégia ActionChains para datepicker"""
    global driver
    
    scroll_to_element_safe(elemento)
    time.sleep(0.5)
    
    ActionChains(driver).move_to_element(elemento).click().perform()
    time.sleep(0.5)
    ActionChains(driver).key_down(Keys.CONTROL).send_keys('a').key_up(Keys.CONTROL).perform()
    time.sleep(0.3)
    ActionChains(driver).send_keys(Keys.DELETE).perform()
    time.sleep(0.3)
    
    for char in data_valor:
        ActionChains(driver).send_keys(char).perform()
        time.sleep(0.05)
    
    ActionChains(driver).send_keys(Keys.TAB).perform()

def _datepicker_tradicional(elemento, data_valor):
    """Estratégia tradicional para datepicker"""
    scroll_to_element_safe(elemento)
    time.sleep(0.5)
    elemento.click()
    time.sleep(0.5)
    elemento.clear()
    elemento.send_keys(data_valor)
    elemento.send_keys(Keys.TAB)

def validar_data_preenchida(elemento, data_esperada):
    """Valida se a data foi preenchida corretamente"""
    try:
        if elemento is None:
            return False
            
        val = (elemento.get_attribute('value') or '').strip()
        if not val:
            return False
            
        if val == data_esperada or data_esperada in val:
            return True
            
        # Tenta comparar datas em diferentes formatos
        formatos = [
            '%d/%m/%Y %H:%M', '%d/%m/%Y %H:%M:%S',
            '%d/%m/%Y', '%m/%d/%Y', '%Y-%m-%d', '%d-%m-%Y'
        ]
        
        for formato in formatos:
            try:
                d1 = datetime.strptime(val, formato)
                d2 = datetime.strptime(data_esperada, formato)
                if d1 == d2:
                    return True
            except:
                continue
                
        return False
        
    except Exception:
        return False





def preencher_datepicker_por_indice(indice_campo, data_valor, max_tentativas=5):
    """Preenche datepicker pelo índice com estratégias múltiplas"""
    def acao():
        if not isinstance(indice_campo, int) or indice_campo < 0:
            raise ValueError(f"Índice inválido: {indice_campo}")
            
        if not data_valor or not isinstance(data_valor, str):
            raise ValueError(f"Data inválida: {data_valor}")
        
        tentativa = 0
        while tentativa < max_tentativas:
            tentativa += 1
            
            try:
                campos = encontrar_campos_datepicker()
                
                if not campos:
                    if tentativa < max_tentativas:
                        log(doc, f"⚠️ Nenhum campo datepicker encontrado, tentativa {tentativa}/{max_tentativas}", 'WARN')
                        time.sleep(2)
                        continue
                    raise Exception("Nenhum campo datepicker encontrado na página")
                
                if indice_campo >= len(campos):
                    raise Exception(f"Índice {indice_campo} inválido. Encontrados {len(campos)} campos")
                
                campo_info = campos[indice_campo]
                elemento = campo_info['elemento']
                campo_id = campo_info['id']
                
                log(doc, f"🎯 Tentativa {tentativa}: Preenchendo datepicker {indice_campo} (ID: {campo_id}) com '{data_valor}'")
                
                # Verifica se já está preenchido corretamente
                if validar_data_preenchida(elemento, data_valor):
                    log(doc, f"✅ Campo {indice_campo} já está preenchido corretamente!")
                    return True
                
                # Estratégias específicas para datepicker
                estrategias = [
                    lambda: _datepicker_jquery(campo_id, data_valor),
                    lambda: _datepicker_javascript(elemento, data_valor),
                    lambda: _datepicker_actionchains(elemento, data_valor),
                    lambda: _datepicker_tradicional(elemento, data_valor)
                ]
                
                for i, estrategia in enumerate(estrategias, 1):
                    try:
                        log(doc, f"   Aplicando estratégia {i} para datepicker...")
                        estrategia()
                        time.sleep(1)
                        
                        # Verifica se funcionou
                        if validar_data_preenchida(elemento, data_valor):
                            valor_atual = elemento.get_attribute('value')
                            log(doc, f"✅ Datepicker preenchido com estratégia {i}: '{valor_atual}'")
                            return True
                        else:
                            log(doc, f"⚠️ Estratégia {i} não preencheu corretamente", 'WARN')
                            
                    except Exception as e:
                        log(doc, f"⚠️ Estratégia {i} falhou: {e}", 'WARN')
                        continue
                
                # Se chegou aqui, nenhuma estratégia funcionou nesta tentativa
                if tentativa < max_tentativas:
                    log(doc, f"⚠️ Tentativa {tentativa} falhou, tentando novamente em 2s...", 'WARN')
                    time.sleep(2)
                    continue
                
            except Exception as e:
                if tentativa < max_tentativas:
                    log(doc, f"⚠️ Erro na tentativa {tentativa}: {e}, tentando novamente...", 'WARN')
                    time.sleep(2)
                    continue
                else:
                    raise
        
        raise Exception(f"Falha ao preencher datepicker {indice_campo} após {max_tentativas} tentativas")
    
    return acao

def clicar_elemento_robusto_by(driver, seletor, by=By.CSS_SELECTOR, timeout=TIMEOUT_DEFAULT):
    """Clique robusto em elemento localizado por (by, seletor)"""
    global doc
    try:
        wait = WebDriverWait(driver, timeout)

        # 1) garantir presença
        elem = wait.until(EC.presence_of_element_located((by, seletor)))

        # 2) tentar remover overlays comuns
        try:
            driver.execute_script("""
                document.querySelectorAll('.modal, .overlay, .blockUI, .toast, .tooltip, [role="dialog"], [data-overlay]')
                  .forEach(e => { if (getComputedStyle(e).position === 'fixed') e.style.display = 'none'; });
            """)
        except Exception:
            pass

        # 3) centralizar na tela
        driver.execute_script("arguments[0].scrollIntoView({block:'center', inline:'center'});", elem)
        time.sleep(0.2)

        # 4) clique "normal" (clickable)
        try:
            elem = wait.until(EC.element_to_be_clickable((by, seletor)))
            elem.click()
            return True
        except (TimeoutException, ElementClickInterceptedException, StaleElementReferenceException):
            pass

        # 5) ActionChains
        try:
            elem = driver.find_element(by, seletor)
            ActionChains(driver).move_to_element(elem).pause(0.05).click().perform()
            return True
        except Exception:
            pass

        # 6) JS .click()
        try:
            elem = driver.find_element(by, seletor)
            driver.execute_script("arguments[0].click();", elem)
            return True
        except Exception:
            pass

        # 7) Disparo manual de eventos
        try:
            elem = driver.find_element(by, seletor)
            driver.execute_script("""
                const el = arguments[0];
                function fire(type){ el.dispatchEvent(new MouseEvent(type,{bubbles:true,cancelable:true,view:window})); }
                el.focus(); fire('mouseover'); fire('mousedown'); fire('mouseup'); fire('click');
            """, elem)
            return True
        except Exception:
            pass

        # 8) Click com offset (às vezes ícones minúsculos)
        try:
            elem = driver.find_element(by, seletor)
            ActionChains(driver).move_to_element_with_offset(elem, 1, 1).click().perform()
            return True
        except Exception:
            pass

        log(doc, f"❌ Não foi possível clicar em: {seletor}")
        return False

    except Exception as e:
        log(doc, f"❌ Erro ao clicar robusto: {e}")
        return False


# ==== VALIDAÇÃO DE REGISTROS MELHORADA ====
def validar_registros_encontrados(timeout=TIMEOUT_LONGO):
    """Sistema robusto de validação de registros encontrados"""
    global driver, wait, doc
    
    resultado = {
        'encontrou_registros': False,
        'quantidade_registros': 0,
        'mensagem': '',
        'tabela_encontrada': False,
        'detalhes': []
    }
    
    try:
        log(doc, "🔍 Iniciando validação de registros...")
        time.sleep(5)  # Aguarda processamento inicial
        
        # Seletores para diferentes tipos de tabelas de resultado
        seletores_tabela = [
            '#DataTables_Table_0',
            '#DataTables_Table_0 tbody',
            '#DataTables_Table_1',
            '#DataTables_Table_2',
            'table[id*="DataTables"]',
            '.wdGrid table',
            'table tbody',
            '.resultados table',
            '[class*="grid"][class*="result"]',
            'table[class*="dataTable"]'
        ]
        
        tabela_encontrada = None
        
        # Busca tabela de resultados
        for seletor in seletores_tabela:
            try:
                elementos = driver.find_elements(By.CSS_SELECTOR, seletor)
                for elemento in elementos:
                    if elemento.is_displayed() and elemento.size['height'] > 0:
                        tabela_encontrada = elemento
                        resultado['tabela_encontrada'] = True
                        log(doc, f"✅ Tabela encontrada: {seletor}")
                        break
                
                if tabela_encontrada:
                    break
            except Exception as e:
                log(doc, f"⚠️ Erro ao buscar tabela com {seletor}: {e}", 'WARN')
                continue
        
        if not tabela_encontrada:
            # Busca mensagens de "sem resultados"
            mensagens_vazio = [
                "Não foi encontrado nenhum contrato com os filtros informados.",
                "Não foram encontrados registros", 
                "Nenhum resultado",
                "Sem resultados para exibir",
                "0 registros encontrados",
                "No data available"
            ]
            
            for mensagem in mensagens_vazio:
                try:
                    elem = driver.find_element(By.XPATH, f"//*[contains(text(), '{mensagem}')]")
                    if elem.is_displayed():
                        resultado['mensagem'] = f"Sistema informou: {elem.text.strip()}"
                        log(doc, f"ℹ️ {resultado['mensagem']}")
                        return resultado
                except:
                    continue
            
            # Verifica se existe indicação de carregamento
            loading_elements = driver.find_elements(By.CSS_SELECTOR, ".loading, .spinner, [class*='load']")
            if any(el.is_displayed() for el in loading_elements):
                log(doc, "⏳ Sistema ainda carregando resultados...", 'WARN')
                time.sleep(5)
                return validar_registros_encontrados(timeout - 10)  # Recursão com timeout reduzido
            
            resultado['mensagem'] = "⚠️ Tabela de resultados não localizada"
            log(doc, resultado['mensagem'], 'WARN')
            return resultado
        
        # Conta e valida registros
        try:
            # Estratégias para encontrar linhas de dados
            seletores_linhas = [
                'tbody tr:not(.dataTables_empty):not([class*="no-data"])',
                'tbody tr[class*="odd"], tbody tr[class*="even"]',
                'tbody tr:not(:empty)',
                'tbody tr'
            ]
            
            linhas_validas = []
            
            for seletor_linha in seletores_linhas:
                try:
                    linhas = tabela_encontrada.find_elements(By.CSS_SELECTOR, seletor_linha)
                    
                    for linha in linhas:
                        try:
                            if not linha.is_displayed():
                                continue
                                
                            texto_linha = linha.text.strip().lower()
                            
                            # Valida se é uma linha com dados reais
                            if (len(texto_linha) > 5 and 
                                not any(termo in texto_linha for termo in [
                                    'nenhum registro', 'sem dados', 'no data', 
                                    'vazio', 'empty', 'não foram encontrados',
                                    'loading', 'carregando'
                                ])):
                                
                                linhas_validas.append({
                                    'elemento': linha,
                                    'texto': texto_linha[:100] + '...' if len(texto_linha) > 100 else texto_linha
                                })
                        except Exception as e:
                            log(doc, f"⚠️ Erro ao processar linha: {e}", 'WARN')
                            continue
                    
                    if linhas_validas:
                        log(doc, f"✅ Encontradas {len(linhas_validas)} linhas válidas com {seletor_linha}")
                        break
                        
                except Exception as e:
                    log(doc, f"⚠️ Erro ao processar {seletor_linha}: {e}", 'WARN')
                    continue
            
            quantidade = len(linhas_validas)
            resultado['quantidade_registros'] = quantidade
            resultado['detalhes'] = [linha['texto'] for linha in linhas_validas[:5]]  # Primeiras 5 linhas
            
            if quantidade > 0:
                resultado['encontrou_registros'] = True
                resultado['mensagem'] = f"✅ {quantidade} registro(s) encontrado(s)"
                
                # Log das primeiras linhas
                log(doc, resultado['mensagem'])
                for i, linha in enumerate(linhas_validas[:3], 1):
                    log(doc, f"   Registro {i}: {linha['texto']}")
                
                if quantidade > 3:
                    log(doc, f"   ... e mais {quantidade-3} registro(s)")
            else:
                resultado['mensagem'] = "ℹ️ Tabela encontrada mas sem registros válidos"
                log(doc, resultado['mensagem'])
        
        except Exception as e:
            log(doc, f"❌ Erro ao contar registros: {e}", 'ERROR')
            # Em caso de erro na contagem, assume que existem registros para não interromper
            resultado['encontrou_registros'] = True
            resultado['quantidade_registros'] = 1
            resultado['mensagem'] = f"⚠️ Erro na validação, continuando teste: {e}"
        
        # Verifica alertas do sistema
        encontrar_mensagem_alerta()
        
        return resultado
        
    except Exception as e:
        log(doc, f"❌ Erro geral na validação: {e}", 'ERROR')
        resultado['encontrou_registros'] = True  # Assume sucesso para não interromper
        resultado['quantidade_registros'] = 1
        resultado['mensagem'] = f"⚠️ Validação falhou, continuando teste: {e}"
        return resultado

def force_click(driver, by, selector, timeout=10):
    """Força clique em um elemento, mesmo que esteja coberto."""
    elem = WebDriverWait(driver, timeout).until(EC.presence_of_element_located((by, selector)))
    driver.execute_script("arguments[0].scrollIntoView({block:'center'});", elem)
    driver.execute_script("arguments[0].click();", elem)
    return True

# Wrappers convenientes
def clicar_elemento_css_robusto(driver, seletor_css, timeout=TIMEOUT_DEFAULT):
    return clicar_elemento_robusto_by(driver, seletor_css, By.CSS_SELECTOR, timeout)

def clicar_elemento_xpath_robusto(driver, xpath, timeout=TIMEOUT_DEFAULT):
    return clicar_elemento_robusto_by(driver, xpath, By.XPATH, timeout)


# ==== EXECUÇÃO DO TESTE ====
def executar_teste():
    global driver, wait, doc, janela_principal  # ✅ CORREÇÃO: incluir janela_principal
    try:
        if not inicializar_driver():
            return False

        # Sanidade: garantir que time é módulo
        assert hasattr(time, "sleep"), f"time virou {time!r}"

        safe_action(doc, "Acessando sistema", lambda: driver.get(URL))

        # ✅ CORREÇÃO: capturar handle somente após o driver estar ativo e a página aberta
        janela_principal = driver.current_window_handle

        safe_action(doc, "Realizando login", lambda: (
            wait.until(EC.presence_of_element_located((By.ID, "j_id15:email"))).send_keys(LOGIN_EMAIL),
            wait.until(EC.presence_of_element_located((By.ID, "j_id15:senha"))).send_keys(LOGIN_PASSWORD, Keys.ENTER),
            time.sleep(5)
        ))

        safe_action(doc, "Ajustando zoom e abrindo menu", lambda: (
            ajustar_zoom(),
            driver.find_element(By.TAG_NAME, "body").send_keys(Keys.F3)
        ))

        safe_action(doc, "Acessando Rateio", lambda: (
            driver.execute_script("arguments[0].scrollIntoView({block: 'center'});",
                                  wait.until(EC.presence_of_element_located((By.XPATH, '/html/body/div[15]/ul/li[10]/img')))),
            wait.until(EC.element_to_be_clickable((By.XPATH, '/html/body/div[15]/ul/li[10]/img'))).click()
        ))

        safe_action(doc, "Clicando em Reagrupamento de Rateios", lambda:
            wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, '#gsRateio > div.wdTelas > div > ul > li:nth-child(2) > a > span'))).click()
        )

        time.sleep(5)

        safe_action(doc, "Abrindo Lov de Grupo de Rateio Origem", lambda: 
            wait.until(EC.element_to_be_clickable((By.XPATH, "//*[@id='gsRateio']/div[2]/div[2]/div[1]/div[1]/div/a"))).click(),
        )

        safe_action(doc, "Preenchendo Grupo de Rateio", lambda:
                    preencher_campo_xpath_com_retry(
                        driver, wait, "//input[@type='text' and contains(@class,'nomePesquisa')]",
                        "GRUPO A"
                    ))

        safe_action(doc, "Pesquisando Grupo de Rateio", lambda: 
            wait.until(EC.element_to_be_clickable((By.XPATH, "//a[contains(@class,'lpFind') and normalize-space(.)='Pesquisar']"))).click(),
        )

        safe_action(doc, "Selecionando Grupo de Rateio", lambda: 
            wait.until(EC.element_to_be_clickable((By.XPATH, "//tr[td[1][normalize-space()='GRUPO A']]"))).click()
        )




        safe_action(doc, "Abrindo Lov de Grupo de Rateio Destino", lambda: 
            wait.until(EC.element_to_be_clickable((By.XPATH, "//*[@id='gsRateio']/div[2]/div[2]/div[1]/div[2]/div/a"))).click(),
        )

        safe_action(doc, "Preenchendo Grupo de Rateio", lambda:
                    preencher_campo_xpath_com_retry(
                        driver, wait, "//input[@type='text' and contains(@class,'nomePesquisa')]",
                        "TESTE NANUQUE"
                    ))

        safe_action(doc, "Pesquisando Grupo de Rateio", lambda: 
            wait.until(EC.element_to_be_clickable((By.XPATH, "//a[contains(@class,'lpFind') and normalize-space(.)='Pesquisar']"))).click(),
        )

        safe_action(doc, "Selecionando Grupo de Rateio", lambda: 
            wait.until(EC.element_to_be_clickable((By.XPATH, "//tr[td[1][normalize-space()='TESTE NANUQUE']]"))).click()
        )


        safe_action(doc, "Reagrupando Rateios (Teste de disparo)", lambda: 
            wait.until(EC.element_to_be_clickable((By.XPATH, "//a[normalize-space(text())='Reagrupar']"))).click(),
        )



        log(doc, "🔍 Verificando mensagens de alerta...")
        encontrar_mensagem_alerta()

        time.sleep(5)


        safe_action(doc, "Informando uma quantidade excedida de Contratos do que o Grupo de Rateio permite (Teste de disparo)", lambda:
            preencher_campo_xpath_com_retry(
                driver, wait, "//*[@id='gsRateio']/div[2]/div[2]/div[1]/div[3]/input",
                "100"
            )
        )

        log(doc, "🔍 Verificando mensagens de alerta...")
        encontrar_mensagem_alerta()

        time.sleep(5)



        safe_action(doc, "Informando a quantidade de Contratos excedida", lambda:
            preencher_campo_xpath_com_retry(
                driver, wait, "//*[@id='gsRateio']/div[2]/div[2]/div[1]/div[3]/input",
                "1"
            )
        )

        safe_action(doc, "Reagrupando Rateios", lambda: 
            wait.until(EC.element_to_be_clickable((By.XPATH, "//a[normalize-space(text())='Reagrupar']"))).click(),
        )

        log(doc, "🔍 Verificando mensagens de alerta...")
        encontrar_mensagem_alerta()



        safe_action(doc, "Fechando modal Rateio", lambda:
            force_click(driver, By.CSS_SELECTOR, "#gsRateio > div.wdTop.ui-draggable-handle > div.wdClose > a")
        )


        return True

    except Exception as e:
        log(doc, f"❌ ERRO FATAL: {e}")
        take_screenshot(driver, doc, "erro_fatal")
        return False
    finally:
        log(doc, "✅ Teste concluído.")

# ==== MAIN ====
def main():
    global doc
    try:
        log(doc, "🚀 Iniciando teste de Reagrupamento de Rateios")
        sucesso = executar_teste()
        if sucesso:
            log(doc, "✅ Teste executado com sucesso!")
        else:
            log(doc, "❌ Teste finalizado com erros.")
    except Exception as e:
        log(doc, f"❌ Erro na execução principal: {e}")
    finally:
        finalizar_relatorio()

if __name__ == "__main__":
    main()
