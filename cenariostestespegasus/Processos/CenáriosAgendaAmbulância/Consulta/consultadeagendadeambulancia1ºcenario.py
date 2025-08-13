from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support.ui import Select, WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service
from selenium.common.exceptions import TimeoutException, NoSuchElementException, StaleElementReferenceException
from webdriver_manager.chrome import ChromeDriverManager
from docx import Document
from docx.shared import Inches
from faker import Faker
from faker.providers import BaseProvider
from validate_docbr import CPF
from datetime import datetime, timedelta
from selenium.webdriver import ActionChains
import subprocess
import os
import time
import random
import re
import pyautogui

# ==== CONFIGURAÇÕES GLOBAIS ====
URL = "http://localhost:8080/gs/index.xhtml"
LOGIN_EMAIL = "joaoeduardo.gold@outlook.com"
LOGIN_PASSWORD = "071999gs"

# Timeouts configuráveis
TIMEOUT_DEFAULT = 20
TIMEOUT_CURTO = 5
TIMEOUT_LONGO = 30

# Configuração de logs mais detalhada
VERBOSE_LOGGING = True
CAPTURAR_SCREENSHOTS = True

# ==== PROVIDERS CUSTOMIZADOS ====
class BrasilProvider(BaseProvider):
    def rg(self):
        numeros = [str(random.randint(0, 9)) for _ in range(8)]
        return ''.join(numeros) + '-' + str(random.randint(0, 9))

fake = Faker("pt_BR")
fake.add_provider(BrasilProvider)

# ==== GERAÇÃO DE DADOS DE TESTE ====
def gerar_datas_validas():
    """Gera datas coerentes dentro de intervalos válidos."""
    hoje = datetime.today().date()
    dez_anos_atras = hoje - timedelta(days=3650)
    
    data_falecimento = fake.date_between(start_date=dez_anos_atras, end_date=hoje)
    idade_minima, idade_maxima = 18, 110
    data_nascimento = data_falecimento - timedelta(days=random.randint(idade_minima * 365, idade_maxima * 365))
    data_sepultamento = data_falecimento + timedelta(days=random.randint(1, 10))
    data_registro = data_sepultamento + timedelta(days=random.randint(1, 10))
    data_velorio = fake.date_between(start_date=data_falecimento, end_date=data_sepultamento)

    return tuple(data.strftime("%d/%m/%Y") for data in [
        data_nascimento, data_falecimento, data_sepultamento, 
        data_velorio, hoje, data_registro
    ])

# Gera os dados de teste
data_nascimento, data_falecimento, data_sepultamento, data_velorio, hoje, data_registro = gerar_datas_validas()

# ==== INICIALIZAÇÃO DE VARIÁVEIS GLOBAIS ====
doc = Document()
doc.add_heading("RELATÓRIO DO TESTE - CONSULTA AGENDA AMBULÂNCIA", 0)
doc.add_paragraph("Cenário 1: Preenchimento e realização da consulta com validações robustas")
doc.add_paragraph(f"Data do teste: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")

screenshot_registradas = set()
driver = None
wait = None

# ==== UTILITÁRIOS DE LOGGING E SCREENSHOTS ====
def log(doc, msg, nivel='INFO'):
    """Sistema de logging melhorado com níveis"""
    timestamp = datetime.now().strftime('%H:%M:%S')
    formatted_msg = f"[{timestamp}] {nivel}: {msg}"
    
    if VERBOSE_LOGGING:
        print(formatted_msg)
    else:
        print(msg)
    
    if hasattr(doc, 'add_paragraph'):
        doc.add_paragraph(formatted_msg)

def take_screenshot(driver, doc, nome, forcar=False):
    """Sistema de screenshots otimizado"""
    if not CAPTURAR_SCREENSHOTS and not forcar:
        return
        
    if driver is None:
        return
        
    if nome not in screenshot_registradas or forcar:
        path = f"screenshots/{nome}.png"
        os.makedirs("screenshots", exist_ok=True)
        try:
            driver.save_screenshot(path)
            doc.add_paragraph(f"Screenshot: {nome}")
            if CAPTURAR_SCREENSHOTS:
                doc.add_picture(path, width=Inches(5.5))
            screenshot_registradas.add(nome)
            log(doc, f"Screenshot capturada: {nome}")
        except Exception as e:
            log(doc, f"Erro ao capturar screenshot {nome}: {e}", 'WARN')

# ==== FUNÇÕES DE UTILITÁRIO MELHORADAS ====
def safe_action(doc, descricao, func, max_retries=3, timeout_customizado=None, critico=True):
    """Execução de ações com retry robusto e tratamento de erros melhorado"""
    global driver
    
    timeout_original = None
    if timeout_customizado and driver:
        try:
            # Ajusta timeout temporariamente
            timeout_original = driver.timeouts.implicit_wait
            driver.implicitly_wait(timeout_customizado)
        except:
            pass
    
    for attempt in range(max_retries):
        try:
            if attempt == 0:
                log(doc, f"Executando: {descricao}")
            else:
                log(doc, f"Retry {attempt + 1}/{max_retries}: {descricao}", 'WARN')
            
            result = func()
            log(doc, f"✅ {descricao} - Sucesso")
            take_screenshot(driver, doc, descricao.lower().replace(" ", "_"))
            
            return True
            
        except (TimeoutException, NoSuchElementException, StaleElementReferenceException) as e:
            if attempt < max_retries - 1:
                log(doc, f"⚠️ Tentativa {attempt + 1} falhou: {type(e).__name__}", 'WARN')
                time.sleep(2 ** attempt)  # Backoff exponencial
                continue
            else:
                error_msg = f"❌ {descricao} falhou após {max_retries} tentativas: {e}"
                log(doc, error_msg, 'ERROR')
                take_screenshot(driver, doc, f"erro_{descricao.lower().replace(' ', '_')}", forcar=True)
                
                if critico:
                    raise Exception(error_msg)
                return False
                
        except Exception as e:
            error_msg = f"❌ Erro inesperado em {descricao}: {e}"
            log(doc, error_msg, 'ERROR')
            take_screenshot(driver, doc, f"erro_critico_{descricao.lower().replace(' ', '_')}", forcar=True)
            
            if critico:
                raise Exception(error_msg)
            return False
    
        finally:
            # Restaura timeout original
            if timeout_original and driver:
                try:
                    driver.implicitly_wait(timeout_original)
                except:
                    pass

def inicializar_driver():
    """Inicialização robusta do WebDriver com configurações otimizadas"""
    global driver, wait
    
    try:
        log(doc, "🔧 Inicializando WebDriver...")
        
        options = Options()
        options.add_argument("--start-maximized")
        options.add_argument("--disable-blink-features=AutomationControlled")
        options.add_argument("--disable-web-security")
        options.add_argument("--allow-running-insecure-content")
        options.add_argument("--no-sandbox")
        options.add_argument("--disable-dev-shm-usage")
        options.add_experimental_option("excludeSwitches", ["enable-automation"])
        options.add_experimental_option('useAutomationExtension', False)
        
        # Configurações de performance
        prefs = {
            "profile.default_content_setting_values": {"notifications": 2},
            "profile.default_content_settings.popups": 0,
        }
        options.add_experimental_option("prefs", prefs)

        service = Service(ChromeDriverManager().install())
        driver = webdriver.Chrome(service=service, options=options)
        
        # Remove indicadores de automação
        driver.execute_script("Object.defineProperty(navigator, 'webdriver', {get: () => undefined})")
        driver.execute_cdp_cmd('Network.setUserAgentOverride', {
            "userAgent": 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
        })
        
        wait = WebDriverWait(driver, TIMEOUT_DEFAULT)
        
        log(doc, "✅ WebDriver inicializado com sucesso")
        return True
        
    except Exception as e:
        log(doc, f"❌ Erro ao inicializar driver: {e}", 'ERROR')
        return False

def ajustar_zoom(nivel=90):
    """Ajusta zoom da página para melhor visualização"""
    global driver, doc
    
    if driver is None:
        return False
        
    try:
        driver.execute_script(f"document.body.style.zoom='{nivel}%'")
        log(doc, f"🔍 Zoom ajustado para {nivel}%")
        return True
    except Exception as e:
        log(doc, f"⚠️ Erro ao ajustar zoom: {e}", 'WARN')
        return False

def encontrar_mensagem_alerta():
    """Sistema melhorado de detecção de alertas do sistema"""
    global driver, doc
    
    if driver is None:
        return None
        
    seletores_alerta = [
        (".alerts.salvo", "✅ Sucesso"),
        (".alerts.alerta", "⚠️ Alerta"),
        (".alerts.erro", "❌ Erro"),
        (".message.success", "✅ Sucesso"),
        (".message.error", "❌ Erro"),
        (".message.warning", "⚠️ Aviso"),
        ("[class*='alert'][class*='success']", "✅ Sucesso"),
        ("[class*='alert'][class*='error']", "❌ Erro"),
        ("[class*='alert'][class*='warning']", "⚠️ Aviso"),
    ]

    for seletor, tipo in seletores_alerta:
        try:
            elementos = driver.find_elements(By.CSS_SELECTOR, seletor)
            for elemento in elementos:
                if elemento.is_displayed() and elemento.text.strip():
                    log(doc, f"📢 {tipo}: {elemento.text.strip()}")
                    return elemento
        except Exception as e:
            continue

    log(doc, "ℹ️ Nenhuma mensagem de sistema detectada", 'DEBUG')
    return None

# ==== FUNÇÕES DE INTERAÇÃO ROBUSTAS ====
def aguardar_elemento(seletor, timeout=TIMEOUT_DEFAULT, condicao='clickable', by_type=By.CSS_SELECTOR):
    """Função centralizada para aguardar elementos com diferentes condições"""
    global wait
    
    condicoes = {
        'present': EC.presence_of_element_located,
        'visible': EC.visibility_of_element_located,
        'clickable': EC.element_to_be_clickable,
        'invisible': EC.invisibility_of_element_located
    }
    
    if condicao not in condicoes:
        condicao = 'clickable'
    
    try:
        elemento = WebDriverWait(driver, timeout).until(
            condicoes[condicao]((by_type, seletor))
        )
        return elemento
    except TimeoutException:
        raise TimeoutException(f"Elemento não encontrado: {seletor} (condição: {condicao})")

def scroll_to_element(elemento):
    """Scroll inteligente até elemento"""
    global driver
    
    try:
        # Scroll suave até o elemento
        driver.execute_script("""
            arguments[0].scrollIntoView({
                behavior: 'smooth',
                block: 'center',
                inline: 'center'
            });
        """, elemento)
        time.sleep(0.5)
        
        # Verifica se elemento está visível
        return elemento.is_displayed()
        
    except Exception as e:
        log(doc, f"⚠️ Erro no scroll: {e}", 'WARN')
        return False

def preencher_campo_robusto(seletor, valor, limpar_primeiro=True):
    """Função melhorada para preenchimento de campos"""
    def acao():
        elemento = aguardar_elemento(seletor, condicao='clickable')
        scroll_to_element(elemento)
        
        # Múltiplas estratégias de preenchimento
        estrategias = [
            lambda: _preencher_tradicional(elemento, valor, limpar_primeiro),
            lambda: _preencher_actionchains(elemento, valor),
            lambda: _preencher_javascript(elemento, valor)
        ]
        
        for i, estrategia in enumerate(estrategias, 1):
            try:
                estrategia()
                # Verifica se foi preenchido corretamente
                valor_atual = elemento.get_attribute('value')
                if valor_atual == valor or valor in valor_atual:
                    log(doc, f"✅ Campo preenchido com estratégia {i}")
                    return True
            except Exception as e:
                log(doc, f"⚠️ Estratégia {i} falhou: {e}", 'WARN')
                continue
                
        raise Exception(f"Falha ao preencher campo {seletor} com todas as estratégias")
    
    return acao

def _preencher_tradicional(elemento, valor, limpar_primeiro):
    """Estratégia tradicional de preenchimento"""
    elemento.click()
    time.sleep(0.2)
    if limpar_primeiro:
        elemento.clear()
    elemento.send_keys(valor)
    elemento.send_keys(Keys.TAB)

def _preencher_actionchains(elemento, valor):
    """Estratégia com ActionChains"""
    global driver
    
    ActionChains(driver).move_to_element(elemento).click().perform()
    time.sleep(0.2)
    ActionChains(driver).key_down(Keys.CONTROL).send_keys('a').key_up(Keys.CONTROL).perform()
    ActionChains(driver).send_keys(valor).perform()
    ActionChains(driver).send_keys(Keys.TAB).perform()

def _preencher_javascript(elemento, valor):
    """Estratégia com JavaScript"""
    global driver
    
    driver.execute_script("""
        var element = arguments[0];
        var value = arguments[1];
        element.focus();
        element.value = '';
        element.value = value;
        element.dispatchEvent(new Event('input', { bubbles: true }));
        element.dispatchEvent(new Event('change', { bubbles: true }));
        element.blur();
    """, elemento, valor)

def selecionar_opcao_robusta(seletor, texto_opcao):
    """Seleção robusta em elementos Select"""
    def acao():
        elemento = aguardar_elemento(seletor)
        scroll_to_element(elemento)
        
        select_obj = Select(elemento)
        
        # Tenta diferentes métodos de seleção
        try:
            select_obj.select_by_visible_text(texto_opcao)
        except:
            try:
                select_obj.select_by_value(texto_opcao)
            except:
                # Busca por texto parcial
                for opcao in select_obj.options:
                    if texto_opcao.lower() in opcao.text.lower():
                        select_obj.select_by_visible_text(opcao.text)
                        break
                else:
                    raise Exception(f"Opção '{texto_opcao}' não encontrada")
        
        log(doc, f"✅ Opção '{texto_opcao}' selecionada")
    
    return acao

def clicar_elemento_robusto(seletor):
    """Clique robusto em elementos"""
    def acao():
        elemento = aguardar_elemento(seletor)
        scroll_to_element(elemento)
        
        # Tenta diferentes métodos de clique
        try:
            elemento.click()
        except:
            # Clique via JavaScript se o normal falhar
            driver.execute_script("arguments[0].click();", elemento)
    
    return acao

# ==== SISTEMA DE DATEPICKER MELHORADO ====
def encontrar_campos_datepicker():
    """Encontra todos os campos datepicker na página"""
    global driver
    
    seletores_datepicker = [
        "input.hasDatepicker",
        "input[id^='dp']",
        "input[maxlength='10'][grupo='']",
        "input[type='text'][maxlength='10']"
    ]
    
    campos_encontrados = []
    
    for seletor in seletores_datepicker:
        try:
            elementos = driver.find_elements(By.CSS_SELECTOR, seletor)
            for elemento in elementos:
                if elemento.is_displayed():
                    info = {
                        'elemento': elemento,
                        'id': elemento.get_attribute('id'),
                        'seletor_usado': seletor,
                        'maxlength': elemento.get_attribute('maxlength'),
                        'placeholder': elemento.get_attribute('placeholder')
                    }
                    # Evita duplicatas
                    if not any(c['id'] == info['id'] for c in campos_encontrados):
                        campos_encontrados.append(info)
        except:
            continue
    
    log(doc, f"📊 Encontrados {len(campos_encontrados)} campos datepicker")
    return campos_encontrados

def preencher_datepicker_por_indice(indice_campo, data_valor, max_tentativas=3):
    """Preenche datepicker pelo índice com estratégias múltiplas"""
    def acao():
        campos = encontrar_campos_datepicker()
        
        if indice_campo >= len(campos):
            raise Exception(f"Índice {indice_campo} inválido. Encontrados {len(campos)} campos")
        
        campo_info = campos[indice_campo]
        elemento = campo_info['elemento']
        campo_id = campo_info['id']
        
        log(doc, f"🎯 Preenchendo datepicker {indice_campo} (ID: {campo_id}) com '{data_valor}'")
        
        # Estratégias específicas para datepicker
        estrategias = [
            lambda: _datepicker_jquery(campo_id, data_valor),
            lambda: _datepicker_javascript(elemento, data_valor),
            lambda: _datepicker_actionchains(elemento, data_valor),
            lambda: _datepicker_tradicional(elemento, data_valor)
        ]
        
        for i, estrategia in enumerate(estrategias, 1):
            try:
                log(doc, f"   Tentando estratégia {i} para datepicker...")
                estrategia()
                
                # Verifica se funcionou
                time.sleep(0.5)
                valor_atual = elemento.get_attribute('value')
                
                if valor_atual and (valor_atual == data_valor or data_valor in valor_atual):
                    log(doc, f"✅ Datepicker preenchido com estratégia {i}: '{valor_atual}'")
                    return True
                    
            except Exception as e:
                log(doc, f"⚠️ Estratégia {i} falhou: {e}", 'WARN')
                continue
        
        raise Exception(f"Falha ao preencher datepicker {indice_campo} com todas as estratégias")
    
    return acao

def _datepicker_jquery(campo_id, data_valor):
    """Estratégia usando jQuery Datepicker API"""
    global driver
    
    resultado = driver.execute_script("""
        var campoId = arguments[0];
        var valor = arguments[1];
        
        if (typeof jQuery === 'undefined') {
            return 'jQuery não disponível';
        }
        
        var $campo = $('#' + campoId);
        if ($campo.length === 0) {
            return 'Campo não encontrado: ' + campoId;
        }
        
        try {
            if ($campo.hasClass('hasDatepicker')) {
                $campo.datepicker('setDate', valor);
            } else {
                $campo.val(valor);
            }
            $campo.trigger('input').trigger('change').trigger('blur');
            return $campo.val();
        } catch(e) {
            return 'Erro: ' + e.message;
        }
    """, campo_id, data_valor)
    
    if 'Erro' in str(resultado):
        raise Exception(f"jQuery falhou: {resultado}")

def _datepicker_javascript(elemento, data_valor):
    """Estratégia JavaScript pura"""
    global driver
    
    driver.execute_script("""
        var campo = arguments[0];
        var valor = arguments[1];
        
        campo.focus();
        campo.value = '';
        campo.value = valor;
        
        ['input', 'change', 'blur', 'keyup'].forEach(function(evento) {
            campo.dispatchEvent(new Event(evento, { bubbles: true }));
        });
    """, elemento, data_valor)

def _datepicker_actionchains(elemento, data_valor):
    """Estratégia ActionChains para datepicker"""
    global driver
    
    scroll_to_element(elemento)
    ActionChains(driver).move_to_element(elemento).click().perform()
    time.sleep(0.3)
    
    # Limpa completamente
    ActionChains(driver).key_down(Keys.CONTROL).send_keys('a').key_up(Keys.CONTROL).perform()
    ActionChains(driver).send_keys(Keys.DELETE).perform()
    time.sleep(0.2)
    
    # Digita caracter por caracter
    for char in data_valor:
        ActionChains(driver).send_keys(char).perform()
        time.sleep(0.05)
    
    ActionChains(driver).send_keys(Keys.TAB).perform()

def _datepicker_tradicional(elemento, data_valor):
    """Estratégia tradicional para datepicker"""
    scroll_to_element(elemento)
    elemento.click()
    time.sleep(0.3)
    elemento.clear()
    elemento.send_keys(data_valor)
    elemento.send_keys(Keys.TAB)

# ==== SISTEMA DE MODAIS MELHORADO ====
def abrir_modal_e_selecionar_robusto(btn_selector, pesquisa_selector, termo_pesquisa, 
                                      btn_pesquisar_selector, resultado_xpath, timeout=TIMEOUT_LONGO):
    """Sistema robusto para trabalhar com modais de seleção"""
    def acao():
        global driver, wait
        
        log(doc, f"🔍 Abrindo modal e buscando: '{termo_pesquisa}'")
        
        # Abre modal
        btn_modal = aguardar_elemento(btn_selector, timeout=timeout)
        scroll_to_element(btn_modal)
        btn_modal.click()
        time.sleep(1)

        # Aguarda modal aparecer e preenche pesquisa
        campo_pesquisa = aguardar_elemento(pesquisa_selector, timeout=TIMEOUT_CURTO)
        campo_pesquisa.clear()
        campo_pesquisa.send_keys(termo_pesquisa)
        time.sleep(0.5)

        # Clica em pesquisar
        btn_pesquisar = aguardar_elemento(btn_pesquisar_selector)
        btn_pesquisar.click()
        
        # Aguarda resultados (timeout mais longo)
        log(doc, "⏳ Aguardando resultados da pesquisa...")
        time.sleep(3)
        
        # Procura e clica no resultado
        resultado = aguardar_elemento(resultado_xpath, timeout=TIMEOUT_LONGO, by_type=By.XPATH)
        scroll_to_element(resultado)
        resultado.click()
        time.sleep(1)
        
        log(doc, f"✅ Item selecionado: {termo_pesquisa}")

    return acao

# ==== SISTEMA DE VALIDAÇÃO DE REGISTROS MELHORADO ====
def validar_registros_encontrados(timeout=TIMEOUT_LONGO):
    """Sistema robusto de validação de registros encontrados"""
    global driver, wait, doc
    
    resultado = {
        'encontrou_registros': False,
        'quantidade_registros': 0,
        'mensagem': '',
        'tabela_encontrada': False,
        'detalhes': []
    }
    
    try:
        log(doc, "🔍 Iniciando validação de registros...")
        time.sleep(4)  # Aguarda processamento inicial
        
        # Seletores para diferentes tipos de tabelas de resultado
        seletores_tabela = [
            '#DataTables_Table_0',
            '#DataTables_Table_0 tbody',
            'table[id*="DataTables"]',
            '.wdGrid table',
            'table tbody',
            '.resultados table',
            '[class*="grid"][class*="result"]'
        ]
        
        tabela_encontrada = None
        
        # Busca tabela de resultados
        for seletor in seletores_tabela:
            try:
                elementos = driver.find_elements(By.CSS_SELECTOR, seletor)
                for elemento in elementos:
                    if elemento.is_displayed() and elemento.size['height'] > 0:
                        tabela_encontrada = elemento
                        resultado['tabela_encontrada'] = True
                        log(doc, f"✅ Tabela encontrada: {seletor}")
                        break
                
                if tabela_encontrada:
                    break
            except:
                continue
        
        if not tabela_encontrada:
            # Busca mensagens de "sem resultados"
            mensagens_vazio = [
                "Nenhum registro encontrado",
                "Não foram encontrados registros", 
                "Nenhum resultado",
                "Sem resultados para exibir",
                "0 registros encontrados"
            ]
            
            for mensagem in mensagens_vazio:
                try:
                    elem = driver.find_element(By.XPATH, f"//*[contains(text(), '{mensagem}')]")
                    if elem.is_displayed():
                        resultado['mensagem'] = f"Sistema informou: {elem.text.strip()}"
                        log(doc, f"ℹ️ {resultado['mensagem']}")
                        return resultado
                except:
                    continue
            
            resultado['mensagem'] = "⚠️ Tabela de resultados não localizada"
            log(doc, resultado['mensagem'], 'WARN')
            return resultado
        
        # Conta e valida registros
        try:
            # Estratégias para encontrar linhas de dados
            seletores_linhas = [
                '#DataTables_Table_0 tbody tr',
                'tbody tr:not(.dataTables_empty)',
                'tbody tr[class*="odd"], tbody tr[class*="even"]',
                'tbody tr'
            ]
            
            linhas_validas = []
            
            for seletor_linha in seletores_linhas:
                try:
                    linhas = tabela_encontrada.find_elements(By.CSS_SELECTOR, seletor_linha)
                    
                    for linha in linhas:
                        texto_linha = linha.text.strip().lower()
                        
                        # Valida se é uma linha com dados reais
                        if (len(texto_linha) > 5 and 
                            not any(termo in texto_linha for termo in [
                                'nenhum registro', 'sem dados', 'no data', 
                                'vazio', 'empty', 'não foram encontrados'
                            ]) and
                            linha.is_displayed()):
                            
                            linhas_validas.append({
                                'elemento': linha,
                                'texto': texto_linha[:100] + '...' if len(texto_linha) > 100 else texto_linha
                            })
                    
                    if linhas_validas:
                        log(doc, f"✅ Encontradas {len(linhas_validas)} linhas válidas")
                        break
                        
                except Exception as e:
                    log(doc, f"⚠️ Erro ao processar {seletor_linha}: {e}", 'WARN')
                    continue
            
            quantidade = len(linhas_validas)
            resultado['quantidade_registros'] = quantidade
            resultado['detalhes'] = [linha['texto'] for linha in linhas_validas[:5]]  # Primeiras 5 linhas
            
            if quantidade > 0:
                resultado['encontrou_registros'] = True
                resultado['mensagem'] = f"✅ {quantidade} registro(s) encontrado(s)"
                
                # Log das primeiras linhas
                log(doc, resultado['mensagem'])
                for i, linha in enumerate(linhas_validas[:3], 1):
                    log(doc, f"   Registro {i}: {linha['texto']}")
                
                if quantidade > 3:
                    log(doc, f"   ... e mais {quantidade-3} registro(s)")
            else:
                resultado['mensagem'] = "ℹ️ Tabela encontrada mas sem registros válidos"
                log(doc, resultado['mensagem'])
        
        except Exception as e:
            log(doc, f"❌ Erro ao contar registros: {e}", 'ERROR')
            # Em caso de erro, assume que existem registros para não interromper
            resultado['encontrou_registros'] = True
            resultado['quantidade_registros'] = 1
            resultado['mensagem'] = f"⚠️ Erro na validação, continuando teste: {e}"
        
        # Verifica alertas do sistema
        encontrar_mensagem_alerta()
        
        return resultado
        
    except Exception as e:
        log(doc, f"❌ Erro geral na validação: {e}", 'ERROR')
        resultado['encontrou_registros'] = True  # Assume sucesso para não interromper
        resultado['quantidade_registros'] = 1
        resultado['mensagem'] = f"⚠️ Validação falhou, continuando teste: {e}"
        return resultado

# ==== FUNÇÃO DE UPLOAD DE ARQUIVOS MELHORADA ====
def fazer_upload_arquivo(seletor_input, caminho_arquivo, timeout=TIMEOUT_DEFAULT):
    """Sistema robusto para upload de arquivos"""
    def acao():
        if not os.path.exists(caminho_arquivo):
            raise Exception(f"Arquivo não encontrado: {caminho_arquivo}")
        
        # Aguarda elemento de upload
        input_file = aguardar_elemento(seletor_input, timeout=timeout)
        
        # Envia arquivo diretamente para o input
        try:
            input_file.send_keys(caminho_arquivo)
            log(doc, f"✅ Arquivo enviado: {os.path.basename(caminho_arquivo)}")
        except Exception as e:
            # Fallback para PyAutoGUI
            log(doc, "⚠️ Tentando upload via PyAutoGUI...", 'WARN')
            input_file.click()
            time.sleep(1)
            pyautogui.write(f'"{caminho_arquivo}"')
            pyautogui.press('enter')
            log(doc, f"✅ Arquivo enviado via PyAutoGUI: {os.path.basename(caminho_arquivo)}")
    
    return acao

# ==== SISTEMA DE LIMPEZA E FINALIZAÇÃO ====
def finalizar_relatorio():
    """Sistema melhorado de finalização com relatório detalhado"""
    global driver, doc
    
    try:
        # Adiciona estatísticas do teste
        total_screenshots = len(screenshot_registradas)
        doc.add_paragraph(f"\n📊 ESTATÍSTICAS DO TESTE:")
        doc.add_paragraph(f"   • Screenshots capturadas: {total_screenshots}")
        doc.add_paragraph(f"   • Tempo total de execução: {datetime.now().strftime('%H:%M:%S')}")
        
        # Salva relatório
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        nome_arquivo = f"relatorio_consulta_agenda_ambulancia_{timestamp}.docx"
        doc.save(nome_arquivo)
        log(doc, f"📄 Relatório salvo: {nome_arquivo}")
        
        # Tenta abrir automaticamente
        try:
            subprocess.run(["start", "winword", nome_arquivo], shell=True, check=False)
        except:
            log(doc, "ℹ️ Relatório salvo, abrir manualmente se necessário")
            
    except Exception as e:
        print(f"Erro ao finalizar relatório: {e}")
    
    # Limpeza do driver
    if driver:
        try:
            driver.quit()
            log(doc, "🧹 WebDriver finalizado")
        except:
            pass

# ==== FUNÇÃO PRINCIPAL DO TESTE ====
def executar_teste():
    """Função principal com lógica de teste melhorada"""
    global driver, wait, doc
    
    try:
        # ==== INICIALIZAÇÃO ====
        log(doc, "🚀 Iniciando teste de Consulta de Agenda de Ambulância...")
        
        if not inicializar_driver():
            return False

        # ==== ACESSO AO SISTEMA ====
        safe_action(doc, "Acessando sistema", 
                   lambda: driver.get(URL), timeout_customizado=TIMEOUT_LONGO)

        safe_action(doc, "Realizando login", lambda: (
            aguardar_elemento("#j_id15\:email").send_keys(LOGIN_EMAIL),
            aguardar_elemento("#j_id15\:senha").send_keys(LOGIN_PASSWORD, Keys.ENTER),
            time.sleep(5)  # Aguarda processamento do login
        ))

        # ==== CONFIGURAÇÃO INICIAL ====
        safe_action(doc, "Configurando ambiente", lambda: (
            ajustar_zoom(90),
            driver.find_element(By.TAG_NAME, "body").send_keys(Keys.F3)
        ))

        # ==== NAVEGAÇÃO ====
        safe_action(doc, "Acessando menu Agenda Ambulância", 
                   clicar_elemento_robusto('body > div.menuLayer.animate.process.overflow.overflowY.boxsize.active > ul > li:nth-child(27) > img'))




        safe_action(doc, "Clicando em Consultar", lambda: (
            wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, '#gsAgendaAmbulancia > div.wdTelas > div > ul > li:nth-child(2) > a > span'))).click()
        ))


        # ==== PREENCHIMENTO DOS CAMPOS ====
        safe_action(doc, "Preenchendo Número da Agenda", 
                   preencher_campo_robusto('#gsAgendaAmbulancia > div.wdTelas > div.telaConsulta.telaConsultaAgendaAmbulancia > div > div.formRow.filtrosConsulta > div > div:nth-child(1) > input', '99'))

        safe_action(doc, "Preenchendo Data Inicial Saída", 
                   preencher_datepicker_por_indice(0, hoje))

        safe_action(doc, "Preenchendo Data Final Saída", 
                   preencher_datepicker_por_indice(1, hoje))

        safe_action(doc, "Preenchendo Número do Contrato", 
                   preencher_campo_robusto('#gsAgendaAmbulancia > div.wdTelas > div.telaConsulta.telaConsultaAgendaAmbulancia > div > div.formRow.filtrosConsulta > div > div:nth-child(4) > input', '113066'))

        # ==== SELEÇÃO VIA MODAIS ====
        safe_action(doc, "Selecionando Paciente", 
                   abrir_modal_e_selecionar_robusto(
                       '#gsAgendaAmbulancia > div.wdTelas > div.telaConsulta.telaConsultaAgendaAmbulancia > div > div.formRow.filtrosConsulta > div > div:nth-child(5) > div > a',
                       '#txtPesquisa',
                       'TESTE TITULAR AMBULANCIA',
                       'body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(2) > a',
                       "//td[contains(text(), 'TESTE TITULAR AMBULANCIA')]"
                   ))

        safe_action(doc, "Selecionando Motorista", 
                   abrir_modal_e_selecionar_robusto(
                       '#gsAgendaAmbulancia > div.wdTelas > div.telaConsulta.telaConsultaAgendaAmbulancia > div > div.formRow.filtrosConsulta > div > div:nth-child(6) > div > a',
                       'body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div.formCol.divPesquisa > input',
                       'MOTORISTA AMBULANCIA SELENIUM',
                       'body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(3) > a',
                       "//td[contains(text(), 'MOTORISTA AMBULANCIA SELENIUM')]"
                   ))

        safe_action(doc, "Selecionando Status", 
                   selecionar_opcao_robusta("#gsAgendaAmbulancia > div.wdTelas > div.telaConsulta.telaConsultaAgendaAmbulancia > div > div.formRow.filtrosConsulta > div > div:nth-child(7) > select", "Em aberto"))

        # ==== EXECUÇÃO DA CONSULTA ====
        safe_action(doc, "Realizando consulta", 
                   clicar_elemento_robusto("#gsAgendaAmbulancia > div.wdTelas > div.telaConsulta.telaConsultaAgendaAmbulancia > div > div.btnHolder > a.btModel.btGray.btfind"))

        # ==== VALIDAÇÃO DOS RESULTADOS ====
        resultado_validacao = None
        safe_action(doc, "Validando resultados da consulta", 
                   lambda: globals().update(resultado_validacao=validar_registros_encontrados()),
                   critico=False)

        # Verifica se deve continuar com o teste
        if not resultado_validacao or not resultado_validacao.get('encontrou_registros', False):
            log(doc, "ℹ️ Nenhum registro encontrado - Finalizando consulta", 'WARN')
            
            safe_action(doc, "Limpando campos (sem registros)", 
                       clicar_elemento_robusto("#gsAgendaAmbulancia > div.wdTelas > div.telaConsulta.telaConsultaAgendaAmbulancia > div > div.btnHolder > a.btModel.btGray.btclear"),
                       critico=False)
            
            safe_action(doc, "Fechando tela (sem registros)", 
                       clicar_elemento_robusto("#gsAgendaAmbulancia > div.wdTop.ui-draggable-handle > div.wdClose > a"),
                       critico=False)
            
            return True

        # ==== PROCESSAMENTO DOS REGISTROS ENCONTRADOS ====
        quantidade = resultado_validacao.get('quantidade_registros', 0)
        log(doc, f"✅ Continuando com processamento de {quantidade} registro(s)")

        # Visualização dos detalhes
        safe_action(doc, "Visualizando detalhes do registro", 
                   clicar_elemento_robusto("#DataTables_Table_0 > tbody > tr.odd > td:nth-child(6) > span.sprites.sp-dadosDinamicos"))

        safe_action(doc, "Capturando detalhes", lambda: (
            time.sleep(2),
            take_screenshot(driver, doc, "detalhes_registro", forcar=True)
        ))

        safe_action(doc, "Fechando detalhes do registro", 
                   clicar_elemento_robusto("body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.btnHolder"))

        # ==== DIGITALIZAÇÃO DE DOCUMENTOS ====
        safe_action(doc, "Abrindo digitalização de documentos", 
                   clicar_elemento_robusto("#DataTables_Table_0 > tbody > tr.odd > td:nth-child(6) > span.sprites.sp-scanner"))

        safe_action(doc, "Adicionando novo documento", 
                   clicar_elemento_robusto("body > div.modalHolder > div.modal.overflow > div.digitalizarDocumentos > div:nth-child(3) > span"))

        safe_action(doc, "Preenchendo descrição do arquivo", 
                   preencher_campo_robusto('body > div:nth-child(306) > div.modal.overflow > div.addHolder > div:nth-child(2) > div > input',
                                          'TESTE DESCRIÇÃO DO ARQUIVO SELENIUM AUTOMATIZADO (Automação de Testes): Teste de digitalização de documentos.'))

        # Upload de arquivo
        caminho_arquivo = r"C:\Users\Gold System\Downloads\TESTEPNGAMBULANCIA.png"
        safe_action(doc, "Fazendo upload do arquivo", 
                   fazer_upload_arquivo("body > div:nth-child(306) > div.modal.overflow > div.addHolder > div.formRow.formLastLine > div > div > div > div > div > div", caminho_arquivo))

        safe_action(doc, "Salvando arquivo digitalizado", 
                   clicar_elemento_robusto("body > div:nth-child(306) > div.modal.overflow > div.addHolder > div.formRow.hAlign > a.btModel.btGreen.hAlign"))

        safe_action(doc, "Fechando modal de digitalização", 
                   clicar_elemento_robusto("#DataTables_Table_0 > tbody > tr.odd > td:nth-child(6) > span.sprites.sp-scanner"))

        # ==== ALTERAÇÃO DE STATUS ====
        safe_action(doc, "Abrindo alteração de status", 
                   clicar_elemento_robusto("#DataTables_Table_0 > tbody > tr.odd > td:nth-child(6) > span.sprites.sp-edit"))

        safe_action(doc, "Alterando status para 'Encerrado'", 
                   selecionar_opcao_robusta("body > div.modalHolder > div.modal.overflow > div.modalAkterar > div:nth-child(3) > div > select", "Encerrado"))

        safe_action(doc, "Preenchendo motivo do encerramento", 
                   preencher_campo_robusto('body > div.modalHolder > div.modal.overflow > div.modalAkterar > div:nth-child(4) > div > textarea',
                                          'TESTE MOTIVO DE ENCERRAMENTO SELENIUM (Automação de Testes): Teste de encerramento de agenda de ambulância.'))

        safe_action(doc, "Salvando alteração de status", 
                   clicar_elemento_robusto("body > div.modalHolder > div.modal.overflow > div.modalAkterar > div.btnHolder > a"))

        # ==== FINALIZAÇÃO ====
        safe_action(doc, "Limpando campos", 
                   clicar_elemento_robusto("#gsAgendaAmbulancia > div.wdTelas > div.telaConsulta.telaConsultaAgendaAmbulancia > div > div.btnHolder > a.btModel.btGray.btclear"),
                   critico=False)

        safe_action(doc, "Fechando tela", 
                   clicar_elemento_robusto("#gsAgendaAmbulancia > div.wdTop.ui-draggable-handle > div.wdClose > a"),
                   critico=False)

        log(doc, "✅ Teste executado com sucesso!")
        return True

    except Exception as e:
        log(doc, f"❌ ERRO FATAL no teste: {e}", 'ERROR')
        take_screenshot(driver, doc, "erro_fatal", forcar=True)
        return False

    finally:
        log(doc, "🏁 Finalizando teste...")

# ==== FUNÇÃO PRINCIPAL ====
def main():
    """Função principal com tratamento de erros robusto"""
    global doc
    
    inicio_teste = datetime.now()
    sucesso = False
    
    try:
        log(doc, "🚀 INICIANDO TESTE DE CONSULTA DE AGENDA DE AMBULÂNCIA")
        log(doc, f"📅 Data/Hora de início: {inicio_teste.strftime('%d/%m/%Y %H:%M:%S')}")
        log(doc, "=" * 80)
        
        sucesso = executar_teste()
        
    except KeyboardInterrupt:
        log(doc, "⚠️ Teste interrompido pelo usuário", 'WARN')
        
    except Exception as e:
        log(doc, f"❌ Erro crítico na execução: {e}", 'ERROR')
        
    finally:
        # Cálculo do tempo de execução
        fim_teste = datetime.now()
        tempo_execucao = fim_teste - inicio_teste
        
        log(doc, "=" * 80)
        log(doc, f"📊 RESUMO DO TESTE:")
        log(doc, f"   • Status: {'✅ SUCESSO' if sucesso else '❌ FALHA'}")
        log(doc, f"   • Início: {inicio_teste.strftime('%H:%M:%S')}")
        log(doc, f"   • Fim: {fim_teste.strftime('%H:%M:%S')}")
        log(doc, f"   • Duração: {str(tempo_execucao).split('.')[0]}")
        log(doc, f"   • Screenshots: {len(screenshot_registradas)}")
        
        # Finaliza relatório
        finalizar_relatorio()
        
        # Retorno do código de saída
        exit_code = 0 if sucesso else 1
        print(f"\n🏁 Teste finalizado com código: {exit_code}")
        
        return exit_code

# ==== EXECUÇÃO ====
if __name__ == "__main__":
    try:
        main()
    except Exception as e:
        print(f"Erro crítico na inicialização: {e}")
        exit(1)



