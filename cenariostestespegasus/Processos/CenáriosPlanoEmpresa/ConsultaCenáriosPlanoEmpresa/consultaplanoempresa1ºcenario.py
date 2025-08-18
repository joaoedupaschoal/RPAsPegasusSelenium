# ==== AUTOMAÇÃO SELENIUM - FECHAMENTO PLANO EMPRESA ====
# Versão melhorada com melhor organização, tratamento de erros e logging

import os
import re
import time
import random
import logging
import traceback
import subprocess
from datetime import datetime, timedelta
from datetime import time as dt_time
from typing import Optional, List, Dict, Any, Tuple
from dataclasses import dataclass

# Selenium imports
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support.ui import Select, WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.common.action_chains import ActionChains
from selenium.webdriver.remote.webelement import WebElement
from selenium.common.exceptions import (
    TimeoutException, NoSuchElementException, StaleElementReferenceException,
    ElementClickInterceptedException, WebDriverException
)

# External libraries
from webdriver_manager.chrome import ChromeDriverManager
from docx import Document
from docx.shared import Inches
from faker import Faker
from faker.providers import BaseProvider
from validate_docbr import CPF


# ==== CONFIGURAÇÕES E CONSTANTES ====
@dataclass
class Config:
    """Configurações centralizadas da aplicação"""
    URL: str = "http://localhost:8080/gs/index.xhtml"
    LOGIN_EMAIL: str = "joaoeduardo.gold@outlook.com"
    LOGIN_PASSWORD: str = "071999gs"
    
    # Timeouts
    TIMEOUT_DEFAULT: int = 30
    TIMEOUT_SHORT: int = 10
    TIMEOUT_LONG: int = 60
    
    # Diretórios
    SCREENSHOTS_DIR: str = "screenshots"
    REPORTS_DIR: str = "reports"
    
    # Configurações do Chrome
    CHROME_OPTIONS: List[str] = None
    
    def __post_init__(self):
        if self.CHROME_OPTIONS is None:
            self.CHROME_OPTIONS = [
                "--start-maximized",
                "--disable-blink-features=AutomationControlled",
                "--disable-extensions",
                "--no-sandbox",
                "--disable-dev-shm-usage"
            ]


# ==== LOGGING CONFIGURADO ====
class TestLogger:
    """Sistema de logging melhorado"""
    
    def __init__(self, doc: Document):
        self.doc = doc
        self.logger = logging.getLogger('selenium_test')
        self.logger.setLevel(logging.INFO)
        
        # Formatter
        formatter = logging.Formatter(
            '%(asctime)s - %(levelname)s - %(message)s',
            datefmt='%Y-%m-%d %H:%M:%S'
        )
        
        # Console handler
        console_handler = logging.StreamHandler()
        console_handler.setFormatter(formatter)
        self.logger.addHandler(console_handler)
        
        # File handler
        os.makedirs('logs', exist_ok=True)
        file_handler = logging.FileHandler(
            f'logs/test_{datetime.now().strftime("%Y%m%d_%H%M%S")}.log'
        )
        file_handler.setFormatter(formatter)
        self.logger.addHandler(file_handler)
    
    def info(self, message: str):
        """Log de informação"""
        self.logger.info(message)
        self.doc.add_paragraph(f"ℹ️ {message}")
    
    def success(self, message: str):
        """Log de sucesso"""
        self.logger.info(f"SUCCESS: {message}")
        self.doc.add_paragraph(f"✅ {message}")
    
    def warning(self, message: str):
        """Log de aviso"""
        self.logger.warning(message)
        self.doc.add_paragraph(f"⚠️ {message}")
    
    def error(self, message: str):
        """Log de erro"""
        self.logger.error(message)
        self.doc.add_paragraph(f"❌ {message}")
    
    def debug(self, message: str):
        """Log de debug"""
        self.logger.debug(message)


# ==== PROVIDER CUSTOMIZADO ====
class BrasilProvider(BaseProvider):
    """Provider customizado para dados brasileiros"""
    
    def rg(self) -> str:
        """Gera RG brasileiro falso"""
        numeros = [str(random.randint(0, 9)) for _ in range(8)]
        return ''.join(numeros) + '-' + str(random.randint(0, 9))
    
    def cnpj(self) -> str:
        """Gera CNPJ brasileiro falso"""
        return f"{random.randint(10, 99)}.{random.randint(100, 999)}.{random.randint(100, 999)}/0001-{random.randint(10, 99)}"


# ==== GERADOR DE DADOS ====
class DataGenerator:
    """Gerador de dados para testes"""
    
    def __init__(self):
        self.fake = Faker("pt_BR")
        self.fake.add_provider(BrasilProvider)
    
    def gerar_datas_validas(self, hora_padrao: str = "00:00", dias_fim: int = 0) -> Tuple[str, ...]:
        """Gera datas coerentes para o teste"""
        hoje_date = datetime.today().date()
        dez_anos_atras = hoje_date - timedelta(days=3650)

        # Falecimento entre 10 anos atrás e hoje
        data_falecimento = self.fake.date_between(start_date=dez_anos_atras, end_date=hoje_date)

        # Nascimento (entre 18 e 110 anos antes do falecimento)
        idade_minima, idade_maxima = 18, 110
        data_nascimento = data_falecimento - timedelta(days=random.randint(idade_minima * 365, idade_maxima * 365))

        # Sepultamento 1..10 dias após o falecimento
        data_sepultamento = data_falecimento + timedelta(days=random.randint(1, 10))

        # Registro 1..10 dias após o sepultamento
        data_registro = data_sepultamento + timedelta(days=random.randint(1, 10))

        # Velório entre o falecimento e o sepultamento
        data_velorio = self.fake.date_between(start_date=data_falecimento, end_date=data_sepultamento)
        
        

        # Início entre 2 e 30 dias no futuro
        data_inicio_date = hoje_date + timedelta(days=random.randint(2, 30))
        h, m = map(int, hora_padrao.split(":"))
        dt_inicio = datetime.combine(data_inicio_date, dt_time(h, m))
        dt_fim = dt_inicio + timedelta(days=dias_fim)

        fmt_data = "%d/%m/%Y"
        fmt_dt = "%d/%m/%Y %H:%M"

        return (
            data_nascimento.strftime(fmt_data),
            data_falecimento.strftime(fmt_data),
            data_sepultamento.strftime(fmt_data),
            data_velorio.strftime(fmt_data),
            dt_inicio.strftime(fmt_dt),
            dt_fim.strftime(fmt_dt),
            data_registro.strftime(fmt_data),
            hoje_date.strftime(fmt_data),
        )


# ==== UTILITÁRIOS ====
class Utils:
    """Utilitários diversos"""
    
    @staticmethod
    def sanitize_filename(name: str, max_length: int = 120) -> str:
        """Sanitiza nome de arquivo"""
        name = name.strip().lower()
        name = re.sub(r"[<>:\"/\\|?*']", "_", name)
        name = re.sub(r"_+", "_", name)
        return name[:max_length]
    
    @staticmethod
    def sanitize_timeout(timeout: Any) -> int:
        """Sanitiza valor de timeout"""
        if not isinstance(timeout, (int, float)) or timeout <= 0:
            return 10
        return int(timeout)


# ==== SCREENSHOT MANAGER ====
class ScreenshotManager:
    """Gerenciador de screenshots"""
    
    def __init__(self, config: Config, logger: TestLogger):
        self.config = config
        self.logger = logger
        self.screenshots_taken = set()
        os.makedirs(config.SCREENSHOTS_DIR, exist_ok=True)
    
    def take_screenshot(self, driver: webdriver.Chrome, name: str) -> bool:
        """Tira screenshot se ainda não foi tirado"""
        if driver is None:
            self.logger.warning("Driver não disponível para screenshot")
            return False
        
        name = Utils.sanitize_filename(name)
        if name in self.screenshots_taken:
            return True
        
        try:
            path = os.path.join(self.config.SCREENSHOTS_DIR, f"{name}.png")
            driver.save_screenshot(path)
            
            self.logger.doc.add_paragraph(f"📸 Screenshot: {name}")
            self.logger.doc.add_picture(path, width=Inches(5.5))
            self.screenshots_taken.add(name)
            
            self.logger.debug(f"Screenshot salva: {path}")
            return True
            
        except Exception as e:
            self.logger.error(f"Erro ao tirar screenshot {name}: {e}")
            return False


# ==== ELEMENTO HANDLER ====
class ElementHandler:
    """Manipulador robusto de elementos"""
    
    def __init__(self, driver: webdriver.Chrome, config: Config, logger: TestLogger):
        self.driver = driver
        self.config = config
        self.logger = logger
        self.wait = WebDriverWait(driver, config.TIMEOUT_DEFAULT)
    
    def wait_for_element(self, locator: Tuple[By, str], timeout: Optional[int] = None, 
                        condition: str = 'clickable') -> WebElement:
        """Aguarda elemento com diferentes condições"""
        timeout = timeout or self.config.TIMEOUT_DEFAULT
        timeout = Utils.sanitize_timeout(timeout)
        
        conditions = {
            'present': EC.presence_of_element_located,
            'visible': EC.visibility_of_element_located,
            'clickable': EC.element_to_be_clickable,
            'invisible': EC.invisibility_of_element_located
        }
        
        condition_func = conditions.get(condition, EC.element_to_be_clickable)
        
        try:
            wait = WebDriverWait(self.driver, timeout)
            element = wait.until(condition_func(locator))
            return element
        except TimeoutException:
            raise TimeoutException(f"Elemento não encontrado: {locator} (condição: {condition}, timeout: {timeout}s)")
    
    def scroll_to_element(self, element: WebElement) -> bool:
        """Faz scroll seguro até elemento"""
        if not element or not element.is_displayed():
            return False
        
        strategies = [
            # Estratégia 1: JavaScript moderno
            lambda: self.driver.execute_script("""
                arguments[0].scrollIntoView({
                    behavior: 'smooth',
                    block: 'center',
                    inline: 'center'
                });
            """, element),
            
            # Estratégia 2: ActionChains
            lambda: ActionChains(self.driver).move_to_element(element).perform(),
            
            # Estratégia 3: JavaScript simples
            lambda: self.driver.execute_script("arguments[0].scrollIntoView();", element)
        ]
        
        for i, strategy in enumerate(strategies, 1):
            try:
                strategy()
                time.sleep(0.5)
                if element.is_displayed():
                    self.logger.debug(f"Scroll realizado com estratégia {i}")
                    return True
            except Exception as e:
                self.logger.debug(f"Estratégia {i} de scroll falhou: {e}")
                continue
        
        return False
    
    def remove_overlays(self):
        """Remove overlays que podem bloquear interações"""
        try:
            self.driver.execute_script("""
                // Remove overlays comuns
                const overlaySelectors = [
                    '.modal-backdrop', '.overlay', '.blockUI', '.loading', '.spinner',
                    '[style*="position: fixed"]', '.toast', '.tooltip'
                ];
                
                overlaySelectors.forEach(selector => {
                    document.querySelectorAll(selector).forEach(el => {
                        if (getComputedStyle(el).position === 'fixed') {
                            el.style.display = 'none';
                        }
                    });
                });
                
                // Remove elementos com z-index muito alto
                document.querySelectorAll('*').forEach(el => {
                    const zIndex = parseInt(getComputedStyle(el).zIndex);
                    if (zIndex > 1000) {
                        const rect = el.getBoundingClientRect();
                        if (rect.width === window.innerWidth && rect.height === window.innerHeight) {
                            el.style.display = 'none';
                        }
                    }
                });
            """)
        except Exception as e:
            self.logger.debug(f"Erro ao remover overlays: {e}")
    
    def robust_click(self, locator: Tuple[By, str], timeout: Optional[int] = None) -> bool:
        """Clique robusto com múltiplas estratégias"""
        timeout = timeout or self.config.TIMEOUT_DEFAULT
        
        try:
            # Aguarda elemento
            element = self.wait_for_element(locator, timeout, 'present')
            
            # Remove overlays
            self.remove_overlays()
            
            # Scroll até elemento
            self.scroll_to_element(element)
            
            # Tenta tornar clicável
            try:
                element = self.wait_for_element(locator, 5, 'clickable')
            except TimeoutException:
                self.logger.warning(f"Elemento não ficou clicável: {locator}")
            
            # Múltiplas estratégias de clique
            click_strategies = [
                lambda: element.click(),
                lambda: ActionChains(self.driver).move_to_element(element).click().perform(),
                lambda: self.driver.execute_script("arguments[0].click();", element),
                lambda: self.driver.execute_script("""
                    const el = arguments[0];
                    el.focus();
                    el.dispatchEvent(new MouseEvent('click', {
                        bubbles: true, 
                        cancelable: true, 
                        view: window
                    }));
                """, element)
            ]
            
            for i, strategy in enumerate(click_strategies, 1):
                try:
                    strategy()
                    time.sleep(0.5)
                    self.logger.debug(f"Clique realizado com estratégia {i}")
                    return True
                except Exception as e:
                    if i == len(click_strategies):
                        raise e
                    self.logger.debug(f"Estratégia {i} de clique falhou: {e}")
                    continue
            
            return False
            
        except Exception as e:
            self.logger.error(f"Falha ao clicar em {locator}: {e}")
            return False
    
    def fill_field(self, locator: Tuple[By, str], value: str, timeout: Optional[int] = None, 
                   clear_first: bool = True) -> bool:
        """Preenche campo com estratégias múltiplas"""
        timeout = timeout or self.config.TIMEOUT_DEFAULT
        max_attempts = 3
        
        for attempt in range(max_attempts):
            try:
                element = self.wait_for_element(locator, timeout)
                self.scroll_to_element(element)
                time.sleep(0.3)
                
                if attempt == 0:
                    # Estratégia 1: Método tradicional
                    if clear_first:
                        element.clear()
                    element.send_keys(value)
                    element.send_keys(Keys.TAB)
                
                elif attempt == 1:
                    # Estratégia 2: ActionChains
                    actions = ActionChains(self.driver)
                    actions.move_to_element(element).click()
                    if clear_first:
                        actions.key_down(Keys.CONTROL).send_keys('a').key_up(Keys.CONTROL)
                    actions.send_keys(value).send_keys(Keys.TAB).perform()
                
                else:
                    # Estratégia 3: JavaScript
                    self.driver.execute_script("""
                        var el = arguments[0], v = arguments[1];
                        el.focus();
                        el.value = v;
                        el.dispatchEvent(new Event('input', { bubbles: true }));
                        el.dispatchEvent(new Event('change', { bubbles: true }));
                        el.blur();
                    """, element, value)
                
                time.sleep(0.5)
                
                # Verifica se foi preenchido
                current_value = element.get_attribute('value') or ''
                if current_value.strip():
                    self.logger.debug(f"Campo preenchido na tentativa {attempt + 1}")
                    return True
                    
            except Exception as e:
                self.logger.debug(f"Tentativa {attempt + 1} falhou: {e}")
                if attempt == max_attempts - 1:
                    self.logger.error(f"Falha ao preencher campo {locator}: {e}")
                time.sleep(0.8)
        
        return False


# ==== DATEPICKER HANDLER ====
class DatepickerHandler:
    """Manipulador especializado para campos de data"""
    
    def __init__(self, driver: webdriver.Chrome, element_handler: ElementHandler, logger: TestLogger):
        self.driver = driver
        self.element_handler = element_handler
        self.logger = logger
    
    def find_datepicker_fields(self) -> List[Dict[str, Any]]:
        """Encontra todos os campos datepicker na página"""
        selectors = [
            "input.hasDatepicker",
            "input[id^='dp']",
            "input[maxlength='10'][grupo='']",
            "input[type='text'][maxlength='10']",
            "input[class*='datepicker']",
            ".hasDatepicker"
        ]
        
        fields = []
        
        for selector in selectors:
            try:
                elements = self.driver.find_elements(By.CSS_SELECTOR, selector)
                for element in elements:
                    if element.is_displayed() and element.is_enabled():
                        field_id = element.get_attribute('id') or f"dp_{len(fields)}"
                        
                        # Evita duplicatas
                        if not any(f['id'] == field_id for f in fields):
                            fields.append({
                                'element': element,
                                'id': field_id,
                                'selector': selector,
                                'maxlength': element.get_attribute('maxlength'),
                                'placeholder': element.get_attribute('placeholder')
                            })
            except Exception as e:
                self.logger.debug(f"Erro ao buscar campos com {selector}: {e}")
        
        self.logger.info(f"Encontrados {len(fields)} campos datepicker")
        return fields
    
    def validate_date_filled(self, element: WebElement, expected_date: str) -> bool:
        """Valida se a data foi preenchida corretamente"""
        try:
            current_value = (element.get_attribute('value') or '').strip()
            if not current_value:
                return False
            
            if current_value == expected_date or expected_date in current_value:
                return True
            
            # Tenta comparar em diferentes formatos
            date_formats = [
                '%d/%m/%Y %H:%M', '%d/%m/%Y %H:%M:%S',
                '%d/%m/%Y', '%m/%d/%Y', '%Y-%m-%d', '%d-%m-%Y'
            ]
            
            for fmt in date_formats:
                try:
                    d1 = datetime.strptime(current_value, fmt)
                    d2 = datetime.strptime(expected_date, fmt)
                    if d1 == d2:
                        return True
                except ValueError:
                    continue
            
            return False
            
        except Exception:
            return False
    
    def fill_datepicker_by_index(self, field_index: int, date_value: str, max_attempts: int = 5) -> bool:
        """Preenche datepicker pelo índice"""
        if not isinstance(field_index, int) or field_index < 0:
            raise ValueError(f"Índice inválido: {field_index}")
        
        if not date_value or not isinstance(date_value, str):
            raise ValueError(f"Data inválida: {date_value}")
        
        for attempt in range(max_attempts):
            try:
                fields = self.find_datepicker_fields()
                
                if not fields:
                    if attempt < max_attempts - 1:
                        self.logger.warning(f"Nenhum campo datepicker encontrado, tentativa {attempt + 1}")
                        time.sleep(2)
                        continue
                    raise Exception("Nenhum campo datepicker encontrado")
                
                if field_index >= len(fields):
                    raise Exception(f"Índice {field_index} inválido. Encontrados {len(fields)} campos")
                
                field_info = fields[field_index]
                element = field_info['element']
                field_id = field_info['id']
                
                self.logger.info(f"Preenchendo datepicker {field_index} (ID: {field_id}) com '{date_value}'")
                
                # Verifica se já está preenchido
                if self.validate_date_filled(element, date_value):
                    self.logger.success(f"Campo {field_index} já está preenchido corretamente!")
                    return True
                
                # Estratégias para datepicker
                strategies = [
                    self._jquery_strategy,
                    self._javascript_strategy,
                    self._actionchains_strategy,
                    self._traditional_strategy
                ]
                
                for i, strategy in enumerate(strategies, 1):
                    try:
                        self.logger.debug(f"Aplicando estratégia {i} para datepicker...")
                        strategy(element, field_id, date_value)
                        time.sleep(1)
                        
                        if self.validate_date_filled(element, date_value):
                            current_value = element.get_attribute('value')
                            self.logger.success(f"Datepicker preenchido com estratégia {i}: '{current_value}'")
                            return True
                        
                    except Exception as e:
                        self.logger.debug(f"Estratégia {i} falhou: {e}")
                        continue
                
                if attempt < max_attempts - 1:
                    self.logger.warning(f"Tentativa {attempt + 1} falhou, tentando novamente...")
                    time.sleep(2)
                
            except Exception as e:
                if attempt < max_attempts - 1:
                    self.logger.warning(f"Erro na tentativa {attempt + 1}: {e}")
                    time.sleep(2)
                else:
                    raise Exception(f"Falha ao preencher datepicker {field_index} após {max_attempts} tentativas") from e
        
        return False
    
    def _jquery_strategy(self, element: WebElement, field_id: str, date_value: str):
        """Estratégia usando jQuery"""
        result = self.driver.execute_script("""
            var fieldId = arguments[0], value = arguments[1];
            if (typeof jQuery === 'undefined') return 'jQuery não disponível';
            var $field = $('#' + fieldId);
            if (!$field.length) return 'Campo não encontrado: ' + fieldId;
            try {
                if ($field.hasClass('hasDatepicker')) { 
                    $field.datepicker('setDate', value); 
                } else { 
                    $field.val(value); 
                }
                $field.trigger('input').trigger('change').trigger('blur');
                return $field.val();
            } catch(e) { 
                return 'Erro: ' + e.message; 
            }
        """, field_id, date_value)
        
        if isinstance(result, str) and ('Erro' in result or 'não disponível' in result):
            raise Exception(f"jQuery falhou: {result}")
    
    def _javascript_strategy(self, element: WebElement, field_id: str, date_value: str):
        """Estratégia usando JavaScript puro"""
        self.driver.execute_script("""
            var field = arguments[0], value = arguments[1];
            field.focus(); 
            field.value = ''; 
            field.value = value;
            ['input','change','blur','keyup'].forEach(ev => 
                field.dispatchEvent(new Event(ev, {bubbles: true}))
            );
        """, element, date_value)
    
    def _actionchains_strategy(self, element: WebElement, field_id: str, date_value: str):
        """Estratégia usando ActionChains"""
        self.element_handler.scroll_to_element(element)
        time.sleep(0.5)
        
        actions = ActionChains(self.driver)
        actions.move_to_element(element).click()
        actions.pause(0.5)
        actions.key_down(Keys.CONTROL).send_keys('a').key_up(Keys.CONTROL)
        actions.pause(0.3)
        actions.send_keys(Keys.DELETE)
        actions.pause(0.3)
        
        for char in date_value:
            actions.send_keys(char).pause(0.05)
        
        actions.send_keys(Keys.TAB).perform()
    
    def _traditional_strategy(self, element: WebElement, field_id: str, date_value: str):
        """Estratégia tradicional"""
        self.element_handler.scroll_to_element(element)
        time.sleep(0.5)
        element.click()
        time.sleep(0.5)
        element.clear()
        element.send_keys(date_value)
        element.send_keys(Keys.TAB)


# ==== MODAL HANDLER ====
class ModalHandler:
    """Manipulador de modais"""
    
    def __init__(self, element_handler: ElementHandler, logger: TestLogger):
        self.element_handler = element_handler
        self.logger = logger
    
    def open_modal_and_select(self, btn_selector: str, search_selector: str, 
                             search_term: str, search_btn_selector: str, 
                             result_xpath: str) -> bool:
        """Abre modal e seleciona item"""
        try:
            # Abre modal
            self.element_handler.robust_click((By.CSS_SELECTOR, btn_selector))
            time.sleep(1)

            # Preenche pesquisa
            self.element_handler.fill_field((By.CSS_SELECTOR, search_selector), search_term)
            time.sleep(0.5)

            # Clica pesquisar
            self.element_handler.robust_click((By.CSS_SELECTOR, search_btn_selector))
            time.sleep(2)

            # Seleciona resultado
            self.element_handler.robust_click((By.XPATH, result_xpath))
            time.sleep(1)
            
            return True
            
        except Exception as e:
            self.logger.error(f"Erro ao manipular modal: {e}")
            return False


# ==== DRIVER MANAGER ====
class DriverManager:
    """Gerenciador do WebDriver"""
    
    def __init__(self, config: Config, logger: TestLogger):
        self.config = config
        self.logger = logger
        self.driver: Optional[webdriver.Chrome] = None
    
    def initialize_driver(self) -> bool:
        """Inicializa o driver do Chrome"""
        try:
            options = Options()
            
            for option in self.config.CHROME_OPTIONS:
                options.add_argument(option)
            
            options.add_experimental_option("excludeSwitches", ["enable-automation"])
            options.add_experimental_option('useAutomationExtension', False)

            service = Service(ChromeDriverManager().install())
            self.driver = webdriver.Chrome(service=service, options=options)
            
            # Remove indicadores de automação
            self.driver.execute_script(
                "Object.defineProperty(navigator, 'webdriver', {get: () => undefined})"
            )
            
            self.logger.success("Driver inicializado com sucesso")
            return True
            
        except Exception as e:
            self.logger.error(f"Erro ao inicializar driver: {e}")
            return False
    
    def quit_driver(self):
        """Encerra o driver"""
        if self.driver:
            try:
                self.driver.quit()
                self.logger.info("Driver encerrado")
            except Exception as e:
                self.logger.warning(f"Erro ao encerrar driver: {e}")


# ==== AÇÃO SEGURA ====
class SafeActionExecutor:
    """Executor de ações seguras com logging e screenshots"""
    
    def __init__(self, screenshot_manager: ScreenshotManager, logger: TestLogger):
        self.screenshot_manager = screenshot_manager
        self.logger = logger
    
    def execute(self, driver: webdriver.Chrome, description: str, action_func) -> bool:
        """Executa ação com tratamento de erro e screenshot"""
        try:
            self.logger.info(f"🔄 {description}...")
            action_func()
            self.logger.success(f"{description} realizada com sucesso")
            self.screenshot_manager.take_screenshot(
                driver, 
                description.lower().replace(" ", "_")
            )
            return True
            
        except Exception as e:
            self.screenshot_manager.take_screenshot(
                driver, 
                f"erro_{description.lower().replace(' ', '_')}"
            )
            self.logger.error(f"{description} falhou: {type(e).__name__}: {e}")
            self.logger.debug(traceback.format_exc())
            return False


# ==== TESTE PRINCIPAL ====
class PlanoEmpresaTest:
    """Classe principal do teste de fechamento de Plano Empresa"""
    
    def __init__(self):
        self.config = Config()
        self.setup_directories()
        self.setup_document()
        self.setup_components()
        self.data_generator = DataGenerator()
    
    def setup_directories(self):
        """Cria diretórios necessários"""
        for directory in [self.config.SCREENSHOTS_DIR, self.config.REPORTS_DIR, 'logs']:
            os.makedirs(directory, exist_ok=True)
    
    def setup_document(self):
        """Configura documento do relatório"""
        self.doc = Document()
        self.doc.add_heading("RELATÓRIO DO TESTE", 0)
        self.doc.add_paragraph(
            "Processo: Consulta Plano Empresa – Cenário 1: "
            "Nesse teste, o usuário irá realizar a Consulta de um Plano Empresa."
        )
        self.doc.add_paragraph(f"Data do teste: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")
    
    def setup_components(self):
        """Configura componentes do teste"""
        self.logger = TestLogger(self.doc)
        self.screenshot_manager = ScreenshotManager(self.config, self.logger)
        self.driver_manager = DriverManager(self.config, self.logger)
        self.safe_action = SafeActionExecutor(self.screenshot_manager, self.logger)
    
    def initialize_handlers(self):
        """Inicializa handlers após driver estar disponível"""
        if not self.driver_manager.driver:
            raise Exception("Driver não foi inicializado")
        
        self.element_handler = ElementHandler(
            self.driver_manager.driver, self.config, self.logger
        )
        self.datepicker_handler = DatepickerHandler(
            self.driver_manager.driver, self.element_handler, self.logger
        )
        self.modal_handler = ModalHandler(self.element_handler, self.logger)
    
    def adjust_zoom(self):
        """Ajusta zoom da página"""
        try:
            self.driver_manager.driver.execute_script("document.body.style.zoom='90%'")
            self.logger.info("Zoom ajustado para 90%")
        except Exception as e:
            self.logger.warning(f"Erro ao ajustar zoom: {e}")
    
    def find_alert_messages(self) -> Optional[WebElement]:
        """Procura por mensagens de alerta na página"""
        selectors = [
            (".alerts.salvo", "✅ Sucesso"),
            (".alerts.alerta", "⚠️ Alerta"),
            (".alerts.erro", "❌ Erro"),
        ]
        
        for selector, msg_type in selectors:
            try:
                element = self.driver_manager.driver.find_element(By.CSS_SELECTOR, selector)
                if element.is_displayed():
                    self.logger.info(f"📢 {msg_type}: {element.text}")
                    return element
            except:
                continue
        
        self.logger.info("ℹ️ Nenhuma mensagem de alerta encontrada")
        return None
    
    def login(self) -> bool:
        """Realiza login no sistema"""
        def login_action():
            # Preenche email
            self.element_handler.fill_field(
                (By.ID, "j_id15:email"), 
                self.config.LOGIN_EMAIL
            )
            
            # Preenche senha e submete
            password_field = self.element_handler.wait_for_element((By.ID, "j_id15:senha"))
            password_field.send_keys(self.config.LOGIN_PASSWORD)
            password_field.send_keys(Keys.ENTER)
            
            time.sleep(5)
        
        return self.safe_action.execute(
            self.driver_manager.driver,
            "Realizando login",
            login_action
        )
    
    def access_plano_empresa_menu(self) -> bool:
        """Acessa o menu Plano Empresa"""
        def menu_action():
            # Abre menu principal
            self.driver_manager.driver.find_element(By.TAG_NAME, "body").send_keys(Keys.F3)
            time.sleep(1)
            
            # Clica em Plano Empresa
            menu_item = self.element_handler.wait_for_element(
                (By.XPATH, '/html/body/div[15]/ul/li[21]/img')
            )
            self.element_handler.scroll_to_element(menu_item)
            menu_item.click()
        
        return self.safe_action.execute(
            self.driver_manager.driver,
            "Acessando menu Plano Empresa",
            menu_action
        )
    
    def access_consulta(self) -> bool:
        """Acessa a funcionalidade de Consulta"""
        def consulta_action():
            self.element_handler.robust_click((
                By.CSS_SELECTOR,
                '#gsPlanoEmpresa > div.wdTelas > div > ul > li:nth-child(2) > a > span'
            ))
            time.sleep(2)
        
        return self.safe_action.execute(
            self.driver_manager.driver,
            "Clicando em Consulta",
            consulta_action
        )
    

    def clique_abas_e_print_contratos(self) -> bool: 
        """Clica na aba e tira print"""
        def click_and_print_action_contratos():
            # Clica na aba "Contratos"
            self.element_handler.robust_click((
                By.XPATH,
                "//*[@id='gsPlanoEmpresa']/div[2]/div[2]/div[2]/ul/li[1]/a"
            ))
            time.sleep(1)
            

        
        return self.safe_action.execute(
            self.driver_manager.driver,
            "Clicando na aba Contratos e tirando print",
            click_and_print_action_contratos
        )

    def clique_abas_e_print_titulos(self) -> bool: 
        """Clica na aba e tira print"""
        def click_and_print_action_titulos():
            # Clica na aba "Títulos"
            self.element_handler.robust_click((
                By.XPATH,
                "//*[@id='gsPlanoEmpresa']/div[2]/div[2]/div[2]/ul/li[2]/a"
            ))
            time.sleep(1)
            

        return self.safe_action.execute(
            self.driver_manager.driver,
            "Clicando na aba Títulos e tirando print",
            click_and_print_action_titulos
        )



    def select_plano_empresa(self) -> bool:
        """Seleciona o Plano Empresa"""
        def selection_action():
            # Abre LOV
            self.element_handler.robust_click((
                By.XPATH,
                "//*[@id='gsPlanoEmpresa']/div[2]/div[2]/div[1]/div/div[1]/div/a"
            ))
            time.sleep(1)
            
            # Seleciona tipo de filtro CNPJ
            tipo_filtro_select = Select(self.element_handler.wait_for_element((
                By.CSS_SELECTOR,
                "body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(1) > select"
            )))
            tipo_filtro_select.select_by_visible_text("CNPJ")
            
            # Preenche CNPJ
            self.element_handler.fill_field((
                By.XPATH,
                "//input[@type='text' and contains(@class,'nomePesquisa')]"
            ), "38.926.740/0001-26")
            
            # Clica pesquisar
            self.element_handler.robust_click((
                By.CSS_SELECTOR,
                'body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(4) > a'
            ))
            time.sleep(2)
            
            # Seleciona resultado
            self.element_handler.robust_click((
                By.XPATH,
                "//td[contains(text(), 'PLANO EMPRESA TESTE SELENIUM AUTOMATIZADO')]"
            ))
        
        return self.safe_action.execute(
            self.driver_manager.driver,
            "Selecionando Plano Empresa",
            selection_action
        )
    
    def search_plano_empresa(self) -> bool:
        """Busca o Plano Empresa selecionado"""
        def search_action():
            self.element_handler.robust_click((
                By.XPATH,
                "//a[contains(@class,'btModel') and contains(@class,'btGray') and contains(normalize-space(.), 'Pesquisar')]"
            ))
        
        return self.safe_action.execute(
            self.driver_manager.driver,
            "Pesquisando Plano Empresa",
            search_action
        )
    
    def fill_fechamento_data(self) -> bool:
        """Preenche dados do fechamento"""
        def fill_data_action():
            # Preenche Data de Vencimento
            success = self.datepicker_handler.fill_datepicker_by_index(0, "17/12/2025")
            if not success:
                raise Exception("Falha ao preencher data de vencimento")
            
            # Seleciona Tipo Valor Base
            tipo_valor_select = Select(self.element_handler.wait_for_element((
                By.XPATH,
                "//select[option[normalize-space(.)='Valor Contrato'] and option[normalize-space(.)='Valor Pacote']]"
            )))
            tipo_valor_select.select_by_visible_text("Valor Contrato")
            
            # Seleciona Tipo de reajuste
            tipo_reajuste_select = Select(self.element_handler.wait_for_element((
                By.XPATH,
                "//select[option[normalize-space(.)='Acréscimo'] and option[normalize-space(.)='Desconto']]"
            )))
            tipo_reajuste_select.select_by_visible_text("Desconto")
            
            # Preenche Taxa de Reajuste
            taxa_value = str(self.data_generator.fake.random_int(min=1, max=10000))
            self.element_handler.fill_field((
                By.XPATH,
                "//input[@ref='pct' and @placeholder='% ' and contains(@style,'width: 70px')]"
            ), taxa_value)
            time.sleep(0.5)
        
        return self.safe_action.execute(
            self.driver_manager.driver,
            "Preenchendo dados do fechamento",
            fill_data_action
        )
    


    def select_tipo_mensalidade(self) -> bool:
        """Seleciona Tipo Mensalidade"""
        def select_action():
            success = self.modal_handler.open_modal_and_select(
                '#gsPlanoEmpresa > div.wdTelas > div.telaConsulta > div.contentHolder.clearfix.overflow.overflowY > div:nth-child(2) > div:nth-child(5) > div > a',
                'body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(2) > input',
                'TESTE TIPO DE MENSALIDADE SELENIUM AUTOMATIZADO',
                'body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(4) > a',
                "//td[contains(text(), 'TESTE TIPO DE MENSALIDADE SELENIUM AUTOMATIZADO')]"
            )
            if not success:
                raise Exception("Falha ao selecionar tipo mensalidade")
        
        return self.safe_action.execute(
            self.driver_manager.driver,
            "Selecionando Tipo Mensalidade",
            select_action
        )
    
    def execute_fechamento(self) -> bool:
        """Executa o fechamento do Plano Empresa"""
        def fechamento_action():
            self.element_handler.robust_click((
                By.CSS_SELECTOR,
                '#gsPlanoEmpresa > div.wdTelas > div.telaConsulta > div.btnHolder > a:nth-child(2)'
            ))
        
        return self.safe_action.execute(
            self.driver_manager.driver,
            "Realizando fechamento de Plano Empresa",
            fechamento_action
        )
    
    def confirm(self) -> bool:
        """Confirma o fechemento do Plano Empresa"""
        def close_action():
            self.element_handler.robust_click((
                By.CSS_SELECTOR,
                '#BtYes'
            ))
        
        return self.safe_action.execute(
            self.driver_manager.driver,
            "Confirmando fechamento de Plano Empresa",
            close_action
        )
    
    def close_modal(self) -> bool:
        """Fecha modal após fechamento"""
        def close_action():
            self.element_handler.robust_click((
                By.CSS_SELECTOR,
                '#gsPlanoEmpresa > div.wdTop.ui-draggable-handle > div.wdClose > a'
            ))
        
        return self.safe_action.execute(
            self.driver_manager.driver,
            "Fechando modal",
            close_action
        )


    def run_test(self) -> bool:
        """Executa o teste completo"""
        try:
            self.logger.info("🚀 Iniciando teste de fechamento de Plano Empresa")
            
            # Inicializa driver
            if not self.driver_manager.initialize_driver():
                return False
            
            # Inicializa handlers
            self.initialize_handlers()
            
            # Executa passos do teste
            steps = [
                ("Acessando sistema", lambda: self.driver_manager.driver.get(self.config.URL)),
                ("Login", self.login),
                ("Ajuste de zoom e menu", lambda: (self.adjust_zoom(), True)[1]),
                ("Acesso ao menu Plano Empresa", self.access_plano_empresa_menu),
                ("Acesso a Consulta", self.access_consulta),
                ("Seleção do Plano Empresa", self.select_plano_empresa),
                ("Busca do Plano Empresa", self.search_plano_empresa),
                ("Clicando na aba Contratos", self.clique_abas_e_print_contratos),
                ("Clicando na aba Títulos", self.clique_abas_e_print_titulos),
                ("Fechamento do modal", self.close_modal)
            ]
            
            success_count = 0
            total_steps = len(steps)
            
            for step_name, step_func in steps:
                if callable(step_func):
                    success = step_func()
                else:
                    success = self.safe_action.execute(
                        self.driver_manager.driver,
                        step_name,
                        step_func
                    )
                
                if success:
                    success_count += 1
                else:
                    self.logger.error(f"Falha na etapa: {step_name}")
            
            # Verifica mensagens de alerta
            self.logger.info("🔍 Verificando mensagens de alerta...")
            self.find_alert_messages()
            
            # Relatório final
            success_rate = (success_count / total_steps) * 100
            self.logger.info(f"📊 Taxa de sucesso: {success_rate:.1f}% ({success_count}/{total_steps} etapas)")
            
            if success_count == total_steps:
                self.logger.success("✅ Teste executado com sucesso!")
                return True
            else:
                self.logger.error("❌ Teste finalizado com erros.")
                return False
                
        except Exception as e:
            self.logger.error(f"❌ ERRO FATAL: {e}")
            self.logger.debug(traceback.format_exc())
            self.screenshot_manager.take_screenshot(
                self.driver_manager.driver, "erro_fatal"
            )
            return False
        
        finally:
            self.logger.info("✅ Teste concluído.")
    
    def finalize_report(self):
        """Finaliza e salva o relatório"""
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"relatorio_consulta_plano_empresa_cenario_1_{timestamp}.docx"
        filepath = os.path.join(self.config.REPORTS_DIR, filename)
        
        try:
            # Adiciona resumo final
            self.doc.add_heading("RESUMO FINAL", 1)
            self.doc.add_paragraph(f"Teste finalizado em: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")
            
            # Salva documento
            self.doc.save(filepath)
            self.logger.success(f"📄 Relatório salvo como: {filepath}")
            
            # Tenta abrir relatório
            try:
                if os.name == 'nt':  # Windows
                    subprocess.run(["start", "winword", filepath], shell=True)
                elif os.name == 'posix':  # Linux/Mac
                    subprocess.run(["xdg-open", filepath])
            except Exception as e:
                self.logger.debug(f"Não foi possível abrir relatório automaticamente: {e}")
                
        except Exception as e:
            self.logger.error(f"Erro ao salvar relatório: {e}")
        
        finally:
            # Encerra driver
            self.driver_manager.quit_driver()


# ==== FUNÇÃO PRINCIPAL ====
def main():
    """Função principal do programa"""
    test = PlanoEmpresaTest()
    
    try:
        success = test.run_test()
        return 0 if success else 1
        
    except KeyboardInterrupt:
        test.logger.warning("❌ Teste interrompido pelo usuário")
        return 1
        
    except Exception as e:
        test.logger.error(f"❌ Erro na execução principal: {e}")
        test.logger.debug(traceback.format_exc())
        return 1
        
    finally:
        test.finalize_report()


# ==== EXECUÇÃO ====
if __name__ == "__main__":
    import sys
    sys.exit(main())