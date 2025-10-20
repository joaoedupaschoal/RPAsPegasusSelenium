# ==== IMPORTS ====
from datetime import datetime, timedelta
from datetime import time as dt_time
import time
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver import ActionChains
from selenium.common.exceptions import *
from selenium.webdriver.support import expected_conditions as EC
from selenium import webdriver
from selenium.webdriver.support.ui import Select, WebDriverWait
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service
from webdriver_manager.chrome import ChromeDriverManager
from docx import Document
from docx.shared import Inches
from faker import Faker
from faker.providers import BaseProvider
from validate_docbr import CPF
import subprocess
import os
import random
import re
import pyautogui

# ==== CONFIGURA√á√ïES GLOBAIS ====
TIMEOUT_DEFAULT = 30
TIMEOUT_CURTO = 10
TIMEOUT_LONGO = 60
CAMINHO_ARQUIVO_UPLOAD = "C:/Users/Gold System/Documents/teste.png"
URL = "https://andromeda.erp-pegasus.com.br/gs/login.xhtml"
LOGIN_EMAIL = "joaoeduardo.gold@outlook.com"
LOGIN_PASSWORD = "071999gs"

# ==== VARI√ÅVEIS GLOBAIS ====
doc = Document()
doc.add_heading("RELAT√ìRIO DO TESTE", 0)
doc.add_paragraph("Gera√ß√£o de Boletos - Gestor Financeiro ‚Äì Cen√°rio 1: Rotina completa de Gera√ß√£o de Boletos")
doc.add_paragraph(f"Data do teste: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")

screenshot_registradas = set()
driver = None
wait = None

# ==== PROVIDERS CUSTOMIZADOS ====
class BrasilProvider(BaseProvider):
    def rg(self):
        numeros = [str(random.randint(0, 9)) for _ in range(8)]
        return ''.join(numeros) + '-' + str(random.randint(0, 9))

fake = Faker("pt_BR")
fake.add_provider(BrasilProvider)

# ==== GERA√á√ÉO DE DATAS ====
def gerar_datas_validas(hora_padrao="00:00", dias_fim=0):
    hoje_date = datetime.today().date()
    dez_anos_atras = hoje_date - timedelta(days=3650)
    data_falecimento = fake.date_between(start_date=dez_anos_atras, end_date=hoje_date)
    idade_minima, idade_maxima = 18, 110
    data_nascimento = data_falecimento - timedelta(days=random.randint(idade_minima * 365, idade_maxima * 365))
    data_sepultamento = data_falecimento + timedelta(days=random.randint(1, 10))
    data_registro = data_sepultamento + timedelta(days=random.randint(1, 10))
    data_velorio = fake.date_between(start_date=data_falecimento, end_date=data_sepultamento)
    data_inicio_date = hoje_date + timedelta(days=random.randint(2, 30))
    h, m = map(int, hora_padrao.split(":"))
    dt_inicio = datetime.combine(data_inicio_date, dt_time(h, m))
    dt_fim = dt_inicio + timedelta(days=dias_fim)
    fmt_data = "%d/%m/%Y"
    fmt_dt = "%d/%m/%Y %H:%M"
    return (
        data_nascimento.strftime(fmt_data),
        data_falecimento.strftime(fmt_data),
        data_sepultamento.strftime(fmt_data),
        data_velorio.strftime(fmt_data),
        dt_inicio.strftime(fmt_dt),
        dt_fim.strftime(fmt_dt),
        data_registro.strftime(fmt_data),
        hoje_date.strftime(fmt_data),
    )

(data_nascimento, data_falecimento, data_sepultamento,
 data_velorio, data_inicio, data_fim, data_registro, hoje) = gerar_datas_validas(
    hora_padrao="08:50",
    dias_fim=0
)

# ==== UTILIT√ÅRIOS DE LOG ====
def log(doc, msg):
    print(msg)
    doc.add_paragraph(msg)

def _sanitize_filename(name: str) -> str:
    name = name.strip().lower()
    name = re.sub(r"[<>:\"/\\|?*']", "_", name)
    name = re.sub(r"_+", "_", name)
    return name[:120]

def take_screenshot(driver, doc, nome):
    if driver is None:
        return
    nome = _sanitize_filename(nome)
    if nome not in screenshot_registradas:
        path = f"screenshots/{nome}.png"
        os.makedirs("screenshots", exist_ok=True)
        try:
            driver.save_screenshot(path)
            doc.add_paragraph(f"Screenshot: {nome}")
            doc.add_picture(path, width=Inches(5.5))
            screenshot_registradas.add(nome)
        except Exception as e:
            log(doc, f"‚ö†Ô∏è Erro ao tirar screenshot {nome}: {e}")


def clicar_cancelar_modal_selecao_conta(js_engine, idx_cancel=0, timeout=10):
    """
    Se existir o modal 'Selecione uma conta banc√°ria', clica no bot√£o 'Cancelar' pelo √≠ndice.
    idx_cancel: 0 = primeiro 'Cancelar' vis√≠vel dentro do modal.
    Retorna True se clicou; False se n√£o encontrou o modal no prazo.
    """
    import time
    driver = js_engine.driver
    fim = time.time() + timeout

    while time.time() < fim:
        try:
            # Marca os bot√µes 'Cancelar' do modal 'payment noClose' vis√≠vel com data-aim tempor√°rio
            qtd = js_engine.execute_js("""
                function vis(el){
                  const s = getComputedStyle(el);
                  return el.offsetParent!==null && s.display!=='none' &&
                         s.visibility!=='hidden' && parseFloat(s.opacity||1)>0.01;
                }
                // procura o container do modal 'payment noClose' vis√≠vel
                var containers = Array.from(document.querySelectorAll("div.modal.overflow > div.payment.noClose"));
                var alvo = containers.find(vis);
                if(!alvo) return 0;

                // dentro do modal, pegue todos os 'Cancelar'
                var btns = Array.from(alvo.parentElement.querySelectorAll("a.btModel.btGray.btcancel"));
                // filtra apenas bot√µes vis√≠veis
                btns = btns.filter(vis);
                // marque com data-aim para sele√ß√£o direta
                btns.forEach((b,i)=> b.setAttribute("data-aim-cancel", "cancel-"+i));
                return btns.length;
            """)

            if qtd and idx_cancel < qtd:
                seletor = f"[data-aim-cancel='cancel-{idx_cancel}']"
                js_engine.force_click(seletor, by_xpath=False)
                js_engine.wait_ajax_complete(5)

                # limpeza do atributo tempor√°rio
                js_engine.execute_js("""
                    document.querySelectorAll("[data-aim-cancel]")
                    .forEach(e=>e.removeAttribute("data-aim-cancel"));
                """)
                return True
        except Exception:
            pass

        time.sleep(0.25)

    return False

def confirmar_envio_e_verificar_alertas(js_engine, doc, indice=1, timeout_alerta=5):
    """
    Clica no bot√£o 'Sim' (confirmar envio de boletos por e-mail)
    e em seguida procura por mensagens ou modais de alerta vis√≠veis no sistema.

    Etapas:
      1) Clica no bot√£o 'Sim' por √≠ndice.
      2) Aguarda retorno √† tela principal.
      3) Verifica se h√° alertas ou mensagens vis√≠veis.
    """
    import time

    # 1) Clique no bot√£o 'Sim'
    safe_action(doc, f"Confirmando o envio de boletos por E-mail", lambda:
        js_engine.force_click(
            f"(//a[@class='btModel btGray btyes' and normalize-space()='Sim'])[{indice}]",
            by_xpath=True
        )
    )

    # 2) Pequeno delay para o sistema responder e retornar √† tela
    js_engine.wait_ajax_complete(8)
    time.sleep(1)

    # 3) Busca mensagens de alerta
    alerta = encontrar_mensagem_alerta()

    if alerta:
        log(doc, f"‚ö†Ô∏è Alerta detectado ap√≥s confirma√ß√£o: {alerta}")
    else:
        log(doc, "‚úÖ Nenhum alerta vis√≠vel ap√≥s confirma√ß√£o.")

    return alerta





def executar_fluxo_boletos(js_engine, doc,
                           indice_ok=5,
                           xpath_select_conta="//select[@class='contaBancaria' and @style='width: 360px;' and @rev='10']",
                           xpath_select_instrucao="//select[@class='instrucaoBoleto' and @style='width: 360px; padding-top: 10px;']",
                           idx_select_conta=1, idx_opcao_conta=1,
                           idx_select_instr=1, idx_opcao_instr=1,
                           screenshot_final=True):
    """
    Fluxo completo com tratamento do modal 'Existe(m) t√≠tulo(s) de Plano Empresa':
      1) Valida tabela; se houver linhas, abre/printa/fecha Detalhes.
      2) Clica em 'Boleto Pegasus'.
         2.1) Se surgir o modal .telaModalTitulosPlanoEmpresa:
              - Fecha SOMENTE esse modal (sem √≠ndice) e segue o fluxo.
      3) Seleciona 'Carn√™', Conta, Instru√ß√£o, clica Ok, confirma e retorna.
    Requisitos: safe_action, log, take_screenshot, validar_resultado_pesquisa,
                selecionar_opcao_por_indice, clicar_ok_e_verificar_modal_confirmacao,
                confirmar_modal_e_retornar_sistema, fechar_modal_com_retry.
    """
    from selenium.webdriver.common.by import By
    from selenium.webdriver.support.ui import WebDriverWait
    from selenium.webdriver.support import expected_conditions as EC
    import time

    driver = js_engine.driver
    wait = WebDriverWait(driver, 10)

    log(doc, "üöÄ Iniciando fluxo completo de gera√ß√£o de boletos...")

    # ===== 1) VALIDA√á√ÉO DO RESULTADO =====
    if not validar_resultado_pesquisa(js_engine):
        log(doc, "‚ö†Ô∏è Nenhum resultado encontrado ‚Äî encerrando o processo.")
        return
    log(doc, "‚úÖ Prosseguindo com o fluxo ‚Äî tabela carregada com sucesso.")

    # Conta linhas
    try:
        linhas = driver.find_elements(
            By.XPATH,
            "//table[(contains(@class,'niceTable padding10') or contains(@id,'tabela') or contains(@class,'dataTable'))]"
            "/tbody/tr[not(contains(@class,'empty')) and not(contains(@style,'display: none'))]"
        )
        qtd_linhas = len(linhas)
        log(doc, f"üìÑ Total de t√≠tulos encontrados: {qtd_linhas}")
    except Exception as e:
        log(doc, f"‚ö†Ô∏è Erro ao contar t√≠tulos: {e}")
        qtd_linhas = 0

    # Detalhes (opcional)
    if qtd_linhas > 0:
        log(doc, "üîç Abrindo modal de detalhes do primeiro t√≠tulo...")
        safe_action(doc, "Visualizando Detalhes", lambda: js_engine.force_click(
            "//i[@class='sprites sp-dadosDinamicos' and @title='Visualizar Detalhes']",
            by_xpath=True
        ))
        try:
            wait.until(EC.presence_of_element_located((
                By.XPATH,
                "//div[contains(@class,'modal') or contains(@class,'detalhes') or contains(@class,'overflow')]"
            )))
            time.sleep(1.2)
            take_screenshot(driver, doc, "Modal_Detalhes_Aberto")
            log(doc, "üì∏ Screenshot capturado com o modal aberto.")
        except Exception as e:
            log(doc, f"‚ö†Ô∏è N√£o foi poss√≠vel confirmar a abertura do modal de detalhes: {e}")
        safe_action(doc, "Fechando modal de Detalhes", lambda: fechar_modal_com_retry(doc, js_engine, wait))
    else:
        log(doc, "‚ö†Ô∏è Nenhum t√≠tulo encontrado ‚Äî prosseguindo sem abrir detalhes.")

    # ===== 2) BOLETO PEGASUS + TRATAMENTO DO MODAL 'PLANO EMPRESA' =====
    safe_action(doc, "Clicando em 'Boleto Pegasus'", lambda: js_engine.force_click(
        "//a[@class='btModel btShort btLight boletoPegasus' and normalize-space(text())='Boleto Pegasus']"
        "/i[@class='fa fa-barcode']/parent::a",
        by_xpath=True
    ))

    time.sleep(0.6)

    plano_empresa_modal = None
    try:
        plano_empresa_modal = WebDriverWait(driver, 2).until(
            EC.presence_of_element_located((
                By.XPATH,
                "//div[contains(@class,'modal') and contains(@class,'overflow')][.//div[contains(@class,'telaModalTitulosPlanoEmpresa')]]"
            ))
        )
    except Exception:
        plano_empresa_modal = None

    if plano_empresa_modal:
        log(doc, "üßæ Detectado modal de 'Plano Empresa' ap√≥s clicar em Boleto Pegasus. Fechando e seguindo o fluxo.")

        fechou = False
        try:
            # üîÑ Substitu√≠do seletor do bot√£o Fechar
            btn_fechar = plano_empresa_modal.find_element(
                By.XPATH,
                ".//a[@class='btModel btGray' and normalize-space(text())='Fechar']"
            )
            driver.execute_script("arguments[0].scrollIntoView({block:'center'});", btn_fechar)
            try:
                btn_fechar.click()
                fechou = True
            except Exception:
                driver.execute_script("arguments[0].click();", btn_fechar)
                fechou = True
        except Exception:
            pass

        # se n√£o conseguir, tenta o 'X' (fa fa-close)
        if not fechou:
            try:
                x_close = plano_empresa_modal.find_element(By.XPATH, ".//a[@class='fa fa-close']")
                driver.execute_script("arguments[0].scrollIntoView({block:'center'});", x_close)
                try:
                    x_close.click()
                except Exception:
                    driver.execute_script("arguments[0].click();", x_close)
            except Exception as e:
                log(doc, f"‚ùå N√£o foi poss√≠vel fechar o modal de Plano Empresa: {e}")

        try:
            WebDriverWait(driver, 5).until_not(EC.presence_of_element_located((
                By.XPATH,
                "//div[contains(@class,'telaModalTitulosPlanoEmpresa')]"
            )))
        except Exception:
            pass

        log(doc, "‚ÑπÔ∏è Modal de Plano Empresa fechado. Prosseguindo com a gera√ß√£o.")

    # ===== 3) FLUXO NORMAL DE GERA√á√ÉO DO BOLETO =====
    safe_action(doc, "Selecionando a op√ß√£o 'Carn√™'", lambda: js_engine.force_click(
        "//li[@tabindex='1' and @ref='carne' and @rel='undefined' and normalize-space(text())='Carne']",
        by_xpath=True
    ))

    safe_action(doc, "Selecionando Conta Banc√°ria",
        selecionar_opcao_por_indice(
            indice_select=idx_select_conta,
            indice_opcao=idx_opcao_conta,
            xpath_customizado=xpath_select_conta
        )
    )

    safe_action(doc, "Selecionando Instru√ß√£o Alternativa",
        selecionar_opcao_por_indice(
            indice_select=idx_select_instr,
            indice_opcao=idx_opcao_instr,
            xpath_customizado=xpath_select_instrucao
        )
    )

    safe_action(doc, "Clicando em 'Ok'",
        clicar_ok_e_verificar_modal_confirmacao(indice_ok=indice_ok)
    )

    time.sleep(10)



    resultado = confirmar_ou_detectar_relatorio(js_engine, doc=doc)

    if resultado:
        log(doc, "‚úÖ Confirma√ß√£o de gera√ß√£o conclu√≠da ou relat√≥rio aberto com sucesso.")
    else:
        log(doc, "‚ùå Nenhum modal de confirma√ß√£o ou relat√≥rio detectado ‚Äî verificar execu√ß√£o.")


        indice = 2
        log(doc, "‚ö° Confirmando o envio de boletos por E-mail (sem espera)...")

        js_engine.instant_click(
            f"(//a[@class='btModel btGray btyes' and normalize-space()='Sim'])[{indice}]",
            by_xpath=True
        )
        take_screenshot(driver, doc, f"confirmacao_envio_boletos_{indice}")
        encontrar_mensagem_alerta()





    time.sleep(3)
    ok_cancelou = clicar_cancelar_modal_selecao_conta(js_engine, idx_cancel=0, timeout=8)
    if ok_cancelou:
        log(doc, "‚úÖ Modal 'Selecione uma conta banc√°ria' detectado ‚Äî clique em Cancelar realizado.")
    else:
        log(doc, "‚ÑπÔ∏è Modal 'Selecione uma conta banc√°ria' n√£o apareceu ‚Äî seguindo o fluxo normalmente.")


    if screenshot_final:
        time.sleep(1.0)
        take_screenshot(driver, doc, "Retorno_Tela_Sistema")

    log(doc, "‚úÖ Fluxo conclu√≠do com sucesso.")


def abrir_modal_e_selecionar_js_puro(
    btn_xpath=None,
    pesquisa_xpath=None,
    termo_pesquisa=None,
    btn_pesquisar_xpath=None,
    resultado_xpath=None,
    timeout=15,
    max_tentativas=3,
    iframe_xpath=None,
    indice_lov=None
):
    """
    Abre modal LOV usando JavaScript puro para TODAS opera√É¬ß√É¬µes.
    Elimina problemas de intercepta√É¬ß√£o, visibilidade e timing do Selenium.
    
    Args:
        btn_xpath: XPath do bot√£o LOV (ignorado se indice_lov fornecido)
        pesquisa_xpath: XPath do campo de pesquisa
        termo_pesquisa: Termo a ser pesquisado
        btn_pesquisar_xpath: XPath do bot√£o Pesquisar
        resultado_xpath: XPath do resultado a clicar
        timeout: Tempo m√É¬°ximo de espera
        max_tentativas: N√É¬∫mero de tentativas
        iframe_xpath: XPath do iframe (se aplic√É¬°vel)
        indice_lov: √É¬çndice do bot√£o LOV na p√É¬°gina (alternativa ao btn_xpath)
    """
    
    def _executar_js(script, *args):
        """Executa JavaScript e retorna resultado"""
        try:
            return driver.execute_script(script, *args)
        except Exception as e:
            log(doc, f"  ‚ùå Erro JS: {str(e)[:100]}")
            raise
    
    def _aguardar_ajax_js(timeout_ajax=10):
        """Aguarda AJAX usando JavaScript"""
        inicio = time.time()
        while time.time() - inicio < timeout_ajax:
            completo = _executar_js("""
                // Verifica jQuery
                var jQueryOk = typeof jQuery === 'undefined' || jQuery.active === 0;
                
                // Verifica fetch/XMLHttpRequest pendentes
                var fetchOk = !window.__pendingRequests || window.__pendingRequests === 0;
                
                // Verifica overlays
                var overlays = document.querySelectorAll(
                    '.blockScreen, .blockUI, .loading, .overlay, [class*="loading"], [class*="spinner"]'
                );
                var overlayOk = true;
                for (var i = 0; i < overlays.length; i++) {
                    var style = window.getComputedStyle(overlays[i]);
                    if (style.display !== 'none' && 
                        style.visibility !== 'hidden' && 
                        parseFloat(style.opacity || 1) > 0.01) {
                        overlayOk = false;
                        break;
                    }
                }
                
                return jQueryOk && fetchOk && overlayOk;
            """)
            
            if completo:
                return True
            time.sleep(0.2)
        return True
    
    def _clicar_js(xpath_ou_elemento, descricao="elemento"):
        """Clica usando JavaScript puro"""
        script = """
            var elemento = arguments[0];
            
            // Se recebeu XPath, resolve
            if (typeof elemento === 'string') {
                var resultado = document.evaluate(
                    elemento, 
                    document, 
                    null, 
                    XPathResult.FIRST_ORDERED_NODE_TYPE, 
                    null
                );
                elemento = resultado.singleNodeValue;
            }
            
            if (!elemento) {
                throw new Error('Elemento n√£o encontrado: ' + arguments[0]);
            }
            
            // For√É¬ßa visibilidade
            elemento.style.display = 'block';
            elemento.style.visibility = 'visible';
            elemento.style.opacity = '1';
            
            // Remove atributos que impedem clique
            elemento.removeAttribute('disabled');
            elemento.removeAttribute('readonly');
            
            // Scroll suave at√© o elemento
            elemento.scrollIntoView({behavior: 'smooth', block: 'center'});
            
            // Dispara eventos completos
            ['mouseover', 'mousedown', 'mouseup', 'click'].forEach(function(eventType) {
                var evt = new MouseEvent(eventType, {
                    bubbles: true,
                    cancelable: true,
                    view: window,
                    detail: 1,
                    clientX: 0,
                    clientY: 0
                });
                elemento.dispatchEvent(evt);
            });
            
            // Tamb√©m tenta click direto
            if (elemento.click) {
                elemento.click();
            }
            
            return true;
        """
        
        try:
            _executar_js(script, xpath_ou_elemento)
            log(doc, f"   üîÑ Clique JS em {descricao}")
            return True
        except Exception as e:
            raise Exception(f"Falha ao clicar em {descricao}: {e}")
    
    def _preencher_campo_js(xpath_campo, valor):
        """Preenche campo usando JavaScript puro"""
        script = """
            var xpath = arguments[0];
            var valor = arguments[1];
            
            // Localiza elemento
            var resultado = document.evaluate(
                xpath, 
                document, 
                null, 
                XPathResult.FIRST_ORDERED_NODE_TYPE, 
                null
            );
            var campo = resultado.singleNodeValue;
            
            if (!campo) {
                throw new Error('Campo n√£o encontrado: ' + xpath);
            }
            
            // For√ßa campo edit√°vel
            campo.removeAttribute('disabled');
            campo.removeAttribute('readonly');
            campo.style.display = 'block';
            campo.style.visibility = 'visible';
            
            // Limpa valor anterior
            campo.value = '';
            
            // Focus no campo
            campo.focus();
            
            // Dispara eventos de input
            ['focus', 'click'].forEach(function(evt) {
                campo.dispatchEvent(new Event(evt, {bubbles: true}));
            });
            
            // Define valor
            campo.value = valor;
            
            // Dispara eventos de mudan√ßa
            ['input', 'change', 'blur'].forEach(function(evt) {
                campo.dispatchEvent(new Event(evt, {bubbles: true}));
            });
            
            // Triggers adicionais para frameworks
            if (typeof campo.onchange === 'function') {
                campo.onchange();
            }
            
            // Trigger jQuery se existir
            if (typeof jQuery !== 'undefined') {
                jQuery(campo).trigger('change');
            }
            
            return campo.value;
        """
        
        valor_final = _executar_js(script, xpath_campo, valor)
        log(doc, f"   ‚úÖ Campo preenchido: '{valor_final}'")
        return valor_final
    
    def _aguardar_elemento_js(xpath, timeout_espera=10, deve_estar_visivel=True):
        """Aguarda elemento usando JavaScript"""
        inicio = time.time()
        while time.time() - inicio < timeout_espera:
            existe = _executar_js("""
                var xpath = arguments[0];
                var deveEstarVisivel = arguments[1];
                
                var resultado = document.evaluate(
                    xpath, 
                    document, 
                    null, 
                    XPathResult.FIRST_ORDERED_NODE_TYPE, 
                    null
                );
                var elemento = resultado.singleNodeValue;
                
                if (!elemento) return false;
                
                if (!deveEstarVisivel) return true;
                
                // Verifica visibilidade real
                var style = window.getComputedStyle(elemento);
                return style.display !== 'none' && 
                       style.visibility !== 'hidden' && 
                       parseFloat(style.opacity || 1) > 0.01 &&
                       elemento.offsetParent !== null;
            """, xpath, deve_estar_visivel)
            
            if existe:
                return True
            time.sleep(0.2)
        
        raise Exception(f"Elemento n√£o encontrado ap√≥s {timeout_espera}s: {xpath}")
    
    def _localizar_botao_lov_js():
        """Localiza bot√£o LOV por √ç¬≠ndice ou XPath usando JS"""
        if indice_lov is not None:
            log(doc, f"   üîÑ  Localizando bot√£o LOV por √≠¬≠ndice: {indice_lov}")
            
            # Conta quantos bot√É¬µes existem
            total = _executar_js("""
                var botoes = document.querySelectorAll("a.sprites.sp-openLov");
                return botoes.length;
            """)
            
            if indice_lov >= total:
                raise Exception(f"√çndice {indice_lov} inv√°lido. Encontrados {total} bot√µes LOV")
            
            xpath_resultado = f"(//a[@class='sprites sp-openLov'])[{indice_lov + 1}]"
            log(doc, f"   ‚úÖ Bot√£o LOV #{indice_lov} localizado")
            return xpath_resultado
            
        else:
            if not btn_xpath:
                raise ValueError("btn_xpath ou indice_lov deve ser fornecido")
            log(doc, f" üîÑ Usando XPath fornecido para bot√£o LOV")
            return btn_xpath
    
    def _trocar_iframe_js(iframe_xpath):
        """Troca para iframe usando Selenium (necess√É¬°rio)"""
        try:
            WebDriverWait(driver, timeout).until(
                EC.frame_to_be_available_and_switch_to_it((By.XPATH, iframe_xpath))
            )
            log(doc, "   ‚úÖ Iframe carregado")
            return True
        except Exception as e:
            raise Exception(f"Falha ao trocar para iframe: {e}")
    
    def acao():
        for tentativa in range(1, max_tentativas + 1):
            try:
                log(doc, f"üîÑ Tentativa {tentativa}/{max_tentativas} - Modal LOV (JS Puro)")
                
                # Volta para conte√É¬∫do principal
                try:
                    driver.switch_to.default_content()
                except:
                    pass
                
                # PASSO 1: Localizar e clicar no bot√£o LOV
                xpath_botao = _localizar_botao_lov_js()
                _aguardar_elemento_js(xpath_botao, timeout, deve_estar_visivel=True)
                _clicar_js(xpath_botao, "bot√£o LOV")
                time.sleep(1.5)
                
                # PASSO 2: Trocar para iframe (se necess√É¬°rio)
                if iframe_xpath:
                    log(doc, " üîÑ Trocando para iframe do modal")
                    _trocar_iframe_js(iframe_xpath)
                    time.sleep(0.5)
                
                # PASSO 3: Aguardar modal abrir
                log(doc, " üîÑ Aguardando modal carregar...")
                _aguardar_ajax_js(timeout_ajax=10)
                
                # PASSO 4: Aguardar e preencher campo de pesquisa
                log(doc, f" üîÑ Aguardando campo de pesquisa")
                _aguardar_elemento_js(pesquisa_xpath, timeout, deve_estar_visivel=True)
                time.sleep(0.5)
                
                log(doc, f" üîÑ Preenchendo: '{termo_pesquisa}'")
                _preencher_campo_js(pesquisa_xpath, termo_pesquisa)
                time.sleep(0.5)
                
                # PASSO 5: Clicar em Pesquisar
                log(doc, "   üîÑ Clicando em Pesquisar")
                _aguardar_elemento_js(btn_pesquisar_xpath, timeout, deve_estar_visivel=True)
                _clicar_js(btn_pesquisar_xpath, "bot√£o Pesquisar")
                
                # PASSO 6: Aguardar resultados
                log(doc, "   üîÑ Aguardando resultados...")
                _aguardar_ajax_js(timeout_ajax=15)
                time.sleep(2)
                
                # PASSO 7: Clicar no resultado
                log(doc, "   üîÑ Clicando no resultado")
                _aguardar_elemento_js(resultado_xpath, timeout, deve_estar_visivel=True)
                _clicar_js(resultado_xpath, "resultado da pesquisa")
                time.sleep(1)
                
                # PASSO 8: Voltar para conte√É¬∫do principal
                if iframe_xpath:
                    try:
                        driver.switch_to.default_content()
                        log(doc, "   √¢≈ì‚Ä¶ Voltou para conte√É¬∫do principal")
                    except:
                        pass
                
                # PASSO 9: Aguardar modal fechar
                time.sleep(1)
                _aguardar_ajax_js()
                
                log(doc, f" ‚úÖ Modal LOV processado (tentativa {tentativa}) - JS Puro")
                return True
                
            except Exception as e:
                log(doc, f"‚ùå Tentativa {tentativa} falhou: {str(e)[:200]}")
                
                # Cleanup em caso de erro
                try:
                    driver.switch_to.default_content()
                except:
                    pass
                
                # Tenta fechar modal via JS
                try:
                    _executar_js("""
                        var botoes = document.querySelectorAll(
                            '.ui-dialog-titlebar-close, .close, [class*="close"]'
                        );
                        botoes.forEach(function(btn) {
                            if (btn.offsetParent !== null) {
                                btn.click();
                            }
                        });
                    """)
                    time.sleep(0.5)
                except:
                    pass
                
                if tentativa < max_tentativas:
                    tempo_espera = 2 + (tentativa * 0.5)
                    log(doc, f" üîÑ Aguardando {tempo_espera}s antes de retentar...")
                    time.sleep(tempo_espera)
                else:
                    raise Exception(
                        f"Falha ap√≥s {max_tentativas} tentativas (JS Puro). "
                        f"√öltimo erro: {e}"
                    )
        
        return False
    
    return acao

class JSForceEngine:
    """Motor de execu√ß√£o JavaScript for√ßado - 100% √† prova de falhas"""
    
    def __init__(self, driver, wait, doc):
        self.driver = driver
        self.wait = wait
        self.doc = doc
    
    def execute_js(self, script, *args):
        """Executa JavaScript com tratamento de erros"""
        try:
            return self.driver.execute_script(script, *args)
        except Exception as e:
            log(self.doc, f"‚ö†Ô∏è Erro JS: {str(e)[:150]}")
            raise
    
    def wait_ajax_complete(self, timeout=15):
        """Vers√£o s√≠ncrona ‚Äî sem Promise"""
        end = time.time() + timeout
        while time.time() < end:
            try:
                done = self.driver.execute_script("""
                    var jQueryOk = (typeof jQuery==='undefined') || (jQuery.active===0);
                    var fetchOk = !window.__pendingRequests || window.__pendingRequests===0;
                    var overlays = document.querySelectorAll(
                      '.blockScreen, .blockUI, .loading, .overlay, [class*="loading"], [class*="spinner"], [class*="wait"]'
                    );
                    var overlayOk = true;
                    for (var i=0;i<overlays.length;i++){
                      var s=window.getComputedStyle(overlays[i]);
                      if(s.display!=='none' && s.visibility!=='hidden' && parseFloat(s.opacity||1)>0.01){ overlayOk=false; break; }
                    }
                    return jQueryOk && fetchOk && overlayOk;
                """)
                if done:
                    return True
            except:
                pass
            time.sleep(0.2)
        return True
    
    def force_click(self, selector, by_xpath=False, max_attempts=5):
        """Clique for√ßado 100% garantido usando JavaScript"""
        log(self.doc, f"üéØ Clique for√ßado em: {selector}")
        
        for attempt in range(max_attempts):
            try:
                # Estrat√©gia progressiva de clique
                strategies = [
                    self._click_strategy_2,  # Eventos MouseEvent completos (melhor compatibilidade)
                    self._click_strategy_1,  # Clique nativo otimizado
                    self._click_strategy_3,  # Dispatch direto
                    self._click_strategy_4,  # For√ßa bruta total
                    self._click_strategy_5,  # √öltimo recurso
                ]

                
                for i, strategy in enumerate(strategies, 1):
                    try:
                        log(self.doc, f"   Tentativa {attempt + 1}.{i}...")
                        result = strategy(selector, by_xpath)
                        if result:
                            log(self.doc, f"‚úÖ Clique bem-sucedido (estrat√©gia {i})")
                            time.sleep(0.5)
                            self.wait_ajax_complete(10)
                            return True
                    except Exception as e:
                        log(self.doc, f"   Estrat√©gia {i} falhou: {str(e)[:80]}")
                        continue
                
                if attempt < max_attempts - 1:
                    time.sleep(1 + attempt * 0.5)
                    
            except Exception as e:
                log(self.doc, f"‚ö†Ô∏è Tentativa {attempt + 1} falhou: {e}")
                if attempt < max_attempts - 1:
                    time.sleep(1.5)
        
        raise Exception(f"Falha ao clicar ap√≥s {max_attempts} tentativas: {selector}")
    
    def _click_strategy_1(self, selector, by_xpath):
        """Estrat√©gia 1: Clique nativo otimizado"""
        script = """
        var selector = arguments[0];
        var byXPath = arguments[1];
        
        var element;
        if (byXPath) {
            var result = document.evaluate(selector, document, null, 
                XPathResult.FIRST_ORDERED_NODE_TYPE, null);
            element = result.singleNodeValue;
        } else {
            element = document.querySelector(selector);
        }
        
        if (!element) throw new Error('Elemento n√£o encontrado');
        
        // Remove obst√°culos
        element.style.pointerEvents = 'auto';
        element.style.display = 'block';
        element.style.visibility = 'visible';
        element.style.opacity = '1';
        element.removeAttribute('disabled');
        
        // Scroll suave
        element.scrollIntoView({behavior: 'smooth', block: 'center'});
        
        // Aguarda scroll
        return new Promise(function(resolve) {
            setTimeout(function() {
                element.click();
                resolve(true);
            }, 300);
        });
        """
        return self.execute_js(script, selector, by_xpath)
    
    def _click_strategy_2(self, selector, by_xpath):
        """Estrat√©gia 2: Eventos MouseEvent completos"""
        script = """
        var selector = arguments[0];
        var byXPath = arguments[1];
        
        var element;
        if (byXPath) {
            var result = document.evaluate(selector, document, null, 
                XPathResult.FIRST_ORDERED_NODE_TYPE, null);
            element = result.singleNodeValue;
        } else {
            element = document.querySelector(selector);
        }
        
        if (!element) throw new Error('Elemento n√£o encontrado');
        
        // Prepara elemento
        element.style.pointerEvents = 'auto';
        element.removeAttribute('disabled');
        element.scrollIntoView({block: 'center'});
        
        // Sequ√™ncia completa de eventos
        var events = ['mouseover', 'mouseenter', 'mousemove', 'mousedown', 'mouseup', 'click'];
        events.forEach(function(eventType) {
            var evt = new MouseEvent(eventType, {
                bubbles: true,
                cancelable: true,
                view: window,
                detail: 1,
                clientX: element.getBoundingClientRect().left + 5,
                clientY: element.getBoundingClientRect().top + 5
            });
            element.dispatchEvent(evt);
        });
        
        // Clique adicional
        if (typeof element.click === 'function') {
            element.click();
        }
        
        return true;
        """
        return self.execute_js(script, selector, by_xpath)
    
    def _click_strategy_3(self, selector, by_xpath):
        """Estrat√©gia 3: Dispatch direto com focus"""
        script = """
        var selector = arguments[0];
        var byXPath = arguments[1];
        
        var element;
        if (byXPath) {
            var result = document.evaluate(selector, document, null, 
                XPathResult.FIRST_ORDERED_NODE_TYPE, null);
            element = result.singleNodeValue;
        } else {
            element = document.querySelector(selector);
        }
        
        if (!element) throw new Error('Elemento n√£o encontrado');
        
        // For√ßa visibilidade total
        element.style.display = 'block';
        element.style.visibility = 'visible';
        element.style.opacity = '1';
        element.style.pointerEvents = 'auto';
        
        // Focus
        element.focus();
        
        // Clique direto
        element.click();
        
        // Dispatch adicional
        element.dispatchEvent(new Event('click', {bubbles: true, cancelable: true}));
        
        return true;
        """
        return self.execute_js(script, selector, by_xpath)
    
    def _click_strategy_4(self, selector, by_xpath):
        """Estrat√©gia 4: For√ßa bruta total"""
        script = """
        var selector = arguments[0];
        var byXPath = arguments[1];
        
        var element;
        if (byXPath) {
            var result = document.evaluate(selector, document, null, 
                XPathResult.FIRST_ORDERED_NODE_TYPE, null);
            element = result.singleNodeValue;
        } else {
            element = document.querySelector(selector);
        }
        
        if (!element) throw new Error('Elemento n√£o encontrado');
        
        // Remove TODOS os bloqueios poss√≠veis
        element.removeAttribute('disabled');
        element.removeAttribute('readonly');
        element.style.pointerEvents = 'auto !important';
        element.style.display = 'block !important';
        element.style.visibility = 'visible !important';
        element.style.opacity = '1 !important';
        
        // Remove overlays globais
        var overlays = document.querySelectorAll('.modal, .overlay, .blockUI, [role="dialog"]');
        overlays.forEach(function(overlay) {
            overlay.style.display = 'none';
            overlay.style.visibility = 'hidden';
        });
        
        // M√∫ltiplos m√©todos de clique
        element.focus();
        element.click();
        
        var clickEvent = new MouseEvent('click', {
            view: window,
            bubbles: true,
            cancelable: true
        });
        element.dispatchEvent(clickEvent);
        
        // Trigger jQuery se existir
        if (typeof jQuery !== 'undefined') {
            jQuery(element).trigger('click');
        }
        
        return true;
        """
        return self.execute_js(script, selector, by_xpath)
    
    def _click_strategy_5(self, selector, by_xpath):
        """Estrat√©gia 5: √öltimo recurso - simula clique no ponto exato"""
        script = """
        var selector = arguments[0];
        var byXPath = arguments[1];
        
        var element;
        if (byXPath) {
            var result = document.evaluate(selector, document, null, 
                XPathResult.FIRST_ORDERED_NODE_TYPE, null);
            element = result.singleNodeValue;
        } else {
            element = document.querySelector(selector);
        }
        
        if (!element) throw new Error('Elemento n√£o encontrado');
        
        // Pega coordenadas
        var rect = element.getBoundingClientRect();
        var x = rect.left + rect.width / 2;
        var y = rect.top + rect.height / 2;
        
        // Cria evento no ponto exato
        var evt = document.createEvent('MouseEvents');
        evt.initMouseEvent('click', true, true, window, 1, x, y, x, y, false, false, false, false, 0, null);
        element.dispatchEvent(evt);
        
        // Fallback: procura por onclick
        if (element.onclick) {
            element.onclick();
        }
        
        // Procura handlers no parent
        var parent = element.parentElement;
        while (parent && parent !== document.body) {
            if (parent.onclick) {
                parent.onclick();
                break;
            }
            parent = parent.parentElement;
        }
        
        return true;
        """
        return self.execute_js(script, selector, by_xpath)
    
    def instant_click(self, selector, by_xpath=False):
        log(self.doc, f"‚ö° Clique instant√¢neo em: {selector}")
        # Usa uma estrat√©gia direta, sem waits e sem sleeps
        return self._click_strategy_2(selector, by_xpath)  # dispara eventos de mouse e retorna

    def force_fill(self, selector, value, by_xpath=False, max_attempts=5):
        """Preenchimento for√ßado 100% garantido"""
        log(self.doc, f"‚úèÔ∏è Preenchimento for√ßado: {selector} = '{value}'")
        
        for attempt in range(max_attempts):
            try:
                strategies = [
                    self._fill_strategy_1,  # Native + Events
                    self._fill_strategy_2,  # React/Angular compatible
                    self._fill_strategy_3,  # jQuery trigger
                    self._fill_strategy_4,  # For√ßa bruta
                    self._fill_strategy_5,  # √öltimo recurso
                ]
                
                for i, strategy in enumerate(strategies, 1):
                    try:
                        log(self.doc, f"   Tentativa {attempt + 1}.{i}...")
                        result = strategy(selector, value, by_xpath)
                        
                        # Valida preenchimento
                        time.sleep(0.3)
                        if self._validate_fill(selector, value, by_xpath):
                            log(self.doc, f"‚úÖ Campo preenchido (estrat√©gia {i})")
                            return True
                    except Exception as e:
                        log(self.doc, f"   Estrat√©gia {i} falhou: {str(e)[:80]}")
                        continue
                
                if attempt < max_attempts - 1:
                    time.sleep(1 + attempt * 0.5)
                    
            except Exception as e:
                log(self.doc, f"‚ö†Ô∏è Tentativa {attempt + 1} falhou: {e}")
        
        raise Exception(f"Falha ao preencher ap√≥s {max_attempts} tentativas: {selector}")
    
    def _fill_strategy_1(self, selector, value, by_xpath):
        """Estrat√©gia 1: Native com eventos"""
        script = """
        var selector = arguments[0];
        var value = arguments[1];
        var byXPath = arguments[2];
        
        var element;
        if (byXPath) {
            var result = document.evaluate(selector, document, null, 
                XPathResult.FIRST_ORDERED_NODE_TYPE, null);
            element = result.singleNodeValue;
        } else {
            element = document.querySelector(selector);
        }
        
        if (!element) throw new Error('Campo n√£o encontrado');
        
        // Prepara campo
        element.removeAttribute('disabled');
        element.removeAttribute('readonly');
        element.style.display = 'block';
        element.style.visibility = 'visible';
        
        // Scroll
        element.scrollIntoView({block: 'center'});
        
        // Focus
        element.focus();
        element.dispatchEvent(new Event('focus', {bubbles: true}));
        
        // Limpa
        element.value = '';
        
        // Preenche
        element.value = value;
        
        // Eventos
        ['input', 'change', 'blur', 'keyup'].forEach(function(evt) {
            element.dispatchEvent(new Event(evt, {bubbles: true}));
        });
        
        return element.value;
        """
        return self.execute_js(script, selector, value, by_xpath)
    
    def _fill_strategy_2(self, selector, value, by_xpath):
        """Estrat√©gia 2: Compat√≠vel com React/Angular"""
        script = """
        var selector = arguments[0];
        var value = arguments[1];
        var byXPath = arguments[2];
        
        var element;
        if (byXPath) {
            var result = document.evaluate(selector, document, null, 
                XPathResult.FIRST_ORDERED_NODE_TYPE, null);
            element = result.singleNodeValue;
        } else {
            element = document.querySelector(selector);
        }
        
        if (!element) throw new Error('Campo n√£o encontrado');
        
        // Setter nativo (React)
        var nativeInputValueSetter = Object.getOwnPropertyDescriptor(
            window.HTMLInputElement.prototype, 'value'
        ).set;
        
        if (nativeInputValueSetter) {
            nativeInputValueSetter.call(element, value);
        } else {
            element.value = value;
        }
        
        // Eventos React
        element.dispatchEvent(new Event('input', {bubbles: true}));
        element.dispatchEvent(new Event('change', {bubbles: true}));
        
        // Eventos adicionais
        element.dispatchEvent(new KeyboardEvent('keydown', {bubbles: true}));
        element.dispatchEvent(new KeyboardEvent('keyup', {bubbles: true}));
        element.dispatchEvent(new Event('blur', {bubbles: true}));
        
        return element.value;
        """
        return self.execute_js(script, selector, value, by_xpath)
    
    def _fill_strategy_3(self, selector, value, by_xpath):
        """Estrat√©gia 3: jQuery trigger"""
        script = """
        var selector = arguments[0];
        var value = arguments[1];
        var byXPath = arguments[2];
        
        var element;
        if (byXPath) {
            var result = document.evaluate(selector, document, null, 
                XPathResult.FIRST_ORDERED_NODE_TYPE, null);
            element = result.singleNodeValue;
        } else {
            element = document.querySelector(selector);
        }
        
        if (!element) throw new Error('Campo n√£o encontrado');
        
        element.value = value;
        
        // Trigger jQuery se dispon√≠vel
        if (typeof jQuery !== 'undefined') {
            jQuery(element).val(value).trigger('input').trigger('change').trigger('blur');
        }
        
        // Eventos nativos como fallback
        ['focus', 'input', 'change', 'blur'].forEach(function(evt) {
            element.dispatchEvent(new Event(evt, {bubbles: true}));
        });
        
        return element.value;
        """
        return self.execute_js(script, selector, value, by_xpath)
    
    def _fill_strategy_4(self, selector, value, by_xpath):
        """Estrat√©gia 4: For√ßa bruta"""
        script = """
        var selector = arguments[0];
        var value = arguments[1];
        var byXPath = arguments[2];
        
        var element;
        if (byXPath) {
            var result = document.evaluate(selector, document, null, 
                XPathResult.FIRST_ORDERED_NODE_TYPE, null);
            element = result.singleNodeValue;
        } else {
            element = document.querySelector(selector);
        }
        
        if (!element) throw new Error('Campo n√£o encontrado');
        
        // Remove TODOS bloqueios
        element.removeAttribute('disabled');
        element.removeAttribute('readonly');
        element.removeAttribute('maxlength');
        
        // Define valor m√∫ltiplas vezes
        element.value = '';
        element.setAttribute('value', value);
        element.value = value;
        
        // For√ßa atualiza√ß√£o visual
        element.style.color = element.style.color;
        
        // TODOS os eventos poss√≠veis
        var events = ['focus', 'click', 'input', 'change', 'keydown', 'keypress', 
                      'keyup', 'blur', 'paste', 'textInput'];
        events.forEach(function(evt) {
            try {
                element.dispatchEvent(new Event(evt, {bubbles: true, cancelable: true}));
            } catch(e) {}
        });
        
        // Handlers diretos
        if (element.oninput) element.oninput();
        if (element.onchange) element.onchange();
        
        return element.value;
        """
        return self.execute_js(script, selector, value, by_xpath)
    
    def _fill_strategy_5(self, selector, value, by_xpath):
        """Estrat√©gia 5: √öltimo recurso - simula digita√ß√£o"""
        script = """
        var selector = arguments[0];
        var value = arguments[1];
        var byXPath = arguments[2];
        
        var element;
        if (byXPath) {
            var result = document.evaluate(selector, document, null, 
                XPathResult.FIRST_ORDERED_NODE_TYPE, null);
            element = result.singleNodeValue;
        } else {
            element = document.querySelector(selector);
        }
        
        if (!element) throw new Error('Campo n√£o encontrado');
        
        element.focus();
        element.value = '';
        
        // Simula digita√ß√£o caractere por caractere
        for (var i = 0; i < value.length; i++) {
            element.value += value[i];
            
            // Evento para cada caractere
            var evt = new KeyboardEvent('keydown', {
                key: value[i],
                code: 'Key' + value[i].toUpperCase(),
                bubbles: true
            });
            element.dispatchEvent(evt);
            
            element.dispatchEvent(new Event('input', {bubbles: true}));
        }
        
        element.dispatchEvent(new Event('change', {bubbles: true}));
        element.dispatchEvent(new Event('blur', {bubbles: true}));
        
        return element.value;
        """
        return self.execute_js(script, selector, value, by_xpath)
    
    def _validate_fill(self, selector, expected_value, by_xpath):
        """Valida se o campo foi preenchido corretamente"""
        script = """
        var selector = arguments[0];
        var expected = arguments[1];
        var byXPath = arguments[2];
        
        var element;
        if (byXPath) {
            var result = document.evaluate(selector, document, null, 
                XPathResult.FIRST_ORDERED_NODE_TYPE, null);
            element = result.singleNodeValue;
        } else {
            element = document.querySelector(selector);
        }
        
        if (!element) return false;
        
        var actual = element.value || '';
        return actual.trim() === expected.trim() || actual.includes(expected);
        """
        try:
            return self.execute_js(script, selector, expected_value, by_xpath)
        except:
            return False
    
    def force_select(self, selector, text, by_xpath=False, max_attempts=5):
        """Sele√ß√£o for√ßada em dropdown/select"""
        log(self.doc, f"üîΩ Sele√ß√£o for√ßada: {selector} = '{text}'")
        
        for attempt in range(max_attempts):
            try:
                script = """
                var selector = arguments[0];
                var text = arguments[1];
                var byXPath = arguments[2];
                
                var element;
                if (byXPath) {
                    var result = document.evaluate(selector, document, null, 
                        XPathResult.FIRST_ORDERED_NODE_TYPE, null);
                    element = result.singleNodeValue;
                } else {
                    element = document.querySelector(selector);
                }
                
                if (!element) throw new Error('Select n√£o encontrado');
                
                // Tenta por texto vis√≠vel
                for (var i = 0; i < element.options.length; i++) {
                    if (element.options[i].text.trim() === text.trim()) {
                        element.selectedIndex = i;
                        element.value = element.options[i].value;
                        element.dispatchEvent(new Event('change', {bubbles: true}));
                        return true;
                    }
                }
                
                // Tenta por value
                element.value = text;
                element.dispatchEvent(new Event('change', {bubbles: true}));
                
                return element.value !== '';
                """
                
                result = self.execute_js(script, selector, text, by_xpath)
                if result:
                    log(self.doc, f"‚úÖ Op√ß√£o selecionada: '{text}'")
                    return True
                    
                if attempt < max_attempts - 1:
                    time.sleep(1)
                    
            except Exception as e:
                log(self.doc, f"‚ö†Ô∏è Tentativa {attempt + 1} falhou: {e}")
                if attempt < max_attempts - 1:
                    time.sleep(1)
        
        raise Exception(f"Falha ao selecionar op√ß√£o ap√≥s {max_attempts} tentativas")
    
    def force_datepicker(self, selector, date_value, by_xpath=False, max_attempts=5):
        """Preenchimento for√ßado de datepicker"""
        log(self.doc, f"üìÖ Datepicker for√ßado: {selector} = '{date_value}'")
        
        for attempt in range(max_attempts):
            try:
                script = """
                var selector = arguments[0];
                var dateValue = arguments[1];
                var byXPath = arguments[2];
                
                var element;
                if (byXPath) {
                    var result = document.evaluate(selector, document, null, 
                        XPathResult.FIRST_ORDERED_NODE_TYPE, null);
                    element = result.singleNodeValue;
                } else {
                    element = document.querySelector(selector);
                }
                
                if (!element) throw new Error('Datepicker n√£o encontrado');
                
                // Remove readonly/disabled
                element.removeAttribute('readonly');
                element.removeAttribute('disabled');
                
                // jQuery datepicker
                if (typeof jQuery !== 'undefined' && jQuery(element).datepicker) {
                    try {
                        jQuery(element).datepicker('setDate', dateValue);
                        jQuery(element).trigger('change');
                        return element.value;
                    } catch(e) {}
                }
                
                // Native
                element.value = '';
                element.value = dateValue;
                
                // Eventos completos
                ['focus', 'input', 'change', 'blur', 'keyup'].forEach(function(evt) {
                    element.dispatchEvent(new Event(evt, {bubbles: true}));
                });
                
                return element.value;
                """
                
                result = self.execute_js(script, selector, date_value, by_xpath)
                
                # Valida
                time.sleep(0.3)
                if self._validate_fill(selector, date_value, by_xpath):
                    log(self.doc, f"‚úÖ Datepicker preenchido: '{date_value}'")
                    return True
                
                if attempt < max_attempts - 1:
                    time.sleep(1)
                    
            except Exception as e:
                log(self.doc, f"‚ö†Ô∏è Tentativa {attempt + 1} falhou: {e}")
                if attempt < max_attempts - 1:
                    time.sleep(1)
        
        raise Exception(f"Falha ao preencher datepicker ap√≥s {max_attempts} tentativas")
    
    def force_modal_open(self, btn_selector, modal_selector, by_xpath=False, max_attempts=5):
        """Abre modal for√ßadamente e aguarda aparecer"""
        log(self.doc, f"üîì Abrindo modal: {btn_selector}")
        
        for attempt in range(max_attempts):
            try:
                # Clica no bot√£o
                self.force_click(btn_selector, by_xpath)
                
                # Aguarda modal aparecer
                time.sleep(1)
                
                script = """
                var selector = arguments[0];
                var byXPath = arguments[1];
                
                var element;
                if (byXPath) {
                    var result = document.evaluate(selector, document, null, 
                        XPathResult.FIRST_ORDERED_NODE_TYPE, null);
                    element = result.singleNodeValue;
                } else {
                    element = document.querySelector(selector);
                }
                
                if (!element) return false;
                
                var style = window.getComputedStyle(element);
                return style.display !== 'none' && 
                       style.visibility !== 'hidden' && 
                       parseFloat(style.opacity || 1) > 0.01;
                """
                
                modal_visible = self.execute_js(script, modal_selector, by_xpath)
                
                if modal_visible:
                    log(self.doc, "‚úÖ Modal aberto com sucesso")
                    self.wait_ajax_complete(5)
                    return True
                
                if attempt < max_attempts - 1:
                    log(self.doc, f"   Modal n√£o apareceu, tentando novamente...")
                    time.sleep(1.5)
                    
            except Exception as e:
                log(self.doc, f"‚ö†Ô∏è Tentativa {attempt + 1} falhou: {e}")
                if attempt < max_attempts - 1:
                    time.sleep(1.5)
        
        raise Exception(f"Falha ao abrir modal ap√≥s {max_attempts} tentativas")
    
    def force_modal_close(self, close_selector=None, by_xpath=False, max_attempts=3):
        """Fecha modal for√ßadamente"""
        log(self.doc, "üîí Fechando modal...")
        
        for attempt in range(max_attempts):
            try:
                script = """
                var closeSelector = arguments[0];
                var byXPath = arguments[1];
                
                // Tenta seletor espec√≠fico
                if (closeSelector) {
                    var element;
                    if (byXPath) {
                        var result = document.evaluate(closeSelector, document, null, 
                            XPathResult.FIRST_ORDERED_NODE_TYPE, null);
                        element = result.singleNodeValue;
                    } else {
                        element = document.querySelector(closeSelector);
                    }
                    
                    if (element) {
                        element.click();
                        return true;
                    }
                }
                
                // Busca bot√µes de fechar comuns
                var selectors = [
                    '.ui-dialog-titlebar-close',
                    '.close',
                    '[data-dismiss="modal"]',
                    '.modal-close',
                    'button[aria-label="Close"]',
                    '.wdClose a'
                ];
                
                for (var i = 0; i < selectors.length; i++) {
                    var btn = document.querySelector(selectors[i]);
                    if (btn && btn.offsetParent !== null) {
                        btn.click();
                        return true;
                    }
                }
                
                // Remove modais diretamente
                var modals = document.querySelectorAll('.modal, .ui-dialog, [role="dialog"]');
                modals.forEach(function(modal) {
                    modal.style.display = 'none';
                    modal.remove();
                });
                
                return modals.length > 0;
                """
                
                result = self.execute_js(script, close_selector, by_xpath)
                
                if result:
                    log(self.doc, "‚úÖ Modal fechado")
                    time.sleep(0.5)
                    return True
                
                if attempt < max_attempts - 1:
                    time.sleep(1)
                    
            except Exception as e:
                log(self.doc, f"‚ö†Ô∏è Tentativa {attempt + 1} falhou: {e}")
                if attempt < max_attempts - 1:
                    time.sleep(1)
        
        log(self.doc, "‚ö†Ô∏è N√£o foi poss√≠vel fechar modal, continuando...")
        return False


class LOVHandler:
    """Handler especializado para modais LOV (List of Values)"""
    
    def __init__(self, js_engine):
        self.js = js_engine
        self.doc = js_engine.doc
    
    def open_and_select(self, btn_index=None, btn_xpath=None, search_text="", 
                       result_text="", iframe_xpath=None, max_attempts=5):
        """Abre LOV, pesquisa e seleciona resultado usando JS for√ßado"""
        
        log(self.doc, f"üîç Processando LOV: '{search_text}' ‚Üí '{result_text}'")
        
        for attempt in range(max_attempts):
            try:
                log(self.doc, f"   Tentativa {attempt + 1}/{max_attempts}")
                
                # PASSO 1: Clica no bot√£o LOV
                if btn_index is not None:
                    btn_selector = f"(//a[@class='sprites sp-openLov'])[{btn_index + 1}]"
                    by_xpath = True
                elif btn_xpath:
                    btn_selector = btn_xpath
                    by_xpath = True
                else:
                    raise ValueError("btn_index ou btn_xpath deve ser fornecido")
                
                log(self.doc, "   üìå Abrindo LOV...")
                self.js.force_click(btn_selector, by_xpath=by_xpath)
                time.sleep(1.5)
                
                # PASSO 2: Troca para iframe se necess√°rio
                if iframe_xpath:
                    log(self.doc, "   üîÑ Entrando no iframe...")
                    try:
                        WebDriverWait(self.js.driver, 10).until(
                            EC.frame_to_be_available_and_switch_to_it((By.XPATH, iframe_xpath))
                        )
                        time.sleep(0.5)
                    except:
                        log(self.doc, "   ‚ö†Ô∏è Iframe n√£o encontrado, continuando...")
                
                # PASSO 3: Aguarda modal carregar
                self.js.wait_ajax_complete(10)
                time.sleep(1)
                
                # PASSO 4: Preenche campo de pesquisa
                log(self.doc, f"   ‚úèÔ∏è Pesquisando: '{search_text}'")
                
                search_selectors = [
                    "//input[@id='txtPesquisa']",
                    "//input[@class='nomePesquisa']",
                ]
                
                search_filled = False
                for selector in search_selectors:
                    try:
                        self.js.force_fill(selector, search_text, by_xpath=True)
                        search_filled = True
                        break
                    except:
                        continue
                
                if not search_filled:
                    raise Exception("Campo de pesquisa n√£o encontrado")
                
                time.sleep(0.5)
                
                # PASSO 5: Clica em Pesquisar
                log(self.doc, "   üîé Executando pesquisa...")
                
                search_btn_selectors = [
                    "//a[contains(@class,'lpFind') and contains(normalize-space(.),'Pesquisar')]",
                    "//button[contains(normalize-space(.),'Pesquisar')]",
                    "//a[contains(normalize-space(.),'Buscar')]"
                ]
                
                search_clicked = False
                for selector in search_btn_selectors:
                    try:
                        self.js.force_click(selector, by_xpath=True)
                        search_clicked = True
                        break
                    except:
                        continue
                
                if not search_clicked:
                    raise Exception("Bot√£o de pesquisa n√£o encontrado")
                
                # PASSO 6: Aguarda resultados
                time.sleep(2)
                self.js.wait_ajax_complete(15)
                
                # PASSO 7: Clica no resultado
                log(self.doc, f"   üéØ Selecionando: '{result_text}'")
                
                result_xpath = f"//tr[td[contains(normalize-space(.), '{result_text}')]]"
                self.js.force_click(result_xpath, by_xpath=True)
                
                time.sleep(1)
                
                # PASSO 8: Volta para conte√∫do principal
                if iframe_xpath:
                    try:
                        self.js.driver.switch_to.default_content()
                        log(self.doc, "   ‚úÖ Voltou para conte√∫do principal")
                    except:
                        pass
                
                # PASSO 9: Aguarda modal fechar
                time.sleep(1)
                self.js.wait_ajax_complete(10)
                
                log(self.doc, f"‚úÖ LOV processado com sucesso!")
                return True
                
            except Exception as e:
                log(self.doc, f"‚ö†Ô∏è Tentativa {attempt + 1} falhou: {str(e)[:150]}")
                
                # Cleanup
                try:
                    self.js.driver.switch_to.default_content()
                except:
                    pass
                
                try:
                    self.js.force_modal_close()
                except:
                    pass
                
                if attempt < max_attempts - 1:
                    time.sleep(2 + attempt * 0.5)
                else:
                    raise Exception(f"Falha ao processar LOV ap√≥s {max_attempts} tentativas: {e}")
        
        return False


# ==== WRAPPERS DE ALTO N√çVEL ====

def safe_action(doc, descricao, func, max_retries=3):
    """Wrapper para a√ß√µes com retry autom√°tico"""
    global driver
    
    for attempt in range(max_retries):
        try:
            log(doc, f"üîÑ {descricao}..." if attempt == 0 else f"üîÑ {descricao}... (Tentativa {attempt + 1})")
            func()
            log(doc, f"‚úÖ {descricao} realizada com sucesso.")
            take_screenshot(driver, doc, _sanitize_filename(descricao))
            return True
        except Exception as e:
            if attempt < max_retries - 1:
                log(doc, f"‚ö†Ô∏è Tentativa {attempt + 1} falhou, tentando novamente...")
                time.sleep(2 + attempt)
                continue
            else:
                log(doc, f"‚ùå Erro ap√≥s {max_retries} tentativas: {e}")
                take_screenshot(driver, doc, _sanitize_filename(f"erro_{descricao}"))
                return False
    
    return False



def fechar_abas_padrao(js_engine, doc):
    """
    Fecha apenas as abas principais (sem boletos gerados).
    Utilizado quando n√£o h√° gera√ß√£o de boletos ou relat√≥rios.
    """
    import time

    try:
        safe_action(doc, "Fechando aba de Gera√ß√£o de Boletos", lambda:
            js_engine.force_click(
                "//a[@class='sprites sp-fecharGrande' and @title='Sair']",
                by_xpath=True
            )
        )
        time.sleep(1)

        safe_action(doc, "Fechando Gestor Financeiro", lambda:
            js_engine.force_click(
                "#gsFinan > div.wdTop.ui-draggable-handle > div.wdClose > a"
            )
        )
    except Exception as e:
        log(doc, f"‚ö†Ô∏è Erro ao fechar abas padr√£o: {e}")


def inicializar_driver():
    """Inicializa WebDriver com configura√ß√µes otimizadas"""
    global driver, wait
    
    try:
        options = Options()
        options.add_argument("--start-maximized")
        options.add_argument("--disable-blink-features=AutomationControlled")
        options.add_argument("--disable-dev-shm-usage")
        options.add_argument("--no-sandbox")
        options.add_experimental_option("excludeSwitches", ["enable-automation"])
        options.add_experimental_option('useAutomationExtension', False)
        
        driver = webdriver.Chrome(
            service=Service(ChromeDriverManager().install()), 
            options=options
        )
        
        driver.execute_script(
            "Object.defineProperty(navigator, 'webdriver', {get: () => undefined})"
        )
        
        wait = WebDriverWait(driver, TIMEOUT_DEFAULT)
        
        log(doc, "‚úÖ Driver inicializado com sucesso")
        return True
        
    except Exception as e:
        log(doc, f"‚ùå Erro ao inicializar driver: {e}")
        return False


def finalizar_relatorio():
    """Salva relat√≥rio e fecha driver"""
    global driver, doc
    
    nome_arquivo = f"relatorio_geracao_titulos_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    
    try:
        doc.save(nome_arquivo)
        log(doc, f"üìÑ Relat√≥rio salvo: {nome_arquivo}")
        
        try:
            subprocess.run(["start", "winword", nome_arquivo], shell=True)
        except:
            pass
            
    except Exception as e:
        print(f"Erro ao salvar relat√≥rio: {e}")
    
    if driver:
        try:
            driver.quit()
            log(doc, "‚úÖ Driver encerrado")
        except:
            pass


def clicar_lov_por_indice(indice_lov: int, max_tentativas: int = 5, timeout: int = 10, scroll: bool = True):
    """
    Clica no √≠cone de LOV <a class="sprites sp-openLov"> pelo √≠ndice (ordem no DOM).
    Retorna uma fun√ß√£o 'acao' para ser usada com safe_action(..., lambda: ...).

    Estrat√©gias de clique:
      1) Espera 'clickable' + click nativo
      2) ScrollIntoView + click nativo
      3) JavaScript click
      4) ActionChains move_to_element + click
    """
    def acao():
        if not isinstance(indice_lov, int) or indice_lov < 0:
            raise ValueError(f"√çndice inv√°lido: {indice_lov}")

        tentativa = 0
        while tentativa < max_tentativas:
            tentativa += 1
            try:
                log(doc, f"üîé Tentativa {tentativa}: Localizando √≠cones LOV ('sp-openLov')...")
                # Coleta atual dos elementos
                elementos = driver.find_elements(By.CSS_SELECTOR, "a.sprites.sp-openLov")

                if not elementos:
                    if tentativa < max_tentativas:
                        log(doc, f"‚ö†Ô∏è Nenhum √≠cone LOV encontrado (tentativa {tentativa}/{max_tentativas}). Reintentando...", "WARN")
                        time.sleep(1.2)
                        continue
                    raise Exception("Nenhum √≠cone LOV ('a.sprites.sp-openLov') foi encontrado na p√°gina.")

                if indice_lov >= len(elementos):
                    raise Exception(f"√çndice {indice_lov} inv√°lido. Encontrados {len(elementos)} √≠cones LOV.")

                # Mant√©m um localizador est√°vel por √≠ndice para waits/refresh
                locator_xpath = f"(//a[contains(@class,'sp-openLov')])[{indice_lov + 1}]"

                # Reaponta o elemento pelo locator (evita stale)
                elemento = driver.find_element(By.XPATH, locator_xpath)

                log(doc, f"üéØ Preparando clique no LOV de √≠ndice {indice_lov} (total: {len(elementos)}).")

                def _wait_clickable():
                    wait.until(EC.element_to_be_clickable((By.XPATH, locator_xpath)))

                estrategias = [
                    # 1) Espera 'clickable' + click nativo
                    lambda: (_wait_clickable(), elemento.click()),
                    # 2) ScrollIntoView + click nativo
                    lambda: (
                        driver.execute_script("arguments[0].scrollIntoView({block:'center', inline:'nearest'});", elemento) if scroll else None,
                        time.sleep(0.2),
                        elemento.click()
                    ),
                    # 3) JavaScript click
                    lambda: driver.execute_script("arguments[0].click();", elemento),
                    # 4) ActionChains
                    lambda: (ActionChains(driver).move_to_element(elemento).pause(0.1).click().perform())
                ]

                for i, estrategia in enumerate(estrategias, 1):
                    try:
                        log(doc, f"   ‚ñ∂Ô∏è Estrat√©gia {i} de clique no LOV...")
                        estrategia()
                        time.sleep(0.3)
                        log(doc, f"‚úÖ Clique no LOV (√≠ndice {indice_lov}) realizado com sucesso (estrat√©gia {i}).")
                        return True
                    except (ElementClickInterceptedException, StaleElementReferenceException, JavascriptException, TimeoutException) as e:
                        log(doc, f"‚ö†Ô∏è Estrat√©gia {i} falhou: {e}", "WARN")
                        # Tenta re-obter o elemento em caso de stale/intercepta√ß√£o
                        try:
                            _ = driver.find_elements(By.CSS_SELECTOR, "a.sprites.sp-openLov")
                            elemento = driver.find_element(By.XPATH, locator_xpath)
                        except Exception:
                            pass
                        continue

                if tentativa < max_tentativas:
                    log(doc, f"‚ö†Ô∏è Tentativa {tentativa} n√£o conseguiu clicar no LOV. Reintentando em 1.2s‚Ä¶", "WARN")
                    time.sleep(1.2)
                    continue

            except Exception as e:
                if tentativa < max_tentativas:
                    log(doc, f"‚ö†Ô∏è Erro na tentativa {tentativa}: {e}. Reintentando em 1.2s‚Ä¶", "WARN")
                    time.sleep(1.2)
                    continue
                raise

        raise Exception(f"Falha ao clicar no LOV de √≠ndice {indice_lov} ap√≥s {max_tentativas} tentativas.")

    return acao



def abrir_modal_e_selecionar_robusto_xpath(
    btn_xpath,
    pesquisa_xpath,
    termo_pesquisa,
    btn_pesquisar_xpath,
    resultado_xpath,
    timeout=12,
    max_tentativas=3,
    iframe_xpath=None,
    indice_lov=None,
    **_ignorar_kwargs
):
    """
    Abre o modal (LOV), pesquisa pelo termo e clica no resultado.
    - Usa clicar_lov_por_indice se o √≠ndice for informado
    - Usa retries, fallback JS e limpeza resistente do input
    - Pode lidar com iframe do modal
    """

    def _js_click(el):
        driver.execute_script("arguments[0].click();", el)

    def _clear_resistente(el):
        try:
            el.clear()
            el.send_keys(Keys.CONTROL, "a")
            el.send_keys(Keys.DELETE)
        except Exception:
            driver.execute_script("arguments[0].value='';", el)

    def _aguardar_ajax_overlay():
        t0 = time.time()
        while time.time() - t0 < 8:
            try:
                ready = driver.execute_script("return document.readyState")
                ajax_ok = driver.execute_script("return window.jQuery ? jQuery.active === 0 : true")
                if ready == "complete" and ajax_ok:
                    break
            except Exception:
                pass
            time.sleep(0.2)

    def acao():
        for tentativa in range(1, max_tentativas + 1):
            try:
                log(doc, f"üîÑ Tentativa {tentativa}/{max_tentativas} - Abrindo LOV robusto")

                driver.switch_to.default_content()
                if indice_lov is not None:
                    log(doc, f"üìå Clicando no LOV de √≠ndice {indice_lov}")
                    clicar_lov_por_indice(indice_lov, max_tentativas=3, timeout=timeout)()
                else:
                    log(doc, "üìå Clicando no bot√£o LOV pelo XPath")
                    botao = WebDriverWait(driver, timeout).until(
                        EC.element_to_be_clickable((By.XPATH, btn_xpath))
                    )
                    try:
                        botao.click()
                    except:
                        _js_click(botao)

                # Troca para iframe, se houver
                if iframe_xpath:
                    WebDriverWait(driver, timeout).until(
                        EC.frame_to_be_available_and_switch_to_it((By.XPATH, iframe_xpath))
                    )

                _aguardar_ajax_overlay()

                # Campo de pesquisa
                campo = WebDriverWait(driver, timeout).until(
                    EC.visibility_of_element_located((By.XPATH, pesquisa_xpath))
                )
                _clear_resistente(campo)
                campo.send_keys(termo_pesquisa)

                # Bot√£o pesquisar
                btn_pesq = WebDriverWait(driver, timeout).until(
                    EC.element_to_be_clickable((By.XPATH, btn_pesquisar_xpath))
                )
                try:
                    btn_pesq.click()
                except:
                    _js_click(btn_pesq)

                _aguardar_ajax_overlay()
                time.sleep(1)

                # Resultado
                resultado = WebDriverWait(driver, timeout).until(
                    EC.element_to_be_clickable((By.XPATH, resultado_xpath))
                )
                driver.execute_script("arguments[0].scrollIntoView({block:'center'});", resultado)
                try:
                    resultado.click()
                except:
                    _js_click(resultado)

                driver.switch_to.default_content()
                log(doc, "‚úÖ Modal LOV processado com sucesso!")
                return True

            except Exception as e:
                log(doc, f"‚ö†Ô∏è Falha na tentativa {tentativa}: {e}")
                driver.switch_to.default_content()
                if tentativa < max_tentativas:
                    time.sleep(2)
                else:
                    raise

    return acao

def validar_resultado_pesquisa(js_engine, tempo_maximo=10):
    """Valida se h√° resultados de t√≠tulos pela presen√ßa da tabela 'niceTable padding10'"""
    try:
        log(doc, "üîç Validando resultado da pesquisa...")

        tempo_inicial = time.time()
        tabela_existe = False

        # Espera ativa at√© a tabela aparecer ou o tempo m√°ximo expirar
        while time.time() - tempo_inicial < tempo_maximo:
            script = """
            var tabela = document.querySelector('table.niceTable.padding10');
            return tabela ? true : false;
            """
            tabela_existe = js_engine.execute_js(script)

            if tabela_existe:
                break
            time.sleep(1)

        if tabela_existe:
            log(doc, "‚úÖ Tabela de T√≠tulos carregada com sucesso, prosseguindo")
            return True
        else:
            log(doc, "‚ùå Nenhum T√≠tulo encontrado ap√≥s aguardar carregamento")
            return False

    except Exception as e:
        log(doc, f"‚ùå Erro ao validar resultado: {e}")
        return False



def aguardar_elemento_disponivel(driver, selector, by_type=By.CSS_SELECTOR, timeout=30):
    """Aguarda elemento estar presente, vis√≠vel e clic√°vel"""
    try:
        wait = WebDriverWait(driver, timeout)
        # Aguarda estar presente
        wait.until(EC.presence_of_element_located((by_type, selector)))
        # Aguarda estar vis√≠vel
        wait.until(EC.visibility_of_element_located((by_type, selector)))
        # Aguarda estar clic√°vel
        element = wait.until(EC.element_to_be_clickable((by_type, selector)))
        return element
    except TimeoutException:
        return None

def safe_click_enhanced(driver, selector, by_type=By.CSS_SELECTOR, timeout=30):
    """Fun√ß√£o de clique ultra-robusta com m√∫ltiplas estrat√©gias"""
    strategies = [
        "aguardar_e_clicar_normal",
        "aguardar_e_clicar_js", 
        "aguardar_e_clicar_action",
        "for√ßa_bruta_js"
    ]
    
    for strategy in strategies:
        try:
            log(doc, f"üîÑ Tentando estrat√©gia: {strategy}")
            
            if strategy == "aguardar_e_clicar_normal":
                element = aguardar_elemento_disponivel(driver, selector, by_type, timeout)
                if element:
                    # Rola para o elemento
                    driver.execute_script("arguments[0].scrollIntoView({block: 'center', behavior: 'smooth'});", element)
                    time.sleep(1)
                    # Relocaliza para evitar stale reference
                    element = driver.find_element(by_type, selector)
                    element.click()
                    return True
                    
            elif strategy == "aguardar_e_clicar_js":
                element = aguardar_elemento_disponivel(driver, selector, by_type, timeout)
                if element:
                    driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", element)
                    time.sleep(0.5)
                    driver.execute_script("arguments[0].click();", element)
                    return True
                    
            elif strategy == "aguardar_e_clicar_action":
                element = aguardar_elemento_disponivel(driver, selector, by_type, timeout)
                if element:
                    actions = ActionChains(driver)
                    actions.move_to_element(element).pause(0.5).click().perform()
                    return True
                    
            elif strategy == "for√ßa_bruta_js":
                # √öltima tentativa: for√ßa bruta com JavaScript
                if by_type == By.CSS_SELECTOR:
                    js_code = f"""
                        var element = document.querySelector('{selector}');
                        if (element) {{
                            element.scrollIntoView({{block: 'center'}});
                            setTimeout(function() {{
                                element.click();
                                console.log('Clique for√ßado executado');
                            }}, 500);
                            return true;
                        }}
                        return false;
                    """
                else:  # XPATH
                    js_code = f"""
                        var element = document.evaluate('{selector}', document, null, XPathResult.FIRST_ORDERED_NODE_TYPE, null).singleNodeValue;
                        if (element) {{
                            element.scrollIntoView({{block: 'center'}});
                            setTimeout(function() {{
                                element.click();
                                console.log('Clique for√ßado XPath executado');
                            }}, 500);
                            return true;
                        }}
                        return false;
                    """
                
                result = driver.execute_script(js_code)
                if result:
                    time.sleep(1)
                    return True
                    
        except Exception as e:
            log(doc, f"‚ö†Ô∏è Estrat√©gia {strategy} falhou: {str(e)[:100]}...")
            continue
    
    return False

def safe_send_keys_enhanced(driver, selector, text, by_type=By.CSS_SELECTOR, clear=True, timeout=20):
    """Fun√ß√£o para envio seguro de texto com retry"""
    for attempt in range(3):
        try:
            element = aguardar_elemento_disponivel(driver, selector, by_type, timeout)
            if element:
                driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", element)
                time.sleep(0.3)
                
                if clear:
                    element.clear()
                    # Fallback para limpar
                    element.send_keys(Keys.CONTROL + "a")
                    element.send_keys(Keys.DELETE)
                
                element.send_keys(text)
                return True
        except Exception as e:
            log(doc, f"‚ö†Ô∏è Tentativa {attempt + 1} de envio de texto falhou: {e}")
            time.sleep(1)
    
    return False

def safe_action_enhanced(driver, doc, descricao, func, max_tentativas=3):
    """Fun√ß√£o safe_action aprimorada"""
    for tentativa in range(max_tentativas):
        try:
            log(doc, f"üîÑ {descricao}... (Tentativa {tentativa + 1})")
            result = func()
            if result is False:  # Se a fun√ß√£o retornou False explicitamente
                raise Exception("Fun√ß√£o retornou False")
            log(doc, f"‚úÖ {descricao} realizada com sucesso.")
            take_screenshot(driver, doc, descricao.lower().replace(" ", "_"))
            return True
        except Exception as e:
            if tentativa < max_tentativas - 1:
                log(doc, f"‚ö†Ô∏è Tentativa {tentativa + 1} falhou: {str(e)[:100]}... Tentando novamente...")
                time.sleep(2)
            else:
                log(doc, f"‚ùå Erro ao {descricao.lower()}: {str(e)[:200]}...")
                take_screenshot(driver, doc, f"erro_{descricao.lower().replace(' ', '_')}")
                return False



def confirmar_modal_generico(driver, doc, seletor="#BtYes", mensagem="Confirmando modal", timeout=5, delay=2):
    """Confirma modal gen√©rico se ele estiver presente"""
    def acao():
        if delay > 0:
            time.sleep(delay)
        
        elementos = driver.find_elements(By.CSS_SELECTOR, seletor)
        if elementos:
            safe_click_enhanced(driver, seletor, timeout=timeout)
            return True
        return False
    
    return safe_action_enhanced(driver, doc, mensagem, acao)


def confirmar_modal_geracao_titulos(js_engine, timeout=12, iframe_xpath=None):
    """
    Clica no 'Sim' da modal 'Confirme a gera√ß√£o de t√≠tulos...' e valida que ela fechou.
    Retorna True somente se a modal deixar de estar vis√≠vel.
    """
    from selenium.webdriver.support.ui import WebDriverWait
    from selenium.webdriver.support import expected_conditions as EC
    from selenium.webdriver.common.by import By
    import time

    driver = js_engine.driver
    wait = WebDriverWait(driver, timeout)

    # Sair para o conte√∫do principal
    try:
        driver.switch_to.default_content()
    except:
        pass

    # Se houver iframe, entrar (se n√£o houver, ignora)
    if iframe_xpath:
        try:
            wait.until(EC.frame_to_be_available_and_switch_to_it((By.XPATH, iframe_xpath)))
        except:
            pass

    # 1) Espera a MODAL alvo ficar vis√≠vel e marca o bot√£o 'Sim' com data-aim tempor√°rio
    found = js_engine.execute_js("""
        // procura modais .modal vis√≠veis e com o texto esperado
        var modals = document.querySelectorAll('div.modal.overflow');
        var alvo = null;
        for (var i=0;i<modals.length;i++){
            var m = modals[i];
            var s = getComputedStyle(m);
            var vis = (m.offsetParent!==null && s.display!=='none' && s.visibility!=='hidden' && parseFloat(s.opacity||1)>0.01);
            if(!vis) continue;
            var txt = (m.innerText||'').toLowerCase();
            if(txt.indexOf('confirme a gera√ß√£o de t√≠tulos')!==-1){
                alvo = m; break;
            }
        }
        if(!alvo) return false;

        // dentro da modal, localizar o 'Sim'
        var btn = alvo.querySelector('a.btModel.btGray.btyes') 
               || alvo.querySelector("a.btyes");
        if(!btn) return false;

        // marca para podermos clicar via CSS
        btn.setAttribute('data-aim','confirm-yes');
        return true;
    """)
    if not found:
        # n√£o achou a modal/btn ‚Äî falhar explicitamente
        return False

    # 2) Clicar com o motor robusto (dispara sequ√™ncia completa de eventos)
    try:
        js_engine.force_click("[data-aim='confirm-yes']", by_xpath=False)
    except Exception:
        # fallback por XPath direto na √¢ncora com texto 'Sim'
        js_engine.force_click("//div[contains(@class,'modal') and contains(.,'Confirme a gera√ß√£o de t√≠tulos')]//a[contains(@class,'btyes')]", by_xpath=True)

    # 3) Aguarda processamento/overlays
    js_engine.wait_ajax_complete(8)

    # 4) Validar que a modal sumiu (n√£o vis√≠vel ou removida)
    for _ in range(int(timeout*2)):  # ~timeout segundos
        closed = js_engine.execute_js("""
            var m = document.querySelector("div.modal.overflow");
            if(!m) return true; // removida do DOM
            var s = getComputedStyle(m);
            var vis = (m.offsetParent!==null && s.display!=='none' && s.visibility!=='hidden' && parseFloat(s.opacity||1)>0.01);
            if(!vis) return true; // oculta
            // se ainda h√° modal vis√≠vel, tenta verificar se a 'loadingContent' est√° ativa
            var ld = m.querySelector('.loadingContent');
            if(ld){
                var sl = getComputedStyle(ld);
                // se loading apareceu e depois some, a modal deve fechar logo
            }
            return false;
        """)
        if closed:
            # remover o atributo tempor√°rio, se ainda existir
            js_engine.execute_js("""
                var b = document.querySelector("[data-aim='confirm-yes']");
                if(b) b.removeAttribute('data-aim');
                return true;
            """)
            return True
        time.sleep(0.5)

    # √∫ltimo recurso: for√ßa evento de mouse completo + tenta esconder overlays e verifica de novo
    js_engine.execute_js("""
        var b = document.querySelector("[data-aim='confirm-yes']") 
             || document.querySelector("div.modal.overflow a.btModel.btGray.btyes");
        if(b){
            ['mouseover','mousedown','mouseup','click'].forEach(function(t){
                b.dispatchEvent(new MouseEvent(t,{bubbles:true,cancelable:true,view:window,detail:1}));
            });
            if(b.click) b.click();
        }
        document.querySelectorAll('.ui-widget-overlay,.blockUI,.modal-backdrop,[class*="overlay"]').forEach(function(o){
            o.style.display='none'; o.style.visibility='hidden'; o.style.opacity='0';
        });
        return true;
    """)
    js_engine.wait_ajax_complete(4)

    closed = js_engine.execute_js("""
        var m = document.querySelector("div.modal.overflow");
        if(!m) return true;
        var s = getComputedStyle(m);
        return !(m.offsetParent!==null && s.display!=='none' && s.visibility!=='hidden' && parseFloat(s.opacity||1)>0.01);
    """)
    return bool(closed)

def selecionar_opcao_select(xpath_select, texto_opcao):
    elemento = wait.until(EC.presence_of_element_located((By.XPATH, xpath_select)))
    driver.execute_script("arguments[0].scrollIntoView({block:'center', inline:'nearest'});", elemento)
    Select(elemento).select_by_visible_text(texto_opcao)


def selecionar_opcao_xpath(xpath, texto):
    def acao():
        select_element = wait.until(EC.presence_of_element_located((By.XPATH, xpath)))
        Select(select_element).select_by_visible_text(texto)
    return acao

def fechar_modal_com_retry(doc, js_engine, wait, max_tentativas=5, pausa=1.5):
    """Fecha o modal clicando no X at√© ele desaparecer."""
    xpath_modal = "//div[contains(@class,'modal') and contains(@style,'z-index')]"
    xpath_fechar = "(//div[contains(@class,'modal') and not(contains(@style,'display: none'))]//a[@class='fa fa-close'])[last()]"

    tentativa = 0
    while tentativa < max_tentativas:
        tentativa += 1
        log(doc, f"üß© Tentativa {tentativa} de fechar modal...")

        try:
            js_engine.force_click(xpath_fechar, by_xpath=True)
            time.sleep(pausa)

            # Verifica se ainda h√° modal vis√≠vel
            modais_visiveis = driver.find_elements(By.XPATH, xpath_modal)
            modais_ativos = [m for m in modais_visiveis if "display: none" not in m.get_attribute("style")]

            if not modais_ativos:
                log(doc, "‚úÖ Modal fechado com sucesso.")
                return True

        except Exception as e:
            log(doc, f"‚ö†Ô∏è Tentativa {tentativa} falhou: {e}")

    log(doc, "‚ùå Modal n√£o foi fechado ap√≥s todas as tentativas.")
    return False


def encontrar_mensagem_alerta():
    seletores = [
        (".alerts.salvo", "‚úÖ Mensagem de Sucesso"),
        (".alerts.alerta", "‚ö†Ô∏è Mensagem de Alerta"),
        (".alerts.erro", "‚ùå Mensagem de Erro"),
    ]

    for seletor, tipo in seletores:
        try:
            elemento = driver.find_element(By.CSS_SELECTOR, seletor)
            if elemento.is_displayed():
                log(doc, f"üì¢ {tipo}: {elemento.text}")
                return elemento
        except:
            continue

    log(doc, "‚ÑπÔ∏è Nenhuma mensagem de alerta encontrada.")
    return None

# =========================
# HELPERS DE FECHAMENTO/CLICK
# =========================
def fechar_modal_geracao_boletos(js_engine, doc):
    """
    Fecha o modal/aba de Gera√ß√£o de Boletos dentro do sistema (√≠cone 'Sair').
    """
    safe_action(doc, "Fechando modal de Gera√ß√£o de Boletos", lambda:
        js_engine.force_click(
            "//a[@class='sprites sp-fecharGrande' and @title='Sair']",
            by_xpath=True
        )
    )

def fechar_modal_gestor_financeiro(js_engine, doc):
    """
    Fecha o modal do Gestor Financeiro (bot√£o X do cabe√ßalho da janela).
    """
    safe_action(doc, "Fechando Gestor Financeiro", lambda:
        js_engine.force_click(
            "#gsFinan > div.wdTop.ui-draggable-handle > div.wdClose > a"
        )
    )

def clicar_cancelar_por_indice(js_engine, doc, indice: int = 1):
    """
    Clica no bot√£o 'Cancelar' pelo √≠ndice (1-based) quando houver m√∫ltiplos.
    """
    from selenium.common.exceptions import WebDriverException
    import time

    xpath = f"(//a[contains(@class,'btModel') and contains(@class,'btGray') and contains(@class,'btcancel')])[{indice}]"
    try:
        safe_action(doc, f"Clicando em 'Cancelar' (√≠ndice {indice})", lambda:
            js_engine.force_click(xpath, by_xpath=True)
        )
        time.sleep(0.8)
    except WebDriverException as e:
        try:
            log(doc, f"‚ö†Ô∏è Falha ao clicar em Cancelar (√≠ndice {indice}): {e}")
        except:
            pass

def confirmar_ou_detectar_relatorio(js_engine, doc=None, timeout_modal=6, timeout_aba=10, fechar_modal_gestores=True, timeout_alerta=6):
    """
    1) Se o modal de confirma√ß√£o existir, clica em 'Sim' e valida o fechamento.
    2) Se o modal N√ÉO existir, valida se abriu uma nova aba/janela (relat√≥rio).
    3) Ap√≥s o retorno ao sistema, verifica se h√° mensagens de alerta.
    Retorna True se confirmou ou abriu relat√≥rio; False se nada ocorreu.
    """
    import time
    driver = js_engine.driver

    # 0) Snapshot dos handles antes
    handles_antes = driver.window_handles[:]

    # 1) Tenta localizar rapidamente o modal de confirma√ß√£o
    try:
        modal_existe = js_engine.execute_js("""
            var modals = document.querySelectorAll('div.modal.overflow');
            for (var i=0;i<modals.length;i++){
                var m = modals[i];
                var s = getComputedStyle(m);
                var vis = (m.offsetParent!==null && s.display!=='none' && s.visibility!=='hidden' && parseFloat(s.opacity||1)>0.01);
                if(!vis) continue;
                var txt = (m.innerText||'').toLowerCase();
                if (txt.indexOf('confirme a gera√ß√£o de t√≠tulos') !== -1
                 || txt.indexOf('deseja prosseguir') !== -1
                 || txt.indexOf('deseja continuar') !== -1) { 
                    // marca o bot√£o 'Sim' para clique
                    var btn = m.querySelector('a.btModel.btGray.btyes') || m.querySelector('a.btyes');
                    if(btn){ btn.setAttribute('data-aim','confirm-yes'); return true; }
                }
            }
            return false;
        """)
    except Exception:
        modal_existe = False

    # 2) Se existe modal: confirma e valida que fechou
    if modal_existe:
        try:
            js_engine.force_click("[data-aim='confirm-yes']", by_xpath=False)
        except Exception:
            js_engine.force_click("//a[contains(@class,'btyes') and normalize-space()='Sim']", by_xpath=True)

        js_engine.wait_ajax_complete(8)

        fim = time.time() + timeout_modal
        while time.time() < fim:
            ainda_visivel = js_engine.execute_js("""
                var m = document.querySelector("div.modal.overflow");
                if(!m) return false;
                var s = getComputedStyle(m);
                return (m.offsetParent!==null && s.display!=='none' && s.visibility!=='hidden' && parseFloat(s.opacity||1)>0.01);
            """)
            if not ainda_visivel:
                break
            time.sleep(0.2)

        js_engine.execute_js("var b=document.querySelector('[data-aim=\"confirm-yes\"]'); if(b) b.removeAttribute('data-aim');")

        # ‚úÖ Ap√≥s o retorno ao sistema, verifica alertas
        alerta = encontrar_mensagem_alerta()
        if alerta:
            log(doc, f"‚ö†Ô∏è Alerta ap√≥s confirma√ß√£o do modal: {alerta}")
        else:
            log(doc, "‚úÖ Nenhum alerta detectado ap√≥s confirma√ß√£o do modal.")

        return True

    # 3) Se N√ÉO existe modal: aguarda abrir nova aba/janela (relat√≥rio)
    fim = time.time() + timeout_aba
    while time.time() < fim:
        handles_depois = driver.window_handles[:]
        if len(handles_depois) > len(handles_antes):
            try:
                nova = list(set(handles_depois) - set(handles_antes))[0]
                driver.switch_to.window(nova)
                driver.switch_to.window(handles_antes[0])
            except Exception:
                pass

            # ‚úÖ Ap√≥s o retorno ao sistema, verifica alertas
            alerta = encontrar_mensagem_alerta()
            if alerta:
                log(doc, f"‚ö†Ô∏è Alerta ap√≥s abertura do relat√≥rio: {alerta}")
            else:
                log(doc, "‚úÖ Nenhum alerta detectado ap√≥s abertura do relat√≥rio.")

            return True
        time.sleep(3)

    # 4) Nem modal nem aba nova? Prossegue
    log(doc, "‚úÖ Nenhum modal de confirma√ß√£o ou relat√≥rio detectado. Prosseguindo execu√ß√£o")
    return True



# =========================
# FLUXO PRINCIPAL COM DOIS CEN√ÅRIOS
# =========================
def confirmar_modal_e_retornar_sistema(
    js_engine,
    botao_xpath: str = "//a[@id='BtYes' and contains(@class,'btyes') and normalize-space()='Sim']",
    esperado_selector: str = "#gsFinan",
    timeout: int = 12,
    iframe_xpath: str = None,
    remove_overlays: bool = False,
    doc=None,
    verificar_nova_aba: bool = True,
    tempo_espera_aba: int = 3,
    indice_cancelar: int = 8  # usado no cen√°rio SEM gera√ß√£o de boleto
):
    """
    Confirma modal (se existir) e processa 2 cen√°rios:

    1) GEROU boleto (nova aba detectada):
       - Volta ao sistema
       - encontrar_mensagem_alerta()
       - Fecha modal de Gera√ß√£o de Boletos
       - Fecha modal de Gestor Financeiro

    2) N√ÉO GEROU:
       - encontrar_mensagem_alerta()
       - Clica em Cancelar por √≠ndice
       - Fecha modal de Gera√ß√£o de Boletos
       - Fecha modal de Gestor Financeiro
    """
    from selenium.webdriver.support.ui import WebDriverWait
    from selenium.webdriver.support import expected_conditions as EC
    from selenium.webdriver.common.by import By
    from selenium.webdriver.common.action_chains import ActionChains
    from selenium.common.exceptions import TimeoutException, StaleElementReferenceException, ElementClickInterceptedException, JavascriptException
    import time

    driver = js_engine.driver
    wait = WebDriverWait(driver, 3)

    def _log(msg):
        try:
            if doc is not None:
                log(doc, msg)
        except:
            pass

    _log("üéØ Iniciando confirma√ß√£o de modal e verifica√ß√£o de nova aba...")

    try:
        # FASE 1: PREPARA√á√ÉO / IFRAME
        try:
            driver.switch_to.default_content()
        except:
            pass

        if iframe_xpath:
            try:
                wait.until(EC.frame_to_be_available_and_switch_to_it((By.XPATH, iframe_xpath)))
                _log(f"   ‚úÖ Entrou no iframe: {iframe_xpath}")
            except Exception as e:
                _log(f"   ‚ö†Ô∏è N√£o foi poss√≠vel entrar no iframe: {e}")

        # Snapshot de abas antes do clique
        abas_iniciais = driver.window_handles
        num_abas_iniciais = len(abas_iniciais)
        _log(f"üìä Abas antes do clique: {num_abas_iniciais} | Handles: {abas_iniciais}")

        # FASE 2: CLIQUE 'SIM'
        clicked = False
        try:
            _log(f"üîé Aguardando bot√£o via XPath: {botao_xpath}")
            btn = WebDriverWait(driver, 3).until(
                EC.presence_of_element_located((By.XPATH, botao_xpath))
            )
            _log("   ‚úÖ Bot√£o localizado (presence). Tentando cliques agressivos...")

            estrategias = [
                lambda el: (WebDriverWait(driver, 3).until(EC.element_to_be_clickable((By.XPATH, botao_xpath))), el.click()),
                lambda el: (driver.execute_script("arguments[0].scrollIntoView({block:'center'});", el), el.click()),
                lambda el: driver.execute_script("arguments[0].click();", el),
                lambda el: driver.execute_script("""
                        var e = arguments[0];
                        ['mouseover','mouseenter','mousedown','mouseup','click'].forEach(function(t){
                            e.dispatchEvent(new MouseEvent(t,{bubbles:true,cancelable:true,view:window,detail:1}));
                        });
                        if(e.click) e.click();
                """, el),
                lambda el: ActionChains(driver).move_to_element(el).pause(0.1).click().perform(),
            ]

            for i, estrategia in enumerate(estrategias, 1):
                try:
                    _log(f"   ‚ñ∂Ô∏è Estrat√©gia {i}/{len(estrategias)}")
                    estrategia(btn)
                    time.sleep(0.25)
                    _log(f"   ‚úÖ Clique executado (estrat√©gia {i})")
                    clicked = True
                    break
                except (ElementClickInterceptedException, StaleElementReferenceException, JavascriptException, TimeoutException) as e:
                    _log(f"   ‚ö†Ô∏è Estrat√©gia {i} falhou: {type(e).__name__}")
                    try:
                        btn = driver.find_element(By.XPATH, botao_xpath)
                    except:
                        pass
                except Exception as e:
                    _log(f"   ‚ö†Ô∏è Erro inesperado na estrat√©gia {i}: {e}")
                    try:
                        btn = driver.find_element(By.XPATH, botao_xpath)
                    except:
                        pass

            if not clicked:
                _log("   ‚ùå N√£o foi poss√≠vel clicar no 'Sim'. Continuando com fallbacks‚Ä¶")

        except TimeoutException:
            _log("   ‚ÑπÔ∏è Bot√£o 'Sim' n√£o encontrado. Prosseguindo (sem clique).")

        # FASE 3: CEN√ÅRIOS (com/sem nova aba)
        if verificar_nova_aba and clicked:
            _log(f"‚è≥ Aguardando {tempo_espera_aba}s para detectar nova aba...")
            nova_aba_detectada, nova_handle = False, None
            t0 = time.time()

            while time.time() - t0 < tempo_espera_aba:
                time.sleep(0.5)
                abas_atuais = driver.window_handles
                if len(abas_atuais) > num_abas_iniciais:
                    nova_aba_detectada = True
                    nova_handle = [h for h in abas_atuais if h not in abas_iniciais]
                    _log(f"‚úÖ Nova aba detectada! Handles atuais: {abas_atuais} | Nova: {nova_handle}")
                    break

            if nova_aba_detectada:
                # ========== CEN√ÅRIO 1: GEROU BOLETO ==========
                _log("üü¢ Cen√°rio GEROU boleto: voltar ao sistema, tratar alerta e fechar modais.")
                try:
                    driver.switch_to.window(abas_iniciais[0])  # volta p/ aba principal
                except:
                    pass
                try:
                    driver.switch_to.default_content()
                except:
                    pass

                # 1) encontrar_mensagem_alerta()
                try:
                    encontrar_mensagem_alerta()  # assume que a fun√ß√£o j√° existe no seu projeto
                except Exception as e:
                    _log(f"‚ö†Ô∏è encontrar_mensagem_alerta() falhou: {e}")

                # 2) Fecha modal de gera√ß√£o de boletos
                try:
                    fechar_modal_geracao_boletos(js_engine, doc)
                except Exception as e:
                    _log(f"‚ö†Ô∏è Falha ao fechar modal de gera√ß√£o: {e}")

                # 3) Fecha modal do gestor financeiro
                try:
                    fechar_modal_gestor_financeiro(js_engine, doc)
                except Exception as e:
                    _log(f"‚ö†Ô∏è Falha ao fechar Gestor Financeiro: {e}")

                # limpeza opcional
                if remove_overlays:
                    try:
                        driver.switch_to.default_content()
                    except:
                        pass
                    try:
                        js_engine.execute_js("""
                            document.querySelectorAll('.ui-widget-overlay,.blockUI,.modal-backdrop,.blockScreen,.overlay,.loading,.spinner,div.modal.overflow,.loadingContent')
                              .forEach(el=>{el.style.display='none';el.style.visibility='hidden';el.style.opacity='0';el.style.pointerEvents='none';});
                            return true;
                        """)
                    except:
                        pass

                # valida retorno
                try:
                    ok = js_engine.execute_js(f"""
                        var el = document.querySelector('{esperado_selector}');
                        if(!el) return false;
                        var s = getComputedStyle(el);
                        return (el.offsetParent!==null && s.display!=='none' && s.visibility!=='hidden' && parseFloat(s.opacity||1)>0.01);
                    """)
                    if ok:
                        _log(f"‚úÖ Retorno √† tela principal detectado: {esperado_selector}")
                        return True
                except:
                    pass

                # espera extra
                _log(f"‚è≥ Aguardando tela principal (m√°x {timeout}s)‚Ä¶")
                t1 = time.time()
                while time.time() - t1 < timeout:
                    try:
                        el = driver.find_element(By.CSS_SELECTOR, esperado_selector)
                        vis = driver.execute_script("""
                            var el=arguments[0],s=getComputedStyle(el);
                            return (s.display!=='none' && s.visibility!=='hidden' && parseFloat(s.opacity||1)>0.01);
                        """, el)
                        if vis:
                            _log("‚úÖ Tela principal vis√≠vel.")
                            return True
                    except:
                        pass
                    time.sleep(0.5)

                _log("‚ùå N√£o foi poss√≠vel confirmar retorno √† tela ap√≥s gerar boleto.")
                return False

            else:
                # ========== CEN√ÅRIO 2: N√ÉO GEROU ==========
                _log("üü° Cen√°rio N√ÉO GEROU boleto: tratar alerta, cancelar por √≠ndice, fechar modais.")
                try:
                    driver.switch_to.default_content()
                except:
                    pass

                # 1) encontrar_mensagem_alerta()
                try:
                    encontrar_mensagem_alerta()
                except Exception as e:
                    _log(f"‚ö†Ô∏è encontrar_mensagem_alerta() falhou: {e}")

                # 2) Cancelar por √≠ndice
                try:
                    clicar_cancelar_por_indice(js_engine, doc, indice=indice_cancelar)
                except Exception as e:
                    _log(f"‚ö†Ô∏è Falha ao clicar em Cancelar por √≠ndice: {e}")

                # 3) Fecha modais (gera√ß√£o ‚Üí gestor)
                try:
                    fechar_modal_geracao_boletos(js_engine, doc)
                except Exception as e:
                    _log(f"‚ö†Ô∏è Falha ao fechar modal de gera√ß√£o: {e}")

                try:
                    fechar_modal_gestor_financeiro(js_engine, doc)
                except Exception as e:
                    _log(f"‚ö†Ô∏è Falha ao fechar Gestor Financeiro: {e}")

                _log("‚ùå Fluxo cancelado: relat√≥rio/aba n√£o foi gerado.")
                return False

        # (Se n√£o pediu para verificar nova aba ou n√£o clicou no 'Sim')
        _log("‚ÑπÔ∏è Sem verifica√ß√£o de nova aba ou sem clique em 'Sim'. Nenhuma a√ß√£o de cen√°rio aplicada.")
        return False

    except Exception as e:
        _log(f"‚ùå Erro em confirmar_modal_e_retornar_sistema: {e}")
        import traceback
        _log(f"   Traceback: {traceback.format_exc()}")
        return False

def clicar_ok_e_verificar_modal_confirmacao(
    max_tentativas: int = 5,
    timeout_click: int = 10,
    timeout_confirmacao: int = 5,
    scroll: bool = True,
    indice_ok: int = None  # üëà novo par√¢metro opcional
):
    """
    Se 'indice_ok' for informado, clica apenas nesse bot√£o.
    Caso contr√°rio, tenta todos at√© achar o modal de confirma√ß√£o.
    """
    def acao():
        locator_ok_base = "//a[@class='btModel btGray btok' and normalize-space()='Ok']"
        locator_modal_root = "//div[contains(@class,'modal') and contains(@class,'overflow')]"
        locator_btn_sim = f"{locator_modal_root}//a[(@id='BtYes' or contains(@class,'btyes')) and normalize-space()='Sim']"

        botoes = driver.find_elements(By.XPATH, locator_ok_base)
        if not botoes:
            raise Exception("Nenhum bot√£o 'Ok' encontrado no DOM.")

        total_botoes = len(botoes)

        # üîπ Se √≠ndice for especificado, clica somente nele
        indices = [indice_ok - 1] if indice_ok is not None else range(total_botoes)

        for idx in indices:
            log(doc, f"üéØ Tentando bot√£o Ok #{idx + 1} de {total_botoes}...")
            locator_ok = f"({locator_ok_base})[{idx + 1}]"
            elemento = driver.find_element(By.XPATH, locator_ok)

            # Estrat√©gias iguais √†s suas
            estrategias_ok = [
                lambda el: (WebDriverWait(driver, timeout_click).until(EC.element_to_be_clickable((By.XPATH, locator_ok))), el.click()),
                lambda el: (driver.execute_script("arguments[0].scrollIntoView({block:'center'});", el), el.click()),
                lambda el: driver.execute_script("arguments[0].click();", el),
                lambda el: ActionChains(driver).move_to_element(el).pause(0.1).click().perform()
            ]

            for i, estrategia in enumerate(estrategias_ok, 1):
                try:
                    log(doc, f"   ‚ñ∂Ô∏è Estrat√©gia {i}/4 de clique...")
                    estrategia(elemento)
                    log(doc, f"‚úÖ Clique executado (estrat√©gia {i}).")
                    break
                except Exception as e:
                    log(doc, f"‚ö†Ô∏è Estrat√©gia {i} falhou: {e}")

            # Ap√≥s clicar, verifica o modal
            try:
                WebDriverWait(driver, timeout_confirmacao).until(
                    EC.presence_of_element_located((By.XPATH, locator_btn_sim))
                )
                log(doc, "‚úÖ Modal de confirma√ß√£o detectado.")
                return True
            except TimeoutException:
                log(doc, "‚ÑπÔ∏è Modal de confirma√ß√£o n√£o detectado.")

        return False

    return acao


def selecionar_opcao_por_indice(
    indice_select: int,
    indice_opcao: int,
    seletor_css: str = "select",
    xpath_customizado: str = None,
    max_tentativas: int = 5,
    timeout: int = 10,
    scroll: bool = True
):
    """
    Seleciona uma op√ß√£o em um dropdown/select pelo √≠ndice do select e √≠ndice da op√ß√£o.
    
    Args:
        indice_select: √çndice do elemento <select> na p√°gina (0-based)
        indice_opcao: √çndice da <option> dentro do select (0-based)
        seletor_css: Seletor CSS para encontrar os selects (padr√£o: "select")
        xpath_customizado: XPath customizado (sobrescreve seletor_css se fornecido)
        max_tentativas: N√∫mero m√°ximo de tentativas
        timeout: Timeout em segundos
        scroll: Se deve fazer scroll at√© o elemento
    
    Returns:
        Fun√ß√£o para usar com safe_action()
    
    Exemplos:
        # Por CSS (padr√£o)
        safe_action(doc, "Selecionando op√ß√£o", 
            selecionar_opcao_por_indice(indice_select=0, indice_opcao=2))
        
        # Por CSS customizado
        safe_action(doc, "Selecionando categoria", 
            selecionar_opcao_por_indice(
                indice_select=1, 
                indice_opcao=3,
                seletor_css="select.categoria"
            ))
        
        # Por XPath customizado
        safe_action(doc, "Selecionando com XPath", 
            selecionar_opcao_por_indice(
                indice_select=0,
                indice_opcao=1,
                xpath_customizado="//div[@class='form']//select"
            ))
    """
    def acao():
        if not isinstance(indice_select, int) or indice_select < 0:
            raise ValueError(f"√çndice do select inv√°lido: {indice_select}")
        
        if not isinstance(indice_opcao, int) or indice_opcao < 0:
            raise ValueError(f"√çndice da op√ß√£o inv√°lido: {indice_opcao}")

        # Define seletor e tipo de busca
        if xpath_customizado:
            seletor = xpath_customizado
            by_type = By.XPATH
            tipo_seletor = "XPath"
        else:
            seletor = seletor_css
            by_type = By.CSS_SELECTOR
            tipo_seletor = "CSS"

        tentativa = 0
        while tentativa < max_tentativas:
            tentativa += 1
            try:
                log(doc, f"üîé Tentativa {tentativa}: Localizando selects ({tipo_seletor}: '{seletor}')...")
                
                # Coleta todos os selects
                selects = driver.find_elements(by_type, seletor)

                if not selects:
                    if tentativa < max_tentativas:
                        log(doc, f"‚ö†Ô∏è Nenhum select encontrado (tentativa {tentativa}/{max_tentativas}). Reintentando...")
                        time.sleep(1.2)
                        continue
                    raise Exception(f"Nenhum select ({tipo_seletor}: '{seletor}') foi encontrado na p√°gina.")

                if indice_select >= len(selects):
                    raise Exception(f"√çndice do select {indice_select} inv√°lido. Encontrados {len(selects)} selects.")

                # Cria localizador est√°vel
                if by_type == By.XPATH:
                    locator_xpath = f"({seletor})[{indice_select + 1}]"
                else:
                    # Converte CSS para XPath para manter consist√™ncia
                    selects_xpath = driver.find_elements(By.XPATH, f"//*[self::select]")
                    # Filtra pelo CSS
                    selects_filtrados = [s for s in selects_xpath if s in selects]
                    locator_xpath = None  # Usa elemento direto
                
                # Reaponta o elemento
                elemento_select = selects[indice_select]

                log(doc, f"üéØ Select #{indice_select} localizado (total: {len(selects)}).")

                # Scroll se necess√°rio
                if scroll:
                    driver.execute_script(
                        "arguments[0].scrollIntoView({block:'center', inline:'nearest'});", 
                        elemento_select
                    )
                    time.sleep(0.3)

                # Aguarda elemento estar vis√≠vel
                wait.until(EC.visibility_of(elemento_select))

                # Cria objeto Select do Selenium
                select_obj = Select(elemento_select)
                
                # Verifica se o √≠ndice da op√ß√£o √© v√°lido
                opcoes = select_obj.options
                if indice_opcao >= len(opcoes):
                    raise Exception(f"√çndice da op√ß√£o {indice_opcao} inv√°lido. O select #{indice_select} tem {len(opcoes)} op√ß√µes.")

                log(doc, f"   üìã Select tem {len(opcoes)} op√ß√µes")
                log(doc, f"   üéØ Selecionando op√ß√£o #{indice_opcao}: '{opcoes[indice_opcao].text}'")

                # Estrat√©gias de sele√ß√£o
                estrategias = [
                    # 1) Select nativo do Selenium
                    lambda: select_obj.select_by_index(indice_opcao),
                    
                    # 2) JavaScript direto
                    lambda: driver.execute_script(f"""
                        var select = arguments[0];
                        select.selectedIndex = {indice_opcao};
                        select.dispatchEvent(new Event('change', {{bubbles: true}}));
                        select.dispatchEvent(new Event('blur', {{bubbles: true}}));
                    """, elemento_select),
                    
                    # 3) JavaScript com value
                    lambda: driver.execute_script(f"""
                        var select = arguments[0];
                        var opcao = select.options[{indice_opcao}];
                        select.value = opcao.value;
                        ['change', 'blur', 'input'].forEach(function(evt) {{
                            select.dispatchEvent(new Event(evt, {{bubbles: true}}));
                        }});
                    """, elemento_select),
                    
                    # 4) JavaScript com trigger jQuery
                    lambda: driver.execute_script(f"""
                        var select = arguments[0];
                        select.selectedIndex = {indice_opcao};
                        if (typeof jQuery !== 'undefined') {{
                            jQuery(select).trigger('change');
                        }} else {{
                            select.dispatchEvent(new Event('change', {{bubbles: true}}));
                        }}
                    """, elemento_select),
                    
                    # 5) Clique direto na op√ß√£o (for√ßa bruta)
                    lambda: (
                        driver.execute_script("""
                            var select = arguments[0];
                            var opcao = select.options[arguments[1]];
                            opcao.selected = true;
                            select.dispatchEvent(new Event('change', {bubbles: true}));
                            select.dispatchEvent(new Event('blur', {bubbles: true}));
                        """, elemento_select, indice_opcao)
                    )
                ]

                for i, estrategia in enumerate(estrategias, 1):
                    try:
                        log(doc, f"   ‚ñ∂Ô∏è Estrat√©gia {i} de sele√ß√£o...")
                        estrategia()
                        time.sleep(0.3)
                        
                        # Valida se a sele√ß√£o funcionou
                        select_obj_validacao = Select(elemento_select)
                        selecionado = select_obj_validacao.first_selected_option
                        
                        if selecionado and selecionado.text == opcoes[indice_opcao].text:
                            log(doc, f"‚úÖ Op√ß√£o selecionada com sucesso (estrat√©gia {i}): '{selecionado.text}'")
                            return True
                        else:
                            log(doc, f"   ‚ö†Ô∏è Estrat√©gia {i} n√£o confirmou sele√ß√£o")
                            continue
                            
                    except (ElementClickInterceptedException, StaleElementReferenceException, 
                           JavascriptException, TimeoutException) as e:
                        log(doc, f"   ‚ö†Ô∏è Estrat√©gia {i} falhou: {e}")
                        # Tenta re-obter o elemento
                        try:
                            selects = driver.find_elements(by_type, seletor)
                            elemento_select = selects[indice_select]
                        except Exception:
                            pass
                        continue

                if tentativa < max_tentativas:
                    log(doc, f"‚ö†Ô∏è Tentativa {tentativa} n√£o conseguiu selecionar. Reintentando em 1.2s...")
                    time.sleep(1.2)
                    continue

            except Exception as e:
                if tentativa < max_tentativas:
                    log(doc, f"‚ö†Ô∏è Erro na tentativa {tentativa}: {e}. Reintentando em 1.2s...")
                    time.sleep(1.2)
                    continue
                raise

        raise Exception(f"Falha ao selecionar op√ß√£o ap√≥s {max_tentativas} tentativas.")

    return acao




def fechar_abas_boletos(js_engine, doc):
    """
    Fecha as abas de Gera√ß√£o de Boletos e Gestor Financeiro.
    Retorna True se ambas foram fechadas com sucesso, False caso contr√°rio.
    """
    sucesso_boletos = False
    sucesso_gestor = False
    
    # Fechar aba de Gera√ß√£o de Boletos
    try:
        safe_action(doc, "Fechando aba de Gera√ß√£o de Boletos", lambda:
                        js_engine.force_click(
                            "//a[@class='sprites sp-fecharGrande' and @title='Sair']",
                            by_xpath=True
                        )
                    )
        time.sleep(1)
        sucesso_boletos = True
        log(doc, "‚úÖ Aba de Gera√ß√£o de Boletos fechada com sucesso.")
    except Exception as e:
        log(doc, f"‚ö†Ô∏è Falha ao fechar aba de Gera√ß√£o de Boletos: {e}")

    # Fechar Gestor Financeiro
    try:
        safe_action(doc, "Fechando Gestor Financeiro", lambda:
            js_engine.force_click(
                "#gsFinan > div.wdTop.ui-draggable-handle > div.wdClose > a"
            )
        )
        sucesso_gestor = True
        log(doc, "‚úÖ Gestor Financeiro fechado com sucesso.")
    except Exception as e:
        log(doc, f"‚ö†Ô∏è Erro ao fechar Gestor Financeiro: {e}")
    
    return sucesso_boletos and sucesso_gestor




# ==== EXECU√á√ÉO DO TESTE ====

def executar_teste():
    """Execu√ß√£o principal do teste com JS for√ßado"""
    global driver, wait, doc
    
    try:
        # Inicializa driver
        if not inicializar_driver():
            return False
        
        # Cria engine JS for√ßado
        js_engine = JSForceEngine(driver, wait, doc)
        lov_handler = LOVHandler(js_engine)
        
        # ===== LOGIN =====
        safe_action(doc, "Acessando sistema", lambda: driver.get(URL))
        
        def fazer_login():
            wait.until(EC.presence_of_element_located((By.ID, "j_id15:email"))).send_keys(LOGIN_EMAIL)
            wait.until(EC.presence_of_element_located((By.ID, "j_id15:senha"))).send_keys(LOGIN_PASSWORD, Keys.ENTER)
            time.sleep(5)

        safe_action(doc, "Realizando login", fazer_login)
        
        # ===== MENU =====
        def abrir_menu():
            driver.execute_script("document.body.style.zoom='90%'")
            driver.find_element(By.TAG_NAME, "body").send_keys(Keys.F3)
            time.sleep(2)
        
        safe_action(doc, "Abrindo menu (F3)", abrir_menu)
        
        # ===== GESTOR FINANCEIRO =====
        safe_action(doc, "Acessando Gestor Financeiro", lambda:
            js_engine.force_click('/html/body/div[15]/ul/li[2]/img', by_xpath=True)
        )
        
        time.sleep(3)
        
        # ===== GERA√á√ÉO DE BOLETOS =====
        safe_action(doc, "Clicando em Impress√£o de T√≠tulos", lambda:
            js_engine.force_click(
                '#gsFinan > div.wdTelas > div > ul > li:nth-child(8) > a > span'
            )
        )
        
        time.sleep(5)
        
        # ===== CONTRATANTE/TITULAR =====
        safe_action(doc, "Selecionando Pessoa", lambda:
            lov_handler.open_and_select(
                btn_index=0,
                search_text="ESTER SARAH BRUNA GON√áALVES",
                result_text="ESTER SARAH BRUNA GON√áALVES"
            )
        )

        safe_action(doc, "Selecionando Op√ß√£o: 'Em aberto'", lambda:
            selecionar_opcao_select("//select[@class='tipoParcela']", "Em aberto")
        )

        # ===== BUSCAR =====
        safe_action(doc, "Clicando em Pesquisar", lambda:
            js_engine.force_click(
                "//a[@class='btModel btGreen btPesquisar boxsize' and normalize-space()='Pesquisar']",
                by_xpath=True
            )
        )

        time.sleep(20)
        
        # ===== VALIDA√á√ÉO DO RESULTADO =====
        executar_fluxo_boletos(js_engine, doc)


        # Fechar as abas
        resultado = fechar_abas_boletos(js_engine, doc)

        if resultado:
            log(doc, "‚úÖ Todas as abas foram fechadas com sucesso.")
        else:
            log(doc, "‚ö†Ô∏è Algumas abas n√£o puderam ser fechadas, mas o fluxo continua.")
                
            log(doc, "üéâ Teste conclu√≠do com sucesso!")
            return True
        
    except Exception as e:
        log(doc, f"‚ùå ERRO FATAL: {e}")
        take_screenshot(driver, doc, "erro_fatal")
        return False


# ==== MAIN ====

def main():
    """Ponto de entrada principal"""
    global doc
    
    try:
        log(doc, "üöÄ Iniciando teste de Gera√ß√£o de T√≠tulos")
        log(doc, "=" * 70)
        
        sucesso = executar_teste()
        
        log(doc, "=" * 70)
        if sucesso:
            log(doc, "‚úÖ TESTE EXECUTADO COM SUCESSO!")
        else:
            log(doc, "‚ùå TESTE FINALIZADO COM ERROS")
            
    except Exception as e:
        log(doc, f"‚ùå Erro na execu√ß√£o principal: {e}")
        
    finally:
        finalizar_relatorio()


if __name__ == "__main__":
    main()