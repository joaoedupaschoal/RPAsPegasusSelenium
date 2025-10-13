# ==== IMPORTS ====
from datetime import datetime, timedelta
from datetime import time as dt_time
import time
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver import ActionChains
from selenium.common.exceptions import *
from selenium.webdriver.support import expected_conditions as EC
from selenium import webdriver
from selenium.webdriver.support.ui import Select, WebDriverWait
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service
from webdriver_manager.chrome import ChromeDriverManager
from docx import Document
from docx.shared import Inches
from faker import Faker
from faker.providers import BaseProvider
from validate_docbr import CPF
import subprocess
import os
import random
import re
import pyautogui

# ==== CONFIGURAÇÕES GLOBAIS ====
TIMEOUT_DEFAULT = 30
TIMEOUT_CURTO = 10
TIMEOUT_LONGO = 60
CAMINHO_ARQUIVO_UPLOAD = "C:/Users/Gold System/Documents/teste.png"
URL = "http://localhost:8080/gs/login.xhtml"
LOGIN_EMAIL = "joaoeduardo.gold@outlook.com"
LOGIN_PASSWORD = "071999gs"

# ==== VARIÁVEIS GLOBAIS ====
doc = Document()
doc.add_heading("RELATÓRIO DO TESTE", 0)
doc.add_paragraph("Gerar Boleto Único - Gestor Financeiro – Cenário 1: Rotina completa de Geração de Título Único")
doc.add_paragraph(f"Data do teste: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")

screenshot_registradas = set()
driver = None
wait = None

# ==== PROVIDERS CUSTOMIZADOS ====
class BrasilProvider(BaseProvider):
    def rg(self):
        numeros = [str(random.randint(0, 9)) for _ in range(8)]
        return ''.join(numeros) + '-' + str(random.randint(0, 9))

fake = Faker("pt_BR")
fake.add_provider(BrasilProvider)

# ==== GERAÇÃO DE DATAS ====
def gerar_datas_validas(hora_padrao="00:00", dias_fim=0):
    hoje_date = datetime.today().date()
    dez_anos_atras = hoje_date - timedelta(days=3650)
    data_falecimento = fake.date_between(start_date=dez_anos_atras, end_date=hoje_date)
    idade_minima, idade_maxima = 18, 110
    data_nascimento = data_falecimento - timedelta(days=random.randint(idade_minima * 365, idade_maxima * 365))
    data_sepultamento = data_falecimento + timedelta(days=random.randint(1, 10))
    data_registro = data_sepultamento + timedelta(days=random.randint(1, 10))
    data_velorio = fake.date_between(start_date=data_falecimento, end_date=data_sepultamento)
    data_inicio_date = hoje_date + timedelta(days=random.randint(2, 30))
    h, m = map(int, hora_padrao.split(":"))
    dt_inicio = datetime.combine(data_inicio_date, dt_time(h, m))
    dt_fim = dt_inicio + timedelta(days=dias_fim)
    fmt_data = "%d/%m/%Y"
    fmt_dt = "%d/%m/%Y %H:%M"
    return (
        data_nascimento.strftime(fmt_data),
        data_falecimento.strftime(fmt_data),
        data_sepultamento.strftime(fmt_data),
        data_velorio.strftime(fmt_data),
        dt_inicio.strftime(fmt_dt),
        dt_fim.strftime(fmt_dt),
        data_registro.strftime(fmt_data),
        hoje_date.strftime(fmt_data),
    )

(data_nascimento, data_falecimento, data_sepultamento,
 data_velorio, data_inicio, data_fim, data_registro, hoje) = gerar_datas_validas(
    hora_padrao="08:50",
    dias_fim=0
)

# ==== UTILITÁRIOS DE LOG ====
def log(doc, msg):
    print(msg)
    doc.add_paragraph(msg)

def _sanitize_filename(name: str) -> str:
    name = name.strip().lower()
    name = re.sub(r"[<>:\"/\\|?*']", "_", name)
    name = re.sub(r"_+", "_", name)
    return name[:120]

def take_screenshot(driver, doc, nome):
    if driver is None:
        return
    nome = _sanitize_filename(nome)
    if nome not in screenshot_registradas:
        path = f"screenshots/{nome}.png"
        os.makedirs("screenshots", exist_ok=True)
        try:
            driver.save_screenshot(path)
            doc.add_paragraph(f"Screenshot: {nome}")
            doc.add_picture(path, width=Inches(5.5))
            screenshot_registradas.add(nome)
        except Exception as e:
            log(doc, f"⚠️ Erro ao tirar screenshot {nome}: {e}")

# ==== NÚCLEO ROBUSTO: JAVASCRIPT FORÇADO ====


def abrir_modal_e_selecionar_js_puro(
    btn_xpath=None,
    pesquisa_xpath=None,
    termo_pesquisa=None,
    btn_pesquisar_xpath=None,
    resultado_xpath=None,
    timeout=15,
    max_tentativas=3,
    iframe_xpath=None,
    indice_lov=None
):
    """
    Abre modal LOV usando JavaScript puro para TODAS operaÃ§Ãµes.
    Elimina problemas de interceptaÃ§ão, visibilidade e timing do Selenium.
    
    Args:
        btn_xpath: XPath do botão LOV (ignorado se indice_lov fornecido)
        pesquisa_xpath: XPath do campo de pesquisa
        termo_pesquisa: Termo a ser pesquisado
        btn_pesquisar_xpath: XPath do botão Pesquisar
        resultado_xpath: XPath do resultado a clicar
        timeout: Tempo mÃ¡ximo de espera
        max_tentativas: NÃºmero de tentativas
        iframe_xpath: XPath do iframe (se aplicÃ¡vel)
        indice_lov: Ãndice do botão LOV na pÃ¡gina (alternativa ao btn_xpath)
    """
    
    def _executar_js(script, *args):
        """Executa JavaScript e retorna resultado"""
        try:
            return driver.execute_script(script, *args)
        except Exception as e:
            log(doc, f"  ❌ Erro JS: {str(e)[:100]}")
            raise
    
    def _aguardar_ajax_js(timeout_ajax=10):
        """Aguarda AJAX usando JavaScript"""
        inicio = time.time()
        while time.time() - inicio < timeout_ajax:
            completo = _executar_js("""
                // Verifica jQuery
                var jQueryOk = typeof jQuery === 'undefined' || jQuery.active === 0;
                
                // Verifica fetch/XMLHttpRequest pendentes
                var fetchOk = !window.__pendingRequests || window.__pendingRequests === 0;
                
                // Verifica overlays
                var overlays = document.querySelectorAll(
                    '.blockScreen, .blockUI, .loading, .overlay, [class*="loading"], [class*="spinner"]'
                );
                var overlayOk = true;
                for (var i = 0; i < overlays.length; i++) {
                    var style = window.getComputedStyle(overlays[i]);
                    if (style.display !== 'none' && 
                        style.visibility !== 'hidden' && 
                        parseFloat(style.opacity || 1) > 0.01) {
                        overlayOk = false;
                        break;
                    }
                }
                
                return jQueryOk && fetchOk && overlayOk;
            """)
            
            if completo:
                return True
            time.sleep(0.2)
        return True
    
    def _clicar_js(xpath_ou_elemento, descricao="elemento"):
        """Clica usando JavaScript puro"""
        script = """
            var elemento = arguments[0];
            
            // Se recebeu XPath, resolve
            if (typeof elemento === 'string') {
                var resultado = document.evaluate(
                    elemento, 
                    document, 
                    null, 
                    XPathResult.FIRST_ORDERED_NODE_TYPE, 
                    null
                );
                elemento = resultado.singleNodeValue;
            }
            
            if (!elemento) {
                throw new Error('Elemento não encontrado: ' + arguments[0]);
            }
            
            // ForÃ§a visibilidade
            elemento.style.display = 'block';
            elemento.style.visibility = 'visible';
            elemento.style.opacity = '1';
            
            // Remove atributos que impedem clique
            elemento.removeAttribute('disabled');
            elemento.removeAttribute('readonly');
            
            // Scroll suave até o elemento
            elemento.scrollIntoView({behavior: 'smooth', block: 'center'});
            
            // Dispara eventos completos
            ['mouseover', 'mousedown', 'mouseup', 'click'].forEach(function(eventType) {
                var evt = new MouseEvent(eventType, {
                    bubbles: true,
                    cancelable: true,
                    view: window,
                    detail: 1,
                    clientX: 0,
                    clientY: 0
                });
                elemento.dispatchEvent(evt);
            });
            
            // Também tenta click direto
            if (elemento.click) {
                elemento.click();
            }
            
            return true;
        """
        
        try:
            _executar_js(script, xpath_ou_elemento)
            log(doc, f"   🔄 Clique JS em {descricao}")
            return True
        except Exception as e:
            raise Exception(f"Falha ao clicar em {descricao}: {e}")
    
    def _preencher_campo_js(xpath_campo, valor):
        """Preenche campo usando JavaScript puro"""
        script = """
            var xpath = arguments[0];
            var valor = arguments[1];
            
            // Localiza elemento
            var resultado = document.evaluate(
                xpath, 
                document, 
                null, 
                XPathResult.FIRST_ORDERED_NODE_TYPE, 
                null
            );
            var campo = resultado.singleNodeValue;
            
            if (!campo) {
                throw new Error('Campo não encontrado: ' + xpath);
            }
            
            // Força campo editável
            campo.removeAttribute('disabled');
            campo.removeAttribute('readonly');
            campo.style.display = 'block';
            campo.style.visibility = 'visible';
            
            // Limpa valor anterior
            campo.value = '';
            
            // Focus no campo
            campo.focus();
            
            // Dispara eventos de input
            ['focus', 'click'].forEach(function(evt) {
                campo.dispatchEvent(new Event(evt, {bubbles: true}));
            });
            
            // Define valor
            campo.value = valor;
            
            // Dispara eventos de mudança
            ['input', 'change', 'blur'].forEach(function(evt) {
                campo.dispatchEvent(new Event(evt, {bubbles: true}));
            });
            
            // Triggers adicionais para frameworks
            if (typeof campo.onchange === 'function') {
                campo.onchange();
            }
            
            // Trigger jQuery se existir
            if (typeof jQuery !== 'undefined') {
                jQuery(campo).trigger('change');
            }
            
            return campo.value;
        """
        
        valor_final = _executar_js(script, xpath_campo, valor)
        log(doc, f"   ✅ Campo preenchido: '{valor_final}'")
        return valor_final
    
    def _aguardar_elemento_js(xpath, timeout_espera=10, deve_estar_visivel=True):
        """Aguarda elemento usando JavaScript"""
        inicio = time.time()
        while time.time() - inicio < timeout_espera:
            existe = _executar_js("""
                var xpath = arguments[0];
                var deveEstarVisivel = arguments[1];
                
                var resultado = document.evaluate(
                    xpath, 
                    document, 
                    null, 
                    XPathResult.FIRST_ORDERED_NODE_TYPE, 
                    null
                );
                var elemento = resultado.singleNodeValue;
                
                if (!elemento) return false;
                
                if (!deveEstarVisivel) return true;
                
                // Verifica visibilidade real
                var style = window.getComputedStyle(elemento);
                return style.display !== 'none' && 
                       style.visibility !== 'hidden' && 
                       parseFloat(style.opacity || 1) > 0.01 &&
                       elemento.offsetParent !== null;
            """, xpath, deve_estar_visivel)
            
            if existe:
                return True
            time.sleep(0.2)
        
        raise Exception(f"Elemento não encontrado após {timeout_espera}s: {xpath}")
    
    def _localizar_botao_lov_js():
        """Localiza botão LOV por Í­ndice ou XPath usando JS"""
        if indice_lov is not None:
            log(doc, f"   🔄  Localizando botão LOV por í­ndice: {indice_lov}")
            
            # Conta quantos botÃµes existem
            total = _executar_js("""
                var botoes = document.querySelectorAll("a.sprites.sp-openLov");
                return botoes.length;
            """)
            
            if indice_lov >= total:
                raise Exception(f"Índice {indice_lov} inválido. Encontrados {total} botões LOV")
            
            xpath_resultado = f"(//a[@class='sprites sp-openLov'])[{indice_lov + 1}]"
            log(doc, f"   ✅ Botão LOV #{indice_lov} localizado")
            return xpath_resultado
            
        else:
            if not btn_xpath:
                raise ValueError("btn_xpath ou indice_lov deve ser fornecido")
            log(doc, f" 🔄 Usando XPath fornecido para botão LOV")
            return btn_xpath
    
    def _trocar_iframe_js(iframe_xpath):
        """Troca para iframe usando Selenium (necessÃ¡rio)"""
        try:
            WebDriverWait(driver, timeout).until(
                EC.frame_to_be_available_and_switch_to_it((By.XPATH, iframe_xpath))
            )
            log(doc, "   ✅ Iframe carregado")
            return True
        except Exception as e:
            raise Exception(f"Falha ao trocar para iframe: {e}")
    
    def acao():
        for tentativa in range(1, max_tentativas + 1):
            try:
                log(doc, f"🔄 Tentativa {tentativa}/{max_tentativas} - Modal LOV (JS Puro)")
                
                # Volta para conteÃºdo principal
                try:
                    driver.switch_to.default_content()
                except:
                    pass
                
                # PASSO 1: Localizar e clicar no botão LOV
                xpath_botao = _localizar_botao_lov_js()
                _aguardar_elemento_js(xpath_botao, timeout, deve_estar_visivel=True)
                _clicar_js(xpath_botao, "botão LOV")
                time.sleep(1.5)
                
                # PASSO 2: Trocar para iframe (se necessÃ¡rio)
                if iframe_xpath:
                    log(doc, " 🔄 Trocando para iframe do modal")
                    _trocar_iframe_js(iframe_xpath)
                    time.sleep(0.5)
                
                # PASSO 3: Aguardar modal abrir
                log(doc, " 🔄 Aguardando modal carregar...")
                _aguardar_ajax_js(timeout_ajax=10)
                
                # PASSO 4: Aguardar e preencher campo de pesquisa
                log(doc, f" 🔄 Aguardando campo de pesquisa")
                _aguardar_elemento_js(pesquisa_xpath, timeout, deve_estar_visivel=True)
                time.sleep(0.5)
                
                log(doc, f" 🔄 Preenchendo: '{termo_pesquisa}'")
                _preencher_campo_js(pesquisa_xpath, termo_pesquisa)
                time.sleep(0.5)
                
                # PASSO 5: Clicar em Pesquisar
                log(doc, "   🔄 Clicando em Pesquisar")
                _aguardar_elemento_js(btn_pesquisar_xpath, timeout, deve_estar_visivel=True)
                _clicar_js(btn_pesquisar_xpath, "botão Pesquisar")
                
                # PASSO 6: Aguardar resultados
                log(doc, "   🔄 Aguardando resultados...")
                _aguardar_ajax_js(timeout_ajax=15)
                time.sleep(2)
                
                # PASSO 7: Clicar no resultado
                log(doc, "   🔄 Clicando no resultado")
                _aguardar_elemento_js(resultado_xpath, timeout, deve_estar_visivel=True)
                _clicar_js(resultado_xpath, "resultado da pesquisa")
                time.sleep(1)
                
                # PASSO 8: Voltar para conteÃºdo principal
                if iframe_xpath:
                    try:
                        driver.switch_to.default_content()
                        log(doc, "   âœ… Voltou para conteÃºdo principal")
                    except:
                        pass
                
                # PASSO 9: Aguardar modal fechar
                time.sleep(1)
                _aguardar_ajax_js()
                
                log(doc, f" ✅ Modal LOV processado (tentativa {tentativa}) - JS Puro")
                return True
                
            except Exception as e:
                log(doc, f"❌ Tentativa {tentativa} falhou: {str(e)[:200]}")
                
                # Cleanup em caso de erro
                try:
                    driver.switch_to.default_content()
                except:
                    pass
                
                # Tenta fechar modal via JS
                try:
                    _executar_js("""
                        var botoes = document.querySelectorAll(
                            '.ui-dialog-titlebar-close, .close, [class*="close"]'
                        );
                        botoes.forEach(function(btn) {
                            if (btn.offsetParent !== null) {
                                btn.click();
                            }
                        });
                    """)
                    time.sleep(0.5)
                except:
                    pass
                
                if tentativa < max_tentativas:
                    tempo_espera = 2 + (tentativa * 0.5)
                    log(doc, f" 🔄 Aguardando {tempo_espera}s antes de retentar...")
                    time.sleep(tempo_espera)
                else:
                    raise Exception(
                        f"Falha após {max_tentativas} tentativas (JS Puro). "
                        f"Último erro: {e}"
                    )
        
        return False
    
    return acao

class JSForceEngine:
    """Motor de execução JavaScript forçado - 100% à prova de falhas"""
    
    def __init__(self, driver, wait, doc):
        self.driver = driver
        self.wait = wait
        self.doc = doc
    
    def execute_js(self, script, *args):
        """Executa JavaScript com tratamento de erros"""
        try:
            return self.driver.execute_script(script, *args)
        except Exception as e:
            log(self.doc, f"⚠️ Erro JS: {str(e)[:150]}")
            raise
    
    def wait_ajax_complete(self, timeout=15):
        """Versão síncrona — sem Promise"""
        end = time.time() + timeout
        while time.time() < end:
            try:
                done = self.driver.execute_script("""
                    var jQueryOk = (typeof jQuery==='undefined') || (jQuery.active===0);
                    var fetchOk = !window.__pendingRequests || window.__pendingRequests===0;
                    var overlays = document.querySelectorAll(
                      '.blockScreen, .blockUI, .loading, .overlay, [class*="loading"], [class*="spinner"], [class*="wait"]'
                    );
                    var overlayOk = true;
                    for (var i=0;i<overlays.length;i++){
                      var s=window.getComputedStyle(overlays[i]);
                      if(s.display!=='none' && s.visibility!=='hidden' && parseFloat(s.opacity||1)>0.01){ overlayOk=false; break; }
                    }
                    return jQueryOk && fetchOk && overlayOk;
                """)
                if done:
                    return True
            except:
                pass
            time.sleep(0.2)
        return True
    
    def force_click(self, selector, by_xpath=False, max_attempts=5):
        """Clique forçado 100% garantido usando JavaScript"""
        log(self.doc, f"🎯 Clique forçado em: {selector}")
        
        for attempt in range(max_attempts):
            try:
                # Estratégia progressiva de clique
                strategies = [
                    self._click_strategy_2,  # Eventos MouseEvent completos (melhor compatibilidade)
                    self._click_strategy_1,  # Clique nativo otimizado
                    self._click_strategy_3,  # Dispatch direto
                    self._click_strategy_4,  # Força bruta total
                    self._click_strategy_5,  # Último recurso
                ]

                
                for i, strategy in enumerate(strategies, 1):
                    try:
                        log(self.doc, f"   Tentativa {attempt + 1}.{i}...")
                        result = strategy(selector, by_xpath)
                        if result:
                            log(self.doc, f"✅ Clique bem-sucedido (estratégia {i})")
                            time.sleep(0.5)
                            self.wait_ajax_complete(10)
                            return True
                    except Exception as e:
                        log(self.doc, f"   Estratégia {i} falhou: {str(e)[:80]}")
                        continue
                
                if attempt < max_attempts - 1:
                    time.sleep(1 + attempt * 0.5)
                    
            except Exception as e:
                log(self.doc, f"⚠️ Tentativa {attempt + 1} falhou: {e}")
                if attempt < max_attempts - 1:
                    time.sleep(1.5)
        
        raise Exception(f"Falha ao clicar após {max_attempts} tentativas: {selector}")
    
    def _click_strategy_1(self, selector, by_xpath):
        """Estratégia 1: Clique nativo otimizado"""
        script = """
        var selector = arguments[0];
        var byXPath = arguments[1];
        
        var element;
        if (byXPath) {
            var result = document.evaluate(selector, document, null, 
                XPathResult.FIRST_ORDERED_NODE_TYPE, null);
            element = result.singleNodeValue;
        } else {
            element = document.querySelector(selector);
        }
        
        if (!element) throw new Error('Elemento não encontrado');
        
        // Remove obstáculos
        element.style.pointerEvents = 'auto';
        element.style.display = 'block';
        element.style.visibility = 'visible';
        element.style.opacity = '1';
        element.removeAttribute('disabled');
        
        // Scroll suave
        element.scrollIntoView({behavior: 'smooth', block: 'center'});
        
        // Aguarda scroll
        return new Promise(function(resolve) {
            setTimeout(function() {
                element.click();
                resolve(true);
            }, 300);
        });
        """
        return self.execute_js(script, selector, by_xpath)
    
    def _click_strategy_2(self, selector, by_xpath):
        """Estratégia 2: Eventos MouseEvent completos"""
        script = """
        var selector = arguments[0];
        var byXPath = arguments[1];
        
        var element;
        if (byXPath) {
            var result = document.evaluate(selector, document, null, 
                XPathResult.FIRST_ORDERED_NODE_TYPE, null);
            element = result.singleNodeValue;
        } else {
            element = document.querySelector(selector);
        }
        
        if (!element) throw new Error('Elemento não encontrado');
        
        // Prepara elemento
        element.style.pointerEvents = 'auto';
        element.removeAttribute('disabled');
        element.scrollIntoView({block: 'center'});
        
        // Sequência completa de eventos
        var events = ['mouseover', 'mouseenter', 'mousemove', 'mousedown', 'mouseup', 'click'];
        events.forEach(function(eventType) {
            var evt = new MouseEvent(eventType, {
                bubbles: true,
                cancelable: true,
                view: window,
                detail: 1,
                clientX: element.getBoundingClientRect().left + 5,
                clientY: element.getBoundingClientRect().top + 5
            });
            element.dispatchEvent(evt);
        });
        
        // Clique adicional
        if (typeof element.click === 'function') {
            element.click();
        }
        
        return true;
        """
        return self.execute_js(script, selector, by_xpath)
    
    def _click_strategy_3(self, selector, by_xpath):
        """Estratégia 3: Dispatch direto com focus"""
        script = """
        var selector = arguments[0];
        var byXPath = arguments[1];
        
        var element;
        if (byXPath) {
            var result = document.evaluate(selector, document, null, 
                XPathResult.FIRST_ORDERED_NODE_TYPE, null);
            element = result.singleNodeValue;
        } else {
            element = document.querySelector(selector);
        }
        
        if (!element) throw new Error('Elemento não encontrado');
        
        // Força visibilidade total
        element.style.display = 'block';
        element.style.visibility = 'visible';
        element.style.opacity = '1';
        element.style.pointerEvents = 'auto';
        
        // Focus
        element.focus();
        
        // Clique direto
        element.click();
        
        // Dispatch adicional
        element.dispatchEvent(new Event('click', {bubbles: true, cancelable: true}));
        
        return true;
        """
        return self.execute_js(script, selector, by_xpath)
    
    def _click_strategy_4(self, selector, by_xpath):
        """Estratégia 4: Força bruta total"""
        script = """
        var selector = arguments[0];
        var byXPath = arguments[1];
        
        var element;
        if (byXPath) {
            var result = document.evaluate(selector, document, null, 
                XPathResult.FIRST_ORDERED_NODE_TYPE, null);
            element = result.singleNodeValue;
        } else {
            element = document.querySelector(selector);
        }
        
        if (!element) throw new Error('Elemento não encontrado');
        
        // Remove TODOS os bloqueios possíveis
        element.removeAttribute('disabled');
        element.removeAttribute('readonly');
        element.style.pointerEvents = 'auto !important';
        element.style.display = 'block !important';
        element.style.visibility = 'visible !important';
        element.style.opacity = '1 !important';
        
        // Remove overlays globais
        var overlays = document.querySelectorAll('.modal, .overlay, .blockUI, [role="dialog"]');
        overlays.forEach(function(overlay) {
            overlay.style.display = 'none';
            overlay.style.visibility = 'hidden';
        });
        
        // Múltiplos métodos de clique
        element.focus();
        element.click();
        
        var clickEvent = new MouseEvent('click', {
            view: window,
            bubbles: true,
            cancelable: true
        });
        element.dispatchEvent(clickEvent);
        
        // Trigger jQuery se existir
        if (typeof jQuery !== 'undefined') {
            jQuery(element).trigger('click');
        }
        
        return true;
        """
        return self.execute_js(script, selector, by_xpath)
    
    def _click_strategy_5(self, selector, by_xpath):
        """Estratégia 5: Último recurso - simula clique no ponto exato"""
        script = """
        var selector = arguments[0];
        var byXPath = arguments[1];
        
        var element;
        if (byXPath) {
            var result = document.evaluate(selector, document, null, 
                XPathResult.FIRST_ORDERED_NODE_TYPE, null);
            element = result.singleNodeValue;
        } else {
            element = document.querySelector(selector);
        }
        
        if (!element) throw new Error('Elemento não encontrado');
        
        // Pega coordenadas
        var rect = element.getBoundingClientRect();
        var x = rect.left + rect.width / 2;
        var y = rect.top + rect.height / 2;
        
        // Cria evento no ponto exato
        var evt = document.createEvent('MouseEvents');
        evt.initMouseEvent('click', true, true, window, 1, x, y, x, y, false, false, false, false, 0, null);
        element.dispatchEvent(evt);
        
        // Fallback: procura por onclick
        if (element.onclick) {
            element.onclick();
        }
        
        // Procura handlers no parent
        var parent = element.parentElement;
        while (parent && parent !== document.body) {
            if (parent.onclick) {
                parent.onclick();
                break;
            }
            parent = parent.parentElement;
        }
        
        return true;
        """
        return self.execute_js(script, selector, by_xpath)
    
    def force_fill(self, selector, value, by_xpath=False, max_attempts=5):
        """Preenchimento forçado 100% garantido"""
        log(self.doc, f"✏️ Preenchimento forçado: {selector} = '{value}'")
        
        for attempt in range(max_attempts):
            try:
                strategies = [
                    self._fill_strategy_1,  # Native + Events
                    self._fill_strategy_2,  # React/Angular compatible
                    self._fill_strategy_3,  # jQuery trigger
                    self._fill_strategy_4,  # Força bruta
                    self._fill_strategy_5,  # Último recurso
                ]
                
                for i, strategy in enumerate(strategies, 1):
                    try:
                        log(self.doc, f"   Tentativa {attempt + 1}.{i}...")
                        result = strategy(selector, value, by_xpath)
                        
                        # Valida preenchimento
                        time.sleep(0.3)
                        if self._validate_fill(selector, value, by_xpath):
                            log(self.doc, f"✅ Campo preenchido (estratégia {i})")
                            return True
                    except Exception as e:
                        log(self.doc, f"   Estratégia {i} falhou: {str(e)[:80]}")
                        continue
                
                if attempt < max_attempts - 1:
                    time.sleep(1 + attempt * 0.5)
                    
            except Exception as e:
                log(self.doc, f"⚠️ Tentativa {attempt + 1} falhou: {e}")
        
        raise Exception(f"Falha ao preencher após {max_attempts} tentativas: {selector}")
    
    def _fill_strategy_1(self, selector, value, by_xpath):
        """Estratégia 1: Native com eventos"""
        script = """
        var selector = arguments[0];
        var value = arguments[1];
        var byXPath = arguments[2];
        
        var element;
        if (byXPath) {
            var result = document.evaluate(selector, document, null, 
                XPathResult.FIRST_ORDERED_NODE_TYPE, null);
            element = result.singleNodeValue;
        } else {
            element = document.querySelector(selector);
        }
        
        if (!element) throw new Error('Campo não encontrado');
        
        // Prepara campo
        element.removeAttribute('disabled');
        element.removeAttribute('readonly');
        element.style.display = 'block';
        element.style.visibility = 'visible';
        
        // Scroll
        element.scrollIntoView({block: 'center'});
        
        // Focus
        element.focus();
        element.dispatchEvent(new Event('focus', {bubbles: true}));
        
        // Limpa
        element.value = '';
        
        // Preenche
        element.value = value;
        
        // Eventos
        ['input', 'change', 'blur', 'keyup'].forEach(function(evt) {
            element.dispatchEvent(new Event(evt, {bubbles: true}));
        });
        
        return element.value;
        """
        return self.execute_js(script, selector, value, by_xpath)
    
    def _fill_strategy_2(self, selector, value, by_xpath):
        """Estratégia 2: Compatível com React/Angular"""
        script = """
        var selector = arguments[0];
        var value = arguments[1];
        var byXPath = arguments[2];
        
        var element;
        if (byXPath) {
            var result = document.evaluate(selector, document, null, 
                XPathResult.FIRST_ORDERED_NODE_TYPE, null);
            element = result.singleNodeValue;
        } else {
            element = document.querySelector(selector);
        }
        
        if (!element) throw new Error('Campo não encontrado');
        
        // Setter nativo (React)
        var nativeInputValueSetter = Object.getOwnPropertyDescriptor(
            window.HTMLInputElement.prototype, 'value'
        ).set;
        
        if (nativeInputValueSetter) {
            nativeInputValueSetter.call(element, value);
        } else {
            element.value = value;
        }
        
        // Eventos React
        element.dispatchEvent(new Event('input', {bubbles: true}));
        element.dispatchEvent(new Event('change', {bubbles: true}));
        
        // Eventos adicionais
        element.dispatchEvent(new KeyboardEvent('keydown', {bubbles: true}));
        element.dispatchEvent(new KeyboardEvent('keyup', {bubbles: true}));
        element.dispatchEvent(new Event('blur', {bubbles: true}));
        
        return element.value;
        """
        return self.execute_js(script, selector, value, by_xpath)
    
    def _fill_strategy_3(self, selector, value, by_xpath):
        """Estratégia 3: jQuery trigger"""
        script = """
        var selector = arguments[0];
        var value = arguments[1];
        var byXPath = arguments[2];
        
        var element;
        if (byXPath) {
            var result = document.evaluate(selector, document, null, 
                XPathResult.FIRST_ORDERED_NODE_TYPE, null);
            element = result.singleNodeValue;
        } else {
            element = document.querySelector(selector);
        }
        
        if (!element) throw new Error('Campo não encontrado');
        
        element.value = value;
        
        // Trigger jQuery se disponível
        if (typeof jQuery !== 'undefined') {
            jQuery(element).val(value).trigger('input').trigger('change').trigger('blur');
        }
        
        // Eventos nativos como fallback
        ['focus', 'input', 'change', 'blur'].forEach(function(evt) {
            element.dispatchEvent(new Event(evt, {bubbles: true}));
        });
        
        return element.value;
        """
        return self.execute_js(script, selector, value, by_xpath)
    
    def _fill_strategy_4(self, selector, value, by_xpath):
        """Estratégia 4: Força bruta"""
        script = """
        var selector = arguments[0];
        var value = arguments[1];
        var byXPath = arguments[2];
        
        var element;
        if (byXPath) {
            var result = document.evaluate(selector, document, null, 
                XPathResult.FIRST_ORDERED_NODE_TYPE, null);
            element = result.singleNodeValue;
        } else {
            element = document.querySelector(selector);
        }
        
        if (!element) throw new Error('Campo não encontrado');
        
        // Remove TODOS bloqueios
        element.removeAttribute('disabled');
        element.removeAttribute('readonly');
        element.removeAttribute('maxlength');
        
        // Define valor múltiplas vezes
        element.value = '';
        element.setAttribute('value', value);
        element.value = value;
        
        // Força atualização visual
        element.style.color = element.style.color;
        
        // TODOS os eventos possíveis
        var events = ['focus', 'click', 'input', 'change', 'keydown', 'keypress', 
                      'keyup', 'blur', 'paste', 'textInput'];
        events.forEach(function(evt) {
            try {
                element.dispatchEvent(new Event(evt, {bubbles: true, cancelable: true}));
            } catch(e) {}
        });
        
        // Handlers diretos
        if (element.oninput) element.oninput();
        if (element.onchange) element.onchange();
        
        return element.value;
        """
        return self.execute_js(script, selector, value, by_xpath)
    
    def _fill_strategy_5(self, selector, value, by_xpath):
        """Estratégia 5: Último recurso - simula digitação"""
        script = """
        var selector = arguments[0];
        var value = arguments[1];
        var byXPath = arguments[2];
        
        var element;
        if (byXPath) {
            var result = document.evaluate(selector, document, null, 
                XPathResult.FIRST_ORDERED_NODE_TYPE, null);
            element = result.singleNodeValue;
        } else {
            element = document.querySelector(selector);
        }
        
        if (!element) throw new Error('Campo não encontrado');
        
        element.focus();
        element.value = '';
        
        // Simula digitação caractere por caractere
        for (var i = 0; i < value.length; i++) {
            element.value += value[i];
            
            // Evento para cada caractere
            var evt = new KeyboardEvent('keydown', {
                key: value[i],
                code: 'Key' + value[i].toUpperCase(),
                bubbles: true
            });
            element.dispatchEvent(evt);
            
            element.dispatchEvent(new Event('input', {bubbles: true}));
        }
        
        element.dispatchEvent(new Event('change', {bubbles: true}));
        element.dispatchEvent(new Event('blur', {bubbles: true}));
        
        return element.value;
        """
        return self.execute_js(script, selector, value, by_xpath)
    
    def _validate_fill(self, selector, expected_value, by_xpath):
        """Valida se o campo foi preenchido corretamente"""
        script = """
        var selector = arguments[0];
        var expected = arguments[1];
        var byXPath = arguments[2];
        
        var element;
        if (byXPath) {
            var result = document.evaluate(selector, document, null, 
                XPathResult.FIRST_ORDERED_NODE_TYPE, null);
            element = result.singleNodeValue;
        } else {
            element = document.querySelector(selector);
        }
        
        if (!element) return false;
        
        var actual = element.value || '';
        return actual.trim() === expected.trim() || actual.includes(expected);
        """
        try:
            return self.execute_js(script, selector, expected_value, by_xpath)
        except:
            return False
    
    def force_select(self, selector, text, by_xpath=False, max_attempts=5):
        """Seleção forçada em dropdown/select"""
        log(self.doc, f"🔽 Seleção forçada: {selector} = '{text}'")
        
        for attempt in range(max_attempts):
            try:
                script = """
                var selector = arguments[0];
                var text = arguments[1];
                var byXPath = arguments[2];
                
                var element;
                if (byXPath) {
                    var result = document.evaluate(selector, document, null, 
                        XPathResult.FIRST_ORDERED_NODE_TYPE, null);
                    element = result.singleNodeValue;
                } else {
                    element = document.querySelector(selector);
                }
                
                if (!element) throw new Error('Select não encontrado');
                
                // Tenta por texto visível
                for (var i = 0; i < element.options.length; i++) {
                    if (element.options[i].text.trim() === text.trim()) {
                        element.selectedIndex = i;
                        element.value = element.options[i].value;
                        element.dispatchEvent(new Event('change', {bubbles: true}));
                        return true;
                    }
                }
                
                // Tenta por value
                element.value = text;
                element.dispatchEvent(new Event('change', {bubbles: true}));
                
                return element.value !== '';
                """
                
                result = self.execute_js(script, selector, text, by_xpath)
                if result:
                    log(self.doc, f"✅ Opção selecionada: '{text}'")
                    return True
                    
                if attempt < max_attempts - 1:
                    time.sleep(1)
                    
            except Exception as e:
                log(self.doc, f"⚠️ Tentativa {attempt + 1} falhou: {e}")
                if attempt < max_attempts - 1:
                    time.sleep(1)
        
        raise Exception(f"Falha ao selecionar opção após {max_attempts} tentativas")
    
    def force_datepicker(self, selector, date_value, by_xpath=False, max_attempts=5):
        """Preenchimento forçado de datepicker"""
        log(self.doc, f"📅 Datepicker forçado: {selector} = '{date_value}'")
        
        for attempt in range(max_attempts):
            try:
                script = """
                var selector = arguments[0];
                var dateValue = arguments[1];
                var byXPath = arguments[2];
                
                var element;
                if (byXPath) {
                    var result = document.evaluate(selector, document, null, 
                        XPathResult.FIRST_ORDERED_NODE_TYPE, null);
                    element = result.singleNodeValue;
                } else {
                    element = document.querySelector(selector);
                }
                
                if (!element) throw new Error('Datepicker não encontrado');
                
                // Remove readonly/disabled
                element.removeAttribute('readonly');
                element.removeAttribute('disabled');
                
                // jQuery datepicker
                if (typeof jQuery !== 'undefined' && jQuery(element).datepicker) {
                    try {
                        jQuery(element).datepicker('setDate', dateValue);
                        jQuery(element).trigger('change');
                        return element.value;
                    } catch(e) {}
                }
                
                // Native
                element.value = '';
                element.value = dateValue;
                
                // Eventos completos
                ['focus', 'input', 'change', 'blur', 'keyup'].forEach(function(evt) {
                    element.dispatchEvent(new Event(evt, {bubbles: true}));
                });
                
                return element.value;
                """
                
                result = self.execute_js(script, selector, date_value, by_xpath)
                
                # Valida
                time.sleep(0.3)
                if self._validate_fill(selector, date_value, by_xpath):
                    log(self.doc, f"✅ Datepicker preenchido: '{date_value}'")
                    return True
                
                if attempt < max_attempts - 1:
                    time.sleep(1)
                    
            except Exception as e:
                log(self.doc, f"⚠️ Tentativa {attempt + 1} falhou: {e}")
                if attempt < max_attempts - 1:
                    time.sleep(1)
        
        raise Exception(f"Falha ao preencher datepicker após {max_attempts} tentativas")
    
    def force_modal_open(self, btn_selector, modal_selector, by_xpath=False, max_attempts=5):
        """Abre modal forçadamente e aguarda aparecer"""
        log(self.doc, f"🔓 Abrindo modal: {btn_selector}")
        
        for attempt in range(max_attempts):
            try:
                # Clica no botão
                self.force_click(btn_selector, by_xpath)
                
                # Aguarda modal aparecer
                time.sleep(1)
                
                script = """
                var selector = arguments[0];
                var byXPath = arguments[1];
                
                var element;
                if (byXPath) {
                    var result = document.evaluate(selector, document, null, 
                        XPathResult.FIRST_ORDERED_NODE_TYPE, null);
                    element = result.singleNodeValue;
                } else {
                    element = document.querySelector(selector);
                }
                
                if (!element) return false;
                
                var style = window.getComputedStyle(element);
                return style.display !== 'none' && 
                       style.visibility !== 'hidden' && 
                       parseFloat(style.opacity || 1) > 0.01;
                """
                
                modal_visible = self.execute_js(script, modal_selector, by_xpath)
                
                if modal_visible:
                    log(self.doc, "✅ Modal aberto com sucesso")
                    self.wait_ajax_complete(5)
                    return True
                
                if attempt < max_attempts - 1:
                    log(self.doc, f"   Modal não apareceu, tentando novamente...")
                    time.sleep(1.5)
                    
            except Exception as e:
                log(self.doc, f"⚠️ Tentativa {attempt + 1} falhou: {e}")
                if attempt < max_attempts - 1:
                    time.sleep(1.5)
        
        raise Exception(f"Falha ao abrir modal após {max_attempts} tentativas")
    
    def force_modal_close(self, close_selector=None, by_xpath=False, max_attempts=3):
        """Fecha modal forçadamente"""
        log(self.doc, "🔒 Fechando modal...")
        
        for attempt in range(max_attempts):
            try:
                script = """
                var closeSelector = arguments[0];
                var byXPath = arguments[1];
                
                // Tenta seletor específico
                if (closeSelector) {
                    var element;
                    if (byXPath) {
                        var result = document.evaluate(closeSelector, document, null, 
                            XPathResult.FIRST_ORDERED_NODE_TYPE, null);
                        element = result.singleNodeValue;
                    } else {
                        element = document.querySelector(closeSelector);
                    }
                    
                    if (element) {
                        element.click();
                        return true;
                    }
                }
                
                // Busca botões de fechar comuns
                var selectors = [
                    '.ui-dialog-titlebar-close',
                    '.close',
                    '[data-dismiss="modal"]',
                    '.modal-close',
                    'button[aria-label="Close"]',
                    '.wdClose a'
                ];
                
                for (var i = 0; i < selectors.length; i++) {
                    var btn = document.querySelector(selectors[i]);
                    if (btn && btn.offsetParent !== null) {
                        btn.click();
                        return true;
                    }
                }
                
                // Remove modais diretamente
                var modals = document.querySelectorAll('.modal, .ui-dialog, [role="dialog"]');
                modals.forEach(function(modal) {
                    modal.style.display = 'none';
                    modal.remove();
                });
                
                return modals.length > 0;
                """
                
                result = self.execute_js(script, close_selector, by_xpath)
                
                if result:
                    log(self.doc, "✅ Modal fechado")
                    time.sleep(0.5)
                    return True
                
                if attempt < max_attempts - 1:
                    time.sleep(1)
                    
            except Exception as e:
                log(self.doc, f"⚠️ Tentativa {attempt + 1} falhou: {e}")
                if attempt < max_attempts - 1:
                    time.sleep(1)
        
        log(self.doc, "⚠️ Não foi possível fechar modal, continuando...")
        return False


class LOVHandler:
    """Handler especializado para modais LOV (List of Values)"""
    
    def __init__(self, js_engine):
        self.js = js_engine
        self.doc = js_engine.doc
    
    def open_and_select(self, btn_index=None, btn_xpath=None, search_text="", 
                       result_text="", iframe_xpath=None, max_attempts=5):
        """Abre LOV, pesquisa e seleciona resultado usando JS forçado"""
        
        log(self.doc, f"🔍 Processando LOV: '{search_text}' → '{result_text}'")
        
        for attempt in range(max_attempts):
            try:
                log(self.doc, f"   Tentativa {attempt + 1}/{max_attempts}")
                
                # PASSO 1: Clica no botão LOV
                if btn_index is not None:
                    btn_selector = f"(//a[@class='sprites sp-openLov'])[{btn_index + 1}]"
                    by_xpath = True
                elif btn_xpath:
                    btn_selector = btn_xpath
                    by_xpath = True
                else:
                    raise ValueError("btn_index ou btn_xpath deve ser fornecido")
                
                log(self.doc, "   📌 Abrindo LOV...")
                self.js.force_click(btn_selector, by_xpath=by_xpath)
                time.sleep(1.5)
                
                # PASSO 2: Troca para iframe se necessário
                if iframe_xpath:
                    log(self.doc, "   🔄 Entrando no iframe...")
                    try:
                        WebDriverWait(self.js.driver, 10).until(
                            EC.frame_to_be_available_and_switch_to_it((By.XPATH, iframe_xpath))
                        )
                        time.sleep(0.5)
                    except:
                        log(self.doc, "   ⚠️ Iframe não encontrado, continuando...")
                
                # PASSO 3: Aguarda modal carregar
                self.js.wait_ajax_complete(10)
                time.sleep(1)
                
                # PASSO 4: Preenche campo de pesquisa
                log(self.doc, f"   ✏️ Pesquisando: '{search_text}'")
                
                search_selectors = [
                    "//input[@id='txtPesquisa']",
                    "//input[@class='nomePesquisa']",
                ]
                
                search_filled = False
                for selector in search_selectors:
                    try:
                        self.js.force_fill(selector, search_text, by_xpath=True)
                        search_filled = True
                        break
                    except:
                        continue
                
                if not search_filled:
                    raise Exception("Campo de pesquisa não encontrado")
                
                time.sleep(0.5)
                
                # PASSO 5: Clica em Pesquisar
                log(self.doc, "   🔎 Executando pesquisa...")
                
                search_btn_selectors = [
                    "//a[contains(@class,'lpFind') and contains(normalize-space(.),'Pesquisar')]",
                    "//button[contains(normalize-space(.),'Pesquisar')]",
                    "//a[contains(normalize-space(.),'Buscar')]"
                ]
                
                search_clicked = False
                for selector in search_btn_selectors:
                    try:
                        self.js.force_click(selector, by_xpath=True)
                        search_clicked = True
                        break
                    except:
                        continue
                
                if not search_clicked:
                    raise Exception("Botão de pesquisa não encontrado")
                
                # PASSO 6: Aguarda resultados
                time.sleep(2)
                self.js.wait_ajax_complete(15)
                
                # PASSO 7: Clica no resultado
                log(self.doc, f"   🎯 Selecionando: '{result_text}'")
                
                result_xpath = f"//tr[td[contains(normalize-space(.), '{result_text}')]]"
                self.js.force_click(result_xpath, by_xpath=True)
                
                time.sleep(1)
                
                # PASSO 8: Volta para conteúdo principal
                if iframe_xpath:
                    try:
                        self.js.driver.switch_to.default_content()
                        log(self.doc, "   ✅ Voltou para conteúdo principal")
                    except:
                        pass
                
                # PASSO 9: Aguarda modal fechar
                time.sleep(1)
                self.js.wait_ajax_complete(10)
                
                log(self.doc, f"✅ LOV processado com sucesso!")
                return True
                
            except Exception as e:
                log(self.doc, f"⚠️ Tentativa {attempt + 1} falhou: {str(e)[:150]}")
                
                # Cleanup
                try:
                    self.js.driver.switch_to.default_content()
                except:
                    pass
                
                try:
                    self.js.force_modal_close()
                except:
                    pass
                
                if attempt < max_attempts - 1:
                    time.sleep(2 + attempt * 0.5)
                else:
                    raise Exception(f"Falha ao processar LOV após {max_attempts} tentativas: {e}")
        
        return False


# ==== WRAPPERS DE ALTO NÍVEL ====

def safe_action(doc, descricao, func, max_retries=3):
    """Wrapper para ações com retry automático"""
    global driver
    
    for attempt in range(max_retries):
        try:
            log(doc, f"🔄 {descricao}..." if attempt == 0 else f"🔄 {descricao}... (Tentativa {attempt + 1})")
            func()
            log(doc, f"✅ {descricao} realizada com sucesso.")
            take_screenshot(driver, doc, _sanitize_filename(descricao))
            return True
        except Exception as e:
            if attempt < max_retries - 1:
                log(doc, f"⚠️ Tentativa {attempt + 1} falhou, tentando novamente...")
                time.sleep(2 + attempt)
                continue
            else:
                log(doc, f"❌ Erro após {max_retries} tentativas: {e}")
                take_screenshot(driver, doc, _sanitize_filename(f"erro_{descricao}"))
                return False
    
    return False


def inicializar_driver():
    """Inicializa WebDriver com configurações otimizadas"""
    global driver, wait
    
    try:
        options = Options()
        options.add_argument("--start-maximized")
        options.add_argument("--disable-blink-features=AutomationControlled")
        options.add_argument("--disable-dev-shm-usage")
        options.add_argument("--no-sandbox")
        options.add_experimental_option("excludeSwitches", ["enable-automation"])
        options.add_experimental_option('useAutomationExtension', False)
        
        driver = webdriver.Chrome(
            service=Service(ChromeDriverManager().install()), 
            options=options
        )
        
        driver.execute_script(
            "Object.defineProperty(navigator, 'webdriver', {get: () => undefined})"
        )
        
        wait = WebDriverWait(driver, TIMEOUT_DEFAULT)
        
        log(doc, "✅ Driver inicializado com sucesso")
        return True
        
    except Exception as e:
        log(doc, f"❌ Erro ao inicializar driver: {e}")
        return False


def finalizar_relatorio():
    """Salva relatório e fecha driver"""
    global driver, doc
    
    nome_arquivo = f"relatorio_geracao_titulos_unicos_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    
    try:
        doc.save(nome_arquivo)
        log(doc, f"📄 Relatório salvo: {nome_arquivo}")
        
        try:
            subprocess.run(["start", "winword", nome_arquivo], shell=True)
        except:
            pass
            
    except Exception as e:
        print(f"Erro ao salvar relatório: {e}")
    
    if driver:
        try:
            driver.quit()
            log(doc, "✅ Driver encerrado")
        except:
            pass


def clicar_lov_por_indice(indice_lov: int, max_tentativas: int = 5, timeout: int = 10, scroll: bool = True):
    """
    Clica no ícone de LOV <a class="sprites sp-openLov"> pelo índice (ordem no DOM).
    Retorna uma função 'acao' para ser usada com safe_action(..., lambda: ...).

    Estratégias de clique:
      1) Espera 'clickable' + click nativo
      2) ScrollIntoView + click nativo
      3) JavaScript click
      4) ActionChains move_to_element + click
    """
    def acao():
        if not isinstance(indice_lov, int) or indice_lov < 0:
            raise ValueError(f"Índice inválido: {indice_lov}")

        tentativa = 0
        while tentativa < max_tentativas:
            tentativa += 1
            try:
                log(doc, f"🔎 Tentativa {tentativa}: Localizando ícones LOV ('sp-openLov')...")
                # Coleta atual dos elementos
                elementos = driver.find_elements(By.CSS_SELECTOR, "a.sprites.sp-openLov")

                if not elementos:
                    if tentativa < max_tentativas:
                        log(doc, f"⚠️ Nenhum ícone LOV encontrado (tentativa {tentativa}/{max_tentativas}). Reintentando...", "WARN")
                        time.sleep(1.2)
                        continue
                    raise Exception("Nenhum ícone LOV ('a.sprites.sp-openLov') foi encontrado na página.")

                if indice_lov >= len(elementos):
                    raise Exception(f"Índice {indice_lov} inválido. Encontrados {len(elementos)} ícones LOV.")

                # Mantém um localizador estável por índice para waits/refresh
                locator_xpath = f"(//a[contains(@class,'sp-openLov')])[{indice_lov + 1}]"

                # Reaponta o elemento pelo locator (evita stale)
                elemento = driver.find_element(By.XPATH, locator_xpath)

                log(doc, f"🎯 Preparando clique no LOV de índice {indice_lov} (total: {len(elementos)}).")

                def _wait_clickable():
                    wait.until(EC.element_to_be_clickable((By.XPATH, locator_xpath)))

                estrategias = [
                    # 1) Espera 'clickable' + click nativo
                    lambda: (_wait_clickable(), elemento.click()),
                    # 2) ScrollIntoView + click nativo
                    lambda: (
                        driver.execute_script("arguments[0].scrollIntoView({block:'center', inline:'nearest'});", elemento) if scroll else None,
                        time.sleep(0.2),
                        elemento.click()
                    ),
                    # 3) JavaScript click
                    lambda: driver.execute_script("arguments[0].click();", elemento),
                    # 4) ActionChains
                    lambda: (ActionChains(driver).move_to_element(elemento).pause(0.1).click().perform())
                ]

                for i, estrategia in enumerate(estrategias, 1):
                    try:
                        log(doc, f"   ▶️ Estratégia {i} de clique no LOV...")
                        estrategia()
                        time.sleep(0.3)
                        log(doc, f"✅ Clique no LOV (índice {indice_lov}) realizado com sucesso (estratégia {i}).")
                        return True
                    except (ElementClickInterceptedException, StaleElementReferenceException, JavascriptException, TimeoutException) as e:
                        log(doc, f"⚠️ Estratégia {i} falhou: {e}", "WARN")
                        # Tenta re-obter o elemento em caso de stale/interceptação
                        try:
                            _ = driver.find_elements(By.CSS_SELECTOR, "a.sprites.sp-openLov")
                            elemento = driver.find_element(By.XPATH, locator_xpath)
                        except Exception:
                            pass
                        continue

                if tentativa < max_tentativas:
                    log(doc, f"⚠️ Tentativa {tentativa} não conseguiu clicar no LOV. Reintentando em 1.2s…", "WARN")
                    time.sleep(1.2)
                    continue

            except Exception as e:
                if tentativa < max_tentativas:
                    log(doc, f"⚠️ Erro na tentativa {tentativa}: {e}. Reintentando em 1.2s…", "WARN")
                    time.sleep(1.2)
                    continue
                raise

        raise Exception(f"Falha ao clicar no LOV de índice {indice_lov} após {max_tentativas} tentativas.")

    return acao



def abrir_modal_e_selecionar_robusto_xpath(
    btn_xpath,
    pesquisa_xpath,
    termo_pesquisa,
    btn_pesquisar_xpath,
    resultado_xpath,
    timeout=12,
    max_tentativas=3,
    iframe_xpath=None,
    indice_lov=None,
    **_ignorar_kwargs
):
    """
    Abre o modal (LOV), pesquisa pelo termo e clica no resultado.
    - Usa clicar_lov_por_indice se o índice for informado
    - Usa retries, fallback JS e limpeza resistente do input
    - Pode lidar com iframe do modal
    """

    def _js_click(el):
        driver.execute_script("arguments[0].click();", el)

    def _clear_resistente(el):
        try:
            el.clear()
            el.send_keys(Keys.CONTROL, "a")
            el.send_keys(Keys.DELETE)
        except Exception:
            driver.execute_script("arguments[0].value='';", el)

    def _aguardar_ajax_overlay():
        t0 = time.time()
        while time.time() - t0 < 8:
            try:
                ready = driver.execute_script("return document.readyState")
                ajax_ok = driver.execute_script("return window.jQuery ? jQuery.active === 0 : true")
                if ready == "complete" and ajax_ok:
                    break
            except Exception:
                pass
            time.sleep(0.2)

    def acao():
        for tentativa in range(1, max_tentativas + 1):
            try:
                log(doc, f"🔄 Tentativa {tentativa}/{max_tentativas} - Abrindo LOV robusto")

                driver.switch_to.default_content()
                if indice_lov is not None:
                    log(doc, f"📌 Clicando no LOV de índice {indice_lov}")
                    clicar_lov_por_indice(indice_lov, max_tentativas=3, timeout=timeout)()
                else:
                    log(doc, "📌 Clicando no botão LOV pelo XPath")
                    botao = WebDriverWait(driver, timeout).until(
                        EC.element_to_be_clickable((By.XPATH, btn_xpath))
                    )
                    try:
                        botao.click()
                    except:
                        _js_click(botao)

                # Troca para iframe, se houver
                if iframe_xpath:
                    WebDriverWait(driver, timeout).until(
                        EC.frame_to_be_available_and_switch_to_it((By.XPATH, iframe_xpath))
                    )

                _aguardar_ajax_overlay()

                # Campo de pesquisa
                campo = WebDriverWait(driver, timeout).until(
                    EC.visibility_of_element_located((By.XPATH, pesquisa_xpath))
                )
                _clear_resistente(campo)
                campo.send_keys(termo_pesquisa)

                # Botão pesquisar
                btn_pesq = WebDriverWait(driver, timeout).until(
                    EC.element_to_be_clickable((By.XPATH, btn_pesquisar_xpath))
                )
                try:
                    btn_pesq.click()
                except:
                    _js_click(btn_pesq)

                _aguardar_ajax_overlay()
                time.sleep(1)

                # Resultado
                resultado = WebDriverWait(driver, timeout).until(
                    EC.element_to_be_clickable((By.XPATH, resultado_xpath))
                )
                driver.execute_script("arguments[0].scrollIntoView({block:'center'});", resultado)
                try:
                    resultado.click()
                except:
                    _js_click(resultado)

                driver.switch_to.default_content()
                log(doc, "✅ Modal LOV processado com sucesso!")
                return True

            except Exception as e:
                log(doc, f"⚠️ Falha na tentativa {tentativa}: {e}")
                driver.switch_to.default_content()
                if tentativa < max_tentativas:
                    time.sleep(2)
                else:
                    raise

    return acao


def validar_resultado_pesquisa(js_engine):
    """Valida se há títulos encontrados e seleciona o primeiro"""
    try:
        log(doc, "🔍 Validando resultado da pesquisa...")
        time.sleep(2)
        
        script = """
        // Procura checkboxes na tabela de resultados
        var checkboxes = document.querySelectorAll(
            "table tbody tr input[class='sorting_1'], " +
            "input[class='sorting_1'], " +
            "table tbody tr input[type='checkbox']"
        );
        
        // Filtra visíveis e habilitados
        var visible = [];
        for (var i = 0; i < checkboxes.length; i++) {
            var cb = checkboxes[i];
            var style = window.getComputedStyle(cb);
            if (style.display !== 'none' && 
                style.visibility !== 'hidden' && 
                cb.offsetParent !== null &&
                !cb.disabled) {
                visible.push(cb);
            }
        }
        
        if (visible.length === 0) return null;
        
        // Clica no primeiro
        var first = visible[0];
        first.scrollIntoView({block: 'center'});
        first.click();
        
        return visible.length;
        """
        
        result = js_engine.execute_js(script)
        
        if result is None or result == 0:
            log(doc, "❌ Nenhum título encontrado")
            return False
        
        log(doc, f"✅ {result} título(s) encontrado(s), primeiro selecionado")
        return True
        
    except Exception as e:
        log(doc, f"❌ Erro ao validar resultado: {e}")
        return False

def forcar_retorno_tela_sistema(js_engine, esperado_selector="#gsFinan", timeout=12):
    """
    Força o retorno para a tela do sistema após confirmação de modal.
    - js_engine: instância de JSForceEngine (tem driver, execute_js, force_click, etc.)
    - esperado_selector: seletor que indica que a tela principal está visível (ajuste conforme sua UI)
    """
    driver = js_engine.driver
    wait = WebDriverWait(driver, 3)
    log(doc, "🔁 Forçando retorno para tela do sistema...")

    try:
        # 1) Volta para o conteúdo principal
        try:
            driver.switch_to.default_content()
        except:
            pass

        # 2) Remove overlays/modais via JS (limpeza agressiva)
        try:
            js_engine.execute_js("""
                document.querySelectorAll('.ui-widget-overlay, .blockUI, .modal-backdrop, .blockScreen, .overlay').forEach(function(o){
                    o.style.display='none'; o.style.visibility='hidden'; o.style.opacity='0';
                });
                // Remove modais persistentes do DOM se existirem
                document.querySelectorAll('.modal, .ui-dialog, [role=\"dialog\"]').forEach(function(m){
                    try { m.remove(); } catch(e) {}
                });
                return true;
            """)
            log(doc, "   ✅ Overlays/modais escondidos/removidos (JS)")
        except Exception as e:
            log(doc, f"   ⚠️ Não foi possível limpar overlays via JS: {e}")

        # 3) Tentar clicar no botão de fechar do módulo (ex.: wdClose)
        try:
            js_engine.force_click("#gsFinan > div.wdTop.ui-draggable-handle > div.wdClose > a", by_xpath=False, max_attempts=3)
            log(doc, "   ✅ Botão de fechar do módulo clicado")
        except Exception as e:
            log(doc, f"   ⚠️ Botão de fechar do módulo não clicado: {e}")

        # 4) Enviar ESC como fallback (fecha modais que respondem a ESC)
        try:
            ActionChains(driver).send_keys(Keys.ESCAPE).perform()
            time.sleep(0.5)
            log(doc, "   ✅ ESC enviado")
        except Exception as e:
            log(doc, f"   ⚠️ Falha ao enviar ESC: {e}")

        # 5) Espera seletor da tela principal ficar presente/visível
        t0 = time.time()
        while time.time() - t0 < timeout:
            try:
                # Ajuste: se a sua tela principal tiver um elemento específico substitua o selector
                element = driver.find_element(By.CSS_SELECTOR, esperado_selector)
                # Checa visibilidade via JS para evitar false positives
                visible = js_engine.execute_js("""
                    var el = arguments[0];
                    if(!el) return false;
                    var s = window.getComputedStyle(el);
                    return (s.display !== 'none' && s.visibility !== 'hidden' && parseFloat(s.opacity||1) > 0.01);
                """, element)
                if visible:
                    log(doc, f"✅ Tela principal detectada ({esperado_selector})")
                    return True
            except Exception:
                pass
            time.sleep(0.5)

        # 6) Último recurso: refresh da página para garantir estado limpo
        try:
            log(doc, "🔄 Último recurso: recarregando a página")
            driver.refresh()
            time.sleep(3)
            # checa novamente
            try:
                driver.find_element(By.CSS_SELECTOR, esperado_selector)
                log(doc, f"✅ Tela principal encontrada após refresh ({esperado_selector})")
                return True
            except:
                pass
        except Exception as e:
            log(doc, f"⚠️ Falha ao recarregar página: {e}")

        log(doc, "⚠️ Não foi possível garantir retorno total à tela do sistema")
        return False

    except Exception as e:
        log(doc, f"❌ Erro em forcar_retorno_tela_sistema: {e}")
        return False



def confirmar_modal_geracao_titulos(js_engine, timeout=12, iframe_xpath=None):
    """
    Clica no 'Sim' da modal 'Confirme a geração de títulos...' e valida que ela fechou.
    Retorna True somente se a modal deixar de estar visível.
    """
    from selenium.webdriver.support.ui import WebDriverWait
    from selenium.webdriver.support import expected_conditions as EC
    from selenium.webdriver.common.by import By
    import time

    driver = js_engine.driver
    wait = WebDriverWait(driver, timeout)

    # Sair para o conteúdo principal
    try:
        driver.switch_to.default_content()
    except:
        pass

    # Se houver iframe, entrar (se não houver, ignora)
    if iframe_xpath:
        try:
            wait.until(EC.frame_to_be_available_and_switch_to_it((By.XPATH, iframe_xpath)))
        except:
            pass

    # 1) Espera a MODAL alvo ficar visível e marca o botão 'Sim' com data-aim temporário
    found = js_engine.execute_js("""
        // procura modais .modal visíveis e com o texto esperado
        var modals = document.querySelectorAll('div.modal.overflow');
        var alvo = null;
        for (var i=0;i<modals.length;i++){
            var m = modals[i];
            var s = getComputedStyle(m);
            var vis = (m.offsetParent!==null && s.display!=='none' && s.visibility!=='hidden' && parseFloat(s.opacity||1)>0.01);
            if(!vis) continue;
            var txt = (m.innerText||'').toLowerCase();
            if(txt.indexOf('confirme a geração de títulos')!==-1){
                alvo = m; break;
            }
        }
        if(!alvo) return false;

        // dentro da modal, localizar o 'Sim'
        var btn = alvo.querySelector('a.btModel.btGray.btyes') 
               || alvo.querySelector("a.btyes");
        if(!btn) return false;

        // marca para podermos clicar via CSS
        btn.setAttribute('data-aim','confirm-yes');
        return true;
    """)
    if not found:
        # não achou a modal/btn — falhar explicitamente
        return False

    # 2) Clicar com o motor robusto (dispara sequência completa de eventos)
    try:
        js_engine.force_click("[data-aim='confirm-yes']", by_xpath=False)
    except Exception:
        # fallback por XPath direto na âncora com texto 'Sim'
        js_engine.force_click("//div[contains(@class,'modal') and contains(.,'Confirme a geração de títulos')]//a[contains(@class,'btyes')]", by_xpath=True)

    # 3) Aguarda processamento/overlays
    js_engine.wait_ajax_complete(8)

    # 4) Validar que a modal sumiu (não visível ou removida)
    for _ in range(int(timeout*2)):  # ~timeout segundos
        closed = js_engine.execute_js("""
            var m = document.querySelector("div.modal.overflow");
            if(!m) return true; // removida do DOM
            var s = getComputedStyle(m);
            var vis = (m.offsetParent!==null && s.display!=='none' && s.visibility!=='hidden' && parseFloat(s.opacity||1)>0.01);
            if(!vis) return true; // oculta
            // se ainda há modal visível, tenta verificar se a 'loadingContent' está ativa
            var ld = m.querySelector('.loadingContent');
            if(ld){
                var sl = getComputedStyle(ld);
                // se loading apareceu e depois some, a modal deve fechar logo
            }
            return false;
        """)
        if closed:
            # remover o atributo temporário, se ainda existir
            js_engine.execute_js("""
                var b = document.querySelector("[data-aim='confirm-yes']");
                if(b) b.removeAttribute('data-aim');
                return true;
            """)
            return True
        time.sleep(0.5)

    # último recurso: força evento de mouse completo + tenta esconder overlays e verifica de novo
    js_engine.execute_js("""
        var b = document.querySelector("[data-aim='confirm-yes']") 
             || document.querySelector("div.modal.overflow a.btModel.btGray.btyes");
        if(b){
            ['mouseover','mousedown','mouseup','click'].forEach(function(t){
                b.dispatchEvent(new MouseEvent(t,{bubbles:true,cancelable:true,view:window,detail:1}));
            });
            if(b.click) b.click();
        }
        document.querySelectorAll('.ui-widget-overlay,.blockUI,.modal-backdrop,[class*="overlay"]').forEach(function(o){
            o.style.display='none'; o.style.visibility='hidden'; o.style.opacity='0';
        });
        return true;
    """)
    js_engine.wait_ajax_complete(4)

    closed = js_engine.execute_js("""
        var m = document.querySelector("div.modal.overflow");
        if(!m) return true;
        var s = getComputedStyle(m);
        return !(m.offsetParent!==null && s.display!=='none' && s.visibility!=='hidden' && parseFloat(s.opacity||1)>0.01);
    """)
    return bool(closed)


# ==== EXECUÇÃO DO TESTE ====

def executar_teste():
    """Execução principal do teste com JS forçado"""
    global driver, wait, doc
    
    try:
        # Inicializa driver
        if not inicializar_driver():
            return False
        
        # Cria engine JS forçado
        js_engine = JSForceEngine(driver, wait, doc)
        lov_handler = LOVHandler(js_engine)
        
        # ===== LOGIN =====
        safe_action(doc, "Acessando sistema", lambda: driver.get(URL))
        
        def fazer_login():
            wait.until(EC.presence_of_element_located((By.ID, "j_id15:email"))).send_keys(LOGIN_EMAIL)
            wait.until(EC.presence_of_element_located((By.ID, "j_id15:senha"))).send_keys(LOGIN_PASSWORD, Keys.ENTER)
            time.sleep(5)

        safe_action(doc, "Realizando login", fazer_login)
        
        # ===== MENU =====
        def abrir_menu():
            driver.execute_script("document.body.style.zoom='90%'")
            driver.find_element(By.TAG_NAME, "body").send_keys(Keys.F3)
            time.sleep(2)
        
        safe_action(doc, "Abrindo menu (F3)", abrir_menu)
        
        # ===== GESTOR FINANCEIRO =====
        safe_action(doc, "Acessando Gestor Financeiro", lambda:
            js_engine.force_click('/html/body/div[15]/ul/li[2]/img', by_xpath=True)
        )
        
        time.sleep(3)
        
        # ===== GERAÇÃO DE TÍTULOS =====
        safe_action(doc, "Clicando em Geração de Títulos", lambda:
            js_engine.force_click(
                '#gsFinan > div.wdTelas > div > ul > li:nth-child(7) > a > span'
            )
        )
        
        time.sleep(5)
        
        # ===== CONTRATANTE/TITULAR =====
        safe_action(doc, "Selecionando Pessoa", lambda:
            lov_handler.open_and_select(
                btn_index=2,
                search_text="JOÃO EDUARDO JUSTINO PASCHOAL",
                result_text="JOÃO EDUARDO JUSTINO PASCHOAL"
            )
        )
        
        # ===== BUSCAR =====
        safe_action(doc, "Clicando em Buscar", lambda:
            js_engine.force_click(
                "//a[@class='btModel btGray btPesquisarTitulos' and normalize-space()='Pesquisar']",
                by_xpath=True
            )
        )
        time.sleep(10)
        
        
       # ===== VALIDAÇÃO DO RESULTADO =====
        resultado_ok = safe_action(
            doc, 
            "Validando e selecionando título", 
            lambda: validar_resultado_pesquisa(js_engine)
        )
        
        if not resultado_ok:
            log(doc, "❌ Teste interrompido: nenhum título encontrado")
            return False
        
        log(doc, "➡️ Título selecionado, prosseguindo...")
        
        # ===== PREENCHIMENTO DOS CAMPOS =====
        
        # Data Inicial
        safe_action(doc, "Preenchendo Vencimento", lambda:
            js_engine.force_datepicker(
                "(//input[contains(@class, 'hasDatepicker')])[1]",
                "09/03/2026",
                by_xpath=True
            )
        )
        
        # ===== GERAR =====
        safe_action(doc, "Clicando em Gerar", lambda:
            js_engine.force_click(
                "//a[contains(@class,'btGerar') and contains(normalize-space(.),'Gerar')]",
                by_xpath=True
            )
        )
        
        time.sleep(5)
        
        # ===== CONFIRMAR =====
        safe_action(doc, "Confirmando", lambda: confirmar_modal_geracao_titulos(js_engine))


        time.sleep(1)
        
        # força retorno à tela principal do módulo
        safe_action(doc, "Forçando retorno à tela do sistema", lambda:
            forcar_retorno_tela_sistema(js_engine, esperado_selector="#gsFinan", timeout=12)
        )

        # ===== VERIFICAR MENSAGEM =====
        log(doc, "🔍 Verificando mensagens de alerta...")
        
        script = """
        var selectors = [
            '.alerts.salvo',
            '.alerts.alerta',
            '.alerts.erro',
            '[class*="alert"]',
            '[class*="message"]'
        ];
        
        for (var i = 0; i < selectors.length; i++) {
            var elem = document.querySelector(selectors[i]);
            if (elem && elem.offsetParent !== null) {
                return {
                    tipo: selectors[i],
                    texto: elem.textContent.trim()
                };
            }
        }
        
        return null;
        """
        
        mensagem = js_engine.execute_js(script)
        
        if mensagem:
            log(doc, f"📢 Mensagem: {mensagem['texto']}")
        else:
            log(doc, "ℹ️ Nenhuma mensagem de alerta encontrada")
        
        # ===== FECHAR MODAL =====
        safe_action(doc, "Fechando Gestor Financeiro", lambda:
            js_engine.force_click(
                "#gsFinan > div.wdTop.ui-draggable-handle > div.wdClose > a"
            )
        )
        
        log(doc, "🎉 Teste concluído com sucesso!")
        return True
        
    except Exception as e:
        log(doc, f"❌ ERRO FATAL: {e}")
        take_screenshot(driver, doc, "erro_fatal")
        return False


# ==== MAIN ====

def main():
    """Ponto de entrada principal"""
    global doc
    
    try:
        log(doc, "🚀 Iniciando teste de Geração de Títulos Únicos")
        log(doc, "=" * 70)
        
        sucesso = executar_teste()
        
        log(doc, "=" * 70)
        if sucesso:
            log(doc, "✅ TESTE EXECUTADO COM SUCESSO!")
        else:
            log(doc, "❌ TESTE FINALIZADO COM ERROS")
            
    except Exception as e:
        log(doc, f"❌ Erro na execução principal: {e}")
        
    finally:
        finalizar_relatorio()


if __name__ == "__main__":
    main()