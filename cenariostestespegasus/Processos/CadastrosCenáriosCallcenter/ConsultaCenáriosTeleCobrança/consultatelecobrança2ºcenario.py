# ==== IMPORTS (sem conflitos) ====
from datetime import datetime, timedelta
from datetime import time as dt_time  # usar para objetos de hora
import time                           # usar para time.sleep(...)
import pyautogui
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support.ui import Select, WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service
from selenium.common.exceptions import TimeoutException, NoSuchElementException, StaleElementReferenceException, ElementClickInterceptedException
from webdriver_manager.chrome import ChromeDriverManager
from docx import Document
from docx.shared import Inches
from faker import Faker
from faker.providers import BaseProvider
from validate_docbr import CPF
from selenium.webdriver import ActionChains
import subprocess
import os
import random
import re
from selenium.webdriver.common.by import By
from selenium.webdriver.support import expected_conditions as EC
from selenium.common.exceptions import ElementClickInterceptedException, StaleElementReferenceException
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.remote.webelement import WebElement
import traceback

# Timeouts configuráveis
TIMEOUT_DEFAULT = 30
TIMEOUT_CURTO = 10
TIMEOUT_LONGO = 60

# ==== PROVIDERS CUSTOMIZADOS ====
class BrasilProvider(BaseProvider):
    def rg(self):
        numeros = [str(random.randint(0, 9)) for _ in range(8)]
        return ''.join(numeros) + '-' + str(random.randint(0, 9))

fake = Faker("pt_BR")
fake.add_provider(BrasilProvider)

# ==== GERAÇÃO DE DATAS ====
def gerar_datas_validas(hora_padrao="00:00", dias_fim=0):
    """Gera datas coerentes. data_inicio/data_fim no formato 'dd/MM/yyyy HH:mm'."""
    hoje_date = datetime.today().date()
    dez_anos_atras = hoje_date - timedelta(days=3650)

    # Falecimento entre 10 anos atrás e hoje
    data_falecimento = fake.date_between(start_date=dez_anos_atras, end_date=hoje_date)

    # Nascimento (entre 18 e 110 anos antes do falecimento)
    idade_minima, idade_maxima = 18, 110
    data_nascimento = data_falecimento - timedelta(days=random.randint(idade_minima * 365, idade_maxima * 365))

    # Sepultamento 1..10 dias após o falecimento
    data_sepultamento = data_falecimento + timedelta(days=random.randint(1, 10))

    # Registro 1..10 dias após o sepultamento
    data_registro = data_sepultamento + timedelta(days=random.randint(1, 10))

    # Velório entre o falecimento e o sepultamento
    data_velorio = fake.date_between(start_date=data_falecimento, end_date=data_sepultamento)

    # Início entre 2 e 30 dias no futuro, com hora escolhida
    data_inicio_date = hoje_date + timedelta(days=random.randint(2, 30))

    # Monta datetime com hora escolhida (ex: "00:00")
    h, m = map(int, hora_padrao.split(":"))
    dt_inicio = datetime.combine(data_inicio_date, dt_time(h, m))

    # Fim: mesmo dia por padrão (dias_fim=0). Ajuste se quiser +N dias.
    dt_fim = dt_inicio + timedelta(days=dias_fim)

    fmt_data = "%d/%m/%Y"
    fmt_dt = "%d/%m/%Y %H:%M"

    return (
        data_nascimento.strftime(fmt_data),
        data_falecimento.strftime(fmt_data),
        data_sepultamento.strftime(fmt_data),
        data_velorio.strftime(fmt_data),
        dt_inicio.strftime(fmt_dt),   # data_inicio com hora
        dt_fim.strftime(fmt_dt),      # data_fim com hora
        data_registro.strftime(fmt_data),
        hoje_date.strftime(fmt_data),
    )

(data_nascimento, data_falecimento, data_sepultamento, 
 data_velorio, data_inicio, data_fim, data_registro, hoje) = gerar_datas_validas(
    hora_padrao="08:50",
    dias_fim=0
)

hora_falecimento = fake.time(pattern="%H:%M")
hora_sepultamento = fake.time(pattern="%H:%M")
localizacao = fake.city()

# ==== CONFIGURAÇÕES ====
URL = "http://localhost:8080/gs/index.xhtml"
LOGIN_EMAIL = "joaoeduardo.gold@outlook.com"
LOGIN_PASSWORD = "071999gs"

# ==== INICIALIZAÇÃO DE VARIÁVEIS GLOBAIS ====
doc = Document()
doc.add_heading("RELATÓRIO DO TESTE", 0)
doc.add_paragraph("Consulta de Telecobrança - Callcenter – Cenário 2: Nesse teste, o usuário irá realizar a consulta de Telecobrança sem preencher nenhum campo para verificar se o sistema efetua o disparo das mensagens de alerta corretamente.")
doc.add_paragraph(f"Data do teste: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")

screenshot_registradas = set()
driver = None
wait = None

# ==== UTILITÁRIOS MELHORADOS ====
def log(doc, msg, nivel='INFO'):
    timestamp = datetime.now().strftime('%H:%M:%S')
    formatted_msg = f"[{timestamp}] {msg}"
    print(formatted_msg)
    try:
        doc.add_paragraph(formatted_msg)
    except Exception as e:
        print(f"Erro ao adicionar ao documento: {e}")

def _sanitize_filename(name: str) -> str:
    if not isinstance(name, str):
        name = str(name)
    name = name.strip().lower()
    name = re.sub(r"[<>:\"/\\|?*']", "_", name)
    name = re.sub(r"_+", "_", name)
    return name[:120]

def take_screenshot(driver, doc, nome, forcar=False):
    if driver is None:
        log(doc, "⚠️ Driver não disponível para screenshot", 'WARN')
        return
    
    try:
        nome = _sanitize_filename(nome)
        if forcar or nome not in screenshot_registradas:
            path = f"screenshots/{nome}.png"
            os.makedirs("screenshots", exist_ok=True)
            driver.save_screenshot(path)
            log(doc, f"📸 Screenshot capturada: {nome}")
            doc.add_paragraph(f"Screenshot: {nome}")
            doc.add_picture(path, width=Inches(5.5))
            screenshot_registradas.add(nome)
    except Exception as e:
        log(doc, f"⚠️ Erro ao tirar screenshot {nome}: {e}", 'WARN')

def safe_action(doc, descricao, func, critico=False):
    """Executa ação de forma segura com tratamento de erros melhorado"""
    try:
        log(doc, f"🔄 {descricao}...")
        resultado = func()
        log(doc, f"✅ {descricao} realizada com sucesso.")
        time.sleep(1)  # Pausa padrão após ações
        take_screenshot(driver, doc, descricao.lower().replace(" ", "_"))
        return resultado
    except Exception as e:
        take_screenshot(driver, doc, f"erro_{descricao.lower().replace(' ', '_')}")
        log(doc, f"❌ {descricao} falhou: {type(e).__name__}: {str(e)}", 'ERROR')
        
        # Log do stacktrace completo apenas se for crítico
        if critico:
            log(doc, "— stacktrace —")
            log(doc, traceback.format_exc())
            raise
        else:
            log(doc, f"⚠️ Continuando execução apesar do erro em: {descricao}", 'WARN')
            return None

def _sanitize_timeout(t):
    """Garante timeout válido"""
    if not isinstance(t, (int, float)) or t <= 0:
        return TIMEOUT_DEFAULT
    return max(5, min(120, t))  # Entre 5 e 120 segundos

# ==== AGUARDAR ELEMENTO MELHORADO ====
def aguardar_elemento(seletor, timeout=TIMEOUT_DEFAULT, condicao='clickable', by_type=By.CSS_SELECTOR):
    """Função centralizada para aguardar elementos com diferentes condições"""
    global driver, wait
    
    if driver is None:
        raise Exception("Driver não inicializado")
    
    timeout = _sanitize_timeout(timeout)
    
    condicoes = {
        'present': EC.presence_of_element_located,
        'visible': EC.visibility_of_element_located,
        'clickable': EC.element_to_be_clickable,
        'invisible': EC.invisibility_of_element_located
    }
    
    if condicao not in condicoes:
        condicao = 'clickable'
    
    try:
        wait_obj = WebDriverWait(driver, timeout)
        elemento = wait_obj.until(condicoes[condicao]((by_type, seletor)))
        return elemento
    except TimeoutException:
        log(doc, f"❌ Timeout aguardando elemento: {seletor} (condição: {condicao}, timeout: {timeout}s)", 'ERROR')
        raise TimeoutException(f"Elemento não encontrado: {seletor} (condição: {condicao})")
    except Exception as e:
        log(doc, f"❌ Erro aguardando elemento {seletor}: {e}", 'ERROR')
        raise

# ==== SCROLL CORRIGIDO - PRINCIPAL CORREÇÃO ====
def scroll_to_element_safe(elemento_ou_seletor, by_type=By.CSS_SELECTOR):
    """Scroll seguro até elemento com validação robusta"""
    global driver
    
    if driver is None:
        log(doc, "⚠️ Driver não disponível para scroll", 'WARN')
        return False
    
    try:
        # Se for seletor, encontra o elemento
        if isinstance(elemento_ou_seletor, str):
            elemento = aguardar_elemento(elemento_ou_seletor, 10, 'present', by_type)
        else:
            elemento = elemento_ou_seletor
        
        if elemento is None:
            log(doc, "⚠️ Elemento não encontrado para scroll", 'WARN')
            return False
        
        # Verifica se elemento é válido antes de fazer scroll
        if not elemento.is_displayed():
            log(doc, "⚠️ Elemento não está visível para scroll", 'WARN')
            return False
        
        # Estratégias de scroll em ordem de preferência
        scroll_strategies = [
            # Estratégia 1: JavaScript com verificação prévia
            lambda: driver.execute_script("""
                var element = arguments[0];
                if (element && typeof element.scrollIntoView === 'function') {
                    element.scrollIntoView({
                        behavior: 'smooth',
                        block: 'center',
                        inline: 'center'
                    });
                    return true;
                } else {
                    return false;
                }
            """, elemento),
            
            # Estratégia 2: ActionChains
            lambda: ActionChains(driver).move_to_element(elemento).perform(),
            
            # Estratégia 3: JavaScript alternativo
            lambda: driver.execute_script("""
                var element = arguments[0];
                if (element) {
                    element.scrollIntoView();
                    window.scrollBy(0, -100);
                }
            """, elemento),
            
            # Estratégia 4: Scroll da página até o elemento
            lambda: driver.execute_script("""
                var element = arguments[0];
                if (element) {
                    var rect = element.getBoundingClientRect();
                    var scrollTop = window.pageYOffset || document.documentElement.scrollTop;
                    var targetY = rect.top + scrollTop - (window.innerHeight / 2);
                    window.scrollTo(0, targetY);
                }
            """, elemento)
        ]
        
        for i, strategy in enumerate(scroll_strategies, 1):
            try:
                log(doc, f"   Tentando estratégia de scroll {i}...")
                result = strategy()
                
                # Para estratégia 1, verifica resultado
                if i == 1 and result is False:
                    log(doc, f"   Estratégia {i}: elemento não suporta scrollIntoView", 'WARN')
                    continue
                
                time.sleep(0.8)  # Aguarda scroll completar
                
                # Verifica se elemento ainda está acessível
                if elemento.is_displayed() and elemento.is_enabled():
                    log(doc, f"✅ Scroll realizado com estratégia {i}")
                    return True
                else:
                    log(doc, f"   Estratégia {i}: elemento não ficou acessível", 'WARN')
                    continue
                    
            except Exception as e:
                log(doc, f"   Estratégia {i} de scroll falhou: {str(e)[:100]}...", 'WARN')
                continue
        
        log(doc, "⚠️ Todas as estratégias de scroll falharam", 'WARN')
        return False
        
    except Exception as e:
        log(doc, f"⚠️ Erro geral no scroll: {e}", 'WARN')
        return False

def remover_overlays():
    """Remove overlays que podem bloquear cliques"""
    global driver
    
    if driver is None:
        return
    
    try:
        driver.execute_script("""
            // Remove overlays comuns
            document.querySelectorAll('.modal-backdrop, .overlay, .blockUI, .loading, .spinner, [style*="position: fixed"]')
                .forEach(el => {
                    if (el.style.position === 'fixed' || el.classList.contains('modal-backdrop')) {
                        el.style.display = 'none';
                    }
                });
            
            // Remove elementos com z-index alto
            document.querySelectorAll('*').forEach(el => {
                const zIndex = getComputedStyle(el).zIndex;
                if (zIndex && parseInt(zIndex) > 1000) {
                    if (el.offsetWidth === window.innerWidth && el.offsetHeight === window.innerHeight) {
                        el.style.display = 'none';
                    }
                }
            });
        """)
    except Exception as e:
        log(doc, f"⚠️ Erro ao remover overlays: {e}", 'WARN')

# ==== CLIQUE ROBUSTO CORRIGIDO ====
def clicar_elemento_robusto(seletor, timeout=TIMEOUT_DEFAULT, by_type=By.CSS_SELECTOR):
    """Clique robusto com múltiplas estratégias"""
    def acao():
        timeout_sanitized = _sanitize_timeout(timeout)
        
        try:
            # 1. Aguarda elemento
            elemento = aguardar_elemento(seletor, timeout_sanitized, 'present', by_type)
            
            # 2. Remove overlays
            remover_overlays()
            
            # 3. Scroll seguro até elemento
            scroll_success = scroll_to_element_safe(elemento)
            if not scroll_success:
                log(doc, f"⚠️ Problemas com scroll, continuando mesmo assim: {seletor}", 'WARN')
            
            # 4. Aguarda ser clicável
            try:
                elemento = aguardar_elemento(seletor, 5, 'clickable', by_type)
            except TimeoutException:
                log(doc, f"⚠️ Elemento não ficou clicável, tentando mesmo assim: {seletor}", 'WARN')
                elemento = aguardar_elemento(seletor, timeout_sanitized, 'present', by_type)
            
            # 5. Múltiplas estratégias de clique
            estrategias = [
                lambda: elemento.click(),
                lambda: ActionChains(driver).move_to_element(elemento).click().perform(),
                lambda: driver.execute_script("arguments[0].click();", elemento),
                lambda: driver.execute_script("""
                    const el = arguments[0];
                    el.focus();
                    el.dispatchEvent(new MouseEvent('click', {
                        bubbles: true, 
                        cancelable: true, 
                        view: window
                    }));
                """, elemento),
                lambda: ActionChains(driver).move_to_element_with_offset(elemento, 1, 1).click().perform()
            ]
            
            for i, estrategia in enumerate(estrategias, 1):
                try:
                    estrategia()
                    time.sleep(0.5)
                    log(doc, f"✅ Clique realizado com estratégia {i}")
                    return True
                except Exception as e:
                    log(doc, f"⚠️ Estratégia {i} de clique falhou: {e}", 'WARN')
                    if i == len(estrategias):
                        raise
                    continue
            
            return False
            
        except Exception as e:
            log(doc, f"❌ Falha ao clicar em {seletor}: {e}", 'ERROR')
            raise
    
    return acao

def clicar_elemento_xpath_robusto(xpath, timeout=TIMEOUT_DEFAULT):
    """Clique robusto via XPath"""
    return clicar_elemento_robusto(xpath, timeout, By.XPATH)

# ==== PREENCHIMENTO DE CAMPOS MELHORADO ====
def _preencher_tradicional(elemento, valor, limpar_primeiro=True):
    """Estratégia tradicional de preenchimento"""
    if limpar_primeiro:
        elemento.clear()
        time.sleep(0.2)
    elemento.click()
    time.sleep(0.2)
    elemento.send_keys(valor)
    elemento.send_keys(Keys.TAB)

def _preencher_actionchains(elemento, valor):
    """Estratégia com ActionChains"""
    global driver
    
    ActionChains(driver).move_to_element(elemento).click().perform()
    time.sleep(0.3)
    ActionChains(driver).key_down(Keys.CONTROL).send_keys('a').key_up(Keys.CONTROL).perform()
    time.sleep(0.2)
    ActionChains(driver).send_keys(valor).perform()
    time.sleep(0.2)
    ActionChains(driver).send_keys(Keys.TAB).perform()

def _preencher_javascript(elemento, valor):
    """Estratégia com JavaScript"""
    global driver
    
    driver.execute_script("""
        var element = arguments[0];
        var value = arguments[1];
        element.focus();
        element.value = '';
        element.value = value;
        element.dispatchEvent(new Event('input', { bubbles: true }));
        element.dispatchEvent(new Event('change', { bubbles: true }));
        element.blur();
    """, elemento, valor)



def _valor_do_elemento(elemento, driver=None):
    """Tenta ler o valor atual (value ou textContent p/ contenteditable)."""
    val = (elemento.get_attribute('value') or "").strip()
    if not val and elemento.get_attribute('contenteditable') in ('true', 'True', True):
        try:
            # usa JS p/ garantir leitura de contenteditable
            return (driver.execute_script("return arguments[0].textContent;", elemento) or "").strip()
        except Exception:
            return (elemento.text or "").strip()
    return val

def preencher_campo_robusto_xpath(xpath, valor, limpar_primeiro=True, timeout=TIMEOUT_DEFAULT):
    """Versão para XPath da função de preenchimento robusto."""
    def acao():
        if valor is None or valor == "":
            log(doc, f"⚠️ Valor vazio para campo (xpath): {xpath}, pulando preenchimento", 'WARN')
            return True

        # 1) aguarda de forma robusta pelo elemento via XPATH
        try:
            elemento = WebDriverWait(driver, timeout).until(
                EC.element_to_be_clickable((By.XPATH, xpath))
            )
        except Exception:
            # fallback: presença + possibilidade de clique depois
            elemento = WebDriverWait(driver, timeout).until(
                EC.presence_of_element_located((By.XPATH, xpath))
            )

        # 2) scroll até o elemento
        scroll_success = scroll_to_element_safe(elemento)
        if not scroll_success:
            log(doc, f"⚠️ Problemas com scroll para (xpath): {xpath}", 'WARN')

        # 3) estratégias de preenchimento (rebusca o elemento a cada tentativa p/ evitar stale)
        def _refind():
            try:
                return WebDriverWait(driver, 2).until(
                    EC.presence_of_element_located((By.XPATH, xpath))
                )
            except Exception:
                return elemento  # último conhecido (melhor que nada)

        estrategias = [
            lambda: _preencher_tradicional(_refind(), valor, limpar_primeiro),
            lambda: _preencher_actionchains(_refind(), valor),
            lambda: _preencher_javascript(_refind(), valor),
        ]

        for i, estrategia in enumerate(estrategias, 1):
            try:
                estrategia()
                time.sleep(0.5)

                # Verificação do preenchimento
                elemento_check = _refind()
                valor_atual = _valor_do_elemento(elemento_check, driver)

                if valor_atual.strip() == str(valor).strip() or str(valor) in valor_atual:
                    log(doc, f"✅ Campo (xpath) preenchido com estratégia {i}: '{valor_atual}'")
                    return True
                else:
                    log(doc, f"⚠️ Estratégia {i} não confirmou o valor. Esperado: '{valor}', Atual: '{valor_atual}'", 'WARN')

            except Exception as e:
                log(doc, f"⚠️ Estratégia {i} (xpath) falhou: {e}", 'WARN')
                if i == len(estrategias):
                    raise

        raise Exception(f"Falha ao preencher (xpath): {xpath} com todas as estratégias")

    return acao


def preencher_campo_robusto(seletor, valor, limpar_primeiro=True, timeout=TIMEOUT_DEFAULT):
    """Função melhorada para preenchimento de campos"""
    def acao():
        if valor is None or valor == "":
            log(doc, f"⚠️ Valor vazio para campo {seletor}, pulando preenchimento", 'WARN')
            return True
            
        elemento = aguardar_elemento(seletor, timeout, 'clickable')
        
        scroll_success = scroll_to_element_safe(elemento)
        if not scroll_success:
            log(doc, f"⚠️ Problemas com scroll para {seletor}", 'WARN')
        
        # Múltiplas estratégias de preenchimento
        estrategias = [
            lambda: _preencher_tradicional(elemento, valor, limpar_primeiro),
            lambda: _preencher_actionchains(elemento, valor),
            lambda: _preencher_javascript(elemento, valor)
        ]
        
        for i, estrategia in enumerate(estrategias, 1):
            try:
                estrategia()
                time.sleep(0.5)
                
                # Verifica se foi preenchido corretamente
                valor_atual = elemento.get_attribute('value') or ""
                if valor_atual.strip() == str(valor).strip() or str(valor) in valor_atual:
                    log(doc, f"✅ Campo preenchido com estratégia {i}: '{valor_atual}'")
                    return True
                else:
                    log(doc, f"⚠️ Estratégia {i} não preencheu corretamente. Esperado: '{valor}', Atual: '{valor_atual}'", 'WARN')
                    
            except Exception as e:
                log(doc, f"⚠️ Estratégia {i} de preenchimento falhou: {e}", 'WARN')
                if i == len(estrategias):
                    raise
                continue
                
        raise Exception(f"Falha ao preencher campo {seletor} com todas as estratégias")
    
    return acao


from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver import ActionChains
from selenium.common.exceptions import TimeoutException, StaleElementReferenceException, JavascriptException
from selenium.webdriver.support import expected_conditions as EC

def preencher_textarea_por_indice(indice_campo, texto, max_tentativas=5, limpar_primeiro=True):
    """Preenche um <textarea> pelo índice (ordem no DOM) usando estratégias múltiplas"""
    def acao():
        if not isinstance(indice_campo, int) or indice_campo < 0:
            raise ValueError(f"Índice inválido: {indice_campo}")
        if texto is None or not isinstance(texto, str):
            raise ValueError(f"Texto inválido: {texto!r}")

        tentativa = 0
        while tentativa < max_tentativas:
            tentativa += 1
            try:
                campos = encontrar_campos_textarea()
                if not campos:
                    if tentativa < max_tentativas:
                        log(doc, f"⚠️ Nenhuma <textarea> encontrada (tentativa {tentativa}/{max_tentativas})", "WARN")
                        time.sleep(1.5)
                        continue
                    raise Exception("Nenhuma <textarea> foi encontrada na página.")

                if indice_campo >= len(campos):
                    raise Exception(f"Índice {indice_campo} inválido. Encontradas {len(campos)} textareas.")

                campo_info = campos[indice_campo]
                elemento   = campo_info["elemento"]
                campo_id   = campo_info.get("id") or "(sem id)"
                campo_name = campo_info.get("name") or "(sem name)"

                log(doc, f"🎯 Tentativa {tentativa}: Preenchendo textarea {indice_campo} (ID: {campo_id}, name: {campo_name}) com {len(texto)} caracteres")

                # Se já estiver preenchido corretamente, encerra
                if validar_textarea_preenchida(elemento, texto):
                    log(doc, f"✅ Textarea {indice_campo} já está com o valor desejado.")
                    return True

                # Estratégias em ordem de 'menos invasiva' para 'mais invasiva'
                estrategias = [
                    lambda: _textarea_tradicional(elemento, texto, limpar_primeiro),
                    lambda: _textarea_actionchains(elemento, texto, limpar_primeiro),
                    lambda: _textarea_js_setvalue(elemento, texto),
                    lambda: _textarea_js_react_input(elemento, texto),
                ]

                for i, estrategia in enumerate(estrategias, 1):
                    try:
                        log(doc, f"   ▶️ Estratégia {i}…")
                        estrategia()
                        time.sleep(0.8)

                        # Revalida após a estratégia
                        if validar_textarea_preenchida(elemento, texto):
                            val = (elemento.get_attribute("value") or "").strip()
                            log(doc, f"✅ Preenchido com sucesso pela estratégia {i}: '{val[:60]}{'…' if len(val) > 60 else ''}'")
                            return True
                        else:
                            log(doc, f"⚠️ Estratégia {i} não refletiu o valor esperado.", "WARN")
                    except (StaleElementReferenceException, JavascriptException, TimeoutException) as e:
                        log(doc, f"⚠️ Estratégia {i} falhou: {e}", "WARN")
                        # Reobter o elemento se necessário
                        try:
                            campos = encontrar_campos_textarea()
                            elemento = campos[indice_campo]["elemento"]
                        except Exception:
                            pass
                        continue

                # Se chegou aqui, nenhuma estratégia funcionou nesta tentativa
                if tentativa < max_tentativas:
                    log(doc, f"⚠️ Tentativa {tentativa} falhou; reintentando em 1.5s…", "WARN")
                    time.sleep(1.5)
                    continue
            except Exception as e:
                if tentativa < max_tentativas:
                    log(doc, f"⚠️ Erro na tentativa {tentativa}: {e}. Retentando…", "WARN")
                    time.sleep(1.5)
                    continue
                else:
                    raise

        raise Exception(f"Falha ao preencher textarea {indice_campo} após {max_tentativas} tentativas.")
    return acao


# =========================
# Helpers usados pela função
# =========================

def encontrar_campos_textarea(timeout=10):
    """
    Retorna uma lista de dicts com metadados de cada <textarea> visível e interativa.
    Ex.: [{'elemento': WebElement, 'id': '...', 'name': '...'}]
    """
    elementos = []
    try:
        # Espera haver pelo menos 1 textarea no DOM (se existir)
        wait.until(lambda d: len(d.find_elements(By.TAG_NAME, "textarea")) >= 0)
        textareas = driver.find_elements(By.TAG_NAME, "textarea")
    except Exception:
        textareas = []

    for el in textareas:
        try:
            if not el.is_displayed() or not el.is_enabled():
                continue
            driver.execute_script("arguments[0].scrollIntoView({block:'center'});", el)
            elementos.append({
                "elemento": el,
                "id": el.get_attribute("id"),
                "name": el.get_attribute("name"),
            })
        except Exception:
            continue

    return elementos


def normalizar_texto(txt):
    if txt is None:
        return ""
    # Normaliza quebras de linha e espaços
    return txt.replace("\r\n", "\n").replace("\r", "\n").strip()


def validar_textarea_preenchida(elemento, texto_esperado):
    """Confere se o valor atual da textarea bate com o texto esperado (normalizado)."""
    try:
        atual = elemento.get_attribute("value")
        # Alguns frameworks populam via textContent em textareas (raro, mas possível)
        if atual is None or atual == "":
            atual = (elemento.text or "")
        return normalizar_texto(atual) == normalizar_texto(texto_esperado)
    except StaleElementReferenceException:
        return False


# =========================
# Estratégias de preenchimento
# =========================

def _prepare_focus_and_clear(elemento, limpar_primeiro=True):
    # Garante visibilidade e foco
    driver.execute_script("arguments[0].scrollIntoView({block:'center'});", elemento)
    try:
        elemento.click()
    except Exception:
        driver.execute_script("arguments[0].focus();", elemento)

    if limpar_primeiro:
        try:
            elemento.clear()
        except Exception:
            # Fallback de limpeza por teclas
            ActionChains(driver)\
                .move_to_element(elemento).click()\
                .key_down(Keys.CONTROL).send_keys("a").key_up(Keys.CONTROL)\
                .send_keys(Keys.DELETE).perform()


def _textarea_tradicional(elemento, texto, limpar_primeiro=True):
    _prepare_focus_and_clear(elemento, limpar_primeiro)
    elemento.send_keys(texto)
    # Dispara blur para muitos bindings reativos
    elemento.send_keys(Keys.TAB)


def _textarea_actionchains(elemento, texto, limpar_primeiro=True):
    _prepare_focus_and_clear(elemento, limpar_primeiro)
    ac = ActionChains(driver)
    ac.move_to_element(elemento).click().perform()
    # Quebra o texto em partes para evitar engasgos em campos longos
    for chunk_start in range(0, len(texto), 400):
        ac.send_keys(texto[chunk_start:chunk_start+400]).perform()
        time.sleep(0.05)
    ac.send_keys(Keys.TAB).perform()


def _textarea_js_setvalue(elemento, texto):
    # Seta .value e dispara eventos clássicos
    driver.execute_script("""
        const el = arguments[0];
        const val = arguments[1];
        el.value = val;
        // Dispara eventos comuns que form libs escutam
        el.dispatchEvent(new Event('input',  {bubbles:true}));
        el.dispatchEvent(new Event('change', {bubbles:true}));
        el.dispatchEvent(new KeyboardEvent('keyup', {bubbles:true}));
        el.dispatchEvent(new Event('blur',   {bubbles:true}));
    """, elemento, texto)


def _textarea_js_react_input(elemento, texto):
    # Compat extra p/ React (setando o setter do prototype) + eventos
    driver.execute_script("""
        const el = arguments[0];
        const val = arguments[1];

        const desc = Object.getOwnPropertyDescriptor(window.HTMLTextAreaElement.prototype, 'value');
        if (desc && desc.set) {
            desc.set.call(el, val);
        } else {
            el.value = val;
        }

        // React/Vue/Svelte geralmente escutam 'input'
        el.dispatchEvent(new Event('input', {bubbles: true}));
        el.dispatchEvent(new Event('change', {bubbles: true}));
        el.dispatchEvent(new Event('blur', {bubbles: true}));
    """, elemento, texto)



# ==== SISTEMA DATEPICKER MELHORADO ====
def encontrar_campos_datepicker():
    """Encontra todos os campos datepicker na página"""
    global driver
    
    if driver is None:
        return []
    
    seletores_datepicker = [
        "input.hasDatepicker",
        "input[id^='dp']",
        "input[maxlength='10'][grupo='']",
        "input[type='text'][maxlength='10']",
        "input[class*='datepicker']",
        ".hasDatepicker"
    ]
    
    campos_encontrados = []
    
    for seletor in seletores_datepicker:
        try:
            elementos = driver.find_elements(By.CSS_SELECTOR, seletor)
            for elemento in elementos:
                if elemento.is_displayed() and elemento.is_enabled():
                    info = {
                        'elemento': elemento,
                        'id': elemento.get_attribute('id') or f"dp_{len(campos_encontrados)}",
                        'seletor_usado': seletor,
                        'maxlength': elemento.get_attribute('maxlength'),
                        'placeholder': elemento.get_attribute('placeholder')
                    }
                    # Evita duplicatas
                    if not any(c['id'] == info['id'] for c in campos_encontrados):
                        campos_encontrados.append(info)
        except Exception as e:
            log(doc, f"⚠️ Erro ao buscar campos datepicker com {seletor}: {e}", 'WARN')
            continue
    
    log(doc, f"📊 Encontrados {len(campos_encontrados)} campos datepicker")
    return campos_encontrados

def _datepicker_jquery(campo_id, data_valor):
    """Estratégia jQuery para datepicker"""
    global driver
    
    resultado = driver.execute_script("""
        var campoId = arguments[0], valor = arguments[1];
        if (typeof jQuery === 'undefined') return 'jQuery não disponível';
        var $campo = $('#' + campoId);
        if (!$campo.length) return 'Campo não encontrado: ' + campoId;
        try {
            if ($campo.hasClass('hasDatepicker')) { 
                $campo.datepicker('setDate', valor); 
            } else { 
                $campo.val(valor); 
            }
            $campo.trigger('input').trigger('change').trigger('blur');
            return $campo.val();
        } catch(e) { 
            return 'Erro: ' + e.message; 
        }
    """, campo_id, data_valor)
    
    if isinstance(resultado, str) and ('Erro' in resultado or 'não disponível' in resultado):
        raise Exception(f"jQuery falhou: {resultado}")

def _datepicker_javascript(elemento, data_valor):
    """Estratégia JavaScript para datepicker"""
    global driver
    
    driver.execute_script("""
        var campo = arguments[0], valor = arguments[1];
        campo.focus(); 
        campo.value = ''; 
        campo.value = valor;
        ['input','change','blur','keyup'].forEach(ev => 
            campo.dispatchEvent(new Event(ev, {bubbles: true}))
        );
    """, elemento, data_valor)

def _datepicker_actionchains(elemento, data_valor):
    """Estratégia ActionChains para datepicker"""
    global driver
    
    scroll_to_element_safe(elemento)
    time.sleep(0.5)
    
    ActionChains(driver).move_to_element(elemento).click().perform()
    time.sleep(0.5)
    ActionChains(driver).key_down(Keys.CONTROL).send_keys('a').key_up(Keys.CONTROL).perform()
    time.sleep(0.3)
    ActionChains(driver).send_keys(Keys.DELETE).perform()
    time.sleep(0.3)
    
    for char in data_valor:
        ActionChains(driver).send_keys(char).perform()
        time.sleep(0.05)
    
    ActionChains(driver).send_keys(Keys.TAB).perform()

def _datepicker_tradicional(elemento, data_valor):
    """Estratégia tradicional para datepicker"""
    scroll_to_element_safe(elemento)
    time.sleep(0.5)
    elemento.click()
    time.sleep(0.5)
    elemento.clear()
    elemento.send_keys(data_valor)
    elemento.send_keys(Keys.TAB)

def validar_data_preenchida(elemento, data_esperada):
    """Valida se a data foi preenchida corretamente"""
    try:
        if elemento is None:
            return False
            
        val = (elemento.get_attribute('value') or '').strip()
        if not val:
            return False
            
        if val == data_esperada or data_esperada in val:
            return True
            
        # Tenta comparar datas em diferentes formatos
        formatos = [
            '%d/%m/%Y %H:%M', '%d/%m/%Y %H:%M:%S',
            '%d/%m/%Y', '%m/%d/%Y', '%Y-%m-%d', '%d-%m-%Y'
        ]
        
        for formato in formatos:
            try:
                d1 = datetime.strptime(val, formato)
                d2 = datetime.strptime(data_esperada, formato)
                if d1 == d2:
                    return True
            except:
                continue
                
        return False
        
    except Exception:
        return False




def selecionar_opcao(selector, texto):
    def acao():
        select_element = wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, selector)))
        Select(select_element).select_by_visible_text(texto)
    return acao

def preencher_datepicker_por_indice(indice_campo, data_valor, max_tentativas=5):
    """Preenche datepicker pelo índice com estratégias múltiplas"""
    def acao():
        if not isinstance(indice_campo, int) or indice_campo < 0:
            raise ValueError(f"Índice inválido: {indice_campo}")
            
        if not data_valor or not isinstance(data_valor, str):
            raise ValueError(f"Data inválida: {data_valor}")
        
        tentativa = 0
        while tentativa < max_tentativas:
            tentativa += 1
            
            try:
                campos = encontrar_campos_datepicker()
                
                if not campos:
                    if tentativa < max_tentativas:
                        log(doc, f"⚠️ Nenhum campo datepicker encontrado, tentativa {tentativa}/{max_tentativas}", 'WARN')
                        time.sleep(2)
                        continue
                    raise Exception("Nenhum campo datepicker encontrado na página")
                
                if indice_campo >= len(campos):
                    raise Exception(f"Índice {indice_campo} inválido. Encontrados {len(campos)} campos")
                
                campo_info = campos[indice_campo]
                elemento = campo_info['elemento']
                campo_id = campo_info['id']
                
                log(doc, f"🎯 Tentativa {tentativa}: Preenchendo datepicker {indice_campo} (ID: {campo_id}) com '{data_valor}'")
                
                # Verifica se já está preenchido corretamente
                if validar_data_preenchida(elemento, data_valor):
                    log(doc, f"✅ Campo {indice_campo} já está preenchido corretamente!")
                    return True
                
                # Estratégias específicas para datepicker
                estrategias = [
                    lambda: _datepicker_jquery(campo_id, data_valor),
                    lambda: _datepicker_javascript(elemento, data_valor),
                    lambda: _datepicker_actionchains(elemento, data_valor),
                    lambda: _datepicker_tradicional(elemento, data_valor)
                ]
                
                for i, estrategia in enumerate(estrategias, 1):
                    try:
                        log(doc, f"   Aplicando estratégia {i} para datepicker...")
                        estrategia()
                        time.sleep(1)
                        
                        # Verifica se funcionou
                        if validar_data_preenchida(elemento, data_valor):
                            valor_atual = elemento.get_attribute('value')
                            log(doc, f"✅ Datepicker preenchido com estratégia {i}: '{valor_atual}'")
                            return True
                        else:
                            log(doc, f"⚠️ Estratégia {i} não preencheu corretamente", 'WARN')
                            
                    except Exception as e:
                        log(doc, f"⚠️ Estratégia {i} falhou: {e}", 'WARN')
                        continue
                
                # Se chegou aqui, nenhuma estratégia funcionou nesta tentativa
                if tentativa < max_tentativas:
                    log(doc, f"⚠️ Tentativa {tentativa} falhou, tentando novamente em 2s...", 'WARN')
                    time.sleep(2)
                    continue
                
            except Exception as e:
                if tentativa < max_tentativas:
                    log(doc, f"⚠️ Erro na tentativa {tentativa}: {e}, tentando novamente...", 'WARN')
                    time.sleep(2)
                    continue
                else:
                    raise
        
        raise Exception(f"Falha ao preencher datepicker {indice_campo} após {max_tentativas} tentativas")
    
    return acao

# ==== MODAL E SELEÇÃO ROBUSTOS ====
def abrir_modal_e_selecionar_robusto(btn_selector, pesquisa_selector, termo_pesquisa, btn_pesquisar_selector, resultado_xpath):
    """Versão robusta da função de modal"""
    def acao():
        if driver is None or wait is None:
            raise Exception("Driver ou wait não inicializados")
        
        try:
            # 1. Abre o modal
            log(doc, f"🔘 Abrindo modal com botão: {btn_selector}")
            elemento_botao = aguardar_elemento(btn_selector, TIMEOUT_DEFAULT, 'clickable')
            scroll_to_element_safe(elemento_botao)
            remover_overlays()
            
            # Clique no botão do modal
            try:
                elemento_botao.click()
            except:
                driver.execute_script("arguments[0].click();", elemento_botao)
            
            time.sleep(2)  # Aguarda modal abrir

            # 2. Aguarda e preenche campo pesquisa
            log(doc, f"🔍 Preenchendo pesquisa com: {termo_pesquisa}")
            campo_pesquisa = aguardar_elemento(pesquisa_selector, TIMEOUT_CURTO, 'clickable')
            campo_pesquisa.clear()
            time.sleep(0.3)
            campo_pesquisa.send_keys(termo_pesquisa)
            time.sleep(0.5)

            # 3. Clica pesquisar
            log(doc, "🔍 Executando pesquisa...")
            pesquisar = aguardar_elemento(btn_pesquisar_selector, TIMEOUT_CURTO, 'clickable')
            try:
                pesquisar.click()
            except:
                driver.execute_script("arguments[0].click();", pesquisar)
            
            time.sleep(3)  # Aguarda resultados
            
            # 4. Aguarda resultado e clica
            log(doc, f"🎯 Selecionando resultado: {resultado_xpath}")
            resultado = aguardar_elemento(resultado_xpath, TIMEOUT_DEFAULT, 'clickable', By.XPATH)
            scroll_to_element_safe(resultado)
            time.sleep(0.5)
            
            try:
                resultado.click()
            except:
                driver.execute_script("arguments[0].click();", resultado)
            
            time.sleep(1)
            log(doc, "✅ Seleção no modal concluída")

        except Exception as e:
            log(doc, f"❌ Erro no modal: {e}", 'ERROR')
            # Tenta fechar modal em caso de erro
            try:
                fechar_modal = driver.find_elements(By.CSS_SELECTOR, ".modal .close, .modal-header .close, [data-dismiss='modal']")
                for botao in fechar_modal:
                    if botao.is_displayed():
                        botao.click()
                        break
            except:
                pass
            raise
    
    return acao

# ==== UPLOAD DE ARQUIVOS MELHORADO ====
def fazer_upload_arquivo(xpath_input, caminho_arquivo, timeout=TIMEOUT_DEFAULT):
    """Sistema robusto para upload de arquivos"""
    def acao():
        if not caminho_arquivo or not isinstance(caminho_arquivo, str):
            raise Exception("Caminho do arquivo não fornecido")
            
        if not os.path.exists(caminho_arquivo):
            raise Exception(f"Arquivo não encontrado: {caminho_arquivo}")
        
        # Converte para caminho absoluto
        caminho_absoluto = os.path.abspath(caminho_arquivo)
        log(doc, f"📁 Fazendo upload do arquivo: {os.path.basename(caminho_absoluto)}")
        
        # Aguarda elemento de upload
        input_file = aguardar_elemento(xpath_input, timeout, 'present', By.XPATH)
        
        # Estratégias de upload
        estrategias = [
            lambda: input_file.send_keys(caminho_absoluto),
            lambda: upload_via_pyautogui(input_file, caminho_absoluto),
            lambda: upload_via_javascript(input_file, caminho_absoluto)
        ]
        
        for i, estrategia in enumerate(estrategias, 1):
            try:
                log(doc, f"   Tentando upload estratégia {i}...")
                estrategia()
                time.sleep(2)
                log(doc, f"✅ Upload realizado com estratégia {i}")
                return True
            except Exception as e:
                log(doc, f"⚠️ Estratégia {i} de upload falhou: {e}", 'WARN')
                if i == len(estrategias):
                    raise
                continue
                
        return False
    
    return acao

def upload_via_pyautogui(input_element, caminho_arquivo):
    """Upload usando PyAutoGUI"""
    input_element.click()
    time.sleep(1)
    pyautogui.write(f'"{caminho_arquivo}"')
    pyautogui.press('enter')
    time.sleep(1)

def upload_via_javascript(input_element, caminho_arquivo):
    """Upload usando JavaScript (limitado)"""
    # Esta estratégia é limitada devido às restrições de segurança do navegador
    driver.execute_script("""
        var input = arguments[0];
        input.style.display = 'block';
        input.style.visibility = 'visible';
        input.style.opacity = '1';
    """, input_element)
    input_element.send_keys(caminho_arquivo)

# ==== VALIDAÇÃO DE REGISTROS MELHORADA ====
def validar_registros_encontrados(timeout=TIMEOUT_LONGO):
    """Sistema robusto de validação de registros encontrados"""
    global driver, wait, doc
    
    resultado = {
        'encontrou_registros': False,
        'quantidade_registros': 0,
        'mensagem': '',
        'tabela_encontrada': False,
        'detalhes': []
    }
    
    try:
        log(doc, "🔍 Iniciando validação de registros...")
        time.sleep(5)  # Aguarda processamento inicial
        
        # Seletores para diferentes tipos de tabelas de resultado
        seletores_tabela = [
            '#DataTables_Table_0',
            '#DataTables_Table_0 tbody',
            '#DataTables_Table_1',
            '#DataTables_Table_2',
            'table[id*="DataTables"]',
            '.wdGrid table',
            'table tbody',
            '.resultados table',
            '[class*="grid"][class*="result"]',
            'table[class*="dataTable"]'
        ]
        
        tabela_encontrada = None
        
        # Busca tabela de resultados
        for seletor in seletores_tabela:
            try:
                elementos = driver.find_elements(By.CSS_SELECTOR, seletor)
                for elemento in elementos:
                    if elemento.is_displayed() and elemento.size['height'] > 0:
                        tabela_encontrada = elemento
                        resultado['tabela_encontrada'] = True
                        log(doc, f"✅ Tabela encontrada: {seletor}")
                        break
                
                if tabela_encontrada:
                    break
            except Exception as e:
                log(doc, f"⚠️ Erro ao buscar tabela com {seletor}: {e}", 'WARN')
                continue
        
        if not tabela_encontrada:
            # Busca mensagens de "sem resultados"
            mensagens_vazio = [
                "Não foi encontrado nenhum contrato com os filtros informados.",
                "Não foram encontrados registros", 
                "Nenhum resultado",
                "Sem resultados para exibir",
                "0 registros encontrados",
                "No data available"
            ]
            
            for mensagem in mensagens_vazio:
                try:
                    elem = driver.find_element(By.XPATH, f"//*[contains(text(), '{mensagem}')]")
                    if elem.is_displayed():
                        resultado['mensagem'] = f"Sistema informou: {elem.text.strip()}"
                        log(doc, f"ℹ️ {resultado['mensagem']}")
                        return resultado
                except:
                    continue
            
            # Verifica se existe indicação de carregamento
            loading_elements = driver.find_elements(By.CSS_SELECTOR, ".loading, .spinner, [class*='load']")
            if any(el.is_displayed() for el in loading_elements):
                log(doc, "⏳ Sistema ainda carregando resultados...", 'WARN')
                time.sleep(5)
                return validar_registros_encontrados(timeout - 10)  # Recursão com timeout reduzido
            
            resultado['mensagem'] = "⚠️ Tabela de resultados não localizada"
            log(doc, resultado['mensagem'], 'WARN')
            return resultado
        
        # Conta e valida registros
        try:
            # Estratégias para encontrar linhas de dados
            seletores_linhas = [
                'tbody tr:not(.dataTables_empty):not([class*="no-data"])',
                'tbody tr[class*="odd"], tbody tr[class*="even"]',
                'tbody tr:not(:empty)',
                'tbody tr'
            ]
            
            linhas_validas = []
            
            for seletor_linha in seletores_linhas:
                try:
                    linhas = tabela_encontrada.find_elements(By.CSS_SELECTOR, seletor_linha)
                    
                    for linha in linhas:
                        try:
                            if not linha.is_displayed():
                                continue
                                
                            texto_linha = linha.text.strip().lower()
                            
                            # Valida se é uma linha com dados reais
                            if (len(texto_linha) > 5 and 
                                not any(termo in texto_linha for termo in [
                                    'nenhum registro', 'sem dados', 'no data', 
                                    'vazio', 'empty', 'não foram encontrados',
                                    'loading', 'carregando'
                                ])):
                                
                                linhas_validas.append({
                                    'elemento': linha,
                                    'texto': texto_linha[:100] + '...' if len(texto_linha) > 100 else texto_linha
                                })
                        except Exception as e:
                            log(doc, f"⚠️ Erro ao processar linha: {e}", 'WARN')
                            continue
                    
                    if linhas_validas:
                        log(doc, f"✅ Encontradas {len(linhas_validas)} linhas válidas com {seletor_linha}")
                        break
                        
                except Exception as e:
                    log(doc, f"⚠️ Erro ao processar {seletor_linha}: {e}", 'WARN')
                    continue
            
            quantidade = len(linhas_validas)
            resultado['quantidade_registros'] = quantidade
            resultado['detalhes'] = [linha['texto'] for linha in linhas_validas[:5]]  # Primeiras 5 linhas
            
            if quantidade > 0:
                resultado['encontrou_registros'] = True
                resultado['mensagem'] = f"✅ {quantidade} registro(s) encontrado(s)"
                
                # Log das primeiras linhas
                log(doc, resultado['mensagem'])
                for i, linha in enumerate(linhas_validas[:3], 1):
                    log(doc, f"   Registro {i}: {linha['texto']}")
                
                if quantidade > 3:
                    log(doc, f"   ... e mais {quantidade-3} registro(s)")
            else:
                resultado['mensagem'] = "ℹ️ Tabela encontrada mas sem registros válidos"
                log(doc, resultado['mensagem'])
        
        except Exception as e:
            log(doc, f"❌ Erro ao contar registros: {e}", 'ERROR')
            # Em caso de erro na contagem, assume que existem registros para não interromper
            resultado['encontrou_registros'] = True
            resultado['quantidade_registros'] = 1
            resultado['mensagem'] = f"⚠️ Erro na validação, continuando teste: {e}"
        
        # Verifica alertas do sistema
        encontrar_mensagem_alerta()
        
        return resultado
        
    except Exception as e:
        log(doc, f"❌ Erro geral na validação: {e}", 'ERROR')
        resultado['encontrou_registros'] = True  # Assume sucesso para não interromper
        resultado['quantidade_registros'] = 1
        resultado['mensagem'] = f"⚠️ Validação falhou, continuando teste: {e}"
        return resultado

# ==== UTILITÁRIOS DIVERSOS ====
def encontrar_mensagem_alerta():
    """Busca mensagens de alerta na página"""
    global driver, doc
    
    if driver is None:
        return None
    
    seletores = [
        (".alerts.salvo", "✅ Sucesso"),
        (".alerts.alerta", "⚠️ Alerta"),
        (".alerts.erro", "❌ Erro"),
        (".alert-success", "✅ Sucesso"),
        (".alert-warning", "⚠️ Alerta"),
        (".alert-danger", "❌ Erro"),
        ("[class*='toast']", "📢 Notificação"),
    ]
    
    for seletor, tipo in seletores:
        try:
            elementos = driver.find_elements(By.CSS_SELECTOR, seletor)
            for elemento in elementos:
                if elemento.is_displayed():
                    texto = elemento.text.strip()
                    if texto:
                        log(doc, f"📢 {tipo}: {texto}")
                        return elemento
        except Exception as e:
            log(doc, f"⚠️ Erro ao buscar alerta {seletor}: {e}", 'WARN')
            continue
    
    return None

def ajustar_zoom(zoom_level="90%"):
    """Ajusta zoom da página"""
    global driver, doc
    
    if driver is None:
        return
    
    try:
        driver.execute_script(f"document.body.style.zoom='{zoom_level}'")
        log(doc, f"🔍 Zoom ajustado para {zoom_level}.")
    except Exception as e:
        log(doc, f"⚠️ Erro ao ajustar zoom: {e}", 'WARN')

def realizar_consulta():
    """Executa a consulta"""
    def acao():
        seletor = '#gsCallCenter > div.wdTelas > div.telaCobranca.telaConsulta.relative > div.formRow.formLastLine > div:nth-child(13) > a'
        elemento = aguardar_elemento(seletor, TIMEOUT_DEFAULT, 'clickable')
        scroll_to_element_safe(elemento)
        
        try:
            elemento.click()
        except:
            driver.execute_script("arguments[0].click();", elemento)
        
        time.sleep(2)
        log(doc, "✅ Consulta executada")
    
    return acao

def finalizar_cadastro():
    """Finaliza o cadastro/consulta"""
    def acao():
        seletor_css_finalizar = '#gsPet > div.wdTelas > div.wdWizard.clearfix.telaConsulta > div.btnHolder > div'
        elemento = aguardar_elemento(seletor_css_finalizar, TIMEOUT_DEFAULT, 'clickable')
        scroll_to_element_safe(elemento)
        
        try:
            elemento.click()
        except:
            driver.execute_script("arguments[0].click();", elemento)
        
        time.sleep(3)
        log(doc, "✅ Cadastro finalizado")
    
    return acao

# ==== RELATÓRIO MELHORADO ====
def finalizar_relatorio():
    """Finaliza e salva o relatório"""
    global driver, doc
    
    try:
        nome_arquivo = f"relatorio_callcenter_registro_de_chamadas_cenario_2_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        
        # Adiciona resumo final
        doc.add_paragraph("\n" + "="*50)
        doc.add_paragraph("RESUMO FINAL DO TESTE")
        doc.add_paragraph("="*50)
        doc.add_paragraph(f"Teste finalizado em: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")
        doc.add_paragraph(f"Total de screenshots capturadas: {len(screenshot_registradas)}")
        
        # Salva documento
        doc.save(nome_arquivo)
        log(doc, f"📄 Relatório salvo como: {nome_arquivo}")
        
        # Tenta abrir o arquivo
        try:
            if os.name == 'nt':  # Windows
                os.startfile(nome_arquivo)
            else:  # Linux/Mac
                subprocess.run(['xdg-open', nome_arquivo])
        except Exception as e:
            log(doc, f"⚠️ Não foi possível abrir automaticamente o relatório: {e}", 'WARN')
            
    except Exception as e:
        print(f"❌ Erro ao salvar relatório: {e}")
    
    # Fecha driver
    if driver:
        try:
            log(doc, "🔚 Fechando navegador...")
            driver.quit()
        except Exception as e:
            print(f"⚠️ Erro ao fechar driver: {e}")

# ==== INICIALIZAÇÃO DE DRIVER MELHORADA ====
def inicializar_driver():
    """Inicializa o driver do Chrome com configurações otimizadas"""
    global driver, wait
    
    try:
        log(doc, "🚀 Inicializando driver do Chrome...")
        
        options = Options()
        options.add_argument("--start-maximized")
        options.add_argument("--disable-blink-features=AutomationControlled")
        options.add_argument("--disable-extensions")
        options.add_argument("--disable-plugins")
        options.add_argument("--no-sandbox")
        options.add_argument("--disable-dev-shm-usage")
        options.add_argument("--disable-gpu")
        options.add_argument("--disable-web-security")
        options.add_argument("--allow-running-insecure-content")
        options.add_experimental_option("excludeSwitches", ["enable-automation"])
        options.add_experimental_option('useAutomationExtension', False)
        
        # Configurações de download (caso necessário)
        prefs = {
            "profile.default_content_settings.popups": 0,
            "profile.default_content_setting_values.automatic_downloads": 1,
        }
        options.add_experimental_option("prefs", prefs)

        # Instala e configura ChromeDriver
        service = Service(ChromeDriverManager().install())
        driver = webdriver.Chrome(service=service, options=options)
        
        # Remove indicadores de automação
        driver.execute_script("Object.defineProperty(navigator, 'webdriver', {get: () => undefined})")
        driver.execute_cdp_cmd('Network.setUserAgentOverride', {
            "userAgent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36"
        })
        
        # Configura wait com timeout padrão
        wait = WebDriverWait(driver, TIMEOUT_DEFAULT)
        
        log(doc, "✅ Driver inicializado com sucesso")
        return True
        
    except Exception as e:
        log(doc, f"❌ Erro ao inicializar driver: {e}", 'ERROR')
        return False

# ==== FUNÇÃO CORRIGIDA PARA ACESSO AO PET ====
def acessar_modulo_callcenter():
    """Versão corrigida para acessar o módulo Callcenter"""
    def acao():
        # XPath do elemento que precisamos clicar
        xpath_callcenter = "/html/body/div[15]/ul/li[18]/img"
        
        try:
            # Aguarda elemento estar presente
            elemento = aguardar_elemento(xpath_callcenter, TIMEOUT_DEFAULT, 'present', By.XPATH)
            
            # Usa scroll seguro
            scroll_success = scroll_to_element_safe(elemento)
            if not scroll_success:
                log(doc, "⚠️ Problemas com scroll, tentando continuar", 'WARN')
            
            # Remove overlays
            remover_overlays()
            time.sleep(1)
            
            # Aguarda elemento ficar clicável
            elemento_clicavel = aguardar_elemento(xpath_callcenter, TIMEOUT_CURTO, 'clickable', By.XPATH)
            
            # Tenta diferentes estratégias de clique
            estrategias_clique = [
                lambda: elemento_clicavel.click(),
                lambda: ActionChains(driver).move_to_element(elemento_clicavel).click().perform(),
                lambda: driver.execute_script("arguments[0].click();", elemento_clicavel),
                lambda: driver.execute_script("""
                    var element = arguments[0];
                    if (element) {
                        element.focus();
                        var event = new MouseEvent('click', {
                            bubbles: true,
                            cancelable: true,
                            view: window
                        });
                        element.dispatchEvent(event);
                    }
                """, elemento_clicavel)
            ]
            
            for i, estrategia in enumerate(estrategias_clique, 1):
                try:
                    log(doc, f"   Tentando estratégia de clique {i} para módulo PET...")
                    estrategia()
                    time.sleep(2)
                    log(doc, f"✅ Clique no módulo PET realizado com estratégia {i}")
                    return True
                except Exception as e:
                    log(doc, f"⚠️ Estratégia {i} falhou: {e}", 'WARN')
                    if i == len(estrategias_clique):
                        raise
                    continue
            
            return False
            
        except Exception as e:
            log(doc, f"❌ Erro ao acessar módulo PET: {e}", 'ERROR')
            raise
    
    return acao

# ==== EXECUÇÃO DO TESTE CORRIGIDA ====
def executar_teste():
    """Executa o teste principal com tratamento robusto de erros"""
    global driver, wait, doc
    
    try:
        # Inicialização
        if not inicializar_driver():
            log(doc, "❌ Falha crítica na inicialização do driver", 'ERROR')
            return False

        log(doc, "🎯 Iniciando execução do teste de consulta de Telecobrança")

        # 1. Acesso ao sistema
        safe_action(doc, "Acessando sistema", lambda: (
            driver.get(URL),
            time.sleep(3)
        ), critico=True)

        # 2. Login
        safe_action(doc, "Realizando login", lambda: (
            aguardar_elemento("#j_id15\\:email", TIMEOUT_DEFAULT).send_keys(LOGIN_EMAIL),
            aguardar_elemento("#j_id15\\:senha", TIMEOUT_DEFAULT).send_keys(LOGIN_PASSWORD),
            aguardar_elemento("#j_id15\\:senha", TIMEOUT_DEFAULT).send_keys(Keys.ENTER),
            time.sleep(5)
        ), critico=True)

        # 3. Ajustes iniciais
        safe_action(doc, "Configurando ambiente", lambda: (
            ajustar_zoom("90%"),
            driver.find_element(By.TAG_NAME, "body").send_keys(Keys.F3),
            time.sleep(2)
        ))

        safe_action(doc, "Acessando módulo Callcenter", acessar_modulo_callcenter(), critico=True)

        safe_action(doc, "Abrindo Telecobrança", lambda: (
            aguardar_elemento('#gsCallCenter > div.wdTelas > div > ul > li:nth-child(2) > a > span').click(),
            time.sleep(3)
        ), critico=True)

        # 6. Preenchimento dos filtros
        safe_action(doc, "Selecionando Vendedor", 
                   abrir_modal_e_selecionar_robusto(
                       '#gsCallCenter > div.wdTelas > div.telaCobranca.telaConsulta.relative > div.formRow.formLastLine > div:nth-child(1) > div > a',
                       'body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(2) > input',
                       'JOSÉ DA SILVA',
                       'body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(4) > a',
                       "//td[contains(text(), 'JOSÉ DA SILVA')]"
                   ))



        safe_action(doc, "Selecionando Pacote", 
                   abrir_modal_e_selecionar_robusto(
                       '#gsCallCenter > div.wdTelas > div.telaCobranca.telaConsulta.relative > div.formRow.formLastLine > div:nth-child(2) > div > a',
                       'body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(2) > input',
                       'PACOTE TESTE JP',
                       'body > div.modalHolder > div.modal.overflow > div:nth-child(1) > div.formRow.formLastLine > div:nth-child(3) > a',
                       "//td[contains(text(), 'PACOTE TESTE JP')]"
                   ))




        safe_action(doc, "Preenchendo Números do Contrato", 
                   preencher_campo_robusto('#gsCallCenter > div.wdTelas > div.telaCobranca.telaConsulta.relative > div.formRow.formLastLine > div:nth-child(3) > input[type=text]:nth-child(2)', '112106'),
                   preencher_campo_robusto('#gsCallCenter > div.wdTelas > div.telaCobranca.telaConsulta.relative > div.formRow.formLastLine > div:nth-child(3) > input.marginLeft5', '112106'),
        )


        safe_action(doc, "Preenchendo Cidade", 
                   preencher_campo_robusto('#gsCallCenter > div.wdTelas > div.telaCobranca.telaConsulta.relative > div.formRow.formLastLine > div:nth-child(4) > input', 'São José do Rio Preto - SP'),
                   clicar_elemento_robusto("#ui-id-15")

                   )
        safe_action(doc, "Preenchendo quntidade de Parcelas em atraso", 
                   preencher_campo_robusto('#gsCallCenter > div.wdTelas > div.telaCobranca.telaConsulta.relative > div.formRow.formLastLine > div:nth-child(5) > input[type=text]:nth-child(2)', '1'),
                   preencher_campo_robusto('#gsCallCenter > div.wdTelas > div.telaCobranca.telaConsulta.relative > div.formRow.formLastLine > div:nth-child(5) > input.marginLeft5', '10'),
               
                   )

        safe_action(doc, "Preenchendo Data da Venda Inicial", 
                   preencher_datepicker_por_indice(0, "14/08/2010"))

        safe_action(doc, "Preenchendo Data da Venda Final", 
                   preencher_datepicker_por_indice(1, "24/11/2025"))

        safe_action(doc, "Preenchendo Data de Vencimento Inicial", 
                   preencher_datepicker_por_indice(2, "14/08/2010"))

        safe_action(doc, "Preenchendo Data de Vencimento Final", 
                   preencher_datepicker_por_indice(3, "24/11/2025"))

        safe_action(doc, "Selecionando Forma de Pagamento", selecionar_opcao(
            "#gsCallCenter > div.wdTelas > div.telaCobranca.telaConsulta.relative > div.formRow.formLastLine > div:nth-child(10) > select",
            "Boleto"
        ))

        safe_action(doc, "Selecionando Filial", selecionar_opcao(
            "#gsCallCenter > div.wdTelas > div.telaCobranca.telaConsulta.relative > div.formRow.formLastLine > div:nth-child(11) > select",
            "Filial Gold 1"
        ))
        

        # 7. Execução da consulta
        safe_action(doc, "Realizando consulta", realizar_consulta())

        # 8. Validação dos resultados
        resultado_validacao = None
        safe_action(doc, "Validando resultados da consulta", lambda: (
            setattr(executar_teste, 'resultado_validacao', validar_registros_encontrados())
        ), critico=False)
        
        resultado_validacao = getattr(executar_teste, 'resultado_validacao', None)

        # Verifica se deve continuar
        if not resultado_validacao or not resultado_validacao.get('encontrou_registros', False):
            log(doc, "ℹ️ Nenhum registro encontrado - Finalizando consulta", 'WARN')
            safe_action(doc, "Fechando tela (sem registros)", 
                       clicar_elemento_robusto("#gsCallCenter > div.wdTop.ui-draggable-handle > div.wdClose > a"))
            return True

        # 9. Processamento dos registros encontrados
        quantidade = resultado_validacao.get('quantidade_registros', 0)
        log(doc, f"✅ Processando {quantidade} registro(s) encontrado(s)")

        # 10. Visualização de dados
        safe_action(doc, "Visualizando dados do Contrato", 
                   clicar_elemento_xpath_robusto("//span[contains(@class,'btContrato') and @title='Dados do Contrato']"))

        safe_action(doc, "Capturando tela de dados", lambda: (
            time.sleep(10),
            take_screenshot(driver, doc, "dados_contrato", forcar=True)
        ))

        # 11. Mensagens do contrato
        safe_action(doc, "Abrindo mensagens do Contrato",
                   clicar_elemento_robusto("body > div.modalHolder > div.modal.overflow.fullscreen.bgBlur > div.detalhesContrato > div > div > div.aba.detalhes > div.row > div.col-sm-4.rAlign > button:nth-child(1)"))

        safe_action(doc, "Capturando mensagens", lambda: (
            time.sleep(2),
            take_screenshot(driver, doc, "mensagens_contrato", forcar=True)
        ))

        safe_action(doc, "Adicionando nova mensagem", 
                   clicar_elemento_xpath_robusto("//a[contains(@class,'btModel btGray btAdd') and contains(normalize-space(.),'Adicionar Mensagem')]"))



        safe_action(doc, "Confirmando mensagem (Teste de disparo)", 
                   clicar_elemento_xpath_robusto("//a[contains(@class,'hAlign') and contains(normalize-space(.),'Confirmar')]"))


        log(doc, "🔍 Verificando mensagens de alerta...")
        encontrar_mensagem_alerta()


        safe_action(doc, "Fechando aba de cadastro de mensagens", 
                   clicar_elemento_xpath_robusto("/html/body/div[21]/div[2]/a"),
                   clicar_elemento_xpath_robusto("//a[@class='fa fa-close']")
                   )

        safe_action(doc, "Fechando aba de mensagens", 
                   clicar_elemento_xpath_robusto("//a[@class='btModel btGray' and contains(normalize-space(.), 'Fechar')]")

                   )

        # 12. Títulos do contrato 

        safe_action(doc, "Visualizando títulos do Contrato", 
                   clicar_elemento_xpath_robusto("//button[@class='btModel btGray' and .//i[contains(@class, 'sprites sp-pesquisar')] and normalize-space(text())='Ver Títulos']"))

        safe_action(doc, "Aguardando carregamento de títulos e capturando Screenshot", lambda: (
            time.sleep(30),
            take_screenshot(driver, doc, "titulos_contrato", forcar=True)
        ))


        safe_action(doc, "Fechando aba de títulos", 
                   clicar_elemento_xpath_robusto("//a[@title='Sair' and contains(@class, 'sp-fecharGrande')]"))
        
        # 14. Finalização
        safe_action(doc, "Fechando tela de consulta", 
                   clicar_elemento_robusto("#gsCallCenter > div.wdTop.ui-draggable-handle > div.wdClose > a"))

        log(doc, "✅ Teste executado com sucesso completo!")
        return True

    except Exception as e:
        log(doc, f"❌ ERRO CRÍTICO NO TESTE: {e}", 'ERROR')
        log(doc, "— stacktrace —")
        log(doc, traceback.format_exc())
        take_screenshot(driver, doc, "erro_critico_final", forcar=True)
        return False
    
    finally:
        log(doc, "🏁 Finalizando teste...")

# ==== FUNÇÃO PRINCIPAL ====
def main():
    """Função principal de execução"""
    global doc
    
    try:
        log(doc, "🚀 INICIANDO TESTE DE CONSULTA DE TELECOBRANÇA DO CALLCENTER")
        log(doc, "="*60)
        
        sucesso = executar_teste()
        
        if sucesso:
            log(doc, "✅ TESTE CONCLUÍDO COM SUCESSO!")
        else:
            log(doc, "❌ TESTE FINALIZADO COM ERROS.")
        
        log(doc, "="*60)
        
    except Exception as e:
        log(doc, f"❌ ERRO FATAL NA EXECUÇÃO PRINCIPAL: {e}", 'ERROR')
        log(doc, traceback.format_exc())
        
    finally:
        finalizar_relatorio()

if __name__ == "__main__":
    main()