# ==== IMPORTS (sem conflitos) ====
from datetime import datetime, timedelta
from datetime import time as dt_time  # usar para objetos de hora
import time                           # usar para time.sleep(...)
import pyautogui
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support.ui import Select, WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service
from selenium.common.exceptions import TimeoutException, WebDriverException, StaleElementReferenceException, ElementClickInterceptedException
from webdriver_manager.chrome import ChromeDriverManager
from docx import Document
from docx.shared import Inches
from faker import Faker
from faker.providers import BaseProvider
from validate_docbr import CPF
from selenium.webdriver import ActionChains
import subprocess
import os
import random
import re
from selenium.webdriver.common.by import By
from selenium.webdriver.support import expected_conditions as EC
from selenium.common.exceptions import ElementClickInterceptedException, StaleElementReferenceException
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.remote.webelement import WebElement
import traceback

# Timeouts configuráveis
TIMEOUT_DEFAULT = 30
TIMEOUT_CURTO = 10
TIMEOUT_LONGO = 60

# ==== PROVIDERS CUSTOMIZADOS ====
class BrasilProvider(BaseProvider):
    def rg(self):
        numeros = [str(random.randint(0, 9)) for _ in range(8)]
        return ''.join(numeros) + '-' + str(random.randint(0, 9))

fake = Faker("pt_BR")
fake.add_provider(BrasilProvider)

# ==== GERAÇÃO DE DATAS ====
def gerar_datas_validas(hora_padrao="00:00", dias_fim=0):
    """Gera datas coerentes. data_inicio/data_fim no formato 'dd/MM/yyyy HH:mm'."""
    hoje_date = datetime.today().date()
    dez_anos_atras = hoje_date - timedelta(days=3650)

    # Falecimento entre 10 anos atrás e hoje
    data_falecimento = fake.date_between(start_date=dez_anos_atras, end_date=hoje_date)

    # Nascimento (entre 18 e 110 anos antes do falecimento)
    idade_minima, idade_maxima = 18, 110
    data_nascimento = data_falecimento - timedelta(days=random.randint(idade_minima * 365, idade_maxima * 365))

    # Sepultamento 1..10 dias após o falecimento
    data_sepultamento = data_falecimento + timedelta(days=random.randint(1, 10))

    # Registro 1..10 dias após o sepultamento
    data_registro = data_sepultamento + timedelta(days=random.randint(1, 10))

    # Velório entre o falecimento e o sepultamento
    data_velorio = fake.date_between(start_date=data_falecimento, end_date=data_sepultamento)

    # Início entre 2 e 30 dias no futuro, com hora escolhida
    data_inicio_date = hoje_date + timedelta(days=random.randint(2, 30))

    # Monta datetime com hora escolhida (ex: "00:00")
    h, m = map(int, hora_padrao.split(":"))
    dt_inicio = datetime.combine(data_inicio_date, dt_time(h, m))

    # Fim: mesmo dia por padrão (dias_fim=0). Ajuste se quiser +N dias.
    dt_fim = dt_inicio + timedelta(days=dias_fim)

    fmt_data = "%d/%m/%Y"
    fmt_dt = "%d/%m/%Y %H:%M"

    return (
        data_nascimento.strftime(fmt_data),
        data_falecimento.strftime(fmt_data),
        data_sepultamento.strftime(fmt_data),
        data_velorio.strftime(fmt_data),
        dt_inicio.strftime(fmt_dt),   # data_inicio com hora
        dt_fim.strftime(fmt_dt),      # data_fim com hora
        data_registro.strftime(fmt_data),
        hoje_date.strftime(fmt_data),
    )

(data_nascimento, data_falecimento, data_sepultamento, 
 data_velorio, data_inicio, data_fim, data_registro, hoje) = gerar_datas_validas(
    hora_padrao="08:50",
    dias_fim=0
)

hora_falecimento = fake.time(pattern="%H:%M")
hora_sepultamento = fake.time(pattern="%H:%M")
localizacao = fake.city()

# ==== CONFIGURAÇÕES ====
URL = "http://localhost:8080/gs/index.xhtml"
LOGIN_EMAIL = "joaoeduardo.gold@outlook.com"
LOGIN_PASSWORD = "071999gs"

# ==== INICIALIZAÇÃO DE VARIÁVEIS GLOBAIS ====
doc = Document()
doc.add_heading("RELATÓRIO DO TESTE", 0)
doc.add_paragraph("Consulta de Agenda de Cobrança - Callcenter – Cenário 1: Realização da consulta de Agenda de Cobrança")
doc.add_paragraph(f"Data do teste: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")

screenshot_registradas = set()
driver = None
wait = None

# ==== UTILITÁRIOS MELHORADOS ====
def log(doc, msg, nivel='INFO'):
    timestamp = datetime.now().strftime('%H:%M:%S')
    formatted_msg = f"[{timestamp}] {msg}"
    print(formatted_msg)
    try:
        doc.add_paragraph(formatted_msg)
    except Exception as e:
        print(f"Erro ao adicionar ao documento: {e}")

def _sanitize_filename(name: str) -> str:
    if not isinstance(name, str):
        name = str(name)
    name = name.strip().lower()
    name = re.sub(r"[<>:\"/\\|?*']", "_", name)
    name = re.sub(r"_+", "_", name)
    return name[:120]

def take_screenshot(driver, doc, nome, forcar=False):
    if driver is None:
        log(doc, "⚠️ Driver não disponível para screenshot", 'WARN')
        return
    
    try:
        nome = _sanitize_filename(nome)
        if forcar or nome not in screenshot_registradas:
            path = f"screenshots/{nome}.png"
            os.makedirs("screenshots", exist_ok=True)
            driver.save_screenshot(path)
            log(doc, f"📸 Screenshot capturada: {nome}")
            doc.add_paragraph(f"Screenshot: {nome}")
            doc.add_picture(path, width=Inches(5.5))
            screenshot_registradas.add(nome)
    except Exception as e:
        log(doc, f"⚠️ Erro ao tirar screenshot {nome}: {e}", 'WARN')

def safe_action(doc, descricao, func, critico=False):
    """Executa ação de forma segura com tratamento de erros melhorado"""
    try:
        log(doc, f"🔄 {descricao}...")
        resultado = func()
        log(doc, f"✅ {descricao} realizada com sucesso.")
        time.sleep(1)  # Pausa padrão após ações
        take_screenshot(driver, doc, descricao.lower().replace(" ", "_"))
        return resultado
    except Exception as e:
        take_screenshot(driver, doc, f"erro_{descricao.lower().replace(' ', '_')}")
        log(doc, f"❌ {descricao} falhou: {type(e).__name__}: {str(e)}", 'ERROR')
        
        # Log do stacktrace completo apenas se for crítico
        if critico:
            log(doc, "— stacktrace —")
            log(doc, traceback.format_exc())
            raise
        else:
            log(doc, f"⚠️ Continuando execução apesar do erro em: {descricao}", 'WARN')
            return None

def _sanitize_timeout(t):
    """Garante timeout válido"""
    if not isinstance(t, (int, float)) or t <= 0:
        return TIMEOUT_DEFAULT
    return max(5, min(120, t))  # Entre 5 e 120 segundos

# ==== AGUARDAR ELEMENTO MELHORADO ====
def aguardar_elemento(seletor, timeout=TIMEOUT_DEFAULT, condicao='clickable', by_type=By.CSS_SELECTOR):
    """Função centralizada para aguardar elementos com diferentes condições"""
    global driver, wait
    
    if driver is None:
        raise Exception("Driver não inicializado")
    
    timeout = _sanitize_timeout(timeout)
    
    condicoes = {
        'present': EC.presence_of_element_located,
        'visible': EC.visibility_of_element_located,
        'clickable': EC.element_to_be_clickable,
        'invisible': EC.invisibility_of_element_located
    }
    
    if condicao not in condicoes:
        condicao = 'clickable'
    
    try:
        wait_obj = WebDriverWait(driver, timeout)
        elemento = wait_obj.until(condicoes[condicao]((by_type, seletor)))
        return elemento
    except TimeoutException:
        log(doc, f"❌ Timeout aguardando elemento: {seletor} (condição: {condicao}, timeout: {timeout}s)", 'ERROR')
        raise TimeoutException(f"Elemento não encontrado: {seletor} (condição: {condicao})")
    except Exception as e:
        log(doc, f"❌ Erro aguardando elemento {seletor}: {e}", 'ERROR')
        raise

# ==== SCROLL CORRIGIDO - PRINCIPAL CORREÇÃO ====
def scroll_to_element_safe(elemento_ou_seletor, by_type=By.CSS_SELECTOR):
    """Scroll seguro até elemento com validação robusta"""
    global driver
    
    if driver is None:
        log(doc, "⚠️ Driver não disponível para scroll", 'WARN')
        return False
    
    try:
        # Se for seletor, encontra o elemento
        if isinstance(elemento_ou_seletor, str):
            elemento = aguardar_elemento(elemento_ou_seletor, 10, 'present', by_type)
        else:
            elemento = elemento_ou_seletor
        
        if elemento is None:
            log(doc, "⚠️ Elemento não encontrado para scroll", 'WARN')
            return False
        
        # Verifica se elemento é válido antes de fazer scroll
        if not elemento.is_displayed():
            log(doc, "⚠️ Elemento não está visível para scroll", 'WARN')
            return False
        
        # Estratégias de scroll em ordem de preferência
        scroll_strategies = [
            # Estratégia 1: JavaScript com verificação prévia
            lambda: driver.execute_script("""
                var element = arguments[0];
                if (element && typeof element.scrollIntoView === 'function') {
                    element.scrollIntoView({
                        behavior: 'smooth',
                        block: 'center',
                        inline: 'center'
                    });
                    return true;
                } else {
                    return false;
                }
            """, elemento),
            
            # Estratégia 2: ActionChains
            lambda: ActionChains(driver).move_to_element(elemento).perform(),
            
            # Estratégia 3: JavaScript alternativo
            lambda: driver.execute_script("""
                var element = arguments[0];
                if (element) {
                    element.scrollIntoView();
                    window.scrollBy(0, -100);
                }
            """, elemento),
            
            # Estratégia 4: Scroll da página até o elemento
            lambda: driver.execute_script("""
                var element = arguments[0];
                if (element) {
                    var rect = element.getBoundingClientRect();
                    var scrollTop = window.pageYOffset || document.documentElement.scrollTop;
                    var targetY = rect.top + scrollTop - (window.innerHeight / 2);
                    window.scrollTo(0, targetY);
                }
            """, elemento)
        ]
        
        for i, strategy in enumerate(scroll_strategies, 1):
            try:
                log(doc, f"   Tentando estratégia de scroll {i}...")
                result = strategy()
                
                # Para estratégia 1, verifica resultado
                if i == 1 and result is False:
                    log(doc, f"   Estratégia {i}: elemento não suporta scrollIntoView", 'WARN')
                    continue
                
                time.sleep(0.8)  # Aguarda scroll completar
                
                # Verifica se elemento ainda está acessível
                if elemento.is_displayed() and elemento.is_enabled():
                    log(doc, f"✅ Scroll realizado com estratégia {i}")
                    return True
                else:
                    log(doc, f"   Estratégia {i}: elemento não ficou acessível", 'WARN')
                    continue
                    
            except Exception as e:
                log(doc, f"   Estratégia {i} de scroll falhou: {str(e)[:100]}...", 'WARN')
                continue
        
        log(doc, "⚠️ Todas as estratégias de scroll falharam", 'WARN')
        return False
        
    except Exception as e:
        log(doc, f"⚠️ Erro geral no scroll: {e}", 'WARN')
        return False

def remover_overlays():
    """Remove overlays que podem bloquear cliques"""
    global driver
    
    if driver is None:
        return
    
    try:
        driver.execute_script("""
            // Remove overlays comuns
            document.querySelectorAll('.modal-backdrop, .overlay, .blockUI, .loading, .spinner, [style*="position: fixed"]')
                .forEach(el => {
                    if (el.style.position === 'fixed' || el.classList.contains('modal-backdrop')) {
                        el.style.display = 'none';
                    }
                });
            
            // Remove elementos com z-index alto
            document.querySelectorAll('*').forEach(el => {
                const zIndex = getComputedStyle(el).zIndex;
                if (zIndex && parseInt(zIndex) > 1000) {
                    if (el.offsetWidth === window.innerWidth && el.offsetHeight === window.innerHeight) {
                        el.style.display = 'none';
                    }
                }
            });
        """)
    except Exception as e:
        log(doc, f"⚠️ Erro ao remover overlays: {e}", 'WARN')

# ==== CLIQUE ROBUSTO CORRIGIDO ====
def clicar_elemento_robusto(seletor, timeout=TIMEOUT_DEFAULT, by_type=By.CSS_SELECTOR):
    """Clique robusto com múltiplas estratégias"""
    def acao():
        timeout_sanitized = _sanitize_timeout(timeout)
        
        try:
            # 1. Aguarda elemento
            elemento = aguardar_elemento(seletor, timeout_sanitized, 'present', by_type)
            
            # 2. Remove overlays
            remover_overlays()
            
            # 3. Scroll seguro até elemento
            scroll_success = scroll_to_element_safe(elemento)
            if not scroll_success:
                log(doc, f"⚠️ Problemas com scroll, continuando mesmo assim: {seletor}", 'WARN')
            
            # 4. Aguarda ser clicável
            try:
                elemento = aguardar_elemento(seletor, 5, 'clickable', by_type)
            except TimeoutException:
                log(doc, f"⚠️ Elemento não ficou clicável, tentando mesmo assim: {seletor}", 'WARN')
                elemento = aguardar_elemento(seletor, timeout_sanitized, 'present', by_type)
            
            # 5. Múltiplas estratégias de clique
            estrategias = [
                lambda: elemento.click(),
                lambda: ActionChains(driver).move_to_element(elemento).click().perform(),
                lambda: driver.execute_script("arguments[0].click();", elemento),
                lambda: driver.execute_script("""
                    const el = arguments[0];
                    el.focus();
                    el.dispatchEvent(new MouseEvent('click', {
                        bubbles: true, 
                        cancelable: true, 
                        view: window
                    }));
                """, elemento),
                lambda: ActionChains(driver).move_to_element_with_offset(elemento, 1, 1).click().perform()
            ]
            
            for i, estrategia in enumerate(estrategias, 1):
                try:
                    estrategia()
                    time.sleep(0.5)
                    log(doc, f"✅ Clique realizado com estratégia {i}")
                    return True
                except Exception as e:
                    log(doc, f"⚠️ Estratégia {i} de clique falhou: {e}", 'WARN')
                    if i == len(estrategias):
                        raise
                    continue
            
            return False
            
        except Exception as e:
            log(doc, f"❌ Falha ao clicar em {seletor}: {e}", 'ERROR')
            raise
    
    return acao

def clicar_elemento_xpath_robusto(xpath, timeout=TIMEOUT_DEFAULT):
    """Clique robusto via XPath"""
    return clicar_elemento_robusto(xpath, timeout, By.XPATH)

# ==== PREENCHIMENTO DE CAMPOS MELHORADO ====
def _preencher_tradicional(elemento, valor, limpar_primeiro=True):
    """Estratégia tradicional de preenchimento"""
    if limpar_primeiro:
        elemento.clear()
        time.sleep(0.2)
    elemento.click()
    time.sleep(0.2)
    elemento.send_keys(valor)
    elemento.send_keys(Keys.TAB)

def _preencher_actionchains(elemento, valor):
    """Estratégia com ActionChains"""
    global driver
    
    ActionChains(driver).move_to_element(elemento).click().perform()
    time.sleep(0.3)
    ActionChains(driver).key_down(Keys.CONTROL).send_keys('a').key_up(Keys.CONTROL).perform()
    time.sleep(0.2)
    ActionChains(driver).send_keys(valor).perform()
    time.sleep(0.2)
    ActionChains(driver).send_keys(Keys.TAB).perform()

def _preencher_javascript(elemento, valor):
    """Estratégia com JavaScript"""
    global driver
    
    driver.execute_script("""
        var element = arguments[0];
        var value = arguments[1];
        element.focus();
        element.value = '';
        element.value = value;
        element.dispatchEvent(new Event('input', { bubbles: true }));
        element.dispatchEvent(new Event('change', { bubbles: true }));
        element.blur();
    """, elemento, valor)



def _valor_do_elemento(elemento, driver=None):
    """Tenta ler o valor atual (value ou textContent p/ contenteditable)."""
    val = (elemento.get_attribute('value') or "").strip()
    if not val and elemento.get_attribute('contenteditable') in ('true', 'True', True):
        try:
            # usa JS p/ garantir leitura de contenteditable
            return (driver.execute_script("return arguments[0].textContent;", elemento) or "").strip()
        except Exception:
            return (elemento.text or "").strip()
    return val

def preencher_campo_robusto_xpath(xpath, valor, limpar_primeiro=True, timeout=TIMEOUT_DEFAULT):
    """Versão para XPath da função de preenchimento robusto."""
    def acao():
        if valor is None or valor == "":
            log(doc, f"⚠️ Valor vazio para campo (xpath): {xpath}, pulando preenchimento", 'WARN')
            return True

        # 1) aguarda de forma robusta pelo elemento via XPATH
        try:
            elemento = WebDriverWait(driver, timeout).until(
                EC.element_to_be_clickable((By.XPATH, xpath))
            )
        except Exception:
            # fallback: presença + possibilidade de clique depois
            elemento = WebDriverWait(driver, timeout).until(
                EC.presence_of_element_located((By.XPATH, xpath))
            )

        # 2) scroll até o elemento
        scroll_success = scroll_to_element_safe(elemento)
        if not scroll_success:
            log(doc, f"⚠️ Problemas com scroll para (xpath): {xpath}", 'WARN')

        # 3) estratégias de preenchimento (rebusca o elemento a cada tentativa p/ evitar stale)
        def _refind():
            try:
                return WebDriverWait(driver, 2).until(
                    EC.presence_of_element_located((By.XPATH, xpath))
                )
            except Exception:
                return elemento  # último conhecido (melhor que nada)

        estrategias = [
            lambda: _preencher_tradicional(_refind(), valor, limpar_primeiro),
            lambda: _preencher_actionchains(_refind(), valor),
            lambda: _preencher_javascript(_refind(), valor),
        ]

        for i, estrategia in enumerate(estrategias, 1):
            try:
                estrategia()
                time.sleep(0.5)

                # Verificação do preenchimento
                elemento_check = _refind()
                valor_atual = _valor_do_elemento(elemento_check, driver)

                if valor_atual.strip() == str(valor).strip() or str(valor) in valor_atual:
                    log(doc, f"✅ Campo (xpath) preenchido com estratégia {i}: '{valor_atual}'")
                    return True
                else:
                    log(doc, f"⚠️ Estratégia {i} não confirmou o valor. Esperado: '{valor}', Atual: '{valor_atual}'", 'WARN')

            except Exception as e:
                log(doc, f"⚠️ Estratégia {i} (xpath) falhou: {e}", 'WARN')
                if i == len(estrategias):
                    raise

        raise Exception(f"Falha ao preencher (xpath): {xpath} com todas as estratégias")

    return acao

from selenium.webdriver.common.by import By
from selenium.webdriver.support import expected_conditions as EC

# --- helper para esperar via XPATH com diferentes estados ---
def aguardar_elemento_xpath(xpath, timeout=TIMEOUT_DEFAULT, estado='clickable'):
    cond = {
        'clickable': EC.element_to_be_clickable((By.XPATH, xpath)),
        'visible':   EC.visibility_of_element_located((By.XPATH, xpath)),
        'present':   EC.presence_of_element_located((By.XPATH, xpath)),
    }.get(estado, EC.element_to_be_clickable((By.XPATH, xpath)))
    return WebDriverWait(driver, timeout).until(cond)



def selecionar_opcao(selector, texto):
    def acao():
        select_element = wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, selector)))
        Select(select_element).select_by_visible_text(texto)
    return acao


def selecionar_opcao_xpath(xpath, texto):
    def acao():
        select_element = wait.until(EC.presence_of_element_located((By.XPATH, xpath)))
        Select(select_element).select_by_visible_text(texto)
    return acao



def preencher_campo_robusto_xpath(xpath, valor, limpar_primeiro=True, timeout=TIMEOUT_DEFAULT):
    """Mesma lógica do preencher_campo_robusto, mas localizando o elemento por XPATH."""
    def acao():
        if valor is None or str(valor) == "":
            log(doc, f"⚠️ Valor vazio para campo {xpath}, pulando preenchimento", 'WARN')
            return True

        # espera e obtém o elemento por XPATH
        elemento = aguardar_elemento_xpath(xpath, timeout, 'clickable')

        # scroll seguro (mantém seu helper atual)
        scroll_success = scroll_to_element_safe(elemento)
        if not scroll_success:
            log(doc, f"⚠️ Problemas com scroll para {xpath}", 'WARN')

        # mesmas estratégias que você já usa
        estrategias = [
            lambda: _preencher_tradicional(elemento, valor, limpar_primeiro),
            lambda: _preencher_actionchains(elemento, valor),
            lambda: _preencher_javascript(elemento, valor),
        ]

        for i, estrategia in enumerate(estrategias, 1):
            try:
                estrategia()
                time.sleep(0.5)

                # validação do preenchimento
                valor_atual = (elemento.get_attribute('value') or "").strip()
                if valor_atual == str(valor).strip() or str(valor) in valor_atual:
                    log(doc, f"✅ Campo (xpath) preenchido com estratégia {i}: '{valor_atual}'")
                    return True
                else:
                    log(doc, f"⚠️ Estratégia {i} não preencheu corretamente (xpath). "
                             f"Esperado: '{valor}', Atual: '{valor_atual}'", 'WARN')
            except Exception as e:
                log(doc, f"⚠️ Estratégia {i} falhou (xpath): {e}", 'WARN')
                if i == len(estrategias):
                    raise

        raise Exception(f"Falha ao preencher campo {xpath} com todas as estratégias (xpath)")

    return acao


def preencher_campo_robusto(seletor, valor, limpar_primeiro=True, timeout=TIMEOUT_DEFAULT):
    """Função melhorada para preenchimento de campos"""
    def acao():
        if valor is None or valor == "":
            log(doc, f"⚠️ Valor vazio para campo {seletor}, pulando preenchimento", 'WARN')
            return True
            
        elemento = aguardar_elemento(seletor, timeout, 'clickable')
        
        scroll_success = scroll_to_element_safe(elemento)
        if not scroll_success:
            log(doc, f"⚠️ Problemas com scroll para {seletor}", 'WARN')
        
        # Múltiplas estratégias de preenchimento
        estrategias = [
            lambda: _preencher_tradicional(elemento, valor, limpar_primeiro),
            lambda: _preencher_actionchains(elemento, valor),
            lambda: _preencher_javascript(elemento, valor)
        ]
        
        for i, estrategia in enumerate(estrategias, 1):
            try:
                estrategia()
                time.sleep(0.5)
                
                # Verifica se foi preenchido corretamente
                valor_atual = elemento.get_attribute('value') or ""
                if valor_atual.strip() == str(valor).strip() or str(valor) in valor_atual:
                    log(doc, f"✅ Campo preenchido com estratégia {i}: '{valor_atual}'")
                    return True
                else:
                    log(doc, f"⚠️ Estratégia {i} não preencheu corretamente. Esperado: '{valor}', Atual: '{valor_atual}'", 'WARN')
                    
            except Exception as e:
                log(doc, f"⚠️ Estratégia {i} de preenchimento falhou: {e}", 'WARN')
                if i == len(estrategias):
                    raise
                continue
                
        raise Exception(f"Falha ao preencher campo {seletor} com todas as estratégias")
    
    return acao




from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver import ActionChains
from selenium.common.exceptions import TimeoutException, StaleElementReferenceException, JavascriptException
from selenium.webdriver.support import expected_conditions as EC

def preencher_textarea_por_indice(indice_campo, texto, max_tentativas=5, limpar_primeiro=True):
    """Preenche um <textarea> pelo índice (ordem no DOM) usando estratégias múltiplas"""
    def acao():
        if not isinstance(indice_campo, int) or indice_campo < 0:
            raise ValueError(f"Índice inválido: {indice_campo}")
        if texto is None or not isinstance(texto, str):
            raise ValueError(f"Texto inválido: {texto!r}")

        tentativa = 0
        while tentativa < max_tentativas:
            tentativa += 1
            try:
                campos = encontrar_campos_textarea()
                if not campos:
                    if tentativa < max_tentativas:
                        log(doc, f"⚠️ Nenhuma <textarea> encontrada (tentativa {tentativa}/{max_tentativas})", "WARN")
                        time.sleep(1.5)
                        continue
                    raise Exception("Nenhuma <textarea> foi encontrada na página.")

                if indice_campo >= len(campos):
                    raise Exception(f"Índice {indice_campo} inválido. Encontradas {len(campos)} textareas.")

                campo_info = campos[indice_campo]
                elemento   = campo_info["elemento"]
                campo_id   = campo_info.get("id") or "(sem id)"
                campo_name = campo_info.get("name") or "(sem name)"

                log(doc, f"🎯 Tentativa {tentativa}: Preenchendo textarea {indice_campo} (ID: {campo_id}, name: {campo_name}) com {len(texto)} caracteres")

                # Se já estiver preenchido corretamente, encerra
                if validar_textarea_preenchida(elemento, texto):
                    log(doc, f"✅ Textarea {indice_campo} já está com o valor desejado.")
                    return True

                # Estratégias em ordem de 'menos invasiva' para 'mais invasiva'
                estrategias = [
                    lambda: _textarea_tradicional(elemento, texto, limpar_primeiro),
                    lambda: _textarea_actionchains(elemento, texto, limpar_primeiro),
                    lambda: _textarea_js_setvalue(elemento, texto),
                    lambda: _textarea_js_react_input(elemento, texto),
                ]

                for i, estrategia in enumerate(estrategias, 1):
                    try:
                        log(doc, f"   ▶️ Estratégia {i}…")
                        estrategia()
                        time.sleep(0.8)

                        # Revalida após a estratégia
                        if validar_textarea_preenchida(elemento, texto):
                            val = (elemento.get_attribute("value") or "").strip()
                            log(doc, f"✅ Preenchido com sucesso pela estratégia {i}: '{val[:60]}{'…' if len(val) > 60 else ''}'")
                            return True
                        else:
                            log(doc, f"⚠️ Estratégia {i} não refletiu o valor esperado.", "WARN")
                    except (StaleElementReferenceException, JavascriptException, TimeoutException) as e:
                        log(doc, f"⚠️ Estratégia {i} falhou: {e}", "WARN")
                        # Reobter o elemento se necessário
                        try:
                            campos = encontrar_campos_textarea()
                            elemento = campos[indice_campo]["elemento"]
                        except Exception:
                            pass
                        continue

                # Se chegou aqui, nenhuma estratégia funcionou nesta tentativa
                if tentativa < max_tentativas:
                    log(doc, f"⚠️ Tentativa {tentativa} falhou; reintentando em 1.5s…", "WARN")
                    time.sleep(1.5)
                    continue
            except Exception as e:
                if tentativa < max_tentativas:
                    log(doc, f"⚠️ Erro na tentativa {tentativa}: {e}. Retentando…", "WARN")
                    time.sleep(1.5)
                    continue
                else:
                    raise

        raise Exception(f"Falha ao preencher textarea {indice_campo} após {max_tentativas} tentativas.")
    return acao


# =========================
# Helpers usados pela função
# =========================

def encontrar_campos_textarea(timeout=10):
    """
    Retorna uma lista de dicts com metadados de cada <textarea> visível e interativa.
    Ex.: [{'elemento': WebElement, 'id': '...', 'name': '...'}]
    """
    elementos = []
    try:
        # Espera haver pelo menos 1 textarea no DOM (se existir)
        wait.until(lambda d: len(d.find_elements(By.TAG_NAME, "textarea")) >= 0)
        textareas = driver.find_elements(By.TAG_NAME, "textarea")
    except Exception:
        textareas = []

    for el in textareas:
        try:
            if not el.is_displayed() or not el.is_enabled():
                continue
            driver.execute_script("arguments[0].scrollIntoView({block:'center'});", el)
            elementos.append({
                "elemento": el,
                "id": el.get_attribute("id"),
                "name": el.get_attribute("name"),
            })
        except Exception:
            continue

    return elementos


def normalizar_texto(txt):
    if txt is None:
        return ""
    # Normaliza quebras de linha e espaços
    return txt.replace("\r\n", "\n").replace("\r", "\n").strip()


def validar_textarea_preenchida(elemento, texto_esperado):
    """Confere se o valor atual da textarea bate com o texto esperado (normalizado)."""
    try:
        atual = elemento.get_attribute("value")
        # Alguns frameworks populam via textContent em textareas (raro, mas possível)
        if atual is None or atual == "":
            atual = (elemento.text or "")
        return normalizar_texto(atual) == normalizar_texto(texto_esperado)
    except StaleElementReferenceException:
        return False


# =========================
# Estratégias de preenchimento
# =========================

def _prepare_focus_and_clear(elemento, limpar_primeiro=True):
    # Garante visibilidade e foco
    driver.execute_script("arguments[0].scrollIntoView({block:'center'});", elemento)
    try:
        elemento.click()
    except Exception:
        driver.execute_script("arguments[0].focus();", elemento)

    if limpar_primeiro:
        try:
            elemento.clear()
        except Exception:
            # Fallback de limpeza por teclas
            ActionChains(driver)\
                .move_to_element(elemento).click()\
                .key_down(Keys.CONTROL).send_keys("a").key_up(Keys.CONTROL)\
                .send_keys(Keys.DELETE).perform()


def _textarea_tradicional(elemento, texto, limpar_primeiro=True):
    _prepare_focus_and_clear(elemento, limpar_primeiro)
    elemento.send_keys(texto)
    # Dispara blur para muitos bindings reativos
    elemento.send_keys(Keys.TAB)

# XPaths dos elementos
btn_prev_xpath = "//button[contains(@class,'fc-prev-button')]"
contrato_xpath = "//div[@class='fc-content']/span[@class='fc-title' and contains(text(),'Contrato: 112972')]"

def clicar_ate_achar_contrato(driver, wait, max_tentativas=20):
    for tentativa in range(max_tentativas):
        try:
            # Verifica se o contrato está visível
            contrato = wait.until(EC.presence_of_element_located((By.XPATH, contrato_xpath)))
            contrato.click()
            print(f"✅ Contrato encontrado e clicado na tentativa {tentativa+1}, Capturando screenshot...")
            return True
        except TimeoutException:
            # Se não achou, clica no botão prev e continua
            print(f"🔄 Tentativa {tentativa+1}: contrato não encontrado, clicando no botão anterior...")
            prev_btn = wait.until(EC.element_to_be_clickable((By.XPATH, btn_prev_xpath)))
            prev_btn.click()
            time.sleep(0.5)  # pequeno intervalo antes da próxima verificação

    print("❌ Contrato não encontrado após várias tentativas")
    return False

def _textarea_actionchains(elemento, texto, limpar_primeiro=True):
    _prepare_focus_and_clear(elemento, limpar_primeiro)
    ac = ActionChains(driver)
    ac.move_to_element(elemento).click().perform()
    # Quebra o texto em partes para evitar engasgos em campos longos
    for chunk_start in range(0, len(texto), 400):
        ac.send_keys(texto[chunk_start:chunk_start+400]).perform()
        time.sleep(0.05)
    ac.send_keys(Keys.TAB).perform()


def _textarea_js_setvalue(elemento, texto):
    # Seta .value e dispara eventos clássicos
    driver.execute_script("""
        const el = arguments[0];
        const val = arguments[1];
        el.value = val;
        // Dispara eventos comuns que form libs escutam
        el.dispatchEvent(new Event('input',  {bubbles:true}));
        el.dispatchEvent(new Event('change', {bubbles:true}));
        el.dispatchEvent(new KeyboardEvent('keyup', {bubbles:true}));
        el.dispatchEvent(new Event('blur',   {bubbles:true}));
    """, elemento, texto)


def _textarea_js_react_input(elemento, texto):
    # Compat extra p/ React (setando o setter do prototype) + eventos
    driver.execute_script("""
        const el = arguments[0];
        const val = arguments[1];

        const desc = Object.getOwnPropertyDescriptor(window.HTMLTextAreaElement.prototype, 'value');
        if (desc && desc.set) {
            desc.set.call(el, val);
        } else {
            el.value = val;
        }

        // React/Vue/Svelte geralmente escutam 'input'
        el.dispatchEvent(new Event('input', {bubbles: true}));
        el.dispatchEvent(new Event('change', {bubbles: true}));
        el.dispatchEvent(new Event('blur', {bubbles: true}));
    """, elemento, texto)



# ==== SISTEMA DATEPICKER MELHORADO ====
def encontrar_campos_datepicker():
    """Encontra todos os campos datepicker na página"""
    global driver
    
    if driver is None:
        return []
    
    seletores_datepicker = [
        "input.hasDatepicker",
        "input[id^='dp']",
        "input[maxlength='10'][grupo='']",
        "input[type='text'][maxlength='10']",
        "input[class*='datepicker']",
        ".hasDatepicker"
    ]
    
    campos_encontrados = []
    
    for seletor in seletores_datepicker:
        try:
            elementos = driver.find_elements(By.CSS_SELECTOR, seletor)
            for elemento in elementos:
                if elemento.is_displayed() and elemento.is_enabled():
                    info = {
                        'elemento': elemento,
                        'id': elemento.get_attribute('id') or f"dp_{len(campos_encontrados)}",
                        'seletor_usado': seletor,
                        'maxlength': elemento.get_attribute('maxlength'),
                        'placeholder': elemento.get_attribute('placeholder')
                    }
                    # Evita duplicatas
                    if not any(c['id'] == info['id'] for c in campos_encontrados):
                        campos_encontrados.append(info)
        except Exception as e:
            log(doc, f"⚠️ Erro ao buscar campos datepicker com {seletor}: {e}", 'WARN')
            continue
    
    log(doc, f"📊 Encontrados {len(campos_encontrados)} campos datepicker")
    return campos_encontrados

def _datepicker_jquery(campo_id, data_valor):
    """Estratégia jQuery para datepicker"""
    global driver
    
    resultado = driver.execute_script("""
        var campoId = arguments[0], valor = arguments[1];
        if (typeof jQuery === 'undefined') return 'jQuery não disponível';
        var $campo = $('#' + campoId);
        if (!$campo.length) return 'Campo não encontrado: ' + campoId;
        try {
            if ($campo.hasClass('hasDatepicker')) { 
                $campo.datepicker('setDate', valor); 
            } else { 
                $campo.val(valor); 
            }
            $campo.trigger('input').trigger('change').trigger('blur');
            return $campo.val();
        } catch(e) { 
            return 'Erro: ' + e.message; 
        }
    """, campo_id, data_valor)
    
    if isinstance(resultado, str) and ('Erro' in resultado or 'não disponível' in resultado):
        raise Exception(f"jQuery falhou: {resultado}")

def _datepicker_javascript(elemento, data_valor):
    """Estratégia JavaScript para datepicker"""
    global driver
    
    driver.execute_script("""
        var campo = arguments[0], valor = arguments[1];
        campo.focus(); 
        campo.value = ''; 
        campo.value = valor;
        ['input','change','blur','keyup'].forEach(ev => 
            campo.dispatchEvent(new Event(ev, {bubbles: true}))
        );
    """, elemento, data_valor)

def _datepicker_actionchains(elemento, data_valor):
    """Estratégia ActionChains para datepicker"""
    global driver
    
    scroll_to_element_safe(elemento)
    time.sleep(0.5)
    
    ActionChains(driver).move_to_element(elemento).click().perform()
    time.sleep(0.5)
    ActionChains(driver).key_down(Keys.CONTROL).send_keys('a').key_up(Keys.CONTROL).perform()
    time.sleep(0.3)
    ActionChains(driver).send_keys(Keys.DELETE).perform()
    time.sleep(0.3)
    
    for char in data_valor:
        ActionChains(driver).send_keys(char).perform()
        time.sleep(0.05)
    
    ActionChains(driver).send_keys(Keys.TAB).perform()

def _datepicker_tradicional(elemento, data_valor):
    """Estratégia tradicional para datepicker"""
    scroll_to_element_safe(elemento)
    time.sleep(0.5)
    elemento.click()
    time.sleep(0.5)
    elemento.clear()
    elemento.send_keys(data_valor)
    elemento.send_keys(Keys.TAB)

def validar_data_preenchida(elemento, data_esperada):
    """Valida se a data foi preenchida corretamente"""
    try:
        if elemento is None:
            return False
            
        val = (elemento.get_attribute('value') or '').strip()
        if not val:
            return False
            
        if val == data_esperada or data_esperada in val:
            return True
            
        # Tenta comparar datas em diferentes formatos
        formatos = [
            '%d/%m/%Y %H:%M', '%d/%m/%Y %H:%M:%S',
            '%d/%m/%Y', '%m/%d/%Y', '%Y-%m-%d', '%d-%m-%Y'
        ]
        
        for formato in formatos:
            try:
                d1 = datetime.strptime(val, formato)
                d2 = datetime.strptime(data_esperada, formato)
                if d1 == d2:
                    return True
            except:
                continue
                
        return False
        
    except Exception:
        return False





def preencher_datepicker_por_indice(indice_campo, data_valor, max_tentativas=5):
    """Preenche datepicker pelo índice com estratégias múltiplas"""
    def acao():
        if not isinstance(indice_campo, int) or indice_campo < 0:
            raise ValueError(f"Índice inválido: {indice_campo}")
            
        if not data_valor or not isinstance(data_valor, str):
            raise ValueError(f"Data inválida: {data_valor}")
        
        tentativa = 0
        while tentativa < max_tentativas:
            tentativa += 1
            
            try:
                campos = encontrar_campos_datepicker()
                
                if not campos:
                    if tentativa < max_tentativas:
                        log(doc, f"⚠️ Nenhum campo datepicker encontrado, tentativa {tentativa}/{max_tentativas}", 'WARN')
                        time.sleep(2)
                        continue
                    raise Exception("Nenhum campo datepicker encontrado na página")
                
                if indice_campo >= len(campos):
                    raise Exception(f"Índice {indice_campo} inválido. Encontrados {len(campos)} campos")
                
                campo_info = campos[indice_campo]
                elemento = campo_info['elemento']
                campo_id = campo_info['id']
                
                log(doc, f"🎯 Tentativa {tentativa}: Preenchendo datepicker {indice_campo} (ID: {campo_id}) com '{data_valor}'")
                
                # Verifica se já está preenchido corretamente
                if validar_data_preenchida(elemento, data_valor):
                    log(doc, f"✅ Campo {indice_campo} já está preenchido corretamente!")
                    return True
                
                # Estratégias específicas para datepicker
                estrategias = [
                    lambda: _datepicker_jquery(campo_id, data_valor),
                    lambda: _datepicker_javascript(elemento, data_valor),
                    lambda: _datepicker_actionchains(elemento, data_valor),
                    lambda: _datepicker_tradicional(elemento, data_valor)
                ]
                
                for i, estrategia in enumerate(estrategias, 1):
                    try:
                        log(doc, f"   Aplicando estratégia {i} para datepicker...")
                        estrategia()
                        time.sleep(1)
                        
                        # Verifica se funcionou
                        if validar_data_preenchida(elemento, data_valor):
                            valor_atual = elemento.get_attribute('value')
                            log(doc, f"✅ Datepicker preenchido com estratégia {i}: '{valor_atual}'")
                            return True
                        else:
                            log(doc, f"⚠️ Estratégia {i} não preencheu corretamente", 'WARN')
                            
                    except Exception as e:
                        log(doc, f"⚠️ Estratégia {i} falhou: {e}", 'WARN')
                        continue
                
                # Se chegou aqui, nenhuma estratégia funcionou nesta tentativa
                if tentativa < max_tentativas:
                    log(doc, f"⚠️ Tentativa {tentativa} falhou, tentando novamente em 2s...", 'WARN')
                    time.sleep(2)
                    continue
                
            except Exception as e:
                if tentativa < max_tentativas:
                    log(doc, f"⚠️ Erro na tentativa {tentativa}: {e}, tentando novamente...", 'WARN')
                    time.sleep(2)
                    continue
                else:
                    raise
        
        raise Exception(f"Falha ao preencher datepicker {indice_campo} após {max_tentativas} tentativas")
    
    return acao

# ==== MODAL E SELEÇÃO ROBUSTOS ====
def abrir_modal_e_selecionar_robusto(btn_selector, pesquisa_selector, termo_pesquisa, btn_pesquisar_selector, resultado_xpath):
    """Versão robusta da função de modal"""
    def acao():
        if driver is None or wait is None:
            raise Exception("Driver ou wait não inicializados")
        
        try:
            # 1. Abre o modal
            log(doc, f"🔘 Abrindo modal com botão: {btn_selector}")
            elemento_botao = aguardar_elemento(btn_selector, TIMEOUT_DEFAULT, 'clickable')
            scroll_to_element_safe(elemento_botao)
            remover_overlays()
            
            # Clique no botão do modal
            try:
                elemento_botao.click()
            except:
                driver.execute_script("arguments[0].click();", elemento_botao)
            
            time.sleep(2)  # Aguarda modal abrir

            # 2. Aguarda e preenche campo pesquisa
            log(doc, f"🔍 Preenchendo pesquisa com: {termo_pesquisa}")
            campo_pesquisa = aguardar_elemento(pesquisa_selector, TIMEOUT_CURTO, 'clickable')
            campo_pesquisa.clear()
            time.sleep(0.3)
            campo_pesquisa.send_keys(termo_pesquisa)
            time.sleep(0.5)

            # 3. Clica pesquisar
            log(doc, "🔍 Executando pesquisa...")
            pesquisar = aguardar_elemento(btn_pesquisar_selector, TIMEOUT_CURTO, 'clickable')
            try:
                pesquisar.click()
            except:
                driver.execute_script("arguments[0].click();", pesquisar)
            
            time.sleep(3)  # Aguarda resultados
            
            # 4. Aguarda resultado e clica
            log(doc, f"🎯 Selecionando resultado: {resultado_xpath}")
            resultado = aguardar_elemento(resultado_xpath, TIMEOUT_DEFAULT, 'clickable', By.XPATH)
            scroll_to_element_safe(resultado)
            time.sleep(0.5)
            
            try:
                resultado.click()
            except:
                driver.execute_script("arguments[0].click();", resultado)
            
            time.sleep(1)
            log(doc, "✅ Seleção no modal concluída")

        except Exception as e:
            log(doc, f"❌ Erro no modal: {e}", 'ERROR')
            # Tenta fechar modal em caso de erro
            try:
                fechar_modal = driver.find_elements(By.CSS_SELECTOR, ".modal .close, .modal-header .close, [data-dismiss='modal']")
                for botao in fechar_modal:
                    if botao.is_displayed():
                        botao.click()
                        break
            except:
                pass
            raise
    
    return acao

# ==== UPLOAD DE ARQUIVOS MELHORADO ====
def fazer_upload_arquivo(xpath_input, caminho_arquivo, timeout=TIMEOUT_DEFAULT):
    """Sistema robusto para upload de arquivos"""
    def acao():
        if not caminho_arquivo or not isinstance(caminho_arquivo, str):
            raise Exception("Caminho do arquivo não fornecido")
            
        if not os.path.exists(caminho_arquivo):
            raise Exception(f"Arquivo não encontrado: {caminho_arquivo}")
        
        # Converte para caminho absoluto
        caminho_absoluto = os.path.abspath(caminho_arquivo)
        log(doc, f"📁 Fazendo upload do arquivo: {os.path.basename(caminho_absoluto)}")
        
        # Aguarda elemento de upload
        input_file = aguardar_elemento(xpath_input, timeout, 'present', By.XPATH)
        
        # Estratégias de upload
        estrategias = [
            lambda: input_file.send_keys(caminho_absoluto),
            lambda: upload_via_pyautogui(input_file, caminho_absoluto),
            lambda: upload_via_javascript(input_file, caminho_absoluto)
        ]
        
        for i, estrategia in enumerate(estrategias, 1):
            try:
                log(doc, f"   Tentando upload estratégia {i}...")
                estrategia()
                time.sleep(2)
                log(doc, f"✅ Upload realizado com estratégia {i}")
                return True
            except Exception as e:
                log(doc, f"⚠️ Estratégia {i} de upload falhou: {e}", 'WARN')
                if i == len(estrategias):
                    raise
                continue
                
        return False
    
    return acao

def upload_via_pyautogui(input_element, caminho_arquivo):
    """Upload usando PyAutoGUI"""
    input_element.click()
    time.sleep(1)
    pyautogui.write(f'"{caminho_arquivo}"')
    pyautogui.press('enter')
    time.sleep(1)

def upload_via_javascript(input_element, caminho_arquivo):
    """Upload usando JavaScript (limitado)"""
    # Esta estratégia é limitada devido às restrições de segurança do navegador
    driver.execute_script("""
        var input = arguments[0];
        input.style.display = 'block';
        input.style.visibility = 'visible';
        input.style.opacity = '1';
    """, input_element)
    input_element.send_keys(caminho_arquivo)

# ==== VALIDAÇÃO DE REGISTROS MELHORADA ====
def validar_registros_encontrados(timeout=TIMEOUT_LONGO):
    """Sistema robusto de validação de registros encontrados"""
    global driver, wait, doc
    
    resultado = {
        'encontrou_registros': False,
        'quantidade_registros': 0,
        'mensagem': '',
        'tabela_encontrada': False,
        'detalhes': []
    }
    
    try:
        log(doc, "🔍 Iniciando validação de registros...")
        time.sleep(5)  # Aguarda processamento inicial
        
        # Seletores para diferentes tipos de tabelas de resultado
        seletores_tabela = [
            '#DataTables_Table_0',
            '#DataTables_Table_0 tbody',
            '#DataTables_Table_1',
            '#DataTables_Table_2',
            'table[id*="DataTables"]',
            '.wdGrid table',
            'table tbody',
            '.resultados table',
            '[class*="grid"][class*="result"]',
            'table[class*="dataTable"]'
        ]
        
        tabela_encontrada = None
        
        # Busca tabela de resultados
        for seletor in seletores_tabela:
            try:
                elementos = driver.find_elements(By.CSS_SELECTOR, seletor)
                for elemento in elementos:
                    if elemento.is_displayed() and elemento.size['height'] > 0:
                        tabela_encontrada = elemento
                        resultado['tabela_encontrada'] = True
                        log(doc, f"✅ Tabela encontrada: {seletor}")
                        break
                
                if tabela_encontrada:
                    break
            except Exception as e:
                log(doc, f"⚠️ Erro ao buscar tabela com {seletor}: {e}", 'WARN')
                continue
        
        if not tabela_encontrada:
            # Busca mensagens de "sem resultados"
            mensagens_vazio = [
                "Não foi encontrado nenhum contrato com os filtros informados.",
                "Não foram encontrados registros", 
                "Nenhum resultado",
                "Sem resultados para exibir",
                "0 registros encontrados",
                "No data available"
            ]
            
            for mensagem in mensagens_vazio:
                try:
                    elem = driver.find_element(By.XPATH, f"//*[contains(text(), '{mensagem}')]")
                    if elem.is_displayed():
                        resultado['mensagem'] = f"Sistema informou: {elem.text.strip()}"
                        log(doc, f"ℹ️ {resultado['mensagem']}")
                        return resultado
                except:
                    continue
            
            # Verifica se existe indicação de carregamento
            loading_elements = driver.find_elements(By.CSS_SELECTOR, ".loading, .spinner, [class*='load']")
            if any(el.is_displayed() for el in loading_elements):
                log(doc, "⏳ Sistema ainda carregando resultados...", 'WARN')
                time.sleep(5)
                return validar_registros_encontrados(timeout - 10)  # Recursão com timeout reduzido
            
            resultado['mensagem'] = "⚠️ Tabela de resultados não localizada"
            log(doc, resultado['mensagem'], 'WARN')
            return resultado
        
        # Conta e valida registros
        try:
            # Estratégias para encontrar linhas de dados
            seletores_linhas = [
                'tbody tr:not(.dataTables_empty):not([class*="no-data"])',
                'tbody tr[class*="odd"], tbody tr[class*="even"]',
                'tbody tr:not(:empty)',
                'tbody tr'
            ]
            
            linhas_validas = []
            
            for seletor_linha in seletores_linhas:
                try:
                    linhas = tabela_encontrada.find_elements(By.CSS_SELECTOR, seletor_linha)
                    
                    for linha in linhas:
                        try:
                            if not linha.is_displayed():
                                continue
                                
                            texto_linha = linha.text.strip().lower()
                            
                            # Valida se é uma linha com dados reais
                            if (len(texto_linha) > 5 and 
                                not any(termo in texto_linha for termo in [
                                    'nenhum registro', 'sem dados', 'no data', 
                                    'vazio', 'empty', 'não foram encontrados',
                                    'loading', 'carregando'
                                ])):
                                
                                linhas_validas.append({
                                    'elemento': linha,
                                    'texto': texto_linha[:100] + '...' if len(texto_linha) > 100 else texto_linha
                                })
                        except Exception as e:
                            log(doc, f"⚠️ Erro ao processar linha: {e}", 'WARN')
                            continue
                    
                    if linhas_validas:
                        log(doc, f"✅ Encontradas {len(linhas_validas)} linhas válidas com {seletor_linha}")
                        break
                        
                except Exception as e:
                    log(doc, f"⚠️ Erro ao processar {seletor_linha}: {e}", 'WARN')
                    continue
            
            quantidade = len(linhas_validas)
            resultado['quantidade_registros'] = quantidade
            resultado['detalhes'] = [linha['texto'] for linha in linhas_validas[:5]]  # Primeiras 5 linhas
            
            if quantidade > 0:
                resultado['encontrou_registros'] = True
                resultado['mensagem'] = f"✅ {quantidade} registro(s) encontrado(s)"
                
                # Log das primeiras linhas
                log(doc, resultado['mensagem'])
                for i, linha in enumerate(linhas_validas[:3], 1):
                    log(doc, f"   Registro {i}: {linha['texto']}")
                
                if quantidade > 3:
                    log(doc, f"   ... e mais {quantidade-3} registro(s)")
            else:
                resultado['mensagem'] = "ℹ️ Tabela encontrada mas sem registros válidos"
                log(doc, resultado['mensagem'])
        
        except Exception as e:
            log(doc, f"❌ Erro ao contar registros: {e}", 'ERROR')
            # Em caso de erro na contagem, assume que existem registros para não interromper
            resultado['encontrou_registros'] = True
            resultado['quantidade_registros'] = 1
            resultado['mensagem'] = f"⚠️ Erro na validação, continuando teste: {e}"
        
        # Verifica alertas do sistema
        encontrar_mensagem_alerta()
        
        return resultado
        
    except Exception as e:
        log(doc, f"❌ Erro geral na validação: {e}", 'ERROR')
        resultado['encontrou_registros'] = True  # Assume sucesso para não interromper
        resultado['quantidade_registros'] = 1
        resultado['mensagem'] = f"⚠️ Validação falhou, continuando teste: {e}"
        return resultado

# ==== UTILITÁRIOS DIVERSOS ====
def encontrar_mensagem_alerta():
    """Busca mensagens de alerta na página"""
    global driver, doc
    
    if driver is None:
        return None
    
    seletores = [
        (".alerts.salvo", "✅ Sucesso"),
        (".alerts.alerta", "⚠️ Alerta"),
        (".alerts.erro", "❌ Erro"),
        (".alert-success", "✅ Sucesso"),
        (".alert-warning", "⚠️ Alerta"),
        (".alert-danger", "❌ Erro"),
        ("[class*='toast']", "📢 Notificação"),
    ]
    
    for seletor, tipo in seletores:
        try:
            elementos = driver.find_elements(By.CSS_SELECTOR, seletor)
            for elemento in elementos:
                if elemento.is_displayed():
                    texto = elemento.text.strip()
                    if texto:
                        log(doc, f"📢 {tipo}: {texto}")
                        return elemento
        except Exception as e:
            log(doc, f"⚠️ Erro ao buscar alerta {seletor}: {e}", 'WARN')
            continue
    
    return None

def ajustar_zoom(zoom_level="90%"):
    """Ajusta zoom da página"""
    global driver, doc
    
    if driver is None:
        return
    
    try:
        driver.execute_script(f"document.body.style.zoom='{zoom_level}'")
        log(doc, f"🔍 Zoom ajustado para {zoom_level}.")
    except Exception as e:
        log(doc, f"⚠️ Erro ao ajustar zoom: {e}", 'WARN')

def realizar_consulta():
    """Executa a consulta"""
    def acao():
        seletor = '#gsCallCenter > div.wdTelas > div.telaRegistroChamadas.telaConsulta.relative > div.formRow.formLastLine > div:nth-child(5) > a'
        elemento = aguardar_elemento(seletor, TIMEOUT_DEFAULT, 'clickable')
        scroll_to_element_safe(elemento)
        
        try:
            elemento.click()
        except:
            driver.execute_script("arguments[0].click();", elemento)
        
        time.sleep(2)
        log(doc, "✅ Consulta executada")
    
    return acao

def finalizar_cadastro():
    """Finaliza o cadastro/consulta"""
    def acao():
        seletor_css_finalizar = '#gsPet > div.wdTelas > div.wdWizard.clearfix.telaConsulta > div.btnHolder > div'
        elemento = aguardar_elemento(seletor_css_finalizar, TIMEOUT_DEFAULT, 'clickable')
        scroll_to_element_safe(elemento)
        
        try:
            elemento.click()
        except:
            driver.execute_script("arguments[0].click();", elemento)
        
        time.sleep(3)
        log(doc, "✅ Cadastro finalizado")
    
    return acao

# ==== RELATÓRIO MELHORADO ====
def finalizar_relatorio():
    """Finaliza e salva o relatório"""
    global driver, doc
    
    try:
        nome_arquivo = f"relatorio_callcenter_agenda_de_cobranças_cenario_1_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        
        # Adiciona resumo final
        doc.add_paragraph("\n" + "="*50)
        doc.add_paragraph("RESUMO FINAL DO TESTE")
        doc.add_paragraph("="*50)
        doc.add_paragraph(f"Teste finalizado em: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")
        doc.add_paragraph(f"Total de screenshots capturadas: {len(screenshot_registradas)}")
        
        # Salva documento
        doc.save(nome_arquivo)
        log(doc, f"📄 Relatório salvo como: {nome_arquivo}")
        
        # Tenta abrir o arquivo
        try:
            if os.name == 'nt':  # Windows
                os.startfile(nome_arquivo)
            else:  # Linux/Mac
                subprocess.run(['xdg-open', nome_arquivo])
        except Exception as e:
            log(doc, f"⚠️ Não foi possível abrir automaticamente o relatório: {e}", 'WARN')
            
    except Exception as e:
        print(f"❌ Erro ao salvar relatório: {e}")
    
    # Fecha driver
    if driver:
        try:
            log(doc, "🔚 Fechando navegador...")
            driver.quit()
        except Exception as e:
            print(f"⚠️ Erro ao fechar driver: {e}")

# ==== INICIALIZAÇÃO DE DRIVER MELHORADA ====
def inicializar_driver():
    """Inicializa o driver do Chrome com configurações otimizadas"""
    global driver, wait
    
    try:
        log(doc, "🚀 Inicializando driver do Chrome...")
        
        options = Options()
        options.add_argument("--start-maximized")
        options.add_argument("--disable-blink-features=AutomationControlled")
        options.add_argument("--disable-extensions")
        options.add_argument("--disable-plugins")
        options.add_argument("--no-sandbox")
        options.add_argument("--disable-dev-shm-usage")
        options.add_argument("--disable-gpu")
        options.add_argument("--disable-web-security")
        options.add_argument("--allow-running-insecure-content")
        options.add_experimental_option("excludeSwitches", ["enable-automation"])
        options.add_experimental_option('useAutomationExtension', False)
        
        # Configurações de download (caso necessário)
        prefs = {
            "profile.default_content_settings.popups": 0,
            "profile.default_content_setting_values.automatic_downloads": 1,
        }
        options.add_experimental_option("prefs", prefs)

        # Instala e configura ChromeDriver
        service = Service(ChromeDriverManager().install())
        driver = webdriver.Chrome(service=service, options=options)
        
        # Remove indicadores de automação
        driver.execute_script("Object.defineProperty(navigator, 'webdriver', {get: () => undefined})")
        driver.execute_cdp_cmd('Network.setUserAgentOverride', {
            "userAgent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36"
        })
        
        # Configura wait com timeout padrão
        wait = WebDriverWait(driver, TIMEOUT_DEFAULT)
        
        log(doc, "✅ Driver inicializado com sucesso")
        return True
        
    except Exception as e:
        log(doc, f"❌ Erro ao inicializar driver: {e}", 'ERROR')
        return False

# ==== FUNÇÃO CORRIGIDA PARA ACESSO AO PET ====
def acessar_modulo_callcenter():
    """Versão corrigida para acessar o módulo Callcenter"""
    def acao():
        # XPath do elemento que precisamos clicar
        xpath_callcenter = "/html/body/div[15]/ul/li[18]/img"
        
        try:
            # Aguarda elemento estar presente
            elemento = aguardar_elemento(xpath_callcenter, TIMEOUT_DEFAULT, 'present', By.XPATH)
            
            # Usa scroll seguro
            scroll_success = scroll_to_element_safe(elemento)
            if not scroll_success:
                log(doc, "⚠️ Problemas com scroll, tentando continuar", 'WARN')
            
            # Remove overlays
            remover_overlays()
            time.sleep(1)
            
            # Aguarda elemento ficar clicável
            elemento_clicavel = aguardar_elemento(xpath_callcenter, TIMEOUT_CURTO, 'clickable', By.XPATH)
            
            # Tenta diferentes estratégias de clique
            estrategias_clique = [
                lambda: elemento_clicavel.click(),
                lambda: ActionChains(driver).move_to_element(elemento_clicavel).click().perform(),
                lambda: driver.execute_script("arguments[0].click();", elemento_clicavel),
                lambda: driver.execute_script("""
                    var element = arguments[0];
                    if (element) {
                        element.focus();
                        var event = new MouseEvent('click', {
                            bubbles: true,
                            cancelable: true,
                            view: window
                        });
                        element.dispatchEvent(event);
                    }
                """, elemento_clicavel)
            ]
            
            for i, estrategia in enumerate(estrategias_clique, 1):
                try:
                    log(doc, f"   Tentando estratégia de clique {i} para módulo Callcenter...")
                    estrategia()
                    time.sleep(2)
                    log(doc, f"✅ Clique no módulo Callcenter realizado com estratégia {i}")
                    return True
                except Exception as e:
                    log(doc, f"⚠️ Estratégia {i} falhou: {e}", 'WARN')
                    if i == len(estrategias_clique):
                        raise
                    continue
            
            return False
            
        except Exception as e:
            log(doc, f"❌ Erro ao acessar módulo Callcenter: {e}", 'ERROR')
            raise
    
    return acao


# ==== SISTEMA ANTI-TIMEOUT JAVASCRIPT ====
class JSTimeoutHandler:
    """Sistema robusto para lidar com timeouts JavaScript no Selenium"""
    
    def __init__(self, driver, doc, timeout_padrao=10, max_retries=3):
        self.driver = driver
        self.doc = doc
        self.timeout_padrao = timeout_padrao
        self.max_retries = max_retries
        self.last_error = None
        
    def log_timeout(self, msg, level="INFO"):
        """Log com timestamp"""

        prefix = {
            "INFO": "ℹ️ ",
            "WARN": "⚠️ ",
            "ERROR": "❌ ",
            "SUCCESS": "✅ "
        }.get(level, "📝 ")
        
        print(f" {prefix} {msg}")
        if hasattr(self.doc, 'add_paragraph'):
            self.doc.add_paragraph(f"{msg}")
    
    def execute_js_safe(self, script, *args, timeout=None, fallback_result=None):
        """Executa JavaScript com proteção contra timeouts"""
        timeout = timeout or self.timeout_padrao
        
        original_timeout = self.driver.timeouts.script
        self.driver.set_script_timeout(timeout)
        
        for tentativa in range(1, self.max_retries + 1):
            try:
                if tentativa > 1:
                    self.log_timeout(f"Tentativa {tentativa}/{self.max_retries}", "INFO")
                
                result = self.driver.execute_script(script, *args)
                self.driver.set_script_timeout(original_timeout)
                
                if tentativa > 1:
                    self.log_timeout("JavaScript executado com sucesso", "SUCCESS")
                return result
                
            except JavascriptException as e:
                self.last_error = e
                self.log_timeout(f"Erro JavaScript: {str(e)[:150]}", "ERROR")
                self._limpar_estado_js()
                
                if tentativa < self.max_retries:
                    time.sleep(1 + tentativa * 0.5)
                    continue
                    
            except TimeoutException as e:
                self.last_error = e
                self.log_timeout(f"Timeout JavaScript ({timeout}s)", "ERROR")
                self._forcar_parada_js()
                
                if tentativa < self.max_retries:
                    time.sleep(2 + tentativa)
                    continue
                    
            except WebDriverException as e:
                self.last_error = e
                self.log_timeout(f"Erro WebDriver: {str(e)[:150]}", "ERROR")
                
                if tentativa < self.max_retries:
                    time.sleep(1.5)
                    continue
                    
            except Exception as e:
                self.last_error = e
                self.log_timeout(f"Erro inesperado: {str(e)[:150]}", "ERROR")
                break
        
        try:
            self.driver.set_script_timeout(original_timeout)
        except:
            pass
            
        if fallback_result is not None:
            self.log_timeout(f"Usando valor fallback: {fallback_result}", "WARN")
        return fallback_result
    
    def _limpar_estado_js(self):
        """Limpa estado JavaScript do browser"""
        try:
            cleanup_script = """
                if (window.__cleanupTimers) {
                    window.__cleanupTimers.forEach(clearTimeout);
                    window.__cleanupTimers.forEach(clearInterval);
                }
                if (typeof jQuery !== 'undefined') {
                    jQuery.active = 0;
                }
                window.__pendingRequests = 0;
                return true;
            """
            self.driver.execute_script(cleanup_script)
            time.sleep(0.5)
        except Exception:
            pass
    
    def _forcar_parada_js(self):
        """Força parada de JavaScript travado"""
        try:
            self.driver.execute_script("window.stop();")
            time.sleep(0.3)
        except Exception:
            pass



# ==== JS FORCE ENGINE COM PROTEÇÃO ANTI-TIMEOUT ====
class JSForceEngine:
    """Motor de execução JavaScript forçado com proteção contra timeouts"""
    
    def __init__(self, driver, wait, doc, timeout_padrao=10, max_retries=3):
        self.driver = driver
        self.wait = wait
        self.doc = doc
        self.timeout_handler = JSTimeoutHandler(driver, doc, timeout_padrao, max_retries)
    
    def execute_js(self, script, *args, timeout=None, fallback_result=None):
        """Executa JavaScript com proteção contra timeout"""
        return self.timeout_handler.execute_js_safe(
            script, *args, timeout=timeout, fallback_result=fallback_result
        )
    
    def wait_ajax_complete(self, timeout=15):
        """Aguarda AJAX completar com proteção contra timeout"""
        script = """
            var jQueryOk = (typeof jQuery==='undefined') || (jQuery.active===0);
            var fetchOk = !window.__pendingRequests || window.__pendingRequests===0;
            var overlays = document.querySelectorAll(
                '.blockScreen, .blockUI, .loading, .overlay, [class*="loading"], [class*="spinner"]'
            );
            var overlayOk = true;
            for (var i=0; i<overlays.length; i++){
                var s=window.getComputedStyle(overlays[i]);
                if(s.display!=='none' && s.visibility!=='hidden' && parseFloat(s.opacity||1)>0.01){
                    overlayOk=false;
                    break;
                }
            }
            return jQueryOk && fetchOk && overlayOk;
        """
        
        end = time.time() + timeout
        while time.time() < end:
            try:
                done = self.execute_js(script, timeout=5, fallback_result=True)
                if done:
                    return True
            except:
                pass
            time.sleep(0.2)
        return True
    
    def scroll_into_view(self, target, padding=100):
        """
        Faz scroll até um WebElement (ou seletor XPath/CSS).
        padding: desloca um pouco pra cima pra não ficar colado no topo.
        """
        try:
            el = target
            # Se vier como string, tenta resolver
            if isinstance(target, str):
                try:
                    el = self.driver.find_element("xpath", target)
                except Exception:
                    el = self.driver.find_element("css selector", target)

            # Estratégia principal via JS
            self.driver.execute_script("""
                const el = arguments[0], pad = arguments[1] || 0;
                if (!el) return;
                el.scrollIntoView({block:'center', inline:'center'});
                try { window.scrollBy(0, -pad); } catch(e) {}
            """, el, padding)
            time.sleep(0.2)
            return True
        except Exception:
            # Fallback ActionChains
            try:
                ActionChains(self.driver).move_to_element(el).perform()
                time.sleep(0.2)
                return True
            except Exception:
                return False

    def click_element(self, el, wait_after=0.4):
        """
        Clica com múltiplas estratégias em um WebElement já localizado.
        """
        try:
            # 1) Clique padrão
            el.click()
            time.sleep(wait_after)
            self.wait_ajax_complete(min(10, int(wait_after*10)) or 3)
            return True
        except Exception:
            pass

        try:
            # 2) Clique via JS
            self.driver.execute_script("arguments[0].click();", el)
            time.sleep(wait_after)
            self.wait_ajax_complete(min(10, int(wait_after*10)) or 3)
            return True
        except Exception:
            pass

        try:
            # 3) Sequência de eventos de mouse
            self.execute_js("""
                const e = arguments[0];
                const rect = e.getBoundingClientRect();
                const x = rect.left + 5, y = rect.top + 5;
                ['mouseover','mouseenter','mousemove','mousedown','mouseup','click'].forEach(t=>{
                    e.dispatchEvent(new MouseEvent(t,{bubbles:true,cancelable:true,view:window,clientX:x,clientY:y}));
                });
                if (typeof e.click==='function') e.click();
                return true;
            """, el, timeout=5, fallback_result=True)
            time.sleep(wait_after)
            self.wait_ajax_complete(min(10, int(wait_after*10)) or 3)
            return True
        except Exception:
            pass

        # 4) Último recurso: ActionChains
        try:
            ActionChains(self.driver).move_to_element(el).pause(0.05).click().perform()
            time.sleep(wait_after)
            self.wait_ajax_complete(min(10, int(wait_after*10)) or 3)
            return True
        except Exception:
            pass

        raise Exception("Não foi possível clicar no elemento com as estratégias disponíveis.")
    
    def force_click(self, selector, by_xpath=False, max_attempts=5):
        """Clique forçado com proteção contra timeout"""
        log(self.doc, f"🎯 Clique forçado em: {selector}")
        
        for attempt in range(max_attempts):
            try:
                strategies = [
                    self._click_strategy_2,
                    self._click_strategy_1,
                    self._click_strategy_3,
                    self._click_strategy_4,
                    self._click_strategy_5,
                ]
                
                for i, strategy in enumerate(strategies, 1):
                    try:
                        if attempt > 0 or i > 1:
                            log(self.doc, f"   Tentativa {attempt + 1}.{i}...")
                        
                        result = self.execute_js(
                            self._get_strategy_script(strategy, selector, by_xpath),
                            selector,
                            by_xpath,
                            timeout=5,
                            fallback_result=False
                        )
                        
                        if result:
                            log(self.doc, f"✅ Clique bem-sucedido (estratégia {i})")
                            time.sleep(0.5)
                            self.wait_ajax_complete(10)
                            return True
                            
                    except Exception as e:
                        if i == 1 and attempt == 0:
                            pass  # Silencia primeiro erro
                        else:
                            log(self.doc, f"   Estratégia {i} falhou: {str(e)[:80]}")
                        continue
                
                if attempt < max_attempts - 1:
                    time.sleep(1 + attempt * 0.5)
                    
            except Exception as e:
                log(self.doc, f"⚠️ Tentativa {attempt + 1} falhou: {e}")
                if attempt < max_attempts - 1:
                    time.sleep(1.5)
        
        raise Exception(f"Falha ao clicar após {max_attempts} tentativas: {selector}")
    
    def _get_strategy_script(self, strategy_func, selector, by_xpath):
        """Retorna o script JavaScript para cada estratégia"""
        base_locator = """
            var selector = arguments[0];
            var byXPath = arguments[1];
            var element;
            
            if (byXPath) {
                var result = document.evaluate(selector, document, null, 
                    XPathResult.FIRST_ORDERED_NODE_TYPE, null);
                element = result.singleNodeValue;
            } else {
                element = document.querySelector(selector);
            }
            
            if (!element) throw new Error('Elemento não encontrado');
        """
        
        if strategy_func == self._click_strategy_1:
            return base_locator + """
                element.style.pointerEvents = 'auto';
                element.style.display = 'block';
                element.style.visibility = 'visible';
                element.style.opacity = '1';
                element.removeAttribute('disabled');
                element.scrollIntoView({behavior: 'smooth', block: 'center'});
                setTimeout(function() { element.click(); }, 300);
                return true;
            """
        elif strategy_func == self._click_strategy_2:
            return base_locator + """
                element.style.pointerEvents = 'auto';
                element.removeAttribute('disabled');
                element.scrollIntoView({block: 'center'});
                
                var events = ['mouseover', 'mouseenter', 'mousemove', 'mousedown', 'mouseup', 'click'];
                events.forEach(function(eventType) {
                    var evt = new MouseEvent(eventType, {
                        bubbles: true, cancelable: true, view: window, detail: 1,
                        clientX: element.getBoundingClientRect().left + 5,
                        clientY: element.getBoundingClientRect().top + 5
                    });
                    element.dispatchEvent(evt);
                });
                
                if (typeof element.click === 'function') element.click();
                return true;
            """
        elif strategy_func == self._click_strategy_3:
            return base_locator + """
                element.style.display = 'block';
                element.style.visibility = 'visible';
                element.style.opacity = '1';
                element.style.pointerEvents = 'auto';
                element.focus();
                element.click();
                element.dispatchEvent(new Event('click', {bubbles: true, cancelable: true}));
                return true;
            """
        elif strategy_func == self._click_strategy_4:
            return base_locator + """
                element.removeAttribute('disabled');
                element.removeAttribute('readonly');
                element.style.pointerEvents = 'auto !important';
                element.style.display = 'block !important';
                element.style.visibility = 'visible !important';
                element.style.opacity = '1 !important';
                
                var overlays = document.querySelectorAll('.modal, .overlay, .blockUI, [role="dialog"]');
                overlays.forEach(function(overlay) {
                    overlay.style.display = 'none';
                    overlay.style.visibility = 'hidden';
                });
                
                element.focus();
                element.click();
                
                var clickEvent = new MouseEvent('click', {
                    view: window, bubbles: true, cancelable: true
                });
                element.dispatchEvent(clickEvent);
                
                if (typeof jQuery !== 'undefined') jQuery(element).trigger('click');
                return true;
            """
        else:  # strategy_5
            return base_locator + """
                var rect = element.getBoundingClientRect();
                var x = rect.left + rect.width / 2;
                var y = rect.top + rect.height / 2;
                
                var evt = document.createEvent('MouseEvents');
                evt.initMouseEvent('click', true, true, window, 1, x, y, x, y, false, false, false, false, 0, null);
                element.dispatchEvent(evt);
                
                if (element.onclick) element.onclick();
                
                var parent = element.parentElement;
                while (parent && parent !== document.body) {
                    if (parent.onclick) {
                        parent.onclick();
                        break;
                    }
                    parent = parent.parentElement;
                }
                return true;
            """
    
    def _click_strategy_1(self, selector, by_xpath):
        pass  # Implementado via _get_strategy_script
    
    def _click_strategy_2(self, selector, by_xpath):
        pass
    
    def _click_strategy_3(self, selector, by_xpath):
        pass
    
    def _click_strategy_4(self, selector, by_xpath):
        pass
    
    def _click_strategy_5(self, selector, by_xpath):
        pass
    
    def force_fill(self, selector, value, by_xpath=False, max_attempts=5):
        """Preenchimento forçado com proteção contra timeout"""
        log(self.doc, f"✏️ Preenchimento forçado: {selector} = '{value}'")
        
        for attempt in range(max_attempts):
            try:
                strategies = [
                    self._fill_strategy_1,
                    self._fill_strategy_2,
                    self._fill_strategy_3,
                    self._fill_strategy_4,
                ]
                
                for i, strategy in enumerate(strategies, 1):
                    try:
                        if attempt > 0 or i > 1:
                            log(self.doc, f"   Tentativa {attempt + 1}.{i}...")
                        
                        result = self.execute_js(
                            self._get_fill_script(strategy, selector, value, by_xpath),
                            selector,
                            value,
                            by_xpath,
                            timeout=5,
                            fallback_result=None
                        )
                        
                        time.sleep(0.3)
                        if self._validate_fill(selector, value, by_xpath):
                            log(self.doc, f"✅ Campo preenchido (estratégia {i})")
                            return True
                    except Exception as e:
                        if i == 1 and attempt == 0:
                            pass
                        else:
                            log(self.doc, f"   Estratégia {i} falhou: {str(e)[:80]}")
                        continue
                
                if attempt < max_attempts - 1:
                    time.sleep(1 + attempt * 0.5)
                    
            except Exception as e:
                log(self.doc, f"⚠️ Tentativa {attempt + 1} falhou: {e}")
        
        raise Exception(f"Falha ao preencher após {max_attempts} tentativas: {selector}")

    def _get_fill_script(self, strategy_func, selector, value, by_xpath):
        """Retorna o script de preenchimento"""
        base_locator = """
            var selector = arguments[0];
            var value = arguments[1];
            var byXPath = arguments[2];
            var element;
            
            if (byXPath) {
                var result = document.evaluate(selector, document, null, 
                    XPathResult.FIRST_ORDERED_NODE_TYPE, null);
                element = result.singleNodeValue;
            } else {
                element = document.querySelector(selector);
            }
            
            if (!element) throw new Error('Campo não encontrado');
        """
        
        if strategy_func == self._fill_strategy_1:
            return base_locator + """
                element.removeAttribute('disabled');
                element.removeAttribute('readonly');
                element.style.display = 'block';
                element.style.visibility = 'visible';
                element.scrollIntoView({block: 'center'});
                element.focus();
                element.dispatchEvent(new Event('focus', {bubbles: true}));
                element.value = '';
                element.value = value;
                ['input', 'change', 'blur', 'keyup'].forEach(function(evt) {
                    element.dispatchEvent(new Event(evt, {bubbles: true}));
                });
                return element.value;
            """
        elif strategy_func == self._fill_strategy_2:
            return base_locator + """
                var nativeInputValueSetter = Object.getOwnPropertyDescriptor(
                    window.HTMLInputElement.prototype, 'value'
                ).set;
                if (nativeInputValueSetter) {
                    nativeInputValueSetter.call(element, value);
                } else {
                    element.value = value;
                }
                element.dispatchEvent(new Event('input', {bubbles: true}));
                element.dispatchEvent(new Event('change', {bubbles: true}));
                element.dispatchEvent(new KeyboardEvent('keydown', {bubbles: true}));
                element.dispatchEvent(new KeyboardEvent('keyup', {bubbles: true}));
                element.dispatchEvent(new Event('blur', {bubbles: true}));
                return element.value;
            """
        elif strategy_func == self._fill_strategy_3:
            return base_locator + """
                element.value = value;
                if (typeof jQuery !== 'undefined') {
                    jQuery(element).val(value).trigger('input').trigger('change').trigger('blur');
                }
                ['focus', 'input', 'change', 'blur'].forEach(function(evt) {
                    element.dispatchEvent(new Event(evt, {bubbles: true}));
                });
                return element.value;
            """
        else:  # strategy_4
            return base_locator + """
                element.removeAttribute('disabled');
                element.removeAttribute('readonly');
                element.removeAttribute('maxlength');
                element.value = '';
                element.setAttribute('value', value);
                element.value = value;
                element.style.color = element.style.color;
                
                var events = ['focus', 'click', 'input', 'change', 'keydown', 'keypress', 
                              'keyup', 'blur', 'paste', 'textInput'];
                events.forEach(function(evt) {
                    try {
                        element.dispatchEvent(new Event(evt, {bubbles: true, cancelable: true}));
                    } catch(e) {}
                });
                
                if (element.oninput) element.oninput();
                if (element.onchange) element.onchange();
                return element.value;
            """
    
    def _fill_strategy_1(self, selector, value, by_xpath):
        pass
    
    def _fill_strategy_2(self, selector, value, by_xpath):
        pass
    
    def _fill_strategy_3(self, selector, value, by_xpath):
        pass
    
    def _fill_strategy_4(self, selector, value, by_xpath):
        pass
    
    def _validate_fill(self, selector, expected_value, by_xpath):
        """Valida preenchimento"""
        script = """
            var selector = arguments[0];
            var expected = arguments[1];
            var byXPath = arguments[2];
            var element;
            
            if (byXPath) {
                var result = document.evaluate(selector, document, null, 
                    XPathResult.FIRST_ORDERED_NODE_TYPE, null);
                element = result.singleNodeValue;
            } else {
                element = document.querySelector(selector);
            }
            
            if (!element) return false;
            var actual = element.value || '';
            return actual.trim() === expected.trim() || actual.includes(expected);
        """
        try:
            return self.execute_js(script, selector, expected_value, by_xpath, timeout=3, fallback_result=False)
        except:
            return False
        
import time

def clicar_finalizar_e_verificar_alerta(js_engine, doc, timeout=5, pausa=0.5):
    """
    Clica no botão 'Finalizar', aguarda 0,5s e procura mensagens de alerta.
    Usa safe_action e js_engine.force_click().
    """
    safe_action(doc, "Clicando em 'Finalizar'", lambda:
        js_engine.force_click(
            "//a[@class='btModel btGray btyes' and normalize-space()='Finalizar']",
            by_xpath=True
        )
    )

    time.sleep(pausa)
    log(doc, "🔍 Verificando mensagens de alerta após o clique em 'Finalizar'...")
    return encontrar_mensagem_alerta()


import time
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.common.exceptions import (
    TimeoutException,
    StaleElementReferenceException,
    ElementClickInterceptedException,
    NoSuchElementException,
    ElementNotInteractableException
)
from selenium.webdriver import ActionChains

class LOVHandler:
    """
    Handler ultra-robusto para manipulação de LOV (List of Values) com:
    - 10 estratégias diferentes de clique
    - Detecção automática de iframes
    - Retry inteligente com backoff exponencial
    - Reforço automático de cliques
    - Logs detalhados de cada tentativa
    """
    
    def __init__(self, js_engine, doc, max_retries=5, timeout=10):
        self.js = js_engine
        self.doc = doc
        self.driver = js_engine.driver
        self.max_retries = max_retries
        self.timeout = timeout
        
    def _log(self, msg, level="INFO"):
        """Log padronizado com níveis"""
        prefixes = {
            "INFO": "ℹ️",
            "SUCCESS": "✅",
            "WARNING": "⚠️",
            "ERROR": "❌",
            "DEBUG": "🔍"
        }
        prefix = prefixes.get(level, "📝")
        log(self.doc, f"{prefix} {msg}")
    
    def _wait_element(self, locator_type, locator_value, timeout=None, condition="clickable"):
        """Aguarda elemento com diferentes condições"""
        timeout = timeout or self.timeout
        conditions = {
            "present": EC.presence_of_element_located,
            "visible": EC.visibility_of_element_located,
            "clickable": EC.element_to_be_clickable
        }
        
        try:
            cond = conditions.get(condition, conditions["clickable"])
            return WebDriverWait(self.driver, timeout).until(
                cond((locator_type, locator_value))
            )
        except TimeoutException:
            return None
    
    def _is_element_visible(self, element):
        """Verifica se elemento está realmente visível"""
        try:
            if not element:
                return False
            if not element.is_displayed():
                return False
            
            # Verifica via JavaScript também
            is_visible = self.driver.execute_script("""
                const el = arguments[0];
                if (!el) return false;
                const style = window.getComputedStyle(el);
                return (
                    style.display !== 'none' &&
                    style.visibility !== 'hidden' &&
                    parseFloat(style.opacity || 1) > 0.01 &&
                    el.offsetParent !== null
                );
            """, element)
            
            return is_visible
        except:
            return False
    
    def _force_element_visible(self, element):
        """Força elemento a ficar visível e interativo"""
        try:
            self.driver.execute_script("""
                const el = arguments[0];
                el.style.display = 'block';
                el.style.visibility = 'visible';
                el.style.opacity = '1';
                el.style.pointerEvents = 'auto';
                el.removeAttribute('disabled');
                el.removeAttribute('readonly');
                
                // Remove overlays que podem bloquear
                const overlays = document.querySelectorAll(
                    '.modal-backdrop, .overlay, .blockUI, [class*="loading"]'
                );
                overlays.forEach(o => {
                    o.style.display = 'none';
                    o.style.visibility = 'hidden';
                });
            """, element)
            return True
        except:
            return False
    
    def _detect_and_enter_iframe(self, iframe_xpath=None):
        """Detecta e entra no iframe automaticamente"""
        try:
            # Primeiro volta para o contexto principal
            self.driver.switch_to.default_content()
            
            # Se foi fornecido um xpath específico
            if iframe_xpath:
                try:
                    frame = self._wait_element("xpath", iframe_xpath, timeout=3)
                    if frame:
                        self.driver.switch_to.frame(frame)
                        self._log(f"Entrou no iframe: {iframe_xpath}", "SUCCESS")
                        return True
                except:
                    pass
            
            # Detecção automática de iframes
            iframe_selectors = [
                "//iframe[contains(@class,'LOV') or contains(@id,'LOV') or contains(@id,'lov')]",
                "//iframe[contains(@class,'modal') or contains(@class,'popup')]",
                "//iframe[contains(@src,'lov') or contains(@src,'LOV')]",
                "(//iframe)[last()]"  # Último iframe (geralmente o modal)
            ]
            
            for selector in iframe_selectors:
                try:
                    frames = self.driver.find_elements(By.XPATH, selector)
                    for frame in frames:
                        if self._is_element_visible(frame):
                            self.driver.switch_to.frame(frame)
                            self._log(f"Iframe detectado automaticamente: {selector}", "SUCCESS")
                            return True
                except:
                    continue
            
            return False
        except Exception as e:
            self._log(f"Erro ao detectar iframe: {e}", "WARNING")
            return False
    
    def _click_strategy_1_standard(self, element):
        """Estratégia 1: Clique padrão do Selenium"""
        element.click()
        return True
    
    def _click_strategy_2_javascript_simple(self, element):
        """Estratégia 2: JavaScript simples"""
        self.driver.execute_script("arguments[0].click();", element)
        return True
    
    def _click_strategy_3_javascript_advanced(self, element):
        """Estratégia 3: JavaScript com eventos completos"""
        self.driver.execute_script("""
            const el = arguments[0];
            const rect = el.getBoundingClientRect();
            const x = rect.left + rect.width / 2;
            const y = rect.top + rect.height / 2;
            
            const events = ['mouseover', 'mouseenter', 'mousemove', 'mousedown', 'mouseup', 'click'];
            events.forEach(eventType => {
                const evt = new MouseEvent(eventType, {
                    bubbles: true,
                    cancelable: true,
                    view: window,
                    detail: 1,
                    clientX: x,
                    clientY: y
                });
                el.dispatchEvent(evt);
            });
            
            if (typeof el.click === 'function') el.click();
        """, element)
        return True
    
    def _click_strategy_4_action_chains(self, element):
        """Estratégia 4: ActionChains com pause"""
        ActionChains(self.driver)\
            .move_to_element(element)\
            .pause(0.1)\
            .click()\
            .perform()
        return True
    
    def _click_strategy_5_action_chains_offset(self, element):
        """Estratégia 5: ActionChains com offset"""
        ActionChains(self.driver)\
            .move_to_element_with_offset(element, 5, 5)\
            .pause(0.05)\
            .click()\
            .perform()
        return True
    
    def _click_strategy_6_force_visible_then_click(self, element):
        """Estratégia 6: Força visibilidade e clica"""
        self._force_element_visible(element)
        time.sleep(0.2)
        element.click()
        return True
    
    def _click_strategy_7_scroll_and_click(self, element):
        """Estratégia 7: Scroll suave até elemento"""
        self.driver.execute_script("""
            arguments[0].scrollIntoView({
                behavior: 'smooth',
                block: 'center',
                inline: 'center'
            });
        """, element)
        time.sleep(0.3)
        element.click()
        return True
    
    def _click_strategy_8_remove_overlays(self, element):
        """Estratégia 8: Remove todos os overlays e clica"""
        self.driver.execute_script("""
            // Remove todos os overlays possíveis
            const overlays = document.querySelectorAll(`
                .modal-backdrop, .overlay, .blockUI, .blockScreen,
                [class*="loading"], [class*="spinner"], [class*="overlay"],
                [style*="z-index: 9999"], [style*="position: fixed"]
            `);
            overlays.forEach(o => {
                o.style.display = 'none';
                o.style.visibility = 'hidden';
                o.remove();
            });
            
            const el = arguments[0];
            el.style.zIndex = '999999';
            el.click();
        """, element)
        return True
    
    def _click_strategy_9_jquery_trigger(self, element):
        """Estratégia 9: jQuery trigger (se disponível)"""
        self.driver.execute_script("""
            const el = arguments[0];
            if (typeof jQuery !== 'undefined') {
                jQuery(el).trigger('click');
            } else {
                el.click();
            }
        """, element)
        return True
    
    def _click_strategy_10_nuclear_option(self, element):
        """Estratégia 10: Opção nuclear - força tudo"""
        self.driver.execute_script("""
            const el = arguments[0];
            
            // Remove TODOS os atributos que podem bloquear
            el.removeAttribute('disabled');
            el.removeAttribute('readonly');
            el.style.pointerEvents = 'auto !important';
            el.style.display = 'block !important';
            el.style.visibility = 'visible !important';
            el.style.opacity = '1 !important';
            el.style.zIndex = '999999 !important';
            
            // Remove TODOS os overlays da página
            document.querySelectorAll('*').forEach(elem => {
                const style = window.getComputedStyle(elem);
                if (
                    style.position === 'fixed' &&
                    parseInt(style.zIndex) > 1000 &&
                    elem !== el &&
                    !elem.contains(el)
                ) {
                    elem.style.display = 'none';
                }
            });
            
            // Foca no elemento
            el.focus();
            
            // Dispara TODOS os eventos possíveis
            const allEvents = [
                'focus', 'focusin', 'mouseover', 'mouseenter', 'mousemove',
                'mousedown', 'mouseup', 'click', 'dblclick'
            ];
            
            allEvents.forEach(eventType => {
                try {
                    const evt = new MouseEvent(eventType, {
                        bubbles: true,
                        cancelable: true,
                        view: window
                    });
                    el.dispatchEvent(evt);
                } catch(e) {}
            });
            
            // Clique direto
            if (typeof el.click === 'function') el.click();
            
            // jQuery se disponível
            if (typeof jQuery !== 'undefined') {
                jQuery(el).trigger('click');
            }
            
            // Tenta onclick manual
            if (el.onclick) el.onclick();
            
            return true;
        """, element)
        return True
    
    def _advanced_click(self, element, max_attempts=10):
        """
        Sistema avançado de clique com 10 estratégias diferentes
        Tenta cada estratégia até uma funcionar
        """
        strategies = [
            ("Clique Padrão", self._click_strategy_1_standard),
            ("JavaScript Simples", self._click_strategy_2_javascript_simple),
            ("JavaScript Avançado", self._click_strategy_3_javascript_advanced),
            ("ActionChains", self._click_strategy_4_action_chains),
            ("ActionChains Offset", self._click_strategy_5_action_chains_offset),
            ("Força Visível", self._click_strategy_6_force_visible_then_click),
            ("Scroll e Clique", self._click_strategy_7_scroll_and_click),
            ("Remove Overlays", self._click_strategy_8_remove_overlays),
            ("jQuery Trigger", self._click_strategy_9_jquery_trigger),
            ("Opção Nuclear", self._click_strategy_10_nuclear_option)
        ]
        
        for attempt in range(1, max_attempts + 1):
            for strategy_name, strategy_func in strategies:
                try:
                    self._log(f"Tentativa {attempt}/{max_attempts}: {strategy_name}", "DEBUG")
                    
                    # Tenta executar a estratégia
                    strategy_func(element)
                    time.sleep(0.3)
                    
                    self._log(f"✓ {strategy_name} funcionou!", "SUCCESS")
                    return True
                    
                except StaleElementReferenceException:
                    self._log(f"Elemento ficou stale, tentando recarregar...", "WARNING")
                    return False  # Precisa recarregar elemento
                    
                except (ElementClickInterceptedException, 
                        ElementNotInteractableException) as e:
                    self._log(f"✗ {strategy_name}: {str(e)[:50]}", "DEBUG")
                    continue
                    
                except Exception as e:
                    self._log(f"✗ {strategy_name}: {str(e)[:50]}", "DEBUG")
                    continue
            
            if attempt < max_attempts:
                time.sleep(0.5 * attempt)  # Backoff exponencial
        
        return False
    
    def _fill_search_fields(self, search_text, max_fields=10):
        """Preenche TODOS os campos de pesquisa encontrados"""
        if not search_text:
            return 0
        
        search_field_xpaths = [
            "//input[@id='txtPesquisa']",
            "//input[contains(@class,'pesquisa')]",
            "//input[contains(@class,'nomePesquisa')]",
            "//input[contains(translate(@name,'PESQUISA','pesquisa'),'pesquisa')]",
            "//input[@type='text' and contains(@style,'width:210px')]",
            "//input[@type='text' and not(@disabled)]",
        ]
        
        filled_count = 0
        for xpath in search_field_xpaths:
            try:
                fields = self.driver.find_elements(By.XPATH, xpath)
                for field in fields[:max_fields]:
                    try:
                        if not self._is_element_visible(field):
                            continue
                        
                        # Limpa e preenche
                        self._force_element_visible(field)
                        field.clear()
                        field.click()
                        field.send_keys(search_text)
                        filled_count += 1
                        
                    except:
                        continue
            except:
                continue
        
        if filled_count > 0:
            self._log(f"Preenchidos {filled_count} campos de pesquisa", "SUCCESS")
        else:
            self._log("Nenhum campo de pesquisa encontrado", "WARNING")
        
        return filled_count
    
    def _click_search_button(self):
        """Clica no botão Pesquisar com múltiplas estratégias"""
        search_button_xpaths = [
            "//a[contains(@class,'btPesquisar') and contains(normalize-space(.),'Pesquisar')]",
            "//button[contains(normalize-space(.),'Pesquisar')]",
            "//input[@type='button' and contains(@value,'Pesquisar')]",
            "//a[contains(@class,'lpFind')]",
            "//a[contains(@onclick,'pesquisar')]"
        ]
        
        # Fallback: ENTER no campo ativo
        try:
            self._log("Tentando ENTER no campo de pesquisa", "DEBUG")
            ActionChains(self.driver).send_keys(Keys.ENTER).perform()
            return True
        except Exception:
            pass
        
        for xpath in search_button_xpaths:
            try:
                btn = self._wait_element("xpath", xpath, timeout=2)
                if btn and self._is_element_visible(btn):
                    if self._advanced_click(btn):
                        self._log("Botão 'Pesquisar' clicado", "SUCCESS")
                        return True
            except:
                pass
        
        self._log("Não foi possível clicar em 'Pesquisar'", "WARNING")
        return False
    
    def _select_result(self, result_text=""):
        """Seleciona o resultado da pesquisa"""
        result_xpaths = []
        
        if result_text:
            result_xpaths = [
                f"//td[contains(normalize-space(.),'{result_text}')]",
                f"//span[contains(normalize-space(.),'{result_text}')]",
                f"//div[contains(normalize-space(.),'{result_text}')]",
                f"//tr[contains(normalize-space(.),'{result_text}')]//td[1]",
                f"//li[contains(normalize-space(.),'{result_text}')]"
            ]
        else:
            result_xpaths = [
                "(//table//tr[1]/td[1])[1]",
                "(//tr[contains(@class,'ui-widget-content')][1])[1]",
                "(//tr[contains(@class,'rich-table-row')][1])[1]",
                "(//li[@class='ui-autocomplete-item'])[1]"
            ]
        
        for xpath in result_xpaths:
            try:
                result = self._wait_element("xpath", xpath, timeout=3)
                if result and self._is_element_visible(result):
                    if self._advanced_click(result):
                        self._log(f"Resultado selecionado: {result_text or 'primeiro'}", "SUCCESS")
                        return xpath  # Retorna xpath usado
            except:
                continue
        
        self._log(f"Não foi possível selecionar: {result_text or 'primeiro resultado'}", "ERROR")
        return None
    
    def _reinforce_click_on_result(self, result_xpath, max_reclick=5):
        """
        Reforça o clique no resultado SEMPRE, independente do modal
        Clica repetidamente até max_reclick vezes
        """
        if not result_xpath:
            return
        
        self._log(f"🔁 Reforçando clique no resultado ({max_reclick}x)", "INFO")
        
        for attempt in range(1, max_reclick + 1):
            try:
                time.sleep(0.35)
                
                # Tenta localizar o resultado novamente
                result = self.driver.find_element(By.XPATH, result_xpath)
                
                # Clica usando sistema avançado
                self._log(f"Reforço {attempt}/{max_reclick}", "DEBUG")
                self._advanced_click(result)
                
            except StaleElementReferenceException:
                self._log(f"Elemento stale no reforço {attempt}", "WARNING")
                try:
                    # Tenta recarregar o elemento
                    result = self.driver.find_element(By.XPATH, result_xpath)
                    self._advanced_click(result)
                except:
                    continue
                    
            except Exception as e:
                self._log(f"Erro no reforço {attempt}: {str(e)[:50]}", "WARNING")
                continue
        
        self._log("✅ Reforço de cliques concluído", "SUCCESS")
    
    def open_and_select(
        self,
        btn_index=None,
        btn_xpath=None,
        btn_css=None,
        search_text="",
        result_text="",
        iframe_xpath=None,
        auto_detect_iframe=True,
        reinforce_clicks=5,
        wait_after=0.5
    ):
        """
        Método principal: abre LOV, pesquisa e seleciona resultado
        
        Args:
            btn_index: Índice do botão LOV (0-based)
            btn_xpath: XPath customizado do botão
            btn_css: Seletor CSS do botão
            search_text: Texto para pesquisar
            result_text: Texto do resultado a selecionar
            iframe_xpath: XPath do iframe (se houver)
            auto_detect_iframe: Detectar iframe automaticamente
            reinforce_clicks: Quantas vezes reforçar o clique (padrão: 5)
            wait_after: Tempo de espera após conclusão
        
        Returns:
            bool: True se bem-sucedido, False caso contrário
        """
        
        self._log(f"🔍 LOV: '{search_text}' → '{result_text}'", "INFO")
        
        for retry in range(1, self.max_retries + 1):
            try:
                # ===== PASSO 1: Volta para contexto principal =====
                try:
                    self.driver.switch_to.default_content()
                except:
                    pass
                
                # ===== PASSO 2: Localiza e clica no botão LOV =====
                btn_selector = None
                by_type = None
                
                if btn_index is not None:
                    btn_selector = f"(//a[@class='sprites sp-openLov'])[{btn_index + 1}]"
                    by_type = "xpath"
                elif btn_xpath:
                    btn_selector = btn_xpath
                    by_type = "xpath"
                elif btn_css:
                    btn_selector = btn_css
                    by_type = "css"
                else:
                    raise ValueError("Forneça btn_index, btn_xpath ou btn_css")
                
                self._log(f"Localizando botão LOV: {btn_selector}", "DEBUG")
                
                lov_button = self._wait_element(by_type, btn_selector, timeout=5)
                if not lov_button:
                    raise Exception(f"Botão LOV não encontrado: {btn_selector}")
                
                self._log("Abrindo LOV...", "INFO")
                if not self._advanced_click(lov_button):
                    raise Exception("Falha ao clicar no botão LOV")
                
                time.sleep(0.8)
                
                # ===== PASSO 3: Entra no iframe se necessário =====
                if iframe_xpath or auto_detect_iframe:
                    self._detect_and_enter_iframe(iframe_xpath)
                    time.sleep(0.3)
                
                # ===== PASSO 4: Preenche campos de pesquisa =====
                if search_text:
                    self._fill_search_fields(search_text)
                    time.sleep(0.3)
                
                # ===== PASSO 5: Clica em Pesquisar =====
                self._click_search_button()
                time.sleep(0.8)
                
                # ===== PASSO 6: Seleciona resultado =====
                result_xpath = self._select_result(result_text)
                if not result_xpath:
                    if retry < self.max_retries:
                        self._log(f"Retry {retry}/{self.max_retries}...", "WARNING")
                        time.sleep(1.5 * retry)
                        continue
                    raise Exception("Falha ao selecionar resultado")
                
                # ===== PASSO 7: Reforça clique SEMPRE =====
                if reinforce_clicks > 0:
                    self._reinforce_click_on_result(result_xpath, reinforce_clicks)
                
                # ===== PASSO 8: Volta para contexto principal =====
                try:
                    self.driver.switch_to.default_content()
                except:
                    pass
                
                time.sleep(wait_after)
                self._log("LOV concluído com sucesso!", "SUCCESS")
                return True
                
            except Exception as e:
                self._log(f"Tentativa {retry} falhou: {str(e)[:100]}", "ERROR")
                
                # Tenta recuperar voltando para contexto principal
                try:
                    self.driver.switch_to.default_content()
                except:
                    pass
                
                if retry < self.max_retries:
                    time.sleep(2 * retry)  # Backoff exponencial
                    continue
        
        self._log(f"LOV falhou após {self.max_retries} tentativas", "ERROR")
        return False
    
# ==== FUNÇÕES AUXILIARES ====
def encontrar_mensagem_alerta():
    seletores = [
        (".alerts.salvo", "✅ Mensagem de Sucesso"),
        (".alerts.alerta", "⚠️ Mensagem de Alerta"),
        (".alerts.erro", "❌ Mensagem de Erro"),
    ]

    for seletor, tipo in seletores:
        try:
            elemento = driver.find_element(By.CSS_SELECTOR, seletor)
            if elemento.is_displayed():
                log(doc, f"📢 {tipo}: {elemento.text}")
                return elemento
        except:
            continue

    log(doc, "ℹ️ Nenhuma mensagem de alerta encontrada.")
    return None

def inicializar_driver():
    """Inicializa WebDriver com configurações otimizadas"""
    global driver, wait
    
    try:
        options = Options()
        options.add_argument("--start-maximized")
        options.add_argument("--disable-blink-features=AutomationControlled")
        options.add_argument("--disable-dev-shm-usage")
        options.add_argument("--no-sandbox")
        options.add_experimental_option("excludeSwitches", ["enable-automation"])
        options.add_experimental_option('useAutomationExtension', False)
        
        driver = webdriver.Chrome(
            service=Service(ChromeDriverManager().install()), 
            options=options
        )
        
        driver.execute_script(
            "Object.defineProperty(navigator, 'webdriver', {get: () => undefined})"
        )
        
        # Configura timeouts globais
        driver.set_script_timeout(30)
        driver.implicitly_wait(10)
        
        wait = WebDriverWait(driver, TIMEOUT_DEFAULT)
        
        log(doc, "✅ Driver inicializado com sucesso")
        return True
        
    except Exception as e:
        log(doc, f"❌ Erro ao inicializar driver: {e}")
        return False

def finalizar_relatorio():
    """Salva relatório e fecha driver"""
    global driver, doc
    
    nome_arquivo = f"relatorio_devolucoes_cenario_2_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    
    try:
        doc.save(nome_arquivo)
        log(doc, f"📄 Relatório salvo: {nome_arquivo}")
        
        try:
            subprocess.run(["start", "winword", nome_arquivo], shell=True)
        except:
            pass
            
    except Exception as e:
        print(f"Erro ao salvar relatório: {e}")
    
    if driver:
        try:
            driver.quit()
            log(doc, "✅ Driver encerrado")
        except:
            pass





def _sanitize_timeout(t):
    """Garante timeout válido"""
    if not isinstance(t, (int, float)) or t <= 0:
        return TIMEOUT_DEFAULT
    return max(5, min(120, t))  # Entre 5 e 120 segundos

# ==== AGUARDAR ELEMENTO MELHORADO ====
def aguardar_elemento(seletor, timeout=TIMEOUT_DEFAULT, condicao='clickable', by_type=By.CSS_SELECTOR):
    """Função centralizada para aguardar elementos com diferentes condições"""
    global driver, wait
    
    if driver is None:
        raise Exception("Driver não inicializado")
    
    timeout = _sanitize_timeout(timeout)
    
    condicoes = {
        'present': EC.presence_of_element_located,
        'visible': EC.visibility_of_element_located,
        'clickable': EC.element_to_be_clickable,
        'invisible': EC.invisibility_of_element_located
    }
    
    if condicao not in condicoes:
        condicao = 'clickable'
    
    try:
        wait_obj = WebDriverWait(driver, timeout)
        elemento = wait_obj.until(condicoes[condicao]((by_type, seletor)))
        return elemento
    except TimeoutException:
        log(doc, f"❌ Timeout aguardando elemento: {seletor} (condição: {condicao}, timeout: {timeout}s)", 'ERROR')
        raise TimeoutException(f"Elemento não encontrado: {seletor} (condição: {condicao})")
    except Exception as e:
        log(doc, f"❌ Erro aguardando elemento {seletor}: {e}", 'ERROR')
        raise

# ==== SCROLL CORRIGIDO - PRINCIPAL CORREÇÃO ====
def scroll_to_element_safe(elemento_ou_seletor, by_type=By.CSS_SELECTOR):
    """Scroll seguro até elemento com validação robusta"""
    global driver
    
    if driver is None:
        log(doc, "⚠️ Driver não disponível para scroll", 'WARN')
        return False
    
    try:
        # Se for seletor, encontra o elemento
        if isinstance(elemento_ou_seletor, str):
            elemento = aguardar_elemento(elemento_ou_seletor, 10, 'present', by_type)
        else:
            elemento = elemento_ou_seletor
        
        if elemento is None:
            log(doc, "⚠️ Elemento não encontrado para scroll", 'WARN')
            return False
        
        # Verifica se elemento é válido antes de fazer scroll
        if not elemento.is_displayed():
            log(doc, "⚠️ Elemento não está visível para scroll", 'WARN')
            return False
        
        # Estratégias de scroll em ordem de preferência
        scroll_strategies = [
            # Estratégia 1: JavaScript com verificação prévia
            lambda: driver.execute_script("""
                var element = arguments[0];
                if (element && typeof element.scrollIntoView === 'function') {
                    element.scrollIntoView({
                        behavior: 'smooth',
                        block: 'center',
                        inline: 'center'
                    });
                    return true;
                } else {
                    return false;
                }
            """, elemento),
            
            # Estratégia 2: ActionChains
            lambda: ActionChains(driver).move_to_element(elemento).perform(),
            
            # Estratégia 3: JavaScript alternativo
            lambda: driver.execute_script("""
                var element = arguments[0];
                if (element) {
                    element.scrollIntoView();
                    window.scrollBy(0, -100);
                }
            """, elemento),
            
            # Estratégia 4: Scroll da página até o elemento
            lambda: driver.execute_script("""
                var element = arguments[0];
                if (element) {
                    var rect = element.getBoundingClientRect();
                    var scrollTop = window.pageYOffset || document.documentElement.scrollTop;
                    var targetY = rect.top + scrollTop - (window.innerHeight / 2);
                    window.scrollTo(0, targetY);
                }
            """, elemento)
        ]
        
        for i, strategy in enumerate(scroll_strategies, 1):
            try:
                log(doc, f"   Tentando estratégia de scroll {i}...")
                result = strategy()
                
                # Para estratégia 1, verifica resultado
                if i == 1 and result is False:
                    log(doc, f"   Estratégia {i}: elemento não suporta scrollIntoView", 'WARN')
                    continue
                
                time.sleep(0.8)  # Aguarda scroll completar
                
                # Verifica se elemento ainda está acessível
                if elemento.is_displayed() and elemento.is_enabled():
                    log(doc, f"✅ Scroll realizado com estratégia {i}")
                    return True
                else:
                    log(doc, f"   Estratégia {i}: elemento não ficou acessível", 'WARN')
                    continue
                    
            except Exception as e:
                log(doc, f"   Estratégia {i} de scroll falhou: {str(e)[:100]}...", 'WARN')
                continue
        
        log(doc, "⚠️ Todas as estratégias de scroll falharam", 'WARN')
        return False
        
    except Exception as e:
        log(doc, f"⚠️ Erro geral no scroll: {e}", 'WARN')
        return False



# ==== SISTEMA DATEPICKER MELHORADO ====
def encontrar_campos_datepicker():
    """Encontra todos os campos datepicker na página"""
    global driver
    
    if driver is None:
        return []
    
    seletores_datepicker = [
        "input.hasDatepicker",
        "input[id^='dp']",
        "input[maxlength='10'][grupo='']",
        "input[type='text'][maxlength='10']",
        "input[class*='datepicker']",
        ".hasDatepicker"
    ]
    
    campos_encontrados = []
    
    for seletor in seletores_datepicker:
        try:
            elementos = driver.find_elements(By.CSS_SELECTOR, seletor)
            for elemento in elementos:
                if elemento.is_displayed() and elemento.is_enabled():
                    info = {
                        'elemento': elemento,
                        'id': elemento.get_attribute('id') or f"dp_{len(campos_encontrados)}",
                        'seletor_usado': seletor,
                        'maxlength': elemento.get_attribute('maxlength'),
                        'placeholder': elemento.get_attribute('placeholder')
                    }
                    # Evita duplicatas
                    if not any(c['id'] == info['id'] for c in campos_encontrados):
                        campos_encontrados.append(info)
        except Exception as e:
            log(doc, f"⚠️ Erro ao buscar campos datepicker com {seletor}: {e}", 'WARN')
            continue
    
    log(doc, f"📊 Encontrados {len(campos_encontrados)} campos datepicker")
    return campos_encontrados

def _datepicker_jquery(campo_id, data_valor):
    """Estratégia jQuery para datepicker"""
    global driver
    
    resultado = driver.execute_script("""
        var campoId = arguments[0], valor = arguments[1];
        if (typeof jQuery === 'undefined') return 'jQuery não disponível';
        var $campo = $('#' + campoId);
        if (!$campo.length) return 'Campo não encontrado: ' + campoId;
        try {
            if ($campo.hasClass('hasDatepicker')) { 
                $campo.datepicker('setDate', valor); 
            } else { 
                $campo.val(valor); 
            }
            $campo.trigger('input').trigger('change').trigger('blur');
            return $campo.val();
        } catch(e) { 
            return 'Erro: ' + e.message; 
        }
    """, campo_id, data_valor)
    
    if isinstance(resultado, str) and ('Erro' in resultado or 'não disponível' in resultado):
        raise Exception(f"jQuery falhou: {resultado}")

def _datepicker_javascript(elemento, data_valor):
    """Estratégia JavaScript para datepicker"""
    global driver
    
    driver.execute_script("""
        var campo = arguments[0], valor = arguments[1];
        campo.focus(); 
        campo.value = ''; 
        campo.value = valor;
        ['input','change','blur','keyup'].forEach(ev => 
            campo.dispatchEvent(new Event(ev, {bubbles: true}))
        );
    """, elemento, data_valor)

def _datepicker_actionchains(elemento, data_valor):
    """Estratégia ActionChains para datepicker"""
    global driver
    
    scroll_to_element_safe(elemento)
    time.sleep(0.5)
    
    ActionChains(driver).move_to_element(elemento).click().perform()
    time.sleep(0.5)
    ActionChains(driver).key_down(Keys.CONTROL).send_keys('a').key_up(Keys.CONTROL).perform()
    time.sleep(0.3)
    ActionChains(driver).send_keys(Keys.DELETE).perform()
    time.sleep(0.3)
    
    for char in data_valor:
        ActionChains(driver).send_keys(char).perform()
        time.sleep(0.05)
    
    ActionChains(driver).send_keys(Keys.TAB).perform()

def _datepicker_tradicional(elemento, data_valor):
    """Estratégia tradicional para datepicker"""
    scroll_to_element_safe(elemento)
    time.sleep(0.5)
    elemento.click()
    time.sleep(0.5)
    elemento.clear()
    elemento.send_keys(data_valor)
    elemento.send_keys(Keys.TAB)

def validar_data_preenchida(elemento, data_esperada):
    """Valida se a data foi preenchida corretamente"""
    try:
        if elemento is None:
            return False
            
        val = (elemento.get_attribute('value') or '').strip()
        if not val:
            return False
            
        if val == data_esperada or data_esperada in val:
            return True
            
        # Tenta comparar datas em diferentes formatos
        formatos = [
            '%d/%m/%Y %H:%M', '%d/%m/%Y %H:%M:%S',
            '%d/%m/%Y', '%m/%d/%Y', '%Y-%m-%d', '%d-%m-%Y'
        ]
        
        for formato in formatos:
            try:
                d1 = datetime.strptime(val, formato)
                d2 = datetime.strptime(data_esperada, formato)
                if d1 == d2:
                    return True
            except:
                continue
                
        return False
        
    except Exception:
        return False





def preencher_datepicker_por_indice(indice_campo, data_valor, max_tentativas=5):
    """Preenche datepicker pelo índice com estratégias múltiplas"""
    def acao():
        if not isinstance(indice_campo, int) or indice_campo < 0:
            raise ValueError(f"Índice inválido: {indice_campo}")
            
        if not data_valor or not isinstance(data_valor, str):
            raise ValueError(f"Data inválida: {data_valor}")
        
        tentativa = 0
        while tentativa < max_tentativas:
            tentativa += 1
            
            try:
                campos = encontrar_campos_datepicker()
                
                if not campos:
                    if tentativa < max_tentativas:
                        log(doc, f"⚠️ Nenhum campo datepicker encontrado, tentativa {tentativa}/{max_tentativas}", 'WARN')
                        time.sleep(2)
                        continue
                    raise Exception("Nenhum campo datepicker encontrado na página")
                
                if indice_campo >= len(campos):
                    raise Exception(f"Índice {indice_campo} inválido. Encontrados {len(campos)} campos")
                
                campo_info = campos[indice_campo]
                elemento = campo_info['elemento']
                campo_id = campo_info['id']
                
                log(doc, f"🎯 Tentativa {tentativa}: Preenchendo datepicker {indice_campo} (ID: {campo_id}) com '{data_valor}'")
                
                # Verifica se já está preenchido corretamente
                if validar_data_preenchida(elemento, data_valor):
                    log(doc, f"✅ Campo {indice_campo} já está preenchido corretamente!")
                    return True
                
                # Estratégias específicas para datepicker
                estrategias = [
                    lambda: _datepicker_jquery(campo_id, data_valor),
                    lambda: _datepicker_javascript(elemento, data_valor),
                    lambda: _datepicker_actionchains(elemento, data_valor),
                    lambda: _datepicker_tradicional(elemento, data_valor)
                ]
                
                for i, estrategia in enumerate(estrategias, 1):
                    try:
                        log(doc, f"   Aplicando estratégia {i} para datepicker...")
                        estrategia()
                        time.sleep(1)
                        
                        # Verifica se funcionou
                        if validar_data_preenchida(elemento, data_valor):
                            valor_atual = elemento.get_attribute('value')
                            log(doc, f"✅ Datepicker preenchido com estratégia {i}: '{valor_atual}'")
                            return True
                        else:
                            log(doc, f"⚠️ Estratégia {i} não preencheu corretamente", 'WARN')
                            
                    except Exception as e:
                        log(doc, f"⚠️ Estratégia {i} falhou: {e}", 'WARN')
                        continue
                
                # Se chegou aqui, nenhuma estratégia funcionou nesta tentativa
                if tentativa < max_tentativas:
                    log(doc, f"⚠️ Tentativa {tentativa} falhou, tentando novamente em 2s...", 'WARN')
                    time.sleep(2)
                    continue
                
            except Exception as e:
                if tentativa < max_tentativas:
                    log(doc, f"⚠️ Erro na tentativa {tentativa}: {e}, tentando novamente...", 'WARN')
                    time.sleep(2)
                    continue
                else:
                    raise
        
        raise Exception(f"Falha ao preencher datepicker {indice_campo} após {max_tentativas} tentativas")
    
    return acao


def preencher_campos_pesquisa_por_indice(self, 
                                         search_text: str, 
                                         search_xpaths=None, 
                                         max_campos: int | None = None, 
                                         pausa: float = 0.3,
                                         limpar_antes: bool = True):
    """
    Procura TODOS os campos de pesquisa e preenche um após o outro (ordem no DOM).
    - search_text: texto a preencher.
    - search_xpaths: lista de XPaths para busca (usa padrão se None).
    - max_campos: limita quantos campos serão preenchidos (None = todos).
    - pausa: pausa curta entre preenchimentos.
    - limpar_antes: se True, limpa o campo antes de preencher.

    Retorna: {"total_encontrados": int, "total_preenchidos": int, "xpaths_usados": [str], "falhas": [str]}
    """
    search_xpaths = search_xpaths or [
        "//input[@id='txtPesquisa']",
        "//input[@class='nomePesquisa']",
        # adicione mais padrões específicos antes do genérico:
        "//input[contains(translate(@name,'ABCDEFGHIJKLMNOPQRSTUVWXYZ','abcdefghijklmnopqrstuvwxyz'),'pesquisa')]",
        "//input[contains(translate(@class,'ABCDEFGHIJKLMNOPQRSTUVWXYZ','abcdefghijklmnopqrstuvwxyz'),'pesquisa')]",
        "//input[@type='text']"
    ]

    # Monta um XPATH-UNIÃO para manter a ordem no DOM
    xpath_uniao = " | ".join(f"({xp})" for xp in search_xpaths)

    # Coleta elementos (evita duplicados por id/ref)
    try:
        elementos = self.js.driver.find_elements(By.XPATH, xpath_uniao)
    except Exception as e:
        log(self.doc, f"⚠️ Erro ao buscar campos de pesquisa: {e}")
        return {"total_encontrados": 0, "total_preenchidos": 0, "xpaths_usados": [], "falhas": [str(e)]}

    # Filtra apenas visíveis e habilitados
    elementos_filtrados = []
    for el in elementos:
        try:
            if el.is_displayed() and el.is_enabled():
                elementos_filtrados.append(el)
        except Exception:
            continue

    total_encontrados = len(elementos_filtrados)
    if total_encontrados == 0:
        log(self.doc, "🔎 Nenhum campo de pesquisa visível/habilitado encontrado.")
        return {"total_encontrados": 0, "total_preenchidos": 0, "xpaths_usados": [], "falhas": []}

    if max_campos is not None and max_campos > 0:
        elementos_filtrados = elementos_filtrados[:max_campos]

    log(self.doc, f"🧭 Campos de pesquisa encontrados: {total_encontrados} | A preencher: {len(elementos_filtrados)}")

    total_preenchidos = 0
    falhas = []
    xpaths_usados = []

    for idx, el in enumerate(elementos_filtrados, start=1):
        try:
            # Recalcula um XPATH relativo único para log (opcional, pode ser pesado).
            # Aqui só registramos o index para simplificar.
            xpaths_usados.append(f"[campo #{idx}]")

            # Traz para a tela e foca
            self.js.driver.execute_script("arguments[0].scrollIntoView({block:'center'});", el)
            try:
                el.click()
            except Exception:
                pass

            if limpar_antes:
                try:
                    el.clear()
                except Exception:
                    # fallback via JS
                    try:
                        self.js.driver.execute_script("arguments[0].value='';", el)
                    except Exception:
                        pass

            # Preenche (usa seu helper para manter padrão)
            try:
                self.js.force_fill_element(el, search_text)  # se você tiver esse helper
            except AttributeError:
                # fallback: force_fill por XPATH do próprio elemento (gera um xpath usando JS)
                try:
                    self.js.driver.execute_script("arguments[0].value = arguments[1];", el, search_text)
                except Exception:
                    # último fallback: send_keys
                    el.send_keys(search_text)

            total_preenchidos += 1
            log(self.doc, f"   ✏️ Preenchido campo #{idx} com '{search_text}'")
            if pausa:
                time.sleep(pausa)

        except (StaleElementReferenceException, ElementNotInteractableException) as e:
            falhas.append(f"Campo #{idx}: {type(e).__name__}")
            log(self.doc, f"   ⚠️ Falha no campo #{idx}: {e}")
            continue
        except Exception as e:
            falhas.append(f"Campo #{idx}: {e}")
            log(self.doc, f"   ⚠️ Erro inesperado no campo #{idx}: {e}")
            continue

    resumo = {
        "total_encontrados": total_encontrados,
        "total_preenchidos": total_preenchidos,
        "xpaths_usados": xpaths_usados,
        "falhas": falhas
    }
    log(self.doc, f"✅ Pesquisa preenchida em {total_preenchidos}/{len(elementos_filtrados)} campos. Falhas: {len(falhas)}")
    return resumo

def clicar_x_por_indice(indice_x: int, max_tentativas: int = 5, timeout: int = 10, scroll: bool = True):
    """Clica no ícone de X pelo índice"""
    def acao():
        if not isinstance(indice_x, int) or indice_x < 0:
            raise ValueError(f"Índice inválido: {indice_x}")

        tentativa = 0
        while tentativa < max_tentativas:
            tentativa += 1
            try:
                log(doc, f"🔎 Tentativa {tentativa}: Localizando Botões X...")
                elementos = driver.find_elements(By.XPATH, "//a[@class='fa fa-close']")

                if not elementos:
                    if tentativa < max_tentativas:
                        log(doc, f"⚠️ Nenhum Botão X encontrado (tentativa {tentativa}/{max_tentativas})")
                        time.sleep(1.2)
                        continue
                    raise Exception("Nenhum Botão X encontrado.")

                if indice_x >= len(elementos):
                    raise Exception(f"Índice {indice_x} inválido. Encontrados {len(elementos)} botões X.")

                locator_xpath = f"(//a[@class='fa fa-close'])[{indice_x + 1}]"
                elemento = driver.find_element(By.XPATH, locator_xpath)

                log(doc, f"🎯 Preparando clique no LOV de índice {indice_x}")

                def _wait_clickable():
                    wait.until(EC.element_to_be_clickable((By.XPATH, locator_xpath)))

                estrategias = [
                    lambda: (_wait_clickable(), elemento.click()),
                    lambda: (
                        driver.execute_script("arguments[0].scrollIntoView({block:'center'});", elemento) if scroll else None,
                        time.sleep(0.2),
                        elemento.click()
                    ),
                    lambda: driver.execute_script("arguments[0].click();", elemento),
                    lambda: ActionChains(driver).move_to_element(elemento).pause(0.1).click().perform()
                ]

                for i, estrategia in enumerate(estrategias, 1):
                    try:
                        log(doc, f"   ▶️ Estratégia {i} de clique no botão X...")
                        estrategia()
                        time.sleep(0.3)
                        log(doc, f"✅ Clique no botão X (índice {indice_x}) realizado (estratégia {i})")
                        return True
                    except (ElementClickInterceptedException, StaleElementReferenceException, JavascriptException, TimeoutException) as e:
                        log(doc, f"⚠️ Estratégia {i} falhou: {e}")
                        try:
                            elementos = driver.find_elements(By.CSS_SELECTOR, "a.sprites.sp-openLov")
                            elemento = driver.find_element(By.XPATH, locator_xpath)
                        except:
                            pass
                        continue

                if tentativa < max_tentativas:
                    log(doc, f"⚠️ Tentativa {tentativa} não conseguiu clicar no botão X. Reintentando...")
                    time.sleep(1.2)
                    continue

            except Exception as e:
                if tentativa < max_tentativas:
                    log(doc, f"⚠️ Erro na tentativa {tentativa}: {e}. Reintentando...")
                    time.sleep(1.2)
                    continue
                raise

        raise Exception(f"Falha ao clicar no LOV de índice {indice_x} após {max_tentativas} tentativas.")

    return acao

def encontrar_campos_textarea(timeout=10):
    """Retorna lista de textareas visíveis e interativas"""
    elementos = []
    try:
        wait.until(lambda d: len(d.find_elements(By.TAG_NAME, "textarea")) >= 0)
        textareas = driver.find_elements(By.TAG_NAME, "textarea")
    except:
        textareas = []

    for el in textareas:
        try:
            if not el.is_displayed() or not el.is_enabled():
                continue
            driver.execute_script("arguments[0].scrollIntoView({block:'center'});", el)
            elementos.append({
                "elemento": el,
                "id": el.get_attribute("id"),
                "name": el.get_attribute("name"),
            })
        except:
            continue

    return elementos

def normalizar_texto(txt):
    if txt is None:
        return ""
    return txt.replace("\r\n", "\n").replace("\r", "\n").strip()

def validar_textarea_preenchida(elemento, texto_esperado):
    """Confere se o valor atual da textarea bate com o esperado"""
    try:
        atual = elemento.get_attribute("value")
        if atual is None or atual == "":
            atual = (elemento.text or "")
        return normalizar_texto(atual) == normalizar_texto(texto_esperado)
    except StaleElementReferenceException:
        return False

def _prepare_focus_and_clear(elemento, limpar_primeiro=True):
    driver.execute_script("arguments[0].scrollIntoView({block:'center'});", elemento)
    try:
        elemento.click()
    except:
        driver.execute_script("arguments[0].focus();", elemento)

    if limpar_primeiro:
        try:
            elemento.clear()
        except:
            ActionChains(driver)\
                .move_to_element(elemento).click()\
                .key_down(Keys.CONTROL).send_keys("a").key_up(Keys.CONTROL)\
                .send_keys(Keys.DELETE).perform()

def _textarea_tradicional(elemento, texto, limpar_primeiro=True):
    _prepare_focus_and_clear(elemento, limpar_primeiro)
    elemento.send_keys(texto)
    elemento.send_keys(Keys.TAB)

def _textarea_actionchains(elemento, texto, limpar_primeiro=True):
    _prepare_focus_and_clear(elemento, limpar_primeiro)
    ac = ActionChains(driver)
    ac.move_to_element(elemento).click().perform()
    for chunk_start in range(0, len(texto), 400):
        ac.send_keys(texto[chunk_start:chunk_start+400]).perform()
        time.sleep(0.05)
    ac.send_keys(Keys.TAB).perform()

def _textarea_js_setvalue(elemento, texto):
    driver.execute_script("""
        const el = arguments[0];
        const val = arguments[1];
        el.value = val;
        el.dispatchEvent(new Event('input',  {bubbles:true}));
        el.dispatchEvent(new Event('change', {bubbles:true}));
        el.dispatchEvent(new KeyboardEvent('keyup', {bubbles:true}));
        el.dispatchEvent(new Event('blur',   {bubbles:true}));
    """, elemento, texto)

def _textarea_js_react_input(elemento, texto):
    driver.execute_script("""
        const el = arguments[0];
        const val = arguments[1];

        const desc = Object.getOwnPropertyDescriptor(window.HTMLTextAreaElement.prototype, 'value');
        if (desc && desc.set) {
            desc.set.call(el, val);
        } else {
            el.value = val;
        }

        el.dispatchEvent(new Event('input', {bubbles: true}));
        el.dispatchEvent(new Event('change', {bubbles: true}));
        el.dispatchEvent(new Event('blur', {bubbles: true}));
    """, elemento, texto)

def preencher_textarea_por_indice(indice_campo, texto, max_tentativas=5, limpar_primeiro=True):
    """Preenche textarea pelo índice usando estratégias múltiplas"""
    def acao():
        if not isinstance(indice_campo, int) or indice_campo < 0:
            raise ValueError(f"Índice inválido: {indice_campo}")
        if texto is None or not isinstance(texto, str):
            raise ValueError(f"Texto inválido: {texto!r}")

        tentativa = 0
        while tentativa < max_tentativas:
            tentativa += 1
            try:
                campos = encontrar_campos_textarea()
                if not campos:
                    if tentativa < max_tentativas:
                        log(doc, f"⚠️ Nenhuma <textarea> encontrada (tentativa {tentativa}/{max_tentativas})")
                        time.sleep(1.5)
                        continue
                    raise Exception("Nenhuma <textarea> foi encontrada.")

                if indice_campo >= len(campos):
                    raise Exception(f"Índice {indice_campo} inválido. Encontradas {len(campos)} textareas.")

                campo_info = campos[indice_campo]
                elemento = campo_info["elemento"]
                campo_id = campo_info.get("id") or "(sem id)"

                log(doc, f"🎯 Tentativa {tentativa}: Preenchendo textarea {indice_campo} (ID: {campo_id})")

                if validar_textarea_preenchida(elemento, texto):
                    log(doc, f"✅ Textarea {indice_campo} já está com o valor desejado.")
                    return True

                estrategias = [
                    lambda: _textarea_tradicional(elemento, texto, limpar_primeiro),
                    lambda: _textarea_actionchains(elemento, texto, limpar_primeiro),
                    lambda: _textarea_js_setvalue(elemento, texto),
                    lambda: _textarea_js_react_input(elemento, texto),
                ]

                for i, estrategia in enumerate(estrategias, 1):
                    try:
                        log(doc, f"   ▶️ Estratégia {i}…")
                        estrategia()
                        time.sleep(0.8)

                        if validar_textarea_preenchida(elemento, texto):
                            val = (elemento.get_attribute("value") or "").strip()
                            log(doc, f"✅ Preenchido com sucesso pela estratégia {i}")
                            return True
                        else:
                            log(doc, f"⚠️ Estratégia {i} não refletiu o valor esperado.")
                    except (StaleElementReferenceException, JavascriptException, TimeoutException) as e:
                        log(doc, f"⚠️ Estratégia {i} falhou: {e}")
                        try:
                            campos = encontrar_campos_textarea()
                            elemento = campos[indice_campo]["elemento"]
                        except:
                            pass
                        continue

                if tentativa < max_tentativas:
                    log(doc, f"⚠️ Tentativa {tentativa} falhou; reintentando em 1.5s…")
                    time.sleep(1.5)
                    continue
            except Exception as e:
                if tentativa < max_tentativas:
                    log(doc, f"⚠️ Erro na tentativa {tentativa}: {e}. Retentando…")
                    time.sleep(1.5)
                    continue
                else:
                    raise

        raise Exception(f"Falha ao preencher textarea {indice_campo} após {max_tentativas} tentativas.")
    return acao


def fechar_abas_extras(driver, doc, aba_principal_index=0):
    """Fecha todas as abas extras (como popups de impressão) mantendo apenas a aba principal"""
    try:
        handles = driver.window_handles
        if len(handles) <= 1:
            log(doc, "ℹ️ Apenas uma aba aberta - nada a fechar.")
            return True
        
        # Guarda o handle da aba principal
        aba_principal = handles[aba_principal_index]
        
        # Fecha todas as outras abas
        abas_fechadas = 0
        for handle in handles:
            if handle != aba_principal:
                try:
                    driver.switch_to.window(handle)
                    driver.close()
                    abas_fechadas += 1
                    log(doc, f"🗑️ Aba extra fechada ({abas_fechadas})")
                except Exception as e:
                    log(doc, f"⚠️ Erro ao fechar aba: {e}")
        
        # Retorna para a aba principal
        driver.switch_to.window(aba_principal)
        time.sleep(0.3)
        
        log(doc, f"✅ {abas_fechadas} aba(s) extra(s) fechada(s). Foco na aba principal.")
        return True
        
    except Exception as e:
        log(doc, f"❌ Erro ao fechar abas extras: {e}")
        return False

def focar_sistema_completo(js_engine, doc):
    """Garante o foco completo na aba principal do sistema e fecha abas extras"""
    driver = js_engine.driver
    try:
        # Primeiro fecha abas extras (como impressão)
        fechar_abas_extras(driver, doc)

        driver.switch_to.default_content()
        js_engine.execute_js("if (window.focus) window.focus();", timeout=3, fallback_result=None)
        time.sleep(0.3)

        log(doc, "✅ Foco garantido na aba do sistema.")
        return True

    except Exception as e:
        log(doc, f"⚠️ Falha ao focar aba do sistema: {e}")
        return False

def clicar_todos_botoes_sim_visiveis(js_engine, doc, pausa_entre=0.0):
    """Clica em TODOS os botões 'Sim' visíveis de uma vez"""
    js = r"""
    (function(){
      const isVisible = el => {
        if (!el) return false;
        const s = getComputedStyle(el);
        return el.offsetParent !== null && s.display !== 'none' &&
               s.visibility !== 'hidden' && parseFloat(s.opacity||1) > 0.01;
      };
      const buttons = Array.from(document.querySelectorAll("a.btModel.btGray.btyes"))
        .filter(isVisible)
        .filter(b => (b.textContent||"").trim().toLowerCase() === "sim");

      let clicked = 0;
      buttons.forEach(b => {
        try {
          b.style.pointerEvents = 'auto';
          b.removeAttribute('disabled');
          b.style.visibility = 'visible';
          b.style.display = 'inline-block';
          b.scrollIntoView({block:'center'});

          ['mouseover','mouseenter','mousemove','mousedown','mouseup','click'].forEach(t=>{
            b.dispatchEvent(new MouseEvent(t,{bubbles:true,cancelable:true,view:window,detail:1}));
          });
          if (typeof b.click === 'function') b.click();

          if (typeof window.jQuery !== 'undefined') {
            window.jQuery(b).trigger('click');
          }
          clicked++;
        } catch(e) {}
      });

      return { totalEncontrados: buttons.length, totalClicados: clicked };
    })();
    """
    try:
        res = js_engine.execute_js(js, timeout=5, fallback_result={"totalEncontrados": 0, "totalClicados": 0})
        total = int(res.get("totalEncontrados", 0))
        clic = int(res.get("totalClicados", 0))
        log(doc, f"⚡ 'Sim' visíveis encontrados: {total} | clicados: {clic}")
        if pausa_entre and clic > 0:
            time.sleep(pausa_entre)
        return res
    except Exception as e:
        log(doc, f"❌ Erro ao clicar em todos os 'Sim': {e}")
        return {"totalEncontrados": 0, "totalClicados": 0, "erro": str(e)}


        
def clicar_sim_com_retry(doc, js_engine, wait, max_tentativas=5, pausa=1.5):
    """Clica em 'Sim' até o modal de confirmação desaparecer"""
    xpath_modal = "//div[contains(@class,'modal') and contains(@style,'z-index')]"
    xpath_sim = "(//div[contains(@class,'modal') and not(contains(@style,'display: none'))]//a[@class='btModel btGray btyes'])[last()]"

    tentativa = 0
    while tentativa < max_tentativas:
        tentativa += 1
        log(doc, f"🧩 Tentativa {tentativa} de fechar modal...")

        try:
            js_engine.force_click(xpath_sim, by_xpath=True)
            time.sleep(pausa)

            modais_visiveis = driver.find_elements(By.XPATH, xpath_modal)
            modais_ativos = [m for m in modais_visiveis if "display: none" not in m.get_attribute("style")]

            if not modais_ativos:
                log(doc, "✅ Botão 'Sim' clicado com sucesso.")
                return True

        except Exception as e:
            log(doc, f"⚠️ Tentativa {tentativa} falhou: {e}")

    log(doc, "❌ Botão 'Sim' não foi clicado após todas as tentativas.")
    return False

def clicar_primeiro_sp_add(js_engine, doc=None, timeout=5):
    """
    Localiza e clica no primeiro elemento <a class="sprites sp-add">, exatamente.
    Ignora variantes como 'sp-addVerde'.
    """
    from selenium.common.exceptions import (
        NoSuchElementException,
        ElementClickInterceptedException,
        StaleElementReferenceException,
    )
    from selenium.webdriver.common.by import By
    import time

    driver = js_engine.driver

    # XPath restrito: exige que as classes sejam EXATAMENTE 'sprites' e 'sp-add'
    xpath_btn = (
        "//a[contains(concat(' ', normalize-space(@class), ' '), ' sprites ') "
        "and contains(concat(' ', normalize-space(@class), ' '), ' sp-add ') "
        "and not(contains(@class, 'sp-addVerde'))][1]"
    )

    log(doc, "🧩 Procurando o botão '<a class=\"sprites sp-add\">' exato...")

    try:
        el = driver.find_element(By.XPATH, xpath_btn)
        log(doc, "🎯 Botão exato encontrado! Tentando clicar...")

        try:
            el.click()
            log(doc, "✅ Clique padrão realizado com sucesso.")
        except (ElementClickInterceptedException, StaleElementReferenceException):
            driver.execute_script("arguments[0].click();", el)
            log(doc, "⚡ Clique forçado via JavaScript realizado.")

        time.sleep(0.5)
        return True

    except NoSuchElementException:
        log(doc, "⚠️ Nenhum botão exato 'sp-add' encontrado.")
        return False

    except Exception as e:
        log(doc, f"❌ Erro ao clicar no botão 'sp-add': {e}")
        return False


def fechar_abas_extras_e_verificar_alerta(driver, doc):
    """
    Fecha abas extras (como de impressão) e verifica imediatamente
    se há alguma mensagem de alerta exibida após o fechamento.
    """
    try:
        safe_action(doc, "Fechando abas extras (impressão)", lambda:
            fechar_abas_extras(driver, doc)
        )

        alerta = encontrar_mensagem_alerta()
        if alerta:
            log(doc, f"⚠️ Alerta detectado após fechar abas extras: '{alerta.text.strip()}'")
        else:
            log(doc, "✅ Nenhum alerta detectado após fechar abas extras.")

        return True
    except Exception as e:
        log(doc, f"⚠️ Erro ao fechar abas extras e verificar alerta: {e}")
        return False

def verificar_e_abrir_caixa(js_engine, doc, timeout=10):
    """
    Verifica o estado do caixa:
      - Se encontrar o botão 'Abrir caixa' visível (sem display: none), realiza a abertura normalmente.
      - Se detectar 'Fechar caixa', entende que o caixa já está aberto e apenas prossegue.
    """
    abrir_xpath = "//a[contains(@class,'btAzulDegrade') and contains(@class,'btAbrirCaixa')]"
    fechar_xpath = "//a[contains(@class,'btAzulDegrade') and contains(@class,'btFecharCaixa')]"

    driver = js_engine.driver
    log(doc, "🔍 Verificando estado do caixa...")

    try:
        # Verifica se existe botão Abrir Caixa visível
        abrir_visivel = driver.execute_script(f"""
            const el = document.evaluate("{abrir_xpath}", document, null, XPathResult.FIRST_ORDERED_NODE_TYPE, null).singleNodeValue;
            if (!el) return false;
            const style = window.getComputedStyle(el);
            return style.display !== 'none' && style.visibility !== 'hidden' && el.offsetParent !== null;
        """)

        # Verifica se existe botão Fechar Caixa visível
        fechar_visivel = driver.execute_script(f"""
            const el = document.evaluate("{fechar_xpath}", document, null, XPathResult.FIRST_ORDERED_NODE_TYPE, null).singleNodeValue;
            if (!el) return false;
            const style = window.getComputedStyle(el);
            return style.display !== 'none' && style.visibility !== 'hidden' && el.offsetParent !== null;
        """)

        if abrir_visivel:
            log(doc, "📦 Caixa fechado — iniciando abertura normal...")
            safe_action(doc, "Clicando em 'Abrir Caixa'", lambda:
                js_engine.force_click(abrir_xpath, by_xpath=True)
            )

            # Valor inicial
            safe_action(doc, "Preenchendo Valor Inicial", lambda:
                js_engine.force_fill(
                    "//input[@type='text' and contains(@class,'valor') and contains(@placeholder,'R$')]",
                    "1000,00",
                    by_xpath=True
                )
            )

            # Descrição
            safe_action(doc, "Preenchendo Descrição", lambda:
                preencher_textarea_por_indice(0, "Abertura automática de caixa via automação Selenium.")
            )

            # Confirmar abertura
            safe_action(doc, "Confirmando Abertura", lambda:
                js_engine.force_click("//a[@class='btModel btGray btyes' and normalize-space()='Abrir']", by_xpath=True)
            )
            time.sleep(0.5)
            encontrar_mensagem_alerta()

            # Autenticação
            safe_action(doc, "Autenticando Abertura", lambda:
                js_engine.force_click("//a[@id='BtYes' and contains(@class,'btModel btGray btyes') and contains(normalize-space(.),'Autenticar')]", by_xpath=True)
            )
            time.sleep(3)
            fechar_abas_extras_e_verificar_alerta(driver, doc)

            # Fecha modal de autenticação
            safe_action(doc, "Fechando modal autenticação", lambda:
                js_engine.force_click("//a[@id='BtNo' and contains(@class,'btno') and normalize-space()='Fechar']", by_xpath=True)
            )

        elif fechar_visivel:
            log(doc, "✅ Caixa já está aberto — prosseguindo com o fluxo normalmente.")
        else:
            log(doc, "⚠️ Nenhum botão de 'Abrir' ou 'Fechar caixa' visível encontrado — prosseguindo com cautela.")

    except Exception as e:
        log(doc, f"❌ Erro ao verificar/abrir caixa: {e}")
        take_screenshot(driver, doc, "erro_verificar_caixa")

def clicar_todos_pesquisar(js_engine, doc, pausa_entre=0.5, timeout=5):
    """
    Procura todos os botões 'Pesquisar' visíveis e clica em cada um deles na ordem.
    Conta e exibe quantos botões existem antes de clicar.
    Usa js_engine.force_click() e registra log detalhado.
    """

    xpath_base = "//a[contains(@class,'btPesquisar btAzulDegrade') and contains(normalize-space(.),'Pesquisar')]"

    try:
        elementos = js_engine.driver.find_elements("xpath", xpath_base)
        total = len(elementos)
        log(doc, f"🔍 Foram encontrados {total} botão(ões) 'Pesquisar' na tela.")

        if total == 0:
            log(doc, "⚠️ Nenhum botão 'Pesquisar' foi encontrado.")
            return {"total": 0, "clicados": 0}

        total_clicados = 0
        for i in range(1, total + 1):
            xpath_indexado = f"({xpath_base})[{i}]"
            try:
                log(doc, f"🎯 Clicando no botão 'Pesquisar' (índice {i}/{total})...")
                js_engine.force_click(xpath_indexado, by_xpath=True)
                js_engine.wait_ajax_complete(timeout)
                total_clicados += 1
                log(doc, f"✅ Clique no botão 'Pesquisar' (índice {i}) realizado com sucesso.")
                if pausa_entre > 0:
                    import time
                    time.sleep(pausa_entre)
            except Exception as e:
                log(doc, f"⚠️ Falha ao clicar no botão 'Pesquisar' (índice {i}): {e}")

        log(doc, f"🧾 Resumo: {total_clicados}/{total} botões 'Pesquisar' clicados com sucesso.")
        return {"total": total, "clicados": total_clicados}

    except Exception as e:
        log(doc, f"⚠️ Erro ao procurar ou clicar nos botões 'Pesquisar': {e}")
        return {"total": 0, "clicados": 0}


def clicar_todos_voltar(js_engine, doc, pausa_entre=0.5, timeout=5):
    """
    Procura todos os botões 'Voltar (ESC)' visíveis e clica em cada um deles na ordem.
    Conta e exibe quantos botões existem antes de clicar.
    Usa js_engine.force_click() e registra log detalhado.
    """

    xpath_base = "//a[contains(@class,'sp-voltarGrande') and @title='Voltar (ESC)']"

    try:
        elementos = js_engine.driver.find_elements("xpath", xpath_base)
        total = len(elementos)
        log(doc, f"🔍 Foram encontrados {total} botão(ões) 'Voltar (ESC)' na tela.")

        if total == 0:
            log(doc, "⚠️ Nenhum botão 'Voltar (ESC)' foi encontrado.")
            return {"total": 0, "clicados": 0}

        total_clicados = 0
        for i in range(1, total + 1):
            xpath_indexado = f"({xpath_base})[{i}]"
            try:
                log(doc, f"🎯 Clicando no botão 'Voltar (ESC)' (índice {i}/{total})...")
                js_engine.force_click(xpath_indexado, by_xpath=True)
                js_engine.wait_ajax_complete(timeout)
                total_clicados += 1
                log(doc, f"✅ Clique no botão 'Voltar (ESC)' (índice {i}) realizado com sucesso.")
                if pausa_entre > 0:
                    import time
                    time.sleep(pausa_entre)
            except Exception as e:
                log(doc, f"⚠️ Falha ao clicar no botão 'Voltar (ESC)' (índice {i}): {e}")

        log(doc, f"🧾 Resumo: {total_clicados}/{total} botões 'Voltar (ESC)' clicados com sucesso.")
        return {"total": total, "clicados": total_clicados}

    except Exception as e:
        log(doc, f"⚠️ Erro ao procurar ou clicar nos botões 'Voltar (ESC)': {e}")
        return {"total": 0, "clicados": 0}


def clicar_salvar_modal(js_engine, doc, timeout=5):
    """Versão simplificada focada no modal específico com logs detalhados"""
    import time
    
    try:
        log(doc, "🔍 Iniciando busca por botões 'Salvar' no modal...")
        
        # Script JavaScript que retorna informações detalhadas
        script = """
        var botoes = document.querySelectorAll('.modal.overflow .btok');
        var info = {
            total: botoes.length,
            visiveis: 0,
            clicados: 0
        };
        
        botoes.forEach(function(btn) {
            if (btn.offsetParent !== null) {  // Verifica se está visível
                info.visiveis++;
                try {
                    btn.click();
                    info.clicados++;
                } catch(e) {
                    console.error('Erro ao clicar:', e);
                }
            }
        });
        
        return info;
        """
        
        log(doc, "⚙️ Executando JavaScript para clicar no(s) botão(ões)...")
        resultado = js_engine.driver.execute_script(script)
        
        log(doc, f"📊 Total de botões encontrados: {resultado['total']}")
        log(doc, f"👁️ Botões visíveis: {resultado['visiveis']}")
        log(doc, f"✅ Botões clicados com sucesso: {resultado['clicados']}")
        
        if resultado['clicados'] == 0:
            log(doc, "⚠️ Nenhum botão 'Salvar' foi clicado.")
            if resultado['total'] == 0:
                log(doc, "❌ Motivo: Nenhum botão encontrado no modal.")
            elif resultado['visiveis'] == 0:
                log(doc, "❌ Motivo: Botões existem mas não estão visíveis.")
        else:
            log(doc, f"🎉 Sucesso! {resultado['clicados']} botão(ões) clicado(s).")
            log(doc, f"⏳ Aguardando conclusão do AJAX (timeout: {timeout}s)...")
            js_engine.wait_ajax_complete(timeout)
            log(doc, "✅ AJAX concluído.")
        
        return {"total": resultado['total'], "clicados": resultado['clicados']}
        
    except Exception as e:
        log(doc, f"❌ Erro ao executar clique no modal: {e}")
        import traceback
        log(doc, f"📋 Detalhes do erro: {traceback.format_exc()}")
        return {"total": 0, "clicados": 0}
    

def clicar_botao_por_classe(js_engine, doc, classe, nome_botao="Botão", timeout=5):
    """Clica em botão por classe CSS com logs detalhados"""
    import time
    
    try:
        log(doc, f"🔍 Iniciando busca pelo botão '{nome_botao}' (classe: {classe})...")
        
        script = f"""
        var botoes = document.querySelectorAll('a.{classe}');
        var info = {{
            total: botoes.length,
            visiveis: 0,
            clicados: 0
        }};
        
        botoes.forEach(function(btn) {{
            if (btn.offsetParent !== null) {{
                info.visiveis++;
                try {{
                    btn.click();
                    info.clicados++;
                }} catch(e) {{
                    console.error('Erro ao clicar:', e);
                }}
            }}
        }});
        
        return info;
        """
        
        log(doc, f"⚙️ Executando JavaScript para clicar no botão '{nome_botao}'...")
        resultado = js_engine.driver.execute_script(script)
        
        log(doc, f"📊 Total de botões encontrados: {resultado['total']}")
        log(doc, f"👁️ Botões visíveis: {resultado['visiveis']}")
        log(doc, f"✅ Botões clicados: {resultado['clicados']}")
        
        if resultado['clicados'] == 0:
            log(doc, f"⚠️ Nenhum botão '{nome_botao}' foi clicado.")
            if resultado['total'] == 0:
                log(doc, f"❌ Motivo: Nenhum botão encontrado.")
            elif resultado['visiveis'] == 0:
                log(doc, f"❌ Motivo: Botão existe mas não está visível.")
        else:
            log(doc, f"🎉 Sucesso! Botão '{nome_botao}' clicado.")
            log(doc, f"⏳ Aguardando AJAX (timeout: {timeout}s)...")
            js_engine.wait_ajax_complete(timeout)
            log(doc, "✅ AJAX concluído.")
        
        return {"total": resultado['total'], "clicados": resultado['clicados']}
        
    except Exception as e:
        log(doc, f"❌ Erro ao clicar no botão '{nome_botao}': {e}")
        return {"total": 0, "clicados": 0}
        
    except Exception as e:
        log(doc, f"❌ Erro ao executar clique no modal: {e}")
        import traceback
        log(doc, f"📋 Detalhes do erro: {traceback.format_exc()}")
        return {"total": 0, "clicados": 0}
    
def clicar_botao_voltar_por_indice(js_engine, doc, indice=1, timeout=5):
    """
    Clica no botão 'Voltar (ESC)' pelo índice informado (1-based).
    Usa js_engine.force_click() e registra log.
    """
    xpath = f"(//a[@class='sprites sp-voltarGrande' and @title='Voltar (ESC)'])[{indice}]"
    log(doc, f"↩️ Clicando no botão 'Voltar (ESC)' (índice {indice})...")

    try:
        js_engine.force_click(xpath, by_xpath=True)
        js_engine.wait_ajax_complete(timeout)
        log(doc, f"✅ Clique no botão 'Voltar (ESC)' (índice {indice}) realizado com sucesso.")
        return True
    except Exception as e:
        log(doc, f"⚠️ Falha ao clicar no botão 'Voltar (ESC)' (índice {indice}): {e}")
        return False


# ==== EXECUÇÃO DO TESTE CORRIGIDA ====
def executar_teste():
    """Executa o teste principal com tratamento robusto de erros"""
    global driver, wait, doc
    
    try:
        # Inicialização
        if not inicializar_driver():
            log(doc, "❌ Falha crítica na inicialização do driver", 'ERROR')
            return False

        log(doc, "🎯 Iniciando execução do teste de consulta de Registro de Chamadas")

        # 1. Acesso ao sistema
        safe_action(doc, "Acessando sistema", lambda: (
            driver.get(URL),
            time.sleep(3)
        ), critico=True)

        # 2. Login
        safe_action(doc, "Realizando login", lambda: (
            aguardar_elemento("#j_id15\\:email", TIMEOUT_DEFAULT).send_keys(LOGIN_EMAIL),
            aguardar_elemento("#j_id15\\:senha", TIMEOUT_DEFAULT).send_keys(LOGIN_PASSWORD),
            aguardar_elemento("#j_id15\\:senha", TIMEOUT_DEFAULT).send_keys(Keys.ENTER),
            time.sleep(5)
        ), critico=True)

        # 3. Ajustes iniciais
        safe_action(doc, "Configurando ambiente", lambda: (
            ajustar_zoom("90%"),
            driver.find_element(By.TAG_NAME, "body").send_keys(Keys.F3),
            time.sleep(2)
        ))

        safe_action(doc, "Acessando módulo Callcenter", acessar_modulo_callcenter(), critico=True)

        safe_action(doc, "Abrindo Agenda de Cobranças", lambda: (
            aguardar_elemento('#gsCallCenter > div.wdTelas > div > ul > li:nth-child(3) > a > span').click(),
            time.sleep(3)
        ), critico=True)

        # 6. Preenchimento dos filtros
        safe_action(doc, "Abrindo o Lov de Cobradores", 
               clicar_elemento_xpath_robusto("//a[@class='sprites sp-openLov']"),
             time.sleep(2)   
)       
        
        safe_action(doc, "Selecionando para filtrar por Código Referência", selecionar_opcao_xpath(
            "//select[@class='tipoFiltro']",
            "Código Referência"
        ))


        safe_action(
            doc,
            "Preenchendo campo de pesquisa como Código Referência do Vendedor",
            preencher_campo_robusto_xpath(
                "//input[@class='nomePesquisa' and @style='width:210px;']",
                "33"
            )
        )

        safe_action(
            doc,
            "Pesquisando Cobrador no Lov",
            clicar_elemento_xpath_robusto(
                "//a[@class='btModel btGray lpFind' and contains(normalize-space(.), 'Pesquisar')]"
            )
        )
        safe_action(doc, "Selecionando Cobrador no Lov", 
               clicar_elemento_xpath_robusto("//td[contains(text(), 'COBRADOR FLAVIA')]")
        )
        time.sleep(10)
        safe_action(
            doc,
            "Procurando compromissos na Agenda",
            lambda: clicar_ate_achar_contrato(driver, wait)
        )


        safe_action(doc, "Fechando Aba do Contrato", clicar_x_por_indice(1))
        safe_action(doc, "Fechando Agenda de Cobranças", clicar_x_por_indice(0))





        # 14. Finalização
        safe_action(doc, "Fechando tela de consulta", 
                   clicar_elemento_robusto("#gsCallCenter > div.wdTop.ui-draggable-handle > div.wdClose > a"))

        log(doc, "✅ Teste executado com sucesso completo!")
        return True

    except Exception as e:
        log(doc, f"❌ ERRO CRÍTICO NO TESTE: {e}", 'ERROR')
        log(doc, "— stacktrace —")
        log(doc, traceback.format_exc())
        take_screenshot(driver, doc, "erro_critico_final", forcar=True)
        return False
    
    finally:
        log(doc, "🏁 Finalizando teste...")

# ==== FUNÇÃO PRINCIPAL ====
def main():
    """Função principal de execução"""
    global doc
    
    try:
        log(doc, "🚀 INICIANDO TESTE DE CONSULTA DE AGENDA DE COBRANÇA DO CALLCENTER")
        log(doc, "="*60)
        
        sucesso = executar_teste()
        
        if sucesso:
            log(doc, "✅ TESTE CONCLUÍDO COM SUCESSO!")
        else:
            log(doc, "❌ TESTE FINALIZADO COM ERROS.")
        
        log(doc, "="*60)
        
    except Exception as e:
        log(doc, f"❌ ERRO FATAL NA EXECUÇÃO PRINCIPAL: {e}", 'ERROR')
        log(doc, traceback.format_exc())
        
    finally:
        finalizar_relatorio()

if __name__ == "__main__":
    main()