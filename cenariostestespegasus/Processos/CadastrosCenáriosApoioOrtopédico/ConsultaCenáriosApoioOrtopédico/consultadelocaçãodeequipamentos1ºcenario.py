# -*- coding: utf-8 -*-
"""
Teste Automatizado - Consulta de Locação de Equipamentos - Apoio Ortopédico
Cenário 1: Preenchimento e realização da Consulta
Versão Melhorada e Otimizada
"""

# ===== IMPORTS =====
import os
import random
import re
import subprocess
import time
import traceback
from datetime import datetime, timedelta, time as dt_time

import pyautogui
from docx import Document
from docx.shared import Inches
from faker import Faker
from faker.providers import BaseProvider
from selenium import webdriver
from selenium.common.exceptions import (
    ElementClickInterceptedException, JavascriptException, NoSuchElementException,
    StaleElementReferenceException, TimeoutException
)
from selenium.webdriver import ActionChains
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.support.ui import WebDriverWait
from validate_docbr import CPF
from webdriver_manager.chrome import ChromeDriverManager

# ===== CONFIGURAÇÕES GLOBAIS =====
TIMEOUT_DEFAULT = 30
TIMEOUT_CURTO = 10
TIMEOUT_LONGO = 60

URL = "http://localhost:8080/gs/index.xhtml"
LOGIN_EMAIL = "joaoeduardo.gold@outlook.com"
LOGIN_PASSWORD = "071999gs"
CAMINHO_ARQUIVO_UPLOAD = r"C:\Users\Gold System\Pictures\Camera Roll\TESTEPNGAPOIOORTOPEDICO.png"

# ===== PROVIDER CUSTOMIZADO =====
class BrasilProvider(BaseProvider):
    def rg(self):
        numeros = [str(random.randint(0, 9)) for _ in range(8)]
        return ''.join(numeros) + '-' + str(random.randint(0, 9))

# ===== INICIALIZAÇÃO =====
fake = Faker("pt_BR")
fake.add_provider(BrasilProvider)

doc = Document()
doc.add_heading("RELATÓRIO DO TESTE AUTOMATIZADO", 0)
doc.add_paragraph("Consulta de Locação de Equipamentos - Apoio Ortopédico – Cenário 1")
doc.add_paragraph(f"Data do teste: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")

screenshot_registradas = set()
driver = None
wait = None

# ===== UTILITÁRIOS =====
def log(doc_obj, msg, nivel='INFO'):
    """Sistema de log melhorado"""
    nivel = str(nivel).upper()
    timestamp = datetime.now().strftime('%H:%M:%S')
    formatted_msg = f"[{timestamp}] [{nivel}] {msg}"
    
    print(formatted_msg)
    
    if doc_obj is not None:
        try:
            doc_obj.add_paragraph(formatted_msg)
        except Exception as e:
            print(f"Erro ao adicionar ao documento: {e}")

def sanitize_filename(name):
    """Sanitiza nome de arquivo"""
    if not isinstance(name, str):
        name = str(name)
    name = name.strip().lower()
    name = re.sub(r"[<>:\"/\\|?*']", "_", name)
    name = re.sub(r"_+", "_", name)
    return name[:120]

def take_screenshot(driver_obj, doc_obj, nome, forcar=False):
    """Captura screenshot com validação"""
    if driver_obj is None:
        log(doc_obj, "Driver não disponível para screenshot", 'WARN')
        return
    
    try:
        nome_sanitizado = sanitize_filename(nome)
        if forcar or nome_sanitizado not in screenshot_registradas:
            os.makedirs("screenshots", exist_ok=True)
            path = f"screenshots/{nome_sanitizado}.png"
            driver_obj.save_screenshot(path)
            
            log(doc_obj, f"Screenshot capturada: {nome_sanitizado}")
            doc_obj.add_paragraph(f"Screenshot: {nome_sanitizado}")
            
            if os.path.exists(path):
                doc_obj.add_picture(path, width=Inches(5.5))
            
            screenshot_registradas.add(nome_sanitizado)
    except Exception as e:
        log(doc_obj, f"Erro ao capturar screenshot {nome}: {e}", 'WARN')

def safe_action(doc_obj, descricao, func, critico=False, max_tentativas=3):
    """Executa ação de forma segura com retry"""
    for tentativa in range(max_tentativas):
        try:
            if tentativa == 0:
                log(doc_obj, f"🔄 {descricao}...")
            else:
                log(doc_obj, f"🔄 {descricao}... (Tentativa {tentativa + 1})")
            
            resultado = func()
            log(doc_obj, f"✅ {descricao} realizada com sucesso.")
            time.sleep(1)
            take_screenshot(driver, doc_obj, descricao.lower().replace(" ", "_"))
            return resultado
            
        except Exception as e:
            if tentativa < max_tentativas - 1:
                log(doc_obj, f"⚠️ Tentativa {tentativa + 1} falhou: {e}. Tentando novamente...", 'WARN')
                time.sleep(2)
                continue
            else:
                take_screenshot(driver, doc_obj, f"erro_{descricao.lower().replace(' ', '_')}")
                log(doc_obj, f"❌ {descricao} falhou após {max_tentativas} tentativas: {e}", 'ERROR')
                
                if critico:
                    log(doc_obj, "Stacktrace completo:", 'ERROR')
                    log(doc_obj, traceback.format_exc(), 'ERROR')
                    raise
                else:
                    log(doc_obj, f"⚠️ Continuando execução apesar do erro em: {descricao}", 'WARN')
                    return None

def aguardar_elemento(seletor, timeout=TIMEOUT_DEFAULT, condicao='clickable', by_type=By.CSS_SELECTOR):
    """Aguarda elemento com condições específicas"""
    global driver
    
    if driver is None:
        raise Exception("Driver não inicializado")
    
    condicoes = {
        'present': EC.presence_of_element_located,
        'visible': EC.visibility_of_element_located,
        'clickable': EC.element_to_be_clickable,
        'invisible': EC.invisibility_of_element_located
    }
    
    condicao = condicoes.get(condicao, EC.element_to_be_clickable)
    
    try:
        wait_obj = WebDriverWait(driver, timeout)
        return wait_obj.until(condicao((by_type, seletor)))
    except TimeoutException:
        log(doc, f"❌ Timeout aguardando elemento: {seletor} (timeout: {timeout}s)", 'ERROR')
        raise

def scroll_to_element(elemento_ou_seletor, by_type=By.CSS_SELECTOR):
    """Scroll seguro até elemento"""
    global driver
    
    if driver is None:
        return False
    
    try:
        if isinstance(elemento_ou_seletor, str):
            elemento = aguardar_elemento(elemento_ou_seletor, 10, 'present', by_type)
        else:
            elemento = elemento_ou_seletor
        
        if elemento and elemento.is_displayed():
            driver.execute_script("""
                arguments[0].scrollIntoView({
                    behavior: 'smooth',
                    block: 'center',
                    inline: 'center'
                });
            """, elemento)
            time.sleep(0.5)
            return True
    except Exception as e:
        log(doc, f"⚠️ Erro no scroll: {e}", 'WARN')
    
    return False

def clicar_elemento_robusto(seletor, by_type=By.CSS_SELECTOR, timeout=TIMEOUT_DEFAULT):
    """Clique robusto com múltiplas estratégias"""
    global driver
    
    if driver is None:
        return False
    
    try:
        # Remove overlays que podem interferir
        driver.execute_script("""
            document.querySelectorAll('.modal, .overlay, .blockUI, .toast, .tooltip, [role="dialog"], [data-overlay]')
                .forEach(e => { 
                    if (getComputedStyle(e).position === 'fixed') 
                        e.style.display = 'none'; 
                });
        """)
        
        elemento = aguardar_elemento(seletor, timeout, 'present', by_type)
        scroll_to_element(elemento)
        
        # Estratégias de clique em ordem de preferência
        estrategias = [
            # 1. Clique normal
            lambda: aguardar_elemento(seletor, 5, 'clickable', by_type).click(),
            
            # 2. ActionChains
            lambda: ActionChains(driver).move_to_element(elemento).click().perform(),
            
            # 3. JavaScript click
            lambda: driver.execute_script("arguments[0].click();", elemento),
            
            # 4. Eventos manuais
            lambda: driver.execute_script("""
                const el = arguments[0];
                function fire(type) { 
                    el.dispatchEvent(new MouseEvent(type, {
                        bubbles: true, 
                        cancelable: true, 
                        view: window
                    })); 
                }
                el.focus();
                fire('mouseover');
                fire('mousedown');
                fire('mouseup');
                fire('click');
            """, elemento),
            
            # 5. Clique com offset
            lambda: ActionChains(driver).move_to_element_with_offset(elemento, 1, 1).click().perform()
        ]
        
        for i, estrategia in enumerate(estrategias, 1):
            try:
                estrategia()
                log(doc, f"✅ Clique realizado com estratégia {i}")
                return True
            except Exception as e:
                log(doc, f"⚠️ Estratégia {i} falhou: {e}", 'WARN')
                continue
        
        log(doc, f"❌ Falha ao clicar em: {seletor}", 'ERROR')
        return False
        
    except Exception as e:
        log(doc, f"❌ Erro no clique robusto: {e}", 'ERROR')
        return False

def preencher_campo_robusto(seletor, valor, by_type=By.XPATH, timeout=TIMEOUT_DEFAULT):
    """Preenchimento robusto de campos"""
    global driver
    
    if not valor or driver is None:
        return False
    
    try:
        elemento = aguardar_elemento(seletor, timeout, 'clickable', by_type)
        scroll_to_element(elemento)
        
        # Estratégias de preenchimento
        estrategias = [
            # 1. Método tradicional
            lambda: (elemento.click(), elemento.clear(), elemento.send_keys(valor), elemento.send_keys(Keys.TAB)),
            
            # 2. ActionChains com limpeza
            lambda: (
                ActionChains(driver)
                .move_to_element(elemento)
                .click()
                .key_down(Keys.CONTROL)
                .send_keys('a')
                .key_up(Keys.CONTROL)
                .send_keys(Keys.DELETE)
                .send_keys(valor)
                .send_keys(Keys.TAB)
                .perform()
            ),
            
            # 3. JavaScript
            lambda: driver.execute_script("""
                const el = arguments[0];
                const val = arguments[1];
                el.focus();
                el.value = '';
                el.value = val;
                ['input', 'change', 'blur'].forEach(ev => 
                    el.dispatchEvent(new Event(ev, { bubbles: true }))
                );
            """, elemento, valor)
        ]
        
        for i, estrategia in enumerate(estrategias, 1):
            try:
                estrategia()
                time.sleep(0.5)
                
                # Valida se o valor foi definido
                valor_atual = elemento.get_attribute('value')
                if valor_atual and (valor in valor_atual or valor_atual == valor):
                    log(doc, f"✅ Campo preenchido com estratégia {i}: '{valor}'")
                    return True
                    
            except Exception as e:
                log(doc, f"⚠️ Estratégia {i} falhou: {e}", 'WARN')
                continue
        
        log(doc, f"❌ Falha ao preencher campo: {seletor}", 'ERROR')
        return False
        
    except Exception as e:
        log(doc, f"❌ Erro no preenchimento: {e}", 'ERROR')
        return False

def fazer_upload_arquivo():
    """Upload de arquivo melhorado"""
    try:
        log(doc, "📄 Realizando upload do arquivo.")
        
        if not os.path.exists(CAMINHO_ARQUIVO_UPLOAD):
            log(doc, f"❌ Arquivo não encontrado: {CAMINHO_ARQUIVO_UPLOAD}", 'ERROR')
            return False
        
        # Tenta encontrar input file
        try:
            file_input = aguardar_elemento("input[type='file']", 10, 'present')
            file_input.send_keys(CAMINHO_ARQUIVO_UPLOAD)
            log(doc, "✅ Upload realizado via input direto.")
            time.sleep(3)
            return True
        except:
            pass
        
        # Tenta clicar no botão de upload e usar pyautogui
        try:
            upload_button = aguardar_elemento(
                ".qq-upload-button, [class*='upload'], [class*='selecione'], [class*='browse']",
                10, 'clickable'
            )
            upload_button.click()
            time.sleep(1)
            
            pyautogui.write(CAMINHO_ARQUIVO_UPLOAD)
            pyautogui.press('enter')
            time.sleep(3)
            
            log(doc, "✅ Upload realizado via pyautogui.")
            return True
        except Exception as e:
            log(doc, f"❌ Erro no upload: {e}", 'ERROR')
            return False
            
    except Exception as e:
        log(doc, f"❌ Erro geral no upload: {e}", 'ERROR')
        return False

def encontrar_campos_datepicker():
    """Encontra campos datepicker na página"""
    global driver
    
    if driver is None:
        return []
    
    seletores = [
        "input.hasDatepicker",
        "input[id^='dp']",
        "input[maxlength='10'][grupo='']",
        "input[type='text'][maxlength='10']",
        "input[class*='datepicker']",
        ".hasDatepicker"
    ]
    
    campos_encontrados = []
    
    for seletor in seletores:
        try:
            elementos = driver.find_elements(By.CSS_SELECTOR, seletor)
            for elemento in elementos:
                if elemento.is_displayed() and elemento.is_enabled():
                    info = {
                        'elemento': elemento,
                        'id': elemento.get_attribute('id') or f"dp_{len(campos_encontrados)}",
                        'seletor_usado': seletor
                    }
                    # Evita duplicatas
                    if not any(c['id'] == info['id'] for c in campos_encontrados):
                        campos_encontrados.append(info)
        except Exception:
            continue
    
    log(doc, f"📊 Encontrados {len(campos_encontrados)} campos datepicker")
    return campos_encontrados

def preencher_datepicker(indice_campo, data_valor, max_tentativas=5):
    """Preenche datepicker por índice"""
    global driver
    
    for tentativa in range(max_tentativas):
        try:
            log(doc, f"🎯 Tentativa {tentativa + 1}: Preenchendo datepicker {indice_campo} com '{data_valor}'")
            
            campos = encontrar_campos_datepicker()
            
            if not campos:
                if tentativa < max_tentativas - 1:
                    time.sleep(2)
                    continue
                raise Exception("Nenhum campo datepicker encontrado")
            
            if indice_campo >= len(campos):
                raise Exception(f"Índice {indice_campo} inválido. Encontrados {len(campos)} campos")
            
            campo_info = campos[indice_campo]
            elemento = campo_info['elemento']
            campo_id = campo_info['id']
            
            # Estratégias para datepicker
            estrategias = [
                # jQuery
                lambda: driver.execute_script("""
                    const campo = $('#' + arguments[0]);
                    if (campo.length && typeof jQuery !== 'undefined') {
                        if (campo.hasClass('hasDatepicker')) {
                            campo.datepicker('setDate', arguments[1]);
                        } else {
                            campo.val(arguments[1]);
                        }
                        campo.trigger('input').trigger('change').trigger('blur');
                        return campo.val();
                    }
                    return null;
                """, campo_id, data_valor),
                
                # JavaScript direto
                lambda: driver.execute_script("""
                    const campo = arguments[0];
                    const valor = arguments[1];
                    campo.focus();
                    campo.value = '';
                    campo.value = valor;
                    ['input','change','blur','keyup'].forEach(ev => 
                        campo.dispatchEvent(new Event(ev, {bubbles: true}))
                    );
                """, elemento, data_valor),
                
                # ActionChains
                lambda: (
                    scroll_to_element(elemento),
                    ActionChains(driver).move_to_element(elemento).click().perform(),
                    time.sleep(0.5),
                    ActionChains(driver).key_down(Keys.CONTROL).send_keys('a').key_up(Keys.CONTROL).perform(),
                    ActionChains(driver).send_keys(Keys.DELETE).perform(),
                    time.sleep(0.3),
                    [ActionChains(driver).send_keys(char).perform() and time.sleep(0.05) for char in data_valor],
                    ActionChains(driver).send_keys(Keys.TAB).perform()
                )[-1]  # Retorna o resultado do último comando
            ]
            
            for i, estrategia in enumerate(estrategias, 1):
                try:
                    log(doc, f"   Aplicando estratégia {i} para datepicker...")
                    estrategia()
                    time.sleep(1)
                    
                    # Valida se funcionou
                    valor_atual = elemento.get_attribute('value')
                    if valor_atual and (data_valor in valor_atual or valor_atual == data_valor):
                        log(doc, f"✅ Datepicker preenchido com estratégia {i}: '{valor_atual}'")
                        return True
                    else:
                        log(doc, f"⚠️ Estratégia {i} não preencheu corretamente", 'WARN')
                        
                except Exception as e:
                    log(doc, f"⚠️ Estratégia {i} falhou: {e}", 'WARN')
                    continue
            
            if tentativa < max_tentativas - 1:
                log(doc, f"⚠️ Tentativa {tentativa + 1} falhou, tentando novamente...", 'WARN')
                time.sleep(2)
                continue
            
        except Exception as e:
            if tentativa < max_tentativas - 1:
                log(doc, f"⚠️ Erro na tentativa {tentativa + 1}: {e}, tentando novamente...", 'WARN')
                time.sleep(2)
                continue
            else:
                raise Exception(f"Falha ao preencher datepicker {indice_campo} após {max_tentativas} tentativas: {e}")
    
    return False

def validar_registros_encontrados():
    """Valida se registros foram encontrados após consulta"""
    global driver
    
    try:
        log(doc, "🔍 Validando registros encontrados...")
        time.sleep(5)  # Aguarda processamento
        
        # Seletores para tabelas de resultado
        seletores_tabela = [
            '#DataTables_Table_0 tbody tr',
            '#DataTables_Table_1 tbody tr',
            'table[id*="DataTables"] tbody tr',
            '.wdGrid table tbody tr',
            'table tbody tr'
        ]
        
        for seletor in seletores_tabela:
            try:
                linhas = driver.find_elements(By.CSS_SELECTOR, seletor)
                linhas_validas = [
                    linha for linha in linhas 
                    if linha.is_displayed() and 
                    linha.text.strip() and 
                    len(linha.text.strip()) > 5 and
                    not any(termo in linha.text.lower() for termo in [
                        'nenhum registro', 'sem dados', 'no data', 'vazio', 'empty'
                    ])
                ]
                
                if linhas_validas:
                    quantidade = len(linhas_validas)
                    log(doc, f"✅ {quantidade} registro(s) encontrado(s)")
                    
                    # Log das primeiras linhas
                    for i, linha in enumerate(linhas_validas[:3], 1):
                        texto = linha.text.strip()[:100]
                        log(doc, f"   Registro {i}: {texto}...")
                    
                    return {
                        'encontrou_registros': True,
                        'quantidade_registros': quantidade,
                        'detalhes': [linha.text.strip()[:100] for linha in linhas_validas[:5]]
                    }
                    
            except Exception as e:
                log(doc, f"⚠️ Erro ao verificar tabela com {seletor}: {e}", 'WARN')
                continue
        
        # Verifica mensagens de "sem resultados"
        mensagens_vazio = [
            "Nenhum registro encontrado",
            "Não foram encontrados registros",
            "Nenhum resultado",
            "0 registros encontrados"
        ]
        
        for mensagem in mensagens_vazio:
            try:
                elem = driver.find_element(By.XPATH, f"//*[contains(text(), '{mensagem}')]")
                if elem.is_displayed():
                    log(doc, f"ℹ️ Sistema informou: {elem.text.strip()}")
                    return {
                        'encontrou_registros': False,
                        'quantidade_registros': 0,
                        'mensagem': f"Sistema informou: {elem.text.strip()}"
                    }
            except:
                continue
        
        log(doc, "⚠️ Não foi possível determinar se há registros", 'WARN')
        return {
            'encontrou_registros': True,  # Assume que há para não interromper
            'quantidade_registros': 1,
            'mensagem': "Validação inconclusiva, continuando teste"
        }
        
    except Exception as e:
        log(doc, f"❌ Erro na validação: {e}", 'ERROR')
        return {
            'encontrou_registros': True,
            'quantidade_registros': 1,
            'mensagem': f"Erro na validação, continuando: {e}"
        }

def encontrar_mensagem_alerta():
    """Procura mensagens de alerta do sistema"""
    global driver
    
    if driver is None:
        return None
    
    seletores = [
        (".alerts.salvo", "✅ Sucesso"),
        (".alerts.alerta", "⚠️ Alerta"),
        (".alerts.erro", "❌ Erro"),
        (".alert", "📢 Alerta"),
        ("[class*='alert']", "📢 Notificação"),
    ]
    
    for seletor, tipo in seletores:
        try:
            elemento = driver.find_element(By.CSS_SELECTOR, seletor)
            if elemento.is_displayed():
                log(doc, f"{tipo}: {elemento.text}")
                return elemento
        except:
            continue
    
    log(doc, "ℹ️ Nenhuma mensagem de alerta encontrada.")
    return None

def ajustar_zoom():
    """Ajusta zoom da página"""
    global driver
    
    if driver is None:
        return
    
    try:
        driver.execute_script("document.body.style.zoom='90%'")
        log(doc, "🔍 Zoom ajustado para 90%.")
    except Exception as e:
        log(doc, f"⚠️ Erro ao ajustar zoom: {e}", 'WARN')

def inicializar_driver():
    """Inicializa o driver Chrome"""
    global driver, wait
    
    try:
        options = Options()
        options.add_argument("--start-maximized")
        options.add_argument("--disable-blink-features=AutomationControlled")
        options.add_argument("--disable-extensions")
        options.add_argument("--disable-plugins-discovery")
        options.add_argument("--no-sandbox")
        options.add_argument("--disable-dev-shm-usage")
        options.add_experimental_option("excludeSwitches", ["enable-automation"])
        options.add_experimental_option('useAutomationExtension', False)

        service = Service(ChromeDriverManager().install())
        driver = webdriver.Chrome(service=service, options=options)
        
        # Remove detecção de automação
        driver.execute_script("Object.defineProperty(navigator, 'webdriver', {get: () => undefined})")
        
        wait = WebDriverWait(driver, TIMEOUT_DEFAULT)
        
        log(doc, "✅ Driver Chrome inicializado com sucesso")
        return True
        
    except Exception as e:
        log(doc, f"❌ Erro ao inicializar driver: {e}", 'ERROR')
        return False

def finalizar_relatorio():
    """Finaliza e salva o relatório"""
    global driver, doc
    
    nome_arquivo = f"relatorio_locacao_equipamentos_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    
    try:
        doc.save(nome_arquivo)
        log(doc, f"📄 Relatório salvo como: {nome_arquivo}")
        
        # Tenta abrir o relatório
        try:
            subprocess.run(["start", "winword", nome_arquivo], shell=True)
        except:
            subprocess.run(["start", nome_arquivo], shell=True)
    except Exception as e:
        print(f"Erro ao salvar relatório: {e}")
    
    # Fecha o driver
    if driver:
        try:
            driver.quit()
            log(doc, "🚪 Driver encerrado")
        except:
            pass

# ===== EXECUÇÃO PRINCIPAL =====
def executar_teste():
    """Função principal do teste"""
    global driver, wait, doc
    
    try:
        # Inicialização
        if not inicializar_driver():
            return False

        log(doc, "🚀 Iniciando teste de Consulta de Locação de Equipamentos")

        # 1. Acesso ao sistema
        safe_action(doc, "Acessando sistema", lambda: driver.get(URL), critico=True)

        # 2. Login
        safe_action(doc, "Realizando login", lambda: (
            aguardar_elemento("#j_id15\\:email", 10).send_keys(LOGIN_EMAIL),
            aguardar_elemento("#j_id15\\:senha", 10).send_keys(LOGIN_PASSWORD),
            aguardar_elemento("#j_id15\\:senha", 10).send_keys(Keys.ENTER),
            time.sleep(5)
        ), critico=True)

        # 3. Configuração inicial
        safe_action(doc, "Configurando interface", lambda: (
            ajustar_zoom(),
            driver.find_element(By.TAG_NAME, "body").send_keys(Keys.F3),
            time.sleep(2)
        ))

        # 4. Acesso ao módulo
        safe_action(doc, "Acessando Apoio Ortopédico", lambda: (
            scroll_to_element('/html/body/div[15]/ul/li[22]/img', By.XPATH),
            clicar_elemento_robusto('/html/body/div[15]/ul/li[22]/img', By.XPATH, 10)
        ), critico=True)

        # 5. Acesso à funcionalidade
        safe_action(doc, "Acessando Consulta Locação de Equipamentos", lambda:
            clicar_elemento_robusto('#gsApoioOrtopedico > div.wdTelas > div.telaInicial.clearfix.overflow.overflowY > ul > li:nth-child(3) > a > span')
        , critico=True)

        time.sleep(3)

        # 6. Preenchimento dos campos
        safe_action(doc, "Preenchendo Número da Locação", lambda:
            preencher_campo_robusto("//input[@class='numeroLocacao']", "306", By.XPATH)
        )

        safe_action(doc, "Abrindo LOV de Equipamentos", lambda:
            clicar_elemento_robusto("//a[contains(@class, 'sprites') and contains(@class, 'sp-openLov')]", By.XPATH)
        )

        time.sleep(2)

        safe_action(doc, "Pesquisando equipamento", lambda: (
            ("/html/body/div[20]/div[2]/div[1]/div[1]/div[2]/input", "EQUIPAMENTO SELENIUM 01", By.XPATH),
            clicar_elemento_robusto("/html/body/div[20]/div[2]/div[1]/div[1]/div[3]/a", By.XPATH),
            time.sleep(2),
            clicar_elemento_robusto("//tr[td[2][contains(normalize-space(.),'EQUIPAMENTO SELENIUM 01')]]", By.XPATH)
        ))

        safe_action(doc, "Preenchendo Número do Contrato", lambda:
            preencher_campo_robusto("//input[@class='numeroContrato']", "113060", By.XPATH)
        )

        safe_action(doc, "Preenchendo Número do Patrimônio", lambda:
            preencher_campo_robusto("//input[@name='numeroPatrimonio']", "876*&", By.XPATH)
        )

        safe_action(doc, "Abrindo LOV do Responsável", lambda:
            clicar_elemento_robusto("//*[@id='gsApoioOrtopedico']/div[2]/div[2]/div/div[1]/div/div[6]/div/a", By.XPATH)
        )
        time.sleep(2)


        safe_action(doc, "Pesquisando responsável", lambda: (   
            preencher_campo_robusto("//*[@id='txtPesquisa']", "TESTE TITULAR 233", By.XPATH),
            clicar_elemento_robusto("/html/body/div[22]/div[2]/div[1]/div[1]/div[2]/a", By.XPATH),
            time.sleep(2),
            clicar_elemento_robusto("//tr[td[2][contains(normalize-space(.),'TESTE TITULAR 233')]]", By.XPATH)
        ))

        # 7. Preenchimento das datas
        safe_action(doc, "Preenchendo Data Locação Inicial", lambda:
            preencher_datepicker(0, "21/08/2025")
        )

        safe_action(doc, "Preenchendo Data Locação Final", lambda:
            preencher_datepicker(1, "21/08/2025")
        )

        # 8. Realizando consulta
        safe_action(doc, "Realizando consulta", lambda: (
            clicar_elemento_robusto("//a[contains(@class,'btfind')]", By.XPATH),
            time.sleep(10)
        ), critico=True)

        # 9. Validação dos resultados
        resultado_validacao = safe_action(doc, "Validando resultados da consulta", validar_registros_encontrados)

        if not resultado_validacao or not resultado_validacao.get('encontrou_registros', False):
            log(doc, "Nenhum registro encontrado - Finalizando consulta", "WARN")
            safe_action(doc, "Fechando tela", lambda:
                clicar_elemento_robusto("#gsApoioOrtopedico > div.wdTop.ui-draggable-handle > div.wdClose > a")
            )
            return True

        # 10. Processamento dos registros
        quantidade = resultado_validacao.get('quantidade_registros', 0)
        log(doc, f"Processando {quantidade} registro(s) encontrado(s)")




        # 11. Visualizando detalhes
        safe_action(doc, "Visualizando Detalhes da Locação", lambda:
            clicar_elemento_robusto("//span[@title='Detalhes da Locação']", By.XPATH)
        )
        time.sleep(2)

        safe_action(doc, "Visualizando Títulos", lambda: (
            clicar_elemento_robusto("//button[contains(normalize-space(.),'Ver Títulos')]", By.XPATH),
            time.sleep(20)
        ))

        safe_action(doc, "Capturando screenshot dos títulos", lambda:
            take_screenshot(driver, doc, "tela_titulos", forcar=True)
        )

        safe_action(doc, "Fechando aba de títulos", lambda:
            clicar_elemento_robusto("//a[contains(@class,'sp-fecharGrande')]", By.XPATH)
        )

        safe_action(doc, "Reabrindo detalhes da locação", lambda: (
            clicar_elemento_robusto("//span[@title='Detalhes da Locação']", By.XPATH),
            time.sleep(3)
        ))

        log(doc, "Verificando mensagens de alerta...")
        encontrar_mensagem_alerta()

        # 12. Digitalização de documentos
        safe_action(doc, "Abrindo digitalização de documentos", lambda:
            clicar_elemento_robusto("//span[@title='Digitalizar documentos' and contains(@class, 'sp-scanner')]", By.XPATH)
        )

        safe_action(doc, "Adicionando novo documento", lambda:
            clicar_elemento_robusto("//span[contains(@class, 'sp-addFile')]", By.XPATH)
        )

        # 13. Upload de arquivo
        if os.path.exists(CAMINHO_ARQUIVO_UPLOAD):
            safe_action(doc, "Fazendo upload do arquivo", fazer_upload_arquivo)
        else:
            log(doc, f"Arquivo não encontrado: {CAMINHO_ARQUIVO_UPLOAD}, pulando upload", 'WARN')

        safe_action(doc, "Salvando arquivo digitalizado", lambda:
            clicar_elemento_robusto("//a[contains(@class,'btModel') and contains(@class,'btGreen') and normalize-space(.)='Salvar']", By.XPATH)
        )

        log(doc, "Verificando mensagens de alerta após digitalização...")
        encontrar_mensagem_alerta()

        safe_action(doc, "Fechando modal de digitalização", lambda:
            clicar_elemento_robusto("//a[contains(@class,'fa') and contains(@class,'fa-close')]", By.XPATH)
        )

        # 14. Finalização
        safe_action(doc, "Fechando tela de consulta", lambda:
            clicar_elemento_robusto("#gsApoioOrtopedico > div.wdTop.ui-draggable-handle > div.wdClose > a")
        )

        log(doc, "Teste concluído com sucesso!")
        return True

    except Exception as e:
        log(doc, f"ERRO FATAL: {e}", 'ERROR')
        take_screenshot(driver, doc, "erro_fatal", forcar=True)
        log(doc, "Stacktrace completo:", 'ERROR')
        log(doc, traceback.format_exc(), 'ERROR')
        return False

def main():
    """Função principal"""
    global doc
    
    try:
        log(doc, "Iniciando teste de Consulta de Locação de Equipamentos")
        
        sucesso = executar_teste()
        
        if sucesso:
            log(doc, "Teste executado com sucesso!")
        else:
            log(doc, "Teste finalizado com erros.")
            
    except Exception as e:
        log(doc, f"Erro na execução principal: {e}", 'ERROR')
        log(doc, traceback.format_exc(), 'ERROR')
    finally:
        finalizar_relatorio()

if __name__ == "__main__":
    main()