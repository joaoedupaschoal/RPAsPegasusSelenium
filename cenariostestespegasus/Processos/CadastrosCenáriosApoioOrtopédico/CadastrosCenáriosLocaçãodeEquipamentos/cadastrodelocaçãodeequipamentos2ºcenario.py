
# ==== IMPORTS (sem conflitos) ====
from datetime import datetime, timedelta
from datetime import time as dt_time  # usar para objetos de hora
import time                           # usar para time.sleep(...)
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver import ActionChains
from selenium.common.exceptions import TimeoutException, WebDriverException, StaleElementReferenceException, JavascriptException
from selenium.webdriver.support import expected_conditions as EC
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support.ui import Select, WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service
from selenium.common.exceptions import TimeoutException, NoSuchElementException, StaleElementReferenceException, ElementClickInterceptedException
from webdriver_manager.chrome import ChromeDriverManager
from docx import Document
from docx.shared import Inches
from faker import Faker
from faker.providers import BaseProvider
from validate_docbr import CPF
from selenium.webdriver import ActionChains
import subprocess
import os
import random
import re

# ==== PROVIDERS CUSTOMIZADOS ====
class BrasilProvider(BaseProvider):
    def rg(self):
        numeros = [str(random.randint(0, 9)) for _ in range(8)]
        return ''.join(numeros) + '-' + str(random.randint(0, 9))

fake = Faker("pt_BR")
fake.add_provider(BrasilProvider)

# ==== GERAÇÃO DE DATAS ====
def gerar_datas_validas(hora_padrao="00:00", dias_fim=0):
    """Gera datas coerentes. data_inicio/data_fim no formato 'dd/MM/yyyy HH:mm'."""
    hoje_date = datetime.today().date()
    dez_anos_atras = hoje_date - timedelta(days=3650)

    # Falecimento entre 10 anos atrás e hoje
    data_falecimento = fake.date_between(start_date=dez_anos_atras, end_date=hoje_date)

    # Nascimento (entre 18 e 110 anos antes do falecimento)
    idade_minima, idade_maxima = 18, 110
    data_nascimento = data_falecimento - timedelta(days=random.randint(idade_minima * 365, idade_maxima * 365))

    # Sepultamento 1..10 dias após o falecimento
    data_sepultamento = data_falecimento + timedelta(days=random.randint(1, 10))

    # Registro 1..10 dias após o sepultamento
    data_registro = data_sepultamento + timedelta(days=random.randint(1, 10))

    # Velório entre o falecimento e o sepultamento
    data_velorio = fake.date_between(start_date=data_falecimento, end_date=data_sepultamento)

    # Início entre 2 e 30 dias no futuro, com hora escolhida
    data_inicio_date = hoje_date + timedelta(days=random.randint(2, 30))

    # Monta datetime com hora escolhida (ex: "00:00")
    h, m = map(int, hora_padrao.split(":"))
    dt_inicio = datetime.combine(data_inicio_date, dt_time(h, m))

    # Fim: mesmo dia por padrão (dias_fim=0). Ajuste se quiser +N dias.
    dt_fim = dt_inicio + timedelta(days=dias_fim)

    fmt_data = "%d/%m/%Y"
    fmt_dt = "%d/%m/%Y %H:%M"

    return (
        data_nascimento.strftime(fmt_data),
        data_falecimento.strftime(fmt_data),
        data_sepultamento.strftime(fmt_data),
        data_velorio.strftime(fmt_data),
        dt_inicio.strftime(fmt_dt),   # data_inicio com hora
        dt_fim.strftime(fmt_dt),      # data_fim com hora
        data_registro.strftime(fmt_data),
        hoje_date.strftime(fmt_data),
    )

(data_nascimento, data_falecimento, data_sepultamento,
 data_velorio, data_inicio, data_fim, data_registro, hoje) = gerar_datas_validas(
    hora_padrao="08:50",
    dias_fim=0
)

hora_falecimento = fake.time(pattern="%H:%M")
hora_sepultamento = fake.time(pattern="%H:%M")
localizacao = fake.city()

# ==== CONFIGURAÇÕES ====
URL = "http://localhost:8080/gs/index.xhtml"
LOGIN_EMAIL = "joaoeduardo.gold@outlook.com"
LOGIN_PASSWORD = "071999gs"

# ==== INICIALIZAÇÃO DE VARIÁVEIS GLOBAIS ====
doc = Document()
doc.add_heading("RELATÓRIO DO TESTE", 0)
doc.add_paragraph("Cadastro de Locação de Equipamentos - Apoio Ortopédico – Cenário 2: Preenchimento completo e cancelamento")
doc.add_paragraph(f"Data do teste: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")

screenshot_registradas = set()
driver = None
wait = None

# ==== UTILITÁRIOS ====
def log(doc, msg):
    print(msg)
    doc.add_paragraph(msg)

def _sanitize_filename(name: str) -> str:
    name = name.strip().lower()
    name = re.sub(r"[<>:\"/\\|?*']", "_", name)
    name = re.sub(r"_+", "_", name)
    return name[:120]

def take_screenshot(driver, doc, nome):
    if driver is None:
        return
    nome = _sanitize_filename(nome)
    if nome not in screenshot_registradas:
        path = f"screenshots/{nome}.png"
        os.makedirs("screenshots", exist_ok=True)
        try:
            driver.save_screenshot(path)
            doc.add_paragraph(f"Screenshot: {nome}")
            doc.add_picture(path, width=Inches(5.5))
            screenshot_registradas.add(nome)
        except Exception as e:
            log(doc, f"⚠️ Erro ao tirar screenshot {nome}: {e}")

def safe_action(doc, descricao, func, max_retries=3):
    global driver
    for attempt in range(max_retries):
        try:
            log(doc, f"🔄 {descricao}..." if attempt == 0 else f"🔄 {descricao}... (Tentativa {attempt + 1})")
            func()
            log(doc, f"✅ {descricao} realizada com sucesso.")
            take_screenshot(driver, doc, _sanitize_filename(descricao))
            return True
        except (TimeoutException, NoSuchElementException, StaleElementReferenceException) as e:
            if attempt < max_retries - 1:
                log(doc, f"⚠️ Tentativa {attempt + 1} falhou para {descricao}, tentando novamente...")
                time.sleep(2)
                continue
            else:
                log(doc, f"❌ Erro ao {descricao.lower()} após {max_retries} tentativas: {e}")
                take_screenshot(driver, doc, _sanitize_filename(f"erro_{descricao}"))
                return False
        except Exception as e:
            log(doc, f"❌ Erro inesperado ao {descricao.lower()}: {e}")
            take_screenshot(driver, doc, _sanitize_filename(f"erro_{descricao}"))
            return False




# ==== SISTEMA ANTI-TIMEOUT JAVASCRIPT ====
class JSTimeoutHandler:
    """Sistema robusto para lidar com timeouts JavaScript no Selenium"""
    
    def __init__(self, driver, doc, timeout_padrao=10, max_retries=3):
        self.driver = driver
        self.doc = doc
        self.timeout_padrao = timeout_padrao
        self.max_retries = max_retries
        self.last_error = None
        
    def log_timeout(self, msg, level="INFO"):
        """Log com timestamp"""

        prefix = {
            "INFO": "ℹ️ ",
            "WARN": "⚠️ ",
            "ERROR": "❌ ",
            "SUCCESS": "✅ "
        }.get(level, "📝 ")
        
        print(f" {prefix} {msg}")
        if hasattr(self.doc, 'add_paragraph'):
            self.doc.add_paragraph(f"{msg}")
    
    def execute_js_safe(self, script, *args, timeout=None, fallback_result=None):
        """Executa JavaScript com proteção contra timeouts"""
        timeout = timeout or self.timeout_padrao
        
        original_timeout = self.driver.timeouts.script
        self.driver.set_script_timeout(timeout)
        
        for tentativa in range(1, self.max_retries + 1):
            try:
                if tentativa > 1:
                    self.log_timeout(f"Tentativa {tentativa}/{self.max_retries}", "INFO")
                
                result = self.driver.execute_script(script, *args)
                self.driver.set_script_timeout(original_timeout)
                
                if tentativa > 1:
                    self.log_timeout("JavaScript executado com sucesso", "SUCCESS")
                return result
                
            except JavascriptException as e:
                self.last_error = e
                self.log_timeout(f"Erro JavaScript: {str(e)[:150]}", "ERROR")
                self._limpar_estado_js()
                
                if tentativa < self.max_retries:
                    time.sleep(1 + tentativa * 0.5)
                    continue
                    
            except TimeoutException as e:
                self.last_error = e
                self.log_timeout(f"Timeout JavaScript ({timeout}s)", "ERROR")
                self._forcar_parada_js()
                
                if tentativa < self.max_retries:
                    time.sleep(2 + tentativa)
                    continue
                    
            except WebDriverException as e:
                self.last_error = e
                self.log_timeout(f"Erro WebDriver: {str(e)[:150]}", "ERROR")
                
                if tentativa < self.max_retries:
                    time.sleep(1.5)
                    continue
                    
            except Exception as e:
                self.last_error = e
                self.log_timeout(f"Erro inesperado: {str(e)[:150]}", "ERROR")
                break
        
        try:
            self.driver.set_script_timeout(original_timeout)
        except:
            pass
            
        if fallback_result is not None:
            self.log_timeout(f"Usando valor fallback: {fallback_result}", "WARN")
        return fallback_result
    
    def _limpar_estado_js(self):
        """Limpa estado JavaScript do browser"""
        try:
            cleanup_script = """
                if (window.__cleanupTimers) {
                    window.__cleanupTimers.forEach(clearTimeout);
                    window.__cleanupTimers.forEach(clearInterval);
                }
                if (typeof jQuery !== 'undefined') {
                    jQuery.active = 0;
                }
                window.__pendingRequests = 0;
                return true;
            """
            self.driver.execute_script(cleanup_script)
            time.sleep(0.5)
        except Exception:
            pass
    
    def _forcar_parada_js(self):
        """Força parada de JavaScript travado"""
        try:
            self.driver.execute_script("window.stop();")
            time.sleep(0.3)
        except Exception:
            pass



# ==== JS FORCE ENGINE COM PROTEÇÃO ANTI-TIMEOUT ====
class JSForceEngine:
    """Motor de execução JavaScript forçado com proteção contra timeouts"""
    
    def __init__(self, driver, wait, doc, timeout_padrao=10, max_retries=3):
        self.driver = driver
        self.wait = wait
        self.doc = doc
        self.timeout_handler = JSTimeoutHandler(driver, doc, timeout_padrao, max_retries)
    
    def execute_js(self, script, *args, timeout=None, fallback_result=None):
        """Executa JavaScript com proteção contra timeout"""
        return self.timeout_handler.execute_js_safe(
            script, *args, timeout=timeout, fallback_result=fallback_result
        )
    
    def wait_ajax_complete(self, timeout=15):
        """Aguarda AJAX completar com proteção contra timeout"""
        script = """
            var jQueryOk = (typeof jQuery==='undefined') || (jQuery.active===0);
            var fetchOk = !window.__pendingRequests || window.__pendingRequests===0;
            var overlays = document.querySelectorAll(
                '.blockScreen, .blockUI, .loading, .overlay, [class*="loading"], [class*="spinner"]'
            );
            var overlayOk = true;
            for (var i=0; i<overlays.length; i++){
                var s=window.getComputedStyle(overlays[i]);
                if(s.display!=='none' && s.visibility!=='hidden' && parseFloat(s.opacity||1)>0.01){
                    overlayOk=false;
                    break;
                }
            }
            return jQueryOk && fetchOk && overlayOk;
        """
        
        end = time.time() + timeout
        while time.time() < end:
            try:
                done = self.execute_js(script, timeout=5, fallback_result=True)
                if done:
                    return True
            except:
                pass
            time.sleep(0.2)
        return True
    
    def scroll_into_view(self, target, padding=100):
        """
        Faz scroll até um WebElement (ou seletor XPath/CSS).
        padding: desloca um pouco pra cima pra não ficar colado no topo.
        """
        try:
            el = target
            # Se vier como string, tenta resolver
            if isinstance(target, str):
                try:
                    el = self.driver.find_element("xpath", target)
                except Exception:
                    el = self.driver.find_element("css selector", target)

            # Estratégia principal via JS
            self.driver.execute_script("""
                const el = arguments[0], pad = arguments[1] || 0;
                if (!el) return;
                el.scrollIntoView({block:'center', inline:'center'});
                try { window.scrollBy(0, -pad); } catch(e) {}
            """, el, padding)
            time.sleep(0.2)
            return True
        except Exception:
            # Fallback ActionChains
            try:
                ActionChains(self.driver).move_to_element(el).perform()
                time.sleep(0.2)
                return True
            except Exception:
                return False

    def click_element(self, el, wait_after=0.4):
        """
        Clica com múltiplas estratégias em um WebElement já localizado.
        """
        try:
            # 1) Clique padrão
            el.click()
            time.sleep(wait_after)
            self.wait_ajax_complete(min(10, int(wait_after*10)) or 3)
            return True
        except Exception:
            pass

        try:
            # 2) Clique via JS
            self.driver.execute_script("arguments[0].click();", el)
            time.sleep(wait_after)
            self.wait_ajax_complete(min(10, int(wait_after*10)) or 3)
            return True
        except Exception:
            pass

        try:
            # 3) Sequência de eventos de mouse
            self.execute_js("""
                const e = arguments[0];
                const rect = e.getBoundingClientRect();
                const x = rect.left + 5, y = rect.top + 5;
                ['mouseover','mouseenter','mousemove','mousedown','mouseup','click'].forEach(t=>{
                    e.dispatchEvent(new MouseEvent(t,{bubbles:true,cancelable:true,view:window,clientX:x,clientY:y}));
                });
                if (typeof e.click==='function') e.click();
                return true;
            """, el, timeout=5, fallback_result=True)
            time.sleep(wait_after)
            self.wait_ajax_complete(min(10, int(wait_after*10)) or 3)
            return True
        except Exception:
            pass

        # 4) Último recurso: ActionChains
        try:
            ActionChains(self.driver).move_to_element(el).pause(0.05).click().perform()
            time.sleep(wait_after)
            self.wait_ajax_complete(min(10, int(wait_after*10)) or 3)
            return True
        except Exception:
            pass

        raise Exception("Não foi possível clicar no elemento com as estratégias disponíveis.")
    
    def force_click(self, selector, by_xpath=False, max_attempts=5):
        """Clique forçado com proteção contra timeout"""
        log(self.doc, f"🎯 Clique forçado em: {selector}")
        
        for attempt in range(max_attempts):
            try:
                strategies = [
                    self._click_strategy_2,
                    self._click_strategy_1,
                    self._click_strategy_3,
                    self._click_strategy_4,
                    self._click_strategy_5,
                ]
                
                for i, strategy in enumerate(strategies, 1):
                    try:
                        if attempt > 0 or i > 1:
                            log(self.doc, f"   Tentativa {attempt + 1}.{i}...")
                        
                        result = self.execute_js(
                            self._get_strategy_script(strategy, selector, by_xpath),
                            selector,
                            by_xpath,
                            timeout=5,
                            fallback_result=False
                        )
                        
                        if result:
                            log(self.doc, f"✅ Clique bem-sucedido (estratégia {i})")
                            time.sleep(0.5)
                            self.wait_ajax_complete(10)
                            return True
                            
                    except Exception as e:
                        if i == 1 and attempt == 0:
                            pass  # Silencia primeiro erro
                        else:
                            log(self.doc, f"   Estratégia {i} falhou: {str(e)[:80]}")
                        continue
                
                if attempt < max_attempts - 1:
                    time.sleep(1 + attempt * 0.5)
                    
            except Exception as e:
                log(self.doc, f"⚠️ Tentativa {attempt + 1} falhou: {e}")
                if attempt < max_attempts - 1:
                    time.sleep(1.5)
        
        raise Exception(f"Falha ao clicar após {max_attempts} tentativas: {selector}")
    
    def _get_strategy_script(self, strategy_func, selector, by_xpath):
        """Retorna o script JavaScript para cada estratégia"""
        base_locator = """
            var selector = arguments[0];
            var byXPath = arguments[1];
            var element;
            
            if (byXPath) {
                var result = document.evaluate(selector, document, null, 
                    XPathResult.FIRST_ORDERED_NODE_TYPE, null);
                element = result.singleNodeValue;
            } else {
                element = document.querySelector(selector);
            }
            
            if (!element) throw new Error('Elemento não encontrado');
        """
        
        if strategy_func == self._click_strategy_1:
            return base_locator + """
                element.style.pointerEvents = 'auto';
                element.style.display = 'block';
                element.style.visibility = 'visible';
                element.style.opacity = '1';
                element.removeAttribute('disabled');
                element.scrollIntoView({behavior: 'smooth', block: 'center'});
                setTimeout(function() { element.click(); }, 300);
                return true;
            """
        elif strategy_func == self._click_strategy_2:
            return base_locator + """
                element.style.pointerEvents = 'auto';
                element.removeAttribute('disabled');
                element.scrollIntoView({block: 'center'});
                
                var events = ['mouseover', 'mouseenter', 'mousemove', 'mousedown', 'mouseup', 'click'];
                events.forEach(function(eventType) {
                    var evt = new MouseEvent(eventType, {
                        bubbles: true, cancelable: true, view: window, detail: 1,
                        clientX: element.getBoundingClientRect().left + 5,
                        clientY: element.getBoundingClientRect().top + 5
                    });
                    element.dispatchEvent(evt);
                });
                
                if (typeof element.click === 'function') element.click();
                return true;
            """
        elif strategy_func == self._click_strategy_3:
            return base_locator + """
                element.style.display = 'block';
                element.style.visibility = 'visible';
                element.style.opacity = '1';
                element.style.pointerEvents = 'auto';
                element.focus();
                element.click();
                element.dispatchEvent(new Event('click', {bubbles: true, cancelable: true}));
                return true;
            """
        elif strategy_func == self._click_strategy_4:
            return base_locator + """
                element.removeAttribute('disabled');
                element.removeAttribute('readonly');
                element.style.pointerEvents = 'auto !important';
                element.style.display = 'block !important';
                element.style.visibility = 'visible !important';
                element.style.opacity = '1 !important';
                
                var overlays = document.querySelectorAll('.modal, .overlay, .blockUI, [role="dialog"]');
                overlays.forEach(function(overlay) {
                    overlay.style.display = 'none';
                    overlay.style.visibility = 'hidden';
                });
                
                element.focus();
                element.click();
                
                var clickEvent = new MouseEvent('click', {
                    view: window, bubbles: true, cancelable: true
                });
                element.dispatchEvent(clickEvent);
                
                if (typeof jQuery !== 'undefined') jQuery(element).trigger('click');
                return true;
            """
        else:  # strategy_5
            return base_locator + """
                var rect = element.getBoundingClientRect();
                var x = rect.left + rect.width / 2;
                var y = rect.top + rect.height / 2;
                
                var evt = document.createEvent('MouseEvents');
                evt.initMouseEvent('click', true, true, window, 1, x, y, x, y, false, false, false, false, 0, null);
                element.dispatchEvent(evt);
                
                if (element.onclick) element.onclick();
                
                var parent = element.parentElement;
                while (parent && parent !== document.body) {
                    if (parent.onclick) {
                        parent.onclick();
                        break;
                    }
                    parent = parent.parentElement;
                }
                return true;
            """
    
    def _click_strategy_1(self, selector, by_xpath):
        pass  # Implementado via _get_strategy_script
    
    def _click_strategy_2(self, selector, by_xpath):
        pass
    
    def _click_strategy_3(self, selector, by_xpath):
        pass
    
    def _click_strategy_4(self, selector, by_xpath):
        pass
    
    def _click_strategy_5(self, selector, by_xpath):
        pass
    
    def force_fill(self, selector, value, by_xpath=False, max_attempts=5):
        """Preenchimento forçado com proteção contra timeout"""
        log(self.doc, f"✏️ Preenchimento forçado: {selector} = '{value}'")
        
        for attempt in range(max_attempts):
            try:
                strategies = [
                    self._fill_strategy_1,
                    self._fill_strategy_2,
                    self._fill_strategy_3,
                    self._fill_strategy_4,
                ]
                
                for i, strategy in enumerate(strategies, 1):
                    try:
                        if attempt > 0 or i > 1:
                            log(self.doc, f"   Tentativa {attempt + 1}.{i}...")
                        
                        result = self.execute_js(
                            self._get_fill_script(strategy, selector, value, by_xpath),
                            selector,
                            value,
                            by_xpath,
                            timeout=5,
                            fallback_result=None
                        )
                        
                        time.sleep(0.3)
                        if self._validate_fill(selector, value, by_xpath):
                            log(self.doc, f"✅ Campo preenchido (estratégia {i})")
                            return True
                    except Exception as e:
                        if i == 1 and attempt == 0:
                            pass
                        else:
                            log(self.doc, f"   Estratégia {i} falhou: {str(e)[:80]}")
                        continue
                
                if attempt < max_attempts - 1:
                    time.sleep(1 + attempt * 0.5)
                    
            except Exception as e:
                log(self.doc, f"⚠️ Tentativa {attempt + 1} falhou: {e}")
        
        raise Exception(f"Falha ao preencher após {max_attempts} tentativas: {selector}")

    def _get_fill_script(self, strategy_func, selector, value, by_xpath):
        """Retorna o script de preenchimento"""
        base_locator = """
            var selector = arguments[0];
            var value = arguments[1];
            var byXPath = arguments[2];
            var element;
            
            if (byXPath) {
                var result = document.evaluate(selector, document, null, 
                    XPathResult.FIRST_ORDERED_NODE_TYPE, null);
                element = result.singleNodeValue;
            } else {
                element = document.querySelector(selector);
            }
            
            if (!element) throw new Error('Campo não encontrado');
        """
        
        if strategy_func == self._fill_strategy_1:
            return base_locator + """
                element.removeAttribute('disabled');
                element.removeAttribute('readonly');
                element.style.display = 'block';
                element.style.visibility = 'visible';
                element.scrollIntoView({block: 'center'});
                element.focus();
                element.dispatchEvent(new Event('focus', {bubbles: true}));
                element.value = '';
                element.value = value;
                ['input', 'change', 'blur', 'keyup'].forEach(function(evt) {
                    element.dispatchEvent(new Event(evt, {bubbles: true}));
                });
                return element.value;
            """
        elif strategy_func == self._fill_strategy_2:
            return base_locator + """
                var nativeInputValueSetter = Object.getOwnPropertyDescriptor(
                    window.HTMLInputElement.prototype, 'value'
                ).set;
                if (nativeInputValueSetter) {
                    nativeInputValueSetter.call(element, value);
                } else {
                    element.value = value;
                }
                element.dispatchEvent(new Event('input', {bubbles: true}));
                element.dispatchEvent(new Event('change', {bubbles: true}));
                element.dispatchEvent(new KeyboardEvent('keydown', {bubbles: true}));
                element.dispatchEvent(new KeyboardEvent('keyup', {bubbles: true}));
                element.dispatchEvent(new Event('blur', {bubbles: true}));
                return element.value;
            """
        elif strategy_func == self._fill_strategy_3:
            return base_locator + """
                element.value = value;
                if (typeof jQuery !== 'undefined') {
                    jQuery(element).val(value).trigger('input').trigger('change').trigger('blur');
                }
                ['focus', 'input', 'change', 'blur'].forEach(function(evt) {
                    element.dispatchEvent(new Event(evt, {bubbles: true}));
                });
                return element.value;
            """
        else:  # strategy_4
            return base_locator + """
                element.removeAttribute('disabled');
                element.removeAttribute('readonly');
                element.removeAttribute('maxlength');
                element.value = '';
                element.setAttribute('value', value);
                element.value = value;
                element.style.color = element.style.color;
                
                var events = ['focus', 'click', 'input', 'change', 'keydown', 'keypress', 
                              'keyup', 'blur', 'paste', 'textInput'];
                events.forEach(function(evt) {
                    try {
                        element.dispatchEvent(new Event(evt, {bubbles: true, cancelable: true}));
                    } catch(e) {}
                });
                
                if (element.oninput) element.oninput();
                if (element.onchange) element.onchange();
                return element.value;
            """
    
    def _fill_strategy_1(self, selector, value, by_xpath):
        pass
    
    def _fill_strategy_2(self, selector, value, by_xpath):
        pass
    
    def _fill_strategy_3(self, selector, value, by_xpath):
        pass
    
    def _fill_strategy_4(self, selector, value, by_xpath):
        pass
    
    def _validate_fill(self, selector, expected_value, by_xpath):
        """Valida preenchimento"""
        script = """
            var selector = arguments[0];
            var expected = arguments[1];
            var byXPath = arguments[2];
            var element;
            
            if (byXPath) {
                var result = document.evaluate(selector, document, null, 
                    XPathResult.FIRST_ORDERED_NODE_TYPE, null);
                element = result.singleNodeValue;
            } else {
                element = document.querySelector(selector);
            }
            
            if (!element) return false;
            var actual = element.value || '';
            return actual.trim() === expected.trim() || actual.includes(expected);
        """
        try:
            return self.execute_js(script, selector, expected_value, by_xpath, timeout=3, fallback_result=False)
        except:
            return False
        
import time

def clicar_finalizar_e_verificar_alerta(js_engine, doc, timeout=5, pausa=0.5):
    """
    Clica no botão 'Finalizar', aguarda 0,5s e procura mensagens de alerta.
    Usa safe_action e js_engine.force_click().
    """
    safe_action(doc, "Clicando em 'Finalizar'", lambda:
        js_engine.force_click(
            "//a[@class='btModel btGray btyes' and normalize-space()='Finalizar']",
            by_xpath=True
        )
    )

    time.sleep(pausa)
    log(doc, "🔍 Verificando mensagens de alerta após o clique em 'Finalizar'...")
    return encontrar_mensagem_alerta()


import time
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.common.exceptions import (
    TimeoutException,
    StaleElementReferenceException,
    ElementClickInterceptedException,
    NoSuchElementException,
    ElementNotInteractableException
)
from selenium.webdriver import ActionChains

class LOVHandler:
    """
    Handler ultra-robusto para manipulação de LOV (List of Values) com:
    - 10 estratégias diferentes de clique
    - Detecção automática de iframes
    - Retry inteligente com backoff exponencial
    - Reforço automático de cliques
    - Logs detalhados de cada tentativa
    """
    
    def __init__(self, js_engine, doc, max_retries=5, timeout=10):
        self.js = js_engine
        self.doc = doc
        self.driver = js_engine.driver
        self.max_retries = max_retries
        self.timeout = timeout

        # índice para seleção sequencial das linhas do LOV
        self.selecao_sequencial_index = 0

    def resetar_selecao_sequencial(self, indice_inicial: int = 0):
        """
        Reseta o índice usado na seleção sequencial de linhas do LOV.
        Usa 0 como padrão (primeira linha).
        """
        try:
            indice = int(indice_inicial)
        except Exception:
            indice = 0

        if indice < 0:
            indice = 0

        self.selecao_sequencial_index = indice
        self._log(
            f"🔁 Índice de seleção sequencial resetado para {self.selecao_sequencial_index}",
            "INFO"
        )

    def _log(self, msg, level="INFO"):
        """Log padronizado com níveis"""
        prefixes = {
            "INFO": "ℹ️",
            "SUCCESS": "✅",
            "WARNING": "⚠️",
            "ERROR": "❌",
            "DEBUG": "🔍"
        }
        prefix = prefixes.get(level, "📝")
        log(self.doc, f"{prefix} {msg}")

    def _select_result_sequencial(self):
        """Seleciona a próxima linha disponível da tabela do LOV, sem repetir."""
        try:
            # Garante que o atributo exista
            if not hasattr(self, "selecao_sequencial_index"):
                self.selecao_sequencial_index = 0

            linhas = self.driver.find_elements(By.XPATH, "//table//tr[td]")

            if not linhas:
                self._log("Nenhuma linha encontrada no LOV.")
                return None

            if self.selecao_sequencial_index >= len(linhas):
                self._log("⚠️ Índice passou do limite, usando última linha disponível.", "WARNING")
                self.selecao_sequencial_index = len(linhas) - 1

            linha_alvo = linhas[self.selecao_sequencial_index]

            self._log(f"Selecionando linha Nº {self.selecao_sequencial_index + 1}", "INFO")

            if self._advanced_click(linha_alvo):
                self._log("Linha selecionada com sucesso!", "SUCCESS")
                self.selecao_sequencial_index += 1
                return get_xpath(linha_alvo)

            return None

        except Exception as e:
            self._log(f"Erro ao selecionar linha sequencial: {e}", "ERROR")
            return None



    def _wait_element(self, locator_type, locator_value, timeout=None, condition="clickable"):
        """Aguarda elemento com diferentes condições"""
        timeout = timeout or self.timeout
        conditions = {
            "present": EC.presence_of_element_located,
            "visible": EC.visibility_of_element_located,
            "clickable": EC.element_to_be_clickable
        }
        
        try:
            cond = conditions.get(condition, conditions["clickable"])
            return WebDriverWait(self.driver, timeout).until(
                cond((locator_type, locator_value))
            )
        except TimeoutException:
            return None
    
    def _is_element_visible(self, element):
        """Verifica se elemento está realmente visível"""
        try:
            if not element:
                return False
            if not element.is_displayed():
                return False
            
            # Verifica via JavaScript também
            is_visible = self.driver.execute_script("""
                const el = arguments[0];
                if (!el) return false;
                const style = window.getComputedStyle(el);
                return (
                    style.display !== 'none' &&
                    style.visibility !== 'hidden' &&
                    parseFloat(style.opacity || 1) > 0.01 &&
                    el.offsetParent !== null
                );
            """, element)
            
            return is_visible
        except:
            return False
    
    def _force_element_visible(self, element):
        """Força elemento a ficar visível e interativo"""
        try:
            self.driver.execute_script("""
                const el = arguments[0];
                el.style.display = 'block';
                el.style.visibility = 'visible';
                el.style.opacity = '1';
                el.style.pointerEvents = 'auto';
                el.removeAttribute('disabled');
                el.removeAttribute('readonly');
                
                // Remove overlays que podem bloquear
                const overlays = document.querySelectorAll(
                    '.modal-backdrop, .overlay, .blockUI, [class*="loading"]'
                );
                overlays.forEach(o => {
                    o.style.display = 'none';
                    o.style.visibility = 'hidden';
                });
            """, element)
            return True
        except:
            return False
    
    def _detect_and_enter_iframe(self, iframe_xpath=None):
        """Detecta e entra no iframe automaticamente"""
        try:
            # Primeiro volta para o contexto principal
            self.driver.switch_to.default_content()
            
            # Se foi fornecido um xpath específico
            if iframe_xpath:
                try:
                    frame = self._wait_element("xpath", iframe_xpath, timeout=3)
                    if frame:
                        self.driver.switch_to.frame(frame)
                        self._log(f"Entrou no iframe: {iframe_xpath}", "SUCCESS")
                        return True
                except:
                    pass
            
            # Detecção automática de iframes
            iframe_selectors = [
                "//iframe[contains(@class,'LOV') or contains(@id,'LOV') or contains(@id,'lov')]",
                "//iframe[contains(@class,'modal') or contains(@class,'popup')]",
                "//iframe[contains(@src,'lov') or contains(@src,'LOV')]",
                "(//iframe)[last()]"  # Último iframe (geralmente o modal)
            ]
            
            for selector in iframe_selectors:
                try:
                    frames = self.driver.find_elements(By.XPATH, selector)
                    for frame in frames:
                        if self._is_element_visible(frame):
                            self.driver.switch_to.frame(frame)
                            self._log(f"Iframe detectado automaticamente: {selector}", "SUCCESS")
                            return True
                except:
                    continue
            
            return False
        except Exception as e:
            self._log(f"Erro ao detectar iframe: {e}", "WARNING")
            return False
    
    def _click_strategy_1_standard(self, element):
        """Estratégia 1: Clique padrão do Selenium"""
        element.click()
        return True
    
    def _click_strategy_2_javascript_simple(self, element):
        """Estratégia 2: JavaScript simples"""
        self.driver.execute_script("arguments[0].click();", element)
        return True
    
    def _click_strategy_3_javascript_advanced(self, element):
        """Estratégia 3: JavaScript com eventos completos"""
        self.driver.execute_script("""
            const el = arguments[0];
            const rect = el.getBoundingClientRect();
            const x = rect.left + rect.width / 2;
            const y = rect.top + rect.height / 2;
            
            const events = ['mouseover', 'mouseenter', 'mousemove', 'mousedown', 'mouseup', 'click'];
            events.forEach(eventType => {
                const evt = new MouseEvent(eventType, {
                    bubbles: true,
                    cancelable: true,
                    view: window,
                    detail: 1,
                    clientX: x,
                    clientY: y
                });
                el.dispatchEvent(evt);
            });
            
            if (typeof el.click === 'function') el.click();
        """, element)
        return True
    
    def _click_strategy_4_action_chains(self, element):
        """Estratégia 4: ActionChains com pause"""
        ActionChains(self.driver)\
            .move_to_element(element)\
            .pause(0.1)\
            .click()\
            .perform()
        return True
    
    def _click_strategy_5_action_chains_offset(self, element):
        """Estratégia 5: ActionChains com offset"""
        ActionChains(self.driver)\
            .move_to_element_with_offset(element, 5, 5)\
            .pause(0.05)\
            .click()\
            .perform()
        return True
    
    def _click_strategy_6_force_visible_then_click(self, element):
        """Estratégia 6: Força visibilidade e clica"""
        self._force_element_visible(element)
        time.sleep(0.2)
        element.click()
        return True
    
    def _click_strategy_7_scroll_and_click(self, element):
        """Estratégia 7: Scroll suave até elemento"""
        self.driver.execute_script("""
            arguments[0].scrollIntoView({
                behavior: 'smooth',
                block: 'center',
                inline: 'center'
            });
        """, element)
        time.sleep(0.3)
        element.click()
        return True
    
    def _click_strategy_8_remove_overlays(self, element):
        """Estratégia 8: Remove todos os overlays e clica"""
        self.driver.execute_script("""
            // Remove todos os overlays possíveis
            const overlays = document.querySelectorAll(`
                .modal-backdrop, .overlay, .blockUI, .blockScreen,
                [class*="loading"], [class*="spinner"], [class*="overlay"],
                [style*="z-index: 9999"], [style*="position: fixed"]
            `);
            overlays.forEach(o => {
                o.style.display = 'none';
                o.style.visibility = 'hidden';
                o.remove();
            });
            
            const el = arguments[0];
            el.style.zIndex = '999999';
            el.click();
        """, element)
        return True
    
    def _click_strategy_9_jquery_trigger(self, element):
        """Estratégia 9: jQuery trigger (se disponível)"""
        self.driver.execute_script("""
            const el = arguments[0];
            if (typeof jQuery !== 'undefined') {
                jQuery(el).trigger('click');
            } else {
                el.click();
            }
        """, element)
        return True
    
    def _click_strategy_10_nuclear_option(self, element):
        """Estratégia 10: Opção nuclear - força tudo"""
        self.driver.execute_script("""
            const el = arguments[0];
            
            // Remove TODOS os atributos que podem bloquear
            el.removeAttribute('disabled');
            el.removeAttribute('readonly');
            el.style.pointerEvents = 'auto !important';
            el.style.display = 'block !important';
            el.style.visibility = 'visible !important';
            el.style.opacity = '1 !important';
            el.style.zIndex = '999999 !important';
            
            // Remove TODOS os overlays da página
            document.querySelectorAll('*').forEach(elem => {
                const style = window.getComputedStyle(elem);
                if (
                    style.position === 'fixed' &&
                    parseInt(style.zIndex) > 1000 &&
                    elem !== el &&
                    !elem.contains(el)
                ) {
                    elem.style.display = 'none';
                }
            });
            
            // Foca no elemento
            el.focus();
            
            // Dispara TODOS os eventos possíveis
            const allEvents = [
                'focus', 'focusin', 'mouseover', 'mouseenter', 'mousemove',
                'mousedown', 'mouseup', 'click', 'dblclick'
            ];
            
            allEvents.forEach(eventType => {
                try {
                    const evt = new MouseEvent(eventType, {
                        bubbles: true,
                        cancelable: true,
                        view: window
                    });
                    el.dispatchEvent(evt);
                } catch(e) {}
            });
            
            // Clique direto
            if (typeof el.click === 'function') el.click();
            
            // jQuery se disponível
            if (typeof jQuery !== 'undefined') {
                jQuery(el).trigger('click');
            }
            
            // Tenta onclick manual
            if (el.onclick) el.onclick();
            
            return true;
        """, element)
        return True
    
    def _advanced_click(self, element, max_attempts=10):
        """
        Sistema avançado de clique com 10 estratégias diferentes
        Tenta cada estratégia até uma funcionar
        """
        strategies = [
            ("Clique Padrão", self._click_strategy_1_standard),
            ("JavaScript Simples", self._click_strategy_2_javascript_simple),
            ("JavaScript Avançado", self._click_strategy_3_javascript_advanced),
            ("ActionChains", self._click_strategy_4_action_chains),
            ("ActionChains Offset", self._click_strategy_5_action_chains_offset),
            ("Força Visível", self._click_strategy_6_force_visible_then_click),
            ("Scroll e Clique", self._click_strategy_7_scroll_and_click),
            ("Remove Overlays", self._click_strategy_8_remove_overlays),
            ("jQuery Trigger", self._click_strategy_9_jquery_trigger),
            ("Opção Nuclear", self._click_strategy_10_nuclear_option)
        ]
        
        for attempt in range(1, max_attempts + 1):
            for strategy_name, strategy_func in strategies:
                try:
                    self._log(f"Tentativa {attempt}/{max_attempts}: {strategy_name}", "DEBUG")
                    
                    # Tenta executar a estratégia
                    strategy_func(element)
                    time.sleep(0.3)
                    
                    self._log(f"✓ {strategy_name} funcionou!", "SUCCESS")
                    return True
                    
                except StaleElementReferenceException:
                    self._log(f"Elemento ficou stale, tentando recarregar...", "WARNING")
                    return False  # Precisa recarregar elemento
                    
                except (ElementClickInterceptedException, 
                        ElementNotInteractableException) as e:
                    self._log(f"✗ {strategy_name}: {str(e)[:50]}", "DEBUG")
                    continue
                    
                except Exception as e:
                    self._log(f"✗ {strategy_name}: {str(e)[:50]}", "DEBUG")
                    continue
            
            if attempt < max_attempts:
                time.sleep(0.5 * attempt)  # Backoff exponencial
        
        return False
        
    def _select_filter_option(self, option_text):
        """
        Seleciona um valor no <select class='tipoFiltro'> ANTES da pesquisa.
        """
        try:
            self._log(f"Selecionando filtro: {option_text}", "INFO")

            # Localiza o SELECT
            select_el = self._wait_element(
                By.CSS_SELECTOR, "select.tipoFiltro",
                timeout=5, condition="visible"
            )

            if not select_el:
                raise Exception("Elemento <select class='tipoFiltro'> não encontrado")

            self._force_element_visible(select_el)
            time.sleep(0.3)

            # Seleciona opção pelo texto visível
            Select(select_el).select_by_visible_text(option_text)

            self._log(f"Filtro selecionado: {option_text}", "SUCCESS")
            time.sleep(0.4)
            return True

        except Exception as e:
            self._log(f"Erro ao selecionar filtro '{option_text}': {e}", "ERROR")
            return False

    
    def _fill_search_fields(self, search_text, max_fields=10):
        """Preenche TODOS os campos de pesquisa encontrados"""
        if not search_text:
            return 0
        
        search_field_xpaths = [
            "//input[@id='txtPesquisa']",
            "//input[contains(@class,'pesquisa')]",
            "//input[contains(@class,'nomePesquisa')]",
            "//input[contains(translate(@name,'PESQUISA','pesquisa'),'pesquisa')]",
            "//input[@type='text' and contains(@style,'width:210px')]",
            "//input[@type='text' and not(@disabled)]",
        ]
        
        filled_count = 0
        for xpath in search_field_xpaths:
            try:
                fields = self.driver.find_elements(By.XPATH, xpath)
                for field in fields[:max_fields]:
                    try:
                        if not self._is_element_visible(field):
                            continue
                        
                        # Limpa e preenche
                        self._force_element_visible(field)
                        field.clear()
                        field.click()
                        field.send_keys(search_text)
                        filled_count += 1
                        
                    except:
                        continue
            except:
                continue
        
        if filled_count > 0:
            self._log(f"Preenchidos {filled_count} campos de pesquisa", "SUCCESS")
        else:
            self._log("Nenhum campo de pesquisa encontrado", "WARNING")
        
        return filled_count
    
    def _click_search_button(self):
        """Clica no botão Pesquisar com múltiplas estratégias"""
        search_button_xpaths = [
            "//a[contains(@class,'btPesquisar') and contains(normalize-space(.),'Pesquisar')]",
            "//button[contains(normalize-space(.),'Pesquisar')]",
            "//input[@type='button' and contains(@value,'Pesquisar')]",
            "//a[contains(@class,'lpFind')]",
            "//a[contains(@onclick,'pesquisar')]"
        ]
        
        # Fallback: ENTER no campo ativo
        try:
            self._log("Tentando ENTER no campo de pesquisa", "DEBUG")
            ActionChains(self.driver).send_keys(Keys.ENTER).perform()
            return True
        except Exception:
            pass
        
        for xpath in search_button_xpaths:
            try:
                btn = self._wait_element("xpath", xpath, timeout=2)
                if btn and self._is_element_visible(btn):
                    if self._advanced_click(btn):
                        self._log("Botão 'Pesquisar' clicado", "SUCCESS")
                        return True
            except:
                pass
        
        self._log("Não foi possível clicar em 'Pesquisar'", "WARNING")
        return False
    
    def _select_result(self, result_text=""):
        """Seleciona o resultado da pesquisa"""
        result_xpaths = []
        
        if result_text:
            result_xpaths = [
                f"//td[contains(normalize-space(.),'{result_text}')]",
                f"//span[contains(normalize-space(.),'{result_text}')]",
                f"//div[contains(normalize-space(.),'{result_text}')]",
                f"//tr[contains(normalize-space(.),'{result_text}')]//td[1]",
                f"//li[contains(normalize-space(.),'{result_text}')]"
            ]
        else:
            result_xpaths = [
                "(//table//tr[1]/td[1])[1]",
                "(//tr[contains(@class,'ui-widget-content')][1])[1]",
                "(//tr[contains(@class,'rich-table-row')][1])[1]",
                "(//li[@class='ui-autocomplete-item'])[1]"
            ]
        
        for xpath in result_xpaths:
            try:
                result = self._wait_element("xpath", xpath, timeout=3)
                if result and self._is_element_visible(result):
                    if self._advanced_click(result):
                        self._log(f"Resultado selecionado: {result_text or 'primeiro'}", "SUCCESS")
                        return xpath  # Retorna xpath usado
            except:
                continue
        
        self._log(f"Não foi possível selecionar: {result_text or 'primeiro resultado'}", "ERROR")
        return None
    
    def _reinforce_click_on_result(self, result_xpath, max_reclick=5):
        """
        Reforça o clique no resultado SEMPRE, independente do modal
        Clica repetidamente até max_reclick vezes
        """
        if not result_xpath:
            return
        
        self._log(f"🔁 Reforçando clique no resultado ({max_reclick}x)", "INFO")
        
        for attempt in range(1, max_reclick + 1):
            try:
                time.sleep(0.35)
                
                # Tenta localizar o resultado novamente
                result = self.driver.find_element(By.XPATH, result_xpath)
                
                # Clica usando sistema avançado
                self._log(f"Reforço {attempt}/{max_reclick}", "DEBUG")
                self._advanced_click(result)
                
            except StaleElementReferenceException:
                self._log(f"Elemento stale no reforço {attempt}", "WARNING")
                try:
                    # Tenta recarregar o elemento
                    result = self.driver.find_element(By.XPATH, result_xpath)
                    self._advanced_click(result)
                except:
                    continue
                    
            except Exception as e:
                self._log(f"Erro no reforço {attempt}: {str(e)[:50]}", "WARNING")
                continue
        
        self._log("✅ Reforço de cliques concluído", "SUCCESS")

    def open_and_select(
        self,
        btn_index=None,
        btn_xpath=None,
        btn_css=None,
        search_text="",
        result_text="",      # agora só pra log
        filter_option=None,
        iframe_xpath=None,
        auto_detect_iframe=True,
        reinforce_clicks=5,
        wait_after=0.5
    ):
        """
        Método principal: abre LOV, aplica filtro/pesquisa e seleciona
        SEMPRE de forma SEQUENCIAL (1ª vez = 1ª linha, 2ª = 2ª, etc.).
        """

        self._log(
            f"🔍 LOV sequencial | filtro='{filter_option}' | pesquisa='{search_text}' | índice atual={self.selecao_sequencial_index}",
            "INFO"
        )

        for retry in range(1, self.max_retries + 1):
            try:
                # ===== PASSO 1: volta p/ contexto principal =====
                try:
                    self.driver.switch_to.default_content()
                except Exception:
                    pass

                # ===== PASSO 2: localizar botão do LOV =====
                btn_selector = None
                by_type = None

                if btn_index is not None:
                    # índice é 0-based aqui, então soma 1 no XPath
                    btn_selector = f"(//a[@class='sprites sp-openLov'])[{btn_index + 1}]"
                    by_type = "xpath"
                elif btn_xpath:
                    btn_selector = btn_xpath
                    by_type = "xpath"
                elif btn_css:
                    btn_selector = btn_css
                    by_type = "css"
                else:
                    raise ValueError("Forneça btn_index, btn_xpath ou btn_css para abrir o LOV.")

                self._log(f"🔎 Localizando botão LOV: {btn_selector}", "DEBUG")
                lov_button = self._wait_element(by_type, btn_selector, timeout=5)

                if not lov_button:
                    raise Exception(f"Botão LOV não encontrado: {btn_selector}")

                self._log("ℹ️ Abrindo LOV...", "INFO")
                if not self._advanced_click(lov_button):
                    raise Exception("Falha ao clicar no botão LOV")

                time.sleep(0.8)

                # ===== PASSO 3: entra no iframe, se houver =====
                if iframe_xpath or auto_detect_iframe:
                    self._detect_and_enter_iframe(iframe_xpath)
                    time.sleep(0.3)

                # ===== PASSO 4: aplica filtro e preenche pesquisa =====
                if filter_option:
                    self._select_filter_option(filter_option)
                    time.sleep(0.3)

                if search_text:
                    self._fill_search_fields(search_text)
                    time.sleep(0.3)

                # ===== PASSO 5: clicar em Pesquisar / ENTER =====
                self._click_search_button()
                time.sleep(0.8)

                # ===== PASSO 6: seleção SEQUENCIAL SEMPRE =====
                result_xpath = self._select_result_sequencial()
                if not result_xpath:
                    raise Exception("Nenhuma linha foi selecionada pelo modo sequencial.")

                # ===== PASSO 7: reforça o clique na linha =====
                if reinforce_clicks > 0:
                    self._reinforce_click_on_result(result_xpath, reinforce_clicks)

                # ===== PASSO 8: volta do iframe p/ contexto principal =====
                try:
                    self.driver.switch_to.default_content()
                except Exception:
                    pass

                time.sleep(wait_after)
                self._log(
                    f"✅ LOV concluído com sucesso (linha índice {self.selecao_sequencial_index})",
                    "SUCCESS"
                )
                return True

            except Exception as e:
                self._log(f"Tentativa {retry} falhou: {str(e)[:120]}", "ERROR")

                try:
                    self.driver.switch_to.default_content()
                except Exception:
                    pass

                if retry < self.max_retries:
                    time.sleep(2 * retry)  # backoff
                    continue

        self._log(f"❌ LOV falhou após {self.max_retries} tentativas.", "ERROR")
        return False



def selecionar_opcao_xpath(xpath, texto):
    def acao():
        select_element = wait.until(EC.presence_of_element_located((By.XPATH, xpath)))
        Select(select_element).select_by_visible_text(texto)
    return acao

def abrir_modal_e_selecionar_robusto(btn_selector, pesquisa_selector, termo_pesquisa, btn_pesquisar_selector, resultado_xpath):
    """Versão robusta da função de modal"""
    global driver, wait, doc
    
    def acao():
        if driver is None or wait is None:
            raise Exception("Driver ou wait não inicializados")
            
        # Abre o modal
        safe_scroll_and_interact(btn_selector, "click")
        time.sleep(1)

        # Aguarda e preenche campo pesquisa
        campo_pesquisa = wait.until(
            EC.element_to_be_clickable((By.CSS_SELECTOR, pesquisa_selector))
        )
        campo_pesquisa.clear()
        campo_pesquisa.send_keys(termo_pesquisa)
        time.sleep(0.5)

        # Clica pesquisar
        pesquisar = wait.until(
            EC.element_to_be_clickable((By.CSS_SELECTOR, btn_pesquisar_selector))
        )
        pesquisar.click()
        time.sleep(2)
        
        # Aguarda resultado e clica
        resultado = wait.until(
            EC.element_to_be_clickable((By.XPATH, resultado_xpath))
        )
        driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", resultado)
        time.sleep(0.5)
        resultado.click()
        time.sleep(1)

    return acao


def finalizar_relatorio():
    global driver, doc
    nome_arquivo = f"relatorio_locacao_equipamentos_cenario_2_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    try:
        doc.save(nome_arquivo)
        log(doc, f"📄 Relatório salvo como: {nome_arquivo}")
        try:
            subprocess.run(["start", "winword", nome_arquivo], shell=True)
        except:
            pass
    except Exception as e:
        print(f"Erro ao salvar relatório: {e}")
    if driver:
        try:
            driver.quit()
        except:
            pass


def encontrar_mensagem_alerta():
    seletores = [
        (".alerts.salvo", "✅ Mensagem de Sucesso"),
        (".alerts.alerta", "⚠️ Mensagem de Alerta"),
        (".alerts.erro", "❌ Mensagem de Erro"),
    ]

    for seletor, tipo in seletores:
        try:
            elemento = driver.find_element(By.CSS_SELECTOR, seletor)
            if elemento.is_displayed():
                log(doc, f"📢 {tipo}: {elemento.text}")
                return elemento
        except:
            continue

    log(doc, "ℹ️ Nenhuma mensagem de alerta encontrada.")
    return None


def ajustar_zoom():
    global driver, doc
    if driver is None:
        return
    try:
        driver.execute_script("document.body.style.zoom='90%'")
        log(doc, "🔍 Zoom ajustado para 90%.")
    except Exception as e:
        log(doc, f"⚠️ Erro ao ajustar zoom: {e}")

def safe_scroll_and_interact(selector, action_type="click", value=None, timeout=10, by_xpath=False):
    """Rola até o elemento e interage com ele de forma robusta."""
    global driver, doc
    if driver is None:
        return None
    try:
        by_type = By.XPATH if by_xpath else By.CSS_SELECTOR
        element = WebDriverWait(driver, timeout).until(EC.presence_of_element_located((by_type, selector)))
        driver.execute_script("arguments[0].scrollIntoView({block: 'center', behavior: 'smooth'});", element)
        time.sleep(0.5)
        if action_type in ["click", "send_keys"]:
            element = WebDriverWait(driver, timeout).until(EC.element_to_be_clickable((by_type, selector)))
        if action_type == "click":
            element.click()
        elif action_type == "send_keys" and value:
            element.clear()
            element.send_keys(value)
        elif action_type == "select" and value:
            Select(element).select_by_visible_text(value)
        return element
    except Exception as e:
        log(doc, f"❌ Erro ao interagir com elemento {selector}: {e}")
        return None

# ==== CLICAR ROBUSTO ====
def clicar_elemento_robusto(driver, wait, seletor_css, timeout=10):
    global doc
    try:
        elem = wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, seletor_css)))
        try:
            driver.execute_script("""
                document.querySelectorAll('.modal, .overlay, .blockUI, .toast, .tooltip, [role="dialog"], [data-overlay]')
                .forEach(e => { if (getComputedStyle(e).position === 'fixed') e.style.display = 'none'; });
            """)
        except Exception:
            pass
        driver.execute_script("arguments[0].scrollIntoView({block:'center', inline:'center'});", elem)
        time.sleep(0.2)
        try:
            elem = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, seletor_css)))
            elem.click()
            return True
        except (TimeoutException, ElementClickInterceptedException, StaleElementReferenceException):
            pass
        try:
            elem = driver.find_element(By.CSS_SELECTOR, seletor_css)
            ActionChains(driver).move_to_element(elem).pause(0.05).click().perform()
            return True
        except Exception:
            pass
        try:
            elem = driver.find_element(By.CSS_SELECTOR, seletor_css)
            driver.execute_script("arguments[0].click();", elem)
            return True
        except Exception:
            pass
        try:
            elem = driver.find_element(By.CSS_SELECTOR, seletor_css)
            driver.execute_script("""
                const el = arguments[0];
                function fire(type){ el.dispatchEvent(new MouseEvent(type,{bubbles:true,cancelable:true,view:window})); }
                el.focus(); fire('mouseover'); fire('mousedown'); fire('mouseup'); fire('click');
            """, elem)
            return True
        except Exception:
            pass
        try:
            elem = driver.find_element(By.CSS_SELECTOR, seletor_css)
            ActionChains(driver).move_to_element_with_offset(elem, 1, 1).click().perform()
            return True
        except Exception:
            pass
        log(doc, f"❌ Não foi possível clicar em: {seletor_css}")
        return False
    except Exception as e:
        log(doc, f"❌ Erro ao clicar robusto: {e}")
        return False

# ==== DATEPICKER ====
def encontrar_campos_datepicker():
    global driver
    seletores = [
        "input.hasDatepicker",
        "input[id^='dp']",
        "input[class*='datepicker']",
        ".hasDatepicker",
        "[data-provide='datepicker']"
    ]
    campos = []
    for seletor in seletores:
        try:
            for elemento in driver.find_elements(By.CSS_SELECTOR, seletor):
                if elemento.is_displayed():
                    cid = elemento.get_attribute('id')
                    if not any(c.get('id') == cid for c in campos):
                        campos.append({'elemento': elemento, 'id': cid})
        except:
            continue
    log(doc, f"📊 Encontrados {len(campos)} campos datepicker")
    return campos

def _datepicker_jquery(campo_id, data_valor):
    global driver
    resultado = driver.execute_script("""
        var campoId = arguments[0], valor = arguments[1];
        if (typeof jQuery === 'undefined') return 'jQuery não disponível';
        var $campo = $('#' + campoId);
        if (!$campo.length) return 'Campo não encontrado: ' + campoId;
        try {
            if ($campo.hasClass('hasDatepicker')) { $campo.datepicker('setDate', valor); }
            else { $campo.val(valor); }
            $campo.trigger('input').trigger('change').trigger('blur');
            return $campo.val();
        } catch(e) { return 'Erro: ' + e.message; }
    """, campo_id, data_valor)
    if isinstance(resultado, str) and 'Erro' in resultado:
        raise Exception(f"jQuery falhou: {resultado}")

def _datepicker_javascript(elemento, data_valor):
    global driver
    driver.execute_script("""
        var campo = arguments[0], valor = arguments[1];
        campo.focus(); campo.value = ''; campo.value = valor;
        ['input','change','blur','keyup'].forEach(ev => campo.dispatchEvent(new Event(ev,{bubbles:true})));
    """, elemento, data_valor)

def _datepicker_actionchains(elemento, data_valor):
    global driver
    driver.execute_script("arguments[0].scrollIntoView({block:'center'});", elemento)
    ActionChains(driver).move_to_element(elemento).click().perform()
    time.sleep(0.3)
    ActionChains(driver).key_down(Keys.CONTROL).send_keys('a').key_up(Keys.CONTROL).perform()
    ActionChains(driver).send_keys(Keys.DELETE).perform()
    time.sleep(0.2)
    for ch in data_valor:
        ActionChains(driver).send_keys(ch).perform()
        time.sleep(0.03)
    ActionChains(driver).send_keys(Keys.TAB).perform()

def _datepicker_tradicional(elemento, data_valor):
    driver.execute_script("arguments[0].scrollIntoView({block:'center'});", elemento)
    elemento.click(); time.sleep(0.2); elemento.clear()
    elemento.send_keys(data_valor); elemento.send_keys(Keys.TAB)

def preencher_datepicker_persistente(indice_campo, data_valor, max_tentativas=10, timeout=30):
    """Preenche datepicker com várias estratégias e valida."""
    inicio_tempo = time.time()
    tentativa = 0

    def validar_data_preenchida(el, data_esperada):
        try:
            val = (el.get_attribute('value') or '').strip()
            if not val: return False
            if val == data_esperada or data_esperada in val: return True
            formatos = [
                '%d/%m/%Y %H:%M', '%d/%m/%Y %H:%M:%S',  # <- COM HORA
                '%d/%m/%Y', '%m/%d/%Y', '%Y-%m-%d', '%d-%m-%Y'
            ]
            for f in formatos:
                try:
                    d1 = datetime.strptime(val, f)
                    d2 = datetime.strptime(data_esperada, f)
                    if d1 == d2: return True
                except:
                    continue
            return False
        except:
            return False

    while tentativa < max_tentativas and (time.time() - inicio_tempo) < timeout:
        tentativa += 1
        try:
            log(doc, f"🔄 Tentativa {tentativa}/{max_tentativas} para campo {indice_campo}")
            campos = encontrar_campos_datepicker()
            if not campos:
                time.sleep(1); continue
            if indice_campo >= len(campos):
                time.sleep(1); continue

            info = campos[indice_campo]
            elemento, campo_id = info['elemento'], info['id']

            if validar_data_preenchida(elemento, data_valor):
                log(doc, f"✅ Campo {indice_campo} já está preenchido corretamente!")
                return True

            estrategias = [
                lambda: _datepicker_jquery(campo_id, data_valor),
                lambda: _datepicker_javascript(elemento, data_valor),
                lambda: _datepicker_actionchains(elemento, data_valor),
                lambda: _datepicker_tradicional(elemento, data_valor),
            ]
            for acao in estrategias:
                try:
                    acao(); time.sleep(0.4)
                    if validar_data_preenchida(elemento, data_valor):
                        log(doc, f"✅ Campo {indice_campo} preenchido!")
                        return True
                except Exception as e:
                    log(doc, f"   ⚠️ Estratégia falhou: {e}")

            log(doc, "❌ Tentativa falhou; tentando novamente...")
            time.sleep(1.2)

        except Exception as e:
            log(doc, f"❌ Erro na tentativa {tentativa}: {e}")
            time.sleep(1)

    raise Exception(f"Falha ao preencher datepicker {indice_campo} após {tentativa} tentativas em {int(time.time()-inicio_tempo)}s")

def preencher_campo_xpath_com_retry(driver, wait, xpath, valor, max_tentativas=3):
    global doc
    if driver is None or wait is None:
        return False
    for tentativa in range(max_tentativas):
        try:
            campo = wait.until(EC.element_to_be_clickable((By.XPATH, xpath)))
            driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", campo)
            time.sleep(0.3)
            if tentativa == 0:
                campo.click(); campo.clear(); campo.send_keys(valor); campo.send_keys(Keys.TAB)
            elif tentativa == 1:
                ActionChains(driver).move_to_element(campo).click().pause(0.1).key_down(Keys.CONTROL).send_keys('a').key_up(Keys.CONTROL).send_keys(valor).send_keys(Keys.TAB).perform()
            else:
                driver.execute_script("""
                    var el = arguments[0], v = arguments[1];
                    el.focus(); el.value = ''; el.value = v;
                    el.dispatchEvent(new Event('input', { bubbles: true }));
                    el.dispatchEvent(new Event('change', { bubbles: true }));
                    el.blur();
                """, campo, valor)
            time.sleep(0.3)
            if (campo.get_attribute('value') or '').strip():
                return True
        except Exception as e:
            log(doc, f"⚠️ Tentativa {tentativa + 1} falhou: {e}")
            time.sleep(0.8)
    return False

# ==== WIZARD (clique por ícone) ====
def get_xpath(we):
    js = """
    function absoluteXPath(element) {
      if (element.tagName.toLowerCase() === 'html') return '/html';
      if (element===document.body) return '/html/body';
      var ix=0; var siblings=element.parentNode.children;
      var same = 0;
      for (var i=0; i<siblings.length; i++) {
        var sib=siblings[i];
        if (sib.tagName===element.tagName) {
          same++;
          if (sib===element) {
            var path = absoluteXPath(element.parentNode) + '/' + element.tagName.toLowerCase();
            if (same>1) path += '['+same+']';
            return path;
          }
        }
      }
      return absoluteXPath(element.parentNode) + '/' + element.tagName.toLowerCase();
    }
    return absoluteXPath(arguments[0]);
    """
    return we.parent.execute_script(js, we)


def _prepare_focus_and_clear(elemento, limpar_primeiro=True):
    # Garante visibilidade e foco
    driver.execute_script("arguments[0].scrollIntoView({block:'center'});", elemento)
    try:
        elemento.click()
    except Exception:
        driver.execute_script("arguments[0].focus();", elemento)

    if limpar_primeiro:
        try:
            elemento.clear()
        except Exception:
            # Fallback de limpeza por teclas
            ActionChains(driver)\
                .move_to_element(elemento).click()\
                .key_down(Keys.CONTROL).send_keys("a").key_up(Keys.CONTROL)\
                .send_keys(Keys.DELETE).perform()





def _textarea_tradicional(elemento, texto, limpar_primeiro=True):
    _prepare_focus_and_clear(elemento, limpar_primeiro)
    elemento.send_keys(texto)
    # Dispara blur para muitos bindings reativos
    elemento.send_keys(Keys.TAB)


def _textarea_actionchains(elemento, texto, limpar_primeiro=True):
    _prepare_focus_and_clear(elemento, limpar_primeiro)
    ac = ActionChains(driver)
    ac.move_to_element(elemento).click().perform()
    # Quebra o texto em partes para evitar engasgos em campos longos
    for chunk_start in range(0, len(texto), 400):
        ac.send_keys(texto[chunk_start:chunk_start+400]).perform()
        time.sleep(0.05)
    ac.send_keys(Keys.TAB).perform()


def _textarea_js_setvalue(elemento, texto):
    # Seta .value e dispara eventos clássicos
    driver.execute_script("""
        const el = arguments[0];
        const val = arguments[1];
        el.value = val;
        // Dispara eventos comuns que form libs escutam
        el.dispatchEvent(new Event('input',  {bubbles:true}));
        el.dispatchEvent(new Event('change', {bubbles:true}));
        el.dispatchEvent(new KeyboardEvent('keyup', {bubbles:true}));
        el.dispatchEvent(new Event('blur',   {bubbles:true}));
    """, elemento, texto)


def _textarea_js_react_input(elemento, texto):
    # Compat extra p/ React (setando o setter do prototype) + eventos
    driver.execute_script("""
        const el = arguments[0];
        const val = arguments[1];

        const desc = Object.getOwnPropertyDescriptor(window.HTMLTextAreaElement.prototype, 'value');
        if (desc && desc.set) {
            desc.set.call(el, val);
        } else {
            el.value = val;
        }

        // React/Vue/Svelte geralmente escutam 'input'
        el.dispatchEvent(new Event('input', {bubbles: true}));
        el.dispatchEvent(new Event('change', {bubbles: true}));
        el.dispatchEvent(new Event('blur', {bubbles: true}));
    """, elemento, texto)


def preencher_textarea_por_indice(indice_campo, texto, max_tentativas=5, limpar_primeiro=True):
    """Preenche um <textarea> pelo índice (ordem no DOM) usando estratégias múltiplas"""
    def acao():
        if not isinstance(indice_campo, int) or indice_campo < 0:
            raise ValueError(f"Índice inválido: {indice_campo}")
        if texto is None or not isinstance(texto, str):
            raise ValueError(f"Texto inválido: {texto!r}")

        tentativa = 0
        while tentativa < max_tentativas:
            tentativa += 1
            try:
                campos = encontrar_campos_textarea()
                if not campos:
                    if tentativa < max_tentativas:
                        log(doc, f"⚠️ Nenhuma <textarea> encontrada (tentativa {tentativa}/{max_tentativas})", "WARN")
                        time.sleep(1.5)
                        continue
                    raise Exception("Nenhuma <textarea> foi encontrada na página.")

                if indice_campo >= len(campos):
                    raise Exception(f"Índice {indice_campo} inválido. Encontradas {len(campos)} textareas.")

                campo_info = campos[indice_campo]
                elemento   = campo_info["elemento"]
                campo_id   = campo_info.get("id") or "(sem id)"
                campo_name = campo_info.get("name") or "(sem name)"

                log(doc, f"🎯 Tentativa {tentativa}: Preenchendo textarea {indice_campo} (ID: {campo_id}, name: {campo_name}) com {len(texto)} caracteres")

                # Se já estiver preenchido corretamente, encerra
                if validar_textarea_preenchida(elemento, texto):
                    log(doc, f"✅ Textarea {indice_campo} já está com o valor desejado.")
                    return True

                # Estratégias em ordem de 'menos invasiva' para 'mais invasiva'
                estrategias = [
                    lambda: _textarea_tradicional(elemento, texto, limpar_primeiro),
                    lambda: _textarea_actionchains(elemento, texto, limpar_primeiro),
                    lambda: _textarea_js_setvalue(elemento, texto),
                    lambda: _textarea_js_react_input(elemento, texto),
                ]

                for i, estrategia in enumerate(estrategias, 1):
                    try:
                        log(doc, f"   ▶️ Estratégia {i}…")
                        estrategia()
                        time.sleep(0.8)

                        # Revalida após a estratégia
                        if validar_textarea_preenchida(elemento, texto):
                            val = (elemento.get_attribute("value") or "").strip()
                            log(doc, f"✅ Preenchido com sucesso pela estratégia {i}: '{val[:60]}{'…' if len(val) > 60 else ''}'")
                            return True
                        else:
                            log(doc, f"⚠️ Estratégia {i} não refletiu o valor esperado.", "WARN")
                    except (StaleElementReferenceException, JavascriptException, TimeoutException) as e:
                        log(doc, f"⚠️ Estratégia {i} falhou: {e}", "WARN")
                        # Reobter o elemento se necessário
                        try:
                            campos = encontrar_campos_textarea()
                            elemento = campos[indice_campo]["elemento"]
                        except Exception:
                            pass
                        continue

                # Se chegou aqui, nenhuma estratégia funcionou nesta tentativa
                if tentativa < max_tentativas:
                    log(doc, f"⚠️ Tentativa {tentativa} falhou; reintentando em 1.5s…", "WARN")
                    time.sleep(1.5)
                    continue
            except Exception as e:
                if tentativa < max_tentativas:
                    log(doc, f"⚠️ Erro na tentativa {tentativa}: {e}. Retentando…", "WARN")
                    time.sleep(1.5)
                    continue
                else:
                    raise

        raise Exception(f"Falha ao preencher textarea {indice_campo} após {max_tentativas} tentativas.")
    return acao


# =========================
# Helpers usados pela função
# =========================

def encontrar_campos_textarea(timeout=10):
    """
    Retorna uma lista de dicts com metadados de cada <textarea> visível e interativa.
    Ex.: [{'elemento': WebElement, 'id': '...', 'name': '...'}]
    """
    elementos = []
    try:
        # Espera haver pelo menos 1 textarea no DOM (se existir)
        wait.until(lambda d: len(d.find_elements(By.TAG_NAME, "textarea")) >= 0)
        textareas = driver.find_elements(By.TAG_NAME, "textarea")
    except Exception:
        textareas = []

    for el in textareas:
        try:
            if not el.is_displayed() or not el.is_enabled():
                continue
            driver.execute_script("arguments[0].scrollIntoView({block:'center'});", el)
            elementos.append({
                "elemento": el,
                "id": el.get_attribute("id"),
                "name": el.get_attribute("name"),
            })
        except Exception:
            continue

    return elementos


def normalizar_texto(txt):
    if txt is None:
        return ""
    # Normaliza quebras de linha e espaços
    return txt.replace("\r\n", "\n").replace("\r", "\n").strip()


def validar_textarea_preenchida(elemento, texto_esperado):
    """Confere se o valor atual da textarea bate com o texto esperado (normalizado)."""
    try:
        atual = elemento.get_attribute("value")
        # Alguns frameworks populam via textContent em textareas (raro, mas possível)
        if atual is None or atual == "":
            atual = (elemento.text or "")
        return normalizar_texto(atual) == normalizar_texto(texto_esperado)
    except StaleElementReferenceException:
        return False





def click_wizard_by_icon(driver, wait, icon_class, expect_selector=None, timeout=12):
    span = wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, f"#gsCrm .btnHolder .sprites.{icon_class}")))
    driver.execute_script("arguments[0].scrollIntoView({block:'center'});", span)
    alvo = driver.execute_script("""
        const span = arguments[0]; let el = span;
        while (el && el !== document && !/^(A|DIV)$/i.test(el.tagName)) el = el.parentElement;
        return el;
    """, span)
    if not alvo:
        raise Exception(f"Não achei ancestral clicável para o ícone .{icon_class}")
    try:
        xp = get_xpath(alvo)
        wait.until(EC.element_to_be_clickable((By.XPATH, xp))).click()
    except Exception:
        try:
            driver.execute_script("arguments[0].click();", alvo)
        except Exception:
            driver.execute_script("""
                const el=arguments[0];
                function fire(t){el.dispatchEvent(new MouseEvent(t,{bubbles:true,cancelable:true,view:window}));}
                el.focus(); fire('mouseover'); fire('mousedown'); fire('mouseup'); fire('click');
            """, alvo)
    if expect_selector:
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, expect_selector)))



# ==== DRIVER ====
def inicializar_driver():
    global driver, wait
    try:
        options = Options()
        options.add_argument("--start-maximized")
        options.add_argument("--disable-blink-features=AutomationControlled")
        options.add_experimental_option("excludeSwitches", ["enable-automation"])
        options.add_experimental_option('useAutomationExtension', False)

        driver = webdriver.Chrome(service=Service(ChromeDriverManager().install()), options=options)
        driver.execute_script("Object.defineProperty(navigator, 'webdriver', {get: () => undefined})")
        wait = WebDriverWait(driver, 20)
        return True
    except Exception as e:
        log(doc, f"❌ Erro ao inicializar driver: {e}")
        return False

# ==== EXECUÇÃO DO TESTE ====
def executar_teste():
    global driver, wait, doc
    try:
        if not inicializar_driver():
            return False
        
        # Cria engine JS forçado COM PROTEÇÃO ANTI-TIMEOUT
        js_engine = JSForceEngine(driver, wait, doc, timeout_padrao=10, max_retries=3)
        lov_handler = LOVHandler(js_engine, doc)
        
        lov_handler.resetar_selecao_sequencial()

        # Sanidade: garantir que time é módulo
        assert hasattr(time, "sleep"), f"time virou {time!r}"

        safe_action(doc, "Acessando sistema", lambda: driver.get(URL))

        safe_action(doc, "Realizando login", lambda: (
            wait.until(EC.presence_of_element_located((By.ID, "j_id15:email"))).send_keys(LOGIN_EMAIL),
            wait.until(EC.presence_of_element_located((By.ID, "j_id15:senha"))).send_keys(LOGIN_PASSWORD, Keys.ENTER),
            time.sleep(5)
        ))

        safe_action(doc, "Ajustando zoom e abrindo menu", lambda: (
            ajustar_zoom(),
            driver.find_element(By.TAG_NAME, "body").send_keys(Keys.F3)
        ))

        safe_action(doc, "Acessando Apoio Ortopédico", lambda: (
            driver.execute_script("arguments[0].scrollIntoView({block: 'center'});",
                                  wait.until(EC.presence_of_element_located((By.XPATH, '/html/body/div[15]/ul/li[22]/img')))),
            wait.until(EC.element_to_be_clickable((By.XPATH, '/html/body/div[15]/ul/li[22]/img'))).click()
        ))

        safe_action(doc, "Clicando em Locação de Equipamentos", lambda:
            wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, '#gsApoioOrtopedico > div.wdTelas > div.telaInicial.clearfix.overflow.overflowY > ul > li:nth-child(1) > a > span'))).click()
        )

        time.sleep(5)

        safe_action(doc, "Preenchendo o Número do Contrato", lambda:
                    preencher_campo_xpath_com_retry(
                        driver, wait, "//input[@class='fc mandatory' and @style='width: 200px;']",
                        "113060"
                    ))

        time.sleep(3)



        safe_action(doc, "Preenchendo Observações", lambda:
            preencher_textarea_por_indice(0,
                 "TESTE OBSERVAÇÃO SELENIUM AUTOMATIZADO (Automação de Testes): Teste de observação longa para validar o campo textarea no sistema. "))

        safe_action(doc, "Avançando para a aba: 'Adicionar Equipamentos'", lambda:
            wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, '#gsApoioOrtopedico > div.wdTelas > div.wdWizard.clearfix.telaConsulta > div.btnHolder > a:nth-child(3)'))).click()
        )

        safe_action(doc, "Selecionando Equipamento", lambda:
            lov_handler.open_and_select(
                btn_index=3,
                filter_option="Status",               # <<< AQUI
                search_text="Disponível",
                result_text="Disponível"
            )
        )
        time.sleep(3)


        safe_action(doc, "Adicionando Equipamento", lambda:
            js_engine.force_click("//a[@class='sprites sp-addVerde']", by_xpath=True)
        )
        time.sleep(3)

        safe_action(doc, "Selecionando Equipamento", lambda:
            lov_handler.open_and_select(
                btn_index=3,
                filter_option="Status",               # <<< AQUI
                search_text="Disponível",
            )
        )
        time.sleep(3)

        safe_action(doc, "Adicionando Equipamento", lambda:
            js_engine.force_click("//a[@class='sprites sp-addVerde']", by_xpath=True)
        )
        time.sleep(3)


        safe_action(doc, "Selecionando Equipamento", lambda:
            lov_handler.open_and_select(
                btn_index=3,
                filter_option="Status",               # <<< AQUI
                search_text="Disponível",
            )
        )
        time.sleep(3)

        safe_action(doc, "Adicionando Equipamento", lambda:
            js_engine.force_click("//a[@class='sprites sp-addVerde']", by_xpath=True)
        )
        time.sleep(3)



        safe_action(doc, "Avançando para a aba: 'Resumo'", lambda:
            wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, '#gsApoioOrtopedico > div.wdTelas > div.wdWizard.clearfix.telaConsulta > div.btnHolder > a:nth-child(3)'))).click()
        )

        safe_action(doc, "Cancelando cadastro", lambda:
            wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, '#gsApoioOrtopedico > div.wdTelas > div.wdWizard.clearfix.telaConsulta > div.btnHolder > a:nth-child(1)'))).click()
        )

        safe_action(doc, "Fechando modal Apoio Ortopédico", lambda:
            clicar_elemento_robusto(driver, wait, '#gsApoioOrtopedico > div.wdTop.ui-draggable-handle > div > a')
        )



        log(doc, "🔍 Verificando mensagens de alerta...")
        encontrar_mensagem_alerta()




        return True

    except Exception as e:
        log(doc, f"❌ ERRO FATAL: {e}")
        take_screenshot(driver, doc, "erro_fatal")
        return False
    finally:
        log(doc, "✅ Teste concluído.")

# ==== MAIN ====
def main():
    global doc
    try:
        log(doc, "🚀 Iniciando teste de Cadastro de Locação de Equipamentos")
        sucesso = executar_teste()
        if sucesso:
            log(doc, "✅ Teste executado com sucesso!")
        else:
            log(doc, "❌ Teste finalizado com erros.")
    except Exception as e:
        log(doc, f"❌ Erro na execução principal: {e}")
    finally:
        finalizar_relatorio()

if __name__ == "__main__":
    main()





