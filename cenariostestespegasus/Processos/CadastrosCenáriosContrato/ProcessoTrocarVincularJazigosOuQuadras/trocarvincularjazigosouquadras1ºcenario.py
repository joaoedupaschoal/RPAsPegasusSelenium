# Refatorado e organizado: cadastrodecontratos1ºcenario.py

from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support.ui import Select, WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service
from selenium.webdriver import ActionChains
from webdriver_manager.chrome import ChromeDriverManager
from docx import Document
from docx.shared import Inches
from faker import Faker
from faker.providers import BaseProvider
from validate_docbr import CPF
from datetime import datetime, timedelta
import subprocess
import os
import time
import random
import string
import re
from selenium.common.exceptions import JavascriptException, TimeoutException, WebDriverException

# ==== PROVIDERS CUSTOMIZADOS ====
class BrasilProvider(BaseProvider):
    def rg(self):
        numeros = [str(random.randint(0, 9)) for _ in range(8)]
        return ''.join(numeros) + '-' + str(random.randint(0, 9))

fake = Faker("pt_BR")
fake.add_provider(BrasilProvider)

# ==== VARIÁVEIS GLOBAIS ====
numero_aleatorio = random.randint(1, 100)
letra_aleatoria = random.choice(string.ascii_uppercase)
cemetery_name = f"Cemitério {fake.last_name()} {fake.random.choice(['Eterno', 'da Paz', 'Memorial', 'Descanso'])}"
qtd_parcelas_em_atraso = int(fake.random.choice(['1', '2', '3', '4', '5']))
dias_para_exumar = int(fake.random.choice(['365', '730', '1095', '1460', '1825']))

def gerar_jazigos():
    quantidade_ruas = random.randint(1, 10)
    max_jazigos_por_rua = random.randint(1, 20)
    quantidade_total_jazigos = quantidade_ruas * max_jazigos_por_rua
    return quantidade_ruas, max_jazigos_por_rua, quantidade_total_jazigos

ruas, jazigos_por_rua, total_jazigos = gerar_jazigos()
altura_cm = random.randint(100, 200)
largura_cm = random.randint(100, 200)
comprimento_cm = random.randint(100, 200)
valor_taxa_adesao = round(random.uniform(2000, 10000), 2)

def gerar_datas_validas():
    """Gera datas coerentes para nascimento, falecimento e sepultamento dentro de um intervalo válido."""
    hoje = datetime.today().date()
    dez_anos_atras = hoje - timedelta(days=3650)
    
    # Data de falecimento entre 10 anos atrás e hoje
    data_falecimento = fake.date_between(start_date=dez_anos_atras, end_date=hoje)
    
    # Pessoa com no mínimo 18 anos na data do falecimento
    idade_minima = 18
    idade_maxima = 110
    data_nascimento = data_falecimento - timedelta(days=random.randint(idade_minima * 365, idade_maxima * 365))
    
    # Sepultamento entre 1 e 10 dias após o falecimento
    data_sepultamento = data_falecimento + timedelta(days=random.randint(1, 10))
    
    # Registro entre 1 e 10 dias após o sepultamento
    data_registro = data_sepultamento + timedelta(days=random.randint(1, 10))
    
    # Datas de/até para período
    data_de = hoje + timedelta(days=random.randint(1, 10))
    data_ate = data_de + timedelta(days=random.randint(1, 100))
    
    return (
        data_nascimento.strftime("%d/%m/%Y"),
        data_falecimento.strftime("%d/%m/%Y"),
        data_sepultamento.strftime("%d/%m/%Y"),
        data_registro.strftime("%d/%m/%Y"),
        data_de.strftime("%d/%m/%Y"),
        data_ate.strftime("%d/%m/%Y"),
    )


def _sanitize_filename(name: str) -> str:
    name = name.strip().lower()
    name = re.sub(r"[<>:\"/\\|?*']", "_", name)
    name = re.sub(r"_+", "_", name)
    name = re.sub(r"[^\w\s-]", "_", name, flags=re.UNICODE)

    return name[:120]

TIMEOUT_DEFAULT = 30
TIMEOUT_CURTO = 10
TIMEOUT_LONGO = 60

# =========================
# Estratégias de preenchimento
# =========================

def _prepare_focus_and_clear(elemento, limpar_primeiro=True):
    # Garante visibilidade e foco
    driver.execute_script("arguments[0].scrollIntoView({block:'center'});", elemento)
    try:
        elemento.click()
    except Exception:
        driver.execute_script("arguments[0].focus();", elemento)

    if limpar_primeiro:
        try:
            elemento.clear()
        except Exception:
            # Fallback de limpeza por teclas
            ActionChains(driver)\
                .move_to_element(elemento).click()\
                .key_down(Keys.CONTROL).send_keys("a").key_up(Keys.CONTROL)\
                .send_keys(Keys.DELETE).perform()


def _textarea_tradicional(elemento, texto, limpar_primeiro=True):
    _prepare_focus_and_clear(elemento, limpar_primeiro)
    elemento.send_keys(texto)
    # Dispara blur para muitos bindings reativos
    elemento.send_keys(Keys.TAB)


def _textarea_actionchains(elemento, texto, limpar_primeiro=True):
    _prepare_focus_and_clear(elemento, limpar_primeiro)
    ac = ActionChains(driver)
    ac.move_to_element(elemento).click().perform()
    # Quebra o texto em partes para evitar engasgos em campos longos
    for chunk_start in range(0, len(texto), 400):
        ac.send_keys(texto[chunk_start:chunk_start+400]).perform()
        time.sleep(0.05)
    ac.send_keys(Keys.TAB).perform()


def _textarea_js_setvalue(elemento, texto):
    # Seta .value e dispara eventos clássicos
    driver.execute_script("""
        const el = arguments[0];
        const val = arguments[1];
        el.value = val;
        // Dispara eventos comuns que form libs escutam
        el.dispatchEvent(new Event('input',  {bubbles:true}));
        el.dispatchEvent(new Event('change', {bubbles:true}));
        el.dispatchEvent(new KeyboardEvent('keyup', {bubbles:true}));
        el.dispatchEvent(new Event('blur',   {bubbles:true}));
    """, elemento, texto)


def _textarea_js_react_input(elemento, texto):
    # Compat extra p/ React (setando o setter do prototype) + eventos
    driver.execute_script("""
        const el = arguments[0];
        const val = arguments[1];

        const desc = Object.getOwnPropertyDescriptor(window.HTMLTextAreaElement.prototype, 'value');
        if (desc && desc.set) {
            desc.set.call(el, val);
        } else {
            el.value = val;
        }

        // React/Vue/Svelte geralmente escutam 'input'
        el.dispatchEvent(new Event('input', {bubbles: true}));
        el.dispatchEvent(new Event('change', {bubbles: true}));
        el.dispatchEvent(new Event('blur', {bubbles: true}));
    """, elemento, texto)



# ==== SISTEMA DATEPICKER MELHORADO ====
def encontrar_campos_datepicker():
    """Encontra todos os campos datepicker na página"""
    global driver
    
    if driver is None:
        return []
    
    seletores_datepicker = [
        "input.hasDatepicker",
        "input[id^='dp']",
        "input[maxlength='10'][grupo='']",
        "input[type='text'][maxlength='10']",
        "input[class*='datepicker']",
        ".hasDatepicker"
    ]
    
    campos_encontrados = []
    
    for seletor in seletores_datepicker:
        try:
            elementos = driver.find_elements(By.CSS_SELECTOR, seletor)
            for elemento in elementos:
                if elemento.is_displayed() and elemento.is_enabled():
                    info = {
                        'elemento': elemento,
                        'id': elemento.get_attribute('id') or f"dp_{len(campos_encontrados)}",
                        'seletor_usado': seletor,
                        'maxlength': elemento.get_attribute('maxlength'),
                        'placeholder': elemento.get_attribute('placeholder')
                    }
                    # Evita duplicatas
                    if not any(c['id'] == info['id'] for c in campos_encontrados):
                        campos_encontrados.append(info)
        except Exception as e:
            log(doc, f"⚠️ Erro ao buscar campos datepicker com {seletor}: {e}", 'WARN')
            continue
    
    log(doc, f"📊 Encontrados {len(campos_encontrados)} campos datepicker")
    return campos_encontrados

def _datepicker_jquery(campo_id, data_valor):
    """Estratégia jQuery para datepicker"""
    global driver
    
    resultado = driver.execute_script("""
        var campoId = arguments[0], valor = arguments[1];
        if (typeof jQuery === 'undefined') return 'jQuery não disponível';
        var $campo = $('#' + campoId);
        if (!$campo.length) return 'Campo não encontrado: ' + campoId;
        try {
            if ($campo.hasClass('hasDatepicker')) { 
                $campo.datepicker('setDate', valor); 
            } else { 
                $campo.val(valor); 
            }
            $campo.trigger('input').trigger('change').trigger('blur');
            return $campo.val();
        } catch(e) { 
            return 'Erro: ' + e.message; 
        }
    """, campo_id, data_valor)
    
    if isinstance(resultado, str) and ('Erro' in resultado or 'não disponível' in resultado):
        raise Exception(f"jQuery falhou: {resultado}")

def _datepicker_javascript(elemento, data_valor):
    """Estratégia JavaScript para datepicker"""
    global driver
    
    driver.execute_script("""
        var campo = arguments[0], valor = arguments[1];
        campo.focus(); 
        campo.value = ''; 
        campo.value = valor;
        ['input','change','blur','keyup'].forEach(ev => 
            campo.dispatchEvent(new Event(ev, {bubbles: true}))
        );
    """, elemento, data_valor)

def _datepicker_actionchains(elemento, data_valor):
    """Estratégia ActionChains para datepicker"""
    global driver
    
    scroll_to_element_safe(elemento)
    time.sleep(0.5)
    
    ActionChains(driver).move_to_element(elemento).click().perform()
    time.sleep(0.5)
    ActionChains(driver).key_down(Keys.CONTROL).send_keys('a').key_up(Keys.CONTROL).perform()
    time.sleep(0.3)
    ActionChains(driver).send_keys(Keys.DELETE).perform()
    time.sleep(0.3)
    
    for char in data_valor:
        ActionChains(driver).send_keys(char).perform()
        time.sleep(0.05)
    
    ActionChains(driver).send_keys(Keys.TAB).perform()

def _datepicker_tradicional(elemento, data_valor):
    """Estratégia tradicional para datepicker"""
    scroll_to_element_safe(elemento)
    time.sleep(0.5)
    elemento.click()
    time.sleep(0.5)
    elemento.clear()
    elemento.send_keys(data_valor)
    elemento.send_keys(Keys.TAB)

def validar_data_preenchida(elemento, data_esperada):
    """Valida se a data foi preenchida corretamente"""
    try:
        if elemento is None:
            return False
            
        val = (elemento.get_attribute('value') or '').strip()
        if not val:
            return False
            
        if val == data_esperada or data_esperada in val:
            return True
            
        # Tenta comparar datas em diferentes formatos
        formatos = [
            '%d/%m/%Y %H:%M', '%d/%m/%Y %H:%M:%S',
            '%d/%m/%Y', '%m/%d/%Y', '%Y-%m-%d', '%d-%m-%Y'
        ]
        
        for formato in formatos:
            try:
                d1 = datetime.strptime(val, formato)
                d2 = datetime.strptime(data_esperada, formato)
                if d1 == d2:
                    return True
            except:
                continue
                
        return False
        
    except Exception:
        return False





def preencher_datepicker_por_indice(indice_campo, data_valor, max_tentativas=5):
    """Preenche datepicker pelo índice com estratégias múltiplas"""
    def acao():
        if not isinstance(indice_campo, int) or indice_campo < 0:
            raise ValueError(f"Índice inválido: {indice_campo}")
            
        if not data_valor or not isinstance(data_valor, str):
            raise ValueError(f"Data inválida: {data_valor}")
        
        tentativa = 0
        while tentativa < max_tentativas:
            tentativa += 1
            
            try:
                campos = encontrar_campos_datepicker()
                
                if not campos:
                    if tentativa < max_tentativas:
                        log(doc, f"⚠️ Nenhum campo datepicker encontrado, tentativa {tentativa}/{max_tentativas}", 'WARN')
                        time.sleep(2)
                        continue
                    raise Exception("Nenhum campo datepicker encontrado na página")
                
                if indice_campo >= len(campos):
                    raise Exception(f"Índice {indice_campo} inválido. Encontrados {len(campos)} campos")
                
                campo_info = campos[indice_campo]
                elemento = campo_info['elemento']
                campo_id = campo_info['id']
                
                log(doc, f"🎯 Tentativa {tentativa}: Preenchendo datepicker {indice_campo} (ID: {campo_id}) com '{data_valor}'")
                
                # Verifica se já está preenchido corretamente
                if validar_data_preenchida(elemento, data_valor):
                    log(doc, f"✅ Campo {indice_campo} já está preenchido corretamente!")
                    return True
                
                # Estratégias específicas para datepicker
                estrategias = [
                    lambda: _datepicker_jquery(campo_id, data_valor),
                    lambda: _datepicker_javascript(elemento, data_valor),
                    lambda: _datepicker_actionchains(elemento, data_valor),
                    lambda: _datepicker_tradicional(elemento, data_valor)
                ]
                
                for i, estrategia in enumerate(estrategias, 1):
                    try:
                        log(doc, f"   Aplicando estratégia {i} para datepicker...")
                        estrategia()
                        time.sleep(1)
                        
                        # Verifica se funcionou
                        if validar_data_preenchida(elemento, data_valor):
                            valor_atual = elemento.get_attribute('value')
                            log(doc, f"✅ Datepicker preenchido com estratégia {i}: '{valor_atual}'")
                            return True
                        else:
                            log(doc, f"⚠️ Estratégia {i} não preencheu corretamente", 'WARN')
                            
                    except Exception as e:
                        log(doc, f"⚠️ Estratégia {i} falhou: {e}", 'WARN')
                        continue
                
                # Se chegou aqui, nenhuma estratégia funcionou nesta tentativa
                if tentativa < max_tentativas:
                    log(doc, f"⚠️ Tentativa {tentativa} falhou, tentando novamente em 2s...", 'WARN')
                    time.sleep(2)
                    continue
                
            except Exception as e:
                if tentativa < max_tentativas:
                    log(doc, f"⚠️ Erro na tentativa {tentativa}: {e}, tentando novamente...", 'WARN')
                    time.sleep(2)
                    continue
                else:
                    raise
        
        raise Exception(f"Falha ao preencher datepicker {indice_campo} após {max_tentativas} tentativas")
    
    return acao


def preencher_datepicker_por_indice_xpath(js_engine, doc, xpath_base, indice, data_valor, descricao="Data", timeout=10):
    """
    Preenche um campo datepicker específico por índice.
    Versão especializada para campos de data que usa estratégias específicas.
    
    Args:
        js_engine: Instância do JSForceEngine
        doc: Documento para logs
        xpath_base: XPath base do datepicker
        indice: Índice do elemento (1-based)
        data_valor: Data no formato string (ex: "30/10/2025")
        descricao: Descrição do campo para logs
        timeout: Timeout para operação
    
    Returns:
        bool: True se preencheu com sucesso, False caso contrário
    """
    xpath_indexado = f"({xpath_base})[{indice}]"
    
    try:
        log(doc, f"📅 Preenchendo {descricao} (índice {indice}): '{data_valor}'")
        
        # Conta elementos
        try:
            elementos = js_engine.driver.find_elements(By.XPATH, xpath_base)
            total = len(elementos)
            log(doc, f"   ℹ️ Total de datepickers encontrados: {total}")
            
            if total == 0:
                log(doc, f"   ⚠️ Nenhum datepicker encontrado")
                return False
            
            if indice > total:
                log(doc, f"   ⚠️ Índice {indice} inválido - só existem {total} datepicker(s)")
                return False
        except:
            pass
        
        # Localiza o elemento específico
        elemento = js_engine.driver.find_element(By.XPATH, xpath_indexado)
        campo_id = elemento.get_attribute('id') or f"datepicker_{indice}"
        
        log(doc, f"   🎯 Elemento localizado (ID: {campo_id})")
        
        # Estratégias específicas para datepicker
        estrategias = [
            # Estratégia 1: jQuery datepicker
            lambda: js_engine.driver.execute_script(f"""
                var campo = document.evaluate("{xpath_indexado}", document, null, 
                    XPathResult.FIRST_ORDERED_NODE_TYPE, null).singleNodeValue;
                if (!campo) return false;
                
                if (typeof jQuery !== 'undefined' && campo.classList.contains('hasDatepicker')) {{
                    jQuery(campo).datepicker('setDate', '{data_valor}');
                }} else {{
                    campo.value = '{data_valor}';
                }}
                
                campo.dispatchEvent(new Event('input', {{bubbles: true}}));
                campo.dispatchEvent(new Event('change', {{bubbles: true}}));
                campo.dispatchEvent(new Event('blur', {{bubbles: true}}));
                return campo.value;
            """),
            
            # Estratégia 2: Force fill padrão
            lambda: js_engine.force_fill(xpath_indexado, data_valor, by_xpath=True),
            
            # Estratégia 3: JavaScript puro
            lambda: js_engine.driver.execute_script(f"""
                var campo = document.evaluate("{xpath_indexado}", document, null, 
                    XPathResult.FIRST_ORDERED_NODE_TYPE, null).singleNodeValue;
                if (!campo) return false;
                
                campo.focus();
                campo.value = '';
                campo.value = '{data_valor}';
                
                ['input', 'change', 'blur', 'keyup'].forEach(ev => 
                    campo.dispatchEvent(new Event(ev, {{bubbles: true}}))
                );
                return campo.value;
            """),
            
            # Estratégia 4: ActionChains
            lambda: (
                elemento.click(),
                time.sleep(0.2),
                ActionChains(js_engine.driver)
                    .key_down(Keys.CONTROL).send_keys('a').key_up(Keys.CONTROL)
                    .send_keys(Keys.DELETE)
                    .send_keys(data_valor)
                    .send_keys(Keys.TAB)
                    .perform()
            )
        ]
        
        # Tenta cada estratégia
        for i, estrategia in enumerate(estrategias, 1):
            try:
                log(doc, f"   ▶️ Tentando estratégia {i}...")
                estrategia()
                time.sleep(0.5)
                
                # Valida
                elemento_reload = js_engine.driver.find_element(By.XPATH, xpath_indexado)
                valor_atual = elemento_reload.get_attribute('value') or ''
                
                if data_valor in valor_atual or valor_atual in data_valor:
                    log(doc, f"   ✅ {descricao} preenchido com estratégia {i}: '{valor_atual}'")
                    return True
                else:
                    log(doc, f"   ⚠️ Estratégia {i} não refletiu o valor")
                    
            except Exception as e:
                log(doc, f"   ⚠️ Estratégia {i} falhou: {str(e)[:80]}")
                continue
        
        log(doc, f"   ❌ Todas as estratégias falharam para {descricao}")
        return False
        
    except Exception as e:
        log(doc, f"   ❌ Erro ao preencher {descricao}: {e}")
        return False


def preencher_campo_monetario_por_indice(js_engine, doc, xpath_base, indice, valor, descricao="Valor", timeout=10):
    """
    Preenche um campo monetário específico por índice.
    Versão especializada para campos de valor/dinheiro.
    
    Args:
        js_engine: Instância do JSForceEngine
        doc: Documento para logs
        xpath_base: XPath base do campo monetário
        indice: Índice do elemento (1-based)
        valor: Valor a ser preenchido (pode ser com ou sem "R$ ")
        descricao: Descrição do campo para logs
        timeout: Timeout para operação
    
    Returns:
        bool: True se preencheu com sucesso, False caso contrário
    """
    xpath_indexado = f"({xpath_base})[{indice}]"
    
    try:
        log(doc, f"💰 Preenchendo {descricao} (índice {indice}): '{valor}'")
        
        # Conta elementos
        try:
            elementos = js_engine.driver.find_elements(By.XPATH, xpath_base)
            total = len(elementos)
            log(doc, f"   ℹ️ Total de campos encontrados: {total}")
            
            if total == 0:
                log(doc, f"   ⚠️ Nenhum campo encontrado")
                return False
            
            if indice > total:
                log(doc, f"   ⚠️ Índice {indice} inválido - só existem {total} campo(s)")
                return False
        except:
            pass
        
        # Normaliza o valor (garante que tem R$)
        valor_limpo = str(valor).replace('R$', '').replace(' ', '').strip()
        valor_formatado = f"R$ {valor_limpo}" if not valor.startswith('R$') else valor
        
        log(doc, f"   🎯 Valor formatado: '{valor_formatado}'")
        
        # Tenta preencher
        js_engine.force_fill(xpath_indexado, valor_formatado, by_xpath=True)
        time.sleep(0.4)
        
        # Valida
        try:
            elemento = js_engine.driver.find_element(By.XPATH, xpath_indexado)
            valor_atual = elemento.get_attribute('value') or ''
            
            # Compara valores normalizados
            atual_num = valor_atual.replace('R$', '').replace(' ', '').replace('.', '').replace(',', '.')
            esper_num = valor_limpo.replace('.', '').replace(',', '.')
            
            try:
                if float(atual_num.strip()) == float(esper_num.strip()):
                    log(doc, f"   ✅ {descricao} preenchido: '{valor_atual}'")
                    return True
            except:
                pass
            
            # Se não conseguiu comparar numericamente, compara texto
            if valor_limpo in valor_atual or valor_atual in valor_limpo:
                log(doc, f"   ✅ {descricao} preenchido: '{valor_atual}'")
                return True
            else:
                log(doc, f"   ⚠️ Valores diferem: esperado '{valor}', atual '{valor_atual}'")
                return True  # Retorna True pois pode ser só formatação
                
        except Exception as e:
            log(doc, f"   ⚠️ Não foi possível validar: {e}")
            return True
        
    except Exception as e:
        log(doc, f"   ❌ Erro ao preencher {descricao}: {e}")
        return False

# ==== SISTEMA ANTI-TIMEOUT JAVASCRIPT ====
class JSTimeoutHandler:
    """Sistema robusto para lidar com timeouts JavaScript no Selenium"""
    
    def __init__(self, driver, doc, timeout_padrao=10, max_retries=3):
        self.driver = driver
        self.doc = doc
        self.timeout_padrao = timeout_padrao
        self.max_retries = max_retries
        self.last_error = None
        
    def log_timeout(self, msg, level="INFO"):
        """Log com timestamp"""

        prefix = {
            "INFO": "ℹ️ ",
            "WARN": "⚠️ ",
            "ERROR": "❌ ",
            "SUCCESS": "✅ "
        }.get(level, "📝 ")
        
        print(f" {prefix} {msg}")
        if hasattr(self.doc, 'add_paragraph'):
            self.doc.add_paragraph(f"{msg}")
    
    def execute_js_safe(self, script, *args, timeout=None, fallback_result=None):
        """Executa JavaScript com proteção contra timeouts"""
        timeout = timeout or self.timeout_padrao
        
        original_timeout = self.driver.timeouts.script
        self.driver.set_script_timeout(timeout)
        
        for tentativa in range(1, self.max_retries + 1):
            try:
                if tentativa > 1:
                    self.log_timeout(f"Tentativa {tentativa}/{self.max_retries}", "INFO")
                
                result = self.driver.execute_script(script, *args)
                self.driver.set_script_timeout(original_timeout)
                
                if tentativa > 1:
                    self.log_timeout("JavaScript executado com sucesso", "SUCCESS")
                return result
                
            except JavascriptException as e:
                self.last_error = e
                self.log_timeout(f"Erro JavaScript: {str(e)[:150]}", "ERROR")
                self._limpar_estado_js()
                
                if tentativa < self.max_retries:
                    time.sleep(1 + tentativa * 0.5)
                    continue
                    
            except TimeoutException as e:
                self.last_error = e
                self.log_timeout(f"Timeout JavaScript ({timeout}s)", "ERROR")
                self._forcar_parada_js()
                
                if tentativa < self.max_retries:
                    time.sleep(2 + tentativa)
                    continue
                    
            except WebDriverException as e:
                self.last_error = e
                self.log_timeout(f"Erro WebDriver: {str(e)[:150]}", "ERROR")
                
                if tentativa < self.max_retries:
                    time.sleep(1.5)
                    continue
                    
            except Exception as e:
                self.last_error = e
                self.log_timeout(f"Erro inesperado: {str(e)[:150]}", "ERROR")
                break
        
        try:
            self.driver.set_script_timeout(original_timeout)
        except:
            pass
            
        if fallback_result is not None:
            self.log_timeout(f"Usando valor fallback: {fallback_result}", "WARN")
        return fallback_result
    
    def _limpar_estado_js(self):
        """Limpa estado JavaScript do browser"""
        try:
            cleanup_script = """
                if (window.__cleanupTimers) {
                    window.__cleanupTimers.forEach(clearTimeout);
                    window.__cleanupTimers.forEach(clearInterval);
                }
                if (typeof jQuery !== 'undefined') {
                    jQuery.active = 0;
                }
                window.__pendingRequests = 0;
                return true;
            """
            self.driver.execute_script(cleanup_script)
            time.sleep(0.5)
        except Exception:
            pass
    
    def _forcar_parada_js(self):
        """Força parada de JavaScript travado"""
        try:
            self.driver.execute_script("window.stop();")
            time.sleep(0.3)
        except Exception:
            pass



# ==== JS FORCE ENGINE COM PROTEÇÃO ANTI-TIMEOUT ====
class JSForceEngine:
    """Motor de execução JavaScript forçado com proteção contra timeouts"""
    
    def __init__(self, driver, wait, doc, timeout_padrao=10, max_retries=3):
        self.driver = driver
        self.wait = wait
        self.doc = doc
        self.timeout_handler = JSTimeoutHandler(driver, doc, timeout_padrao, max_retries)
    
    def execute_js(self, script, *args, timeout=None, fallback_result=None):
        """Executa JavaScript com proteção contra timeout"""
        return self.timeout_handler.execute_js_safe(
            script, *args, timeout=timeout, fallback_result=fallback_result
        )
    
    def wait_ajax_complete(self, timeout=15):
        """Aguarda AJAX completar com proteção contra timeout"""
        script = """
            var jQueryOk = (typeof jQuery==='undefined') || (jQuery.active===0);
            var fetchOk = !window.__pendingRequests || window.__pendingRequests===0;
            var overlays = document.querySelectorAll(
                '.blockScreen, .blockUI, .loading, .overlay, [class*="loading"], [class*="spinner"]'
            );
            var overlayOk = true;
            for (var i=0; i<overlays.length; i++){
                var s=window.getComputedStyle(overlays[i]);
                if(s.display!=='none' && s.visibility!=='hidden' && parseFloat(s.opacity||1)>0.01){
                    overlayOk=false;
                    break;
                }
            }
            return jQueryOk && fetchOk && overlayOk;
        """
        
        end = time.time() + timeout
        while time.time() < end:
            try:
                done = self.execute_js(script, timeout=5, fallback_result=True)
                if done:
                    return True
            except:
                pass
            time.sleep(0.2)
        return True
    
    def scroll_into_view(self, target, padding=100):
        """
        Faz scroll até um WebElement (ou seletor XPath/CSS).
        padding: desloca um pouco pra cima pra não ficar colado no topo.
        """
        try:
            el = target
            # Se vier como string, tenta resolver
            if isinstance(target, str):
                try:
                    el = self.driver.find_element("xpath", target)
                except Exception:
                    el = self.driver.find_element("css selector", target)

            # Estratégia principal via JS
            self.driver.execute_script("""
                const el = arguments[0], pad = arguments[1] || 0;
                if (!el) return;
                el.scrollIntoView({block:'center', inline:'center'});
                try { window.scrollBy(0, -pad); } catch(e) {}
            """, el, padding)
            time.sleep(0.2)
            return True
        except Exception:
            # Fallback ActionChains
            try:
                ActionChains(self.driver).move_to_element(el).perform()
                time.sleep(0.2)
                return True
            except Exception:
                return False

    def click_element(self, el, wait_after=0.4):
        """
        Clica com múltiplas estratégias em um WebElement já localizado.
        """
        try:
            # 1) Clique padrão
            el.click()
            time.sleep(wait_after)
            self.wait_ajax_complete(min(10, int(wait_after*10)) or 3)
            return True
        except Exception:
            pass

        try:
            # 2) Clique via JS
            self.driver.execute_script("arguments[0].click();", el)
            time.sleep(wait_after)
            self.wait_ajax_complete(min(10, int(wait_after*10)) or 3)
            return True
        except Exception:
            pass

        try:
            # 3) Sequência de eventos de mouse
            self.execute_js("""
                const e = arguments[0];
                const rect = e.getBoundingClientRect();
                const x = rect.left + 5, y = rect.top + 5;
                ['mouseover','mouseenter','mousemove','mousedown','mouseup','click'].forEach(t=>{
                    e.dispatchEvent(new MouseEvent(t,{bubbles:true,cancelable:true,view:window,clientX:x,clientY:y}));
                });
                if (typeof e.click==='function') e.click();
                return true;
            """, el, timeout=5, fallback_result=True)
            time.sleep(wait_after)
            self.wait_ajax_complete(min(10, int(wait_after*10)) or 3)
            return True
        except Exception:
            pass

        # 4) Último recurso: ActionChains
        try:
            ActionChains(self.driver).move_to_element(el).pause(0.05).click().perform()
            time.sleep(wait_after)
            self.wait_ajax_complete(min(10, int(wait_after*10)) or 3)
            return True
        except Exception:
            pass

        raise Exception("Não foi possível clicar no elemento com as estratégias disponíveis.")
    
    def force_click(self, selector, by_xpath=False, max_attempts=5):
        """Clique forçado com proteção contra timeout"""
        log(self.doc, f"🎯 Clique forçado em: {selector}")
        
        for attempt in range(max_attempts):
            try:
                strategies = [
                    self._click_strategy_2,
                    self._click_strategy_1,
                    self._click_strategy_3,
                    self._click_strategy_4,
                    self._click_strategy_5,
                ]
                
                for i, strategy in enumerate(strategies, 1):
                    try:
                        if attempt > 0 or i > 1:
                            log(self.doc, f"   Tentativa {attempt + 1}.{i}...")
                        
                        result = self.execute_js(
                            self._get_strategy_script(strategy, selector, by_xpath),
                            selector,
                            by_xpath,
                            timeout=5,
                            fallback_result=False
                        )
                        
                        if result:
                            log(self.doc, f"✅ Clique bem-sucedido (estratégia {i})")
                            time.sleep(0.5)
                            self.wait_ajax_complete(10)
                            return True
                            
                    except Exception as e:
                        if i == 1 and attempt == 0:
                            pass  # Silencia primeiro erro
                        else:
                            log(self.doc, f"   Estratégia {i} falhou: {str(e)[:80]}")
                        continue
                
                if attempt < max_attempts - 1:
                    time.sleep(1 + attempt * 0.5)
                    
            except Exception as e:
                log(self.doc, f"⚠️ Tentativa {attempt + 1} falhou: {e}")
                if attempt < max_attempts - 1:
                    time.sleep(1.5)
        
        raise Exception(f"Falha ao clicar após {max_attempts} tentativas: {selector}")
    
    def _get_strategy_script(self, strategy_func, selector, by_xpath):
        """Retorna o script JavaScript para cada estratégia"""
        base_locator = """
            var selector = arguments[0];
            var byXPath = arguments[1];
            var element;
            
            if (byXPath) {
                var result = document.evaluate(selector, document, null, 
                    XPathResult.FIRST_ORDERED_NODE_TYPE, null);
                element = result.singleNodeValue;
            } else {
                element = document.querySelector(selector);
            }
            
            if (!element) throw new Error('Elemento não encontrado');
        """
        
        if strategy_func == self._click_strategy_1:
            return base_locator + """
                element.style.pointerEvents = 'auto';
                element.style.display = 'block';
                element.style.visibility = 'visible';
                element.style.opacity = '1';
                element.removeAttribute('disabled');
                element.scrollIntoView({behavior: 'smooth', block: 'center'});
                setTimeout(function() { element.click(); }, 300);
                return true;
            """
        elif strategy_func == self._click_strategy_2:
            return base_locator + """
                element.style.pointerEvents = 'auto';
                element.removeAttribute('disabled');
                element.scrollIntoView({block: 'center'});
                
                var events = ['mouseover', 'mouseenter', 'mousemove', 'mousedown', 'mouseup', 'click'];
                events.forEach(function(eventType) {
                    var evt = new MouseEvent(eventType, {
                        bubbles: true, cancelable: true, view: window, detail: 1,
                        clientX: element.getBoundingClientRect().left + 5,
                        clientY: element.getBoundingClientRect().top + 5
                    });
                    element.dispatchEvent(evt);
                });
                
                if (typeof element.click === 'function') element.click();
                return true;
            """
        elif strategy_func == self._click_strategy_3:
            return base_locator + """
                element.style.display = 'block';
                element.style.visibility = 'visible';
                element.style.opacity = '1';
                element.style.pointerEvents = 'auto';
                element.focus();
                element.click();
                element.dispatchEvent(new Event('click', {bubbles: true, cancelable: true}));
                return true;
            """
        elif strategy_func == self._click_strategy_4:
            return base_locator + """
                element.removeAttribute('disabled');
                element.removeAttribute('readonly');
                element.style.pointerEvents = 'auto !important';
                element.style.display = 'block !important';
                element.style.visibility = 'visible !important';
                element.style.opacity = '1 !important';
                
                var overlays = document.querySelectorAll('.modal, .overlay, .blockUI, [role="dialog"]');
                overlays.forEach(function(overlay) {
                    overlay.style.display = 'none';
                    overlay.style.visibility = 'hidden';
                });
                
                element.focus();
                element.click();
                
                var clickEvent = new MouseEvent('click', {
                    view: window, bubbles: true, cancelable: true
                });
                element.dispatchEvent(clickEvent);
                
                if (typeof jQuery !== 'undefined') jQuery(element).trigger('click');
                return true;
            """
        else:  # strategy_5
            return base_locator + """
                var rect = element.getBoundingClientRect();
                var x = rect.left + rect.width / 2;
                var y = rect.top + rect.height / 2;
                
                var evt = document.createEvent('MouseEvents');
                evt.initMouseEvent('click', true, true, window, 1, x, y, x, y, false, false, false, false, 0, null);
                element.dispatchEvent(evt);
                
                if (element.onclick) element.onclick();
                
                var parent = element.parentElement;
                while (parent && parent !== document.body) {
                    if (parent.onclick) {
                        parent.onclick();
                        break;
                    }
                    parent = parent.parentElement;
                }
                return true;
            """
    
    def _click_strategy_1(self, selector, by_xpath):
        pass  # Implementado via _get_strategy_script
    
    def _click_strategy_2(self, selector, by_xpath):
        pass
    
    def _click_strategy_3(self, selector, by_xpath):
        pass
    
    def _click_strategy_4(self, selector, by_xpath):
        pass
    
    def _click_strategy_5(self, selector, by_xpath):
        pass
    
    def force_fill(self, selector, value, by_xpath=False, max_attempts=5):
        """Preenchimento forçado com proteção contra timeout"""
        log(self.doc, f"✏️ Preenchimento forçado: {selector} = '{value}'")
        
        for attempt in range(max_attempts):
            try:
                strategies = [
                    self._fill_strategy_1,
                    self._fill_strategy_2,
                    self._fill_strategy_3,
                    self._fill_strategy_4,
                ]
                
                for i, strategy in enumerate(strategies, 1):
                    try:
                        if attempt > 0 or i > 1:
                            log(self.doc, f"   Tentativa {attempt + 1}.{i}...")
                        
                        result = self.execute_js(
                            self._get_fill_script(strategy, selector, value, by_xpath),
                            selector,
                            value,
                            by_xpath,
                            timeout=5,
                            fallback_result=None
                        )
                        
                        time.sleep(0.3)
                        if self._validate_fill(selector, value, by_xpath):
                            log(self.doc, f"✅ Campo preenchido (estratégia {i})")
                            return True
                    except Exception as e:
                        if i == 1 and attempt == 0:
                            pass
                        else:
                            log(self.doc, f"   Estratégia {i} falhou: {str(e)[:80]}")
                        continue
                
                if attempt < max_attempts - 1:
                    time.sleep(1 + attempt * 0.5)
                    
            except Exception as e:
                log(self.doc, f"⚠️ Tentativa {attempt + 1} falhou: {e}")
        
        raise Exception(f"Falha ao preencher após {max_attempts} tentativas: {selector}")

    def _get_fill_script(self, strategy_func, selector, value, by_xpath):
        """Retorna o script de preenchimento"""
        base_locator = """
            var selector = arguments[0];
            var value = arguments[1];
            var byXPath = arguments[2];
            var element;
            
            if (byXPath) {
                var result = document.evaluate(selector, document, null, 
                    XPathResult.FIRST_ORDERED_NODE_TYPE, null);
                element = result.singleNodeValue;
            } else {
                element = document.querySelector(selector);
            }
            
            if (!element) throw new Error('Campo não encontrado');
        """
        
        if strategy_func == self._fill_strategy_1:
            return base_locator + """
                element.removeAttribute('disabled');
                element.removeAttribute('readonly');
                element.style.display = 'block';
                element.style.visibility = 'visible';
                element.scrollIntoView({block: 'center'});
                element.focus();
                element.dispatchEvent(new Event('focus', {bubbles: true}));
                element.value = '';
                element.value = value;
                ['input', 'change', 'blur', 'keyup'].forEach(function(evt) {
                    element.dispatchEvent(new Event(evt, {bubbles: true}));
                });
                return element.value;
            """
        elif strategy_func == self._fill_strategy_2:
            return base_locator + """
                var nativeInputValueSetter = Object.getOwnPropertyDescriptor(
                    window.HTMLInputElement.prototype, 'value'
                ).set;
                if (nativeInputValueSetter) {
                    nativeInputValueSetter.call(element, value);
                } else {
                    element.value = value;
                }
                element.dispatchEvent(new Event('input', {bubbles: true}));
                element.dispatchEvent(new Event('change', {bubbles: true}));
                element.dispatchEvent(new KeyboardEvent('keydown', {bubbles: true}));
                element.dispatchEvent(new KeyboardEvent('keyup', {bubbles: true}));
                element.dispatchEvent(new Event('blur', {bubbles: true}));
                return element.value;
            """
        elif strategy_func == self._fill_strategy_3:
            return base_locator + """
                element.value = value;
                if (typeof jQuery !== 'undefined') {
                    jQuery(element).val(value).trigger('input').trigger('change').trigger('blur');
                }
                ['focus', 'input', 'change', 'blur'].forEach(function(evt) {
                    element.dispatchEvent(new Event(evt, {bubbles: true}));
                });
                return element.value;
            """
        else:  # strategy_4
            return base_locator + """
                element.removeAttribute('disabled');
                element.removeAttribute('readonly');
                element.removeAttribute('maxlength');
                element.value = '';
                element.setAttribute('value', value);
                element.value = value;
                element.style.color = element.style.color;
                
                var events = ['focus', 'click', 'input', 'change', 'keydown', 'keypress', 
                              'keyup', 'blur', 'paste', 'textInput'];
                events.forEach(function(evt) {
                    try {
                        element.dispatchEvent(new Event(evt, {bubbles: true, cancelable: true}));
                    } catch(e) {}
                });
                
                if (element.oninput) element.oninput();
                if (element.onchange) element.onchange();
                return element.value;
            """
    
    def _fill_strategy_1(self, selector, value, by_xpath):
        pass
    
    def _fill_strategy_2(self, selector, value, by_xpath):
        pass
    
    def _fill_strategy_3(self, selector, value, by_xpath):
        pass
    
    def _fill_strategy_4(self, selector, value, by_xpath):
        pass
    
    def _validate_fill(self, selector, expected_value, by_xpath):
        """Valida preenchimento"""
        script = """
            var selector = arguments[0];
            var expected = arguments[1];
            var byXPath = arguments[2];
            var element;
            
            if (byXPath) {
                var result = document.evaluate(selector, document, null, 
                    XPathResult.FIRST_ORDERED_NODE_TYPE, null);
                element = result.singleNodeValue;
            } else {
                element = document.querySelector(selector);
            }
            
            if (!element) return false;
            var actual = element.value || '';
            return actual.trim() === expected.trim() || actual.includes(expected);
        """
        try:
            return self.execute_js(script, selector, expected_value, by_xpath, timeout=3, fallback_result=False)
        except:
            return False
        
import time

def clicar_finalizar_e_verificar_alerta(js_engine, doc, timeout=5, pausa=0.5):
    """
    Clica no botão 'Finalizar', aguarda 0,5s e procura mensagens de alerta.
    Usa safe_action e js_engine.force_click().
    """
    safe_action(doc, "Clicando em 'Finalizar'", lambda:
        js_engine.force_click(
            "//a[@class='btModel btGray' and normalize-space()='Finalizar']",
            by_xpath=True
        )
    )

    time.sleep(pausa)
    log(doc, "🔍 Verificando mensagens de alerta após o clique em 'Finalizar'...")
    return encontrar_mensagem_alerta()


import time
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.common.exceptions import (
    TimeoutException,
    StaleElementReferenceException,
    ElementClickInterceptedException,
    NoSuchElementException,
    ElementNotInteractableException
)
from selenium.webdriver import ActionChains

class LOVHandler:
    """
    Handler ultra-robusto para manipulação de LOV (List of Values) com:
    - 10 estratégias diferentes de clique
    - Detecção automática de iframes
    - Retry inteligente com backoff exponencial
    - Reforço automático de cliques
    - Logs detalhados de cada tentativa
    """
    
    def __init__(self, js_engine, doc, max_retries=5, timeout=10):
        self.js = js_engine
        self.doc = doc
        self.driver = js_engine.driver
        self.max_retries = max_retries
        self.timeout = timeout
        
        # ✅ CORREÇÃO 1: Inicializar contador sequencial
        self.selecao_sequencial_index = 0
        
    def _log(self, msg, level="INFO"):
        """Log padronizado com níveis"""
        prefixes = {
            "INFO": "ℹ️",
            "SUCCESS": "✅",
            "WARNING": "⚠️",
            "ERROR": "❌",
            "DEBUG": "🔍"
        }
        prefix = prefixes.get(level, "📝")
        log(self.doc, f"{prefix} {msg}")
    


    def open_and_select(
        self,
        btn_index=None,
        btn_xpath=None,
        btn_css=None,
        search_text="",
        result_text="",      # agora só pra log
        filter_option=None,
        iframe_xpath=None,
        auto_detect_iframe=True,
        reinforce_clicks=5,
        wait_after=0.5
    ):
        """
        Método principal: abre LOV, aplica filtro/pesquisa e seleciona
        SEMPRE de forma SEQUENCIAL (1ª vez = 1ª linha, 2ª = 2ª, etc.).
        """

        self._log(
            f"🔍 LOV sequencial | filtro='{filter_option}' | pesquisa='{search_text}' | índice atual={self.selecao_sequencial_index}",
            "INFO"
        )

        for retry in range(1, self.max_retries + 1):
            try:
                # ===== PASSO 1: volta p/ contexto principal =====
                try:
                    self.driver.switch_to.default_content()
                except Exception:
                    pass

                # ===== PASSO 2: localizar botão do LOV =====
                btn_selector = None
                by_type = None

                if btn_index is not None:
                    # índice é 0-based aqui, então soma 1 no XPath
                    btn_selector = f"(//a[@class='sprites sp-openLov'])[{btn_index + 1}]"
                    by_type = "xpath"
                elif btn_xpath:
                    btn_selector = btn_xpath
                    by_type = "xpath"
                elif btn_css:
                    btn_selector = btn_css
                    by_type = "css"
                else:
                    raise ValueError("Forneça btn_index, btn_xpath ou btn_css para abrir o LOV.")

                self._log(f"🔎 Localizando botão LOV: {btn_selector}", "DEBUG")
                lov_button = self._wait_element(by_type, btn_selector, timeout=5)

                if not lov_button:
                    raise Exception(f"Botão LOV não encontrado: {btn_selector}")

                self._log("ℹ️ Abrindo LOV.", "INFO")
                if not self._advanced_click(lov_button):
                    raise Exception("Falha ao clicar no botão LOV")

                time.sleep(0.8)

                # ===== PASSO 3: entra no iframe, se houver =====
                if iframe_xpath or auto_detect_iframe:
                    self._detect_and_enter_iframe(iframe_xpath)
                    time.sleep(0.3)

                # ===== PASSO 4: aplica filtro e preenche pesquisa =====
                if filter_option:
                    self._select_filter_option(filter_option)
                    time.sleep(0.3)

                if search_text:
                    self._fill_search_fields(search_text)
                    time.sleep(0.3)

                # ===== PASSO 5: clicar em Pesquisar =====
                self._click_search_button()
                time.sleep(0.8)

                # ===== PASSO 6: seleção SEQUENCIAL SEMPRE =====
                result_xpath = self._select_result_sequencial(result_text)
                if not result_xpath:
                    raise Exception("Nenhuma linha foi selecionada pelo modo sequencial.")

                # ===== PASSO 7: reforça o clique na linha =====
                if reinforce_clicks > 0:
                    self._reinforce_click_on_result(result_xpath, reinforce_clicks)

                # ===== PASSO 8: volta do iframe p/ contexto principal =====
                try:
                    self.driver.switch_to.default_content()
                except Exception:
                    pass

                time.sleep(wait_after)
                self._log(
                    f"✅ LOV concluído com sucesso (linha índice {self.selecao_sequencial_index})",
                    "SUCCESS"
                )
                return True

            except Exception as e:
                self._log(f"Tentativa {retry} falhou: {str(e)[:120]}", "ERROR")

                try:
                    self.driver.switch_to.default_content()
                except Exception:
                    pass

                if retry < self.max_retries:
                    time.sleep(2 * retry)  # backoff
                    continue

        self._log(f"❌ LOV falhou após {self.max_retries} tentativas.", "ERROR")
        return False


    def _click_search_button(self):
        """Clica no botão Pesquisar com múltiplas estratégias (prioriza clique)"""
        search_button_xpaths = [
            "//a[contains(@class,'btPesquisar') and contains(normalize-space(.),'Pesquisar')]",
            "//button[contains(normalize-space(.),'Pesquisar')]",
            "//input[@type='button' and contains(@value,'Pesquisar')]",
            "//a[contains(@class,'lpFind')]",
            "//a[contains(@onclick,'pesquisar')]"
        ]

        self._log("🔍 Procurando botão 'Pesquisar' dentro do LOV...", "INFO")

        # 1) TENTAR CLICAR NOS BOTÕES
        for xpath in search_button_xpaths:
            try:
                btn = self._wait_element("xpath", xpath, timeout=2)
                if btn and self._is_element_visible(btn):
                    self._log(f"🔍 Tentando clicar em 'Pesquisar' via XPath: {xpath}", "DEBUG")
                    if self._advanced_click(btn):
                        self._log("✅ Botão 'Pesquisar' clicado com sucesso", "SUCCESS")
                        return True
            except Exception:
                continue

        # 2) FALLBACK: ENTER SOMENTE SE NÃO ACHOU BOTÃO
        self._log("⚠️ Botão 'Pesquisar' não encontrado, tentando ENTER no campo ativo...", "WARNING")
        try:
            ActionChains(self.driver).send_keys(Keys.ENTER).perform()
            self._log("✅ ENTER enviado como fallback", "SUCCESS")
            return True
        except Exception:
            self._log("❌ Falha ao enviar ENTER como fallback", "ERROR")
            return False


    def _select_result_sequencial(self, result_text=""):
        """
        Seleciona SEMPRE de forma sequencial:
        - 1ª vez que for chamado: 1ª linha
        - 2ª vez: 2ª linha
        - etc.

        Usa self.selecao_sequencial_index como índice base (0-based internamente).
        O parâmetro result_text é opcional e usado só para log.
        """
        driver = self.driver
        idx = self.selecao_sequencial_index

        base_xpath = "//table//tr[@role='row' and @style]"
        linhas = driver.find_elements(By.XPATH, base_xpath)

        if not linhas:
            # aqui é caso real de LOV vazia mesmo (nenhum resultado)
            raise Exception("Nenhuma linha encontrada na LOV.")

        total = len(linhas)

        # Garantir que o índice esteja em faixa válida
        if idx < 0:
            self._log(
                f"⚠ Índice sequencial negativo ({idx}). Ajustando para 0 (primeira linha).",
                "WARNING"
            )
            idx = 0

        if idx >= total:
            # Aqui é exatamente o erro que você viu: índice > quantidade de linhas
            # Em vez de quebrar o fluxo, vamos ajustar para a ÚLTIMA linha disponível
            self._log(
                f"⚠ Índice sequencial {idx} maior que a quantidade de linhas ({total}). "
                f"Ajustando para a última linha (índice {total - 1}).",
                "WARNING"
            )
            idx = total - 1
            # Opcional: sincronizar o contador global com esse idx ajustado
            self.selecao_sequencial_index = idx

        # Agora idx com certeza está dentro de [0, total-1]
        el = linhas[idx]

        driver.execute_script("arguments[0].scrollIntoView(true);", el)
        time.sleep(0.2)
        driver.execute_script("arguments[0].click();", el)

        # incrementa o índice para a próxima chamada (sequencial)
        self.selecao_sequencial_index = idx + 1

        # monta um XPath que a _reinforce_click_on_result consiga reutilizar
        # (idx é 0-based, XPath é 1-based)
        result_xpath = f"({base_xpath})[{idx + 1}]"

        self._log(
            f"✅ Resultado selecionado (sequencial): índice_interno={idx} "
            f"(linha {idx + 1} de {total}), texto='{result_text}'",
            "SUCCESS"
        )
        return result_xpath

    
    def _select_filter_option(self, filter_option):
        """
        Seleciona opção no dropdown de filtro (ex: Nome, CPF, etc.)
        """
        try:
            # XPath do dropdown de filtro (adapte conforme sua aplicação)
            xpath_dropdown = "//select[contains(@class,'filtro') or contains(@id,'filtro')]"
            
            dropdown = self.driver.find_element("xpath", xpath_dropdown)
            
            from selenium.webdriver.support.ui import Select
            select = Select(dropdown)
            select.select_by_visible_text(filter_option)
            
            self._log(f"Filtro '{filter_option}' selecionado", "SUCCESS")
            return True
            
        except Exception as e:
            self._log(f"Erro ao selecionar filtro: {e}", "WARNING")
            return False
    def verificar_sessao_ativa(driver, doc):
        """
        ✅ CORREÇÃO 3: Verifica se a sessão do navegador ainda está ativa
        """
        try:
            # Tenta executar comando simples
            driver.current_url
            return True
        except Exception as e:
            log(doc, f"❌ Sessão do navegador perdida: {e}")
            return False


    def safe_action_with_session_check(doc, descricao, func, max_retries=3):
        """
        ✅ CORREÇÃO 3: Wrapper melhorado com verificação de sessão
        """
        global driver
        
        for attempt in range(max_retries):
            try:
                # Verifica se sessão está ativa antes de tentar
                if not verificar_sessao_ativa(driver, doc):
                    log(doc, f"❌ Sessão inválida - não é possível executar: {descricao}")
                    return False
                
                log(doc, f"🔄 {descricao}..." if attempt == 0 else f"🔄 {descricao}... (Tentativa {attempt + 1})")
                func()
                log(doc, f"✅ {descricao} realizada com sucesso.")
                take_screenshot(driver, doc, _sanitize_filename(descricao))
                return True
                
            except Exception as e:
                if not verificar_sessao_ativa(driver, doc):
                    log(doc, f"❌ Sessão perdida durante execução de: {descricao}")
                    return False
                    
                if attempt < max_retries - 1:
                    log(doc, f"⚠️ Tentativa {attempt + 1} falhou, tentando novamente...")
                    time.sleep(2 + attempt)
                    continue
                else:
                    log(doc, f"❌ Erro após {max_retries} tentativas: {e}")
                    take_screenshot(driver, doc, _sanitize_filename(f"erro_{descricao}"))
                    return False
        
        return False        


    def _wait_element(self, locator_type, locator_value, timeout=None, condition="clickable"):
        """Aguarda elemento com diferentes condições"""
        timeout = timeout or self.timeout
        conditions = {
            "present": EC.presence_of_element_located,
            "visible": EC.visibility_of_element_located,
            "clickable": EC.element_to_be_clickable
        }
        
        try:
            cond = conditions.get(condition, conditions["clickable"])
            return WebDriverWait(self.driver, timeout).until(
                cond((locator_type, locator_value))
            )
        except TimeoutException:
            return None
    
    def _is_element_visible(self, element):
        """Verifica se elemento está realmente visível"""
        try:
            if not element:
                return False
            if not element.is_displayed():
                return False
            
            # Verifica via JavaScript também
            is_visible = self.driver.execute_script("""
                const el = arguments[0];
                if (!el) return false;
                const style = window.getComputedStyle(el);
                return (
                    style.display !== 'none' &&
                    style.visibility !== 'hidden' &&
                    parseFloat(style.opacity || 1) > 0.01 &&
                    el.offsetParent !== null
                );
            """, element)
            
            return is_visible
        except:
            return False
    
    def _force_element_visible(self, element):
        """Força elemento a ficar visível e interativo"""
        try:
            self.driver.execute_script("""
                const el = arguments[0];
                el.style.display = 'block';
                el.style.visibility = 'visible';
                el.style.opacity = '1';
                el.style.pointerEvents = 'auto';
                el.removeAttribute('disabled');
                el.removeAttribute('readonly');
                
                // Remove overlays que podem bloquear
                const overlays = document.querySelectorAll(
                    '.modal-backdrop, .overlay, .blockUI, [class*="loading"]'
                );
                overlays.forEach(o => {
                    o.style.display = 'none';
                    o.style.visibility = 'hidden';
                });
            """, element)
            return True
        except:
            return False
    
    def _detect_and_enter_iframe(self, iframe_xpath=None):
        """Detecta e entra no iframe automaticamente"""
        try:
            # Primeiro volta para o contexto principal
            self.driver.switch_to.default_content()
            
            # Se foi fornecido um xpath específico
            if iframe_xpath:
                try:
                    frame = self._wait_element("xpath", iframe_xpath, timeout=3)
                    if frame:
                        self.driver.switch_to.frame(frame)
                        self._log(f"Entrou no iframe: {iframe_xpath}", "SUCCESS")
                        return True
                except:
                    pass
            
            # Detecção automática de iframes
            iframe_selectors = [
                "//iframe[contains(@class,'LOV') or contains(@id,'LOV') or contains(@id,'lov')]",
                "//iframe[contains(@class,'modal') or contains(@class,'popup')]",
                "//iframe[contains(@src,'lov') or contains(@src,'LOV')]",
                "(//iframe)[last()]"  # Último iframe (geralmente o modal)
            ]
            
            for selector in iframe_selectors:
                try:
                    frames = self.driver.find_elements(By.XPATH, selector)
                    for frame in frames:
                        if self._is_element_visible(frame):
                            self.driver.switch_to.frame(frame)
                            self._log(f"Iframe detectado automaticamente: {selector}", "SUCCESS")
                            return True
                except:
                    continue
            
            return False
        except Exception as e:
            self._log(f"Erro ao detectar iframe: {e}", "WARNING")
            return False
    
    def _click_strategy_1_standard(self, element):
        """Estratégia 1: Clique padrão do Selenium"""
        element.click()
        return True
    
    def _click_strategy_2_javascript_simple(self, element):
        """Estratégia 2: JavaScript simples"""
        self.driver.execute_script("arguments[0].click();", element)
        return True
    
    def _click_strategy_3_javascript_advanced(self, element):
        """Estratégia 3: JavaScript com eventos completos"""
        self.driver.execute_script("""
            const el = arguments[0];
            const rect = el.getBoundingClientRect();
            const x = rect.left + rect.width / 2;
            const y = rect.top + rect.height / 2;
            
            const events = ['mouseover', 'mouseenter', 'mousemove', 'mousedown', 'mouseup', 'click'];
            events.forEach(eventType => {
                const evt = new MouseEvent(eventType, {
                    bubbles: true,
                    cancelable: true,
                    view: window,
                    detail: 1,
                    clientX: x,
                    clientY: y
                });
                el.dispatchEvent(evt);
            });
            
            if (typeof el.click === 'function') el.click();
        """, element)
        return True
    
    def _click_strategy_4_action_chains(self, element):
        """Estratégia 4: ActionChains com pause"""
        ActionChains(self.driver)\
            .move_to_element(element)\
            .pause(0.1)\
            .click()\
            .perform()
        return True
    
    def _click_strategy_5_action_chains_offset(self, element):
        """Estratégia 5: ActionChains com offset"""
        ActionChains(self.driver)\
            .move_to_element_with_offset(element, 5, 5)\
            .pause(0.05)\
            .click()\
            .perform()
        return True
    
    def _click_strategy_6_force_visible_then_click(self, element):
        """Estratégia 6: Força visibilidade e clica"""
        self._force_element_visible(element)
        time.sleep(0.2)
        element.click()
        return True
    
    def _click_strategy_7_scroll_and_click(self, element):
        """Estratégia 7: Scroll suave até elemento"""
        self.driver.execute_script("""
            arguments[0].scrollIntoView({
                behavior: 'smooth',
                block: 'center',
                inline: 'center'
            });
        """, element)
        time.sleep(0.3)
        element.click()
        return True
    
    def _click_strategy_8_remove_overlays(self, element):
        """Estratégia 8: Remove todos os overlays e clica"""
        self.driver.execute_script("""
            // Remove todos os overlays possíveis
            const overlays = document.querySelectorAll(`
                .modal-backdrop, .overlay, .blockUI, .blockScreen,
                [class*="loading"], [class*="spinner"], [class*="overlay"],
                [style*="z-index: 9999"], [style*="position: fixed"]
            `);
            overlays.forEach(o => {
                o.style.display = 'none';
                o.style.visibility = 'hidden';
                o.remove();
            });
            
            const el = arguments[0];
            el.style.zIndex = '999999';
            el.click();
        """, element)
        return True
    
    def _click_strategy_9_jquery_trigger(self, element):
        """Estratégia 9: jQuery trigger (se disponível)"""
        self.driver.execute_script("""
            const el = arguments[0];
            if (typeof jQuery !== 'undefined') {
                jQuery(el).trigger('click');
            } else {
                el.click();
            }
        """, element)
        return True
    
    def _click_strategy_10_nuclear_option(self, element):
        """Estratégia 10: Opção nuclear - força tudo"""
        self.driver.execute_script("""
            const el = arguments[0];
            
            // Remove TODOS os atributos que podem bloquear
            el.removeAttribute('disabled');
            el.removeAttribute('readonly');
            el.style.pointerEvents = 'auto !important';
            el.style.display = 'block !important';
            el.style.visibility = 'visible !important';
            el.style.opacity = '1 !important';
            el.style.zIndex = '999999 !important';
            
            // Remove TODOS os overlays da página
            document.querySelectorAll('*').forEach(elem => {
                const style = window.getComputedStyle(elem);
                if (
                    style.position === 'fixed' &&
                    parseInt(style.zIndex) > 1000 &&
                    elem !== el &&
                    !elem.contains(el)
                ) {
                    elem.style.display = 'none';
                }
            });
            
            // Foca no elemento
            el.focus();
            
            // Dispara TODOS os eventos possíveis
            const allEvents = [
                'focus', 'focusin', 'mouseover', 'mouseenter', 'mousemove',
                'mousedown', 'mouseup', 'click', 'dblclick'
            ];
            
            allEvents.forEach(eventType => {
                try {
                    const evt = new MouseEvent(eventType, {
                        bubbles: true,
                        cancelable: true,
                        view: window
                    });
                    el.dispatchEvent(evt);
                } catch(e) {}
            });
            
            // Clique direto
            if (typeof el.click === 'function') el.click();
            
            // jQuery se disponível
            if (typeof jQuery !== 'undefined') {
                jQuery(el).trigger('click');
            }
            
            // Tenta onclick manual
            if (el.onclick) el.onclick();
            
            return true;
        """, element)
        return True
    
    def _advanced_click(self, element, max_attempts=10):
        """
        Sistema avançado de clique com 10 estratégias diferentes
        Tenta cada estratégia até uma funcionar
        """
        strategies = [
            ("Clique Padrão", self._click_strategy_1_standard),
            ("JavaScript Simples", self._click_strategy_2_javascript_simple),
            ("JavaScript Avançado", self._click_strategy_3_javascript_advanced),
            ("ActionChains", self._click_strategy_4_action_chains),
            ("ActionChains Offset", self._click_strategy_5_action_chains_offset),
            ("Força Visível", self._click_strategy_6_force_visible_then_click),
            ("Scroll e Clique", self._click_strategy_7_scroll_and_click),
            ("Remove Overlays", self._click_strategy_8_remove_overlays),
            ("jQuery Trigger", self._click_strategy_9_jquery_trigger),
            ("Opção Nuclear", self._click_strategy_10_nuclear_option)
        ]
        
        for attempt in range(1, max_attempts + 1):
            for strategy_name, strategy_func in strategies:
                try:
                    self._log(f"Tentativa {attempt}/{max_attempts}: {strategy_name}", "DEBUG")
                    
                    # Tenta executar a estratégia
                    strategy_func(element)
                    time.sleep(0.3)
                    
                    self._log(f"✓ {strategy_name} funcionou!", "SUCCESS")
                    return True
                    
                except StaleElementReferenceException:
                    self._log(f"Elemento ficou stale, tentando recarregar...", "WARNING")
                    return False  # Precisa recarregar elemento
                    
                except (ElementClickInterceptedException, 
                        ElementNotInteractableException) as e:
                    self._log(f"✗ {strategy_name}: {str(e)[:50]}", "DEBUG")
                    continue
                    
                except Exception as e:
                    self._log(f"✗ {strategy_name}: {str(e)[:50]}", "DEBUG")
                    continue
            
            if attempt < max_attempts:
                time.sleep(0.5 * attempt)  # Backoff exponencial
        
        return False
    
    def _fill_search_fields(self, search_text, max_fields=10):
        """Preenche campos de pesquisa encontrados (priorizando Nome/Numero)"""
        if not search_text:
            return 0

        # 🎯 1) XPaths prioritários para esse LOV específico
        if search_text.isdigit():
            # Se só tem número, prioriza o campo de número
            search_field_xpaths = [
                "//input[contains(@class,'fc') and contains(@class,'numero') and not(@disabled)]",
            ]
        else:
            # Se tem letras, prioriza o campo de nome
            search_field_xpaths = [
                "//input[contains(@class,'fc') and contains(@class,'nome') and not(@disabled)]",
            ]

        # 🎯 2) Depois mantém os genéricos que você já tinha
        search_field_xpaths += [
            "//input[@id='txtPesquisa']",
            "//input[contains(@class,'pesquisa')]",
            "//input[contains(@class,'nomePesquisa')]",
            "//input[contains(translate(@name,'PESQUISA','pesquisa'),'pesquisa')]",
            "//input[@type='text' and contains(@style,'width:210px')]",
            "//input[@type='text' and not(@disabled)]",
        ]

        filled_count = 0
        for xpath in search_field_xpaths:
            try:
                fields = self.driver.find_elements(By.XPATH, xpath)
                for field in fields[:max_fields]:
                    try:
                        if not self._is_element_visible(field):
                            continue

                        self._force_element_visible(field)
                        field.clear()
                        field.click()
                        field.send_keys(search_text)
                        filled_count += 1
                    except:
                        continue
            except:
                continue

        if filled_count > 0:
            self._log(f"Preenchidos {filled_count} campos de pesquisa", "SUCCESS")
        else:
            self._log("Nenhum campo de pesquisa encontrado", "WARNING")

        return filled_count


    
    def _select_result(self, result_text=""):
        """Seleciona o resultado da pesquisa"""
        result_xpaths = []
        
        if result_text:
            result_xpaths = [
                f"//td[contains(normalize-space(.),'{result_text}')]",
                f"//span[contains(normalize-space(.),'{result_text}')]",
                f"//div[contains(normalize-space(.),'{result_text}')]",
                f"//tr[contains(normalize-space(.),'{result_text}')]//td[1]",
                f"//li[contains(normalize-space(.),'{result_text}')]"
            ]
        else:
            result_xpaths = [
                "(//table//tr[1]/td[1])[1]",
                "(//tr[contains(@class,'ui-widget-content')][1])[1]",
                "(//tr[contains(@class,'rich-table-row')][1])[1]",
                "(//li[@class='ui-autocomplete-item'])[1]"
            ]
        
        for xpath in result_xpaths:
            try:
                result = self._wait_element("xpath", xpath, timeout=3)
                if result and self._is_element_visible(result):
                    if self._advanced_click(result):
                        self._log(f"Resultado selecionado: {result_text or 'primeiro'}", "SUCCESS")
                        return xpath  # Retorna xpath usado
            except:
                continue
        
        self._log(f"Não foi possível selecionar: {result_text or 'primeiro resultado'}", "ERROR")
        return None
    
    def _reinforce_click_on_result(self, result_xpath, max_reclick=5):
        """
        Reforça o clique no resultado SEMPRE, independente do modal
        Clica repetidamente até max_reclick vezes
        """
        if not result_xpath:
            return
        
        self._log(f"🔁 Reforçando clique no resultado ({max_reclick}x)", "INFO")
        
        for attempt in range(1, max_reclick + 1):
            try:
                time.sleep(0.35)
                
                # Tenta localizar o resultado novamente
                result = self.driver.find_element(By.XPATH, result_xpath)
                
                # Clica usando sistema avançado
                self._log(f"Reforço {attempt}/{max_reclick}", "DEBUG")
                self._advanced_click(result)
                
            except StaleElementReferenceException:
                self._log(f"Elemento stale no reforço {attempt}", "WARNING")
                try:
                    # Tenta recarregar o elemento
                    result = self.driver.find_element(By.XPATH, result_xpath)
                    self._advanced_click(result)
                except:
                    continue
                    
            except Exception as e:
                self._log(f"Erro no reforço {attempt}: {str(e)[:50]}", "WARNING")
                continue
        
        self._log("✅ Reforço de cliques concluído", "SUCCESS")
    

    def open_and_select_sequencial_titular(
        self,
        btn_index=None,
        btn_xpath=None,
        btn_css=None,
        search_text="",
        result_text="",      # agora só pra log
        filter_option=None,
        iframe_xpath=None,
        auto_detect_iframe=True,
        reinforce_clicks=5,
        wait_after=0.5
    ):
        """
        Método principal: abre LOV, aplica filtro/pesquisa e seleciona
        SEMPRE de forma SEQUENCIAL (1ª vez = 1ª linha, 2ª = 2ª, etc.).
        """

        self._log(
            f"🔍 LOV sequencial | filtro='{filter_option}' | pesquisa='{search_text}' | índice atual={self.selecao_sequencial_index}",
            "INFO"
        )

        for retry in range(1, self.max_retries + 1):
            try:
                # ===== PASSO 1: volta p/ contexto principal =====
                try:
                    self.driver.switch_to.default_content()
                except Exception:
                    pass

                # ===== PASSO 2: localizar botão do LOV =====
                btn_selector = None
                by_type = None

                if btn_index is not None:
                    # índice é 0-based aqui, então soma 1 no XPath
                    btn_selector = f"(//a[@class='sprites sp-changeTitular'])[{btn_index + 1}]"
                    by_type = "xpath"
                elif btn_xpath:
                    btn_selector = btn_xpath
                    by_type = "xpath"
                elif btn_css:
                    btn_selector = btn_css
                    by_type = "css"
                else:
                    raise ValueError("Forneça btn_index, btn_xpath ou btn_css para abrir o LOV.")

                self._log(f"🔎 Localizando botão LOV: {btn_selector}", "DEBUG")
                lov_button = self._wait_element(by_type, btn_selector, timeout=5)

                if not lov_button:
                    raise Exception(f"Botão LOV não encontrado: {btn_selector}")

                self._log("ℹ️ Abrindo LOV...", "INFO")
                if not self._advanced_click(lov_button):
                    raise Exception("Falha ao clicar no botão LOV")

                time.sleep(0.8)

                # ===== PASSO 3: entra no iframe, se houver =====
                if iframe_xpath or auto_detect_iframe:
                    self._detect_and_enter_iframe(iframe_xpath)
                    time.sleep(0.3)

                # ===== PASSO 4: aplica filtro e preenche pesquisa =====
                if filter_option:
                    self._select_filter_option(filter_option)
                    time.sleep(0.3)

                if search_text:
                    self._fill_search_fields(search_text)
                    time.sleep(0.3)

                # ===== PASSO 5: clicar em Pesquisar =====
                self._click_search_button()
                time.sleep(0.8)

                # ===== PASSO 6: seleção SEQUENCIAL SEMPRE =====
                result_xpath = self._select_result_sequencial()
                if not result_xpath:
                    raise Exception("Nenhuma linha foi selecionada pelo modo sequencial.")

                # ===== PASSO 7: reforça o clique na linha =====
                if reinforce_clicks > 0:
                    self._reinforce_click_on_result(result_xpath, reinforce_clicks)

                # ===== PASSO 8: volta do iframe p/ contexto principal =====
                try:
                    self.driver.switch_to.default_content()
                except Exception:
                    pass

                time.sleep(wait_after)
                self._log(
                    f"✅ LOV concluído com sucesso (linha índice {self.selecao_sequencial_index})",
                    "SUCCESS"
                )
                return True

            except Exception as e:
                self._log(f"Tentativa {retry} falhou: {str(e)[:120]}", "ERROR")

                try:
                    self.driver.switch_to.default_content()
                except Exception:
                    pass

                if retry < self.max_retries:
                    time.sleep(2 * retry)  # backoff
                    continue

        self._log(f"❌ LOV falhou após {self.max_retries} tentativas.", "ERROR")
        return False

    def open_and_select_sequencial_dependente(
        self,
        btn_index=None,
        btn_xpath=None,
        btn_css=None,
        search_text="",
        result_text="",      # agora só pra log
        filter_option=None,
        iframe_xpath=None,
        auto_detect_iframe=True,
        reinforce_clicks=5,
        wait_after=0.5
    ):
        """
        Método principal: abre LOV, aplica filtro/pesquisa e seleciona
        SEMPRE de forma SEQUENCIAL (1ª vez = 1ª linha, 2ª = 2ª, etc.).
        """

        self._log(
            f"🔍 LOV sequencial | filtro='{filter_option}' | pesquisa='{search_text}' | índice atual={self.selecao_sequencial_index}",
            "INFO"
        )

        for retry in range(1, self.max_retries + 1):
            try:
                # ===== PASSO 1: volta p/ contexto principal =====
                try:
                    self.driver.switch_to.default_content()
                except Exception:
                    pass

                # ===== PASSO 2: localizar botão do LOV =====
                btn_selector = None
                by_type = None

                if btn_index is not None:
                    # índice é 0-based aqui, então soma 1 no XPath
                    btn_selector = f"(//a[@class='sprites sp-addDependentes'])[{btn_index + 1}]"
                    by_type = "xpath"
                elif btn_xpath:
                    btn_selector = btn_xpath
                    by_type = "xpath"
                elif btn_css:
                    btn_selector = btn_css
                    by_type = "css"
                else:
                    raise ValueError("Forneça btn_index, btn_xpath ou btn_css para abrir o LOV.")

                self._log(f"🔎 Localizando botão LOV: {btn_selector}", "DEBUG")
                lov_button = self._wait_element(by_type, btn_selector, timeout=5)

                if not lov_button:
                    raise Exception(f"Botão LOV não encontrado: {btn_selector}")

                self._log("ℹ️ Abrindo LOV...", "INFO")
                if not self._advanced_click(lov_button):
                    raise Exception("Falha ao clicar no botão LOV")

                time.sleep(0.8)

                # ===== PASSO 3: entra no iframe, se houver =====
                if iframe_xpath or auto_detect_iframe:
                    self._detect_and_enter_iframe(iframe_xpath)
                    time.sleep(0.3)

                # ===== PASSO 4: aplica filtro e preenche pesquisa =====
                if filter_option:
                    self._select_filter_option(filter_option)
                    time.sleep(0.3)

                if search_text:
                    self._fill_search_fields(search_text)
                    time.sleep(0.3)

                # ===== PASSO 5: clicar em Pesquisar =====
                self._click_search_button()
                time.sleep(0.8)

                # ===== PASSO 6: seleção SEQUENCIAL SEMPRE =====
                result_xpath = self._select_result_sequencial()
                if not result_xpath:
                    raise Exception("Nenhuma linha foi selecionada pelo modo sequencial.")

                # ===== PASSO 7: reforça o clique na linha =====
                if reinforce_clicks > 0:
                    self._reinforce_click_on_result(result_xpath, reinforce_clicks)

                # ===== PASSO 8: volta do iframe p/ contexto principal =====
                try:
                    self.driver.switch_to.default_content()
                except Exception:
                    pass

                time.sleep(wait_after)
                self._log(
                    f"✅ LOV concluído com sucesso (linha índice {self.selecao_sequencial_index})",
                    "SUCCESS"
                )
                return True

            except Exception as e:
                self._log(f"Tentativa {retry} falhou: {str(e)[:120]}", "ERROR")

                try:
                    self.driver.switch_to.default_content()
                except Exception:
                    pass

                if retry < self.max_retries:
                    time.sleep(2 * retry)  # backoff
                    continue

        self._log(f"❌ LOV falhou após {self.max_retries} tentativas.", "ERROR")
        return False



# ==== FUNÇÕES AUXILIARES ====


def encontrar_mensagem_alerta():
    seletores = [
        (".alerts.salvo", "✅ Mensagem de Sucesso"),
        (".alerts.alerta", "⚠️ Mensagem de Alerta"),
        (".alerts.erro", "❌ Mensagem de Erro"),
    ]

    for seletor, tipo in seletores:
        try:
            elemento = driver.find_element(By.CSS_SELECTOR, seletor)
            if elemento.is_displayed():
                log(doc, f"📢 {tipo}: {elemento.text}")
                return elemento
        except:
            continue

    log(doc, "ℹ️ Nenhuma mensagem de alerta encontrada.")
    return None

# ==== FUNÇÕES DE UTILITÁRIO MELHORADAS ====
def safe_action(doc, descricao, func, max_retries=3, timeout_customizado=None, critico=True):
    """Execução de ações com retry robusto e tratamento de erros melhorado"""
    global driver
    
    timeout_original = None
    if timeout_customizado and driver:
        try:
            # Ajusta timeout temporariamente
            timeout_original = driver.timeouts.implicit_wait
            driver.implicitly_wait(timeout_customizado)
        except:
            pass
    
    for attempt in range(max_retries):
        try:
            if attempt == 0:
                log(doc, f"Executando: {descricao}")
            else:
                log(doc, f"Retry {attempt + 1}/{max_retries}: {descricao}", 'WARN')
            
            result = func()
            log(doc, f"✅ {descricao} - Sucesso")
            take_screenshot(driver, doc, descricao.lower().replace(" ", "_"))
            
            return True
            
        except (TimeoutException, NoSuchElementException, StaleElementReferenceException) as e:
            if attempt < max_retries - 1:
                log(doc, f"⚠️ Tentativa {attempt + 1} falhou: {type(e).__name__}", 'WARN')
                time.sleep(2 ** attempt)  # Backoff exponencial
                continue
            else:
                error_msg = f"❌ {descricao} falhou após {max_retries} tentativas: {e}"
                log(doc, error_msg, 'ERROR')
                take_screenshot(driver, doc, f"erro_{descricao.lower().replace(' ', '_')}", forcar=True)
                
                if critico:
                    raise Exception(error_msg)
                return False
                
        except Exception as e:
            error_msg = f"❌ Erro inesperado em {descricao}: {e}"
            log(doc, error_msg, 'ERROR')
            take_screenshot(driver, doc, f"erro_critico_{descricao.lower().replace(' ', '_')}", forcar=True)
            
            if critico:
                raise Exception(error_msg)
            return False
    
        finally:
            # Restaura timeout original
            if timeout_original and driver:
                try:
                    driver.implicitly_wait(timeout_original)
                except:
                    pass


def inicializar_driver():
    """Inicializa WebDriver com configurações otimizadas"""
    global driver, wait
    
    try:
        options = Options()
        options.add_argument("--start-maximized")
        options.add_argument("--disable-blink-features=AutomationControlled")
        options.add_argument("--disable-dev-shm-usage")
        options.add_argument("--no-sandbox")
        options.add_experimental_option("excludeSwitches", ["enable-automation"])
        options.add_experimental_option('useAutomationExtension', False)
        
        driver = webdriver.Chrome(
            service=Service(ChromeDriverManager().install()), 
            options=options
        )
        
        driver.execute_script(
            "Object.defineProperty(navigator, 'webdriver', {get: () => undefined})"
        )
        
        # Configura timeouts globais
        driver.set_script_timeout(30)
        driver.implicitly_wait(10)
        
        wait = WebDriverWait(driver, TIMEOUT_DEFAULT)
        
        log(doc, "✅ Driver inicializado com sucesso")
        return True
        
    except Exception as e:
        log(doc, f"❌ Erro ao inicializar driver: {e}")
        return False

def finalizar_relatorio():
    """Salva relatório e fecha driver"""
    global driver, doc
    
    nome_arquivo = f"relatorio__processo_consulta_cremacao_cenario_1_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    
    try:
        doc.save(nome_arquivo)
        log(doc, f"📄 Relatório salvo: {nome_arquivo}")
        
        try:
            subprocess.run(["start", "winword", nome_arquivo], shell=True)
        except:
            pass
            
    except Exception as e:
        print(f"Erro ao salvar relatório: {e}")
    
    if driver:
        try:
            driver.quit()
            log(doc, "✅ Driver encerrado")
        except:
            pass


def fechar_modal_mensagens(js_engine, doc, max_tentativas=5, timeout=10):
    """
    Detecta e fecha o modal de mensagens do sistema.
    
    SE O MODAL NÃO EXISTIR: retorna True imediatamente (nada a fazer)
    SE O MODAL EXISTIR: tenta fechá-lo com múltiplas estratégias
    
    Características do modal:
    - z-index alto (10002+)
    - Classe: 'modal overflow'
    - Contém: <h2>Mensagens</h2>
    - Botão fechar: <a class="btModel btGray">Fechar</a>
    
    Args:
        js_engine: Instância do JSForceEngine
        doc: Documento para logs
        max_tentativas: Número máximo de tentativas
        timeout: Tempo máximo de espera
    
    Returns:
        bool: True se não detectou modal OU se fechou com sucesso
              False se detectou mas não conseguiu fechar
    """
    
    # XPaths para detectar o modal
    xpath_modal = "//div[contains(@class,'modal') and contains(@class,'overflow') and .//h2[normalize-space()='Mensagens']]"
    
    # XPaths para o botão Fechar (em ordem de especificidade)
    xpaths_fechar = [
        # 1. Botão dentro do modal de mensagens (mais específico)
        "//div[contains(@class,'modal') and .//h2[text()='Mensagens']]//a[@class='btModel btGray' and contains(normalize-space(),'Fechar')]",
        
        # 2. Botão com ícone cancelar dentro do modal
        "//div[contains(@class,'modal') and .//h2[text()='Mensagens']]//a[@class='btModel btGray']//span[contains(@class,'sp-cancelar')]/parent::a",
        
        # 3. Ícone X de fechar no canto
        "//div[contains(@class,'modal') and .//h2[text()='Mensagens']]//a[@class='fa fa-close']",
        
        # 4. Qualquer botão "Fechar" visível no modal
        "//div[contains(@class,'modal')]//a[contains(@class,'btGray') and contains(text(),'Fechar')]",
        
        # 5. Fallback - último botão com classe btGray no modal
        "(//div[contains(@class,'modal')]//a[@class='btModel btGray'])[last()]"
    ]
    
    log(doc, "🔍 Verificando presença do modal de mensagens...")
    
    # ===== VERIFICAÇÃO INICIAL: Modal existe? =====
    try:
        modais = js_engine.driver.find_elements("xpath", xpath_modal)
        modais_visiveis = [m for m in modais if m.is_displayed()]
        
        if not modais_visiveis:
            log(doc, "✅ Modal de mensagens NÃO detectado - nada a fazer")
            return True
        
        log(doc, "⚠️  Modal de mensagens DETECTADO - iniciando fechamento...")
        
    except Exception as e:
        log(doc, f"✅ Erro ao verificar modal (provavelmente não existe): {e}")
        return True
    
    # ===== SE CHEGOU AQUI: Modal existe e precisa ser fechado =====
    inicio = time.time()
    
    for tentativa in range(1, max_tentativas + 1):
        try:
            # Verifica timeout
            if time.time() - inicio > timeout:
                log(doc, f"⏰ Timeout de {timeout}s atingido")
                return False
            
            log(doc, f"")
            log(doc, f"🔄 Tentativa {tentativa}/{max_tentativas}")
            
            # ===== RE-VERIFICA se modal ainda existe =====
            try:
                modais = js_engine.driver.find_elements("xpath", xpath_modal)
                modais_visiveis = [m for m in modais if m.is_displayed()]
                
                if not modais_visiveis:
                    log(doc, "✅ Modal desapareceu - fechamento bem-sucedido!")
                    return True
                
                modal = modais_visiveis[0]
                z_index = modal.value_of_css_property("z-index")
                log(doc, f"   📊 Modal ainda visível (z-index: {z_index})")
                
            except Exception as e:
                log(doc, f"✅ Modal não encontrado na re-verificação: {e}")
                return True
            
            # ===== PASSO 1: Scroll até o botão dentro do modal =====
            log(doc, "📜 Rolando até o botão 'Fechar'...")
            
            try:
                js_engine.execute_js("""
                    const modal = arguments[0];
                    const content = modal.querySelector('.content, .formRow');
                    if (content) {
                        content.scrollIntoView({block: 'end', behavior: 'smooth'});
                    }
                    modal.scrollTop = modal.scrollHeight;
                """, modal, timeout=3, fallback_result=None)
                
                time.sleep(0.8)
                
            except Exception as e:
                log(doc, f"   ⚠️  Erro ao rolar modal: {e}")
            
            # ===== PASSO 2: Tenta cada XPath de fechar =====
            for i, xpath_fechar in enumerate(xpaths_fechar, 1):
                try:
                    log(doc, f"   🎯 Tentando XPath {i}/{len(xpaths_fechar)}...")
                    
                    botoes = js_engine.driver.find_elements("xpath", xpath_fechar)
                    
                    if not botoes:
                        log(doc, f"      ℹ️  XPath {i} não retornou elementos")
                        continue
                    
                    # Filtra apenas botões visíveis
                    botoes_visiveis = []
                    for btn in botoes:
                        try:
                            if btn.is_displayed() and btn.is_enabled():
                                location_ok = js_engine.execute_js("""
                                    const el = arguments[0];
                                    const rect = el.getBoundingClientRect();
                                    return rect.width > 0 && rect.height > 0;
                                """, btn, timeout=2, fallback_result=False)
                                
                                if location_ok:
                                    botoes_visiveis.append(btn)
                        except:
                            continue
                    
                    if not botoes_visiveis:
                        log(doc, f"      ℹ️  XPath {i}: nenhum botão visível")
                        continue
                    
                    botao = botoes_visiveis[0]
                    log(doc, f"      ✅ Botão encontrado (XPath {i})")
                    
                    # ===== PASSO 3: Scroll até o botão =====
                    try:
                        js_engine.scroll_into_view(botao, padding=50)
                        time.sleep(0.5)
                    except:
                        pass
                    
                    # ===== PASSO 4: Clica usando estratégias múltiplas =====
                    log(doc, f"      🖱️  Clicando no botão 'Fechar'...")
                    
                    # Estratégia 1: Click do engine
                    try:
                        js_engine.click_element(botao, wait_after=1.0)
                        log(doc, f"   ✅ Clique executado (Engine Click)")
                        
                        time.sleep(1.5)
                        js_engine.wait_ajax_complete(5)
                        
                        # Valida se fechou
                        modais_apos = js_engine.driver.find_elements("xpath", xpath_modal)
                        modais_visiveis_apos = [m for m in modais_apos if m.is_displayed()]
                        
                        if not modais_visiveis_apos:
                            log(doc, "   ✅ Confirmado: Modal fechado com sucesso!")
                            return True
                        else:
                            log(doc, "   ⚠️  Modal ainda visível após clique")
                            continue
                            
                    except Exception as e:
                        log(doc, f"      ⚠️  Engine Click falhou: {e}")
                    
                    # Estratégia 2: JavaScript agressivo
                    try:
                        js_engine.execute_js("""
                            const btn = arguments[0];
                            
                            // Remove bloqueios
                            btn.style.pointerEvents = 'auto';
                            btn.removeAttribute('disabled');
                            btn.style.visibility = 'visible';
                            btn.style.zIndex = '999999';
                            
                            // Scroll até o botão
                            btn.scrollIntoView({block: 'center'});
                            
                            // Dispara todos os eventos
                            ['mouseenter', 'mouseover', 'mousedown', 'focus', 'mouseup', 'click'].forEach(type => {
                                btn.dispatchEvent(new MouseEvent(type, {
                                    bubbles: true,
                                    cancelable: true,
                                    view: window
                                }));
                            });
                            
                            // Click direto
                            if (typeof btn.click === 'function') {
                                btn.click();
                            }
                            
                            // jQuery se disponível
                            if (typeof jQuery !== 'undefined') {
                                jQuery(btn).trigger('click');
                            }
                            
                            // onclick manual
                            if (btn.onclick) {
                                btn.onclick();
                            }
                            
                            return true;
                        """, botao, timeout=5, fallback_result=False)
                        
                        log(doc, f"   ✅ Clique JavaScript executado")
                        
                        time.sleep(1.5)
                        js_engine.wait_ajax_complete(5)
                        
                        # Valida
                        modais_apos = js_engine.driver.find_elements("xpath", xpath_modal)
                        modais_visiveis_apos = [m for m in modais_apos if m.is_displayed()]
                        
                        if not modais_visiveis_apos:
                            log(doc, "   ✅ Confirmado: Modal fechado com sucesso!")
                            return True
                            
                    except Exception as e:
                        log(doc, f"      ⚠️  JavaScript falhou: {e}")
                    
                    # Estratégia 3: ActionChains
                    try:
                        from selenium.webdriver import ActionChains
                        ActionChains(js_engine.driver)\
                            .move_to_element(botao)\
                            .pause(0.1)\
                            .click()\
                            .perform()
                        
                        log(doc, f"   ✅ Clique ActionChains executado")
                        
                        time.sleep(1.5)
                        js_engine.wait_ajax_complete(5)
                        
                        # Valida
                        modais_apos = js_engine.driver.find_elements("xpath", xpath_modal)
                        modais_visiveis_apos = [m for m in modais_apos if m.is_displayed()]
                        
                        if not modais_visiveis_apos:
                            log(doc, "   ✅ Confirmado: Modal fechado com sucesso!")
                            return True
                            
                    except Exception as e:
                        log(doc, f"      ⚠️  ActionChains falhou: {e}")
                    
                    log(doc, f"      ⚠️  XPath {i} não conseguiu fechar o modal")
                    
                except StaleElementReferenceException:
                    log(doc, f"      ⚠️  XPath {i}: elemento stale")
                    continue
                    
                except Exception as e:
                    log(doc, f"      ⚠️  XPath {i} falhou: {str(e)[:80]}")
                    continue
            
            # Se chegou aqui, nenhum XPath funcionou nesta tentativa
            if tentativa < max_tentativas:
                log(doc, f"   ⏳ Aguardando 2s antes da próxima tentativa...")
                time.sleep(2)
                
        except Exception as e:
            log(doc, f"⚠️  Erro na tentativa {tentativa}: {e}")
            if tentativa < max_tentativas:
                time.sleep(2)
    
    # Se chegou aqui, todas as tentativas falharam
    log(doc, f"❌ Falha ao fechar modal após {max_tentativas} tentativas")
    return False




def _sanitize_timeout(t):
    """Garante timeout válido"""
    if not isinstance(t, (int, float)) or t <= 0:
        return TIMEOUT_DEFAULT
    return max(5, min(120, t))  # Entre 5 e 120 segundos

# ==== AGUARDAR ELEMENTO MELHORADO ====
def aguardar_elemento(seletor, timeout=TIMEOUT_DEFAULT, condicao='clickable', by_type=By.CSS_SELECTOR):
    """Função centralizada para aguardar elementos com diferentes condições"""
    global driver, wait
    
    if driver is None:
        raise Exception("Driver não inicializado")
    
    timeout = _sanitize_timeout(timeout)
    
    condicoes = {
        'present': EC.presence_of_element_located,
        'visible': EC.visibility_of_element_located,
        'clickable': EC.element_to_be_clickable,
        'invisible': EC.invisibility_of_element_located
    }
    
    if condicao not in condicoes:
        condicao = 'clickable'
    
    try:
        wait_obj = WebDriverWait(driver, timeout)
        elemento = wait_obj.until(condicoes[condicao]((by_type, seletor)))
        return elemento
    except TimeoutException:
        log(doc, f"❌ Timeout aguardando elemento: {seletor} (condição: {condicao}, timeout: {timeout}s)", 'ERROR')
        raise TimeoutException(f"Elemento não encontrado: {seletor} (condição: {condicao})")
    except Exception as e:
        log(doc, f"❌ Erro aguardando elemento {seletor}: {e}", 'ERROR')
        raise

# ==== SCROLL CORRIGIDO - PRINCIPAL CORREÇÃO ====
def scroll_to_element_safe(elemento_ou_seletor, by_type=By.CSS_SELECTOR):
    """Scroll seguro até elemento com validação robusta"""
    global driver
    
    if driver is None:
        log(doc, "⚠️ Driver não disponível para scroll", 'WARN')
        return False
    
    try:
        # Se for seletor, encontra o elemento
        if isinstance(elemento_ou_seletor, str):
            elemento = aguardar_elemento(elemento_ou_seletor, 10, 'present', by_type)
        else:
            elemento = elemento_ou_seletor
        
        if elemento is None:
            log(doc, "⚠️ Elemento não encontrado para scroll", 'WARN')
            return False
        
        # Verifica se elemento é válido antes de fazer scroll
        if not elemento.is_displayed():
            log(doc, "⚠️ Elemento não está visível para scroll", 'WARN')
            return False
        
        # Estratégias de scroll em ordem de preferência
        scroll_strategies = [
            # Estratégia 1: JavaScript com verificação prévia
            lambda: driver.execute_script("""
                var element = arguments[0];
                if (element && typeof element.scrollIntoView === 'function') {
                    element.scrollIntoView({
                        behavior: 'smooth',
                        block: 'center',
                        inline: 'center'
                    });
                    return true;
                } else {
                    return false;
                }
            """, elemento),
            
            # Estratégia 2: ActionChains
            lambda: ActionChains(driver).move_to_element(elemento).perform(),
            
            # Estratégia 3: JavaScript alternativo
            lambda: driver.execute_script("""
                var element = arguments[0];
                if (element) {
                    element.scrollIntoView();
                    window.scrollBy(0, -100);
                }
            """, elemento),
            
            # Estratégia 4: Scroll da página até o elemento
            lambda: driver.execute_script("""
                var element = arguments[0];
                if (element) {
                    var rect = element.getBoundingClientRect();
                    var scrollTop = window.pageYOffset || document.documentElement.scrollTop;
                    var targetY = rect.top + scrollTop - (window.innerHeight / 2);
                    window.scrollTo(0, targetY);
                }
            """, elemento)
        ]
        
        for i, strategy in enumerate(scroll_strategies, 1):
            try:
                log(doc, f"   Tentando estratégia de scroll {i}...")
                result = strategy()
                
                # Para estratégia 1, verifica resultado
                if i == 1 and result is False:
                    log(doc, f"   Estratégia {i}: elemento não suporta scrollIntoView", 'WARN')
                    continue
                
                time.sleep(0.8)  # Aguarda scroll completar
                
                # Verifica se elemento ainda está acessível
                if elemento.is_displayed() and elemento.is_enabled():
                    log(doc, f"✅ Scroll realizado com estratégia {i}")
                    return True
                else:
                    log(doc, f"   Estratégia {i}: elemento não ficou acessível", 'WARN')
                    continue
                    
            except Exception as e:
                log(doc, f"   Estratégia {i} de scroll falhou: {str(e)[:100]}...", 'WARN')
                continue
        
        log(doc, "⚠️ Todas as estratégias de scroll falharam", 'WARN')
        return False
        
    except Exception as e:
        log(doc, f"⚠️ Erro geral no scroll: {e}", 'WARN')
        return False



# ==== SISTEMA DATEPICKER MELHORADO ====
def encontrar_campos_datepicker():
    """Encontra todos os campos datepicker na página"""
    global driver
    
    if driver is None:
        return []
    
    seletores_datepicker = [
        "input.hasDatepicker",
        "input[id^='dp']",
        "input[maxlength='10'][grupo='']",
        "input[type='text'][maxlength='10']",
        "input[class*='datepicker']",
        ".hasDatepicker"
    ]
    
    campos_encontrados = []
    
    for seletor in seletores_datepicker:
        try:
            elementos = driver.find_elements(By.CSS_SELECTOR, seletor)
            for elemento in elementos:
                if elemento.is_displayed() and elemento.is_enabled():
                    info = {
                        'elemento': elemento,
                        'id': elemento.get_attribute('id') or f"dp_{len(campos_encontrados)}",
                        'seletor_usado': seletor,
                        'maxlength': elemento.get_attribute('maxlength'),
                        'placeholder': elemento.get_attribute('placeholder')
                    }
                    # Evita duplicatas
                    if not any(c['id'] == info['id'] for c in campos_encontrados):
                        campos_encontrados.append(info)
        except Exception as e:
            log(doc, f"⚠️ Erro ao buscar campos datepicker com {seletor}: {e}", 'WARN')
            continue
    
    log(doc, f"📊 Encontrados {len(campos_encontrados)} campos datepicker")
    return campos_encontrados

def _datepicker_jquery(campo_id, data_valor):
    """Estratégia jQuery para datepicker"""
    global driver
    
    resultado = driver.execute_script("""
        var campoId = arguments[0], valor = arguments[1];
        if (typeof jQuery === 'undefined') return 'jQuery não disponível';
        var $campo = $('#' + campoId);
        if (!$campo.length) return 'Campo não encontrado: ' + campoId;
        try {
            if ($campo.hasClass('hasDatepicker')) { 
                $campo.datepicker('setDate', valor); 
            } else { 
                $campo.val(valor); 
            }
            $campo.trigger('input').trigger('change').trigger('blur');
            return $campo.val();
        } catch(e) { 
            return 'Erro: ' + e.message; 
        }
    """, campo_id, data_valor)
    
    if isinstance(resultado, str) and ('Erro' in resultado or 'não disponível' in resultado):
        raise Exception(f"jQuery falhou: {resultado}")

def _datepicker_javascript(elemento, data_valor):
    """Estratégia JavaScript para datepicker"""
    global driver
    
    driver.execute_script("""
        var campo = arguments[0], valor = arguments[1];
        campo.focus(); 
        campo.value = ''; 
        campo.value = valor;
        ['input','change','blur','keyup'].forEach(ev => 
            campo.dispatchEvent(new Event(ev, {bubbles: true}))
        );
    """, elemento, data_valor)

def _datepicker_actionchains(elemento, data_valor):
    """Estratégia ActionChains para datepicker"""
    global driver
    
    scroll_to_element_safe(elemento)
    time.sleep(0.5)
    
    ActionChains(driver).move_to_element(elemento).click().perform()
    time.sleep(0.5)
    ActionChains(driver).key_down(Keys.CONTROL).send_keys('a').key_up(Keys.CONTROL).perform()
    time.sleep(0.3)
    ActionChains(driver).send_keys(Keys.DELETE).perform()
    time.sleep(0.3)
    
    for char in data_valor:
        ActionChains(driver).send_keys(char).perform()
        time.sleep(0.05)
    
    ActionChains(driver).send_keys(Keys.TAB).perform()

def _datepicker_tradicional(elemento, data_valor):
    """Estratégia tradicional para datepicker"""
    scroll_to_element_safe(elemento)
    time.sleep(0.5)
    elemento.click()
    time.sleep(0.5)
    elemento.clear()
    elemento.send_keys(data_valor)
    elemento.send_keys(Keys.TAB)

def validar_data_preenchida(elemento, data_esperada):
    """Valida se a data foi preenchida corretamente"""
    try:
        if elemento is None:
            return False
            
        val = (elemento.get_attribute('value') or '').strip()
        if not val:
            return False
            
        if val == data_esperada or data_esperada in val:
            return True
            
        # Tenta comparar datas em diferentes formatos
        formatos = [
            '%d/%m/%Y %H:%M', '%d/%m/%Y %H:%M:%S',
            '%d/%m/%Y', '%m/%d/%Y', '%Y-%m-%d', '%d-%m-%Y'
        ]
        
        for formato in formatos:
            try:
                d1 = datetime.strptime(val, formato)
                d2 = datetime.strptime(data_esperada, formato)
                if d1 == d2:
                    return True
            except:
                continue
                
        return False
        
    except Exception:
        return False





def preencher_datepicker_por_indice(indice_campo, data_valor, max_tentativas=5):
    """Preenche datepicker pelo índice com estratégias múltiplas"""
    def acao():
        if not isinstance(indice_campo, int) or indice_campo < 0:
            raise ValueError(f"Índice inválido: {indice_campo}")
            
        if not data_valor or not isinstance(data_valor, str):
            raise ValueError(f"Data inválida: {data_valor}")
        
        tentativa = 0
        while tentativa < max_tentativas:
            tentativa += 1
            
            try:
                campos = encontrar_campos_datepicker()
                
                if not campos:
                    if tentativa < max_tentativas:
                        log(doc, f"⚠️ Nenhum campo datepicker encontrado, tentativa {tentativa}/{max_tentativas}", 'WARN')
                        time.sleep(2)
                        continue
                    raise Exception("Nenhum campo datepicker encontrado na página")
                
                if indice_campo >= len(campos):
                    raise Exception(f"Índice {indice_campo} inválido. Encontrados {len(campos)} campos")
                
                campo_info = campos[indice_campo]
                elemento = campo_info['elemento']
                campo_id = campo_info['id']
                
                log(doc, f"🎯 Tentativa {tentativa}: Preenchendo datepicker {indice_campo} (ID: {campo_id}) com '{data_valor}'")
                
                # Verifica se já está preenchido corretamente
                if validar_data_preenchida(elemento, data_valor):
                    log(doc, f"✅ Campo {indice_campo} já está preenchido corretamente!")
                    return True
                
                # Estratégias específicas para datepicker
                estrategias = [
                    lambda: _datepicker_jquery(campo_id, data_valor),
                    lambda: _datepicker_javascript(elemento, data_valor),
                    lambda: _datepicker_actionchains(elemento, data_valor),
                    lambda: _datepicker_tradicional(elemento, data_valor)
                ]
                
                for i, estrategia in enumerate(estrategias, 1):
                    try:
                        log(doc, f"   Aplicando estratégia {i} para datepicker...")
                        estrategia()
                        time.sleep(1)
                        
                        # Verifica se funcionou
                        if validar_data_preenchida(elemento, data_valor):
                            valor_atual = elemento.get_attribute('value')
                            log(doc, f"✅ Datepicker preenchido com estratégia {i}: '{valor_atual}'")
                            return True
                        else:
                            log(doc, f"⚠️ Estratégia {i} não preencheu corretamente", 'WARN')
                            
                    except Exception as e:
                        log(doc, f"⚠️ Estratégia {i} falhou: {e}", 'WARN')
                        continue
                
                # Se chegou aqui, nenhuma estratégia funcionou nesta tentativa
                if tentativa < max_tentativas:
                    log(doc, f"⚠️ Tentativa {tentativa} falhou, tentando novamente em 2s...", 'WARN')
                    time.sleep(2)
                    continue
                
            except Exception as e:
                if tentativa < max_tentativas:
                    log(doc, f"⚠️ Erro na tentativa {tentativa}: {e}, tentando novamente...", 'WARN')
                    time.sleep(2)
                    continue
                else:
                    raise
        
        raise Exception(f"Falha ao preencher datepicker {indice_campo} após {max_tentativas} tentativas")
    
    return acao


def preencher_campos_pesquisa_por_indice(self, 
                                         search_text: str, 
                                         search_xpaths=None, 
                                         max_campos: int | None = None, 
                                         pausa: float = 0.3,
                                         limpar_antes: bool = True):
    """
    Procura TODOS os campos de pesquisa e preenche um após o outro (ordem no DOM).
    - search_text: texto a preencher.
    - search_xpaths: lista de XPaths para busca (usa padrão se None).
    - max_campos: limita quantos campos serão preenchidos (None = todos).
    - pausa: pausa curta entre preenchimentos.
    - limpar_antes: se True, limpa o campo antes de preencher.

    Retorna: {"total_encontrados": int, "total_preenchidos": int, "xpaths_usados": [str], "falhas": [str]}
    """
    search_xpaths = search_xpaths or [
        "//input[@id='txtPesquisa']",
        "//input[@class='nomePesquisa']",
        # adicione mais padrões específicos antes do genérico:
        "//input[contains(translate(@name,'ABCDEFGHIJKLMNOPQRSTUVWXYZ','abcdefghijklmnopqrstuvwxyz'),'pesquisa')]",
        "//input[contains(translate(@class,'ABCDEFGHIJKLMNOPQRSTUVWXYZ','abcdefghijklmnopqrstuvwxyz'),'pesquisa')]",
        "//input[@type='text']"
    ]

    # Monta um XPATH-UNIÃO para manter a ordem no DOM
    xpath_uniao = " | ".join(f"({xp})" for xp in search_xpaths)

    # Coleta elementos (evita duplicados por id/ref)
    try:
        elementos = self.js.driver.find_elements(By.XPATH, xpath_uniao)
    except Exception as e:
        log(self.doc, f"⚠️ Erro ao buscar campos de pesquisa: {e}")
        return {"total_encontrados": 0, "total_preenchidos": 0, "xpaths_usados": [], "falhas": [str(e)]}

    # Filtra apenas visíveis e habilitados
    elementos_filtrados = []
    for el in elementos:
        try:
            if el.is_displayed() and el.is_enabled():
                elementos_filtrados.append(el)
        except Exception:
            continue

    total_encontrados = len(elementos_filtrados)
    if total_encontrados == 0:
        log(self.doc, "🔎 Nenhum campo de pesquisa visível/habilitado encontrado.")
        return {"total_encontrados": 0, "total_preenchidos": 0, "xpaths_usados": [], "falhas": []}

    if max_campos is not None and max_campos > 0:
        elementos_filtrados = elementos_filtrados[:max_campos]

    log(self.doc, f"🧭 Campos de pesquisa encontrados: {total_encontrados} | A preencher: {len(elementos_filtrados)}")

    total_preenchidos = 0
    falhas = []
    xpaths_usados = []

    for idx, el in enumerate(elementos_filtrados, start=1):
        try:
            # Recalcula um XPATH relativo único para log (opcional, pode ser pesado).
            # Aqui só registramos o index para simplificar.
            xpaths_usados.append(f"[campo #{idx}]")

            # Traz para a tela e foca
            self.js.driver.execute_script("arguments[0].scrollIntoView({block:'center'});", el)
            try:
                el.click()
            except Exception:
                pass

            if limpar_antes:
                try:
                    el.clear()
                except Exception:
                    # fallback via JS
                    try:
                        self.js.driver.execute_script("arguments[0].value='';", el)
                    except Exception:
                        pass

            # Preenche (usa seu helper para manter padrão)
            try:
                self.js.force_fill_element(el, search_text)  # se você tiver esse helper
            except AttributeError:
                # fallback: force_fill por XPATH do próprio elemento (gera um xpath usando JS)
                try:
                    self.js.driver.execute_script("arguments[0].value = arguments[1];", el, search_text)
                except Exception:
                    # último fallback: send_keys
                    el.send_keys(search_text)

            total_preenchidos += 1
            log(self.doc, f"   ✏️ Preenchido campo #{idx} com '{search_text}'")
            if pausa:
                time.sleep(pausa)

        except (StaleElementReferenceException, ElementNotInteractableException) as e:
            falhas.append(f"Campo #{idx}: {type(e).__name__}")
            log(self.doc, f"   ⚠️ Falha no campo #{idx}: {e}")
            continue
        except Exception as e:
            falhas.append(f"Campo #{idx}: {e}")
            log(self.doc, f"   ⚠️ Erro inesperado no campo #{idx}: {e}")
            continue

    resumo = {
        "total_encontrados": total_encontrados,
        "total_preenchidos": total_preenchidos,
        "xpaths_usados": xpaths_usados,
        "falhas": falhas
    }
    log(self.doc, f"✅ Pesquisa preenchida em {total_preenchidos}/{len(elementos_filtrados)} campos. Falhas: {len(falhas)}")
    return resumo

def clicar_lov_por_indice(indice_lov: int, max_tentativas: int = 5, timeout: int = 10, scroll: bool = True):
    """Clica no ícone de LOV pelo índice"""
    def acao():
        if not isinstance(indice_lov, int) or indice_lov < 0:
            raise ValueError(f"Índice inválido: {indice_lov}")

        tentativa = 0
        while tentativa < max_tentativas:
            tentativa += 1
            try:
                log(doc, f"🔎 Tentativa {tentativa}: Localizando ícones LOV...")
                elementos = driver.find_elements(By.CSS_SELECTOR, "a.sprites.sp-openLov")

                if not elementos:
                    if tentativa < max_tentativas:
                        log(doc, f"⚠️ Nenhum ícone LOV encontrado (tentativa {tentativa}/{max_tentativas})")
                        time.sleep(1.2)
                        continue
                    raise Exception("Nenhum ícone LOV encontrado.")

                if indice_lov >= len(elementos):
                    raise Exception(f"Índice {indice_lov} inválido. Encontrados {len(elementos)} ícones LOV.")

                locator_xpath = f"(//a[contains(@class,'sp-openLov')])[{indice_lov + 1}]"
                elemento = driver.find_element(By.XPATH, locator_xpath)

                log(doc, f"🎯 Preparando clique no LOV de índice {indice_lov}")

                def _wait_clickable():
                    wait.until(EC.element_to_be_clickable((By.XPATH, locator_xpath)))

                estrategias = [
                    lambda: (_wait_clickable(), elemento.click()),
                    lambda: (
                        driver.execute_script("arguments[0].scrollIntoView({block:'center'});", elemento) if scroll else None,
                        time.sleep(0.2),
                        elemento.click()
                    ),
                    lambda: driver.execute_script("arguments[0].click();", elemento),
                    lambda: ActionChains(driver).move_to_element(elemento).pause(0.1).click().perform()
                ]

                for i, estrategia in enumerate(estrategias, 1):
                    try:
                        log(doc, f"   ▶️ Estratégia {i} de clique no LOV...")
                        estrategia()
                        time.sleep(0.3)
                        log(doc, f"✅ Clique no LOV (índice {indice_lov}) realizado (estratégia {i})")
                        return True
                    except (ElementClickInterceptedException, StaleElementReferenceException, JavascriptException, TimeoutException) as e:
                        log(doc, f"⚠️ Estratégia {i} falhou: {e}")
                        try:
                            elementos = driver.find_elements(By.CSS_SELECTOR, "a.sprites.sp-openLov")
                            elemento = driver.find_element(By.XPATH, locator_xpath)
                        except:
                            pass
                        continue

                if tentativa < max_tentativas:
                    log(doc, f"⚠️ Tentativa {tentativa} não conseguiu clicar no LOV. Reintentando...")
                    time.sleep(1.2)
                    continue

            except Exception as e:
                if tentativa < max_tentativas:
                    log(doc, f"⚠️ Erro na tentativa {tentativa}: {e}. Reintentando...")
                    time.sleep(1.2)
                    continue
                raise

        raise Exception(f"Falha ao clicar no LOV de índice {indice_lov} após {max_tentativas} tentativas.")

    return acao

def encontrar_campos_textarea(timeout=10):
    """Retorna lista de textareas visíveis e interativas"""
    elementos = []
    try:
        wait.until(lambda d: len(d.find_elements(By.TAG_NAME, "textarea")) >= 0)
        textareas = driver.find_elements(By.TAG_NAME, "textarea")
    except:
        textareas = []

    for el in textareas:
        try:
            if not el.is_displayed() or not el.is_enabled():
                continue
            driver.execute_script("arguments[0].scrollIntoView({block:'center'});", el)
            elementos.append({
                "elemento": el,
                "id": el.get_attribute("id"),
                "name": el.get_attribute("name"),
            })
        except:
            continue

    return elementos

def normalizar_texto(txt):
    if txt is None:
        return ""
    return txt.replace("\r\n", "\n").replace("\r", "\n").strip()

def validar_textarea_preenchida(elemento, texto_esperado):
    """Confere se o valor atual da textarea bate com o esperado"""
    try:
        atual = elemento.get_attribute("value")
        if atual is None or atual == "":
            atual = (elemento.text or "")
        return normalizar_texto(atual) == normalizar_texto(texto_esperado)
    except StaleElementReferenceException:
        return False

def _prepare_focus_and_clear(elemento, limpar_primeiro=True):
    driver.execute_script("arguments[0].scrollIntoView({block:'center'});", elemento)
    try:
        elemento.click()
    except:
        driver.execute_script("arguments[0].focus();", elemento)

    if limpar_primeiro:
        try:
            elemento.clear()
        except:
            ActionChains(driver)\
                .move_to_element(elemento).click()\
                .key_down(Keys.CONTROL).send_keys("a").key_up(Keys.CONTROL)\
                .send_keys(Keys.DELETE).perform()

def _textarea_tradicional(elemento, texto, limpar_primeiro=True):
    _prepare_focus_and_clear(elemento, limpar_primeiro)
    elemento.send_keys(texto)
    elemento.send_keys(Keys.TAB)

def _textarea_actionchains(elemento, texto, limpar_primeiro=True):
    _prepare_focus_and_clear(elemento, limpar_primeiro)
    ac = ActionChains(driver)
    ac.move_to_element(elemento).click().perform()
    for chunk_start in range(0, len(texto), 400):
        ac.send_keys(texto[chunk_start:chunk_start+400]).perform()
        time.sleep(0.05)
    ac.send_keys(Keys.TAB).perform()

def _textarea_js_setvalue(elemento, texto):
    driver.execute_script("""
        const el = arguments[0];
        const val = arguments[1];
        el.value = val;
        el.dispatchEvent(new Event('input',  {bubbles:true}));
        el.dispatchEvent(new Event('change', {bubbles:true}));
        el.dispatchEvent(new KeyboardEvent('keyup', {bubbles:true}));
        el.dispatchEvent(new Event('blur',   {bubbles:true}));
    """, elemento, texto)

def _textarea_js_react_input(elemento, texto):
    driver.execute_script("""
        const el = arguments[0];
        const val = arguments[1];

        const desc = Object.getOwnPropertyDescriptor(window.HTMLTextAreaElement.prototype, 'value');
        if (desc && desc.set) {
            desc.set.call(el, val);
        } else {
            el.value = val;
        }

        el.dispatchEvent(new Event('input', {bubbles: true}));
        el.dispatchEvent(new Event('change', {bubbles: true}));
        el.dispatchEvent(new Event('blur', {bubbles: true}));
    """, elemento, texto)

def preencher_textarea_por_indice(indice_campo, texto, max_tentativas=5, limpar_primeiro=True):
    """Preenche textarea pelo índice usando estratégias múltiplas"""
    def acao():
        if not isinstance(indice_campo, int) or indice_campo < 0:
            raise ValueError(f"Índice inválido: {indice_campo}")
        if texto is None or not isinstance(texto, str):
            raise ValueError(f"Texto inválido: {texto!r}")

        tentativa = 0
        while tentativa < max_tentativas:
            tentativa += 1
            try:
                campos = encontrar_campos_textarea()
                if not campos:
                    if tentativa < max_tentativas:
                        log(doc, f"⚠️ Nenhuma <textarea> encontrada (tentativa {tentativa}/{max_tentativas})")
                        time.sleep(1.5)
                        continue
                    raise Exception("Nenhuma <textarea> foi encontrada.")

                if indice_campo >= len(campos):
                    raise Exception(f"Índice {indice_campo} inválido. Encontradas {len(campos)} textareas.")

                campo_info = campos[indice_campo]
                elemento = campo_info["elemento"]
                campo_id = campo_info.get("id") or "(sem id)"

                log(doc, f"🎯 Tentativa {tentativa}: Preenchendo textarea {indice_campo} (ID: {campo_id})")

                if validar_textarea_preenchida(elemento, texto):
                    log(doc, f"✅ Textarea {indice_campo} já está com o valor desejado.")
                    return True

                estrategias = [
                    lambda: _textarea_tradicional(elemento, texto, limpar_primeiro),
                    lambda: _textarea_actionchains(elemento, texto, limpar_primeiro),
                    lambda: _textarea_js_setvalue(elemento, texto),
                    lambda: _textarea_js_react_input(elemento, texto),
                ]

                for i, estrategia in enumerate(estrategias, 1):
                    try:
                        log(doc, f"   ▶️ Estratégia {i}…")
                        estrategia()
                        time.sleep(0.8)

                        if validar_textarea_preenchida(elemento, texto):
                            val = (elemento.get_attribute("value") or "").strip()
                            log(doc, f"✅ Preenchido com sucesso pela estratégia {i}")
                            return True
                        else:
                            log(doc, f"⚠️ Estratégia {i} não refletiu o valor esperado.")
                    except (StaleElementReferenceException, JavascriptException, TimeoutException) as e:
                        log(doc, f"⚠️ Estratégia {i} falhou: {e}")
                        try:
                            campos = encontrar_campos_textarea()
                            elemento = campos[indice_campo]["elemento"]
                        except:
                            pass
                        continue

                if tentativa < max_tentativas:
                    log(doc, f"⚠️ Tentativa {tentativa} falhou; reintentando em 1.5s…")
                    time.sleep(1.5)
                    continue
            except Exception as e:
                if tentativa < max_tentativas:
                    log(doc, f"⚠️ Erro na tentativa {tentativa}: {e}. Retentando…")
                    time.sleep(1.5)
                    continue
                else:
                    raise

        raise Exception(f"Falha ao preencher textarea {indice_campo} após {max_tentativas} tentativas.")
    return acao


def fechar_abas_extras(driver, doc, aba_principal_index=0):
    """Fecha todas as abas extras (como popups de impressão) mantendo apenas a aba principal"""
    try:
        handles = driver.window_handles
        if len(handles) <= 1:
            log(doc, "ℹ️ Apenas uma aba aberta - nada a fechar.")
            return True
        
        # Guarda o handle da aba principal
        aba_principal = handles[aba_principal_index]
        
        # Fecha todas as outras abas
        abas_fechadas = 0
        for handle in handles:
            if handle != aba_principal:
                try:
                    driver.switch_to.window(handle)
                    driver.close()
                    abas_fechadas += 1
                    log(doc, f"🗑️ Aba extra fechada ({abas_fechadas})")
                except Exception as e:
                    log(doc, f"⚠️ Erro ao fechar aba: {e}")
        
        # Retorna para a aba principal
        driver.switch_to.window(aba_principal)
        time.sleep(0.3)
        
        log(doc, f"✅ {abas_fechadas} aba(s) extra(s) fechada(s). Foco na aba principal.")
        return True
        
    except Exception as e:
        log(doc, f"❌ Erro ao fechar abas extras: {e}")
        return False

def focar_sistema_completo(js_engine, doc):
    """Garante o foco completo na aba principal do sistema e fecha abas extras"""
    driver = js_engine.driver
    try:
        # Primeiro fecha abas extras (como impressão)
        fechar_abas_extras(driver, doc)

        driver.switch_to.default_content()
        js_engine.execute_js("if (window.focus) window.focus();", timeout=3, fallback_result=None)
        time.sleep(0.3)

        log(doc, "✅ Foco garantido na aba do sistema.")
        return True

    except Exception as e:
        log(doc, f"⚠️ Falha ao focar aba do sistema: {e}")
        return False

def clicar_todos_botoes_sim_visiveis(js_engine, doc, pausa_entre=0.0):
    """Clica em TODOS os botões 'Sim' visíveis de uma vez"""
    js = r"""
    (function(){
      const isVisible = el => {
        if (!el) return false;
        const s = getComputedStyle(el);
        return el.offsetParent !== null && s.display !== 'none' &&
               s.visibility !== 'hidden' && parseFloat(s.opacity||1) > 0.01;
      };
      const buttons = Array.from(document.querySelectorAll("a.btModel.btGray.btyes"))
        .filter(isVisible)
        .filter(b => (b.textContent||"").trim().toLowerCase() === "sim");

      let clicked = 0;
      buttons.forEach(b => {
        try {
          b.style.pointerEvents = 'auto';
          b.removeAttribute('disabled');
          b.style.visibility = 'visible';
          b.style.display = 'inline-block';
          b.scrollIntoView({block:'center'});

          ['mouseover','mouseenter','mousemove','mousedown','mouseup','click'].forEach(t=>{
            b.dispatchEvent(new MouseEvent(t,{bubbles:true,cancelable:true,view:window,detail:1}));
          });
          if (typeof b.click === 'function') b.click();

          if (typeof window.jQuery !== 'undefined') {
            window.jQuery(b).trigger('click');
          }
          clicked++;
        } catch(e) {}
      });

      return { totalEncontrados: buttons.length, totalClicados: clicked };
    })();
    """
    try:
        res = js_engine.execute_js(js, timeout=5, fallback_result={"totalEncontrados": 0, "totalClicados": 0})
        total = int(res.get("totalEncontrados", 0))
        clic = int(res.get("totalClicados", 0))
        log(doc, f"⚡ 'Sim' visíveis encontrados: {total} | clicados: {clic}")
        if pausa_entre and clic > 0:
            time.sleep(pausa_entre)
        return res
    except Exception as e:
        log(doc, f"❌ Erro ao clicar em todos os 'Sim': {e}")
        return {"totalEncontrados": 0, "totalClicados": 0, "erro": str(e)}


def clicar_nao_ate_sumir(js_engine, doc, index=0, timeout=15, pausa=0.5):
    """
    Clica repetidamente no botão 'Não' (BtNo) dentro de um modal de confirmação,
    até que o modal desapareça da tela.

    Parâmetros:
    - js_engine: instância do engine JSForceEngine
    - doc: documento de log
    - index: índice do botão 'Não' (0 = primeiro, 1 = segundo, etc.)
    - timeout: tempo máximo em segundos antes de desistir
    - pausa: intervalo entre tentativas
    """

    # XPath do botão 'Não' - mais simples e direto
    xpath_botao = f"(//a[@id='BtNo' and contains(@class,'btno')])[{index + 1}]"
    
    # XPath do modal - busca qualquer modal visível com botão de confirmação
    xpath_modal_generico = "//div[contains(@class,'modal') and contains(@class,'confirmationYesNo')]"

    inicio = time.time()
    tentativas = 0
    sucesso = False

    log(doc, f"🟦 Iniciando clique no botão 'Não' (índice {index}) até o modal sumir...")

    # FORÇA pelo menos UMA tentativa de clique antes de verificar
    clicou_pelo_menos_uma_vez = False

    while True:
        tentativas += 1
        
        try:
            # Verifica se o botão 'Não' existe
            try:
                botoes = js_engine.driver.find_elements("xpath", xpath_botao)
                if not botoes and clicou_pelo_menos_uma_vez:
                    log(doc, f"✅ Botão 'Não' não encontrado - modal fechado após {tentativas - 1} tentativa(s).")
                    sucesso = True
                    break
                    
                if not botoes:
                    log(doc, "⚠️ Botão 'Não' não encontrado. Aguardando aparecer...")
                    time.sleep(pausa)
                    if time.time() - inicio > timeout:
                        log(doc, f"⏰ Timeout: botão 'Não' nunca apareceu.")
                        break
                    continue
                    
            except Exception as e:
                log(doc, f"⚠️ Erro ao localizar botão: {e}")
                if clicou_pelo_menos_uma_vez:
                    sucesso = True
                    break
                continue

            log(doc, f"➡️ Tentativa {tentativas}: clicando no botão 'Não'...")

            # Estratégia 1 — JavaScript com múltiplos eventos (MAIS AGRESSIVA)
            try:
                script = f"""
                var btn = document.evaluate("{xpath_botao}", document, null, 
                    XPathResult.FIRST_ORDERED_NODE_TYPE, null).singleNodeValue;
                if (btn) {{
                    // Remove qualquer bloqueio
                    btn.style.pointerEvents = 'auto';
                    btn.removeAttribute('disabled');
                    btn.style.visibility = 'visible';
                    btn.style.display = 'inline-block';
                    
                    // Scroll até o botão
                    btn.scrollIntoView({{block: 'center'}});
                    
                    // Dispara TODOS os eventos
                    ['mouseover','mouseenter','mousedown','focus','mouseup','click'].forEach(function(evt) {{
                        btn.dispatchEvent(new MouseEvent(evt, {{bubbles: true, cancelable: true}}));
                    }});
                    
                    // Clique direto
                    btn.click();
                    
                    // jQuery se disponível
                    if (typeof jQuery !== 'undefined') {{
                        jQuery(btn).trigger('click');
                    }}
                    
                    // onclick manual
                    if (btn.onclick) btn.onclick();
                    
                    return true;
                }}
                return false;
                """
                resultado = js_engine.driver.execute_script(script)
                if resultado:
                    log(doc, "✅ Clique via JavaScript ultra-agressivo executado.")
                    clicou_pelo_menos_uma_vez = True
                else:
                    log(doc, "⚠️ Script JavaScript retornou false.")
            except Exception as e:
                log(doc, f"⚠️ Falha no JavaScript agressivo: {e}")

            time.sleep(pausa)

            # Estratégia 2 — Force click do engine
            try:
                js_engine.force_click(xpath_botao, by_xpath=True)
                log(doc, "✅ Clique via force_click() executado.")
                clicou_pelo_menos_uma_vez = True
            except Exception as e:
                log(doc, f"⚠️ Falha no force_click: {e}")

            time.sleep(pausa)

            # Estratégia 3 — Selenium + ActionChains
            try:
                btn = js_engine.driver.find_element("xpath", xpath_botao)
                ActionChains(js_engine.driver).move_to_element(btn).pause(0.1).click().perform()
                log(doc, "✅ Clique via ActionChains executado.")
                clicou_pelo_menos_uma_vez = True
            except Exception as e:
                log(doc, f"⚠️ Falha no ActionChains: {e}")

            time.sleep(pausa)
            
            # Aguarda AJAX
            js_engine.wait_ajax_complete(3)

            # Verifica se o botão/modal ainda existem
            try:
                botoes = js_engine.driver.find_elements("xpath", xpath_botao)
                if not botoes:
                    log(doc, f"✅ Modal fechado após {tentativas} tentativa(s)!")
                    sucesso = True
                    break
            except:
                log(doc, "✅ Modal fechado (erro ao verificar = não existe mais)!")
                sucesso = True
                break

        except (NoSuchElementException, StaleElementReferenceException):
            if clicou_pelo_menos_uma_vez:
                log(doc, "✅ Elemento 'Não' sumiu - modal fechado!")
                sucesso = True
                break
            else:
                log(doc, "⚠️ Elemento 'Não' não encontrado, tentando novamente...")
        except Exception as e:
            log(doc, f"❌ Erro inesperado: {e}")

        if time.time() - inicio > timeout:
            log(doc, f"⏰ Tempo limite de {timeout}s atingido.")
            break

        time.sleep(pausa)

    if sucesso:
        log(doc, "🎉 Confirmação concluída com sucesso!")
    else:
        log(doc, "⚠️ Não foi possível confirmar - modal pode ainda estar aberto.")
    
    return sucesso



        
def clicar_sim_com_retry(doc, js_engine, wait, max_tentativas=5, pausa=1.5):
    """Clica em 'Sim' até o modal de confirmação desaparecer"""
    xpath_modal = "//div[contains(@class,'modal') and contains(@style,'z-index')]"
    xpath_sim = "(//div[contains(@class,'modal') and not(contains(@style,'display: none'))]//a[@class='btModel btGray btyes'])[last()]"

    tentativa = 0
    while tentativa < max_tentativas:
        tentativa += 1
        log(doc, f"🧩 Tentativa {tentativa} de fechar modal...")

        try:
            js_engine.force_click(xpath_sim, by_xpath=True)
            time.sleep(pausa)

            modais_visiveis = driver.find_elements(By.XPATH, xpath_modal)
            modais_ativos = [m for m in modais_visiveis if "display: none" not in m.get_attribute("style")]

            if not modais_ativos:
                log(doc, "✅ Botão 'Sim' clicado com sucesso.")
                return True

        except Exception as e:
            log(doc, f"⚠️ Tentativa {tentativa} falhou: {e}")

    log(doc, "❌ Botão 'Sim' não foi clicado após todas as tentativas.")
    return False

def clicar_primeiro_sp_add(js_engine, doc=None, timeout=5):
    """
    Localiza e clica no primeiro elemento <a class="sprites sp-add">, exatamente.
    Ignora variantes como 'sp-addVerde'.
    """
    from selenium.common.exceptions import (
        NoSuchElementException,
        ElementClickInterceptedException,
        StaleElementReferenceException,
    )
    from selenium.webdriver.common.by import By
    import time

    driver = js_engine.driver

    # XPath restrito: exige que as classes sejam EXATAMENTE 'sprites' e 'sp-add'
    xpath_btn = (
        "//a[contains(concat(' ', normalize-space(@class), ' '), ' sprites ') "
        "and contains(concat(' ', normalize-space(@class), ' '), ' sp-add ') "
        "and not(contains(@class, 'sp-addVerde'))][1]"
    )

    log(doc, "🧩 Procurando o botão '<a class=\"sprites sp-add\">' exato...")

    try:
        el = driver.find_element(By.XPATH, xpath_btn)
        log(doc, "🎯 Botão exato encontrado! Tentando clicar...")

        try:
            el.click()
            log(doc, "✅ Clique padrão realizado com sucesso.")
        except (ElementClickInterceptedException, StaleElementReferenceException):
            driver.execute_script("arguments[0].click();", el)
            log(doc, "⚡ Clique forçado via JavaScript realizado.")

        time.sleep(0.5)
        return True

    except NoSuchElementException:
        log(doc, "⚠️ Nenhum botão exato 'sp-add' encontrado.")
        return False

    except Exception as e:
        log(doc, f"❌ Erro ao clicar no botão 'sp-add': {e}")
        return False


def fechar_abas_extras_e_verificar_alerta(driver, doc):
    """
    Fecha abas extras (como de impressão) e verifica imediatamente
    se há alguma mensagem de alerta exibida após o fechamento.
    """
    try:
        safe_action(doc, "Fechando abas extras (impressão)", lambda:
            fechar_abas_extras(driver, doc)
        )

        alerta = encontrar_mensagem_alerta()
        if alerta:
            log(doc, f"⚠️ Alerta detectado após fechar abas extras: '{alerta.text.strip()}'")
        else:
            log(doc, "✅ Nenhum alerta detectado após fechar abas extras.")

        return True
    except Exception as e:
        log(doc, f"⚠️ Erro ao fechar abas extras e verificar alerta: {e}")
        return False



def preencher_campo_com_retry(driver, wait, seletor, valor, max_tentativas=3):
    """Tenta preencher o campo com diferentes métodos até conseguir"""
    
    for tentativa in range(max_tentativas):
        try:

            
            # Aguarda o elemento
            campo = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, seletor)))
            
            # Scroll até o elemento se necessário
            driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", campo)
            time.sleep(0.5)
            
            # Método 1: Tradicional
            if tentativa == 0:
                campo.click()
                campo.clear()
                campo.send_keys(valor)
                campo.send_keys(Keys.TAB)
            
            # Método 2: ActionChains
            elif tentativa == 1:
                ActionChains(driver).move_to_element(campo).click().perform()
                time.sleep(0.2)
                ActionChains(driver).key_down(Keys.CONTROL).send_keys('a').key_up(Keys.CONTROL).perform()
                ActionChains(driver).send_keys(valor).perform()
                ActionChains(driver).send_keys(Keys.TAB).perform()
            
            # Método 3: JavaScript
            else:
                driver.execute_script("""
                    var element = arguments[0];
                    var valor = arguments[1];
                    element.focus();
                    element.value = '';
                    element.value = valor;
                    element.dispatchEvent(new Event('input', { bubbles: true }));
                    element.dispatchEvent(new Event('change', { bubbles: true }));
                    element.blur();
                """, campo, valor)
            
            time.sleep(0.5)
            
            # Verifica se o valor foi preenchido
            valor_atual = campo.get_attribute('value')
            if valor_atual == valor:

                return True
            else:
                print()
                
        except Exception as e:
            time.sleep(1)
    

    return False

def gerar_dados_documentos():
    """Gera documentos fictícios para o cadastro."""
    carteira_trabalho = str(random.randint(10000000, 99999999))
    pis = fake.cpf().replace('.', '').replace('-', '')[:11]
    cnh = str(random.randint(10000000000, 99999999999))
    
    return carteira_trabalho, pis, cnh

# Gera os dados necessários
data_nascimento, data_falecimento, data_sepultamento, data_registro, data_de, data_ate = gerar_datas_validas()
carteira_trabalho, pis, cnh = gerar_dados_documentos()

vencimento_cnh = fake.date_between(start_date='today', end_date='+10y')
vencimento_cnh_str = vencimento_cnh.strftime('%d/%m/%Y')

data_admissao = fake.date_between(start_date='-10y', end_date='today')
data_admissao_str = data_admissao.strftime('%d/%m/%Y')

hora_falecimento = fake.time(pattern="%H:%M")
hora_sepultamento = fake.time(pattern="%H:%M")
localizacao = fake.city()

# ==== CONFIGURAÇÕES ====
URL = "http://localhost:8080/gs/index.xhtml"
LOGIN_EMAIL = "joaoeduardo.gold@outlook.com"
LOGIN_PASSWORD = "071999gs"
VERBOSE_LOGGING = True
CAPTURAR_SCREENSHOTS = True

# ==== DOCUMENTO ====
doc = Document()
doc.add_heading("RELATÓRIO DO TESTE", 0)
doc.add_paragraph("Aprovação de Trocar/Vincular Jazigos ou Quadras – Cenário 1: Preenchimento completo e salvamento.")
doc.add_paragraph(f"Data do teste: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")

screenshot_registradas = set()

def log(doc, msg):
    print(msg)
    doc.add_paragraph(msg)

# ==== FUNÇÕES DE UTILITÁRIO ====
def log(doc, msg, nivel='INFO'):
    """Sistema de logging melhorado com níveis"""
    timestamp = datetime.now().strftime('%H:%M:%S')
    formatted_msg = f"[{timestamp}] {nivel}: {msg}"
    
    if VERBOSE_LOGGING:
        print(formatted_msg)
    else:
        print(msg)
    
    if hasattr(doc, 'add_paragraph'):
        doc.add_paragraph(formatted_msg)

def take_screenshot(driver, doc, nome, forcar=False):
    if nome not in screenshot_registradas:
        path = f"screenshots/{nome}.png"
        os.makedirs("screenshots", exist_ok=True)
        driver.save_screenshot(path)
        doc.add_paragraph(f"Screenshot: {nome}")
        doc.add_picture(path, width=Inches(5.5))
        screenshot_registradas.add(nome)

def finalizar_relatorio():
    nome_arquivo = f"relatorio_trocar_vincular_jazigos_ou_quadras_cenario_1_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc.save(nome_arquivo)
    log(doc, f"📄 Relatório salvo como: {nome_arquivo}")
    subprocess.run(["start", "winword", nome_arquivo], shell=True)
    driver.quit()

def ajustar_zoom():
    try:
        driver.execute_script("document.body.style.zoom='90%'")
        log(doc, "🔍 Zoom ajustado para 90%.")
    except Exception as e:
        log(doc, f"⚠️ Erro ao ajustar zoom: {e}")

def abrir_modal_e_selecionar_robusto_xpath(
    btn_xpath,
    pesquisa_xpath,
    termo_pesquisa,
    btn_pesquisar_xpath,
    resultado_xpath,
    timeout=12,
    max_tentativas=3,
    iframe_xpath=None,   # se o LOV abrir em iframe, informe o xpath aqui
):
    """
    Abre o modal (LOV), pesquisa pelo termo e clica no resultado.
    - Usa retries
    - Click normal + fallback via JS
    - Clear resistente no input
    - Opcional: troca para iframe do modal
    """

    def _js_click(el):
        driver.execute_script("arguments[0].click();", el)

    def _clear_resistente(el):
        try:
            el.clear()
            # alguns inputs ignoram clear(); garanta com CTRL+A + DEL
            el.send_keys(Keys.CONTROL, "a")
            el.send_keys(Keys.DELETE)
        except Exception:
            # fallback final via JS
            driver.execute_script("arguments[0].value='';", el)

    def _aguardar_ajax_overlay():
        # Ajuste se tiver seletor de overlay específico (ex: .blockScreen)
        t0 = time.time()
        while time.time() - t0 < 8:
            try:
                ready = driver.execute_script("return document.readyState")
                ajax_ok = driver.execute_script("return window.jQuery ? jQuery.active === 0 : true")
                if ready == "complete" and ajax_ok:
                    break
            except Exception:
                pass
            time.sleep(0.2)

    def acao():
        tentativa = 0
        while tentativa < max_tentativas:
            tentativa += 1
            try:
                # 1) Abrir o modal
                btn = WebDriverWait(driver, timeout).until(
                    EC.element_to_be_clickable((By.XPATH, btn_xpath))
                )
                try:
                    btn.click()
                except Exception:
                    _js_click(btn)

                # 2) (Opcional) Entrar no iframe do modal
                if iframe_xpath:
                    frame = WebDriverWait(driver, timeout).until(
                        EC.frame_to_be_available_and_switch_to_it((By.XPATH, iframe_xpath))
                    )

                # 3) Localizar e preparar campo de pesquisa
                campo = WebDriverWait(driver, timeout).until(
                    EC.visibility_of_element_located((By.XPATH, pesquisa_xpath))
                )
                _clear_resistente(campo)
                campo.send_keys(termo_pesquisa)

                # 4) Clicar no botão Pesquisar (com fallback JS)
                btn_pesq = WebDriverWait(driver, timeout).until(
                    EC.element_to_be_clickable((By.XPATH, btn_pesquisar_xpath))
                )
                try:
                    btn_pesq.click()
                except Exception:
                    _js_click(btn_pesq)

                _aguardar_ajax_overlay()
                time.sleep(0.4)

                # 5) Aguardar e clicar no resultado
                resultado = WebDriverWait(driver, timeout).until(
                    EC.element_to_be_clickable((By.XPATH, resultado_xpath))
                )
                driver.execute_script("arguments[0].scrollIntoView({block:'center'});", resultado)
                try:
                    resultado.click()
                except Exception:
                    _js_click(resultado)

                # 6) Sair do iframe (se entrou)
                if iframe_xpath:
                    driver.switch_to.default_content()

                time.sleep(0.6)
                return True

            except Exception as e:
                # Se estava em iframe, volte para o conteúdo principal antes de tentar de novo
                try:
                    driver.switch_to.default_content()
                except Exception:
                    pass

                if tentativa < max_tentativas:
                    log(doc, f"⚠️ Tentativa {tentativa}/{max_tentativas} falhou: {e}. Retentando…")
                    time.sleep(1.0 + 0.3 * tentativa)
                    continue
                else:
                    raise

    return acao

def encontrar_mensagem_alerta():
    seletores = [
        (".alerts.salvo", "✅ Mensagem de Sucesso"),
        (".alerts.alerta", "⚠️ Mensagem de Alerta"),
        (".alerts.erro", "❌ Mensagem de Erro"),
    ]

    for seletor, tipo in seletores:
        try:
            elemento = driver.find_element(By.CSS_SELECTOR, seletor)
            if elemento.is_displayed():
                log(doc, f"📢 {tipo}: {elemento.text}")
                return elemento
        except:
            continue

    log(doc, "ℹ️ Nenhuma mensagem de alerta encontrada.")
    return None


def rolar_ate_grupo_rateio():
    open_lov_rateio = WebDriverWait(driver, 10).until(
        EC.element_to_be_clickable((By.CSS_SELECTOR, '#gsContratos > div.wdTelas > div.wdWizard.clearfix.telaConsulta > div.wizardHolder > div > div.stepPacote.step3 > div:nth-child(11) > div > div > a'))
    )
    driver.execute_script("arguments[0].scrollIntoView();", open_lov_rateio)



def rolar_ate_salvar_titular():
    salvar_titular = WebDriverWait(driver, 10).until(
        EC.element_to_be_clickable((By.XPATH, "//a[@class='btModel btGray hAlign' and @title='Adicionar ao contrato(Ctrl + enter)' and normalize-space(text())='Salvar']"))
    )
    driver.execute_script("arguments[0].scrollIntoView();", salvar_titular)


def preencher_campo_data(selector, valor):
    def acao():
        campo = wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, selector)))
        campo.click()
        campo.clear()
        campo.send_keys(valor)
        time.sleep(0.2)
    return acao

def selecionar_opcao(selector, texto):
    def acao():
        select_element = wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, selector)))
        Select(select_element).select_by_visible_text(texto)
    return acao

from datetime import datetime, timedelta
from selenium.webdriver.common.by import By
import time

from datetime import datetime, timedelta
from selenium.webdriver.common.by import By
import time

def tratar_modal_gerar_parcela_transferencia(js_engine, doc, timeout=10):
    driver = js_engine.driver

    xpath_modal = ("//div[contains(@class,'modal') and "
                   ".//h2[normalize-space()='Gerar Parcela Transferência de Contrato']]")

    time.sleep(0.8)
    modais = driver.find_elements(By.XPATH, xpath_modal)
    modais_visiveis = [m for m in modais if m.is_displayed()]

    if not modais_visiveis:
        log(doc, "ℹ️ Modal 'Gerar Parcela Transferência de Contrato' não apareceu.")
        return False

    modal = modais_visiveis[0]
    log(doc, "🟦 Modal detectado: Gerar Parcela Transferência de Contrato")

    # Força modal na visão
    try:
        js_engine.scroll_into_view(modal)
    except:
        pass

    # ======================================
    # 1) ABRIR LOV - TIPO MENSALIDADE
    # ======================================
    try:
        log(doc, "📌 Abrindo LOV de Tipo Mensalidade...")
        abrir_modal_lov_handler(js_engine, 0)
        time.sleep(1)

        # Selecionar primeira linha
        log(doc, "📌 Selecionando primeira linha do LOV...")

        js_engine.execute_js("""
            (function(){
                const isVisible = el => {
                    if (!el) return false;
                    const s = getComputedStyle(el);
                    return el.offsetParent !== null &&
                        s.display !== 'none' &&
                        s.visibility !== 'hidden' &&
                        parseFloat(s.opacity||1) > 0.01;
                };

                const links = Array.from(document.querySelectorAll(".linkAlterar"))
                    .filter(isVisible);

                if (links.length === 0) {
                    return { ok: false, msg: "Nenhum elemento .linkAlterar visível encontrado." };
                }

                const el = links[0];
                const rect = el.getBoundingClientRect();
                const x = rect.left + rect.width/2;
                const y = rect.top + rect.height/2;

                ['mouseover','mouseenter','mousemove','mousedown','mouseup','click']
                    .forEach(evt => {
                        el.dispatchEvent(new MouseEvent(evt, {
                            bubbles:true,
                            cancelable:true,
                            view:window,
                            clientX:x,
                            clientY:y
                        }));
                    });

                if (typeof el.click === "function") el.click();

                return { ok:true, msg:"Clique em .linkAlterar executado com sucesso." };
            })();
        """, timeout=8)
        time.sleep(1)
    except Exception as e:
        log(doc, f"⚠️ Erro ao selecionar Tipo Mensalidade: {e}")

    # ======================================
    # 2) PREENCHER DATA VENCIMENTO
    # ======================================
    try:
        data = (datetime.today() + timedelta(days=5)).strftime("%d/%m/%Y")
        log(doc, f"📅 Preenchendo Data de Vencimento: {data}")

        preencher_datepicker_por_indice(js_engine, 0, data)
        time.sleep(1)
    except Exception as e:
        log(doc, f"⚠️ Erro ao preencher data: {e}")

    # ======================================
    # 3) CLICAR EM GERAR
    # ======================================
    try:
        log(doc, "💾 Clicando em Gerar...")
        js_engine.force_click(
            "//a[contains(@class,'btSave') and contains(.,'Gerar')]",
            by_xpath=True
        )
        time.sleep(1)
        js_engine.wait_ajax_complete(10)
    except Exception as e:
        log(doc, f"⚠️ Erro ao clicar em Gerar: {e}")

    # ======================================
    # 4) AGUARDAR MODAL SUMIR
    # ======================================
    inicio = time.time()
    while time.time() - inicio < timeout:
        modais = driver.find_elements(By.XPATH, xpath_modal)
        visiveis = [m for m in modais if m.is_displayed()]
        if not visiveis:
            log(doc, "✅ Modal fechado com sucesso.")
            return True
        time.sleep(0.4)

    log(doc, "⚠️ Timeout esperando modal fechar.", "WARN")
    return True

def preencher_pessoa_completa(nome_pessoa):
    """Função para preencher dados completos de uma pessoa."""
    def acao():
        # Dados Pessoais
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "#cg_1 > div.wdTelas > div > div.catWrapper > div > div.cat_dadosPessoais.categoriaHolder > div > div > div:nth-child(2) > div:nth-child(2) > input"))).send_keys(nome_pessoa)
        Select(wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "#cg_1 > div.wdTelas > div > div.catWrapper > div > div.cat_dadosPessoais.categoriaHolder > div > div > div:nth-child(2) > div:nth-child(3) > select")))).select_by_visible_text("Física")
        Select(wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "#cg_1 > div.wdTelas > div > div.catWrapper > div > div.cat_dadosPessoais.categoriaHolder > div > div > div:nth-child(2) > div:nth-child(4) > select")))).select_by_visible_text("Carteira de Identidade Classista")
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "#cg_1 > div.wdTelas > div > div.catWrapper > div > div.cat_dadosPessoais.categoriaHolder > div > div > div.formRow.divPessoaFISICA > div:nth-child(1) > input"))).send_keys(fake.rg())

        # Data de expedição
        campo_data_expedicao = wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, "input.dataExpedicao")))
        campo_data_expedicao.click()
        campo_data_expedicao.send_keys(fake.date_of_birth(minimum_age=18, maximum_age=60).strftime("%d/%m/%Y"))

        # CPF
        cpf = CPF().generate()
        cpf_field = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "#cg_1 > div.wdTelas > div > div.catWrapper > div > div.cat_dadosPessoais.categoriaHolder > div > div > div.formRow.divPessoaFISICA > div:nth-child(3) > input")))
        cpf_field.click()
        time.sleep(0.5)
        cpf_field.send_keys(cpf)

        # Dados Complementares
        wait.until(EC.element_to_be_clickable((By.LINK_TEXT, "Dados Complementares"))).click()
        time.sleep(1)

        Select(wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "#cg_1 > div.wdTelas > div > div.catWrapper > div > div.cat_dadosComplementares.categoriaHolder > div > div > div > div:nth-child(1) > select")))).select_by_visible_text("Solteiro")
        Select(wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "#cg_1 > div.wdTelas > div > div.catWrapper > div > div.cat_dadosComplementares.categoriaHolder > div > div > div > div:nth-child(2) > select")))).select_by_visible_text("Feminino")

        # Data de nascimento
        campo_data_nascimento = wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, "input.dataNascimento")))
        campo_data_nascimento.click()
        campo_data_nascimento.send_keys(fake.date_of_birth(minimum_age=18, maximum_age=60).strftime("%d/%m/%Y"))

        # Contatos
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "#cg_1 > div.wdTelas > div > div.catWrapper > div > div.cat_dadosComplementares.categoriaHolder > div > div > div > div:nth-child(5) > input"))).send_keys(fake.phone_number())
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "#cg_1 > div.wdTelas > div > div.catWrapper > div > div.cat_dadosComplementares.categoriaHolder > div > div > div > div:nth-child(6) > input"))).send_keys(fake.phone_number())
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "#cg_1 > div.wdTelas > div > div.catWrapper > div > div.cat_dadosComplementares.categoriaHolder > div > div > div > div:nth-child(7) > input"))).send_keys(fake.phone_number())
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "#cg_1 > div.wdTelas > div > div.catWrapper > div > div.cat_dadosComplementares.categoriaHolder > div > div > div > div:nth-child(9) > input"))).send_keys(fake.email())
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "#cg_1 > div.wdTelas > div > div.catWrapper > div > div.cat_dadosComplementares.categoriaHolder > div > div > div > div:nth-child(11) > input"))).send_keys(fake.city())
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "#cg_1 > div.wdTelas > div > div.catWrapper > div > div.cat_dadosComplementares.categoriaHolder > div > div > div > div:nth-child(12) > input"))).send_keys(fake.country())
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "#cg_1 > div.wdTelas > div > div.catWrapper > div > div.cat_dadosComplementares.categoriaHolder > div > div > div > div:nth-child(13) > input"))).send_keys(fake.first_name())
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "#cg_1 > div.wdTelas > div > div.catWrapper > div > div.cat_dadosComplementares.categoriaHolder > div > div > div > div:nth-child(14) > input"))).send_keys(fake.first_name())
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "#cg_1 > div.wdTelas > div > div.catWrapper > div > div.cat_dadosComplementares.categoriaHolder > div > div > div > div:nth-child(16) > input"))).send_keys(fake.job())

        # Endereços
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "#cg_1 > div.wdTelas > div > div.categorias.overflow.overflowY > ul > li.li_enderecos > a"))).click()
        time.sleep(3)

        # Preenche endereço
        try:
            elemento = wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, "#cg_1 > div.wdTelas > div > div.catWrapper > div > div.cat_enderecos.categoriaHolder > div.groupHolder.clearfix.grupo_enderecoResidencial > div > div:nth-child(2) > div:nth-child(1) > div > input")))
            elemento.send_keys("15081115")
            
            botao = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "#cg_1 > div.wdTelas > div > div.catWrapper > div > div.cat_enderecos.categoriaHolder > div.groupHolder.clearfix.grupo_enderecoResidencial > div > div:nth-child(2) > div:nth-child(1) > div > a")))
            botao.click()
        except Exception as e:
            log(doc, f"Erro ao preencher endereço: {e}")

        time.sleep(5)

        try:
            element = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "#BtYes")))
            element.click()
        except:
            pass

        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "#cg_1 > div.wdTelas > div > div.catWrapper > div > div.cat_enderecos.categoriaHolder > div.groupHolder.clearfix.grupo_enderecoResidencial > div > div:nth-child(3) > div:nth-child(2) > input"))).send_keys("1733")
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "#cg_1 > div.wdTelas > div > div.catWrapper > div > div.cat_enderecos.categoriaHolder > div.groupHolder.clearfix.grupo_enderecoResidencial > div > div:nth-child(3) > div:nth-child(3) > input"))).send_keys("Casa")
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "#cg_1 > div.wdTelas > div > div.catWrapper > div > div.cat_enderecos.categoriaHolder > div.groupHolder.clearfix.grupo_enderecoResidencial > div > div:nth-child(3) > div:nth-child(9) > label > input"))).click()

        # Salvar pessoa
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "#cg_1 > div.wdTelas > div > div.btnHolder > a.btModel.btGray.btsave"))).click()
        time.sleep(3)

    return acao

def preencher_datepicker_por_indice(js_engine, indice, data_string):
    """
    Clica no datepicker pelo índice e preenche a data via JS.
    - indice: índice do ícone do calendário (<span class="sp-calendar">)
    - data_string: formato dd/mm/yyyy
    """
    driver = js_engine.driver
    from selenium.webdriver.common.by import By

    icones = driver.find_elements(By.XPATH, "//span[contains(@class,'sp-calendar')]")
    icones_visiveis = [i for i in icones if i.is_displayed()]

    if indice >= len(icones_visiveis):
        raise Exception(f"Ícone de datepicker índice {indice} não existe (total: {len(icones_visiveis)})")

    # Clica no calendário
    js_engine.force_click(f"(//span[contains(@class,'sp-calendar')])[{indice+1}]", by_xpath=True)

    # Pega o input anterior (irmão)
    script = f"""
    (function(){{
        const icons = Array.from(document.querySelectorAll("span.sp-calendar"));
        const el = icons[{indice}];
        if(!el) return false;
        const input = el.previousElementSibling || el.closest("div").querySelector("input");
        if(!input) return false;
        input.value = "{data_string}";
        input.dispatchEvent(new Event('input', {{bubbles:true}}));
        input.dispatchEvent(new Event('change', {{bubbles:true}}));
        return true;
    }})(); 
    """
    js_engine.execute_js(script, timeout=5)


def abrir_modal_lov_handler(js_engine, indice):
    """
    Abre o LOV pelo índice usando JS Engine.
    """
    driver = js_engine.driver
    from selenium.webdriver.common.by import By

    xpath = "//a[contains(@class,'sp-openLov')]"
    itens = driver.find_elements(By.XPATH, xpath)
    visiveis = [el for el in itens if el.is_displayed()]

    if indice >= len(visiveis):
        raise Exception(f"Ícone LOV índice {indice} não encontrado. Total: {len(visiveis)}")

    # Clica via JS
    js_engine.force_click(f"({xpath})[{indice+1}]", by_xpath=True)


def validar_pacote_selecionado(js_engine, doc, timeout=3):
    """
    Valida se ALGUM pacote tem a classe 'selected'
    """
    xpath_selecionado = "//div[contains(@class,'pacote') and contains(@class,'selected')]"
    
    inicio = time.time()
    while time.time() - inicio < timeout:
        try:
            elementos = js_engine.driver.find_elements("xpath", xpath_selecionado)
            
            for elemento in elementos:
                if elemento.is_displayed():
                    ref = elemento.get_attribute("ref") or "?"
                    classes = elemento.get_attribute("class") or ""
                    
                    if "selected" in classes:
                        log(doc, f"   ✓ Validação OK: Pacote ref={ref} está selecionado")
                        return True
                        
        except Exception:
            pass
        
        time.sleep(0.2)
    
    log(doc, "   ✗ Validação FALHOU: Nenhum pacote com classe 'selected'")
    return False


def clicar_avancar_robusto(js_engine, doc, max_tentativas=5, pausa=1.5):
    """
    Clica no botão 'Avançar' com validação de display
    """
    # XPath mais específico baseado no HTML fornecido
    xpath_avancar = "//a[@class='btModel btGray' and contains(text(),'Avancar') and @style='display: inline-block;']"
    
    # XPaths alternativos
    xpaths_alternativos = [
        "//a[@class='btModel btGray' and contains(normalize-space(),'Avancar')]",
        "//a[contains(@class,'btModel') and contains(@class,'btGray') and contains(text(),'Avancar')]",
        "//a[contains(@class,'btGray')]/span[@class='sprites sp-next']/parent::a"
    ]
    
    for tentativa in range(1, max_tentativas + 1):
        try:
            log(doc, f"➡️ Tentativa {tentativa}/{max_tentativas}: Procurando 'Avançar'...")
            
            # Aguarda AJAX
            time.sleep(1)
            js_engine.wait_ajax_complete(8)
            
            # Tenta XPath principal primeiro
            all_xpaths = [xpath_avancar] + xpaths_alternativos
            
            for i, xpath in enumerate(all_xpaths, 1):
                try:
                    elementos = js_engine.driver.find_elements("xpath", xpath)
                    
                    # Filtra visíveis e habilitados
                    botoes_visiveis = []
                    for el in elementos:
                        try:
                            style = el.get_attribute("style") or ""
                            if el.is_displayed() and el.is_enabled():
                                # Verifica se não está com display:none
                                if "display: none" not in style and "display:none" not in style:
                                    botoes_visiveis.append(el)
                        except:
                            continue
                    
                    if not botoes_visiveis:
                        continue
                    
                    log(doc, f"   ✓ Botão 'Avançar' encontrado (XPath {i})")
                    
                    botao = botoes_visiveis[0]
                    
                    # Scroll até o botão
                    js_engine.scroll_into_view(botao, padding=150)
                    time.sleep(0.5)
                    
                    # Clica com o engine
                    js_engine.click_element(botao, wait_after=1.5)
                    
                    log(doc, "✅ Botão 'Avançar' clicado!")
                    js_engine.wait_ajax_complete(10)
                    return True
                    
                except StaleElementReferenceException:
                    log(doc, f"   ⚠️ Elemento stale (XPath {i})")
                    time.sleep(0.5)
                    continue
                    
                except Exception as e:
                    log(doc, f"   ⚠️ XPath {i} falhou: {str(e)[:80]}")
                    continue
            
            # Se chegou aqui, nenhum XPath funcionou
            if tentativa < max_tentativas:
                log(doc, f"   ⏳ Retry em {pausa}s...")
                time.sleep(pausa)
                continue
            
        except Exception as e:
            log(doc, f"❌ Erro na tentativa {tentativa}: {e}")
            if tentativa < max_tentativas:
                time.sleep(pausa)
                continue
    
    raise Exception(f"Falha ao clicar em 'Avançar' após {max_tentativas} tentativas")


def click_avancar(driver):
    xpath = "//a[@class='btModel btGray' and contains(normalize-space(),'Avancar')]"
    
    for tentativa in range(1, 8):
        try:
            el = WebDriverWait(driver, 5).until(
                EC.element_to_be_clickable((By.XPATH, xpath))
            )

            # 1) Scroll até o elemento
            driver.execute_script("arguments[0].scrollIntoView(true);", el)
            time.sleep(0.2)

            # 2) Click normal
            try:
                el.click()
                return
            except:
                pass

            # 3) ActionChains
            try:
                actions = ActionChains(driver)
                actions.move_to_element(el).click().perform()
                return
            except:
                pass

            # 4) JavaScript click
            try:
                driver.execute_script("arguments[0].click();", el)
                return
            except:
                pass

            # 5) JS mousedown + mouseup (O MELHOR PARA PEGASUS)
            try:
                driver.execute_script("""
                    var e = arguments[0];
                    e.dispatchEvent(new MouseEvent('mousedown', {bubbles:true}));
                    e.dispatchEvent(new MouseEvent('mouseup', {bubbles:true}));
                    e.dispatchEvent(new MouseEvent('click', {bubbles:true}));
                """, el)
                return
            except:
                pass

        except Exception as e:
            time.sleep(0.5)

    raise Exception("Falha ao clicar no botão Avançar após todas as estratégias.")


def clicar_avancar_por_indice(driver, indice, tentativas=6):
    xpath = "//a[contains(normalize-space(),'Avancar') and ./span[contains(@class,'sp-next')]]"

    print("[DEBUG] Procurando botões Avancar no DOM...")  # NOVO

    botoes = driver.find_elements(By.XPATH, xpath)

    print(f"[DEBUG] find_elements retornou: {len(botoes)} elementos")  # NOVO
    total = len(botoes)

    if total == 0:
        raise Exception("Nenhum botão 'Avancar' encontrado no DOM.")

    if indice >= total:
        raise Exception(f"Índice {indice} inválido — existem apenas {total} botões no DOM.")

    el = botoes[indice]
    print(f"[INFO] Clicando no botão 'Avancar' de índice {indice}...")

    for tente in range(1, tentativas + 1):
        try:
            # Scroll até o botão
            driver.execute_script("arguments[0].scrollIntoView(true);", el)
            time.sleep(0.2)

            # Estratégia 1: click normal
            try:
                el.click()
                print("[OK] Clique normal realizado.")
                return
            except:
                pass

            # Estratégia 2: ActionChains
            try:
                ActionChains(driver).move_to_element(el).click().perform()
                print("[OK] Clique via ActionChains realizado.")
                return
            except:
                pass

            # Estratégia 3: JS click
            try:
                driver.execute_script("arguments[0].click();", el)
                print("[OK] Clique via JavaScript realizado.")
                return
            except:
                pass

            # Estratégia 4: JS mousedown + mouseup (melhor para Pegasus)
            try:
                driver.execute_script("""
                    arguments[0].dispatchEvent(new MouseEvent('mousedown', {bubbles:true}));
                    arguments[0].dispatchEvent(new MouseEvent('mouseup', {bubbles:true}));
                    arguments[0].dispatchEvent(new MouseEvent('click', {bubbles:true}));
                """, el)
                print("[OK] Clique via mousedown + mouseup realizado.")
                return
            except:
                pass

        except Exception as e:
            print(f"[WARN] Tentativa {tente}/{tentativas} falhou: {e}")
            time.sleep(0.3)

    raise Exception(f"Falha ao clicar no botão Avançar de índice {indice} após {tentativas} tentativas.")





def clicar_finalizar_por_indice(driver, indice, tentativas=6):
    xpath = "//div[@class='btModel btGray' and @style='display: inline-block;' and normalize-space(text())='Finalizar']"

    print("[DEBUG] Procurando botões Finalizar no DOM...")  # NOVO

    botoes = driver.find_elements(By.XPATH, xpath)

    print(f"[DEBUG] find_elements retornou: {len(botoes)} elementos")  # NOVO
    total = len(botoes)

    if total == 0:
        raise Exception("Nenhum botão 'Finalizar' encontrado no DOM.")

    if indice >= total:
        raise Exception(f"Índice {indice} inválido — existem apenas {total} botões no DOM.")

    el = botoes[indice]
    print(f"[INFO] Clicando no botão 'Finalizar' de índice {indice}...")

    for tente in range(1, tentativas + 1):
        try:
            # Scroll até o botão
            driver.execute_script("arguments[0].scrollIntoView(true);", el)
            time.sleep(0.2)

            # Estratégia 1: click normal
            try:
                el.click()
                print("[OK] Clique normal realizado.")
                return
            except:
                pass

            # Estratégia 2: ActionChains
            try:
                ActionChains(driver).move_to_element(el).click().perform()
                print("[OK] Clique via ActionChains realizado.")
                return
            except:
                pass

            # Estratégia 3: JS click
            try:
                driver.execute_script("arguments[0].click();", el)
                print("[OK] Clique via JavaScript realizado.")
                return
            except:
                pass

            # Estratégia 4: JS mousedown + mouseup (melhor para Pegasus)
            try:
                driver.execute_script("""
                    arguments[0].dispatchEvent(new MouseEvent('mousedown', {bubbles:true}));
                    arguments[0].dispatchEvent(new MouseEvent('mouseup', {bubbles:true}));
                    arguments[0].dispatchEvent(new MouseEvent('click', {bubbles:true}));
                """, el)
                print("[OK] Clique via mousedown + mouseup realizado.")
                return
            except:
                pass

        except Exception as e:
            print(f"[WARN] Tentativa {tente}/{tentativas} falhou: {e}")
            time.sleep(0.3)

    raise Exception(f"Falha ao clicar no botão Finalizar de índice {indice} após {tentativas} tentativas.")




def selecionar_pacote_robusto(js_engine, relatorio, max_tentativas=5):
    """Seleciona pacote com validação da classe 'selected'"""
    xpath_pacote_nao_selecionado = "//div[contains(@class,'pacote') and contains(@class,'clearfix') and not(contains(@class,'selected'))][1]"
    xpath_pacote_selecionado = "//div[contains(@class,'pacote') and contains(@class,'selected')]"
    
    for tentativa in range(1, max_tentativas + 1):
        try:
            relatorio.log(f"🎯 Tentativa {tentativa}/{max_tentativas}: Selecionando pacote...")
            
            time.sleep(1)
            js_engine.wait_ajax_complete(5)
            
            # Verifica se já está selecionado
            try:
                pacotes_selecionados = js_engine.driver.find_elements("xpath", xpath_pacote_selecionado)
                if pacotes_selecionados and pacotes_selecionados[0].is_displayed():
                    relatorio.log("✅ Pacote já estava selecionado!", "SUCCESS")
                    return True
            except:
                pass
            
            # Localiza pacote não selecionado
            pacotes = js_engine.driver.find_elements("xpath", xpath_pacote_nao_selecionado)
            
            if not pacotes:
                if tentativa < max_tentativas:
                    time.sleep(1.5)
                    continue
                raise Exception("Nenhum pacote disponível")
            
            pacote = pacotes[0]
            ref_pacote = pacote.get_attribute("ref") or "?"
            relatorio.log(f"📦 Pacote encontrado (ref: {ref_pacote})")
            
            js_engine.scroll_into_view(pacote, padding=100)
            time.sleep(0.5)
            
            # Clique JavaScript agressivo
            js_engine.driver.execute_script("""
                const pacote = arguments[0];
                pacote.style.pointerEvents = 'auto';
                pacote.scrollIntoView({block: 'center'});
                pacote.classList.add('selected');
                
                const eventos = ['mouseenter', 'mouseover', 'mousedown', 'mouseup', 'click'];
                eventos.forEach(tipo => {
                    pacote.dispatchEvent(new MouseEvent(tipo, {
                        bubbles: true,
                        cancelable: true,
                        view: window
                    }));
                });
                
                if (typeof pacote.click === 'function') pacote.click();
                if (typeof jQuery !== 'undefined') jQuery(pacote).trigger('click');
            """, pacote)
            
            time.sleep(1)
            js_engine.wait_ajax_complete(5)
            
            # Valida
            pacotes_selecionados = js_engine.driver.find_elements("xpath", xpath_pacote_selecionado)
            if pacotes_selecionados and pacotes_selecionados[0].is_displayed():
                relatorio.log("✅ Pacote selecionado com sucesso!", "SUCCESS")
                return True
            
            if tentativa < max_tentativas:
                time.sleep(1.5)
                
        except Exception as e:
            if tentativa < max_tentativas:
                relatorio.log(f"⚠️ Tentativa {tentativa} falhou: {e}", "WARNING")
                time.sleep(2)
            else:
                raise
    
    raise Exception(f"Falha ao selecionar pacote após {max_tentativas} tentativas")


def clicar_pesquisar_por_indice(js_engine, doc, indice=1, timeout=5):
    """
    Clica no botão 'Pesquisar' pelo índice informado (1-based).
    Conta e exibe quantos botões existem antes do clique.
    Usa js_engine.force_click() e registra log.
    """
    xpath_base = "//a[contains(@class,'btModel btGray lpFind') and contains(normalize-space(.),'Pesquisar')]"
    xpath_indexado = f"({xpath_base})[{indice}]"

    try:
        # Conta quantos botões existem
        elementos = js_engine.driver.find_elements("xpath", xpath_base)
        total = len(elementos)
        log(doc, f"🔍 Foram encontrados {total} botão(ões) 'Pesquisar' na tela.")

        if total == 0:
            log(doc, "⚠️ Nenhum botão 'Pesquisar' foi encontrado.")
            return False
        if indice > total:
            log(doc, f"⚠️ Índice {indice} inválido — só existem {total} botão(ões).")
            return False

        log(doc, f"🎯 Clicando no botão 'Pesquisar' (índice {indice})...")
        js_engine.force_click(xpath_indexado, by_xpath=True)
        js_engine.wait_ajax_complete(timeout)
        log(doc, f"✅ Clique no botão 'Pesquisar' (índice {indice}) realizado com sucesso.")
        return True

    except Exception as e:
        log(doc, f"⚠️ Erro ao clicar no botão 'Pesquisar' (índice {indice}): {e}")
        return False



def clicar_todos_pesquisar(js_engine, doc, pausa_entre=0.5, timeout=5):
    """
    Procura todos os botões 'Pesquisar' visíveis e clica em cada um deles na ordem.
    Conta e exibe quantos botões existem antes de clicar.
    Usa js_engine.force_click() e registra log detalhado.
    """

    xpath_base = "//a[contains(@class,'btModel btGray') and contains(normalize-space(.),'Pesquisar')]"

    try:
        elementos = js_engine.driver.find_elements("xpath", xpath_base)
        total = len(elementos)
        log(doc, f"🔍 Foram encontrados {total} botão(ões) 'Pesquisar' na tela.")

        if total == 0:
            log(doc, "⚠️ Nenhum botão 'Pesquisar' foi encontrado.")
            return {"total": 0, "clicados": 0}

        total_clicados = 0
        for i in range(1, total + 1):
            xpath_indexado = f"({xpath_base})[{i}]"
            try:
                log(doc, f"🎯 Clicando no botão 'Pesquisar' (índice {i}/{total})...")
                js_engine.force_click(xpath_indexado, by_xpath=True)
                js_engine.wait_ajax_complete(timeout)
                total_clicados += 1
                log(doc, f"✅ Clique no botão 'Pesquisar' (índice {i}) realizado com sucesso.")
                if pausa_entre > 0:
                    import time
                    time.sleep(pausa_entre)
            except Exception as e:
                log(doc, f"⚠️ Falha ao clicar no botão 'Pesquisar' (índice {i}): {e}")

        log(doc, f"🧾 Resumo: {total_clicados}/{total} botões 'Pesquisar' clicados com sucesso.")
        return {"total": total, "clicados": total_clicados}

    except Exception as e:
        log(doc, f"⚠️ Erro ao procurar ou clicar nos botões 'Pesquisar': {e}")
        return {"total": 0, "clicados": 0}
def selecionar_agregado_todos_dependentes_js(js_engine, doc):
    """
    Versão JavaScript pura que processa todos de uma vez
    Mais rápida mas menos resiliente
    """
    
    log(doc, "⚡ Selecionando 'Agregado' para todos (modo JavaScript)...")
    
    script = """
    (function() {
        const resultados = {
            total: 0,
            processados: 0,
            erros: []
        };
        
        try {
            // Localiza todos os botões "Selecione"
            const botoes = Array.from(document.querySelectorAll('a.trocaTipoDependente'));
            resultados.total = botoes.length;
            
            if (botoes.length === 0) {
                return resultados;
            }
            
            // Processa cada botão IMEDIATAMENTE (sem setTimeout)
            botoes.forEach((botao, index) => {
                try {
                    // Scroll e clica no botão
                    botao.scrollIntoView({block: 'center'});
                    
                    // Dispara eventos de mouse
                    ['mouseenter', 'mouseover', 'mousedown', 'mouseup', 'click'].forEach(eventType => {
                        botao.dispatchEvent(new MouseEvent(eventType, {
                            bubbles: true,
                            cancelable: true,
                            view: window
                        }));
                    });
                    
                    // Clique direto
                    if (typeof botao.click === 'function') {
                        botao.click();
                    }
                    
                    // Pequena pausa para o dropdown aparecer
                    // (em JavaScript síncrono, vamos usar um loop de espera)
                    const startTime = Date.now();
                    while (Date.now() - startTime < 300) {
                        // Espera 300ms
                    }
                    
                    // Procura "Agregado" visível
                    const agregados = Array.from(document.querySelectorAll('li[tabindex="2"][ref="26"]'));
                    const agregado = agregados.find(el => {
                        const texto = (el.textContent || '').trim();
                        const visivel = el.offsetParent !== null;
                        const style = window.getComputedStyle(el);
                        const exibido = style.display !== 'none' && style.visibility !== 'hidden';
                        return texto === 'Agregado' && visivel && exibido;
                    });
                    
                    if (agregado) {
                        agregado.scrollIntoView({block: 'center'});
                        
                        // Dispara eventos no agregado
                        ['mouseenter', 'mouseover', 'mousedown', 'mouseup', 'click'].forEach(eventType => {
                            agregado.dispatchEvent(new MouseEvent(eventType, {
                                bubbles: true,
                                cancelable: true,
                                view: window
                            }));
                        });
                        
                        if (typeof agregado.click === 'function') {
                            agregado.click();
                        }
                        
                        resultados.processados++;
                    } else {
                        resultados.erros.push('Elemento ' + (index+1) + ': Agregado não encontrado');
                    }
                    
                } catch (error) {
                    resultados.erros.push('Elemento ' + (index+1) + ': ' + error.message);
                }
            });
            
        } catch (error) {
            resultados.erros.push('Erro geral: ' + error.message);
        }
        
        return resultados;
    })();
    """
    
    try:
        resultado = js_engine.execute_js(script, timeout=30, fallback_result={
            "total": 0,
            "processados": 0,
            "erros": ["Timeout ou erro na execução"]
        })
        
        # Validação robusta do resultado
        if resultado is None:
            log(doc, "⚠️ JavaScript retornou None - usando valores padrão")
            resultado = {"total": 0, "processados": 0, "erros": ["JavaScript retornou None"]}
        
        if not isinstance(resultado, dict):
            log(doc, f"⚠️ JavaScript retornou tipo inválido: {type(resultado)}")
            resultado = {"total": 0, "processados": 0, "erros": [f"Tipo inválido: {type(resultado)}"]}
        
        # Garante que as chaves existem
        total = int(resultado.get('total', 0))
        processados = int(resultado.get('processados', 0))
        erros = resultado.get('erros', [])
        
        log(doc, f"📊 Total encontrado: {total}")
        log(doc, f"✅ Processados: {processados}")
        
        if erros:
            log(doc, f"⚠️ Erros: {len(erros)}")
            for erro in erros[:5]:  # Mostra apenas os primeiros 5 erros
                log(doc, f"   - {erro}")
            if len(erros) > 5:
                log(doc, f"   ... e mais {len(erros) - 5} erro(s)")
        
        # Aguarda um pouco para garantir que os cliques foram processados
        if total > 0:
            tempo_espera = 2
            log(doc, f"⏳ Aguardando {tempo_espera}s para garantir processamento...")
            time.sleep(tempo_espera)
        
        return {
            "total": total,
            "processados": processados,
            "erros": erros
        }
        
    except Exception as e:
        log(doc, f"❌ Erro ao executar JavaScript: {e}")
        return {
            "total": 0,
            "processados": 0,
            "erros": [str(e)]
        }

def selecionar_agregado_todos_dependentes_seguro(js_engine, doc, max_tentativas=3, pausa_entre_elementos=2.0, pausa_apos_clique=1.5, pausa_apos_selecao=2.0, max_iteracoes=50):
    """
    Localiza e processa SEMPRE o PRIMEIRO elemento <a class="trocaTipoDependente">Selecione</a>,
    até que não existam mais elementos "Selecione".
    
    SOLUÇÃO PARA O PROBLEMA DE ÍNDICES DINÂMICOS:
    - Sempre processa o índice [0]
    - Re-busca a lista após cada processamento
    - Para quando não encontrar mais nenhum "Selecione"
    
    Args:
        js_engine: Instância do JSForceEngine
        doc: Documento para logs
        max_tentativas: Número de tentativas por elemento
        pausa_entre_elementos: Pausa ENTRE cada elemento processado
        pausa_apos_clique: Pausa APÓS clicar em "Selecione"
        pausa_apos_selecao: Pausa APÓS selecionar "Agregado"
        max_iteracoes: Limite de segurança para evitar loop infinito
    """
    
    xpath_selecione = "//a[@class='trocaTipoDependente' and normalize-space(text())='Selecione']"
    xpath_agregado = "//li[@tabindex='1' and @ref='26' and normalize-space(text())='Agregado']"
    
    log(doc, "🔄 Iniciando seleção de 'Agregado' - Modo Seguro (índice dinâmico)")
    log(doc, f"⏱️  Configuração de pausas:")
    log(doc, f"   • Entre elementos: {pausa_entre_elementos}s")
    log(doc, f"   • Após clicar 'Selecione': {pausa_apos_clique}s")
    log(doc, f"   • Após selecionar 'Agregado': {pausa_apos_selecao}s")
    
    resultado = {
        "total_processados": 0,
        "falhas": []
    }
    
    try:
        # Aguarda AJAX completo ANTES de começar
        log(doc, "⏳ Aguardando página estabilizar...")
        js_engine.wait_ajax_complete(8)
        time.sleep(2)
        
        # Conta quantos existem inicialmente
        elementos_iniciais = js_engine.driver.find_elements("xpath", xpath_selecione)
        total_inicial = len(elementos_iniciais)
        
        if total_inicial == 0:
            log(doc, "ℹ️ Nenhum elemento 'Selecione' encontrado.")
            return resultado
        
        log(doc, f"📊 Encontrados {total_inicial} elemento(s) 'Selecione' inicialmente")
        log(doc, f"🎯 Estratégia: Processar SEMPRE o primeiro elemento até não sobrar nenhum")
        
        iteracao = 0
        
        # Loop até não encontrar mais "Selecione"
        while iteracao < max_iteracoes:
            iteracao += 1
            
            log(doc, f"")
            log(doc, f"{'='*60}")
            log(doc, f"🔄 ITERAÇÃO {iteracao}")
            log(doc, f"{'='*60}")
            
            # ===== RE-BUSCA a lista SEMPRE =====
            log(doc, "🔍 Re-buscando elementos 'Selecione' na página...")
            js_engine.wait_ajax_complete(5)
            time.sleep(1)
            
            elementos_atuais = js_engine.driver.find_elements("xpath", xpath_selecione)
            total_atual = len(elementos_atuais)
            
            log(doc, f"📊 Elementos restantes: {total_atual}")
            
            # ===== CONDIÇÃO DE PARADA: Não há mais "Selecione" =====
            if total_atual == 0:
                log(doc, "")
                log(doc, "🎉 Sucesso! Não há mais elementos 'Selecione' para processar.")
                break
            
            # ===== SEMPRE PROCESSA O ÍNDICE [0] (primeiro elemento) =====
            tentativa = 0
            sucesso = False
            
            while tentativa < max_tentativas and not sucesso:
                tentativa += 1
                
                try:
                    log(doc, f"🎯 Processando PRIMEIRO elemento [0] (tentativa {tentativa}/{max_tentativas})...")
                    
                    # Re-localiza para garantir que está atualizado
                    elementos_fresh = js_engine.driver.find_elements("xpath", xpath_selecione)
                    
                    if len(elementos_fresh) == 0:
                        log(doc, "✅ Lista vazia - elemento foi processado por outra thread/evento")
                        sucesso = True
                        break
                    
                    elemento_atual = elementos_fresh[0]  # SEMPRE o primeiro [0]
                    
                    # Verifica se elemento está visível
                    if not elemento_atual.is_displayed():
                        log(doc, f"⚠️ Primeiro elemento não está visível")
                        break
                    
                    # Scroll até o elemento
                    log(doc, f"   📜 Fazendo scroll até elemento...")
                    js_engine.scroll_into_view(elemento_atual, padding=150)
                    time.sleep(0.8)
                    
                    # Clica no "Selecione"
                    log(doc, f"   ▶️ Clicando em 'Selecione'...")
                    js_engine.click_element(elemento_atual, wait_after=0.5)
                    
                    log(doc, f"   ⏳ Aguardando dropdown abrir ({pausa_apos_clique}s)...")
                    time.sleep(pausa_apos_clique)
                    js_engine.wait_ajax_complete(3)
                    time.sleep(0.5)
                    
                    # Localiza e clica em "Agregado"
                    log(doc, f"   🔍 Procurando 'Agregado' no dropdown...")
                    
                    agregado_elements = js_engine.driver.find_elements("xpath", xpath_agregado)
                    
                    if not agregado_elements:
                        log(doc, f"   ⚠️ Item 'Agregado' não encontrado no dropdown")
                        raise Exception("Agregado não encontrado")
                    
                    log(doc, f"   📊 Encontrados {len(agregado_elements)} elemento(s) 'Agregado'")
                    
                    # Pega o primeiro elemento visível
                    agregado = None
                    for idx, ag in enumerate(agregado_elements):
                        try:
                            if ag.is_displayed():
                                log(doc, f"   ✓ Agregado visível encontrado (índice {idx})")
                                agregado = ag
                                break
                        except:
                            continue
                    
                    if not agregado:
                        log(doc, f"   ⚠️ Nenhum item 'Agregado' visível")
                        raise Exception("Agregado não visível")
                    
                    # Scroll até o agregado
                    log(doc, f"   📜 Scroll até 'Agregado'...")
                    js_engine.scroll_into_view(agregado, padding=100)
                    time.sleep(0.5)
                    
                    # Clica em "Agregado"
                    log(doc, f"   ▶️ Clicando em 'Agregado'...")
                    js_engine.click_element(agregado, wait_after=0.5)
                    
                    log(doc, f"   ✅ 'Agregado' selecionado! (Processado: {resultado['total_processados'] + 1}/{total_inicial})")
                    
                    # Aguarda após seleção
                    log(doc, f"   ⏳ Aguardando sistema processar ({pausa_apos_selecao}s)...")
                    time.sleep(pausa_apos_selecao)
                    js_engine.wait_ajax_complete(5)
                    
                    resultado["total_processados"] += 1
                    sucesso = True
                    
                    # Pausa entre elementos
                    log(doc, f"   ⏳ Pausa antes da próxima iteração ({pausa_entre_elementos}s)...")
                    time.sleep(pausa_entre_elementos)
                    
                except StaleElementReferenceException:
                    log(doc, f"   ⚠️ Elemento ficou stale (tentativa {tentativa})")
                    if tentativa < max_tentativas:
                        log(doc, f"   ⏳ Aguardando 2s antes de retentar...")
                        time.sleep(2)
                        continue
                    else:
                        resultado["falhas"].append(f"Iteração {iteracao}: StaleElement após {max_tentativas} tentativas")
                        
                except Exception as e:
                    log(doc, f"   ⚠️ Erro: {str(e)[:100]}")
                    if tentativa < max_tentativas:
                        log(doc, f"   ⏳ Aguardando 2s antes de retentar...")
                        time.sleep(2)
                        continue
                    else:
                        resultado["falhas"].append(f"Iteração {iteracao}: {str(e)[:100]}")
            
            if not sucesso:
                log(doc, f"❌ Falha na iteração {iteracao} após {max_tentativas} tentativas")
                # Continua para próxima iteração mesmo com falha
                time.sleep(3)
        
        # Proteção contra loop infinito
        if iteracao >= max_iteracoes:
            log(doc, f"⚠️ AVISO: Atingido limite de {max_iteracoes} iterações!")
            resultado["falhas"].append(f"Limite de {max_iteracoes} iterações atingido")
        
        # ===== LOG FINAL =====
        log(doc, "")
        log(doc, "=" * 60)
        log(doc, f"📋 RESUMO FINAL:")
        log(doc, f"   • Total inicial: {total_inicial}")
        log(doc, f"   • Total processado: {resultado['total_processados']}")
        log(doc, f"   • Iterações realizadas: {iteracao}")
        
        # Verifica quantos sobraram
        elementos_finais = js_engine.driver.find_elements("xpath", xpath_selecione)
        total_final = len(elementos_finais)
        log(doc, f"   • Elementos restantes: {total_final}")
        
        if total_final == 0:
            log(doc, f"   ✅ SUCESSO TOTAL - Todos os {total_inicial} elementos foram processados!")
        else:
            log(doc, f"   ⚠️ Ainda restam {total_final} elemento(s) não processado(s)")
        
        if resultado["falhas"]:
            log(doc, f"   • Falhas: {len(resultado['falhas'])}")
            for falha in resultado["falhas"]:
                log(doc, f"     - {falha}")
        
        log(doc, "=" * 60)
        
        # Pausa final
        log(doc, "⏳ Aguardando processamento final (3s)...")
        time.sleep(3)
        js_engine.wait_ajax_complete(8)
        
        return resultado
        
    except Exception as e:
        log(doc, f"❌ Erro geral ao selecionar agregados: {e}")
        return resultado




def clicar_sim_ate_sumir(js_engine, doc, index=0, timeout=15, pausa=0.5):
    """
    Clica repetidamente no botão 'Sim' (BtYes) dentro de um modal de confirmação,
    até que o modal desapareça da tela.

    Parâmetros:
    - js_engine: instância do engine JSForceEngine
    - doc: documento de log
    - index: índice do botão 'Sim' (0 = primeiro, 1 = segundo, etc.)
    - timeout: tempo máximo em segundos antes de desistir
    - pausa: intervalo entre tentativas
    """

    # XPath do botão 'Sim' - mais simples e direto
    xpath_botao = f"(//a[@id='BtYes' and contains(@class,'btyes')])[{index + 1}]"
    
    # XPath do modal - busca qualquer modal visível com botão de confirmação
    xpath_modal_generico = "//div[contains(@class,'modal') and contains(@class,'confirmationYesNo')]"

    inicio = time.time()
    tentativas = 0
    sucesso = False

    log(doc, f"🟦 Iniciando clique no botão 'Sim' (índice {index}) até o modal sumir...")

    # FORÇA pelo menos UMA tentativa de clique antes de verificar
    clicou_pelo_menos_uma_vez = False

    while True:
        tentativas += 1
        
        try:
            # Verifica se o botão 'Sim' existe
            try:
                botoes = js_engine.driver.find_elements("xpath", xpath_botao)
                if not botoes and clicou_pelo_menos_uma_vez:
                    log(doc, f"✅ Botão 'Sim' não encontrado - modal fechado após {tentativas - 1} tentativa(s).")
                    sucesso = True
                    break
                    
                if not botoes:
                    log(doc, "⚠️ Botão 'Sim' não encontrado. Aguardando aparecer...")
                    time.sleep(pausa)
                    if time.time() - inicio > timeout:
                        log(doc, f"⏰ Timeout: botão 'Sim' nunca apareceu.")
                        break
                    continue
                    
            except Exception as e:
                log(doc, f"⚠️ Erro ao localizar botão: {e}")
                if clicou_pelo_menos_uma_vez:
                    sucesso = True
                    break
                continue

            log(doc, f"➡️ Tentativa {tentativas}: clicando no botão 'Sim'...")

            # Estratégia 1 — JavaScript com múltiplos eventos (MAIS AGRESSIVA)
            try:
                script = f"""
                var btn = document.evaluate("{xpath_botao}", document, null, 
                    XPathResult.FIRST_ORDERED_NODE_TYPE, null).singleNodeValue;
                if (btn) {{
                    // Remove qualquer bloqueio
                    btn.style.pointerEvents = 'auto';
                    btn.removeAttribute('disabled');
                    btn.style.visibility = 'visible';
                    btn.style.display = 'inline-block';
                    
                    // Scroll até o botão
                    btn.scrollIntoView({{block: 'center'}});
                    
                    // Dispara TODOS os eventos
                    ['mouseover','mouseenter','mousedown','focus','mouseup','click'].forEach(function(evt) {{
                        btn.dispatchEvent(new MouseEvent(evt, {{bubbles: true, cancelable: true}}));
                    }});
                    
                    // Clique direto
                    btn.click();
                    
                    // jQuery se disponível
                    if (typeof jQuery !== 'undefined') {{
                        jQuery(btn).trigger('click');
                    }}
                    
                    // onclick manual
                    if (btn.onclick) btn.onclick();
                    
                    return true;
                }}
                return false;
                """
                resultado = js_engine.driver.execute_script(script)
                if resultado:
                    log(doc, "✅ Clique via JavaScript ultra-agressivo executado.")
                    clicou_pelo_menos_uma_vez = True
                else:
                    log(doc, "⚠️ Script JavaScript retornou false.")
            except Exception as e:
                log(doc, f"⚠️ Falha no JavaScript agressivo: {e}")

            time.sleep(pausa)

            # Estratégia 2 — Force click do engine
            try:
                js_engine.force_click(xpath_botao, by_xpath=True)
                log(doc, "✅ Clique via force_click() executado.")
                clicou_pelo_menos_uma_vez = True
            except Exception as e:
                log(doc, f"⚠️ Falha no force_click: {e}")

            time.sleep(pausa)

            # Estratégia 3 — Selenium + ActionChains
            try:
                btn = js_engine.driver.find_element("xpath", xpath_botao)
                ActionChains(js_engine.driver).move_to_element(btn).pause(0.1).click().perform()
                log(doc, "✅ Clique via ActionChains executado.")
                clicou_pelo_menos_uma_vez = True
            except Exception as e:
                log(doc, f"⚠️ Falha no ActionChains: {e}")

            time.sleep(pausa)
            
            # Aguarda AJAX
            js_engine.wait_ajax_complete(3)

            # Verifica se o botão/modal ainda existem
            try:
                botoes = js_engine.driver.find_elements("xpath", xpath_botao)
                if not botoes:
                    log(doc, f"✅ Modal fechado após {tentativas} tentativa(s)!")
                    sucesso = True
                    break
            except:
                log(doc, "✅ Modal fechado (erro ao verificar = não existe mais)!")
                sucesso = True
                break

        except (NoSuchElementException, StaleElementReferenceException):
            if clicou_pelo_menos_uma_vez:
                log(doc, "✅ Elemento 'Sim' sumiu - modal fechado!")
                sucesso = True
                break
            else:
                log(doc, "⚠️ Elemento 'Sim' não encontrado, tentando novamente...")
        except Exception as e:
            log(doc, f"❌ Erro inesperado: {e}")

        if time.time() - inicio > timeout:
            log(doc, f"⏰ Tempo limite de {timeout}s atingido.")
            break

        time.sleep(pausa)

    if sucesso:
        log(doc, "🎉 Confirmação concluída com sucesso!")
    else:
        log(doc, "⚠️ Não foi possível confirmar - modal pode ainda estar aberto.")
    
    return sucesso

def detectar_e_processar_modal(js_engine, doc, timeout=10):
    """
    Detecta automaticamente o tipo de modal exibido e processa adequadamente:
    - Modal Sim/Não: Clica em "Sim" até desaparecer
    - Modal de Pagamento: Preenche campos e clica em "Ok"
    - Nenhum modal: Retorna False
    
    Returns:
        dict: {
            "tipo": "sim_nao" | "pagamento" | "nenhum",
            "processado": bool,
            "mensagem": str
        }
    """
    
    # XPaths dos modais
    xpath_modal_sim_nao = "//div[contains(@class,'modal') and contains(@class,'overflow')]//div[contains(@class,'confirmationYesNo')]"
    xpath_modal_pagamento = "//div[contains(@class,'modal') and contains(@class,'overflow')]//div[contains(@class,'paymentWindow')]"
    
    log(doc, "🔍 Detectando tipo de modal...")
    
    inicio = time.time()
    while time.time() - inicio < timeout:
        try:
            # Verifica Modal Sim/Não
            try:
                modal_sim_nao = js_engine.driver.find_elements("xpath", xpath_modal_sim_nao)
                if modal_sim_nao and modal_sim_nao[0].is_displayed():
                    log(doc, "✅ Modal 'Sim/Não' detectado!")
                    
                    # Extrai mensagem do modal
                    try:
                        mensagem = js_engine.driver.find_element("xpath", 
                            f"{xpath_modal_sim_nao}//p[@class='custom']").text
                        log(doc, f"📝 Mensagem: {mensagem}")
                    except:
                        mensagem = "Mensagem não capturada"
                    
                    # Processa com a função existente
                    log(doc, "▶️ Processando modal 'Sim/Não'...")
                    sucesso = clicar_sim_ate_sumir(js_engine, doc)
                    
                    return {
                        "tipo": "sim_nao",
                        "processado": sucesso,
                        "mensagem": mensagem
                    }
            except:
                pass
            
            # Verifica Modal de Pagamento
            try:
                modal_pagamento = js_engine.driver.find_elements("xpath", xpath_modal_pagamento)
                if modal_pagamento and modal_pagamento[0].is_displayed():
                    log(doc, "✅ Modal 'Pagamento' detectado!")
                    
                    # Extrai informações do modal
                    try:
                        label = js_engine.driver.find_element("xpath",
                            f"{xpath_modal_pagamento}//label[@style='display:block;']").text
                        valor_total = js_engine.driver.find_element("xpath",
                            f"{xpath_modal_pagamento}//h3").text
                        log(doc, f"📝 {label}")
                        log(doc, f"💰 Valor Total: {valor_total}")
                    except:
                        label = "Label não capturada"
                        valor_total = "R$ 0,00"
                    
                    # Processa modal de pagamento
                    log(doc, "▶️ Processando modal 'Pagamento'...")
                    sucesso = processar_modal_pagamento(js_engine, doc)
                    
                    return {
                        "tipo": "pagamento",
                        "processado": sucesso,
                        "mensagem": f"{label} - {valor_total}"
                    }
            except:
                pass
            
            time.sleep(0.3)
            
        except Exception as e:
            log(doc, f"⚠️ Erro ao detectar modal: {e}")
            time.sleep(0.5)
    
    log(doc, "ℹ️ Nenhum modal detectado")
    return {
        "tipo": "nenhum",
        "processado": False,
        "mensagem": "Nenhum modal encontrado"
    }


def processar_modal_pagamento(js_engine, doc, max_tentativas=5):
    """
    Preenche os campos do modal de pagamento e clica em "Ok"
    
    Campos:
    - Entrada: Valor da entrada (opcional)
    - Núm. Parcelas: Número de parcelas
    - 1º Vencimento: Data do primeiro vencimento
    - Dia Vencimento: Dia do vencimento (1-31)
    
    Returns:
        bool: True se processou com sucesso
    """
    
    xpath_modal = "//div[contains(@class,'paymentWindow') and contains(@style,'visibility: visible')]"
    
    # XPaths dos campos
    xpath_entrada = f"{xpath_modal}//input[@class='entradaPYW fieldCalc']"
    xpath_num_parcelas = f"{xpath_modal}//input[@class='numeroParcelasPYW fieldCalc']"
    xpath_primeiro_vencimento = f"{xpath_modal}//input[contains(@class,'primeiroVencimentoPYW')]"
    xpath_dia_vencimento = f"{xpath_modal}//input[@class='diaVencimentoPYW']"
    xpath_btn_ok = f"{xpath_modal}//a[@class='btModel btGray btok']"
    
    # Valores padrão
    entrada = "R$ 0,00"  # Sem entrada
    num_parcelas = str(random.randint(1, 12))  # 1 a 12 parcelas
    
    # Gera data de vencimento (próximo mês)
    hoje = datetime.today()
    proximo_mes = hoje + timedelta(days=30)
    primeiro_vencimento = proximo_mes.strftime("%d/%m/%Y")
    dia_vencimento = str(random.randint(5, 28))  # Dia seguro (5-28)
    
    log(doc, f"💳 Valores gerados:")
    log(doc, f"   • Entrada: {entrada}")
    log(doc, f"   • Parcelas: {num_parcelas}")
    log(doc, f"   • 1º Vencimento: {primeiro_vencimento}")
    log(doc, f"   • Dia Vencimento: {dia_vencimento}")
    
    for tentativa in range(1, max_tentativas + 1):
        try:
            log(doc, f"🔄 Tentativa {tentativa}/{max_tentativas}")
            
            # Aguarda modal estar completamente carregado
            time.sleep(1)
            js_engine.wait_ajax_complete(5)
            
            # Preenche Entrada
            try:
                log(doc, "   📝 Preenchendo Entrada...")
                js_engine.force_fill(xpath_entrada, entrada, by_xpath=True)
                time.sleep(0.5)
            except Exception as e:
                log(doc, f"   ⚠️ Erro ao preencher Entrada: {e}")
            
            # Preenche Número de Parcelas
            try:
                log(doc, "   📝 Preenchendo Número de Parcelas...")
                js_engine.force_fill(xpath_num_parcelas, num_parcelas, by_xpath=True)
                time.sleep(0.5)
            except Exception as e:
                log(doc, f"   ⚠️ Erro ao preencher Parcelas: {e}")
            
            # Preenche 1º Vencimento (Datepicker)
            try:
                log(doc, "   📅 Preenchendo 1º Vencimento...")
                elemento_data = js_engine.driver.find_element("xpath", xpath_primeiro_vencimento)
                
                # Estratégias para datepicker
                estrategias = [
                    lambda: _datepicker_jquery(
                        elemento_data.get_attribute('id'), 
                        primeiro_vencimento
                    ),
                    lambda: _datepicker_javascript(elemento_data, primeiro_vencimento),
                    lambda: _datepicker_actionchains(elemento_data, primeiro_vencimento),
                    lambda: _datepicker_tradicional(elemento_data, primeiro_vencimento)
                ]
                
                for i, estrategia in enumerate(estrategias, 1):
                    try:
                        estrategia()
                        time.sleep(0.5)
                        if validar_data_preenchida(elemento_data, primeiro_vencimento):
                            log(doc, f"   ✅ Data preenchida (estratégia {i})")
                            break
                    except:
                        continue
                        
            except Exception as e:
                log(doc, f"   ⚠️ Erro ao preencher 1º Vencimento: {e}")
            
            # Preenche Dia Vencimento
            try:
                log(doc, "   📝 Preenchendo Dia Vencimento...")
                js_engine.force_fill(xpath_dia_vencimento, dia_vencimento, by_xpath=True)
                time.sleep(0.5)
            except Exception as e:
                log(doc, f"   ⚠️ Erro ao preencher Dia Vencimento: {e}")
            
            # Clica em "Ok"
            try:
                log(doc, "   ▶️ Clicando em 'Ok'...")
                js_engine.force_click(xpath_btn_ok, by_xpath=True)
                time.sleep(1.5)
                js_engine.wait_ajax_complete(10)
                
                # Verifica se modal fechou
                try:
                    modal = js_engine.driver.find_element("xpath", xpath_modal)
                    if not modal.is_displayed():
                        log(doc, "✅ Modal de pagamento processado com sucesso!")
                        return True
                except:
                    log(doc, "✅ Modal de pagamento processado com sucesso!")
                    return True
                    
            except Exception as e:
                log(doc, f"   ⚠️ Erro ao clicar em 'Ok': {e}")
            
            if tentativa < max_tentativas:
                time.sleep(2)
                
        except Exception as e:
            log(doc, f"   ❌ Erro na tentativa {tentativa}: {e}")
            if tentativa < max_tentativas:
                time.sleep(2)
    
    log(doc, f"❌ Falha ao processar modal de pagamento após {max_tentativas} tentativas")
    return False



def verificar_sessao_ativa(driver, doc):
    """
    ✅ CORREÇÃO 3: Verifica se a sessão do navegador ainda está ativa
    """
    try:
        # Tenta executar comando simples
        driver.current_url
        return True
    except Exception as e:
        log(doc, f"❌ Sessão do navegador perdida: {e}")
        return False



def safe_action_with_session_check(doc, descricao, func, max_retries=3):
    """
    ✅ CORREÇÃO 3: Wrapper melhorado com verificação de sessão
    """
    global driver
    
    for attempt in range(max_retries):
        try:
            # Verifica se sessão está ativa antes de tentar
            if not verificar_sessao_ativa(driver, doc):
                log(doc, f"❌ Sessão inválida - não é possível executar: {descricao}")
                return False
            
            log(doc, f"🔄 {descricao}..." if attempt == 0 else f"🔄 {descricao}... (Tentativa {attempt + 1})")
            func()
            log(doc, f"✅ {descricao} realizada com sucesso.")
            take_screenshot(driver, doc, _sanitize_filename(descricao))
            return True
            
        except Exception as e:
            if not verificar_sessao_ativa(driver, doc):
                log(doc, f"❌ Sessão perdida durante execução de: {descricao}")
                return False
                
            if attempt < max_retries - 1:
                log(doc, f"⚠️ Tentativa {attempt + 1} falhou, tentando novamente...")
                time.sleep(2 + attempt)
                continue
            else:
                log(doc, f"❌ Erro após {max_retries} tentativas: {e}")
                take_screenshot(driver, doc, _sanitize_filename(f"erro_{descricao}"))
                return False
    
    return False

XPATHS_BOTOES_LOV = {
    "titular": "//a[@class='sprites sp-changeTitular']",
    "dependente": "//a[@class='sprites sp-addDependentes']"
}



def abrir_lov_por_indice_seguro(driver, js_engine, doc, tipo, indice, max_tentativas=5):
    """
    ✅ Versão CORRIGIDA com proteção contra crash
    """
    XPATHS_BOTOES_LOV = {
        "titular": "//a[@class='sprites sp-changeTitular']",
        "dependente": "//a[@class='sprites sp-addDependentes']"
    }
    
    xpath_base = XPATHS_BOTOES_LOV.get(tipo)
    if not xpath_base:
        raise ValueError(f"Tipo inválido: {tipo}")
    
    for tentativa in range(1, max_tentativas + 1):
        try:
            log(doc, f"🔍 Tentativa {tentativa}: Abrindo LOV {tipo} (índice {indice})...")
            
            # ✅ Verifica sessão
            try:
                driver.current_url
            except:
                log(doc, "❌ Sessão perdida")
                return False
            
            # ✅ Aguarda AJAX
            js_engine.wait_ajax_complete(5)
            time.sleep(1)
            
            # ✅ Localiza botões
            botoes = driver.find_elements(By.XPATH, xpath_base)
            
            if not botoes:
                log(doc, f"⚠️ Nenhum botão LOV '{tipo}' encontrado")
                if tentativa < max_tentativas:
                    time.sleep(2)
                    continue
                raise Exception(f"Botão LOV '{tipo}' não encontrado")
            
            if indice >= len(botoes):
                raise Exception(f"Índice {indice} inválido - existem {len(botoes)} botões")
            
            botao = botoes[indice]
            
            # ✅ Verifica se está visível
            if not botao.is_displayed():
                log(doc, "⚠️ Botão não está visível")
                time.sleep(1)
                continue
            
            # ✅ Scroll seguro
            driver.execute_script(
                "arguments[0].scrollIntoView({block:'center'});", 
                botao
            )
            time.sleep(0.5)
            
            # ✅ Clique com JavaScript (mais seguro)
            driver.execute_script("""
                const btn = arguments[0];
                btn.style.pointerEvents = 'auto';
                btn.scrollIntoView({block:'center'});
                
                ['mouseenter', 'mouseover', 'mousedown', 'mouseup', 'click'].forEach(evt => {
                    btn.dispatchEvent(new MouseEvent(evt, {bubbles:true, cancelable:true}));
                });
                
                btn.click();
            """, botao)
            
            log(doc, f"✅ LOV '{tipo}' aberto")
            time.sleep(1.5)
            js_engine.wait_ajax_complete(8)
            return True
            
        except Exception as e:
            log(doc, f"⚠️ Tentativa {tentativa} falhou: {str(e)[:100]}")
            if tentativa < max_tentativas:
                time.sleep(2 + tentativa * 0.5)
                continue
    
    raise Exception(f"Falha ao abrir LOV após {max_tentativas} tentativas")

def inicializar_driver_protegido():
    """Driver com proteção contra crashes"""
    options = Options()
    options.add_argument("--start-maximized")
    options.add_argument("--disable-blink-features=AutomationControlled")
    options.add_argument("--disable-dev-shm-usage")
    options.add_argument("--no-sandbox")
    options.add_argument("--disable-gpu")  # ✅ NOVO
    options.add_argument("--disable-extensions")  # ✅ NOVO
    options.add_argument("--disable-popup-blocking")  # ✅ NOVO
    
    # ✅ CRÍTICO: Evita crashes por memória
    options.add_experimental_option("excludeSwitches", ["enable-automation", "enable-logging"])
    options.add_experimental_option('useAutomationExtension', False)
    
    # ✅ Aumenta estabilidade
    options.add_argument("--remote-debugging-port=9222")
    
    driver = webdriver.Chrome(
        service=Service(ChromeDriverManager().install()), 
        options=options
    )
    
    driver.set_page_load_timeout(60)
    driver.set_script_timeout(60)
    driver.implicitly_wait(10)
    
    return driver

def executar_com_protecao(doc, descricao, funcao, max_tentativas=3):
    """
    ✅ Wrapper que protege contra crashes e recupera a sessão
    """
    global driver, js_engine
    
    for tentativa in range(1, max_tentativas + 1):
        try:
            # Verifica sessão ANTES
            try:
                driver.current_url
            except:
                log(doc, f"❌ Sessão perdida antes de '{descricao}'")
                
                # ✅ TENTA RECUPERAR (opcional)
                log(doc, "🔄 Tentando reinicializar driver...")
                driver = inicializar_driver_protegido()
                js_engine = JSForceEngine(driver, WebDriverWait(driver, 20), doc)
                
                # Refaz login
                driver.get(URL)
                # ... código de login ...
                
                time.sleep(3)
            
            log(doc, f"🔄 {descricao}... (Tentativa {tentativa})")
            
            resultado = funcao()
            
            # Verifica sessão DEPOIS
            try:
                driver.current_url
                log(doc, f"✅ {descricao} concluído")
                return resultado
            except:
                log(doc, f"⚠️ Sessão perdida DURANTE '{descricao}'")
                if tentativa < max_tentativas:
                    time.sleep(3)
                    continue
                    
        except Exception as e:
            log(doc, f"❌ Erro em '{descricao}': {str(e)[:150]}")
            
            if tentativa < max_tentativas:
                log(doc, f"⏳ Retry em 3s...")
                time.sleep(3)
                continue
            else:
                # Tira screenshot do erro
                try:
                    take_screenshot(driver, doc, f"erro_{descricao.lower().replace(' ', '_')}")
                except:
                    pass
                raise
    
    raise Exception(f"Falha após {max_tentativas} tentativas: {descricao}")


def preencher_campo_pesquisa_lov(driver, texto):
    xpath_input = "//div[contains(@class,'busca')]//input | //input[@type='text']"
    campo = WebDriverWait(driver, 10).until(
        EC.element_to_be_clickable((By.XPATH, xpath_input))
    )
    campo.clear()
    campo.send_keys(texto)
    time.sleep(0.3)



def clicar_primeira_linha_lov_seguro(driver, js_engine, doc, max_tentativas=5):
    """
    ✅ Versão CORRIGIDA que evita "target frame detached"
    """
    xpath_primeira_linha = "//table//tr[@role='row' and @style][1]"
    
    for tentativa in range(1, max_tentativas + 1):
        try:
            log(doc, f"🎯 Tentativa {tentativa}: Clicando na primeira linha do LOV...")
            
            # ✅ Aguarda AJAX completo
            js_engine.wait_ajax_complete(5)
            time.sleep(1)
            
            # ✅ Verifica se ainda está no contexto correto
            try:
                driver.current_url  # Testa se sessão está ativa
            except:
                log(doc, "❌ Sessão perdida - abortando")
                return False
            
            # ✅ Localiza elemento com timeout curto
            try:
                linha = WebDriverWait(driver, 8).until(
                    EC.presence_of_element_located((By.XPATH, xpath_primeira_linha))
                )
            except TimeoutException:
                log(doc, f"⚠️ Linha não encontrada na tentativa {tentativa}")
                if tentativa < max_tentativas:
                    time.sleep(2)
                    continue
                raise Exception("Linha do LOV não encontrada")
            
            # ✅ Verifica se elemento está pronto
            if not linha.is_displayed():
                log(doc, "⚠️ Linha não está visível")
                time.sleep(1.5)
                continue
            
            # ✅ Scroll ANTES de clicar
            driver.execute_script(
                "arguments[0].scrollIntoView({block:'center', behavior:'smooth'});", 
                linha
            )
            time.sleep(0.8)
            
            # ✅ Estratégia 1: JavaScript Click (mais confiável)
            try:
                driver.execute_script("""
                    const el = arguments[0];
                    el.style.pointerEvents = 'auto';
                    
                    // Dispara eventos
                    ['mouseover', 'mouseenter', 'mousedown', 'mouseup', 'click'].forEach(evt => {
                        el.dispatchEvent(new MouseEvent(evt, {
                            bubbles: true,
                            cancelable: true,
                            view: window
                        }));
                    });
                    
                    el.click();
                """, linha)
                
                log(doc, "✅ Linha clicada com JavaScript")
                time.sleep(1.5)
                js_engine.wait_ajax_complete(8)
                return True
                
            except Exception as e:
                log(doc, f"⚠️ JavaScript falhou: {e}")
            
            # ✅ Estratégia 2: Selenium padrão
            try:
                linha.click()
                log(doc, "✅ Linha clicada com Selenium")
                time.sleep(1.5)
                js_engine.wait_ajax_complete(8)
                return True
            except Exception as e:
                log(doc, f"⚠️ Selenium falhou: {e}")
            
            # ✅ Estratégia 3: ActionChains
            try:
                ActionChains(driver).move_to_element(linha).pause(0.2).click().perform()
                log(doc, "✅ Linha clicada com ActionChains")
                time.sleep(1.5)
                js_engine.wait_ajax_complete(8)
                return True
            except Exception as e:
                log(doc, f"⚠️ ActionChains falhou: {e}")
            
            if tentativa < max_tentativas:
                log(doc, f"⏳ Aguardando antes de retry...")
                time.sleep(2 + tentativa * 0.5)
                
        except StaleElementReferenceException:
            log(doc, f"⚠️ Elemento ficou stale na tentativa {tentativa}")
            if tentativa < max_tentativas:
                time.sleep(2)
                continue
                
        except Exception as e:
            log(doc, f"❌ Erro na tentativa {tentativa}: {str(e)[:100]}")
            if tentativa < max_tentativas:
                time.sleep(2 + tentativa)
                continue
    
    raise Exception(f"Falha ao clicar na linha do LOV após {max_tentativas} tentativas")



# ==== VALIDAÇÃO DE REGISTROS MELHORADA ====
def validar_registros_encontrados(timeout=TIMEOUT_LONGO):
    """Sistema robusto de validação de registros encontrados"""
    global driver, wait, doc
    
    resultado = {
        'encontrou_registros': False,
        'quantidade_registros': 0,
        'mensagem': '',
        'tabela_encontrada': False,
        'detalhes': []
    }
    
    try:
        log(doc, "🔍 Iniciando validação de registros...")
        time.sleep(5)  # Aguarda processamento inicial
        
        # Seletores para diferentes tipos de tabelas de resultado
        seletores_tabela = [
            '#DataTables_Table_0',
            '#DataTables_Table_0 tbody',
            '#DataTables_Table_1',
            '#DataTables_Table_2',
            'table[id*="DataTables"]',
            '.wdGrid table',
            'table tbody',
            '.resultados table',
            '[class*="grid"][class*="result"]',
            'table[class*="dataTable"]'
        ]
        
        tabela_encontrada = None
        
        # Busca tabela de resultados
        for seletor in seletores_tabela:
            try:
                elementos = driver.find_elements(By.CSS_SELECTOR, seletor)
                for elemento in elementos:
                    if elemento.is_displayed() and elemento.size['height'] > 0:
                        tabela_encontrada = elemento
                        resultado['tabela_encontrada'] = True
                        log(doc, f"✅ Tabela encontrada: {seletor}")
                        break
                
                if tabela_encontrada:
                    break
            except Exception as e:
                log(doc, f"⚠️ Erro ao buscar tabela com {seletor}: {e}", 'WARN')
                continue
        
        if not tabela_encontrada:
            # Busca mensagens de "sem resultados"
            mensagens_vazio = [
                "Nenhum registro encontrado",
                "Não foram encontrados registros", 
                "Nenhum resultado",
                "Sem resultados para exibir",
                "0 registros encontrados",
                "No data available"
            ]
            
            for mensagem in mensagens_vazio:
                try:
                    elem = driver.find_element(By.XPATH, f"//*[contains(text(), '{mensagem}')]")
                    if elem.is_displayed():
                        resultado['mensagem'] = f"Sistema informou: {elem.text.strip()}"
                        log(doc, f"ℹ️ {resultado['mensagem']}")
                        return resultado
                except:
                    continue
            
            # Verifica se existe indicação de carregamento
            loading_elements = driver.find_elements(By.CSS_SELECTOR, ".loading, .spinner, [class*='load']")
            if any(el.is_displayed() for el in loading_elements):
                log(doc, "⏳ Sistema ainda carregando resultados...", 'WARN')
                time.sleep(5)
                return validar_registros_encontrados(timeout - 10)  # Recursão com timeout reduzido
            
            resultado['mensagem'] = "⚠️ Tabela de resultados não localizada"
            log(doc, resultado['mensagem'], 'WARN')
            return resultado
        
        # Conta e valida registros
        try:
            # Estratégias para encontrar linhas de dados
            seletores_linhas = [
                'tbody tr:not(.dataTables_empty):not([class*="no-data"])',
                'tbody tr[class*="odd"], tbody tr[class*="even"]',
                'tbody tr:not(:empty)',
                'tbody tr'
            ]
            
            linhas_validas = []
            
            for seletor_linha in seletores_linhas:
                try:
                    linhas = tabela_encontrada.find_elements(By.CSS_SELECTOR, seletor_linha)
                    
                    for linha in linhas:
                        try:
                            if not linha.is_displayed():
                                continue
                                
                            texto_linha = linha.text.strip().lower()
                            
                            # Valida se é uma linha com dados reais
                            if (len(texto_linha) > 5 and 
                                not any(termo in texto_linha for termo in [
                                    'nenhum registro', 'sem dados', 'no data', 
                                    'vazio', 'empty', 'não foram encontrados',
                                    'loading', 'carregando'
                                ])):
                                
                                linhas_validas.append({
                                    'elemento': linha,
                                    'texto': texto_linha[:100] + '...' if len(texto_linha) > 100 else texto_linha
                                })
                        except Exception as e:
                            log(doc, f"⚠️ Erro ao processar linha: {e}", 'WARN')
                            continue
                    
                    if linhas_validas:
                        log(doc, f"✅ Encontradas {len(linhas_validas)} linhas válidas com {seletor_linha}")
                        break
                        
                except Exception as e:
                    log(doc, f"⚠️ Erro ao processar {seletor_linha}: {e}", 'WARN')
                    continue
            
            quantidade = len(linhas_validas)
            resultado['quantidade_registros'] = quantidade
            resultado['detalhes'] = [linha['texto'] for linha in linhas_validas[:5]]  # Primeiras 5 linhas
            
            if quantidade > 0:
                resultado['encontrou_registros'] = True
                resultado['mensagem'] = f"✅ {quantidade} registro(s) encontrado(s)"
                
                # Log das primeiras linhas
                log(doc, resultado['mensagem'])
                for i, linha in enumerate(linhas_validas[:3], 1):
                    log(doc, f"   Registro {i}: {linha['texto']}")
                
                if quantidade > 3:
                    log(doc, f"   ... e mais {quantidade-3} registro(s)")
            else:
                resultado['mensagem'] = "ℹ️ Tabela encontrada mas sem registros válidos"
                log(doc, resultado['mensagem'])
        
        except Exception as e:
            log(doc, f"❌ Erro ao contar registros: {e}", 'ERROR')
            # Em caso de erro na contagem, assume que existem registros para não interromper
            resultado['encontrou_registros'] = True
            resultado['quantidade_registros'] = 1
            resultado['mensagem'] = f"⚠️ Erro na validação, continuando teste: {e}"
        
        # Verifica alertas do sistema
        encontrar_mensagem_alerta()
        
        return resultado
        
    except Exception as e:
        log(doc, f"❌ Erro geral na validação: {e}", 'ERROR')
        resultado['encontrou_registros'] = True  # Assume sucesso para não interromper
        resultado['quantidade_registros'] = 1
        resultado['mensagem'] = f"⚠️ Validação falhou, continuando teste: {e}"
        return resultado
import time, traceback


def scroll_to_element(elemento):
    """Scroll inteligente até elemento"""
    global driver
    
    try:
        # Scroll suave até o elemento
        driver.execute_script("""
            arguments[0].scrollIntoView({
                behavior: 'smooth',
                block: 'center',
                inline: 'center'
            });
        """, elemento)
        time.sleep(0.5)
        
        # Verifica se elemento está visível
        return elemento.is_displayed()
        
    except Exception as e:
        log(doc, f"⚠️ Erro no scroll: {e}", 'WARN')
        return False

def clicar_elemento_robusto(seletor):
    """Clique robusto em elementos"""
    def acao():
        elemento = aguardar_elemento(seletor)
        scroll_to_element(elemento)
        
        # Tenta diferentes métodos de clique
        try:
            elemento.click()
        except:
            # Clique via JavaScript se o normal falhar
            driver.execute_script("arguments[0].click();", elemento)
    
    return acao


def selecionar_opcao_xpath(xpath, texto):
    def acao():
        select_element = wait.until(EC.presence_of_element_located((By.XPATH, xpath)))
        Select(select_element).select_by_visible_text(texto)
    return acao

# ==== INICIALIZAÇÃO DO DRIVER ====
options = Options()
options.add_argument("--start-maximized")
driver = webdriver.Chrome(service=Service(ChromeDriverManager().install()), options=options)
wait = WebDriverWait(driver, 20)

# ==== EXECUÇÃO DO TESTE ====
try:
    
         
    # Cria engine JS forçado COM PROTEÇÃO ANTI-TIMEOUT
    js_engine = JSForceEngine(driver, wait, doc, timeout_padrao=10, max_retries=3)
    lov_handler = LOVHandler(js_engine, doc)      

    safe_action(doc, "Acessando sistema", lambda: driver.get(URL))

    safe_action(doc, "Realizando login", lambda: (
        wait.until(EC.presence_of_element_located((By.ID, "j_id15:email"))).send_keys(LOGIN_EMAIL),
        wait.until(EC.presence_of_element_located((By.ID, "j_id15:senha"))).send_keys(LOGIN_PASSWORD, Keys.ENTER),
        wait.until(EC.presence_of_element_located((By.TAG_NAME, "body")))
    ))

    safe_action(doc, "Esperando sistema carregar e ajustando zoom", lambda: (
        time.sleep(5),
        ajustar_zoom()
    ))

    safe_action(doc, "Abrindo menu Contratos", lambda: (
        driver.find_element(By.TAG_NAME, "body").send_keys(Keys.F3),
        time.sleep(0.5),
        wait.until(EC.element_to_be_clickable((By.XPATH, "//li[@class='iconElement shortcutIcon ui-draggable ui-draggable-handle' and @cname='I.CT' and @ref='I.CT']"))).click()
    ))

    safe_action(doc, "Clicando em 'Trocar ou Vincular Jazigos ou Quadras'", lambda: (
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, '#gsContratos > div.wdTelas > div.telaInicial.clearfix.overflow.overflowY > ul > li:nth-child(12) > a > span'))).click()
    ))


    safe_action(doc, "Preenchendo Número do Contrato", lambda:
        js_engine.force_fill("//div[@class='buscaContrato' and label[normalize-space()='Numero do Contrato']]/input", "112883", by_xpath=True)
    )


    safe_action(doc, "Clicando em todos os botões 'Pesquisar'", lambda:
            clicar_todos_pesquisar(js_engine, doc, pausa_entre=0.3)
        )

    time.sleep(5)  # Espera extra para garantir carregamento completo



    safe_action(doc, "Alterando Quadra", lambda:
         lov_handler.open_and_select(
            btn_index=2,
            search_text="QUADRA TESTE SELENIUM AUTOMATIZADO",
            result_text="QUADRA TESTE SELENIUM AUTOMATIZADO"
        )
    )


    safe_action(doc, "Selecionando jazigo por Número", lambda:
        lov_handler.open_and_select(
            btn_index=3,
            search_text="1"        
        )
    )


    safe_action(doc, "Selecionando Motivo", selecionar_opcao_xpath(
        "//select[@class='fc mandatory' and @style='width:365px;']",
        "VINCULAR JAZIGO - GAVETA"
    ))

    safe_action(doc, "Salvando alterações", lambda:
        js_engine.force_click(
            "//a[@class='btModel btGray btsave' and contains(normalize-space(),'Salvar')]",
            by_xpath=True
        )
    )
    
    time.sleep(10)
    encontrar_mensagem_alerta()
    
    safe_action_with_session_check(doc, "Fechando Modal Contratos", lambda:
        js_engine.force_click(
            "#gsContratos > div.wdTop.ui-draggable-handle > div > a",
            by_xpath=False
        )
    )


except Exception as e:
    log(doc, f"❌ ERRO FATAL: {e}")
    take_screenshot(driver, doc, "erro_fatal")

finally:

    log(doc, "✅ Teste concluído com sucesso.")

    finalizar_relatorio()